"""Period-comparison chart helpers for the Codex plugin."""

from __future__ import annotations

import argparse
import calendar
import json
import logging
import math
import shutil
import sys
import warnings
import zipfile
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from html import escape
from pathlib import Path
from typing import Any, Iterable, Sequence

import polars as pl
from legacy_charting import (
    write_legacy_actual_vs_previous_year_chart,
    write_legacy_dot_chart,
    write_legacy_horizontal_waterfall_chart,
    write_legacy_multitier_column_chart,
    write_legacy_slope_chart,
)

SCRIPT_DIR = Path(__file__).resolve().parent


def _ensure_local_review_session_import() -> None:
    """Use this plugin's review-session module in multi-plugin test runs."""

    script_dir = str(SCRIPT_DIR)
    if script_dir in sys.path:
        sys.path.remove(script_dir)
    sys.path.insert(0, script_dir)
    module = sys.modules.get("review_session")
    module_file = getattr(module, "__file__", None) if module is not None else None
    if module_file and Path(module_file).resolve().is_relative_to(SCRIPT_DIR.resolve()):
        return
    if module is not None:
        del sys.modules["review_session"]


_ensure_local_review_session_import()
from review_session import write_review_session_artifacts, write_run_intake

__all__ = [
    "InspectionResult",
    "PeriodComparisonRunResult",
    "add_common_args",
    "configure_logging",
    "inspect_period_comparison_inputs",
    "run_period_comparison",
]

LOGGER = logging.getLogger(__name__)
SCHEMA_VERSION = "1.0"
PLUGIN_ROOT = Path(__file__).resolve().parents[1]
VENDOR_ROOT = PLUGIN_ROOT / "vendor"
REPO_ROOT = Path(__file__).resolve().parents[3]
SHARED_VENDOR_ROOT = REPO_ROOT / "plugins" / "_shared" / "vendor"


def _ensure_shared_modules_path() -> None:
    """Make shared plugin harness modules importable in repo and ZIP runs."""

    shared_parent = SHARED_VENDOR_ROOT if SHARED_VENDOR_ROOT.exists() else VENDOR_ROOT
    shared_text = str(shared_parent)
    if shared_text in sys.path:
        sys.path.remove(shared_text)
    sys.path.insert(0, shared_text)
    module_root = (shared_parent / "modules").resolve()
    for name, module in list(sys.modules.items()):
        if name == "modules" or name.startswith("modules."):
            module_file = getattr(module, "__file__", None)
            if not module_file or not Path(module_file).resolve().is_relative_to(
                module_root
            ):
                del sys.modules[name]


_ensure_shared_modules_path()
from modules.chart_harness import (  # noqa: E402  # isort: skip
    apply_recipe_filters,
    apply_recipe_cohorts,
    available_analysis_context,
    PERIOD_GRAIN_MONTH,
    PERIOD_GRAIN_QUARTER,
    PERIOD_GRAIN_WEEK,
    PERIOD_GRAIN_YEAR,
    PERIOD_TYPE_CALENDAR,
    PERIOD_TYPE_FISCAL,
    PERIOD_TYPE_ROLLING,
    PERIOD_TYPE_TO_DATE,
    period_contract_options,
    preserve_recipe_cohorts,
    recipe_cohort_dimension_names,
    recipe_cohort_source_dimensions,
    preserve_recipe_filters,
    reporting_period_line_from_recipe,
    reporting_subject_label_from_recipe,
    write_prepared_data_manifest,
)

CANONICAL_DATE = "Date"
CANONICAL_PERIOD = "Period"
CURRENT_PERIOD = "AC"
PREVIOUS_PERIOD = "PY"
MONTH_ORDER = [
    "Jan",
    "Feb",
    "Mar",
    "Apr",
    "May",
    "Jun",
    "Jul",
    "Aug",
    "Sep",
    "Oct",
    "Nov",
    "Dec",
]
MONTH_LABELS = {index + 1: label for index, label in enumerate(MONTH_ORDER)}
TOLERANCE = 1e-9
ARTIFACT_MODE_DATA_ONLY = "data_only"
ARTIFACT_MODE_DATA_AND_RENDER = "data_and_render"
ARTIFACT_MODES = {
    ARTIFACT_MODE_DATA_ONLY,
    ARTIFACT_MODE_DATA_AND_RENDER,
}
NATIVE_TABLE_CAPABILITIES = {
    "comparison_table": "period_comparison.comparison_table",
    "time_series_table": "period_comparison.time_series_table",
}
PERIOD_CHART_CANDIDATE_IDS = (
    "year_over_year_column",
    "year_over_year_line",
    "year_over_year_by_period",
    "year_over_year_slope",
    "year_over_year_dot",
    "year_over_year_waterfall",
)
PERIOD_SMALL_MULTIPLES_CANDIDATE_ID = "year_over_year_small_multiples"
LEGACY_CHART_ERROR_TYPES = (
    ImportError,
    ModuleNotFoundError,
    OSError,
    RuntimeError,
    TypeError,
    ValueError,
    KeyError,
    pl.exceptions.PolarsError,
)


@dataclass(frozen=True)
class InspectionResult:
    """Inspection result and suggested recipe paths."""

    payload: dict[str, Any]
    recipe: dict[str, Any]
    output_dir: Path


@dataclass(frozen=True)
class PeriodComparisonRunResult:
    """Period-comparison run result."""

    monthly_frame: pl.DataFrame
    audit: dict[str, Any]
    summary_markdown: str
    artifact_paths: list[str]


def configure_logging(verbose: bool = False) -> None:
    """Configure command-line logging."""

    logging.basicConfig(
        level=logging.DEBUG if verbose else logging.WARNING,
        format="%(levelname)s:%(name)s:%(message)s",
    )


def add_common_args(parser: argparse.ArgumentParser) -> None:
    """Add common plugin CLI arguments."""

    parser.add_argument("input_file", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--recipe", type=Path)
    parser.add_argument("--language", default="en")
    parser.add_argument("--currency", default=None)
    parser.add_argument(
        "--artifact-mode",
        choices=sorted(ARTIFACT_MODES),
        default=ARTIFACT_MODE_DATA_AND_RENDER,
        help=(
            "Write chart data/context only or keep the legacy data-and-render behavior."
        ),
    )
    parser.add_argument("--verbose", action="store_true")


def utc_now() -> str:
    """Return an ISO timestamp for audit files."""

    return datetime.now(timezone.utc).isoformat()


def json_safe(value: Any) -> Any:
    """Return JSON-safe values for common analysis objects."""

    if isinstance(value, dict):
        return {str(key): json_safe(item) for key, item in value.items()}
    if isinstance(value, list):
        return [json_safe(item) for item in value]
    if isinstance(value, tuple):
        return [json_safe(item) for item in value]
    if isinstance(value, (Path, datetime, date)):
        return str(value)
    return value


def write_json(path: Path, payload: dict[str, Any]) -> None:
    """Write stable JSON."""

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(json_safe(payload), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def read_json(path: Path) -> dict[str, Any]:
    """Read a JSON object."""

    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"Expected JSON object: {path}")
    return payload


def _collect_csv_scan(path: Path, *, separator: str) -> pl.DataFrame:
    """Read delimited input through a lazy scan and collect once."""

    lf = pl.scan_csv(path, separator=separator, infer_schema_length=10000)
    try:
        return lf.collect(engine="streaming")
    except pl.exceptions.PolarsError:
        return lf.collect()


def read_table(path: Path) -> pl.DataFrame:
    """Read a supported CSV or Excel file."""

    suffix = path.suffix.lower()
    if suffix in {".csv", ".txt", ".tsv", ".psv"}:
        separator = {".tsv": "\t", ".psv": "|"}.get(suffix, ",")
        return _collect_csv_scan(path, separator=separator)
    if suffix in {".xlsx", ".xlsm"}:
        return pl.read_excel(path)
    raise ValueError(
        f"Unsupported input file type '{suffix}'. Use CSV, TSV, PSV, XLSX, or XLSM."
    )


def get_schema_and_column_names(df: pl.DataFrame) -> tuple[list[str], dict[str, str]]:
    """Return column names and JSON-safe schema."""

    return list(df.schema.keys()), {
        name: str(dtype) for name, dtype in df.schema.items()
    }


def normalize_name(name: str) -> str:
    """Normalize a column name for matching."""

    return " ".join(name.replace("_", " ").replace("-", " ").lower().split())


def compact_name(name: str) -> str:
    """Return a compact normalized name."""

    return "".join(normalize_name(name).split())


def first_matching_column(columns: Iterable[str], hints: Iterable[str]) -> str | None:
    """Return the first column whose normalized name matches the hints."""

    normalized = [
        (column, normalize_name(column), compact_name(column)) for column in columns
    ]
    hint_values = [(normalize_name(hint), compact_name(hint)) for hint in hints]
    for hint, compact_hint in hint_values:
        for column, column_name, column_compact in normalized:
            if column_name == hint or column_compact == compact_hint:
                return column
    for hint, compact_hint in hint_values:
        for column, column_name, column_compact in normalized:
            if hint in column_name or compact_hint in column_compact:
                return column
    return None


def numeric_columns(df: pl.DataFrame) -> list[str]:
    """Return numeric columns."""

    return [
        name
        for name, dtype in df.schema.items()
        if dtype.is_numeric() and not normalize_name(name).endswith("id")
    ]


def parse_date_expression(column: str) -> pl.Expr:
    """Return a permissive date parsing expression."""

    expr = pl.col(column)
    return expr.cast(pl.Date, strict=False).fill_null(
        expr.cast(pl.Utf8)
        .str.strptime(pl.Date, strict=False)
        .fill_null(
            expr.cast(pl.Utf8).str.strptime(pl.Datetime, strict=False).cast(pl.Date)
        )
    )


def infer_date_column(df: pl.DataFrame, columns: list[str]) -> str | None:
    """Infer the most likely date column."""

    direct = first_matching_column(
        columns, ["date", "order date", "orderdate", "month"]
    )
    candidates = [direct] if direct else []
    candidates.extend(column for column in columns if column not in candidates)
    best_column: str | None = None
    best_count = 0
    for column in candidates:
        if column is None:
            continue
        try:
            parsed = df.select(
                parse_date_expression(column).is_not_null().sum().alias("parsed")
            ).item()
        except (pl.exceptions.PolarsError, TypeError, ValueError):
            parsed = 0
        if int(parsed or 0) > best_count:
            best_column = column
            best_count = int(parsed or 0)
    return best_column if best_count > 0 else None


def infer_amount_column(df: pl.DataFrame, columns: list[str]) -> str | None:
    """Infer the primary amount column."""

    preferred = first_matching_column(
        columns,
        [
            "salesamount",
            "sales amount",
            "sales",
            "revenue",
            "amount",
            "net sales",
            "value",
        ],
    )
    if preferred and preferred in numeric_columns(df):
        return preferred
    numeric = numeric_columns(df)
    return numeric[0] if numeric else None


def infer_dimensions(
    columns: list[str], date_column: str | None, amount_column: str | None
) -> list[str]:
    """Infer useful reporting dimensions."""

    excluded = {
        date_column,
        amount_column,
        CANONICAL_DATE,
        CANONICAL_PERIOD,
        "Period",
        "Scenario",
    }
    dimension_hints = [
        "category",
        "subcategory",
        "productline",
        "product line",
        "region",
        "country",
        "customer",
        "segment",
        "channel",
        "product",
    ]
    dimensions: list[str] = []
    for hint in dimension_hints:
        column = first_matching_column(columns, [hint])
        if column and column not in excluded and column not in dimensions:
            dimensions.append(column)
    for column in columns:
        if column not in excluded and column not in dimensions:
            dimensions.append(column)
    return dimensions[:5]


def build_recipe(
    input_path: Path,
    df: pl.DataFrame,
    *,
    language: str,
    existing_recipe: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Build a default recipe, merging an existing recipe if supplied."""

    columns, schema = get_schema_and_column_names(df)
    date_column = infer_date_column(df, columns)
    amount_column = infer_amount_column(df, columns)
    dimensions = infer_dimensions(columns, date_column, amount_column)
    recipe: dict[str, Any] = {
        "schema_version": SCHEMA_VERSION,
        "source_file": str(input_path),
        "language": language,
        "mappings": {
            "date_column": date_column,
            "amount_column": amount_column,
            "dimensions": dimensions,
        },
        "options": {
            "currency": "EUR",
            "comparison_basis": "period",
            "period_comparison_mode": "previous_year",
            "period_type": PERIOD_TYPE_TO_DATE,
            "period_grain": PERIOD_GRAIN_MONTH,
            "fiscal_start_month": 1,
            "current_period_label": CURRENT_PERIOD,
            "previous_period_label": PREVIOUS_PERIOD,
            "charts": [
                "year_over_year_column",
                "year_over_year_line",
                "year_over_year_by_period",
                "year_over_year_slope",
                "year_over_year_dot",
                "year_over_year_waterfall",
            ],
            "small_multiples": bool(dimensions),
            "small_multiples_dimension": None,
            "max_small_multiples": 6,
            "max_chart_items": 12,
        },
        "inspection": {
            "columns": columns,
            "schema": schema,
        },
    }
    if existing_recipe:
        recipe["mappings"].update(existing_recipe.get("mappings") or {})
        recipe["options"].update(existing_recipe.get("options") or {})
        if existing_recipe.get("language"):
            recipe["language"] = existing_recipe["language"]
    return validate_recipe(df, recipe)


def _period_derivation_sources(options: dict[str, Any]) -> list[str]:
    """Return raw source columns required by period derivation options."""

    return recipe_cohort_source_dimensions({"options": options})


def _period_derivation_dimension_names(options: dict[str, Any]) -> set[str]:
    """Return dimension names that will be created during canonical preparation."""

    return recipe_cohort_dimension_names({"options": options})


def _unique_texts(values: Iterable[str]) -> list[str]:
    """Return unique strings while preserving order."""

    unique: list[str] = []
    for value in values:
        if value not in unique:
            unique.append(value)
    return unique


def _normalized_period_label(value: Any) -> str:
    """Return a normalized label for mechanical period-comparison checks."""

    return str(value).strip().casefold()


def validate_recipe(df: pl.DataFrame, recipe: dict[str, Any]) -> dict[str, Any]:
    """Validate the recipe against the input frame."""

    columns, _ = get_schema_and_column_names(df)
    mappings = recipe.setdefault("mappings", {})
    options = recipe.setdefault("options", {})
    date_column = mappings.get("date_column")
    amount_column = mappings.get("amount_column")
    if not date_column or date_column not in columns:
        raise ValueError("A valid date column is required for period comparison.")
    if not amount_column or amount_column not in columns:
        raise ValueError("A valid amount column is required for period comparison.")
    if amount_column not in numeric_columns(df):
        raise ValueError(f"Amount column must be numeric: {amount_column}")
    current_period_label = options.get("current_period_label") or CURRENT_PERIOD
    previous_period_label = options.get("previous_period_label") or PREVIOUS_PERIOD
    # This deterministic guard is justified because same-label comparisons are
    # mechanically invalid: AC-vs-AC, PY-vs-PY, etc. have no contrast.
    if _normalized_period_label(current_period_label) == _normalized_period_label(
        previous_period_label
    ):
        raise ValueError(
            "Period comparison requires distinct current and previous period "
            f"labels; both were {current_period_label!r}."
        )
    for source_dimension in _period_derivation_sources(options):
        if source_dimension not in columns:
            raise ValueError(
                "Period derivation source column is not present in input: "
                f"{source_dimension}"
            )
        if source_dimension in {date_column, amount_column}:
            raise ValueError(
                "Period derivation source column must be a business dimension: "
                f"{source_dimension}"
            )
    derived_dimension_names = _period_derivation_dimension_names(options)
    dimensions = [
        dimension
        for dimension in mappings.get("dimensions") or []
        if (dimension in columns or dimension in derived_dimension_names)
        and dimension not in {date_column, amount_column}
    ]
    mappings["dimensions"] = dimensions
    contract_source = dict(options)
    if (
        "period_type" not in contract_source
        and contract_source.get("period_comparison_mode") == "previous_year"
    ):
        contract_source.pop("period_comparison_mode", None)
    normalized_period_contract = period_contract_options(
        contract_source,
        default_type=PERIOD_TYPE_TO_DATE,
        default_grain=PERIOD_GRAIN_MONTH,
    )
    options["currency"] = options.get("currency") or "EUR"
    options["period_type"] = normalized_period_contract["period_type"]
    options["period_grain"] = normalized_period_contract["period_grain"]
    options["fiscal_start_month"] = normalized_period_contract["fiscal_start_month"]
    options["small_multiples"] = bool(options.get("small_multiples", bool(dimensions)))
    return recipe


def inspect_period_comparison_inputs(
    input_path: Path,
    output_dir: Path,
    *,
    language: str = "en",
) -> InspectionResult:
    """Inspect inputs and write suggested period-comparison recipe files."""

    df = read_table(input_path)
    recipe = build_recipe(input_path, df, language=language)
    columns, schema = get_schema_and_column_names(df)
    payload = {
        "schema_version": SCHEMA_VERSION,
        "input_file": str(input_path),
        "row_count": df.height,
        "column_count": df.width,
        "columns": columns,
        "schema": schema,
        "available_analysis_context": available_analysis_context(df),
        "suggested_mappings": recipe["mappings"],
        "suggested_options": recipe["options"],
        "warnings": [],
    }
    output_dir.mkdir(parents=True, exist_ok=True)
    write_json(output_dir / "inspection.json", payload)
    write_json(output_dir / "suggested_recipe.json", recipe)
    return InspectionResult(payload=payload, recipe=recipe, output_dir=output_dir)


def _ensure_legacy_import_path() -> None:
    """Prioritize shared plugin modules in dev, otherwise plugin vendored modules."""

    legacy_parent = (
        SHARED_VENDOR_ROOT
        if (SHARED_VENDOR_ROOT / "modules" / "__init__.py").exists()
        else VENDOR_ROOT
    )
    legacy = str(legacy_parent)
    if legacy in sys.path:
        sys.path.remove(legacy)
    sys.path.insert(0, legacy)
    module_root = (legacy_parent / "modules").resolve()
    for name, module in list(sys.modules.items()):
        if name == "modules" or name.startswith("modules."):
            module_file = getattr(module, "__file__", None)
            if not module_file or not Path(module_file).resolve().is_relative_to(
                module_root
            ):
                del sys.modules[name]


def cleanup_legacy_imports() -> None:
    """Remove shared/vendored ``modules`` imports loaded from this plugin."""

    module_roots = [
        SHARED_VENDOR_ROOT.resolve(),
        VENDOR_ROOT.resolve(),
    ]
    for name, module in list(sys.modules.items()):
        if name == "modules" or name.startswith("modules."):
            module_file = getattr(module, "__file__", None)
            module_path = Path(module_file).resolve() if module_file else None
            if module_path and any(
                module_path.is_relative_to(root) for root in module_roots
            ):
                del sys.modules[name]
    for vendor_str in (str(SHARED_VENDOR_ROOT), str(VENDOR_ROOT)):
        while vendor_str in sys.path:
            sys.path.remove(vendor_str)


def _collect_safe(lf: pl.LazyFrame, *, prefer_streaming: bool = False) -> pl.DataFrame:
    """Collect a lazy frame with optional streaming fallback."""

    if not prefer_streaming:
        return lf.collect()
    try:
        return lf.collect(engine="streaming")
    except pl.exceptions.PolarsError:
        return lf.collect()


def _date_bounds(df: pl.DataFrame, date_column: str) -> tuple[date, date]:
    """Return min and max parsed dates."""

    parsed = df.select(parse_date_expression(date_column).alias("_date")).drop_nulls()
    if parsed.is_empty():
        raise ValueError(f"Could not parse dates in column: {date_column}")
    min_date = parsed.select(pl.col("_date").min()).item()
    max_date = parsed.select(pl.col("_date").max()).item()
    return min_date, max_date


def _add_months(value: date, months: int) -> date:
    """Shift a date by whole months, preserving the closest valid day."""

    month_index = value.month - 1 + months
    year = value.year + month_index // 12
    month = month_index % 12 + 1
    day = min(value.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


def _month_end(year: int, month: int) -> date:
    return date(year, month, calendar.monthrange(year, month)[1])


def _positive_int_option(options: dict[str, Any], key: str, default: int) -> int:
    try:
        value = int(options.get(key, default))
    except (TypeError, ValueError):
        value = default
    return value if value > 0 else default


def _rolling_window_length_options(
    options: dict[str, Any], period_grain: str
) -> dict[str, int]:
    """Return a rolling window length from explicit options or period grain.

    The grain mapping is deterministic because these boundaries are mechanical:
    week=7 days, month=1 month, quarter=3 months, year=12 months.
    """

    if options.get("rolling_window_days") is not None:
        return {"days": _positive_int_option(options, "rolling_window_days", 7)}
    if options.get("rolling_window_months") is not None:
        return {"months": _positive_int_option(options, "rolling_window_months", 12)}
    if period_grain == PERIOD_GRAIN_WEEK:
        return {"days": 7}
    if period_grain == PERIOD_GRAIN_MONTH:
        return {"months": 1}
    if period_grain == PERIOD_GRAIN_QUARTER:
        return {"months": 3}
    if period_grain == PERIOD_GRAIN_YEAR:
        return {"months": 12}
    return {"months": 12}


def _calendar_period_comparison_window(
    max_date: date, period_grain: str
) -> tuple[date, date, date, date]:
    """Return latest calendar period and same prior-year period bounds."""

    if period_grain == PERIOD_GRAIN_YEAR:
        current_start = date(max_date.year, 1, 1)
        current_end = date(max_date.year, 12, 31)
    elif period_grain == PERIOD_GRAIN_QUARTER:
        start_month = ((max_date.month - 1) // 3) * 3 + 1
        current_start = date(max_date.year, start_month, 1)
        current_end = _add_months(current_start, 3) - timedelta(days=1)
    elif period_grain == PERIOD_GRAIN_WEEK:
        current_start = max_date - timedelta(days=max_date.weekday())
        current_end = current_start + timedelta(days=6)
        iso = max_date.isocalendar()
        try:
            previous_start = date.fromisocalendar(iso.year - 1, iso.week, 1)
        except ValueError:
            previous_start = current_start - timedelta(weeks=52)
        previous_end = previous_start + timedelta(days=6)
        return current_start, current_end, previous_start, previous_end
    else:
        current_start = date(max_date.year, max_date.month, 1)
        current_end = _month_end(max_date.year, max_date.month)
    previous_start = _add_months(current_start, -12)
    previous_end = _add_months(current_end, -12)
    return current_start, current_end, previous_start, previous_end


def _rolling_period_comparison_window(
    max_date: date, options: dict[str, Any], period_grain: str
) -> tuple[date, date, date, date, dict[str, int]]:
    """Return current rolling window and prior-year comparison bounds."""

    window_length = _rolling_window_length_options(options, period_grain)
    if "days" in window_length:
        window_days = window_length["days"]
        current_start = max_date - timedelta(days=window_days - 1)
        previous_end = _add_months(max_date, -12)
        previous_start = previous_end - timedelta(days=window_days - 1)
        return (
            current_start,
            max_date,
            previous_start,
            previous_end,
            {
                "rolling_window_days": window_days,
            },
        )
    window_months = window_length["months"]
    current_start = _add_months(max_date, -(window_months - 1)).replace(day=1)
    previous_start = _add_months(current_start, -12)
    previous_end = _add_months(max_date, -12)
    return (
        current_start,
        max_date,
        previous_start,
        previous_end,
        {
            "rolling_window_months": window_months,
        },
    )


def _period_filter_expr(date_column: str, year: int, cutoff_month: int) -> pl.Expr:
    return (pl.col(date_column).dt.year() == year) & (
        pl.col(date_column).dt.month() <= cutoff_month
    )


def prepare_canonical_frame(
    df: pl.DataFrame, recipe: dict[str, Any]
) -> tuple[pl.DataFrame, dict[str, Any]]:
    """Build the AC/PY canonical frame consumed by legacy chart preparation."""

    mappings = recipe["mappings"]
    options = recipe.setdefault("options", {})
    date_column = str(mappings["date_column"])
    amount_column = str(mappings["amount_column"])
    dimensions = [str(item) for item in mappings.get("dimensions") or []]
    raw_columns, _schema = get_schema_and_column_names(df)
    derivation_sources = _period_derivation_sources(options)
    base_dimensions = _unique_texts(
        [
            *[dimension for dimension in dimensions if dimension in raw_columns],
            *derivation_sources,
        ]
    )
    min_date, max_date = _date_bounds(df, date_column)
    contract = period_contract_options(
        options,
        default_type=PERIOD_TYPE_TO_DATE,
        default_grain=PERIOD_GRAIN_MONTH,
    )
    period_type = str(contract["period_type"])
    period_grain = str(contract["period_grain"])
    fiscal_start_month = int(contract["fiscal_start_month"])
    current_year = max_date.year
    previous_year = current_year - 1
    cutoff_month = max_date.month
    parsed_date = pl.col(CANONICAL_DATE)
    if period_type == PERIOD_TYPE_FISCAL:
        fiscal_year = (
            max_date.year if max_date.month >= fiscal_start_month else max_date.year - 1
        )
        current_start = date(fiscal_year, fiscal_start_month, 1)
        current_end = max_date
        previous_start = _add_months(current_start, -12)
        previous_end = _add_months(max_date, -12)
        period_label_expr = (
            pl.when((parsed_date >= current_start) & (parsed_date <= max_date))
            .then(pl.lit(CURRENT_PERIOD))
            .when((parsed_date >= previous_start) & (parsed_date <= previous_end))
            .then(pl.lit(PREVIOUS_PERIOD))
            .otherwise(pl.lit(None))
        )
    elif period_type == PERIOD_TYPE_CALENDAR:
        (
            current_start,
            current_end,
            previous_start,
            previous_end,
        ) = _calendar_period_comparison_window(max_date, period_grain)
        period_label_expr = (
            pl.when((parsed_date >= current_start) & (parsed_date <= current_end))
            .then(pl.lit(CURRENT_PERIOD))
            .when((parsed_date >= previous_start) & (parsed_date <= previous_end))
            .then(pl.lit(PREVIOUS_PERIOD))
            .otherwise(pl.lit(None))
        )
    elif period_type == PERIOD_TYPE_ROLLING:
        (
            current_start,
            current_end,
            previous_start,
            previous_end,
            rolling_options,
        ) = _rolling_period_comparison_window(max_date, options, period_grain)
        period_label_expr = (
            pl.when((parsed_date >= current_start) & (parsed_date <= current_end))
            .then(pl.lit(CURRENT_PERIOD))
            .when((parsed_date >= previous_start) & (parsed_date <= previous_end))
            .then(pl.lit(PREVIOUS_PERIOD))
            .otherwise(pl.lit(None))
        )
        options.update(rolling_options)
    else:
        current_start = date(current_year, 1, 1)
        current_end = max_date
        previous_start = date(previous_year, 1, 1)
        previous_end = _add_months(max_date, -12)
        period_label_expr = (
            pl.when(_period_filter_expr(CANONICAL_DATE, current_year, cutoff_month))
            .then(pl.lit(CURRENT_PERIOD))
            .when(_period_filter_expr(CANONICAL_DATE, previous_year, cutoff_month))
            .then(pl.lit(PREVIOUS_PERIOD))
            .otherwise(pl.lit(None))
        )
    prepared = (
        df.with_columns(parse_date_expression(date_column).alias(CANONICAL_DATE))
        .filter(pl.col(CANONICAL_DATE).is_not_null())
        .with_columns(period_label_expr.alias(CANONICAL_PERIOD))
        .filter(pl.col(CANONICAL_PERIOD).is_not_null())
        .select(
            [
                CANONICAL_DATE,
                CANONICAL_PERIOD,
                pl.col(amount_column).cast(pl.Float64).alias(amount_column),
                *[
                    pl.col(dimension).cast(pl.Utf8).alias(dimension)
                    for dimension in base_dimensions
                ],
            ]
        )
    )
    prepared, derivation_audit = apply_recipe_cohorts(
        prepared,
        recipe,
        period_column=CANONICAL_PERIOD,
        value_column=amount_column,
        current_period=CURRENT_PERIOD,
        previous_period=PREVIOUS_PERIOD,
    )
    prepared_columns, _prepared_schema = get_schema_and_column_names(prepared)
    final_select = [
        CANONICAL_DATE,
        CANONICAL_PERIOD,
        amount_column,
        *[dimension for dimension in dimensions if dimension in prepared_columns],
    ]
    prepared = prepared.select(final_select)
    if prepared.filter(pl.col(CANONICAL_PERIOD) == CURRENT_PERIOD).is_empty():
        raise ValueError(
            "No rows found for current year/period window "
            f"{current_start.isoformat()}..{current_end.isoformat()}."
        )
    if prepared.filter(pl.col(CANONICAL_PERIOD) == PREVIOUS_PERIOD).is_empty():
        raise ValueError(
            "No rows found for previous year/period window "
            f"{previous_start.isoformat()}..{previous_end.isoformat()}."
        )
    period_window = {
        "current": {
            "label": CURRENT_PERIOD,
            "year": current_start.year,
            "start_date": str(current_start),
            "end_date": str(current_end),
            "month_cutoff": cutoff_month,
        },
        "previous": {
            "label": PREVIOUS_PERIOD,
            "year": previous_start.year,
            "start_date": str(previous_start),
            "end_date": str(previous_end),
            "month_cutoff": cutoff_month,
        },
        "period_type": period_type,
        "period_grain": period_grain,
        "fiscal_start_month": fiscal_start_month,
        "source_min_date": str(min_date),
        "source_max_date": str(max_date),
    }
    options["period_window"] = period_window
    options["period_derivation_audit"] = derivation_audit
    return prepared, period_window


def legacy_period_monthly_table(
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
) -> tuple[pl.DataFrame, dict[str, Any]]:
    """Prepare the monthly AC/PY table through legacy multitier-column preparation."""

    _ensure_legacy_import_path()
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        from modules.data.misc_charts_data_prep import (
            prepare_data_for_multitier_column_plot,
        )
        from modules.utilities.config import get_naming_params

        naming = get_naming_params()
        metric = str(recipe["mappings"]["amount_column"])
        options = recipe.get("options") or {}
        window = options.get("period_window") or {}
        current = window.get("current") or {}
        previous = window.get("previous") or {}
        most_recent = date(
            int(current.get("year")),
            int(current.get("month_cutoff")),
            28,
        )
        least_recent = date(
            int(previous.get("year")),
            1,
            1,
        )
        chart_dict = {
            naming["chosenChart"]: naming["multitierColumnChart"],
            naming["selectedPeriods"]: [CURRENT_PERIOD, PREVIOUS_PERIOD],
            naming["plotSmallMultiplesOtherCharts"]: False,
        }
        param_dict = {
            naming["columnHash"]: {},
            naming["mostRecentDate"]: most_recent,
            naming["leastRecentDate"]: least_recent,
            naming["periodLengthInMonths"]: 12,
        }
        table = prepare_data_for_multitier_column_plot(
            canonical.lazy(),
            None,
            metric,
            chart_dict,
            param_dict,
        )
    monthly = _collect_safe(table.lazy() if isinstance(table, pl.DataFrame) else table)
    keep_columns = [
        column
        for column in [
            CANONICAL_DATE,
            CURRENT_PERIOD,
            PREVIOUS_PERIOD,
            "difference in value",
            "difference in %",
        ]
        if column in monthly.columns
    ]
    monthly = monthly.select(keep_columns).filter(pl.col(CANONICAL_DATE).is_not_null())
    monthly = monthly.filter(pl.col(CANONICAL_DATE).is_in(MONTH_ORDER))
    monthly = monthly.with_columns(
        pl.col(CANONICAL_DATE)
        .replace({label: index + 1 for index, label in enumerate(MONTH_ORDER)})
        .cast(pl.Int64)
        .alias("_month_index")
    ).sort("_month_index")
    audit = {
        "status": "written",
        "source_function": "modules.data.misc_charts_data_prep.prepare_data_for_multitier_column_plot",
        "row_count": monthly.height,
    }
    return monthly, audit


def period_totals(canonical: pl.DataFrame, metric: str) -> dict[str, float]:
    """Return AC/PY totals and delta."""

    grouped = canonical.group_by(CANONICAL_PERIOD).agg(
        pl.col(metric).sum().alias("value")
    )
    values = {
        row[CANONICAL_PERIOD]: float(row["value"] or 0.0) for row in grouped.to_dicts()
    }
    ac_value = values.get(CURRENT_PERIOD, 0.0)
    py_value = values.get(PREVIOUS_PERIOD, 0.0)
    return {
        "current": ac_value,
        "previous": py_value,
        "delta": ac_value - py_value,
        "delta_percent": (
            ((ac_value - py_value) / py_value * 100)
            if abs(py_value) > TOLERANCE
            else None
        ),
    }


def monthly_context(monthly: pl.DataFrame) -> list[dict[str, Any]]:
    """Return JSON-friendly month rows."""

    rows: list[dict[str, Any]] = []
    for row in monthly.to_dicts():
        ac_value = float(row.get(CURRENT_PERIOD) or 0.0)
        py_value = float(row.get(PREVIOUS_PERIOD) or 0.0)
        delta = ac_value - py_value
        rows.append(
            {
                "month": row[CANONICAL_DATE],
                "month_index": int(row["_month_index"]),
                "current_amount": ac_value,
                "previous_amount": py_value,
                "delta": delta,
                "delta_percent": (
                    (delta / py_value * 100) if abs(py_value) > TOLERANCE else None
                ),
            }
        )
    return rows


def by_period_table(canonical: pl.DataFrame, metric: str) -> pl.DataFrame:
    """Build year/semester/quarter/month comparison rows."""

    max_date = canonical.select(pl.col(CANONICAL_DATE).max()).item()
    windows = [
        ("Last 12M", 12),
        ("Last 6M", 6),
        ("Last 3M", 3),
        ("Last 1M", 1),
    ]
    rows: list[dict[str, Any]] = []
    for label, months in windows:
        cutoff = max_date.month
        min_month = max(1, cutoff - months + 1)
        subset = canonical.filter(
            (pl.col(CANONICAL_DATE).dt.month() >= min_month)
            & (pl.col(CANONICAL_DATE).dt.month() <= cutoff)
        )
        totals = period_totals(subset, metric)
        rows.append({"window": label, "months": months, **totals})
    return pl.DataFrame(rows)


def select_small_multiples_dimension(
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
) -> dict[str, Any]:
    """Choose a dimension for small multiples."""

    dimensions = [str(item) for item in recipe["mappings"].get("dimensions") or []]
    metric = str(recipe["mappings"]["amount_column"])
    max_panels = int((recipe.get("options") or {}).get("max_small_multiples") or 6)
    candidates: list[dict[str, Any]] = []
    for dimension in dimensions:
        cardinality = canonical.select(pl.col(dimension).n_unique()).item()
        if cardinality is None or int(cardinality) < 2:
            continue
        if int(cardinality) > 30:
            continue
        grouped = (
            canonical.group_by([dimension, CANONICAL_PERIOD])
            .agg(pl.col(metric).sum().alias("value"))
            .pivot(
                index=dimension,
                on=CANONICAL_PERIOD,
                values="value",
                aggregate_function="sum",
            )
            .fill_null(0)
        )
        if CURRENT_PERIOD in grouped.columns and PREVIOUS_PERIOD in grouped.columns:
            score = grouped.select(
                (pl.col(CURRENT_PERIOD) - pl.col(PREVIOUS_PERIOD)).abs().sum()
            ).item()
        else:
            score = 0.0
        candidates.append(
            {
                "dimension": dimension,
                "cardinality": int(cardinality),
                "total_abs_delta": float(score or 0.0),
                "readability_rank": _small_multiples_readability_rank(
                    dimension,
                    int(cardinality),
                    max_panels,
                ),
                "label_quality_rank": _dimension_value_label_quality_rank(
                    canonical,
                    dimension,
                ),
            }
        )
    if not candidates:
        return {"status": "not_selected", "dimension": None, "candidates": []}
    candidates.sort(
        key=lambda item: (
            item["readability_rank"],
            item["label_quality_rank"],
            _dimension_semantic_rank(item["dimension"]),
            -item["total_abs_delta"],
            item["dimension"],
        )
    )
    return {
        "status": "selected_ranked_candidate",
        "dimension": candidates[0]["dimension"],
        "candidates": candidates,
    }


def _small_multiples_readability_rank(
    _dimension: str,
    cardinality: int,
    max_panels: int,
) -> int:
    """Prefer dimensions that can be shown without an unreadable panel grid."""

    if 3 <= cardinality <= max_panels:
        return 0
    if 2 <= cardinality <= max_panels:
        return 1
    if cardinality <= max_panels + 2:
        return 2
    return 3


def _dimension_semantic_rank(dimension: str) -> int:
    """Prefer familiar commercial hierarchy dimensions for default panels."""

    compact = compact_name(dimension)
    priorities = [
        ("productline", 0),
        ("producttype", 1),
        ("subcategory", 2),
        ("category", 3),
        ("country", 4),
        ("region", 5),
        ("customer", 6),
    ]
    for needle, rank in priorities:
        if needle in compact:
            return rank
    return 20


def _dimension_value_label_quality_rank(
    canonical: pl.DataFrame,
    dimension: str,
) -> int:
    """Prefer human-readable labels over one-letter or coded labels."""

    max_label_length = canonical.select(
        pl.col(dimension).cast(pl.Utf8).str.len_chars().max()
    ).item()
    if max_label_length is None:
        return 1
    return 0 if int(max_label_length) > 2 else 1


def _format_millions(value: float) -> str:
    """Format values compactly."""

    abs_value = abs(value)
    sign = "-" if value < 0 else ""
    if abs_value >= 1_000_000:
        return f"{sign}{abs_value / 1_000_000:.1f}M"
    if abs_value >= 1_000:
        return f"{sign}{abs_value / 1_000:.1f}K"
    return f"{value:.0f}"


def _table_rows_from_frame(
    frame: pl.DataFrame,
    *,
    label_column: str,
    baseline_column: str,
    comparison_column: str,
) -> list[dict[str, Any]]:
    """Return normalized rows for two-scenario reporting tables."""

    rows: list[dict[str, Any]] = []
    for row in frame.to_dicts():
        baseline = float(row.get(baseline_column) or 0.0)
        comparison = float(row.get(comparison_column) or 0.0)
        absolute_variance = comparison - baseline
        relative_variance = (
            (absolute_variance / baseline * 100) if abs(baseline) > TOLERANCE else None
        )
        rows.append(
            {
                "row_label": str(row.get(label_column) or ""),
                "baseline_value": baseline,
                "comparison_value": comparison,
                "absolute_variance": absolute_variance,
                "relative_variance": relative_variance,
            }
        )
    return rows


def _write_table_rows_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    """Write normalized reporting-table rows to CSV."""

    pl.DataFrame(rows).write_csv(path)


def _table_value_scale(rows: list[dict[str, Any]]) -> tuple[float, str]:
    """Return a compact display scale and label for table values."""

    values = [
        abs(float(row.get(key) or 0.0))
        for row in rows
        for key in ("baseline_value", "comparison_value", "absolute_variance")
    ]
    max_value = max(values or [0.0])
    if max_value >= 1_000_000:
        return 1_000_000.0, "m"
    if max_value >= 1_000:
        return 1_000.0, "k"
    return 1.0, "units"


def _reporting_metric_label(recipe: dict[str, Any], metric: str) -> str:
    """Return a business-facing metric label for a raw source column."""

    options = recipe.get("options") or {}
    for key in (
        "reporting_metric_label",
        "metric_label",
        "measure_label",
        "value_label",
    ):
        label = str(options.get(key) or "").strip()
        if label:
            return label

    normalized = "".join(ch for ch in metric.casefold() if ch.isalnum())
    sales_names = {
        "amount",
        "sales",
        "salesamount",
        "salesvalue",
        "netsales",
        "netrevenue",
        "revenue",
        "turnover",
        "valuelc",
        "valueusd",
        "valueeur",
    }
    unit_names = {"unit", "units", "quantity", "qty", "volume"}
    if normalized in sales_names or (
        normalized.startswith("value")
        and any(token in normalized for token in ("lc", "usd", "eur", "gbp"))
    ):
        return "Sales"
    if normalized in unit_names:
        return "Units"
    return metric.replace("_", " ").replace("-", " ").strip().title() or "Value"


def _format_scaled_value(value: float, scale: float, *, signed: bool = False) -> str:
    """Format table values with tabular, compact notation."""

    scaled = value / scale
    prefix = "+" if signed and scaled > 0 else ""
    if scale == 1.0:
        return f"{prefix}{scaled:,.0f}"
    return f"{prefix}{scaled:,.1f}"


def _format_relative_percent(value: float | None) -> str:
    """Format a signed relative variance."""

    if value is None or math.isnan(value):
        return "n/a"
    return f"{value:+.1f}%"


def _variance_marker_html(
    *,
    value: float | None,
    label: str,
    max_abs_value: float,
    marker_style: str = "bar",
) -> str:
    """Return centered positive/negative variance marker markup."""

    if value is None or math.isnan(float(value)):
        return (
            '<span class="variance-marker">'
            '<span class="variance-lane negative-lane"></span>'
            '<span class="variance-axis"></span>'
            '<span class="variance-lane positive-lane"></span>'
            '<span class="variance-label muted">n/a</span>'
            "</span>"
        )
    value = float(value)
    width = 0.0
    if max_abs_value > TOLERANCE:
        width = min(100.0, abs(value) / max_abs_value * 100)
    sign_class = (
        "positive"
        if value > TOLERANCE
        else "negative" if value < -TOLERANCE else "neutral"
    )
    negative_width = width if sign_class == "negative" else 0.0
    positive_width = width if sign_class == "positive" else 0.0
    if marker_style == "pin":
        negative_pin = (
            f'<span class="variance-pin-line negative" style="width:{negative_width:.1f}%">'
            '<span class="variance-pin"></span></span>'
            if sign_class == "negative"
            else ""
        )
        positive_pin = (
            f'<span class="variance-pin-line positive" style="width:{positive_width:.1f}%">'
            '<span class="variance-pin"></span></span>'
            if sign_class == "positive"
            else ""
        )
        neutral_pin = (
            '<span class="variance-zero-pin"><span class="variance-pin"></span></span>'
            if sign_class == "neutral"
            else ""
        )
        return (
            '<span class="variance-marker">'
            f'<span class="variance-lane negative-lane">{negative_pin}</span>'
            f'<span class="variance-axis">{neutral_pin}</span>'
            f'<span class="variance-lane positive-lane">{positive_pin}</span>'
            f'<span class="variance-label {sign_class}">{escape(label)}</span>'
            "</span>"
        )
    return (
        '<span class="variance-marker">'
        '<span class="variance-lane negative-lane">'
        f'<span class="variance-bar negative" style="width:{negative_width:.1f}%"></span>'
        "</span>"
        '<span class="variance-axis"></span>'
        '<span class="variance-lane positive-lane">'
        f'<span class="variance-bar positive" style="width:{positive_width:.1f}%"></span>'
        "</span>"
        f'<span class="variance-label {sign_class}">{escape(label)}</span>'
        "</span>"
    )


def _render_reporting_table_html(
    *,
    row_header: str,
    rows: list[dict[str, Any]],
    baseline_label: str,
    comparison_label: str,
    entity_label: str,
    comparison_caption: str,
    metric: str,
    source_label: str,
    output_path: Path,
) -> None:
    """Write a compact scenario and variance table artifact."""

    scale, scale_label = _table_value_scale(rows)
    measure_suffix = "" if scale_label == "units" else f" in {scale_label}"
    max_abs_absolute = max(
        (abs(float(row["absolute_variance"])) for row in rows),
        default=0.0,
    )
    max_abs_relative = max(
        (
            abs(float(row["relative_variance"]))
            for row in rows
            if row.get("relative_variance") is not None
        ),
        default=0.0,
    )
    row_label_width = 92 if row_header == "Period" else 124
    value_width = 88
    variance_width = 172
    table_width = row_label_width + (value_width * 2) + (variance_width * 2)
    page_width = table_width + 56
    table_rows: list[str] = []
    for index, row in enumerate(rows):
        absolute_variance = float(row["absolute_variance"])
        sign_class = (
            "positive"
            if absolute_variance > TOLERANCE
            else "negative" if absolute_variance < -TOLERANCE else "neutral"
        )
        relative_variance = row.get("relative_variance")
        relative_value = None if relative_variance is None else float(relative_variance)
        table_rows.append(
            '<tr class="'
            + ("summary-row" if index == 0 and row_header == "Window" else "detail-row")
            + '">'
            f"<th>{escape(str(row['row_label']))}</th>"
            f"<td>{escape(_format_scaled_value(float(row['baseline_value']), scale))}</td>"
            f"<td>{escape(_format_scaled_value(float(row['comparison_value']), scale))}</td>"
            f'<td class="variance-cell group-start {sign_class}">'
            f"{_variance_marker_html(value=absolute_variance, label=_format_scaled_value(absolute_variance, scale, signed=True), max_abs_value=max_abs_absolute)}</td>"
            f'<td class="variance-cell">'
            f"{_variance_marker_html(value=relative_value, label=_format_relative_percent(relative_value), max_abs_value=max_abs_relative, marker_style='pin')}</td>"
            "</tr>"
        )
    html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{escape(entity_label)} - {escape(metric)}</title>
  <style>
    :root {{
      --ink: #0f1114;
      --muted: #5d6670;
      --rule: #c9cdd1;
      --heavy: #111;
      --pin: #222;
      --soft: #f3f4f5;
      --negative: #e22a1d;
      --positive: #86ad00;
      --neutral: #444;
    }}
    * {{
      box-sizing: border-box;
    }}
    body {{
      margin: 0;
      background: #fff;
      color: var(--ink);
      font: 15px/1.2 Arial, Helvetica, sans-serif;
      font-variant-numeric: tabular-nums;
    }}
    .page {{
      width: {page_width}px;
      padding: 24px 28px 20px;
      background: #fff;
    }}
    .title-block {{
      border-bottom: 2px solid #858585;
      padding-bottom: 24px;
      width: {table_width}px;
    }}
    .title-line {{
      margin: 0;
      color: var(--ink);
      font-size: 12px;
      line-height: 1.18;
    }}
    .title-metric strong {{
      font-weight: 700;
    }}
    table {{
      border-collapse: collapse;
      margin-top: 14px;
      table-layout: fixed;
      width: {table_width}px;
    }}
    col.row-label {{ width: {row_label_width}px; }}
    col.value {{ width: {value_width}px; }}
    col.variance {{ width: {variance_width}px; }}
    thead tr.group th {{
      border-bottom: 0;
      color: var(--ink);
      font-size: 13px;
      font-weight: 700;
      padding: 0 8px 4px;
      text-align: center;
    }}
    thead tr.labels th {{
      border-bottom: 2px solid var(--heavy);
      font-size: 12px;
      font-weight: 700;
      padding: 5px 8px 7px;
      text-align: right;
    }}
    thead tr.labels th:first-child,
    tbody th {{
      text-align: left;
    }}
    tbody th,
    tbody td {{
      border-bottom: 1px solid var(--rule);
      font-size: 15px;
      height: 31px;
      padding: 5px 7px;
      text-align: right;
      vertical-align: middle;
      white-space: nowrap;
    }}
    tbody th {{
      font-weight: 500;
      text-align: left;
    }}
    tbody tr.summary-row th,
    tbody tr.summary-row td {{
      border-bottom: 2px solid var(--heavy);
      font-weight: 700;
    }}
    .group-start {{
      border-left: 2px solid var(--heavy);
    }}
    .negative {{
      color: var(--negative);
    }}
    .positive {{
      color: var(--positive);
    }}
    .neutral {{
      color: var(--neutral);
    }}
    .muted {{
      color: var(--muted);
    }}
    .variance-cell {{
      padding-left: 8px;
    }}
    .variance-marker {{
      display: grid;
      gap: 0;
      grid-template-columns: 50px 2px 50px 56px;
    }}
    .variance-lane {{
      align-items: center;
      display: flex;
      height: 18px;
    }}
    .negative-lane {{
      justify-content: flex-end;
    }}
    .positive-lane {{
      justify-content: flex-start;
    }}
    .variance-axis {{
      background: var(--heavy);
      display: block;
      height: 22px;
      margin-top: -2px;
      width: 2px;
    }}
    .variance-bar {{
      display: block;
      height: 11px;
      min-width: 0;
    }}
    .variance-bar.negative {{
      background: var(--negative);
    }}
    .variance-bar.positive {{
      background: var(--positive);
    }}
    .variance-pin-line {{
      align-items: center;
      display: flex;
      height: 2px;
      min-width: 0;
    }}
    .variance-pin-line.negative {{
      background: var(--negative);
      justify-content: flex-start;
    }}
    .variance-pin-line.positive {{
      background: var(--positive);
      justify-content: flex-end;
    }}
    .variance-pin {{
      background: var(--pin);
      display: block;
      flex: 0 0 auto;
      height: 7px;
      width: 7px;
    }}
    .variance-zero-pin {{
      align-items: center;
      display: flex;
      height: 22px;
      justify-content: center;
      left: -3px;
      position: relative;
      width: 7px;
    }}
    .variance-label {{
      font-size: 13px;
      line-height: 18px;
      padding-left: 8px;
      text-align: right;
    }}
    .source {{
      border-top: 1px solid var(--rule);
      color: var(--muted);
      font-size: 12px;
      margin-top: 14px;
      padding-top: 8px;
    }}
  </style>
</head>
<body>
  <main class="page" data-gallery-screenshot>
    <header class="title-block">
      <p class="title-line">{escape(entity_label)}</p>
      <p class="title-line title-metric"><strong>{escape(metric)}</strong>{escape(measure_suffix)}</p>
      <p class="title-line">{escape(comparison_caption)}</p>
    </header>
    <table>
      <colgroup>
        <col class="row-label">
        <col class="value">
        <col class="value">
        <col class="variance">
        <col class="variance">
      </colgroup>
      <thead>
        <tr class="group">
          <th></th>
          <th colspan="2">Scenario</th>
          <th class="group-start" colspan="2">Change</th>
        </tr>
        <tr class="labels">
          <th>{escape(row_header)}</th>
          <th>{escape(baseline_label)}</th>
          <th>{escape(comparison_label)}</th>
          <th class="group-start">{escape(comparison_label)}-{escape(baseline_label)}</th>
          <th>% change</th>
        </tr>
      </thead>
      <tbody>
        {''.join(table_rows)}
      </tbody>
    </table>
    <div class="source">{escape(source_label)}</div>
  </main>
</body>
</html>
"""
    output_path.write_text(html, encoding="utf-8")


def write_native_reporting_tables(
    monthly: pl.DataFrame,
    by_period: pl.DataFrame,
    recipe: dict[str, Any],
    output_dir: Path,
) -> tuple[list[str], dict[str, Any]]:
    """Write native reporting-table artifacts for period comparison."""

    options = recipe.get("options") or {}
    mappings = recipe.get("mappings") or {}
    metric = str(mappings.get("amount_column") or "Value")
    metric_label = _reporting_metric_label(recipe, metric)
    unit = str(options.get("currency") or "EUR")
    baseline_label = str(options.get("previous_period_label") or PREVIOUS_PERIOD)
    comparison_label = str(options.get("current_period_label") or CURRENT_PERIOD)
    entity_label = reporting_subject_label_from_recipe(recipe) or "Period comparison"
    comparison_caption = reporting_period_line_from_recipe(
        recipe,
        current_label=comparison_label,
        previous_label=baseline_label,
    )
    period_window = options.get("period_window") or {}
    artifacts: list[str] = []
    audit_tables: dict[str, Any] = {}

    table_specs = [
        {
            "table_key": "comparison_table",
            "row_header": "Window",
            "frame": by_period,
            "label_column": "window",
            "baseline_column": "previous",
            "comparison_column": "current",
            "source_table": "period_comparison_by_period.csv",
        },
        {
            "table_key": "time_series_table",
            "row_header": "Period",
            "frame": monthly,
            "label_column": CANONICAL_DATE,
            "baseline_column": PREVIOUS_PERIOD,
            "comparison_column": CURRENT_PERIOD,
            "source_table": "period_comparison_monthly.csv",
        },
    ]
    for spec in table_specs:
        table_key = str(spec["table_key"])
        rows = _table_rows_from_frame(
            spec["frame"],
            label_column=str(spec["label_column"]),
            baseline_column=str(spec["baseline_column"]),
            comparison_column=str(spec["comparison_column"]),
        )
        csv_path = output_dir / f"{table_key}_chart_data.csv"
        html_path = output_dir / f"{table_key}.html"
        context_path = output_dir / f"{table_key}_chart_context.json"
        _write_table_rows_csv(csv_path, rows)
        _scale, scale_label = _table_value_scale(rows)
        measure_suffix = "" if scale_label == "units" else f" in {scale_label}"
        chart_title_lines = [
            entity_label,
            f"{metric_label}{measure_suffix}",
            comparison_caption,
        ]
        write_json(
            context_path,
            {
                "schema_version": SCHEMA_VERSION,
                "object_type": "table",
                "table_key": table_key,
                "capability_id": NATIVE_TABLE_CAPABILITIES[table_key],
                "metric": metric,
                "metric_label": metric_label,
                "unit": unit,
                "dimensions": mappings.get("dimensions") or [],
                "selected_periods": [baseline_label, comparison_label],
                "chart_title_lines": chart_title_lines,
                "title_contract": {
                    "who": chart_title_lines[0],
                    "what": chart_title_lines[1],
                    "when": chart_title_lines[2],
                },
                "source_table": spec["source_table"],
                "table_rows": len(rows),
                "period_window": period_window,
                "comparison": {
                    "baseline_period": baseline_label,
                    "comparison_period": comparison_label,
                    "period_window": period_window,
                },
            },
        )
        _render_reporting_table_html(
            row_header=str(spec["row_header"]),
            rows=rows,
            baseline_label=baseline_label,
            comparison_label=comparison_label,
            entity_label=entity_label,
            comparison_caption=comparison_caption,
            metric=metric_label,
            source_label=f"Source: {spec['source_table']}",
            output_path=html_path,
        )
        artifacts.extend([str(csv_path), str(html_path), str(context_path)])
        audit_tables[table_key] = {
            "status": "written",
            "row_count": len(rows),
            "data": csv_path.name,
            "html": html_path.name,
            "context": context_path.name,
        }
    return artifacts, {"status": "written", "tables": audit_tables}


def write_column_chart(
    _monthly: pl.DataFrame,
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    output_dir: Path,
    *,
    render: bool = True,
) -> tuple[list[str], dict[str, Any]]:
    """Write year-over-year column chart."""

    export = write_legacy_multitier_column_chart(
        canonical,
        recipe,
        output_dir,
        artifact_name="year_over_year_column.png",
        render=render,
    )
    return export.paths, export.audit


def _enabled_chart_names(recipe: dict[str, Any]) -> set[str]:
    """Return enabled chart names from the recipe, preserving default richness."""

    options = recipe.get("options") or {}
    if "charts" not in options:
        return {
            "year_over_year_column",
            "year_over_year_line",
            "year_over_year_by_period",
            "year_over_year_slope",
            "year_over_year_dot",
            "year_over_year_waterfall",
        }
    configured = options.get("charts")
    if not isinstance(configured, list):
        return set()
    return {str(name) for name in configured}


def _failed_legacy_chart_audit(chart_name: str, exc: BaseException) -> dict[str, Any]:
    """Return an auditable failure record for an isolated legacy chart failure."""

    LOGGER.warning("Legacy period chart failed: %s: %s", chart_name, exc)
    return {
        "status": "failed",
        "error_type": type(exc).__name__,
        "error": str(exc),
    }


def write_line_chart(
    _monthly: pl.DataFrame,
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    output_dir: Path,
    *,
    render: bool = True,
) -> tuple[list[str], dict[str, Any]]:
    """Write year-over-year line chart."""

    export = write_legacy_actual_vs_previous_year_chart(
        canonical,
        recipe,
        output_dir,
        artifact_name="year_over_year_line.png",
        by_period=False,
        render=render,
    )
    return export.paths, export.audit


def write_by_period_chart(
    _table: pl.DataFrame,
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    output_dir: Path,
    *,
    render: bool = True,
) -> tuple[list[str], dict[str, Any]]:
    """Write by-period comparison chart."""

    export = write_legacy_actual_vs_previous_year_chart(
        canonical,
        recipe,
        output_dir,
        artifact_name="year_over_year_by_period.png",
        by_period=True,
        render=render,
    )
    return export.paths, export.audit


def write_slope_chart(
    _monthly: pl.DataFrame,
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    output_dir: Path,
    *,
    render: bool = True,
) -> tuple[list[str], dict[str, Any]]:
    """Write legacy slope chart for selected period-comparison dimension."""

    export = write_legacy_slope_chart(
        canonical,
        recipe,
        output_dir,
        artifact_name="year_over_year_slope.png",
        render=render,
    )
    return export.paths, export.audit


def write_dot_chart(
    _monthly: pl.DataFrame,
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    output_dir: Path,
    *,
    render: bool = True,
) -> tuple[list[str], dict[str, Any]]:
    """Write legacy dot chart for selected period-comparison dimension."""

    dimensions = [str(item) for item in recipe["mappings"].get("dimensions") or []]
    primary_dimension = dimensions[0] if dimensions else None
    export = write_legacy_dot_chart(
        canonical,
        recipe,
        output_dir,
        artifact_name="year_over_year_dot.png",
        dimension=primary_dimension,
        render=render,
    )
    return export.paths, export.audit


def write_waterfall_chart(
    monthly: pl.DataFrame,
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    output_dir: Path,
    *,
    render: bool = True,
) -> tuple[list[str], dict[str, Any]]:
    """Write year-over-year horizontal waterfall chart."""

    rows = monthly_context(monthly)
    py_total = sum(row["previous_amount"] for row in rows)
    ac_total = sum(row["current_amount"] for row in rows)
    export = write_legacy_horizontal_waterfall_chart(
        canonical,
        recipe,
        output_dir,
        artifact_name="year_over_year_waterfall.png",
        render=render,
    )
    audit = dict(export.audit)
    audit["reconciliation_delta"] = (
        ac_total - py_total - sum(row["delta"] for row in rows)
    )
    return export.paths, audit


def small_multiples_table(
    canonical: pl.DataFrame, recipe: dict[str, Any], selection: dict[str, Any]
) -> pl.DataFrame:
    """Build small-multiples panel data by selected dimension."""

    dimension = selection.get("dimension")
    metric = str(recipe["mappings"]["amount_column"])
    if not dimension:
        return pl.DataFrame()
    grouped = (
        canonical.group_by([dimension, CANONICAL_PERIOD])
        .agg(pl.col(metric).sum().alias("value"))
        .pivot(
            index=dimension,
            on=CANONICAL_PERIOD,
            values="value",
            aggregate_function="sum",
        )
        .fill_null(0)
    )
    if CURRENT_PERIOD not in grouped.columns:
        grouped = grouped.with_columns(pl.lit(0.0).alias(CURRENT_PERIOD))
    if PREVIOUS_PERIOD not in grouped.columns:
        grouped = grouped.with_columns(pl.lit(0.0).alias(PREVIOUS_PERIOD))
    grouped = grouped.with_columns(
        (pl.col(CURRENT_PERIOD) - pl.col(PREVIOUS_PERIOD)).alias("delta")
    )
    ordered_columns = [dimension, PREVIOUS_PERIOD, CURRENT_PERIOD, "delta"]
    return (
        grouped.with_columns(
            pl.max_horizontal(
                pl.col(CURRENT_PERIOD).abs(),
                pl.col(PREVIOUS_PERIOD).abs(),
            ).alias("_panel_size")
        )
        .sort(["_panel_size", dimension], descending=[True, False])
        .select([*ordered_columns, "_panel_size"])
        .drop("_panel_size")
    )


def write_small_multiples_chart(
    panel_table: pl.DataFrame,
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    selection: dict[str, Any],
    output_dir: Path,
    *,
    render: bool = True,
) -> tuple[list[str], dict[str, Any], dict[str, Any]]:
    """Write small-multiples period-comparison chart and context."""

    if panel_table.is_empty() or not selection.get("dimension"):
        return [], {"status": "not_written_no_dimension"}, {}

    max_panels = int((recipe.get("options") or {}).get("max_small_multiples") or 6)
    dimension = str(selection["dimension"])
    rows = panel_table.head(max_panels).to_dicts()
    other = panel_table.slice(max_panels)
    has_other = not other.is_empty()
    if has_other:
        rows.append(
            {
                dimension: "Other",
                CURRENT_PERIOD: other.select(pl.col(CURRENT_PERIOD).sum()).item(),
                PREVIOUS_PERIOD: other.select(pl.col(PREVIOUS_PERIOD).sum()).item(),
                "delta": other.select(pl.col("delta").sum()).item(),
            }
        )
    row_count = len(rows)
    top_values = [str(row[dimension]) for row in rows if row[dimension] != "Other"]
    repeat_values = [str(row[dimension]) for row in rows]
    chart_frame = canonical
    if has_other:
        chart_frame = canonical.with_columns(
            pl.when(pl.col(dimension).cast(pl.Utf8).is_in(top_values))
            .then(pl.col(dimension).cast(pl.Utf8))
            .otherwise(pl.lit("Other"))
            .alias(dimension)
        )
    chart_frame = _ensure_small_multiple_period_rows(
        chart_frame,
        recipe,
        dimension,
        repeat_values,
    )
    dimensions = [str(item) for item in recipe["mappings"].get("dimensions") or []]
    dot_dimension = next(
        (candidate for candidate in dimensions if candidate != dimension),
        dimension,
    )
    exports = {
        "column": write_legacy_multitier_column_chart(
            chart_frame,
            recipe,
            output_dir,
            artifact_name="year_over_year_column_small_multiples.png",
            small_multiples_dimension=dimension,
            repeat_values=repeat_values,
            render=render,
        ),
        "line": write_legacy_actual_vs_previous_year_chart(
            chart_frame,
            recipe,
            output_dir,
            artifact_name="year_over_year_line_small_multiples.png",
            by_period=False,
            small_multiples_dimension=dimension,
            repeat_values=repeat_values,
            render=render,
        ),
        "by_period": write_legacy_actual_vs_previous_year_chart(
            chart_frame,
            recipe,
            output_dir,
            artifact_name="year_over_year_by_period_small_multiples.png",
            by_period=True,
            small_multiples_dimension=dimension,
            repeat_values=repeat_values,
            render=render,
        ),
        "dot": write_legacy_dot_chart(
            chart_frame,
            recipe,
            output_dir,
            artifact_name="year_over_year_dot_small_multiples.png",
            dimension=dot_dimension,
            small_multiples_dimension=dimension,
            repeat_values=repeat_values,
            render=render,
        ),
        "slope": write_legacy_slope_chart(
            chart_frame,
            recipe,
            output_dir,
            artifact_name="year_over_year_slope_small_multiples.png",
            small_multiples_dimension=dimension,
            repeat_values=repeat_values,
            render=render,
        ),
        "waterfall": write_legacy_horizontal_waterfall_chart(
            chart_frame,
            recipe,
            output_dir,
            artifact_name="year_over_year_waterfall_small_multiples.png",
            small_multiples_dimension=dimension,
            repeat_values=repeat_values,
            render=render,
        ),
    }
    chart_paths = [path for export in exports.values() for path in export.paths]
    line_png = output_dir / "year_over_year_line_small_multiples.png"
    line_paths = [Path(path) for path in exports["line"].paths]
    alias_source = line_png if line_png.exists() else next(iter(line_paths), None)
    alias_path = output_dir / (
        f"year_over_year_small_multiples{alias_source.suffix}"
        if alias_source
        else "year_over_year_small_multiples.png"
    )
    if render and alias_source and alias_source.exists():
        shutil.copy2(alias_source, alias_path)
        chart_paths.append(str(alias_path))
    chart_audits = {name: export.audit for name, export in exports.items()}
    line_audit = exports["line"].audit
    source_functions = sorted(
        {
            source
            for export in exports.values()
            for source in export.audit.get("source_functions", [])
        }
    )
    summary_path = output_dir / "year_over_year_small_multiples_summary.csv"
    context_path = output_dir / "year_over_year_small_multiples_context.json"
    panel_table.write_csv(summary_path)
    context = {
        "schema_version": SCHEMA_VERSION,
        "analysis_type": "period_comparison_small_multiples",
        "status": "written" if render else "data_written",
        "dimension": dimension,
        "panel_count": row_count,
        "has_other_panel": has_other,
        "panels": rows,
        "chart_artifact": alias_path.name if alias_path.exists() else None,
        "chart_artifacts": [Path(path).name for path in chart_paths],
        "codex_interpretation_contract": {
            "must_review_when_written": True,
            "required_points": [
                "Explain whether movement is concentrated in the largest panels.",
                "Use shared AC/PY totals; do not infer variance components from this chart.",
            ],
        },
    }
    write_json(context_path, context)
    audit = {
        "status": "written" if render else "data_written",
        "artifact": alias_path.name if render else None,
        "artifacts": [Path(path).name for path in chart_paths],
        "charts": chart_audits,
        "renderer": line_audit.get("renderer"),
        "plotly_export_error": line_audit.get("plotly_export_error"),
        "source_functions": source_functions,
        "dimension": dimension,
        "panel_count": row_count,
        "has_other_panel": has_other,
    }
    return [*chart_paths, str(summary_path), str(context_path)], audit, context


def _ensure_small_multiple_period_rows(
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    dimension: str,
    repeat_values: list[str],
) -> pl.DataFrame:
    """Add zero-value rows so every legacy small-multiple panel has both periods."""

    metric = str(recipe["mappings"]["amount_column"])
    date_by_period = {
        row[CANONICAL_PERIOD]: row[CANONICAL_DATE]
        for row in canonical.group_by(CANONICAL_PERIOD)
        .agg(pl.col(CANONICAL_DATE).min().alias(CANONICAL_DATE))
        .to_dicts()
    }
    filler_rows: list[dict[str, Any]] = []
    for value in repeat_values:
        for period in (PREVIOUS_PERIOD, CURRENT_PERIOD):
            exists = not canonical.filter(
                (pl.col(dimension).cast(pl.Utf8) == value)
                & (pl.col(CANONICAL_PERIOD) == period)
            ).is_empty()
            if exists:
                continue
            filler_rows.append(
                {
                    column: (
                        date_by_period.get(period)
                        if column == CANONICAL_DATE
                        else (
                            period
                            if column == CANONICAL_PERIOD
                            else (
                                0.0
                                if column == metric
                                else value if column == dimension else None
                            )
                        )
                    )
                    for column in canonical.columns
                }
            )
    if not filler_rows:
        return canonical
    filler = pl.DataFrame(filler_rows, schema=canonical.schema)
    return pl.concat([canonical, filler], how="vertical")


def build_summary_markdown(
    recipe: dict[str, Any], totals: dict[str, float], selection: dict[str, Any]
) -> str:
    """Build deterministic markdown summary."""

    window = (recipe.get("options") or {}).get("period_window") or {}
    lines = [
        "# Period Comparison Source Data",
        "",
        f"- Source file: `{recipe.get('source_file')}`",
        f"- Metric: `{recipe['mappings']['amount_column']}`",
        f"- Comparison: `{(window.get('current') or {}).get('year')}` vs `{(window.get('previous') or {}).get('year')}`",
        f"- Previous period total: `{totals['previous']:,.2f}`",
        f"- Current period total: `{totals['current']:,.2f}`",
        f"- Delta: `{totals['delta']:,.2f}`",
        f"- Small multiples dimension: `{selection.get('dimension') or 'not selected'}`",
        "",
        "The charts compare current period and previous-year period totals. They are not a price-volume-mix decomposition.",
        "",
    ]
    return "\n".join(lines)


def write_client_report(
    recipe: dict[str, Any],
    totals: dict[str, float],
    chart_paths: list[str],
    output_dir: Path,
) -> tuple[list[str], dict[str, Any]]:
    """Write Markdown and DOCX client report."""

    md_path = output_dir / "period_comparison_client_report.md"
    docx_path = output_dir / "period_comparison_client_report.docx"
    lines = [
        "# Period Comparison",
        "",
        f"Current period totals are {_format_millions(totals['current'])} versus {_format_millions(totals['previous'])} in the previous-year period.",
        f"The movement is {_format_millions(totals['delta'])}.",
        "",
        "## Source Files",
        "",
        *[
            f"- `{Path(path).name}`"
            for path in chart_paths
            if Path(path).suffix.lower() == ".png"
        ],
        "",
    ]
    md_path.write_text("\n".join(lines), encoding="utf-8")
    try:
        from docx import Document
        from docx.shared import Inches

        document = Document()
        document.add_heading("Period Comparison", level=1)
        document.add_paragraph(
            f"Current period totals are {_format_millions(totals['current'])} versus "
            f"{_format_millions(totals['previous'])} in the previous-year period. "
            f"The movement is {_format_millions(totals['delta'])}."
        )
        for chart_path in chart_paths:
            path = Path(chart_path)
            if path.suffix.lower() == ".png" and path.exists():
                document.add_picture(str(path), width=Inches(6.4))
        document.save(docx_path)
        status = "written"
        error = None
    except (
        ImportError,
        ModuleNotFoundError,
        OSError,
        ValueError,
        zipfile.BadZipFile,
    ) as exc:
        status = "not_written"
        error = str(exc)
    paths = [str(md_path)]
    if docx_path.exists():
        paths.append(str(docx_path))
    return paths, {
        "status": status,
        "markdown": md_path.name,
        "docx": docx_path.name if docx_path.exists() else None,
        "error": error,
    }


def _relative_path(path: Path, base: Path) -> str:
    try:
        return path.relative_to(base).as_posix()
    except ValueError:
        return path.as_posix()


def _normalize_artifact_mode(artifact_mode: str) -> str:
    """Return a supported artifact mode or raise for invalid contract input."""

    normalized = str(artifact_mode or ARTIFACT_MODE_DATA_AND_RENDER).strip().lower()
    if normalized not in ARTIFACT_MODES:
        allowed = ", ".join(sorted(ARTIFACT_MODES))
        raise ValueError(f"Unsupported artifact_mode {artifact_mode!r}; use {allowed}.")
    return normalized


def _load_json_if_exists(path: Path) -> dict[str, Any] | None:
    if not path.exists():
        return None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    return payload if isinstance(payload, dict) else None


def run_period_comparison(
    input_path: Path,
    output_dir: Path,
    recipe_path: Path | None = None,
    *,
    language: str = "en",
    currency: str | None = None,
    artifact_mode: str = ARTIFACT_MODE_DATA_AND_RENDER,
) -> PeriodComparisonRunResult:
    """Run period-comparison charts and write outputs."""

    artifact_mode = _normalize_artifact_mode(artifact_mode)
    try:
        df = read_table(input_path)
        existing_recipe = read_json(recipe_path) if recipe_path else None
        recipe = build_recipe(
            input_path,
            df,
            language=language,
            existing_recipe=existing_recipe,
        )
        recipe = preserve_recipe_filters(recipe, existing_recipe)
        recipe = preserve_recipe_cohorts(recipe, existing_recipe)
        if currency:
            recipe["options"]["currency"] = currency
        recipe = validate_recipe(df, recipe)
        df, filter_audit = apply_recipe_filters(df, recipe)
        recipe.setdefault("options", {})["recipe_filter_audit"] = filter_audit
        output_dir.mkdir(parents=True, exist_ok=True)
        run_intake = write_run_intake(
            output_dir,
            input_path,
            recipe_path=recipe_path,
            recipe=recipe,
            source_row_count=df.height,
        )
        canonical, period_window = prepare_canonical_frame(df, recipe)
        monthly, legacy_monthly_audit = legacy_period_monthly_table(canonical, recipe)
        metric = str(recipe["mappings"]["amount_column"])
        totals = period_totals(canonical, metric)
        selection = select_small_multiples_dimension(canonical, recipe)
        by_period = by_period_table(canonical, metric)
        context = {
            "schema_version": SCHEMA_VERSION,
            "analysis_type": "period_comparison",
            "source_file": str(input_path),
            "artifact_mode": artifact_mode,
            "metric": metric,
            "unit": recipe["options"].get("currency") or "EUR",
            "comparison": period_window,
            "totals": totals,
            "monthly": monthly_context(monthly),
            "recipe_filters": filter_audit,
            "period_derivations": recipe["options"].get("period_derivation_audit"),
            "recipe_cohorts": recipe["options"].get("recipe_cohort_audit"),
            "small_multiples_selection": selection,
            "legacy_functions_used": [
                "modules.data.misc_charts_data_prep.prepare_data_for_multitier_column_plot",
                "modules.charting.chart_helpers.prepare_actual_vs_year_ago_dataframe",
                "modules.data.time_series_data_prep.prepare_data_for_slope_plot",
                "modules.data.waterfall_data_prep.prepare_data_for_horizontal_waterfall_plot",
                "modules.charting.plot_charts.plot_dot_chart",
                "modules.charting.plot_charts.plot_slope_charts",
            ],
            "codex_interpretation_contract": {
                "must_review_when_written": True,
                "required_points": [
                    "Explain whether the current period is above or below previous year.",
                    "Identify the months and panels driving the movement.",
                    "Do not describe this as price-volume-mix or root-cause variance.",
                ],
            },
        }
        artifact_paths: list[str] = []
        chart_audits: dict[str, Any] = {}
        enabled_charts = _enabled_chart_names(recipe)
        requested_small_multiples = bool(
            (recipe.get("options") or {}).get("small_multiples", True)
        )
        render_charts = artifact_mode != ARTIFACT_MODE_DATA_ONLY
        for name, writer_args in (
            (
                "year_over_year_column",
                (write_column_chart, monthly, canonical, recipe, output_dir),
            ),
            (
                "year_over_year_line",
                (write_line_chart, monthly, canonical, recipe, output_dir),
            ),
            (
                "year_over_year_by_period",
                (write_by_period_chart, by_period, canonical, recipe, output_dir),
            ),
            (
                "year_over_year_slope",
                (write_slope_chart, monthly, canonical, recipe, output_dir),
            ),
            (
                "year_over_year_dot",
                (write_dot_chart, monthly, canonical, recipe, output_dir),
            ),
            (
                "year_over_year_waterfall",
                (write_waterfall_chart, monthly, canonical, recipe, output_dir),
            ),
        ):
            if name not in enabled_charts:
                chart_audits[name] = {
                    "status": "skipped",
                    "reason": "not enabled in recipe options.charts",
                }
                continue
            writer = writer_args[0]
            try:
                paths, chart_audit = writer(*writer_args[1:], render=render_charts)
            except LEGACY_CHART_ERROR_TYPES as exc:
                paths = []
                chart_audit = _failed_legacy_chart_audit(name, exc)
            artifact_paths.extend(paths)
            chart_audits[name] = chart_audit
        if requested_small_multiples:
            try:
                small_paths, small_audit, small_context = write_small_multiples_chart(
                    small_multiples_table(canonical, recipe, selection),
                    canonical,
                    recipe,
                    selection,
                    output_dir,
                    render=render_charts,
                )
            except LEGACY_CHART_ERROR_TYPES as exc:
                small_paths = []
                small_context = {}
                small_audit = _failed_legacy_chart_audit("small_multiples", exc)
            artifact_paths.extend(small_paths)
            if small_context:
                context["small_multiples"] = small_context
        else:
            small_audit = {
                "status": "skipped",
                "reason": "not enabled in recipe options.small_multiples",
            }
        monthly.write_csv(output_dir / "period_comparison_monthly.csv")
        by_period.write_csv(output_dir / "period_comparison_by_period.csv")
        table_paths, table_audit = write_native_reporting_tables(
            monthly,
            by_period,
            recipe,
            output_dir,
        )
        artifact_paths.extend(table_paths)
        canonical_path = output_dir / "period_comparison_canonical.csv"
        canonical.write_csv(canonical_path)
        prepared_manifest_path = write_prepared_data_manifest(
            output_dir=output_dir,
            plugin="period-comparison",
            chart_family="period_comparison",
            source_file=input_path,
            prepared_path=canonical_path,
            frame=canonical,
            recipe=recipe,
            preparation_audit={
                "status": "prepared",
                "recipe_filters": filter_audit,
                "recipe_cohorts": recipe["options"].get("recipe_cohort_audit"),
                "period_window": period_window,
                "period_derivations": recipe["options"].get("period_derivation_audit"),
                "legacy_monthly": legacy_monthly_audit,
            },
        )
        artifact_paths.append(str(prepared_manifest_path))
        try:
            monthly.write_excel(output_dir / "period_comparison_results.xlsx")
            xlsx_status = "written"
        except (ImportError, ModuleNotFoundError, OSError, ValueError) as exc:
            xlsx_status = f"not_written: {exc}"
        write_json(output_dir / "period_comparison_context.json", context)
        summary = build_summary_markdown(recipe, totals, selection)
        (output_dir / "period_comparison_summary.md").write_text(
            summary, encoding="utf-8"
        )
        report_paths, report_audit = write_client_report(
            recipe, totals, artifact_paths, output_dir
        )
        artifact_paths.extend(report_paths)
        write_json(output_dir / "used_recipe.json", recipe)
        audit: dict[str, Any] = {
            "schema_version": SCHEMA_VERSION,
            "created_at": utc_now(),
            "input_file": str(input_path),
            "recipe": recipe,
            "legacy_runtime": {
                "monthly_preparation": legacy_monthly_audit,
                "chart_audits": chart_audits,
                "small_multiples": small_audit,
                "native_tables": table_audit,
                "client_report": report_audit,
                "artifact_mode": artifact_mode,
            },
            "checks": {
                "monthly_row_count": monthly.height,
                "canonical_row_count": canonical.height,
                "delta_reconciliation": totals["current"]
                - totals["previous"]
                - totals["delta"],
                "legacy_chart_attempt_count": len(
                    [
                        item
                        for item in chart_audits.values()
                        if item.get("status") != "skipped"
                    ]
                )
                + (0 if small_audit.get("status") == "skipped" else 1),
                "legacy_chart_written_count": sum(
                    1
                    for item in chart_audits.values()
                    if item.get("status") == "written"
                )
                + (1 if small_audit.get("status") == "written" else 0),
                "legacy_chart_data_count": sum(
                    1
                    for item in chart_audits.values()
                    if item.get("status") in {"written", "data_written"}
                )
                + (
                    1 if small_audit.get("status") in {"written", "data_written"} else 0
                ),
            },
            "outputs": {
                "period_comparison_monthly.csv": "written",
                "period_comparison_by_period.csv": "written",
                "period_comparison_canonical.csv": "written",
                "prepared_data_manifest.json": "written",
                "period_comparison_results.xlsx": xlsx_status,
                "period_comparison_context.json": "written",
                "period_comparison_summary.md": "written",
                "used_recipe.json": "written",
            },
        }
        for path in artifact_paths:
            audit["outputs"][Path(path).name] = "written"
        write_json(output_dir / "period_comparison_audit.json", audit)
        audit["outputs"]["period_comparison_audit.json"] = "written"
        review_session = write_review_session_artifacts(
            output_dir,
            input_path,
            run_id=run_intake.run_id,
            run_intake_path=run_intake.path,
            recipe_path=recipe_path,
            recipe=recipe,
            monthly_rows=monthly.to_dicts(),
            by_period_rows=by_period.to_dicts(),
            audit=audit,
        )
        audit["review_session"] = {
            "run_intake_path": review_session.run_intake_path.name,
            "review_payload_path": review_session.review_payload_path.name,
            "ui_decisions_path": review_session.ui_decisions_path.name,
            "final_artifacts_path": review_session.final_artifacts_path.name,
            "review_item_count": review_session.review_item_count,
        }
        for path in (
            review_session.run_intake_path,
            review_session.review_payload_path,
            review_session.ui_decisions_path,
            review_session.final_artifacts_path,
        ):
            audit["outputs"][_relative_path(path, output_dir)] = "written"
        artifact_paths = [
            *artifact_paths,
            str(review_session.run_intake_path),
            str(review_session.review_payload_path),
            str(review_session.ui_decisions_path),
            str(review_session.final_artifacts_path),
        ]
        write_json(output_dir / "period_comparison_audit.json", audit)
        return PeriodComparisonRunResult(
            monthly_frame=monthly,
            audit=audit,
            summary_markdown=summary,
            artifact_paths=artifact_paths,
        )
    finally:
        cleanup_legacy_imports()
