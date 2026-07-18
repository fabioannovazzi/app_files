"""Mix and contribution analysis helpers for the Codex plugin."""

from __future__ import annotations

import argparse
import json
import logging
import re
import sys
import zipfile
from dataclasses import dataclass
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Sequence

import polars as pl
from legacy_mix_charting import (
    LegacyPreparedDataCache,
    cleanup_legacy_imports,
    write_legacy_mix_chart,
)

SCRIPT_DIR = Path(__file__).resolve().parent


def _ensure_local_review_session_import() -> None:
    """Use this plugin's review-session module in multi-plugin test runs."""

    script_dir = str(SCRIPT_DIR)
    if script_dir in sys.path:
        sys.path.remove(script_dir)
    sys.path.insert(0, script_dir)
    module = sys.modules.get("review_session")
    module_file = getattr(module, "__file__", None) if module is not None else None
    if module_file and Path(module_file).resolve().is_relative_to(SCRIPT_DIR.resolve()):
        return
    if module is not None:
        del sys.modules["review_session"]


_ensure_local_review_session_import()
from review_session import write_review_session_artifacts, write_run_intake

PLUGIN_ROOT = Path(__file__).resolve().parents[1]
VENDOR_ROOT = PLUGIN_ROOT / "vendor"
REPO_ROOT = Path(__file__).resolve().parents[3]
SHARED_VENDOR_ROOT = REPO_ROOT / "plugins" / "_shared" / "vendor"


def _ensure_shared_modules_path() -> None:
    """Make shared plugin harness modules importable in repo and ZIP runs."""

    shared_parent = SHARED_VENDOR_ROOT if SHARED_VENDOR_ROOT.exists() else VENDOR_ROOT
    shared_text = str(shared_parent)
    if shared_text in sys.path:
        sys.path.remove(shared_text)
    sys.path.insert(0, shared_text)
    module_root = (shared_parent / "modules").resolve()
    for name, module in list(sys.modules.items()):
        if name == "modules" or name.startswith("modules."):
            module_file = getattr(module, "__file__", None)
            if not module_file or not Path(module_file).resolve().is_relative_to(
                module_root
            ):
                del sys.modules[name]


_ensure_shared_modules_path()
from modules.chart_harness import (  # noqa: E402  # isort: skip
    apply_recipe_filters,
    apply_recipe_cohorts,
    available_analysis_context,
    default_scenario_comparison_pair,
    json_safe,
    normalize_recipe_cohort_contract,
    period_contract_options,
    preserve_recipe_cohorts,
    preserve_recipe_filters,
    recipe_cohort_dimension_names,
    recipe_cohort_period_labels,
    recipe_cohort_source_dimensions,
    reporting_entity_label_from_recipe,
    reporting_subject_label_from_recipe,
    write_prepared_data_manifest,
)
from modules.charting.chart_primitives import (  # noqa: E402  # isort: skip
    FOCUS_ITEM_HIGHLIGHT_MAX_ITEMS,
)

__all__ = [
    "InspectionResult",
    "MixContributionRunResult",
    "add_common_args",
    "cleanup_retired_chart_artifacts",
    "configure_logging",
    "inspect_mix_inputs",
    "run_mix_contribution",
]

LOGGER = logging.getLogger(__name__)
SCHEMA_VERSION = "1.0"
CANONICAL_DATE = "Date"
CANONICAL_PERIOD = "Period"
CURRENT_PERIOD = "AC"
IBCS_SCENARIO_ABBREVIATIONS = {
    "AC": "Actual",
    "PY": "Previous year",
    "PM": "Previous month",
    "PQ": "Previous quarter",
    "PL": "Plan",
    "FC": "Forecast",
}
TOLERANCE = 1e-9
LEGACY_AMOUNT_COLUMN = "Sales"
LEGACY_UNITS_COLUMN = "Units"
LEGACY_UNIT_PRICE_COLUMN = "Unit Price"
LEGACY_MULTIPLIED_COLUMN = f"{LEGACY_UNITS_COLUMN} x {LEGACY_UNIT_PRICE_COLUMN}"
LEGACY_MARGIN_COLUMN = "Margin"
LEGACY_MARGIN_PERCENT_COLUMN = "Margin in %"
LEGACY_SALES_GROWTH_COLUMN = "Sales Growth Rate"
LEGACY_TOTAL_COLUMN_DIMENSION = "Total View"
LEGACY_TOTAL_COLUMN_LABEL = "Total"
PARETO_DEFAULT_SHOW_ONLY = "All"
PARETO_SHOW_ONLY_ALIASES = {
    "all": "All",
    "all_items": "All",
    "all-ranked-items": "All",
    "all_ranked_items": "All",
    "bottom": "Bottom",
    "bottom_items": "Bottom",
    "bottom-ranked-items": "Bottom",
    "bottom_ranked_items": "Bottom",
    "top": "Top",
    "top_items": "Top",
    "top-ranked-items": "Top",
    "top_ranked_items": "Top",
}
AREA_CHART_MAX_NAMED_ITEMS = 6
STACKED_COLUMN_MAX_NAMED_ITEMS = 6
MEKKO_MAX_NAMED_ITEMS = 8
MEKKO_MIN_NAMED_ITEMS = 2
MEKKO_MIN_ITEM_SHARE = 0.02
MEKKO_CUMULATIVE_SHARE = 0.95
STACKED_PARETO_MAX_RANKED_ITEMS = 80
STACKED_BAR_SMALL_MULTIPLES_MIN_PANELS = 2
STACKED_BAR_SMALL_MULTIPLES_MIN_MULTI_ROW_SHARE = 0.5
PERIOD_GRAIN_YEAR = "year"
PERIOD_GRAIN_QUARTER = "quarter"
PERIOD_GRAIN_MONTH = "month"
PERIOD_GRAIN_WEEK = "week"
COHORT_VISIBLE_PERIOD_COUNT = 3
RETIRED_CHART_ARTIFACT_STEMS = (
    "horizontal_waterfall",
    "horizontal_waterfall_small_multiples",
    "stacked_column_with_cagr",
    "stacked_column_with_overlay",
)
CAGR_BLOCKED_METRIC_TERMS = (
    "%",
    "cwd",
    "distribution",
    "growth",
    "percent",
    "price",
    "rate",
    "ratio",
    "share",
)
CAGR_BLOCKED_METRICS = {
    LEGACY_MARGIN_PERCENT_COLUMN,
    LEGACY_MULTIPLIED_COLUMN,
    LEGACY_SALES_GROWTH_COLUMN,
    LEGACY_UNIT_PRICE_COLUMN,
}
CHART_NAME_ALIASES = {
    "column": "column_total",
    "column_plus_marker": "column_total_with_overlay",
    "column_with_marker": "column_total_with_overlay",
    "simple_column": "column_total",
    "stacked_column_with_cagr": "stacked_column",
}
FOCUS_ITEM_RENDER_ELIGIBLE_CHARTS = {
    "area_absolute",
    "area_share",
    "line",
    "marimekko",
    "marimekko_small_multiples",
    "stacked_bar",
    "stacked_bar_small_multiples",
    "stacked_column",
}
FOCUS_ITEM_UNSUPPORTED_CHARTS = {
    "barmekko",
    "barmekko_small_multiples",
    "line_small_multiples",
    "stacked_column_synthesis",
}
ARTIFACT_MODE_DATA_ONLY = "data_only"
ARTIFACT_MODE_DATA_AND_RENDER = "data_and_render"
ARTIFACT_MODES = {
    ARTIFACT_MODE_DATA_ONLY,
    ARTIFACT_MODE_DATA_AND_RENDER,
}


def _normalize_requested_chart_names(charts: Iterable[Any]) -> list[str]:
    """Map retired chart request names to canonical source producers."""

    normalized: list[str] = []
    seen: set[str] = set()
    for chart in charts:
        name = CHART_NAME_ALIASES.get(str(chart), str(chart))
        if name in seen:
            continue
        seen.add(name)
        normalized.append(name)
    return normalized


@dataclass(frozen=True)
class InspectionResult:
    """Inspection result and suggested recipe paths."""

    payload: dict[str, Any]
    recipe: dict[str, Any]
    output_dir: Path


@dataclass(frozen=True)
class MixContributionRunResult:
    """Mix-contribution run result."""

    canonical_frame: pl.DataFrame
    audit: dict[str, Any]
    summary_markdown: str
    artifact_paths: list[str]


def configure_logging(verbose: bool = False) -> None:
    """Configure command-line logging."""

    logging.basicConfig(
        level=logging.DEBUG if verbose else logging.WARNING,
        format="%(levelname)s:%(name)s:%(message)s",
    )


def add_common_args(parser: argparse.ArgumentParser) -> None:
    """Add common plugin CLI arguments."""

    parser.add_argument("input_file", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--recipe", type=Path)
    parser.add_argument("--language", default="en")
    parser.add_argument("--currency", default=None)
    parser.add_argument(
        "--artifact-mode",
        choices=sorted(ARTIFACT_MODES),
        default=ARTIFACT_MODE_DATA_AND_RENDER,
        help=(
            "Control chart artifact generation: data_only writes chart data/context "
            "without PNG/HTML, and data_and_render keeps legacy behavior."
        ),
    )
    parser.add_argument("--verbose", action="store_true")


def cleanup_retired_chart_artifacts(output_dir: Path) -> dict[str, Any]:
    """Remove stale artifacts for chart specs the plugin no longer generates."""

    removed: list[str] = []
    failed: list[dict[str, str]] = []
    for stem in RETIRED_CHART_ARTIFACT_STEMS:
        for path in sorted(output_dir.glob(f"{stem}*")):
            if not path.is_file():
                continue
            try:
                path.unlink()
                removed.append(_relative_path(path, output_dir))
            except OSError as exc:
                failed.append(
                    {
                        "path": _relative_path(path, output_dir),
                        "error": str(exc),
                    }
                )
    return {
        "retired_chart_artifact_stems": list(RETIRED_CHART_ARTIFACT_STEMS),
        "removed": removed,
        "failed": failed,
    }


def utc_now() -> str:
    """Return an ISO timestamp for audit files."""

    return datetime.now(timezone.utc).isoformat()


def json_safe(value: Any) -> Any:
    """Return JSON-safe values for common analysis objects."""

    if isinstance(value, dict):
        return {str(key): json_safe(item) for key, item in value.items()}
    if isinstance(value, list):
        return [json_safe(item) for item in value]
    if isinstance(value, tuple):
        return [json_safe(item) for item in value]
    if isinstance(value, (Path, datetime, date)):
        return str(value)
    return value


def write_json(path: Path, payload: dict[str, Any]) -> None:
    """Write stable JSON."""

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(json_safe(payload), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def read_json(path: Path) -> dict[str, Any]:
    """Read a JSON object."""

    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"Expected JSON object: {path}")
    return payload


def _collect_csv_scan(path: Path, *, separator: str) -> pl.DataFrame:
    """Read delimited input through a lazy scan and collect once."""

    lf = pl.scan_csv(path, separator=separator, infer_schema_length=10000)
    try:
        return lf.collect(engine="streaming")
    except pl.exceptions.PolarsError:
        return lf.collect()


def read_table(path: Path) -> pl.DataFrame:
    """Read a supported CSV or Excel file."""

    suffix = path.suffix.lower()
    if suffix in {".csv", ".txt", ".tsv", ".psv"}:
        separator = {".tsv": "\t", ".psv": "|"}.get(suffix, ",")
        return _collect_csv_scan(path, separator=separator)
    if suffix in {".xlsx", ".xlsm"}:
        return pl.read_excel(path)
    raise ValueError(
        f"Unsupported input file type '{suffix}'. Use CSV, TSV, PSV, XLSX, or XLSM."
    )


def get_schema_and_column_names(df: pl.DataFrame) -> tuple[list[str], dict[str, str]]:
    """Return column names and JSON-safe schema."""

    return list(df.schema.keys()), {
        name: str(dtype) for name, dtype in df.schema.items()
    }


def normalize_name(name: str) -> str:
    """Normalize a column name for matching."""

    return " ".join(name.replace("_", " ").replace("-", " ").lower().split())


def compact_name(name: str) -> str:
    """Return a compact normalized name."""

    return "".join(normalize_name(name).split())


def first_matching_column(columns: Iterable[str], hints: Iterable[str]) -> str | None:
    """Return the first column whose normalized name matches the hints."""

    normalized = [
        (column, normalize_name(column), compact_name(column)) for column in columns
    ]
    hint_values = [(normalize_name(hint), compact_name(hint)) for hint in hints]
    for hint, compact_hint in hint_values:
        for column, column_name, column_compact in normalized:
            if column_name == hint or column_compact == compact_hint:
                return column
    for hint, compact_hint in hint_values:
        for column, column_name, column_compact in normalized:
            if hint in column_name or compact_hint in column_compact:
                return column
    return None


def numeric_columns(df: pl.DataFrame) -> list[str]:
    """Return numeric columns that look like measures, not ids."""

    return [
        name
        for name, dtype in df.schema.items()
        if dtype.is_numeric() and not normalize_name(name).endswith("id")
    ]


def parse_date_expression(column: str) -> pl.Expr:
    """Return a permissive date parsing expression."""

    expr = pl.col(column)
    return expr.cast(pl.Date, strict=False).fill_null(
        expr.cast(pl.Utf8)
        .str.strptime(pl.Date, strict=False)
        .fill_null(
            expr.cast(pl.Utf8).str.strptime(pl.Datetime, strict=False).cast(pl.Date)
        )
    )


def _normalize_period_grain(value: Any) -> str | None:
    """Return a supported canonical period grain."""

    if value is None:
        return None
    normalized = normalize_name(str(value))
    if normalized in {"year", "annual", "yearly"}:
        return PERIOD_GRAIN_YEAR
    if normalized in {"quarter", "quarterly", "qtr"}:
        return PERIOD_GRAIN_QUARTER
    if normalized in {"month", "monthly"}:
        return PERIOD_GRAIN_MONTH
    if normalized in {"week", "weekly"}:
        return PERIOD_GRAIN_WEEK
    return None


def _requires_raw_date_period_window(options: dict[str, Any]) -> bool:
    """Return whether annual comparison mode must keep raw dates until plotting."""

    raw_mode = str(options.get("period_comparison_mode") or "").strip().lower()
    normalized_mode = raw_mode.replace("-", "_")
    if normalized_mode in {
        "calendar",
        "calendar_period",
        "calendar_year",
        "calendar_years",
        "complete_calendar_year",
        "complete_calendar_years",
        "rolling",
        "rolling_period",
        "rolling_window",
        "r12m",
        "year_to_date",
        "ytd",
    }:
        return True
    return bool(options.get("period_to_date") or options.get("rolling_comparison"))


def _coerce_period_label_to_grain(value: Any, period_grain: str | None) -> str:
    """Coerce a period label to the canonical period grain."""

    text = str(value)
    if period_grain != PERIOD_GRAIN_YEAR:
        return text
    if isinstance(value, (date, datetime)):
        return str(value.year)
    if len(text) >= 4 and text[:4].isdigit():
        return text[:4]
    try:
        return str(datetime.fromisoformat(text).year)
    except ValueError:
        return text


def _coerce_recipe_period_labels_to_grain(
    recipe: dict[str, Any], period_grain: str | None
) -> dict[str, Any]:
    """Coerce recipe period labels when canonical periods are grain-derived."""

    if period_grain is None:
        return {"status": "skipped", "reason": "raw_periods"}
    options = recipe.setdefault("options", {})
    changed: dict[str, dict[str, str]] = {}
    for key in (
        "current_period_label",
        "cohort_current_period",
        "cohort_previous_period",
        "current_period",
        "previous_period",
        "comparison_period",
        "baseline_period",
    ):
        if key in options and options[key] not in (None, ""):
            before = str(options[key])
            after = _coerce_period_label_to_grain(before, period_grain)
            options[key] = after
            if after != before:
                changed[key] = {"before": before, "after": after}
    cohort_definition = options.get("cohort_definition")
    if isinstance(cohort_definition, dict):
        periods = cohort_definition.get("periods")
        if isinstance(periods, dict):
            for key in ("current_period", "previous_period"):
                if key in periods and periods[key] not in (None, ""):
                    before = str(periods[key])
                    after = _coerce_period_label_to_grain(before, period_grain)
                    periods[key] = after
                    if after != before:
                        changed[f"cohort_definition.periods.{key}"] = {
                            "before": before,
                            "after": after,
                        }
    return {
        "status": "written",
        "period_grain": period_grain,
        "changed": changed,
    }


def infer_date_column(df: pl.DataFrame, columns: list[str]) -> str | None:
    """Infer the most likely date column."""

    direct = first_matching_column(
        columns, ["date", "order date", "orderdate", "month"]
    )
    candidates = [direct] if direct else []
    candidates.extend(column for column in columns if column not in candidates)
    best_column: str | None = None
    best_count = 0
    for column in candidates:
        if column is None:
            continue
        try:
            parsed = df.select(
                parse_date_expression(column).is_not_null().sum().alias("parsed")
            ).item()
        except (pl.exceptions.PolarsError, TypeError, ValueError):
            parsed = 0
        if int(parsed or 0) > best_count:
            best_column = column
            best_count = int(parsed or 0)
    return best_column if best_count > 0 else None


def infer_period_column(columns: list[str]) -> str | None:
    """Infer an optional period/scenario column."""

    return first_matching_column(
        columns,
        [
            "period",
            "scenario",
            "version",
            "actual plan",
            "actual vs plan",
            "ac py",
        ],
    )


def infer_amount_column(df: pl.DataFrame, columns: list[str]) -> str | None:
    """Infer the primary amount column."""

    preferred = first_matching_column(
        columns,
        [
            "salesamount",
            "sales amount",
            "sales",
            "revenue",
            "amount",
            "net sales",
            "value",
            "turnover",
        ],
    )
    if preferred and preferred in numeric_columns(df):
        return preferred
    numeric = numeric_columns(df)
    return numeric[0] if numeric else None


def infer_width_metric_column(
    df: pl.DataFrame, columns: list[str], amount_column: str | None
) -> str | None:
    """Infer a second numeric metric suitable for BarMekko width."""

    preferred = first_matching_column(
        columns,
        [
            "units",
            "unit",
            "quantity",
            "qty",
            "volume",
            "order quantity",
        ],
    )
    if preferred and preferred != amount_column and preferred in numeric_columns(df):
        return preferred
    return None


def _looks_like_ratio_metric_name(column: str | None) -> bool:
    """Return whether a metric name appears to be a ratio/rate/percent."""

    if not column:
        return False
    normalized = normalize_name(str(column))
    compact = compact_name(str(column))
    return (
        "%" in str(column)
        or "pct" in compact
        or any(
            term in normalized
            for term in ("percent", "percentage", "rate", "ratio", "share")
        )
    )


def _non_ratio_numeric_metric_columns(
    df: pl.DataFrame, columns: list[str], excluded: set[str | None]
) -> list[str]:
    """Return numeric metric columns excluding ids and ratio-like names."""

    numeric = set(numeric_columns(df))
    return [
        column
        for column in columns
        if column in numeric
        and column not in excluded
        and not _looks_like_ratio_metric_name(column)
    ]


def infer_margin_column(
    df: pl.DataFrame,
    columns: list[str],
    amount_column: str | None,
    width_metric_column: str | None,
) -> str | None:
    """Infer an absolute margin/profit column for derived margin percent."""

    candidates = _non_ratio_numeric_metric_columns(
        df, columns, {amount_column, width_metric_column}
    )
    return first_matching_column(
        candidates,
        [
            "gross margin",
            "margin amount",
            "contribution margin",
            "gross profit",
            "profit",
            "margin",
        ],
    )


def infer_margin_percent_column(
    df: pl.DataFrame,
    columns: list[str],
    amount_column: str | None,
    width_metric_column: str | None,
    margin_column: str | None,
) -> str | None:
    """Infer an already calculated margin percent column when present."""

    numeric = set(numeric_columns(df))
    candidates = [
        column
        for column in columns
        if column in numeric
        and column not in {amount_column, width_metric_column, margin_column}
        and _looks_like_ratio_metric_name(column)
    ]
    return first_matching_column(
        candidates,
        [
            "margin in %",
            "gross margin in %",
            "margin %",
            "gross margin %",
            "margin percent",
            "gross margin percent",
            "margin percentage",
            "gross margin percentage",
        ],
    )


def infer_dimensions(
    columns: list[str],
    date_column: str | None,
    amount_column: str | None,
    period_column: str | None,
) -> list[str]:
    """Infer useful business dimensions for mix and contribution charts."""

    excluded = {
        date_column,
        amount_column,
        period_column,
        CANONICAL_DATE,
        CANONICAL_PERIOD,
        "Scenario",
    }
    dimension_hints = [
        "productline",
        "product line",
        "category",
        "subcategory",
        "product type",
        "product",
        "region",
        "country",
        "customer",
        "segment",
        "channel",
        "brand",
    ]
    dimensions: list[str] = []
    for hint in dimension_hints:
        column = first_matching_column(columns, [hint])
        if column and column not in excluded and column not in dimensions:
            dimensions.append(column)
    for column in columns:
        if column not in excluded and column not in dimensions:
            dimensions.append(column)
    return dimensions[:6]


def build_recipe(
    input_path: Path,
    df: pl.DataFrame,
    *,
    language: str,
    existing_recipe: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Build a default recipe, merging an existing recipe if supplied."""

    columns, schema = get_schema_and_column_names(df)
    date_column = infer_date_column(df, columns)
    period_column = infer_period_column(columns)
    amount_column = infer_amount_column(df, columns)
    width_metric_column = infer_width_metric_column(df, columns, amount_column)
    margin_column = infer_margin_column(df, columns, amount_column, width_metric_column)
    margin_percent_column = infer_margin_percent_column(
        df, columns, amount_column, width_metric_column, margin_column
    )
    dimensions = infer_dimensions(columns, date_column, amount_column, period_column)
    metric_columns = {
        item
        for item in (width_metric_column, margin_column, margin_percent_column)
        if item
    }
    if metric_columns:
        dimensions = [
            dimension for dimension in dimensions if dimension not in metric_columns
        ]
    recipe: dict[str, Any] = {
        "schema_version": SCHEMA_VERSION,
        "source_file": str(input_path),
        "language": language,
        "mappings": {
            "date_column": date_column,
            "period_column": period_column,
            "amount_column": amount_column,
            "width_metric_column": width_metric_column,
            "margin_column": margin_column,
            "margin_percent_column": margin_percent_column,
            "dimensions": dimensions,
        },
        "options": {
            "currency": "EUR",
            "current_period_label": CURRENT_PERIOD,
            "period_selection": "infer_current_or_all",
            "charts": [
                "marimekko",
                "marimekko_small_multiples",
                "barmekko",
                "barmekko_small_multiples",
                "stacked_bar",
                "stacked_bar_small_multiples",
                "bar",
                "bar_small_multiples",
                "related_metrics_bar",
                "related_metrics_bar_small_multiples",
                "stacked_column",
                "stacked_column_small_multiples",
                "stacked_column_synthesis",
                "line",
                "line_small_multiples",
                "area_absolute",
                "area_share",
                "pareto",
                "stacked_pareto_abc",
                "stacked_pareto_by_dimension",
                "multitier_bar",
                "multitier_bar_dimension_panels",
            ],
            "small_multiples": bool(len(dimensions) >= 3),
            "small_multiples_dimension": None,
            "max_chart_items": 12,
            "small_multiples_max_panels": 6,
        },
        "inspection": {
            "columns": columns,
            "schema": schema,
        },
    }
    if existing_recipe:
        recipe["mappings"].update(existing_recipe.get("mappings") or {})
        recipe["options"].update(existing_recipe.get("options") or {})
        if existing_recipe.get("language"):
            recipe["language"] = existing_recipe["language"]
    return validate_recipe(df, recipe)


def _normalized_period_label(value: Any) -> str:
    """Return a normalized label for mechanical period-comparison checks."""

    return str(value).strip().casefold()


def _first_option_value(options: dict[str, Any], keys: Iterable[str]) -> Any:
    """Return the first non-empty recipe option value for the given keys."""

    for key in keys:
        value = options.get(key)
        if value not in (None, ""):
            return value
    return None


def _validate_distinct_period_options(options: dict[str, Any]) -> None:
    """Reject explicitly requested same-label period comparisons."""

    current_value = _first_option_value(
        options,
        ("current_period_label", "current_period", "cohort_current_period"),
    )
    if current_value in (None, ""):
        current_value = CURRENT_PERIOD
    current = _normalized_period_label(current_value)
    for key in (
        "previous_period_label",
        "previous_period",
        "comparison_period",
        "baseline_period",
        "cohort_previous_period",
    ):
        value = options.get(key)
        if value in (None, ""):
            continue
        # This deterministic guard is justified because same-label comparisons
        # are mechanically invalid: AC-vs-AC, PY-vs-PY, etc. have no contrast.
        if _normalized_period_label(value) == current:
            raise ValueError(
                "Mix comparison requires distinct current and comparison "
                f"period labels; {key} and the current period are both "
                f"{value!r}."
            )

    selected_periods = options.get("selected_periods")
    if isinstance(selected_periods, list):
        normalized_periods = [
            _normalized_period_label(value)
            for value in selected_periods
            if value not in (None, "")
        ]
        if len(normalized_periods) != len(set(normalized_periods)):
            raise ValueError(
                "Mix comparison selected_periods must not contain duplicate "
                "period labels."
            )


def validate_recipe(df: pl.DataFrame, recipe: dict[str, Any]) -> dict[str, Any]:
    """Validate the recipe against the input frame."""

    columns, _ = get_schema_and_column_names(df)
    mappings = recipe.setdefault("mappings", {})
    options = recipe.setdefault("options", {})
    amount_column = mappings.get("amount_column")
    source_amount_column = mappings.get("source_amount_column")
    if amount_column not in columns and source_amount_column in columns:
        amount_column = source_amount_column
        mappings["amount_column"] = amount_column
    if not amount_column or amount_column not in columns:
        raise ValueError("A valid amount column is required for mix analysis.")
    if amount_column not in numeric_columns(df):
        raise ValueError(f"Amount column must be numeric: {amount_column}")
    _validate_distinct_period_options(options)
    for optional_column in ("date_column", "period_column"):
        value = mappings.get(optional_column)
        if value and value not in columns:
            mappings[optional_column] = None
    width_metric_column = mappings.get("width_metric_column")
    if width_metric_column and (
        width_metric_column not in columns
        or width_metric_column == amount_column
        or width_metric_column not in numeric_columns(df)
    ):
        mappings["width_metric_column"] = None
        width_metric_column = None
    margin_column = mappings.get("margin_column")
    if margin_column and _looks_like_ratio_metric_name(str(margin_column)):
        if not mappings.get("margin_percent_column"):
            mappings["margin_percent_column"] = margin_column
        mappings["margin_column"] = None
        margin_column = None
    if margin_column and (
        margin_column not in columns
        or margin_column == amount_column
        or margin_column == width_metric_column
        or margin_column not in numeric_columns(df)
    ):
        mappings["margin_column"] = None
        margin_column = None
    margin_percent_column = mappings.get("margin_percent_column")
    if margin_percent_column and (
        margin_percent_column not in columns
        or margin_percent_column == amount_column
        or margin_percent_column == width_metric_column
        or margin_percent_column == margin_column
        or margin_percent_column not in numeric_columns(df)
    ):
        mappings["margin_percent_column"] = None
        margin_percent_column = None
    related_marker_metric_column = mappings.get("related_marker_metric_column")
    if related_marker_metric_column and (
        related_marker_metric_column not in columns
        or related_marker_metric_column == amount_column
        or related_marker_metric_column not in numeric_columns(df)
    ):
        mappings["related_marker_metric_column"] = None
        related_marker_metric_column = None
    if (
        related_marker_metric_column
        and related_marker_metric_column == width_metric_column
    ):
        mappings["width_metric_column"] = None
        width_metric_column = None
    excluded = {
        mappings.get("date_column"),
        mappings.get("period_column"),
        amount_column,
        mappings.get("width_metric_column"),
        mappings.get("margin_column"),
        mappings.get("margin_percent_column"),
        mappings.get("related_marker_metric_column"),
    }
    cohort_source_dimensions = recipe_cohort_source_dimensions(recipe)
    missing_cohort_sources = [
        column for column in cohort_source_dimensions if column not in columns
    ]
    if missing_cohort_sources:
        raise ValueError(
            "Recipe cohort source columns are missing: "
            + ", ".join(missing_cohort_sources)
        )
    cohort_dimension_names = recipe_cohort_dimension_names(recipe)
    dimensions = [
        dimension
        for dimension in mappings.get("dimensions") or []
        if (dimension in columns or dimension in cohort_dimension_names)
        and dimension not in excluded
    ]
    if not dimensions:
        raise ValueError("At least one business dimension is required.")
    mappings["dimensions"] = dimensions
    options["currency"] = options.get("currency") or "EUR"
    normalized_period_contract = period_contract_options(options)
    if normalize_name(str(options.get("period_grain") or "")) in {
        "raw",
        "scenario",
        "scenario_period",
    }:
        normalized_period_contract["period_grain"] = "raw"
    options["period_type"] = normalized_period_contract["period_type"]
    options["period_grain"] = normalized_period_contract["period_grain"]
    options["fiscal_start_month"] = normalized_period_contract["fiscal_start_month"]
    options["small_multiples"] = bool(
        options.get("small_multiples", len(dimensions) >= 3)
    )
    options["max_chart_items"] = int(options.get("max_chart_items") or 12)
    options["small_multiples_max_panels"] = max(
        2, int(options.get("small_multiples_max_panels") or 6)
    )
    options["charts"] = _normalize_requested_chart_names(options.get("charts") or [])
    options["reporting_entity_label"] = options.get(
        "reporting_entity_label"
    ) or reporting_entity_label_from_recipe(recipe)
    return recipe


def inspect_mix_inputs(
    input_path: Path,
    output_dir: Path,
    *,
    language: str = "en",
) -> InspectionResult:
    """Inspect inputs and write suggested mix-contribution recipe files."""

    df = read_table(input_path)
    recipe = build_recipe(input_path, df, language=language)
    columns, schema = get_schema_and_column_names(df)
    payload = {
        "schema_version": SCHEMA_VERSION,
        "input_file": str(input_path),
        "row_count": df.height,
        "column_count": df.width,
        "columns": columns,
        "schema": schema,
        "available_analysis_context": available_analysis_context(df),
        "suggested_mappings": recipe["mappings"],
        "suggested_options": recipe["options"],
        "warnings": [],
    }
    output_dir.mkdir(parents=True, exist_ok=True)
    write_json(output_dir / "inspection.json", payload)
    write_json(output_dir / "suggested_recipe.json", recipe)
    return InspectionResult(payload=payload, recipe=recipe, output_dir=output_dir)


def _select_period_value(df: pl.DataFrame, period_column: str) -> str | None:
    values = [
        str(value)
        for value in df.select(
            pl.col(period_column).cast(pl.Utf8).drop_nulls().unique()
        )
        .to_series()
        .to_list()
    ]
    if not values:
        return None
    normalized = {normalize_name(value): value for value in values}
    for candidate in ("ac", "actual", "current", "current period", "period one"):
        if candidate in normalized:
            return normalized[candidate]
    return None


def prepare_canonical_frame(
    df: pl.DataFrame, recipe: dict[str, Any]
) -> tuple[pl.DataFrame, dict[str, Any]]:
    """Build the AC canonical frame consumed by legacy chart preparation."""

    mappings = recipe["mappings"]
    options = recipe.setdefault("options", {})
    date_column = mappings.get("date_column")
    period_column = mappings.get("period_column")
    period_grain = _normalize_period_grain(options.get("period_grain"))
    source_amount_column = str(mappings["amount_column"])
    amount_column = LEGACY_AMOUNT_COLUMN
    mappings["source_amount_column"] = source_amount_column
    mappings["amount_column"] = amount_column
    dimensions = [str(item) for item in mappings.get("dimensions") or []]
    raw_columns, _schema = get_schema_and_column_names(df)
    base_dimensions = list(
        dict.fromkeys(
            [
                *[dimension for dimension in dimensions if dimension in raw_columns],
                *[
                    dimension
                    for dimension in recipe_cohort_source_dimensions(recipe)
                    if dimension in raw_columns
                ],
            ]
        )
    )
    period_value = None
    width_metric_column = mappings.get("width_metric_column")
    margin_column = mappings.get("margin_column")
    margin_percent_column = mappings.get("margin_percent_column")
    related_marker_metric_column = mappings.get("related_marker_metric_column")
    lf = df.lazy()
    if date_column:
        lf = lf.with_columns(
            parse_date_expression(str(date_column)).alias(CANONICAL_DATE)
        )
    else:
        lf = lf.with_columns(pl.lit(date.today()).cast(pl.Date).alias(CANONICAL_DATE))
    period_source_column = period_column
    if period_column:
        period_value = _select_period_value(df, str(period_column))
        keep_raw_date_periods_for_window = (
            period_grain == PERIOD_GRAIN_YEAR
            and date_column
            and str(period_column) == str(date_column)
            and _requires_raw_date_period_window(options)
        )
        if period_grain == PERIOD_GRAIN_YEAR and not keep_raw_date_periods_for_window:
            period_source_column = date_column or period_column
            lf = lf.with_columns(
                parse_date_expression(str(period_source_column))
                .dt.year()
                .cast(pl.Int64)
                .cast(pl.Utf8)
                .fill_null(CURRENT_PERIOD)
                .alias(CANONICAL_PERIOD)
            )
        else:
            lf = lf.with_columns(
                pl.col(str(period_column))
                .cast(pl.Utf8)
                .fill_null(CURRENT_PERIOD)
                .alias(CANONICAL_PERIOD)
            )
    else:
        lf = lf.with_columns(pl.lit(CURRENT_PERIOD).alias(CANONICAL_PERIOD))
    select_exprs: list[pl.Expr] = [
        pl.col(CANONICAL_DATE),
        pl.col(CANONICAL_PERIOD),
        pl.col(source_amount_column).cast(pl.Float64).alias(amount_column),
    ]
    legacy_metric_columns = [amount_column]
    legacy_width_metric_column = None
    if width_metric_column and width_metric_column != source_amount_column:
        legacy_width_metric_column = LEGACY_UNITS_COLUMN
        legacy_metric_columns.append(legacy_width_metric_column)
        select_exprs.append(
            pl.col(str(width_metric_column))
            .cast(pl.Float64)
            .alias(legacy_width_metric_column)
        )
        legacy_metric_columns.extend(
            [LEGACY_UNIT_PRICE_COLUMN, LEGACY_MULTIPLIED_COLUMN]
        )
        units_expr = pl.col(str(width_metric_column)).cast(pl.Float64)
        sales_expr = pl.col(source_amount_column).cast(pl.Float64)
        select_exprs.extend(
            [
                pl.when(units_expr != 0.0)
                .then(sales_expr / units_expr)
                .otherwise(None)
                .alias(LEGACY_UNIT_PRICE_COLUMN),
                sales_expr.alias(LEGACY_MULTIPLIED_COLUMN),
            ]
        )
    legacy_margin_column = None
    legacy_margin_percent_column = None
    if margin_column and margin_column != source_amount_column:
        legacy_margin_column = LEGACY_MARGIN_COLUMN
        legacy_margin_percent_column = LEGACY_MARGIN_PERCENT_COLUMN
        legacy_metric_columns.extend(
            [legacy_margin_column, legacy_margin_percent_column]
        )
        margin_expr = pl.col(str(margin_column)).cast(pl.Float64)
        sales_expr = pl.col(source_amount_column).cast(pl.Float64)
        select_exprs.extend(
            [
                margin_expr.alias(legacy_margin_column),
                pl.when(sales_expr != 0.0)
                .then(margin_expr / sales_expr)
                .otherwise(0.0)
                .alias(legacy_margin_percent_column),
            ]
        )
    elif margin_percent_column and margin_percent_column != source_amount_column:
        legacy_margin_percent_column = LEGACY_MARGIN_PERCENT_COLUMN
        legacy_metric_columns.append(legacy_margin_percent_column)
        margin_percent_expr = pl.col(str(margin_percent_column)).cast(pl.Float64)
        select_exprs.append(
            pl.when(margin_percent_expr.abs() > 1.5)
            .then(margin_percent_expr / 100.0)
            .otherwise(margin_percent_expr)
            .alias(legacy_margin_percent_column)
        )
    if related_marker_metric_column == margin_column:
        related_marker_metric_column = legacy_margin_column
    elif related_marker_metric_column == margin_percent_column:
        related_marker_metric_column = legacy_margin_percent_column
    if (
        related_marker_metric_column
        and related_marker_metric_column != source_amount_column
        and related_marker_metric_column != width_metric_column
        and related_marker_metric_column not in legacy_metric_columns
    ):
        related_marker_metric_column = str(related_marker_metric_column)
        legacy_metric_columns.append(related_marker_metric_column)
        select_exprs.append(
            pl.col(related_marker_metric_column)
            .cast(pl.Float64)
            .alias(related_marker_metric_column)
        )
    select_exprs.extend(
        [
            pl.col(dimension).cast(pl.Utf8).fill_null("Unspecified").alias(dimension)
            for dimension in base_dimensions
        ]
    )
    canonical = (
        lf.filter(pl.col(CANONICAL_DATE).is_not_null()).select(select_exprs).collect()
    )
    if canonical.is_empty():
        raise ValueError("No rows remain after canonical period selection.")
    sort_columns = [
        column
        for column in [CANONICAL_DATE, CANONICAL_PERIOD, *base_dimensions]
        if column in canonical.columns
    ]
    canonical = canonical.sort(sort_columns)
    audit = {
        "status": "written",
        "date_column": date_column,
        "period_column": period_column,
        "period_source_column": period_source_column,
        "period_grain": period_grain or "raw",
        "period_label_coercion": _coerce_recipe_period_labels_to_grain(
            recipe, period_grain
        ),
        "selected_period_value": period_value,
        "period_policy": (
            "selected_current_period_value"
            if period_value is not None
            else "all_rows_as_current_period"
        ),
        "source_amount_column": source_amount_column,
        "legacy_amount_column": amount_column,
        "source_width_metric_column": width_metric_column,
        "legacy_width_metric_column": legacy_width_metric_column,
        "source_margin_column": margin_column,
        "legacy_margin_column": legacy_margin_column,
        "source_margin_percent_column": margin_percent_column,
        "legacy_margin_percent_column": legacy_margin_percent_column,
        "related_marker_metric_column": related_marker_metric_column,
        "legacy_metric_columns": legacy_metric_columns,
        "row_count": canonical.height,
        "dimension_count": len(dimensions),
        "cohort_source_dimensions": [
            dimension
            for dimension in recipe_cohort_source_dimensions(recipe)
            if dimension in raw_columns
        ],
    }
    mappings["legacy_metric_columns"] = legacy_metric_columns
    mappings["legacy_width_metric_column"] = legacy_width_metric_column
    mappings["legacy_margin_column"] = legacy_margin_column
    mappings["legacy_margin_percent_column"] = legacy_margin_percent_column
    mappings["related_marker_metric_column"] = related_marker_metric_column
    return canonical, audit


def _dimension_cardinality(
    canonical: pl.DataFrame, dimensions: list[str]
) -> dict[str, int]:
    return {
        dimension: int(
            canonical.select(pl.col(dimension).n_unique().alias("n")).item() or 0
        )
        for dimension in dimensions
    }


def _useful_dimensions(canonical: pl.DataFrame, dimensions: list[str]) -> list[str]:
    cardinality = _dimension_cardinality(canonical, dimensions)
    useful = [
        dimension for dimension in dimensions if cardinality.get(dimension, 0) > 1
    ]
    return useful or dimensions


def _is_observed_hierarchy_or_duplicate_pair(
    canonical: pl.DataFrame, left: str, right: str
) -> bool:
    """Return True when two dimensions do not form a useful Mekko cross."""

    if left == right:
        return True
    cardinality = _dimension_cardinality(canonical, [left, right])
    left_count = cardinality.get(left, 0)
    right_count = cardinality.get(right, 0)
    if left_count < 2 or right_count < 2:
        return True
    pair_count = int(
        canonical.select([pl.col(left), pl.col(right)]).unique().height or 0
    )
    return pair_count <= max(left_count, right_count)


def _is_residual_bucket_label(value: Any) -> bool:
    """Return True for visual residual buckets such as Other/Others rank."""

    normalized = normalize_name(str(value))
    return (
        normalized in {"other", "others"}
        or normalized.startswith("other rank")
        or normalized.startswith("others rank")
        or normalized.startswith("all other")
    )


def _stacked_bar_small_multiples_period_frame(
    canonical: pl.DataFrame, selected_periods: list[str] | None
) -> pl.DataFrame:
    """Return the period slice used to judge stacked-bar panel usefulness."""

    columns, _schema = get_schema_and_column_names(canonical)
    if CANONICAL_PERIOD not in columns or not selected_periods:
        return canonical
    period_values = [str(period) for period in selected_periods if period is not None]
    if not period_values:
        return canonical
    period_frame = canonical.filter(
        pl.col(CANONICAL_PERIOD).cast(pl.Utf8).is_in(period_values)
    )
    return canonical if period_frame.is_empty() else period_frame


def _is_admissible_stacked_bar_small_multiples_dimension(
    canonical: pl.DataFrame,
    *,
    x_dimension: str,
    y_dimension: str | None,
    facet_dimension: str,
    metric: str,
    selected_periods: list[str] | None,
) -> bool:
    """Reject facets that mechanically collapse stacked-bar panels.

    This is deterministic because a panel with only one visible row is not a
    useful stacked-bar small multiple: the plugin can measure that directly from
    the prepared data before exposing the artifact to the reporting layer.
    """

    if (
        not y_dimension
        or facet_dimension in {x_dimension, y_dimension}
        or _is_observed_hierarchy_or_duplicate_pair(
            canonical, y_dimension, facet_dimension
        )
    ):
        return False
    columns, _schema = get_schema_and_column_names(canonical)
    required_columns = {facet_dimension, x_dimension, y_dimension, metric}
    if not required_columns.issubset(set(columns)):
        return False
    frame = _stacked_bar_small_multiples_period_frame(canonical, selected_periods)
    panel_rows = (
        frame.group_by([facet_dimension, x_dimension, y_dimension])
        .agg(pl.col(metric).sum().abs().alias("__value"))
        .filter(pl.col("__value") > TOLERANCE)
        .group_by(facet_dimension)
        .agg(
            [
                pl.col(x_dimension).n_unique().alias("__visible_x"),
                pl.col(y_dimension).n_unique().alias("__visible_y"),
                pl.col("__value").sum().alias("__total"),
            ]
        )
    )
    usable_panels = [
        (
            int(row["__visible_x"]),
            int(row["__visible_y"]),
            float(row["__total"] or 0.0),
        )
        for row in panel_rows.to_dicts()
        if not _is_residual_bucket_label(row[facet_dimension])
    ]
    if len(usable_panels) < STACKED_BAR_SMALL_MULTIPLES_MIN_PANELS:
        return False
    panel_total = sum(total for _visible_x, _visible_y, total in usable_panels)
    if panel_total <= TOLERANCE:
        return False
    multi_axis_panels = sum(
        1
        for visible_x, visible_y, _total in usable_panels
        if visible_x >= 2 and visible_y >= 2
    )
    return (
        multi_axis_panels >= STACKED_BAR_SMALL_MULTIPLES_MIN_PANELS
        and multi_axis_panels / len(usable_panels)
        >= STACKED_BAR_SMALL_MULTIPLES_MIN_MULTI_ROW_SHARE
    )


def _find_mekko_dimension_pair(
    canonical: pl.DataFrame, dimensions: list[str]
) -> tuple[str, str] | None:
    """Choose the first non-hierarchical pair suitable for Mekko/BarMekko."""

    if len(dimensions) < 2:
        return None
    first = dimensions[0]
    preferred_pairs = [
        (first, dimension) for dimension in dimensions[1:] if dimension != first
    ]
    fallback_pairs = [
        (left, right)
        for left_index, left in enumerate(dimensions)
        for right in dimensions[left_index + 1 :]
        if left != right and (left, right) not in preferred_pairs
    ]
    for left, right in [*preferred_pairs, *fallback_pairs]:
        if not _is_observed_hierarchy_or_duplicate_pair(canonical, left, right):
            return left, right
    return None


def _is_valid_mekko_dimension_triple(
    canonical: pl.DataFrame, left: str, right: str, facet: str
) -> bool:
    """Return True when axes and facet are pairwise suitable for Mekko."""

    return not any(
        _is_observed_hierarchy_or_duplicate_pair(canonical, first, second)
        for first, second in ((left, right), (left, facet), (right, facet))
    )


def _find_mekko_small_multiple_dimensions(
    canonical: pl.DataFrame, dimensions: list[str]
) -> tuple[str, str, str] | None:
    """Return ``x``, ``y`` and facet dimensions for Mekko small multiples."""

    if len(dimensions) < 3:
        return None
    for left_index, left in enumerate(dimensions):
        for right in dimensions[left_index + 1 :]:
            if _is_observed_hierarchy_or_duplicate_pair(canonical, left, right):
                continue
            for facet in dimensions:
                if facet in {left, right}:
                    continue
                if _is_valid_mekko_dimension_triple(canonical, left, right, facet):
                    return left, right, facet
    return None


def _is_observed_child_parent_pair(
    canonical: pl.DataFrame, child: str, parent: str
) -> bool:
    """Return True when each observed child belongs to one observed parent."""

    if child == parent:
        return False
    cardinality = _dimension_cardinality(canonical, [child, parent])
    child_count = cardinality.get(child, 0)
    parent_count = cardinality.get(parent, 0)
    if child_count < 2 or parent_count < 2 or child_count < parent_count:
        return False
    pair_count = int(
        canonical.select([pl.col(child), pl.col(parent)]).unique().height or 0
    )
    return pair_count == child_count


def _find_pareto_dimensions(
    canonical: pl.DataFrame, dimensions: list[str]
) -> tuple[str, str | None, str]:
    """Choose item and optional parent dimensions for legacy Pareto charts."""

    if not dimensions:
        raise ValueError("Pareto charts require at least one dimension.")
    cardinality = _dimension_cardinality(canonical, dimensions)
    ranked_dimensions = sorted(
        dimensions,
        key=lambda dimension: (
            cardinality.get(dimension, 0),
            -dimensions.index(dimension),
        ),
        reverse=True,
    )
    for child in ranked_dimensions:
        for parent in dimensions:
            if _is_observed_child_parent_pair(canonical, child, parent):
                return child, parent, "observed_child_parent_dimension_pair"
    return ranked_dimensions[0], None, "largest_observed_dimension"


def _find_stacked_pareto_dimensions(
    canonical: pl.DataFrame, dimensions: list[str]
) -> tuple[str, str | None, str]:
    """Choose readable item and optional parent dimensions for stacked Pareto."""

    if not dimensions:
        raise ValueError("Stacked Pareto charts require at least one dimension.")
    cardinality = _dimension_cardinality(canonical, dimensions)
    hierarchy_candidates: list[tuple[int, int, int, str, str]] = []
    for child in dimensions:
        child_count = cardinality.get(child, 0)
        if child_count < 2 or child_count > STACKED_PARETO_MAX_RANKED_ITEMS:
            continue
        for parent in dimensions:
            parent_count = cardinality.get(parent, 0)
            if _is_observed_child_parent_pair(canonical, child, parent):
                hierarchy_candidates.append(
                    (
                        child_count,
                        -parent_count,
                        -dimensions.index(child),
                        child,
                        parent,
                    )
                )
    if hierarchy_candidates:
        _child_count, _parent_score, _order, child, parent = max(hierarchy_candidates)
        return child, parent, "bounded_observed_child_parent_dimension_pair"

    bounded_dimensions = [
        dimension
        for dimension in dimensions
        if 2 <= cardinality.get(dimension, 0) <= STACKED_PARETO_MAX_RANKED_ITEMS
    ]
    if bounded_dimensions:
        dimension = max(
            bounded_dimensions,
            key=lambda item: (cardinality.get(item, 0), -dimensions.index(item)),
        )
        return dimension, None, "bounded_ranked_dimension"
    return _find_pareto_dimensions(canonical, dimensions)


def _pareto_metrics(
    canonical: pl.DataFrame, metric: str, width_metric: str | None
) -> list[str]:
    """Return additive metrics suitable for legacy Pareto charting."""

    columns, _schema = get_schema_and_column_names(canonical)
    numeric_metric_columns = set(numeric_columns(canonical))
    excluded = {LEGACY_UNIT_PRICE_COLUMN, LEGACY_MULTIPLIED_COLUMN}
    metrics: list[str] = []
    for candidate in [metric, width_metric, *columns]:
        if not candidate:
            continue
        candidate_name = str(candidate)
        if (
            candidate_name in numeric_metric_columns
            and candidate_name not in excluded
            and candidate_name not in metrics
        ):
            metrics.append(candidate_name)
        if len(metrics) >= 2:
            break
    return metrics or [metric]


def _metric_supports_cagr(metric: str) -> bool:
    """Return whether a metric can sensibly carry a CAGR annotation."""

    metric_name = str(metric or "").strip()
    if not metric_name or metric_name in CAGR_BLOCKED_METRICS:
        return False
    normalized = metric_name.lower()
    return not any(term in normalized for term in CAGR_BLOCKED_METRIC_TERMS)


def _can_show_cagr(metric: str, selected_periods: Sequence[str] | None) -> bool:
    """Return whether a chart has both a valid metric and enough periods for CAGR."""

    if not _metric_supports_cagr(metric):
        return False
    unique_periods = {str(period) for period in (selected_periods or []) if period}
    return len(unique_periods) >= 2


def _can_show_total_cagr(metric: str, selected_periods: Sequence[str] | None) -> bool:
    """Return whether a total CAGR label is meaningful for a period window."""

    if not _metric_supports_cagr(metric):
        return False
    unique_periods = {str(period) for period in (selected_periods or []) if period}
    return len(unique_periods) >= 2


def _has_comparison_periods(selected_periods: Sequence[str] | None) -> bool:
    """Return whether a spec can support before/after comparison logic."""

    unique_periods = {str(period) for period in (selected_periods or []) if period}
    return len(unique_periods) >= 2


def _related_marker_metric(
    canonical: pl.DataFrame,
    metric: str,
    width_metric: str | None,
    selected_periods: list[str],
) -> str | None:
    """Return the default marker metric for related-metrics bars."""

    if len(selected_periods) >= 2 and metric == LEGACY_AMOUNT_COLUMN:
        return LEGACY_SALES_GROWTH_COLUMN
    for derived_metric in (LEGACY_UNIT_PRICE_COLUMN, LEGACY_MARGIN_PERCENT_COLUMN):
        if derived_metric in canonical.columns and derived_metric != metric:
            return derived_metric
    if width_metric and width_metric in canonical.columns and width_metric != metric:
        return str(width_metric)
    numeric_metric_columns = set(numeric_columns(canonical))
    for column in canonical.columns:
        if column in {
            metric,
            LEGACY_MARGIN_PERCENT_COLUMN,
            LEGACY_MULTIPLIED_COLUMN,
            LEGACY_UNIT_PRICE_COLUMN,
        }:
            continue
        if column in numeric_metric_columns:
            return column
    return None


def _related_metric_value_cols(metric: str, marker_metric: str) -> list[str]:
    """Return source value columns needed by the legacy overlay bar path."""

    if marker_metric == LEGACY_SALES_GROWTH_COLUMN:
        return [metric]
    if marker_metric == LEGACY_UNIT_PRICE_COLUMN:
        return list(dict.fromkeys([metric, LEGACY_UNITS_COLUMN]))
    if marker_metric == LEGACY_MARGIN_PERCENT_COLUMN:
        return list(dict.fromkeys([metric, LEGACY_MARGIN_COLUMN]))
    return list(dict.fromkeys([metric, marker_metric]))


def _default_related_metrics_dimension(
    canonical: pl.DataFrame, dimensions: list[str], fallback: str
) -> str:
    """Return the ranking dimension for bar charts with marker overlays."""

    cardinality = _dimension_cardinality(canonical, dimensions)
    detailed = [
        dimension for dimension in dimensions if cardinality.get(dimension, 0) >= 5
    ]
    if not detailed:
        return fallback
    preferred_names = [
        "brand",
        "customer",
        "retailer",
        "store",
        "productname",
        "product_name",
        "product",
        "sku",
        "item",
    ]
    for preferred in preferred_names:
        match = next(
            (
                dimension
                for dimension in detailed
                if normalize_name(dimension) == preferred
            ),
            None,
        )
        if match:
            return match
    return max(detailed, key=lambda dimension: cardinality.get(dimension, 0))


def _unique_dimensions(
    dimensions: Iterable[str | None], *, limit: int = 3
) -> list[str]:
    """Return unique non-empty dimensions while preserving order."""

    unique: list[str] = []
    for dimension in dimensions:
        if dimension and dimension not in unique:
            unique.append(dimension)
        if len(unique) >= limit:
            break
    return unique


def _explicit_dimension_pair(
    dimensions: list[str],
    options: dict[str, Any],
    *,
    panel_keys: tuple[str, ...],
    item_keys: tuple[str, ...],
) -> tuple[str, str] | None:
    """Return a requested panel/item dimension pair when both are valid."""

    panel_dimension = next(
        (str(options[key]) for key in panel_keys if options.get(key) in dimensions),
        "",
    )
    item_dimension = next(
        (str(options[key]) for key in item_keys if options.get(key) in dimensions),
        "",
    )
    if panel_dimension and item_dimension and panel_dimension != item_dimension:
        return panel_dimension, item_dimension
    return None


def _period_values(canonical: pl.DataFrame) -> list[str]:
    """Return stable period values from a canonical frame."""

    if CANONICAL_PERIOD not in canonical.columns:
        return []
    return [
        str(value)
        for value in canonical.select(pl.col(CANONICAL_PERIOD).cast(pl.Utf8).unique())
        .sort(CANONICAL_PERIOD)
        .to_series()
        .to_list()
    ]


def _period_values_for_grain(
    canonical: pl.DataFrame, period_grain: str | None
) -> list[str]:
    """Return period values after applying a requested display grain."""

    normalized_grain = _normalize_period_grain(period_grain)
    if normalized_grain == PERIOD_GRAIN_YEAR and CANONICAL_DATE in canonical.columns:
        return [
            str(value)
            for value in canonical.select(
                pl.col(CANONICAL_DATE)
                .dt.year()
                .cast(pl.Int64)
                .cast(pl.Utf8)
                .unique()
                .sort()
            )
            .to_series()
            .to_list()
            if value is not None
        ]
    return _period_values(canonical)


def _actual_date_range_period_label(
    canonical: pl.DataFrame, selected_periods: Sequence[str] | None
) -> str | None:
    """Return a reader-facing AC period label for date-backed actuals."""

    if [str(period) for period in selected_periods or [] if period] != [CURRENT_PERIOD]:
        return None
    columns, _schema = get_schema_and_column_names(canonical)
    if CANONICAL_DATE not in columns:
        return None
    min_date, max_date = canonical.select(
        pl.col(CANONICAL_DATE).min().alias("min_date"),
        pl.col(CANONICAL_DATE).max().alias("max_date"),
    ).row(0)
    if min_date is None or max_date is None:
        return None
    if min_date == max_date:
        return f"{CURRENT_PERIOD}, {min_date.isoformat()}"
    return f"{CURRENT_PERIOD}, {min_date.isoformat()} to {max_date.isoformat()}"


def _date_axis_point_count(
    canonical: pl.DataFrame, selected_periods: list[str] | None
) -> int:
    """Return resolved date-axis cardinality for time-chart grammar validation."""

    columns, _schema = get_schema_and_column_names(canonical)
    if CANONICAL_DATE not in columns:
        return 0
    frame = canonical
    if selected_periods:
        frame = frame.filter(
            pl.col(CANONICAL_PERIOD).cast(pl.Utf8).is_in(selected_periods)
        )
    if frame.is_empty():
        return 0
    return int(frame.select(pl.col(CANONICAL_DATE).n_unique()).item())


def _selected_periods_for_time_charts(
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
    selected_current_period: list[str],
) -> list[str]:
    """Return a valid period window for time charts, or an empty list.

    Time charts need at least two resolved dates. When periods are raw dates, a
    current-period filter would collapse the x-axis to one point, so use the
    full available date window instead.
    """

    values = _period_values(canonical)
    mappings = recipe.get("mappings") or {}
    date_column = mappings.get("date_column")
    period_column = mappings.get("period_column")
    period_values_are_dates = _period_values_look_like_dates(values)
    if period_values_are_dates or (
        date_column and period_column and str(date_column) == str(period_column)
    ):
        return values if _date_axis_point_count(canonical, values) >= 2 else []
    if _date_axis_point_count(canonical, selected_current_period) >= 2:
        return selected_current_period
    return values if _date_axis_point_count(canonical, values) >= 2 else []


def _period_values_look_like_dates(values: list[str]) -> bool:
    """Return whether period labels appear to be raw date values."""

    if not values:
        return False
    parsed_count = 0
    for value in values[: min(len(values), 20)]:
        try:
            datetime.fromisoformat(str(value))
        except ValueError:
            continue
        parsed_count += 1
    return parsed_count >= max(1, min(len(values), 20) // 2)


def _configured_comparison_period(options: dict[str, Any]) -> bool:
    """Return whether the recipe names a real non-current comparison period."""

    current = str(
        options.get("current_period_label")
        or options.get("current_period")
        or options.get("cohort_current_period")
        or CURRENT_PERIOD
    )
    for key in (
        "previous_period_label",
        "previous_period",
        "comparison_period",
        "baseline_period",
        "cohort_previous_period",
    ):
        value = options.get(key)
        if value in (None, ""):
            continue
        if str(value) != current:
            return True
    return False


def _default_period_grain_for_stacked_column(
    canonical: pl.DataFrame, recipe: dict[str, Any]
) -> str | None:
    """Return a safe default grain for discrete period-axis stacked columns."""

    options = recipe.get("options") or {}
    explicit_grain = _normalize_period_grain(
        options.get("stacked_column_period_grain") or options.get("period_grain")
    )
    if explicit_grain is not None:
        return explicit_grain
    mappings = recipe.get("mappings") or {}
    values = _period_values(canonical)
    if (
        CANONICAL_DATE in canonical.columns
        and not _configured_comparison_period(options)
        and values == [str(options.get("current_period_label") or CURRENT_PERIOD)]
        and _period_values_for_grain(canonical, PERIOD_GRAIN_YEAR)
    ):
        return PERIOD_GRAIN_YEAR
    if str(mappings.get("date_column") or "") != str(
        mappings.get("period_column") or ""
    ):
        return None
    # Raw date labels are mechanically identifiable when the date column is the
    # period column. Default them to a year grain for discrete composition
    # charts; recipes that intentionally need weeks must set period_grain.
    if _period_values_look_like_dates(values):
        return PERIOD_GRAIN_YEAR
    return None


def _configured_period_value(
    options: dict[str, Any],
    values: list[str],
    keys: Iterable[str],
) -> str | None:
    """Return the first configured period value that exists in canonical data."""

    available = set(values)
    for key in keys:
        value = options.get(key)
        if value in (None, ""):
            continue
        text = str(value)
        if text in available:
            return text
    return None


def _selected_periods_for_current(
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
) -> list[str]:
    """Return the configured or inferred current period."""

    values = _period_values(canonical)
    options = recipe.get("options") or {}
    configured = _configured_period_value(
        options,
        values,
        ("current_period_label", "current_period", "cohort_current_period"),
    )
    if configured is not None:
        return [configured]
    if CURRENT_PERIOD in values:
        return [CURRENT_PERIOD]
    return [values[-1]] if values else []


def _selected_periods_for_comparison(
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
) -> list[str]:
    """Return a legacy-friendly baseline/current period pair when available."""

    values = _period_values(canonical)
    options = recipe.get("options") or {}
    current = _configured_period_value(
        options,
        values,
        ("current_period_label", "current_period", "cohort_current_period"),
    )
    if current is None:
        current = CURRENT_PERIOD if CURRENT_PERIOD in values else None
    if current is None:
        current = values[-1] if values else CURRENT_PERIOD
    baseline = _configured_period_value(
        options,
        values,
        (
            "previous_period_label",
            "previous_period",
            "comparison_period",
            "baseline_period",
            "cohort_previous_period",
        ),
    )
    scenario_baseline, scenario_current = default_scenario_comparison_pair(values)
    if baseline is None and scenario_baseline and scenario_current == current:
        baseline = scenario_baseline
    if baseline is None:
        baseline = next(
            (
                value
                for value in values
                if normalize_name(value)
                in {"pl", "py", "plan", "fc", "forecast", "budget"}
            ),
            None,
        )
    if baseline is None:
        prior_values = [value for value in values if value != current]
        baseline = prior_values[-1] if prior_values else None
    return [baseline, current] if baseline and baseline != current else [current]


def _cohort_chart_periods(
    recipe: dict[str, Any],
    *,
    current_fallback: str,
    previous_fallback: str,
) -> tuple[str, str]:
    """Return the effective periods used by deterministic cohort chart specs."""

    contract = normalize_recipe_cohort_contract(recipe)
    periods = contract.get("periods") if isinstance(contract, dict) else {}
    if not isinstance(periods, dict):
        periods = {}
    current = periods.get("current_period") or current_fallback
    previous = periods.get("previous_period") or previous_fallback
    return str(current), str(previous)


def _ordered_cohort_period_values(
    canonical: pl.DataFrame,
    *,
    current_period: str,
    previous_period: str,
    period_grain: str | None,
) -> list[str]:
    values = _period_values_for_grain(canonical, period_grain)
    if (
        current_period in values
        and previous_period in values
        and values.index(previous_period) > values.index(current_period)
    ):
        other_values = [
            value for value in values if value not in {previous_period, current_period}
        ]
        return [*other_values, previous_period, current_period]
    return values


def _cohort_display_periods(
    canonical: pl.DataFrame,
    *,
    current_period: str,
    previous_period: str,
    period_grain: str | None,
) -> tuple[list[str], list[str], str | None]:
    """Return current + two prior periods, with older periods bucketed.

    This deterministic rule is mechanically verifiable from period activity:
    cohort charts need a stable, bounded display window rather than semantic
    judgment about which historical years matter.
    """

    periods = _ordered_cohort_period_values(
        canonical,
        current_period=current_period,
        previous_period=previous_period,
        period_grain=period_grain,
    )
    if not periods:
        return [current_period], [current_period], None
    if current_period in periods:
        end_index = periods.index(current_period) + 1
    else:
        end_index = len(periods)
    start_index = max(0, end_index - COHORT_VISIBLE_PERIOD_COUNT)
    visible_periods = periods[start_index:end_index]
    if not visible_periods:
        visible_periods = [current_period]
    before_label = (
        f"Before {visible_periods[0]}" if start_index > 0 and visible_periods else None
    )
    selected_periods = [*([before_label] if before_label else []), *visible_periods]
    return selected_periods, visible_periods, before_label


def _recent_period_window(
    periods: Sequence[str],
    *,
    current_period: str,
    window_size: int,
) -> list[str]:
    """Return a bounded recent period window ending at the current period."""

    values = [str(period) for period in periods if period]
    if not values:
        return []
    if current_period in values:
        end_index = values.index(current_period) + 1
    else:
        end_index = len(values)
    start_index = max(0, end_index - max(1, window_size))
    return values[start_index:end_index]


def _cohort_derived_chart_specs(
    canonical: pl.DataFrame,
    recipe: dict[str, Any],
) -> list[dict[str, str]]:
    """Return source/derived pairs for formal since/lost cohort charts."""

    contract = normalize_recipe_cohort_contract(recipe)
    columns, _schema = get_schema_and_column_names(canonical)
    chart_specs: list[dict[str, str]] = []
    for item in contract.get("derived_dimensions") or []:
        if not isinstance(item, dict):
            continue
        kind = str(item.get("kind") or item.get("cohort_mode") or "")
        if kind not in {"since", "lost"}:
            continue
        source_dimension = str(item.get("source_dimension") or "")
        cohort_dimension = str(item.get("name") or item.get("output_column") or "")
        if not cohort_dimension:
            cohort_dimension = (
                f"{source_dimension}_Since"
                if kind == "since"
                else f"{source_dimension}_Lost"
            )
        if source_dimension not in columns:
            continue
        chart_specs.append(
            {
                "kind": kind,
                "source_dimension": source_dimension,
                "cohort_dimension": cohort_dimension,
            }
        )
    return chart_specs


def _default_small_multiples_dimension(
    dimensions: list[str], primary: str, secondary: str | None
) -> str | None:
    """Return a useful default facet dimension for small-multiple charts."""

    available = [
        dimension for dimension in dimensions if dimension not in {primary, secondary}
    ]
    preferred_hints = [
        "region",
        "country",
        "market",
        "channel",
        "segment",
        "customer",
        "brand",
    ]
    for hint in preferred_hints:
        match = next(
            (dimension for dimension in available if hint in normalize_name(dimension)),
            None,
        )
        if match:
            return match
    return available[0] if available else None


def _default_stacked_bar_small_multiples_dimension(
    canonical: pl.DataFrame,
    dimensions: list[str],
    x_dimension: str,
    y_dimension: str | None,
    metric: str,
    selected_periods: list[str] | None,
) -> str | None:
    """Return a facet dimension that does not collapse stacked segments."""

    available = [
        dimension
        for dimension in dimensions
        if dimension not in {x_dimension, y_dimension}
    ]
    available = [
        dimension
        for dimension in available
        if _is_admissible_stacked_bar_small_multiples_dimension(
            canonical,
            x_dimension=x_dimension,
            y_dimension=y_dimension,
            facet_dimension=dimension,
            metric=metric,
            selected_periods=selected_periods,
        )
    ]
    if not available:
        return None
    preferred = _default_small_multiples_dimension(
        [*available, x_dimension, *([y_dimension] if y_dimension else [])],
        x_dimension,
        y_dimension,
    )
    return (
        preferred if preferred in available else (available[0] if available else None)
    )


def _current_period_frame(canonical: pl.DataFrame) -> pl.DataFrame:
    """Return the current period slice used for contribution facts."""

    values = _period_values(canonical)
    if CURRENT_PERIOD in values:
        return canonical.filter(
            pl.col(CANONICAL_PERIOD).cast(pl.Utf8) == CURRENT_PERIOD
        )
    return canonical


def _smart_mekko_named_item_count(
    canonical: pl.DataFrame,
    dimension: str,
    metric: str,
    selected_periods: Sequence[str],
    max_items: int,
) -> int:
    """Choose named Mekko items from exact contribution shares.

    This deterministic rule is appropriate because label capacity is a
    mechanical layout constraint: visible area share, not semantic judgment,
    determines whether another category can carry a readable label.
    """

    if (
        not dimension
        or dimension not in canonical.columns
        or metric not in canonical.columns
    ):
        return max(1, min(max_items, MEKKO_MAX_NAMED_ITEMS))
    max_named = max(1, min(int(max_items), MEKKO_MAX_NAMED_ITEMS))
    frame = canonical
    if CANONICAL_PERIOD in frame.columns and selected_periods:
        frame = frame.filter(
            pl.col(CANONICAL_PERIOD)
            .cast(pl.Utf8)
            .is_in([str(period) for period in selected_periods])
        )
    if frame.is_empty():
        return max_named

    totals = (
        frame.group_by(dimension)
        .agg(pl.col(metric).sum().abs().alias("__metric_total"))
        .sort("__metric_total", descending=True)
    )
    if totals.is_empty():
        return max_named
    values = [
        float(value or 0.0)
        for value in totals.get_column("__metric_total").to_list()
        if float(value or 0.0) > 0.0
    ]
    if not values:
        return min(max_named, max(MEKKO_MIN_NAMED_ITEMS, 1))

    total = sum(values)
    if total <= 0:
        return min(max_named, len(values))
    keep_floor = min(MEKKO_MIN_NAMED_ITEMS, len(values), max_named)
    count = 0
    cumulative = 0.0
    for value in values:
        if count >= max_named:
            break
        share = value / total
        if count >= keep_floor and (
            cumulative >= MEKKO_CUMULATIVE_SHARE or share < MEKKO_MIN_ITEM_SHARE
        ):
            break
        cumulative += share
        count += 1
    return max(keep_floor, count)


def _mekko_axis_limit_overrides(
    canonical: pl.DataFrame,
    *,
    x_dimension: str,
    w_dimension: str,
    metric: str,
    selected_periods: Sequence[str],
    max_items: int,
) -> dict[str, int]:
    """Return axis-specific Mekko limits for row and segment dimensions."""

    return {
        "x_max_items": _smart_mekko_named_item_count(
            canonical, x_dimension, metric, selected_periods, max_items
        ),
        "w_max_items": _smart_mekko_named_item_count(
            canonical, w_dimension, metric, selected_periods, max_items
        ),
    }


def _focus_request_items(value: Any) -> list[str]:
    """Return capped, non-empty requested focus items as strings."""

    if value in (None, "", [], {}):
        return []
    if isinstance(value, (list, tuple, set)):
        raw_items = list(value)
    else:
        raw_items = [value]
    items: list[str] = []
    for item in raw_items:
        text = str(item).strip()
        if text:
            items.append(text)
    return items[:FOCUS_ITEM_HIGHLIGHT_MAX_ITEMS]


def _focus_dimension_request(value: Any) -> str | None:
    """Return the requested focus dimension when provided."""

    if value in (None, "", [], {}):
        return None
    text = str(value).strip()
    return text or None


def _focus_candidate_dimensions(
    spec: dict[str, Any],
    *,
    requested_dimension: str | None,
    primary_dimension: str,
) -> list[str]:
    """Return chart dimensions whose items can receive focus highlighting."""

    chart_name = str(spec.get("name") or "")
    if chart_name in {"stacked_bar", "stacked_bar_small_multiples"}:
        return _unique_dimensions([spec.get("y_dimension")])
    if chart_name in {"marimekko", "marimekko_small_multiples"}:
        return _unique_dimensions([spec.get("y_dimension")])
    if chart_name == "stacked_column":
        return _unique_dimensions([spec.get("x_dimension")])
    if chart_name in {"area_absolute", "area_share"}:
        return _unique_dimensions([*(spec.get("dimensions") or []), primary_dimension])
    if chart_name == "line":
        return _unique_dimensions(
            [
                requested_dimension,
                *(spec.get("dimensions") or []),
                spec.get("x_dimension"),
                primary_dimension,
            ]
        )
    return []


def _dimension_values_as_strings(canonical: pl.DataFrame, dimension: str) -> list[str]:
    """Return canonical dimension values as stable strings for focus matching."""

    columns, _schema = get_schema_and_column_names(canonical)
    if dimension not in columns:
        return []
    values = (
        canonical.select(pl.col(dimension).drop_nulls().cast(pl.Utf8).unique())
        .get_column(dimension)
        .to_list()
    )
    deduped: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = str(value).strip()
        if text and text not in seen:
            deduped.append(text)
            seen.add(text)
    return deduped


def _resolve_focus_item(
    canonical: pl.DataFrame,
    *,
    item: str,
    dimension: str,
) -> tuple[str | None, str]:
    """Resolve a requested focus item to the canonical value for a dimension."""

    candidates = _dimension_values_as_strings(canonical, dimension)
    if item in candidates:
        return item, "matched_exact"
    folded_item = item.casefold()
    folded_matches = [
        candidate for candidate in candidates if candidate.casefold() == folded_item
    ]
    unique_matches = list(dict.fromkeys(folded_matches))
    if len(unique_matches) == 1:
        return unique_matches[0], "matched_case_insensitive"
    if len(unique_matches) > 1:
        return None, "focus_item_ambiguous_case_insensitive_match"
    return None, "focus_item_not_found"


def _focus_metadata_for_spec(
    canonical: pl.DataFrame,
    spec: dict[str, Any],
    *,
    focus_items: Sequence[str],
    focus_dimension: str | None,
    available_dimensions: Sequence[str],
    primary_dimension: str,
) -> dict[str, Any]:
    """Return focus metadata for one spec when a focus item was requested."""

    if not focus_items:
        return {}
    requested_item = focus_items[0]
    chart_name = str(spec.get("name") or "")
    if chart_name in FOCUS_ITEM_UNSUPPORTED_CHARTS:
        return {
            "focus_item": requested_item,
            "focus_dimension": focus_dimension or spec.get("y_dimension"),
            "focus_status": "unsupported",
            "focus_reason": "legacy_chart_has_no_focus_highlight",
        }
    if chart_name not in FOCUS_ITEM_RENDER_ELIGIBLE_CHARTS:
        return {}
    candidates = _focus_candidate_dimensions(
        spec,
        requested_dimension=focus_dimension,
        primary_dimension=primary_dimension,
    )
    if focus_dimension and focus_dimension not in available_dimensions:
        return {
            "focus_item": requested_item,
            "focus_dimension": focus_dimension,
            "focus_status": "unresolved",
            "focus_reason": "focus_dimension_not_available",
        }
    if not candidates:
        return {
            "focus_item": requested_item,
            "focus_dimension": focus_dimension,
            "focus_status": "unresolved",
            "focus_reason": "focus_dimension_not_active",
        }
    if focus_dimension:
        if focus_dimension not in candidates:
            return {
                "focus_item": requested_item,
                "focus_dimension": focus_dimension,
                "focus_status": "unresolved",
                "focus_reason": "focus_dimension_not_active",
            }
        resolved_dimension = focus_dimension
    else:
        resolved_dimension = candidates[0]
    resolved_item, reason = _resolve_focus_item(
        canonical,
        item=requested_item,
        dimension=resolved_dimension,
    )
    if resolved_item is None:
        return {
            "focus_item": requested_item,
            "focus_dimension": resolved_dimension,
            "focus_status": "unresolved",
            "focus_reason": reason,
        }
    return {
        "focus_item": resolved_item,
        "focus_dimension": resolved_dimension,
        "focus_status": "resolved",
        "focus_reason": reason,
    }


def _pareto_show_only_option(options: dict[str, Any]) -> str:
    """Return the legacy Pareto display mode requested by the recipe."""

    raw_value = (
        options.get("pareto_show_only")
        or options.get("pareto_display_mode")
        or options.get("show_only")
    )
    if raw_value in (None, "", [], {}):
        return PARETO_DEFAULT_SHOW_ONLY
    normalized = str(raw_value).strip().lower().replace(" ", "_")
    return PARETO_SHOW_ONLY_ALIASES.get(normalized, PARETO_DEFAULT_SHOW_ONLY)


def build_chart_specs(
    canonical: pl.DataFrame, recipe: dict[str, Any]
) -> list[dict[str, Any]]:
    """Build the complete legacy chart attempt list."""

    options = recipe.get("options") or {}
    requested = set(_normalize_requested_chart_names(options.get("charts") or []))
    focus_items = _focus_request_items(options.get("focus_item"))
    focus_dimension = _focus_dimension_request(options.get("focus_dimension"))
    dimensions = _useful_dimensions(
        canonical, [str(item) for item in recipe["mappings"].get("dimensions") or []]
    )
    if not dimensions:
        return []
    primary = dimensions[0]
    secondary = dimensions[1] if len(dimensions) > 1 else None
    mekko_pair = _find_mekko_dimension_pair(canonical, dimensions)
    mekko_small_dimensions = _find_mekko_small_multiple_dimensions(
        canonical, dimensions
    )
    selected_current_period = _selected_periods_for_current(canonical, recipe)
    actual_period_display_label = _actual_date_range_period_label(
        canonical, selected_current_period
    )
    selected_comparison_periods = _selected_periods_for_comparison(canonical, recipe)
    selected_time_chart_periods = _selected_periods_for_time_charts(
        canonical,
        recipe,
        selected_current_period,
    )
    stacked_column_period_grain = _default_period_grain_for_stacked_column(
        canonical, recipe
    )
    stacked_column_period_values = _period_values_for_grain(
        canonical, stacked_column_period_grain
    )
    metric = str(recipe["mappings"]["amount_column"])
    width_metric = recipe["mappings"].get("legacy_width_metric_column")
    pareto_dimension, pareto_parent_dimension, pareto_dimension_selection = (
        _find_pareto_dimensions(canonical, dimensions)
    )
    (
        stacked_pareto_dimension,
        stacked_pareto_parent_dimension,
        stacked_pareto_dimension_selection,
    ) = _find_stacked_pareto_dimensions(canonical, dimensions)
    pareto_metric_columns = _pareto_metrics(
        canonical, metric, str(width_metric) if width_metric else None
    )
    related_marker_metric = _related_marker_metric(
        canonical,
        metric,
        str(width_metric) if width_metric else None,
        selected_comparison_periods,
    )
    explicit_marker_metric = recipe["mappings"].get("related_marker_metric_column")
    canonical_columns, _canonical_schema = get_schema_and_column_names(canonical)
    if explicit_marker_metric and explicit_marker_metric in canonical_columns:
        related_marker_metric = str(explicit_marker_metric)
    related_dimension = _default_related_metrics_dimension(
        canonical, dimensions, primary
    )
    related_small_dimension = options.get("small_multiples_dimension")
    if (
        related_small_dimension not in dimensions
        or related_small_dimension == related_dimension
    ):
        related_small_dimension = _default_small_multiples_dimension(
            dimensions, related_dimension, None
        )
    small_dimension = options.get("small_multiples_dimension")
    if small_dimension not in dimensions:
        small_dimension = _default_small_multiples_dimension(
            dimensions, primary, secondary
        )
    enable_small = bool(options.get("small_multiples") and small_dimension)
    if mekko_pair is not None:
        stacked_x_dimension, stacked_y_dimension = mekko_pair
    else:
        stacked_x_dimension, stacked_y_dimension = primary, secondary
    stacked_small_dimension = options.get("small_multiples_dimension")
    if stacked_small_dimension not in dimensions or stacked_small_dimension in {
        stacked_x_dimension,
        stacked_y_dimension,
    }:
        stacked_small_dimension = _default_stacked_bar_small_multiples_dimension(
            canonical,
            dimensions,
            stacked_x_dimension,
            stacked_y_dimension,
            metric,
            selected_current_period,
        )
    elif not _is_admissible_stacked_bar_small_multiples_dimension(
        canonical,
        x_dimension=stacked_x_dimension,
        y_dimension=stacked_y_dimension,
        facet_dimension=str(stacked_small_dimension),
        metric=metric,
        selected_periods=selected_current_period,
    ):
        stacked_small_dimension = None
    enable_stacked_small = bool(
        options.get("small_multiples") and stacked_small_dimension
    )
    max_items = int(options.get("max_chart_items") or 12)
    area_max_items = min(max_items, AREA_CHART_MAX_NAMED_ITEMS)
    stacked_column_max_items = min(max_items, STACKED_COLUMN_MAX_NAMED_ITEMS)
    mekko_axis_overrides: dict[str, int] = {}
    if mekko_pair is not None:
        mekko_axis_overrides = _mekko_axis_limit_overrides(
            canonical,
            x_dimension=mekko_pair[0],
            w_dimension=mekko_pair[1],
            metric=metric,
            selected_periods=selected_current_period or [CURRENT_PERIOD],
            max_items=max_items,
        )
    small_multiples_max_panels = int(options.get("small_multiples_max_panels") or 6)
    dimension_display_labels = options.get("dimension_display_labels") or options.get(
        "display_dimension_labels"
    )
    if isinstance(dimension_display_labels, dict):
        dimension_display_labels = {
            str(key): str(value)
            for key, value in dimension_display_labels.items()
            if str(key).strip() and str(value).strip()
        }
    else:
        dimension_display_labels = {}
    reporting_entity_label = reporting_entity_label_from_recipe(recipe)
    reporting_subject_label = reporting_subject_label_from_recipe(recipe)
    common = {
        "max_items": max_items,
        "small_multiples_max_panels": small_multiples_max_panels,
        "dimensions": dimensions[:3],
        "x_dimension": primary,
        "y_dimension": secondary,
        "metric": metric,
        "metrics": [metric],
        "selected_periods": selected_current_period or [CURRENT_PERIOD],
        "period_grain": options.get("period_grain"),
        "period_window": options.get("period_window"),
        "period_display_label": actual_period_display_label,
        "period_to_date": bool(options.get("period_to_date", False)),
        "reporting_entity_label": reporting_entity_label,
        "reporting_subject_label": reporting_subject_label,
        "dimension_display_labels": dimension_display_labels,
    }
    pareto_show_only = _pareto_show_only_option(options)
    specs: list[dict[str, Any]] = []

    def add(
        name: str,
        plotter: str,
        chart_key: str,
        artifact: str,
        *,
        base_chart: str | None = None,
        select_with_base_chart: bool = True,
        **extra: Any,
    ) -> None:
        if (
            requested
            and name not in requested
            and not (
                select_with_base_chart
                and base_chart is not None
                and base_chart in requested
            )
        ):
            return
        spec = {
            "name": name,
            "plotter": plotter,
            "legacy_chart_key": chart_key,
            "artifact_name": artifact,
            **common,
            **extra,
        }
        if base_chart is not None:
            spec["base_chart"] = base_chart
        focus_metadata = _focus_metadata_for_spec(
            canonical,
            spec,
            focus_items=focus_items,
            focus_dimension=focus_dimension,
            available_dimensions=dimensions,
            primary_dimension=primary,
        )
        if focus_metadata:
            spec.update(focus_metadata)
        specs.append(spec)

    current_period, previous_period = _cohort_chart_periods(
        recipe,
        current_fallback=(selected_current_period or [CURRENT_PERIOD])[-1],
        previous_fallback=(
            selected_comparison_periods[0]
            if selected_comparison_periods
            else CURRENT_PERIOD
        ),
    )
    cohort_contract = normalize_recipe_cohort_contract(recipe)
    has_like_for_like = isinstance(cohort_contract.get("like_for_like"), dict)
    like_for_like_source_dimension = None
    if has_like_for_like:
        like_for_like_contract = cohort_contract.get("like_for_like") or {}
        like_for_like_source_dimension = (
            str(
                like_for_like_contract.get("source_dimension")
                or like_for_like_contract.get("entity_dimension")
                or ""
            ).strip()
            or None
        )
    period_selection_mode = str(options.get("period_selection") or "").strip().lower()
    uses_explicit_comparison_periods = period_selection_mode in {
        "explicit_comparison_periods",
        "explicit_cohort_periods",
    }
    explicit_current_period = str(
        options.get("current_period_label")
        or options.get("current_period")
        or options.get("cohort_current_period")
        or current_period
    )
    explicit_previous_period = str(
        options.get("previous_period_label")
        or options.get("previous_period")
        or options.get("comparison_period")
        or options.get("baseline_period")
        or options.get("cohort_previous_period")
        or previous_period
    )
    explicit_comparison_periods = (
        [explicit_previous_period, explicit_current_period]
        if explicit_previous_period
        and explicit_previous_period != explicit_current_period
        else [explicit_current_period]
    )
    stacked_column_period_selection = (
        explicit_comparison_periods
        if uses_explicit_comparison_periods
        else (
            stacked_column_period_values
            if stacked_column_period_grain
            else selected_comparison_periods
        )
    )
    stacked_column_period_choice = stacked_column_period_grain or options.get(
        "period_grain"
    )
    like_for_like_stacked_period_selection = stacked_column_period_selection
    if len(stacked_column_period_values) >= 3:
        like_for_like_stacked_period_selection = _recent_period_window(
            stacked_column_period_values,
            current_period=str(
                (stacked_column_period_selection or stacked_column_period_values)[-1]
            ),
            window_size=3,
        )
    stacked_column_period_mode = (
        "comparison_periods"
        if uses_explicit_comparison_periods
        else (
            "all_periods_at_grain"
            if stacked_column_period_grain
            else "comparison_periods"
        )
    )
    stacked_column_dimension_selection = (
        "primary_dimension_comparison_period_axis_at_grain"
        if uses_explicit_comparison_periods and stacked_column_period_grain
        else (
            "primary_dimension_period_axis_at_grain"
            if stacked_column_period_grain
            else "primary_dimension_period_axis"
        )
    )
    stacked_column_show_cagr = _can_show_cagr(metric, stacked_column_period_selection)
    stacked_column_show_total_cagr = _can_show_total_cagr(
        metric, stacked_column_period_selection
    )
    (
        cohort_period_selection,
        cohort_visible_periods,
        cohort_before_period_label,
    ) = _cohort_display_periods(
        canonical,
        current_period=current_period,
        previous_period=previous_period,
        period_grain=stacked_column_period_grain,
    )
    for cohort_spec in _cohort_derived_chart_specs(canonical, recipe):
        kind = cohort_spec["kind"]
        source_dimension = cohort_spec["source_dimension"]
        cohort_dimension = (
            f"{source_dimension}_Since"
            if kind == "since"
            else f"{source_dimension}_Lost"
        )
        selected_period = previous_period if kind == "lost" else current_period
        add(
            f"cohort_{kind}_stacked_column",
            "plot_stacked_column_charts",
            "stackedColumnChart",
            f"cohort_{kind}_stacked_column.png",
            dimensions=[cohort_dimension],
            x_dimension=CANONICAL_PERIOD,
            y_dimension=cohort_dimension,
            selected_periods=cohort_period_selection,
            period_grain=stacked_column_period_grain,
            period_selection_mode="cohort_recent_periods_with_before_bucket",
            capture_chart_data=True,
            dimension_selection=f"{kind}_cohort_stacked_column",
            show_cagr=False,
            suppress_stacked_percentage_annotations=True,
            format_stacked_value_labels_like_totals=True,
            suppress_zero_rounded_stacked_labels=(kind == "lost"),
            suppress_single_active_value_label=(kind == "lost"),
            display_dimension_label=(
                "First active year" if kind == "since" else "Last active year"
            ),
            cohort_kind=kind,
            cohort_reference_period=selected_period,
            cohort_dimension=cohort_dimension,
            cohort_source_dimension=source_dimension,
            cohort_activity_metric=metric,
            cohort_visible_period_count=COHORT_VISIBLE_PERIOD_COUNT,
            cohort_visible_periods=cohort_visible_periods,
            cohort_before_period_label=cohort_before_period_label,
            chosen_cohort_column=source_dimension if kind == "since" else None,
            lost_and_dropped_column=source_dimension if kind == "lost" else None,
        )

    add(
        "column_total",
        "plot_stacked_column_charts",
        "stackedColumnChart",
        "column_total.png",
        dimensions=[LEGACY_TOTAL_COLUMN_DIMENSION],
        x_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
        y_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
        selected_periods=stacked_column_period_selection,
        period_grain=stacked_column_period_choice,
        period_selection_mode=stacked_column_period_mode,
        capture_chart_data=True,
        capture_figure="last",
        dimension_selection="legacy_total_column_period_axis",
        show_cagr=stacked_column_show_cagr,
        total_column_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
        total_column_label=LEGACY_TOTAL_COLUMN_LABEL,
    )
    if related_marker_metric:
        related_value_cols = _related_metric_value_cols(metric, related_marker_metric)
        add(
            "column_total_with_overlay",
            "plot_stacked_column_charts",
            "stackedColumnChart",
            "column_total_with_overlay.png",
            base_chart="column_total",
            select_with_base_chart=False,
            dimensions=[LEGACY_TOTAL_COLUMN_DIMENSION],
            x_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
            y_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
            selected_periods=stacked_column_period_selection,
            period_grain=stacked_column_period_choice,
            period_selection_mode=stacked_column_period_mode,
            metrics=[metric, related_marker_metric],
            value_cols=related_value_cols,
            plot_overlay_chart=True,
            highlight_overlay_chart=True,
            capture_chart_data=True,
            capture_figure="last",
            dimension_selection="legacy_total_column_period_axis_with_overlay",
            show_cagr=stacked_column_show_cagr,
            total_column_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
            total_column_label=LEGACY_TOTAL_COLUMN_LABEL,
        )
    if has_like_for_like:
        add(
            "like_for_like_column_total",
            "plot_stacked_column_charts",
            "stackedColumnChart",
            "like_for_like_column_total.png",
            base_chart="column_total",
            select_with_base_chart=False,
            dimensions=[LEGACY_TOTAL_COLUMN_DIMENSION],
            x_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
            y_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
            selected_periods=stacked_column_period_selection,
            period_grain=stacked_column_period_choice,
            period_selection_mode=stacked_column_period_mode,
            capture_chart_data=True,
            capture_figure="last",
            dimension_selection="like_for_like_total_column_period_axis",
            show_cagr=stacked_column_show_cagr,
            population_mode="like_for_like",
            population_dimension=like_for_like_source_dimension,
            total_column_dimension=LEGACY_TOTAL_COLUMN_DIMENSION,
            total_column_label=LEGACY_TOTAL_COLUMN_LABEL,
        )

    add(
        "bar",
        "plot_stacked_bar_charts",
        "stackedBarChart",
        "bar.png",
        dimensions=[related_dimension],
        x_dimension=related_dimension,
        y_dimension=None,
        selected_periods=selected_comparison_periods,
        capture_chart_data=True,
        dimension_selection="detailed_dimension_plain_bar",
        show_average_value=True,
    )
    if (
        options.get("small_multiples")
        and related_small_dimension
        and related_small_dimension != related_dimension
    ):
        add(
            "bar_small_multiples",
            "plot_stacked_bar_charts",
            "stackedBarChart",
            "bar_small_multiples.png",
            base_chart="bar",
            dimensions=_unique_dimensions([related_dimension, related_small_dimension]),
            x_dimension=related_dimension,
            y_dimension=None,
            selected_periods=selected_comparison_periods,
            small_multiples_dimension=related_small_dimension,
            small_multiples_panel_axis="Y",
            capture_chart_data=True,
            capture_figure="last",
            dimension_selection="detailed_dimension_plain_bar_by_panel",
        )
    if related_marker_metric:
        related_value_cols = _related_metric_value_cols(metric, related_marker_metric)
        add(
            "related_metrics_bar",
            "plot_stacked_bar_charts",
            "stackedBarChart",
            "related_metrics_bar.png",
            dimensions=[related_dimension],
            x_dimension=related_dimension,
            y_dimension=None,
            selected_periods=selected_comparison_periods,
            metrics=[metric, related_marker_metric],
            value_cols=related_value_cols,
            plot_overlay_chart=True,
            related_metrics_bar=True,
            capture_chart_data=True,
            dimension_selection="detailed_dimension_with_related_marker_metric",
            show_average_value=True,
        )
        if (
            options.get("small_multiples")
            and related_small_dimension
            and related_small_dimension != related_dimension
        ):
            add(
                "related_metrics_bar_small_multiples",
                "plot_stacked_bar_charts",
                "stackedBarChart",
                "related_metrics_bar_small_multiples.png",
                base_chart="related_metrics_bar",
                dimensions=_unique_dimensions(
                    [related_dimension, related_small_dimension]
                ),
                x_dimension=related_dimension,
                y_dimension=None,
                selected_periods=selected_comparison_periods,
                metrics=[metric, related_marker_metric],
                value_cols=related_value_cols,
                small_multiples_dimension=related_small_dimension,
                small_multiples_panel_axis="Y",
                plot_overlay_chart=True,
                related_metrics_bar=True,
                capture_chart_data=True,
                capture_figure="last",
                dimension_selection=(
                    "detailed_dimension_with_related_marker_metric_by_panel"
                ),
            )

    if secondary:
        if mekko_pair is not None:
            mekko_primary, mekko_secondary = mekko_pair
            mekko_dimensions = [
                mekko_primary,
                mekko_secondary,
                *[
                    dimension
                    for dimension in dimensions
                    if dimension not in {mekko_primary, mekko_secondary}
                ],
            ][:3]
            add(
                "marimekko",
                "plot_mekko_charts",
                "marimekkoChart",
                "marimekko.png",
                dimensions=mekko_dimensions,
                x_dimension=mekko_primary,
                y_dimension=mekko_secondary,
                dimension_selection=("first_non_hierarchical_observed_dimension_pair"),
                show_legend_mode="inside",
                value_label_mode="absolute",
                **mekko_axis_overrides,
            )
        if width_metric:
            if mekko_pair is not None:
                add(
                    "barmekko",
                    "plot_mekko_charts",
                    "barmekkoChart",
                    "barmekko.png",
                    dimensions=mekko_dimensions,
                    x_dimension=mekko_primary,
                    y_dimension=mekko_secondary,
                    dimension_selection=(
                        "first_non_hierarchical_observed_dimension_pair"
                    ),
                    x_metric=str(width_metric),
                    y_metric=LEGACY_UNIT_PRICE_COLUMN,
                    multiplied_metric=LEGACY_MULTIPLIED_COLUMN,
                    metrics=[LEGACY_UNIT_PRICE_COLUMN, str(width_metric)],
                    value_cols=[metric, str(width_metric), LEGACY_MULTIPLIED_COLUMN],
                    show_legend_mode="inside",
                    value_label_mode="absolute",
                    **mekko_axis_overrides,
                )
        add(
            "stacked_bar",
            "plot_stacked_bar_charts",
            "stackedBarChart",
            "stacked_bar.png",
            dimensions=_unique_dimensions([stacked_x_dimension, stacked_y_dimension]),
            x_dimension=stacked_x_dimension,
            y_dimension=stacked_y_dimension,
            capture_chart_data=True,
            dimension_selection="first_non_hierarchical_observed_dimension_pair",
            show_legend_mode="inside",
        )
        if enable_stacked_small and stacked_small_dimension:
            add(
                "stacked_bar_small_multiples",
                "plot_stacked_bar_charts",
                "stackedBarChart",
                "stacked_bar_small_multiples.png",
                base_chart="stacked_bar",
                dimensions=_unique_dimensions(
                    [
                        stacked_x_dimension,
                        stacked_y_dimension,
                        stacked_small_dimension,
                    ]
                ),
                x_dimension=stacked_x_dimension,
                y_dimension=stacked_y_dimension,
                small_multiples_dimension=stacked_small_dimension,
                small_multiples_panel_axis="Y",
                capture_chart_data=True,
                capture_figure="last",
                dimension_selection="two_dimension_stacked_bar_by_panel",
                show_top_for_each_item=True,
                show_legend_mode="inside",
            )
        add(
            "stacked_pareto_abc",
            "plot_stacked_pareto_chart",
            "stackedParetoChart",
            "stacked_pareto_abc.png",
            base_chart="stacked_pareto",
            dimensions=_unique_dimensions([stacked_pareto_dimension, *dimensions]),
            x_dimension=stacked_pareto_dimension,
            y_dimension=None,
            count_dimension=stacked_pareto_dimension,
            aggregate_uniques_by_dimension=False,
            metrics=pareto_metric_columns,
            value_cols=pareto_metric_columns,
            capture_chart_data=True,
            dimension_selection=f"abc_classes_{stacked_pareto_dimension_selection}",
            stacked_pareto_mode="abc_classes",
        )
        if stacked_pareto_parent_dimension:
            add(
                "stacked_pareto_by_dimension",
                "plot_stacked_pareto_chart",
                "stackedParetoChart",
                "stacked_pareto_by_dimension.png",
                base_chart="stacked_pareto",
                dimensions=_unique_dimensions(
                    [
                        stacked_pareto_dimension,
                        stacked_pareto_parent_dimension,
                        *dimensions,
                    ]
                ),
                x_dimension=stacked_pareto_dimension,
                y_dimension=stacked_pareto_parent_dimension,
                count_dimension=stacked_pareto_dimension,
                aggregate_uniques_by_dimension=True,
                aggregate_uniques_dimension=stacked_pareto_parent_dimension,
                aggregate_other_items=True,
                metrics=pareto_metric_columns,
                value_cols=pareto_metric_columns,
                capture_chart_data=True,
                dimension_selection=stacked_pareto_dimension_selection,
                stacked_pareto_mode="aggregate_by_dimension",
            )
        if enable_small and mekko_small_dimensions is not None:
            (
                mekko_small_primary,
                mekko_small_secondary,
                mekko_small_dimension,
            ) = mekko_small_dimensions
            mekko_small_dimensions_array = [
                mekko_small_primary,
                mekko_small_secondary,
                mekko_small_dimension,
            ]
            add(
                "marimekko_small_multiples",
                "plot_mekko_charts",
                "marimekkoChart",
                "marimekko_small_multiples.png",
                base_chart="marimekko",
                dimensions=mekko_small_dimensions_array,
                x_dimension=mekko_small_primary,
                y_dimension=mekko_small_secondary,
                small_multiples_dimension=mekko_small_dimension,
                small_multiples_panel_axis="Y",
                dimension_selection=(
                    "small_multiple_facet_with_non_hierarchical_observed_axes"
                ),
                show_legend_mode="inside",
                value_label_mode="absolute",
                **mekko_axis_overrides,
            )
            if width_metric:
                add(
                    "barmekko_small_multiples",
                    "plot_mekko_charts",
                    "barmekkoChart",
                    "barmekko_small_multiples.png",
                    base_chart="barmekko",
                    dimensions=mekko_small_dimensions_array,
                    x_dimension=mekko_small_primary,
                    y_dimension=mekko_small_secondary,
                    small_multiples_dimension=mekko_small_dimension,
                    small_multiples_panel_axis="Y",
                    dimension_selection=(
                        "small_multiple_facet_with_non_hierarchical_observed_axes"
                    ),
                    x_metric=str(width_metric),
                    y_metric=LEGACY_UNIT_PRICE_COLUMN,
                    multiplied_metric=LEGACY_MULTIPLIED_COLUMN,
                    metrics=[LEGACY_UNIT_PRICE_COLUMN, str(width_metric)],
                    value_cols=[metric, str(width_metric), LEGACY_MULTIPLIED_COLUMN],
                    show_legend_mode="inside",
                    value_label_mode="absolute",
                    **mekko_axis_overrides,
                )
    add(
        "stacked_column",
        "plot_stacked_column_charts",
        "stackedColumnChart",
        "stacked_column.png",
        dimensions=[primary],
        x_dimension=primary,
        y_dimension=primary,
        selected_periods=stacked_column_period_selection,
        period_grain=stacked_column_period_choice,
        period_selection_mode=stacked_column_period_mode,
        max_items=stacked_column_max_items,
        capture_chart_data=True,
        dimension_selection=stacked_column_dimension_selection,
        show_cagr=stacked_column_show_cagr,
        show_total_cagr=stacked_column_show_total_cagr,
        suppress_stacked_percentage_annotations=True,
        format_stacked_value_labels_like_totals=True,
    )
    if options.get("small_multiples") or "stacked_column_small_multiples" in requested:
        add(
            "stacked_column_small_multiples",
            "plot_stacked_column_charts",
            "stackedColumnChart",
            "stacked_column_small_multiples.png",
            base_chart="stacked_column",
            dimensions=[primary],
            x_dimension=primary,
            y_dimension=primary,
            selected_periods=stacked_column_period_selection,
            period_grain=stacked_column_period_choice,
            period_selection_mode=stacked_column_period_mode,
            max_items=stacked_column_max_items,
            small_multiples_dimension=primary,
            small_multiples_panel_axis="X",
            capture_chart_data=True,
            capture_figure="last",
            dimension_selection=f"{stacked_column_dimension_selection}_by_panel",
            show_cagr=stacked_column_show_cagr,
            show_total_cagr=stacked_column_show_total_cagr,
            suppress_stacked_percentage_annotations=True,
            format_stacked_value_labels_like_totals=True,
        )
    if has_like_for_like:
        like_for_like_stacked_show_cagr = _can_show_cagr(
            metric, like_for_like_stacked_period_selection
        )
        add(
            "like_for_like_stacked_column",
            "plot_stacked_column_charts",
            "stackedColumnChart",
            "like_for_like_stacked_column.png",
            base_chart="stacked_column",
            select_with_base_chart=False,
            dimensions=[primary],
            x_dimension=primary,
            y_dimension=primary,
            selected_periods=like_for_like_stacked_period_selection,
            period_grain=stacked_column_period_choice,
            period_selection_mode=stacked_column_period_mode,
            max_items=stacked_column_max_items,
            capture_chart_data=True,
            dimension_selection=f"like_for_like_{stacked_column_dimension_selection}",
            show_cagr=like_for_like_stacked_show_cagr,
            show_total_cagr=_can_show_total_cagr(
                metric, like_for_like_stacked_period_selection
            ),
            suppress_stacked_percentage_annotations=True,
            format_stacked_value_labels_like_totals=True,
            population_mode="like_for_like",
            population_dimension=like_for_like_source_dimension,
        )
    if len(dimensions) >= 2:
        add(
            "stacked_column_synthesis",
            "plot_stacked_column_charts",
            "stackedColumnChart",
            "stacked_column_synthesis.png",
            dimensions=dimensions[:6],
            x_dimension=primary,
            y_dimension=primary,
            # Synthesis has no period axis, but it must still use the grained
            # selection so raw weekly/daily dates do not become report labels.
            selected_periods=(
                stacked_column_period_selection
                if stacked_column_period_grain
                else selected_current_period
            )
            or [CURRENT_PERIOD],
            period_grain=stacked_column_period_choice,
            period_selection_mode=stacked_column_period_mode,
            synthesis_plot=True,
            capture_figure="last",
            capture_chart_data=True,
            synthesis_uniform_palette=True,
            dimension_selection=("legacy_summary_stacked_column_multi_dimension"),
        )
    add(
        "pareto",
        "plot_pareto_chart",
        "paretoChart",
        "pareto.png",
        dimensions=_unique_dimensions([pareto_dimension, *dimensions]),
        x_dimension=pareto_dimension,
        y_dimension=None,
        metrics=pareto_metric_columns,
        value_cols=pareto_metric_columns,
        dimension_selection=pareto_dimension_selection,
        show_absolute_values=False,
        show_rank=pareto_show_only == "All",
        show_only=pareto_show_only,
    )
    if selected_time_chart_periods:
        add(
            "area_absolute",
            "plot_area_charts",
            "areaChart",
            "area_absolute.png",
            dimensions=[primary],
            y_dimension=None,
            selected_periods=selected_time_chart_periods,
            max_items=area_max_items,
            dimension_selection="ranked_area_with_fragmentation_other",
        )
        add(
            "area_share",
            "plot_area_charts",
            "areaChart",
            "area_share.png",
            dimensions=[primary],
            y_dimension=None,
            selected_periods=selected_time_chart_periods,
            max_items=area_max_items,
            dimension_selection="ranked_area_with_fragmentation_other",
            share_view=True,
        )
        add(
            "line",
            "plot_timeline_charts",
            "timelineChart",
            "line.png",
            dimensions=[],
            x_dimension=None,
            y_dimension=None,
            selected_periods=selected_time_chart_periods,
        )
    if selected_time_chart_periods and (
        enable_small or "line_small_multiples" in requested
    ):
        add(
            "line_small_multiples",
            "plot_timeline_charts",
            "timelineChart",
            "line_small_multiples.png",
            base_chart="line",
            dimensions=[primary],
            x_dimension=primary,
            y_dimension=None,
            selected_periods=selected_time_chart_periods,
            small_multiples_dimension=primary,
            small_multiples_panel_axis="X",
            capture_figure="last",
            dimension_selection="primary_dimension_timeline_facets",
        )
    if _has_comparison_periods(selected_comparison_periods):
        add(
            "multitier_bar",
            "plot_multitier_bar_chart",
            "multitierBarChart",
            "multitier_bar.png",
            dimensions=[primary],
            x_dimension=primary,
            y_dimension=None,
            selected_periods=selected_comparison_periods,
            capture_chart_data=True,
            capture_figure="last",
            dimension_selection="ranked_primary_dimension_multitier_bar",
        )
        small_multitier_dimensions: list[str] = []
        if enable_small:
            small_multitier_dimensions = _unique_dimensions(
                dimensions,
                limit=small_multiples_max_panels,
            )
            dimension_panel_max_items = int(
                options.get("dimension_panel_max_items")
                or options.get("dimension_panel_x_max_items")
                or min(max_items, 6)
            )
            dimension_panel_aggregate_other_items = bool(
                options.get(
                    "dimension_panel_aggregate_other_items",
                    options.get("aggregate_other_items", False),
                )
            )
            add(
                "multitier_bar_dimension_panels",
                "plot_multitier_bar_chart",
                "multitierBarChart",
                "multitier_bar_dimension_panels.png",
                base_chart="multitier_bar",
                dimensions=small_multitier_dimensions,
                x_dimension=None,
                y_dimension=None,
                selected_periods=selected_comparison_periods,
                small_multiples_dimension=None,
                x_max_items=max(1, dimension_panel_max_items),
                aggregate_other_items=dimension_panel_aggregate_other_items,
                capture_figure="last",
                dimension_panel_chart=True,
            )
        two_dimension_pair = _explicit_dimension_pair(
            dimensions,
            options,
            panel_keys=(
                "multitier_bar_panel_dimension",
                "multitier_bar_small_multiples_dimension",
                "multitier_bar_by_dimension",
            ),
            item_keys=(
                "multitier_bar_item_dimension",
                "multitier_bar_x_dimension",
                "multitier_bar_bar_dimension",
            ),
        )
        if two_dimension_pair is not None and (
            "multitier_bar_two_dimension" in requested
            or bool(options.get("multitier_bar_two_dimension"))
        ):
            panel_dimension, item_dimension = two_dimension_pair
            add(
                "multitier_bar_two_dimension",
                "plot_multitier_bar_chart",
                "multitierBarChart",
                "multitier_bar_two_dimension.png",
                base_chart="multitier_bar",
                dimensions=[panel_dimension, item_dimension],
                x_dimension=item_dimension,
                y_dimension=None,
                selected_periods=selected_comparison_periods,
                small_multiples_dimension=panel_dimension,
                x_max_items=int(
                    options.get("multitier_bar_item_max_items")
                    or options.get("dimension_panel_max_items")
                    or min(max_items, 6)
                ),
                aggregate_other_items=bool(
                    options.get(
                        "multitier_bar_aggregate_other_items",
                        options.get("dimension_panel_aggregate_other_items", False),
                    )
                ),
                capture_figure="last",
                dimension_selection="panel_dimension_item_dimension_multitier_bar",
            )
    return specs


def build_contribution_summary(
    canonical: pl.DataFrame, recipe: dict[str, Any]
) -> pl.DataFrame:
    """Return a top-item contribution table for the primary dimension."""

    metric = str(recipe["mappings"]["amount_column"])
    dimensions = _useful_dimensions(
        canonical, [str(item) for item in recipe["mappings"].get("dimensions") or []]
    )
    primary = dimensions[0]
    total = float(canonical.select(pl.col(metric).sum()).item() or 0.0)
    return (
        canonical.group_by(primary)
        .agg(pl.col(metric).sum().alias(metric))
        .sort(metric, descending=True)
        .with_columns(
            (
                pl.col(metric) / total
                if abs(total) > TOLERANCE
                else pl.lit(None, dtype=pl.Float64)
            ).alias("share_of_total")
        )
    )


def _format_millions(value: float | None) -> str:
    if value is None:
        return "n/a"
    return f"{value / 1_000_000:,.1f}m"


def total_context(canonical: pl.DataFrame, recipe: dict[str, Any]) -> dict[str, Any]:
    """Return total and top contribution facts."""

    metric = str(recipe["mappings"]["amount_column"])
    current = _current_period_frame(canonical)
    summary = build_contribution_summary(current, recipe)
    total = float(current.select(pl.col(metric).sum()).item() or 0.0)
    top_rows = summary.head(10).to_dicts()
    return {
        "metric": metric,
        "total": total,
        "top_items": [
            {
                "item": row.get(summary.columns[0]),
                "value": float(row.get(metric) or 0.0),
                "share_of_total": row.get("share_of_total"),
            }
            for row in top_rows
        ],
    }


def build_summary_markdown(
    recipe: dict[str, Any],
    contribution: dict[str, Any],
    chart_audits: dict[str, Any],
) -> str:
    """Build deterministic markdown summary."""

    written = [
        name for name, audit in chart_audits.items() if audit.get("status") == "written"
    ]
    data_written = [
        name
        for name, audit in chart_audits.items()
        if audit.get("status") == "data_written"
    ]
    failed = [
        name
        for name, audit in chart_audits.items()
        if audit.get("status") not in {"written", "data_written"}
    ]
    lines = [
        "# Mix & Contribution Source Data",
        "",
        f"- Source file: `{recipe.get('source_file')}`",
        f"- Metric: `{contribution['metric']}`",
        f"- Total: `{contribution['total']:,.2f}`",
        f"- Dimensions: `{', '.join(recipe['mappings'].get('dimensions') or [])}`",
        f"- Legacy charts written: `{len(written)}`",
        f"- Legacy chart data candidates written: `{len(data_written)}`",
        f"- Legacy chart attempts not written: `{len(failed)}`",
        "",
        "## Top Contribution Items",
        "",
    ]
    for row in contribution["top_items"][:5]:
        share = row["share_of_total"]
        share_label = f"{share:.1%}" if isinstance(share, float) else "n/a"
        lines.append(
            f"- {row['item']}: {_format_millions(row['value'])} ({share_label})"
        )
    lines.extend(
        [
            "",
            "Charts are generated by the vendored legacy charting pipeline. "
            "Attempts that did not produce a legacy Plotly figure are listed in the audit.",
            "",
        ]
    )
    return "\n".join(lines)


def write_client_report(
    recipe: dict[str, Any],
    contribution: dict[str, Any],
    chart_paths: list[str],
    output_dir: Path,
) -> tuple[list[str], dict[str, Any]]:
    """Write Markdown and DOCX client report."""

    md_path = output_dir / "mix_contribution_client_report.md"
    docx_path = output_dir / "mix_contribution_client_report.docx"
    top_items = contribution["top_items"][:5]
    lines = [
        "# Mix & Contribution Analysis",
        "",
        f"Total {contribution['metric']} is {_format_millions(contribution['total'])}.",
        "",
        "## Main Contribution Items",
        "",
        *[f"- {item['item']}: {_format_millions(item['value'])}" for item in top_items],
        "",
        "## Source Files",
        "",
        *[
            f"- `{Path(path).name}`"
            for path in chart_paths
            if Path(path).suffix.lower() == ".png"
        ],
        "",
    ]
    md_path.write_text("\n".join(lines), encoding="utf-8")
    try:
        from docx import Document
        from docx.shared import Inches

        document = Document()
        document.add_heading("Mix & Contribution Analysis", level=1)
        document.add_paragraph(
            f"Total {contribution['metric']} is "
            f"{_format_millions(contribution['total'])}."
        )
        if top_items:
            document.add_heading("Main Contribution Items", level=2)
            for item in top_items:
                document.add_paragraph(
                    f"{item['item']}: {_format_millions(item['value'])}",
                    style="List Bullet",
                )
        document.add_heading("Charts", level=2)
        for chart_path in chart_paths:
            path = Path(chart_path)
            if path.suffix.lower() == ".png" and path.exists():
                document.add_picture(str(path), width=Inches(6.4))
        document.save(docx_path)
        status = "written"
        error = None
    except (
        ImportError,
        ModuleNotFoundError,
        OSError,
        ValueError,
        zipfile.BadZipFile,
    ) as exc:
        status = "not_written"
        error = str(exc)
    paths = [str(md_path)]
    if docx_path.exists():
        paths.append(str(docx_path))
    return paths, {
        "status": status,
        "markdown": md_path.name,
        "docx": docx_path.name if docx_path.exists() else None,
        "error": error,
    }


def _relative_path(path: Path, base: Path) -> str:
    try:
        return str(path.relative_to(base))
    except ValueError:
        return str(path)


def _artifact_kind(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".png", ".html"}:
        return "chart"
    if suffix in {".csv", ".xlsx"}:
        return "table"
    if suffix == ".json":
        return "context"
    if suffix in {".md", ".docx"}:
        return "report"
    return "artifact"


def write_chart_context_artifacts(
    chart_name: str,
    chart_context: dict[str, Any],
    output_dir: Path,
) -> tuple[list[str], dict[str, Any]]:
    """Write model-readable chart data captured from the legacy chart path."""

    safe_name = "".join(
        character if character.isalnum() or character in {"_", "-"} else "_"
        for character in chart_name
    )
    context_path = output_dir / f"{safe_name}_chart_context.json"
    table_path = output_dir / f"{safe_name}_chart_data.csv"
    write_json(context_path, chart_context)
    paths = [str(context_path)]
    table_status = "not_written_no_rows"
    rows = chart_context.get("related_metric_rows")
    if not rows:
        rows = chart_context.get("series_by_dimension")
    if not rows:
        rows = chart_context.get("waterfall_rows")
    if not rows:
        data_frame = chart_context.get("data_frame") or {}
        rows = data_frame.get("rows") if isinstance(data_frame, dict) else None
    if isinstance(rows, list) and rows:
        pl.DataFrame(rows).write_csv(table_path)
        paths.append(str(table_path))
        table_status = "written"
    return paths, {
        "status": "written",
        "context_path": context_path.name,
        "table_path": table_path.name if table_path.exists() else None,
        "table_status": table_status,
        "source": chart_context.get("chart_data_source"),
    }


def _normalize_artifact_mode(artifact_mode: str) -> str:
    """Return a supported artifact mode or raise for invalid contract input."""

    normalized = str(artifact_mode or ARTIFACT_MODE_DATA_AND_RENDER).strip().lower()
    if normalized not in ARTIFACT_MODES:
        allowed = ", ".join(sorted(ARTIFACT_MODES))
        raise ValueError(f"Unsupported artifact_mode {artifact_mode!r}; use {allowed}.")
    return normalized


def _chart_artifact_id(spec: dict[str, Any]) -> str:
    """Return the chart artifact ID for generated chart data."""

    return str(spec.get("name") or Path(str(spec.get("artifact_name") or "")).stem)


def _spec_for_data_only_artifacts(spec: dict[str, Any]) -> dict[str, Any]:
    """Force chart-data capture for data-only artifact generation."""

    return {**spec, "capture_chart_data": True}


def _slim_data_only_context(chart_context: dict[str, Any]) -> dict[str, Any]:
    """Remove render-heavy payloads from data-only chart context."""

    return {
        key: value
        for key, value in chart_context.items()
        if key not in {"plotly_figures", "exports"}
    }


def run_mix_contribution(
    input_path: Path,
    output_dir: Path,
    recipe_path: Path | None = None,
    *,
    language: str = "en",
    currency: str | None = None,
    artifact_mode: str = ARTIFACT_MODE_DATA_AND_RENDER,
) -> MixContributionRunResult:
    """Run the full mix-contribution workflow."""

    artifact_mode = _normalize_artifact_mode(artifact_mode)
    df = read_table(input_path)
    existing_recipe = read_json(recipe_path) if recipe_path else None
    recipe = build_recipe(
        input_path, df, language=language, existing_recipe=existing_recipe
    )
    recipe = preserve_recipe_filters(recipe, existing_recipe)
    recipe = preserve_recipe_cohorts(recipe, existing_recipe)
    if currency:
        recipe.setdefault("options", {})["currency"] = currency
    output_dir.mkdir(parents=True, exist_ok=True)
    retired_artifact_cleanup = cleanup_retired_chart_artifacts(output_dir)
    try:
        df, filter_audit = apply_recipe_filters(df, recipe)
        recipe.setdefault("options", {})["recipe_filter_audit"] = filter_audit
        run_intake = write_run_intake(
            output_dir,
            input_path,
            recipe_path=recipe_path,
            recipe=recipe,
            source_row_count=df.height,
        )
        canonical, preparation_audit = prepare_canonical_frame(df, recipe)
        preparation_audit["recipe_filters"] = filter_audit
        current_period, previous_period = recipe_cohort_period_labels(
            recipe,
            default_current=CURRENT_PERIOD,
            default_previous="PY",
        )
        canonical, cohort_audit = apply_recipe_cohorts(
            canonical,
            recipe,
            period_column=CANONICAL_PERIOD,
            value_column=LEGACY_AMOUNT_COLUMN,
            current_period=current_period,
            previous_period=previous_period,
        )
        preparation_audit["recipe_cohorts"] = cohort_audit
        contribution = total_context(canonical, recipe)
        specs = build_chart_specs(canonical, recipe)
        prepared_data_cache = LegacyPreparedDataCache.empty()
        render_charts = artifact_mode != ARTIFACT_MODE_DATA_ONLY
        artifact_paths: list[str] = []
        chart_audits: dict[str, Any] = {}
        chart_context_artifacts: dict[str, Any] = {}
        for spec in specs:
            run_spec = (
                _spec_for_data_only_artifacts(spec)
                if artifact_mode == ARTIFACT_MODE_DATA_ONLY
                else spec
            )
            export = write_legacy_mix_chart(
                canonical,
                recipe,
                output_dir,
                run_spec,
                prepared_data_cache=prepared_data_cache,
                render=render_charts,
            )
            artifact_paths.extend(export.paths)
            spec_name = str(spec["name"])
            chart_audits[spec_name] = export.audit
            chart_context = getattr(export, "chart_context", None)
            if chart_context:
                if artifact_mode == ARTIFACT_MODE_DATA_ONLY:
                    chart_context = _slim_data_only_context(chart_context)
                context_paths, context_audit = write_chart_context_artifacts(
                    spec_name, chart_context, output_dir
                )
                artifact_paths.extend(context_paths)
                chart_context_artifacts[spec_name] = context_audit
                chart_audits[spec_name]["chart_context"] = context_audit
        current = _current_period_frame(canonical)
        summary_table = build_contribution_summary(current, recipe)
        canonical_path = output_dir / "mix_contribution_canonical.csv"
        canonical.write_csv(canonical_path)
        prepared_manifest_path = write_prepared_data_manifest(
            output_dir=output_dir,
            plugin="mix-contribution-analysis",
            chart_family="mix_contribution",
            source_file=input_path,
            prepared_path=canonical_path,
            frame=canonical,
            recipe=recipe,
            preparation_audit=preparation_audit,
        )
        summary_table.write_csv(output_dir / "mix_contribution_summary.csv")
        artifact_paths.extend(
            [
                str(canonical_path),
                str(prepared_manifest_path),
                str(output_dir / "mix_contribution_summary.csv"),
            ]
        )
        try:
            summary_table.write_excel(output_dir / "mix_contribution_results.xlsx")
            xlsx_status = "written"
            artifact_paths.append(str(output_dir / "mix_contribution_results.xlsx"))
        except (ImportError, ModuleNotFoundError, OSError, ValueError) as exc:
            xlsx_status = f"not_written: {exc}"
        context = {
            "schema_version": SCHEMA_VERSION,
            "source_file": str(input_path),
            "artifact_mode": artifact_mode,
            "mappings": recipe["mappings"],
            "options": recipe["options"],
            "prepared_data_manifest": prepared_manifest_path.name,
            "contribution": contribution,
            "chart_audits": chart_audits,
            "chart_context_artifacts": chart_context_artifacts,
        }
        write_json(output_dir / "mix_contribution_context.json", context)
        artifact_paths.append(str(output_dir / "mix_contribution_context.json"))
        summary = build_summary_markdown(recipe, contribution, chart_audits)
        (output_dir / "mix_contribution_summary.md").write_text(
            summary, encoding="utf-8"
        )
        artifact_paths.append(str(output_dir / "mix_contribution_summary.md"))
        report_paths, report_audit = write_client_report(
            recipe, contribution, artifact_paths, output_dir
        )
        artifact_paths.extend(report_paths)
        write_json(output_dir / "used_recipe.json", recipe)
        artifact_paths.append(str(output_dir / "used_recipe.json"))
        audit: dict[str, Any] = {
            "schema_version": SCHEMA_VERSION,
            "created_at": utc_now(),
            "input_file": str(input_path),
            "recipe": recipe,
            "legacy_runtime": {
                "preparation": preparation_audit,
                "chart_audits": chart_audits,
                "chart_context_artifacts": chart_context_artifacts,
                "client_report": report_audit,
                "retired_artifact_cleanup": retired_artifact_cleanup,
                "artifact_mode": artifact_mode,
            },
            "checks": {
                "canonical_row_count": canonical.height,
                "summary_row_count": summary_table.height,
                "legacy_chart_attempt_count": len(chart_audits),
                "legacy_chart_written_count": sum(
                    1
                    for audit_item in chart_audits.values()
                    if audit_item.get("status") == "written"
                ),
                "legacy_chart_data_count": sum(
                    1
                    for audit_item in chart_audits.values()
                    if audit_item.get("status") in {"written", "data_written"}
                ),
            },
            "outputs": {
                "mix_contribution_canonical.csv": "written",
                "prepared_data_manifest.json": "written",
                "mix_contribution_summary.csv": "written",
                "mix_contribution_results.xlsx": xlsx_status,
                "mix_contribution_context.json": "written",
                "mix_contribution_summary.md": "written",
                "used_recipe.json": "written",
            },
        }
        for path in artifact_paths:
            audit["outputs"][_relative_path(Path(path), output_dir)] = "written"
        write_json(output_dir / "mix_contribution_audit.json", audit)
        audit["outputs"]["mix_contribution_audit.json"] = "written"
        review_session = write_review_session_artifacts(
            output_dir,
            input_path,
            run_id=run_intake.run_id,
            run_intake_path=run_intake.path,
            recipe_path=recipe_path,
            recipe=recipe,
            summary_rows=summary_table.to_dicts(),
            audit=audit,
            generated_paths=[
                *artifact_paths,
                output_dir / "mix_contribution_audit.json",
            ],
        )
        audit["review_session"] = {
            "run_intake_path": review_session.run_intake_path.name,
            "review_payload_path": review_session.review_payload_path.name,
            "ui_decisions_path": review_session.ui_decisions_path.name,
            "final_artifacts_path": review_session.final_artifacts_path.name,
            "review_item_count": review_session.review_item_count,
        }
        for path in (
            review_session.run_intake_path,
            review_session.review_payload_path,
            review_session.ui_decisions_path,
            review_session.final_artifacts_path,
        ):
            audit["outputs"][_relative_path(path, output_dir)] = "written"
        artifact_paths = [
            *artifact_paths,
            str(review_session.run_intake_path),
            str(review_session.review_payload_path),
            str(review_session.ui_decisions_path),
            str(review_session.final_artifacts_path),
        ]
        write_json(output_dir / "mix_contribution_audit.json", audit)
        return MixContributionRunResult(
            canonical_frame=canonical,
            audit=audit,
            summary_markdown=summary,
            artifact_paths=artifact_paths,
        )
    finally:
        cleanup_legacy_imports()


if __name__ == "__main__":
    raise SystemExit("Use run_mix_contribution.py or inspect_inputs.py")
