"""Local, scope-bound studio archive indexing and retrieval."""

from __future__ import annotations

import atexit
import errno
import hashlib
import io
import json
import os
import re
import secrets
import shutil
import sqlite3
import stat

# Native folder selection uses fixed executable names and never invokes a shell.
import subprocess  # nosec B404
import sys
import tempfile
import warnings
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from email import policy
from email.parser import BytesParser
from email.utils import getaddresses
from functools import wraps
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Callable, Iterator, Mapping, Sequence


def _add_vera_assurance_module_path() -> None:
    """Load the packaged shared contract from source or an installed plugin."""

    candidates = (
        Path(__file__).resolve().parents[1] / "vendor" / "modules",
        Path(__file__).resolve().parents[2] / "_shared" / "vendor" / "modules",
    )
    for module_root in candidates:
        if (module_root / "vera_assurance").is_dir():
            if str(module_root) not in sys.path:
                sys.path.insert(0, str(module_root))
            return
    raise RuntimeError("The required vera_assurance module is not available.")


_add_vera_assurance_module_path()

_SCRIPT_DIRECTORY = Path(__file__).resolve().parent
if str(_SCRIPT_DIRECTORY) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIRECTORY))

import client_ledger as ledger  # noqa: E402
import google_drive as drive  # noqa: E402
from vera_assurance import (  # noqa: E402
    JOURNAL_SAMPLING_CHECK_ENTRIES_HANDOFF,
    VERA_CLIENT_WORKFLOW_IDS,
    build_studio_client_folder_binding,
)

CLIENT_WORKFLOW_IDS = (*VERA_CLIENT_WORKFLOW_IDS, "apertura-pratica")

__all__ = [
    "ArchiveAccessError",
    "ArchiveError",
    "ArchiveFolderPickerUnavailableError",
    "ArchiveNotConfiguredError",
    "SourceChangedError",
    "cancel_studio_client_workflow",
    "close_studio_client_engagement",
    "complete_studio_client_workflow",
    "configure_archive",
    "create_studio_client_engagement",
    "create_studio_client",
    "diagnose_archive_access",
    "fail_studio_client_workflow",
    "finalize_studio_client_workflow",
    "get_studio_client_folder",
    "get_studio_archive_organization_inventory",
    "authorize_studio_google_drive",
    "bind_studio_client_google_drive",
    "import_studio_client_document",
    "list_studio_client_engagements",
    "list_studio_client_identities",
    "list_studio_clients",
    "match_studio_email_client",
    "open_archive_source",
    "open_studio_archive_organization_item",
    "open_studio_google_drive_source",
    "plan_gmail_client_search",
    "prepare_studio_client_workflow",
    "recover_studio_client_ledger",
    "refresh_archive",
    "report_studio_client_retention",
    "resolve_studio_client_identity",
    "search_archive",
    "setup_archive_with_folder_picker",
    "set_studio_client_identity",
    "snapshot_studio_client_folder",
    "snapshot_studio_client_google_drive",
    "start_check_entries_from_sample",
    "start_studio_client_workflow",
    "studio_archive_status",
    "studio_google_drive_status",
]

SCHEMA_VERSION = "2"
CONFIG_SCHEMA_VERSION = 1
CLIENT_IDENTITIES_SCHEMA_VERSION = 2
STATE_ENV = "VERA_STUDIO_ARCHIVE_STATE_DIR"
DEFAULT_STATE_SUBDIR = Path(".mparanza") / "vera-studio-archive"
CONFIG_FILENAME = "config.json"
SESSION_ENV = "VERA_STUDIO_ARCHIVE_SESSION_ID"
_PROCESS_SESSION = f"{os.getpid()}-{secrets.token_hex(16)}"
_STATE_LEASES: dict[Path, tuple[Any, bytes | None]] = {}
_CONFIGURATION_CHANGED = (
    "Studio Archive configuration changed during this run, "
    "re-validate before continuing."
)


def _session_id() -> str:
    """Share an identity across CLI commands only when the host supplies one."""
    return (
        os.environ.get(SESSION_ENV, "").strip()
        or os.environ.get("CODEX_THREAD_ID", "").strip()
        or _PROCESS_SESSION
    )


def _check_configuration(state: Path) -> None:
    """Reject changes to a configuration already observed by this process."""
    expected = _STATE_LEASES[state][1]
    path = state / CONFIG_FILENAME
    actual = path.read_bytes() if path.exists() else None
    if actual != expected:
        raise ArchiveError(_CONFIGURATION_CHANGED)


def _claim_state(state: Path) -> None:
    """Hold an OS lock until process exit, covering the entire read/use window."""
    state = state.resolve()
    if state in _STATE_LEASES:
        _check_configuration(state)
        return
    state.mkdir(parents=True, exist_ok=True, mode=0o700)
    handle = (state / ".config.lock").open("a+b")
    try:
        if os.name == "nt":
            import msvcrt

            handle.seek(0)
            handle.write(b"0")
            handle.flush()
            handle.seek(0)
            msvcrt.locking(handle.fileno(), msvcrt.LK_NBLCK, 1)
        else:
            import fcntl

            fcntl.flock(handle.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
    except OSError as exc:
        handle.close()
        raise ArchiveError(
            "Studio Archive state is in use by another process; "
            "use an isolated session state directory or retry after it finishes."
        ) from exc
    path = state / CONFIG_FILENAME
    _STATE_LEASES[state] = (handle, path.read_bytes() if path.exists() else None)


def _release_state_leases() -> None:
    for handle, _ in _STATE_LEASES.values():
        handle.close()
    _STATE_LEASES.clear()


atexit.register(_release_state_leases)
CLIENT_IDENTITIES_FILENAME = "client-identities.json"
GOOGLE_DRIVE_BINDINGS_FILENAME = "google-drive-bindings.json"
GOOGLE_DRIVE_AUTH_FILENAME = "google-drive-token.json"
GOOGLE_DRIVE_BINDINGS_SCHEMA = "vera.studio_archive_google_drive_bindings.v1"
ARCHIVE_ORGANIZATION_INVENTORY_SCHEMA = "vera.archive_organization_model_inventory.v1"
ARCHIVE_ORGANIZATION_INVENTORY_REF_PREFIX = "archive_inventory_"
ARCHIVE_ORGANIZATION_ITEM_REF_PREFIX = "archive_item_"
ARCHIVE_ORGANIZATION_DUPLICATE_REF_PREFIX = "exact_group_"
GOOGLE_DRIVE_EXPORTS = {
    "application/vnd.google-apps.document": ("text/plain", ".txt"),
    "application/vnd.google-apps.spreadsheet": (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xlsx",
    ),
    "application/vnd.google-apps.presentation": ("application/pdf", ".pdf"),
    "application/vnd.google-apps.drawing": ("application/pdf", ".pdf"),
}
DATABASE_FILENAME = "archive.sqlite3"
MANAGED_ENGAGEMENTS_DIRECTORY = ledger.LEDGER_DIRECTORY

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".xml"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
XLSX_SUFFIXES = {".xlsx"}
EMAIL_SUFFIXES = {".eml"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff"}
SUPPORTED_SUFFIXES = (
    TEXT_SUFFIXES
    | PDF_SUFFIXES
    | DOCX_SUFFIXES
    | XLSX_SUFFIXES
    | EMAIL_SUFFIXES
    | IMAGE_SUFFIXES
)

MAX_FILES = 50_000
MAX_TOTAL_BYTES = 50 * 1024 * 1024 * 1024
MAX_FILE_BYTES = 256 * 1024 * 1024
MAX_TEXT_BYTES = 20 * 1024 * 1024
MAX_EMAIL_BYTES = 30 * 1024 * 1024
MAX_PDF_BYTES = 100 * 1024 * 1024
MAX_PDF_PAGES = 500
MAX_PDF_TEXT_CHARS = 20_000_000
MAX_IMAGE_FRAMES = 100
MAX_IMAGE_TOTAL_PIXELS = 20_000_000
MAX_WORKBOOK_SHEETS = 100
MAX_WORKSHEET_ROWS = 20_000
MAX_WORKSHEET_COLUMNS = 512
MAX_OOXML_MEMBER_BYTES = 20 * 1024 * 1024
MAX_OOXML_TOTAL_BYTES = 100 * 1024 * 1024
MAX_OOXML_MEMBERS = 5_000
MAX_OOXML_COMPRESSION_RATIO = 200
MAX_CHUNK_CHARS = 6_000
MAX_CHUNK_LINES = 120
MAX_SEARCH_TOKENS = 24
MAX_OPEN_CHARS = 24_000
MAX_STATUS_SCAN_ISSUES = 200
MAX_STATUS_DOCUMENT_ISSUES = 200
MAX_CLIENT_IDENTITIES = 5_000
MAX_CLIENT_EMAIL_ADDRESSES = 20
MAX_CLIENT_LEGAL_NAMES = 20
MAX_CLIENT_TAX_IDENTIFIERS = 20
MAX_GMAIL_QUERY_IDENTITIES = 10
MAX_GMAIL_TOPIC_CHARS = 200
SUPPORTED_ENGAGEMENT_IMPORT_ROLES = {"journal", "source", "support"}
IGNORED_NAMES = {
    ".DS_Store",
    ".git",
    ".hg",
    ".svn",
    "Thumbs.db",
    "__pycache__",
    "desktop.ini",
}

ARCHIVE_HOST_PERMISSION_REQUIRED = "MPARANZA_ARCHIVE_HOST_PERMISSION_REQUIRED"
_SMB_CREDENTIAL_WINERRORS = frozenset({86, 1219, 1326, 1909})
_NETWORK_UNAVAILABLE_WINERRORS = frozenset({53, 64, 67, 121, 1231, 1232})
_ACCESS_DENIED_WINERRORS = frozenset({5})
_ACCESS_DENIED_ERRNOS = frozenset({errno.EACCES, errno.EPERM})


class ArchiveError(RuntimeError):
    """Base class for bounded archive workflow errors."""

    code = "archive_error"


class ArchiveAccessError(ArchiveError):
    """Describe one mechanically observed archive-root access failure safely."""

    def __init__(
        self,
        message: str,
        *,
        code: str,
        stage: str,
        path_kind: str,
        recommended_action: str,
        host_access_approved: bool,
        error: OSError | None = None,
    ) -> None:
        super().__init__(message)
        self.code = code
        details: dict[str, Any] = {
            "category": code,
            "stage": stage,
            "path_kind": path_kind,
            "recommended_action": recommended_action,
            "host_access_approved": host_access_approved,
        }
        if error is not None:
            if isinstance(error.errno, int):
                details["errno"] = error.errno
            winerror = getattr(error, "winerror", None)
            if isinstance(winerror, int):
                details["winerror"] = winerror
        self.details = details


class ArchiveNotConfiguredError(ArchiveError):
    """Raised when an operation needs a configured archive."""

    code = "archive_not_configured"

    def __init__(self, message: str) -> None:
        super().__init__(message)
        self.details = _studio_archive_setup_contract()


class ArchiveFolderPickerUnavailableError(ArchiveError):
    """Raised when this desktop runtime cannot open a native folder picker."""

    code = "archive_folder_picker_unavailable"

    def __init__(self) -> None:
        super().__init__(
            "The native Studio Archive folder chooser is unavailable on this runtime."
        )
        self.details = {
            **_studio_archive_setup_contract(),
            "manual_path_fallback_allowed": True,
        }


class SourceChangedError(ArchiveError):
    """Raised when an indexed source no longer matches its recorded bytes."""

    code = "source_changed_refresh_required"


@dataclass(frozen=True)
class Scope:
    """One mechanically selected archive-relative search boundary."""

    scope_id: str
    relative_dir: str
    display_name: str

    def as_json(self) -> dict[str, str]:
        """Return a JSON-safe scope record."""

        return {
            "scope_id": self.scope_id,
            "relative_dir": self.relative_dir,
            "display_name": self.display_name,
        }


@dataclass(frozen=True)
class ArchiveConfig:
    """One user's local archive configuration."""

    archive_root: Path
    scopes: tuple[Scope, ...]
    configured_at: str

    def as_json(self) -> dict[str, Any]:
        """Return the persisted configuration shape."""

        return {
            "schema_version": CONFIG_SCHEMA_VERSION,
            "archive_root": str(self.archive_root),
            "configured_at": self.configured_at,
            "scopes": [scope.as_json() for scope in self.scopes],
        }


@dataclass(frozen=True)
class ClientIdentity:
    """One private client identity record bound to an exact archive scope."""

    client_id: str
    scope_id: str
    email_addresses: tuple[str, ...]
    legal_names: tuple[str, ...]
    tax_identifiers: tuple[str, ...]
    updated_at: str

    def as_json(self) -> dict[str, Any]:
        """Return the persisted private-registry shape."""

        return {
            "client_id": self.client_id,
            "scope_id": self.scope_id,
            "email_addresses": list(self.email_addresses),
            "legal_names": list(self.legal_names),
            "tax_identifiers": list(self.tax_identifiers),
            "updated_at": self.updated_at,
        }


@dataclass(frozen=True)
class DiscoveredFile:
    """One regular source file discovered inside a configured scope."""

    scope_id: str
    relative_path: str
    path: Path
    size_bytes: int
    mtime_ns: int


@dataclass(frozen=True)
class ScanIssue:
    """One skipped archive entry that may limit search completeness."""

    scope_id: str
    relative_path: str
    reason: str
    size_bytes: int | None

    def as_json(self) -> dict[str, Any]:
        """Return a JSON-safe scan issue."""

        return {
            "scope_id": self.scope_id,
            "relative_path": self.relative_path,
            "reason": self.reason,
            "size_bytes": self.size_bytes,
        }


@dataclass(frozen=True)
class ExtractedChunk:
    """One bounded, citable source fragment."""

    ordinal: int
    locator_kind: str
    locator_value: str
    text: str


@dataclass(frozen=True)
class ExtractionResult:
    """Mechanically extracted chunks and limitations for one source."""

    chunks: tuple[ExtractedChunk, ...]
    extraction_method: str
    status: str
    needs_ocr: bool
    limitations: tuple[str, ...]


class _HtmlTextExtractor(HTMLParser):
    """Collect visible text from an HTML email part."""

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        """Keep non-empty visible text."""

        if data.strip():
            self.parts.append(data)


def _now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    flags = os.O_RDONLY
    if hasattr(os, "O_NOFOLLOW"):
        flags |= os.O_NOFOLLOW
    descriptor = os.open(path, flags)
    with os.fdopen(descriptor, "rb") as handle:
        if not stat.S_ISREG(os.fstat(handle.fileno()).st_mode):
            raise ArchiveError("Studio Archive source is not a regular file.")
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _stable_id(prefix: str, *values: str) -> str:
    payload = "\x1f".join(values).encode("utf-8")
    return f"{prefix}_{hashlib.sha256(payload).hexdigest()[:24]}"


def _new_private_id(prefix: str, existing: set[str]) -> str:
    """Create an opaque persistent ID; identity must not depend on a folder name."""

    for _ in range(100):
        candidate = f"{prefix}_{secrets.token_hex(12)}"
        if candidate not in existing:
            return candidate
    raise ArchiveError(f"Could not allocate a unique {prefix} identifier.")


def _state_dir(
    explicit: Path | None = None,
    *,
    create: bool = False,
) -> Path:
    selected = explicit
    if selected is None:
        environment_value = os.environ.get(STATE_ENV, "").strip()
        selected = (
            Path(environment_value).expanduser()
            if environment_value
            else Path.home()
            / DEFAULT_STATE_SUBDIR
            / "sessions"
            / hashlib.sha256(_session_id().encode("utf-8")).hexdigest()[:32]
        )
    selected = Path(selected).expanduser()
    if not selected.is_absolute():
        raise ArchiveError("Studio Archive state directory must be absolute.")
    if selected.is_symlink():
        raise ArchiveError("Studio Archive state directory cannot be a symbolic link.")
    if create:
        selected.mkdir(parents=True, exist_ok=True, mode=0o700)
        try:
            selected.chmod(0o700)
        except OSError as exc:
            raise ArchiveError(f"Could not secure Studio Archive state: {exc}") from exc
    elif selected.exists() and not selected.is_dir():
        raise ArchiveError("Studio Archive state path must be a directory.")
    elif (
        selected.exists()
        and os.name == "posix"
        and stat.S_IMODE(selected.stat().st_mode) & 0o077
    ):
        raise ArchiveError(
            "Studio Archive state directory must not be accessible by group or others."
        )
    return selected.resolve()


def _config_path(state_dir: Path) -> Path:
    return state_dir / CONFIG_FILENAME


def _client_identities_path(state_dir: Path) -> Path:
    return state_dir / CLIENT_IDENTITIES_FILENAME


def _google_drive_bindings_path(state_dir: Path) -> Path:
    return state_dir / GOOGLE_DRIVE_BINDINGS_FILENAME


def _google_drive_token_path(state_dir: Path) -> Path:
    return state_dir / GOOGLE_DRIVE_AUTH_FILENAME


def _database_path(state_dir: Path) -> Path:
    return state_dir / DATABASE_FILENAME


def _assert_private_file(path: Path, label: str) -> None:
    if os.name != "posix" or not path.exists():
        return
    if stat.S_IMODE(path.stat().st_mode) & 0o077:
        raise ArchiveError(
            f"Studio Archive {label} must not be accessible by group or others."
        )


def _write_private_json(path: Path, payload: dict[str, Any]) -> None:
    if path.name == CONFIG_FILENAME:
        _claim_state(path.parent)
        existing = _STATE_LEASES[path.parent.resolve()][1]
        if existing is not None:
            owner = json.loads(existing).get("session_id")
            if owner is not None and owner != _session_id():
                raise ArchiveError(_CONFIGURATION_CHANGED)
        payload = {**payload, "session_id": _session_id()}
    path.parent.mkdir(parents=True, exist_ok=True, mode=0o700)
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{path.name}.",
        suffix=".tmp",
        dir=path.parent,
    )
    temporary = Path(temporary_name)
    try:
        with os.fdopen(descriptor, "w", encoding="utf-8") as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2, sort_keys=True)
            handle.write("\n")
            handle.flush()
            os.fsync(handle.fileno())
        temporary.chmod(0o600)
        written_bytes = temporary.read_bytes() if path.name == CONFIG_FILENAME else None
        temporary.replace(path)
        path.chmod(0o600)
        if path.name == CONFIG_FILENAME:
            state = path.parent.resolve()
            _STATE_LEASES[state] = (_STATE_LEASES[state][0], written_bytes)
    finally:
        if temporary.exists():
            temporary.unlink()


def _path_is_within(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
    except ValueError:
        return False
    return True


def _archive_path_kind(path: Path) -> str:
    """Classify only path syntax needed for access diagnostics."""

    raw_path = os.fspath(path)
    if raw_path.startswith(("\\\\", "//")):
        return "unc_share"
    return "local_or_mounted"


def _archive_access_error(
    error: OSError,
    *,
    stage: str,
    path_kind: str,
    host_access_approved: bool,
) -> ArchiveAccessError:
    """Map stable OS error codes to auditable access categories.

    Fixed codes are appropriate here because they are mechanically emitted by
    the operating system. Unknown combinations remain unclassified instead of
    inferring a credential, sandbox, or filesystem-permission cause.
    """

    winerror = getattr(error, "winerror", None)
    if winerror in _SMB_CREDENTIAL_WINERRORS:
        return ArchiveAccessError(
            "The operating system rejected the SMB session or credentials for the "
            "selected network share. Connect the share in the signed-in desktop "
            "session and retry; Vera never requests or stores SMB credentials.",
            code="archive_smb_credentials_required",
            stage=stage,
            path_kind=path_kind,
            recommended_action="connect_smb_session",
            host_access_approved=host_access_approved,
            error=error,
        )
    if winerror in _NETWORK_UNAVAILABLE_WINERRORS:
        return ArchiveAccessError(
            "The selected network share is unreachable from this desktop session. "
            "Verify the server and share connection, then retry the same path.",
            code="archive_network_share_unreachable",
            stage=stage,
            path_kind=path_kind,
            recommended_action="verify_share_connection",
            host_access_approved=host_access_approved,
            error=error,
        )
    access_denied = (
        error.errno in _ACCESS_DENIED_ERRNOS or winerror in _ACCESS_DENIED_WINERRORS
    )
    if access_denied and os.environ.get("CODEX_SANDBOX") and not host_access_approved:
        return ArchiveAccessError(
            f"{ARCHIVE_HOST_PERMISSION_REQUIRED}: Codex could not read the selected "
            "archive root from its current sandbox. Retry the same access diagnostic "
            "with host folder access approval before creating a client or engagement.",
            code="archive_host_access_permission_required",
            stage=stage,
            path_kind=path_kind,
            recommended_action="retry_with_host_folder_access",
            host_access_approved=False,
            error=error,
        )
    if access_denied:
        return ArchiveAccessError(
            "The operating system denied filesystem access to the selected archive "
            "root. Grant the signed-in user share and filesystem read/list permissions "
            "and write permission before client registration, then retry.",
            code="archive_filesystem_access_denied",
            stage=stage,
            path_kind=path_kind,
            recommended_action="grant_share_and_filesystem_permissions",
            host_access_approved=host_access_approved,
            error=error,
        )
    if error.errno == errno.ENOENT:
        return ArchiveAccessError(
            "The selected archive root does not exist in this desktop session.",
            code="archive_root_not_found",
            stage=stage,
            path_kind=path_kind,
            recommended_action="verify_archive_root",
            host_access_approved=host_access_approved,
            error=error,
        )
    return ArchiveAccessError(
        "The selected archive root is unavailable for an unclassified operating-system "
        "reason. The diagnostic has preserved the numeric error code without exposing "
        "the private path.",
        code="archive_root_unavailable",
        stage=stage,
        path_kind=path_kind,
        recommended_action="review_os_error_code",
        host_access_approved=host_access_approved,
        error=error,
    )


def _validate_archive_root(
    root: Path,
    state_dir: Path,
    *,
    host_access_approved: bool = False,
) -> Path:
    candidate = Path(root).expanduser()
    path_kind = _archive_path_kind(candidate)
    if not candidate.is_absolute():
        if path_kind == "unc_share":
            raise ArchiveAccessError(
                "UNC syntax is not a local absolute path on this runtime. Mount the "
                "share as a local drive or folder and provide that mounted absolute "
                "path; Vera never requests SMB credentials.",
                code="archive_unc_requires_local_mount",
                stage="path_validation",
                path_kind=path_kind,
                recommended_action="mount_share_locally",
                host_access_approved=host_access_approved,
            )
        raise ArchiveError("Archive root must be an absolute path.")
    if candidate.is_symlink():
        raise ArchiveError("Archive root cannot be a symbolic link.")
    try:
        resolved = candidate.resolve(strict=True)
    except OSError as exc:
        raise _archive_access_error(
            exc,
            stage="path_resolution",
            path_kind=path_kind,
            host_access_approved=host_access_approved,
        ) from exc
    try:
        resolved_mode = resolved.stat().st_mode
    except OSError as exc:
        raise _archive_access_error(
            exc,
            stage="root_metadata",
            path_kind=path_kind,
            host_access_approved=host_access_approved,
        ) from exc
    if not stat.S_ISDIR(resolved_mode):
        raise ArchiveError("Archive root must be a directory.")
    if _path_is_within(state_dir, resolved) or _path_is_within(resolved, state_dir):
        raise ArchiveError(
            "The private Studio Archive state directory and source archive must "
            "not contain one another."
        )
    return resolved


def _scope_from_relative(relative_dir: str, root: Path) -> Scope:
    if relative_dir == ".":
        display_name = root.name or "Studio Archive"
    else:
        display_name = Path(relative_dir).name
    return Scope(
        scope_id=_stable_id("scope", relative_dir.casefold()),
        relative_dir=relative_dir,
        display_name=display_name,
    )


def _discover_top_level_scopes(
    root: Path,
    *,
    host_access_approved: bool = False,
) -> tuple[Scope, ...]:
    directories: list[Path] = []
    root_files = 0
    path_kind = _archive_path_kind(root)
    try:
        entries = sorted(root.iterdir(), key=lambda path: path.name.casefold())
    except OSError as exc:
        raise _archive_access_error(
            exc,
            stage="root_listing",
            path_kind=path_kind,
            host_access_approved=host_access_approved,
        ) from exc
    for path in entries:
        if path.name in IGNORED_NAMES:
            continue
        if path.is_symlink():
            root_files += 1
            continue
        try:
            mode = path.lstat().st_mode
        except OSError as exc:
            raise _archive_access_error(
                exc,
                stage="root_entry_metadata",
                path_kind=path_kind,
                host_access_approved=host_access_approved,
            ) from exc
        if stat.S_ISDIR(mode):
            directories.append(path)
        elif stat.S_ISREG(mode):
            root_files += 1
    if directories:
        directory_scopes = tuple(
            _scope_from_relative(path.relative_to(root).as_posix(), root)
            for path in directories
        )
        scopes = (
            (*directory_scopes, _scope_from_relative(".", root))
            if root_files
            else directory_scopes
        )
        if len({scope.scope_id for scope in scopes}) != len(scopes):
            raise ArchiveError(
                "Top-level archive directory names collide when case is ignored."
            )
        return scopes
    if root_files:
        return (_scope_from_relative(".", root),)
    return (_scope_from_relative(".", root),)


def _config_matches(
    path: Path,
    *,
    archive_root: Path,
    scopes: tuple[Scope, ...],
) -> bool:
    """Compare a persisted config without requiring its old paths to exist."""

    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return False
    return (
        payload.get("schema_version") == CONFIG_SCHEMA_VERSION
        and payload.get("archive_root") == str(archive_root)
        and payload.get("scopes") == [scope.as_json() for scope in scopes]
    )


def configure_archive(
    archive_root: Path,
    *,
    state_dir: Path | None = None,
    host_access_approved: bool = False,
) -> dict[str, Any]:
    """Configure one local archive and mechanically discover its top-level scopes."""

    planned_state = _state_dir(state_dir)
    root = _validate_archive_root(
        archive_root,
        planned_state,
        host_access_approved=host_access_approved,
    )
    private_state = _state_dir(state_dir, create=True)
    scopes = _discover_top_level_scopes(
        root,
        host_access_approved=host_access_approved,
    )
    _claim_state(private_state)
    config_path = _config_path(private_state)
    if config_path.is_file() and _config_matches(
        config_path,
        archive_root=root,
        scopes=scopes,
    ):
        if os.name == "posix":
            config_path.chmod(0o600)
        return studio_archive_status(state_dir=private_state)
    config = ArchiveConfig(
        archive_root=root,
        scopes=scopes,
        configured_at=_now_iso(),
    )
    _write_private_json(config_path, config.as_json())
    status = studio_archive_status(state_dir=private_state)
    status["index_requires_refresh"] = True
    return status


def _studio_archive_setup_contract() -> dict[str, Any]:
    """Describe the fixed local recovery path for first-time configuration."""

    return {
        "setup_required": True,
        "guided_setup": {
            "tool_name": "setup_studio_archive",
            "action": "select_folder_and_configure",
            "selection_mode": "native_directory_picker",
        },
        "manual_path_fallback": {
            "diagnose_tool_name": "diagnose_studio_archive_access",
            "configure_tool_name": "configure_studio_archive",
            "allowed_after": "archive_folder_picker_unavailable",
        },
    }


def _native_folder_picker_command() -> tuple[str, ...]:
    """Build fixed argv for the mechanically known desktop platform.

    Platform selection is deterministic because it enforces a local, no-shell
    UI boundary; it does not interpret the user's archive or client semantics.
    """

    if sys.platform == "darwin":
        executable = shutil.which("osascript")
        if executable is None:
            raise ArchiveFolderPickerUnavailableError()
        script = (
            'try\nPOSIX path of (choose folder with prompt "Select the Vera Studio '
            'Archive folder")\non error number -128\nreturn ""\nend try'
        )
        return (executable, "-e", script)
    if sys.platform == "win32":
        executable = (
            shutil.which("powershell.exe")
            or shutil.which("powershell")
            or shutil.which("pwsh")
        )
        if executable is None:
            raise ArchiveFolderPickerUnavailableError()
        script = (
            "Add-Type -AssemblyName System.Windows.Forms; "
            "$dialog = New-Object System.Windows.Forms.FolderBrowserDialog; "
            "$dialog.Description = 'Select the Vera Studio Archive folder'; "
            "$dialog.ShowNewFolderButton = $false; "
            "if ($dialog.ShowDialog() -eq "
            "[System.Windows.Forms.DialogResult]::OK) { "
            "[Console]::Out.Write($dialog.SelectedPath) }"
        )
        return (executable, "-NoProfile", "-STA", "-Command", script)
    for executable_name, arguments in (
        (
            "zenity",
            (
                "--file-selection",
                "--directory",
                "--title=Select the Vera Studio Archive folder",
            ),
        ),
        (
            "kdialog",
            (
                "--getexistingdirectory",
                ".",
                "--title",
                "Select the Vera Studio Archive folder",
            ),
        ),
    ):
        executable = shutil.which(executable_name)
        if executable is not None:
            return (executable, *arguments)
    raise ArchiveFolderPickerUnavailableError()


def _select_archive_root_with_native_picker() -> Path | None:
    """Return the directory selected by the user, or ``None`` after cancellation."""

    command = _native_folder_picker_command()
    try:
        # The command is fixed per supported platform; selected paths are output only.
        completed = subprocess.run(  # nosec B603
            command,
            capture_output=True,
            text=True,
            check=False,
            timeout=300,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise ArchiveFolderPickerUnavailableError() from exc
    selected = completed.stdout.strip()
    if completed.returncode == 0:
        return Path(selected) if selected else None
    if completed.returncode == 1 and not completed.stderr.strip():
        return None
    raise ArchiveFolderPickerUnavailableError()


def setup_archive_with_folder_picker(
    *,
    state_dir: Path | None = None,
    folder_selector: Callable[[], Path | None] | None = None,
) -> dict[str, Any]:
    """Select, diagnose, and configure one archive root through a local picker."""

    selector = folder_selector or _select_archive_root_with_native_picker
    selected_root = selector()
    if selected_root is None:
        return {
            "configured": False,
            "setup_status": "cancelled",
            **_studio_archive_setup_contract(),
        }
    diagnostic = diagnose_archive_access(selected_root, state_dir=state_dir)
    result = configure_archive(selected_root, state_dir=state_dir)
    result.pop("archive_root", None)
    result.update(
        {
            "setup_status": "configured",
            "setup_required": False,
            "archive_root_returned": False,
            "access_diagnostic": diagnostic,
        }
    )
    return result


def diagnose_archive_access(
    archive_root: Path,
    *,
    state_dir: Path | None = None,
    host_access_approved: bool = False,
) -> dict[str, Any]:
    """Verify archive-root path resolution and listing without persisting state."""

    planned_state = _state_dir(state_dir)
    root = _validate_archive_root(
        archive_root,
        planned_state,
        host_access_approved=host_access_approved,
    )
    scopes = _discover_top_level_scopes(
        root,
        host_access_approved=host_access_approved,
    )
    return {
        "ok": True,
        "path_kind": _archive_path_kind(root),
        "path_resolution": "available",
        "root_listing": "readable",
        "scope_count": len(scopes),
        "host_access_approved": host_access_approved,
        "client_ledger_write_access": "not_tested",
        "private_path_returned": False,
    }


def _load_config(
    state_dir: Path,
    *,
    validate_scope_roots: bool = True,
) -> ArchiveConfig:
    _claim_state(state_dir)
    path = _config_path(state_dir)
    if not path.is_file():
        raise ArchiveNotConfiguredError(
            "Studio Archive is not configured. Start the guided folder setup first."
        )
    _assert_private_file(path, "configuration")
    try:
        payload = json.loads(_STATE_LEASES[state_dir.resolve()][1] or b"null")
    except (OSError, json.JSONDecodeError) as exc:
        raise ArchiveError(
            f"Studio Archive configuration is unreadable: {exc}"
        ) from exc
    if payload.get("session_id", _session_id()) != _session_id():
        raise ArchiveError(_CONFIGURATION_CHANGED)
    if payload.get("schema_version") != CONFIG_SCHEMA_VERSION:
        raise ArchiveError("Studio Archive configuration version is unsupported.")
    raw_root = payload.get("archive_root")
    raw_scopes = payload.get("scopes")
    if not isinstance(raw_root, str) or not isinstance(raw_scopes, list):
        raise ArchiveError("Studio Archive configuration is malformed.")
    root = _validate_archive_root(Path(raw_root), state_dir)
    scopes: list[Scope] = []
    for item in raw_scopes:
        if not isinstance(item, dict):
            raise ArchiveError("Studio Archive scope configuration is malformed.")
        values = tuple(
            item.get(key) for key in ("scope_id", "relative_dir", "display_name")
        )
        if not all(isinstance(value, str) and value for value in values):
            raise ArchiveError("Studio Archive scope configuration is incomplete.")
        scope = Scope(
            scope_id=str(values[0]),
            relative_dir=str(values[1]),
            display_name=str(values[2]),
        )
        expected = _scope_from_relative(scope.relative_dir, root)
        if scope.scope_id != expected.scope_id:
            raise ArchiveError(
                "Studio Archive scope identifier does not match its path."
            )
        if scope.relative_dir != ".":
            relative = Path(scope.relative_dir)
            if (
                relative.is_absolute()
                or scope.relative_dir != relative.as_posix()
                or any(part in {"", ".", ".."} for part in relative.parts)
            ):
                raise ArchiveError(
                    "Configured archive scope is not a normalized relative path."
                )
        if validate_scope_roots:
            _resolve_scope_root(root, scope)
        scopes.append(scope)
    if not scopes or len({scope.scope_id for scope in scopes}) != len(scopes):
        raise ArchiveError("Studio Archive scopes must be non-empty and unique.")
    configured_at = str(payload.get("configured_at") or "")
    return ArchiveConfig(root, tuple(scopes), configured_at)


def _normalize_email_address(value: str) -> str:
    if not isinstance(value, str):
        raise ArchiveError("Client email addresses must be strings.")
    raw = value.strip()
    parsed = getaddresses([raw])
    if len(parsed) != 1:
        raise ArchiveError(f"Client email address is invalid: {value!r}.")
    address = parsed[0][1].strip().casefold()
    if (
        len(address) > 254
        or address.count("@") != 1
        or re.fullmatch(
            r"[A-Za-z0-9.!#$%&'*+/=?^_`|~-]+@[A-Za-z0-9.-]+",
            address,
        )
        is None
    ):
        raise ArchiveError(f"Client email address is invalid: {value!r}.")
    local_part, domain = address.rsplit("@", maxsplit=1)
    if (
        not local_part
        or not domain
        or domain.startswith((".", "-"))
        or domain.endswith((".", "-"))
        or ".." in domain
    ):
        raise ArchiveError(f"Client email address is invalid: {value!r}.")
    return address


def _normalize_email_addresses(values: Sequence[str]) -> tuple[str, ...]:
    if isinstance(values, (str, bytes)) or len(values) > MAX_CLIENT_EMAIL_ADDRESSES:
        raise ArchiveError(
            "A client may have at most "
            f"{MAX_CLIENT_EMAIL_ADDRESSES} confirmed email addresses."
        )
    normalized = {_normalize_email_address(value) for value in values}
    return tuple(sorted(normalized))


def _normalize_legal_names(values: Sequence[str]) -> tuple[str, ...]:
    if isinstance(values, (str, bytes)) or len(values) > MAX_CLIENT_LEGAL_NAMES:
        raise ArchiveError(
            f"A client may have at most {MAX_CLIENT_LEGAL_NAMES} legal names."
        )
    normalized: dict[str, str] = {}
    for value in values:
        if not isinstance(value, str):
            raise ArchiveError("Client legal names must be strings.")
        name = re.sub(r"\s+", " ", value).strip()
        if (
            not name
            or len(name) > 160
            or re.search(r"[\x00-\x1f\x7f]", name) is not None
        ):
            raise ArchiveError("Client legal names must contain 1 to 160 characters.")
        normalized.setdefault(name.casefold(), name)
    return tuple(normalized[key] for key in sorted(normalized))


def _normalize_tax_identifiers(values: Sequence[str]) -> tuple[str, ...]:
    if isinstance(values, (str, bytes)) or len(values) > MAX_CLIENT_TAX_IDENTIFIERS:
        raise ArchiveError(
            "A client may have at most "
            f"{MAX_CLIENT_TAX_IDENTIFIERS} tax identifiers."
        )
    normalized: set[str] = set()
    for value in values:
        if not isinstance(value, str):
            raise ArchiveError("Client tax identifiers must be strings.")
        identifier = re.sub(r"\s+", "", value).upper()
        if re.fullmatch(r"[A-Z0-9]{5,32}", identifier) is None:
            raise ArchiveError(
                "Client tax identifiers must contain 5 to 32 letters or digits."
            )
        normalized.add(identifier)
    return tuple(sorted(normalized))


def _validate_identity_uniqueness(records: Sequence[ClientIdentity]) -> None:
    client_ids: set[str] = set()
    scope_ids: set[str] = set()
    email_owners: dict[str, str] = {}
    tax_owners: dict[str, str] = {}
    for record in records:
        if record.client_id in client_ids:
            raise ArchiveError("A stable client ID is assigned more than once.")
        if record.scope_id in scope_ids:
            raise ArchiveError("An archive scope is assigned to more than one client.")
        client_ids.add(record.client_id)
        scope_ids.add(record.scope_id)
        for email_address in record.email_addresses:
            previous_client = email_owners.setdefault(email_address, record.client_id)
            if previous_client != record.client_id:
                raise ArchiveError(
                    "Client email address is assigned to more than one client: "
                    f"{email_address}."
                )
        for tax_identifier in record.tax_identifiers:
            previous_client = tax_owners.setdefault(tax_identifier, record.client_id)
            if previous_client != record.client_id:
                raise ArchiveError(
                    "Client tax identifier is assigned to more than one client: "
                    f"{tax_identifier}."
                )


def _load_client_identities(state_dir: Path) -> tuple[ClientIdentity, ...]:
    path = _client_identities_path(state_dir)
    if not path.is_file():
        return ()
    _assert_private_file(path, "client identity registry")
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ArchiveError(
            f"Studio Archive client identity registry is unreadable: {exc}"
        ) from exc
    if (
        not isinstance(payload, dict)
        or payload.get("schema_version") != CLIENT_IDENTITIES_SCHEMA_VERSION
        or not isinstance(payload.get("clients"), list)
    ):
        raise ArchiveError("Studio Archive client identity registry is malformed.")
    raw_clients = payload["clients"]
    if len(raw_clients) > MAX_CLIENT_IDENTITIES:
        raise ArchiveError("Studio Archive client identity registry is too large.")
    records: list[ClientIdentity] = []
    seen_client_ids: set[str] = set()
    seen_scope_ids: set[str] = set()
    required_keys = {
        "client_id",
        "scope_id",
        "email_addresses",
        "legal_names",
        "tax_identifiers",
        "updated_at",
    }
    for item in raw_clients:
        if not isinstance(item, dict) or set(item) != required_keys:
            raise ArchiveError("Studio Archive client identity record is malformed.")
        client_id = item["client_id"]
        scope_id = item["scope_id"]
        updated_at = item["updated_at"]
        if (
            not isinstance(client_id, str)
            or re.fullmatch(r"client_[0-9a-f]{24}", client_id) is None
            or client_id in seen_client_ids
            or not isinstance(scope_id, str)
            or re.fullmatch(r"scope_[0-9a-f]{24}", scope_id) is None
            or scope_id in seen_scope_ids
            or not isinstance(updated_at, str)
            or not updated_at
        ):
            raise ArchiveError("Studio Archive client identity record is invalid.")
        raw_emails = item["email_addresses"]
        raw_names = item["legal_names"]
        raw_tax_ids = item["tax_identifiers"]
        if not all(
            isinstance(values, list) for values in (raw_emails, raw_names, raw_tax_ids)
        ):
            raise ArchiveError("Studio Archive client identity values are malformed.")
        records.append(
            ClientIdentity(
                client_id=client_id,
                scope_id=scope_id,
                email_addresses=_normalize_email_addresses(raw_emails),
                legal_names=_normalize_legal_names(raw_names),
                tax_identifiers=_normalize_tax_identifiers(raw_tax_ids),
                updated_at=updated_at,
            )
        )
        seen_client_ids.add(client_id)
        seen_scope_ids.add(scope_id)
    _validate_identity_uniqueness(records)
    return tuple(sorted(records, key=lambda record: record.client_id))


def _write_client_identities(
    state_dir: Path,
    records: Sequence[ClientIdentity],
) -> None:
    _validate_identity_uniqueness(records)
    _write_private_json(
        _client_identities_path(state_dir),
        {
            "schema_version": CLIENT_IDENTITIES_SCHEMA_VERSION,
            "clients": [
                record.as_json()
                for record in sorted(records, key=lambda item: item.client_id)
            ],
        },
    )


def _normalize_engagement_label(value: str) -> str:
    if not isinstance(value, str):
        raise ArchiveError("Engagement label must be text.")
    label = re.sub(r"\s+", " ", value).strip()
    if not label or len(label) > 160 or re.search(r"[\x00-\x1f\x7f]", label):
        raise ArchiveError("Engagement label must contain 1 to 160 characters.")
    return label


def _discover_ledger_clients(config: ArchiveConfig) -> tuple[dict[str, Any], ...]:
    """Read portable client IDs from exact top-level customer folders."""

    scoped_roots = [
        (scope.scope_id, _resolve_scope_root(config.archive_root, scope))
        for scope in config.scopes
        if scope.relative_dir != "."
    ]
    try:
        return ledger.find_client_manifests(scoped_roots)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Customer-folder ledger is invalid: {exc}") from exc


def _synchronize_client_identities(
    state_dir: Path,
    config: ArchiveConfig,
) -> tuple[ClientIdentity, ...]:
    """Rebuild the private scope pointer from portable customer manifests.

    Email, PEC, legal-name, and tax-identifier values stay in the private
    registry.  Only the opaque client ID is recovered from the shared folder.
    """

    records = list(_load_client_identities(state_dir))
    by_client = {record.client_id: record for record in records}
    by_scope = {record.scope_id: record for record in records}
    changed = False
    for found in _discover_ledger_clients(config):
        client_id = found["client_id"]
        scope_id = found["scope_id"]
        scope_owner = by_scope.get(scope_id)
        if scope_owner is not None and scope_owner.client_id != client_id:
            raise ArchiveError(
                "A customer-folder manifest conflicts with the private client registry."
            )
        existing = by_client.get(client_id)
        if existing is None:
            replacement = ClientIdentity(
                client_id=client_id,
                scope_id=scope_id,
                email_addresses=(),
                legal_names=(),
                tax_identifiers=(),
                updated_at=_now_iso(),
            )
            records.append(replacement)
            by_client[client_id] = replacement
            by_scope[scope_id] = replacement
            changed = True
            continue
        if existing.scope_id == scope_id:
            continue
        replacement = ClientIdentity(
            client_id=existing.client_id,
            scope_id=scope_id,
            email_addresses=existing.email_addresses,
            legal_names=existing.legal_names,
            tax_identifiers=existing.tax_identifiers,
            updated_at=_now_iso(),
        )
        records = [
            replacement if record.client_id == client_id else record
            for record in records
        ]
        by_client[client_id] = replacement
        by_scope.pop(existing.scope_id, None)
        by_scope[scope_id] = replacement
        changed = True
    if changed:
        _write_client_identities(state_dir, records)
    return tuple(sorted(records, key=lambda record: record.client_id))


def _client_record(
    record: ClientIdentity | None,
    *,
    scope: Scope,
) -> dict[str, Any]:
    if record is None:
        return {
            **scope.as_json(),
            "client_id": None,
            "registration_status": "unregistered",
            "profile_status": "alias_only",
            "email_addresses": [],
            "legal_names": [],
            "tax_identifiers": [],
            "updated_at": None,
        }
    profile_status = "configured" if record.email_addresses else "candidate_only"
    return {
        **scope.as_json(),
        "client_id": record.client_id,
        "registration_status": "registered",
        "profile_status": profile_status,
        "email_addresses": list(record.email_addresses),
        "legal_names": list(record.legal_names),
        "tax_identifiers": list(record.tax_identifiers),
        "updated_at": record.updated_at,
    }


def list_studio_client_identities(
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """List exact archive scopes and their private Gmail identity profiles."""

    private_state = _state_dir(state_dir)
    stored_config = _load_config(private_state, validate_scope_roots=False)
    config, scopes_changed = _current_scope_view(stored_config)
    records = (
        _load_client_identities(private_state)
        if scopes_changed
        else _synchronize_client_identities(private_state, config)
    )
    records_by_scope = {record.scope_id: record for record in records}
    active_scope_ids = {scope.scope_id for scope in config.scopes}
    clients = [
        _client_record(records_by_scope.get(scope.scope_id), scope=scope)
        for scope in config.scopes
    ]
    orphaned = [
        {
            **record.as_json(),
            "profile_status": "orphaned",
        }
        for record in records
        if record.scope_id not in active_scope_ids
    ]
    return {
        "scope_configuration_changed": scopes_changed,
        "registered_client_count": sum(
            client["registration_status"] == "registered" for client in clients
        ),
        "unregistered_scope_count": sum(
            client["registration_status"] == "unregistered" for client in clients
        ),
        "configured_profile_count": sum(
            client["profile_status"] == "configured" for client in clients
        ),
        "candidate_only_profile_count": sum(
            client["profile_status"] == "candidate_only" for client in clients
        ),
        "alias_only_profile_count": sum(
            client["profile_status"] == "alias_only" for client in clients
        ),
        "orphaned_profile_count": len(orphaned),
        "clients": clients,
        "orphaned_profiles": orphaned,
        "gmail_connector_called": False,
    }


def _client_directory_record(
    record: ClientIdentity | None,
    *,
    scope: Scope | None,
) -> dict[str, Any]:
    """Project one private identity record onto its model-safe directory row.

    Exact identity values stay in the owner-only registry.  Counts preserve the
    operational distinction between an empty and configured profile without
    disclosing another client's email, legal-name aliases, or tax identifiers.
    """

    if scope is None:
        if record is None:
            raise ArchiveError("An orphaned client directory row requires a profile.")
        return {
            "client_id": record.client_id,
            "scope_id": record.scope_id,
            "registration_status": "orphaned",
            "profile_status": "orphaned",
            "identity_counts": {
                "email_addresses": len(record.email_addresses),
                "legal_names": len(record.legal_names),
                "tax_identifiers": len(record.tax_identifiers),
            },
        }
    if record is None:
        return {
            **scope.as_json(),
            "client_id": None,
            "registration_status": "unregistered",
            "profile_status": "alias_only",
            "identity_counts": {
                "email_addresses": 0,
                "legal_names": 0,
                "tax_identifiers": 0,
            },
        }
    return {
        **scope.as_json(),
        "client_id": record.client_id,
        "registration_status": "registered",
        "profile_status": (
            "configured" if record.email_addresses else "candidate_only"
        ),
        "identity_counts": {
            "email_addresses": len(record.email_addresses),
            "legal_names": len(record.legal_names),
            "tax_identifiers": len(record.tax_identifiers),
        },
    }


def list_studio_clients(
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """List client scopes without exposing the private identity registry."""

    private_state = _state_dir(state_dir)
    if not _config_path(private_state).is_file():
        return {
            "configured": False,
            "scope_configuration_changed": False,
            "registered_client_count": 0,
            "unregistered_scope_count": 0,
            "configured_profile_count": 0,
            "candidate_only_profile_count": 0,
            "alias_only_profile_count": 0,
            "orphaned_profile_count": 0,
            "clients": [],
            "orphaned_profiles": [],
            "private_identity_values_returned": False,
            "gmail_connector_called": False,
            **_studio_archive_setup_contract(),
        }
    stored_config = _load_config(private_state, validate_scope_roots=False)
    config, scopes_changed = _current_scope_view(stored_config)
    records = (
        _load_client_identities(private_state)
        if scopes_changed
        else _synchronize_client_identities(private_state, config)
    )
    records_by_scope = {record.scope_id: record for record in records}
    active_scope_ids = {scope.scope_id for scope in config.scopes}
    clients = [
        _client_directory_record(records_by_scope.get(scope.scope_id), scope=scope)
        for scope in config.scopes
    ]
    orphaned = [
        _client_directory_record(record, scope=None)
        for record in records
        if record.scope_id not in active_scope_ids
    ]
    return {
        "configured": True,
        "scope_configuration_changed": scopes_changed,
        "registered_client_count": sum(
            client["registration_status"] == "registered" for client in clients
        ),
        "unregistered_scope_count": sum(
            client["registration_status"] == "unregistered" for client in clients
        ),
        "configured_profile_count": sum(
            client["profile_status"] == "configured" for client in clients
        ),
        "candidate_only_profile_count": sum(
            client["profile_status"] == "candidate_only" for client in clients
        ),
        "alias_only_profile_count": sum(
            client["profile_status"] == "alias_only" for client in clients
        ),
        "orphaned_profile_count": len(orphaned),
        "clients": clients,
        "orphaned_profiles": orphaned,
        "private_identity_values_returned": False,
        "gmail_connector_called": False,
    }


def resolve_studio_client_identity(
    identity_kind: str,
    identity_value: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Resolve one user-supplied identity by exact local equality.

    Exact normalized equality is deterministic because it is an identity and
    security boundary.  This function does not rank semantic name similarity
    and never returns the stored registry values.
    """

    if identity_kind == "email_address":
        normalized: str = _normalize_email_address(identity_value)
        field = "email_addresses"
    elif identity_kind == "legal_name":
        normalized = _normalize_legal_names([identity_value])[0].casefold()
        field = "legal_names"
    elif identity_kind == "tax_identifier":
        normalized = _normalize_tax_identifiers([identity_value])[0]
        field = "tax_identifiers"
    else:
        raise ArchiveError(
            "identity_kind must be email_address, legal_name, or tax_identifier."
        )

    private_state = _state_dir(state_dir)
    stored_config = _load_config(private_state, validate_scope_roots=False)
    config, scopes_changed = _current_scope_view(stored_config)
    records = (
        _load_client_identities(private_state)
        if scopes_changed
        else _synchronize_client_identities(private_state, config)
    )
    scopes_by_id = {scope.scope_id: scope for scope in config.scopes}
    matches: list[dict[str, Any]] = []
    for record in records:
        values = getattr(record, field)
        comparison_values = (
            {value.casefold() for value in values}
            if identity_kind == "legal_name"
            else set(values)
        )
        if normalized not in comparison_values:
            continue
        matches.append(
            _client_directory_record(record, scope=scopes_by_id.get(record.scope_id))
        )
    return {
        "resolution_status": (
            "exact_match"
            if len(matches) == 1
            else "ambiguous_exact_match" if matches else "no_exact_match"
        ),
        "identity_kind": identity_kind,
        "match_count": len(matches),
        "matches": matches,
        "private_identity_values_returned": False,
        "scope_configuration_changed": scopes_changed,
        "gmail_connector_called": False,
    }


def get_studio_client_folder(
    client_id: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Return a portable binding for one registered, stable client ID."""

    private_state = _state_dir(state_dir)
    stored_config = _load_config(private_state, validate_scope_roots=False)
    config, scopes_changed = _current_scope_view(stored_config)
    if scopes_changed:
        raise ArchiveError(
            "Top-level archive scopes changed; refresh before selecting a client folder."
        )
    identities = _synchronize_client_identities(private_state, config)
    identity = next(
        (item for item in identities if item.client_id == client_id),
        None,
    )
    if identity is None:
        raise ArchiveError("Stable client ID is not registered.")
    scope = next(
        (item for item in config.scopes if item.scope_id == identity.scope_id),
        None,
    )
    if scope is None:
        raise ArchiveError(
            "The registered client folder is missing; relink it before continuing."
        )
    if scope.relative_dir == ".":
        raise ArchiveError(
            "The archive root scope is not a client folder; select one immediate "
            "top-level client directory."
        )
    client_root = _resolve_scope_root(config.archive_root, scope)
    binding = build_studio_client_folder_binding(
        studio_client_id=identity.client_id,
        scope_id=scope.scope_id,
        archive_root=config.archive_root,
        scope_relative_dir=scope.relative_dir,
        client_root=client_root,
        display_name=scope.display_name,
    )
    return {
        "status": "ready",
        "client_folder": binding,
        "source_archive_mutated": False,
        "gmail_connector_called": False,
    }


def set_studio_client_identity(
    scope_id: str,
    *,
    email_addresses: Sequence[str] = (),
    legal_names: Sequence[str] = (),
    tax_identifiers: Sequence[str] = (),
    replace_orphaned_scope_id: str | None = None,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Replace one scope's confirmed private Gmail identity profile."""

    private_state = _state_dir(state_dir)
    config = _load_config(private_state, validate_scope_roots=False)
    _, scopes_changed = _current_scope_view(config)
    if scopes_changed:
        raise ArchiveError(
            "Top-level archive scopes changed; refresh before configuring clients."
        )
    scopes_by_id = {scope.scope_id: scope for scope in config.scopes}
    scope = scopes_by_id.get(scope_id)
    if scope is None:
        raise ArchiveError("Client identity scope is not configured.")
    if scope.relative_dir == ".":
        raise ArchiveError(
            "The archive root scope is not a client folder; select one immediate "
            "top-level client directory."
        )
    records = list(_synchronize_client_identities(private_state, config))
    existing = next(
        (record for record in records if record.scope_id == scope_id),
        None,
    )
    if replace_orphaned_scope_id is not None:
        if (
            not isinstance(replace_orphaned_scope_id, str)
            or re.fullmatch(
                r"scope_[0-9a-f]{24}",
                replace_orphaned_scope_id,
            )
            is None
        ):
            raise ArchiveError("Replacement client scope identifier is invalid.")
        if replace_orphaned_scope_id in scopes_by_id:
            raise ArchiveError(
                "Only an orphaned client profile can be explicitly rebound."
            )
        orphaned = next(
            (
                record
                for record in records
                if record.scope_id == replace_orphaned_scope_id
            ),
            None,
        )
        if orphaned is None:
            raise ArchiveError("Orphaned client profile was not found.")
        if existing is not None:
            raise ArchiveError(
                "The target client scope already has an identity profile."
            )
        if email_addresses or legal_names or tax_identifiers:
            raise ArchiveError(
                "Do not supply identity values while rebinding an orphaned profile."
            )
        replacement = ClientIdentity(
            client_id=orphaned.client_id,
            scope_id=scope_id,
            email_addresses=orphaned.email_addresses,
            legal_names=orphaned.legal_names,
            tax_identifiers=orphaned.tax_identifiers,
            updated_at=_now_iso(),
        )
        updated_records = [
            record for record in records if record.scope_id != replace_orphaned_scope_id
        ] + [replacement]
        try:
            ledger.create_client_manifest(
                _resolve_scope_root(config.archive_root, scope),
                replacement.client_id,
            )
        except ledger.LedgerError as exc:
            raise ArchiveError(f"Customer-folder ledger is invalid: {exc}") from exc
        _write_client_identities(private_state, updated_records)
        return {
            "status": "rebound",
            "client": _client_directory_record(replacement, scope=scope),
            "replaced_orphaned_scope_id": replace_orphaned_scope_id,
            "gmail_connector_called": False,
            "gmail_credentials_stored": False,
        }
    normalized_emails = _normalize_email_addresses(email_addresses)
    normalized_names = _normalize_legal_names(legal_names)
    normalized_tax_ids = _normalize_tax_identifiers(tax_identifiers)
    if not (normalized_emails or normalized_names or normalized_tax_ids):
        raise ArchiveError(
            "Configure at least one confirmed email address, legal name, "
            "or tax identifier."
        )
    unchanged = existing is not None and (
        existing.email_addresses == normalized_emails
        and existing.legal_names == normalized_names
        and existing.tax_identifiers == normalized_tax_ids
    )
    replacement = ClientIdentity(
        client_id=(
            existing.client_id
            if existing is not None
            else _new_private_id("client", {record.client_id for record in records})
        ),
        scope_id=scope_id,
        email_addresses=normalized_emails,
        legal_names=normalized_names,
        tax_identifiers=normalized_tax_ids,
        updated_at=existing.updated_at if unchanged else _now_iso(),
    )
    updated_records = [record for record in records if record.scope_id != scope_id] + [
        replacement
    ]
    _validate_identity_uniqueness(updated_records)
    try:
        ledger.create_client_manifest(
            _resolve_scope_root(config.archive_root, scope),
            replacement.client_id,
        )
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Customer-folder ledger is invalid: {exc}") from exc
    if not unchanged:
        _write_client_identities(private_state, updated_records)
    return {
        "status": "unchanged" if unchanged else "configured",
        "client": _client_directory_record(replacement, scope=scope),
        "gmail_connector_called": False,
        "gmail_credentials_stored": False,
    }


def _safe_client_directory_name(legal_name: str) -> str:
    """Derive a portable display folder; the stable client ID remains separate."""

    normalized = _normalize_legal_names([legal_name])[0]
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f\x7f]', " ", normalized)
    name = re.sub(r"\s+", " ", name).strip(" .")[:120].rstrip(" .")
    reserved = {
        "CON",
        "PRN",
        "AUX",
        "NUL",
        *{f"COM{number}" for number in range(1, 10)},
        *{f"LPT{number}" for number in range(1, 10)},
    }
    if not name or name.upper() in reserved or name in IGNORED_NAMES:
        raise ArchiveError(
            "The confirmed legal name cannot produce a safe client folder label."
        )
    return name


def create_studio_client(
    legal_name: str,
    *,
    email_addresses: Sequence[str] = (),
    tax_identifiers: Sequence[str] = (),
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Create one registered client folder after the user chooses New client."""

    private_state = _state_dir(state_dir)
    config = _load_config(private_state, validate_scope_roots=False)
    _, scopes_changed = _current_scope_view(config)
    if scopes_changed:
        raise ArchiveError(
            "Top-level archive scopes changed; refresh before creating a client."
        )
    normalized_names = _normalize_legal_names([legal_name])
    normalized_emails = _normalize_email_addresses(email_addresses)
    normalized_tax_ids = _normalize_tax_identifiers(tax_identifiers)
    directory_name = _safe_client_directory_name(normalized_names[0])
    existing_by_name = {scope.relative_dir.casefold(): scope for scope in config.scopes}
    if directory_name.casefold() in existing_by_name:
        existing_scope = existing_by_name[directory_name.casefold()]
        raise ArchiveError(
            "A top-level folder already has the generated client label; select and "
            f"register existing scope {existing_scope.scope_id} instead."
        )
    records = list(_load_client_identities(private_state))
    client_id = _new_private_id("client", {record.client_id for record in records})
    client_root = config.archive_root / directory_name
    if client_root.exists() or client_root.is_symlink():
        raise ArchiveError("The generated client folder already exists.")
    client_root.mkdir()
    try:
        current_scopes = _discover_top_level_scopes(config.archive_root)
        scope = next(
            (item for item in current_scopes if item.relative_dir == directory_name),
            None,
        )
        if scope is None:
            raise ArchiveError("The new client folder was not discovered exactly.")
        record = ClientIdentity(
            client_id=client_id,
            scope_id=scope.scope_id,
            email_addresses=normalized_emails,
            legal_names=normalized_names,
            tax_identifiers=normalized_tax_ids,
            updated_at=_now_iso(),
        )
        try:
            ledger.create_client_manifest(client_root, client_id)
        except ledger.LedgerError as exc:
            raise ArchiveError(f"Customer-folder ledger is invalid: {exc}") from exc
        updated_records = [*records, record]
        _validate_identity_uniqueness(updated_records)
        updated_config = ArchiveConfig(
            archive_root=config.archive_root,
            scopes=current_scopes,
            configured_at=_now_iso(),
        )
        _write_private_json(_config_path(private_state), updated_config.as_json())
        try:
            _write_client_identities(private_state, updated_records)
        except (ArchiveError, OSError):
            _write_private_json(_config_path(private_state), config.as_json())
            raise
    except (ArchiveError, OSError):
        shutil.rmtree(client_root, ignore_errors=True)
        raise
    folder = get_studio_client_folder(client_id, state_dir=private_state)
    return {
        "status": "created",
        "client": _client_directory_record(record, scope=scope),
        "client_folder": folder["client_folder"],
        "relationship_setup_status": "new_client_workflow_pending",
        "source_archive_mutated": True,
        "next_workflow": "new-client",
    }


def create_studio_client_engagement(
    client_id: str,
    engagement_label: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Create one durable engagement inside the selected customer folder."""

    private_state = _state_dir(state_dir)
    folder = get_studio_client_folder(client_id, state_dir=private_state)[
        "client_folder"
    ]
    client_root = Path(folder["client_root"])
    try:
        engagement = ledger.create_engagement(
            client_root,
            client_id,
            _normalize_engagement_label(engagement_label),
        )
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Customer-folder engagement is invalid: {exc}") from exc
    input_root = (
        client_root
        / ledger.LEDGER_DIRECTORY
        / "engagements"
        / engagement["engagement_id"]
        / "inputs"
    )
    return {
        "status": "created",
        "client_id": client_id,
        "engagement": {**engagement, "imports": []},
        "input_dir": str(input_root.resolve(strict=True)),
        "source_archive_mutated": True,
    }


def _archive_organization_inventory_ref(snapshot_sha256: str) -> str:
    digest = hashlib.sha256(
        f"archive-organization-inventory-v1\0{snapshot_sha256}".encode("utf-8")
    ).hexdigest()
    return ARCHIVE_ORGANIZATION_INVENTORY_REF_PREFIX + digest[:24]


def _archive_organization_item_ref(
    snapshot_sha256: str,
    relative_path: str,
) -> str:
    digest = hashlib.sha256(
        ("archive-organization-item-v1\0" f"{snapshot_sha256}\0{relative_path}").encode(
            "utf-8"
        )
    ).hexdigest()
    return ARCHIVE_ORGANIZATION_ITEM_REF_PREFIX + digest[:24]


def _project_drive_relative_path(snapshot_sha256: str, relative_path: str) -> str:
    """Replace Drive-ID suffixes with stable unlinkable display references."""

    projected: list[str] = []
    for index, component in enumerate(relative_path.split("/")):
        match = re.fullmatch(
            r"(?P<label>.*)~(?P<drive_ref>[A-Za-z0-9_-]{12})", component
        )
        if match is None:
            projected.append(component)
            continue
        opaque = hashlib.sha256(
            (
                "archive-organization-path-component-v1\0"
                f"{snapshot_sha256}\0{index}\0{component}"
            ).encode("utf-8")
        ).hexdigest()[:10]
        projected.append(f"{match.group('label')}~ref_{opaque}")
    return "/".join(projected)


def _archive_organization_duplicate_key(
    item: Mapping[str, Any],
    *,
    storage_kind: str,
) -> str | None:
    if storage_kind == "local_filesystem":
        return str(item["sha256"])
    value = item.get("sha256_checksum")
    return str(value) if value else None


def _archive_organization_open_supported(
    item: Mapping[str, Any],
    *,
    storage_kind: str,
) -> bool:
    if storage_kind == "local_filesystem":
        return Path(str(item["relative_path"])).suffix.lower() in SUPPORTED_SUFFIXES
    capabilities = item.get("capabilities")
    if (
        not isinstance(capabilities, Mapping)
        or capabilities.get("can_download") is not True
    ):
        return False
    mime_type = str(item.get("mime_type") or "")
    if mime_type in GOOGLE_DRIVE_EXPORTS:
        return True
    if mime_type.startswith("application/vnd.google-apps."):
        return False
    return Path(
        str(item.get("name") or "")
    ).suffix.lower() in SUPPORTED_SUFFIXES or mime_type.startswith("text/")


def _project_archive_organization_inventory(
    snapshot: Mapping[str, Any],
) -> dict[str, Any]:
    """Return the full semantic inventory without execution-only identifiers."""

    snapshot_sha256 = str(snapshot["content_sha256"])
    storage_kind = (
        "google_drive"
        if snapshot.get("schema_version") == drive.DRIVE_SNAPSHOT_SCHEMA
        else "local_filesystem"
    )
    files = snapshot.get("files")
    if not isinstance(files, list):
        raise ArchiveError("Archive organization snapshot files are invalid.")
    refs_by_path = {
        str(item["relative_path"]): _archive_organization_item_ref(
            snapshot_sha256,
            str(item["relative_path"]),
        )
        for item in files
    }
    duplicate_paths: dict[str, list[str]] = {}
    for item in files:
        duplicate_key = _archive_organization_duplicate_key(
            item,
            storage_kind=storage_kind,
        )
        if duplicate_key is None:
            continue
        duplicate_paths.setdefault(duplicate_key, []).append(str(item["relative_path"]))
    canonical_by_path: dict[str, str] = {}
    duplicate_ref_by_path: dict[str, str] = {}
    for duplicate_key, paths in duplicate_paths.items():
        if len(paths) < 2:
            continue
        canonical = min(
            paths,
            key=lambda value: (
                len(Path(value).parts),
                value.casefold(),
                value,
            ),
        )
        group_ref = (
            ARCHIVE_ORGANIZATION_DUPLICATE_REF_PREFIX
            + hashlib.sha256(
                (
                    "archive-organization-exact-group-v1\0"
                    f"{snapshot_sha256}\0{duplicate_key}"
                ).encode("utf-8")
            ).hexdigest()[:20]
        )
        for relative_path in paths:
            canonical_by_path[relative_path] = canonical
            duplicate_ref_by_path[relative_path] = group_ref

    projected_files: list[dict[str, Any]] = []
    for item in files:
        relative_path = str(item["relative_path"])
        display_path = (
            _project_drive_relative_path(snapshot_sha256, relative_path)
            if storage_kind == "google_drive"
            else relative_path
        )
        canonical = canonical_by_path.get(relative_path)
        if storage_kind == "local_filesystem":
            modified_at = datetime.fromtimestamp(
                int(item["modified_ns"]) / 1_000_000_000,
                tz=timezone.utc,
            ).isoformat()
            size_bytes = int(item["byte_count"])
            mime_type = None
            name = Path(relative_path).name
        else:
            modified_at = str(item["modified_time"])
            size_bytes = item["size_bytes"]
            mime_type = str(item["mime_type"])
            name = str(item["name"])
        projected_files.append(
            {
                "item_ref": refs_by_path[relative_path],
                "relative_path": display_path,
                "name": name,
                "file_extension": Path(name).suffix.lower() or None,
                "mime_type": mime_type,
                "size_bytes": size_bytes,
                "modified_at": modified_at,
                "evidence_access": (
                    "available"
                    if _archive_organization_open_supported(
                        item,
                        storage_kind=storage_kind,
                    )
                    else "unsupported"
                ),
                "exact_duplicate_group": duplicate_ref_by_path.get(relative_path),
                "exact_duplicate_of": (
                    refs_by_path[canonical]
                    if canonical is not None and canonical != relative_path
                    else None
                ),
            }
        )
    excluded = snapshot.get("excluded")
    if not isinstance(excluded, list):
        raise ArchiveError("Archive organization snapshot exclusions are invalid.")
    projected_excluded = []
    for item in excluded:
        relative_path = str(item["relative_path"])
        projected_excluded.append(
            {
                "relative_path": (
                    _project_drive_relative_path(snapshot_sha256, relative_path)
                    if storage_kind == "google_drive"
                    else relative_path
                ),
                "reason": str(item["reason"]),
            }
        )
    return {
        "schema_version": ARCHIVE_ORGANIZATION_INVENTORY_SCHEMA,
        "inventory_ref": _archive_organization_inventory_ref(snapshot_sha256),
        "storage_kind": storage_kind,
        "root_name": snapshot.get("root_name"),
        "captured_at": str(snapshot["captured_at"]),
        "file_count": len(projected_files),
        "known_total_bytes": (
            int(snapshot["total_bytes"])
            if storage_kind == "local_filesystem"
            else int(snapshot["known_total_bytes"])
        ),
        "files": projected_files,
        "excluded": projected_excluded,
        "raw_hashes_returned": False,
        "drive_ids_returned": False,
        "absolute_paths_returned": False,
    }


def snapshot_studio_client_folder(
    client_id: str,
    engagement_id: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Capture and import a bounded file-identity snapshot for one client folder."""

    private_state = _state_dir(state_dir)
    root = _selected_ledger_root(client_id, engagement_id, state_dir=private_state)
    try:
        result = ledger.snapshot_client_folder(root, client_id, engagement_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Client-folder snapshot failed: {exc}") from exc
    return {
        "status": result["status"],
        "input_id": result["input_id"],
        "model_inventory": _project_archive_organization_inventory(result["snapshot"]),
        "documents_copied": False,
        "source_archive_mutated": result["source_archive_mutated"],
    }


def _drive_bindings_digest(content: Mapping[str, Any]) -> str:
    return hashlib.sha256(
        json.dumps(
            content,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()


def _load_google_drive_bindings(state_dir: Path) -> list[dict[str, Any]]:
    path = _google_drive_bindings_path(state_dir)
    if not path.is_file():
        return []
    _assert_private_file(path, "Google Drive bindings")
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ArchiveError(f"Google Drive bindings are unreadable: {exc}") from exc
    if not isinstance(payload, dict) or set(payload) != {
        "schema_version",
        "bindings",
        "content_sha256",
    }:
        raise ArchiveError("Google Drive bindings are malformed.")
    content = {
        "schema_version": payload["schema_version"],
        "bindings": payload["bindings"],
    }
    if (
        payload["schema_version"] != GOOGLE_DRIVE_BINDINGS_SCHEMA
        or payload["content_sha256"] != _drive_bindings_digest(content)
        or not isinstance(payload["bindings"], list)
    ):
        raise ArchiveError("Google Drive bindings are invalid or stale.")
    normalized: list[dict[str, Any]] = []
    seen_clients: set[str] = set()
    seen_folders: set[str] = set()
    for item in payload["bindings"]:
        if not isinstance(item, dict) or set(item) != {
            "client_id",
            "folder_id",
            "drive_id",
            "display_name",
            "bound_at",
        }:
            raise ArchiveError("Google Drive binding shape is invalid.")
        client_id = str(item["client_id"])
        folder_id = str(item["folder_id"])
        if client_id in seen_clients or folder_id in seen_folders:
            raise ArchiveError("Google Drive bindings are not one-to-one.")
        if re.fullmatch(r"client_[0-9a-f]{24}", client_id) is None:
            raise ArchiveError("Google Drive binding client ID is invalid.")
        if re.fullmatch(r"[A-Za-z0-9_-]{3,256}", folder_id) is None:
            raise ArchiveError("Google Drive binding folder ID is invalid.")
        drive_id = item["drive_id"]
        if drive_id is not None and (
            not isinstance(drive_id, str)
            or re.fullmatch(r"[A-Za-z0-9_-]{3,256}", drive_id) is None
        ):
            raise ArchiveError("Google Drive binding Shared Drive ID is invalid.")
        display_name = str(item["display_name"]).strip()
        bound_at = str(item["bound_at"]).strip()
        if not display_name or not bound_at:
            raise ArchiveError("Google Drive binding metadata is invalid.")
        normalized.append(
            {
                "client_id": client_id,
                "folder_id": folder_id,
                "drive_id": drive_id,
                "display_name": display_name,
                "bound_at": bound_at,
            }
        )
        seen_clients.add(client_id)
        seen_folders.add(folder_id)
    return normalized


def _write_google_drive_bindings(
    state_dir: Path, bindings: Sequence[Mapping[str, Any]]
) -> None:
    ordered = sorted(
        (dict(item) for item in bindings), key=lambda item: item["client_id"]
    )
    content = {
        "schema_version": GOOGLE_DRIVE_BINDINGS_SCHEMA,
        "bindings": ordered,
    }
    _write_private_json(
        _google_drive_bindings_path(state_dir),
        {**content, "content_sha256": _drive_bindings_digest(content)},
    )


def authorize_studio_google_drive(
    client_secrets_path: Path,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Run the explicit Drive OAuth flow and store its token in private state."""

    private_state = _state_dir(state_dir, create=True)
    try:
        drive.authorize_google_drive(
            client_secrets_path,
            _google_drive_token_path(private_state),
        )
    except drive.DriveError as exc:
        raise ArchiveError(str(exc)) from exc
    return {
        "status": "authorized",
        "scope": drive.DRIVE_SCOPE,
        "token_path": str(_google_drive_token_path(private_state)),
        "credentials_persisted_privately": True,
        "external_service": "google-drive",
    }


def studio_google_drive_status(
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Report local authorization and client-to-folder bindings without an API call."""

    private_state = _state_dir(state_dir)
    token_path = _google_drive_token_path(private_state)
    if token_path.exists():
        _assert_private_file(token_path, "Google Drive token")
    bindings = _load_google_drive_bindings(private_state)
    return {
        "status": "ready" if token_path.is_file() else "authorization_required",
        "oauth_scope": drive.DRIVE_SCOPE,
        "restricted_scope": True,
        "token_present": token_path.is_file(),
        "binding_count": len(bindings),
        "bindings": bindings,
        "google_drive_api_called": False,
    }


def _drive_gateway(
    state_dir: Path,
    gateway: drive.DriveGateway | None,
) -> drive.DriveGateway:
    if gateway is not None:
        return gateway
    try:
        return drive.load_google_drive_gateway(_google_drive_token_path(state_dir))
    except drive.DriveError as exc:
        raise ArchiveError(str(exc)) from exc


def bind_studio_client_google_drive(
    client_id: str,
    folder_id: str,
    *,
    state_dir: Path | None = None,
    gateway: drive.DriveGateway | None = None,
) -> dict[str, Any]:
    """Bind one registered Studio Archive client to one exact Drive folder ID."""

    private_state = _state_dir(state_dir, create=True)
    get_studio_client_folder(client_id, state_dir=private_state)
    selected_gateway = _drive_gateway(private_state, gateway)
    try:
        folder = selected_gateway.get_file(folder_id)
    except drive.DriveError as exc:
        raise ArchiveError(str(exc)) from exc
    if (
        folder.get("mimeType") != drive.DRIVE_FOLDER_MIME_TYPE
        or folder.get("trashed") is True
    ):
        raise ArchiveError("Selected Google Drive item is not an active folder.")
    observed_id = folder.get("id")
    display_name = folder.get("name")
    if (
        observed_id != folder_id
        or not isinstance(display_name, str)
        or not display_name.strip()
        or len(display_name.strip()) > 768
    ):
        raise ArchiveError("Google Drive returned invalid root-folder metadata.")
    drive_id = folder.get("driveId")
    if drive_id is not None and (
        not isinstance(drive_id, str)
        or re.fullmatch(r"[A-Za-z0-9_-]{3,256}", drive_id) is None
    ):
        raise ArchiveError("Google Drive returned an invalid Shared Drive ID.")
    existing_bindings = _load_google_drive_bindings(private_state)
    if any(
        item["folder_id"] == folder_id and item["client_id"] != client_id
        for item in existing_bindings
    ):
        raise ArchiveError(
            "This Google Drive folder is already bound to another Vera client."
        )
    bindings = [item for item in existing_bindings if item["client_id"] != client_id]
    binding = {
        "client_id": client_id,
        "folder_id": folder_id,
        "drive_id": drive_id,
        "display_name": display_name.strip(),
        "bound_at": _now_iso(),
    }
    bindings.append(binding)
    _write_google_drive_bindings(private_state, bindings)
    return {
        "status": "bound",
        "binding": binding,
        "google_drive_api_called": True,
        "source_archive_mutated": False,
    }


def _google_drive_binding(client_id: str, state_dir: Path) -> dict[str, Any]:
    matches = [
        item
        for item in _load_google_drive_bindings(state_dir)
        if item["client_id"] == client_id
    ]
    if len(matches) != 1:
        raise ArchiveError(
            "Select and bind exactly one Google Drive folder for this client first."
        )
    return matches[0]


def snapshot_studio_client_google_drive(
    client_id: str,
    engagement_id: str,
    *,
    state_dir: Path | None = None,
    gateway: drive.DriveGateway | None = None,
) -> dict[str, Any]:
    """Snapshot one bound Drive tree and import only its immutable JSON receipt."""

    private_state = _state_dir(state_dir)
    client_root = _selected_ledger_root(
        client_id, engagement_id, state_dir=private_state
    )
    binding = _google_drive_binding(client_id, private_state)
    selected_gateway = _drive_gateway(private_state, gateway)
    try:
        snapshot = drive.snapshot_google_drive_folder(
            selected_gateway,
            binding["folder_id"],
            client_id,
            engagement_id,
        )
    except drive.DriveError as exc:
        raise ArchiveError(str(exc)) from exc
    if (
        snapshot["root_name"] != binding["display_name"]
        or snapshot["drive_id"] != binding["drive_id"]
    ):
        raise ArchiveError(
            "The bound Google Drive root changed identity; review the binding."
        )
    engagement_root = (
        client_root / ledger.LEDGER_DIRECTORY / "engagements" / engagement_id
    )
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=".vera-google-drive-snapshot-",
        suffix=".json",
        dir=engagement_root,
    )
    temporary = Path(temporary_name)
    try:
        with os.fdopen(descriptor, "w", encoding="utf-8") as handle:
            json.dump(snapshot, handle, ensure_ascii=False, indent=2, sort_keys=True)
            handle.write("\n")
            handle.flush()
            os.fsync(handle.fileno())
        imported = ledger.import_document(
            client_root,
            client_id,
            engagement_id,
            temporary,
            "source",
        )
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Google Drive snapshot import failed: {exc}") from exc
    finally:
        temporary.unlink(missing_ok=True)
    return {
        "status": imported["status"],
        "input_id": imported["receipt"]["input_id"],
        "model_inventory": _project_archive_organization_inventory(snapshot),
        "documents_copied": False,
        "google_drive_api_called": True,
        "remote_archive_mutated": False,
        "studio_ledger_mutated": True,
    }


def _load_archive_organization_snapshot(
    client_id: str,
    engagement_id: str,
    snapshot_input_id: str,
    *,
    state_dir: Path,
) -> tuple[Path, dict[str, Any]]:
    client_root = _selected_ledger_root(
        client_id,
        engagement_id,
        state_dir=state_dir,
    )
    try:
        receipt = ledger.load_input_receipt(
            client_root,
            engagement_id,
            snapshot_input_id,
        )
    except ledger.LedgerError as exc:
        raise ArchiveError("Archive snapshot input is invalid.") from exc
    try:
        snapshot = json.loads(Path(receipt["path"]).read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ArchiveError("Archive snapshot is unreadable.") from exc
    if (
        not isinstance(snapshot, dict)
        or snapshot.get("schema_version")
        not in {ledger.FOLDER_SNAPSHOT_SCHEMA, drive.DRIVE_SNAPSHOT_SCHEMA}
        or snapshot.get("client_id") != client_id
        or snapshot.get("engagement_id") != engagement_id
    ):
        raise ArchiveError("Archive snapshot identity is invalid.")
    content = {key: value for key, value in snapshot.items() if key != "content_sha256"}
    if snapshot.get("content_sha256") != _drive_bindings_digest(content):
        raise ArchiveError("Archive snapshot digest is stale.")
    if snapshot["schema_version"] == drive.DRIVE_SNAPSHOT_SCHEMA:
        binding = _google_drive_binding(client_id, state_dir)
        if (
            snapshot.get("root_folder_id") != binding["folder_id"]
            or snapshot.get("root_name") != binding["display_name"]
            or snapshot.get("drive_id") != binding["drive_id"]
        ):
            raise ArchiveError("Google Drive snapshot binding is stale.")
    return client_root, snapshot


def get_studio_archive_organization_inventory(
    client_id: str,
    engagement_id: str,
    snapshot_input_id: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Return the resumable full-population semantic projection for one snapshot."""

    private_state = _state_dir(state_dir)
    _, snapshot = _load_archive_organization_snapshot(
        client_id,
        engagement_id,
        snapshot_input_id,
        state_dir=private_state,
    )
    return {
        "status": "ready",
        "input_id": snapshot_input_id,
        "model_inventory": _project_archive_organization_inventory(snapshot),
        "source_archive_mutated": False,
        "google_drive_api_called": False,
    }


def _bounded_extraction_payload(extraction: ExtractionResult) -> dict[str, Any]:
    text_parts: list[str] = []
    locators: list[dict[str, str]] = []
    consumed = 0
    for chunk in extraction.chunks:
        remaining = MAX_OPEN_CHARS - consumed
        if remaining <= 0:
            break
        selected_text = chunk.text[:remaining]
        text_parts.append(selected_text)
        consumed += len(selected_text)
        locators.append(
            {
                "kind": chunk.locator_kind,
                "value": chunk.locator_value,
            }
        )
    limitations = list(extraction.limitations)
    if len(extraction.chunks) > len(locators):
        limitations.append("open_text_truncated")
    return {
        "status": extraction.status,
        "text": "\n\n".join(text_parts),
        "locators": locators,
        "extraction_method": extraction.extraction_method,
        "limitations": limitations,
    }


def open_studio_archive_organization_item(
    client_id: str,
    engagement_id: str,
    snapshot_input_id: str,
    item_ref: str,
    *,
    state_dir: Path | None = None,
    gateway: drive.DriveGateway | None = None,
) -> dict[str, Any]:
    """Open one snapshot item through an opaque model-facing reference."""

    if re.fullmatch(r"archive_item_[0-9a-f]{24}", item_ref) is None:
        raise ArchiveError("Archive organization item reference is invalid.")
    private_state = _state_dir(state_dir, create=True)
    client_root, snapshot = _load_archive_organization_snapshot(
        client_id,
        engagement_id,
        snapshot_input_id,
        state_dir=private_state,
    )
    snapshot_sha256 = str(snapshot["content_sha256"])
    files = snapshot.get("files")
    if not isinstance(files, list):
        raise ArchiveError("Archive snapshot files are invalid.")
    matches = [
        item
        for item in files
        if _archive_organization_item_ref(
            snapshot_sha256,
            str(item.get("relative_path") or ""),
        )
        == item_ref
    ]
    if len(matches) != 1:
        raise ArchiveError(
            "Archive organization item is not present exactly once in the snapshot."
        )
    expected = matches[0]
    storage_kind = (
        "google_drive"
        if snapshot["schema_version"] == drive.DRIVE_SNAPSHOT_SCHEMA
        else "local_filesystem"
    )
    relative_path = str(expected["relative_path"])
    display_path = (
        _project_drive_relative_path(snapshot_sha256, relative_path)
        if storage_kind == "google_drive"
        else relative_path
    )
    if storage_kind == "google_drive":
        opened = open_studio_google_drive_source(
            client_id,
            engagement_id,
            snapshot_input_id,
            str(expected["file_id"]),
            state_dir=private_state,
            gateway=gateway,
        )
        return {
            "status": opened["status"],
            "item_ref": item_ref,
            "storage_kind": storage_kind,
            "relative_path": display_path,
            "name": opened["name"],
            "mime_type": opened["mime_type"],
            "citation": f"archive-item:{item_ref} ({display_path})",
            "text": opened["text"],
            "locators": opened["locators"],
            "extraction_method": opened["extraction_method"],
            "evidence_mode": opened["evidence_mode"],
            "limitations": opened["limitations"],
            "source_identity_revalidated": True,
            "google_drive_api_called": True,
            "remote_archive_mutated": False,
            "temporary_content_deleted": True,
        }

    source = _resolve_source_file(client_root, relative_path)
    if _sha256_file(source) != str(expected["sha256"]):
        raise SourceChangedError(
            "Local archive source changed after the selected snapshot."
        )
    extraction = _extract_document(source, enable_ocr=False)
    bounded = _bounded_extraction_payload(extraction)
    return {
        **bounded,
        "item_ref": item_ref,
        "storage_kind": storage_kind,
        "relative_path": display_path,
        "name": source.name,
        "mime_type": None,
        "citation": f"archive-item:{item_ref} ({display_path})",
        "evidence_mode": "local_read",
        "source_identity_revalidated": True,
        "google_drive_api_called": False,
        "source_archive_mutated": False,
    }


def open_studio_google_drive_source(
    client_id: str,
    engagement_id: str,
    snapshot_input_id: str,
    file_id: str,
    *,
    state_dir: Path | None = None,
    gateway: drive.DriveGateway | None = None,
) -> dict[str, Any]:
    """Revalidate and extract one bounded file from an immutable Drive snapshot."""

    private_state = _state_dir(state_dir, create=True)
    client_root = _selected_ledger_root(
        client_id, engagement_id, state_dir=private_state
    )
    binding = _google_drive_binding(client_id, private_state)
    try:
        receipt = ledger.load_input_receipt(
            client_root,
            engagement_id,
            snapshot_input_id,
        )
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Google Drive snapshot input is invalid: {exc}") from exc
    try:
        snapshot = json.loads(Path(receipt["path"]).read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ArchiveError(f"Google Drive snapshot is unreadable: {exc}") from exc
    if (
        not isinstance(snapshot, dict)
        or snapshot.get("schema_version") != drive.DRIVE_SNAPSHOT_SCHEMA
        or snapshot.get("client_id") != client_id
        or snapshot.get("engagement_id") != engagement_id
        or snapshot.get("root_folder_id") != binding["folder_id"]
    ):
        raise ArchiveError("Google Drive snapshot identity is invalid.")
    content = {key: value for key, value in snapshot.items() if key != "content_sha256"}
    if snapshot.get("content_sha256") != _drive_bindings_digest(content):
        raise ArchiveError("Google Drive snapshot digest is stale.")
    files = snapshot.get("files")
    if not isinstance(files, list):
        raise ArchiveError("Google Drive snapshot files are invalid.")
    matches = [item for item in files if item.get("file_id") == file_id]
    if len(matches) != 1:
        raise ArchiveError("Drive file ID is not present exactly once in the snapshot.")
    expected = matches[0]
    selected_gateway = _drive_gateway(private_state, gateway)
    try:
        current = drive.normalize_file_metadata(
            selected_gateway.get_file(file_id), expected["parent_id"]
        )
    except drive.DriveError as exc:
        raise ArchiveError(str(exc)) from exc
    for key in (
        "file_id",
        "parent_id",
        "name",
        "mime_type",
        "version",
        "md5_checksum",
        "sha256_checksum",
        "drive_id",
    ):
        if current[key] != expected[key]:
            raise SourceChangedError(
                "Google Drive source changed after the selected snapshot."
            )
    if not current["capabilities"]["can_download"]:
        raise ArchiveError("Google Drive does not permit this file to be downloaded.")
    mime_type = current["mime_type"]
    export = GOOGLE_DRIVE_EXPORTS.get(mime_type)
    try:
        if export is not None:
            payload = selected_gateway.export_bytes(file_id, export[0])
            suffix = export[1]
            evidence_mode = "export"
        elif mime_type.startswith("application/vnd.google-apps."):
            raise ArchiveError(
                "This Google-native file type has no supported evidence export."
            )
        else:
            payload = selected_gateway.download_bytes(file_id)
            suffix = Path(current["name"]).suffix.lower()
            if suffix not in SUPPORTED_SUFFIXES:
                if mime_type.startswith("text/"):
                    suffix = ".txt"
                else:
                    raise ArchiveError(
                        "This Drive binary file type is not supported for extraction."
                    )
            evidence_mode = "download"
    except drive.DriveError as exc:
        raise ArchiveError(str(exc)) from exc
    if len(payload) > drive.MAX_EVIDENCE_BYTES:
        raise ArchiveError("Google Drive evidence exceeds the 100 MB read boundary.")
    try:
        current_after_read = drive.normalize_file_metadata(
            selected_gateway.get_file(file_id), expected["parent_id"]
        )
    except drive.DriveError as exc:
        raise ArchiveError(str(exc)) from exc
    for key in (
        "file_id",
        "parent_id",
        "name",
        "mime_type",
        "version",
        "md5_checksum",
        "sha256_checksum",
        "drive_id",
    ):
        if current_after_read[key] != expected[key]:
            raise SourceChangedError(
                "Google Drive source changed while the evidence was being read."
            )
    if evidence_mode == "download":
        if current["sha256_checksum"] is not None and (
            hashlib.sha256(payload).hexdigest() != current["sha256_checksum"]
        ):
            raise SourceChangedError(
                "Downloaded Google Drive evidence does not match its SHA-256."
            )
        if current["md5_checksum"] is not None and (
            hashlib.md5(payload, usedforsecurity=False).hexdigest()
            != current["md5_checksum"]
        ):
            raise SourceChangedError(
                "Downloaded Google Drive evidence does not match its MD5 checksum."
            )
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=".vera-drive-evidence-",
        suffix=suffix,
        dir=private_state,
    )
    temporary = Path(temporary_name)
    try:
        with os.fdopen(descriptor, "wb") as handle:
            handle.write(payload)
            handle.flush()
            os.fsync(handle.fileno())
        extraction = _extract_document(temporary, enable_ocr=False)
    finally:
        temporary.unlink(missing_ok=True)
    text_parts: list[str] = []
    locators: list[dict[str, str]] = []
    consumed = 0
    for chunk in extraction.chunks:
        remaining = MAX_OPEN_CHARS - consumed
        if remaining <= 0:
            break
        selected_text = chunk.text[:remaining]
        text_parts.append(selected_text)
        consumed += len(selected_text)
        locators.append(
            {
                "kind": chunk.locator_kind,
                "value": chunk.locator_value,
            }
        )
    limitations = list(extraction.limitations)
    if len(extraction.chunks) > len(locators):
        limitations.append("open_text_truncated")
    return {
        "status": extraction.status,
        "client_id": client_id,
        "engagement_id": engagement_id,
        "snapshot_input_id": snapshot_input_id,
        "file_id": file_id,
        "relative_path": expected["relative_path"],
        "name": current["name"],
        "mime_type": mime_type,
        "version": current["version"],
        "citation": (
            f"gdrive:{file_id}@v{current['version']} " f"({expected['relative_path']})"
        ),
        "text": "\n\n".join(text_parts),
        "locators": locators,
        "extraction_method": extraction.extraction_method,
        "evidence_mode": evidence_mode,
        "limitations": limitations,
        "google_drive_api_called": True,
        "remote_archive_mutated": False,
        "temporary_content_deleted": True,
    }


def _workflow_version(workflow_id: str) -> str:
    """Read the selected component version without inventing one in a run."""

    component_root = Path(__file__).resolve().parents[2] / workflow_id
    manifest_path = component_root / ".codex-plugin" / "plugin.json"
    if not manifest_path.is_file():
        return "unversioned"
    try:
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ArchiveError(f"Workflow manifest is unreadable: {exc}") from exc
    version = payload.get("version") if isinstance(payload, dict) else None
    if not isinstance(version, str) or not version.strip():
        raise ArchiveError("Workflow manifest has no valid version.")
    return version.strip()


def prepare_studio_client_workflow(
    engagement_id: str,
    workflow_id: str,
    *,
    input_ids: Sequence[str] = (),
    upstream_artifacts: Sequence[Mapping[str, Any]] = (),
    label: str | None = None,
    purpose: str | None = None,
    idempotency_key: str | None = None,
    new_run: bool = False,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Prepare an exact, recoverable, idempotent customer-folder run."""

    if workflow_id not in CLIENT_WORKFLOW_IDS:
        raise ArchiveError("Workflow is not supported by the client engagement gate.")
    private_state = _state_dir(state_dir)
    config = _load_config(private_state, validate_scope_roots=False)
    current, scopes_changed = _current_scope_view(config)
    if scopes_changed:
        raise ArchiveError(
            "Top-level archive scopes changed; refresh before preparing a run."
        )
    identities = _synchronize_client_identities(private_state, current)
    matches: list[tuple[ClientIdentity, Scope]] = []
    for identity in identities:
        scope = next(
            (item for item in current.scopes if item.scope_id == identity.scope_id),
            None,
        )
        if scope is None:
            continue
        client_root = _resolve_scope_root(current.archive_root, scope)
        try:
            ledger.load_engagement_manifest(client_root, engagement_id)
        except ledger.LedgerError:
            continue
        matches.append((identity, scope))
    if not matches:
        raise ArchiveError("Client engagement was not found.")
    if len(matches) != 1:
        raise ArchiveError(
            "Client engagement identity is ambiguous across customer folders."
        )
    identity, scope = matches[0]
    client_root = _resolve_scope_root(current.archive_root, scope)
    try:
        prepared = ledger.prepare_run(
            client_root,
            identity.client_id,
            engagement_id,
            workflow_id,
            _workflow_version(workflow_id),
            input_ids=input_ids,
            upstream_artifacts=upstream_artifacts,
            label=label,
            purpose=purpose,
            idempotency_key=idempotency_key,
            new_run=new_run,
        )
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Client workflow run is invalid: {exc}") from exc
    return {
        "status": prepared["status"],
        "client_id": identity.client_id,
        "engagement_id": engagement_id,
        "run": prepared["run"],
        "input_manifest": prepared["input_manifest"],
        "client_engagement": prepared["context"],
        "client_engagement_path": prepared["context_path"],
    }


def _journal_sampling_handoff_references(
    client_root: Path,
    engagement_id: str,
    sample_run_id: str,
) -> list[dict[str, str]]:
    """Resolve the exact closed artifacts for one mechanical workflow handoff."""

    try:
        loaded = ledger.load_run(client_root, engagement_id, sample_run_id)
        artifact_manifest = ledger.validate_run_artifacts(
            client_root,
            engagement_id,
            sample_run_id,
        )
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Journal Sampling run is unavailable: {exc}") from exc
    run = loaded["run"]
    if run["workflow_id"] != "journal-sampling":
        raise ArchiveError("Selected sample run is not a Journal Sampling run.")
    if run["status"] not in {"ready_for_review", "completed"}:
        raise ArchiveError(
            "Journal Sampling must be review-ready or completed before Check Entries."
        )
    artifact_by_path = {
        artifact["path"]: artifact for artifact in artifact_manifest["artifacts"]
    }
    expected_paths = {
        path for path, _artifact_id, _role in JOURNAL_SAMPLING_CHECK_ENTRIES_HANDOFF
    }
    missing_paths = sorted(expected_paths - set(artifact_by_path))
    if missing_paths:
        raise ArchiveError(
            "Journal Sampling has no complete Check Entries handoff; "
            f"missing={missing_paths}."
        )
    references: list[dict[str, str]] = []
    for path, required_artifact_id, role in JOURNAL_SAMPLING_CHECK_ENTRIES_HANDOFF:
        artifact = artifact_by_path[path]
        if (
            required_artifact_id is not None
            and artifact["artifact_id"] != required_artifact_id
        ):
            raise ArchiveError(
                "Journal Sampling handoff has the wrong semantic artifact identity "
                f"for {path}."
            )
        references.append(
            {
                "run_id": sample_run_id,
                "artifact_id": artifact["artifact_id"],
                "role": role,
            }
        )
    return references


def start_check_entries_from_sample(
    client_id: str,
    engagement_id: str,
    sample_run_id: str,
    *,
    support_input_ids: Sequence[str],
    label: str | None = None,
    purpose: str | None = None,
    idempotency_key: str | None = None,
    new_run: bool = False,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Prepare and start Check Entries from one exact sample and support batch."""

    private_state = _state_dir(state_dir)
    root = _selected_ledger_root(client_id, engagement_id, state_dir=private_state)
    if isinstance(support_input_ids, (str, bytes)) or not support_input_ids:
        raise ArchiveError("Select at least one support input for Check Entries.")
    normalized_support_ids = list(support_input_ids)
    if (
        not all(isinstance(input_id, str) for input_id in normalized_support_ids)
        or len(normalized_support_ids) > 10_000
        or len(set(normalized_support_ids)) != len(normalized_support_ids)
    ):
        raise ArchiveError("Check Entries support input selection is invalid.")
    try:
        support_receipts = [
            ledger.load_input_receipt(root, engagement_id, input_id)
            for input_id in normalized_support_ids
        ]
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Check Entries support input is invalid: {exc}") from exc
    if any(receipt["role"] != "support" for receipt in support_receipts):
        raise ArchiveError(
            "Check Entries accepts only inputs imported with the support role."
        )
    upstream_artifacts = _journal_sampling_handoff_references(
        root,
        engagement_id,
        sample_run_id,
    )
    try:
        prepared = ledger.prepare_run(
            root,
            client_id,
            engagement_id,
            "check-entries",
            _workflow_version("check-entries"),
            input_ids=normalized_support_ids,
            upstream_artifacts=upstream_artifacts,
            label=label or "Check sampled journal entries",
            purpose=(
                purpose
                or "Check one exact Journal Sampling sample against one support batch."
            ),
            idempotency_key=idempotency_key,
            new_run=new_run,
        )
        current_status = prepared["run"]["status"]
        if current_status in {"prepared", "failed"}:
            active = ledger.start_run(root, engagement_id, prepared["run"]["run_id"])
        elif current_status in {"running", "ready_for_review", "completed"}:
            active = prepared
        else:
            raise ArchiveError(
                "The existing Check Entries run is cancelled; request a new run."
            )
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Check Entries handoff could not start: {exc}") from exc
    return {
        "status": active["run"]["status"],
        "preparation_status": prepared["status"],
        "client_id": client_id,
        "engagement_id": engagement_id,
        "sample_run_id": sample_run_id,
        "support_input_ids": normalized_support_ids,
        "run": active["run"],
        "input_manifest": active["input_manifest"],
        "client_engagement": active["context"],
        "client_engagement_path": active["context_path"],
    }


def import_studio_client_document(
    client_id: str,
    source_path: Path,
    role: str,
    *,
    engagement_id: str | None = None,
    engagement_label: str | None = None,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Copy one authorized file into one explicit, immutable engagement input."""

    if role not in SUPPORTED_ENGAGEMENT_IMPORT_ROLES:
        raise ArchiveError("Import role must be journal, source, or support.")
    if engagement_id is None:
        raise ArchiveError(
            "Select or create an engagement before importing a document."
        )
    if engagement_label is not None:
        raise ArchiveError(
            "Engagement creation and document import are separate actions."
        )
    private_state = _state_dir(state_dir)
    folder = get_studio_client_folder(client_id, state_dir=private_state)[
        "client_folder"
    ]
    client_root = Path(folder["client_root"])
    try:
        engagement = ledger.load_engagement_manifest(client_root, engagement_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Selected client engagement is invalid: {exc}") from exc
    if engagement["client_id"] != client_id:
        raise ArchiveError("Selected engagement belongs to another client.")
    try:
        resolved_source = source_path.expanduser().resolve(strict=True)
    except OSError as exc:
        raise ArchiveError(f"Selected import source is unavailable: {exc}") from exc
    archive_root = Path(folder["archive_root"])
    if _path_is_within(resolved_source, archive_root) and not _path_is_within(
        resolved_source, client_root
    ):
        raise ArchiveError(
            "Selected import source belongs to another Studio Archive scope."
        )
    try:
        imported = ledger.import_document(
            client_root,
            client_id,
            engagement_id,
            source_path,
            role,
        )
        receipts = ledger.list_inputs(client_root, engagement_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Controlled document import failed: {exc}") from exc
    return {
        "status": imported["status"],
        "client_id": client_id,
        "engagement": {**engagement, "imports": list(receipts)},
        "input_receipt": imported["receipt"],
        "input_id": imported["receipt"]["input_id"],
        "imported_path": imported["imported_path"],
        "original_preserved": True,
        "source_archive_mutated": imported["source_archive_mutated"],
    }


def list_studio_client_engagements(
    client_id: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """List recoverable engagements, exact receipts, lifecycle, and artifacts."""

    private_state = _state_dir(state_dir)
    folder = get_studio_client_folder(client_id, state_dir=private_state)[
        "client_folder"
    ]
    client_root = Path(folder["client_root"])
    try:
        stored_engagements = ledger.list_engagements(client_root, client_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Customer-folder ledger is invalid: {exc}") from exc
    engagements: list[dict[str, Any]] = []
    for engagement in stored_engagements:
        engagement_id = engagement["engagement_id"]
        try:
            imports = list(ledger.list_inputs(client_root, engagement_id))
            stored_runs = ledger.list_runs(
                client_root, engagement_id, verify_inputs=False
            )
        except ledger.LedgerError as exc:
            raise ArchiveError(f"Customer-folder ledger is invalid: {exc}") from exc
        workflow_runs: list[dict[str, Any]] = []
        for loaded in stored_runs:
            run = loaded["run"]
            input_issue: str | None = None
            artifact_issue: str | None = None
            try:
                ledger.load_run(client_root, engagement_id, run["run_id"])
                inputs_valid = True
            except ledger.LedgerError as exc:
                inputs_valid = False
                input_issue = str(exc)
            artifacts = None
            artifacts_valid = False
            if run["status"] in {"ready_for_review", "completed"}:
                try:
                    artifacts = ledger.validate_run_artifacts(
                        client_root, engagement_id, run["run_id"]
                    )
                    artifacts_valid = True
                except ledger.LedgerError as exc:
                    artifact_issue = str(exc)
            available = (
                inputs_valid
                and artifacts_valid
                and run["status"] in {"ready_for_review", "completed"}
            )
            context = loaded["context"]
            output_dir = Path(loaded["output_dir"])
            record: dict[str, Any] = {
                "workflow_id": run["workflow_id"],
                "workflow_version": run["workflow_version"],
                "run_id": run["run_id"],
                "label": run["label"],
                "purpose": run["purpose"],
                "status": run["status"],
                "created_at": run["created_at"],
                "updated_at": run["updated_at"],
                "input_manifest": loaded["input_manifest"],
                "inputs_valid": inputs_valid,
                "input_issue": input_issue,
                "artifacts_available": available,
                "artifact_issue": artifact_issue,
                "artifact_manifest": artifacts,
                "run_output_dir": str(output_dir),
                "run_output_available": available,
                "client_engagement_path": loaded["context_path"],
                "client_engagement": context,
            }
            if run["workflow_id"] == "journal-sampling":
                record.update(
                    {
                        "normalized_journal_path": str(
                            output_dir / "normalization" / "normalized_journal.csv"
                        ),
                        "normalization_diagnostics_path": str(
                            output_dir
                            / "normalization"
                            / "normalization_diagnostics.json"
                        ),
                        "normalization_available": available
                        and (
                            output_dir / "normalization" / "normalized_journal.csv"
                        ).is_file()
                        and (
                            output_dir
                            / "normalization"
                            / "normalization_diagnostics.json"
                        ).is_file(),
                        "sample_output_dir": str(output_dir / "sample"),
                        "sample_available": available
                        and (output_dir / "sample" / "journal_sample.csv").is_file(),
                    }
                )
            elif run["workflow_id"] == "check-entries":
                record.update(
                    {
                        "checks_output_dir": str(output_dir / "checks"),
                        "checks_available": available
                        and (output_dir / "checks" / "check_audit.json").is_file(),
                    }
                )
            workflow_runs.append(record)
        engagements.append(
            {
                **engagement,
                "imports": imports,
                "workflow_run_count": len(workflow_runs),
                "workflow_runs": workflow_runs,
            }
        )
    return {
        "client_id": client_id,
        "engagement_count": len(engagements),
        "engagements": engagements,
    }


def _selected_ledger_root(
    client_id: str,
    engagement_id: str,
    *,
    state_dir: Path,
) -> Path:
    folder = get_studio_client_folder(client_id, state_dir=state_dir)["client_folder"]
    client_root = Path(folder["client_root"])
    try:
        engagement = ledger.load_engagement_manifest(client_root, engagement_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Selected client engagement is invalid: {exc}") from exc
    if engagement["client_id"] != client_id:
        raise ArchiveError("Selected engagement belongs to another client.")
    return client_root


def start_studio_client_workflow(
    client_id: str,
    engagement_id: str,
    run_id: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Mark one prepared run as running before executing helper scripts."""

    private_state = _state_dir(state_dir)
    root = _selected_ledger_root(client_id, engagement_id, state_dir=private_state)
    try:
        loaded = ledger.start_run(root, engagement_id, run_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Workflow run could not start: {exc}") from exc
    return {"status": loaded["run"]["status"], "run": loaded["run"]}


def fail_studio_client_workflow(
    client_id: str,
    engagement_id: str,
    run_id: str,
    reason: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Record a failed run while retaining its evidence and diagnostics."""

    private_state = _state_dir(state_dir)
    root = _selected_ledger_root(client_id, engagement_id, state_dir=private_state)
    try:
        loaded = ledger.fail_run(root, engagement_id, run_id, reason)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Workflow failure could not be recorded: {exc}") from exc
    return {"status": loaded["run"]["status"], "run": loaded["run"]}


def cancel_studio_client_workflow(
    client_id: str,
    engagement_id: str,
    run_id: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Cancel one abandoned run without deleting it."""

    private_state = _state_dir(state_dir)
    root = _selected_ledger_root(client_id, engagement_id, state_dir=private_state)
    try:
        loaded = ledger.cancel_run(root, engagement_id, run_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Workflow run could not be cancelled: {exc}") from exc
    return {"status": loaded["run"]["status"], "run": loaded["run"]}


def finalize_studio_client_workflow(
    client_id: str,
    engagement_id: str,
    run_id: str,
    artifacts: Sequence[Mapping[str, Any]],
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Declare the purpose of every output and seal its exact bytes."""

    private_state = _state_dir(state_dir)
    root = _selected_ledger_root(client_id, engagement_id, state_dir=private_state)
    try:
        loaded = ledger.finalize_run(root, engagement_id, run_id, artifacts)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Workflow artifacts could not be finalized: {exc}") from exc
    return {
        "status": loaded["run"]["status"],
        "run": loaded["run"],
        "artifact_manifest": loaded["artifact_manifest"],
    }


def complete_studio_client_workflow(
    client_id: str,
    engagement_id: str,
    run_id: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Complete one review-ready run whose artifacts still validate."""

    private_state = _state_dir(state_dir)
    root = _selected_ledger_root(client_id, engagement_id, state_dir=private_state)
    try:
        loaded = ledger.complete_run(root, engagement_id, run_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Workflow run could not be completed: {exc}") from exc
    return {"status": loaded["run"]["status"], "run": loaded["run"]}


def close_studio_client_engagement(
    client_id: str,
    engagement_id: str,
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Close one engagement after every active run is resolved."""

    private_state = _state_dir(state_dir)
    root = _selected_ledger_root(client_id, engagement_id, state_dir=private_state)
    try:
        engagement = ledger.close_engagement(root, engagement_id)
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Engagement could not be closed: {exc}") from exc
    return {"status": engagement["status"], "engagement": engagement}


def report_studio_client_retention(
    client_id: str,
    *,
    older_than_days: int | None = None,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Return a non-destructive retention inventory for professional review."""

    private_state = _state_dir(state_dir)
    folder = get_studio_client_folder(client_id, state_dir=private_state)[
        "client_folder"
    ]
    try:
        return ledger.retention_report(
            Path(folder["client_root"]), older_than_days=older_than_days
        )
    except ledger.LedgerError as exc:
        raise ArchiveError(f"Retention report could not be built: {exc}") from exc


def recover_studio_client_ledger(
    *,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Rebuild private client pointers and verify all portable ledger records."""

    private_state = _state_dir(state_dir)
    stored = _load_config(private_state, validate_scope_roots=False)
    current, scopes_changed = _current_scope_view(stored)
    if scopes_changed:
        raise ArchiveError("Refresh the archive before recovering customer folders.")
    identities = _synchronize_client_identities(private_state, current)
    engagement_count = 0
    input_count = 0
    run_count = 0
    for identity in identities:
        scope = next(
            (item for item in current.scopes if item.scope_id == identity.scope_id),
            None,
        )
        if scope is None:
            continue
        root = _resolve_scope_root(current.archive_root, scope)
        client_manifest = root / ledger.LEDGER_DIRECTORY / "client.json"
        if not client_manifest.is_file():
            continue
        try:
            engagements = ledger.list_engagements(root, identity.client_id)
            for engagement in engagements:
                engagement_count += 1
                input_count += len(
                    ledger.list_inputs(root, engagement["engagement_id"])
                )
                runs = ledger.list_runs(root, engagement["engagement_id"])
                for loaded in runs:
                    if loaded["run"]["status"] in {
                        "ready_for_review",
                        "completed",
                    }:
                        ledger.validate_run_artifacts(
                            root,
                            engagement["engagement_id"],
                            loaded["run"]["run_id"],
                        )
                run_count += len(runs)
        except ledger.LedgerError as exc:
            raise ArchiveError(f"Customer-folder recovery failed: {exc}") from exc
    return {
        "status": "recovered",
        "client_count": len(_discover_ledger_clients(current)),
        "engagement_count": engagement_count,
        "input_count": input_count,
        "run_count": run_count,
        "private_identity_values_recovered": False,
    }


def _gmail_safe_phrase(value: str) -> str:
    normalized = re.sub(r'["{}\\():\[\]]', " ", value)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    if not normalized:
        raise ArchiveError("Gmail search phrase contains no safe characters.")
    return f'"{normalized}"'


def _gmail_date(value: str | None, label: str) -> str | None:
    if value is None:
        return None
    if not isinstance(value, str) or re.fullmatch(r"\d{4}-\d{2}-\d{2}", value) is None:
        raise ArchiveError(f"{label} must use YYYY-MM-DD.")
    try:
        datetime.strptime(value, "%Y-%m-%d")
    except ValueError as exc:
        raise ArchiveError(f"{label} is not a valid calendar date.") from exc
    return value.replace("-", "/")


def _gmail_query_prefix(
    *,
    after: str | None,
    before: str | None,
) -> str:
    parts = ["in:anywhere", "-in:spam", "-in:trash"]
    normalized_after = _gmail_date(after, "after")
    normalized_before = _gmail_date(before, "before")
    if (
        normalized_after is not None
        and normalized_before is not None
        and normalized_after >= normalized_before
    ):
        raise ArchiveError("after must be earlier than before.")
    if normalized_after is not None:
        parts.append(f"after:{normalized_after}")
    if normalized_before is not None:
        parts.append(f"before:{normalized_before}")
    return " ".join(parts)


def _chunked(values: Sequence[str], size: int) -> Iterator[tuple[str, ...]]:
    for offset in range(0, len(values), size):
        yield tuple(values[offset : offset + size])


def plan_gmail_client_search(
    scope_id: str,
    *,
    topic: str | None = None,
    after: str | None = None,
    before: str | None = None,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Build bounded Gmail-native searches without calling the connector."""

    private_state = _state_dir(state_dir)
    config = _load_config(private_state, validate_scope_roots=False)
    _, scopes_changed = _current_scope_view(config)
    if scopes_changed:
        raise ArchiveError(
            "Top-level archive scopes changed; refresh before planning Gmail search."
        )
    if scope_id == "all":
        raise ArchiveError("Studio-wide Gmail search is not supported.")
    scope = next((item for item in config.scopes if item.scope_id == scope_id), None)
    if scope is None:
        raise ArchiveError("Gmail search scope is not configured.")
    records = _load_client_identities(private_state)
    record = next((item for item in records if item.scope_id == scope_id), None)
    if topic is not None:
        if not isinstance(topic, str) or not topic.strip():
            raise ArchiveError("Gmail search topic must be non-empty when supplied.")
        if len(topic) > MAX_GMAIL_TOPIC_CHARS:
            raise ArchiveError(
                "Gmail search topic must contain at most "
                f"{MAX_GMAIL_TOPIC_CHARS} characters."
            )
        topic_phrase = _gmail_safe_phrase(topic)
    else:
        topic_phrase = None
    prefix = _gmail_query_prefix(after=after, before=before)
    topic_suffix = "" if topic_phrase is None else f" {topic_phrase}"
    queries: list[dict[str, Any]] = []
    if record is not None:
        for query_index, addresses in enumerate(
            _chunked(record.email_addresses, MAX_GMAIL_QUERY_IDENTITIES),
            start=1,
        ):
            participant_terms = " ".join(
                term
                for address in addresses
                for term in (
                    f"from:{address}",
                    f"to:{address}",
                    f"cc:{address}",
                )
            )
            queries.append(
                {
                    "query_id": f"direct-{query_index}",
                    "kind": "confirmed_participant",
                    "query": f"{prefix} {{{participant_terms}}}{topic_suffix}",
                    "max_results": 20,
                    "routing_rule": "exact_unique_address_match_required",
                }
            )
        candidate_values = tuple(
            dict.fromkeys(
                (
                    scope.display_name,
                    *record.legal_names,
                    *record.tax_identifiers,
                )
            )
        )
    else:
        candidate_values = (scope.display_name,)
    for query_index, identities in enumerate(
        _chunked(candidate_values, MAX_GMAIL_QUERY_IDENTITIES),
        start=1,
    ):
        identity_terms = " ".join(_gmail_safe_phrase(value) for value in identities)
        queries.append(
            {
                "query_id": f"candidate-{query_index}",
                "kind": "identity_candidate",
                "query": f"{prefix} {{{identity_terms}}}{topic_suffix}",
                "max_results": 20,
                "routing_rule": "message_read_and_semantic_review_required",
            }
        )
    if record is None:
        profile_status = "alias_only"
    elif record.email_addresses:
        profile_status = "configured"
    else:
        profile_status = "candidate_only"
    return {
        "connector": "gmail",
        "scope_id": scope.scope_id,
        "display_name": scope.display_name,
        "profile_status": profile_status,
        "queries": queries,
        "requires_connector_profile_check": True,
        "requires_message_read_before_use": True,
        "gmail_connector_called": False,
        "warnings": (
            []
            if record is not None and record.email_addresses
            else [
                "No confirmed participant address is configured. Candidate "
                "results must be reviewed and an address confirmed before "
                "automatic client routing."
            ]
        ),
    }


def match_studio_email_client(
    header_addresses: Sequence[str],
    *,
    headers_complete: bool = False,
    expected_scope_id: str | None = None,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Match Gmail headers only by unique, confirmed full email addresses."""

    if (
        isinstance(header_addresses, (str, bytes))
        or not header_addresses
        or len(header_addresses) > 100
    ):
        raise ArchiveError("Provide between 1 and 100 Gmail header address values.")
    if not isinstance(headers_complete, bool):
        raise ArchiveError("headers_complete must be a boolean.")
    private_state = _state_dir(state_dir)
    config = _load_config(private_state, validate_scope_roots=False)
    _, scopes_changed = _current_scope_view(config)
    if scopes_changed:
        raise ArchiveError(
            "Top-level archive scopes changed; refresh before matching Gmail."
        )
    scopes_by_id = {scope.scope_id: scope for scope in config.scopes}
    if expected_scope_id is not None and expected_scope_id not in scopes_by_id:
        raise ArchiveError("Expected Gmail client scope is not configured.")
    parsed_addresses: set[str] = set()
    unparsed_headers: list[str] = []
    for raw_header in header_addresses:
        if not isinstance(raw_header, str) or len(raw_header) > 2_000:
            raise ArchiveError("Gmail header address values must be bounded strings.")
        parsed = getaddresses([raw_header])
        accepted = False
        parsed_completely = bool(parsed)
        for _, address in parsed:
            if not address:
                parsed_completely = False
                continue
            try:
                parsed_addresses.add(_normalize_email_address(address))
            except ArchiveError:
                parsed_completely = False
                continue
            accepted = True
        if not accepted or not parsed_completely:
            unparsed_headers.append(raw_header)
    records = _load_client_identities(private_state)
    owners = {
        email_address: record.scope_id
        for record in records
        for email_address in record.email_addresses
    }
    matched: dict[str, list[str]] = {}
    for address in sorted(parsed_addresses):
        owner = owners.get(address)
        if owner is not None:
            matched.setdefault(owner, []).append(address)
    candidate_scope_ids = sorted(matched)
    header_coverage_complete = headers_complete and not unparsed_headers
    if len(candidate_scope_ids) > 1:
        routing_status = "ambiguous"
        matched_scope_id = None
    elif not header_coverage_complete:
        routing_status = "incomplete"
        matched_scope_id = None
    elif len(candidate_scope_ids) == 1:
        routing_status = "exact"
        matched_scope_id: str | None = candidate_scope_ids[0]
    else:
        routing_status = "unassigned"
        matched_scope_id = None
    belongs_to_expected_scope: bool | None
    if expected_scope_id is None or matched_scope_id is None:
        belongs_to_expected_scope = None
    else:
        belongs_to_expected_scope = matched_scope_id == expected_scope_id
    return {
        "routing_status": routing_status,
        "matched_scope_id": matched_scope_id,
        "candidate_scope_ids": candidate_scope_ids,
        "matches": [
            {
                "scope_id": scope_id,
                "display_name": scopes_by_id[scope_id].display_name,
                "email_addresses": matched[scope_id],
                "match_method": "exact_email_address",
            }
            for scope_id in candidate_scope_ids
        ],
        "belongs_to_expected_scope": belongs_to_expected_scope,
        "may_use_in_scoped_answer": belongs_to_expected_scope is True,
        "requires_semantic_review": routing_status != "exact",
        "parsed_email_addresses": sorted(parsed_addresses),
        "unparsed_header_count": len(unparsed_headers),
        "header_coverage_complete": header_coverage_complete,
        "gmail_connector_called": False,
        "gmail_data_persisted": False,
    }


def _config_fingerprint(config: ArchiveConfig) -> str:
    payload = json.dumps(
        {
            "schema_version": CONFIG_SCHEMA_VERSION,
            "archive_root": str(config.archive_root),
            "scopes": [scope.as_json() for scope in config.scopes],
        },
        sort_keys=True,
        separators=(",", ":"),
    )
    return _sha256_bytes(payload.encode("utf-8"))


def _connect(
    state_dir: Path,
    *,
    readonly: bool = False,
) -> sqlite3.Connection:
    path = _database_path(state_dir)
    if readonly:
        if not path.is_file():
            raise ArchiveError("Studio Archive index is missing; refresh it first.")
        _assert_private_file(path, "index")
        connection = sqlite3.connect(f"{path.as_uri()}?mode=ro", uri=True)
    else:
        connection = sqlite3.connect(path)
    connection.row_factory = sqlite3.Row
    connection.execute("PRAGMA foreign_keys = ON")
    if readonly:
        connection.execute("PRAGMA query_only = ON")
        try:
            schema_row = connection.execute(
                "SELECT value FROM metadata WHERE key = 'schema_version'"
            ).fetchone()
        except sqlite3.OperationalError as exc:
            connection.close()
            raise ArchiveError(
                "Studio Archive index is invalid; rebuild it before searching."
            ) from exc
        if schema_row is None or schema_row["value"] != SCHEMA_VERSION:
            connection.close()
            raise ArchiveError("Studio Archive database schema is unsupported.")
        return connection
    connection.execute("PRAGMA secure_delete = ON")
    connection.execute("PRAGMA journal_mode = DELETE")
    connection.executescript("""
        CREATE TABLE IF NOT EXISTS metadata (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS documents (
            document_id TEXT PRIMARY KEY,
            scope_id TEXT NOT NULL,
            relative_path TEXT NOT NULL UNIQUE,
            extension TEXT NOT NULL,
            size_bytes INTEGER NOT NULL,
            mtime_ns INTEGER NOT NULL,
            sha256 TEXT NOT NULL,
            extraction_method TEXT NOT NULL,
            status TEXT NOT NULL,
            needs_ocr INTEGER NOT NULL,
            limitations_json TEXT NOT NULL,
            indexed_at TEXT NOT NULL,
            last_seen_generation INTEGER NOT NULL
        );
        CREATE TABLE IF NOT EXISTS chunks (
            source_id TEXT PRIMARY KEY,
            document_id TEXT NOT NULL REFERENCES documents(document_id)
                ON DELETE CASCADE,
            ordinal INTEGER NOT NULL,
            locator_kind TEXT NOT NULL,
            locator_value TEXT NOT NULL,
            text TEXT NOT NULL,
            text_sha256 TEXT NOT NULL,
            UNIQUE(document_id, ordinal)
        );
        CREATE TABLE IF NOT EXISTS scan_issues (
            relative_path TEXT PRIMARY KEY,
            scope_id TEXT NOT NULL,
            reason TEXT NOT NULL,
            size_bytes INTEGER
        );
        CREATE VIRTUAL TABLE IF NOT EXISTS chunk_fts USING fts5(
            source_id UNINDEXED,
            text,
            tokenize = 'unicode61 remove_diacritics 2'
        );
        CREATE INDEX IF NOT EXISTS documents_scope_idx
            ON documents(scope_id, relative_path);
        CREATE INDEX IF NOT EXISTS chunks_document_idx
            ON chunks(document_id, ordinal);
        """)
    try:
        connection.execute(
            "INSERT INTO chunk_fts(chunk_fts, rank) VALUES('secure-delete', 1)"
        )
    except sqlite3.OperationalError as exc:
        connection.close()
        raise ArchiveError(
            "The active SQLite FTS5 runtime does not support secure deletion."
        ) from exc
    connection.execute(
        "INSERT OR IGNORE INTO metadata(key, value) VALUES('schema_version', ?)",
        (SCHEMA_VERSION,),
    )
    schema_row = connection.execute(
        "SELECT value FROM metadata WHERE key = 'schema_version'"
    ).fetchone()
    if schema_row is None or schema_row["value"] != SCHEMA_VERSION:
        connection.close()
        raise ArchiveError("Studio Archive database schema is unsupported.")
    connection.commit()
    try:
        path.chmod(0o600)
    except OSError:
        connection.close()
        raise
    return connection


def _metadata_get(connection: sqlite3.Connection, key: str) -> str | None:
    row = connection.execute(
        "SELECT value FROM metadata WHERE key = ?",
        (key,),
    ).fetchone()
    return None if row is None else str(row["value"])


def _metadata_set(connection: sqlite3.Connection, key: str, value: str) -> None:
    connection.execute(
        """
        INSERT INTO metadata(key, value) VALUES(?, ?)
        ON CONFLICT(key) DO UPDATE SET value = excluded.value
        """,
        (key, value),
    )


def _clear_index(connection: sqlite3.Connection) -> None:
    connection.execute("DELETE FROM chunk_fts")
    connection.execute("DELETE FROM chunks")
    connection.execute("DELETE FROM documents")
    connection.execute("DELETE FROM scan_issues")


def _replace_scan_issues(
    connection: sqlite3.Connection,
    issues: Sequence[ScanIssue],
) -> None:
    connection.execute("DELETE FROM scan_issues")
    connection.executemany(
        """
        INSERT INTO scan_issues(relative_path, scope_id, reason, size_bytes)
        VALUES (?, ?, ?, ?)
        """,
        (
            (
                issue.relative_path,
                issue.scope_id,
                issue.reason,
                issue.size_bytes,
            )
            for issue in issues
        ),
    )


def _scan_issue_status(
    connection: sqlite3.Connection,
    *,
    scope_id: str | None = None,
) -> dict[str, Any]:
    issue_count = int(
        connection.execute(
            """
            SELECT COUNT(*) AS count
            FROM scan_issues
            WHERE ? IS NULL OR scope_id = ?
            """,
            (scope_id, scope_id),
        ).fetchone()["count"]
    )
    rows = connection.execute(
        """
        SELECT scope_id, relative_path, reason, size_bytes
        FROM scan_issues
        WHERE ? IS NULL OR scope_id = ?
        ORDER BY relative_path
        LIMIT ?
        """,
        (scope_id, scope_id, MAX_STATUS_SCAN_ISSUES),
    ).fetchall()
    return {
        "scan_issue_count": issue_count,
        "scan_issues": [
            {
                "scope_id": str(row["scope_id"]),
                "relative_path": str(row["relative_path"]),
                "reason": str(row["reason"]),
                "size_bytes": (
                    None if row["size_bytes"] is None else int(row["size_bytes"])
                ),
            }
            for row in rows
        ],
        "scan_issues_truncated": issue_count > len(rows),
    }


def _document_issue_status(
    connection: sqlite3.Connection,
    *,
    scope_id: str | None = None,
) -> dict[str, Any]:
    """Return a bounded inventory of indexed documents with evidence limits."""

    issue_count = int(
        connection.execute(
            """
            SELECT COUNT(*) AS count
            FROM (
                SELECT d.document_id
                FROM documents AS d
                LEFT JOIN chunks AS c ON c.document_id = d.document_id
                WHERE ? IS NULL OR d.scope_id = ?
                GROUP BY d.document_id
                HAVING d.status != 'indexed'
                    OR d.needs_ocr = 1
                    OR d.limitations_json != '[]'
                    OR COUNT(c.source_id) = 0
            )
            """,
            (scope_id, scope_id),
        ).fetchone()["count"]
    )
    rows = connection.execute(
        """
        SELECT
            d.scope_id,
            d.relative_path,
            d.status,
            d.needs_ocr,
            d.limitations_json,
            COUNT(c.source_id) AS chunk_count
        FROM documents AS d
        LEFT JOIN chunks AS c ON c.document_id = d.document_id
        WHERE ? IS NULL OR d.scope_id = ?
        GROUP BY
            d.document_id,
            d.scope_id,
            d.relative_path,
            d.status,
            d.needs_ocr,
            d.limitations_json
        HAVING d.status != 'indexed'
            OR d.needs_ocr = 1
            OR d.limitations_json != '[]'
            OR COUNT(c.source_id) = 0
        ORDER BY d.relative_path
        LIMIT ?
        """,
        (scope_id, scope_id, MAX_STATUS_DOCUMENT_ISSUES),
    ).fetchall()
    return {
        "document_issue_count": issue_count,
        "document_issues": [
            {
                "scope_id": str(row["scope_id"]),
                "relative_path": str(row["relative_path"]),
                "document_status": str(row["status"]),
                "needs_ocr": bool(row["needs_ocr"]),
                "limitations": _decode_limitations(str(row["limitations_json"])),
                "chunk_count": int(row["chunk_count"]),
            }
            for row in rows
        ],
        "document_issues_truncated": issue_count > len(rows),
    }


def _resolve_scope_root(root: Path, scope: Scope) -> Path:
    if scope.relative_dir == ".":
        return root
    relative = Path(scope.relative_dir)
    if (
        relative.is_absolute()
        or scope.relative_dir != relative.as_posix()
        or any(part in {"", ".", ".."} for part in relative.parts)
    ):
        raise ArchiveError(
            "Configured archive scope is not a normalized relative path."
        )
    candidate = root
    for part in relative.parts:
        candidate /= part
        if candidate.is_symlink():
            raise ArchiveError("Configured archive scope contains a symbolic link.")
    try:
        resolved = candidate.resolve(strict=True)
    except OSError as exc:
        raise ArchiveError(f"Configured archive scope is unavailable: {exc}") from exc
    if not _path_is_within(resolved, root) or not resolved.is_dir():
        raise ArchiveError("Configured archive scope escapes the archive root.")
    return resolved


def _walk_archive_entries(directory: Path) -> Iterator[tuple[Path, str | None]]:
    try:
        entries = sorted(
            os.scandir(directory),
            key=lambda entry: entry.name.casefold(),
        )
    except OSError as exc:
        raise ArchiveError(
            f"Archive enumeration failed at {directory.name}: {exc}"
        ) from exc
    for entry in entries:
        if entry.name in IGNORED_NAMES:
            continue
        try:
            if entry.is_symlink():
                yield Path(entry.path), "symbolic_link_not_followed"
            elif entry.is_dir(follow_symlinks=False):
                yield from _walk_archive_entries(Path(entry.path))
            elif entry.is_file(follow_symlinks=False):
                yield Path(entry.path), None
        except OSError as exc:
            raise ArchiveError(
                f"Archive enumeration failed at {entry.name}: {exc}"
            ) from exc


def _scope_entries(
    scope_root: Path,
    scope: Scope,
) -> Iterator[tuple[Path, str | None]]:
    """Yield source files for one non-overlapping configured scope."""

    if scope.relative_dir != ".":
        yield from _walk_archive_entries(scope_root)
        return
    try:
        entries = sorted(scope_root.iterdir(), key=lambda path: path.name.casefold())
    except OSError as exc:
        raise ArchiveError(f"Archive enumeration failed at root: {exc}") from exc
    for path in entries:
        if path.name in IGNORED_NAMES:
            continue
        if path.is_symlink():
            yield path, "symbolic_link_not_followed"
            continue
        try:
            if stat.S_ISREG(path.lstat().st_mode):
                yield path, None
        except OSError as exc:
            raise ArchiveError(
                f"Archive enumeration failed at {path.name}: {exc}"
            ) from exc


def _excluded_ledger_search_path(path: Path, scope_root: Path) -> bool:
    """Exclude technical manifests/runs while retaining canonical input evidence."""

    try:
        parts = path.relative_to(scope_root).parts
    except ValueError:
        return False
    if not parts or parts[0] != ledger.LEDGER_DIRECTORY:
        return False
    if (
        len(parts) >= 6
        and parts[1] == "engagements"
        and re.fullmatch(r"eng_[0-9a-f]{24}", parts[2]) is not None
        and parts[3] == "inputs"
        and re.fullmatch(r"input_[0-9a-f]{24}", parts[4]) is not None
    ):
        return parts[5] == "receipt.json" or len(parts) != 6
    return True


def _discover_files(
    config: ArchiveConfig,
) -> tuple[list[DiscoveredFile], list[ScanIssue]]:
    discovered: list[DiscoveredFile] = []
    issues: list[ScanIssue] = []
    total_bytes = 0
    scanned_files = 0
    seen_paths: set[str] = set()
    seen_casefolded_paths: set[str] = set()
    for scope in config.scopes:
        scope_root = _resolve_scope_root(config.archive_root, scope)
        for path, skip_reason in _scope_entries(scope_root, scope):
            if _excluded_ledger_search_path(path, scope_root):
                continue
            relative_path = path.relative_to(config.archive_root).as_posix()
            if relative_path in seen_paths:
                raise ArchiveError(
                    f"Configured scopes overlap at archive path: {relative_path}"
                )
            seen_paths.add(relative_path)
            scanned_files += 1
            if scanned_files > MAX_FILES:
                raise ArchiveError(
                    f"Archive exceeds the {MAX_FILES:,}-file first-version limit."
                )
            if skip_reason is not None:
                issues.append(
                    ScanIssue(
                        scope_id=scope.scope_id,
                        relative_path=relative_path,
                        reason=skip_reason,
                        size_bytes=None,
                    )
                )
                continue
            try:
                metadata = path.stat(follow_symlinks=False)
            except OSError:
                issues.append(
                    ScanIssue(
                        scope_id=scope.scope_id,
                        relative_path=relative_path,
                        reason="source_metadata_unavailable",
                        size_bytes=None,
                    )
                )
                continue
            if not stat.S_ISREG(metadata.st_mode):
                continue
            suffix = path.suffix.lower()
            if suffix not in SUPPORTED_SUFFIXES:
                issues.append(
                    ScanIssue(
                        scope_id=scope.scope_id,
                        relative_path=relative_path,
                        reason="unsupported_extension",
                        size_bytes=metadata.st_size,
                    )
                )
                continue
            casefolded_path = relative_path.casefold()
            if casefolded_path in seen_casefolded_paths:
                raise ArchiveError(
                    "Archive contains paths that collide when case is ignored: "
                    f"{relative_path}"
                )
            seen_casefolded_paths.add(casefolded_path)
            if metadata.st_size > MAX_FILE_BYTES:
                issues.append(
                    ScanIssue(
                        scope_id=scope.scope_id,
                        relative_path=relative_path,
                        reason="file_size_limit_exceeded",
                        size_bytes=metadata.st_size,
                    )
                )
                continue
            total_bytes += metadata.st_size
            if total_bytes > MAX_TOTAL_BYTES:
                raise ArchiveError(
                    "Archive exceeds the first-version total byte limit."
                )
            discovered.append(
                DiscoveredFile(
                    scope_id=scope.scope_id,
                    relative_path=relative_path,
                    path=path,
                    size_bytes=metadata.st_size,
                    mtime_ns=metadata.st_mtime_ns,
                )
            )
    return discovered, issues


def _resolve_source_file(root: Path, relative_path: str) -> Path:
    """Resolve one indexed regular file without following symbolic links."""

    relative = Path(relative_path)
    if (
        relative.is_absolute()
        or relative_path != relative.as_posix()
        or any(part in {"", ".", ".."} for part in relative.parts)
    ):
        raise ArchiveError("Indexed source path is not normalized and relative.")
    candidate = root
    for part in relative.parts:
        candidate /= part
        if candidate.is_symlink():
            raise ArchiveError("Indexed source path contains a symbolic link.")
    try:
        resolved = candidate.resolve(strict=True)
    except OSError as exc:
        raise SourceChangedError(f"Indexed source is unavailable: {exc}") from exc
    if not _path_is_within(resolved, root):
        raise ArchiveError("Indexed source path escapes the archive root.")
    if not stat.S_ISREG(resolved.lstat().st_mode):
        raise ArchiveError("Indexed source is not a regular file.")
    return resolved


def _normalize_text(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    lines = [
        re.sub(r"[ \t]+", " ", line).strip()
        for line in text.replace("\r", "\n").split("\n")
    ]
    return "\n".join(line for line in lines if line).strip()


def _text_is_useful(text: str) -> bool:
    compact = _normalize_text(text)
    if len(compact) < 40:
        return False
    alpha_count = sum(character.isalpha() for character in compact)
    return alpha_count >= max(20, len(compact) // 10)


def _chunks_from_numbered_lines(
    rows: Sequence[tuple[int, str]],
    *,
    locator_kind: str,
    locator_prefix: str = "",
) -> tuple[ExtractedChunk, ...]:
    chunks: list[ExtractedChunk] = []
    current: list[tuple[int, str]] = []
    current_chars = 0

    def flush() -> None:
        nonlocal current, current_chars
        if not current:
            return
        start = current[0][0]
        end = current[-1][0]
        range_value = str(start) if start == end else f"{start}-{end}"
        locator_value = f"{locator_prefix}{range_value}"
        chunks.append(
            ExtractedChunk(
                ordinal=len(chunks),
                locator_kind=locator_kind,
                locator_value=locator_value,
                text="\n".join(value for _, value in current),
            )
        )
        current = []
        current_chars = 0

    for number, raw_text in rows:
        text = _normalize_text(raw_text)
        if not text:
            continue
        if current and (
            len(current) >= MAX_CHUNK_LINES
            or current_chars + len(text) + 1 > MAX_CHUNK_CHARS
        ):
            flush()
        if len(text) > MAX_CHUNK_CHARS:
            flush()
            for offset in range(0, len(text), MAX_CHUNK_CHARS):
                piece = text[offset : offset + MAX_CHUNK_CHARS]
                chunks.append(
                    ExtractedChunk(
                        ordinal=len(chunks),
                        locator_kind=locator_kind,
                        locator_value=f"{locator_prefix}{number}",
                        text=piece,
                    )
                )
            continue
        current.append((number, text))
        current_chars += len(text) + 1
    flush()
    return tuple(chunks)


def _extract_plain_text(path: Path) -> ExtractionResult:
    if path.stat().st_size > MAX_TEXT_BYTES:
        return ExtractionResult((), "text", "error", False, ("text_file_too_large",))
    payload = path.read_bytes()
    text = payload.decode("utf-8-sig", errors="replace")
    rows = tuple(enumerate(text.splitlines(), start=1))
    chunks = _chunks_from_numbered_lines(rows, locator_kind="lines")
    status = "indexed" if chunks else "partial"
    limitations = () if chunks else ("no_extractable_text",)
    return ExtractionResult(chunks, "plain_text", status, False, limitations)


def _validate_ooxml_archive(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        members = [info for info in archive.infolist() if not info.is_dir()]
        if len(members) > MAX_OOXML_MEMBERS:
            raise ArchiveError("OOXML archive contains too many members.")
        names = [info.filename for info in members]
        if len(names) != len(set(names)):
            raise ArchiveError("OOXML archive contains duplicate member names.")
        if sum(info.file_size for info in members) > MAX_OOXML_TOTAL_BYTES:
            raise ArchiveError("Expanded OOXML archive exceeds the size limit.")
        for info in members:
            if info.flag_bits & 0x1:
                raise ArchiveError("Encrypted OOXML members are unsupported.")
            ratio = info.file_size / max(info.compress_size, 1)
            if info.file_size > MAX_OOXML_MEMBER_BYTES:
                raise ArchiveError("OOXML member exceeds the size limit.")
            if ratio > MAX_OOXML_COMPRESSION_RATIO:
                raise ArchiveError("OOXML member compression ratio is unsafe.")


def _extract_docx(path: Path) -> ExtractionResult:
    _validate_ooxml_archive(path)
    from docx import Document

    document = Document(path)
    chunks: list[ExtractedChunk] = list(
        _chunks_from_numbered_lines(
            tuple(
                (index, paragraph.text)
                for index, paragraph in enumerate(document.paragraphs, start=1)
            ),
            locator_kind="paragraphs",
        )
    )
    for table_number, table in enumerate(document.tables, start=1):
        rows = [
            "\t".join(_normalize_text(cell.text) for cell in row.cells)
            for row in table.rows
        ]
        for chunk in _chunks_from_numbered_lines(
            tuple(enumerate(rows, start=1)),
            locator_kind="table",
            locator_prefix=f"{table_number}, rows ",
        ):
            chunks.append(
                ExtractedChunk(
                    ordinal=len(chunks),
                    locator_kind=chunk.locator_kind,
                    locator_value=chunk.locator_value,
                    text=chunk.text,
                )
            )
    status = "indexed" if chunks else "partial"
    return ExtractionResult(
        tuple(chunks),
        "docx",
        status,
        False,
        () if chunks else ("no_extractable_text",),
    )


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    return _normalize_text(str(value))


def _extract_xlsx(path: Path) -> ExtractionResult:
    _validate_ooxml_archive(path)
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter

    workbook = load_workbook(path, read_only=True, data_only=False)
    chunks: list[ExtractedChunk] = []
    limitations: list[str] = []
    try:
        for sheet_index, worksheet in enumerate(workbook.worksheets, start=1):
            if sheet_index > MAX_WORKBOOK_SHEETS:
                limitations.append("workbook_sheet_limit_reached")
                break
            rows: list[tuple[int, str]] = []
            if int(worksheet.max_column or 0) > MAX_WORKSHEET_COLUMNS:
                limitations.append(f"worksheet_column_limit_reached:{worksheet.title}")
            for row_number, row in enumerate(
                worksheet.iter_rows(max_col=MAX_WORKSHEET_COLUMNS),
                start=1,
            ):
                if row_number > MAX_WORKSHEET_ROWS:
                    limitations.append(f"worksheet_row_limit_reached:{worksheet.title}")
                    break
                values = []
                for column_number, cell in enumerate(row, start=1):
                    value = _cell_text(cell.value)
                    if value:
                        values.append(
                            f"{get_column_letter(column_number)}{row_number}={value}"
                        )
                if values:
                    rows.append((row_number, "\t".join(values)))
            for chunk in _chunks_from_numbered_lines(
                rows,
                locator_kind="sheet",
                locator_prefix=f"{worksheet.title}!rows ",
            ):
                chunks.append(
                    ExtractedChunk(
                        ordinal=len(chunks),
                        locator_kind=chunk.locator_kind,
                        locator_value=chunk.locator_value,
                        text=chunk.text,
                    )
                )
    finally:
        workbook.close()
    status = "indexed" if chunks and not limitations else "partial"
    if not chunks:
        limitations.append("no_extractable_text")
    return ExtractionResult(
        tuple(chunks),
        "xlsx",
        status,
        False,
        tuple(dict.fromkeys(limitations)),
    )


def _html_to_text(value: str) -> str:
    parser = _HtmlTextExtractor()
    parser.feed(value)
    parser.close()
    return " ".join(parser.parts)


def _extract_eml(path: Path) -> ExtractionResult:
    if path.stat().st_size > MAX_EMAIL_BYTES:
        return ExtractionResult((), "eml", "error", False, ("email_too_large",))
    message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    lines = [
        f"Subject: {message.get('subject', '')}",
        f"From: {message.get('from', '')}",
        f"To: {message.get('to', '')}",
        f"Date: {message.get('date', '')}",
    ]
    attachment_count = 0
    for part in message.walk():
        if part.is_multipart():
            continue
        if part.get_content_disposition() == "attachment" or part.get_filename():
            attachment_count += 1
            continue
        if part.get_content_type() not in {"text/plain", "text/html"}:
            continue
        try:
            content = part.get_content()
        except (LookupError, UnicodeError):
            content = (part.get_payload(decode=True) or b"").decode(
                "utf-8",
                errors="replace",
            )
        text = content if isinstance(content, str) else str(content)
        if part.get_content_type() == "text/html":
            text = _html_to_text(text)
        lines.extend(text.splitlines())
    chunks = _chunks_from_numbered_lines(
        tuple(enumerate(lines, start=1)),
        locator_kind="message lines",
    )
    limitations = (
        (f"attachments_not_indexed:{attachment_count}",) if attachment_count else ()
    )
    status = "indexed" if chunks and not limitations else "partial"
    if not chunks:
        limitations = (*limitations, "no_extractable_text")
    return ExtractionResult(chunks, "eml", status, False, limitations)


def _ensure_vendor_import_path() -> None:
    component_root = Path(__file__).resolve().parents[1]
    candidates = (
        component_root / "vendor" / "modules",
        component_root.parent / "_shared" / "vendor" / "modules",
    )
    for candidate in candidates:
        if candidate.is_dir() and str(candidate) not in sys.path:
            sys.path.insert(0, str(candidate))


def _run_local_ocr(image_bytes: bytes) -> tuple[str, tuple[str, ...], bool]:
    _ensure_vendor_import_path()
    try:
        from vera_ocr import extract_text_from_image_bytes
    except (ImportError, ModuleNotFoundError):
        return "", ("ocr_runtime_unavailable",), False
    result = extract_text_from_image_bytes(
        image_bytes,
        language="it",
        allow_model_download=False,
    )
    if result.network_used:
        raise ArchiveError("Local OCR unexpectedly reported network use.")
    warnings = tuple(str(value) for value in result.warnings)
    if result.status != "ok":
        return "", (f"ocr_{result.status}", *warnings), False
    return _normalize_text(result.text), warnings, True


def _render_pdf_page(path: Path, page_index: int) -> bytes:
    import fitz

    with fitz.open(path) as document:
        page = document.load_page(page_index)
        scale = 200.0 / 72.0
        rendered_pixels = int(page.rect.width * scale) * int(page.rect.height * scale)
        if rendered_pixels > MAX_IMAGE_TOTAL_PIXELS:
            raise ArchiveError("Rendered PDF page exceeds the OCR pixel limit.")
        pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
        return bytes(pixmap.tobytes("png"))


def _page_chunks(
    text: str,
    *,
    page_number: int,
    ordinal_start: int,
) -> tuple[ExtractedChunk, ...]:
    normalized = _normalize_text(text)
    if not normalized:
        return ()
    pieces = [
        normalized[offset : offset + MAX_CHUNK_CHARS]
        for offset in range(0, len(normalized), MAX_CHUNK_CHARS)
    ]
    return tuple(
        ExtractedChunk(
            ordinal=ordinal_start + index,
            locator_kind="page",
            locator_value=str(page_number),
            text=piece,
        )
        for index, piece in enumerate(pieces)
    )


def _extract_pdf(path: Path, *, enable_ocr: bool) -> ExtractionResult:
    if path.stat().st_size > MAX_PDF_BYTES:
        return ExtractionResult((), "pdf", "error", False, ("pdf_too_large",))
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted and not reader.decrypt(""):
            return ExtractionResult(
                (),
                "pdf",
                "partial",
                False,
                ("password_protected_pdf",),
            )
        total_pages = len(reader.pages)
    except (OSError, PyPdfError, RecursionError, TypeError, ValueError) as exc:
        return ExtractionResult(
            (),
            "pdf",
            "error",
            False,
            (f"pdf_open_failed:{type(exc).__name__}:{exc}",),
        )
    chunks: list[ExtractedChunk] = []
    limitations: list[str] = []
    needs_ocr = False
    used_ocr = False
    unresolved_ocr = False
    extracted_chars = 0
    page_limit = min(total_pages, MAX_PDF_PAGES)
    for page_index in range(page_limit):
        page_number = page_index + 1
        if extracted_chars >= MAX_PDF_TEXT_CHARS:
            limitations.append("pdf_text_character_limit_reached")
            break
        try:
            page = reader.pages[page_index]
            text = page.extract_text() or ""
        except (
            AttributeError,
            KeyError,
            PyPdfError,
            RecursionError,
            TypeError,
            ValueError,
        ):
            text = ""
            limitations.append(f"page_{page_number}_text_extraction_failed")
        if not _text_is_useful(text):
            needs_ocr = True
            native_text = _normalize_text(text)
            if enable_ocr:
                try:
                    image_bytes = _render_pdf_page(path, page_index)
                except (
                    ArchiveError,
                    ImportError,
                    ModuleNotFoundError,
                    OSError,
                    RuntimeError,
                    ValueError,
                ):
                    limitations.append(f"page_{page_number}_ocr_render_unavailable")
                else:
                    ocr_text, warnings, succeeded = _run_local_ocr(image_bytes)
                    limitations.extend(
                        f"page_{page_number}_{warning}" for warning in warnings
                    )
                    if succeeded and ocr_text:
                        text_parts = [native_text] if native_text else []
                        if ocr_text.casefold() not in native_text.casefold():
                            text_parts.append(ocr_text)
                        text = "\n".join(text_parts)
                        used_ocr = True
                        limitations.append(
                            f"page_{page_number}_ocr_text_requires_visual_confirmation"
                        )
            if not _text_is_useful(text):
                unresolved_ocr = True
                limitations.append(f"page_{page_number}_no_extractable_text")
        remaining_chars = MAX_PDF_TEXT_CHARS - extracted_chars
        if len(text) > remaining_chars:
            text = text[:remaining_chars]
            limitations.append("pdf_text_character_limit_reached")
        extracted_chars += len(text)
        chunks.extend(
            _page_chunks(
                text,
                page_number=page_number,
                ordinal_start=len(chunks),
            )
        )
    if total_pages > MAX_PDF_PAGES:
        limitations.append(f"pdf_page_limit_reached:{MAX_PDF_PAGES}/{total_pages}")
    method = "pdf_text+local_ocr" if used_ocr else "pdf_text"
    if not chunks:
        status = "partial"
    elif limitations:
        status = "partial"
    else:
        status = "indexed"
    return ExtractionResult(
        tuple(chunks),
        method,
        status,
        needs_ocr and unresolved_ocr,
        tuple(dict.fromkeys(limitations)),
    )


def _image_frames(path: Path) -> tuple[tuple[bytes, ...], bool]:
    from PIL import Image

    frames: list[bytes] = []
    total_pixels = 0
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(path) as image:
                source_frame_count = max(
                    1,
                    int(getattr(image, "n_frames", 1)),
                )
                frame_count = min(
                    source_frame_count,
                    MAX_IMAGE_FRAMES,
                )
                for frame_index in range(frame_count):
                    image.seek(frame_index)
                    frame_pixels = int(image.width) * int(image.height)
                    total_pixels += frame_pixels
                    if total_pixels > MAX_IMAGE_TOTAL_PIXELS:
                        raise ArchiveError("Image frames exceed the OCR pixel limit.")
                    frame = image.convert("RGB")
                    buffer = io.BytesIO()
                    frame.save(buffer, format="PNG")
                    frames.append(buffer.getvalue())
    except (Image.DecompressionBombError, Image.DecompressionBombWarning) as exc:
        raise ArchiveError("Image exceeds the safe decoding limit.") from exc
    return tuple(frames), source_frame_count > MAX_IMAGE_FRAMES


def _extract_image(path: Path, *, enable_ocr: bool) -> ExtractionResult:
    if not enable_ocr:
        return ExtractionResult((), "image", "partial", True, ("ocr_disabled",))
    chunks: list[ExtractedChunk] = []
    limitations: list[str] = []
    unresolved_ocr = False
    try:
        frames, frames_truncated = _image_frames(path)
    except (ArchiveError, ImportError, ModuleNotFoundError, OSError, ValueError):
        return ExtractionResult(
            (),
            "image",
            "partial",
            True,
            ("image_frame_extraction_failed",),
        )
    if frames_truncated:
        limitations.append(f"image_frame_limit_reached:{MAX_IMAGE_FRAMES}")
    for page_number, image_bytes in enumerate(frames, start=1):
        text, warnings, succeeded = _run_local_ocr(image_bytes)
        limitations.extend(f"page_{page_number}_{warning}" for warning in warnings)
        if succeeded and text:
            limitations.append(
                f"page_{page_number}_ocr_text_requires_visual_confirmation"
            )
        else:
            unresolved_ocr = True
            limitations.append(f"page_{page_number}_no_extractable_text")
        chunks.extend(
            _page_chunks(
                text,
                page_number=page_number,
                ordinal_start=len(chunks),
            )
        )
    status = "indexed" if chunks and not limitations else "partial"
    return ExtractionResult(
        tuple(chunks),
        "local_ocr",
        status,
        unresolved_ocr or not bool(chunks),
        tuple(dict.fromkeys(limitations)),
    )


def _extract_document(path: Path, *, enable_ocr: bool) -> ExtractionResult:
    suffix = path.suffix.lower()
    try:
        if suffix in TEXT_SUFFIXES:
            return _extract_plain_text(path)
        if suffix in PDF_SUFFIXES:
            return _extract_pdf(path, enable_ocr=enable_ocr)
        if suffix in DOCX_SUFFIXES:
            return _extract_docx(path)
        if suffix in XLSX_SUFFIXES:
            return _extract_xlsx(path)
        if suffix in EMAIL_SUFFIXES:
            return _extract_eml(path)
        if suffix in IMAGE_SUFFIXES:
            return _extract_image(path, enable_ocr=enable_ocr)
    except (
        ArchiveError,
        AttributeError,
        ImportError,
        KeyError,
        ModuleNotFoundError,
        NotImplementedError,
        OSError,
        RecursionError,
        RuntimeError,
        TypeError,
        UnicodeError,
        ValueError,
        zipfile.BadZipFile,
        zipfile.LargeZipFile,
    ) as exc:
        return ExtractionResult(
            (),
            suffix.removeprefix(".") or "unknown",
            "error",
            suffix in IMAGE_SUFFIXES,
            (f"extraction_failed:{type(exc).__name__}:{exc}",),
        )
    return ExtractionResult((), "unsupported", "error", False, ("unsupported",))


def _delete_document_chunks(
    connection: sqlite3.Connection,
    document_id: str,
) -> None:
    source_rows = connection.execute(
        "SELECT source_id FROM chunks WHERE document_id = ?",
        (document_id,),
    ).fetchall()
    source_ids = [str(row["source_id"]) for row in source_rows]
    if source_ids:
        connection.executemany(
            "DELETE FROM chunk_fts WHERE source_id = ?",
            ((source_id,) for source_id in source_ids),
        )
    connection.execute(
        "DELETE FROM chunks WHERE document_id = ?",
        (document_id,),
    )


def _replace_document(
    connection: sqlite3.Connection,
    *,
    item: DiscoveredFile,
    source_sha256: str,
    extraction: ExtractionResult,
    generation: int,
    indexed_at: str,
) -> None:
    document_id = _stable_id("doc", item.relative_path.casefold())
    _delete_document_chunks(connection, document_id)
    connection.execute(
        """
        INSERT INTO documents(
            document_id, scope_id, relative_path, extension, size_bytes,
            mtime_ns, sha256, extraction_method, status, needs_ocr,
            limitations_json, indexed_at, last_seen_generation
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(document_id) DO UPDATE SET
            scope_id = excluded.scope_id,
            relative_path = excluded.relative_path,
            extension = excluded.extension,
            size_bytes = excluded.size_bytes,
            mtime_ns = excluded.mtime_ns,
            sha256 = excluded.sha256,
            extraction_method = excluded.extraction_method,
            status = excluded.status,
            needs_ocr = excluded.needs_ocr,
            limitations_json = excluded.limitations_json,
            indexed_at = excluded.indexed_at,
            last_seen_generation = excluded.last_seen_generation
        """,
        (
            document_id,
            item.scope_id,
            item.relative_path,
            item.path.suffix.lower(),
            item.size_bytes,
            item.mtime_ns,
            source_sha256,
            extraction.extraction_method,
            extraction.status,
            int(extraction.needs_ocr),
            json.dumps(extraction.limitations, ensure_ascii=False),
            indexed_at,
            generation,
        ),
    )
    for chunk in extraction.chunks:
        text_sha256 = _sha256_bytes(chunk.text.encode("utf-8"))
        source_id = _stable_id(
            "src",
            document_id,
            source_sha256,
            chunk.locator_kind,
            chunk.locator_value,
            str(chunk.ordinal),
            text_sha256,
        )
        connection.execute(
            """
            INSERT INTO chunks(
                source_id, document_id, ordinal, locator_kind, locator_value,
                text, text_sha256
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                source_id,
                document_id,
                chunk.ordinal,
                chunk.locator_kind,
                chunk.locator_value,
                chunk.text,
                text_sha256,
            ),
        )
        connection.execute(
            "INSERT INTO chunk_fts(source_id, text) VALUES (?, ?)",
            (source_id, chunk.text),
        )


def refresh_archive(
    *,
    rebuild: bool = False,
    enable_ocr: bool = False,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Incrementally refresh the private local index without writing source files."""

    private_state = _state_dir(state_dir)
    stored_config = _load_config(private_state, validate_scope_roots=False)
    current_scopes = _discover_top_level_scopes(stored_config.archive_root)
    scopes_changed = current_scopes != stored_config.scopes
    if scopes_changed:
        config = ArchiveConfig(
            archive_root=stored_config.archive_root,
            scopes=current_scopes,
            configured_at=_now_iso(),
        )
        _write_private_json(_config_path(private_state), config.as_json())
    else:
        config = stored_config
    discovered, scan_issues = _discover_files(config)
    connection = _connect(private_state)
    indexed_at = _now_iso()
    try:
        previous_generation = int(_metadata_get(connection, "scan_generation") or "0")
        generation = previous_generation + 1
        fingerprint = _config_fingerprint(config)
        indexed_root = _metadata_get(connection, "archive_root")
        root_changed = indexed_root is not None and indexed_root != str(
            config.archive_root
        )
        if rebuild or root_changed:
            _clear_index(connection)
        existing = {
            str(row["relative_path"]): row
            for row in connection.execute("SELECT * FROM documents").fetchall()
        }
        _replace_scan_issues(connection, scan_issues)
        counts = {
            "discovered_files": len(discovered),
            "indexed_files": 0,
            "unchanged_files": 0,
            "metadata_only_files": 0,
            "removed_files": 0,
            "partial_files": 0,
            "failed_files": 0,
            "needs_ocr_files": 0,
            "unsupported_files": sum(
                issue.reason == "unsupported_extension" for issue in scan_issues
            ),
            "oversized_files": sum(
                issue.reason == "file_size_limit_exceeded" for issue in scan_issues
            ),
        }
        for item in discovered:
            previous = existing.get(item.relative_path)
            requires_reindex = previous is not None and (
                str(previous["status"]) == "error"
                or (enable_ocr and bool(previous["needs_ocr"]))
            )
            try:
                source_path = _resolve_source_file(
                    config.archive_root,
                    item.relative_path,
                )
                source_sha256 = _sha256_file(source_path)
            except (ArchiveError, OSError) as exc:
                extraction = ExtractionResult(
                    (),
                    item.path.suffix.lower().removeprefix(".") or "unknown",
                    "error",
                    item.path.suffix.lower() in IMAGE_SUFFIXES,
                    (f"source_unavailable_before_extraction:{exc}",),
                )
                _replace_document(
                    connection,
                    item=item,
                    source_sha256="",
                    extraction=extraction,
                    generation=generation,
                    indexed_at=indexed_at,
                )
                counts["indexed_files"] += 1
                counts["failed_files"] += 1
                counts["needs_ocr_files"] += int(extraction.needs_ocr)
                continue
            if (
                previous is not None
                and str(previous["sha256"]) == source_sha256
                and not requires_reindex
            ):
                connection.execute(
                    """
                    UPDATE documents
                    SET size_bytes = ?, mtime_ns = ?, last_seen_generation = ?
                    WHERE relative_path = ?
                    """,
                    (
                        item.size_bytes,
                        item.mtime_ns,
                        generation,
                        item.relative_path,
                    ),
                )
                metadata_unchanged = (
                    int(previous["size_bytes"]) == item.size_bytes
                    and int(previous["mtime_ns"]) == item.mtime_ns
                )
                count_key = (
                    "unchanged_files" if metadata_unchanged else "metadata_only_files"
                )
                counts[count_key] += 1
                counts["partial_files"] += int(previous["status"] == "partial")
                counts["failed_files"] += int(previous["status"] == "error")
                counts["needs_ocr_files"] += int(previous["needs_ocr"])
                continue

            extraction = _extract_document(source_path, enable_ocr=enable_ocr)
            try:
                post_metadata = source_path.stat(follow_symlinks=False)
                post_sha256 = _sha256_file(source_path)
            except (ArchiveError, OSError) as exc:
                extraction = ExtractionResult(
                    (),
                    extraction.extraction_method,
                    "error",
                    extraction.needs_ocr,
                    (
                        *extraction.limitations,
                        f"source_unavailable_after_extraction:{exc}",
                    ),
                )
                post_metadata = None
                post_sha256 = ""
            if (
                post_metadata is None
                or post_metadata.st_size != item.size_bytes
                or post_metadata.st_mtime_ns != item.mtime_ns
                or post_sha256 != source_sha256
            ):
                extraction = ExtractionResult(
                    (),
                    extraction.extraction_method,
                    "error",
                    extraction.needs_ocr,
                    (*extraction.limitations, "source_changed_during_refresh"),
                )
            _replace_document(
                connection,
                item=item,
                source_sha256=source_sha256,
                extraction=extraction,
                generation=generation,
                indexed_at=indexed_at,
            )
            counts["indexed_files"] += 1
            counts["partial_files"] += int(extraction.status == "partial")
            counts["failed_files"] += int(extraction.status == "error")
            counts["needs_ocr_files"] += int(extraction.needs_ocr)

        removed_rows = connection.execute(
            """
            SELECT document_id FROM documents
            WHERE last_seen_generation != ?
            """,
            (generation,),
        ).fetchall()
        for row in removed_rows:
            document_id = str(row["document_id"])
            _delete_document_chunks(connection, document_id)
            connection.execute(
                "DELETE FROM documents WHERE document_id = ?",
                (document_id,),
            )
        counts["removed_files"] = len(removed_rows)
        _metadata_set(connection, "scan_generation", str(generation))
        _metadata_set(connection, "last_refresh_at", indexed_at)
        _metadata_set(connection, "config_fingerprint", fingerprint)
        _metadata_set(connection, "archive_root", str(config.archive_root))
        _metadata_set(connection, "ocr_enabled_last_refresh", json.dumps(enable_ocr))
        connection.commit()
        document_count = int(
            connection.execute("SELECT COUNT(*) AS count FROM documents").fetchone()[
                "count"
            ]
        )
        chunk_count = int(
            connection.execute("SELECT COUNT(*) AS count FROM chunks").fetchone()[
                "count"
            ]
        )
        recovered_clients = _synchronize_client_identities(private_state, config)
        return {
            "status": "refreshed",
            "last_refresh_at": indexed_at,
            "rebuild": bool(rebuild or root_changed),
            "scope_configuration_changed": scopes_changed,
            "scopes": _scope_records(config),
            "ocr_enabled": enable_ocr,
            "document_count": document_count,
            "chunk_count": chunk_count,
            "recovered_client_count": len(
                [
                    record
                    for record in recovered_clients
                    if any(scope.scope_id == record.scope_id for scope in config.scopes)
                ]
            ),
            **counts,
            **_scan_issue_status(connection),
            **_document_issue_status(connection),
        }
    finally:
        connection.close()


def _scope_records(config: ArchiveConfig) -> list[dict[str, str]]:
    return [scope.as_json() for scope in config.scopes]


def _current_scope_view(config: ArchiveConfig) -> tuple[ArchiveConfig, bool]:
    current_scopes = _discover_top_level_scopes(config.archive_root)
    changed = current_scopes != config.scopes
    if not changed:
        return config, False
    return (
        ArchiveConfig(
            archive_root=config.archive_root,
            scopes=current_scopes,
            configured_at=config.configured_at,
        ),
        True,
    )


def studio_archive_status(*, state_dir: Path | None = None) -> dict[str, Any]:
    """Return configuration and derived-index status without changing state."""

    private_state = _state_dir(state_dir)
    config_path = _config_path(private_state)
    if not config_path.is_file():
        return {
            "configured": False,
            "document_count": 0,
            "chunk_count": 0,
            "last_refresh_at": None,
            "scopes": [],
            "needs_ocr_document_count": 0,
            "partial_document_count": 0,
            "failed_document_count": 0,
            "scan_issue_count": 0,
            "scan_issues": [],
            "scan_issues_truncated": False,
            "document_issue_count": 0,
            "document_issues": [],
            "document_issues_truncated": False,
            **_studio_archive_setup_contract(),
        }
    stored_config = _load_config(private_state, validate_scope_roots=False)
    config, scopes_changed = _current_scope_view(stored_config)
    database_path = _database_path(private_state)
    if not database_path.is_file():
        return {
            "configured": True,
            "archive_root": str(config.archive_root),
            "document_count": 0,
            "chunk_count": 0,
            "last_refresh_at": None,
            "scopes": _scope_records(config),
            "index_requires_refresh": True,
            "scope_configuration_changed": scopes_changed,
            "needs_ocr_document_count": 0,
            "partial_document_count": 0,
            "failed_document_count": 0,
            "scan_issue_count": 0,
            "scan_issues": [],
            "scan_issues_truncated": False,
            "document_issue_count": 0,
            "document_issues": [],
            "document_issues_truncated": False,
        }
    connection = _connect(private_state, readonly=True)
    try:
        document_count = int(
            connection.execute("SELECT COUNT(*) AS count FROM documents").fetchone()[
                "count"
            ]
        )
        chunk_count = int(
            connection.execute("SELECT COUNT(*) AS count FROM chunks").fetchone()[
                "count"
            ]
        )
        fingerprint_matches = _metadata_get(
            connection, "config_fingerprint"
        ) == _config_fingerprint(stored_config)
        return {
            "configured": True,
            "archive_root": str(config.archive_root),
            "document_count": document_count,
            "chunk_count": chunk_count,
            "last_refresh_at": _metadata_get(connection, "last_refresh_at"),
            "scopes": _scope_records(config),
            "index_requires_refresh": scopes_changed or not fingerprint_matches,
            "scope_configuration_changed": scopes_changed,
            "needs_ocr_document_count": int(
                connection.execute(
                    "SELECT COUNT(*) AS count FROM documents WHERE needs_ocr = 1"
                ).fetchone()["count"]
            ),
            "partial_document_count": int(
                connection.execute(
                    "SELECT COUNT(*) AS count FROM documents WHERE status = 'partial'"
                ).fetchone()["count"]
            ),
            "failed_document_count": int(
                connection.execute(
                    "SELECT COUNT(*) AS count FROM documents WHERE status = 'error'"
                ).fetchone()["count"]
            ),
            **_scan_issue_status(connection),
            **_document_issue_status(connection),
        }
    finally:
        connection.close()


def _fts_query(value: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ArchiveError("Search query must be non-empty.")
    if len(value) > 500:
        raise ArchiveError("Search query must contain at most 500 characters.")
    tokens: list[str] = []
    seen: set[str] = set()
    for match in re.finditer(r"[^\W_]+", value, flags=re.UNICODE):
        token = match.group(0)
        key = token.casefold()
        if len(key) < 2 or key in seen:
            continue
        seen.add(key)
        tokens.append(token)
        if len(tokens) >= MAX_SEARCH_TOKENS:
            break
    if not tokens:
        raise ArchiveError("Search query contains no searchable terms.")
    return " OR ".join(f'"{token.replace(chr(34), chr(34) * 2)}"' for token in tokens)


def _citation(relative_path: str, locator_kind: str, locator_value: str) -> str:
    if locator_kind == "page":
        return f"{relative_path}, p. {locator_value}"
    return f"{relative_path}, {locator_kind} {locator_value}"


def _decode_limitations(value: str) -> list[str]:
    try:
        payload = json.loads(value)
    except json.JSONDecodeError:
        return ["invalid_stored_extraction_limitations"]
    if not isinstance(payload, list) or not all(
        isinstance(item, str) for item in payload
    ):
        return ["invalid_stored_extraction_limitations"]
    return payload


def search_archive(
    query: str,
    *,
    scope_id: str,
    limit: int = 10,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Search one exact configured scope, or ``all`` when explicitly requested."""

    if not 1 <= int(limit) <= 20:
        raise ArchiveError("Search limit must be between 1 and 20.")
    private_state = _state_dir(state_dir)
    config = _load_config(private_state, validate_scope_roots=False)
    _, scopes_changed = _current_scope_view(config)
    if scopes_changed:
        raise ArchiveError(
            "Top-level archive scopes changed; refresh before searching."
        )
    allowed_scopes = {scope.scope_id for scope in config.scopes}
    if scope_id != "all" and scope_id not in allowed_scopes:
        raise ArchiveError("Search scope is not configured.")
    expression = _fts_query(query)
    connection = _connect(private_state, readonly=True)
    try:
        if _metadata_get(connection, "config_fingerprint") != _config_fingerprint(
            config
        ):
            raise ArchiveError(
                "Archive configuration changed; refresh before searching."
            )
        # Deduplicate before LIMIT so one document with many high-scoring chunks
        # cannot mechanically hide other matching documents.
        if scope_id == "all":
            rows = connection.execute(
                """
                WITH matches AS (
                    SELECT
                        c.source_id,
                        c.document_id,
                        c.ordinal,
                        c.locator_kind,
                        c.locator_value,
                        d.scope_id,
                        d.relative_path,
                        d.sha256,
                        d.extraction_method,
                        d.status,
                        d.needs_ocr,
                        d.limitations_json,
                        d.indexed_at,
                        bm25(chunk_fts) AS score,
                        snippet(chunk_fts, 1, '[[', ']]', ' … ', 28) AS snippet
                    FROM chunk_fts
                    JOIN chunks AS c ON c.source_id = chunk_fts.source_id
                    JOIN documents AS d ON d.document_id = c.document_id
                    WHERE chunk_fts MATCH ?
                ), ranked AS (
                    SELECT
                        *,
                        ROW_NUMBER() OVER (
                            PARTITION BY scope_id, sha256
                            ORDER BY score, relative_path, ordinal
                        ) AS content_rank
                    FROM matches
                )
                SELECT *
                FROM ranked
                WHERE content_rank = 1
                ORDER BY score, relative_path, ordinal
                LIMIT ?
                """,
                (expression, int(limit)),
            ).fetchall()
        else:
            rows = connection.execute(
                """
                WITH matches AS (
                    SELECT
                        c.source_id,
                        c.document_id,
                        c.ordinal,
                        c.locator_kind,
                        c.locator_value,
                        d.scope_id,
                        d.relative_path,
                        d.sha256,
                        d.extraction_method,
                        d.status,
                        d.needs_ocr,
                        d.limitations_json,
                        d.indexed_at,
                        bm25(chunk_fts) AS score,
                        snippet(chunk_fts, 1, '[[', ']]', ' … ', 28) AS snippet
                    FROM chunk_fts
                    JOIN chunks AS c ON c.source_id = chunk_fts.source_id
                    JOIN documents AS d ON d.document_id = c.document_id
                    WHERE chunk_fts MATCH ?
                      AND d.scope_id = ?
                ), ranked AS (
                    SELECT
                        *,
                        ROW_NUMBER() OVER (
                            PARTITION BY scope_id, sha256
                            ORDER BY score, relative_path, ordinal
                        ) AS content_rank
                    FROM matches
                )
                SELECT *
                FROM ranked
                WHERE content_rank = 1
                ORDER BY score, relative_path, ordinal
                LIMIT ?
                """,
                (expression, scope_id, int(limit)),
            ).fetchall()
        results: list[dict[str, Any]] = []
        seen_documents: set[tuple[str, str]] = set()
        for row in rows:
            content_key = (str(row["scope_id"]), str(row["sha256"]))
            if content_key in seen_documents:
                continue
            seen_documents.add(content_key)
            results.append(
                {
                    "rank": len(results) + 1,
                    "source_id": str(row["source_id"]),
                    "document_id": str(row["document_id"]),
                    "scope_id": str(row["scope_id"]),
                    "relative_path": str(row["relative_path"]),
                    "locator_kind": str(row["locator_kind"]),
                    "locator_value": str(row["locator_value"]),
                    "citation": _citation(
                        str(row["relative_path"]),
                        str(row["locator_kind"]),
                        str(row["locator_value"]),
                    ),
                    "snippet": str(row["snippet"]),
                    "extraction_method": str(row["extraction_method"]),
                    "document_status": str(row["status"]),
                    "needs_ocr": bool(row["needs_ocr"]),
                    "limitations": _decode_limitations(str(row["limitations_json"])),
                    "source_sha256": str(row["sha256"]),
                    "indexed_at": str(row["indexed_at"]),
                    "score": float(row["score"]),
                    "verification_required": True,
                }
            )
            if len(results) == int(limit):
                break
        return {
            "query": query,
            "scope_id": scope_id,
            "result_count": len(results),
            "results": results,
            **_scan_issue_status(
                connection,
                scope_id=None if scope_id == "all" else scope_id,
            ),
            **_document_issue_status(
                connection,
                scope_id=None if scope_id == "all" else scope_id,
            ),
        }
    finally:
        connection.close()


def open_archive_source(
    source_id: str,
    *,
    context_chunks: int = 0,
    state_dir: Path | None = None,
) -> dict[str, Any]:
    """Open one indexed source ID after re-verifying the current source bytes."""

    if not isinstance(source_id, str) or not re.fullmatch(
        r"src_[0-9a-f]{24}", source_id
    ):
        raise ArchiveError("Source ID is invalid.")
    if not 0 <= int(context_chunks) <= 2:
        raise ArchiveError("context_chunks must be between 0 and 2.")
    private_state = _state_dir(state_dir)
    config = _load_config(private_state, validate_scope_roots=False)
    _, scopes_changed = _current_scope_view(config)
    if scopes_changed:
        raise ArchiveError(
            "Top-level archive scopes changed; refresh before opening sources."
        )
    connection = _connect(private_state, readonly=True)
    try:
        if _metadata_get(connection, "config_fingerprint") != _config_fingerprint(
            config
        ):
            raise ArchiveError(
                "Archive configuration changed; refresh before opening sources."
            )
        row = connection.execute(
            """
            SELECT c.*, d.relative_path, d.sha256, d.scope_id,
                   d.extraction_method, d.status, d.needs_ocr,
                   d.limitations_json, d.indexed_at
            FROM chunks AS c
            JOIN documents AS d ON d.document_id = c.document_id
            WHERE c.source_id = ?
            """,
            (source_id,),
        ).fetchone()
        if row is None:
            raise ArchiveError("Source ID is not present in the current index.")
        source_path = _resolve_source_file(
            config.archive_root, str(row["relative_path"])
        )
        current_sha256 = _sha256_file(source_path)
        if current_sha256 != str(row["sha256"]):
            raise SourceChangedError(
                "The source file changed after indexing. Refresh before using this "
                "citation; rebuild if the ordinary refresh still reports it as stale."
            )
        context = int(context_chunks)
        context_rows = connection.execute(
            """
            SELECT source_id, ordinal, locator_kind, locator_value, text
            FROM chunks
            WHERE document_id = ? AND ordinal BETWEEN ? AND ?
            ORDER BY ordinal
            """,
            (
                str(row["document_id"]),
                max(0, int(row["ordinal"]) - context),
                int(row["ordinal"]) + context,
            ),
        ).fetchall()
        fragments = [
            {
                "source_id": str(context_row["source_id"]),
                "ordinal": int(context_row["ordinal"]),
                "locator_kind": str(context_row["locator_kind"]),
                "locator_value": str(context_row["locator_value"]),
                "citation": _citation(
                    str(row["relative_path"]),
                    str(context_row["locator_kind"]),
                    str(context_row["locator_value"]),
                ),
                "text": str(context_row["text"])[:MAX_OPEN_CHARS],
            }
            for context_row in context_rows
        ]
        return {
            "source_id": source_id,
            "document_id": str(row["document_id"]),
            "scope_id": str(row["scope_id"]),
            "relative_path": str(row["relative_path"]),
            "locator_kind": str(row["locator_kind"]),
            "locator_value": str(row["locator_value"]),
            "citation": _citation(
                str(row["relative_path"]),
                str(row["locator_kind"]),
                str(row["locator_value"]),
            ),
            "source_sha256": str(row["sha256"]),
            "source_verified": True,
            "extraction_method": str(row["extraction_method"]),
            "document_status": str(row["status"]),
            "needs_ocr": bool(row["needs_ocr"]),
            "limitations": _decode_limitations(str(row["limitations_json"])),
            "indexed_at": str(row["indexed_at"]),
            "fragments": fragments,
            **_scan_issue_status(connection, scope_id=str(row["scope_id"])),
            **_document_issue_status(connection, scope_id=str(row["scope_id"])),
        }
    finally:
        connection.close()


def _verify_run_configuration(function: Callable[..., Any]) -> Callable[..., Any]:
    """Verify the pinned configuration before returning any public operation."""

    @wraps(function)
    def checked(*args: Any, **kwargs: Any) -> Any:
        result = function(*args, **kwargs)
        state = _state_dir(kwargs.get("state_dir"))
        if state in _STATE_LEASES:
            _check_configuration(state)
        return result

    return checked


for _operation_name in __all__:
    _operation = globals()[_operation_name]
    if callable(_operation) and not isinstance(_operation, type):
        globals()[_operation_name] = _verify_run_configuration(_operation)
