from __future__ import annotations

import sys as _bootstrap_sys  # isort: skip

_bootstrap_sys.dont_write_bytecode = True
_bootstrap_sys.pycache_prefix = (
    r"Z:\__concordato_no_bytecode__"
    if _bootstrap_sys.platform == "win32"
    else "/dev/null/concordato-plan-review"
)

import os as _bootstrap_os  # isort: skip

_BOOTSTRAP_PATH = _bootstrap_os.path.join(
    _bootstrap_os.path.dirname(_bootstrap_os.path.abspath(__file__)),
    "implementation_bootstrap.py",
)
_BOOTSTRAP_ENTRY = _bootstrap_os.lstat(_BOOTSTRAP_PATH)
if _BOOTSTRAP_ENTRY.st_mode & 0o170000 != 0o100000 or _BOOTSTRAP_ENTRY.st_nlink != 1:
    raise RuntimeError("Concordato implementation bootstrap is not a real file.")
with open(_BOOTSTRAP_PATH, "rb") as _bootstrap_handle:
    _BOOTSTRAP_BEFORE = _bootstrap_os.fstat(_bootstrap_handle.fileno())
    _BOOTSTRAP_BYTES = _bootstrap_handle.read()
    _BOOTSTRAP_AFTER = _bootstrap_os.fstat(_bootstrap_handle.fileno())
_BOOTSTRAP_IDENTITY = (
    _BOOTSTRAP_ENTRY.st_dev,
    _BOOTSTRAP_ENTRY.st_ino,
    _BOOTSTRAP_ENTRY.st_size,
    _BOOTSTRAP_ENTRY.st_mtime_ns,
)
if (
    _BOOTSTRAP_IDENTITY
    != (
        _BOOTSTRAP_BEFORE.st_dev,
        _BOOTSTRAP_BEFORE.st_ino,
        _BOOTSTRAP_BEFORE.st_size,
        _BOOTSTRAP_BEFORE.st_mtime_ns,
    )
    or _BOOTSTRAP_IDENTITY
    != (
        _BOOTSTRAP_AFTER.st_dev,
        _BOOTSTRAP_AFTER.st_ino,
        _BOOTSTRAP_AFTER.st_size,
        _BOOTSTRAP_AFTER.st_mtime_ns,
    )
    or len(_BOOTSTRAP_BYTES) != _BOOTSTRAP_AFTER.st_size
):
    raise RuntimeError("Concordato implementation bootstrap changed while read.")
_BOOTSTRAP_NAMESPACE = {
    "__file__": _BOOTSTRAP_PATH,
    "__name__": "_concordato_implementation_bootstrap",
}
# The exact stable single-link bootstrap source is verified above.
exec(  # nosec B102
    compile(_BOOTSTRAP_BYTES, _BOOTSTRAP_PATH, "exec"), _BOOTSTRAP_NAMESPACE
)
_BOOTSTRAP_NAMESPACE["activate_implementation_boundary"]()
_SCRIPTS_DIR = _bootstrap_os.path.dirname(_bootstrap_os.path.abspath(__file__))
if _SCRIPTS_DIR not in _bootstrap_sys.path:
    _bootstrap_sys.path.insert(0, _SCRIPTS_DIR)

import argparse
import json
import re
import shutil
import sys
from pathlib import Path
from typing import Any

from docx import Document
from replay_assurance import replay_assurance
from vera_assurance import (  # noqa: E402
    AssuranceContractError,
    load_client_engagement_context_file,
)

__all__ = ["apply_review_edits", "main"]

SUMMARY_DOCX = "concordato_preventivo_review_summary.docx"
REGENERATE_NATIVE_OUTPUT_ACTION = (
    "Regenerate native DOCX/XLSX/PDF outputs before final handoff."
)
FINAL_HANDOFF_ACTION = (
    "Review is recorded; professional conclusion and publication remain withheld."
)
COMPLETE_REVIEW_ACTION = "Complete remaining review decisions before final handoff."


def clean_text(value: object) -> str:
    """Return a stripped string for safe JSON field comparison."""

    return value.strip() if isinstance(value, str) else ""


def _read_json(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"Expected JSON object: {path}")
    return payload


def _write_json(path: Path, payload: dict[str, Any]) -> None:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _load_cli_customer_context(
    *,
    client_engagement: Path,
    output_dir: Path,
    persistent_output_dir: Path | None,
    input_paths: list[Path],
) -> dict[str, Any]:
    """Authorize either the canonical output or its MCP transaction copy."""

    context = load_client_engagement_context_file(
        client_engagement,
        expected_workflow_id="concordato-plan-review",
    )
    expected_output = Path(str(context["output_dir"])).resolve()
    persistent_output = (
        persistent_output_dir.expanduser().resolve()
        if persistent_output_dir is not None
        else expected_output
    )
    if persistent_output != expected_output:
        raise AssuranceContractError(
            "persistent Concordato output must be the customer run output root"
        )
    actual_output = output_dir.expanduser().resolve(strict=True)
    if actual_output == expected_output or actual_output.is_relative_to(
        expected_output
    ):
        return load_client_engagement_context_file(
            client_engagement,
            expected_workflow_id="concordato-plan-review",
            input_paths=input_paths,
            output_dir=actual_output,
        )
    try:
        relative = actual_output.relative_to(Path(str(context["run_root"])))
    except ValueError as exc:
        raise AssuranceContractError(
            "Concordato output is outside the customer run"
        ) from exc
    if not (
        len(relative.parts) == 2
        and relative.parts[0].startswith(".generated-review-transaction-")
        and relative.parts[1] == "working"
        and all(
            path.expanduser().resolve().is_relative_to(actual_output)
            for path in input_paths
        )
    ):
        raise AssuranceContractError(
            "Concordato output is outside the customer run and its review transaction"
        )
    return context


def _safe_item_id(value: object) -> str:
    text = clean_text(value) or "item"
    cleaned = "".join(char if char.isalnum() or char in "._-" else "-" for char in text)
    return cleaned.strip("-") or "item"


def _backup_file(output_dir: Path, item_id: str, target_name: str) -> dict[str, Any]:
    source = output_dir / target_name
    if not source.exists():
        return {}
    suffix = source.suffix or ".docx"
    relative = (
        Path("revisions")
        / "originals"
        / f"{source.stem}__{_safe_item_id(item_id)}{suffix}"
    )
    target = output_dir / relative
    target.parent.mkdir(parents=True, exist_ok=True)
    if not target.exists():
        shutil.copy2(source, target)
    return {
        "path": relative.as_posix(),
        "kind": suffix.lstrip(".") or "file",
        "status": "backup_original",
        "source_artifact": target_name,
        "item_id": item_id,
    }


def _upsert_output(outputs: list[dict[str, Any]], record: dict[str, Any]) -> None:
    path = record.get("path")
    for index, output in enumerate(outputs):
        if isinstance(output, dict) and output.get("path") == path:
            outputs[index] = {**output, **record}
            return
    outputs.append(record)


def _memo_effects(effects: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        effect
        for effect in effects
        if effect.get("action") == "edit"
        and clean_text(effect.get("item_id")) == "codex-review-memo"
        and clean_text(effect.get("edit_value"))
    ]


def _summary_docx_requested(output_dir: Path, final_artifacts: dict[str, Any]) -> bool:
    if (output_dir / SUMMARY_DOCX).exists():
        return True
    outputs = final_artifacts.get("outputs")
    if not isinstance(outputs, list):
        return False
    return any(
        isinstance(output, dict) and clean_text(output.get("path")) == SUMMARY_DOCX
        for output in outputs
    )


def _visible_memo_lines(markdown_text: str) -> list[str]:
    lines: list[str] = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"^[-*]\s+", "", line)
        if line:
            lines.append(line)
    return lines


def _required_docx_text(memo_text: str) -> list[str]:
    fragments = ["Memo revisore Codex"]
    fragments.extend(_visible_memo_lines(memo_text))
    return list(dict.fromkeys(fragments))


def _append_memo_to_summary_docx(path: Path, memo_text: str) -> None:
    document = Document(path)
    document.add_heading("Memo revisore Codex", level=1)
    for raw_line in memo_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading_match = re.match(r"^#{1,6}\s+(.*)$", line)
        if heading_match:
            document.add_heading(heading_match.group(1).strip(), level=2)
            continue
        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue
        document.add_paragraph(line)
    document.save(path)


def _write_review_memo(
    output_dir: Path,
    effect: dict[str, Any],
    memo_text: str,
) -> tuple[dict[str, Any], dict[str, Any] | None]:
    target_name = clean_text(effect.get("target_artifact")) or "codex_run_review.md"
    target_path = output_dir / target_name
    backup_output = _backup_file(
        output_dir,
        clean_text(effect.get("item_id")),
        target_name,
    )
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_existed = target_path.exists()
    target_path.write_text(memo_text, encoding="utf-8")
    revision_artifact = clean_text(effect.get("revision_artifact"))
    effect["target_artifact"] = target_name
    effect["artifact_update"] = (
        "target_artifact_updated" if target_existed else "target_artifact_created"
    )
    if revision_artifact:
        effect["promoted_from_revision"] = revision_artifact
    if backup_output:
        effect["original_artifact_backup"] = backup_output["path"]
    return (
        {
            "path": target_name,
            "kind": target_path.suffix.lstrip(".") or "md",
            "status": "updated_from_review",
            "item_id": clean_text(effect.get("item_id")),
            "size_bytes": target_path.stat().st_size,
            "required_text": [memo_text],
            "qa_checks": ["nonempty_text", "required_text"],
        },
        backup_output or None,
    )


def _pending_native_paths(effects: list[dict[str, Any]]) -> list[str]:
    paths: list[str] = []
    for effect in effects:
        if not effect.get("requires_native_regeneration"):
            continue
        raw_paths = effect.get("derived_native_regeneration_paths")
        if not isinstance(raw_paths, list):
            raw_paths = [effect.get("target_artifact")]
        for raw_path in raw_paths:
            text = clean_text(raw_path)
            if text:
                paths.append(text)
    return list(dict.fromkeys(paths))


def _application_status(applied: dict[str, Any]) -> str:
    if int(applied.get("blocker_count") or 0) > 0:
        return "blocked"
    if int(applied.get("native_regeneration_count") or 0) > 0:
        return "partial_review_applied"
    if int(applied.get("decision_count") or 0) < int(applied.get("item_count") or 0):
        return "partial_review_applied"
    return "review_applied_assurance_withheld"


def _next_actions(current: list[Any], status: str) -> list[str]:
    next_actions = [
        clean_text(action)
        for action in current
        if clean_text(action) != REGENERATE_NATIVE_OUTPUT_ACTION
    ]
    if status == "review_applied_assurance_withheld":
        next_actions.append(FINAL_HANDOFF_ACTION)
    elif status == "partial_review_applied":
        next_actions.append(COMPLETE_REVIEW_ACTION)
    return list(dict.fromkeys(action for action in next_actions if action))


def apply_review_edits(
    output_dir: Path,
    applied_decisions_path: Path,
    final_artifacts_path: Path,
    *,
    client_context: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Refresh Concordato native handoff artifacts after reviewed memo edits."""

    output_dir = output_dir.resolve()
    applied_decisions_path = applied_decisions_path.resolve()
    final_artifacts_path = final_artifacts_path.resolve()
    applied = _read_json(applied_decisions_path)
    final_artifacts = _read_json(final_artifacts_path)
    # The parent already replayed the predecessor closure before opening the
    # bounded transaction. At this point authorized review files exist, so the
    # predecessor whole-tree seal is intentionally stale until finalization.
    assurance_replay = replay_assurance(
        output_dir,
        require_output_closure=False,
        client_context=client_context,
    )
    persisted_review = _read_json(output_dir / "review_payload.json")
    applied_review = applied.get("review_payload")
    if not isinstance(applied_review, dict) or applied_review.get(
        "content_sha256"
    ) != persisted_review.get("content_sha256"):
        raise ValueError("Applied decisions are not bound to review_payload.json")
    effects = [
        effect for effect in applied.get("effects", []) if isinstance(effect, dict)
    ]
    candidate_effects = _memo_effects(effects)
    if not candidate_effects:
        return {
            "ok": True,
            "updated_effect_count": 0,
            "message": "No Concordato memo refresh was required.",
            "applied_decisions": applied,
            "final_artifacts": final_artifacts,
        }
    if not _summary_docx_requested(output_dir, final_artifacts):
        return {
            "ok": True,
            "updated_effect_count": 0,
            "message": "No Concordato summary DOCX is available to refresh.",
            "applied_decisions": applied,
            "final_artifacts": final_artifacts,
        }

    docx_path = output_dir / SUMMARY_DOCX
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    memo_text = clean_text(candidate_effects[-1].get("edit_value"))
    memo_output, memo_backup_output = _write_review_memo(
        output_dir,
        candidate_effects[-1],
        memo_text,
    )
    backup_output = _backup_file(
        output_dir,
        clean_text(candidate_effects[-1].get("item_id")),
        SUMMARY_DOCX,
    )
    _append_memo_to_summary_docx(docx_path, memo_text)

    native_regenerated_paths = [SUMMARY_DOCX]
    native_pending = [
        path
        for path in _pending_native_paths(effects)
        if path not in set(native_regenerated_paths)
    ]
    for effect in candidate_effects:
        effect["native_regeneration_status"] = "regenerated"
        effect["native_regenerated_paths"] = native_regenerated_paths

    applied["effects"] = effects
    applied["native_regeneration_count"] = len(native_pending)
    applied["native_regeneration_paths"] = native_pending
    applied["native_regenerated_count"] = len(native_regenerated_paths)
    applied["native_regenerated_paths"] = native_regenerated_paths
    target_update_paths = list(applied.get("target_update_paths") or [])
    if memo_output["path"] not in target_update_paths:
        target_update_paths.append(memo_output["path"])
    applied["target_update_paths"] = target_update_paths
    applied["target_update_count"] = len(target_update_paths)
    original_backup_paths = list(applied.get("original_backup_paths") or [])
    if memo_backup_output and memo_backup_output["path"] not in original_backup_paths:
        original_backup_paths.append(memo_backup_output["path"])
    if backup_output and backup_output["path"] not in original_backup_paths:
        original_backup_paths.append(backup_output["path"])
    applied["original_backup_paths"] = original_backup_paths
    applied["application_status"] = _application_status(applied)

    outputs = [
        output
        for output in final_artifacts.get("outputs", [])
        if isinstance(output, dict)
    ]
    _upsert_output(outputs, memo_output)
    if memo_backup_output:
        _upsert_output(outputs, memo_backup_output)
    _upsert_output(
        outputs,
        {
            "path": SUMMARY_DOCX,
            "kind": "docx",
            "status": "updated_from_review",
            "native_regenerated": True,
            "source_artifact": "codex_run_review.md",
            "size_bytes": docx_path.stat().st_size,
            "required_text": _required_docx_text(memo_text),
            "qa_checks": ["nonempty_text", "required_text"],
        },
    )
    if backup_output:
        _upsert_output(outputs, backup_output)
    final_artifacts["outputs"] = outputs
    final_artifacts["status"] = applied["application_status"]
    final_artifacts["review_status"] = applied["application_status"]
    final_artifacts["final_ready"] = False
    review_application = final_artifacts.setdefault("review_application", {})
    if isinstance(review_application, dict):
        review_application["application_status"] = applied["application_status"]
        review_application["native_regeneration_count"] = applied[
            "native_regeneration_count"
        ]
        review_application["native_regeneration_paths"] = native_pending
        review_application["native_regenerated_count"] = applied[
            "native_regenerated_count"
        ]
        review_application["native_regenerated_paths"] = native_regenerated_paths
        review_application["target_update_count"] = applied["target_update_count"]
        review_application["target_update_paths"] = target_update_paths
        review_application["original_backup_paths"] = original_backup_paths
    final_artifacts["next_actions"] = _next_actions(
        list(final_artifacts.get("next_actions") or []),
        applied["application_status"],
    )

    _write_json(applied_decisions_path, applied)
    _write_json(final_artifacts_path, final_artifacts)
    return {
        "ok": True,
        "updated_effect_count": len(candidate_effects),
        "native_regenerated_paths": native_regenerated_paths,
        "backup_paths": [backup_output["path"]] if backup_output else [],
        "application_status": applied["application_status"],
        "assurance_replay": assurance_replay,
        "applied_decisions": applied,
        "final_artifacts": final_artifacts,
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Apply Concordato review edits to downstream artifacts."
    )
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--applied-decisions", type=Path, required=True)
    parser.add_argument("--final-artifacts", type=Path, required=True)
    parser.add_argument("--client-engagement", type=Path, required=True)
    parser.add_argument("--persistent-output-dir", type=Path, help=argparse.SUPPRESS)
    args = parser.parse_args(argv)
    try:
        client_context = _load_cli_customer_context(
            client_engagement=args.client_engagement,
            output_dir=args.output_dir,
            persistent_output_dir=args.persistent_output_dir,
            input_paths=[args.applied_decisions, args.final_artifacts],
        )
    except AssuranceContractError as exc:
        parser.error(str(exc))
    result = apply_review_edits(
        args.output_dir,
        args.applied_decisions,
        args.final_artifacts,
        client_context=client_context,
    )
    sys.stdout.write(json.dumps(result, ensure_ascii=False) + "\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
