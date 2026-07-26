"""Reviewed semantic model and mechanical schedules for Concordato Preventivo.

This module validates and renders reviewer-supplied meaning. It never infers a
document's legal role, a creditor's priority, statutory compliance, plan
feasibility, or evidence sufficiency. Deterministic logic is limited to output
shape, source binding, exact arithmetic, and reproducible rendering because
those operations are mechanically verifiable and audit-relevant.
"""

from __future__ import annotations

import csv
import json
import re
import sys
from collections import defaultdict
from dataclasses import dataclass
from decimal import Decimal, localcontext
from pathlib import Path
from typing import Any, Mapping, Sequence

import openpyxl
from docx import Document
from openpyxl.styles import Alignment, Font, PatternFill


def _ensure_vendor_import_path() -> None:
    """Expose component-local or repository-shared Vera vendor modules."""

    component_root = Path(__file__).resolve().parent.parent
    candidates = (
        component_root / "vendor" / "modules",
        component_root.parent / "_shared" / "vendor" / "modules",
    )
    for candidate in candidates:
        if candidate.is_dir() and str(candidate) not in sys.path:
            sys.path.insert(0, str(candidate))


_ensure_vendor_import_path()
from vera_assurance import (  # noqa: E402
    MoneyValidationError,
    build_reviewed_decision_receipt,
    canonical_json_sha256,
    decimal_text,
    parse_canonical_decimal,
    validate_reviewed_decision_receipt,
)

__all__ = [
    "CASE_ADAPTER_ID",
    "CASE_ADAPTER_VERSION",
    "CASE_SCHEMA_VERSION",
    "REQUIRED_REVIEW_AREAS",
    "SemanticArtifacts",
    "build_case_model_template",
    "derive_case_schedules",
    "review_concordato_case_model",
    "validate_semantic_recipe",
    "write_semantic_artifacts",
]

CASE_SCHEMA_VERSION = "concordato.preventivo.case.v1"
CASE_RECIPE_SCHEMA_VERSION = "concordato.preventivo.semantic_recipe.v1"
CASE_OUTPUT_SCHEMA_VERSION = "concordato.preventivo.case_output.v1"
CASE_ADAPTER_ID = "concordato_preventivo_semantic_case"
CASE_ADAPTER_VERSION = "v1"
CASE_DECISION_TYPE = "semantic_review"
DECIMAL_PRECISION = 96

PLAN_TYPES = {
    "continuity_direct",
    "continuity_indirect",
    "liquidation",
    "mixed",
    "unclear",
}
PROCEDURE_STAGES = {
    "draft",
    "filed",
    "opened",
    "voting",
    "homologated",
    "execution",
    "closed",
    "unclear",
}
PERIMETER_STATUSES = {"complete", "partial", "missing", "unclear"}
DOCUMENT_ROLES = {
    "proposal",
    "plan",
    "attestation",
    "creditor_schedule",
    "accounting_records",
    "business_plan",
    "financial_model",
    "liquidation_analysis",
    "tax_social_security_schedule",
    "valuation",
    "professional_report",
    "judicial_commissioner_report",
    "court_filing",
    "court_order",
    "other_support",
    "excluded",
}
CLAIM_STATUSES = {"asserted", "admitted", "contested", "conditional", "unclear"}
PRIORITIES = {
    "prededuction",
    "secured",
    "privileged",
    "unsecured",
    "subordinated",
    "unclear",
}
TREATMENT_FORMS = {
    "cash",
    "non_cash",
    "mixed",
    "continuing_relationship",
    "unclear",
}
VOTING_TREATMENTS = {"voting", "non_voting", "partially_voting", "unclear"}
QUESTION_ASSESSMENTS = {"addressed", "gap", "unclear", "not_applicable"}
ASSUMPTION_STATUSES = {"supported", "unsupported", "unclear"}
ISSUE_STATUSES = {"open", "resolved", "unclear"}
SEVERITIES = {"critical", "high", "medium", "low", "not_assessed"}
MILESTONE_STATUSES = {"planned", "completed", "delayed", "unclear"}
REQUIRED_REVIEW_AREAS = (
    "procedure_identity",
    "proposal_plan_consistency",
    "document_perimeter",
    "creditor_perimeter",
    "creditor_treatment",
    "voting_homologation",
    "liquidation_alternative",
    "feasibility_liquidity",
    "attestation",
    "accounting_consistency",
    "tax_social_security",
)
REVIEW_AREAS = set(REQUIRED_REVIEW_AREAS) | {
    "continuity_economics",
    "sources_and_uses",
    "milestones",
    "other",
}
DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
HTTP_URL_RE = re.compile(r"^https?://", re.IGNORECASE)

TEXT = {
    "it": {
        "title": "Revisione del concordato preventivo",
        "status": "Stato del modello semantico",
        "withheld": "Modello semantico non ancora riesaminato",
        "reviewed": "Modello semantico riesaminato",
        "procedure": "Procedura",
        "deterministic": "Riepilogo aritmetico",
        "questions": "Domande di revisione",
        "issues": "Questioni aperte",
        "numeric_appendix": "Appendice di tie-out numerico",
        "limitation": (
            "Il documento organizza evidenze e controlli aritmetici; non costituisce "
            "parere legale né attestazione del piano."
        ),
        "no_rows": "Nessuna riga riesaminata",
    },
    "en": {
        "title": "Concordato Preventivo Review",
        "status": "Semantic model status",
        "withheld": "Semantic model not yet reviewed",
        "reviewed": "Semantic model reviewed",
        "procedure": "Procedure",
        "deterministic": "Arithmetic summary",
        "questions": "Review questions",
        "issues": "Open issues",
        "numeric_appendix": "Numerical tie-out appendix",
        "limitation": (
            "This document organizes evidence and arithmetic controls; it is not a "
            "legal opinion or plan attestation."
        ),
        "no_rows": "No reviewed rows",
    },
    "fr": {
        "title": "Revue du concordato preventivo",
        "status": "Statut du modèle sémantique",
        "withheld": "Modèle sémantique non encore revu",
        "reviewed": "Modèle sémantique revu",
        "procedure": "Procédure",
        "deterministic": "Synthèse arithmétique",
        "questions": "Questions de revue",
        "issues": "Points ouverts",
        "numeric_appendix": "Annexe de rapprochement numérique",
        "limitation": (
            "Ce document organise les éléments probants et les contrôles "
            "arithmétiques; il ne constitue ni un avis juridique ni une attestation."
        ),
        "no_rows": "Aucune ligne revue",
    },
    "de": {
        "title": "Prüfung des Concordato Preventivo",
        "status": "Status des semantischen Modells",
        "withheld": "Semantisches Modell noch nicht geprüft",
        "reviewed": "Semantisches Modell geprüft",
        "procedure": "Verfahren",
        "deterministic": "Arithmetische Übersicht",
        "questions": "Prüffragen",
        "issues": "Offene Punkte",
        "numeric_appendix": "Anhang zum Zahlenabgleich",
        "limitation": (
            "Dieses Dokument strukturiert Nachweise und Rechenkontrollen; es ist "
            "weder ein Rechtsgutachten noch eine Planbescheinigung."
        ),
        "no_rows": "Keine geprüften Zeilen",
    },
    "es": {
        "title": "Revisión del concordato preventivo",
        "status": "Estado del modelo semántico",
        "withheld": "Modelo semántico aún no revisado",
        "reviewed": "Modelo semántico revisado",
        "procedure": "Procedimiento",
        "deterministic": "Resumen aritmético",
        "questions": "Preguntas de revisión",
        "issues": "Cuestiones abiertas",
        "numeric_appendix": "Anexo de conciliación numérica",
        "limitation": (
            "Este documento organiza evidencias y controles aritméticos; no es un "
            "dictamen jurídico ni una atestación del plan."
        ),
        "no_rows": "No hay filas revisadas",
    },
}

QUESTION_TEXT = {
    "it": {
        "procedure_identity": "La procedura, lo stato e la versione del piano sono identificati?",
        "proposal_plan_consistency": "Proposta e piano sono coerenti tra loro?",
        "document_perimeter": "Il perimetro documentale è completo e aggiornato?",
        "creditor_perimeter": "Il perimetro dei creditori è completo alla data di riferimento?",
        "creditor_treatment": "Classi, prelazioni, trattamento e tempi sono ricostruiti?",
        "voting_homologation": "Voto, maggioranze, eventuali contestazioni e stato dell'omologazione sono ricostruiti?",
        "liquidation_alternative": "L'alternativa liquidatoria è ricostruita e confrontabile?",
        "feasibility_liquidity": "Fonti, impieghi e liquidità rendono il piano eseguibile?",
        "attestation": "L'attestazione è presente e coerente con proposta e piano?",
        "accounting_consistency": "Piano e proposta sono coerenti con i dati contabili?",
        "tax_social_security": "Il trattamento tributario e previdenziale è ricostruito?",
    },
    "en": {
        "procedure_identity": "Are the procedure, stage, and plan version identified?",
        "proposal_plan_consistency": "Are the proposal and plan mutually consistent?",
        "document_perimeter": "Is the document perimeter complete and current?",
        "creditor_perimeter": "Is the creditor perimeter complete at the cut-off?",
        "creditor_treatment": "Are classes, priorities, treatment, and timing reconstructed?",
        "voting_homologation": "Are voting, majorities, objections, and homologation status reconstructed?",
        "liquidation_alternative": "Is the liquidation alternative reconstructed and comparable?",
        "feasibility_liquidity": "Do sources, uses, and liquidity support execution of the plan?",
        "attestation": "Is the attestation present and consistent with proposal and plan?",
        "accounting_consistency": "Are the proposal and plan consistent with accounting data?",
        "tax_social_security": "Is tax and social-security treatment reconstructed?",
    },
    "fr": {
        "procedure_identity": "La procédure, son stade et la version du plan sont-ils identifiés ?",
        "proposal_plan_consistency": "La proposition et le plan sont-ils cohérents entre eux ?",
        "document_perimeter": "Le périmètre documentaire est-il complet et à jour ?",
        "creditor_perimeter": "Le périmètre des créanciers est-il complet à la date de référence ?",
        "creditor_treatment": "Les classes, priorités, traitements et échéances sont-ils reconstitués ?",
        "voting_homologation": "Le vote, les majorités, les objections et le statut de l'homologation sont-ils reconstitués ?",
        "liquidation_alternative": "L'alternative liquidative est-elle reconstituée et comparable ?",
        "feasibility_liquidity": "Les sources, emplois et liquidités permettent-ils l'exécution du plan ?",
        "attestation": "L'attestation est-elle présente et cohérente avec la proposition et le plan ?",
        "accounting_consistency": "La proposition et le plan sont-ils cohérents avec les données comptables ?",
        "tax_social_security": "Le traitement fiscal et social est-il reconstitué ?",
    },
    "de": {
        "procedure_identity": "Sind Verfahren, Stadium und Planversion identifiziert?",
        "proposal_plan_consistency": "Stimmen Vorschlag und Plan miteinander überein?",
        "document_perimeter": "Ist der Dokumentenumfang vollständig und aktuell?",
        "creditor_perimeter": "Ist der Gläubigerumfang zum Stichtag vollständig?",
        "creditor_treatment": "Sind Klassen, Rang, Behandlung und Zeitplan rekonstruiert?",
        "voting_homologation": "Sind Abstimmung, Mehrheiten, Einwände und Homologationsstatus rekonstruiert?",
        "liquidation_alternative": "Ist die Liquidationsalternative rekonstruiert und vergleichbar?",
        "feasibility_liquidity": "Tragen Quellen, Verwendungen und Liquidität die Durchführung des Plans?",
        "attestation": "Liegt die Bescheinigung vor und stimmt sie mit Vorschlag und Plan überein?",
        "accounting_consistency": "Stimmen Vorschlag und Plan mit den Rechnungslegungsdaten überein?",
        "tax_social_security": "Ist die steuer- und sozialversicherungsrechtliche Behandlung rekonstruiert?",
    },
    "es": {
        "procedure_identity": "¿Están identificados el procedimiento, la fase y la versión del plan?",
        "proposal_plan_consistency": "¿Son coherentes entre sí la propuesta y el plan?",
        "document_perimeter": "¿Está completo y actualizado el perímetro documental?",
        "creditor_perimeter": "¿Está completo el perímetro de acreedores a la fecha de referencia?",
        "creditor_treatment": "¿Se han reconstruido clases, prioridades, tratamiento y plazos?",
        "voting_homologation": "¿Se han reconstruido la votación, las mayorías, las objeciones y el estado de homologación?",
        "liquidation_alternative": "¿Se ha reconstruido la alternativa de liquidación y es comparable?",
        "feasibility_liquidity": "¿Respaldan las fuentes, los usos y la liquidez la ejecución del plan?",
        "attestation": "¿Está presente la atestación y es coherente con la propuesta y el plan?",
        "accounting_consistency": "¿Son coherentes la propuesta y el plan con los datos contables?",
        "tax_social_security": "¿Se ha reconstruido el tratamiento fiscal y de seguridad social?",
    },
}


@dataclass(frozen=True)
class SemanticArtifacts:
    """Semantic output paths and reviewed mechanical results."""

    status: str
    decision: dict[str, Any] | None
    case_model: dict[str, Any] | None
    derived: dict[str, Any]
    artifact_paths: tuple[str, ...]
    error: str | None = None


def _language(value: str) -> str:
    code = str(value or "it").lower().replace("_", "-").split("-", 1)[0]
    return code if code in TEXT else "it"


def _object(value: object, label: str) -> Mapping[str, Any]:
    if not isinstance(value, Mapping):
        raise ValueError(f"{label} must be an object")
    return value


def _array(value: object, label: str) -> list[Any]:
    if not isinstance(value, list):
        raise ValueError(f"{label} must be an array")
    return value


def _exact_keys(
    value: Mapping[str, Any],
    *,
    required: set[str],
    optional: set[str] = frozenset(),
    label: str,
) -> None:
    keys = set(value)
    missing = sorted(required - keys)
    unexpected = sorted(keys - required - optional)
    if missing or unexpected:
        raise ValueError(
            f"{label} fields are invalid; missing={missing}, unexpected={unexpected}"
        )


def _text(
    value: object,
    label: str,
    *,
    allow_empty: bool = False,
    max_length: int = 20_000,
) -> str:
    if not isinstance(value, str):
        raise ValueError(f"{label} must be text")
    normalized = value.strip()
    if not allow_empty and not normalized:
        raise ValueError(f"{label} must not be empty")
    if len(normalized) > max_length:
        raise ValueError(f"{label} exceeds {max_length} characters")
    return normalized


def _enum(value: object, allowed: set[str], label: str) -> str:
    normalized = _text(value, label)
    if normalized not in allowed:
        raise ValueError(f"{label} is unsupported: {normalized}")
    return normalized


def _iso_date(value: object, label: str, *, allow_empty: bool = True) -> str:
    normalized = _text(value, label, allow_empty=allow_empty, max_length=10)
    if normalized and not DATE_RE.fullmatch(normalized):
        raise ValueError(f"{label} must use YYYY-MM-DD")
    return normalized


def _currency(value: object, label: str) -> str:
    normalized = _text(value, label, max_length=3).upper()
    if not re.fullmatch(r"[A-Z]{3}", normalized):
        raise ValueError(f"{label} must be a three-letter currency code")
    return normalized


def _money(
    value: object,
    label: str,
    *,
    nonnegative: bool = False,
) -> str:
    if (
        not isinstance(value, str)
        or re.fullmatch(
            r"-?(?:0|[1-9]\d*)(?:\.\d+)?",
            value,
        )
        is None
    ):
        raise ValueError(f"{label} must be unambiguous decimal text")
    parsed = Decimal(value)
    if not parsed.is_finite():
        raise ValueError(f"{label} must be finite")
    if nonnegative and parsed < 0:
        raise ValueError(f"{label} must not be negative")
    return decimal_text(parsed)


def _identifier(value: object, label: str) -> str:
    normalized = _text(value, label, max_length=200)
    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9._:-]*", normalized):
        raise ValueError(f"{label} must be a stable identifier")
    return normalized


def _known_sources(
    inventory: Sequence[Mapping[str, Any]],
) -> dict[str, str]:
    sources: dict[str, str] = {}
    for row in inventory:
        source_ref = str(row.get("source_artifact_ref") or "").strip()
        relative_path = str(row.get("relative_path") or "").strip()
        if source_ref and relative_path and row.get("supported"):
            if source_ref in sources:
                raise ValueError("Inventory source artifact references must be unique")
            sources[source_ref] = relative_path
    return sources


def _evidence_refs(
    value: object,
    *,
    known_sources: Mapping[str, str],
    label: str,
) -> list[dict[str, str]]:
    normalized: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for index, raw in enumerate(_array(value, label)):
        item = _object(raw, f"{label}[{index}]")
        _exact_keys(
            item,
            required={"source_artifact_ref", "locator"},
            optional={"excerpt", "relative_path"},
            label=f"{label}[{index}]",
        )
        source_ref = _text(
            item["source_artifact_ref"],
            f"{label}[{index}].source_artifact_ref",
            max_length=300,
        )
        if source_ref not in known_sources:
            raise ValueError(
                f"{label}[{index}] references a source outside the perimeter"
            )
        supplied_path = _text(
            item.get("relative_path", ""),
            f"{label}[{index}].relative_path",
            allow_empty=True,
            max_length=1_000,
        )
        if supplied_path and supplied_path != known_sources[source_ref]:
            raise ValueError(f"{label}[{index}] source path is stale")
        locator = _text(item["locator"], f"{label}[{index}].locator", max_length=500)
        identity = (source_ref, locator)
        if identity in seen:
            raise ValueError(f"{label} contains a duplicate evidence locator")
        seen.add(identity)
        normalized.append(
            {
                "source_artifact_ref": source_ref,
                "relative_path": known_sources[source_ref],
                "locator": locator,
                "excerpt": _text(
                    item.get("excerpt", ""),
                    f"{label}[{index}].excerpt",
                    allow_empty=True,
                    max_length=2_000,
                ),
            }
        )
    return normalized


def _unique_rows(
    rows: Sequence[dict[str, Any]],
    *,
    id_field: str,
    label: str,
) -> None:
    identities = [str(row[id_field]) for row in rows]
    if len(identities) != len(set(identities)):
        raise ValueError(f"{label} identifiers must be unique")


def build_case_model_template(
    inventory: Sequence[Mapping[str, Any]],
    *,
    reference_date: str = "",
    language: str = "it",
) -> dict[str, Any]:
    """Build an explicitly unreviewed semantic template from captured sources."""

    known_sources = _known_sources(inventory)
    code = _language(language)
    questions = QUESTION_TEXT.get(code, QUESTION_TEXT["en"])
    return {
        "schema_version": CASE_SCHEMA_VERSION,
        "legal_framework": {
            "jurisdiction": "IT",
            "instrument": "concordato_preventivo",
            "framework_name": "Codice della crisi d'impresa e dell'insolvenza",
            "as_of_date": "",
            "authority_refs": [],
            "judgment_basis": "",
        },
        "procedure": {
            "identification_status": "unclear",
            "debtor_name": "",
            "court": "",
            "procedure_reference": "",
            "stage": "unclear",
            "plan_type": "unclear",
            "reference_date": reference_date,
            "currency": "EUR",
            "judgment_basis": "",
        },
        "document_perimeter": {
            "status": "unclear",
            "judgment_basis": "",
            "documents": [
                {
                    "source_artifact_ref": source_ref,
                    "relative_path": relative_path,
                    "roles": ["needs_review"],
                    "authoritative_for": [],
                    "version_date": "",
                    "judgment_basis": "",
                }
                for source_ref, relative_path in sorted(
                    known_sources.items(), key=lambda item: item[1]
                )
            ],
        },
        "creditor_population": {
            "status": "missing",
            "cutoff_date": reference_date,
            "currency": "EUR",
            "judgment_basis": "",
            "creditors": [],
        },
        "sources_and_uses": {
            "status": "missing",
            "currency": "EUR",
            "balance_tolerance": "0.01",
            "judgment_basis": "",
            "items": [],
        },
        "liquidity": {
            "status": "missing",
            "currency": "EUR",
            "bridge_tolerance": "0.01",
            "judgment_basis": "",
            "periods": [],
        },
        "milestones": [],
        "review_questions": [
            {
                "question_id": f"rq-{index:02d}",
                "area": area,
                "question": questions.get(area, QUESTION_TEXT["en"][area]),
                "assessment": "unclear",
                "evidence_refs": [],
                "judgment_basis": "",
                "follow_up": "",
            }
            for index, area in enumerate(REQUIRED_REVIEW_AREAS, start=1)
        ],
        "assumptions": [],
        "issues": [],
    }


def _normalize_authority_refs(value: object) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    for index, raw in enumerate(_array(value, "legal_framework.authority_refs")):
        item = _object(raw, f"legal_framework.authority_refs[{index}]")
        _exact_keys(
            item,
            required={"title", "url", "provisions"},
            label=f"legal_framework.authority_refs[{index}]",
        )
        url = _text(
            item["url"],
            f"legal_framework.authority_refs[{index}].url",
            max_length=2_000,
        )
        if not HTTP_URL_RE.match(url):
            raise ValueError("Legal authority URLs must use HTTP or HTTPS")
        provisions = [
            _text(
                provision,
                f"legal_framework.authority_refs[{index}].provisions",
                max_length=200,
            )
            for provision in _array(
                item["provisions"],
                f"legal_framework.authority_refs[{index}].provisions",
            )
        ]
        normalized.append(
            {
                "title": _text(
                    item["title"],
                    f"legal_framework.authority_refs[{index}].title",
                    max_length=500,
                ),
                "url": url,
                "provisions": provisions,
            }
        )
    return normalized


def _normalize_document_perimeter(
    value: object,
    *,
    known_sources: Mapping[str, str],
) -> dict[str, Any]:
    perimeter = _object(value, "document_perimeter")
    _exact_keys(
        perimeter,
        required={"status", "judgment_basis", "documents"},
        label="document_perimeter",
    )
    documents: list[dict[str, Any]] = []
    for index, raw in enumerate(
        _array(perimeter["documents"], "document_perimeter.documents")
    ):
        item = _object(raw, f"document_perimeter.documents[{index}]")
        _exact_keys(
            item,
            required={
                "source_artifact_ref",
                "relative_path",
                "roles",
                "authoritative_for",
                "version_date",
                "judgment_basis",
            },
            label=f"document_perimeter.documents[{index}]",
        )
        source_ref = _text(
            item["source_artifact_ref"],
            f"document_perimeter.documents[{index}].source_artifact_ref",
            max_length=300,
        )
        if source_ref not in known_sources:
            raise ValueError(
                "Document perimeter references a source outside the capture"
            )
        relative_path = _text(
            item["relative_path"],
            f"document_perimeter.documents[{index}].relative_path",
            max_length=1_000,
        )
        if relative_path != known_sources[source_ref]:
            raise ValueError(
                "Document perimeter path does not match its source receipt"
            )
        roles = [
            _enum(
                role,
                DOCUMENT_ROLES,
                f"document_perimeter.documents[{index}].roles",
            )
            for role in _array(
                item["roles"], f"document_perimeter.documents[{index}].roles"
            )
        ]
        if not roles or len(roles) != len(set(roles)):
            raise ValueError("Every document requires unique reviewed semantic roles")
        if "excluded" in roles and len(roles) != 1:
            raise ValueError("An excluded document cannot have another semantic role")
        authoritative_for = [
            _enum(
                role,
                DOCUMENT_ROLES - {"excluded"},
                f"document_perimeter.documents[{index}].authoritative_for",
            )
            for role in _array(
                item["authoritative_for"],
                f"document_perimeter.documents[{index}].authoritative_for",
            )
        ]
        if len(authoritative_for) != len(set(authoritative_for)) or any(
            role not in roles for role in authoritative_for
        ):
            raise ValueError("Authoritative document roles must be unique role subsets")
        documents.append(
            {
                "source_artifact_ref": source_ref,
                "relative_path": relative_path,
                "roles": roles,
                "authoritative_for": authoritative_for,
                "version_date": _iso_date(
                    item["version_date"],
                    f"document_perimeter.documents[{index}].version_date",
                ),
                "judgment_basis": _text(
                    item["judgment_basis"],
                    f"document_perimeter.documents[{index}].judgment_basis",
                ),
            }
        )
    document_refs = [row["source_artifact_ref"] for row in documents]
    if set(document_refs) != set(known_sources) or len(document_refs) != len(
        known_sources
    ):
        raise ValueError(
            "Document perimeter must classify every captured supported source exactly once"
        )
    return {
        "status": _enum(
            perimeter["status"], PERIMETER_STATUSES, "document_perimeter.status"
        ),
        "judgment_basis": _text(
            perimeter["judgment_basis"], "document_perimeter.judgment_basis"
        ),
        "documents": sorted(documents, key=lambda row: row["relative_path"]),
    }


def _normalize_creditors(
    value: object,
    *,
    known_sources: Mapping[str, str],
) -> dict[str, Any]:
    population = _object(value, "creditor_population")
    _exact_keys(
        population,
        required={
            "status",
            "cutoff_date",
            "currency",
            "judgment_basis",
            "creditors",
        },
        label="creditor_population",
    )
    currency = _currency(population["currency"], "creditor_population.currency")
    creditors: list[dict[str, Any]] = []
    for index, raw in enumerate(
        _array(population["creditors"], "creditor_population.creditors")
    ):
        item = _object(raw, f"creditor_population.creditors[{index}]")
        _exact_keys(
            item,
            required={
                "creditor_id",
                "creditor_name",
                "claim_amount",
                "claim_status",
                "priority",
                "class_id",
                "treatment_form",
                "proposed_cash_amount",
                "proposed_non_cash_amount",
                "liquidation_recovery_amount",
                "payment_start",
                "payment_end",
                "voting_treatment",
                "evidence_refs",
                "judgment_basis",
            },
            optional={"currency"},
            label=f"creditor_population.creditors[{index}]",
        )
        supplied_currency = item.get("currency")
        if (
            supplied_currency is not None
            and _currency(
                supplied_currency,
                f"creditor_population.creditors[{index}].currency",
            )
            != currency
        ):
            raise ValueError("Creditor row currency differs from the population")
        creditors.append(
            {
                "creditor_id": _identifier(
                    item["creditor_id"],
                    f"creditor_population.creditors[{index}].creditor_id",
                ),
                "creditor_name": _text(
                    item["creditor_name"],
                    f"creditor_population.creditors[{index}].creditor_name",
                    max_length=500,
                ),
                "claim_amount": _money(
                    item["claim_amount"],
                    f"creditor_population.creditors[{index}].claim_amount",
                    nonnegative=True,
                ),
                "claim_status": _enum(
                    item["claim_status"],
                    CLAIM_STATUSES,
                    f"creditor_population.creditors[{index}].claim_status",
                ),
                "priority": _enum(
                    item["priority"],
                    PRIORITIES,
                    f"creditor_population.creditors[{index}].priority",
                ),
                "class_id": _identifier(
                    item["class_id"],
                    f"creditor_population.creditors[{index}].class_id",
                ),
                "treatment_form": _enum(
                    item["treatment_form"],
                    TREATMENT_FORMS,
                    f"creditor_population.creditors[{index}].treatment_form",
                ),
                "proposed_cash_amount": _money(
                    item["proposed_cash_amount"],
                    f"creditor_population.creditors[{index}].proposed_cash_amount",
                    nonnegative=True,
                ),
                "proposed_non_cash_amount": _money(
                    item["proposed_non_cash_amount"],
                    f"creditor_population.creditors[{index}].proposed_non_cash_amount",
                    nonnegative=True,
                ),
                "liquidation_recovery_amount": _money(
                    item["liquidation_recovery_amount"],
                    (
                        "creditor_population.creditors"
                        f"[{index}].liquidation_recovery_amount"
                    ),
                    nonnegative=True,
                ),
                "payment_start": _iso_date(
                    item["payment_start"],
                    f"creditor_population.creditors[{index}].payment_start",
                ),
                "payment_end": _iso_date(
                    item["payment_end"],
                    f"creditor_population.creditors[{index}].payment_end",
                ),
                "voting_treatment": _enum(
                    item["voting_treatment"],
                    VOTING_TREATMENTS,
                    f"creditor_population.creditors[{index}].voting_treatment",
                ),
                "evidence_refs": _evidence_refs(
                    item["evidence_refs"],
                    known_sources=known_sources,
                    label=f"creditor_population.creditors[{index}].evidence_refs",
                ),
                "judgment_basis": _text(
                    item["judgment_basis"],
                    f"creditor_population.creditors[{index}].judgment_basis",
                ),
                "currency": currency,
            }
        )
    _unique_rows(creditors, id_field="creditor_id", label="Creditors")
    status = _enum(
        population["status"], PERIMETER_STATUSES, "creditor_population.status"
    )
    if status == "complete" and not creditors:
        raise ValueError("A complete creditor population cannot be empty")
    return {
        "status": status,
        "cutoff_date": _iso_date(
            population["cutoff_date"],
            "creditor_population.cutoff_date",
            allow_empty=False,
        ),
        "currency": currency,
        "judgment_basis": _text(
            population["judgment_basis"], "creditor_population.judgment_basis"
        ),
        "creditors": sorted(creditors, key=lambda row: row["creditor_id"]),
    }


def _normalize_sources_and_uses(
    value: object,
    *,
    known_sources: Mapping[str, str],
) -> dict[str, Any]:
    schedule = _object(value, "sources_and_uses")
    _exact_keys(
        schedule,
        required={
            "status",
            "currency",
            "balance_tolerance",
            "judgment_basis",
            "items",
        },
        label="sources_and_uses",
    )
    currency = _currency(schedule["currency"], "sources_and_uses.currency")
    items: list[dict[str, Any]] = []
    for index, raw in enumerate(_array(schedule["items"], "sources_and_uses.items")):
        item = _object(raw, f"sources_and_uses.items[{index}]")
        _exact_keys(
            item,
            required={
                "item_id",
                "side",
                "category",
                "description",
                "amount",
                "period",
                "evidence_refs",
                "judgment_basis",
            },
            optional={"currency"},
            label=f"sources_and_uses.items[{index}]",
        )
        supplied_currency = item.get("currency")
        if (
            supplied_currency is not None
            and _currency(
                supplied_currency,
                f"sources_and_uses.items[{index}].currency",
            )
            != currency
        ):
            raise ValueError("Sources-and-uses row currency differs from the schedule")
        items.append(
            {
                "item_id": _identifier(
                    item["item_id"], f"sources_and_uses.items[{index}].item_id"
                ),
                "side": _enum(
                    item["side"],
                    {"source", "use"},
                    f"sources_and_uses.items[{index}].side",
                ),
                "category": _text(
                    item["category"],
                    f"sources_and_uses.items[{index}].category",
                    max_length=300,
                ),
                "description": _text(
                    item["description"],
                    f"sources_and_uses.items[{index}].description",
                    max_length=2_000,
                ),
                "amount": _money(
                    item["amount"],
                    f"sources_and_uses.items[{index}].amount",
                    nonnegative=True,
                ),
                "period": _text(
                    item["period"],
                    f"sources_and_uses.items[{index}].period",
                    allow_empty=True,
                    max_length=100,
                ),
                "evidence_refs": _evidence_refs(
                    item["evidence_refs"],
                    known_sources=known_sources,
                    label=f"sources_and_uses.items[{index}].evidence_refs",
                ),
                "judgment_basis": _text(
                    item["judgment_basis"],
                    f"sources_and_uses.items[{index}].judgment_basis",
                ),
                "currency": currency,
            }
        )
    _unique_rows(items, id_field="item_id", label="Sources and uses")
    return {
        "status": _enum(
            schedule["status"], PERIMETER_STATUSES, "sources_and_uses.status"
        ),
        "currency": currency,
        "balance_tolerance": _money(
            schedule["balance_tolerance"],
            "sources_and_uses.balance_tolerance",
            nonnegative=True,
        ),
        "judgment_basis": _text(
            schedule["judgment_basis"], "sources_and_uses.judgment_basis"
        ),
        "items": sorted(items, key=lambda row: row["item_id"]),
    }


def _normalize_liquidity(
    value: object,
    *,
    known_sources: Mapping[str, str],
) -> dict[str, Any]:
    schedule = _object(value, "liquidity")
    _exact_keys(
        schedule,
        required={
            "status",
            "currency",
            "bridge_tolerance",
            "judgment_basis",
            "periods",
        },
        label="liquidity",
    )
    currency = _currency(schedule["currency"], "liquidity.currency")
    periods: list[dict[str, Any]] = []
    money_fields = (
        "opening_cash",
        "operating_inflows",
        "other_inflows",
        "new_finance_inflows",
        "operating_outflows",
        "procedure_costs",
        "creditor_distributions",
        "financing_outflows",
        "other_outflows",
        "reported_closing_cash",
    )
    nonnegative_fields = set(money_fields) - {"opening_cash", "reported_closing_cash"}
    for index, raw in enumerate(_array(schedule["periods"], "liquidity.periods")):
        item = _object(raw, f"liquidity.periods[{index}]")
        _exact_keys(
            item,
            required={
                "period_id",
                "period",
                *money_fields,
                "evidence_refs",
                "judgment_basis",
            },
            optional={"currency"},
            label=f"liquidity.periods[{index}]",
        )
        supplied_currency = item.get("currency")
        if (
            supplied_currency is not None
            and _currency(
                supplied_currency,
                f"liquidity.periods[{index}].currency",
            )
            != currency
        ):
            raise ValueError("Liquidity row currency differs from the schedule")
        normalized = {
            "period_id": _identifier(
                item["period_id"], f"liquidity.periods[{index}].period_id"
            ),
            "period": _text(
                item["period"], f"liquidity.periods[{index}].period", max_length=100
            ),
            "evidence_refs": _evidence_refs(
                item["evidence_refs"],
                known_sources=known_sources,
                label=f"liquidity.periods[{index}].evidence_refs",
            ),
            "judgment_basis": _text(
                item["judgment_basis"],
                f"liquidity.periods[{index}].judgment_basis",
            ),
            "currency": currency,
        }
        normalized.update(
            {
                field: _money(
                    item[field],
                    f"liquidity.periods[{index}].{field}",
                    nonnegative=field in nonnegative_fields,
                )
                for field in money_fields
            }
        )
        periods.append(normalized)
    _unique_rows(periods, id_field="period_id", label="Liquidity periods")
    return {
        "status": _enum(schedule["status"], PERIMETER_STATUSES, "liquidity.status"),
        "currency": currency,
        "bridge_tolerance": _money(
            schedule["bridge_tolerance"],
            "liquidity.bridge_tolerance",
            nonnegative=True,
        ),
        "judgment_basis": _text(schedule["judgment_basis"], "liquidity.judgment_basis"),
        "periods": sorted(periods, key=lambda row: row["period_id"]),
    }


def _normalize_milestones(
    value: object,
    *,
    known_sources: Mapping[str, str],
) -> list[dict[str, Any]]:
    milestones: list[dict[str, Any]] = []
    for index, raw in enumerate(_array(value, "milestones")):
        item = _object(raw, f"milestones[{index}]")
        _exact_keys(
            item,
            required={
                "milestone_id",
                "date_or_period",
                "description",
                "status",
                "evidence_refs",
                "judgment_basis",
            },
            label=f"milestones[{index}]",
        )
        milestones.append(
            {
                "milestone_id": _identifier(
                    item["milestone_id"], f"milestones[{index}].milestone_id"
                ),
                "date_or_period": _text(
                    item["date_or_period"],
                    f"milestones[{index}].date_or_period",
                    max_length=100,
                ),
                "description": _text(
                    item["description"],
                    f"milestones[{index}].description",
                    max_length=2_000,
                ),
                "status": _enum(
                    item["status"], MILESTONE_STATUSES, f"milestones[{index}].status"
                ),
                "evidence_refs": _evidence_refs(
                    item["evidence_refs"],
                    known_sources=known_sources,
                    label=f"milestones[{index}].evidence_refs",
                ),
                "judgment_basis": _text(
                    item["judgment_basis"], f"milestones[{index}].judgment_basis"
                ),
            }
        )
    _unique_rows(milestones, id_field="milestone_id", label="Milestones")
    return sorted(milestones, key=lambda row: row["milestone_id"])


def _normalize_review_questions(
    value: object,
    *,
    known_sources: Mapping[str, str],
) -> list[dict[str, Any]]:
    questions: list[dict[str, Any]] = []
    for index, raw in enumerate(_array(value, "review_questions")):
        item = _object(raw, f"review_questions[{index}]")
        _exact_keys(
            item,
            required={
                "question_id",
                "area",
                "question",
                "assessment",
                "evidence_refs",
                "judgment_basis",
                "follow_up",
            },
            label=f"review_questions[{index}]",
        )
        questions.append(
            {
                "question_id": _identifier(
                    item["question_id"], f"review_questions[{index}].question_id"
                ),
                "area": _enum(
                    item["area"], REVIEW_AREAS, f"review_questions[{index}].area"
                ),
                "question": _text(
                    item["question"],
                    f"review_questions[{index}].question",
                    max_length=2_000,
                ),
                "assessment": _enum(
                    item["assessment"],
                    QUESTION_ASSESSMENTS,
                    f"review_questions[{index}].assessment",
                ),
                "evidence_refs": _evidence_refs(
                    item["evidence_refs"],
                    known_sources=known_sources,
                    label=f"review_questions[{index}].evidence_refs",
                ),
                "judgment_basis": _text(
                    item["judgment_basis"],
                    f"review_questions[{index}].judgment_basis",
                ),
                "follow_up": _text(
                    item["follow_up"],
                    f"review_questions[{index}].follow_up",
                    allow_empty=True,
                    max_length=2_000,
                ),
            }
        )
    _unique_rows(questions, id_field="question_id", label="Review questions")
    covered = {row["area"] for row in questions}
    missing = sorted(set(REQUIRED_REVIEW_AREAS) - covered)
    if missing:
        raise ValueError(f"Review questions omit required semantic areas: {missing}")
    return sorted(questions, key=lambda row: row["question_id"])


def _normalize_assumptions(
    value: object,
    *,
    known_sources: Mapping[str, str],
) -> list[dict[str, Any]]:
    assumptions: list[dict[str, Any]] = []
    for index, raw in enumerate(_array(value, "assumptions")):
        item = _object(raw, f"assumptions[{index}]")
        _exact_keys(
            item,
            required={
                "assumption_id",
                "area",
                "statement",
                "status",
                "materiality",
                "evidence_refs",
                "judgment_basis",
            },
            label=f"assumptions[{index}]",
        )
        assumptions.append(
            {
                "assumption_id": _identifier(
                    item["assumption_id"], f"assumptions[{index}].assumption_id"
                ),
                "area": _enum(item["area"], REVIEW_AREAS, f"assumptions[{index}].area"),
                "statement": _text(
                    item["statement"],
                    f"assumptions[{index}].statement",
                    max_length=5_000,
                ),
                "status": _enum(
                    item["status"],
                    ASSUMPTION_STATUSES,
                    f"assumptions[{index}].status",
                ),
                "materiality": _enum(
                    item["materiality"],
                    SEVERITIES,
                    f"assumptions[{index}].materiality",
                ),
                "evidence_refs": _evidence_refs(
                    item["evidence_refs"],
                    known_sources=known_sources,
                    label=f"assumptions[{index}].evidence_refs",
                ),
                "judgment_basis": _text(
                    item["judgment_basis"], f"assumptions[{index}].judgment_basis"
                ),
            }
        )
    _unique_rows(assumptions, id_field="assumption_id", label="Assumptions")
    return sorted(assumptions, key=lambda row: row["assumption_id"])


def _normalize_issues(
    value: object,
    *,
    known_sources: Mapping[str, str],
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for index, raw in enumerate(_array(value, "issues")):
        item = _object(raw, f"issues[{index}]")
        _exact_keys(
            item,
            required={
                "issue_id",
                "area",
                "statement",
                "status",
                "severity",
                "evidence_refs",
                "owner",
                "next_action",
                "judgment_basis",
            },
            label=f"issues[{index}]",
        )
        issues.append(
            {
                "issue_id": _identifier(item["issue_id"], f"issues[{index}].issue_id"),
                "area": _enum(item["area"], REVIEW_AREAS, f"issues[{index}].area"),
                "statement": _text(
                    item["statement"], f"issues[{index}].statement", max_length=5_000
                ),
                "status": _enum(
                    item["status"], ISSUE_STATUSES, f"issues[{index}].status"
                ),
                "severity": _enum(
                    item["severity"], SEVERITIES, f"issues[{index}].severity"
                ),
                "evidence_refs": _evidence_refs(
                    item["evidence_refs"],
                    known_sources=known_sources,
                    label=f"issues[{index}].evidence_refs",
                ),
                "owner": _text(
                    item["owner"],
                    f"issues[{index}].owner",
                    allow_empty=True,
                    max_length=500,
                ),
                "next_action": _text(
                    item["next_action"],
                    f"issues[{index}].next_action",
                    allow_empty=True,
                    max_length=2_000,
                ),
                "judgment_basis": _text(
                    item["judgment_basis"], f"issues[{index}].judgment_basis"
                ),
            }
        )
    _unique_rows(issues, id_field="issue_id", label="Issues")
    return sorted(issues, key=lambda row: row["issue_id"])


def _normalize_case_model(
    inventory: Sequence[Mapping[str, Any]],
    case_model: Mapping[str, Any],
    *,
    reference_date: str = "",
) -> dict[str, Any]:
    """Normalize a reviewer-authored model without making semantic decisions."""

    _exact_keys(
        case_model,
        required={
            "schema_version",
            "legal_framework",
            "procedure",
            "document_perimeter",
            "creditor_population",
            "sources_and_uses",
            "liquidity",
            "milestones",
            "review_questions",
            "assumptions",
            "issues",
        },
        label="case_model",
    )
    if case_model["schema_version"] != CASE_SCHEMA_VERSION:
        raise ValueError("Unsupported Concordato Preventivo case-model schema")
    known_sources = _known_sources(inventory)
    if not known_sources:
        raise ValueError("A semantic model requires at least one captured source")

    framework = _object(case_model["legal_framework"], "legal_framework")
    _exact_keys(
        framework,
        required={
            "jurisdiction",
            "instrument",
            "framework_name",
            "as_of_date",
            "authority_refs",
            "judgment_basis",
        },
        label="legal_framework",
    )
    if _text(framework["jurisdiction"], "legal_framework.jurisdiction") != "IT":
        raise ValueError("This capability is scoped to the Italian jurisdiction")
    if (
        _text(framework["instrument"], "legal_framework.instrument")
        != "concordato_preventivo"
    ):
        raise ValueError("The semantic model must identify concordato_preventivo")
    authority_refs = _normalize_authority_refs(framework["authority_refs"])
    if not authority_refs:
        raise ValueError("Reviewed legal framework requires at least one authority")

    procedure = _object(case_model["procedure"], "procedure")
    _exact_keys(
        procedure,
        required={
            "identification_status",
            "debtor_name",
            "court",
            "procedure_reference",
            "stage",
            "plan_type",
            "reference_date",
            "currency",
            "judgment_basis",
        },
        label="procedure",
    )
    procedure_reference_date = _iso_date(
        procedure["reference_date"], "procedure.reference_date", allow_empty=False
    )
    if reference_date and procedure_reference_date != reference_date:
        raise ValueError("Semantic model reference date differs from the run")

    normalized = {
        "schema_version": CASE_SCHEMA_VERSION,
        "legal_framework": {
            "jurisdiction": "IT",
            "instrument": "concordato_preventivo",
            "framework_name": _text(
                framework["framework_name"], "legal_framework.framework_name"
            ),
            "as_of_date": _iso_date(
                framework["as_of_date"],
                "legal_framework.as_of_date",
                allow_empty=False,
            ),
            "authority_refs": authority_refs,
            "judgment_basis": _text(
                framework["judgment_basis"], "legal_framework.judgment_basis"
            ),
        },
        "procedure": {
            "identification_status": _enum(
                procedure["identification_status"],
                {"complete", "partial", "unclear"},
                "procedure.identification_status",
            ),
            "debtor_name": _text(
                procedure["debtor_name"],
                "procedure.debtor_name",
                allow_empty=True,
                max_length=500,
            ),
            "court": _text(
                procedure["court"],
                "procedure.court",
                allow_empty=True,
                max_length=500,
            ),
            "procedure_reference": _text(
                procedure["procedure_reference"],
                "procedure.procedure_reference",
                allow_empty=True,
                max_length=500,
            ),
            "stage": _enum(procedure["stage"], PROCEDURE_STAGES, "procedure.stage"),
            "plan_type": _enum(
                procedure["plan_type"], PLAN_TYPES, "procedure.plan_type"
            ),
            "reference_date": procedure_reference_date,
            "currency": _currency(procedure["currency"], "procedure.currency"),
            "judgment_basis": _text(
                procedure["judgment_basis"], "procedure.judgment_basis"
            ),
        },
        "document_perimeter": _normalize_document_perimeter(
            case_model["document_perimeter"], known_sources=known_sources
        ),
        "creditor_population": _normalize_creditors(
            case_model["creditor_population"], known_sources=known_sources
        ),
        "sources_and_uses": _normalize_sources_and_uses(
            case_model["sources_and_uses"], known_sources=known_sources
        ),
        "liquidity": _normalize_liquidity(
            case_model["liquidity"], known_sources=known_sources
        ),
        "milestones": _normalize_milestones(
            case_model["milestones"], known_sources=known_sources
        ),
        "review_questions": _normalize_review_questions(
            case_model["review_questions"], known_sources=known_sources
        ),
        "assumptions": _normalize_assumptions(
            case_model["assumptions"], known_sources=known_sources
        ),
        "issues": _normalize_issues(case_model["issues"], known_sources=known_sources),
    }
    currencies = {
        normalized["procedure"]["currency"],
        normalized["creditor_population"]["currency"],
        normalized["sources_and_uses"]["currency"],
        normalized["liquidity"]["currency"],
    }
    if len(currencies) != 1:
        raise ValueError("Procedure and mechanical schedules must use one currency")
    return normalized


def review_concordato_case_model(
    inventory: Sequence[Mapping[str, Any]],
    case_model: Mapping[str, Any],
    *,
    reviewer_ref: str,
    reviewed_on: str,
    reference_date: str = "",
) -> dict[str, Any]:
    """Seal a professional's semantic case model against captured sources."""

    normalized = _normalize_case_model(
        inventory, case_model, reference_date=reference_date
    )
    source_refs = [
        row["source_artifact_ref"]
        for row in normalized["document_perimeter"]["documents"]
    ]
    digest = canonical_json_sha256(normalized)
    decision = build_reviewed_decision_receipt(
        decision_id=f"decision.concordato_semantic_case.{digest}",
        decision_type=CASE_DECISION_TYPE,
        status="reviewed",
        reviewer_ref=_text(reviewer_ref, "reviewer_ref", max_length=500),
        reviewed_on=_text(reviewed_on, "reviewed_on", max_length=100),
        adapter_id=CASE_ADAPTER_ID,
        adapter_version=CASE_ADAPTER_VERSION,
        source_artifact_refs=source_refs,
        content=normalized,
    )
    return {
        "schema_version": CASE_RECIPE_SCHEMA_VERSION,
        "semantic_case_decision": decision,
    }


def _load_recipe(
    recipe: Path | Mapping[str, Any] | None,
) -> dict[str, Any] | None:
    if recipe is None:
        return None
    if isinstance(recipe, Mapping):
        return dict(recipe)
    payload = json.loads(recipe.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError("Semantic recipe must be a JSON object")
    return payload


def validate_semantic_recipe(
    inventory: Sequence[Mapping[str, Any]],
    recipe: Path | Mapping[str, Any] | None,
    *,
    reference_date: str = "",
) -> tuple[dict[str, Any] | None, dict[str, Any] | None]:
    """Validate a reviewed, source-bound semantic case recipe."""

    payload = _load_recipe(recipe)
    if payload is None:
        return None, None
    _exact_keys(
        payload,
        required={"schema_version", "semantic_case_decision"},
        label="semantic_recipe",
    )
    if payload["schema_version"] != CASE_RECIPE_SCHEMA_VERSION:
        raise ValueError("Unsupported semantic recipe schema")
    raw_decision = _object(
        payload["semantic_case_decision"], "semantic_recipe.semantic_case_decision"
    )
    decision = validate_reviewed_decision_receipt(
        raw_decision,
        expected_decision_type=CASE_DECISION_TYPE,
        expected_adapter_id=CASE_ADAPTER_ID,
        expected_adapter_version=CASE_ADAPTER_VERSION,
        require_reviewed=True,
    )
    content = _object(decision["content"], "semantic_case_decision.content")
    normalized = _normalize_case_model(
        inventory, content, reference_date=reference_date
    )
    if dict(content) != normalized:
        raise ValueError("Semantic case decision content is not canonical")
    expected_refs = [
        row["source_artifact_ref"]
        for row in normalized["document_perimeter"]["documents"]
    ]
    if list(decision["source_artifact_refs"]) != expected_refs:
        raise ValueError("Semantic case decision source perimeter is stale")
    return dict(decision), normalized


def _decimal(value: object, label: str) -> Decimal:
    if not isinstance(value, str):
        raise ValueError(f"{label} must be canonical decimal text")
    try:
        return parse_canonical_decimal(value, label=label)
    except (MoneyValidationError, ValueError) as exc:
        raise ValueError(str(exc)) from exc


def _percentage(numerator: Decimal, denominator: Decimal) -> str:
    if denominator == 0:
        return ""
    with localcontext() as context:
        context.prec = DECIMAL_PRECISION
        value = (numerator / denominator * Decimal("100")).quantize(Decimal("0.0001"))
    return decimal_text(value)


def derive_case_schedules(case_model: Mapping[str, Any]) -> dict[str, Any]:
    """Compute exact schedules from an already reviewed semantic model."""

    creditor_rows: list[dict[str, Any]] = []
    class_totals: dict[tuple[str, str], dict[str, Decimal | int]] = defaultdict(
        lambda: {
            "creditor_count": 0,
            "claim_amount": Decimal("0"),
            "proposed_cash_amount": Decimal("0"),
            "proposed_non_cash_amount": Decimal("0"),
            "liquidation_recovery_amount": Decimal("0"),
        }
    )
    for creditor in case_model["creditor_population"]["creditors"]:
        claim = _decimal(creditor["claim_amount"], "creditor claim")
        plan_cash = _decimal(creditor["proposed_cash_amount"], "plan cash")
        plan_non_cash = _decimal(creditor["proposed_non_cash_amount"], "plan non-cash")
        liquidation = _decimal(
            creditor["liquidation_recovery_amount"], "liquidation recovery"
        )
        plan_total = plan_cash + plan_non_cash
        creditor_rows.append(
            {
                **dict(creditor),
                "proposed_total_amount": decimal_text(plan_total),
                "proposed_recovery_pct": _percentage(plan_total, claim),
                "liquidation_recovery_pct": _percentage(liquidation, claim),
                "plan_vs_liquidation_delta": decimal_text(plan_total - liquidation),
                "evidence_ref_count": len(creditor["evidence_refs"]),
            }
        )
        aggregate = class_totals[(creditor["class_id"], creditor["priority"])]
        aggregate["creditor_count"] = int(aggregate["creditor_count"]) + 1
        aggregate["claim_amount"] = Decimal(aggregate["claim_amount"]) + claim
        aggregate["proposed_cash_amount"] = (
            Decimal(aggregate["proposed_cash_amount"]) + plan_cash
        )
        aggregate["proposed_non_cash_amount"] = (
            Decimal(aggregate["proposed_non_cash_amount"]) + plan_non_cash
        )
        aggregate["liquidation_recovery_amount"] = (
            Decimal(aggregate["liquidation_recovery_amount"]) + liquidation
        )

    class_rows: list[dict[str, Any]] = []
    for (class_id, priority), aggregate in sorted(class_totals.items()):
        claim = Decimal(aggregate["claim_amount"])
        plan_cash = Decimal(aggregate["proposed_cash_amount"])
        plan_non_cash = Decimal(aggregate["proposed_non_cash_amount"])
        liquidation = Decimal(aggregate["liquidation_recovery_amount"])
        plan_total = plan_cash + plan_non_cash
        class_rows.append(
            {
                "class_id": class_id,
                "priority": priority,
                "creditor_count": int(aggregate["creditor_count"]),
                "claim_amount": decimal_text(claim),
                "proposed_cash_amount": decimal_text(plan_cash),
                "proposed_non_cash_amount": decimal_text(plan_non_cash),
                "proposed_total_amount": decimal_text(plan_total),
                "proposed_recovery_pct": _percentage(plan_total, claim),
                "liquidation_recovery_amount": decimal_text(liquidation),
                "liquidation_recovery_pct": _percentage(liquidation, claim),
                "plan_vs_liquidation_delta": decimal_text(plan_total - liquidation),
            }
        )

    source_use_rows = [dict(row) for row in case_model["sources_and_uses"]["items"]]
    source_total = sum(
        (
            _decimal(row["amount"], "source amount")
            for row in source_use_rows
            if row["side"] == "source"
        ),
        Decimal("0"),
    )
    use_total = sum(
        (
            _decimal(row["amount"], "use amount")
            for row in source_use_rows
            if row["side"] == "use"
        ),
        Decimal("0"),
    )
    surplus_shortfall = source_total - use_total
    balance_tolerance = _decimal(
        case_model["sources_and_uses"]["balance_tolerance"],
        "sources and uses balance tolerance",
    )

    liquidity_rows: list[dict[str, Any]] = []
    bridge_tolerance = _decimal(
        case_model["liquidity"]["bridge_tolerance"],
        "liquidity bridge tolerance",
    )
    for row in case_model["liquidity"]["periods"]:
        opening = _decimal(row["opening_cash"], "opening cash")
        inflows = sum(
            (
                _decimal(row[field], field)
                for field in (
                    "operating_inflows",
                    "other_inflows",
                    "new_finance_inflows",
                )
            ),
            Decimal("0"),
        )
        outflows = sum(
            (
                _decimal(row[field], field)
                for field in (
                    "operating_outflows",
                    "procedure_costs",
                    "creditor_distributions",
                    "financing_outflows",
                    "other_outflows",
                )
            ),
            Decimal("0"),
        )
        calculated_closing = opening + inflows - outflows
        reported_closing = _decimal(
            row["reported_closing_cash"], "reported closing cash"
        )
        bridge_difference = reported_closing - calculated_closing
        liquidity_rows.append(
            {
                **dict(row),
                "total_inflows": decimal_text(inflows),
                "total_outflows": decimal_text(outflows),
                "calculated_closing_cash": decimal_text(calculated_closing),
                "bridge_difference": decimal_text(bridge_difference),
                "bridge_within_tolerance": abs(bridge_difference) <= bridge_tolerance,
                "evidence_ref_count": len(row["evidence_refs"]),
            }
        )

    total_claim = sum(
        (_decimal(row["claim_amount"], "claim amount") for row in creditor_rows),
        Decimal("0"),
    )
    total_plan_cash = sum(
        (_decimal(row["proposed_cash_amount"], "plan cash") for row in creditor_rows),
        Decimal("0"),
    )
    total_plan_non_cash = sum(
        (
            _decimal(row["proposed_non_cash_amount"], "plan non-cash")
            for row in creditor_rows
        ),
        Decimal("0"),
    )
    total_liquidation = sum(
        (
            _decimal(row["liquidation_recovery_amount"], "liquidation recovery")
            for row in creditor_rows
        ),
        Decimal("0"),
    )
    total_plan = total_plan_cash + total_plan_non_cash
    minimum_cash = (
        min(
            _decimal(row["reported_closing_cash"], "reported closing cash")
            for row in liquidity_rows
        )
        if liquidity_rows
        else None
    )
    question_areas = {row["area"] for row in case_model["review_questions"]}
    open_material_issues = [
        row
        for row in case_model["issues"]
        if row["status"] == "open" and row["severity"] in {"critical", "high"}
    ]
    checks = [
        {
            "check_id": "document_perimeter_status",
            "status": (
                "passed"
                if case_model["document_perimeter"]["status"] == "complete"
                else "attention"
            ),
            "observation": case_model["document_perimeter"]["status"],
            "mechanical_scope": "Reports the reviewer-assigned perimeter status.",
        },
        {
            "check_id": "creditor_population_status",
            "status": (
                "passed"
                if case_model["creditor_population"]["status"] == "complete"
                else "attention"
            ),
            "observation": case_model["creditor_population"]["status"],
            "mechanical_scope": "Reports the reviewer-assigned population status.",
        },
        {
            "check_id": "review_question_coverage",
            "status": (
                "passed"
                if set(REQUIRED_REVIEW_AREAS) <= question_areas
                else "attention"
            ),
            "observation": (
                f"{len(question_areas)} areas represented; "
                f"{len(REQUIRED_REVIEW_AREAS)} required."
            ),
            "mechanical_scope": "Checks presence only, not adequacy of assessments.",
        },
        {
            "check_id": "sources_and_uses_balance",
            "status": (
                "not_assessed"
                if case_model["sources_and_uses"]["status"] == "missing"
                else (
                    "passed"
                    if abs(surplus_shortfall) <= balance_tolerance
                    else "attention"
                )
            ),
            "observation": decimal_text(surplus_shortfall),
            "mechanical_scope": "Sources minus uses, compared with reviewed tolerance.",
        },
        {
            "check_id": "liquidity_bridge",
            "status": (
                "not_assessed"
                if not liquidity_rows
                else (
                    "passed"
                    if all(row["bridge_within_tolerance"] for row in liquidity_rows)
                    else "attention"
                )
            ),
            "observation": (
                f"{sum(1 for row in liquidity_rows if not row['bridge_within_tolerance'])} "
                "periods outside tolerance"
            ),
            "mechanical_scope": "Opening cash plus inflows minus outflows.",
        },
        {
            "check_id": "open_material_issues",
            "status": "attention" if open_material_issues else "passed",
            "observation": f"{len(open_material_issues)} open critical/high issues",
            "mechanical_scope": "Counts reviewer-classified issues without reclassifying them.",
        },
    ]
    for check in checks:
        check["limitation"] = (
            "This mechanical observation is not a legal, tax, feasibility, "
            "materiality, or evidence-sufficiency conclusion."
        )
    return {
        "schema_version": "concordato.preventivo.derived_schedules.v1",
        "currency": case_model["procedure"]["currency"],
        "summary": {
            "creditor_count": len(creditor_rows),
            "class_count": len(class_rows),
            "total_claim_amount": decimal_text(total_claim),
            "total_proposed_cash_amount": decimal_text(total_plan_cash),
            "total_proposed_non_cash_amount": decimal_text(total_plan_non_cash),
            "total_proposed_recovery_amount": decimal_text(total_plan),
            "total_proposed_recovery_pct": _percentage(total_plan, total_claim),
            "total_liquidation_recovery_amount": decimal_text(total_liquidation),
            "total_liquidation_recovery_pct": _percentage(
                total_liquidation, total_claim
            ),
            "plan_vs_liquidation_delta": decimal_text(total_plan - total_liquidation),
            "source_total": decimal_text(source_total),
            "use_total": decimal_text(use_total),
            "surplus_shortfall": decimal_text(surplus_shortfall),
            "funding_gap": decimal_text(use_total - source_total),
            "minimum_reported_closing_cash": (
                decimal_text(minimum_cash) if minimum_cash is not None else ""
            ),
            "open_issue_count": sum(
                1 for row in case_model["issues"] if row["status"] == "open"
            ),
            "open_critical_high_issue_count": len(open_material_issues),
        },
        "creditors": creditor_rows,
        "classes": class_rows,
        "sources_and_uses": source_use_rows,
        "liquidity": liquidity_rows,
        "checks": checks,
    }


def _write_json(path: Path, payload: object) -> None:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _write_csv(
    path: Path,
    rows: Sequence[Mapping[str, Any]],
    fieldnames: Sequence[str],
) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(
            handle, fieldnames=list(fieldnames), extrasaction="ignore"
        )
        writer.writeheader()
        writer.writerows(rows)


def _cell_value(value: object) -> object:
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return value


def _add_sheet(
    workbook: openpyxl.Workbook,
    *,
    title: str,
    rows: Sequence[Mapping[str, Any]],
    headers: Sequence[str],
    no_rows: str,
) -> None:
    sheet = workbook.create_sheet(title=title)
    sheet.append(list(headers))
    if rows:
        for row in rows:
            sheet.append([_cell_value(row.get(header, "")) for header in headers])
    else:
        sheet.append([no_rows, *("" for _ in headers[1:])])
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="163A5F")
        cell.alignment = Alignment(vertical="top", wrap_text=True)
    for column in sheet.columns:
        max_length = max(len(str(cell.value or "")) for cell in column)
        sheet.column_dimensions[column[0].column_letter].width = min(
            max(max_length + 2, 12), 48
        )
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions


def _write_workbook(
    path: Path,
    *,
    status: str,
    case_model: Mapping[str, Any] | None,
    derived: Mapping[str, Any],
    matches: Sequence[Mapping[str, Any]],
    language: str,
) -> None:
    text = TEXT[_language(language)]
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)
    overview = (
        [
            {"metric": "semantic_model_status", "value": status},
            *[
                {"metric": key, "value": value}
                for key, value in derived.get("summary", {}).items()
            ],
        ]
        if case_model is not None
        else [{"metric": "semantic_model_status", "value": status}]
    )
    _add_sheet(
        workbook,
        title="Overview",
        rows=overview,
        headers=("metric", "value"),
        no_rows=text["no_rows"],
    )
    documents = (
        case_model["document_perimeter"]["documents"] if case_model is not None else []
    )
    _add_sheet(
        workbook,
        title="Documents",
        rows=documents,
        headers=(
            "relative_path",
            "roles",
            "authoritative_for",
            "version_date",
            "judgment_basis",
            "source_artifact_ref",
        ),
        no_rows=text["no_rows"],
    )
    _add_sheet(
        workbook,
        title="Creditors",
        rows=derived.get("creditors", []),
        headers=(
            "creditor_id",
            "creditor_name",
            "claim_amount",
            "claim_status",
            "priority",
            "class_id",
            "treatment_form",
            "proposed_cash_amount",
            "proposed_non_cash_amount",
            "proposed_total_amount",
            "proposed_recovery_pct",
            "liquidation_recovery_amount",
            "liquidation_recovery_pct",
            "plan_vs_liquidation_delta",
            "payment_start",
            "payment_end",
            "voting_treatment",
            "judgment_basis",
            "evidence_ref_count",
        ),
        no_rows=text["no_rows"],
    )
    _add_sheet(
        workbook,
        title="Classes",
        rows=derived.get("classes", []),
        headers=(
            "class_id",
            "priority",
            "creditor_count",
            "claim_amount",
            "proposed_cash_amount",
            "proposed_non_cash_amount",
            "proposed_total_amount",
            "proposed_recovery_pct",
            "liquidation_recovery_amount",
            "liquidation_recovery_pct",
            "plan_vs_liquidation_delta",
        ),
        no_rows=text["no_rows"],
    )
    _add_sheet(
        workbook,
        title="Sources Uses",
        rows=derived.get("sources_and_uses", []),
        headers=(
            "item_id",
            "side",
            "category",
            "description",
            "amount",
            "period",
            "currency",
            "judgment_basis",
            "evidence_refs",
        ),
        no_rows=text["no_rows"],
    )
    _add_sheet(
        workbook,
        title="Liquidity",
        rows=derived.get("liquidity", []),
        headers=(
            "period_id",
            "period",
            "opening_cash",
            "total_inflows",
            "total_outflows",
            "reported_closing_cash",
            "calculated_closing_cash",
            "bridge_difference",
            "bridge_within_tolerance",
            "judgment_basis",
            "evidence_ref_count",
        ),
        no_rows=text["no_rows"],
    )
    _add_sheet(
        workbook,
        title="Review Questions",
        rows=case_model["review_questions"] if case_model is not None else [],
        headers=(
            "question_id",
            "area",
            "question",
            "assessment",
            "judgment_basis",
            "follow_up",
            "evidence_refs",
        ),
        no_rows=text["no_rows"],
    )
    _add_sheet(
        workbook,
        title="Issues",
        rows=case_model["issues"] if case_model is not None else [],
        headers=(
            "issue_id",
            "area",
            "statement",
            "status",
            "severity",
            "owner",
            "next_action",
            "judgment_basis",
            "evidence_refs",
        ),
        no_rows=text["no_rows"],
    )
    _add_sheet(
        workbook,
        title="Mechanical Checks",
        rows=derived.get("checks", []),
        headers=(
            "check_id",
            "status",
            "observation",
            "mechanical_scope",
            "limitation",
        ),
        no_rows=text["no_rows"],
    )
    _add_sheet(
        workbook,
        title="Numeric Tie-Out",
        rows=matches,
        headers=(
            "plan_source_file",
            "plan_location",
            "plan_amount",
            "support_source_file",
            "support_location",
            "support_amount",
            "difference",
            "match_status",
        ),
        no_rows=text["no_rows"],
    )
    workbook.save(path)


def _write_markdown(
    path: Path,
    *,
    status: str,
    case_model: Mapping[str, Any] | None,
    derived: Mapping[str, Any],
    language: str,
    match_count: int,
) -> None:
    text = TEXT[_language(language)]
    lines = [
        f"# {text['title']}",
        "",
        f"## {text['status']}",
        "",
        f"- {text['reviewed'] if status == 'reviewed' else text['withheld']}",
        f"- {text['limitation']}",
    ]
    if case_model is not None:
        procedure = case_model["procedure"]
        lines.extend(
            [
                "",
                f"## {text['procedure']}",
                "",
                f"- Debtor: {procedure['debtor_name'] or '—'}",
                f"- Court: {procedure['court'] or '—'}",
                f"- Procedure reference: {procedure['procedure_reference'] or '—'}",
                f"- Stage: `{procedure['stage']}`",
                f"- Plan type: `{procedure['plan_type']}`",
                f"- Reference date: `{procedure['reference_date']}`",
                "",
                f"## {text['deterministic']}",
                "",
            ]
        )
        for key, value in derived["summary"].items():
            lines.append(f"- `{key}`: {value}")
        lines.extend(["", f"## {text['questions']}", ""])
        for row in case_model["review_questions"]:
            lines.append(
                f"- **{row['area']} — {row['assessment']}**: {row['question']}"
            )
            if row["follow_up"]:
                lines.append(f"  - Follow-up: {row['follow_up']}")
        lines.extend(["", f"## {text['issues']}", ""])
        open_issues = [
            row for row in case_model["issues"] if row["status"] != "resolved"
        ]
        if open_issues:
            for row in open_issues:
                lines.append(
                    f"- **{row['severity']} · {row['area']}**: {row['statement']}"
                )
        else:
            lines.append(f"- {text['no_rows']}")
    lines.extend(
        [
            "",
            f"## {text['numeric_appendix']}",
            "",
            f"- Candidate amount matches: {match_count}",
            "- Equal amounts remain candidate evidence and require contextual review.",
            "",
        ]
    )
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_summary_docx(
    path: Path,
    *,
    status: str,
    case_model: Mapping[str, Any] | None,
    derived: Mapping[str, Any],
    language: str,
    candidate_count: int,
    match_count: int,
) -> None:
    text = TEXT[_language(language)]
    document = Document()
    document.add_heading(text["title"], level=0)
    document.add_paragraph(text["limitation"])
    document.add_heading(text["status"], level=1)
    document.add_paragraph(
        text["reviewed"] if status == "reviewed" else text["withheld"]
    )
    if case_model is not None:
        procedure = case_model["procedure"]
        document.add_heading(text["procedure"], level=1)
        procedure_table = document.add_table(rows=0, cols=2)
        procedure_table.style = "Table Grid"
        for label, value in (
            ("Debtor", procedure["debtor_name"]),
            ("Court", procedure["court"]),
            ("Procedure reference", procedure["procedure_reference"]),
            ("Stage", procedure["stage"]),
            ("Plan type", procedure["plan_type"]),
            ("Reference date", procedure["reference_date"]),
        ):
            cells = procedure_table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value or "—")
        document.add_heading(text["deterministic"], level=1)
        summary_table = document.add_table(rows=1, cols=2)
        summary_table.style = "Table Grid"
        summary_table.rows[0].cells[0].text = "Metric"
        summary_table.rows[0].cells[1].text = "Value"
        for key, value in derived["summary"].items():
            cells = summary_table.add_row().cells
            cells[0].text = key
            cells[1].text = str(value)
        document.add_heading(text["questions"], level=1)
        question_table = document.add_table(rows=1, cols=4)
        question_table.style = "Table Grid"
        for cell, value in zip(
            question_table.rows[0].cells,
            ("Area", "Assessment", "Question", "Follow-up"),
            strict=True,
        ):
            cell.text = value
        for row in case_model["review_questions"]:
            cells = question_table.add_row().cells
            cells[0].text = row["area"]
            cells[1].text = row["assessment"]
            cells[2].text = row["question"]
            cells[3].text = row["follow_up"]
        document.add_heading(text["issues"], level=1)
        issue_table = document.add_table(rows=1, cols=5)
        issue_table.style = "Table Grid"
        for cell, value in zip(
            issue_table.rows[0].cells,
            ("Severity", "Area", "Status", "Issue", "Next action"),
            strict=True,
        ):
            cell.text = value
        for row in case_model["issues"]:
            cells = issue_table.add_row().cells
            cells[0].text = row["severity"]
            cells[1].text = row["area"]
            cells[2].text = row["status"]
            cells[3].text = row["statement"]
            cells[4].text = row["next_action"]
    document.add_heading(text["numeric_appendix"], level=1)
    document.add_paragraph(f"Candidate plan amounts: {candidate_count}")
    document.add_paragraph(f"Candidate amount matches: {match_count}")
    document.add_paragraph(
        "Equal amounts are candidate evidence only; contextual support remains "
        "professional judgment."
    )
    document.save(path)


def write_semantic_artifacts(
    output_dir: Path,
    inventory: Sequence[Mapping[str, Any]],
    *,
    semantic_decision: Mapping[str, Any] | None,
    case_model: Mapping[str, Any] | None,
    reference_date: str,
    language: str,
    matches: Sequence[Mapping[str, Any]],
    candidate_count: int,
    error: str | None = None,
) -> SemanticArtifacts:
    """Write primary semantic outputs and the subordinate numeric appendix."""

    output_dir.mkdir(parents=True, exist_ok=True)
    template = build_case_model_template(
        inventory, reference_date=reference_date, language=language
    )
    _write_json(output_dir / "suggested_concordato_case_model.json", template)
    status = (
        "reviewed"
        if semantic_decision is not None and case_model is not None
        else ("invalid_semantic_recipe" if error else "needs_semantic_review")
    )
    derived = (
        derive_case_schedules(case_model)
        if case_model is not None
        else {
            "schema_version": "concordato.preventivo.derived_schedules.v1",
            "currency": "EUR",
            "summary": {},
            "creditors": [],
            "classes": [],
            "sources_and_uses": [],
            "liquidity": [],
            "checks": [
                {
                    "check_id": "semantic_model_reviewed",
                    "status": "attention",
                    "observation": error or "No reviewed semantic case model supplied.",
                    "mechanical_scope": "Requires a reviewed, source-bound case decision.",
                    "limitation": (
                        "No legal, tax, feasibility, materiality, or "
                        "evidence-sufficiency conclusion was made."
                    ),
                }
            ],
        }
    )
    case_output = {
        "schema_version": CASE_OUTPUT_SCHEMA_VERSION,
        "status": status,
        "semantic_case_decision_ref": (
            semantic_decision.get("decision_id")
            if semantic_decision is not None
            else None
        ),
        "semantic_case_content_sha256": (
            semantic_decision.get("content_sha256")
            if semantic_decision is not None
            else None
        ),
        "case_model": dict(case_model) if case_model is not None else None,
        "error": error,
        "limitations": [
            "The case model records reviewed professional judgments; code did not infer them.",
            "Mechanical checks are not legal, tax, feasibility, materiality, or evidence-sufficiency conclusions.",
            "Reviewer identity is recorded but not cryptographically authenticated.",
        ],
    }
    _write_json(output_dir / "concordato_case_model.json", case_output)
    _write_json(
        output_dir / "concordato_semantic_checks.json",
        {
            "schema_version": "concordato.preventivo.semantic_checks.v1",
            "status": status,
            "checks": derived["checks"],
            "summary": derived["summary"],
        },
    )
    _write_csv(
        output_dir / "creditor_treatment.csv",
        derived["creditors"],
        (
            "creditor_id",
            "creditor_name",
            "claim_amount",
            "claim_status",
            "priority",
            "class_id",
            "treatment_form",
            "proposed_cash_amount",
            "proposed_non_cash_amount",
            "proposed_total_amount",
            "proposed_recovery_pct",
            "liquidation_recovery_amount",
            "liquidation_recovery_pct",
            "plan_vs_liquidation_delta",
            "payment_start",
            "payment_end",
            "voting_treatment",
            "currency",
            "evidence_ref_count",
            "judgment_basis",
        ),
    )
    _write_csv(
        output_dir / "creditor_class_summary.csv",
        derived["classes"],
        (
            "class_id",
            "priority",
            "creditor_count",
            "claim_amount",
            "proposed_cash_amount",
            "proposed_non_cash_amount",
            "proposed_total_amount",
            "proposed_recovery_pct",
            "liquidation_recovery_amount",
            "liquidation_recovery_pct",
            "plan_vs_liquidation_delta",
        ),
    )
    _write_csv(
        output_dir / "sources_and_uses.csv",
        derived["sources_and_uses"],
        (
            "item_id",
            "side",
            "category",
            "description",
            "amount",
            "period",
            "currency",
            "judgment_basis",
        ),
    )
    _write_csv(
        output_dir / "liquidity_schedule.csv",
        derived["liquidity"],
        (
            "period_id",
            "period",
            "opening_cash",
            "operating_inflows",
            "other_inflows",
            "new_finance_inflows",
            "total_inflows",
            "operating_outflows",
            "procedure_costs",
            "creditor_distributions",
            "financing_outflows",
            "other_outflows",
            "total_outflows",
            "reported_closing_cash",
            "calculated_closing_cash",
            "bridge_difference",
            "bridge_within_tolerance",
            "currency",
            "judgment_basis",
        ),
    )
    _write_workbook(
        output_dir / "concordato_review_workpaper.xlsx",
        status=status,
        case_model=case_model,
        derived=derived,
        matches=matches,
        language=language,
    )
    _write_markdown(
        output_dir / "concordato_semantic_review.md",
        status=status,
        case_model=case_model,
        derived=derived,
        language=language,
        match_count=len(matches),
    )
    _write_summary_docx(
        output_dir / "concordato_preventivo_review_summary.docx",
        status=status,
        case_model=case_model,
        derived=derived,
        language=language,
        candidate_count=candidate_count,
        match_count=len(matches),
    )
    artifact_paths = (
        "suggested_concordato_case_model.json",
        "concordato_case_model.json",
        "concordato_semantic_checks.json",
        "creditor_treatment.csv",
        "creditor_class_summary.csv",
        "sources_and_uses.csv",
        "liquidity_schedule.csv",
        "concordato_review_workpaper.xlsx",
        "concordato_semantic_review.md",
        "concordato_preventivo_review_summary.docx",
    )
    return SemanticArtifacts(
        status=status,
        decision=dict(semantic_decision) if semantic_decision is not None else None,
        case_model=dict(case_model) if case_model is not None else None,
        derived=derived,
        artifact_paths=artifact_paths,
        error=error,
    )
