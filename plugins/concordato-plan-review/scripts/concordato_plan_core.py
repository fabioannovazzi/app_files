from __future__ import annotations

import csv
import json
import logging
import re
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

import openpyxl
from docx import Document
from pypdf import PdfReader

SCRIPT_DIR = Path(__file__).resolve().parent


def _ensure_local_review_session_import() -> None:
    """Use this plugin's review-session module in multi-plugin test runs."""

    script_dir = str(SCRIPT_DIR)
    if script_dir in sys.path:
        sys.path.remove(script_dir)
    sys.path.insert(0, script_dir)
    module = sys.modules.get("review_session")
    module_file = getattr(module, "__file__", None) if module is not None else None
    if module_file and Path(module_file).resolve().is_relative_to(SCRIPT_DIR.resolve()):
        return
    if module is not None:
        del sys.modules["review_session"]


_ensure_local_review_session_import()
from review_session import write_review_session_artifacts, write_run_intake

LOGGER = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".pdf", ".xlsx", ".xlsm", ".csv"}
TEXT_SUFFIXES = {".csv", ".txt", ".md"}
PDF_SUFFIXES = {".pdf"}
WORKBOOK_SUFFIXES = {".xlsx", ".xlsm"}
SUPPORTED_LANGUAGES = ("it", "en", "fr", "de", "es")
STOPWORDS = {
    "al",
    "alla",
    "and",
    "con",
    "da",
    "dal",
    "del",
    "della",
    "di",
    "e",
    "for",
    "il",
    "in",
    "la",
    "of",
    "per",
    "the",
    "to",
}
ROLE_RULES: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("debt_tax_social_security_detail", ("debiti tributari", "previdenziali")),
    ("adjusted_db", ("db_31", "db 31", "db_31.03", "rettificat")),
    ("trial_balance", ("bilancio", "bive", "situazione")),
    ("general_ledger", ("mastrini", "mastrino", "mastro")),
    ("concordato_plan", ("piano cp", "piano", "concordato")),
)
DATE_LIKE_RE = re.compile(r"^\d{1,2}[./-]\d{1,2}[./-]\d{2,4}$")
AMOUNT_TOKEN_RE = re.compile(
    r"(?<![\w/])-?\(?€\s*\d+(?:[.,]\d{1,2})?\)?(?![\w/])"
    r"|(?<![\w/])-?\(?\d{1,3}(?:[.\s]\d{3})+(?:,\d{1,2})?\)?(?![\w/])"
    r"|(?<![\w/])-?\(?\d+(?:[.,]\d{2})\)?(?![\w/])"
)
WORD_RE = re.compile(r"[a-z0-9]+")

__all__ = [
    "AmountCandidate",
    "ReviewRun",
    "classify_source_role",
    "extract_amount_candidates_from_text",
    "find_exact_amount_matches",
    "normalize_language",
    "parse_amount_token",
    "run_concordato_review",
    "write_json",
]


@dataclass(frozen=True)
class AmountCandidate:
    """Mechanically extracted numeric candidate from a source document."""

    source_file: str
    source_role: str
    location: str
    amount: float
    token: str
    context: str


@dataclass(frozen=True)
class ReviewRun:
    """Output paths and high-level counts for one review run."""

    output_dir: Path
    inventory: list[dict[str, Any]]
    amount_candidates: list[AmountCandidate]
    exact_matches: list[dict[str, Any]]
    audit: dict[str, Any]


def configure_logging(verbose: bool = False) -> None:
    """Configure script logging without changing imported callers."""

    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(level=level, format="%(levelname)s: %(message)s")


def normalize_language(
    language: object | None,
    *,
    default: str = "en",
    allow_auto: bool = False,
) -> str:
    """Normalize a language tag to a supported plugin locale."""

    text = str(language or default).strip().lower().replace("_", "-")
    code = text.split("-", 1)[0]
    if allow_auto and code == "auto":
        return "auto"
    return code if code in SUPPORTED_LANGUAGES else default


def write_json(path: Path, payload: Any) -> None:
    """Write JSON with stable UTF-8 formatting."""

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def _clean_text(value: object | None) -> str:
    if value is None:
        return ""
    return str(value).replace("\u00a0", " ").replace("\u202f", " ").strip()


def _normalize_text(value: object | None) -> str:
    text = unicodedata.normalize("NFKD", _clean_text(value))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower().replace("_", " ")
    return re.sub(r"\s+", " ", text)


def _tokens(value: object | None) -> set[str]:
    return {
        token
        for token in WORD_RE.findall(_normalize_text(value))
        if len(token) > 2 and token not in STOPWORDS and not token.isdigit()
    }


def classify_source_role(path: Path) -> str:
    """Suggest a source role from stable filename cues for intake only."""

    text = _normalize_text(path.stem)
    for role, needles in ROLE_RULES:
        if any(needle in text for needle in needles):
            return role
    return "unclassified"


def parse_amount_token(token: str) -> float | None:
    """Parse common Italian/European amount tokens into floats."""

    cleaned = (
        token.replace("€", "").replace("\u00a0", " ").replace("\u202f", " ").strip()
    )
    if not cleaned or DATE_LIKE_RE.match(cleaned):
        return None
    is_negative = cleaned.startswith("-") or (
        cleaned.startswith("(") and cleaned.endswith(")")
    )
    cleaned = cleaned.strip("-() ").replace(" ", "")
    if not re.search(r"\d", cleaned):
        return None

    comma_pos = cleaned.rfind(",")
    dot_pos = cleaned.rfind(".")
    if comma_pos >= 0 and dot_pos >= 0:
        if comma_pos > dot_pos:
            cleaned = cleaned.replace(".", "").replace(",", ".")
        else:
            cleaned = cleaned.replace(",", "")
    elif comma_pos >= 0:
        whole, _, decimals = cleaned.partition(",")
        cleaned = whole.replace(".", "") + ("." + decimals if decimals else "")
    elif dot_pos >= 0:
        parts = cleaned.split(".")
        if len(parts[-1]) == 3 and all(len(part) == 3 for part in parts[1:]):
            cleaned = "".join(parts)

    try:
        amount = float(cleaned)
    except ValueError:
        return None
    if is_negative:
        amount = -amount
    return amount


def _compact_context(text: str, start: int, end: int, window: int = 180) -> str:
    context_start = max(0, start - window // 2)
    context_end = min(len(text), end + window // 2)
    context = re.sub(r"\s+", " ", text[context_start:context_end]).strip()
    return context[:window]


def extract_amount_candidates_from_text(
    text: str,
    *,
    source_file: str,
    source_role: str,
    location: str,
) -> list[AmountCandidate]:
    """Extract numeric candidates without deciding whether they are material."""

    candidates: list[AmountCandidate] = []
    for match in AMOUNT_TOKEN_RE.finditer(text):
        token = match.group(0)
        amount = parse_amount_token(token)
        if amount is None:
            continue
        if abs(amount) < 0.005:
            continue
        candidates.append(
            AmountCandidate(
                source_file=source_file,
                source_role=source_role,
                location=location,
                amount=round(amount, 2),
                token=token.strip(),
                context=_compact_context(text, match.start(), match.end()),
            )
        )
    return candidates


def _file_inventory(input_dir: Path) -> list[dict[str, Any]]:
    files = []
    for path in sorted(input_dir.rglob("*")):
        if not path.is_file() or path.name == ".DS_Store":
            continue
        suffix = path.suffix.lower()
        files.append(
            {
                "path": str(path),
                "relative_path": path.relative_to(input_dir).as_posix(),
                "name": path.name,
                "suffix": suffix,
                "size_bytes": path.stat().st_size,
                "supported": suffix in SUPPORTED_SUFFIXES,
                "suggested_role": classify_source_role(path),
            }
        )
    return files


def _read_pdf_pages(path: Path) -> list[dict[str, Any]]:
    reader = PdfReader(str(path))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(
            {
                "source_file": path.name,
                "page": page_number,
                "method": "pdf_text",
                "text_length": len(text),
                "line_count": len(text.splitlines()),
                "text": text,
            }
        )
    return pages


def _excel_column_name(index: int) -> str:
    letters: list[str] = []
    idx = index
    while idx:
        idx, rem = divmod(idx - 1, 26)
        letters.append(chr(65 + rem))
    return "".join(reversed(letters))


def _workbook_candidates(
    path: Path,
    *,
    source_role: str,
    max_rows_per_sheet: int,
) -> tuple[list[dict[str, Any]], list[AmountCandidate]]:
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheet_summaries: list[dict[str, Any]] = []
    candidates: list[AmountCandidate] = []
    for sheet in workbook.worksheets:
        scanned_rows = min(sheet.max_row, max_rows_per_sheet)
        sheet_summaries.append(
            {
                "source_file": path.name,
                "sheet": sheet.title,
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
                "scanned_rows": scanned_rows,
            }
        )
        for row_idx, row in enumerate(
            sheet.iter_rows(
                min_row=1,
                max_row=scanned_rows,
                max_col=sheet.max_column,
                values_only=True,
            ),
            start=1,
        ):
            row_values = [_clean_text(value) for value in row]
            row_context = " | ".join(value for value in row_values if value)
            if not row_context:
                continue
            for col_idx, value in enumerate(row, start=1):
                location = f"{sheet.title}!{_excel_column_name(col_idx)}{row_idx}"
                if isinstance(value, (int, float)) and abs(value) >= 0.005:
                    candidates.append(
                        AmountCandidate(
                            source_file=path.name,
                            source_role=source_role,
                            location=location,
                            amount=round(float(value), 2),
                            token=str(value),
                            context=row_context[:180],
                        )
                    )
                elif isinstance(value, str):
                    candidates.extend(
                        extract_amount_candidates_from_text(
                            value,
                            source_file=path.name,
                            source_role=source_role,
                            location=location,
                        )
                    )
    return sheet_summaries, candidates


def _text_file_candidates(path: Path, *, source_role: str) -> list[AmountCandidate]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return extract_amount_candidates_from_text(
        text,
        source_file=path.name,
        source_role=source_role,
        location="text",
    )


def _candidate_rows(candidates: Iterable[AmountCandidate]) -> list[dict[str, Any]]:
    return [
        {
            "source_file": item.source_file,
            "source_role": item.source_role,
            "location": item.location,
            "amount": item.amount,
            "token": item.token,
            "context": item.context,
        }
        for item in candidates
    ]


def _context_overlap(left: str, right: str) -> float:
    left_tokens = _tokens(left)
    right_tokens = _tokens(right)
    if not left_tokens or not right_tokens:
        return 0.0
    return round(len(left_tokens & right_tokens) / len(left_tokens | right_tokens), 4)


def find_exact_amount_matches(
    candidates: list[AmountCandidate],
    *,
    tolerance: float,
    plan_role: str = "concordato_plan",
    max_matches_per_plan_amount: int = 20,
) -> list[dict[str, Any]]:
    """Find mechanically matching amounts; this is not semantic support."""

    plan_candidates = [item for item in candidates if item.source_role == plan_role]
    support_candidates = [item for item in candidates if item.source_role != plan_role]
    matches: list[dict[str, Any]] = []
    for plan in plan_candidates:
        row_matches = []
        for support in support_candidates:
            difference = round(plan.amount - support.amount, 2)
            abs_difference = abs(difference)
            if abs_difference > tolerance:
                continue
            row_matches.append(
                {
                    "plan_source_file": plan.source_file,
                    "plan_location": plan.location,
                    "plan_amount": plan.amount,
                    "plan_context": plan.context,
                    "support_source_file": support.source_file,
                    "support_role": support.source_role,
                    "support_location": support.location,
                    "support_amount": support.amount,
                    "support_context": support.context,
                    "difference": difference,
                    "abs_difference": abs_difference,
                    "context_token_overlap": _context_overlap(
                        plan.context, support.context
                    ),
                    "match_status": "candidate_amount_match",
                }
            )
        row_matches.sort(
            key=lambda row: (row["abs_difference"], -row["context_token_overlap"])
        )
        matches.extend(row_matches[:max_matches_per_plan_amount])
    return matches


def _write_csv(path: Path, rows: list[dict[str, Any]], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fieldnames})


def _write_workbook(
    path: Path,
    *,
    inventory: list[dict[str, Any]],
    candidates: list[dict[str, Any]],
    matches: list[dict[str, Any]],
    language: str,
) -> None:
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)
    sheet_names = (
        ("Inventario", "Importes candidatos", "Coincidencias candidatas")
        if language == "es"
        else ("Inventory", "Amount candidates", "Candidate matches")
    )
    empty_message = "No se generaron filas" if language == "es" else "No rows generated"
    for sheet_name, rows in zip(sheet_names, (inventory, candidates, matches)):
        sheet = workbook.create_sheet(sheet_name)
        headers = list(rows[0]) if rows else ["message"]
        sheet.append(headers)
        if rows:
            for row in rows:
                sheet.append([row.get(header, "") for header in headers])
        else:
            sheet.append([empty_message])
        sheet.freeze_panes = "A2"
        for column_cells in sheet.columns:
            header = str(column_cells[0].value or "")
            width = min(max(len(header) + 4, 14), 45)
            sheet.column_dimensions[column_cells[0].column_letter].width = width
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def _write_review_packet(
    path: Path,
    *,
    input_dir: Path,
    reference_date: str,
    language: str,
    document_language: str,
    tolerance: float,
    inventory: list[dict[str, Any]],
    candidates: list[AmountCandidate],
    matches: list[dict[str, Any]],
) -> None:
    role_counts: dict[str, int] = {}
    for item in inventory:
        role = str(item["suggested_role"])
        role_counts[role] = role_counts.get(role, 0) + 1
    plan_amounts = [
        item for item in candidates if item.source_role == "concordato_plan"
    ]
    supported_plan_keys = {
        (row["plan_source_file"], row["plan_location"], row["plan_amount"])
        for row in matches
    }
    unmatched_plan_count = sum(
        1
        for item in plan_amounts
        if (item.source_file, item.location, item.amount) not in supported_plan_keys
    )
    if language == "es":
        lines = [
            "# Paquete de revisión del plan de concordato",
            "",
            f"- Carpeta de entrada: `{input_dir}`",
            f"- Fecha de referencia: `{reference_date or 'no facilitada'}`",
            f"- Idioma: `{language}`",
            f"- Idioma de los documentos: `{document_language}`",
            f"- Tolerancia de importes: `{tolerance:,.2f}`",
            "",
            "## Roles de origen sugeridos a partir de los nombres de archivo",
            "",
        ]
    else:
        lines = [
            "# Concordato plan review packet",
            "",
            f"- Input folder: `{input_dir}`",
            f"- Reference date: `{reference_date or 'not provided'}`",
            f"- Language: `{language}`",
            f"- Document language: `{document_language}`",
            f"- Amount tolerance: `{tolerance:,.2f}`",
            "",
            "## Source roles suggested from file names",
            "",
        ]
    for role, count in sorted(role_counts.items()):
        lines.append(f"- `{role}`: {count}")
    if language == "es":
        lines.extend(
            [
                "",
                "## Recuentos deterministas",
                "",
                f"- Importes candidatos extraídos: {len(candidates)}",
                f"- Importes candidatos del plan: {len(plan_amounts)}",
                f"- Coincidencias candidatas por importe: {len(matches)}",
                f"- Importes candidatos del plan sin coincidencia por importe: {unmatched_plan_count}",
                "",
                "## Revisión requerida por Codex",
                "",
                "- Trate las coincidencias exactas de importe solo como candidatas, no como justificantes definitivos.",
                "- Clasifique las cifras del plan como datos históricos, rectificaciones, reclasificaciones, hipótesis prospectivas o partidas sin justificar o poco claras.",
                "- Revise los importes elevados del plan sin coincidencia y los importes con múltiples coincidencias del mismo valor.",
                "- Revise por separado los saldos de deudas tributarias y con la Seguridad Social frente al detalle específico correspondiente.",
                "- Redacte las cuestiones críticas orientadas al auditor en `codex_run_review.md` después de revisar el contexto.",
            ]
        )
    else:
        lines.extend(
            [
                "",
                "## Deterministic counts",
                "",
                f"- Amount candidates extracted: {len(candidates)}",
                f"- Plan amount candidates: {len(plan_amounts)}",
                f"- Candidate amount matches: {len(matches)}",
                f"- Plan amount candidates without an amount match: {unmatched_plan_count}",
                "",
                "## Codex review required",
                "",
                "- Treat exact amount matches as candidates only, not final support.",
                "- Classify plan numbers as historical data, rectification, reclassification, prospective assumption, or unsupported/unclear.",
                "- Review large unmatched plan amounts and plan amounts with many equal-value matches.",
                "- Inspect tax/social-security debt balances separately against the dedicated detail schedule.",
                "- Write auditor-oriented criticalities in `codex_run_review.md` after reviewing context.",
            ]
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _format_amount(value: object) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    return f"{number:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _unique_plan_candidates(
    candidates: list[AmountCandidate],
) -> list[AmountCandidate]:
    seen: set[tuple[str, str, float]] = set()
    unique: list[AmountCandidate] = []
    for item in candidates:
        if item.source_role != "concordato_plan":
            continue
        key = (item.source_file, item.location, item.amount)
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)
    return unique


def _best_match_by_plan_key(
    matches: list[dict[str, Any]],
) -> dict[tuple[str, str, float], dict[str, Any]]:
    best: dict[tuple[str, str, float], dict[str, Any]] = {}
    for row in matches:
        key = (
            str(row["plan_source_file"]),
            str(row["plan_location"]),
            float(row["plan_amount"]),
        )
        current = best.get(key)
        if current is None:
            best[key] = row
            continue
        row_score = (
            float(row.get("abs_difference", 0)),
            -float(row.get("context_token_overlap", 0)),
        )
        current_score = (
            float(current.get("abs_difference", 0)),
            -float(current.get("context_token_overlap", 0)),
        )
        if row_score < current_score:
            best[key] = row
    return best


def _add_docx_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def _write_summary_docx(
    path: Path,
    *,
    input_dir: Path,
    reference_date: str,
    tolerance: float,
    inventory: list[dict[str, Any]],
    candidates: list[AmountCandidate],
    matches: list[dict[str, Any]],
    extraction_errors: list[dict[str, str]],
) -> None:
    """Write a readable Word summary with match / no-match status."""

    plan_candidates = _unique_plan_candidates(candidates)
    best_matches = _best_match_by_plan_key(matches)
    matched_plan = [
        item
        for item in plan_candidates
        if (item.source_file, item.location, item.amount) in best_matches
    ]
    unmatched_plan = [
        item
        for item in plan_candidates
        if (item.source_file, item.location, item.amount) not in best_matches
    ]
    role_counts: dict[str, int] = {}
    for item in inventory:
        role = str(item["suggested_role"])
        role_counts[role] = role_counts.get(role, 0) + 1

    document = Document()
    document.add_heading("Revisione piano concordato - sintesi tie-out", 0)
    document.add_paragraph(f"Cartella analizzata: {input_dir}")
    document.add_paragraph(f"Data di riferimento: {reference_date or 'non indicata'}")
    document.add_paragraph(f"Tolleranza importi: {_format_amount(tolerance)} euro")

    document.add_heading("Conclusione operativa", level=1)
    document.add_paragraph(
        "Il documento distingue i numeri del piano che hanno almeno un match "
        "meccanico per importo nei file di supporto dai numeri per cui non e "
        "stato trovato un importo corrispondente entro la tolleranza. Il match "
        "per importo e una prova candidata: prima di chiudere la verifica serve "
        "controllare contesto, voce e fonte."
    )

    _add_docx_table(
        document,
        ["Area", "Esito"],
        [
            ["File analizzati", str(len(inventory))],
            ["Numeri candidati nel piano", str(len(plan_candidates))],
            ["Numeri del piano con match per importo", str(len(matched_plan))],
            ["Numeri del piano senza match per importo", str(len(unmatched_plan))],
            ["Match candidati complessivi", str(len(matches))],
        ],
    )

    document.add_heading("Fonti riconosciute", level=1)
    source_rows = [[role, str(count)] for role, count in sorted(role_counts.items())]
    _add_docx_table(document, ["Ruolo suggerito", "File"], source_rows)

    document.add_heading("Esempi di numeri che battono per importo", level=1)
    matched_rows: list[list[str]] = []
    for item in sorted(matched_plan, key=lambda row: abs(row.amount), reverse=True)[
        :12
    ]:
        match = best_matches[(item.source_file, item.location, item.amount)]
        matched_rows.append(
            [
                item.location,
                _format_amount(item.amount),
                str(match["support_role"]),
                f"{match['support_source_file']} - {match['support_location']}",
                _format_amount(match["difference"]),
            ]
        )
    if matched_rows:
        _add_docx_table(
            document,
            ["Piano", "Importo piano", "Fonte", "Riferimento fonte", "Differenza"],
            matched_rows,
        )
    else:
        document.add_paragraph("Nessun match per importo trovato.")

    document.add_heading("Esempi di numeri che non battono", level=1)
    unmatched_rows = [
        [item.location, _format_amount(item.amount), item.context[:180]]
        for item in sorted(
            unmatched_plan, key=lambda row: abs(row.amount), reverse=True
        )[:12]
    ]
    if unmatched_rows:
        _add_docx_table(
            document,
            ["Piano", "Importo piano", "Contesto"],
            unmatched_rows,
        )
    else:
        document.add_paragraph(
            "Tutti i numeri candidati del piano hanno almeno un match per importo."
        )

    document.add_heading("Da spiegare nel memo del revisore", level=1)
    for sentence in (
        "Il passaggio fonte contabile -> rettifica -> valore di piano deve essere documentato voce per voce.",
        "I match per importo uguale non bastano quando la stessa cifra compare in piu prospetti.",
        "Le rettifiche del revisore richiamate dal piano devono essere collegate a un supporto o a un memo autonomo.",
        "Le riclassifiche e compensazioni concordatarie devono essere separate dalle rettifiche contabili.",
        "I numeri prospettici non devono battere sui saldi storici, ma devono avere assunzioni esplicite e verificabili.",
    ):
        document.add_paragraph(sentence, style="List Bullet")

    if extraction_errors:
        document.add_heading("Errori di estrazione", level=1)
        _add_docx_table(
            document,
            ["File", "Errore"],
            [[str(row["source_file"]), str(row["error"])] for row in extraction_errors],
        )

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def run_concordato_review(
    input_dir: Path,
    output_dir: Path,
    *,
    reference_date: str = "",
    language: str = "it",
    document_language: str = "auto",
    tolerance: float = 1.0,
    max_rows_per_sheet: int = 5000,
) -> ReviewRun:
    """Run deterministic intake and candidate tie-out for a concordato plan."""

    input_dir = input_dir.resolve()
    output_dir = output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    language = normalize_language(language, default="it")
    document_language = normalize_language(
        document_language, default=language, allow_auto=True
    )

    inventory = _file_inventory(input_dir)
    run_intake = write_run_intake(
        output_dir,
        input_dir,
        reference_date=reference_date,
        language=language,
        document_language=document_language,
        tolerance=tolerance,
        max_rows_per_sheet=max_rows_per_sheet,
        inventory=inventory,
    )
    source_pages: list[dict[str, Any]] = []
    sheet_summaries: list[dict[str, Any]] = []
    candidates: list[AmountCandidate] = []
    extraction_errors: list[dict[str, str]] = []

    for item in inventory:
        if not item["supported"]:
            continue
        path = Path(item["path"])
        role = str(item["suggested_role"])
        suffix = path.suffix.lower()
        try:
            if suffix in PDF_SUFFIXES:
                pages = _read_pdf_pages(path)
                source_pages.extend(pages)
                for page in pages:
                    candidates.extend(
                        extract_amount_candidates_from_text(
                            str(page["text"]),
                            source_file=path.name,
                            source_role=role,
                            location=f"page {page['page']}",
                        )
                    )
            elif suffix in WORKBOOK_SUFFIXES:
                sheets, workbook_candidates = _workbook_candidates(
                    path,
                    source_role=role,
                    max_rows_per_sheet=max_rows_per_sheet,
                )
                sheet_summaries.extend(sheets)
                candidates.extend(workbook_candidates)
            elif suffix in TEXT_SUFFIXES:
                candidates.extend(_text_file_candidates(path, source_role=role))
        except (OSError, ValueError, RuntimeError, KeyError) as exc:
            extraction_errors.append(
                {"source_file": path.name, "error": f"{type(exc).__name__}: {exc}"}
            )
            LOGGER.warning("Failed to inspect %s: %s", path, exc)

    candidate_rows = _candidate_rows(candidates)
    matches = find_exact_amount_matches(candidates, tolerance=tolerance)
    audit = {
        "run_id": run_intake.run_id,
        "input_dir": str(input_dir),
        "reference_date": reference_date,
        "language": language,
        "document_language": document_language,
        "tolerance": tolerance,
        "max_rows_per_sheet": max_rows_per_sheet,
        "file_count": len(inventory),
        "supported_file_count": sum(1 for item in inventory if item["supported"]),
        "source_page_count": len(source_pages),
        "sheet_count": len(sheet_summaries),
        "amount_candidate_count": len(candidates),
        "candidate_match_count": len(matches),
        "extraction_errors": extraction_errors,
        "deterministic_boundary": (
            "Inventory, extraction, number parsing, arithmetic and candidate amount "
            "matching are deterministic for auditability. Semantic support, "
            "omissions, legal/tax relevance and going-concern criticalities are "
            "Codex/reviewer judgment."
        ),
    }

    write_json(output_dir / "inventory.json", inventory)
    write_json(output_dir / "source_pages.json", source_pages)
    write_json(output_dir / "workbook_sheets.json", sheet_summaries)
    write_json(output_dir / "run_audit.json", audit)
    _write_csv(
        output_dir / "amount_candidates.csv",
        candidate_rows,
        [
            "source_file",
            "source_role",
            "location",
            "amount",
            "token",
            "context",
        ],
    )
    _write_csv(
        output_dir / "exact_amount_matches.csv",
        matches,
        [
            "plan_source_file",
            "plan_location",
            "plan_amount",
            "plan_context",
            "support_source_file",
            "support_role",
            "support_location",
            "support_amount",
            "support_context",
            "difference",
            "abs_difference",
            "context_token_overlap",
            "match_status",
        ],
    )
    _write_workbook(
        output_dir / "concordato_tie_out_workpaper.xlsx",
        inventory=inventory,
        candidates=candidate_rows,
        matches=matches,
        language=language,
    )
    _write_review_packet(
        output_dir / "review_packet.md",
        input_dir=input_dir,
        reference_date=reference_date,
        language=language,
        document_language=document_language,
        tolerance=tolerance,
        inventory=inventory,
        candidates=candidates,
        matches=matches,
    )
    _write_summary_docx(
        output_dir / "concordato_review_summary.docx",
        input_dir=input_dir,
        reference_date=reference_date,
        tolerance=tolerance,
        inventory=inventory,
        candidates=candidates,
        matches=matches,
        extraction_errors=extraction_errors,
    )
    review_session = write_review_session_artifacts(
        output_dir,
        input_dir,
        run_id=run_intake.run_id,
        run_intake_path=run_intake.path,
        reference_date=reference_date,
        language=language,
        document_language=document_language,
        tolerance=tolerance,
        max_rows_per_sheet=max_rows_per_sheet,
        inventory=inventory,
        candidates=candidates,
        matches=matches,
        extraction_errors=extraction_errors,
        audit=audit,
    )
    audit["review_session"] = {
        "run_intake_path": review_session.run_intake_path.name,
        "review_payload_path": review_session.review_payload_path.name,
        "ui_decisions_path": review_session.ui_decisions_path.name,
        "final_artifacts_path": review_session.final_artifacts_path.name,
        "review_item_count": review_session.review_item_count,
    }
    write_json(output_dir / "run_audit.json", audit)
    return ReviewRun(
        output_dir=output_dir,
        inventory=inventory,
        amount_candidates=candidates,
        exact_matches=matches,
        audit=audit,
    )
