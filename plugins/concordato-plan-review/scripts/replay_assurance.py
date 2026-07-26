"""Replay persisted Concordato assurance before any review write."""

from __future__ import annotations

import sys as _bootstrap_sys  # isort: skip

_bootstrap_sys.dont_write_bytecode = True
_bootstrap_sys.pycache_prefix = (
    r"Z:\__concordato_no_bytecode__"
    if _bootstrap_sys.platform == "win32"
    else "/dev/null/concordato-plan-review"
)

import os as _bootstrap_os  # isort: skip

_BOOTSTRAP_PATH = _bootstrap_os.path.join(
    _bootstrap_os.path.dirname(_bootstrap_os.path.abspath(__file__)),
    "implementation_bootstrap.py",
)
_BOOTSTRAP_ENTRY = _bootstrap_os.lstat(_BOOTSTRAP_PATH)
if _BOOTSTRAP_ENTRY.st_mode & 0o170000 != 0o100000 or _BOOTSTRAP_ENTRY.st_nlink != 1:
    raise RuntimeError("Concordato implementation bootstrap is not a real file.")
with open(_BOOTSTRAP_PATH, "rb") as _bootstrap_handle:
    _BOOTSTRAP_BEFORE = _bootstrap_os.fstat(_bootstrap_handle.fileno())
    _BOOTSTRAP_BYTES = _bootstrap_handle.read()
    _BOOTSTRAP_AFTER = _bootstrap_os.fstat(_bootstrap_handle.fileno())
_BOOTSTRAP_IDENTITY = (
    _BOOTSTRAP_ENTRY.st_dev,
    _BOOTSTRAP_ENTRY.st_ino,
    _BOOTSTRAP_ENTRY.st_size,
    _BOOTSTRAP_ENTRY.st_mtime_ns,
)
if (
    _BOOTSTRAP_IDENTITY
    != (
        _BOOTSTRAP_BEFORE.st_dev,
        _BOOTSTRAP_BEFORE.st_ino,
        _BOOTSTRAP_BEFORE.st_size,
        _BOOTSTRAP_BEFORE.st_mtime_ns,
    )
    or _BOOTSTRAP_IDENTITY
    != (
        _BOOTSTRAP_AFTER.st_dev,
        _BOOTSTRAP_AFTER.st_ino,
        _BOOTSTRAP_AFTER.st_size,
        _BOOTSTRAP_AFTER.st_mtime_ns,
    )
    or len(_BOOTSTRAP_BYTES) != _BOOTSTRAP_AFTER.st_size
):
    raise RuntimeError("Concordato implementation bootstrap changed while read.")
_BOOTSTRAP_NAMESPACE = {
    "__file__": _BOOTSTRAP_PATH,
    "__name__": "_concordato_implementation_bootstrap",
}
# The exact stable single-link bootstrap source is verified above.
exec(  # nosec B102
    compile(_BOOTSTRAP_BYTES, _BOOTSTRAP_PATH, "exec"), _BOOTSTRAP_NAMESPACE
)
_BOOTSTRAP_NAMESPACE["activate_implementation_boundary"]()
_SCRIPTS_DIR = _bootstrap_os.path.dirname(_bootstrap_os.path.abspath(__file__))
if _SCRIPTS_DIR not in _bootstrap_sys.path:
    _bootstrap_sys.path.insert(0, _SCRIPTS_DIR)

import argparse
import copy
import json
import re
import shutil
import sys
import tempfile
from collections.abc import Mapping
from pathlib import Path, PurePosixPath
from typing import Any

import openpyxl
from concordato_plan_core import (
    ASSURANCE_IMPLEMENTATION_ROOT,
    COMPONENT_ROOT,
    SOURCE_ROLE_RECIPE_SCHEMA,
    run_concordato_review,
    validate_implementation_contract,
    validate_numeric_evidence_closure,
)
from docx import Document
from output_closure import (
    validate_final_artifact_index,
    validate_output_closure,
)
from vera_assurance import canonical_json_sha256, validate_assurance_envelope


def _read_object(path: Path) -> dict[str, Any]:
    observed = path.lstat()
    if path.is_symlink() or not path.is_file() or observed.st_nlink != 1:
        raise ValueError(f"{path.name} must be one unlinked regular file")
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"{path.name} must be a JSON object")
    return payload


def _read_regular_bytes(path: Path) -> bytes:
    observed = path.lstat()
    if path.is_symlink() or not path.is_file() or observed.st_nlink != 1:
        raise ValueError(f"{path.name} must be one unlinked regular file")
    return path.read_bytes()


def _decision_by_type(
    decisions: list[object],
    decision_type: str,
) -> dict[str, Any] | None:
    matches = [
        dict(item)
        for item in decisions
        if isinstance(item, Mapping) and item.get("decision_type") == decision_type
    ]
    if len(matches) > 1:
        raise ValueError(f"Reviewed {decision_type} decision is not unique")
    return matches[0] if matches else None


def _recipe_from_reviewed_decisions(
    reviewed: Mapping[str, Any],
) -> tuple[dict[str, Any] | None, dict[str, Any] | None]:
    raw_decisions = reviewed.get("decisions")
    if not isinstance(raw_decisions, list):
        raise ValueError("reviewed_decisions.decisions must be a list")
    source_decision = _decision_by_type(raw_decisions, "source_role_mapping")
    calculation_decision = _decision_by_type(
        raw_decisions,
        "calculation_formula_authority",
    )
    if source_decision is None and calculation_decision is None:
        return None, None
    if source_decision is None or calculation_decision is None:
        raise ValueError("Reviewed source and calculation decisions must be paired")
    content = source_decision.get("content")
    if not isinstance(content, Mapping):
        raise ValueError("Reviewed source-role content is invalid")
    raw_sources = content.get("source_roles")
    raw_dispositions = content.get("candidate_dispositions")
    if not isinstance(raw_sources, list) or not isinstance(raw_dispositions, list):
        raise ValueError("Reviewed source-role perimeter is invalid")
    source_roles: dict[str, dict[str, str]] = {}
    for raw_source in raw_sources:
        if not isinstance(raw_source, Mapping):
            raise ValueError("Reviewed source-role entry is invalid")
        relative = str(raw_source.get("relative_path") or "")
        if not relative or relative in source_roles:
            raise ValueError("Reviewed source-role paths are invalid")
        source_roles[relative] = {
            "role": str(raw_source.get("role") or ""),
            "currency": str(raw_source.get("currency") or ""),
            "unit": str(raw_source.get("unit") or ""),
        }
    candidate_dispositions: dict[str, str] = {}
    for raw_disposition in raw_dispositions:
        if not isinstance(raw_disposition, Mapping):
            raise ValueError("Reviewed candidate disposition is invalid")
        identity = str(raw_disposition.get("candidate_id") or "")
        if not identity or identity in candidate_dispositions:
            raise ValueError("Reviewed candidate disposition identities are invalid")
        candidate_dispositions[identity] = str(raw_disposition.get("disposition") or "")
    return (
        {
            "schema_version": SOURCE_ROLE_RECIPE_SCHEMA,
            "source_roles": source_roles,
            "candidate_dispositions": candidate_dispositions,
            "source_role_decision": source_decision,
            "calculation_decision": calculation_decision,
        },
        calculation_decision,
    )


def _workbook_image(path: Path) -> list[dict[str, Any]]:
    workbook = openpyxl.load_workbook(path, data_only=False, read_only=False)
    try:
        image: list[dict[str, Any]] = []
        for sheet in workbook.worksheets:
            cells = []
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is None and cell.style_id == 0:
                        continue
                    cells.append(
                        {
                            "coordinate": cell.coordinate,
                            "value": (
                                cell.value
                                if isinstance(
                                    cell.value,
                                    (str, int, bool, type(None)),
                                )
                                else str(cell.value)
                            ),
                            "data_type": cell.data_type,
                            "number_format": cell.number_format,
                            "style_id": cell.style_id,
                        }
                    )
            image.append(
                {
                    "title": sheet.title,
                    "state": sheet.sheet_state,
                    "freeze_panes": (
                        None if sheet.freeze_panes is None else str(sheet.freeze_panes)
                    ),
                    "merged_cells": sorted(str(item) for item in sheet.merged_cells),
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column,
                    "cells": cells,
                }
            )
        return image
    finally:
        workbook.close()


def _docx_image(path: Path) -> dict[str, Any]:
    document = Document(path)
    return {
        "paragraphs": [
            {
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style is not None else None,
            }
            for paragraph in document.paragraphs
        ],
        "tables": [
            [
                [
                    {
                        "text": cell.text,
                        "paragraph_styles": [
                            (
                                paragraph.style.name
                                if paragraph.style is not None
                                else None
                            )
                            for paragraph in cell.paragraphs
                        ],
                    }
                    for cell in row.cells
                ]
                for row in table.rows
            ]
            for table in document.tables
        ],
    }


def _memo_text_from_applied_state(
    output_dir: Path,
    *,
    require_target: bool,
) -> str | None:
    applied_path = output_dir / "applied_decisions.json"
    if not applied_path.exists():
        return None
    applied = _read_object(applied_path)
    effects = applied.get("effects")
    if not isinstance(effects, list):
        raise ValueError("applied_decisions.effects must be a list")
    memo_effects = [
        effect
        for effect in effects
        if isinstance(effect, Mapping)
        and effect.get("item_id") == "codex-review-memo"
        and effect.get("action") == "edit"
        and isinstance(effect.get("edit_value"), str)
        and str(effect["edit_value"]).strip()
    ]
    if not memo_effects:
        return None
    memo_text = str(memo_effects[-1]["edit_value"]).strip()
    target = str(memo_effects[-1].get("target_artifact") or "codex_run_review.md")
    if target != "codex_run_review.md":
        raise ValueError("Concordato review memo target is unsupported")
    target_path = output_dir / target
    if not target_path.exists() and not require_target:
        return memo_text
    if _read_regular_bytes(target_path) != memo_text.encode("utf-8"):
        raise ValueError("Concordato review memo is not derived from applied state")
    return memo_text


def _append_expected_memo(path: Path, memo_text: str) -> None:
    document = Document(path)
    document.add_heading("Memo revisore Codex", level=1)
    for raw_line in memo_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^#{1,6}\s+(.*)$", line)
        if heading:
            document.add_heading(heading.group(1).strip(), level=2)
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", line)
        if bullet:
            document.add_paragraph(bullet.group(1).strip(), style="List Bullet")
            continue
        document.add_paragraph(line)
    document.save(path)


def _without_keys(value: object, keys: set[str]) -> object:
    if isinstance(value, list):
        return [_without_keys(item, keys) for item in value]
    if isinstance(value, Mapping):
        return {
            str(key): _without_keys(item, keys)
            for key, item in value.items()
            if key not in keys
        }
    return value


_ACTION_STATUSES = {
    "accept": "accepted",
    "reject": "rejected",
    "edit": "edited",
    "mark_unclear": "needs_evidence",
    "request_more_documents": "needs_evidence",
    "skip": "skipped",
}
_FOLLOWUP_ACTIONS = {
    "reject",
    "mark_unclear",
    "request_more_documents",
}
_UI_DECISION_FIELDS = {
    "action",
    "decided_at",
    "edit_value",
    "followup_context",
    "item_id",
    "item_type",
    "requested_documents",
    "reviewer_note",
    "status",
    "title",
}


def _review_items(review_payload: Mapping[str, Any]) -> dict[str, dict[str, Any]]:
    raw_items = review_payload.get("items")
    if not isinstance(raw_items, list):
        raise ValueError("review_payload.items must be a list")
    if review_payload.get("item_count") != len(raw_items):
        raise ValueError("review_payload.item_count is stale")
    items: dict[str, dict[str, Any]] = {}
    for raw_item in raw_items:
        if not isinstance(raw_item, Mapping):
            raise ValueError("review_payload item is invalid")
        item = dict(raw_item)
        identity = item.get("id")
        if not isinstance(identity, str) or not identity or identity in items:
            raise ValueError("review_payload item identities are invalid")
        items[identity] = item
    return items


def _validate_ui_decisions(
    ui_decisions: Mapping[str, Any],
    review_payload: Mapping[str, Any],
) -> tuple[list[dict[str, Any]], dict[str, dict[str, Any]]]:
    """Validate reviewer authority input without inventing reviewer judgment."""

    allowed_fields = {
        "schema_version",
        "plugin",
        "workflow",
        "run_id",
        "decided_at",
        "decision_source",
        "review_payload_path",
        "review_payload_content_sha256",
        "decisions",
        "decision_count",
        "item_count",
        "status",
        "reviewer",
    }
    if set(ui_decisions) - allowed_fields:
        raise ValueError("ui_decisions.json contains unsupported fields")
    for field in ("schema_version", "plugin", "workflow", "run_id"):
        if ui_decisions.get(field) != review_payload.get(field):
            raise ValueError(f"ui_decisions.{field} is not bound to review_payload")
    if ui_decisions.get("review_payload_path") != "review_payload.json":
        raise ValueError("ui_decisions review payload path is invalid")
    if ui_decisions.get("review_payload_content_sha256") != review_payload.get(
        "content_sha256"
    ):
        raise ValueError("ui_decisions review payload digest is stale")
    decision_source = ui_decisions.get("decision_source")
    if not isinstance(decision_source, str) or not decision_source.strip():
        raise ValueError("ui_decisions.decision_source is invalid")
    reviewer = ui_decisions.get("reviewer")
    if reviewer is not None and (not isinstance(reviewer, str) or not reviewer.strip()):
        raise ValueError("ui_decisions.reviewer is invalid")

    items = _review_items(review_payload)
    raw_decisions = ui_decisions.get("decisions")
    if not isinstance(raw_decisions, list):
        raise ValueError("ui_decisions.decisions must be a list")
    if ui_decisions.get("decision_count") != len(raw_decisions):
        raise ValueError("ui_decisions.decision_count is stale")
    if "item_count" in ui_decisions and ui_decisions.get("item_count") != len(items):
        raise ValueError("ui_decisions.item_count is stale")
    expected_status = (
        "pending_review"
        if not raw_decisions
        else "reviewed" if len(raw_decisions) == len(items) else "partial_review"
    )
    if ui_decisions.get("status") != expected_status:
        raise ValueError("ui_decisions.status is not derived from its decisions")
    decided_at = ui_decisions.get("decided_at")
    if raw_decisions:
        if not isinstance(decided_at, str) or not decided_at.strip():
            raise ValueError("ui_decisions.decided_at is invalid")
    elif decided_at is not None:
        raise ValueError("pending ui_decisions cannot claim a decision timestamp")

    decisions: list[dict[str, Any]] = []
    seen: set[str] = set()
    for raw_decision in raw_decisions:
        if not isinstance(raw_decision, Mapping):
            raise ValueError("ui_decisions decision is invalid")
        if set(raw_decision) - _UI_DECISION_FIELDS:
            raise ValueError("ui_decisions decision contains unsupported fields")
        decision = dict(raw_decision)
        identity = decision.get("item_id")
        if not isinstance(identity, str) or identity not in items or identity in seen:
            raise ValueError("ui_decisions decision identity is invalid")
        seen.add(identity)
        item = items[identity]
        action = decision.get("action")
        item_actions = item.get("allowed_actions")
        if (
            action not in _ACTION_STATUSES
            or not isinstance(item_actions, list)
            or action not in item_actions
        ):
            raise ValueError("ui_decisions action is not allowed for its review item")
        if (
            decision.get("item_type") != item.get("item_type")
            or decision.get("title") != item.get("title")
            or decision.get("status") != _ACTION_STATUSES[action]
            or decision.get("decided_at") != decided_at
        ):
            raise ValueError(
                "ui_decisions decision is not derived from its review item"
            )
        edit_value = decision.get("edit_value")
        if action == "edit" and (
            not isinstance(edit_value, str) or not edit_value.strip()
        ):
            raise ValueError("ui_decisions edit action requires edit_value")
        for field in ("edit_value", "reviewer_note"):
            value = decision.get(field)
            if value is not None and not isinstance(value, str):
                raise ValueError(f"ui_decisions decision {field} is invalid")
        requested_documents = decision.get("requested_documents")
        if requested_documents is not None and (
            not isinstance(requested_documents, list)
            or any(
                not isinstance(value, str) or not value.strip()
                for value in requested_documents
            )
        ):
            raise ValueError("ui_decisions requested_documents is invalid")
        followup_context = decision.get("followup_context")
        if followup_context is not None and not isinstance(
            followup_context,
            Mapping,
        ):
            raise ValueError("ui_decisions followup_context is invalid")
        decisions.append(decision)
    return decisions, items


def _short_item_value(value: object) -> str | None:
    return value.strip() if isinstance(value, str) and value.strip() else None


def _expected_effect_targets(item: Mapping[str, Any]) -> dict[str, str | None]:
    data = item.get("data")
    item_data = dict(data) if isinstance(data, Mapping) else {}

    def first(*values: object) -> str | None:
        return next(
            (text for value in values if (text := _short_item_value(value))),
            None,
        )

    return {
        "target_artifact": first(
            item_data.get("target_artifact"),
            item.get("output_path"),
            item_data.get("path"),
        ),
        "target_path": first(
            item_data.get("target_path"),
            item_data.get("field_path"),
            item_data.get("field"),
        ),
        "target_id_field": first(
            item_data.get("target_id_field"),
            item_data.get("record_id_field"),
        ),
        "target_record_id": first(
            item_data.get("target_record_id"),
            item_data.get("record_id"),
        ),
        "target_field": first(
            item_data.get("target_field"),
            item_data.get("edit_field"),
        ),
        "target_records_key": first(
            item_data.get("target_records_key"),
            item_data.get("records_key"),
        ),
        "source_path": _short_item_value(item.get("source_path")),
    }


def _canonical_review_paths(value: object, field: str) -> list[str]:
    if not isinstance(value, list):
        raise ValueError(f"applied_decisions.{field} must be a list")
    paths: list[str] = []
    for raw_path in value:
        if not isinstance(raw_path, str):
            raise ValueError(f"applied_decisions.{field} contains a non-path")
        parsed = PurePosixPath(raw_path)
        if (
            not raw_path
            or parsed.is_absolute()
            or parsed.as_posix() != raw_path
            or ".." in parsed.parts
            or "\\" in raw_path
        ):
            raise ValueError(f"applied_decisions.{field} path is not canonical")
        paths.append(raw_path)
    if len(paths) != len(set(paths)):
        raise ValueError(f"applied_decisions.{field} paths are not unique")
    return paths


def _validate_applied_decisions(
    applied: Mapping[str, Any],
    ui_decisions: Mapping[str, Any],
    review_payload: Mapping[str, Any],
    decisions: list[dict[str, Any]],
    items: Mapping[str, dict[str, Any]],
) -> tuple[list[dict[str, Any]], dict[str, list[str]]]:
    """Derive application state from persisted reviewer authority."""

    for field in ("schema_version", "plugin", "workflow", "run_id"):
        if applied.get(field) != review_payload.get(field):
            raise ValueError(f"applied_decisions.{field} is stale")
    if applied.get("decision_source") != ui_decisions.get("decision_source"):
        raise ValueError("applied_decisions.decision_source is stale")
    if applied.get("reviewer") != ui_decisions.get("reviewer"):
        raise ValueError("applied_decisions.reviewer is stale")
    expected_review = {
        "path": "review_payload.json",
        "content_sha256": review_payload.get("content_sha256"),
        "item_count": len(items),
        "review_type": review_payload.get("review_type"),
    }
    if applied.get("review_payload") != expected_review:
        raise ValueError("applied_decisions review payload binding is stale")
    if applied.get("decisions") != decisions:
        raise ValueError("applied_decisions are not the persisted UI decisions")
    if applied.get("decision_count") != len(decisions) or applied.get(
        "item_count"
    ) != len(items):
        raise ValueError("applied_decisions counts are stale")
    applied_at = applied.get("applied_at")
    if not isinstance(applied_at, str) or not applied_at.strip():
        raise ValueError("applied_decisions.applied_at is invalid")

    raw_effects = applied.get("effects")
    if not isinstance(raw_effects, list) or len(raw_effects) != len(decisions):
        raise ValueError("applied_decisions.effects are incomplete")
    effects: list[dict[str, Any]] = []
    for decision, raw_effect in zip(decisions, raw_effects, strict=True):
        if not isinstance(raw_effect, Mapping):
            raise ValueError("applied_decisions effect is invalid")
        effect = dict(raw_effect)
        item = items[str(decision["item_id"])]
        expected_core: dict[str, Any] = {
            "item_id": decision["item_id"],
            "item_type": decision["item_type"],
            "title": decision["title"],
            "action": decision["action"],
            "status": decision["status"],
            "applied_at": applied_at,
            "applied": True,
            "requires_followup": decision["action"] in _FOLLOWUP_ACTIONS,
            **_expected_effect_targets(item),
        }
        if any(effect.get(field) != value for field, value in expected_core.items()):
            raise ValueError(
                "applied_decisions effect is not derived from its UI decision"
            )
        for field in (
            "edit_value",
            "followup_context",
            "requested_documents",
            "reviewer_note",
        ):
            if effect.get(field) != decision.get(field):
                raise ValueError(
                    "applied_decisions effect does not preserve review content"
                )
        effects.append(effect)

    path_fields = {
        "revision_paths": "revision_count",
        "target_update_paths": "target_update_count",
        "structured_update_paths": "structured_update_count",
        "native_regeneration_paths": "native_regeneration_count",
    }
    paths: dict[str, list[str]] = {}
    for path_field, count_field in path_fields.items():
        paths[path_field] = _canonical_review_paths(
            applied.get(path_field),
            path_field,
        )
        if applied.get(count_field) != len(paths[path_field]):
            raise ValueError(f"applied_decisions.{count_field} is stale")
    for path_field in ("native_regenerated_paths", "original_backup_paths"):
        raw_paths = applied.get(path_field, [])
        paths[path_field] = _canonical_review_paths(raw_paths, path_field)
    if "native_regenerated_count" in applied and applied.get(
        "native_regenerated_count"
    ) != len(paths["native_regenerated_paths"]):
        raise ValueError("applied_decisions.native_regenerated_count is stale")

    blocker_count = sum(1 for effect in effects if effect["requires_followup"] is True)
    if applied.get("blocker_count") != blocker_count:
        raise ValueError("applied_decisions.blocker_count is stale")
    expected_status = (
        "pending_review"
        if not effects
        else (
            "blocked"
            if blocker_count
            else (
                "partial_review_applied"
                if paths["native_regeneration_paths"] or len(decisions) < len(items)
                else "review_applied_assurance_withheld"
            )
        )
    )
    if applied.get("application_status") != expected_status:
        raise ValueError("applied_decisions.application_status is stale")
    return effects, paths


def _expected_blockers(effects: list[dict[str, Any]]) -> list[dict[str, Any]]:
    blockers: list[dict[str, Any]] = []
    for effect in effects:
        if not effect["requires_followup"]:
            continue
        blocker = {
            "item_id": effect["item_id"],
            "item_type": effect["item_type"],
            "title": effect["title"],
            "action": effect["action"],
            "status": effect["status"],
            "reviewer_note": effect.get("reviewer_note"),
            "requested_documents": effect.get("requested_documents", []),
        }
        context = effect.get("followup_context")
        if isinstance(context, Mapping) and context:
            blocker["followup_context"] = dict(context)
        blockers.append(blocker)
    return blockers


def _expected_next_actions(
    baseline: Mapping[str, Any],
    applied: Mapping[str, Any],
    effects: list[dict[str, Any]],
    *,
    language: str,
) -> list[str]:
    raw_actions = baseline.get("next_actions")
    if not isinstance(raw_actions, list) or any(
        not isinstance(action, str) for action in raw_actions
    ):
        raise ValueError("baseline final_artifacts next_actions are invalid")
    actions = list(raw_actions)
    spanish = language.lower().startswith("es")
    if any(effect["requires_followup"] for effect in effects):
        actions.append(
            "Resuelva las decisiones de revisión bloqueadas antes de considerar "
            "listos los artefactos finales."
            if spanish
            else "Resolve blocked review decisions before treating final artifacts "
            "as ready."
        )
    elif applied.get("native_regeneration_count"):
        actions.append(
            "Vuelva a generar las salidas nativas DOCX, XLSX o PDF antes de la "
            "entrega final."
            if spanish
            else "Regenerate native DOCX/XLSX/PDF outputs before final handoff."
        )
    elif applied.get("application_status") == "review_applied_assurance_withheld":
        actions.append(
            "La revisión se registró, pero la conclusión profesional y la "
            "publicación siguen retenidas."
            if spanish
            else "Review was recorded, but professional conclusion and publication "
            "remain withheld."
        )
    elif applied.get("application_status") == "partial_review_applied":
        actions.append(
            "Complete las decisiones de revisión restantes antes de la entrega final."
            if spanish
            else "Complete remaining review decisions before final handoff."
        )
    return list(dict.fromkeys(actions))


def _validate_applied_final_artifacts(
    final_artifacts: Mapping[str, Any],
    baseline: Mapping[str, Any],
    applied: Mapping[str, Any],
    review_payload: Mapping[str, Any],
    effects: list[dict[str, Any]],
    paths: Mapping[str, list[str]],
) -> None:
    expected_fields = {
        "schema_version",
        "plugin",
        "workflow",
        "run_id",
        "outputs",
        "caveats",
        "blockers",
        "next_actions",
        "status",
        "review_status",
        "final_ready",
        "assurance",
        "review_application",
    }
    if set(final_artifacts) != expected_fields:
        raise ValueError("applied final_artifacts.json has invalid fields")
    for field in ("schema_version", "plugin", "workflow", "run_id"):
        if final_artifacts.get(field) != review_payload.get(field):
            raise ValueError(f"final_artifacts.{field} is stale")
    status = applied.get("application_status")
    if (
        final_artifacts.get("status") != status
        or final_artifacts.get("review_status") != status
        or final_artifacts.get("final_ready") is not False
        or final_artifacts.get("assurance") != review_payload.get("assurance")
    ):
        raise ValueError("final_artifacts review status is not derived")
    if final_artifacts.get("caveats") != baseline.get("caveats"):
        raise ValueError("final_artifacts caveats are not independently reproducible")
    if final_artifacts.get("blockers") != _expected_blockers(effects):
        raise ValueError("final_artifacts blockers are not derived from review effects")
    if final_artifacts.get("next_actions") != _expected_next_actions(
        baseline,
        applied,
        effects,
        language=str(review_payload.get("language") or "en"),
    ):
        raise ValueError("final_artifacts next_actions are not derived")

    application = final_artifacts.get("review_application")
    if not isinstance(application, Mapping):
        raise ValueError("final_artifacts.review_application is invalid")
    application_fields = {
        "applied_at",
        "application_status",
        "decision_count",
        "item_count",
        "blocker_count",
        "revision_count",
        "revision_paths",
        "target_update_count",
        "target_update_paths",
        "structured_update_count",
        "structured_update_paths",
        "native_regeneration_count",
        "native_regeneration_paths",
        "original_backup_paths",
    }
    for field in application_fields:
        if application.get(field) != applied.get(field):
            raise ValueError(f"final_artifacts.review_application.{field} is stale")
    if application.get("applied_decisions_path") != "applied_decisions.json":
        raise ValueError("final_artifacts applied decision path is stale")
    for field in ("native_regenerated_count", "native_regenerated_paths"):
        if field in applied or field in application:
            if application.get(field) != applied.get(field):
                raise ValueError(f"final_artifacts.review_application.{field} is stale")

    raw_outputs = final_artifacts.get("outputs")
    baseline_outputs = baseline.get("outputs")
    if not isinstance(raw_outputs, list) or not isinstance(baseline_outputs, list):
        raise ValueError("final_artifacts outputs are invalid")
    outputs_by_path = {
        str(output.get("path")): dict(output)
        for output in raw_outputs
        if isinstance(output, Mapping)
    }
    if len(outputs_by_path) != len(raw_outputs):
        raise ValueError("final_artifacts output paths are invalid or duplicated")
    mutable_paths = {
        "ui_decisions.json",
        "applied_decisions.json",
        *(path for path_list in paths.values() for path in path_list),
    }
    baseline_by_path = {
        str(output.get("path")): dict(output)
        for output in baseline_outputs
        if isinstance(output, Mapping)
    }
    for relative, expected_output in baseline_by_path.items():
        actual_output = outputs_by_path.get(relative)
        if actual_output is None:
            raise ValueError("final_artifacts dropped a source-generated output")
        if relative not in mutable_paths and _without_keys(
            actual_output,
            {"sha256", "size_bytes"},
        ) != _without_keys(expected_output, {"sha256", "size_bytes"}):
            raise ValueError("final_artifacts output metadata is not reproducible")
    if set(outputs_by_path) - set(baseline_by_path) - mutable_paths:
        raise ValueError("final_artifacts contains an unauthorized output")
    if not mutable_paths.issubset(outputs_by_path):
        raise ValueError("final_artifacts omits a review-owned output")
    if outputs_by_path["ui_decisions.json"].get("status") != "written_reviewed":
        raise ValueError("final_artifacts UI decision status is stale")
    if outputs_by_path["applied_decisions.json"].get("status") != status:
        raise ValueError("final_artifacts applied decision status is stale")


def _validate_independent_outputs(
    output_dir: Path,
    source_root: Path,
    run_intake: Mapping[str, Any],
    reviewed: Mapping[str, Any],
    *,
    require_final_state: bool,
) -> None:
    """Regenerate deterministic outputs and compare independent representations."""

    recipe, _ = _recipe_from_reviewed_decisions(reviewed)
    assumptions = run_intake.get("assumptions")
    if not isinstance(assumptions, Mapping):
        raise ValueError("run_intake.assumptions is invalid")
    with tempfile.TemporaryDirectory(prefix="concordato-replay-") as temporary:
        baseline = Path(temporary) / "baseline"
        run_concordato_review(
            source_root,
            baseline,
            reference_date=str(assumptions.get("reference_date") or ""),
            language=str(run_intake.get("language") or "it"),
            document_language=str(run_intake.get("document_language") or "auto"),
            tolerance=assumptions.get("tolerance"),
            max_rows_per_sheet=int(assumptions.get("max_rows_per_sheet") or 0),
            recipe=recipe,
        )

        exact_paths = {
            "amount_candidates.csv",
            "assurance_gates.json",
            "exact_amount_matches.csv",
            "inventory.json",
            "numeric_evidence_ledger.json",
            "raw_amount_candidates.csv",
            "review_packet.md",
            "reviewed_decisions.json",
            "source_pages.json",
            "source_qualifications.json",
            "source_receipts.json",
            "suggested_source_role_recipe.json",
            "workbook_sheets.json",
        }
        for relative in sorted(exact_paths):
            if _read_regular_bytes(output_dir / relative) != _read_regular_bytes(
                baseline / relative
            ):
                raise ValueError(
                    f"{relative} is not independently reproducible from source authority"
                )
        if _workbook_image(
            output_dir / "concordato_tie_out_workpaper.xlsx"
        ) != _workbook_image(baseline / "concordato_tie_out_workpaper.xlsx"):
            raise ValueError("Concordato workpaper is not independently reproducible")

        baseline_docx = baseline / "concordato_review_summary.docx"
        expected_docx = Path(temporary) / "expected-summary.docx"
        shutil.copy2(baseline_docx, expected_docx)
        memo_text = _memo_text_from_applied_state(
            output_dir,
            require_target=require_final_state,
        )
        if memo_text is not None:
            _append_expected_memo(expected_docx, memo_text)
        actual_docx_image = _docx_image(output_dir / "concordato_review_summary.docx")
        allowed_docx_images = [_docx_image(expected_docx)]
        if not require_final_state and memo_text is not None:
            allowed_docx_images.append(_docx_image(baseline_docx))
        if actual_docx_image not in allowed_docx_images:
            raise ValueError(
                "Concordato summary is not independently reproducible from "
                "deterministic facts and applied review state"
            )

        actual_audit = _read_object(output_dir / "run_audit.json")
        expected_audit = _read_object(baseline / "run_audit.json")
        for audit in (actual_audit, expected_audit):
            audit["run_id"] = "<run>"
            envelope_ref = audit.get("assurance_envelope")
            if isinstance(envelope_ref, dict):
                envelope_ref["content_sha256"] = "<envelope>"
        if actual_audit != expected_audit:
            raise ValueError("run_audit.json is not independently reproducible")

        actual_review = _read_object(output_dir / "review_payload.json")
        expected_review = _read_object(baseline / "review_payload.json")
        normalized_actual_review = copy.deepcopy(actual_review)
        normalized_expected_review = copy.deepcopy(expected_review)
        for review in (normalized_actual_review, normalized_expected_review):
            review["run_id"] = "<run>"
            review["created_at"] = "<time>"
            review["content_sha256"] = "<digest>"
            assurance = review.get("assurance")
            if isinstance(assurance, dict):
                assurance["envelope_content_sha256"] = "<envelope>"
        if _without_keys(normalized_actual_review, {"size_bytes"}) != _without_keys(
            normalized_expected_review,
            {"size_bytes"},
        ):
            raise ValueError("review_payload.json is not independently reproducible")

        actual_intake = copy.deepcopy(dict(run_intake))
        expected_intake = _read_object(baseline / "run_intake.json")
        for intake in (actual_intake, expected_intake):
            for key in ("run_id", "created_at", "output_dir", "execution_trace"):
                intake.pop(key, None)
        if actual_intake != expected_intake:
            raise ValueError("run_intake.json is not independently reproducible")

        actual_handoff = _read_regular_bytes(output_dir / "review_handoff.md").decode(
            "utf-8"
        )
        expected_handoff = _read_regular_bytes(baseline / "review_handoff.md").decode(
            "utf-8"
        )
        actual_handoff = actual_handoff.replace(
            str(run_intake.get("run_id")),
            "<run>",
        )
        expected_baseline_intake = _read_object(baseline / "run_intake.json")
        expected_handoff = expected_handoff.replace(
            str(expected_baseline_intake.get("run_id")),
            "<run>",
        )
        if actual_handoff != expected_handoff:
            raise ValueError("review_handoff.md is not independently reproducible")

        actual_ui = _read_object(output_dir / "ui_decisions.json")
        decisions, items = _validate_ui_decisions(actual_ui, actual_review)
        actual_final = _read_object(output_dir / "final_artifacts.json")
        expected_final = _read_object(baseline / "final_artifacts.json")
        applied_path = output_dir / "applied_decisions.json"
        if not applied_path.exists():
            normalization_keys = {
                "completed_at",
                "content_sha256",
                "envelope_content_sha256",
                "run_id",
                "sha256",
                "size_bytes",
            }
            if _without_keys(
                actual_final,
                normalization_keys,
            ) != _without_keys(expected_final, normalization_keys):
                raise ValueError(
                    "final_artifacts.json is not independently reproducible"
                )
        else:
            applied = _read_object(applied_path)
            effects, review_paths = _validate_applied_decisions(
                applied,
                actual_ui,
                actual_review,
                decisions,
                items,
            )
            _validate_applied_final_artifacts(
                actual_final,
                expected_final,
                applied,
                actual_review,
                effects,
                review_paths,
            )


def replay_assurance(
    output_dir: Path,
    *,
    require_output_closure: bool = True,
) -> dict[str, Any]:
    """Replay receipts and optionally require the current whole-output seal."""

    output_dir = output_dir.resolve()
    run_intake = _read_object(output_dir / "run_intake.json")
    review_payload = _read_object(output_dir / "review_payload.json")
    envelope = _read_object(output_dir / "assurance_envelope.json")
    reviewed = _read_object(output_dir / "reviewed_decisions.json")
    source_paths = run_intake.get("input_paths")
    if not isinstance(source_paths, list) or len(source_paths) != 1:
        raise ValueError("run_intake.input_paths must contain exactly one source root")
    content_digest = review_payload.pop("content_sha256", None)
    expected_review_digest = canonical_json_sha256(review_payload)
    if content_digest != expected_review_digest:
        raise ValueError("Persisted review_payload content digest is stale")
    validated = validate_assurance_envelope(
        envelope,
        artifact_roots={
            "source": Path(str(source_paths[0])).resolve(),
            "run": output_dir,
            "implementation": COMPONENT_ROOT,
            "assurance_implementation": ASSURANCE_IMPLEMENTATION_ROOT,
        },
    )
    recipe, calculation_decision = _recipe_from_reviewed_decisions(reviewed)
    del recipe
    if calculation_decision is not None:
        implementation_receipts = validate_implementation_contract(calculation_decision)
        expected_refs = [
            str(receipt["artifact_id"]) for receipt in implementation_receipts
        ]
        if envelope.get("implementation_artifact_refs") != expected_refs:
            raise ValueError("Assurance envelope implementation set is stale")
        envelope_implementation = [
            receipt
            for receipt in envelope.get("artifact_receipts", [])
            if isinstance(receipt, Mapping)
            and receipt.get("root_id") in {"implementation", "assurance_implementation"}
        ]
        if envelope_implementation != implementation_receipts:
            raise ValueError("Assurance envelope implementation receipts are stale")
    if envelope.get("reviewed_decisions") != reviewed.get("decisions"):
        raise ValueError("Persisted reviewed decisions are outside the envelope")
    assurance = review_payload.get("assurance")
    if not isinstance(assurance, dict):
        raise ValueError("review_payload.assurance is missing")
    if assurance.get("envelope_content_sha256") != validated["content_sha256"]:
        raise ValueError(
            "Review payload is not bound to the current assurance envelope"
        )
    if assurance.get("final_ready") is not False:
        raise ValueError(
            "Concordato review payload cannot claim deterministic final readiness"
        )
    numeric_ledger = _read_object(output_dir / "numeric_evidence_ledger.json")
    if numeric_ledger.get("schema_version") == "concordato.numeric_evidence_ledger.v2":
        validate_numeric_evidence_closure(output_dir, numeric_ledger)
    _validate_independent_outputs(
        output_dir,
        Path(str(source_paths[0])).resolve(),
        run_intake,
        reviewed,
        require_final_state=require_output_closure,
    )
    result = {
        "ok": True,
        "run_id": review_payload.get("run_id"),
        "review_payload_content_sha256": expected_review_digest,
        "assurance_envelope_content_sha256": validated["content_sha256"],
        "report_ready": validated["gate_register"]["report_ready"],
    }
    if not require_output_closure:
        return result
    validate_final_artifact_index(output_dir)
    closure = validate_output_closure(output_dir)
    if closure["run_id"] != review_payload.get("run_id"):
        raise ValueError("Workflow output closure run identity is stale")
    return {
        **result,
        "workflow_output_closure_content_sha256": closure["content_sha256"],
        "workflow_output_closure_phase": closure["phase"],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    try:
        result = replay_assurance(args.output_dir)
    except (OSError, ValueError) as exc:
        sys.stdout.write(json.dumps({"ok": False, "error": str(exc)}) + "\n")
        return 1
    sys.stdout.write(json.dumps(result) + "\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
