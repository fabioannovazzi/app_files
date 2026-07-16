"""Package Codex-written Deep Research validation outputs."""

from __future__ import annotations

import argparse
import json
import re
import sys
import tempfile
from pathlib import Path
from typing import Any

try:
    import pypandoc  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - optional dependency
    pypandoc = None  # type: ignore[assignment]

try:
    from docx import Document  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore[assignment]

try:
    from .review_session import write_review_session_artifacts, write_run_intake
except ImportError:  # pragma: no cover - supports direct script imports
    import importlib.util

    _review_session_path = Path(__file__).resolve().parent / "review_session.py"
    _review_session_spec = importlib.util.spec_from_file_location(
        "mparanza_deep_research_validator_review_session",
        _review_session_path,
    )
    assert _review_session_spec and _review_session_spec.loader
    _review_session = importlib.util.module_from_spec(_review_session_spec)
    sys.modules[_review_session_spec.name] = _review_session
    _review_session_spec.loader.exec_module(_review_session)
    write_review_session_artifacts = _review_session.write_review_session_artifacts
    write_run_intake = _review_session.write_run_intake

__all__ = [
    "build_audit",
    "render_validation_package",
    "try_write_docx",
    "write_validation_package",
]

ALLOWED_VERDICTS = {
    "supported",
    "partially_supported",
    "not_supported",
    "contradicted",
    "uncertain",
}


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _write_json(path: Path, payload: dict[str, Any]) -> None:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _find_quote_in_sources(quote: str, sources: list[dict[str, Any]]) -> bool:
    target = re.sub(r"\s+", " ", str(quote or "")).strip().casefold()
    if not target:
        return False
    for source in sources:
        excerpt = re.sub(r"\s+", " ", str(source.get("excerpt", ""))).casefold()
        if target in excerpt:
            return True
    return False


def build_audit(
    document_inventory: dict[str, Any],
    source_inventory: dict[str, Any],
    claims_review: dict[str, Any],
) -> dict[str, Any]:
    """Return deterministic audit for Codex-written review JSON."""

    claims = claims_review.get("claims", [])
    if not isinstance(claims, list):
        claims = []
    sources = source_inventory.get("sources", [])
    if not isinstance(sources, list):
        sources = []

    failed_checks: list[str] = []
    if int(document_inventory.get("character_count", 0) or 0) <= 0:
        failed_checks.append("document_text_present")
    if not claims:
        failed_checks.append("claims_review_present")

    invalid_claims: list[int] = []
    missing_claim_text: list[int] = []
    missing_review: list[int] = []
    quote_matches: list[dict[str, Any]] = []

    for position, claim in enumerate(claims, start=1):
        claim_index = int(claim.get("claim_index") or position)
        verdict = str(claim.get("verdict") or "").strip()
        if verdict not in ALLOWED_VERDICTS:
            invalid_claims.append(claim_index)
        if not str(claim.get("claim_text") or "").strip():
            missing_claim_text.append(claim_index)
        if not (
            str(claim.get("source_support") or "").strip()
            or str(claim.get("reasoning_review") or "").strip()
            or str(claim.get("proposed_fix") or "").strip()
        ):
            missing_review.append(claim_index)
        source_quote = str(claim.get("source_quote") or "").strip()
        if source_quote:
            quote_matches.append(
                {
                    "claim_index": claim_index,
                    "matched": _find_quote_in_sources(source_quote, sources),
                }
            )

    if invalid_claims:
        failed_checks.append("valid_verdicts")
    if missing_claim_text:
        failed_checks.append("claim_text_present")
    if missing_review:
        failed_checks.append("review_text_present")

    attention_claims = [
        int(claim.get("claim_index") or index)
        for index, claim in enumerate(claims, start=1)
        if str(claim.get("verdict") or "")
        in {"partially_supported", "not_supported", "contradicted", "uncertain"}
    ]
    return {
        "status": "pass" if not failed_checks else "fail",
        "failed_checks": failed_checks,
        "claim_count": len(claims),
        "attention_claim_indices": attention_claims,
        "invalid_claim_indices": invalid_claims,
        "missing_claim_text_indices": missing_claim_text,
        "missing_review_indices": missing_review,
        "quote_matches": quote_matches,
        "source_count": len(sources),
        "document_url_count": len(document_inventory.get("urls", []) or []),
    }


def _package_markdown(
    document_inventory: dict[str, Any],
    source_inventory: dict[str, Any],
    claims_review: dict[str, Any],
    audit: dict[str, Any],
    validated_document: str,
) -> str:
    claims = claims_review.get("claims", [])
    claim_lines: list[str] = []
    if isinstance(claims, list):
        for claim in claims:
            claim_lines.append(
                "\n".join(
                    [
                        f"### Claim {claim.get('claim_index', '')}".strip(),
                        f"Verdict: {claim.get('verdict', '')}",
                        str(claim.get("claim_text", "")).strip(),
                        f"Source support: {claim.get('source_support', '')}",
                        f"Reasoning: {claim.get('reasoning_review', '')}",
                        f"Proposed fix: {claim.get('proposed_fix', '')}",
                    ]
                ).strip()
            )
    sections = [
        "# Deep Research Validation Package",
        f"Audit status: {audit.get('status')}",
        f"Claims reviewed: {audit.get('claim_count')}",
        f"Sources inspected: {audit.get('source_count')}",
        "## Document Inventory",
        f"Words: {document_inventory.get('word_count', 0)}",
        f"URLs: {document_inventory.get('urls', [])}",
        "## Source Inventory",
        json.dumps(source_inventory, ensure_ascii=False, indent=2),
        "## Claims Review",
        "\n\n".join(claim_lines) if claim_lines else "No claims reviewed.",
    ]
    if validated_document.strip():
        sections.extend(["## Validated Document", validated_document.strip()])
    return "\n\n".join(sections).strip() + "\n"


def render_validation_package(
    document_inventory: dict[str, Any],
    source_inventory: dict[str, Any],
    claims_review: dict[str, Any],
    audit: dict[str, Any],
    validated_document: str,
) -> str:
    """Render package Markdown from already-reviewed validation records."""

    return _package_markdown(
        document_inventory,
        source_inventory,
        claims_review,
        audit,
        validated_document,
    )


def _write_docx_fallback(markdown_text: str, output_path: Path) -> bool:
    if Document is None:
        return False
    document = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading_match = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading_match:
            document.add_heading(
                heading_match.group(2).strip(),
                level=min(len(heading_match.group(1)), 4),
            )
            continue
        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue
        document.add_paragraph(line)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return True


def try_write_docx(markdown_text: str, output_path: Path) -> bool:
    """Write a DOCX from Markdown using local deterministic renderers only."""

    if pypandoc is None:
        return _write_docx_fallback(markdown_text, output_path)
    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmpfile:
            temp_path = Path(tmpfile.name)
        pypandoc.convert_text(
            markdown_text,
            to="docx",
            format="md",
            outputfile=str(temp_path),
        )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(temp_path.read_bytes())
        return True
    except (OSError, RuntimeError):
        return _write_docx_fallback(markdown_text, output_path)
    finally:
        if temp_path and temp_path.exists():
            temp_path.unlink()


def write_validation_package(
    document_inventory_path: Path,
    source_inventory_path: Path,
    claims_review_path: Path,
    output_dir: Path,
    *,
    validated_document_path: Path | None = None,
    write_docx: bool = False,
) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    document_inventory = _read_json(document_inventory_path)
    source_inventory = _read_json(source_inventory_path)
    claims_review = _read_json(claims_review_path)
    run_intake = write_run_intake(
        output_dir,
        document_inventory_path=document_inventory_path,
        source_inventory_path=source_inventory_path,
        claims_review_path=claims_review_path,
        document_inventory=document_inventory,
        source_inventory=source_inventory,
        claims_review=claims_review,
    )
    validated_document = (
        validated_document_path.read_text(encoding="utf-8")
        if validated_document_path and validated_document_path.exists()
        else str(claims_review.get("validated_document", "") or "")
    )

    audit = build_audit(document_inventory, source_inventory, claims_review)
    audit_path = output_dir / "validation_audit.json"
    review_out_path = output_dir / "claims_review.json"
    validated_md_path = output_dir / "validated_document.md"
    package_path = output_dir / "validation_package.md"
    _write_json(audit_path, audit)
    _write_json(review_out_path, claims_review)
    validated_md_path.write_text(validated_document.strip() + "\n", encoding="utf-8")
    package_text = render_validation_package(
        document_inventory,
        source_inventory,
        claims_review,
        audit,
        validated_document,
    )
    package_path.write_text(package_text, encoding="utf-8")

    paths = {
        "claims_review": review_out_path,
        "validation_audit": audit_path,
        "validated_document": validated_md_path,
        "validation_package": package_path,
    }
    if write_docx:
        docx_path = output_dir / "validated_document.docx"
        if try_write_docx(validated_document or package_text, docx_path):
            paths["validated_document_docx"] = docx_path
    review_session = write_review_session_artifacts(
        output_dir,
        run_id=run_intake.run_id,
        run_intake_path=run_intake.path,
        document_inventory_path=document_inventory_path,
        source_inventory_path=source_inventory_path,
        claims_review_path=claims_review_path,
        document_inventory=document_inventory,
        source_inventory=source_inventory,
        claims_review=claims_review,
        audit=audit,
        paths=paths,
    )
    audit["review_session"] = {
        "run_id": review_session.run_id,
        "run_intake_path": str(review_session.run_intake_path),
        "review_payload_path": str(review_session.review_payload_path),
        "ui_decisions_path": str(review_session.ui_decisions_path),
        "final_artifacts_path": str(review_session.final_artifacts_path),
        "review_item_count": review_session.review_item_count,
    }
    _write_json(audit_path, audit)
    return paths


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document_inventory", type=Path)
    parser.add_argument("source_inventory", type=Path)
    parser.add_argument("claims_review", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--validated-document", type=Path)
    parser.add_argument("--docx", action="store_true")
    args = parser.parse_args()
    write_validation_package(
        args.document_inventory,
        args.source_inventory,
        args.claims_review,
        args.output_dir,
        validated_document_path=args.validated_document,
        write_docx=args.docx,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
