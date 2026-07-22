"""Output helpers for Codex audit reconciliation workpapers.

These functions are support utilities for a Codex skill. They are not a
standalone CLI. A case-specific reconciliation script can import them to produce
the required Excel and Word deliverables from deterministic reconciliation rows.
"""

from __future__ import annotations

import json
import sys
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

try:
    from .locale_support import output_labels
except ImportError:  # pragma: no cover - direct import support
    scripts_dir = Path(__file__).resolve().parent
    if str(scripts_dir) not in sys.path:
        sys.path.insert(0, str(scripts_dir))
    from locale_support import output_labels  # type: ignore


NAVY = "17365D"
LIGHT_BLUE = "E8EEF7"
INK = "1F2937"
MUTED = "64748B"
GRAY_FILL = "F3F6FA"
AMBER_FILL = "FFF7E6"


REPORT_TEXT: dict[str, dict[str, str]] = {
    "it": {
        "executive_summary": "Sintesi esecutiva",
        "scope_method": "Perimetro e metodo",
        "how_to_read": "Come leggere gli esiti",
        "external_evidence": "Evidenza esterna considerata",
        "rollforward": "Controllo saldo / roll-forward",
        "account_rollforward": "Controllo saldi da mastro e giornale",
        "post_cutoff": "Evidenze successive al cut-off",
        "post_cutoff_detail": "Principali candidati",
        "additional_analyses": "Analisi deterministiche aggiuntive",
        "aging": "Aging partite aperte",
        "evidence_concentration": "Concentrazione per tipo evidenza",
        "review_signals": "Righe prioritarie da leggere",
        "document_source_map": "Mappa documento-fonti",
        "reversal_candidates": "Possibili storni, giroconti o compensazioni",
        "cutoff_window": "Movimenti vicino al cut-off",
        "checks": "Controlli automatici",
        "review": "Revisione manuale Codex",
        "limitations": "Limiti della procedura",
        "excel_reference": "Rinvio al file Excel",
        "scope_copy": (
            "La procedura riconcilia le partite aperte con le evidenze contabili, "
            "bancarie e documentali disponibili. Le regole sono deterministiche: "
            "una riga viene chiusa solo quando esiste una regola riproducibile e "
            "un riferimento documentale conservato nel workbook."
        ),
        "executive_copy": (
            "Questo documento riassume il risultato della riconciliazione e spiega "
            "come interpretare le categorie. Non sostituisce il dettaglio riga per "
            "riga: il file Excel resta il workpaper principale per verificare fonti, "
            "regole applicate, importi e note di revisione."
        ),
        "status_closed": (
            "Chiusa da evidenza qualificata: la riga risulta superata o chiusa "
            "secondo le regole e le assunzioni indicate."
        ),
        "status_probable_payment": (
            "Pagamento probabile da verificare: esiste un movimento bancario "
            "probabile collegato alla riga, ma manca ancora conferma "
            "dell'allocazione riga/documento per trattarlo come chiusura definitiva."
        ),
        "status_open_supported": (
            "Aperta ma supportata: la posizione resta supportata dai documenti "
            "disponibili; non è stata trovata evidenza esterna specifica che la chiuda."
        ),
        "status_needs_evidence": (
            "Serve evidenza aggiuntiva: esiste una possibile spiegazione, ma manca "
            "ancora un documento necessario per trattarla come evidenza forte."
        ),
        "status_unresolved": (
            "Non risolta: la procedura non ha trovato collegamenti sufficienti; "
            "serve revisione manuale o documentazione aggiuntiva."
        ),
        "external_copy": (
            "Questa sezione riassume le evidenze esterne aggregate, ad esempio banca, "
            "factor, operatori di pagamento, distinte o compensazioni. Gli importi "
            "aggregati aiutano a capire il quadro complessivo, ma la conclusione "
            "auditabile resta riga per riga nel workbook."
        ),
        "rollforward_copy": (
            "Questa sezione confronta i movimenti contabili disponibili con il saldo "
            "ricostruito. Serve come controllo di coerenza aggregato e non come prova "
            "sostitutiva dei singoli match."
        ),
        "account_rollforward_copy": (
            "Questa sezione confronta, per ciascun conto significativo, saldo di apertura "
            "e saldo finale da mastro con i movimenti del giornale. Uno scostamento indica "
            "che il perimetro dei movimenti non spiega interamente il saldo contabile e va "
            "indagato prima di trarre conclusioni."
        ),
        "post_cutoff_copy": (
            "Questa sezione riassume le evidenze successive alla data di riferimento che "
            "condividono il riferimento documento con una partita aperta. Sono utili per "
            "spiegare incassi, pagamenti o chiusure successive, ma non chiudono da sole la "
            "posizione alla data di cut-off."
        ),
        "additional_analyses_copy": (
            "Le analisi seguenti sono controlli esplorativi deterministici: evidenziano "
            "concentrazioni, eccezioni e collegamenti tra fonti. Non modificano da sole "
            "lo stato di riconciliazione delle singole righe."
        ),
        "checks_copy": (
            "I controlli automatici verificano completezza tecnica, coerenza delle "
            "righe elaborate e presenza del pacchetto di revisione. Eventuali WARN o "
            "FAIL devono essere risolti o motivati prima di usare il lavoro come base conclusiva."
        ),
        "review_copy": (
            "Il pacchetto di revisione seleziona righe ad alto valore, righe dubbie "
            "e un campione riproducibile. Le righe PENDING richiedono conferma o nota "
            "manuale prima di considerare la revisione completa."
        ),
        "limitations_copy": (
            "Una riga non riconciliata non dimostra automaticamente un errore o una "
            "irregolarità. Significa che, nei documenti disponibili e con le regole "
            "applicate, non è stata trovata evidenza sufficiente. Per chiuderla servono "
            "documenti sorgente, banca/factor, compensazioni o dettagli di allocazione."
        ),
        "excel_copy": (
            "Il dettaglio auditabile è nel file Excel: ogni riga conserva fonte, regola, "
            "stato, importi, evidenza e motivazione del trattamento."
        ),
        "no_rows": "Nessuna riga disponibile per questa sezione.",
        "rows": "righe",
        "amount": "importo",
    },
    "es": {
        "executive_summary": "Resumen ejecutivo",
        "scope_method": "Alcance y método",
        "how_to_read": "Cómo interpretar los resultados",
        "external_evidence": "Evidencia externa considerada",
        "rollforward": "Control de saldo / roll-forward",
        "account_rollforward": "Control de saldos del mayor y el diario",
        "post_cutoff": "Evidencia posterior a la fecha de corte",
        "post_cutoff_detail": "Principales candidatos",
        "additional_analyses": "Análisis deterministas adicionales",
        "aging": "Antigüedad de partidas abiertas",
        "evidence_concentration": "Concentración por tipo de evidencia",
        "review_signals": "Líneas prioritarias para revisión",
        "document_source_map": "Mapa de documentos y fuentes",
        "reversal_candidates": "Posibles reversiones, traspasos o compensaciones",
        "cutoff_window": "Movimientos próximos a la fecha de corte",
        "checks": "Controles automáticos",
        "review": "Revisión manual de Codex",
        "limitations": "Limitaciones del procedimiento",
        "excel_reference": "Referencia al archivo Excel",
        "scope_copy": (
            "El procedimiento concilia las partidas abiertas con la evidencia "
            "contable, bancaria y documental disponible. Las reglas son "
            "deterministas: una partida solo se cierra cuando se conserva en el "
            "libro una regla reproducible y una referencia a la fuente."
        ),
        "executive_copy": (
            "Este documento resume el resultado de la conciliación y explica cómo "
            "interpretar las categorías. No sustituye la revisión línea por línea: "
            "el libro Excel sigue siendo el papel de trabajo principal para "
            "comprobar fuentes, reglas, importes y notas de revisión."
        ),
        "status_closed": (
            "Cerrada por evidencia cualificada: la línea queda saldada o cerrada "
            "conforme a las reglas y los supuestos indicados."
        ),
        "status_probable_payment": (
            "Pago probable pendiente de verificación: existe un movimiento bancario "
            "probable vinculado a la línea, pero aún debe confirmarse su asignación "
            "a la línea o al documento antes de considerarla cerrada definitivamente."
        ),
        "status_open_supported": (
            "Abierta pero respaldada: la posición sigue respaldada por los documentos "
            "disponibles; no se encontró evidencia externa específica que la cierre."
        ),
        "status_needs_evidence": (
            "Requiere evidencia adicional: existe una posible explicación, pero aún "
            "falta un documento necesario para considerarla evidencia sólida."
        ),
        "status_unresolved": (
            "No resuelta: el procedimiento no encontró vínculos suficientes; se "
            "requiere revisión manual o documentación adicional."
        ),
        "external_copy": (
            "Esta sección resume la evidencia externa agregada, por ejemplo de bancos, "
            "factores, operadores de pago, lotes de pago o compensaciones. Los importes "
            "agregados ayudan a entender el conjunto, pero la conclusión auditable "
            "permanece línea por línea en el libro."
        ),
        "rollforward_copy": (
            "Esta sección compara los movimientos contables disponibles con el saldo "
            "reconstruido. Es un control de coherencia agregado y no sustituye la "
            "conciliación de cada línea."
        ),
        "account_rollforward_copy": (
            "Esta sección compara, para cada cuenta significativa, los saldos inicial "
            "y final del mayor con los movimientos del diario. Una diferencia indica "
            "que el perímetro de movimientos no explica por completo el saldo contable "
            "y debe investigarse antes de extraer conclusiones."
        ),
        "post_cutoff_copy": (
            "Esta sección resume la evidencia posterior a la fecha de referencia que "
            "comparte una referencia documental con una partida abierta. Puede explicar "
            "cobros, pagos o cancelaciones posteriores, pero no cierra por sí sola la "
            "posición en la fecha de corte."
        ),
        "additional_analyses_copy": (
            "Los análisis siguientes son controles exploratorios deterministas: "
            "señalan concentraciones, excepciones y vínculos entre fuentes. Por sí "
            "solos no modifican el estado de conciliación de las líneas."
        ),
        "checks_copy": (
            "Los controles automáticos comprueban la integridad técnica, la coherencia "
            "de las líneas procesadas y la presencia del paquete de revisión. Los "
            "elementos con AVISO o ERROR deben resolverse o explicarse antes de usar "
            "el trabajo como base para una conclusión."
        ),
        "review_copy": (
            "El paquete de revisión selecciona líneas de importe elevado, líneas "
            "dudosas y una muestra reproducible. Las líneas PENDIENTES requieren "
            "confirmación o una nota manual antes de completar la revisión."
        ),
        "limitations_copy": (
            "Una línea no conciliada no demuestra automáticamente un error o una "
            "irregularidad. Significa que los documentos disponibles y las reglas "
            "aplicadas no aportaron evidencia suficiente. Pueden ser necesarios "
            "documentos fuente, evidencia bancaria o del factor, acuerdos de "
            "compensación o detalles de asignación."
        ),
        "excel_copy": (
            "El detalle auditable línea por línea se encuentra en el libro Excel: "
            "cada línea conserva la fuente, la regla, el estado, el importe, la "
            "evidencia y la justificación del tratamiento."
        ),
        "no_rows": "No hay líneas disponibles para esta sección.",
        "rows": "líneas",
        "amount": "importe",
    },
    "en": {
        "executive_summary": "Executive Summary",
        "scope_method": "Scope and Method",
        "how_to_read": "How to Read the Results",
        "external_evidence": "External Evidence Considered",
        "rollforward": "Balance / Roll-Forward Check",
        "account_rollforward": "Ledger and Journal Balance Check",
        "post_cutoff": "Post-Cutoff Evidence",
        "post_cutoff_detail": "Main Candidates",
        "additional_analyses": "Additional Deterministic Analyses",
        "aging": "Open-Item Aging",
        "evidence_concentration": "Evidence-Type Concentration",
        "review_signals": "Priority Rows to Review",
        "document_source_map": "Document Source Map",
        "reversal_candidates": "Possible Reversals, Transfers or Set-Offs",
        "cutoff_window": "Movements Near Cutoff",
        "checks": "Automated Checks",
        "review": "Codex Manual Review",
        "limitations": "Procedure Limits",
        "excel_reference": "Excel Reference",
        "scope_copy": (
            "The procedure reconciles open items against the available accounting, "
            "banking and documentary evidence. Rules are deterministic: an item is "
            "closed only when a reproducible rule and source reference are retained "
            "in the workbook."
        ),
        "executive_copy": (
            "This document summarizes the reconciliation result and explains how to "
            "interpret the categories. It does not replace row-level review: the Excel "
            "workbook remains the main audit workpaper for sources, rules, amounts and notes."
        ),
        "status_closed": "closed: the row is cleared by qualified evidence under the stated rules and assumptions.",
        "status_probable_payment": "probable_payment: a likely bank movement is linked to the row, but row/document allocation still needs confirmation before treating it as definitively closed.",
        "status_open_supported": "open_supported: the open position remains supported by the available documents.",
        "status_needs_evidence": "needs_evidence: a possible explanation exists, but required evidence is still missing.",
        "status_unresolved": "unresolved: no sufficient link was found; manual review or additional documents are required.",
        "external_copy": (
            "This section summarizes external evidence such as bank, factor, payment "
            "operators, payment batches or documented set-offs. Aggregates help explain "
            "the overall picture, but conclusions remain row-level in the workbook."
        ),
        "rollforward_copy": (
            "This section compares available accounting movements with the reconstructed "
            "balance. It is an aggregate consistency check, not a substitute for row-level matching."
        ),
        "account_rollforward_copy": (
            "This section compares, for each significant account, opening and closing ledger "
            "balances to journal movements. A difference means the movement perimeter does not "
            "fully explain the accounting balance and should be investigated before drawing conclusions."
        ),
        "post_cutoff_copy": (
            "This section summarizes evidence after the reporting date that shares a document "
            "reference with an open item. It can explain later receipts, payments or clearings, "
            "but it does not clear the position at the cutoff date by itself."
        ),
        "additional_analyses_copy": (
            "The following analyses are deterministic exploratory controls: they highlight "
            "concentrations, exceptions and cross-source links. They do not change row-level "
            "reconciliation status by themselves."
        ),
        "checks_copy": (
            "Automated checks test technical completeness, row consistency and review packet "
            "presence. WARN or FAIL items should be resolved or explained before relying on the output."
        ),
        "review_copy": (
            "The review packet selects high-value, doubtful and reproducible sample rows. "
            "PENDING rows require manual confirmation or notes before the review is complete."
        ),
        "limitations_copy": (
            "An unreconciled row is not automatic proof of error or irregularity. It means "
            "the available documents and rules did not provide sufficient evidence. Source "
            "documents, bank/factor evidence, set-off agreements or allocation detail may be required."
        ),
        "excel_copy": (
            "The auditable row-level detail is in the Excel workbook: each row retains source, "
            "rule, status, amount, evidence and treatment rationale."
        ),
        "no_rows": "No rows available for this section.",
        "rows": "rows",
        "amount": "amount",
    },
}

STANDARD_SHEET_ORDER = [
    "Index",
    "Assumptions",
    "Source inventory",
    "Normalized records",
    "Reconciliation detail",
    "Bank allocation candidates",
    "External evidence aggregate",
    "External evidence detail",
    "Ledger balance check",
    "Account rollforward check",
    "Journal rollforward",
    "Journal detail",
    "Post-cutoff candidates",
    "Open item aging",
    "Evidence concentration",
    "Review signals",
    "Document source map",
    "Reversal candidates",
    "Cutoff window movements",
    "Summary",
    "Checks",
    "Review",
]

ITALIAN_SHEET_LABELS = {
    "Index": "Indice",
    "Assumptions": "Assunzioni",
    "Source inventory": "Inventario fonti",
    "Normalized records": "Righe normalizzate",
    "Reconciliation detail": "Dettaglio riconciliazione",
    "Bank allocation candidates": "Candidati banca",
    "External evidence aggregate": "Evidenze esterne agg.",
    "External evidence detail": "Dettaglio evidenze esterne",
    "Ledger balance check": "Controllo saldi mastro",
    "Account rollforward check": "Roll-forward conti",
    "Journal rollforward": "Roll-forward giornale",
    "Journal detail": "Dettaglio giornale",
    "Post-cutoff candidates": "Candidati post cut-off",
    "Open item aging": "Aging partite",
    "Evidence concentration": "Concentrazione evidenze",
    "Review signals": "Segnali revisione",
    "Document source map": "Mappa documento-fonti",
    "Reversal candidates": "Candidati storni",
    "Cutoff window movements": "Movimenti cut-off",
    "Summary": "Sintesi",
    "Checks": "Controlli",
    "Review": "Revisione Codex",
}

ITALIAN_FIELD_LABELS = {
    "account": "Conto",
    "account_name": "Nome conto",
    "age_days_at_reference": "Giorni alla data di riferimento",
    "aging_bucket": "Fascia aging",
    "amount": "Importo",
    "amount_abs_total": "Importo assoluto",
    "bank_rows": "Righe banca",
    "candidate_id": "ID candidato",
    "candidate_reasons": "Motivi candidato",
    "candidate_type": "Tipo candidato",
    "check": "Controllo",
    "closing_balance_signed_debit_minus_credit": "Saldo finale dare meno avere",
    "closing_difference_journal_minus_ledger": "Differenza finale giornale-mastro",
    "closing_net_debit_minus_credit": "Saldo finale netto dare meno avere",
    "cutoff_date": "Data di cut-off",
    "cutoff_window_timing": "Periodo cut-off",
    "description": "Descrizione",
    "document_date": "Data documento",
    "document_key": "Chiave documento",
    "document_no": "Documento",
    "document_no_examples": "Esempi documento",
    "evidence_amount": "Importo evidenza",
    "evidence_concentration": "Concentrazione evidenze",
    "evidence_date": "Data evidenza",
    "evidence_level": "Forza evidenza",
    "evidence_source_file": "File fonte evidenza",
    "evidence_type": "Tipo evidenza",
    "exact_amount_match": "Importo esatto",
    "expected": "Atteso",
    "factoring_pro_soluto_closes_item": "Factoring pro-soluto chiude con evidenza esterna allocata",
    "factoring_rows": "Righe factor",
    "files": "File",
    "journal_rows": "Righe giornale",
    "ledger_rows": "Righe mastro",
    "matched_evidence_amounts": "Importi evidenza abbinata",
    "matched_evidence_id": "ID evidenza abbinata",
    "matched_evidence_reference": "Riferimento evidenza abbinata",
    "matched_evidence_type": "Tipo evidenza abbinata",
    "message": "Messaggio",
    "metadata_field": "Campo",
    "metadata_value": "Valore",
    "missing_evidence": "Evidenza mancante",
    "open_amount": "Importo aperto",
    "open_amount_total": "Importo aperto totale",
    "open_item_rows": "Righe partite aperte",
    "opening_difference_journal_minus_ledger": "Differenza apertura giornale-mastro",
    "payment_order_rows": "Righe distinte",
    "post_cutoff_events_excluded": "Eventi post cut-off esclusi dalla chiusura",
    "prior_reconciliation_status": "Esito precedente",
    "record_id": "ID riga",
    "reconciliation_status": "Esito riconciliazione",
    "reconciliation_status_counts": "Esiti riconciliazione",
    "review_note": "Nota revisione",
    "review_result": "Esito revisione",
    "review_signal_rank": "Priorita revisione",
    "review_signals": "Segnali revisione",
    "review_status": "Esito revisione",
    "rows": "Righe",
    "rule_applied": "Regola applicata",
    "share_of_abs_amount_percent": "Peso %",
    "sheet": "Foglio",
    "source_file": "File origine",
    "source_page": "Pagina origine",
    "source_role": "Ruolo fonte",
    "source_row": "Riga origine",
    "source_sheet": "Foglio origine",
    "status": "Esito",
    "support_bucket": "Tipo supporto",
    "value": "Valore",
}

ITALIAN_RECONCILIATION_STATUS_LABELS = {
    "closed": "Chiusa da evidenza",
    "probable_payment": "Pagamento probabile da verificare",
    "open_supported": "Aperta ma supportata",
    "needs_evidence": "Serve evidenza aggiuntiva",
    "unresolved": "Non risolta",
    "out_of_scope": "Fuori perimetro",
    "unknown": "Non classificata",
}

ITALIAN_RULE_LABELS = {
    "direct_external_or_documented": "Evidenza esterna o documentata diretta",
    "documented_compensation": "Compensazione documentata",
    "external_bank_match": "Movimento bancario esterno abbinato",
    "factoring_bridge_only": "Solo ponte factor/anticipo",
    "factoring_or_advance_match": "Factor o anticipo abbinato",
    "factoring_with_bank_or_external_support": "Factor con supporto bancario o esterno",
    "grouped_open_amount_internal_booking_support": "Supporto interno aggregato su documento",
    "grouped_payment_external_match": "Pagamento aggregato con evidenza esterna",
    "internal_accounting_only": "Solo scrittura contabile interna",
    "internal_booking_open_support": "Supporto interno a partita aperta",
    "internal_closure_without_external": "Chiusura interna senza evidenza esterna",
    "out_of_scope": "Fuori perimetro",
    "payment_order_amount_mismatch": "Distinta con importo non allineato",
    "payment_order_only": "Solo distinta o ordine di pagamento",
    "probable_bank_payment_candidate": "Pagamento bancario probabile",
    "unallocated_external_bank_requires_allocation": "Banca esterna non allocata",
    "unresolved": "Non risolta",
    "unknown": "Non classificata",
}

ITALIAN_SOURCE_ROLE_LABELS = {
    "bank_statement": "Estratto conto bancario",
    "factoring_statement": "Estratto factor / anticipo",
    "journal": "Libro giornale",
    "ledger": "Mastrino",
    "open_items": "Partite aperte",
    "payment_order": "Distinta / ordine di pagamento",
    "unknown": "Non classificata",
}

ITALIAN_EVIDENCE_LEVEL_LABELS = {
    "strong_external": "Evidenza esterna forte",
    "external": "Evidenza esterna",
    "bridge": "Evidenza ponte",
    "internal": "Evidenza interna",
    "none": "Nessuna evidenza",
    "unknown": "Non classificata",
}

ITALIAN_REVIEW_SIGNAL_LABELS = {
    "before_cutoff": "Prima del cut-off",
    "exact_amount_match": "Importo esatto",
    "high_value": "Importo elevato",
    "needs_evidence": "Serve evidenza",
    "open_supported": "Aperta supportata",
    "opposite_sign_amount": "Importo di segno opposto",
    "probable_payment": "Pagamento probabile",
    "stale_open_item": "Partita datata",
    "unresolved": "Non risolta",
}

ITALIAN_CUTOFF_TIMING_LABELS = {
    "before_cutoff": "Prima del cut-off",
    "after_cutoff": "Dopo il cut-off",
    "on_cutoff": "Alla data di cut-off",
    "unknown": "Non classificato",
}

ITALIAN_REVIEW_STATUS_LABELS = {
    "APPROVED": "Approvata",
    "CHALLENGED": "Da chiarire",
    "FAIL": "Errore",
    "PASS": "OK",
    "PENDING": "Da completare",
    "UNKNOWN": "Non indicata",
}

SPANISH_SHEET_LABELS = {
    "Index": "Índice",
    "Assumptions": "Supuestos",
    "Source inventory": "Inventario de fuentes",
    "Normalized records": "Registros normalizados",
    "Reconciliation detail": "Detalle de conciliación",
    "Bank allocation candidates": "Candidatos bancarios",
    "External evidence aggregate": "Resumen evidencia externa",
    "External evidence detail": "Detalle evidencia externa",
    "Ledger balance check": "Control de saldos del mayor",
    "Account rollforward check": "Roll-forward de cuentas",
    "Journal rollforward": "Roll-forward del diario",
    "Journal detail": "Detalle del diario",
    "Post-cutoff candidates": "Candidatos post-corte",
    "Open item aging": "Antigüedad de partidas",
    "Evidence concentration": "Concentración de evidencias",
    "Review signals": "Señales de revisión",
    "Document source map": "Mapa documento-fuentes",
    "Reversal candidates": "Candidatos a reversión",
    "Cutoff window movements": "Movimientos cerca del corte",
    "Summary": "Resumen",
    "Checks": "Controles",
    "Review": "Revisión Codex",
}

SPANISH_FIELD_LABELS = {
    "account": "Cuenta",
    "account_name": "Nombre de la cuenta",
    "age_days_at_reference": "Días a la fecha de referencia",
    "aging_bucket": "Tramo de antigüedad",
    "amount": "Importe",
    "amount_abs_total": "Importe absoluto",
    "bank_rows": "Líneas bancarias",
    "candidate_id": "ID del candidato",
    "candidate_reasons": "Motivos del candidato",
    "candidate_type": "Tipo de candidato",
    "check": "Control",
    "compensation_requires_bank": "La compensación requiere evidencia bancaria",
    "closing_balance_signed_debit_minus_credit": "Saldo final debe menos haber",
    "closing_difference_journal_minus_ledger": "Diferencia final diario-mayor",
    "closing_net_debit_minus_credit": "Saldo final neto debe menos haber",
    "cutoff_date": "Fecha de corte",
    "cutoff_window_timing": "Periodo respecto al corte",
    "description": "Descripción",
    "document_date": "Fecha del documento",
    "document_key": "Clave del documento",
    "document_no": "Documento",
    "document_no_examples": "Ejemplos de documentos",
    "evidence_amount": "Importe de la evidencia",
    "evidence_concentration": "Concentración de evidencias",
    "evidence_date": "Fecha de la evidencia",
    "evidence_level": "Solidez de la evidencia",
    "evidence_source_file": "Archivo fuente de la evidencia",
    "evidence_type": "Tipo de evidencia",
    "exact_amount_match": "Importe exacto",
    "expected": "Esperado",
    "factoring_pro_soluto_closes_item": "El factoring sin recurso cierra con evidencia externa asignada",
    "factoring_rows": "Líneas del factor",
    "files": "Archivos",
    "currency": "Moneda",
    "document_date_tolerance_days": "Tolerancia de fecha del documento en días",
    "document_language": "Idioma del documento",
    "factoring_operator_keywords": "Términos del factor u operador",
    "journal_rows": "Líneas del diario",
    "ledger_rows": "Líneas del mayor",
    "matched_evidence_amounts": "Importes de evidencia conciliada",
    "matched_evidence_id": "ID de evidencia conciliada",
    "matched_evidence_reference": "Referencia de evidencia conciliada",
    "matched_evidence_type": "Tipo de evidencia conciliada",
    "message": "Mensaje",
    "metadata_field": "Campo",
    "metadata_value": "Valor",
    "missing_evidence": "Evidencia pendiente",
    "open_amount": "Importe abierto",
    "open_amount_total": "Importe abierto total",
    "open_item_rows": "Líneas de partidas abiertas",
    "opening_difference_journal_minus_ledger": "Diferencia inicial diario-mayor",
    "payment_order_rows": "Líneas de órdenes de pago",
    "payment_orders_are_bank_evidence": "Las órdenes de pago constituyen evidencia bancaria",
    "post_cutoff_events_excluded": "Eventos posteriores al corte excluidos del cierre",
    "probable_bank_exact_matches_close": "Las coincidencias bancarias probables exactas cierran la partida",
    "promote_probable_bank_payments": "Promover los pagos bancarios probables",
    "prior_reconciliation_status": "Estado anterior",
    "record_id": "ID de línea",
    "reconciliation_status": "Estado de conciliación",
    "reconciliation_status_counts": "Estados de conciliación",
    "review_note": "Nota de revisión",
    "review_result": "Resultado de la revisión",
    "review_signal_rank": "Prioridad de revisión",
    "review_signals": "Señales de revisión",
    "review_status": "Estado de la revisión",
    "rows": "Líneas",
    "rule_applied": "Regla aplicada",
    "scope_year": "Año del alcance",
    "share_of_abs_amount_percent": "Peso %",
    "sheet": "Hoja",
    "source_file": "Archivo de origen",
    "source_page": "Página de origen",
    "source_role": "Función de la fuente",
    "source_row": "Línea de origen",
    "source_sheet": "Hoja de origen",
    "status": "Estado",
    "support_bucket": "Tipo de respaldo",
    "report_language": "Idioma del informe",
    "amount_tolerance": "Tolerancia de importe",
    "value": "Valor",
}

SPANISH_RECONCILIATION_STATUS_LABELS = {
    "closed": "Cerrada por evidencia",
    "probable_payment": "Pago probable pendiente de verificación",
    "open_supported": "Abierta pero respaldada",
    "needs_evidence": "Requiere evidencia adicional",
    "unresolved": "No resuelta",
    "out_of_scope": "Fuera de alcance",
    "unknown": "Sin clasificar",
}

SPANISH_RULE_LABELS = {
    "direct_external_or_documented": "Evidencia directa externa o documentada",
    "documented_compensation": "Compensación documentada",
    "external_bank_match": "Movimiento bancario externo conciliado",
    "factoring_bridge_only": "Solo evidencia puente del factor o anticipo",
    "factoring_or_advance_match": "Factor o anticipo conciliado",
    "factoring_with_bank_or_external_support": "Factor con respaldo bancario o externo",
    "grouped_open_amount_internal_booking_support": "Respaldo interno agregado por documento",
    "grouped_payment_external_match": "Pago agregado con evidencia externa",
    "internal_accounting_only": "Solo asiento contable interno",
    "internal_booking_open_support": "Respaldo interno de partida abierta",
    "internal_closure_without_external": "Cierre interno sin evidencia externa",
    "out_of_scope": "Fuera de alcance",
    "payment_order_amount_mismatch": "Orden de pago con importe no coincidente",
    "payment_order_only": "Solo orden o lote de pago",
    "probable_bank_payment_candidate": "Pago bancario probable",
    "unallocated_external_bank_requires_allocation": "Movimiento bancario externo sin asignar",
    "unresolved": "No resuelta",
    "unknown": "Sin clasificar",
}

SPANISH_SOURCE_ROLE_LABELS = {
    "bank_statement": "Extracto bancario",
    "factoring_statement": "Extracto del factor o anticipo",
    "journal": "Libro diario",
    "ledger": "Libro mayor",
    "open_items": "Partidas abiertas",
    "payment_order": "Orden o lote de pago",
    "unknown": "Sin clasificar",
}

SPANISH_EVIDENCE_LEVEL_LABELS = {
    "strong_external": "Evidencia externa sólida",
    "external": "Evidencia externa",
    "bridge": "Evidencia puente",
    "internal": "Evidencia interna",
    "none": "Sin evidencia",
    "unknown": "Sin clasificar",
}

SPANISH_REVIEW_SIGNAL_LABELS = {
    "before_cutoff": "Antes de la fecha de corte",
    "exact_amount_match": "Importe exacto",
    "high_value": "Importe elevado",
    "needs_evidence": "Requiere evidencia",
    "open_supported": "Abierta pero respaldada",
    "opposite_sign_amount": "Importe de signo opuesto",
    "probable_payment": "Pago probable",
    "stale_open_item": "Partida antigua",
    "unresolved": "No resuelta",
}

SPANISH_CUTOFF_TIMING_LABELS = {
    "before_cutoff": "Antes de la fecha de corte",
    "after_cutoff": "Después de la fecha de corte",
    "on_cutoff": "En la fecha de corte",
    "unknown": "Sin clasificar",
}

SPANISH_REVIEW_STATUS_LABELS = {
    "APPROVED": "Aprobada",
    "CHALLENGED": "Pendiente de aclaración",
    "FAIL": "Error",
    "PASS": "Correcta",
    "PENDING": "Pendiente",
    "UNKNOWN": "No indicada",
}


def _headers(
    rows: list[dict[str, Any]], preferred: list[str] | None = None
) -> list[str]:
    if preferred:
        tail = sorted({key for row in rows for key in row.keys()} - set(preferred))
        return preferred + tail
    return sorted({key for row in rows for key in row.keys()}) if rows else ["message"]


def _decimal(value: object) -> Decimal:
    text = str(value or "").strip().replace(",", ".")
    if not text:
        return Decimal("0.00")
    try:
        return Decimal(text)
    except InvalidOperation:
        return Decimal("0.00")


def _language_code(language: object | None) -> str:
    text = str(language or "it").strip().lower().replace("_", "-")
    code = text.split("-", 1)[0]
    return {
        "esp": "es",
        "espanol": "es",
        "español": "es",
        "spa": "es",
        "spanish": "es",
    }.get(code, code)


def _is_italian(language: object | None) -> bool:
    return _language_code(language) == "it"


def _is_spanish(language: object | None) -> bool:
    return _language_code(language) == "es"


def _is_localized(language: object | None) -> bool:
    return _language_code(language) in {"it", "es"}


def _locale_catalog(
    language: object | None,
    italian: dict[str, str],
    spanish: dict[str, str],
) -> dict[str, str]:
    if _is_italian(language):
        return italian
    if _is_spanish(language):
        return spanish
    return {}


def _labels(language: str) -> dict[str, str]:
    return output_labels(_language_code(language))


def _sheet_label(sheet_name: object, language: str) -> str:
    value = str(sheet_name or "")
    labels = _locale_catalog(language, ITALIAN_SHEET_LABELS, SPANISH_SHEET_LABELS)
    return labels.get(value, value)


def _fallback_label(value: object) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    label = text.replace("_", " ")
    return label[:1].upper() + label[1:]


def _field_label(field: object, language: str) -> str:
    value = str(field or "")
    labels = _locale_catalog(language, ITALIAN_FIELD_LABELS, SPANISH_FIELD_LABELS)
    if not labels:
        return value
    return labels.get(value.lower(), _fallback_label(value))


def _mapped_label(value: object, labels: dict[str, str]) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    return labels.get(text.lower(), _fallback_label(text))


def _reconciliation_status_label(status: object, language: str) -> str:
    labels = _locale_catalog(
        language,
        ITALIAN_RECONCILIATION_STATUS_LABELS,
        SPANISH_RECONCILIATION_STATUS_LABELS,
    )
    if not labels:
        return str(status or "")
    return _mapped_label(status, labels)


def _rule_label(rule: object, language: str) -> str:
    labels = _locale_catalog(language, ITALIAN_RULE_LABELS, SPANISH_RULE_LABELS)
    if not labels:
        return str(rule or "")
    return _mapped_label(rule, labels)


def _source_role_label(source_role: object, language: str) -> str:
    labels = _locale_catalog(
        language, ITALIAN_SOURCE_ROLE_LABELS, SPANISH_SOURCE_ROLE_LABELS
    )
    if not labels:
        return str(source_role or "")
    return _mapped_label(source_role, labels)


def _evidence_level_label(evidence_level: object, language: str) -> str:
    labels = _locale_catalog(
        language, ITALIAN_EVIDENCE_LEVEL_LABELS, SPANISH_EVIDENCE_LEVEL_LABELS
    )
    if not labels:
        return str(evidence_level or "")
    return _mapped_label(evidence_level, labels)


def _review_status_label(review_status: object, language: str) -> str:
    text = str(review_status or "").strip()
    labels = _locale_catalog(
        language, ITALIAN_REVIEW_STATUS_LABELS, SPANISH_REVIEW_STATUS_LABELS
    )
    if not labels:
        return text
    return labels.get(text.upper(), _fallback_label(text))


def _review_signal_label(signal: object, language: str) -> str:
    labels = _locale_catalog(
        language, ITALIAN_REVIEW_SIGNAL_LABELS, SPANISH_REVIEW_SIGNAL_LABELS
    )
    if not labels:
        return str(signal or "")
    return _mapped_label(signal, labels)


def _cutoff_timing_label(timing: object, language: str) -> str:
    labels = _locale_catalog(
        language, ITALIAN_CUTOFF_TIMING_LABELS, SPANISH_CUTOFF_TIMING_LABELS
    )
    if not labels:
        return str(timing or "")
    return _mapped_label(timing, labels)


def _localized_semicolon_list(value: object, language: str, labeler: Any) -> str:
    text = str(value or "").strip()
    if not text or not _is_localized(language):
        return text
    parts = [part.strip() for part in text.split(";") if part.strip()]
    if not parts:
        return ""
    return "; ".join(labeler(part, language) for part in parts)


def _localized_status_counts(value: object, language: str) -> str:
    text = str(value or "").strip()
    if not text or not _is_localized(language):
        return text
    parts: list[str] = []
    for part in text.split(";"):
        chunk = part.strip()
        if not chunk:
            continue
        if ":" not in chunk:
            parts.append(_reconciliation_status_label(chunk, language))
            continue
        status, count = chunk.split(":", 1)
        parts.append(
            f"{_reconciliation_status_label(status.strip(), language)}: {count.strip()}"
        )
    return "; ".join(parts)


def _localize_value(field: object, value: object, language: str) -> object:
    if not _is_localized(language):
        return value
    if value is None:
        return ""

    field_name = str(field or "").lower()
    if isinstance(value, bool):
        return _yes_no_label("YES" if value else "NO", language)
    if field_name in {
        "reconciliation_status",
        "prior_reconciliation_status",
        "deterministic_status",
    }:
        return _reconciliation_status_label(value, language)
    if field_name == "reconciliation_status_counts":
        return _localized_status_counts(value, language)
    if field_name in {"rule_applied", "deterministic_rule"}:
        return _rule_label(value, language)
    if field_name in {
        "evidence_type",
        "matched_evidence_type",
        "prior_matched_evidence_type",
    }:
        return _evidence_type_label(value, language)
    if field_name in {"source_role", "role"}:
        return _source_role_label(value, language)
    if field_name in {"evidence_level", "support_level"}:
        return _evidence_level_label(value, language)
    if field_name == "support_bucket":
        return _support_bucket_label(value, language)
    if field_name == "status":
        return _status_label(value, language)
    if field_name in {"review_status", "review_result"}:
        return _review_status_label(value, language)
    if field_name in {"review_signals", "candidate_reasons"}:
        return _localized_semicolon_list(value, language, _review_signal_label)
    if field_name == "cutoff_window_timing":
        return _cutoff_timing_label(value, language)
    if field_name in {"exact_amount_match", "has_external_evidence"}:
        return _yes_no_label(value, language)
    return value


def _localize_row(row: dict[str, Any], language: str) -> dict[str, Any]:
    if not _is_localized(language):
        return dict(row)
    return {
        _field_label(key, language): _localize_value(key, value, language)
        for key, value in row.items()
    }


def _localize_rows(rows: list[dict[str, Any]], language: str) -> list[dict[str, Any]]:
    return [_localize_row(row, language) for row in rows]


def _localized_excel_sheets(
    sheets: dict[str, list[dict[str, Any]]],
    ordered_names: list[str],
    language: str,
) -> dict[str, list[dict[str, Any]]]:
    if not _is_localized(language):
        return {}

    output = _labels(language)
    field_header = output["metadata_field"]
    value_header = output["metadata_value"]
    sheet_header = _field_label("sheet", language)
    rows_header = _field_label("rows", language)
    localized: dict[str, list[dict[str, Any]]] = {}
    for sheet_name in ordered_names:
        display_name = _sheet_label(sheet_name, language)
        rows = sheets[sheet_name]
        if sheet_name == "Index":
            localized[display_name] = [
                {
                    sheet_header: _sheet_label(row.get("Sheet"), language),
                    rows_header: row.get("Rows", ""),
                }
                for row in rows
            ]
        elif sheet_name == "Assumptions":
            assumption_rows: list[dict[str, Any]] = []
            for row in rows:
                field = next(
                    (
                        row[key]
                        for key in (field_header, "Field", "Campo", "Champ", "Feld")
                        if key in row
                    ),
                    "",
                )
                value = next(
                    (
                        row[key]
                        for key in (
                            value_header,
                            "Value",
                            "Valore",
                            "Valor",
                            "Valeur",
                            "Wert",
                        )
                        if key in row
                    ),
                    "",
                )
                assumption_rows.append(
                    {
                        field_header: _field_label(field, language),
                        value_header: _localize_value(field, value, language),
                    }
                )
            localized[display_name] = assumption_rows
        else:
            localized[display_name] = _localize_rows(rows, language)
    return localized


def _set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_paragraph_spacing(
    paragraph: Any, *, before: float = 0, after: float = 6, line_spacing: float = 1.1
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _configure_word_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in [
        ("Title", 22, INK, 0, 4),
        ("Subtitle", 11, MUTED, 0, 14),
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 13, NAVY, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def _paragraph_border_bottom(
    paragraph: Any, color: str = NAVY, size: str = "12"
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def _cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(
    cell: Any, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _column_widths(headers: list[str]) -> list[float]:
    if not headers:
        return [6.4]
    weights = []
    for header in headers:
        lower = str(header).lower()
        if any(
            token in lower
            for token in ("rows", "righe", "líneas", "files", "archivos", "checks")
        ):
            weights.append(0.65)
        elif any(
            token in lower
            for token in ("amount", "importo", "importe", "status", "stato", "estado")
        ):
            weights.append(1.0)
        elif any(
            token in lower
            for token in (
                "rule",
                "regola",
                "regla",
                "source",
                "fuente",
                "evidence",
                "evidencia",
                "note",
                "nota",
                "expected",
                "esperado",
            )
        ):
            weights.append(2.0)
        else:
            weights.append(1.35)
    total = sum(weights)
    return [round((weight / total) * 6.4, 2) for weight in weights]


def _set_table_geometry(table: Any, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def _format_table(
    table: Any, headers: list[str], *, header_fill: str = GRAY_FILL
) -> None:
    _set_table_geometry(table, _column_widths(headers))
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                _set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.1)
            if row_idx == 0:
                _cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run, bold=True, color=INK)


def _add_callout(
    doc: Document, title: str, body: str, *, fill: str = LIGHT_BLUE
) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _cell_shading(cell, fill)
    _set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    title_p = cell.paragraphs[0]
    title_p.text = title
    _set_paragraph_spacing(title_p, after=3)
    for run in title_p.runs:
        _set_run_font(run, bold=True, color=NAVY)
    body_p = cell.add_paragraph(body)
    _set_paragraph_spacing(body_p, after=0)
    _set_table_geometry(table, [6.4])


def _report_text(language: str) -> dict[str, str]:
    code = _language_code(language)
    return REPORT_TEXT.get(code, REPORT_TEXT["en"])


def _format_decimal(value: object) -> str:
    return f"{_decimal(value):,.2f}"


def _add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        if text:
            paragraph = doc.add_paragraph(text)
            _set_paragraph_spacing(paragraph)


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        if item:
            paragraph = doc.add_paragraph(item, style="List Bullet")
            _set_paragraph_spacing(paragraph, after=4, line_spacing=1.167)


def _add_table(
    doc: Document,
    rows: list[dict[str, Any]],
    preferred: list[str] | None = None,
    limit: int | None = None,
) -> None:
    if not rows:
        return
    visible_rows = rows[:limit] if limit is not None else rows
    headers = _headers(visible_rows, preferred)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for row in visible_rows:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            cells[idx].text = str(row.get(header, ""))
    _format_table(table, headers)


def _summary_by_status(
    summary_rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    text = _report_text(language)
    status_header = _field_label("status", language)
    buckets: dict[str, dict[str, Any]] = {}
    for row in summary_rows:
        status = str(
            row.get("reconciliation_status")
            or row.get("Esito")
            or row.get("Estado")
            or row.get("Status")
            or "unknown"
        )
        bucket = buckets.setdefault(
            status, {"status": status, text["rows"]: 0, text["amount"]: Decimal("0.00")}
        )
        bucket[text["rows"]] += int(
            _decimal(
                row.get("rows")
                or row.get("Righe")
                or row.get("Líneas")
                or row.get("Rows")
                or 0
            )
        )
        bucket[text["amount"]] += _decimal(
            row.get("amount")
            or row.get("Importo")
            or row.get("Importe")
            or row.get("Amount")
        )
    return [
        {
            status_header: _reconciliation_status_label(row["status"], language),
            text["rows"]: row[text["rows"]],
            text["amount"]: _format_decimal(row[text["amount"]]),
        }
        for row in sorted(buckets.values(), key=lambda item: str(item["status"]))
    ]


def _check_summary(
    checks: list[dict[str, Any]], language: str = "en"
) -> list[dict[str, Any]]:
    buckets: dict[str, int] = {}
    for row in checks:
        status = str(row.get("status") or "UNKNOWN")
        buckets[status] = buckets.get(status, 0) + 1
    status_header = _field_label("status", language)
    checks_header = (
        "Controlli"
        if _is_italian(language)
        else ("Controles" if _is_spanish(language) else "checks")
    )
    return [
        {
            status_header: _status_label(status, language),
            checks_header: count,
        }
        for status, count in sorted(buckets.items())
    ]


def _status_label(status: object, language: str) -> str:
    value = str(status or "UNKNOWN").upper()
    if _is_italian(language):
        return {
            "PASS": "OK",
            "DIFFERENCE": "Differenza",
            "FAIL": "Errore",
            "MISSING_JOURNAL_OR_LEDGER": "Manca giornale o mastro",
            "UNKNOWN": "Non indicato",
            "WARN": "Avviso",
        }.get(value, value)
    if _is_spanish(language):
        return {
            "PASS": "Correcto",
            "DIFFERENCE": "Diferencia",
            "FAIL": "Error",
            "MISSING_JOURNAL_OR_LEDGER": "Falta el diario o el mayor",
            "UNKNOWN": "No indicado",
            "WARN": "Aviso",
        }.get(value, value)
    return value


def _yes_no_label(value: object, language: str) -> str:
    normalized = str(value or "").strip().upper()
    if _is_italian(language):
        return (
            "Sì"
            if normalized == "YES"
            else ("No" if normalized == "NO" else normalized)
        )
    if _is_spanish(language):
        return (
            "Sí"
            if normalized == "YES"
            else ("No" if normalized == "NO" else normalized)
        )
    return normalized


def _evidence_type_label(evidence_type: object, language: str) -> str:
    value = str(evidence_type or "unknown")
    if _is_italian(language):
        return {
            "external_bank": "Banca",
            "probable_external_bank": "Banca probabile",
            "unallocated_external_bank": "Banca non allocata",
            "external_factoring": "Factor",
            "factoring_bridge": "Factor / anticipo",
            "internal_accounting": "Prima nota",
            "internal_booking": "Scrittura interna",
            "internal_closure": "Scrittura di chiusura",
            "compensation": "Compensazione",
            "payment_order": "Ordine di pagamento",
            "payment_order_bridge": "Distinta / ordine di pagamento",
            "open_item": "Partita aperta",
            "unknown": "Non classificata",
        }.get(value, value)
    if _is_spanish(language):
        return {
            "external_bank": "Banco",
            "probable_external_bank": "Movimiento bancario probable",
            "unallocated_external_bank": "Movimiento bancario sin asignar",
            "external_factoring": "Factor",
            "factoring_bridge": "Factor o anticipo",
            "internal_accounting": "Contabilidad interna",
            "internal_booking": "Asiento interno",
            "internal_closure": "Asiento de cierre",
            "compensation": "Compensación",
            "payment_order": "Orden de pago",
            "payment_order_bridge": "Lote u orden de pago",
            "open_item": "Partida abierta",
            "unknown": "Sin clasificar",
        }.get(value, value)
    return value


def _account_rollforward_report_rows(
    rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    if _is_italian(language):
        labels = {
            "account": "Conto",
            "account_name": "Nome conto",
            "status": "Esito",
            "opening_diff": "Differenza apertura",
            "closing_diff": "Differenza chiusura",
            "review_note": "Nota",
        }
    elif _is_spanish(language):
        labels = {
            "account": "Cuenta",
            "account_name": "Nombre de la cuenta",
            "status": "Estado",
            "opening_diff": "Diferencia inicial",
            "closing_diff": "Diferencia final",
            "review_note": "Nota",
        }
    else:
        labels = {
            "account": "Account",
            "account_name": "Account name",
            "status": "Status",
            "opening_diff": "Opening difference",
            "closing_diff": "Closing difference",
            "review_note": "Note",
        }

    def sort_key(row: dict[str, Any]) -> tuple[int, int, Decimal, str]:
        account = str(row.get("account") or "")
        status = str(row.get("status") or "").upper()
        closing_diff = abs(_decimal(row.get("closing_difference_journal_minus_ledger")))
        return (
            0 if account == "TOTAL" else 1,
            0 if status != "PASS" else 1,
            -closing_diff,
            account,
        )

    report_rows: list[dict[str, Any]] = []
    for row in sorted(rows, key=sort_key):
        opening_diff = row.get("opening_difference_journal_minus_ledger")
        closing_diff = row.get("closing_difference_journal_minus_ledger")
        report_rows.append(
            {
                labels["account"]: row.get("account", ""),
                labels["account_name"]: row.get("account_name", ""),
                labels["status"]: _status_label(row.get("status"), language),
                labels["opening_diff"]: (
                    "" if opening_diff in (None, "") else _format_decimal(opening_diff)
                ),
                labels["closing_diff"]: (
                    "" if closing_diff in (None, "") else _format_decimal(closing_diff)
                ),
                labels["review_note"]: row.get("review_note", ""),
            }
        )
    return report_rows


def _post_cutoff_summary_rows(
    rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    if _is_italian(language):
        labels = {
            "evidence_type": "Tipo evidenza",
            "exact_match": "Importo esatto",
            "rows": "Righe",
            "open_amount": "Importo aperto",
            "evidence_amount": "Importo evidenza",
        }
    elif _is_spanish(language):
        labels = {
            "evidence_type": "Tipo de evidencia",
            "exact_match": "Importe exacto",
            "rows": "Líneas",
            "open_amount": "Importe abierto",
            "evidence_amount": "Importe de la evidencia",
        }
    else:
        labels = {
            "evidence_type": "Evidence type",
            "exact_match": "Exact amount",
            "rows": "Rows",
            "open_amount": "Open amount",
            "evidence_amount": "Evidence amount",
        }
    buckets: dict[tuple[str, str], dict[str, Any]] = {}
    for row in rows:
        evidence_type = _evidence_type_label(row.get("evidence_type"), language)
        exact_match = _yes_no_label(row.get("exact_amount_match"), language)
        key = (evidence_type, exact_match)
        bucket = buckets.setdefault(
            key,
            {
                labels["evidence_type"]: evidence_type,
                labels["exact_match"]: exact_match,
                labels["rows"]: 0,
                labels["open_amount"]: Decimal("0.00"),
                labels["evidence_amount"]: Decimal("0.00"),
            },
        )
        bucket[labels["rows"]] += 1
        bucket[labels["open_amount"]] += _decimal(row.get("open_amount"))
        bucket[labels["evidence_amount"]] += _decimal(row.get("evidence_amount"))

    return [
        {
            labels["evidence_type"]: row[labels["evidence_type"]],
            labels["exact_match"]: row[labels["exact_match"]],
            labels["rows"]: row[labels["rows"]],
            labels["open_amount"]: _format_decimal(row[labels["open_amount"]]),
            labels["evidence_amount"]: _format_decimal(row[labels["evidence_amount"]]),
        }
        for row in sorted(
            buckets.values(),
            key=lambda item: (
                str(item[labels["evidence_type"]]),
                str(item[labels["exact_match"]]),
            ),
        )
    ]


def _post_cutoff_detail_rows(
    rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    if _is_italian(language):
        labels = {
            "document_no": "Documento",
            "document_date": "Data documento",
            "open_amount": "Importo aperto",
            "evidence_date": "Data evidenza",
            "evidence_type": "Tipo evidenza",
            "exact_match": "Importo esatto",
            "source": "Fonte",
        }
    elif _is_spanish(language):
        labels = {
            "document_no": "Documento",
            "document_date": "Fecha del documento",
            "open_amount": "Importe abierto",
            "evidence_date": "Fecha de la evidencia",
            "evidence_type": "Tipo de evidencia",
            "exact_match": "Importe exacto",
            "source": "Fuente",
        }
    else:
        labels = {
            "document_no": "Document",
            "document_date": "Document date",
            "open_amount": "Open amount",
            "evidence_date": "Evidence date",
            "evidence_type": "Evidence type",
            "exact_match": "Exact amount",
            "source": "Source",
        }

    def sort_key(row: dict[str, Any]) -> tuple[Decimal, str, str]:
        return (
            -abs(_decimal(row.get("open_amount"))),
            str(row.get("document_key") or ""),
            str(row.get("evidence_date") or ""),
        )

    report_rows: list[dict[str, Any]] = []
    for row in sorted(rows, key=sort_key):
        report_rows.append(
            {
                labels["document_no"]: row.get("document_no", ""),
                labels["document_date"]: row.get("document_date", ""),
                labels["open_amount"]: _format_decimal(row.get("open_amount")),
                labels["evidence_date"]: row.get("evidence_date", ""),
                labels["evidence_type"]: _evidence_type_label(
                    row.get("evidence_type"), language
                ),
                labels["exact_match"]: _yes_no_label(
                    row.get("exact_amount_match"), language
                ),
                labels["source"]: row.get("evidence_source_file", ""),
            }
        )
    return report_rows


def _support_bucket_label(value: object, language: str) -> str:
    bucket = str(value or "unknown")
    if _is_italian(language):
        return {
            "bank": "Banca",
            "bank_probable": "Banca probabile",
            "factor_or_advance": "Factor / anticipo",
            "payment_order": "Distinta / ordine di pagamento",
            "compensation": "Compensazione",
            "internal_accounting": "Supporto interno",
            "bridge_only": "Documento ponte",
            "no_evidence": "Nessuna evidenza",
            "unknown": "Non classificata",
        }.get(bucket, bucket)
    if _is_spanish(language):
        return {
            "bank": "Banco",
            "bank_probable": "Movimiento bancario probable",
            "factor_or_advance": "Factor o anticipo",
            "payment_order": "Lote u orden de pago",
            "compensation": "Compensación",
            "internal_accounting": "Respaldo interno",
            "bridge_only": "Documento puente",
            "no_evidence": "Sin evidencia",
            "unknown": "Sin clasificar",
        }.get(bucket, bucket)
    return bucket


def _review_signal_report_rows(
    rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    if _is_italian(language):
        labels = {
            "rank": "Priorita",
            "document": "Documento",
            "amount": "Importo",
            "age": "Giorni",
            "status": "Esito",
            "signals": "Motivo",
        }
    elif _is_spanish(language):
        labels = {
            "rank": "Prioridad",
            "document": "Documento",
            "amount": "Importe",
            "age": "Días",
            "status": "Estado",
            "signals": "Motivo",
        }
    else:
        labels = {
            "rank": "Priority",
            "document": "Document",
            "amount": "Amount",
            "age": "Days",
            "status": "Status",
            "signals": "Reason",
        }
    return [
        {
            labels["rank"]: row.get("review_signal_rank", ""),
            labels["document"]: row.get("document_no") or row.get("document_key", ""),
            labels["amount"]: _format_decimal(row.get("amount")),
            labels["age"]: row.get("age_days_at_reference", ""),
            labels["status"]: _reconciliation_status_label(
                row.get("reconciliation_status"), language
            ),
            labels["signals"]: _localized_semicolon_list(
                row.get("review_signals", ""), language, _review_signal_label
            ),
        }
        for row in rows
    ]


def _document_source_report_rows(
    rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    if _is_italian(language):
        labels = {
            "document": "Documento",
            "open_amount": "Importo aperto",
            "open_rows": "Righe aperte",
            "ledger": "Mastro",
            "journal": "Giornale",
            "bank": "Banca",
            "payment": "Distinte",
            "factor": "Factor",
            "note": "Nota",
        }
    elif _is_spanish(language):
        labels = {
            "document": "Documento",
            "open_amount": "Importe abierto",
            "open_rows": "Líneas abiertas",
            "ledger": "Mayor",
            "journal": "Diario",
            "bank": "Banco",
            "payment": "Órdenes de pago",
            "factor": "Factor",
            "note": "Nota",
        }
    else:
        labels = {
            "document": "Document",
            "open_amount": "Open amount",
            "open_rows": "Open rows",
            "ledger": "Ledger",
            "journal": "Journal",
            "bank": "Bank",
            "payment": "Payment orders",
            "factor": "Factor",
            "note": "Note",
        }
    return [
        {
            labels["document"]: row.get("document_no_examples")
            or row.get("document_key", ""),
            labels["open_amount"]: _format_decimal(row.get("open_amount_total")),
            labels["open_rows"]: row.get("open_item_rows", ""),
            labels["ledger"]: row.get("ledger_rows", ""),
            labels["journal"]: row.get("journal_rows", ""),
            labels["bank"]: row.get("bank_rows", ""),
            labels["payment"]: row.get("payment_order_rows", ""),
            labels["factor"]: row.get("factoring_rows", ""),
            labels["note"]: row.get("review_note", ""),
        }
        for row in rows
    ]


def _reversal_report_rows(
    rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    if _is_italian(language):
        labels = {
            "document": "Documento",
            "open_amount": "Importo aperto",
            "evidence_amount": "Importo evidenza",
            "evidence_type": "Tipo evidenza",
            "date": "Data evidenza",
            "reason": "Motivo",
        }
    elif _is_spanish(language):
        labels = {
            "document": "Documento",
            "open_amount": "Importe abierto",
            "evidence_amount": "Importe de la evidencia",
            "evidence_type": "Tipo de evidencia",
            "date": "Fecha de la evidencia",
            "reason": "Motivo",
        }
    else:
        labels = {
            "document": "Document",
            "open_amount": "Open amount",
            "evidence_amount": "Evidence amount",
            "evidence_type": "Evidence type",
            "date": "Evidence date",
            "reason": "Reason",
        }
    return [
        {
            labels["document"]: row.get("document_no") or row.get("document_key", ""),
            labels["open_amount"]: _format_decimal(row.get("open_amount")),
            labels["evidence_amount"]: _format_decimal(row.get("evidence_amount")),
            labels["evidence_type"]: _evidence_type_label(
                row.get("evidence_type"), language
            ),
            labels["date"]: row.get("evidence_date", ""),
            labels["reason"]: _localized_semicolon_list(
                row.get("candidate_reasons", ""), language, _review_signal_label
            ),
        }
        for row in rows
    ]


def _cutoff_window_summary_rows(
    rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    if _is_italian(language):
        labels = {
            "timing": "Periodo",
            "evidence_type": "Tipo evidenza",
            "rows": "Righe",
            "amount": "Importo",
        }
    elif _is_spanish(language):
        labels = {
            "timing": "Periodo",
            "evidence_type": "Tipo de evidencia",
            "rows": "Líneas",
            "amount": "Importe",
        }
    else:
        labels = {
            "timing": "Timing",
            "evidence_type": "Evidence type",
            "rows": "Rows",
            "amount": "Amount",
        }
    buckets: dict[tuple[str, str], dict[str, Any]] = {}
    for row in rows:
        timing = _cutoff_timing_label(row.get("cutoff_window_timing"), language)
        evidence_type = _evidence_type_label(row.get("evidence_type"), language)
        key = (timing, evidence_type)
        bucket = buckets.setdefault(
            key,
            {
                labels["timing"]: timing,
                labels["evidence_type"]: evidence_type,
                labels["rows"]: 0,
                labels["amount"]: Decimal("0.00"),
            },
        )
        bucket[labels["rows"]] += 1
        bucket[labels["amount"]] += _decimal(row.get("amount"))
    return [
        {
            labels["timing"]: row[labels["timing"]],
            labels["evidence_type"]: row[labels["evidence_type"]],
            labels["rows"]: row[labels["rows"]],
            labels["amount"]: _format_decimal(row[labels["amount"]]),
        }
        for row in sorted(
            buckets.values(),
            key=lambda item: (
                str(item[labels["timing"]]),
                str(item[labels["evidence_type"]]),
            ),
        )
    ]


def _evidence_concentration_report_rows(
    rows: list[dict[str, Any]], language: str
) -> list[dict[str, Any]]:
    if _is_italian(language):
        labels = {
            "bucket": "Tipo evidenza",
            "status": "Esito",
            "rows": "Righe",
            "amount": "Importo assoluto",
            "share": "Peso %",
        }
    elif _is_spanish(language):
        labels = {
            "bucket": "Tipo de evidencia",
            "status": "Estado",
            "rows": "Líneas",
            "amount": "Importe absoluto",
            "share": "Peso %",
        }
    else:
        labels = {
            "bucket": "Evidence type",
            "status": "Status",
            "rows": "Rows",
            "amount": "Absolute amount",
            "share": "Share %",
        }
    return [
        {
            labels["bucket"]: _support_bucket_label(
                row.get("support_bucket"), language
            ),
            labels["status"]: _reconciliation_status_label(
                row.get("reconciliation_status"), language
            ),
            labels["rows"]: row.get("rows", ""),
            labels["amount"]: _format_decimal(row.get("amount_abs_total")),
            labels["share"]: row.get("share_of_abs_amount_percent", ""),
        }
        for row in rows
    ]


def rows_from_mapping(
    mapping: dict[str, Any], key_name: str = "Campo", value_name: str = "Valore"
) -> list[dict[str, Any]]:
    def render(value: Any) -> Any:
        if isinstance(value, (dict, list, tuple, set)):
            return json.dumps(value, ensure_ascii=False, sort_keys=True)
        return value

    return [
        {key_name: key, value_name: render(value)} for key, value in mapping.items()
    ]


def summary_from_reconciliation(
    reconciliation_rows: list[dict[str, Any]],
    amount_field: str = "amount",
) -> list[dict[str, Any]]:
    buckets: dict[tuple[str, str], dict[str, Any]] = {}
    for row in reconciliation_rows:
        status = str(row.get("reconciliation_status") or "unknown")
        rule = str(row.get("rule_applied") or "unknown")
        key = (status, rule)
        bucket = buckets.setdefault(
            key,
            {
                "reconciliation_status": status,
                "rule_applied": rule,
                "rows": 0,
                "amount": Decimal("0.00"),
            },
        )
        bucket["rows"] += 1
        bucket["amount"] += _decimal(row.get(amount_field))
    return [
        {
            "reconciliation_status": row["reconciliation_status"],
            "rule_applied": row["rule_applied"],
            "rows": row["rows"],
            "amount": f"{row['amount']:.2f}",
        }
        for row in sorted(
            buckets.values(),
            key=lambda item: (item["reconciliation_status"], item["rule_applied"]),
        )
    ]


def build_audit_workbook_sheets(
    *,
    assumptions: dict[str, Any],
    source_inventory: list[dict[str, Any]],
    normalized_records: list[dict[str, Any]],
    reconciliation_rows: list[dict[str, Any]],
    checks: list[dict[str, Any]],
    bank_allocation_candidates: list[dict[str, Any]] | None = None,
    external_evidence_summary: list[dict[str, Any]] | None = None,
    external_evidence_detail: list[dict[str, Any]] | None = None,
    ledger_balance_rows: list[dict[str, Any]] | None = None,
    account_rollforward_check: list[dict[str, Any]] | None = None,
    aggregate_rollforward_rows: list[dict[str, Any]] | None = None,
    aggregate_rollforward_summary: list[dict[str, Any]] | None = None,
    post_cutoff_candidates: list[dict[str, Any]] | None = None,
    aging_summary: list[dict[str, Any]] | None = None,
    evidence_concentration: list[dict[str, Any]] | None = None,
    review_signals: list[dict[str, Any]] | None = None,
    document_source_map: list[dict[str, Any]] | None = None,
    reversal_candidates: list[dict[str, Any]] | None = None,
    cutoff_window_movements: list[dict[str, Any]] | None = None,
    review_rows: list[dict[str, Any]] | None = None,
    language: str = "it",
) -> dict[str, list[dict[str, Any]]]:
    labels = _labels(language)
    sheets: dict[str, list[dict[str, Any]]] = {
        "Assumptions": rows_from_mapping(
            assumptions,
            labels["metadata_field"],
            labels["metadata_value"],
        ),
        "Source inventory": source_inventory,
        "Normalized records": normalized_records,
        "Reconciliation detail": reconciliation_rows,
        "Bank allocation candidates": bank_allocation_candidates or [],
        "External evidence aggregate": external_evidence_summary or [],
        "External evidence detail": external_evidence_detail or [],
        "Ledger balance check": ledger_balance_rows or [],
        "Account rollforward check": account_rollforward_check or [],
        "Journal rollforward": aggregate_rollforward_summary or [],
        "Journal detail": aggregate_rollforward_rows or [],
        "Post-cutoff candidates": post_cutoff_candidates or [],
        "Open item aging": aging_summary or [],
        "Evidence concentration": evidence_concentration or [],
        "Review signals": review_signals or [],
        "Document source map": document_source_map or [],
        "Reversal candidates": reversal_candidates or [],
        "Cutoff window movements": cutoff_window_movements or [],
        "Summary": summary_from_reconciliation(reconciliation_rows),
        "Checks": checks,
        "Review": review_rows or [],
    }
    sheets["Index"] = [
        {"Sheet": name, "Rows": len(rows)} for name, rows in sheets.items()
    ]
    return {name: sheets.get(name, []) for name in STANDARD_SHEET_ORDER}


def write_excel_workpaper(
    output_path: str | Path,
    sheets: dict[str, list[dict[str, Any]]],
    preferred_headers: dict[str, list[str]] | None = None,
    language: str = "en",
) -> Path:
    """Write an audit-ready Excel workbook with filters, frozen headers and wrapping."""

    path = Path(output_path)
    wb = Workbook()
    wb.remove(wb.active)
    header_fill = PatternFill("solid", fgColor=NAVY)
    preferred_headers = preferred_headers or {}

    ordered_names = [name for name in STANDARD_SHEET_ORDER if name in sheets]
    ordered_names.extend(name for name in sheets if name not in ordered_names)
    localized_sheets = _localized_excel_sheets(sheets, ordered_names, language)

    def write_sheet(
        sheet_name: str,
        rows: list[dict[str, Any]],
        *,
        preferred: list[str] | None = None,
        hidden: bool = False,
    ) -> None:
        ws = wb.create_sheet(sheet_name[:31])
        headers = _headers(rows, preferred)
        ws.append(headers)
        for cell in ws[1]:
            cell.font = Font(color="FFFFFF", bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        for row in rows:
            ws.append([row.get(header, "") for header in headers])
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for idx, header in enumerate(headers, start=1):
            width = 14
            lower = header.lower()
            if any(
                token in lower
                for token in [
                    "reason",
                    "next",
                    "description",
                    "descripción",
                    "reference",
                    "referencia",
                    "source",
                    "fuente",
                    "evidence",
                    "evidencia",
                ]
            ):
                width = 42
            if any(
                token in lower
                for token in ["amount", "importe", "balance", "saldo", "value", "valor"]
            ):
                width = 16
            if any(token in lower for token in ["date", "fecha"]):
                width = 14
            ws.column_dimensions[get_column_letter(idx)].width = width
            for cell in ws[get_column_letter(idx)]:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        if hidden:
            ws.sheet_state = "hidden"

    for sheet_name, rows in localized_sheets.items():
        write_sheet(sheet_name, rows)

    for sheet_name in ordered_names:
        write_sheet(
            sheet_name,
            sheets[sheet_name],
            preferred=preferred_headers.get(sheet_name),
            hidden=bool(localized_sheets),
        )

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    return path


def write_word_report(
    output_path: str | Path,
    title: str,
    metadata: dict[str, Any],
    summary_rows: list[dict[str, Any]],
    assumptions: dict[str, Any],
    next_steps: list[str],
    narrative: str = "",
    source_inventory: list[dict[str, Any]] | None = None,
    external_evidence_summary: list[dict[str, Any]] | None = None,
    account_rollforward_check: list[dict[str, Any]] | None = None,
    aggregate_rollforward_summary: list[dict[str, Any]] | None = None,
    post_cutoff_candidates: list[dict[str, Any]] | None = None,
    aging_summary: list[dict[str, Any]] | None = None,
    evidence_concentration: list[dict[str, Any]] | None = None,
    review_signals: list[dict[str, Any]] | None = None,
    document_source_map: list[dict[str, Any]] | None = None,
    reversal_candidates: list[dict[str, Any]] | None = None,
    cutoff_window_movements: list[dict[str, Any]] | None = None,
    checks: list[dict[str, Any]] | None = None,
    review_rows: list[dict[str, Any]] | None = None,
    language: str = "it",
) -> Path:
    """Write an explanatory Word report; detailed row evidence belongs in Excel."""

    path = Path(output_path)
    labels = _labels(language)
    text = _report_text(language)
    doc = Document()
    _configure_word_styles(doc)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = text["executive_summary"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(header, after=0)
    for run in header.runs:
        _set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = labels["excel_authority"]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(footer, after=0)
    for run in footer.runs:
        _set_run_font(run, size=8.5, color=MUTED)

    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.add_run(title)
    _set_paragraph_spacing(title_paragraph, after=3)
    for run in title_paragraph.runs:
        _set_run_font(run, size=22, color=INK, bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Workpaper di riconciliazione deterministica"
        if _is_italian(language)
        else (
            "Papel de trabajo de conciliación determinista"
            if _is_spanish(language)
            else "Deterministic reconciliation workpaper"
        )
    )
    _set_paragraph_spacing(subtitle, after=10)
    for run in subtitle.runs:
        _set_run_font(run, size=11, color=MUTED)

    rule = doc.add_paragraph()
    _paragraph_border_bottom(rule)
    _set_paragraph_spacing(rule, after=12)

    doc.add_heading(text["executive_summary"], level=1)
    _add_callout(
        doc,
        labels["conclusion"],
        " ".join([narrative or labels["fallback_narrative"], text["executive_copy"]]),
        fill=LIGHT_BLUE,
    )

    if metadata:
        doc.add_heading(
            (
                "Metadati run"
                if _is_italian(language)
                else (
                    "Metadatos de la ejecución"
                    if _is_spanish(language)
                    else "Run Metadata"
                )
            ),
            level=1,
        )
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = labels["metadata_field"]
        table.rows[0].cells[1].text = labels["metadata_value"]
        for key, value in metadata.items():
            cells = table.add_row().cells
            cells[0].text = _field_label(key, language)
            cells[1].text = str(_localize_value(key, value, language))
        _format_table(table, [labels["metadata_field"], labels["metadata_value"]])

    doc.add_heading(text["scope_method"], level=1)
    scope = doc.add_paragraph(text["scope_copy"])
    _set_paragraph_spacing(scope)
    if source_inventory:
        source_summary: dict[str, int] = {}
        for row in source_inventory:
            role = str(row.get("source_role") or row.get("role") or "unknown")
            source_summary[role] = source_summary.get(role, 0) + 1
        _add_table(
            doc,
            [
                {
                    _field_label("source_role", language): _source_role_label(
                        role, language
                    ),
                    _field_label("files", language): count,
                }
                for role, count in sorted(source_summary.items())
            ],
        )

    doc.add_heading(labels["summary"], level=1)
    status_summary = _summary_by_status(summary_rows, language)
    if status_summary:
        _add_table(doc, status_summary)
    else:
        doc.add_paragraph(text["no_rows"])
    if summary_rows:
        doc.add_heading(
            (
                "Dettaglio regole"
                if _is_italian(language)
                else ("Detalle de reglas" if _is_spanish(language) else "Rule Detail")
            ),
            level=2,
        )
        _add_table(doc, _localize_rows(summary_rows, language))

    doc.add_heading(text["how_to_read"], level=1)
    _add_bullet_list(
        doc,
        [
            text["status_closed"],
            text["status_probable_payment"],
            text["status_open_supported"],
            text["status_needs_evidence"],
            text["status_unresolved"],
        ],
    )

    if external_evidence_summary:
        doc.add_heading(text["external_evidence"], level=1)
        external = doc.add_paragraph(text["external_copy"])
        _set_paragraph_spacing(external)
        _add_table(doc, _localize_rows(external_evidence_summary, language))

    if account_rollforward_check:
        doc.add_heading(text["account_rollforward"], level=1)
        account_rollforward = doc.add_paragraph(text["account_rollforward_copy"])
        _set_paragraph_spacing(account_rollforward)
        _add_table(
            doc,
            _account_rollforward_report_rows(account_rollforward_check, language),
            limit=12,
        )

    if aggregate_rollforward_summary:
        doc.add_heading(text["rollforward"], level=1)
        rollforward = doc.add_paragraph(text["rollforward_copy"])
        _set_paragraph_spacing(rollforward)
        _add_table(doc, _localize_rows(aggregate_rollforward_summary, language))

    if post_cutoff_candidates:
        doc.add_heading(text["post_cutoff"], level=1)
        post_cutoff = doc.add_paragraph(text["post_cutoff_copy"])
        _set_paragraph_spacing(post_cutoff)
        _add_table(doc, _post_cutoff_summary_rows(post_cutoff_candidates, language))
        doc.add_heading(text["post_cutoff_detail"], level=2)
        _add_table(
            doc,
            _post_cutoff_detail_rows(post_cutoff_candidates, language),
            limit=10,
        )

    additional_sections = [
        aging_summary,
        evidence_concentration,
        review_signals,
        document_source_map,
        reversal_candidates,
        cutoff_window_movements,
    ]
    if any(additional_sections):
        doc.add_heading(text["additional_analyses"], level=1)
        additional = doc.add_paragraph(text["additional_analyses_copy"])
        _set_paragraph_spacing(additional)
        if aging_summary:
            doc.add_heading(text["aging"], level=2)
            _add_table(doc, _localize_rows(aging_summary, language))
        if evidence_concentration:
            doc.add_heading(text["evidence_concentration"], level=2)
            _add_table(
                doc,
                _evidence_concentration_report_rows(evidence_concentration, language),
            )
        if review_signals:
            doc.add_heading(text["review_signals"], level=2)
            _add_table(
                doc,
                _review_signal_report_rows(review_signals, language),
                limit=10,
            )
        if document_source_map:
            doc.add_heading(text["document_source_map"], level=2)
            _add_table(
                doc,
                _document_source_report_rows(document_source_map, language),
                limit=10,
            )
        if reversal_candidates:
            doc.add_heading(text["reversal_candidates"], level=2)
            _add_table(
                doc,
                _reversal_report_rows(reversal_candidates, language),
                limit=10,
            )
        if cutoff_window_movements:
            doc.add_heading(text["cutoff_window"], level=2)
            _add_table(
                doc,
                _cutoff_window_summary_rows(cutoff_window_movements, language),
            )

    if checks:
        doc.add_heading(text["checks"], level=1)
        checks_copy = doc.add_paragraph(text["checks_copy"])
        _set_paragraph_spacing(checks_copy)
        _add_table(doc, _check_summary(checks, language))
        flagged = [
            row for row in checks if str(row.get("status") or "").upper() != "PASS"
        ]
        if flagged:
            _add_table(doc, _localize_rows(flagged, language), limit=10)

    if review_rows:
        doc.add_heading(text["review"], level=1)
        review = doc.add_paragraph(text["review_copy"])
        _set_paragraph_spacing(review)
        review_summary: dict[str, int] = {}
        for row in review_rows:
            status = str(row.get("review_status") or "UNKNOWN")
            review_summary[status] = review_summary.get(status, 0) + 1
        _add_table(
            doc,
            [
                {
                    _field_label("review_status", language): _review_status_label(
                        status, language
                    ),
                    _field_label("rows", language): count,
                }
                for status, count in sorted(review_summary.items())
            ],
        )

    doc.add_heading(labels["assumptions"], level=1)
    for key, value in assumptions.items():
        paragraph = doc.add_paragraph(
            f"{_field_label(key, language)}: {_localize_value(key, value, language)}",
            style="List Bullet",
        )
        _set_paragraph_spacing(paragraph, after=4, line_spacing=1.167)

    doc.add_heading(text["limitations"], level=1)
    _add_callout(doc, text["limitations"], text["limitations_copy"], fill=AMBER_FILL)

    doc.add_heading(labels["next_steps"], level=1)
    for step in next_steps:
        paragraph = doc.add_paragraph(step, style="List Bullet")
        _set_paragraph_spacing(paragraph, after=4, line_spacing=1.167)

    doc.add_heading(text["excel_reference"], level=1)
    excel_copy = doc.add_paragraph(text["excel_copy"])
    _set_paragraph_spacing(excel_copy)
    authority = doc.add_paragraph()
    authority.add_run(labels["excel_authority"]).bold = True
    _set_paragraph_spacing(authority, before=8, after=0)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path
