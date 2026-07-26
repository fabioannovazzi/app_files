"""Replayable mechanical assurance controls for Audit Reconciliation.

The controls in this module are deterministic because byte identity, reviewed
receipt binding, exact Decimal arithmetic, file-set equality, and gate
dependencies are mechanically verifiable. They do not decide source meaning,
materiality, evidence sufficiency, or an accounting conclusion.
"""

from __future__ import annotations

import sys as _bootstrap_sys

_bootstrap_sys.dont_write_bytecode = True
_bootstrap_sys.pycache_prefix = (
    r"Z:\__audit_reconciliation_no_bytecode__"
    if _bootstrap_sys.platform == "win32"
    else "/dev/null/audit-reconciliation"
)

import os as _bootstrap_os

_BOOTSTRAP_PATH = _bootstrap_os.path.join(
    _bootstrap_os.path.dirname(_bootstrap_os.path.abspath(__file__)),
    "implementation_bootstrap.py",
)
_BOOTSTRAP_NAMESPACE = {
    "__file__": _BOOTSTRAP_PATH,
    "__name__": "_audit_reconciliation_implementation_bootstrap",
}
_bootstrap_stat = _bootstrap_os.lstat(_BOOTSTRAP_PATH)
if _bootstrap_stat.st_mode & 0o170000 != 0o100000 or _bootstrap_stat.st_nlink != 1:
    raise RuntimeError(
        "implementation bootstrap must be an ordinary single-link regular file"
    )
_bootstrap_descriptor = _bootstrap_os.open(
    _BOOTSTRAP_PATH,
    _bootstrap_os.O_RDONLY | getattr(_bootstrap_os, "O_NOFOLLOW", 0),
)
try:
    _bootstrap_open_stat = _bootstrap_os.fstat(_bootstrap_descriptor)
    _bootstrap_identity = (
        _bootstrap_stat.st_dev,
        _bootstrap_stat.st_ino,
        _bootstrap_stat.st_size,
        _bootstrap_stat.st_mtime_ns,
        _bootstrap_stat.st_nlink,
    )
    if _bootstrap_identity != (
        _bootstrap_open_stat.st_dev,
        _bootstrap_open_stat.st_ino,
        _bootstrap_open_stat.st_size,
        _bootstrap_open_stat.st_mtime_ns,
        _bootstrap_open_stat.st_nlink,
    ):
        raise RuntimeError("implementation bootstrap changed before it was read")
    with _bootstrap_os.fdopen(
        _bootstrap_descriptor,
        "rb",
        closefd=False,
    ) as _bootstrap_handle:
        _bootstrap_source = _bootstrap_handle.read()
    _bootstrap_after_stat = _bootstrap_os.fstat(_bootstrap_descriptor)
    if (
        _bootstrap_identity
        != (
            _bootstrap_after_stat.st_dev,
            _bootstrap_after_stat.st_ino,
            _bootstrap_after_stat.st_size,
            _bootstrap_after_stat.st_mtime_ns,
            _bootstrap_after_stat.st_nlink,
        )
        or len(_bootstrap_source) != _bootstrap_after_stat.st_size
    ):
        raise RuntimeError("implementation bootstrap changed while it was read")
finally:
    _bootstrap_os.close(_bootstrap_descriptor)
# Execute only the pre-opened single-link bootstrap source.
exec(  # nosec B102
    compile(_bootstrap_source, _BOOTSTRAP_PATH, "exec"),
    _BOOTSTRAP_NAMESPACE,
)
_BOOTSTRAP_ROOTS = _BOOTSTRAP_NAMESPACE["activate_implementation_boundary"](
    ("locale_support", "reconciliation_helpers")
)
_BOOTSTRAP_VALIDATE_IMPLEMENTATION = _BOOTSTRAP_NAMESPACE[
    "validate_implementation_tree"
]
IMPLEMENTATION_CONTRACT = tuple(_BOOTSTRAP_NAMESPACE["IMPLEMENTATION_CONTRACT"])

import argparse
import csv
import hashlib
import json
import mimetypes
import os
import re
import shutil
import stat
import sys
import tempfile
from collections.abc import Callable, Mapping, Sequence
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

_SCRIPTS_DIRECTORY = Path(__file__).resolve().parent
if str(_SCRIPTS_DIRECTORY) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIRECTORY))

from reconciliation_helpers import (  # noqa: E402
    closed_bank_allocation_controls,
    reconcile_open_items,
    reconciliation_checks,
)


def _add_vera_assurance_module_path() -> Path:
    """Use the shared module root admitted by the pre-import boundary."""

    shared_root = Path(str(_BOOTSTRAP_ROOTS["shared_assurance"]))
    module_root = shared_root.parent
    if str(module_root) not in sys.path:
        sys.path.insert(0, str(module_root))
    return shared_root


VERA_ASSURANCE_ROOT = _add_vera_assurance_module_path()

from vera_assurance import (  # noqa: E402
    artifact_receipt,
    build_allocation_ledger,
    build_gate_register,
    build_numeric_evidence_ledger,
    build_reviewed_decision_receipt,
    canonical_json_sha256,
    decimal_text,
    file_snapshot,
    parse_canonical_decimal,
    validate_allocation_ledger,
    validate_artifact_receipt,
    validate_gate_register,
    validate_numeric_evidence_ledger,
    validate_reviewed_decision_receipt,
    validate_source_qualification,
    write_json,
)

__all__ = [
    "AssuranceRunError",
    "build_implementation_receipts",
    "build_applied_review_authority",
    "build_professional_review_authority",
    "build_review_payload_mapping",
    "build_reviewed_source_decisions",
    "build_source_receipts",
    "capture_review_transition_predecessor",
    "finalize_assurance_run",
    "prepare_assurance_run",
    "retain_review_transition",
    "review_decision_fingerprint",
    "rollback_assurance_run",
    "reviewed_date_convention",
    "reviewed_money_convention",
    "run_review_output_transaction",
    "validate_assurance_run",
    "validate_final_output_inventory",
    "validate_professional_review_authority",
    "validate_receipt_set",
    "validate_review_transition_history",
]

ASSURANCE_SCHEMA_VERSION = "audit_reconciliation.assurance.v1"
SOURCE_DECISION_TYPE = "audit_reconciliation_source_mapping"
SOURCE_DECISION_ADAPTER_VERSION = "2"
FINAL_OUTPUT_DIRECTORY = "assurance_final_outputs"
FINAL_OUTPUT_INVENTORY = "final_output_inventory.json"
RUN_TREE_SCHEMA_VERSION = "audit_reconciliation.run_tree.v1"
REVIEW_TRANSITION_SCHEMA_VERSION = "audit_reconciliation.review_transition.v1"
REVIEW_PAYLOAD_MAPPING_SCHEMA_VERSION = "audit_reconciliation.review_payload_mapping.v1"
REVIEW_TRANSITION_HISTORY_DIRECTORY = "assurance_transition_history"
PREDECESSOR_RUN_SNAPSHOT_DIRECTORY = "predecessor_run"
REVIEW_TRANSITION_EVIDENCE_FILES = (
    (
        "predecessor_assurance_receipts.json",
        "transition.predecessor_assurance",
        "predecessor_assurance",
    ),
    (
        "predecessor_professional_review.json",
        "transition.predecessor_professional_review",
        "predecessor_professional_review",
    ),
    (
        "predecessor_reconciliation_results.json",
        "transition.predecessor_reconciliation_results",
        "predecessor_reconciliation_results",
    ),
    (
        "predecessor_review_payload.json",
        "transition.predecessor_review_payload",
        "predecessor_review_payload",
    ),
    (
        "review_payload_mapping.json",
        "transition.review_payload_mapping",
        "review_payload_mapping",
    ),
    (
        "applied_decisions.json",
        "transition.applied_decisions",
        "applied_decisions",
    ),
    (
        "successor_professional_review.json",
        "transition.successor_professional_review",
        "successor_professional_review",
    ),
)
REVIEW_TRANSITION_RECEIPT_FILE = "transition_receipt.json"
ASSURANCE_SEAL_FIELDS = frozenset(
    {
        "schema_version",
        "run_id",
        "run_date",
        "source_root",
        "source_receipts",
        "reviewed_source_decisions",
        "source_qualifications",
        "implementation_receipts",
        "prepared_receipt",
        "professional_review_receipt",
        "professional_review_authority",
        "review_transition_receipts",
        "workflow_output_contract",
        "run_tree_contract",
        "final_output_inventory",
        "numeric_evidence_ledger",
        "rendered_value_addresses",
        "allocation_value_addresses",
        "allocation_ledgers",
        "gate_register",
        "limitations",
        "content_sha256",
    }
)

_CHECKPOINT_PATTERN = re.compile(r"[0-9a-f]{64}")
SUPPORTED_SOURCE_ADAPTER_VERSIONS = {
    "bank_statement_text_v1": "2",
    "journal_header_columns_v1": "2",
    "legacy_it_accounting_export_v1": "2",
    "open_items_text_v1": "2",
    "payment_order_html_zip_v1": "2",
}
TRUSTED_NATIVE_OUTPUT_PROFILES = (
    frozenset(
        {
            "riconciliazione_audit.xlsx",
            "scheda_operativa_commercialista.xlsx",
            "relazione_riconciliazione_audit.docx",
        }
    ),
    frozenset(
        {
            "riconciliazione_audit.xlsx",
            "scheda_operativa_commercialista.xlsx",
            "relazione_riconciliazione_audit.docx",
            "richieste_mirate_evidenze.xlsx",
        }
    ),
    # These compact profiles are retained for the low-level assurance API and
    # its self-contained replay tests. They are code-owned and cannot be
    # extended by editing a run's JSON contract.
    frozenset({"audit.xlsx"}),
    frozenset({"audit.xlsx", "report.txt"}),
    frozenset({"audit.xlsx", "report.docx"}),
    frozenset({"audit.xlsx", "summary.json"}),
)
RUN_CONTROL_PATHS = (
    "assurance_gates.json",
    "assurance_receipts.json",
    FINAL_OUTPUT_INVENTORY,
    "prepared_records.json",
    "reconciliation_results.json",
)
RUN_ROOT_FIXED_FILES = frozenset(
    {
        "account_rollforward_check.json",
        "aging_summary.json",
        "applied_decisions.json",
        "artifact_card.md",
        "assurance_gates.json",
        "assurance_receipts.json",
        "bank_allocation_candidates.json",
        "codex_review_packet.json",
        "cutoff_window_movements.json",
        "document_source_map.json",
        "evidence_concentration.json",
        "external_evidence_detail.json",
        "external_evidence_summary.json",
        "extraction_errors.json",
        "final_artifacts.json",
        FINAL_OUTPUT_INVENTORY,
        "journal_rollforward_rows.json",
        "journal_rollforward_summary.json",
        "ledger_balance_rows.json",
        "normalized_records.json",
        "numeric_evidence_ledger.json",
        "post_cutoff_candidates.json",
        "prepared_records.json",
        "professional_review.json",
        "reconciliation_results.json",
        "relationship_allocation_ledgers.json",
        "reversal_candidates.json",
        "review_handoff.md",
        "review_payload.json",
        "review_signals.json",
        "review_ui.html",
        "run_intake.json",
        "run_manifest.json",
        "source_pages.json",
        "source_qualifications.json",
        "ui_decisions.json",
    }
)
RUN_CACHE_DIRECTORY = ".audit_reconciliation_cache"
PERIMETER_FIELDS = {
    "entity_ref",
    "party_ref",
    "currency",
    "unit",
    "direction_policy",
    "allocation_policy",
}
MONEY_FIELDS = {
    "decimal_separator",
    "thousands_separator",
    "reported_unit",
    "reported_increment",
}
DATE_FIELDS = {"order"}
DATE_ORDERS = {"day_first", "month_first"}
SUPPORTED_REPORTED_INCREMENT = "0.01"
REVIEWER_REF_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]*$")
REVIEWER_REF_TRUST = "unsigned_untrusted_label"
MATERIAL_VALUE_HEADERS = {
    "amount",
    "importo",
    "importe",
    "balance",
    "saldo",
    "total",
    "difference",
    "differenza",
    "diferencia",
    "residual",
    "residuo",
    "debit",
    "credit",
    "dare",
    "avere",
}
MATERIAL_VALUE_KEY_PARTS = (
    "amount",
    "importo",
    "importe",
    "balance",
    "saldo",
    "total",
    "difference",
    "differenza",
    "diferencia",
    "residual",
    "residuo",
    "allocation",
    "allocated",
)
MATERIAL_TEXT_RE = re.compile(r"(?<![\d.])-?\d+(?:[.,]\d{2,})(?![\d.])")
CURRENCY_TEXT_RE = re.compile(
    r"(?i)(?:EUR|USD|GBP|CHF|€|\$|£)\s*([-+]?\d+(?:[.,]\d+)?)"
    r"|([-+]?\d+(?:[.,]\d+)?)\s*(?:EUR|USD|GBP|CHF|€|\$|£)"
)
SOURCE_ROLES = {
    "open_items",
    "counterparty_open_items",
    "ledger",
    "journal",
    "bank_statement",
    "payment_order",
    "factoring_statement",
    "compensation_support",
}
DIRECTION_POLICIES = {
    "customer",
    "supplier",
    "receivable",
    "payable",
    "debit",
    "credit",
    "source_signed",
    "not_applicable",
}
ALLOCATION_POLICIES = {
    "one_to_one",
    "one_to_many",
    "many_to_one",
    "many_to_many",
    "not_applicable",
}


class AssuranceRunError(ValueError):
    """Raised when an assurance boundary cannot be replayed exactly."""


_ACTIVE_RUN_TRANSACTIONS: dict[str, dict[str, Any]] = {}


def _canonical_iso_date(value: object, *, label: str) -> str:
    """Return exact ISO date text without coercion."""

    if not isinstance(value, str) or not value or value != value.strip():
        raise AssuranceRunError(f"{label} must be non-empty trimmed ISO date text")
    try:
        parsed = date.fromisoformat(value)
    except ValueError as exc:
        raise AssuranceRunError(f"{label} must be an ISO date") from exc
    if parsed.isoformat() != value:
        raise AssuranceRunError(f"{label} must be canonical ISO date text")
    return value


def _canonical_reviewer_ref(value: object, *, label: str) -> str:
    """Validate an unsigned label; this does not authenticate a reviewer."""

    if (
        not isinstance(value, str)
        or not value
        or value != value.strip()
        or REVIEWER_REF_RE.fullmatch(value) is None
    ):
        raise AssuranceRunError(f"{label} must be a canonical reviewer identifier")
    return value


def _run_date(assumptions: Mapping[str, Any]) -> str:
    """Seal one stable run date, defaulting only at intake time."""

    supplied = assumptions.get("assurance_run_date")
    if supplied is None:
        return date.today().isoformat()
    return _canonical_iso_date(supplied, label="assurance_run_date")


def _review_metadata(
    *,
    reviewer_ref: object,
    reviewed_on: object,
    run_date: str,
    label: str,
) -> tuple[str, str]:
    reviewer = _canonical_reviewer_ref(reviewer_ref, label=f"{label}.reviewer_ref")
    reviewed = _canonical_iso_date(reviewed_on, label=f"{label}.reviewed_on")
    if reviewed > run_date:
        raise AssuranceRunError(
            f"{label}.reviewed_on cannot be after the sealed run date"
        )
    return reviewer, reviewed


def _validated_tolerance(value: object, *, label: str = "amount_tolerance") -> str:
    """Return a canonical non-negative tolerance, preserving exact zero."""

    if isinstance(value, (bool, float)) or not isinstance(value, (str, int, Decimal)):
        raise AssuranceRunError(f"{label} must be canonical Decimal text")
    normalized = (
        decimal_text(Decimal(value))
        if isinstance(value, (int, Decimal)) and not isinstance(value, bool)
        else value
    )
    try:
        parsed = parse_canonical_decimal(normalized, label=label)
    except ValueError as exc:
        raise AssuranceRunError(str(exc)) from exc
    if parsed < 0:
        raise AssuranceRunError(f"{label} cannot be negative")
    return decimal_text(parsed)


def _json_safe(value: Any, *, label: str = "value") -> Any:
    """Return an exact JSON value, rejecting binary floating-point inputs."""

    if value is None or isinstance(value, (str, bool, int)):
        return value
    if isinstance(value, Path):
        return value.as_posix()
    if isinstance(value, Decimal):
        return decimal_text(value)
    if isinstance(value, float):
        raise AssuranceRunError(
            f"{label} contains a binary float; authoritative precision is unproven"
        )
    if isinstance(value, Mapping):
        return {
            str(key): _json_safe(item, label=f"{label}.{key}")
            for key, item in value.items()
        }
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
        return [
            _json_safe(item, label=f"{label}[{index}]")
            for index, item in enumerate(value)
        ]
    raise AssuranceRunError(f"{label} contains unsupported type {type(value).__name__}")


def _artifact_id(prefix: str, relative_path: str) -> str:
    digest = canonical_json_sha256({"path": relative_path})
    return f"{prefix}.{digest}"


def _media_type(path: Path) -> str:
    guessed, _ = mimetypes.guess_type(path.name)
    return guessed or "application/octet-stream"


def build_source_receipts(
    input_root: Path,
    paths: Sequence[Path],
) -> list[dict[str, Any]]:
    """Capture full current-byte receipts for every source regular file."""

    root = Path(input_root).resolve()
    receipts = []
    for path in sorted((Path(item).resolve() for item in paths), key=str):
        relative = path.relative_to(root).as_posix()
        _, digest = file_snapshot(path)
        receipts.append(
            artifact_receipt(
                root,
                path,
                artifact_id=(
                    "source."
                    + canonical_json_sha256({"path": relative, "sha256": digest})
                ),
                role="source",
                root_id="source",
                media_type=_media_type(path),
            )
        )
    return receipts


def validate_receipt_set(
    roots: Mapping[str, Path],
    receipts: Sequence[Mapping[str, Any]],
) -> list[dict[str, Any]]:
    """Replay every receipt against ordinary, single-link current file bytes.

    Link-count and file-type checks are mechanical filesystem assertions. They
    prevent a receipt from appearing current while another pathname can mutate
    the same inode or while a special file is substituted for an artifact.
    """

    normalized = []
    seen_paths: set[tuple[str, str]] = set()
    for receipt in receipts:
        root_id = receipt.get("root_id")
        relative = receipt.get("path")
        if (
            not isinstance(root_id, str)
            or root_id not in roots
            or not isinstance(relative, str)
        ):
            raise AssuranceRunError("artifact receipt root or path is unavailable")
        path_key = (root_id, relative)
        if path_key in seen_paths:
            raise AssuranceRunError(
                "artifact receipts cannot assign multiple identities to one path"
            )
        seen_paths.add(path_key)
        unresolved = Path(roots[root_id]).resolve() / relative
        try:
            file_stat = unresolved.lstat()
        except FileNotFoundError as exc:
            raise AssuranceRunError(f"artifact does not exist: {unresolved}") from exc
        if (
            unresolved.is_symlink()
            or not stat.S_ISREG(file_stat.st_mode)
            or file_stat.st_nlink != 1
        ):
            raise AssuranceRunError(
                "assurance artifacts must be ordinary single-link regular files"
            )
        normalized.append(validate_artifact_receipt(roots, receipt))
    identifiers = [str(item["artifact_id"]) for item in normalized]
    if len(identifiers) != len(set(identifiers)):
        raise AssuranceRunError("artifact receipt identities must be unique")
    return normalized


def _validate_source_boundary(
    source_root: Path,
    receipts: Sequence[Mapping[str, Any]],
) -> None:
    """Require receipts for the exact visible regular-file intake set."""

    root = Path(source_root).resolve()
    physical: set[str] = set()
    for path in root.iterdir():
        if path.name.startswith("."):
            continue
        file_stat = path.lstat()
        if (
            path.is_symlink()
            or not stat.S_ISREG(file_stat.st_mode)
            or file_stat.st_nlink != 1
        ):
            raise AssuranceRunError(
                "source boundary entries must be ordinary single-link regular files"
            )
        physical.add(path.relative_to(root).as_posix())
    declared = {str(receipt["path"]) for receipt in receipts}
    if physical != declared:
        raise AssuranceRunError(
            "source receipts do not cover the exact current source-file set"
        )
    for receipt in receipts:
        relative_path = str(receipt["path"])
        expected_artifact_id = "source." + canonical_json_sha256(
            {
                "path": relative_path,
                "sha256": str(receipt["sha256"]),
            }
        )
        if receipt["artifact_id"] != expected_artifact_id:
            raise AssuranceRunError(
                f"source receipt identity is stale for {relative_path}"
            )


def _source_decision_input(
    assumptions: Mapping[str, Any],
    relative_path: str,
) -> Mapping[str, Any] | None:
    decisions = assumptions.get("reviewed_source_decisions")
    if not isinstance(decisions, Mapping):
        return None
    value = decisions.get(relative_path)
    if value is None:
        value = decisions.get(Path(relative_path).name)
    return value if isinstance(value, Mapping) else None


def _validate_source_decision_content(
    value: Mapping[str, Any],
    *,
    relative_path: str,
    adapter_family: str,
    run_date: str,
) -> dict[str, Any]:
    required = {
        "role",
        "adapter_family",
        "reviewer_ref",
        "reviewed_on",
        "perimeter",
        "money",
        "date",
    }
    if set(value) != required:
        raise AssuranceRunError(
            f"reviewed source decision for {relative_path} has invalid fields"
        )
    if value["adapter_family"] != adapter_family:
        raise AssuranceRunError(f"reviewed source adapter for {relative_path} is stale")
    perimeter = value["perimeter"]
    money = value["money"]
    date_authority = value["date"]
    if not isinstance(perimeter, Mapping) or set(perimeter) != PERIMETER_FIELDS:
        raise AssuranceRunError(
            f"reviewed source perimeter for {relative_path} is incomplete"
        )
    if not isinstance(money, Mapping) or set(money) != MONEY_FIELDS:
        raise AssuranceRunError(
            f"reviewed money convention for {relative_path} is incomplete"
        )
    if (
        not isinstance(date_authority, Mapping)
        or set(date_authority) != DATE_FIELDS
        or date_authority["order"] not in DATE_ORDERS
    ):
        raise AssuranceRunError(
            f"reviewed date convention for {relative_path} is incomplete"
        )
    if value["role"] not in SOURCE_ROLES:
        raise AssuranceRunError(f"reviewed source role for {relative_path} is invalid")
    if not isinstance(adapter_family, str) or not adapter_family:
        raise AssuranceRunError(
            f"reviewed source adapter for {relative_path} is missing"
        )
    if (
        adapter_family not in SUPPORTED_SOURCE_ADAPTER_VERSIONS
        or SUPPORTED_SOURCE_ADAPTER_VERSIONS[adapter_family]
        != SOURCE_DECISION_ADAPTER_VERSION
    ):
        raise AssuranceRunError(
            f"reviewed source adapter binding is stale for {relative_path}: "
            "unsupported at the assurance boundary"
        )
    for field in ("entity_ref", "party_ref"):
        if perimeter[field] is not None and (
            not isinstance(perimeter[field], str) or not perimeter[field].strip()
        ):
            raise AssuranceRunError(f"perimeter.{field} must be text or null")
    for field in ("currency", "unit"):
        if not isinstance(perimeter[field], str) or not perimeter[field].strip():
            raise AssuranceRunError(f"perimeter.{field} must be non-empty text")
    if perimeter["direction_policy"] not in DIRECTION_POLICIES:
        raise AssuranceRunError("perimeter.direction_policy is unsupported")
    if perimeter["allocation_policy"] not in ALLOCATION_POLICIES:
        raise AssuranceRunError("perimeter.allocation_policy is unsupported")
    if (
        value["role"] in {"open_items", "counterparty_open_items"}
        and perimeter["direction_policy"] == "not_applicable"
    ):
        raise AssuranceRunError(
            "open-item sources require an explicit direction policy"
        )
    decimal_separator = money["decimal_separator"]
    thousands_separator = money["thousands_separator"]
    if decimal_separator not in {None, ".", ","}:
        raise AssuranceRunError("decimal_separator must be '.', ',', or null")
    if thousands_separator not in {None, ".", ","}:
        raise AssuranceRunError("thousands_separator must be '.', ',', or null")
    if (
        decimal_separator is not None
        and thousands_separator is not None
        and decimal_separator == thousands_separator
    ):
        raise AssuranceRunError("decimal and thousands separators must differ")
    increment = money["reported_increment"]
    if (
        not isinstance(money["reported_unit"], str)
        or not money["reported_unit"].strip()
    ):
        raise AssuranceRunError("reported_unit must be non-empty text")
    if increment != SUPPORTED_REPORTED_INCREMENT:
        raise AssuranceRunError(
            "reported_increment must be exactly 0.01; other increments are unsupported"
        )
    _review_metadata(
        reviewer_ref=value["reviewer_ref"],
        reviewed_on=value["reviewed_on"],
        run_date=run_date,
        label=f"reviewed_source_decisions.{relative_path}",
    )
    content = {
        "source_path": relative_path,
        "role": value["role"],
        "adapter_family": value["adapter_family"],
        "perimeter": dict(perimeter),
        "money": dict(money),
        "date": dict(date_authority),
    }
    return _json_safe(content, label=f"reviewed_source_decisions.{relative_path}")


def build_reviewed_source_decisions(
    *,
    input_root: Path,
    source_receipts: Sequence[Mapping[str, Any]],
    adapter_families: Mapping[str, str],
    assumptions: Mapping[str, Any],
) -> tuple[dict[str, dict[str, Any]], dict[str, str]]:
    """Build or replay explicit reviewed decisions bound to current sources."""

    root = Path(input_root).resolve()
    run_date = _run_date(assumptions)
    receipt_by_path = {str(item["path"]): item for item in source_receipts}
    decisions: dict[str, dict[str, Any]] = {}
    errors: dict[str, str] = {}
    for relative_path, receipt in sorted(receipt_by_path.items()):
        supplied = _source_decision_input(assumptions, relative_path)
        if supplied is None:
            errors[relative_path] = (
                "A reviewed source-decision receipt with an explicit perimeter "
                "and money convention is required."
            )
            continue
        adapter_family = adapter_families.get(
            relative_path,
            adapter_families.get(Path(relative_path).name, ""),
        )
        try:
            if supplied.get("schema_version") == "vera.reviewed_decision_receipt.v1":
                normalized = validate_reviewed_decision_receipt(
                    supplied,
                    expected_source_artifact_refs=[str(receipt["artifact_id"])],
                    expected_adapter_id=adapter_family,
                    expected_adapter_version=SOURCE_DECISION_ADAPTER_VERSION,
                    require_reviewed=True,
                )
                content = normalized["content"]
                if (
                    not isinstance(content, Mapping)
                    or content.get("source_path") != relative_path
                    or content.get("adapter_family") != adapter_family
                ):
                    raise AssuranceRunError("reviewed source decision content is stale")
                replay_input = {
                    "role": content.get("role"),
                    "adapter_family": content.get("adapter_family"),
                    "reviewer_ref": normalized["reviewer_ref"],
                    "reviewed_on": normalized["reviewed_on"],
                    "perimeter": content.get("perimeter"),
                    "money": content.get("money"),
                    "date": content.get("date"),
                }
                if _validate_source_decision_content(
                    replay_input,
                    relative_path=relative_path,
                    adapter_family=adapter_family,
                    run_date=run_date,
                ) != dict(content):
                    raise AssuranceRunError("reviewed source decision content is stale")
            else:
                content = _validate_source_decision_content(
                    supplied,
                    relative_path=relative_path,
                    adapter_family=adapter_family,
                    run_date=run_date,
                )
                reviewer_ref, reviewed_on = _review_metadata(
                    reviewer_ref=supplied["reviewer_ref"],
                    reviewed_on=supplied["reviewed_on"],
                    run_date=run_date,
                    label=f"reviewed_source_decisions.{relative_path}",
                )
                normalized = build_reviewed_decision_receipt(
                    decision_id=_artifact_id("source_mapping", relative_path),
                    decision_type=SOURCE_DECISION_TYPE,
                    status="reviewed",
                    reviewer_ref=reviewer_ref,
                    reviewed_on=reviewed_on,
                    adapter_id=adapter_family,
                    adapter_version=SOURCE_DECISION_ADAPTER_VERSION,
                    source_artifact_refs=[str(receipt["artifact_id"])],
                    content=content,
                )
            decisions[relative_path] = normalized
        except (ValueError, AssuranceRunError) as exc:
            errors[relative_path] = str(exc)
    validate_receipt_set({"source": root}, source_receipts)
    return decisions, errors


def reviewed_money_convention(
    decision: Mapping[str, Any] | None,
) -> dict[str, Any] | None:
    """Return an exact reviewed money convention, if the decision is current."""

    if not isinstance(decision, Mapping):
        return None
    content = decision.get("content")
    money = content.get("money") if isinstance(content, Mapping) else None
    if not isinstance(money, Mapping) or set(money) != MONEY_FIELDS:
        return None
    return dict(money)


def reviewed_date_convention(
    decision: Mapping[str, Any] | None,
) -> dict[str, Any] | None:
    """Return the exact source-bound date authority, if current."""

    if not isinstance(decision, Mapping):
        return None
    content = decision.get("content")
    date_authority = content.get("date") if isinstance(content, Mapping) else None
    if not isinstance(date_authority, Mapping) or set(date_authority) != DATE_FIELDS:
        return None
    if date_authority.get("order") not in DATE_ORDERS:
        return None
    return dict(date_authority)


def build_implementation_receipts(
    plugin_root: Path | None = None,
) -> list[dict[str, Any]]:
    """Capture the exact ordered transitive implementation contract."""

    root = (plugin_root or Path(__file__).resolve().parents[1]).resolve()
    roots = _implementation_roots(root)
    _validate_implementation_tree(root)
    receipts = []
    for root_id, relative_path in IMPLEMENTATION_CONTRACT:
        selected_root = roots[root_id]
        path = selected_root / relative_path
        receipts.append(
            artifact_receipt(
                selected_root,
                path,
                artifact_id=_artifact_id(
                    "implementation",
                    f"{root_id}/{relative_path}",
                ),
                role="implementation",
                root_id=root_id,
                media_type=_media_type(path),
            )
        )
    return receipts


def _implementation_roots(plugin_root: Path) -> dict[str, Path]:
    return {
        "plugin": Path(plugin_root).resolve(),
        "shared_assurance": VERA_ASSURANCE_ROOT,
    }


def _validate_real_directory(path: Path, *, label: str) -> None:
    try:
        current = path.lstat()
    except FileNotFoundError as exc:
        raise AssuranceRunError(f"{label} is unavailable") from exc
    if path.is_symlink() or not stat.S_ISDIR(current.st_mode):
        raise AssuranceRunError(f"{label} must be a real directory")


def _validate_implementation_tree(plugin_root: Path) -> None:
    """Replay the same physical closure enforced before local imports."""

    try:
        _BOOTSTRAP_VALIDATE_IMPLEMENTATION(
            str(Path(plugin_root).absolute()),
            shared_assurance_root=str(VERA_ASSURANCE_ROOT),
        )
    except RuntimeError as exc:
        raise AssuranceRunError(str(exc)) from exc


def _write_exact_json(path: Path, payload: Mapping[str, Any]) -> Path:
    destination = Path(path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    handle, temporary_name = tempfile.mkstemp(
        prefix=f".{destination.name}.",
        suffix=".tmp",
        dir=destination.parent,
    )
    os.close(handle)
    temporary = Path(temporary_name)
    try:
        write_json(temporary, _json_safe(payload))
        os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)
    return destination


def _begin_run_transaction(output_dir: Path) -> dict[str, Any]:
    """Snapshot the complete pre-run output state for whole-run rollback."""

    out_dir = Path(output_dir).resolve()
    if out_dir.as_posix() in _ACTIVE_RUN_TRANSACTIONS:
        _rollback_run_transaction(
            out_dir,
            _ACTIVE_RUN_TRANSACTIONS.get(out_dir.as_posix()),
        )
    transaction_root = Path(
        tempfile.mkdtemp(prefix=".audit-run-transaction-", dir=out_dir.parent)
    )
    snapshot = transaction_root / "snapshot"
    existed = out_dir.exists() or out_dir.is_symlink()
    if existed:
        if out_dir.is_symlink() or not out_dir.is_dir():
            shutil.rmtree(transaction_root, ignore_errors=True)
            raise AssuranceRunError("output_dir must be a real directory")
        shutil.copytree(out_dir, snapshot, symlinks=True)
    else:
        snapshot.mkdir()
    transaction = {
        "transaction_root": transaction_root.as_posix(),
        "snapshot_path": snapshot.as_posix(),
        "output_existed": existed,
        "output_dir": out_dir.as_posix(),
    }
    _ACTIVE_RUN_TRANSACTIONS[out_dir.as_posix()] = transaction
    return transaction


def _rollback_run_transaction(
    output_dir: Path,
    transaction: Mapping[str, Any] | None,
) -> None:
    """Restore the exact output directory image captured before preparation."""

    if not isinstance(transaction, Mapping):
        return
    transaction_root = Path(str(transaction.get("transaction_root") or ""))
    snapshot = Path(str(transaction.get("snapshot_path") or ""))
    out_dir = Path(output_dir).resolve()
    if not transaction_root.is_dir() or not snapshot.is_dir():
        return
    failed = transaction_root / "failed"
    output_existed = bool(transaction.get("output_existed"))
    moved_current = False
    try:
        if out_dir.exists() or out_dir.is_symlink():
            os.replace(out_dir, failed)
            moved_current = True
        if output_existed:
            os.replace(snapshot, out_dir)
    except OSError:
        if moved_current and not out_dir.exists() and failed.exists():
            os.replace(failed, out_dir)
        raise
    finally:
        if failed.is_dir() and not failed.is_symlink():
            shutil.rmtree(failed, ignore_errors=True)
        elif failed.exists() or failed.is_symlink():
            failed.unlink(missing_ok=True)
        shutil.rmtree(transaction_root, ignore_errors=True)
        _ACTIVE_RUN_TRANSACTIONS.pop(out_dir.as_posix(), None)


def _commit_run_transaction(transaction: Mapping[str, Any] | None) -> None:
    if not isinstance(transaction, Mapping):
        return
    transaction_root = Path(str(transaction.get("transaction_root") or ""))
    if transaction_root.is_dir():
        shutil.rmtree(transaction_root, ignore_errors=True)
    output_dir = str(transaction.get("output_dir") or "")
    if output_dir:
        _ACTIVE_RUN_TRANSACTIONS.pop(output_dir, None)


def run_review_output_transaction(
    output_dir: Path,
    operation: Callable[[Path], dict[str, Any]],
    *,
    expected_predecessor_checkpoint: str | None = None,
) -> dict[str, Any]:
    """Run one browser write on a copy and atomically promote it after checks."""

    requested_dir = Path(output_dir)
    _validate_real_directory(requested_dir, label="review output root")
    out_dir = requested_dir.resolve()
    _physical_tree_entries(out_dir)
    assurance_path = out_dir / "assurance_receipts.json"
    assurance = (
        validate_assurance_run(
            out_dir,
            expected_predecessor_checkpoint=expected_predecessor_checkpoint,
        )
        if assurance_path.exists() or assurance_path.is_symlink()
        else None
    )
    trusted_fingerprint = _tree_fingerprint(out_dir)
    transaction_root = Path(
        tempfile.mkdtemp(prefix=".audit-review-transaction-", dir=out_dir.parent)
    )
    transaction_root.chmod(0o700)
    working_dir = transaction_root / "working"
    shutil.copytree(out_dir, working_dir)
    try:
        result = operation(working_dir)
        entries = _physical_tree_entries(working_dir)
        if assurance is not None:
            output_contract = _validate_workflow_output_contract(
                assurance["workflow_output_contract"]
            )
            _validate_run_tree_allowlist(
                output_dir=working_dir,
                entries=entries,
                output_contract=output_contract,
                expected_predecessor_checkpoint=expected_predecessor_checkpoint,
            )
        working_fingerprint = _tree_fingerprint(working_dir)
        if _tree_fingerprint(out_dir) != trusted_fingerprint:
            raise AssuranceRunError(
                "review output changed during the browser transaction"
            )
        rewritten_result = _rewrite_transaction_paths(
            result,
            working_root=working_dir,
            canonical_root=out_dir,
        )
        backup_dir = transaction_root / "trusted-backup"
        failed_dir = transaction_root / "failed-candidate"
        detached = False
        installed = False
        verified = False
        try:
            os.replace(out_dir, backup_dir)
            detached = True
            if _tree_fingerprint(backup_dir) != trusted_fingerprint:
                raise AssuranceRunError(
                    "review output changed before transaction commit"
                )
            os.replace(working_dir, out_dir)
            installed = True
            if _tree_fingerprint(out_dir) != working_fingerprint:
                raise AssuranceRunError(
                    "review output changed during transaction commit"
                )
            verified = True
        finally:
            if detached and not verified:
                if installed and (out_dir.exists() or out_dir.is_symlink()):
                    os.replace(out_dir, failed_dir)
                os.replace(backup_dir, out_dir)
        return rewritten_result
    finally:
        shutil.rmtree(transaction_root, ignore_errors=True)


def _tree_fingerprint(root: Path) -> tuple[tuple[object, ...], ...]:
    """Return a byte/mode identity for transaction concurrency checks."""

    rows: list[tuple[object, ...]] = []
    root_stat = root.lstat()
    rows.append((".", "directory", stat.S_IMODE(root_stat.st_mode)))
    for entry in _physical_tree_entries(root):
        path = root / entry["path"]
        current = path.lstat()
        if entry["kind"] == "directory":
            rows.append(
                (
                    entry["path"],
                    "directory",
                    stat.S_IMODE(current.st_mode),
                )
            )
        else:
            rows.append(
                (
                    entry["path"],
                    "file",
                    stat.S_IMODE(current.st_mode),
                    current.st_size,
                    hashlib.sha256(path.read_bytes()).hexdigest(),
                )
            )
    return tuple(rows)


def _rewrite_transaction_paths(
    value: object,
    *,
    working_root: Path,
    canonical_root: Path,
) -> Any:
    if isinstance(value, Mapping):
        return {
            str(key): _rewrite_transaction_paths(
                item,
                working_root=working_root,
                canonical_root=canonical_root,
            )
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [
            _rewrite_transaction_paths(
                item,
                working_root=working_root,
                canonical_root=canonical_root,
            )
            for item in value
        ]
    if isinstance(value, str):
        working_text = working_root.as_posix()
        if value == working_text:
            return canonical_root.as_posix()
        if value.startswith(f"{working_text}/"):
            return canonical_root.as_posix() + value[len(working_text) :]
    return value


def rollback_assurance_run(output_dir: Path) -> None:
    """Rollback an active prepared run after any downstream workflow failure."""

    out_dir = Path(output_dir).resolve()
    _rollback_run_transaction(
        out_dir,
        _ACTIVE_RUN_TRANSACTIONS.get(out_dir.as_posix()),
    )


def _validated_source_decisions(
    *,
    source_receipts: Sequence[Mapping[str, Any]],
    reviewed_source_decisions: Sequence[Mapping[str, Any]],
    run_date: str,
) -> list[dict[str, Any]]:
    """Require one exact v2 mapping decision for every source artifact."""

    receipt_by_id = {
        str(receipt["artifact_id"]): receipt for receipt in source_receipts
    }
    if len(reviewed_source_decisions) != len(receipt_by_id):
        raise AssuranceRunError(
            "exactly one current reviewed source decision is required per source"
        )
    normalized: list[dict[str, Any]] = []
    seen_source_refs: set[str] = set()
    for raw_decision in reviewed_source_decisions:
        preliminary = validate_reviewed_decision_receipt(
            raw_decision,
            expected_decision_type=SOURCE_DECISION_TYPE,
            expected_adapter_version=SOURCE_DECISION_ADAPTER_VERSION,
            require_reviewed=True,
        )
        refs = list(preliminary["source_artifact_refs"])
        if len(refs) != 1:
            raise AssuranceRunError(
                "each reviewed source decision must bind exactly one source"
            )
        source_ref = str(refs[0])
        receipt = receipt_by_id.get(source_ref)
        if receipt is None or source_ref in seen_source_refs:
            raise AssuranceRunError(
                "reviewed source decision references a stale source identity"
            )
        relative_path = str(receipt["path"])
        adapter_family = str(preliminary["adapter_id"])
        current = validate_reviewed_decision_receipt(
            preliminary,
            expected_decision_id=_artifact_id("source_mapping", relative_path),
            expected_decision_type=SOURCE_DECISION_TYPE,
            expected_source_artifact_refs=[source_ref],
            expected_adapter_id=adapter_family,
            expected_adapter_version=SOURCE_DECISION_ADAPTER_VERSION,
            require_reviewed=True,
        )
        _review_metadata(
            reviewer_ref=current["reviewer_ref"],
            reviewed_on=current["reviewed_on"],
            run_date=run_date,
            label=f"reviewed_source_decisions.{relative_path}",
        )
        content = current["content"]
        replay_input = {
            "role": content.get("role"),
            "adapter_family": content.get("adapter_family"),
            "reviewer_ref": current["reviewer_ref"],
            "reviewed_on": current["reviewed_on"],
            "perimeter": content.get("perimeter"),
            "money": content.get("money"),
            "date": content.get("date"),
        }
        expected_content = _validate_source_decision_content(
            replay_input,
            relative_path=relative_path,
            adapter_family=adapter_family,
            run_date=run_date,
        )
        if dict(content) != expected_content:
            raise AssuranceRunError("reviewed source decision content is stale")
        normalized.append(current)
        seen_source_refs.add(source_ref)
    if seen_source_refs != set(receipt_by_id):
        raise AssuranceRunError(
            "reviewed source decisions do not cover the exact source set"
        )
    return sorted(normalized, key=lambda item: str(item["source_artifact_refs"][0]))


def _source_receipt_for_row(
    row: Mapping[str, Any],
    source_receipts: Sequence[Mapping[str, Any]],
) -> Mapping[str, Any] | None:
    source_file = str(row.get("source_file") or "").split("!", 1)[0]
    exact = [
        receipt for receipt in source_receipts if str(receipt["path"]) == source_file
    ]
    if len(exact) == 1:
        return exact[0]
    by_name = [
        receipt
        for receipt in source_receipts
        if Path(str(receipt["path"])).name == source_file
    ]
    return by_name[0] if len(by_name) == 1 else None


def _validate_prepared_population(
    *,
    open_items: Sequence[Mapping[str, Any]],
    evidence_rows: Sequence[Mapping[str, Any]],
    source_receipts: Sequence[Mapping[str, Any]],
    reviewed_source_decisions: Sequence[Mapping[str, Any]],
    source_qualifications: Sequence[Mapping[str, Any]],
) -> None:
    """Close source identity, qualification, perimeter, and cent precision."""

    rows = [*open_items, *evidence_rows]
    record_ids: list[str] = []
    decision_by_ref = {
        str(decision["source_artifact_refs"][0]): decision
        for decision in reviewed_source_decisions
    }
    row_count_by_ref = {str(receipt["artifact_id"]): 0 for receipt in source_receipts}
    for index, row in enumerate(rows):
        record_id = row.get("record_id")
        if (
            not isinstance(record_id, str)
            or not record_id
            or record_id != record_id.strip()
        ):
            raise AssuranceRunError(
                f"prepared record {index} requires a canonical record_id"
            )
        record_ids.append(record_id)
        if not source_receipts:
            continue
        receipt = _source_receipt_for_row(row, source_receipts)
        if receipt is None:
            raise AssuranceRunError(
                f"prepared record {record_id} lacks one unambiguous current source"
            )
        source_ref = str(receipt["artifact_id"])
        row_count_by_ref[source_ref] += 1
        decision = decision_by_ref[source_ref]
        content = decision["content"]
        perimeter = content["perimeter"]
        money = content["money"]
        exact_fields = {
            "source_role": content["role"],
            "entity_ref": perimeter["entity_ref"],
            "party_ref": perimeter["party_ref"],
            "currency": perimeter["currency"],
            "unit": perimeter["unit"],
            "direction_policy": perimeter["direction_policy"],
            "allocation_policy": perimeter["allocation_policy"],
            "reported_unit": money["reported_unit"],
            "reported_increment": SUPPORTED_REPORTED_INCREMENT,
        }
        for field, expected in exact_fields.items():
            if row.get(field) != expected:
                raise AssuranceRunError(
                    f"prepared record {record_id} has stale reviewed {field}"
                )
        amount = _canonical_money(row.get("amount"))
        raw_amount = row.get("amount")
        if raw_amount is not None and raw_amount != "" and amount is None:
            raise AssuranceRunError(
                f"prepared record {record_id} has an invalid populated amount"
            )
        if amount is not None and Decimal(amount) % Decimal(
            SUPPORTED_REPORTED_INCREMENT
        ):
            raise AssuranceRunError(
                f"prepared record {record_id} exceeds supported 0.01 precision"
            )
        for field in ("document_date", "posting_date", "value_date"):
            populated_date = row.get(field)
            if populated_date is None or populated_date == "":
                continue
            try:
                _canonical_iso_date(
                    populated_date,
                    label=f"prepared record {record_id}.{field}",
                )
            except AssuranceRunError as exc:
                raise AssuranceRunError(
                    f"prepared record {record_id} has an invalid populated critical date"
                ) from exc
    if len(record_ids) != len(set(record_ids)):
        raise AssuranceRunError("prepared record identities must be unique")
    if not source_receipts:
        if source_qualifications:
            raise AssuranceRunError(
                "source qualifications require corresponding source receipts"
            )
        return
    qualification_by_ref = {
        str(qualification["source_artifact_refs"][0]): qualification
        for qualification in source_qualifications
    }
    for source_ref, row_count in row_count_by_ref.items():
        qualification = qualification_by_ref[source_ref]
        if qualification["reviewed_mapping_ref"] is None:
            raise AssuranceRunError(
                "every source qualification requires a reviewed mapping reference"
            )
        if qualification["status"] != "qualified" and row_count:
            raise AssuranceRunError(
                "unqualified sources cannot emit plausible prepared rows"
            )
        if int(qualification["emitted_row_count"]) != row_count:
            raise AssuranceRunError("source qualification emitted row count is stale")


def prepare_assurance_run(
    *,
    output_dir: Path,
    open_items: Sequence[Mapping[str, Any]],
    evidence_rows: Sequence[Mapping[str, Any]],
    assumptions: Mapping[str, Any],
    source_root: Path | None = None,
    source_receipts: Sequence[Mapping[str, Any]] = (),
    reviewed_source_decisions: Sequence[Mapping[str, Any]] = (),
    source_qualifications: Sequence[Mapping[str, Any]] = (),
    professional_review_authority: Mapping[str, Any] | None = None,
    expected_predecessor_checkpoint: str | None = None,
) -> dict[str, Any]:
    """Replay inputs and implementation, then seal the prepared population."""

    out_dir = Path(output_dir).resolve()
    plugin_root = Path(__file__).resolve().parents[1]
    implementation_receipts = build_implementation_receipts(plugin_root)
    roots = _implementation_roots(plugin_root)
    if source_receipts:
        if source_root is None:
            raise AssuranceRunError("source_root is required for source receipts")
        roots["source"] = Path(source_root).resolve()
    normalized_sources = validate_receipt_set(roots, source_receipts)
    if source_receipts and source_root is not None:
        _validate_source_boundary(Path(source_root), normalized_sources)
    normalized_implementation = validate_receipt_set(roots, implementation_receipts)
    run_date = _run_date(assumptions)
    normalized_decisions = _validated_source_decisions(
        source_receipts=normalized_sources,
        reviewed_source_decisions=reviewed_source_decisions,
        run_date=run_date,
    )
    normalized_qualifications = _validated_source_qualifications(
        source_receipts=normalized_sources,
        reviewed_source_decisions=normalized_decisions,
        source_qualifications=source_qualifications,
    )
    prepared_assumptions = dict(assumptions)
    prepared_assumptions["assurance_run_date"] = run_date
    prepared_assumptions["amount_tolerance"] = _validated_tolerance(
        prepared_assumptions.get("amount_tolerance", "0.01")
    )
    _validate_prepared_population(
        open_items=open_items,
        evidence_rows=evidence_rows,
        source_receipts=normalized_sources,
        reviewed_source_decisions=normalized_decisions,
        source_qualifications=normalized_qualifications,
    )
    supplied_review_authority = professional_review_authority
    existing_review_path = out_dir / "professional_review.json"
    if supplied_review_authority is None and existing_review_path.is_file():
        _require_single_link_regular(
            existing_review_path,
            label="professional_review.json",
        )
        supplied_review_authority = _read_json_mapping(existing_review_path)
    normalized_review_authority = (
        validate_professional_review_authority(supplied_review_authority)
        if supplied_review_authority is not None
        else None
    )
    review_transition_receipts = validate_review_transition_history(
        out_dir,
        current_professional_review=normalized_review_authority,
        expected_predecessor_checkpoint=expected_predecessor_checkpoint,
    )
    prepared_payload = {
        "schema_version": "audit_reconciliation.prepared_records.v1",
        "open_items": list(open_items),
        "evidence_rows": list(evidence_rows),
        "assumptions": prepared_assumptions,
    }
    transaction = _begin_run_transaction(out_dir)
    try:
        prepared_path = _write_exact_json(
            out_dir / "prepared_records.json",
            prepared_payload,
        )
        prepared_receipt = artifact_receipt(
            out_dir,
            prepared_path,
            artifact_id="prepared.records",
            role="prepared",
            root_id="run",
            media_type="application/json",
        )
        validate_receipt_set(
            {**roots, "run": out_dir},
            [*normalized_sources, *normalized_implementation, prepared_receipt],
        )
    except (OSError, ValueError):
        _rollback_run_transaction(out_dir, transaction)
        raise
    return {
        "schema_version": ASSURANCE_SCHEMA_VERSION,
        "source_root": str(Path(source_root).resolve()) if source_root else None,
        "source_receipts": normalized_sources,
        "reviewed_source_decisions": normalized_decisions,
        "source_qualifications": normalized_qualifications,
        "implementation_receipts": normalized_implementation,
        "prepared_receipt": prepared_receipt,
        "run_date": run_date,
        "professional_review_authority": normalized_review_authority,
        "review_transition_receipts": review_transition_receipts,
        "expected_predecessor_checkpoint": expected_predecessor_checkpoint,
        "_run_transaction": transaction,
    }


def _regular_file_set(root: Path) -> set[str]:
    _validate_real_directory(root, label="final output boundary")
    paths: set[str] = set()
    for path in root.rglob("*"):
        current = path.lstat()
        if path.is_symlink() or not stat.S_ISREG(current.st_mode):
            raise AssuranceRunError("final output boundary contains an unsafe entry")
        if current.st_nlink != 1:
            raise AssuranceRunError(
                "final output boundary files must not be hardlinked"
            )
        paths.add(path.relative_to(root).as_posix())
    return paths


def _physical_tree_entries(root: Path) -> list[dict[str, str]]:
    """Capture every ordinary file and directory without following links."""

    _validate_real_directory(root, label="assurance run root")
    entries: list[dict[str, str]] = []
    for path in sorted(root.rglob("*"), key=lambda item: item.as_posix()):
        current = path.lstat()
        relative = path.relative_to(root).as_posix()
        if path.is_symlink():
            raise AssuranceRunError("assurance run tree contains a symlink")
        if stat.S_ISDIR(current.st_mode):
            kind = "directory"
        elif stat.S_ISREG(current.st_mode):
            if current.st_nlink != 1:
                raise AssuranceRunError(
                    "assurance run files must be ordinary single-link files"
                )
            kind = "file"
        else:
            raise AssuranceRunError(
                "assurance run tree contains a non-file, non-directory entry"
            )
        entries.append({"path": relative, "kind": kind})
    return entries


def _safe_declared_relative_path(value: object) -> str | None:
    if not isinstance(value, str) or not value or "\\" in value:
        return None
    candidate = Path(value)
    if candidate.is_absolute() or any(
        part in {"", ".", ".."} for part in candidate.parts
    ):
        return None
    return candidate.as_posix()


def _declared_revision_paths(output_dir: Path) -> set[str]:
    """Return code-shaped review revision paths from the persisted application."""

    applied_path = output_dir / "applied_decisions.json"
    if not applied_path.exists():
        return set()
    _require_single_link_regular(applied_path, label="applied_decisions.json")
    applied = _read_json_mapping(applied_path)
    declared: set[str] = set()
    for field in (
        "revision_paths",
        "original_backup_paths",
        "backup_paths",
    ):
        values = applied.get(field)
        if isinstance(values, list):
            for value in values:
                relative = _safe_declared_relative_path(value)
                if relative is not None and relative.startswith("revisions/"):
                    declared.add(relative)
    effects = applied.get("effects")
    if isinstance(effects, list):
        for effect in effects:
            if not isinstance(effect, Mapping):
                continue
            for field in ("revision_artifact", "original_artifact_backup"):
                relative = _safe_declared_relative_path(effect.get(field))
                if relative is not None and relative.startswith("revisions/"):
                    declared.add(relative)
    return declared


def _validate_run_tree_allowlist(
    *,
    output_dir: Path,
    entries: Sequence[Mapping[str, str]],
    output_contract: Mapping[str, Any],
    expected_predecessor_checkpoint: str | None = None,
    externally_anchored_run_sha256: str | None = None,
) -> None:
    """Reject every root entry not owned by the sealed workflow contract."""

    native_paths = {str(path) for path in output_contract.get("native_paths", [])}
    declared_final_paths = {
        str(path) for path in output_contract.get("declared_paths", [])
    }
    allowed_root_files = RUN_ROOT_FIXED_FILES | native_paths
    revision_files = _declared_revision_paths(output_dir)
    revision_directories = {
        parent.as_posix()
        for relative in revision_files
        for parent in Path(relative).parents
        if parent.as_posix() not in {".", ""}
    }
    transition_receipts = validate_review_transition_history(
        output_dir,
        expected_predecessor_checkpoint=expected_predecessor_checkpoint,
        _externally_anchored_run_sha256=externally_anchored_run_sha256,
    )
    transition_digests = {
        str(receipt["predecessor_assurance_sha256"]) for receipt in transition_receipts
    }
    transition_directories = (
        {
            REVIEW_TRANSITION_HISTORY_DIRECTORY,
            *(
                (Path(REVIEW_TRANSITION_HISTORY_DIRECTORY) / digest).as_posix()
                for digest in transition_digests
            ),
        }
        if transition_digests
        else set()
    )
    transition_snapshot_files: set[str] = set()
    for digest in transition_digests:
        snapshot_prefix = (
            Path(REVIEW_TRANSITION_HISTORY_DIRECTORY)
            / digest
            / PREDECESSOR_RUN_SNAPSHOT_DIRECTORY
        )
        transition_directories.add(snapshot_prefix.as_posix())
        snapshot_root = output_dir / snapshot_prefix
        for snapshot_entry in _physical_tree_entries(snapshot_root):
            retained_path = (snapshot_prefix / str(snapshot_entry["path"])).as_posix()
            if snapshot_entry["kind"] == "directory":
                transition_directories.add(retained_path)
            else:
                transition_snapshot_files.add(retained_path)
    transition_files = {
        (Path(REVIEW_TRANSITION_HISTORY_DIRECTORY) / digest / filename).as_posix()
        for digest in transition_digests
        for filename in [
            *(item[0] for item in REVIEW_TRANSITION_EVIDENCE_FILES),
            REVIEW_TRANSITION_RECEIPT_FILE,
        ]
    }
    for entry in entries:
        relative = str(entry.get("path") or "")
        kind = str(entry.get("kind") or "")
        parts = Path(relative).parts
        if not parts:
            raise AssuranceRunError("assurance run tree has an invalid path")
        if len(parts) == 1:
            if kind == "file" and relative in allowed_root_files:
                continue
            if kind == "directory" and relative in {
                FINAL_OUTPUT_DIRECTORY,
                RUN_CACHE_DIRECTORY,
                *revision_directories,
                *transition_directories,
            }:
                continue
            raise AssuranceRunError(
                f"assurance run root contains an unrelated entry: {relative}"
            )
        if parts[0] == FINAL_OUTPUT_DIRECTORY:
            nested = Path(*parts[1:]).as_posix()
            if kind != "file" or nested not in declared_final_paths:
                raise AssuranceRunError(
                    "final output boundary contains an unrelated entry"
                )
            continue
        if parts[0] == RUN_CACHE_DIRECTORY:
            continue
        if parts[0] == REVIEW_TRANSITION_HISTORY_DIRECTORY:
            if (
                kind == "file"
                and relative in transition_files | transition_snapshot_files
            ) or (kind == "directory" and relative in transition_directories):
                continue
        if parts[0] == "revisions":
            if (kind == "file" and relative in revision_files) or (
                kind == "directory" and relative in revision_directories
            ):
                continue
        raise AssuranceRunError(
            f"assurance run tree contains an unrelated entry: {relative}"
        )


def _build_run_tree_contract(
    output_dir: Path,
    output_contract: Mapping[str, Any],
    *,
    expected_predecessor_checkpoint: str | None = None,
) -> dict[str, Any]:
    entries = _physical_tree_entries(output_dir)
    _validate_run_tree_allowlist(
        output_dir=output_dir,
        entries=entries,
        output_contract=output_contract,
        expected_predecessor_checkpoint=expected_predecessor_checkpoint,
    )
    content = {
        "schema_version": RUN_TREE_SCHEMA_VERSION,
        "entries": entries,
    }
    return {
        **content,
        "content_sha256": canonical_json_sha256(content),
    }


def _validate_run_tree_contract(
    output_dir: Path,
    value: object,
    output_contract: Mapping[str, Any],
    *,
    expected_predecessor_checkpoint: str | None = None,
    externally_anchored_run_sha256: str | None = None,
) -> dict[str, Any]:
    if not isinstance(value, Mapping) or set(value) != {
        "schema_version",
        "entries",
        "content_sha256",
    }:
        raise AssuranceRunError("assurance run tree contract has invalid fields")
    if value["schema_version"] != RUN_TREE_SCHEMA_VERSION:
        raise AssuranceRunError("unsupported assurance run tree contract")
    raw_entries = value["entries"]
    if not isinstance(raw_entries, list):
        raise AssuranceRunError("assurance run tree entries must be a list")
    entries: list[dict[str, str]] = []
    for entry in raw_entries:
        if (
            not isinstance(entry, Mapping)
            or set(entry) != {"path", "kind"}
            or not isinstance(entry["path"], str)
            or entry["kind"] not in {"file", "directory"}
        ):
            raise AssuranceRunError("assurance run tree entry is invalid")
        entries.append({"path": entry["path"], "kind": entry["kind"]})
    if entries != sorted(entries, key=lambda item: item["path"]):
        raise AssuranceRunError("assurance run tree entries must be ordered")
    paths = [entry["path"] for entry in entries]
    if len(paths) != len(set(paths)):
        raise AssuranceRunError("assurance run tree entries must be unique")
    content = {
        "schema_version": RUN_TREE_SCHEMA_VERSION,
        "entries": entries,
    }
    if value["content_sha256"] != canonical_json_sha256(content):
        raise AssuranceRunError("assurance run tree digest is stale")
    if entries != _physical_tree_entries(output_dir):
        raise AssuranceRunError(
            "assurance run physical file and directory closure is stale"
        )
    _validate_run_tree_allowlist(
        output_dir=output_dir,
        entries=entries,
        output_contract=output_contract,
        expected_predecessor_checkpoint=expected_predecessor_checkpoint,
        externally_anchored_run_sha256=externally_anchored_run_sha256,
    )
    return {**content, "content_sha256": value["content_sha256"]}


def _require_single_link_regular(path: Path, *, label: str) -> None:
    """Reject missing, linked, or special assurance-control paths."""

    try:
        current = Path(path).lstat()
    except FileNotFoundError as exc:
        raise AssuranceRunError(f"{label} is unavailable") from exc
    if (
        Path(path).is_symlink()
        or not stat.S_ISREG(current.st_mode)
        or current.st_nlink != 1
    ):
        raise AssuranceRunError(f"{label} must be an ordinary single-link regular file")


def validate_final_output_inventory(
    output_dir: Path,
    inventory: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Replay exact hashes and declared-vs-physical file-set equality."""

    out_dir = Path(output_dir).resolve()
    payload = (
        dict(inventory)
        if inventory is not None
        else json.loads((out_dir / FINAL_OUTPUT_INVENTORY).read_text(encoding="utf-8"))
    )
    required = {
        "schema_version",
        "boundary_root",
        "declared_paths",
        "artifact_receipts",
        "content_sha256",
    }
    if set(payload) != required:
        raise AssuranceRunError("final output inventory has invalid fields")
    if payload["schema_version"] != "audit_reconciliation.final_outputs.v1":
        raise AssuranceRunError("unsupported final output inventory schema")
    boundary_name = payload["boundary_root"]
    if boundary_name != FINAL_OUTPUT_DIRECTORY:
        raise AssuranceRunError("final output boundary is stale")
    declared = payload["declared_paths"]
    if (
        not isinstance(declared, list)
        or any(not isinstance(item, str) for item in declared)
        or declared != sorted(set(declared))
    ):
        raise AssuranceRunError("declared final output paths must be sorted and unique")
    physical = _regular_file_set(out_dir / boundary_name)
    if physical != set(declared):
        raise AssuranceRunError(
            "declared and physical final output file sets do not match"
        )
    receipts = payload["artifact_receipts"]
    if not isinstance(receipts, list):
        raise AssuranceRunError("artifact_receipts must be a list")
    normalized_receipts = validate_receipt_set(
        {"final": out_dir / boundary_name},
        receipts,
    )
    if {str(item["path"]) for item in normalized_receipts} != physical:
        raise AssuranceRunError("final output receipts do not cover the full boundary")
    content = {
        "schema_version": payload["schema_version"],
        "boundary_root": boundary_name,
        "declared_paths": declared,
        "artifact_receipts": normalized_receipts,
    }
    if payload["content_sha256"] != canonical_json_sha256(content):
        raise AssuranceRunError("final output inventory content digest is stale")
    return {**content, "content_sha256": payload["content_sha256"]}


def _publish_final_outputs(
    output_dir: Path,
    declared_outputs: Sequence[Path],
    output_contract: Mapping[str, Any],
) -> dict[str, Any]:
    boundary = output_dir / FINAL_OUTPUT_DIRECTORY
    names = [Path(path).name for path in declared_outputs]
    if len(names) != len(set(names)):
        raise AssuranceRunError("final output filenames must be unique")
    expected = set(output_contract["declared_paths"])
    if expected != {*names, "reconciliation_results.json"}:
        raise AssuranceRunError(
            "declared outputs do not match the closed workflow output contract"
        )
    if boundary.exists():
        existing = _regular_file_set(boundary)
        if existing - expected:
            raise AssuranceRunError(
                "unexpected file exists in the final output boundary"
            )
    else:
        boundary.mkdir(parents=True)
    for source in declared_outputs:
        current = Path(source).resolve()
        _require_single_link_regular(
            current,
            label=f"declared final output {current.name}",
        )
        shutil.copyfile(current, boundary / current.name)
    physical = _regular_file_set(boundary)
    if physical != expected:
        raise AssuranceRunError(
            "declared and physical final output file sets do not match"
        )
    receipts = [
        artifact_receipt(
            boundary,
            boundary / name,
            artifact_id=_artifact_id("final_output", name),
            role=(
                "workpaper"
                if Path(name).suffix.lower() in {".xlsx", ".xlsm"}
                else "report" if Path(name).suffix.lower() == ".docx" else "output"
            ),
            root_id="final",
            media_type=_media_type(Path(name)),
        )
        for name in sorted(expected)
    ]
    content = {
        "schema_version": "audit_reconciliation.final_outputs.v1",
        "boundary_root": FINAL_OUTPUT_DIRECTORY,
        "declared_paths": sorted(expected),
        "artifact_receipts": receipts,
    }
    inventory = {**content, "content_sha256": canonical_json_sha256(content)}
    _write_exact_json(output_dir / FINAL_OUTPUT_INVENTORY, inventory)
    return validate_final_output_inventory(output_dir, inventory)


def _workflow_output_contract(
    declared_outputs: Sequence[Path],
) -> dict[str, Any]:
    """Close one code-owned output profile before inspecting the boundary."""

    native_paths = sorted(Path(path).name for path in declared_outputs)
    if len(native_paths) != len(set(native_paths)):
        raise AssuranceRunError("native workflow output filenames must be unique")
    if "reconciliation_results.json" in native_paths:
        raise AssuranceRunError(
            "native workflow outputs cannot replace the reconciliation control output"
        )
    if frozenset(native_paths) not in TRUSTED_NATIVE_OUTPUT_PROFILES:
        raise AssuranceRunError(
            "native workflow outputs do not match a trusted workflow-owned "
            "profile or exact material record-id set"
        )
    declared_paths = sorted({*native_paths, "reconciliation_results.json"})
    content = {
        "schema_version": "audit_reconciliation.output_contract.v1",
        "native_paths": native_paths,
        "required_control_paths": ["reconciliation_results.json"],
        "declared_paths": declared_paths,
    }
    return {**content, "content_sha256": canonical_json_sha256(content)}


def _validate_workflow_output_contract(value: object) -> dict[str, Any]:
    if not isinstance(value, Mapping):
        raise AssuranceRunError("workflow output contract must be an object")
    required = {
        "schema_version",
        "native_paths",
        "required_control_paths",
        "declared_paths",
        "content_sha256",
    }
    if set(value) != required:
        raise AssuranceRunError("workflow output contract has invalid fields")
    if value["schema_version"] != "audit_reconciliation.output_contract.v1":
        raise AssuranceRunError("unsupported workflow output contract")
    native_paths = value["native_paths"]
    control_paths = value["required_control_paths"]
    declared_paths = value["declared_paths"]
    for label, paths in (
        ("native_paths", native_paths),
        ("required_control_paths", control_paths),
        ("declared_paths", declared_paths),
    ):
        if (
            not isinstance(paths, list)
            or any(
                not isinstance(item, str) or not item or Path(item).name != item
                for item in paths
            )
            or paths != sorted(set(paths))
        ):
            raise AssuranceRunError(
                f"workflow output contract {label} must be sorted basenames"
            )
    if control_paths != ["reconciliation_results.json"]:
        raise AssuranceRunError("workflow control output contract is stale")
    if frozenset(native_paths) not in TRUSTED_NATIVE_OUTPUT_PROFILES:
        raise AssuranceRunError(
            "workflow output contract is not a trusted workflow-owned profile"
        )
    if set(declared_paths) != {*native_paths, *control_paths}:
        raise AssuranceRunError("workflow declared output contract does not close")
    content = {
        "schema_version": value["schema_version"],
        "native_paths": native_paths,
        "required_control_paths": control_paths,
        "declared_paths": declared_paths,
    }
    if value["content_sha256"] != canonical_json_sha256(content):
        raise AssuranceRunError("workflow output contract digest is stale")
    return {**content, "content_sha256": value["content_sha256"]}


def _review_gate_status(
    review_rows: Sequence[Mapping[str, Any]],
    reconciliation_rows: Sequence[Mapping[str, Any]],
    *,
    run_date: str,
) -> str:
    """Require exact reviewed-record closure at one sealed run date."""

    reviewable_rows = [
        row
        for row in reconciliation_rows
        if str(row.get("reconciliation_status") or "").strip().lower() != "out_of_scope"
    ]
    if not reviewable_rows:
        return "not_applicable" if not review_rows else "failed"
    if not review_rows:
        return "withheld"
    expected_ids = [row.get("record_id") for row in reviewable_rows]
    if any(
        not isinstance(record_id, str)
        or not record_id
        or record_id != record_id.strip()
        for record_id in expected_ids
    ) or len(expected_ids) != len(set(expected_ids)):
        return "failed"
    actual_ids = [row.get("record_id") for row in review_rows]
    if (
        any(
            not isinstance(record_id, str)
            or not record_id
            or record_id != record_id.strip()
            for record_id in actual_ids
        )
        or len(actual_ids) != len(set(actual_ids))
        or set(actual_ids) != set(expected_ids)
    ):
        return "failed"
    statuses = {
        str(row.get("review_status") or "").strip().upper() for row in review_rows
    }
    if not statuses <= {"PASS", "FAIL", "PENDING", "UNRESOLVED"}:
        return "failed"
    if "FAIL" in statuses:
        return "failed"
    for row in review_rows:
        if not row.get("reviewer_ref") or not row.get("reviewed_on"):
            return "withheld"
    try:
        for index, row in enumerate(review_rows):
            _review_metadata(
                reviewer_ref=row.get("reviewer_ref"),
                reviewed_on=row.get("reviewed_on"),
                run_date=run_date,
                label=f"review_rows[{index}]",
            )
    except AssuranceRunError:
        return "failed"
    if statuses == {"PASS"}:
        return "passed"
    return "withheld"


def _review_projection(
    review_rows: Sequence[Mapping[str, Any]],
) -> list[dict[str, Any]]:
    """Project mutable display rows onto the professional authority contract."""

    return [
        {
            "record_id": row.get("record_id"),
            "review_status": row.get("review_status"),
            "reviewer_ref": row.get("reviewer_ref"),
            "reviewed_on": row.get("reviewed_on"),
        }
        for row in review_rows
    ]


def build_professional_review_authority(
    review_rows: Sequence[Mapping[str, Any]],
    *,
    origin: str = "direct_persisted_review",
    run_id: str | None = None,
    decision_fingerprint: str | None = None,
    predecessor_assurance_sha256: str | None = None,
) -> dict[str, Any]:
    """Persist review claims; reviewer_ref remains unsigned and untrusted."""

    if origin not in {"direct_persisted_review", "applied_decisions"}:
        raise AssuranceRunError("professional review authority origin is unsupported")
    records = _json_safe(
        _review_projection(review_rows),
        label="professional_review.records",
    )
    fingerprint = decision_fingerprint or canonical_json_sha256(
        {"origin": origin, "records": records}
    )
    if re.fullmatch(r"[0-9a-f]{64}", fingerprint) is None:
        raise AssuranceRunError(
            "professional review decision fingerprint must be SHA-256 text"
        )
    if (
        predecessor_assurance_sha256 is not None
        and re.fullmatch(
            r"[0-9a-f]{64}",
            predecessor_assurance_sha256,
        )
        is None
    ):
        raise AssuranceRunError(
            "professional review predecessor digest must be SHA-256 text"
        )
    content = {
        "schema_version": "audit_reconciliation.professional_review.v1",
        "origin": origin,
        "run_id": run_id,
        "records": records,
        "reviewer_ref_trust": REVIEWER_REF_TRUST,
        "decision_fingerprint": fingerprint,
        "predecessor_assurance_sha256": predecessor_assurance_sha256,
    }
    return {**content, "content_sha256": canonical_json_sha256(content)}


def validate_professional_review_authority(value: object) -> dict[str, Any]:
    """Replay the persisted review authority without granting semantic status."""

    if not isinstance(value, Mapping):
        raise AssuranceRunError("professional review authority must be an object")
    required = {
        "schema_version",
        "origin",
        "run_id",
        "records",
        "reviewer_ref_trust",
        "decision_fingerprint",
        "predecessor_assurance_sha256",
        "content_sha256",
    }
    if set(value) != required:
        raise AssuranceRunError("professional review authority has invalid fields")
    if value["reviewer_ref_trust"] != REVIEWER_REF_TRUST:
        raise AssuranceRunError(
            "professional review reviewer_ref trust boundary is stale"
        )
    records = value["records"]
    if not isinstance(records, list) or any(
        not isinstance(record, Mapping)
        or set(record) != {"record_id", "review_status", "reviewer_ref", "reviewed_on"}
        for record in records
    ):
        raise AssuranceRunError("professional review authority records are invalid")
    normalized = build_professional_review_authority(
        records,
        origin=str(value["origin"]),
        run_id=(
            str(value["run_id"])
            if isinstance(value["run_id"], str) and value["run_id"]
            else None
        ),
        decision_fingerprint=str(value["decision_fingerprint"]),
        predecessor_assurance_sha256=(
            str(value["predecessor_assurance_sha256"])
            if value["predecessor_assurance_sha256"] is not None
            else None
        ),
    )
    if value["content_sha256"] != normalized["content_sha256"]:
        raise AssuranceRunError("professional review authority digest is stale")
    return normalized


def _review_item_record_id(item: Mapping[str, Any]) -> str | None:
    data = item.get("data")
    if not isinstance(data, Mapping):
        return None
    record_id = data.get("record_id")
    if (
        not isinstance(record_id, str)
        or not record_id
        or record_id != record_id.strip()
    ):
        record_id = None
    if record_id is None and data.get("target_id_field") == "record_id":
        candidate = data.get("target_record_id")
        if isinstance(candidate, str) and candidate and candidate == candidate.strip():
            record_id = candidate
    return record_id


def build_review_payload_mapping(
    review_payload: Mapping[str, Any],
) -> dict[str, Any]:
    """Bind the exact ordered review item namespace to accounting records."""

    run_id = review_payload.get("run_id")
    items = review_payload.get("items")
    if (
        not isinstance(run_id, str)
        or not run_id
        or run_id != run_id.strip()
        or not isinstance(items, list)
    ):
        raise AssuranceRunError("review payload mapping source is invalid")
    mapped_items: list[dict[str, Any]] = []
    seen: set[str] = set()
    for index, item in enumerate(items):
        if not isinstance(item, Mapping):
            raise AssuranceRunError(f"review payload item {index} is invalid")
        item_id = item.get("id")
        if (
            not isinstance(item_id, str)
            or not item_id
            or item_id != item_id.strip()
            or item_id in seen
        ):
            raise AssuranceRunError("review payload item identities are invalid")
        seen.add(item_id)
        mapped_items.append(
            {
                "item_id": item_id,
                "record_id": _review_item_record_id(item),
            }
        )
    content = {
        "schema_version": REVIEW_PAYLOAD_MAPPING_SCHEMA_VERSION,
        "run_id": run_id,
        "items": mapped_items,
    }
    return {**content, "content_sha256": canonical_json_sha256(content)}


def _validate_review_payload_mapping(value: object) -> dict[str, Any]:
    if not isinstance(value, Mapping) or set(value) != {
        "schema_version",
        "run_id",
        "items",
        "content_sha256",
    }:
        raise AssuranceRunError("review payload mapping has invalid fields")
    if value["schema_version"] != REVIEW_PAYLOAD_MAPPING_SCHEMA_VERSION:
        raise AssuranceRunError("unsupported review payload mapping")
    run_id = value["run_id"]
    raw_items = value["items"]
    if (
        not isinstance(run_id, str)
        or not run_id
        or run_id != run_id.strip()
        or not isinstance(raw_items, list)
    ):
        raise AssuranceRunError("review payload mapping is invalid")
    items: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in raw_items:
        if not isinstance(item, Mapping) or set(item) != {"item_id", "record_id"}:
            raise AssuranceRunError("review payload mapping item is invalid")
        item_id = item["item_id"]
        record_id = item["record_id"]
        if (
            not isinstance(item_id, str)
            or not item_id
            or item_id != item_id.strip()
            or item_id in seen
            or (
                record_id is not None
                and (
                    not isinstance(record_id, str)
                    or not record_id
                    or record_id != record_id.strip()
                )
            )
        ):
            raise AssuranceRunError("review payload mapping item is invalid")
        seen.add(item_id)
        items.append({"item_id": item_id, "record_id": record_id})
    content = {
        "schema_version": REVIEW_PAYLOAD_MAPPING_SCHEMA_VERSION,
        "run_id": run_id,
        "items": items,
    }
    if value["content_sha256"] != canonical_json_sha256(content):
        raise AssuranceRunError("review payload mapping digest is stale")
    return {**content, "content_sha256": value["content_sha256"]}


def review_decision_fingerprint(
    review_payload_mapping: Mapping[str, Any],
    effects: Sequence[Mapping[str, Any]],
) -> str:
    """Rederive the decision fingerprint from retained mapping and effects."""

    mapping = _validate_review_payload_mapping(review_payload_mapping)
    record_by_item = {
        str(item["item_id"]): item["record_id"] for item in mapping["items"]
    }
    records: list[dict[str, Any]] = []
    seen: set[str] = set()
    for effect in effects:
        if not isinstance(effect, Mapping):
            raise AssuranceRunError("applied review effect is invalid")
        item_id = effect.get("item_id")
        if (
            not isinstance(item_id, str)
            or item_id not in record_by_item
            or item_id in seen
        ):
            raise AssuranceRunError("applied review effect item is invalid")
        seen.add(item_id)
        requested_documents = effect.get("requested_documents", [])
        if not isinstance(requested_documents, list):
            raise AssuranceRunError("applied review requested_documents must be a list")
        records.append(
            {
                "item_id": item_id,
                "record_id": record_by_item[item_id],
                "action": effect.get("action"),
                "edit_value": effect.get("edit_value"),
                "requested_documents": requested_documents,
            }
        )
    records.sort(key=lambda item: str(item["item_id"]))
    return canonical_json_sha256(
        {
            "run_id": mapping["run_id"],
            "decisions": _json_safe(records, label="review transition decisions"),
        }
    )


def _validated_assurance_envelope(value: object) -> dict[str, Any]:
    if not isinstance(value, Mapping) or set(value) != ASSURANCE_SEAL_FIELDS:
        raise AssuranceRunError("assurance receipt seal has invalid fields")
    if value["schema_version"] != ASSURANCE_SCHEMA_VERSION:
        raise AssuranceRunError("unsupported assurance receipt seal")
    content = {key: value[key] for key in ASSURANCE_SEAL_FIELDS - {"content_sha256"}}
    if value["content_sha256"] != canonical_json_sha256(content):
        raise AssuranceRunError("assurance receipt seal digest is stale")
    run_id = value["run_id"]
    if run_id is not None and (
        not isinstance(run_id, str) or not run_id or run_id != run_id.strip()
    ):
        raise AssuranceRunError("assurance run identity is invalid")
    _canonical_iso_date(value["run_date"], label="run_date")
    return dict(value)


def _validated_expected_predecessor_checkpoint(value: object) -> str:
    """Validate an externally retained predecessor digest.

    The digest is a mechanical equality checkpoint only. Its provenance and
    authority depend entirely on the separate channel from which the caller
    supplies it; no local run-tree value is allowed to substitute for it.
    """

    if (
        not isinstance(value, str)
        or value != value.strip()
        or _CHECKPOINT_PATTERN.fullmatch(value) is None
    ):
        raise AssuranceRunError(
            "an external expected predecessor checkpoint is required"
        )
    return value


def _require_expected_predecessor_checkpoint(
    expected: object,
    observed: object,
) -> str:
    checkpoint = _validated_expected_predecessor_checkpoint(expected)
    if checkpoint != observed:
        raise AssuranceRunError(
            "external expected predecessor checkpoint does not match"
        )
    return checkpoint


def build_applied_review_authority(
    *,
    predecessor_assurance: Mapping[str, Any],
    predecessor_professional_review: Mapping[str, Any],
    predecessor_reconciliation: Mapping[str, Any],
    review_payload_mapping: Mapping[str, Any],
    effects: Sequence[Mapping[str, Any]],
    reviewer_ref: str | None,
) -> dict[str, Any]:
    """Rederive successor review records from retained predecessor evidence."""

    predecessor = _validated_assurance_envelope(predecessor_assurance)
    prior_authority = validate_professional_review_authority(
        predecessor_professional_review
    )
    if prior_authority != predecessor["professional_review_authority"]:
        raise AssuranceRunError(
            "retained predecessor professional review is not seal-bound"
        )
    mapping = _validate_review_payload_mapping(review_payload_mapping)
    rows = predecessor_reconciliation.get("reconciliation_rows")
    if not isinstance(rows, list):
        raise AssuranceRunError(
            "retained predecessor reconciliation population is invalid"
        )
    expected_ids = [
        str(row.get("record_id"))
        for row in rows
        if isinstance(row, Mapping)
        and isinstance(row.get("record_id"), str)
        and row.get("record_id")
        and str(row.get("reconciliation_status") or "").strip().lower()
        != "out_of_scope"
    ]
    if len(expected_ids) != len(set(expected_ids)):
        raise AssuranceRunError(
            "professional review source-row identities are not unique"
        )
    if (
        prior_authority["run_id"] is not None
        and mapping["run_id"] != prior_authority["run_id"]
    ):
        raise AssuranceRunError("retained review payload run identity is stale")
    mapped_record_ids = {
        str(item["record_id"])
        for item in mapping["items"]
        if item["record_id"] is not None
    }
    if not mapped_record_ids.issubset(expected_ids):
        raise AssuranceRunError(
            "retained review mapping references an unknown source row"
        )
    metadata_by_id = {
        str(record["record_id"]): record
        for record in prior_authority["records"]
        if isinstance(record, Mapping)
    }
    if not set(expected_ids).issubset(metadata_by_id):
        raise AssuranceRunError(
            "retained predecessor professional review population is stale"
        )
    record_by_item = {
        str(item["item_id"]): item["record_id"] for item in mapping["items"]
    }
    effect_by_record: dict[str, Mapping[str, Any]] = {}
    for effect in effects:
        if not isinstance(effect, Mapping):
            raise AssuranceRunError("applied review effect is invalid")
        item_id = effect.get("item_id")
        if not isinstance(item_id, str) or item_id not in record_by_item:
            raise AssuranceRunError("applied review effect item is invalid")
        record_id = record_by_item[item_id]
        if record_id is None:
            continue
        if record_id in effect_by_record:
            raise AssuranceRunError(
                f"professional review has duplicate decisions for {record_id}"
            )
        effect_by_record[record_id] = effect
    fingerprint = review_decision_fingerprint(mapping, effects)
    records: list[dict[str, Any]] = []
    for record_id in expected_ids:
        prior = metadata_by_id[record_id]
        effect = effect_by_record.get(record_id)
        action = effect.get("action") if effect is not None else None
        records.append(
            {
                "record_id": record_id,
                "review_status": (
                    "PASS"
                    if action == "accept"
                    else "FAIL" if action == "reject" else "UNRESOLVED"
                ),
                "reviewer_ref": (
                    reviewer_ref or prior.get("reviewer_ref") or "reviewer.local"
                ),
                "reviewed_on": (
                    prior.get("reviewed_on") or str(predecessor["run_date"])[:10]
                ),
            }
        )
    return build_professional_review_authority(
        records,
        origin="applied_decisions",
        run_id=str(mapping["run_id"]),
        decision_fingerprint=fingerprint,
        predecessor_assurance_sha256=str(predecessor["content_sha256"]),
    )


def _stable_regular_bytes(path: Path, *, label: str) -> bytes:
    source = Path(path)
    try:
        path_stat = source.lstat()
    except FileNotFoundError as exc:
        raise AssuranceRunError(f"{label} is unavailable") from exc
    if (
        source.is_symlink()
        or not stat.S_ISREG(path_stat.st_mode)
        or path_stat.st_nlink != 1
    ):
        raise AssuranceRunError(f"{label} must be an ordinary single-link regular file")
    flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    descriptor = os.open(source, flags)
    try:
        before = os.fstat(descriptor)
        with os.fdopen(descriptor, "rb", closefd=False) as handle:
            payload = handle.read()
        after = os.fstat(descriptor)
    finally:
        os.close(descriptor)
    identity = (
        path_stat.st_dev,
        path_stat.st_ino,
        path_stat.st_size,
        path_stat.st_mtime_ns,
        path_stat.st_nlink,
    )
    before_identity = (
        before.st_dev,
        before.st_ino,
        before.st_size,
        before.st_mtime_ns,
        before.st_nlink,
    )
    after_identity = (
        after.st_dev,
        after.st_ino,
        after.st_size,
        after.st_mtime_ns,
        after.st_nlink,
    )
    if (
        identity != before_identity
        or before_identity != after_identity
        or len(payload) != after.st_size
    ):
        raise AssuranceRunError(f"{label} changed while it was read")
    return payload


def _write_exact_bytes(path: Path, payload: bytes) -> Path:
    destination = Path(path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    handle, temporary_name = tempfile.mkstemp(
        prefix=f".{destination.name}.",
        suffix=".tmp",
        dir=destination.parent,
    )
    temporary = Path(temporary_name)
    try:
        with os.fdopen(handle, "wb") as stream:
            stream.write(payload)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)
    return destination


def capture_review_transition_predecessor(
    output_dir: Path,
    capture_dir: Path,
    *,
    expected_predecessor_checkpoint: str | None = None,
) -> dict[str, Any]:
    """Capture and replay the complete predecessor outside the candidate tree."""

    out_dir = Path(output_dir)
    trusted_fingerprint = _tree_fingerprint(out_dir)
    checkpoint = _validated_expected_predecessor_checkpoint(
        expected_predecessor_checkpoint
    )
    assurance = validate_assurance_run(
        out_dir,
        _externally_anchored_run_sha256=checkpoint,
    )
    _require_expected_predecessor_checkpoint(
        checkpoint,
        assurance["content_sha256"],
    )
    capture = Path(capture_dir)
    _validate_real_directory(capture.parent, label="transition capture parent")
    if capture.exists() or capture.is_symlink():
        _validate_real_directory(capture, label="transition capture directory")
        if any(capture.iterdir()):
            raise AssuranceRunError("transition capture directory must be empty")
        capture.chmod(0o700)
    else:
        capture.mkdir(mode=0o700)
    snapshot = capture / PREDECESSOR_RUN_SNAPSHOT_DIRECTORY
    try:
        shutil.copytree(out_dir, snapshot)
        if (
            _tree_fingerprint(out_dir) != trusted_fingerprint
            or _tree_fingerprint(snapshot) != trusted_fingerprint
        ):
            raise AssuranceRunError(
                "predecessor changed while its complete run was captured"
            )
        captured_assurance = validate_assurance_run(
            snapshot,
            _externally_anchored_run_sha256=checkpoint,
        )
        if captured_assurance != assurance:
            raise AssuranceRunError("captured predecessor assurance changed")
    except (OSError, ValueError):
        shutil.rmtree(capture, ignore_errors=True)
        raise
    sources = {
        "predecessor_assurance_receipts.json": snapshot / "assurance_receipts.json",
        "predecessor_professional_review.json": snapshot / "professional_review.json",
        "predecessor_reconciliation_results.json": snapshot
        / FINAL_OUTPUT_DIRECTORY
        / "reconciliation_results.json",
        "predecessor_review_payload.json": snapshot / "review_payload.json",
    }
    try:
        for filename, source in sources.items():
            _write_exact_bytes(
                capture / filename,
                _stable_regular_bytes(source, label=source.name),
            )
        selected_assurance = _read_json_mapping(
            capture / "predecessor_assurance_receipts.json"
        )
        if selected_assurance != captured_assurance:
            raise AssuranceRunError(
                "selected predecessor seal differs from the complete snapshot"
            )
        captured_professional = _read_json_mapping(
            capture / "predecessor_professional_review.json"
        )
        if captured_professional != assurance["professional_review_authority"]:
            raise AssuranceRunError("captured predecessor professional review changed")
        review_payload = _read_json_mapping(capture / "predecessor_review_payload.json")
        build_review_payload_mapping(review_payload)
    except (OSError, ValueError):
        shutil.rmtree(capture, ignore_errors=True)
        raise
    return {
        "capture_dir": capture.as_posix(),
        "predecessor_assurance_sha256": assurance["content_sha256"],
    }


def _receipt_matches_archived_bytes(
    receipt: object,
    archived_path: Path,
    *,
    expected_artifact_id: str,
    expected_path: str,
    expected_role: str,
) -> None:
    if not isinstance(receipt, Mapping):
        raise AssuranceRunError("retained predecessor artifact receipt is invalid")
    required = {
        "schema_version",
        "artifact_id",
        "root_id",
        "role",
        "path",
        "byte_count",
        "sha256",
    }
    if set(receipt) != required and set(receipt) != required | {"media_type"}:
        raise AssuranceRunError("retained predecessor artifact receipt is invalid")
    byte_count, digest = file_snapshot(archived_path)
    if (
        receipt.get("schema_version") != "vera.artifact_receipt.v1"
        or receipt.get("artifact_id") != expected_artifact_id
        or receipt.get("path") != expected_path
        or receipt.get("role") != expected_role
        or receipt.get("byte_count") != byte_count
        or receipt.get("sha256") != digest
    ):
        raise AssuranceRunError(
            "retained predecessor artifact does not match its sealed receipt"
        )


def _validate_archived_applied_decisions(
    applied: object,
    *,
    mapping: Mapping[str, Any],
) -> tuple[list[dict[str, Any]], str | None]:
    if not isinstance(applied, Mapping):
        raise AssuranceRunError("retained applied decisions must be an object")
    decisions = applied.get("decisions")
    effects = applied.get("effects")
    if not isinstance(decisions, list) or not isinstance(effects, list):
        raise AssuranceRunError(
            "retained applied decisions require decisions and effects"
        )
    if (
        applied.get("run_id") != mapping["run_id"]
        or applied.get("decision_count") != len(decisions)
        or applied.get("item_count") != len(mapping["items"])
        or len(decisions) != len(effects)
    ):
        raise AssuranceRunError("retained applied decision counts are stale")
    mapped_ids = {str(item["item_id"]) for item in mapping["items"]}
    normalized_effects: list[dict[str, Any]] = []
    seen: set[str] = set()
    for decision, effect in zip(decisions, effects, strict=True):
        if not isinstance(decision, Mapping) or not isinstance(effect, Mapping):
            raise AssuranceRunError("retained applied decision pair is invalid")
        item_id = decision.get("item_id")
        if (
            not isinstance(item_id, str)
            or item_id not in mapped_ids
            or item_id in seen
            or effect.get("item_id") != item_id
            or effect.get("action") != decision.get("action")
        ):
            raise AssuranceRunError("retained applied decision pair is stale")
        seen.add(item_id)
        normalized_effects.append(dict(effect))
    professional = applied.get("professional_review")
    if (
        not isinstance(professional, Mapping)
        or professional.get("path") != "professional_review.json"
        or professional.get("successor_assurance_replayed") is not False
        or applied.get("application_status") == "final_ready"
    ):
        raise AssuranceRunError(
            "retained first-apply professional review state is invalid"
        )
    reviewer = applied.get("reviewer")
    if reviewer is not None:
        reviewer = _canonical_reviewer_ref(
            reviewer,
            label="retained applied reviewer",
        )
    return normalized_effects, reviewer


def _transition_receipt_content(
    *,
    predecessor_assurance_sha256: str,
    decision_fingerprint: str,
    successor_professional_review_sha256: str,
    artifact_receipts: Sequence[Mapping[str, Any]],
) -> dict[str, Any]:
    return {
        "schema_version": REVIEW_TRANSITION_SCHEMA_VERSION,
        "transition_id": f"review_transition.{predecessor_assurance_sha256}",
        "predecessor_assurance_sha256": predecessor_assurance_sha256,
        "decision_fingerprint": decision_fingerprint,
        "successor_professional_review_sha256": (successor_professional_review_sha256),
        "artifact_receipts": list(artifact_receipts),
    }


def _rederive_predecessor_reconciliation(
    *,
    predecessor: Mapping[str, Any],
    predecessor_snapshot: Path,
    predecessor_reconciliation: Mapping[str, Any],
) -> None:
    """Replay deterministic predecessor rows, allocations, and core checks.

    Reconciliation classification and exact allocation arithmetic are mechanical
    outputs of the retained population, reviewed assumptions, and current
    receipted implementation. Semantic sufficiency remains reviewer judgment.
    """

    prepared = _read_json_mapping(predecessor_snapshot / "prepared_records.json")
    open_items = prepared.get("open_items")
    evidence_rows = prepared.get("evidence_rows")
    assumptions = prepared.get("assumptions")
    if (
        not isinstance(open_items, list)
        or any(not isinstance(row, dict) for row in open_items)
        or not isinstance(evidence_rows, list)
        or any(not isinstance(row, dict) for row in evidence_rows)
        or not isinstance(assumptions, dict)
    ):
        raise AssuranceRunError(
            "retained predecessor prepared population cannot be replayed"
        )
    assumption_currency = assumptions.get("currency")
    if assumption_currency is not None:
        if (
            not isinstance(assumption_currency, str)
            or not assumption_currency
            or assumption_currency != assumption_currency.strip()
        ):
            raise AssuranceRunError("retained predecessor currency is invalid")
        contradictory = [
            str(row.get("record_id") or "")
            for row in [*open_items, *evidence_rows]
            if isinstance(row.get("currency"), str)
            and row["currency"]
            and row["currency"] != assumption_currency
        ]
        if contradictory:
            raise AssuranceRunError(
                "retained predecessor currency contradicts its prepared population"
            )
    expected_rows = reconcile_open_items(
        [dict(row) for row in open_items],
        [dict(row) for row in evidence_rows],
        dict(assumptions),
    )
    retained_rows = predecessor_reconciliation.get("reconciliation_rows")
    if retained_rows != expected_rows:
        raise AssuranceRunError(
            "retained predecessor reconciliation cannot be rederived"
        )
    expected_allocations, failures = closed_bank_allocation_controls(
        expected_rows,
        [dict(row) for row in evidence_rows],
        dict(assumptions),
    )
    normalized_allocations = [
        validate_allocation_ledger(ledger) for ledger in expected_allocations
    ]
    if failures:
        raise AssuranceRunError(
            "retained predecessor allocation controls cannot be rederived"
        )
    if (
        predecessor_reconciliation.get("allocation_ledgers") != normalized_allocations
        or predecessor.get("allocation_ledgers") != normalized_allocations
    ):
        raise AssuranceRunError("retained predecessor allocations cannot be rederived")
    expected_checks = reconciliation_checks(open_items, expected_rows)
    retained_checks = predecessor_reconciliation.get("checks")
    if not isinstance(retained_checks, list):
        raise AssuranceRunError("retained predecessor checks are invalid")
    retained_by_name: dict[str, Mapping[str, Any]] = {}
    for check in retained_checks:
        if not isinstance(check, Mapping):
            raise AssuranceRunError("retained predecessor check is invalid")
        name = check.get("check")
        if not isinstance(name, str) or not name or name in retained_by_name:
            raise AssuranceRunError("retained predecessor check identities are invalid")
        retained_by_name[name] = check
    for expected in expected_checks:
        if retained_by_name.get(str(expected["check"])) != expected:
            raise AssuranceRunError(
                "retained predecessor mechanical checks cannot be rederived"
            )


def _validate_review_transition_directory(
    transition_dir: Path,
    *,
    externally_anchored_predecessor_sha256: str | None = None,
) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    _validate_real_directory(transition_dir, label="review transition directory")
    digest = transition_dir.name
    if re.fullmatch(r"[0-9a-f]{64}", digest) is None:
        raise AssuranceRunError("review transition history path is invalid")
    expected_names = {
        *(filename for filename, _, _ in REVIEW_TRANSITION_EVIDENCE_FILES),
        REVIEW_TRANSITION_RECEIPT_FILE,
        PREDECESSOR_RUN_SNAPSHOT_DIRECTORY,
    }
    observed_names: set[str] = set()
    for path in transition_dir.iterdir():
        observed_names.add(path.name)
        if path.name == PREDECESSOR_RUN_SNAPSHOT_DIRECTORY:
            _validate_real_directory(path, label="predecessor run snapshot")
        else:
            _require_single_link_regular(path, label="review transition artifact")
    if observed_names != expected_names:
        raise AssuranceRunError(
            "review transition directory does not match the exact file contract"
        )
    raw_receipt = _read_json_mapping(transition_dir / REVIEW_TRANSITION_RECEIPT_FILE)
    required = {
        "schema_version",
        "transition_id",
        "predecessor_assurance_sha256",
        "decision_fingerprint",
        "successor_professional_review_sha256",
        "artifact_receipts",
        "content_sha256",
    }
    if set(raw_receipt) != required:
        raise AssuranceRunError("review transition receipt has invalid fields")
    raw_artifact_receipts = raw_receipt["artifact_receipts"]
    if not isinstance(raw_artifact_receipts, list):
        raise AssuranceRunError("review transition artifact receipts must be a list")
    artifact_receipts = validate_receipt_set(
        {"transition": transition_dir},
        raw_artifact_receipts,
    )
    expected_metadata = [
        (filename, artifact_id, role)
        for filename, artifact_id, role in REVIEW_TRANSITION_EVIDENCE_FILES
    ]
    actual_metadata = [
        (
            receipt.get("path"),
            receipt.get("artifact_id"),
            receipt.get("role"),
        )
        for receipt in artifact_receipts
    ]
    if actual_metadata != expected_metadata or any(
        receipt.get("root_id") != "transition" for receipt in artifact_receipts
    ):
        raise AssuranceRunError(
            "review transition artifact receipts are missing, expanded, or reordered"
        )
    predecessor = _validated_assurance_envelope(
        _read_json_mapping(transition_dir / "predecessor_assurance_receipts.json")
    )
    if predecessor["content_sha256"] != digest:
        raise AssuranceRunError(
            "review transition path does not match predecessor seal content"
        )
    predecessor_snapshot = transition_dir / PREDECESSOR_RUN_SNAPSHOT_DIRECTORY
    snapshot_fingerprint = _tree_fingerprint(predecessor_snapshot)
    replayed_predecessor = validate_assurance_run(
        predecessor_snapshot,
        _externally_anchored_run_sha256=externally_anchored_predecessor_sha256,
    )
    if replayed_predecessor != predecessor:
        raise AssuranceRunError(
            "retained predecessor seal differs from its complete replay"
        )
    archived_snapshot_pairs = (
        (
            "predecessor_assurance_receipts.json",
            predecessor_snapshot / "assurance_receipts.json",
        ),
        (
            "predecessor_professional_review.json",
            predecessor_snapshot / "professional_review.json",
        ),
        (
            "predecessor_reconciliation_results.json",
            predecessor_snapshot
            / FINAL_OUTPUT_DIRECTORY
            / "reconciliation_results.json",
        ),
        (
            "predecessor_review_payload.json",
            predecessor_snapshot / "review_payload.json",
        ),
    )
    for archived_name, snapshot_path in archived_snapshot_pairs:
        if _stable_regular_bytes(
            transition_dir / archived_name,
            label=archived_name,
        ) != _stable_regular_bytes(
            snapshot_path,
            label=f"snapshot {snapshot_path.name}",
        ):
            raise AssuranceRunError(
                "selected predecessor evidence differs from its complete replay"
            )
    if _tree_fingerprint(predecessor_snapshot) != snapshot_fingerprint:
        raise AssuranceRunError("retained predecessor changed during complete replay")
    if predecessor["implementation_receipts"] != build_implementation_receipts():
        raise AssuranceRunError(
            "retained predecessor implementation receipts are stale"
        )
    predecessor_professional = validate_professional_review_authority(
        _read_json_mapping(transition_dir / "predecessor_professional_review.json")
    )
    if predecessor_professional != predecessor["professional_review_authority"]:
        raise AssuranceRunError("retained predecessor professional review is stale")
    _receipt_matches_archived_bytes(
        predecessor["professional_review_receipt"],
        transition_dir / "predecessor_professional_review.json",
        expected_artifact_id="professional.review",
        expected_path="professional_review.json",
        expected_role="review_authority",
    )
    predecessor_reconciliation = _read_json_mapping(
        transition_dir / "predecessor_reconciliation_results.json"
    )
    _rederive_predecessor_reconciliation(
        predecessor=predecessor,
        predecessor_snapshot=predecessor_snapshot,
        predecessor_reconciliation=predecessor_reconciliation,
    )
    final_inventory = predecessor.get("final_output_inventory")
    final_receipts = (
        final_inventory.get("artifact_receipts")
        if isinstance(final_inventory, Mapping)
        else None
    )
    reconciliation_receipts = [
        receipt
        for receipt in final_receipts or []
        if isinstance(receipt, Mapping)
        and receipt.get("path") == "reconciliation_results.json"
    ]
    if len(reconciliation_receipts) != 1:
        raise AssuranceRunError(
            "retained predecessor seal lacks one reconciliation result receipt"
        )
    _receipt_matches_archived_bytes(
        reconciliation_receipts[0],
        transition_dir / "predecessor_reconciliation_results.json",
        expected_artifact_id=str(reconciliation_receipts[0].get("artifact_id")),
        expected_path="reconciliation_results.json",
        expected_role="output",
    )
    predecessor_review_payload = _read_json_mapping(
        transition_dir / "predecessor_review_payload.json"
    )
    mapping = _validate_review_payload_mapping(
        _read_json_mapping(transition_dir / "review_payload_mapping.json")
    )
    if mapping != build_review_payload_mapping(predecessor_review_payload):
        raise AssuranceRunError(
            "retained review-payload item-to-record mapping is stale"
        )
    applied = _read_json_mapping(transition_dir / "applied_decisions.json")
    effects, reviewer = _validate_archived_applied_decisions(
        applied,
        mapping=mapping,
    )
    fingerprint = review_decision_fingerprint(mapping, effects)
    applied_professional = applied["professional_review"]
    if applied_professional.get("decision_fingerprint") != fingerprint:
        raise AssuranceRunError("retained applied decision fingerprint is stale")
    successor_professional = validate_professional_review_authority(
        _read_json_mapping(transition_dir / "successor_professional_review.json")
    )
    expected_successor = build_applied_review_authority(
        predecessor_assurance=predecessor,
        predecessor_professional_review=predecessor_professional,
        predecessor_reconciliation=predecessor_reconciliation,
        review_payload_mapping=mapping,
        effects=effects,
        reviewer_ref=reviewer,
    )
    if successor_professional != expected_successor:
        raise AssuranceRunError(
            "retained successor professional review cannot be rederived"
        )
    content = _transition_receipt_content(
        predecessor_assurance_sha256=digest,
        decision_fingerprint=fingerprint,
        successor_professional_review_sha256=str(
            successor_professional["content_sha256"]
        ),
        artifact_receipts=artifact_receipts,
    )
    if (
        raw_receipt.get("content_sha256") != canonical_json_sha256(content)
        or {key: raw_receipt[key] for key in content} != content
    ):
        raise AssuranceRunError("review transition receipt is stale")
    normalized_receipt = {
        **content,
        "content_sha256": raw_receipt["content_sha256"],
    }
    return normalized_receipt, predecessor, successor_professional


def validate_review_transition_history(
    output_dir: Path,
    expected_receipts: Sequence[Mapping[str, Any]] | None = None,
    *,
    current_professional_review: Mapping[str, Any] | None = None,
    expected_predecessor_checkpoint: str | None = None,
    _externally_anchored_run_sha256: str | None = None,
) -> list[dict[str, Any]]:
    """Replay exact physical closure and every retained review transition."""

    out_dir = Path(output_dir)
    if current_professional_review is None:
        professional_path = out_dir / "professional_review.json"
        if professional_path.exists() or professional_path.is_symlink():
            current_professional_review = _read_json_mapping(professional_path)
    current = (
        validate_professional_review_authority(current_professional_review)
        if current_professional_review is not None
        else None
    )
    immediate_predecessor = (
        current.get("predecessor_assurance_sha256")
        if isinstance(current, Mapping) and current.get("origin") == "applied_decisions"
        else None
    )
    if immediate_predecessor is not None:
        if _externally_anchored_run_sha256 is None:
            _require_expected_predecessor_checkpoint(
                expected_predecessor_checkpoint,
                immediate_predecessor,
            )
        elif _CHECKPOINT_PATTERN.fullmatch(_externally_anchored_run_sha256) is None:
            raise AssuranceRunError("externally anchored assurance digest is invalid")

    history_root = out_dir / REVIEW_TRANSITION_HISTORY_DIRECTORY
    if not history_root.exists() and not history_root.is_symlink():
        receipts: list[dict[str, Any]] = []
        bundles: list[tuple[dict[str, Any], dict[str, Any], dict[str, Any]]] = []
    else:
        _validate_real_directory(history_root, label="review transition history")
        transition_dirs = sorted(
            history_root.iterdir(),
            key=lambda path: path.name,
        )
        if not transition_dirs:
            raise AssuranceRunError("review transition history cannot be empty")
        transition_by_digest = {path.name: path for path in transition_dirs}
        bundles = []
        anchored_prior_digests: set[str] = set()
        if isinstance(immediate_predecessor, str):
            immediate_path = transition_by_digest.get(immediate_predecessor)
            if immediate_path is None:
                raise AssuranceRunError(
                    "applied professional review lacks its exact retained transition"
                )
            immediate_bundle = _validate_review_transition_directory(
                immediate_path,
                externally_anchored_predecessor_sha256=immediate_predecessor,
            )
            bundles.append(immediate_bundle)
            predecessor_history = immediate_bundle[1].get("review_transition_receipts")
            if isinstance(predecessor_history, list):
                anchored_prior_digests = {
                    str(item.get("predecessor_assurance_sha256"))
                    for item in predecessor_history
                    if isinstance(item, Mapping)
                }
        for path in transition_dirs:
            if path.name == immediate_predecessor:
                continue
            bundles.append(
                _validate_review_transition_directory(
                    path,
                    externally_anchored_predecessor_sha256=(
                        path.name if path.name in anchored_prior_digests else None
                    ),
                )
            )
        bundles.sort(key=lambda bundle: str(bundle[0]["predecessor_assurance_sha256"]))
        receipts = [bundle[0] for bundle in bundles]
    if expected_receipts is not None and list(expected_receipts) != receipts:
        raise AssuranceRunError(
            "review transition receipts are missing, expanded, changed, or reordered"
        )
    receipt_by_digest = {
        str(receipt["predecessor_assurance_sha256"]): receipt for receipt in receipts
    }
    for _, predecessor, _ in bundles:
        predecessor_history = predecessor.get("review_transition_receipts")
        if (
            not isinstance(predecessor_history, list)
            or any(
                receipt_by_digest.get(str(item.get("predecessor_assurance_sha256")))
                != item
                for item in predecessor_history
                if isinstance(item, Mapping)
            )
            or any(not isinstance(item, Mapping) for item in predecessor_history)
        ):
            raise AssuranceRunError(
                "retained predecessor transition history is unavailable"
            )
    if current is None:
        if receipts:
            raise AssuranceRunError(
                "review transition history requires professional review authority"
            )
        return receipts
    if current["origin"] == "direct_persisted_review":
        if receipts:
            raise AssuranceRunError(
                "direct professional review cannot claim transition history"
            )
        return receipts
    predecessor_digest = current.get("predecessor_assurance_sha256")
    matches = [
        successor
        for receipt, _, successor in bundles
        if receipt["predecessor_assurance_sha256"] == predecessor_digest
    ]
    if len(matches) != 1 or matches[0] != current:
        raise AssuranceRunError(
            "applied professional review lacks its exact retained transition"
        )
    return receipts


def retain_review_transition(
    output_dir: Path,
    capture_dir: Path,
    *,
    expected_predecessor_checkpoint: str | None = None,
) -> dict[str, Any]:
    """Retain and receipt the first-apply transition inside the COW candidate."""

    out_dir = Path(output_dir)
    capture = Path(capture_dir)
    _validate_real_directory(capture, label="transition capture directory")
    capture_names = {
        "predecessor_assurance_receipts.json",
        "predecessor_professional_review.json",
        "predecessor_reconciliation_results.json",
        "predecessor_review_payload.json",
        PREDECESSOR_RUN_SNAPSHOT_DIRECTORY,
    }
    if {path.name for path in capture.iterdir()} != capture_names:
        raise AssuranceRunError("transition capture does not match its exact contract")
    for path in capture.iterdir():
        if path.name == PREDECESSOR_RUN_SNAPSHOT_DIRECTORY:
            _validate_real_directory(path, label="captured predecessor run")
        else:
            _require_single_link_regular(path, label="transition capture artifact")
    predecessor = _validated_assurance_envelope(
        _read_json_mapping(capture / "predecessor_assurance_receipts.json")
    )
    checkpoint = _require_expected_predecessor_checkpoint(
        expected_predecessor_checkpoint,
        predecessor["content_sha256"],
    )
    captured_snapshot = capture / PREDECESSOR_RUN_SNAPSHOT_DIRECTORY
    captured_fingerprint = _tree_fingerprint(captured_snapshot)
    replayed_predecessor = validate_assurance_run(
        captured_snapshot,
        _externally_anchored_run_sha256=checkpoint,
    )
    if replayed_predecessor != predecessor:
        raise AssuranceRunError(
            "captured predecessor seal differs from its complete replay"
        )
    review_payload = _read_json_mapping(capture / "predecessor_review_payload.json")
    mapping = build_review_payload_mapping(review_payload)
    applied_path = out_dir / "applied_decisions.json"
    successor_path = out_dir / "professional_review.json"
    applied = _read_json_mapping(applied_path)
    effects, reviewer = _validate_archived_applied_decisions(
        applied,
        mapping=mapping,
    )
    successor = validate_professional_review_authority(
        _read_json_mapping(successor_path)
    )
    expected_successor = build_applied_review_authority(
        predecessor_assurance=predecessor,
        predecessor_professional_review=_read_json_mapping(
            capture / "predecessor_professional_review.json"
        ),
        predecessor_reconciliation=_read_json_mapping(
            capture / "predecessor_reconciliation_results.json"
        ),
        review_payload_mapping=mapping,
        effects=effects,
        reviewer_ref=reviewer,
    )
    if successor != expected_successor:
        raise AssuranceRunError(
            "generated professional review does not close against predecessor evidence"
        )
    predecessor_digest = str(predecessor["content_sha256"])
    history_root = out_dir / REVIEW_TRANSITION_HISTORY_DIRECTORY
    if history_root.exists() or history_root.is_symlink():
        _validate_real_directory(history_root, label="review transition history")
    else:
        history_root.mkdir(mode=0o700)
    transition_dir = history_root / predecessor_digest
    try:
        transition_dir.mkdir(mode=0o700)
    except FileExistsError as exc:
        raise AssuranceRunError("review transition already exists") from exc
    for filename in capture_names - {PREDECESSOR_RUN_SNAPSHOT_DIRECTORY}:
        _write_exact_bytes(
            transition_dir / filename,
            _stable_regular_bytes(
                capture / filename,
                label=f"captured {filename}",
            ),
        )
    retained_snapshot = transition_dir / PREDECESSOR_RUN_SNAPSHOT_DIRECTORY
    shutil.copytree(captured_snapshot, retained_snapshot)
    if _tree_fingerprint(retained_snapshot) != captured_fingerprint:
        raise AssuranceRunError("retained predecessor snapshot changed during copy")
    _write_exact_json(
        transition_dir / "review_payload_mapping.json",
        mapping,
    )
    _write_exact_bytes(
        transition_dir / "applied_decisions.json",
        _stable_regular_bytes(applied_path, label="applied_decisions.json"),
    )
    _write_exact_bytes(
        transition_dir / "successor_professional_review.json",
        _stable_regular_bytes(successor_path, label="professional_review.json"),
    )
    artifact_receipts = [
        artifact_receipt(
            transition_dir,
            transition_dir / filename,
            artifact_id=artifact_id,
            role=role,
            root_id="transition",
            media_type="application/json",
        )
        for filename, artifact_id, role in REVIEW_TRANSITION_EVIDENCE_FILES
    ]
    fingerprint = review_decision_fingerprint(mapping, effects)
    content = _transition_receipt_content(
        predecessor_assurance_sha256=predecessor_digest,
        decision_fingerprint=fingerprint,
        successor_professional_review_sha256=str(successor["content_sha256"]),
        artifact_receipts=artifact_receipts,
    )
    receipt = {**content, "content_sha256": canonical_json_sha256(content)}
    _write_exact_json(
        transition_dir / REVIEW_TRANSITION_RECEIPT_FILE,
        receipt,
    )
    receipts = validate_review_transition_history(
        out_dir,
        current_professional_review=successor,
        expected_predecessor_checkpoint=checkpoint,
    )
    if receipt not in receipts:
        raise AssuranceRunError("retained review transition did not replay")
    return {
        "predecessor_assurance_sha256": predecessor_digest,
        "transition_receipt": receipt,
        "history_paths": [
            (
                Path(REVIEW_TRANSITION_HISTORY_DIRECTORY)
                / predecessor_digest
                / str(entry["path"])
            ).as_posix()
            for entry in _physical_tree_entries(transition_dir)
            if entry["kind"] == "file"
        ],
    }


def _locator(kind: str, **fields: object) -> str:
    return json.dumps({"kind": kind, **fields}, sort_keys=True, separators=(",", ":"))


def _parsed_locator(value: str) -> dict[str, Any]:
    try:
        locator = json.loads(value)
    except json.JSONDecodeError as exc:
        raise AssuranceRunError("assurance locator must be canonical JSON") from exc
    if (
        not isinstance(locator, dict)
        or not isinstance(locator.get("kind"), str)
        or value != json.dumps(locator, sort_keys=True, separators=(",", ":"))
    ):
        raise AssuranceRunError("assurance locator must be canonical JSON")
    return locator


def _header_index(
    headers: Sequence[object],
    candidates: Sequence[str],
) -> int | None:
    normalized = {
        str(value or "").strip().lower(): index for index, value in enumerate(headers)
    }
    return next((normalized[name] for name in candidates if name in normalized), None)


def _material_header(value: object) -> bool:
    return str(value or "").strip().lower() in MATERIAL_VALUE_HEADERS


def _material_text_values(value: object) -> set[str]:
    if not isinstance(value, str):
        return set()
    values = set()
    for token in MATERIAL_TEXT_RE.findall(value):
        normalized = token.replace(",", ".")
        canonical = _canonical_money(normalized)
        if canonical is not None:
            values.add(canonical)
    for match in CURRENCY_TEXT_RE.finditer(value):
        token = match.group(1) or match.group(2)
        normalized = token.replace(",", ".")
        canonical = _canonical_money(normalized)
        if canonical is not None:
            values.add(canonical)
    return values


def _material_cell_value(value: object, *, material_column: bool) -> bool:
    if value is None or isinstance(value, (bool, date, datetime)):
        return False
    if _material_text_values(value):
        return True
    if material_column and isinstance(value, (int, float, Decimal)):
        return True
    return material_column and _canonical_money(value) is not None


def _workbook_primary_record_addresses(
    workbook_path: Path,
) -> tuple[dict[str, str], set[tuple[str, str]]]:
    """Return record-bound amount cells from the workflow detail sheet."""

    workbook = load_workbook(workbook_path, read_only=True, data_only=True)
    try:
        sheet = next(
            (
                workbook[name]
                for name in (
                    "Reconciliation detail",
                    "Dettaglio riconciliazione",
                    "Detalle de conciliación",
                )
                if name in workbook.sheetnames
            ),
            None,
        )
        addresses: dict[str, str] = {}
        addressed_cells: set[tuple[str, str]] = set()
        if sheet is not None:
            header_cells = next(
                sheet.iter_rows(min_row=1, max_row=1),
                (),
            )
            headers = [cell.value for cell in header_cells]
            record_column = _header_index(
                headers,
                ("record_id", "id record", "id registro"),
            )
            value_column = _header_index(headers, ("amount", "importo", "importe"))
            if record_column is not None and value_column is not None:
                for row_number, cells in enumerate(
                    sheet.iter_rows(min_row=2),
                    start=2,
                ):
                    record_id = str(cells[record_column].value or "").strip()
                    if not record_id:
                        continue
                    if record_id in addresses:
                        raise AssuranceRunError(
                            f"{workbook_path.name} contains duplicate record identities"
                        )
                    value_cell = cells[value_column]
                    addresses[record_id] = _locator(
                        "xlsx_record_value",
                        sheet=sheet.title,
                        record_cell=(
                            f"{get_column_letter(record_column + 1)}{row_number}"
                        ),
                        value_cell=(
                            f"{get_column_letter(value_column + 1)}{row_number}"
                        ),
                        record_id=record_id,
                    )
                    addressed_cells.add((sheet.title, value_cell.coordinate))
        return addresses, addressed_cells
    finally:
        workbook.close()


def _workbook_record_addresses(workbook_path: Path) -> dict[str, str]:
    """Address value and record identity together in a declared workbook."""

    addresses, addressed_cells = _workbook_primary_record_addresses(workbook_path)
    workbook = load_workbook(workbook_path, read_only=True, data_only=True)
    try:
        for current_sheet in workbook.worksheets:
            material_columns: dict[int, int] = {}
            for cells in current_sheet.iter_rows():
                for cell in cells:
                    if _material_header(cell.value):
                        material_columns[cell.column] = min(
                            cell.row,
                            material_columns.get(cell.column, cell.row),
                        )
            for cells in current_sheet.iter_rows():
                for cell in cells:
                    if (current_sheet.title, cell.coordinate) in addressed_cells:
                        continue
                    header_row = material_columns.get(cell.column)
                    if _material_cell_value(
                        cell.value,
                        material_column=(
                            header_row is not None and cell.row > header_row
                        ),
                    ):
                        raise AssuranceRunError(
                            f"{workbook_path.name} contains a material figure "
                            f"without a record address at "
                            f"{current_sheet.title}!{cell.coordinate}"
                        )
        return addresses
    finally:
        workbook.close()


def _docx_record_addresses(
    document_path: Path,
) -> dict[str, str]:
    """Address record/value table rows or reject unaddressed material figures."""

    document = Document(document_path)
    addresses: dict[str, str] = {}
    addressed_cells: set[tuple[int, int, int]] = set()
    for table_index, table in enumerate(document.tables):
        if not table.rows:
            continue
        headers = [cell.text for cell in table.rows[0].cells]
        record_column = _header_index(
            headers,
            ("record_id", "id record", "id registro"),
        )
        value_column = _header_index(headers, ("amount", "importo", "importe"))
        if record_column is None or value_column is None:
            continue
        for row_index, row in enumerate(table.rows[1:], start=1):
            record_id = row.cells[record_column].text.strip()
            if not record_id:
                continue
            if record_id in addresses:
                raise AssuranceRunError(
                    f"{document_path.name} contains duplicate record identities"
                )
            addresses[record_id] = _locator(
                "docx_record_value",
                table=table_index,
                row=row_index,
                record_cell=record_column,
                value_cell=value_column,
                record_id=record_id,
            )
            addressed_cells.add((table_index, row_index, value_column))
    for paragraph in document.paragraphs:
        if _material_text_values(paragraph.text):
            raise AssuranceRunError(
                f"{document_path.name} contains material figures without record addresses"
            )
    for table_index, table in enumerate(document.tables):
        material_columns = {
            index
            for index, cell in enumerate(table.rows[0].cells if table.rows else [])
            if _material_header(cell.text)
        }
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if (table_index, row_index, cell_index) in addressed_cells:
                    continue
                if _material_cell_value(
                    cell.text,
                    material_column=(row_index > 0 and cell_index in material_columns),
                ):
                    raise AssuranceRunError(
                        f"{document_path.name} contains a material figure "
                        "without a record address"
                    )
    return addresses


def _rendered_money(value: object, *, material_column: bool) -> str | None:
    """Return one exact material value represented by an Office cell."""

    if not _material_cell_value(value, material_column=material_column):
        return None
    embedded = _material_text_values(value)
    if len(embedded) > 1:
        raise AssuranceRunError(
            "one rendered cell cannot contain multiple material figures"
        )
    if embedded:
        return next(iter(embedded))
    return _canonical_money(value)


def _semantic_token(value: object) -> str:
    return re.sub(r"[^a-z0-9]+", "_", str(value or "").casefold()).strip("_")


def _row_record_refs(
    values: Sequence[object],
    record_rows: Sequence[Mapping[str, Any]],
) -> list[str]:
    """Resolve a rendered row to sealed records using identity, not amount."""

    cell_tokens = [str(value or "").strip().casefold() for value in values]
    direct_refs: set[str] = set()
    for row in record_rows:
        record_id = str(row.get("record_id") or "").strip()
        if not record_id:
            continue
        record_token = record_id.casefold()
        if any(
            value == record_token
            or re.search(
                rf"(?<![a-z0-9]){re.escape(record_token)}(?![a-z0-9])",
                value,
            )
            for value in cell_tokens
        ):
            direct_refs.add(record_id)
    if direct_refs:
        return sorted(direct_refs)
    document_refs: set[str] = set()
    for row in record_rows:
        record_id = str(row.get("record_id") or "").strip()
        if not record_id:
            continue
        document_tokens = {
            str(row.get(field) or "").strip().casefold()
            for field in ("document_key", "document_no")
            if str(row.get(field) or "").strip()
        }
        if document_tokens.intersection(cell_tokens):
            document_refs.add(record_id)
    return sorted(document_refs)


def _row_numeric_field_total(
    row: Mapping[str, Any],
    field: str,
    *,
    absolute: bool = False,
) -> Decimal:
    raw = row.get(field)
    values: Sequence[object]
    if isinstance(raw, Sequence) and not isinstance(raw, (str, bytes, bytearray)):
        values = raw
    else:
        values = [raw]
    total = Decimal("0")
    for item in values:
        parsed = _canonical_money(item)
        if parsed is None and isinstance(item, str):
            parsed_values = _material_text_values(item)
            if len(parsed_values) == 1:
                parsed = next(iter(parsed_values))
        if parsed is not None:
            number = Decimal(parsed)
            total += abs(number) if absolute else number
    return total


def _status_context(
    header: object,
    values: Sequence[object],
    reconciliation_rows: Sequence[Mapping[str, Any]],
) -> str | None:
    aliases = {
        "closed": {"chiusa_da_evidenza", "cerrada_por_evidencia"},
        "needs_evidence": {
            "serve_evidenza_aggiuntiva",
            "requiere_evidencia_adicional",
        },
        "open_supported": {
            "aperta_ma_supportata",
            "abierta_pero_respaldada",
        },
        "probable_payment": {
            "pagamento_probabile_da_verificare",
            "pago_probable_pendiente_de_verificacion",
        },
        "unresolved": {"non_risolta", "no_resuelta"},
    }
    statuses = sorted(
        {
            "closed",
            "needs_evidence",
            "open_supported",
            "probable_payment",
            "unresolved",
        }
        | {
            str(row.get("reconciliation_status") or "").strip()
            for row in reconciliation_rows
            if str(row.get("reconciliation_status") or "").strip()
        }
    )
    header_token = _semantic_token(header)
    for status in statuses:
        token = _semantic_token(status)
        if token and token in header_token:
            return status
    row_tokens = {_semantic_token(value) for value in values}
    matches = [
        status
        for status in statuses
        if _semantic_token(status) in row_tokens
        or bool(aliases.get(status, set()).intersection(row_tokens))
    ]
    return matches[0] if len(matches) == 1 else None


def _rendered_value_formula(
    *,
    value: str,
    header: object,
    values: Sequence[object],
    owner_refs: Sequence[str],
    reconciliation_rows: Sequence[Mapping[str, Any]],
    support_rows: Sequence[Mapping[str, Any]],
    external_detail_rows: Sequence[Mapping[str, Any]],
) -> tuple[list[str], dict[str, Any]]:
    """Bind a rendered value to a mechanically replayable record formula."""

    rows_by_id = {str(row.get("record_id") or ""): row for row in support_rows}
    rows_by_id.update(
        {str(row.get("record_id") or ""): row for row in reconciliation_rows}
    )
    owners = [rows_by_id[record_id] for record_id in owner_refs]
    number = Decimal(value)
    header_token = _semantic_token(header)
    amount_header = any(
        token in header_token
        for token in (
            "amount",
            "importo",
            "importe",
            "balance",
            "saldo",
            "debit",
            "credit",
            "net",
        )
    )
    absolute_header = (
        any(token in header_token for token in ("absolute", "assoluto", "absoluto"))
        or "amount_abs" in header_token
    )
    difference_header = any(
        token in header_token
        for token in ("difference", "differenza", "diferencia", "residual", "residuo")
    )
    share_header = any(token in header_token for token in ("share", "peso", "percent"))
    external_field = next(
        (
            field
            for field in (
                "settlement_effect_signed_net_debit_minus_credit",
                "cash_flow_signed",
                "amount",
            )
            if (
                field == "amount"
                and header_token in {"amount", "amount_total", "importo", "importe"}
            )
            or (field != "amount" and field in header_token)
        ),
        None,
    )
    external_categories = {
        str(row.get("external_category") or "") for row in external_detail_rows
    }
    value_tokens = {_semantic_token(item) for item in values}
    external_category = next(
        (
            category
            for category in sorted(external_categories)
            if _semantic_token(category) in value_tokens
        ),
        None,
    )
    if "total" in value_tokens:
        external_category = "TOTAL"
    if external_field is not None and external_category is not None:
        selected_external = [
            row
            for row in external_detail_rows
            if external_category == "TOTAL"
            or str(row.get("external_category") or "") == external_category
        ]
        expected = sum(
            (
                _row_numeric_field_total(row, external_field)
                for row in selected_external
            ),
            Decimal("0"),
        )
        if number == expected:
            selected_refs = sorted(
                str(row["record_id"])
                for row in selected_external
                if str(row.get("record_id") or "")
            )
            return selected_refs, {
                "kind": "external_evidence_sum",
                "field": external_field,
                "category": external_category,
            }

    if len(owners) > 1:
        reconciliation_ids = {
            str(row.get("record_id") or "") for row in reconciliation_rows
        }
        evidence_tokens = ("evidence", "evidenza", "evidencia", "supporting", "bank")
        preferred = [
            record_id
            for record_id in owner_refs
            if (
                record_id not in reconciliation_ids
                if any(token in header_token for token in evidence_tokens)
                else record_id in reconciliation_ids
            )
            and number == _row_numeric_field_total(rows_by_id[record_id], "amount")
        ]
        if preferred:
            return [preferred[0]], {
                "kind": "record_field",
                "field": "amount",
                "absolute": False,
            }

    if len(owners) == 1:
        owner = owners[0]
        amount = _row_numeric_field_total(owner, "amount")
        if number == amount:
            return list(owner_refs), {
                "kind": "record_field",
                "field": "amount",
                "absolute": False,
            }
        if absolute_header and number == abs(amount):
            return list(owner_refs), {
                "kind": "record_field",
                "field": "amount",
                "absolute": True,
            }
        for field in ("reported_increment", "balance"):
            field_value = _row_numeric_field_total(owner, field)
            if field in header_token and number == field_value:
                return list(owner_refs), {
                    "kind": "record_field",
                    "field": field,
                    "absolute": False,
                }
        if (
            "evidence" in header_token
            or "evidenza" in header_token
            or "evidencia" in header_token
        ):
            evidence_total = _row_numeric_field_total(
                owner,
                "matched_evidence_amounts",
            )
            if number == evidence_total:
                return list(owner_refs), {
                    "kind": "record_field_sum",
                    "field": "matched_evidence_amounts",
                    "absolute": False,
                }
        if difference_header and number == Decimal("0"):
            return list(owner_refs), {
                "kind": "record_self_difference",
                "field": "amount",
            }

    if (
        owners
        and any(token in header_token for token in ("open", "aperto", "abierto"))
        and amount_header
    ):
        open_refs = sorted(
            str(row.get("record_id") or "")
            for row in owners
            if str(row.get("source_role") or "").casefold() == "open_items"
            or str(row.get("evidence_type") or "").casefold() == "open_item"
        )
        open_total = sum(
            (
                _row_numeric_field_total(rows_by_id[record_id], "amount")
                for record_id in open_refs
            ),
            Decimal("0"),
        )
        if number == open_total:
            return open_refs, {
                "kind": "source_role_sum",
                "field": "amount",
                "source_bucket": "open_items",
            }

    if owners and amount_header:
        owner_total = sum(
            (_row_numeric_field_total(row, "amount") for row in owners),
            Decimal("0"),
        )
        if number == (
            sum(
                (abs(_row_numeric_field_total(row, "amount")) for row in owners),
                Decimal("0"),
            )
            if absolute_header
            else owner_total
        ):
            return list(owner_refs), {
                "kind": "record_set_sum",
                "field": "amount",
                "absolute": absolute_header,
            }

    status = _status_context(header, values, reconciliation_rows)
    selected = [
        row
        for row in reconciliation_rows
        if status is None or str(row.get("reconciliation_status") or "") == status
    ]
    selected_refs = sorted(str(row["record_id"]) for row in selected)
    total = sum(
        (_row_numeric_field_total(row, "amount") for row in selected),
        Decimal("0"),
    )
    absolute_total = sum(
        (abs(_row_numeric_field_total(row, "amount")) for row in selected),
        Decimal("0"),
    )
    if share_header and status is not None:
        denominator = sum(
            (
                abs(_row_numeric_field_total(row, "amount"))
                for row in reconciliation_rows
            ),
            Decimal("0"),
        )
        expected = (
            Decimal("0")
            if denominator == 0
            else (absolute_total / denominator) * Decimal("100")
        )
        if number == expected:
            return selected_refs, {
                "kind": "status_absolute_share_percent",
                "field": "amount",
                "status": status,
            }
    if amount_header:
        expected = absolute_total if absolute_header else total
        if number == expected:
            return selected_refs, {
                "kind": "status_sum" if status is not None else "population_sum",
                "field": "amount",
                "absolute": absolute_header,
                "status": status,
            }
    raise AssuranceRunError(
        "rendered material figure has no current record/source formula: "
        f"header={header!r}, value={value}, owner_record_refs={list(owner_refs)!r}"
    )


def _rendered_value_address(
    *,
    artifact_ref: str,
    locator: str,
    value: str,
    owner_refs: Sequence[str],
    formula: Mapping[str, Any],
) -> dict[str, Any]:
    return {
        "address_id": _artifact_id(
            "rendered_value",
            f"{artifact_ref}/{locator}",
        ),
        "artifact_ref": artifact_ref,
        "locator": locator,
        "value": value,
        "owner_record_refs": list(owner_refs),
        "source_record_refs": list(owner_refs),
        "formula": dict(formula),
    }


def _native_rendered_value_addresses(
    *,
    final_receipts: Sequence[Mapping[str, Any]],
    final_root: Path,
    reconciliation_rows: Sequence[Mapping[str, Any]],
    prepared: Mapping[str, Any],
) -> list[dict[str, Any]]:
    """Address every material figure in the code-owned three-file package."""

    by_name = {Path(str(item["path"])).name: item for item in final_receipts}
    native_names = frozenset(by_name) - {"reconciliation_results.json"}
    if native_names not in TRUSTED_NATIVE_OUTPUT_PROFILES[:2]:
        return []
    support_rows = [
        *list(prepared.get("open_items") or []),
        *list(prepared.get("evidence_rows") or []),
    ]
    try:
        from .reconciliation_helpers import external_evidence_detail_rows
    except ImportError:  # pragma: no cover - direct script import
        scripts_dir = Path(__file__).resolve().parent
        if str(scripts_dir) not in sys.path:
            sys.path.insert(0, str(scripts_dir))
        from reconciliation_helpers import external_evidence_detail_rows  # type: ignore

    external_detail = external_evidence_detail_rows(
        list(prepared.get("evidence_rows") or []),
        dict(prepared.get("assumptions") or {}),
    )
    record_rows = [*support_rows, *reconciliation_rows]
    addresses: list[dict[str, Any]] = []
    for name in sorted(native_names):
        receipt = by_name[name]
        artifact_ref = str(receipt["artifact_id"])
        path = final_root / name
        if path.suffix.lower() in {".xlsx", ".xlsm"}:
            workbook = load_workbook(path, read_only=True, data_only=True)
            try:
                for sheet in workbook.worksheets:
                    rows = list(sheet.iter_rows())
                    if not rows:
                        continue
                    headers = [cell.value for cell in rows[0]]
                    material_columns = {
                        index
                        for index, header in enumerate(headers)
                        if _material_header(header)
                    }
                    for row in rows[1:]:
                        values = [cell.value for cell in row]
                        owner_refs = _row_record_refs(values, record_rows)
                        for index, cell in enumerate(row):
                            value = _rendered_money(
                                cell.value,
                                material_column=index in material_columns,
                            )
                            if value is None:
                                continue
                            header = headers[index] if index < len(headers) else ""
                            owners, formula = _rendered_value_formula(
                                value=value,
                                header=header,
                                values=values,
                                owner_refs=owner_refs,
                                reconciliation_rows=reconciliation_rows,
                                support_rows=support_rows,
                                external_detail_rows=external_detail,
                            )
                            addresses.append(
                                _rendered_value_address(
                                    artifact_ref=artifact_ref,
                                    locator=_locator(
                                        "xlsx_material_value",
                                        sheet=sheet.title,
                                        cell=cell.coordinate,
                                        header_cell=(
                                            f"{get_column_letter(index + 1)}1"
                                        ),
                                    ),
                                    value=value,
                                    owner_refs=owners,
                                    formula=formula,
                                )
                            )
            finally:
                workbook.close()
        elif path.suffix.lower() == ".docx":
            document = Document(path)
            for paragraph_index, paragraph in enumerate(document.paragraphs):
                if _material_text_values(paragraph.text):
                    raise AssuranceRunError(
                        f"{name} contains a material figure outside a "
                        "workflow-owned table address"
                    )
            for table_index, table in enumerate(document.tables):
                if not table.rows:
                    continue
                headers = [cell.text for cell in table.rows[0].cells]
                material_columns = {
                    index
                    for index, header in enumerate(headers)
                    if _material_header(header)
                }
                for row_index, row in enumerate(table.rows[1:], start=1):
                    values = [cell.text for cell in row.cells]
                    owner_refs = _row_record_refs(values, record_rows)
                    for cell_index, cell in enumerate(row.cells):
                        value = _rendered_money(
                            cell.text,
                            material_column=cell_index in material_columns,
                        )
                        if value is None:
                            continue
                        header = (
                            headers[cell_index] if cell_index < len(headers) else ""
                        )
                        owners, formula = _rendered_value_formula(
                            value=value,
                            header=header,
                            values=values,
                            owner_refs=owner_refs,
                            reconciliation_rows=reconciliation_rows,
                            support_rows=support_rows,
                            external_detail_rows=external_detail,
                        )
                        addresses.append(
                            _rendered_value_address(
                                artifact_ref=artifact_ref,
                                locator=_locator(
                                    "docx_material_value",
                                    table=table_index,
                                    row=row_index,
                                    cell=cell_index,
                                    header_cell=cell_index,
                                ),
                                value=value,
                                owner_refs=owners,
                                formula=formula,
                            )
                        )
    addresses.sort(key=lambda item: (item["artifact_ref"], item["locator"]))
    return addresses


def _material_json_key(value: object) -> bool:
    normalized = str(value).strip().lower()
    return any(part in normalized for part in MATERIAL_VALUE_KEY_PARTS)


def _assert_json_material_values_addressed(
    payload: object,
    *,
    path_name: str,
    allowed_paths: set[tuple[object, ...]],
) -> None:
    def walk(value: object, location: tuple[object, ...]) -> None:
        if isinstance(value, Mapping):
            for key, child in value.items():
                child_location = (*location, str(key))
                if (
                    _material_json_key(key)
                    and child_location not in allowed_paths
                    and (
                        _canonical_money(child) is not None or isinstance(child, float)
                    )
                ):
                    raise AssuranceRunError(
                        f"{path_name} contains a material figure without "
                        "a record address"
                    )
                walk(child, child_location)
        elif isinstance(value, list):
            for index, child in enumerate(value):
                walk(child, (*location, index))

    walk(payload, ())


def _reconciliation_json_material_paths(
    reconciliation_rows: Sequence[Mapping[str, Any]],
    allocation_ledgers: Sequence[Mapping[str, Any]],
) -> set[tuple[object, ...]]:
    paths = {
        ("reconciliation_rows", index, "amount")
        for index, _ in enumerate(reconciliation_rows)
    }
    for ledger_index, ledger in enumerate(allocation_ledgers):
        for population in ("source_records", "target_records"):
            paths.update(
                {
                    ("allocation_ledgers", ledger_index, population, index, "amount")
                    for index, _ in enumerate(ledger[population])
                }
            )
        paths.update(
            {
                ("allocation_ledgers", ledger_index, "allocations", index, "amount")
                for index, _ in enumerate(ledger["allocations"])
            }
        )
        for population in ("source_residuals", "target_residuals"):
            paths.update(
                {
                    (
                        "allocation_ledgers",
                        ledger_index,
                        population,
                        index,
                        "residual",
                    )
                    for index, _ in enumerate(ledger[population])
                }
            )
    return paths


def _json_record_addresses(
    path: Path,
) -> dict[str, str]:
    """Address a native reconciliation JSON or reject unowned material values."""

    payload = _read_json_mapping(path)
    rows = payload.get("reconciliation_rows")
    if isinstance(rows, list):
        addresses: dict[str, str] = {}
        allowed_paths: set[tuple[object, ...]] = set()
        for index, row in enumerate(rows):
            if not isinstance(row, Mapping):
                raise AssuranceRunError(f"{path.name} has an invalid record row")
            record_id = str(row.get("record_id") or "")
            if not record_id or record_id in addresses:
                raise AssuranceRunError(
                    f"{path.name} has missing or duplicate record identities"
                )
            addresses[record_id] = _output_record_locator(index, record_id)
            allowed_paths.add(("reconciliation_rows", index, "amount"))
        _assert_json_material_values_addressed(
            payload,
            path_name=path.name,
            allowed_paths=allowed_paths,
        )
        return addresses
    _assert_json_material_values_addressed(
        payload,
        path_name=path.name,
        allowed_paths=set(),
    )
    return {}


def _record_source_artifact_ref(
    row: Mapping[str, Any],
    source_receipts: Sequence[Mapping[str, Any]],
) -> str | None:
    receipt = _source_receipt_for_row(row, source_receipts)
    return str(receipt["artifact_id"]) if receipt is not None else None


def _canonical_money(value: object) -> str | None:
    if isinstance(value, bool) or isinstance(value, float) or value is None:
        return None
    if isinstance(value, Decimal):
        return decimal_text(value) if value.is_finite() else None
    if isinstance(value, int):
        return decimal_text(Decimal(value))
    if not isinstance(value, str):
        return None
    text = value.strip()
    if re.fullmatch(r"-?\d+(?:\.\d+)?", text) is None:
        return None
    return decimal_text(Decimal(text))


def _source_money(value: object, convention: Mapping[str, Any]) -> str | None:
    """Parse one reviewed source value under the cent-only convention."""

    if isinstance(value, (bool, float)) or value is None:
        return None
    text = str(value).strip()
    decimal_separator = convention["decimal_separator"]
    thousands_separator = convention["thousands_separator"]
    if decimal_separator is None:
        if re.fullmatch(r"-?\d+", text) is None:
            return None
        normalized = text
    else:
        if thousands_separator:
            text = text.replace(str(thousands_separator), "")
        normalized = text.replace(str(decimal_separator), ".")
    canonical = _canonical_money(normalized)
    if canonical is None or Decimal(canonical) % Decimal(SUPPORTED_REPORTED_INCREMENT):
        return None
    return canonical


def _source_document_tokens(row: Mapping[str, Any]) -> set[str]:
    """Return exact, mechanically comparable document identifiers."""

    values: set[str] = set()
    for field in ("document_no", "document_number", "invoice_no", "document_key"):
        raw = row.get(field)
        if not isinstance(raw, str):
            continue
        for token in re.split(r"[;,\n]", raw):
            cleaned = token.strip()
            if not cleaned:
                continue
            values.add(cleaned.casefold())
            if "|" in cleaned:
                prefix, suffix = cleaned.rsplit("|", 1)
                if re.fullmatch(r"\d{4}", suffix.strip()):
                    values.add(prefix.strip().casefold())
    return values


def _source_date(
    value: object,
    *,
    convention: Mapping[str, Any],
) -> str | None:
    """Parse a native/ISO or source-ordered date without semantic guessing."""

    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if not isinstance(value, str):
        return None
    text = value.strip()
    if not text:
        return None
    try:
        return date.fromisoformat(text).isoformat()
    except ValueError:
        pass
    match = re.fullmatch(r"(\d{1,2})[./-](\d{1,2})[./-](\d{2}|\d{4})", text)
    if match is None:
        return None
    first, second, year_text = (int(part) for part in match.groups())
    year = year_text + 2000 if year_text < 100 else year_text
    order = convention.get("order")
    if order == "day_first":
        day, month = first, second
    elif order == "month_first":
        month, day = first, second
    else:
        return None
    try:
        return date(year, month, day).isoformat()
    except ValueError:
        return None


def _validate_tabular_source_identity(
    *,
    row: Mapping[str, Any],
    headers: Sequence[object],
    values: Sequence[object],
    date_convention: Mapping[str, Any],
    strict: bool = True,
) -> tuple[int | None, str | None]:
    """Bind a prepared row to its document identity and critical source dates."""

    record_column = _header_index(
        headers,
        ("record_id", "id record", "id registro"),
    )
    document_column = _header_index(
        headers,
        (
            "document",
            "document_no",
            "document number",
            "invoice",
            "invoice_no",
            "fattura",
            "documento",
            "partita",
        ),
    )
    source_record_id = (
        str(values[record_column] or "").strip()
        if record_column is not None and record_column < len(values)
        else ""
    )
    record_id = str(row.get("record_id") or "")
    if source_record_id:
        if strict and source_record_id != record_id:
            raise AssuranceRunError("source row record identity is stale")
        identity_column = record_column
        identity_value = source_record_id
    else:
        document_tokens = _source_document_tokens(row)
        if not document_tokens:
            raise AssuranceRunError(
                "source row lacks a document or record identity binding"
            )
        if document_column is None or document_column >= len(values):
            raise AssuranceRunError("source document identity column is unavailable")
        source_document = str(values[document_column] or "").strip()
        if strict and source_document.casefold() not in document_tokens:
            raise AssuranceRunError("source row document identity is stale")
        identity_column = document_column
        identity_value = source_document

    date_headers = {
        "document_date": (
            "document_date",
            "document date",
            "invoice_date",
            "data documento",
            "fecha documento",
            "date",
        ),
        "posting_date": (
            "posting_date",
            "posting date",
            "data registrazione",
            "fecha contabilización",
        ),
        "value_date": (
            "value_date",
            "value date",
            "data valuta",
            "fecha valor",
        ),
    }
    for field, candidates in date_headers.items():
        prepared_date = row.get(field)
        if prepared_date in {None, ""}:
            continue
        date_column = _header_index(headers, candidates)
        if date_column is None or date_column >= len(values):
            raise AssuranceRunError(f"source-bound {field} column is unavailable")
        located_date = _source_date(
            values[date_column],
            convention=date_convention,
        )
        if strict and located_date != prepared_date:
            raise AssuranceRunError(f"prepared {field} is stale against source")
    return identity_column, identity_value


def _source_value_locator(
    *,
    row: Mapping[str, Any],
    source_root: Path,
    receipt: Mapping[str, Any],
    decision: Mapping[str, Any],
    strict_identity: bool = True,
) -> tuple[str, str]:
    """Dereference one bounded current-source amount or abstain."""

    path = source_root / str(receipt["path"])
    amount = _canonical_money(row.get("amount"))
    if amount is None:
        raise AssuranceRunError("source-backed material value is not canonical")
    money = decision["content"]["money"]
    date_convention = decision["content"]["date"]
    source_row = row.get("source_row")
    if (
        isinstance(source_row, bool)
        or not isinstance(source_row, int)
        or source_row < 1
    ):
        raise AssuranceRunError("source-backed row lacks a bounded source_row")
    suffix = path.suffix.lower()
    if suffix == ".csv":
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            sample = handle.read(4096)
            handle.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            except csv.Error as exc:
                raise AssuranceRunError("CSV source delimiter is not bounded") from exc
            rows = list(csv.reader(handle, dialect))
        if len(rows) < source_row or not rows:
            raise AssuranceRunError("CSV source_row locator is stale")
        amount_column = _header_index(
            rows[0],
            ("amount", "importo", "importe", "balance", "saldo"),
        )
        if amount_column is None or amount_column >= len(rows[source_row - 1]):
            raise AssuranceRunError("CSV amount column locator is unavailable")
        located_row = rows[source_row - 1]
        located = _source_money(located_row[amount_column], money)
        identity_column, identity_value = _validate_tabular_source_identity(
            row=row,
            headers=rows[0],
            values=located_row,
            date_convention=date_convention,
            strict=strict_identity,
        )
        locator = _locator(
            "csv_record_value",
            row=source_row,
            value_column=amount_column,
            identity_column=identity_column,
            identity_value=identity_value,
            record_id=str(row["record_id"]),
        )
    elif suffix in {".xlsx", ".xlsm"}:
        workbook = load_workbook(path, read_only=True, data_only=True)
        try:
            sheet_name = row.get("source_sheet")
            if not isinstance(sheet_name, str) or sheet_name not in workbook.sheetnames:
                raise AssuranceRunError("workbook source_sheet locator is stale")
            sheet = workbook[sheet_name]
            headers = [
                cell.value for cell in next(sheet.iter_rows(min_row=1, max_row=1))
            ]
            amount_column = _header_index(
                headers,
                ("amount", "importo", "importe", "balance", "saldo"),
            )
            if amount_column is None:
                raise AssuranceRunError("workbook amount column locator is unavailable")
            located_values = [
                sheet.cell(row=source_row, column=index + 1).value
                for index in range(len(headers))
            ]
            located = _source_money(
                located_values[amount_column],
                money,
            )
            identity_column, identity_value = _validate_tabular_source_identity(
                row=row,
                headers=headers,
                values=located_values,
                date_convention=date_convention,
                strict=strict_identity,
            )
            locator = _locator(
                "xlsx_source_record_value",
                sheet=sheet_name,
                row=source_row,
                value_column=amount_column + 1,
                identity_column=(
                    identity_column + 1 if identity_column is not None else None
                ),
                identity_value=identity_value,
                record_id=str(row["record_id"]),
            )
        finally:
            workbook.close()
    elif suffix == ".pdf":
        try:
            import fitz
        except ImportError as exc:
            raise AssuranceRunError(
                "PDF source locators cannot be dereferenced"
            ) from exc
        source_page = row.get("source_page")
        if (
            isinstance(source_page, bool)
            or not isinstance(source_page, int)
            or source_page < 1
        ):
            raise AssuranceRunError("PDF source_page locator is stale")
        document = fitz.open(path)
        try:
            if source_page > document.page_count:
                raise AssuranceRunError("PDF source_page locator is stale")
            lines = [
                line.strip()
                for line in document[source_page - 1].get_text("text").splitlines()
                if line.strip()
            ]
        finally:
            document.close()
        source_value_row = row.get("source_value_row", source_row)
        if (
            isinstance(source_value_row, bool)
            or not isinstance(source_value_row, int)
            or source_value_row < 1
            or source_value_row > len(lines)
        ):
            raise AssuranceRunError("PDF source_row locator is stale")
        source_line = lines[source_value_row - 1]
        candidates = [
            _source_money(token, money)
            for token in re.findall(
                r"[-+]?\d[\d.,]*",
                source_line,
            )
        ]
        located = amount if amount in candidates else None
        document_tokens = _source_document_tokens(row)
        if (
            strict_identity
            and document_tokens
            and not any(token in source_line.casefold() for token in document_tokens)
        ):
            raise AssuranceRunError("PDF source row document identity is stale")
        for field in ("document_date", "posting_date", "value_date"):
            prepared_date = row.get(field)
            if prepared_date in {None, ""}:
                continue
            source_dates = {
                parsed
                for token in re.findall(
                    r"\d{1,4}[./-]\d{1,2}[./-]\d{1,4}",
                    source_line,
                )
                if (
                    parsed := _source_date(
                        token,
                        convention=date_convention,
                    )
                )
            }
            if strict_identity and prepared_date not in source_dates:
                raise AssuranceRunError(f"prepared {field} is stale against PDF source")
        locator = _locator(
            "pdf_text_record_value",
            page=source_page,
            line=source_value_row,
            document_tokens=sorted(document_tokens),
            record_id=str(row["record_id"]),
        )
    else:
        raise AssuranceRunError(
            f"{path.name} has no bounded source-value locator; source must abstain"
        )
    if located != amount:
        raise AssuranceRunError("prepared amount does not equal current source value")
    return locator, amount


def _prepared_record_locator(
    population: str,
    index: int,
    record_id: str,
) -> str:
    return _locator(
        "json_record_value",
        record_pointer=f"/{population}/{index}/record_id",
        value_pointer=f"/{population}/{index}/amount",
        record_id=record_id,
    )


def _output_record_locator(index: int, record_id: str) -> str:
    return _locator(
        "json_record_value",
        record_pointer=f"/reconciliation_rows/{index}/record_id",
        value_pointer=f"/reconciliation_rows/{index}/amount",
        record_id=record_id,
    )


def _numeric_entries(
    *,
    reconciliation_rows: Sequence[Mapping[str, Any]],
    allocation_ledgers: Sequence[Mapping[str, Any]],
    source_receipts: Sequence[Mapping[str, Any]],
    reviewed_source_decisions: Sequence[Mapping[str, Any]],
    prepared_receipt: Mapping[str, Any],
    final_receipts: Sequence[Mapping[str, Any]],
    workbook_name: str,
    final_root: Path,
    prepared: Mapping[str, Any],
    source_root: Path,
    strict_source_identity: bool = False,
) -> list[dict[str, Any]]:
    final_by_name = {
        Path(str(receipt["path"])).name: receipt for receipt in final_receipts
    }
    reconciliation_receipt = final_by_name.get("reconciliation_results.json")
    if reconciliation_receipt is None:
        return []
    if not source_receipts:
        return []
    material_rows = [
        row
        for row in reconciliation_rows
        if _canonical_money(row.get("amount")) is not None
    ]
    material_ids = {str(row.get("record_id") or "") for row in material_rows}
    workbook_addresses: dict[str, dict[str, str]] = {}
    docx_addresses: dict[str, dict[str, str]] = {}
    json_addresses: dict[str, dict[str, str]] = {}
    native_names = frozenset(final_by_name) - {"reconciliation_results.json"}
    production_profile = native_names in TRUSTED_NATIVE_OUTPUT_PROFILES[:2]
    for name, receipt in final_by_name.items():
        suffix = Path(name).suffix.lower()
        if suffix in {".xlsx", ".xlsm"}:
            if production_profile:
                addresses = (
                    _workbook_primary_record_addresses(final_root / name)[0]
                    if name == "riconciliazione_audit.xlsx"
                    else {}
                )
            else:
                addresses = _workbook_record_addresses(final_root / name)
            if (
                material_rows
                and (not production_profile or addresses)
                and set(addresses) != material_ids
            ):
                raise AssuranceRunError(
                    f"{name} does not address the exact material record-id set"
                )
            workbook_addresses[str(receipt["artifact_id"])] = addresses
        elif suffix == ".docx":
            addresses = (
                {} if production_profile else _docx_record_addresses(final_root / name)
            )
            if addresses and set(addresses) != material_ids:
                raise AssuranceRunError(
                    f"{name} does not address the exact material record-id set"
                )
            docx_addresses[str(receipt["artifact_id"])] = addresses
        elif suffix == ".json" and name != "reconciliation_results.json":
            addresses = _json_record_addresses(final_root / name)
            if addresses and set(addresses) != material_ids:
                raise AssuranceRunError(
                    f"{name} does not address the exact material record-id set"
                )
            json_addresses[str(receipt["artifact_id"])] = addresses
    entries: list[dict[str, Any]] = []
    decision_by_source_ref = {
        str(source_ref): str(decision["decision_id"])
        for decision in reviewed_source_decisions
        for source_ref in decision["source_artifact_refs"]
    }
    prepared_by_record: dict[str, tuple[str, int, Mapping[str, Any]]] = {}
    for population in ("open_items", "evidence_rows"):
        for prepared_index, prepared_row in enumerate(prepared[population]):
            record_id = str(prepared_row.get("record_id") or "")
            if record_id in prepared_by_record:
                raise AssuranceRunError("prepared record identities must be unique")
            prepared_by_record[record_id] = (population, prepared_index, prepared_row)
    decision_by_source_ref_full = {
        str(source_ref): decision
        for decision in reviewed_source_decisions
        for source_ref in decision["source_artifact_refs"]
    }
    receipt_by_id = {
        str(receipt["artifact_id"]): receipt for receipt in source_receipts
    }
    for index, row in enumerate(reconciliation_rows):
        value = _canonical_money(row.get("amount"))
        source_ref = _record_source_artifact_ref(row, source_receipts)
        if value is None or source_ref is None:
            continue
        record_id = str(row.get("record_id") or "")
        prepared_match = prepared_by_record.get(record_id)
        if prepared_match is None:
            raise AssuranceRunError(
                f"reconciliation record {record_id} is absent from prepared records"
            )
        population, prepared_index, prepared_row = prepared_match
        if _canonical_money(prepared_row.get("amount")) != value:
            raise AssuranceRunError(
                "reconciliation amount is stale against prepared row"
            )
        source_locator, source_value = _source_value_locator(
            row=prepared_row,
            source_root=source_root,
            receipt=receipt_by_id[source_ref],
            decision=decision_by_source_ref_full[source_ref],
            strict_identity=strict_source_identity,
        )
        outputs = [
            {
                "artifact_ref": str(reconciliation_receipt["artifact_id"]),
                "locator": _output_record_locator(index, record_id),
                "value": value,
            }
        ]
        for artifact_ref, addresses in {
            **workbook_addresses,
            **docx_addresses,
            **json_addresses,
        }.items():
            if record_id in addresses:
                outputs.append(
                    {
                        "artifact_ref": artifact_ref,
                        "locator": addresses[record_id],
                        "value": value,
                    }
                )
        entries.append(
            {
                "evidence_id": _artifact_id(
                    "material_value",
                    f"reconciliation/{index}/{record_id}",
                ),
                "value": value,
                "unit": str(row.get("unit") or "currency_amount"),
                "currency": str(row.get("currency") or "EUR"),
                "source": {
                    "artifact_ref": source_ref,
                    "locator": source_locator,
                    "value": source_value,
                },
                "prepared": {
                    "artifact_ref": str(prepared_receipt["artifact_id"]),
                    "locator": _prepared_record_locator(
                        population,
                        prepared_index,
                        record_id,
                    ),
                    "value": value,
                },
                "outputs": outputs,
                "calculation_ref": None,
                "decision_ref": decision_by_source_ref.get(source_ref),
                "limitations": [],
            }
        )
    return entries


def _allocation_value_addresses(
    allocation_ledgers: Sequence[Mapping[str, Any]],
    *,
    reconciliation_artifact_ref: str,
) -> list[dict[str, Any]]:
    """Address every allocation and residual in the sealed result artifact."""

    addresses = []
    for ledger_index, raw_ledger in enumerate(allocation_ledgers):
        ledger = validate_allocation_ledger(raw_ledger)
        source_index_by_ref = {
            str(record["record_id"]): index
            for index, record in enumerate(ledger["source_records"])
        }
        target_index_by_ref = {
            str(record["record_id"]): index
            for index, record in enumerate(ledger["target_records"])
        }
        for allocation_index, allocation in enumerate(ledger["allocations"]):
            source_ref = str(allocation["source_record_ref"])
            target_ref = str(allocation["target_record_ref"])
            addresses.append(
                {
                    "value_id": str(allocation["allocation_id"]),
                    "kind": "allocation",
                    "value": str(allocation["amount"]),
                    "currency": str(allocation["currency"]),
                    "unit": str(allocation["unit"]),
                    "artifact_ref": reconciliation_artifact_ref,
                    "source_record_ref": source_ref,
                    "target_record_ref": target_ref,
                    "source_population_locator": (
                        f"/allocation_ledgers/{ledger_index}/source_records/"
                        f"{source_index_by_ref[source_ref]}/amount"
                    ),
                    "target_population_locator": (
                        f"/allocation_ledgers/{ledger_index}/target_records/"
                        f"{target_index_by_ref[target_ref]}/amount"
                    ),
                    "reconciliation_locator": (
                        f"/allocation_ledgers/{ledger_index}/allocations/"
                        f"{allocation_index}/amount"
                    ),
                }
            )
        for residual_kind in ("source_residuals", "target_residuals"):
            for residual_index, residual in enumerate(ledger[residual_kind]):
                addresses.append(
                    {
                        "value_id": _artifact_id(
                            "residual",
                            (
                                f"{ledger['ledger_id']}/{residual_kind}/"
                                f"{residual['record_ref']}"
                            ),
                        ),
                        "kind": residual_kind.removesuffix("s"),
                        "value": str(residual["residual"]),
                        "currency": None,
                        "unit": None,
                        "record_ref": str(residual["record_ref"]),
                        "artifact_ref": reconciliation_artifact_ref,
                        "reconciliation_locator": (
                            f"/allocation_ledgers/{ledger_index}/{residual_kind}/"
                            f"{residual_index}/residual"
                        ),
                        "limitation": (
                            "Residual is preserved exactly and is not forced to zero."
                        ),
                    }
                )
    return addresses


def _validated_source_qualifications(
    *,
    source_receipts: Sequence[Mapping[str, Any]],
    reviewed_source_decisions: Sequence[Mapping[str, Any]],
    source_qualifications: Sequence[Mapping[str, Any]],
) -> list[dict[str, Any]]:
    """Validate one qualification per source and replay its mapping decision."""

    source_ids = {str(receipt["artifact_id"]) for receipt in source_receipts}
    normalized_decisions = [
        validate_reviewed_decision_receipt(
            decision,
            require_reviewed=True,
        )
        for decision in reviewed_source_decisions
    ]
    decision_ids = [str(decision["decision_id"]) for decision in normalized_decisions]
    if len(decision_ids) != len(set(decision_ids)):
        raise AssuranceRunError("reviewed source decision identities must be unique")
    decisions = {
        str(decision["decision_id"]): decision for decision in normalized_decisions
    }
    normalized = [
        validate_source_qualification(qualification)
        for qualification in source_qualifications
    ]
    qualification_ids = [
        str(qualification["qualification_id"]) for qualification in normalized
    ]
    if len(qualification_ids) != len(set(qualification_ids)):
        raise AssuranceRunError("source qualification identities must be unique")
    qualified_refs: list[str] = []
    for qualification in normalized:
        refs = [str(ref) for ref in qualification["source_artifact_refs"]]
        if len(refs) != 1:
            raise AssuranceRunError(
                "each source qualification must bind exactly one source artifact"
            )
        qualified_refs.extend(refs)
        if not set(refs) <= source_ids:
            raise AssuranceRunError(
                "source qualification references a stale source artifact"
            )
        mapping_ref = qualification["reviewed_mapping_ref"]
        if mapping_ref is None:
            raise AssuranceRunError(
                "every source qualification requires a reviewed mapping reference"
            )
        decision = decisions.get(str(mapping_ref))
        if decision is None:
            raise AssuranceRunError(
                "source qualification references an unknown reviewed decision"
            )
        expected_version = SUPPORTED_SOURCE_ADAPTER_VERSIONS.get(
            str(qualification["adapter_id"])
        )
        if expected_version is None:
            raise AssuranceRunError("source qualification uses an unsupported adapter")
        validate_reviewed_decision_receipt(
            decision,
            expected_source_artifact_refs=refs,
            expected_adapter_id=str(qualification["adapter_id"]),
            expected_adapter_version=expected_version,
            require_reviewed=True,
        )
        if qualification["adapter_version"] != expected_version:
            raise AssuranceRunError("source qualification adapter version is stale")
    if source_ids:
        if len(qualified_refs) != len(set(qualified_refs)):
            raise AssuranceRunError(
                "a source artifact is covered by multiple qualifications"
            )
        if set(qualified_refs) != source_ids:
            raise AssuranceRunError(
                "source qualifications do not cover the exact source set"
            )
    elif normalized:
        raise AssuranceRunError(
            "source qualifications require corresponding source receipts"
        )
    return normalized


def _validate_reconciliation_population(
    prepared_open_items: Sequence[Mapping[str, Any]],
    reconciliation_rows: Sequence[Mapping[str, Any]],
) -> None:
    """Require one reconciliation successor for every prepared open item.

    This comparison is intentionally exact for all prepared fields. Mechanical
    classification may add fields, but it may not drop or rewrite source
    identity, locators, dates, perimeter, or amounts.
    """

    prepared_by_id: dict[str, Mapping[str, Any]] = {}
    for index, row in enumerate(prepared_open_items):
        record_id = row.get("record_id")
        if not isinstance(record_id, str) or not record_id:
            raise AssuranceRunError(
                f"prepared open item {index} lacks a canonical record identity"
            )
        if record_id in prepared_by_id:
            raise AssuranceRunError("prepared open-item identities must be unique")
        prepared_by_id[record_id] = row
    reconciliation_by_id: dict[str, Mapping[str, Any]] = {}
    for index, row in enumerate(reconciliation_rows):
        record_id = row.get("record_id")
        if not isinstance(record_id, str) or not record_id:
            raise AssuranceRunError(
                f"reconciliation row {index} lacks a canonical record identity"
            )
        if record_id in reconciliation_by_id:
            raise AssuranceRunError("reconciliation record identities must be unique")
        reconciliation_by_id[record_id] = row
    if set(reconciliation_by_id) != set(prepared_by_id):
        raise AssuranceRunError(
            "a reconciliation record is absent from prepared records, or a "
            "prepared record is absent from reconciliation; exact sets/cardinality "
            "do not match"
        )
    for record_id, prepared_row in prepared_by_id.items():
        reconciliation_row = reconciliation_by_id[record_id]
        for field, prepared_value in prepared_row.items():
            if field not in reconciliation_row or _json_safe(
                reconciliation_row[field],
                label=f"reconciliation_rows.{record_id}.{field}",
            ) != _json_safe(
                prepared_value,
                label=f"prepared_records.{record_id}.{field}",
            ):
                raise AssuranceRunError(
                    f"reconciliation record {record_id} rewrites prepared field {field}"
                )


def _closed_status(value: object) -> bool:
    status_text = str(value or "").strip().lower()
    return status_text == "closed" or status_text.startswith("closed_")


def _validate_closed_relationships(
    *,
    reconciliation_rows: Sequence[Mapping[str, Any]],
    evidence_rows: Sequence[Mapping[str, Any]],
    allocation_ledgers: Sequence[Mapping[str, Any]],
) -> None:
    """Require current evidence plus one balanced relationship for every closure."""

    evidence_by_id = {
        str(row.get("record_id")): row
        for row in evidence_rows
        if isinstance(row.get("record_id"), str) and row.get("record_id")
    }
    if len(evidence_by_id) != len(
        [
            row
            for row in evidence_rows
            if isinstance(row.get("record_id"), str) and row.get("record_id")
        ]
    ):
        raise AssuranceRunError("prepared evidence record identities must be unique")
    ledgers = [validate_allocation_ledger(ledger) for ledger in allocation_ledgers]
    for row in reconciliation_rows:
        if not _closed_status(row.get("reconciliation_status")):
            continue
        record_id = str(row.get("record_id") or "")
        evidence_id = str(
            row.get("matched_evidence_id") or row.get("matched_evidence_ref") or ""
        )
        evidence = evidence_by_id.get(evidence_id)
        if evidence is None:
            raise AssuranceRunError(
                f"closed reconciliation record {record_id} lacks current evidence"
            )
        for field in ("currency", "unit", "entity_ref", "party_ref"):
            left = row.get(field)
            right = evidence.get(field)
            if (left is not None or right is not None) and left != right:
                raise AssuranceRunError(
                    f"closed reconciliation record {record_id} has an "
                    f"incompatible evidence {field}"
                )
        open_amount = _canonical_money(row.get("amount"))
        evidence_amount = _canonical_money(evidence.get("amount"))
        if open_amount is None or evidence_amount is None:
            raise AssuranceRunError(
                f"closed reconciliation record {record_id} lacks exact amounts"
            )
        relationship_ledger_id = str(row.get("relationship_allocation_ledger_id") or "")
        candidates = [
            ledger
            for ledger in ledgers
            if ledger["balanced"]
            and (
                not relationship_ledger_id
                or ledger["ledger_id"] == relationship_ledger_id
            )
        ]
        relationship_found = False
        for ledger in candidates:
            source_refs = {
                str(record["record_id"]) for record in ledger["source_records"]
            }
            target_refs = {
                str(record["record_id"]) for record in ledger["target_records"]
            }
            allocation_refs = {
                (
                    str(allocation["source_record_ref"]),
                    str(allocation["target_record_ref"]),
                )
                for allocation in ledger["allocations"]
            }
            if not allocation_refs:
                continue
            evidence_digest_ref = (
                "source." + hashlib.sha256(evidence_id.encode("utf-8")).hexdigest()[:24]
            )
            record_digest_ref = (
                "target." + hashlib.sha256(record_id.encode("utf-8")).hexdigest()[:24]
            )
            if (
                evidence_digest_ref in source_refs
                and record_digest_ref in target_refs
                and (evidence_digest_ref, record_digest_ref) in allocation_refs
            ):
                relationship_found = True
                break
        if not relationship_found:
            raise AssuranceRunError(
                f"closed reconciliation record {record_id} lacks a balanced "
                "evidence allocation relationship"
            )


def _reconciliation_gate_failed(
    checks: Sequence[Mapping[str, Any]],
    allocation_ledgers: Sequence[Mapping[str, Any]],
    reconciliation_rows: Sequence[Mapping[str, Any]],
) -> bool:
    if reconciliation_rows and not checks:
        return True
    if any(
        not isinstance(row.get("check"), str)
        or not str(row["check"]).strip()
        or not isinstance(row.get("status"), str)
        or not str(row["status"]).strip()
        for row in checks
    ):
        return True
    return any(
        str(row["status"]).upper() != "PASS"
        and not str(row["check"]).startswith("codex_review")
        for row in checks
    ) or any(
        not validate_allocation_ledger(ledger)["balanced"]
        for ledger in allocation_ledgers
    )


def _sealed_run_id(
    output_dir: Path,
    professional_review: Mapping[str, Any],
) -> str | None:
    """Bind the persisted review run identity into the assurance digest."""

    review_payload_path = output_dir / "review_payload.json"
    run_id: object = professional_review.get("run_id")
    if review_payload_path.exists() or review_payload_path.is_symlink():
        _require_single_link_regular(
            review_payload_path,
            label="review_payload.json",
        )
        run_id = _read_json_mapping(review_payload_path).get("run_id")
    if run_id is None:
        return None
    if not isinstance(run_id, str) or not run_id or run_id != run_id.strip():
        raise AssuranceRunError("review run identity is invalid")
    professional_run_id = professional_review.get("run_id")
    if professional_run_id is not None and professional_run_id != run_id:
        raise AssuranceRunError(
            "professional review and review payload run identities differ"
        )
    return run_id


def _finalize_assurance_run_in_place(
    *,
    output_dir: Path,
    context: Mapping[str, Any],
    reconciliation_rows: Sequence[Mapping[str, Any]],
    allocation_ledgers: Sequence[Mapping[str, Any]],
    checks: Sequence[Mapping[str, Any]],
    review_rows: Sequence[Mapping[str, Any]],
    source_qualifications: Sequence[Mapping[str, Any]],
    declared_outputs: Sequence[Path],
    workbook_name: str,
) -> dict[str, Any]:
    """Replay the run, publish an exact final set, and write independent gates."""

    out_dir = Path(output_dir).resolve()
    expected_predecessor_checkpoint = context.get("expected_predecessor_checkpoint")
    output_contract = _workflow_output_contract(declared_outputs)
    _validate_run_tree_allowlist(
        output_dir=out_dir,
        entries=_physical_tree_entries(out_dir),
        output_contract=output_contract,
        expected_predecessor_checkpoint=(
            str(expected_predecessor_checkpoint)
            if expected_predecessor_checkpoint is not None
            else None
        ),
    )
    plugin_root = Path(__file__).resolve().parents[1]
    roots: dict[str, Path] = {
        **_implementation_roots(plugin_root),
        "run": out_dir,
    }
    source_root = context.get("source_root")
    if source_root:
        roots["source"] = Path(str(source_root)).resolve()
    source_receipts = list(context.get("source_receipts") or [])
    run_date = _canonical_iso_date(
        context.get("run_date"),
        label="context.run_date",
    )
    source_decisions = _validated_source_decisions(
        source_receipts=source_receipts,
        reviewed_source_decisions=list(context.get("reviewed_source_decisions") or []),
        run_date=run_date,
    )
    implementation_receipts = list(context.get("implementation_receipts") or [])
    prepared_receipt = dict(context["prepared_receipt"])
    validate_receipt_set(
        roots,
        [*source_receipts, *implementation_receipts, prepared_receipt],
    )
    if source_receipts and source_root:
        _validate_source_boundary(Path(str(source_root)), source_receipts)
    normalized_qualifications = _validated_source_qualifications(
        source_receipts=source_receipts,
        reviewed_source_decisions=source_decisions,
        source_qualifications=source_qualifications,
    )
    prepared_qualifications = list(context.get("source_qualifications") or [])
    if normalized_qualifications != prepared_qualifications:
        raise AssuranceRunError(
            "source qualifications changed after the prepared boundary"
        )
    if any(
        qualification["status"] != "qualified"
        for qualification in normalized_qualifications
    ):
        raise AssuranceRunError(
            "unqualified sources cannot produce reconciliation or final artifacts"
        )
    prepared = _read_json_mapping(out_dir / "prepared_records.json")
    _validate_prepared_population(
        open_items=prepared["open_items"],
        evidence_rows=prepared["evidence_rows"],
        source_receipts=source_receipts,
        reviewed_source_decisions=source_decisions,
        source_qualifications=normalized_qualifications,
    )
    _validate_reconciliation_population(
        prepared["open_items"],
        reconciliation_rows,
    )
    normalized_allocations = [
        validate_allocation_ledger(ledger) for ledger in allocation_ledgers
    ]
    _validate_closed_relationships(
        reconciliation_rows=reconciliation_rows,
        evidence_rows=prepared["evidence_rows"],
        allocation_ledgers=normalized_allocations,
    )
    context_review_authority = context.get("professional_review_authority")
    if context_review_authority is None:
        professional_review = build_professional_review_authority(
            review_rows,
        )
    else:
        professional_review = validate_professional_review_authority(
            context_review_authority
        )
    if professional_review["records"] != _review_projection(review_rows):
        raise AssuranceRunError(
            "rendered review rows do not match persisted professional review authority"
        )
    professional_review_path = _write_exact_json(
        out_dir / "professional_review.json",
        professional_review,
    )
    professional_review_receipt = artifact_receipt(
        out_dir,
        professional_review_path,
        artifact_id="professional.review",
        role="review_authority",
        root_id="run",
        media_type="application/json",
    )
    review_transition_receipts = validate_review_transition_history(
        out_dir,
        current_professional_review=professional_review,
        expected_predecessor_checkpoint=(
            str(expected_predecessor_checkpoint)
            if expected_predecessor_checkpoint is not None
            else None
        ),
    )
    if review_transition_receipts != list(
        context.get("review_transition_receipts") or []
    ):
        raise AssuranceRunError(
            "review transition history changed after the prepared boundary"
        )
    if source_receipts:
        unaddressed_rows = [
            index
            for index, row in enumerate(reconciliation_rows)
            if _canonical_money(row.get("amount")) is not None
            and _record_source_artifact_ref(row, source_receipts) is None
        ]
        if unaddressed_rows:
            raise AssuranceRunError(
                "material reconciliation rows lack a current source artifact "
                f"reference: {unaddressed_rows[:10]}"
            )
    reconciliation_payload = {
        "schema_version": "audit_reconciliation.reconciliation_results.v1",
        "reconciliation_rows": list(reconciliation_rows),
        "allocation_ledgers": normalized_allocations,
        "checks": list(checks),
        "review_rows": list(review_rows),
        "source_qualifications": normalized_qualifications,
    }
    if source_receipts:
        _assert_json_material_values_addressed(
            reconciliation_payload,
            path_name="reconciliation_results.json",
            allowed_paths=_reconciliation_json_material_paths(
                reconciliation_rows,
                normalized_allocations,
            ),
        )
    reconciliation_path = _write_exact_json(
        out_dir / "reconciliation_results.json",
        reconciliation_payload,
    )
    published = _publish_final_outputs(
        out_dir,
        [*declared_outputs, reconciliation_path],
        output_contract,
    )
    final_receipts = published["artifact_receipts"]
    final_by_name = {
        Path(str(receipt["path"])).name: receipt for receipt in final_receipts
    }
    reconciliation_receipt = final_by_name["reconciliation_results.json"]
    reconciliation_ref = str(reconciliation_receipt["artifact_id"])
    entries = _numeric_entries(
        reconciliation_rows=reconciliation_rows,
        allocation_ledgers=normalized_allocations,
        source_receipts=source_receipts,
        reviewed_source_decisions=source_decisions,
        prepared_receipt=prepared_receipt,
        final_receipts=final_receipts,
        workbook_name=workbook_name,
        final_root=out_dir / FINAL_OUTPUT_DIRECTORY,
        prepared=prepared,
        source_root=Path(str(source_root)).resolve() if source_root else out_dir,
        strict_source_identity=True,
    )
    rendered_value_addresses = _native_rendered_value_addresses(
        final_receipts=final_receipts,
        final_root=out_dir / FINAL_OUTPUT_DIRECTORY,
        reconciliation_rows=reconciliation_rows,
        prepared=prepared,
    )
    material_row_count = sum(
        _canonical_money(row.get("amount")) is not None for row in reconciliation_rows
    )
    if source_receipts and len(entries) != material_row_count:
        raise AssuranceRunError(
            "every material reconciliation amount requires a source-prepared-output address"
        )
    if entries:
        numeric_ledger = build_numeric_evidence_ledger(
            entries,
            ledger_id="audit_reconciliation_material_values",
        )
        validate_numeric_evidence_ledger(numeric_ledger)
    else:
        numeric_ledger = None

    source_status = (
        "not_applicable"
        if not source_receipts
        else (
            "passed"
            if normalized_qualifications
            and all(
                row.get("status") == "qualified" for row in normalized_qualifications
            )
            else "failed"
        )
    )
    reconciliation_failed = _reconciliation_gate_failed(
        checks,
        normalized_allocations,
        reconciliation_rows,
    )
    reconciliation_status = "failed" if reconciliation_failed else "passed"
    semantic_status = _review_gate_status(
        professional_review["records"],
        reconciliation_rows,
        run_date=run_date,
    )
    reporting_status = (
        "passed"
        if source_status in {"passed", "not_applicable"}
        and reconciliation_status == "passed"
        and semantic_status in {"passed", "not_applicable"}
        else "blocked"
    )
    source_refs = [str(item["artifact_id"]) for item in source_receipts]
    implementation_refs = [str(item["artifact_id"]) for item in implementation_receipts]
    final_refs = [str(item["artifact_id"]) for item in final_receipts]
    numeric_refs = (
        [str(numeric_ledger["ledger_id"])] if numeric_ledger is not None else []
    )
    gates = build_gate_register(
        {
            "source": {
                "status": source_status,
                "evidence_refs": (
                    [*source_refs, reconciliation_ref]
                    if source_status == "passed"
                    else source_refs
                ),
                "limitations": (
                    []
                    if source_status in {"passed", "not_applicable"}
                    else ["At least one source is unqualified or stale."]
                ),
            },
            "preparation": {
                "status": "passed",
                "evidence_refs": [
                    str(prepared_receipt["artifact_id"]),
                    *implementation_refs,
                ],
                "limitations": [],
            },
            "reconciliation": {
                "status": reconciliation_status,
                "evidence_refs": [
                    reconciliation_ref,
                    *numeric_refs,
                    *(str(ledger["ledger_id"]) for ledger in normalized_allocations),
                ],
                "limitations": (
                    []
                    if reconciliation_status == "passed"
                    else ["At least one mechanical reconciliation check failed."]
                ),
            },
            "semantic_review": {
                "status": semantic_status,
                "evidence_refs": (
                    [str(professional_review_receipt["artifact_id"])]
                    if semantic_status == "passed"
                    else []
                ),
                "limitations": (
                    []
                    if semantic_status in {"passed", "not_applicable"}
                    else ["Required professional review is pending or failed."]
                ),
            },
            "reporting": {
                "status": reporting_status,
                "evidence_refs": final_refs if reporting_status == "passed" else [],
                "limitations": (
                    []
                    if reporting_status == "passed"
                    else ["Reporting is blocked by an upstream gate."]
                ),
            },
            "publication": {
                "status": "withheld",
                "evidence_refs": [],
                "limitations": [
                    "Publication is a separate action and was not performed."
                ],
            },
        }
    )
    final_artifacts_path = out_dir / "final_artifacts.json"
    if final_artifacts_path.is_file():
        final_artifacts = json.loads(final_artifacts_path.read_text(encoding="utf-8"))
        if not isinstance(final_artifacts, dict):
            raise AssuranceRunError("final_artifacts.json must contain an object")
        failed_gate_names = [
            name
            for name, gate in gates["gates"].items()
            if gate["status"] in {"failed", "blocked"}
        ]
        final_artifacts["assurance"] = {
            "receipts_path": "assurance_receipts.json",
            "gates_path": "assurance_gates.json",
            "final_output_inventory_path": FINAL_OUTPUT_INVENTORY,
            "final_output_boundary": FINAL_OUTPUT_DIRECTORY,
            "report_ready": gates["report_ready"],
            "publication_status": gates["gates"]["publication"]["status"],
        }
        if failed_gate_names:
            final_artifacts["status"] = "blocked"
            final_artifacts["blockers"] = [
                {
                    "kind": "assurance_gate_failed",
                    "detail": f"Assurance gate did not pass: {name}.",
                }
                for name in failed_gate_names
            ]
        elif not gates["report_ready"]:
            final_artifacts["status"] = "written_pending_review"
        elif professional_review["origin"] == "applied_decisions":
            final_artifacts["status"] = "final_ready"
            final_artifacts["blockers"] = []
            review_application = final_artifacts.get("review_application")
            if not isinstance(review_application, dict):
                review_application = {}
            review_application.update(
                {
                    "application_status": "final_ready",
                    "successor_assurance_replayed": True,
                    "decision_fingerprint": professional_review["decision_fingerprint"],
                    "professional_review_path": "professional_review.json",
                }
            )
            final_artifacts["review_application"] = review_application
        elif final_artifacts.get("status") == "final_ready":
            final_artifacts["status"] = "ready_for_professional_review"
        _write_exact_json(final_artifacts_path, final_artifacts)
    _write_exact_json(out_dir / "assurance_gates.json", gates)
    numeric_path = out_dir / "numeric_evidence_ledger.json"
    if numeric_ledger is not None:
        _write_exact_json(numeric_path, numeric_ledger)
    elif numeric_path.exists() or numeric_path.is_symlink():
        _require_single_link_regular(
            numeric_path,
            label="numeric_evidence_ledger.json",
        )
        numeric_path.unlink()
    _write_exact_json(
        out_dir / "assurance_receipts.json",
        {"status": "pending_run_tree_seal"},
    )
    run_tree_contract = _build_run_tree_contract(
        out_dir,
        output_contract,
        expected_predecessor_checkpoint=(
            str(expected_predecessor_checkpoint)
            if expected_predecessor_checkpoint is not None
            else None
        ),
    )
    control_content = {
        "schema_version": ASSURANCE_SCHEMA_VERSION,
        "run_id": _sealed_run_id(out_dir, professional_review),
        "run_date": run_date,
        "source_root": str(source_root) if source_root else None,
        "source_receipts": source_receipts,
        "reviewed_source_decisions": source_decisions,
        "source_qualifications": normalized_qualifications,
        "implementation_receipts": implementation_receipts,
        "prepared_receipt": prepared_receipt,
        "professional_review_receipt": professional_review_receipt,
        "professional_review_authority": professional_review,
        "review_transition_receipts": review_transition_receipts,
        "workflow_output_contract": output_contract,
        "run_tree_contract": run_tree_contract,
        "final_output_inventory": published,
        "numeric_evidence_ledger": numeric_ledger,
        "rendered_value_addresses": rendered_value_addresses,
        "allocation_value_addresses": _allocation_value_addresses(
            normalized_allocations,
            reconciliation_artifact_ref=reconciliation_ref,
        ),
        "allocation_ledgers": normalized_allocations,
        "gate_register": gates,
        "limitations": [
            (
                "Source meaning and accounting perimeter are reviewer decisions; "
                "the assurance layer verifies only their current-source binding."
            ),
            (
                "Materiality, evidence sufficiency, and accounting conclusions "
                "remain professional judgments and are not inferred here."
            ),
            (
                "reviewer_ref is an unsigned, unauthenticated label and must be "
                "treated as untrusted metadata, not identity proof."
            ),
            ("Publication remains withheld until a separate authorized action."),
        ],
    }
    control_payload = {
        **control_content,
        "content_sha256": canonical_json_sha256(control_content),
    }
    _write_exact_json(out_dir / "assurance_receipts.json", control_payload)

    return validate_assurance_run(
        out_dir,
        expected_predecessor_checkpoint=(
            str(expected_predecessor_checkpoint)
            if expected_predecessor_checkpoint is not None
            else None
        ),
    )


def _replace_assurance_tree(
    *,
    staged_dir: Path,
    output_dir: Path,
    expected_predecessor_checkpoint: str | None = None,
) -> dict[str, Any]:
    """Commit a replayed assurance tree and restore the prior tree on failure."""

    target_names = (
        FINAL_OUTPUT_DIRECTORY,
        FINAL_OUTPUT_INVENTORY,
        "reconciliation_results.json",
        "assurance_receipts.json",
        "assurance_gates.json",
        "numeric_evidence_ledger.json",
        "professional_review.json",
        "final_artifacts.json",
    )
    backup_dir = Path(
        tempfile.mkdtemp(prefix=".audit-assurance-backup-", dir=output_dir.parent)
    )
    moved_backups: list[str] = []
    installed: list[str] = []
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        for name in target_names:
            target = output_dir / name
            staged = staged_dir / name
            if target.exists() or target.is_symlink():
                os.replace(target, backup_dir / name)
                moved_backups.append(name)
            if staged.exists() or staged.is_symlink():
                os.replace(staged, target)
                installed.append(name)
        return validate_assurance_run(
            output_dir,
            expected_predecessor_checkpoint=expected_predecessor_checkpoint,
        )
    except (OSError, ValueError):
        for name in reversed(installed):
            target = output_dir / name
            if target.is_dir() and not target.is_symlink():
                shutil.rmtree(target)
            elif target.exists() or target.is_symlink():
                target.unlink()
        for name in reversed(moved_backups):
            backup = backup_dir / name
            if backup.exists() or backup.is_symlink():
                os.replace(backup, output_dir / name)
        raise
    finally:
        shutil.rmtree(backup_dir, ignore_errors=True)


def finalize_assurance_run(
    *,
    output_dir: Path,
    context: Mapping[str, Any],
    reconciliation_rows: Sequence[Mapping[str, Any]],
    allocation_ledgers: Sequence[Mapping[str, Any]],
    checks: Sequence[Mapping[str, Any]],
    review_rows: Sequence[Mapping[str, Any]],
    source_qualifications: Sequence[Mapping[str, Any]],
    declared_outputs: Sequence[Path],
    workbook_name: str,
) -> dict[str, Any]:
    """Build and replay the exact assurance tree under whole-run rollback."""

    out_dir = Path(output_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        result = _finalize_assurance_run_in_place(
            output_dir=out_dir,
            context=context,
            reconciliation_rows=reconciliation_rows,
            allocation_ledgers=allocation_ledgers,
            checks=checks,
            review_rows=review_rows,
            source_qualifications=source_qualifications,
            declared_outputs=declared_outputs,
            workbook_name=workbook_name,
        )
    except (OSError, ValueError):
        _rollback_run_transaction(
            out_dir,
            context.get("_run_transaction") if isinstance(context, Mapping) else None,
        )
        raise
    else:
        _commit_run_transaction(
            context.get("_run_transaction") if isinstance(context, Mapping) else None,
        )
        return result


def _read_json_mapping(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise AssuranceRunError(f"{path.name} is not readable JSON") from exc
    if not isinstance(value, dict):
        raise AssuranceRunError(f"{path.name} must contain an object")
    return value


def _json_pointer_value(document: object, locator: str) -> object:
    if not locator.startswith("/"):
        raise AssuranceRunError("JSON locator must be an absolute pointer")
    current = document
    for raw_token in locator[1:].split("/"):
        token = raw_token.replace("~1", "/").replace("~0", "~")
        if isinstance(current, list):
            if not token.isdigit() or int(token) >= len(current):
                raise AssuranceRunError("JSON locator list index is stale")
            current = current[int(token)]
        elif isinstance(current, Mapping):
            if token not in current:
                raise AssuranceRunError("JSON locator object key is stale")
            current = current[token]
        else:
            raise AssuranceRunError("JSON locator traverses a scalar")
    return current


def _workbook_locator_value(path: Path, locator: str) -> object:
    parsed = _parsed_locator(locator)
    if (
        set(parsed)
        != {
            "kind",
            "sheet",
            "record_cell",
            "value_cell",
            "record_id",
        }
        or parsed["kind"] != "xlsx_record_value"
    ):
        raise AssuranceRunError("workbook locator has invalid fields")
    sheet_name = parsed["sheet"]
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        if sheet_name not in workbook.sheetnames:
            raise AssuranceRunError("workbook locator sheet is stale")
        sheet = workbook[sheet_name]
        if (
            str(sheet[str(parsed["record_cell"])].value or "").strip()
            != parsed["record_id"]
        ):
            raise AssuranceRunError("workbook locator record identity is stale")
        return sheet[str(parsed["value_cell"])].value
    finally:
        workbook.close()


def _docx_locator_value(path: Path, locator: str) -> object:
    parsed = _parsed_locator(locator)
    if (
        set(parsed)
        != {
            "kind",
            "table",
            "row",
            "record_cell",
            "value_cell",
            "record_id",
        }
        or parsed["kind"] != "docx_record_value"
    ):
        raise AssuranceRunError("Word locator has invalid fields")
    document = Document(path)
    try:
        table = document.tables[int(parsed["table"])]
        row = table.rows[int(parsed["row"])]
        if row.cells[int(parsed["record_cell"])].text.strip() != parsed["record_id"]:
            raise AssuranceRunError("Word locator record identity is stale")
        return row.cells[int(parsed["value_cell"])].text.strip()
    except (IndexError, TypeError, ValueError) as exc:
        raise AssuranceRunError("Word locator is stale") from exc


def _json_record_locator_value(document: object, locator: str) -> object:
    parsed = _parsed_locator(locator)
    if (
        set(parsed)
        != {
            "kind",
            "record_pointer",
            "value_pointer",
            "record_id",
        }
        or parsed["kind"] != "json_record_value"
    ):
        raise AssuranceRunError("JSON record locator has invalid fields")
    record_id = _json_pointer_value(document, str(parsed["record_pointer"]))
    if record_id != parsed["record_id"]:
        raise AssuranceRunError("JSON record locator identity is stale")
    return _json_pointer_value(document, str(parsed["value_pointer"]))


def validate_assurance_run(
    output_dir: Path,
    *,
    expected_predecessor_checkpoint: str | None = None,
    _externally_anchored_run_sha256: str | None = None,
) -> dict[str, Any]:
    """Freshly replay the latest sealed run before readiness can advance."""

    requested_dir = Path(output_dir)
    _validate_real_directory(requested_dir, label="assurance run root")
    out_dir = requested_dir.resolve()
    for relative_path in RUN_CONTROL_PATHS:
        _require_single_link_regular(
            out_dir / relative_path,
            label=relative_path,
        )
    for optional_path in (
        "final_artifacts.json",
        "numeric_evidence_ledger.json",
        "professional_review.json",
    ):
        path = out_dir / optional_path
        if path.exists() or path.is_symlink():
            _require_single_link_regular(path, label=optional_path)
    control_path = out_dir / "assurance_receipts.json"
    payload = _validated_assurance_envelope(_read_json_mapping(control_path))
    if _externally_anchored_run_sha256 is not None:
        _require_expected_predecessor_checkpoint(
            _externally_anchored_run_sha256,
            payload["content_sha256"],
        )
    run_date = _canonical_iso_date(payload["run_date"], label="run_date")
    output_contract = _validate_workflow_output_contract(
        payload["workflow_output_contract"]
    )

    plugin_root = Path(__file__).resolve().parents[1]
    roots: dict[str, Path] = {
        **_implementation_roots(plugin_root),
        "run": out_dir,
    }
    source_root = payload["source_root"]
    if source_root is not None:
        if not isinstance(source_root, str) or not source_root:
            raise AssuranceRunError("source_root must be absolute text or null")
        roots["source"] = Path(source_root).resolve()
    source_receipts = validate_receipt_set(
        roots,
        payload["source_receipts"],
    )
    if source_receipts:
        if source_root is None:
            raise AssuranceRunError("source receipts require source_root")
        _validate_source_boundary(Path(source_root), source_receipts)
    elif source_root is not None:
        raise AssuranceRunError("source_root is not allowed without source receipts")

    implementation_receipts = validate_receipt_set(
        roots,
        payload["implementation_receipts"],
    )
    if implementation_receipts != build_implementation_receipts(plugin_root):
        raise AssuranceRunError(
            "implementation receipts do not cover the current workflow contract"
        )
    prepared_receipt = validate_receipt_set(
        roots,
        [payload["prepared_receipt"]],
    )[0]
    if (
        prepared_receipt["artifact_id"] != "prepared.records"
        or prepared_receipt["path"] != "prepared_records.json"
        or prepared_receipt["role"] != "prepared"
    ):
        raise AssuranceRunError("prepared record receipt is not the workflow boundary")
    professional_review_receipt = validate_receipt_set(
        roots,
        [payload["professional_review_receipt"]],
    )[0]
    if (
        professional_review_receipt["artifact_id"] != "professional.review"
        or professional_review_receipt["path"] != "professional_review.json"
        or professional_review_receipt["role"] != "review_authority"
    ):
        raise AssuranceRunError(
            "professional review receipt is not the authority boundary"
        )
    professional_review = validate_professional_review_authority(
        payload["professional_review_authority"]
    )
    if _read_json_mapping(out_dir / "professional_review.json") != professional_review:
        raise AssuranceRunError("professional review authority file is stale")
    if payload["run_id"] != _sealed_run_id(out_dir, professional_review):
        raise AssuranceRunError("sealed review run identity is stale")
    review_transition_receipts = validate_review_transition_history(
        out_dir,
        payload["review_transition_receipts"],
        current_professional_review=professional_review,
        expected_predecessor_checkpoint=expected_predecessor_checkpoint,
        _externally_anchored_run_sha256=_externally_anchored_run_sha256,
    )
    prepared = _read_json_mapping(out_dir / "prepared_records.json")
    if (
        set(prepared)
        != {"schema_version", "open_items", "evidence_rows", "assumptions"}
        or prepared["schema_version"] != "audit_reconciliation.prepared_records.v1"
        or not isinstance(prepared["open_items"], list)
        or not isinstance(prepared["evidence_rows"], list)
        or not isinstance(prepared["assumptions"], dict)
    ):
        raise AssuranceRunError("prepared record boundary has invalid fields")
    if prepared["assumptions"].get("assurance_run_date") != run_date:
        raise AssuranceRunError("prepared run date is stale")
    if _validated_tolerance(
        prepared["assumptions"].get("amount_tolerance"),
    ) != prepared["assumptions"].get("amount_tolerance"):
        raise AssuranceRunError("prepared amount tolerance is stale")

    reviewed_decisions = _validated_source_decisions(
        source_receipts=source_receipts,
        reviewed_source_decisions=payload["reviewed_source_decisions"],
        run_date=run_date,
    )

    final_inventory = validate_final_output_inventory(
        out_dir,
        payload["final_output_inventory"],
    )
    if final_inventory["declared_paths"] != output_contract["declared_paths"]:
        raise AssuranceRunError(
            "final output inventory does not match the sealed workflow contract"
        )
    if _read_json_mapping(out_dir / FINAL_OUTPUT_INVENTORY) != final_inventory:
        raise AssuranceRunError("final output inventory file is stale")
    _validate_run_tree_contract(
        out_dir,
        payload["run_tree_contract"],
        output_contract,
        expected_predecessor_checkpoint=expected_predecessor_checkpoint,
        externally_anchored_run_sha256=_externally_anchored_run_sha256,
    )

    result_path = out_dir / FINAL_OUTPUT_DIRECTORY / "reconciliation_results.json"
    result = _read_json_mapping(result_path)
    result_fields = {
        "schema_version",
        "reconciliation_rows",
        "allocation_ledgers",
        "checks",
        "review_rows",
        "source_qualifications",
    }
    if (
        set(result) != result_fields
        or result["schema_version"] != "audit_reconciliation.reconciliation_results.v1"
    ):
        raise AssuranceRunError("sealed reconciliation result has invalid fields")
    for field in (
        "reconciliation_rows",
        "allocation_ledgers",
        "checks",
        "review_rows",
        "source_qualifications",
    ):
        if not isinstance(result[field], list):
            raise AssuranceRunError(
                f"sealed reconciliation result {field} must be a list"
            )
    if _review_projection(result["review_rows"]) != professional_review["records"]:
        raise AssuranceRunError(
            "sealed review rows do not match persisted professional review authority"
        )
    qualifications = _validated_source_qualifications(
        source_receipts=source_receipts,
        reviewed_source_decisions=reviewed_decisions,
        source_qualifications=result["source_qualifications"],
    )
    if qualifications != payload["source_qualifications"]:
        raise AssuranceRunError("sealed source qualifications are stale")
    _validate_prepared_population(
        open_items=prepared["open_items"],
        evidence_rows=prepared["evidence_rows"],
        source_receipts=source_receipts,
        reviewed_source_decisions=reviewed_decisions,
        source_qualifications=qualifications,
    )
    _validate_reconciliation_population(
        prepared["open_items"],
        result["reconciliation_rows"],
    )
    allocations = [
        validate_allocation_ledger(ledger) for ledger in result["allocation_ledgers"]
    ]
    if allocations != payload["allocation_ledgers"]:
        raise AssuranceRunError("allocation ledgers are stale")
    _validate_closed_relationships(
        reconciliation_rows=result["reconciliation_rows"],
        evidence_rows=prepared["evidence_rows"],
        allocation_ledgers=allocations,
    )

    reconciliation_receipt = next(
        (
            receipt
            for receipt in final_inventory["artifact_receipts"]
            if receipt["path"] == "reconciliation_results.json"
        ),
        None,
    )
    if reconciliation_receipt is None:
        raise AssuranceRunError(
            "final output inventory omits reconciliation_results.json"
        )
    expected_addresses = _allocation_value_addresses(
        allocations,
        reconciliation_artifact_ref=str(reconciliation_receipt["artifact_id"]),
    )
    if payload["allocation_value_addresses"] != expected_addresses:
        raise AssuranceRunError("allocation value addresses are stale")
    expected_rendered_addresses = _native_rendered_value_addresses(
        final_receipts=final_inventory["artifact_receipts"],
        final_root=out_dir / FINAL_OUTPUT_DIRECTORY,
        reconciliation_rows=result["reconciliation_rows"],
        prepared=prepared,
    )
    if payload["rendered_value_addresses"] != expected_rendered_addresses:
        raise AssuranceRunError("rendered material value addresses are stale")

    numeric = payload["numeric_evidence_ledger"]
    numeric_path = out_dir / "numeric_evidence_ledger.json"
    if numeric is None:
        if numeric_path.exists():
            raise AssuranceRunError("unexpected numeric evidence ledger file")
        if source_receipts and any(
            _canonical_money(row.get("amount")) is not None
            for row in result["reconciliation_rows"]
        ):
            raise AssuranceRunError("material rows require a numeric evidence ledger")
    else:
        normalized_numeric = validate_numeric_evidence_ledger(numeric)
        if _read_json_mapping(numeric_path) != normalized_numeric:
            raise AssuranceRunError("numeric evidence ledger file is stale")
        material_row_count = sum(
            _canonical_money(row.get("amount")) is not None
            for row in result["reconciliation_rows"]
        )
        if source_receipts and len(normalized_numeric["entries"]) != material_row_count:
            raise AssuranceRunError(
                "numeric evidence ledger does not cover every material row"
            )
        final_receipt_by_id = {
            str(receipt["artifact_id"]): receipt
            for receipt in final_inventory["artifact_receipts"]
        }
        decision_id_set = {
            str(decision["decision_id"]) for decision in reviewed_decisions
        }
        source_id_set = {str(receipt["artifact_id"]) for receipt in source_receipts}
        expected_entries = _numeric_entries(
            reconciliation_rows=result["reconciliation_rows"],
            allocation_ledgers=allocations,
            source_receipts=source_receipts,
            reviewed_source_decisions=reviewed_decisions,
            prepared_receipt=prepared_receipt,
            final_receipts=final_inventory["artifact_receipts"],
            workbook_name="",
            final_root=out_dir / FINAL_OUTPUT_DIRECTORY,
            prepared=prepared,
            source_root=(
                Path(str(source_root)).resolve() if source_root is not None else out_dir
            ),
            strict_source_identity=True,
        )
        expected_numeric = build_numeric_evidence_ledger(
            expected_entries,
            ledger_id="audit_reconciliation_material_values",
        )
        if normalized_numeric != expected_numeric:
            raise AssuranceRunError(
                "numeric evidence ledger does not match current record/source/output closure"
            )
        for entry in normalized_numeric["entries"]:
            if entry["source"]["artifact_ref"] not in source_id_set:
                raise AssuranceRunError(
                    "numeric evidence source locator is not current"
                )
            if entry["decision_ref"] not in decision_id_set:
                raise AssuranceRunError(
                    "numeric evidence decision locator is not current"
                )
            if entry["prepared"]["artifact_ref"] != prepared_receipt["artifact_id"]:
                raise AssuranceRunError(
                    "numeric evidence prepared artifact locator is stale"
                )
            prepared_value = _json_record_locator_value(
                prepared,
                str(entry["prepared"]["locator"]),
            )
            if _canonical_money(prepared_value) != entry["value"]:
                raise AssuranceRunError(
                    "numeric evidence prepared value locator is stale"
                )
            for output in entry["outputs"]:
                output_receipt = final_receipt_by_id.get(output["artifact_ref"])
                if output_receipt is None:
                    raise AssuranceRunError(
                        "numeric evidence output locator is not current"
                    )
                output_path = (
                    out_dir / FINAL_OUTPUT_DIRECTORY / str(output_receipt["path"])
                )
                if output_path.suffix.lower() == ".json":
                    output_document = (
                        result
                        if output_receipt["path"] == "reconciliation_results.json"
                        else _read_json_mapping(output_path)
                    )
                    located_value = _json_record_locator_value(
                        output_document,
                        str(output["locator"]),
                    )
                elif output_path.suffix.lower() in {".xlsx", ".xlsm"}:
                    located_value = _workbook_locator_value(
                        output_path,
                        str(output["locator"]),
                    )
                elif output_path.suffix.lower() == ".docx":
                    located_value = _docx_locator_value(
                        output_path,
                        str(output["locator"]),
                    )
                else:
                    raise AssuranceRunError(
                        "numeric evidence output locator uses an unsupported artifact"
                    )
                if _canonical_money(located_value) != entry["value"]:
                    raise AssuranceRunError(
                        "numeric evidence output value locator is stale"
                    )

    gates = validate_gate_register(payload["gate_register"])
    if _read_json_mapping(out_dir / "assurance_gates.json") != gates:
        raise AssuranceRunError("assurance gate file is stale")
    source_status = (
        "not_applicable"
        if not source_receipts
        else (
            "passed"
            if qualifications
            and all(row["status"] == "qualified" for row in qualifications)
            else "failed"
        )
    )
    reconciliation_status = (
        "failed"
        if _reconciliation_gate_failed(
            result["checks"],
            allocations,
            result["reconciliation_rows"],
        )
        else "passed"
    )
    semantic_status = _review_gate_status(
        professional_review["records"],
        result["reconciliation_rows"],
        run_date=run_date,
    )
    reporting_status = (
        "passed"
        if source_status in {"passed", "not_applicable"}
        and reconciliation_status == "passed"
        and semantic_status in {"passed", "not_applicable"}
        else "blocked"
    )
    expected_statuses = {
        "source": source_status,
        "preparation": "passed",
        "reconciliation": reconciliation_status,
        "semantic_review": semantic_status,
        "reporting": reporting_status,
        "publication": "withheld",
    }
    actual_statuses = {name: gate["status"] for name, gate in gates["gates"].items()}
    if actual_statuses != expected_statuses:
        raise AssuranceRunError(
            "assurance gate statuses do not match the latest sealed result"
        )
    source_refs = [str(item["artifact_id"]) for item in source_receipts]
    implementation_refs = [str(item["artifact_id"]) for item in implementation_receipts]
    final_refs = [
        str(item["artifact_id"]) for item in final_inventory["artifact_receipts"]
    ]
    numeric_refs = [str(numeric["ledger_id"])] if isinstance(numeric, Mapping) else []
    expected_gates = build_gate_register(
        {
            "source": {
                "status": source_status,
                "evidence_refs": (
                    [*source_refs, str(reconciliation_receipt["artifact_id"])]
                    if source_status == "passed"
                    else source_refs
                ),
                "limitations": (
                    []
                    if source_status in {"passed", "not_applicable"}
                    else ["At least one source is unqualified or stale."]
                ),
            },
            "preparation": {
                "status": "passed",
                "evidence_refs": [
                    str(prepared_receipt["artifact_id"]),
                    *implementation_refs,
                ],
                "limitations": [],
            },
            "reconciliation": {
                "status": reconciliation_status,
                "evidence_refs": [
                    str(reconciliation_receipt["artifact_id"]),
                    *numeric_refs,
                    *(str(ledger["ledger_id"]) for ledger in allocations),
                ],
                "limitations": (
                    []
                    if reconciliation_status == "passed"
                    else ["At least one mechanical reconciliation check failed."]
                ),
            },
            "semantic_review": {
                "status": semantic_status,
                "evidence_refs": (
                    [str(professional_review_receipt["artifact_id"])]
                    if semantic_status == "passed"
                    else []
                ),
                "limitations": (
                    []
                    if semantic_status in {"passed", "not_applicable"}
                    else ["Required professional review is pending or failed."]
                ),
            },
            "reporting": {
                "status": reporting_status,
                "evidence_refs": final_refs if reporting_status == "passed" else [],
                "limitations": (
                    []
                    if reporting_status == "passed"
                    else ["Reporting is blocked by an upstream gate."]
                ),
            },
            "publication": {
                "status": "withheld",
                "evidence_refs": [],
                "limitations": [
                    "Publication is a separate action and was not performed."
                ],
            },
        }
    )
    if gates != expected_gates:
        raise AssuranceRunError(
            "assurance gate evidence and limitations do not match current authority"
        )
    return payload


def _command_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Replay Audit Reconciliation assurance controls."
    )
    subparsers = parser.add_subparsers(dest="command", required=True)
    replay = subparsers.add_parser(
        "validate-run-json",
        help="Replay a sealed run and emit the bounded MCP bridge result.",
    )
    replay.add_argument("output_dir")
    replay.add_argument("--expected-predecessor-checkpoint")
    capture = subparsers.add_parser(
        "capture-review-transition-json",
        help="Capture a fully replayed predecessor outside the candidate tree.",
    )
    capture.add_argument("output_dir")
    capture.add_argument("capture_dir")
    capture.add_argument("--expected-predecessor-checkpoint")
    retain = subparsers.add_parser(
        "retain-review-transition-json",
        help="Retain and replay a first-apply review transition.",
    )
    retain.add_argument("output_dir")
    retain.add_argument("capture_dir")
    retain.add_argument("--expected-predecessor-checkpoint")
    return parser


def _main(argv: Sequence[str] | None = None) -> int:
    args = _command_parser().parse_args(argv)
    if args.command in {
        "capture-review-transition-json",
        "retain-review-transition-json",
    }:
        try:
            result = (
                capture_review_transition_predecessor(
                    Path(args.output_dir),
                    Path(args.capture_dir),
                    expected_predecessor_checkpoint=(
                        args.expected_predecessor_checkpoint
                    ),
                )
                if args.command == "capture-review-transition-json"
                else retain_review_transition(
                    Path(args.output_dir),
                    Path(args.capture_dir),
                    expected_predecessor_checkpoint=(
                        args.expected_predecessor_checkpoint
                    ),
                )
            )
        except (OSError, ValueError):
            sys.stdout.write('{"ok":false,"error":"review transition replay failed"}\n')
            return 1
        sys.stdout.write(
            json.dumps(
                {"ok": True, **result},
                ensure_ascii=False,
                separators=(",", ":"),
            )
            + "\n"
        )
        return 0
    if args.command != "validate-run-json":
        return 2
    try:
        assurance = validate_assurance_run(
            Path(args.output_dir),
            expected_predecessor_checkpoint=args.expected_predecessor_checkpoint,
        )
        reconciliation = _read_json_mapping(
            Path(args.output_dir).resolve()
            / FINAL_OUTPUT_DIRECTORY
            / "reconciliation_results.json"
        )
    except (OSError, ValueError):
        sys.stdout.write('{"ok":false,"error":"complete assurance replay failed"}\n')
        return 1
    bounded_assurance = {
        "run_date": assurance["run_date"],
        "content_sha256": assurance["content_sha256"],
        "professional_review_authority": assurance["professional_review_authority"],
        "gate_register": assurance["gate_register"],
    }
    response = {
        "ok": True,
        "assurance": bounded_assurance,
        "result": {
            "reconciliation_rows": reconciliation["reconciliation_rows"],
        },
    }
    sys.stdout.write(
        json.dumps(response, ensure_ascii=False, separators=(",", ":")) + "\n"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(_main())
