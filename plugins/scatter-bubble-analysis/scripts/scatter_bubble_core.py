"""Core workflow for the scatter-bubble-analysis Codex plugin."""

from __future__ import annotations

import argparse
import calendar
import json
import logging
import re
import sys
import zipfile
from dataclasses import dataclass
from datetime import UTC, date, datetime, timedelta
from pathlib import Path
from typing import Any, Sequence

import polars as pl
from legacy_scatter_bubble_charting import (
    CANONICAL_DATE,
    CANONICAL_PERIOD,
    CURRENT_PERIOD,
    LegacyPreparedDataCache,
    cleanup_legacy_imports,
    ensure_legacy_import_path,
    write_legacy_scatter_bubble_chart,
)

SCRIPT_DIR = Path(__file__).resolve().parent


def _ensure_local_review_session_import() -> None:
    """Use this plugin's review-session module in multi-plugin test runs."""

    script_dir = str(SCRIPT_DIR)
    if script_dir in sys.path:
        sys.path.remove(script_dir)
    sys.path.insert(0, script_dir)
    module = sys.modules.get("review_session")
    module_file = getattr(module, "__file__", None) if module is not None else None
    if module_file and Path(module_file).resolve().is_relative_to(SCRIPT_DIR.resolve()):
        return
    if module is not None:
        del sys.modules["review_session"]


_ensure_local_review_session_import()
from review_session import write_review_session_artifacts, write_run_intake

ensure_legacy_import_path()
from modules.chart_harness import (  # noqa: E402  # isort: skip
    apply_recipe_filters,
    apply_recipe_cohorts,
    available_analysis_context,
    json_safe,
    PERIOD_TYPE_CALENDAR,
    PERIOD_TYPE_FISCAL,
    PERIOD_TYPE_ROLLING,
    PERIOD_TYPE_TO_DATE,
    period_contract_options,
    period_label_expression,
    preserve_recipe_cohorts,
    preserve_recipe_filters,
    recipe_cohort_dimension_names,
    recipe_cohort_period_labels,
    recipe_cohort_source_dimensions,
    reporting_entity_label_from_recipe,
    reporting_subject_label_from_recipe,
    write_prepared_data_manifest,
)
from modules.utilities.config import get_metric_array_params, get_naming_params
from modules.utilities.helpers import get_schema_and_column_names

__all__ = [
    "SCHEMA_VERSION",
    "ScatterBubbleRunResult",
    "add_common_args",
    "build_chart_specs",
    "build_recipe",
    "configure_logging",
    "prepare_canonical_frame",
    "read_json",
    "read_table",
    "run_scatter_bubble",
    "write_json",
]

LOGGER = logging.getLogger(__name__)
SCHEMA_VERSION = "1.0"
DEFAULT_DATE = datetime(2026, 1, 1, tzinfo=UTC).date()
ARTIFACT_MODE_DATA_ONLY = "data_only"
ARTIFACT_MODE_DATA_AND_RENDER = "data_and_render"
ARTIFACT_MODE_RENDER_ONLY = "render_only"
ARTIFACT_MODES = {
    ARTIFACT_MODE_DATA_ONLY,
    ARTIFACT_MODE_DATA_AND_RENDER,
    ARTIFACT_MODE_RENDER_ONLY,
}
CSV_EXTENSIONS = {".csv", ".tsv", ".psv", ".txt"}
EXCEL_EXTENSIONS = {".xlsx", ".xlsm", ".xls"}
DISPLAY_VALUE_PREFIX_DIVISORS = (
    ("t", 1_000_000_000_000),
    ("b", 1_000_000_000),
    ("m", 1_000_000),
    ("k", 1_000),
)
METRIC_TOKEN_PATTERN = re.compile(r"[a-z0-9]+")
ROLLING_PERIOD_DERIVATION = "latest_rolling_year_vs_prior_year"


@dataclass(frozen=True)
class ScatterBubbleRunResult:
    """Result object returned by ``run_scatter_bubble``."""

    canonical_frame: pl.DataFrame
    audit: dict[str, Any]
    summary_markdown: str
    artifact_paths: list[str]


def utc_now() -> str:
    """Return the current UTC timestamp."""

    return datetime.now(UTC).isoformat()


def configure_logging(verbose: bool = False) -> None:
    """Configure plugin logging."""

    logging.basicConfig(
        level=logging.DEBUG if verbose else logging.INFO,
        format="%(levelname)s:%(name)s:%(message)s",
    )


def add_common_args(parser: argparse.ArgumentParser) -> None:
    """Add common CLI arguments for plugin scripts."""

    parser.add_argument("input_file", type=Path)
    parser.add_argument("--output-dir", type=Path)
    parser.add_argument("--recipe", type=Path)
    parser.add_argument("--language", default="en")
    parser.add_argument("--currency", default=None)
    parser.add_argument(
        "--artifact-mode",
        choices=sorted(ARTIFACT_MODES),
        default=ARTIFACT_MODE_DATA_AND_RENDER,
        help=(
            "Write chart data/context only or keep the legacy data-and-render behavior."
        ),
    )
    parser.add_argument("--verbose", action="store_true")


def default_output_dir(input_file: Path) -> Path:
    """Return the safe default output directory next to the input file."""

    return (
        input_file.expanduser().resolve().parent / "output" / "scatter-bubble-analysis"
    )


def read_json(path: Path | None) -> dict[str, Any] | None:
    """Read a JSON object from ``path`` when provided."""

    if path is None:
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    """Write a JSON file with deterministic formatting."""

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, default=str) + "\n", encoding="utf-8")


def _collect_csv_scan(path: Path, *, separator: str) -> pl.DataFrame:
    """Read delimited input through a lazy scan and collect once."""

    lf = pl.scan_csv(path, separator=separator, infer_schema_length=10000)
    try:
        return lf.collect(engine="streaming")
    except pl.exceptions.PolarsError:
        return lf.collect()


def read_table(path: Path) -> pl.DataFrame:
    """Read CSV/TSV/PSV or Excel data into a Polars DataFrame."""

    suffix = path.suffix.lower()
    if suffix in CSV_EXTENSIONS:
        separator = ","
        if suffix == ".tsv":
            separator = "\t"
        elif suffix == ".psv":
            separator = "|"
        return _collect_csv_scan(path, separator=separator)
    if suffix in EXCEL_EXTENSIONS:
        return pl.read_excel(path)
    raise ValueError(f"Unsupported input extension: {suffix}")


def _is_numeric_dtype(dtype: pl.DataType) -> bool:
    return bool(getattr(dtype, "is_numeric", lambda: False)())


def _is_date_dtype(dtype: pl.DataType) -> bool:
    return dtype in {pl.Date, pl.Datetime}


def _column_score(column: str, keywords: tuple[str, ...]) -> int:
    lowered = column.lower()
    return sum(1 for keyword in keywords if keyword in lowered)


def _first_matching_column(
    columns: list[str],
    keywords: tuple[str, ...],
    *,
    exclude: set[str] | None = None,
) -> str | None:
    excluded = exclude or set()
    candidates = [column for column in columns if column not in excluded]
    ranked = sorted(
        candidates,
        key=lambda column: (_column_score(column, keywords), column.lower()),
        reverse=True,
    )
    if ranked and _column_score(ranked[0], keywords) > 0:
        return ranked[0]
    return candidates[0] if candidates else None


def _unique_count(frame: pl.DataFrame, column: str) -> int:
    return int(frame.select(pl.col(column).n_unique()).item() or 0)


def _dimension_candidates(frame: pl.DataFrame, numeric_columns: set[str]) -> list[str]:
    columns, schema = get_schema_and_column_names(frame)
    result: list[str] = []
    for column in columns:
        if column in numeric_columns or column in {CANONICAL_DATE, CANONICAL_PERIOD}:
            continue
        dtype = schema[column]
        if _is_date_dtype(dtype):
            continue
        unique_count = _unique_count(frame, column)
        if 1 < unique_count <= max(frame.height, 2):
            result.append(column)
    preferred = ("brand", "product", "item", "company", "customer", "retailer")
    return sorted(
        result,
        key=lambda column: (
            0 if _column_score(column, preferred) else 1,
            _unique_count(frame, column),
            column.lower(),
        ),
    )


def _period_column(columns: list[str]) -> str | None:
    for column in columns:
        lowered = column.lower()
        if lowered in {"period", "scenario", "month", "year_month"}:
            return column
    return None


def _date_column(columns: list[str], schema: dict[str, pl.DataType]) -> str | None:
    for column in columns:
        if _is_date_dtype(schema[column]):
            return column
    for column in columns:
        lowered = column.lower()
        if "date" in lowered or lowered in {"month", "year_month"}:
            return column
    return None


def _coalesce_mapping(
    mappings: dict[str, Any],
    key: str,
    fallback: str | None,
) -> str | None:
    value = mappings[key] if key in mappings else None
    return str(value) if value else fallback


def build_recipe(
    input_path: Path,
    frame: pl.DataFrame,
    *,
    language: str = "en",
    existing_recipe: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Infer or merge the recipe needed by the scatter/bubble workflow."""

    columns, schema = get_schema_and_column_names(frame)
    numeric_columns = [
        column for column in columns if _is_numeric_dtype(schema[column])
    ]
    if len(numeric_columns) < 2:
        raise ValueError("Scatter charts require at least two numeric metric columns.")
    dimensions = _dimension_candidates(frame, set(numeric_columns))
    if not dimensions:
        raise ValueError("Scatter/bubble charts require at least one dimension column.")

    existing_mappings = dict((existing_recipe or {}).get("mappings") or {})
    existing_options = dict((existing_recipe or {}).get("options") or {})
    x_metric = _coalesce_mapping(
        existing_mappings,
        "x_metric_column",
        _first_matching_column(
            numeric_columns,
            ("price", "rate", "margin", "unit"),
        ),
    )
    y_metric = _coalesce_mapping(
        existing_mappings,
        "y_metric_column",
        _first_matching_column(
            numeric_columns,
            ("unit", "volume", "qty", "quantity", "sales", "revenue"),
            exclude={str(x_metric)} if x_metric else set(),
        ),
    )
    if y_metric is None:
        y_metric = numeric_columns[1]
    bubble_size = _coalesce_mapping(
        existing_mappings,
        "bubble_size_metric_column",
        _first_matching_column(
            numeric_columns,
            ("sales", "revenue", "amount", "value"),
            exclude={str(x_metric), str(y_metric)},
        ),
    )
    if bubble_size is None:
        bubble_size = y_metric
    mapped_dimensions = (
        existing_mappings["dimensions"] if "dimensions" in existing_mappings else None
    )
    if mapped_dimensions:
        cohort_dimension_names = recipe_cohort_dimension_names(existing_recipe)
        dimensions = [
            str(item)
            for item in mapped_dimensions
            if str(item) in columns or str(item) in cohort_dimension_names
        ]
    dot_dimension = _coalesce_mapping(
        existing_mappings,
        "dot_dimension",
        dimensions[0],
    )
    color_dimension = _coalesce_mapping(
        existing_mappings,
        "color_dimension",
        dimensions[1] if len(dimensions) > 1 else None,
    )
    small_multiples_dimension = _coalesce_mapping(
        existing_mappings,
        "small_multiples_dimension",
        next(
            (
                dimension
                for dimension in dimensions
                if dimension not in {dot_dimension, color_dimension}
            ),
            color_dimension,
        ),
    )
    recipe = {
        "schema_version": SCHEMA_VERSION,
        "source_file": str(input_path),
        "language": language,
        "mappings": {
            "x_metric_column": x_metric,
            "y_metric_column": y_metric,
            "bubble_size_metric_column": bubble_size,
            "dimensions": dimensions,
            "dot_dimension": dot_dimension,
            "color_dimension": color_dimension,
            "small_multiples_dimension": small_multiples_dimension,
            "period_column": (
                existing_mappings["period_column"]
                if "period_column" in existing_mappings
                else _period_column(columns)
            ),
            "date_column": (
                existing_mappings["date_column"]
                if "date_column" in existing_mappings
                else _date_column(columns, schema)
            ),
        },
        "options": {
            "currency": (
                existing_options["currency"]
                if "currency" in existing_options
                else "EUR"
            ),
            "color_palette": (
                existing_options["color_palette"]
                if "color_palette" in existing_options
                else "bain"
            ),
            "charts": (
                existing_options["charts"] if "charts" in existing_options else []
            ),
            "small_multiples": (
                existing_options["small_multiples"]
                if "small_multiples" in existing_options
                else bool(small_multiples_dimension)
            ),
            "max_chart_items": (
                existing_options["max_chart_items"]
                if "max_chart_items" in existing_options
                else 12
            ),
            "small_multiples_max_panels": (
                existing_options["small_multiples_max_panels"]
                if "small_multiples_max_panels" in existing_options
                else 6
            ),
        },
    }
    if existing_options.get("reporting_entity_label"):
        recipe["options"]["reporting_entity_label"] = str(
            existing_options["reporting_entity_label"]
        )
    if existing_options.get("metric_aliases"):
        recipe["options"]["metric_aliases"] = dict(existing_options["metric_aliases"])
    if existing_options.get("period_derivation"):
        recipe["options"]["period_derivation"] = dict(
            existing_options["period_derivation"]
        )
    period_option_keys = (
        "period_type",
        "period_grain",
        "period_comparison_mode",
        "fiscal_start_month",
        "rolling_window_months",
        "rolling_window_days",
        "current_period_label",
        "previous_period_label",
    )
    for key in period_option_keys:
        if key in existing_options:
            recipe["options"][key] = existing_options[key]
    if _period_contract_requested(recipe["options"]):
        default_period_type = (
            PERIOD_TYPE_ROLLING
            if "rolling_window_months" in recipe["options"]
            or "rolling_window_days" in recipe["options"]
            else PERIOD_TYPE_CALENDAR
        )
        contract = period_contract_options(
            recipe["options"], default_type=default_period_type
        )
        recipe["options"]["period_type"] = contract["period_type"]
        recipe["options"]["period_grain"] = contract["period_grain"]
        recipe["options"]["fiscal_start_month"] = contract["fiscal_start_month"]
    return recipe


def _recipe_metric_aliases(recipe: dict[str, Any]) -> dict[str, str]:
    """Return target legacy metric name -> source column aliases from the recipe."""

    options = dict(recipe.get("options") or {})
    aliases = options["metric_aliases"] if "metric_aliases" in options else {}
    return {str(target): str(source) for target, source in dict(aliases).items()}


def _legacy_metric_dependencies(metric: str) -> list[str]:
    """Return raw legacy metric columns required to prepare ``metric``."""

    naming_params = get_naming_params()
    amount = naming_params["monetaryLocalCurrencyName"]
    units = naming_params["unitsName"]
    volume = naming_params["volumeName"]
    margin = naming_params["marginName"]
    metric_array_params = get_metric_array_params()
    growth_metrics = set(metric_array_params[naming_params["growthMetricArray"]])
    if metric == naming_params["pricePerUnitName"]:
        return [amount, units]
    if metric == naming_params["pricePerVolumeName"]:
        return [amount, volume]
    if metric == naming_params["salesGrowthName"]:
        return [amount]
    if metric == naming_params["unitsGrowthName"]:
        return [units]
    if metric == naming_params["volumeGrowthName"]:
        return [volume]
    if metric == naming_params["marginGrowthName"]:
        return [margin]
    if metric in growth_metrics:
        return [amount]
    return []


def _canonical_metrics_for_requested_metrics(
    metrics: list[str],
    raw_columns: set[str] | None = None,
) -> list[str]:
    """Return metric columns that must exist before legacy chart preparation."""

    result: list[str] = []
    raw_columns = raw_columns or set()
    for metric in metrics:
        if metric in raw_columns:
            result.append(metric)
            continue
        dependencies = _legacy_metric_dependencies(metric)
        if dependencies:
            result.extend(dependencies)
        else:
            result.append(metric)
    return list(dict.fromkeys(result))


def _legacy_chart_value_columns(metrics: list[str]) -> list[str]:
    """Return value columns passed to legacy, including derived metric inputs."""

    result = list(metrics)
    for metric in metrics:
        result.extend(_legacy_metric_dependencies(metric))
    return list(dict.fromkeys(result))


def _metric_source_column(
    metric: str,
    raw_columns: set[str],
    metric_aliases: dict[str, str],
) -> str | None:
    """Return the source column that provides ``metric`` before legacy derivation."""

    if metric in raw_columns:
        return metric
    alias = metric_aliases.get(metric)
    if alias and alias in raw_columns:
        return alias
    return None


def _source_date_expr(frame: pl.DataFrame, source_column: str) -> pl.Expr:
    _columns, schema = get_schema_and_column_names(frame)
    dtype = schema[source_column]
    if dtype == pl.Date:
        return pl.col(source_column)
    if dtype == pl.Datetime:
        return pl.col(source_column).dt.date()
    return (
        pl.col(source_column)
        .cast(pl.Utf8, strict=False)
        .str.strptime(pl.Date, strict=False)
        .fill_null(DEFAULT_DATE)
    )


def _date_expr(frame: pl.DataFrame, source_column: str | None) -> pl.Expr:
    if not source_column:
        return pl.lit(DEFAULT_DATE).alias(CANONICAL_DATE)
    return _source_date_expr(frame, source_column).alias(CANONICAL_DATE)


def _to_date(value: Any) -> date:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    return datetime.fromisoformat(str(value)[:10]).date()


def _add_months(value: date, months: int) -> date:
    """Shift a date by whole months, preserving the closest valid day."""

    month_index = value.month - 1 + months
    year = value.year + month_index // 12
    month = month_index % 12 + 1
    day = min(value.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


def _period_contract_requested(options: dict[str, Any]) -> bool:
    """Return whether options explicitly request date period derivation."""

    return any(
        key in options
        for key in (
            "period_type",
            "period_grain",
            "period_comparison_mode",
            "fiscal_start_month",
            "rolling_window_months",
            "rolling_window_days",
        )
    )


def _rolling_window_days(options: dict[str, Any], period_grain: str) -> int:
    """Return a rolling window length from explicit options or period grain."""

    for key, multiplier in (("rolling_window_days", 1), ("rolling_window_months", 30)):
        if key not in options:
            continue
        try:
            value = int(options[key])
        except (TypeError, ValueError):
            continue
        if value > 0:
            return value * multiplier
    if period_grain == "week":
        return 7
    if period_grain == "month":
        return 30
    if period_grain == "quarter":
        return 91
    return 364


def _to_date_period_expression(
    frame: pl.DataFrame,
    source_column: str,
    options: dict[str, Any],
) -> tuple[pl.Expr, dict[str, Any]]:
    """Return AC/PY labels for a current/prior-year to-date window."""

    contract = period_contract_options(options, default_type=PERIOD_TYPE_TO_DATE)
    latest_value = frame.select(_source_date_expr(frame, source_column).max()).item()
    latest_date = _to_date(latest_value)
    fiscal_start_month = int(contract["fiscal_start_month"])
    current_year = (
        latest_date.year
        if latest_date.month >= fiscal_start_month
        else latest_date.year - 1
    )
    current_start = date(current_year, fiscal_start_month, 1)
    previous_start = _add_months(current_start, -12)
    previous_end = _add_months(latest_date, -12)
    current_label = str(options.get("current_period_label") or CURRENT_PERIOD)
    previous_label = str(options.get("previous_period_label") or "PY")
    source_date = _source_date_expr(frame, source_column)
    expression = (
        pl.when((source_date >= current_start) & (source_date <= latest_date))
        .then(pl.lit(current_label))
        .when((source_date >= previous_start) & (source_date <= previous_end))
        .then(pl.lit(previous_label))
        .otherwise(None)
        .alias(CANONICAL_PERIOD)
    )
    return expression, {
        "type": "date_period_contract",
        "period_type": PERIOD_TYPE_TO_DATE,
        "period_grain": contract["period_grain"],
        "fiscal_start_month": fiscal_start_month,
        "date_column": source_column,
        "current_period": current_label,
        "previous_period": previous_label,
        "latest_date": latest_date.isoformat(),
        "current_start": current_start.isoformat(),
        "previous_start": previous_start.isoformat(),
        "previous_end": previous_end.isoformat(),
        "filters_to_selected_windows": True,
    }


def _period_expr(
    frame: pl.DataFrame,
    recipe: dict[str, Any],
    period_column: str | None,
    date_column: str | None,
) -> tuple[pl.Expr, dict[str, Any]]:
    options = dict(recipe.get("options") or {})
    derivation = (
        dict(options["period_derivation"]) if "period_derivation" in options else {}
    )
    if derivation:
        derivation_type = str(derivation.get("type", ""))
        if derivation_type != ROLLING_PERIOD_DERIVATION:
            raise ValueError(f"Unsupported period derivation type: {derivation_type}")
        source_column = str(derivation.get("date_column") or date_column or "")
        if not source_column:
            raise ValueError("period_derivation requires a date column")
        latest_value = frame.select(
            _source_date_expr(frame, source_column).max()
        ).item()
        latest_date = _to_date(latest_value)
        window_days = int(derivation.get("window_days", 364))
        current_start = latest_date - timedelta(days=window_days - 1)
        previous_start = current_start - timedelta(days=window_days)
        current_label = str(derivation.get("current_period", CURRENT_PERIOD))
        previous_label = str(derivation.get("previous_period", "PY"))
        source_date = _source_date_expr(frame, source_column)
        expression = (
            pl.when((source_date >= current_start) & (source_date <= latest_date))
            .then(pl.lit(current_label))
            .when((source_date >= previous_start) & (source_date < current_start))
            .then(pl.lit(previous_label))
            .otherwise(None)
            .alias(CANONICAL_PERIOD)
        )
        return expression, {
            "type": derivation_type,
            "date_column": source_column,
            "current_period": current_label,
            "previous_period": previous_label,
            "latest_date": latest_date.isoformat(),
            "current_start": current_start.isoformat(),
            "previous_start": previous_start.isoformat(),
            "window_days": window_days,
            "filters_to_selected_windows": True,
        }
    if date_column and _period_contract_requested(options):
        source_column = str(date_column)
        default_period_type = (
            PERIOD_TYPE_ROLLING
            if "rolling_window_months" in options or "rolling_window_days" in options
            else PERIOD_TYPE_CALENDAR
        )
        contract = period_contract_options(options, default_type=default_period_type)
        period_type = str(contract["period_type"])
        source_date = _source_date_expr(frame, source_column)
        if period_type == PERIOD_TYPE_ROLLING:
            latest_value = frame.select(source_date.max()).item()
            latest_date = _to_date(latest_value)
            window_days = _rolling_window_days(options, str(contract["period_grain"]))
            current_start = latest_date - timedelta(days=window_days - 1)
            previous_start = current_start - timedelta(days=window_days)
            current_label = str(options.get("current_period_label") or CURRENT_PERIOD)
            previous_label = str(options.get("previous_period_label") or "PY")
            expression = (
                pl.when((source_date >= current_start) & (source_date <= latest_date))
                .then(pl.lit(current_label))
                .when((source_date >= previous_start) & (source_date < current_start))
                .then(pl.lit(previous_label))
                .otherwise(None)
                .alias(CANONICAL_PERIOD)
            )
            return expression, {
                "type": "date_period_contract",
                "period_type": PERIOD_TYPE_ROLLING,
                "period_grain": contract["period_grain"],
                "date_column": source_column,
                "current_period": current_label,
                "previous_period": previous_label,
                "latest_date": latest_date.isoformat(),
                "current_start": current_start.isoformat(),
                "previous_start": previous_start.isoformat(),
                "window_days": window_days,
                "filters_to_selected_windows": True,
            }
        if period_type == PERIOD_TYPE_TO_DATE:
            return _to_date_period_expression(frame, source_column, options)
        if period_type in {PERIOD_TYPE_CALENDAR, PERIOD_TYPE_FISCAL}:
            fiscal_start_month = (
                int(contract["fiscal_start_month"])
                if period_type == PERIOD_TYPE_FISCAL
                else 1
            )
            return (
                period_label_expression(
                    source_date,
                    period_grain=str(contract["period_grain"]),
                    fiscal_start_month=fiscal_start_month,
                ).alias(CANONICAL_PERIOD),
                {
                    "type": "date_period_contract",
                    "period_type": period_type,
                    "period_grain": contract["period_grain"],
                    "fiscal_start_month": fiscal_start_month,
                    "date_column": source_column,
                    "filters_to_selected_windows": False,
                },
            )
    if period_column:
        return (
            pl.col(str(period_column))
            .cast(pl.Utf8, strict=False)
            .fill_null(CURRENT_PERIOD)
            .alias(CANONICAL_PERIOD),
            {"type": "source_column", "period_column": period_column},
        )
    return (
        pl.lit(CURRENT_PERIOD).alias(CANONICAL_PERIOD),
        {"type": "constant", "period": CURRENT_PERIOD},
    )


def prepare_canonical_frame(
    frame: pl.DataFrame, recipe: dict[str, Any]
) -> tuple[pl.DataFrame, dict[str, Any]]:
    """Return canonical data used by the legacy charting adapter."""

    mappings = recipe["mappings"]
    metrics = [
        str(mappings["x_metric_column"]),
        str(mappings["y_metric_column"]),
        str(mappings["bubble_size_metric_column"]),
    ]
    raw_columns, _schema = get_schema_and_column_names(frame)
    raw_column_set = set(raw_columns)
    canonical_metrics = _canonical_metrics_for_requested_metrics(
        metrics, raw_column_set
    )
    legacy_derived_metrics = [
        metric
        for metric in metrics
        if metric not in raw_column_set and _legacy_metric_dependencies(metric)
    ]
    dimensions = [str(item) for item in mappings["dimensions"]]
    metric_aliases = _recipe_metric_aliases(recipe)
    metric_sources: dict[str, str] = {}
    missing_metrics: list[str] = []
    for metric in canonical_metrics:
        source_column = _metric_source_column(metric, raw_column_set, metric_aliases)
        if source_column is None:
            missing_metrics.append(metric)
        else:
            metric_sources[metric] = source_column
    if missing_metrics:
        missing = ", ".join(missing_metrics)
        raise ValueError(f"Missing source columns for legacy metrics: {missing}")
    base_dimensions = list(
        dict.fromkeys(
            [
                *[dimension for dimension in dimensions if dimension in raw_columns],
                *[
                    dimension
                    for dimension in recipe_cohort_source_dimensions(recipe)
                    if dimension in raw_columns
                ],
            ]
        )
    )
    period_column = mappings["period_column"]
    date_column = mappings["date_column"]
    period_expression, period_audit = _period_expr(
        frame,
        recipe,
        str(period_column) if period_column else None,
        str(date_column) if date_column else None,
    )
    period_derivation_date_column = (
        str(period_audit["date_column"]) if "date_column" in period_audit else None
    )
    date_source_column = (
        str(date_column) if date_column else period_derivation_date_column
    )
    select_columns = list(
        dict.fromkeys(
            [
                *metric_sources.values(),
                *base_dimensions,
                *([str(period_column)] if period_column else []),
                *([date_source_column] if date_source_column else []),
                *(
                    [period_derivation_date_column]
                    if period_derivation_date_column
                    else []
                ),
            ]
        )
    )
    canonical = frame.select(select_columns)
    expressions: list[pl.Expr] = [
        pl.col(source).cast(pl.Float64, strict=False).fill_null(0.0).alias(metric)
        for metric, source in metric_sources.items()
    ]
    expressions.extend(
        [
            pl.col(dimension)
            .cast(pl.Utf8, strict=False)
            .fill_null("Undefined")
            .alias(dimension)
            for dimension in base_dimensions
        ]
    )
    expressions.append(period_expression)
    expressions.append(_date_expr(frame, date_source_column))
    canonical = canonical.with_columns(expressions).select(
        list(
            dict.fromkeys(
                [
                    CANONICAL_DATE,
                    CANONICAL_PERIOD,
                    *base_dimensions,
                    *canonical_metrics,
                ]
            )
        )
    )
    if period_audit.get("filters_to_selected_windows"):
        canonical = canonical.filter(pl.col(CANONICAL_PERIOD).is_not_null())
    preparation_audit = {
        "status": "prepared",
        "input_row_count": frame.height,
        "canonical_row_count": canonical.height,
        "metrics": metrics,
        "canonical_metrics": canonical_metrics,
        "legacy_derived_metrics": legacy_derived_metrics,
        "metric_aliases": metric_aliases,
        "metric_sources": metric_sources,
        "dimensions": dimensions,
        "cohort_source_dimensions": [
            dimension
            for dimension in recipe_cohort_source_dimensions(recipe)
            if dimension in raw_columns
        ],
        "period_column": period_column,
        "date_column": date_column,
        "period_derivation": period_audit,
    }
    return canonical, preparation_audit


def _period_values(canonical: pl.DataFrame) -> list[str]:
    values = (
        canonical.select(pl.col(CANONICAL_PERIOD).cast(pl.Utf8).unique())
        .to_series()
        .to_list()
    )
    return sorted(str(value) for value in values if value is not None)


def _metrics_require_growth_periods(metrics: list[str]) -> bool:
    naming_params = get_naming_params()
    metric_array_params = get_metric_array_params()
    growth_metrics = set(metric_array_params[naming_params["growthMetricArray"]])
    return any(metric in growth_metrics for metric in metrics)


def _selected_periods_for_metrics(
    period_values: list[str],
    metrics: list[str],
) -> tuple[list[str], str]:
    if not period_values:
        return [CURRENT_PERIOD], CURRENT_PERIOD
    current_period = (
        CURRENT_PERIOD if CURRENT_PERIOD in period_values else period_values[-1]
    )
    if _metrics_require_growth_periods(metrics) and len(period_values) >= 2:
        previous_period = "PY" if "PY" in period_values else period_values[-2]
        return [previous_period, current_period], current_period
    return [current_period], current_period


def _metric_name_suggests_product(
    x_metric: str, y_metric: str, product_metric: str
) -> bool:
    """Return whether metric names suggest a product relationship."""

    has_price_and_units = (
        _metric_name_suggests_price(x_metric) and _metric_name_suggests_units(y_metric)
    ) or (
        _metric_name_suggests_price(y_metric) and _metric_name_suggests_units(x_metric)
    )
    return has_price_and_units and _metric_name_suggests_value(product_metric)


def _metric_tokens(metric: str) -> set[str]:
    """Return normalized tokens for metric-name semantic checks."""

    return set(METRIC_TOKEN_PATTERN.findall(metric.lower()))


def _metric_values_support_product(
    canonical: pl.DataFrame,
    x_metric: str,
    y_metric: str,
    product_metric: str,
) -> bool:
    """Return whether row values show ``product_metric`` is ``x_metric * y_metric``."""

    if product_metric in {x_metric, y_metric}:
        return False
    metrics = [x_metric, y_metric, product_metric]
    if any(metric not in canonical.collect_schema().names() for metric in metrics):
        return False
    stats = (
        canonical.select(
            (pl.col(x_metric) * pl.col(y_metric)).alias("__expected"),
            pl.col(product_metric).alias("__actual"),
        )
        .select(
            (pl.col("__actual") - pl.col("__expected")).abs().sum().alias("abs_error"),
            pl.col("__actual").abs().sum().alias("abs_actual"),
            pl.col("__expected").abs().sum().alias("abs_expected"),
        )
        .row(0, named=True)
    )
    denominator = max(
        float(stats["abs_actual"] or 0.0),
        float(stats["abs_expected"] or 0.0),
        1.0,
    )
    relative_error = float(stats["abs_error"] or 0.0) / denominator
    return relative_error <= 0.05


def _product_metric_for_isolines(
    canonical: pl.DataFrame,
    x_metric: str,
    y_metric: str,
    product_metric: str,
) -> str | None:
    """Return the metric to use for legacy isolines when x*y generates it."""

    if _metric_values_support_product(canonical, x_metric, y_metric, product_metric):
        return product_metric
    if _metric_name_suggests_product(x_metric, y_metric, product_metric):
        return product_metric
    return None


def _metric_name_suggests_value(metric: str) -> bool:
    tokens = _metric_tokens(metric)
    return bool(tokens & {"sales", "revenue", "amount", "value", "turnover"})


def _metric_name_suggests_price(metric: str) -> bool:
    tokens = _metric_tokens(metric)
    return "price" in tokens


def _metric_name_suggests_units(metric: str) -> bool:
    tokens = _metric_tokens(metric)
    if tokens & {"price", "rate", "cost"}:
        return False
    return bool(tokens & {"unit", "units", "volume", "qty", "quantity"})


def _metric_name_suggests_ratio(metric: str) -> bool:
    tokens = _metric_tokens(metric)
    if tokens & {"growth", "change", "variance"}:
        return False
    return bool(tokens & {"price", "rate", "avg", "average"})


def _weighted_x_metric_operands(
    x_metric: str, y_metric: str, bubble_size_metric: str
) -> tuple[str, str] | None:
    """Return numerator and denominator for weighted x-axis summaries."""

    if not _metric_name_suggests_ratio(x_metric):
        return None
    if _metric_name_suggests_value(y_metric) and _metric_name_suggests_units(
        bubble_size_metric
    ):
        return y_metric, bubble_size_metric
    if _metric_name_suggests_value(bubble_size_metric) and _metric_name_suggests_units(
        y_metric
    ):
        return bubble_size_metric, y_metric
    return None


def _display_value_prefix(canonical: pl.DataFrame, metric: str) -> str:
    """Return the legacy value-prefix token for a metric, if useful."""

    if metric not in canonical.collect_schema().names():
        return ""
    max_value = canonical.select(pl.col(metric).abs().max()).item()
    max_abs = abs(float(max_value or 0.0))
    for prefix, divisor in DISPLAY_VALUE_PREFIX_DIVISORS:
        if max_abs > divisor:
            return prefix
    return ""


def _is_hierarchical_dimension_pair(
    canonical: pl.DataFrame,
    child_dimension: str | None,
    parent_dimension: str | None,
) -> bool:
    """Return whether each child item belongs to one parent item."""

    if (
        not child_dimension
        or not parent_dimension
        or child_dimension == parent_dimension
    ):
        return False
    columns = set(canonical.collect_schema().names())
    if child_dimension not in columns or parent_dimension not in columns:
        return False
    max_parent_count = (
        canonical.group_by(child_dimension)
        .agg(pl.col(parent_dimension).n_unique().alias("__parent_count"))
        .select(pl.col("__parent_count").max())
        .item()
    )
    return int(max_parent_count or 0) <= 1


def build_chart_specs(
    canonical: pl.DataFrame, recipe: dict[str, Any]
) -> list[dict[str, Any]]:
    """Build the legacy scatter/bubble chart attempt list."""

    mappings = recipe["mappings"]
    options = recipe["options"]
    requested = set(str(item) for item in options["charts"])
    period_values = _period_values(canonical)
    x_metric = str(mappings["x_metric_column"])
    y_metric = str(mappings["y_metric_column"])
    bubble_size = str(mappings["bubble_size_metric_column"])
    chart_metrics = [x_metric, y_metric, bubble_size]
    selected_periods, selected_period = _selected_periods_for_metrics(
        period_values,
        chart_metrics,
    )
    dot_dimension = str(mappings["dot_dimension"])
    color_dimension = (
        str(mappings["color_dimension"]) if mappings["color_dimension"] else None
    )
    small_multiples_dimension = (
        str(mappings["small_multiples_dimension"])
        if mappings["small_multiples_dimension"]
        else None
    )
    dimensions = [str(item) for item in mappings["dimensions"]]
    bubble_color_dimension = (
        color_dimension
        if _is_hierarchical_dimension_pair(canonical, dot_dimension, color_dimension)
        else None
    )
    metrics = _legacy_chart_value_columns(chart_metrics)
    reporting_entity_label = reporting_entity_label_from_recipe(recipe)
    reporting_subject_label = reporting_subject_label_from_recipe(recipe)
    isoline_metric = _product_metric_for_isolines(
        canonical,
        x_metric,
        y_metric,
        bubble_size,
    )
    display_value_prefix = _display_value_prefix(canonical, bubble_size)
    display_value_prefixes = (
        {bubble_size: display_value_prefix} if display_value_prefix else {}
    )
    common = {
        "dimensions": dimensions,
        "dot_dimension": dot_dimension,
        "color_dimension": color_dimension,
        "x_metric": x_metric,
        "y_metric": y_metric,
        "bubble_size_metric": bubble_size,
        "metrics": metrics,
        "selected_periods": selected_periods,
        "to_plot_period": selected_period,
        "reporting_entity_label": reporting_entity_label,
        "reporting_subject_label": reporting_subject_label,
        "max_items": int(options["max_chart_items"]),
        "small_multiples_max_panels": int(options["small_multiples_max_panels"]),
        "color_palette": (
            str(options["color_palette"]) if "color_palette" in options else "bain"
        ),
        "aggregate_other_items": True,
        "capture_chart_data": True,
        "display_value_prefix_metric": bubble_size,
        "display_value_prefixes": display_value_prefixes,
    }
    specs: list[dict[str, Any]] = []

    def add(
        name: str,
        plotter: str,
        legacy_chart_key: str,
        artifact_name: str,
        *,
        base_chart: str | None = None,
        **extra: Any,
    ) -> None:
        if (
            requested
            and name not in requested
            and (base_chart is None or base_chart not in requested)
        ):
            return
        spec = {
            "name": name,
            "plotter": plotter,
            "legacy_chart_key": legacy_chart_key,
            "artifact_name": artifact_name,
            **common,
            **extra,
        }
        if base_chart:
            spec["base_chart"] = base_chart
        specs.append(spec)

    add(
        "scatter",
        "plot_scatter_charts",
        "scatterChart",
        "scatter.png",
        dimension_selection="dot_dimension_total",
        show_iso_line=bool(isoline_metric),
        isoline_metric=isoline_metric,
        adjust_bubble_labels=True,
    )
    if options["small_multiples"] and small_multiples_dimension:
        add(
            "scatter_small_multiples",
            "plot_scatter_charts",
            "scatterChart",
            "scatter_small_multiples.png",
            base_chart="scatter",
            dot_dimension=dot_dimension,
            color_dimension=None,
            small_multiples_dimension=small_multiples_dimension,
            capture_figure="last",
            dimension_selection="dot_dimension_by_panel",
            show_iso_line=bool(isoline_metric),
            isoline_metric=isoline_metric,
            adjust_bubble_labels=True,
        )
    add(
        "bubble",
        "plot_bubble_charts",
        "bubbleChart",
        "bubble.png",
        capture_figure="last",
        dimension_selection="dot_dimension_total",
        color_dimension=bubble_color_dimension,
        plot_total_bubble=True,
        adjust_bubble_labels=True,
    )
    if options["small_multiples"] and small_multiples_dimension:
        add(
            "bubble_small_multiples",
            "plot_bubble_charts",
            "bubbleChart",
            "bubble_small_multiples.png",
            base_chart="bubble",
            dot_dimension=dot_dimension,
            color_dimension=bubble_color_dimension,
            small_multiples_dimension=small_multiples_dimension,
            capture_figure="last",
            dimension_selection="dot_dimension_by_panel",
            plot_total_bubble=True,
            adjust_bubble_labels=True,
        )
    return specs


def build_relationship_summary(
    canonical: pl.DataFrame, recipe: dict[str, Any]
) -> pl.DataFrame:
    """Return a compact deterministic summary table for interpretation."""

    mappings = recipe["mappings"]
    dot_dimension = str(mappings["dot_dimension"])
    metrics = list(
        dict.fromkeys(
            [
                str(mappings["x_metric_column"]),
                str(mappings["y_metric_column"]),
                str(mappings["bubble_size_metric_column"]),
            ]
        )
    )
    x_metric = str(mappings["x_metric_column"])
    y_metric = str(mappings["y_metric_column"])
    bubble_size_metric = str(mappings["bubble_size_metric_column"])
    canonical_columns = set(canonical.collect_schema().names())
    summary_metrics = [
        metric
        for metric in _canonical_metrics_for_requested_metrics(metrics)
        if metric in canonical_columns
    ]
    weighted_x_operands = _weighted_x_metric_operands(
        x_metric,
        y_metric,
        bubble_size_metric,
    )
    period_values = _period_values(canonical)
    selected_period = (
        CURRENT_PERIOD if CURRENT_PERIOD in period_values else period_values[-1]
    )
    summary_frame = canonical.filter(pl.col(CANONICAL_PERIOD) == selected_period)
    aggregated_metrics = [
        metric
        for metric in summary_metrics
        if weighted_x_operands is None
        or metric != x_metric
        or x_metric not in canonical_columns
    ]
    aggregations = [pl.col(metric).sum().alias(metric) for metric in aggregated_metrics]
    grouped = summary_frame.group_by(dot_dimension).agg(aggregations)
    if weighted_x_operands is not None and all(
        metric in grouped.collect_schema().names() for metric in weighted_x_operands
    ):
        numerator, denominator = weighted_x_operands
        grouped = grouped.with_columns(
            pl.when(pl.col(denominator).abs() > 0)
            .then(pl.col(numerator) / pl.col(denominator))
            .otherwise(0.0)
            .alias(x_metric)
        )
        summary_metrics = list(dict.fromkeys([*summary_metrics, x_metric]))
    size_metric = str(mappings["bubble_size_metric_column"])
    output_metrics = [
        metric
        for metric in list(dict.fromkeys([*metrics, *summary_metrics]))
        if metric in grouped.collect_schema().names()
    ]
    sort_metric = (
        size_metric
        if size_metric in grouped.collect_schema().names()
        else output_metrics[-1]
    )
    return grouped.select([dot_dimension, *output_metrics]).sort(
        sort_metric,
        descending=True,
    )


def build_summary_markdown(
    recipe: dict[str, Any],
    summary_table: pl.DataFrame,
    chart_audits: dict[str, Any],
) -> str:
    """Build deterministic markdown summary."""

    language = _output_language(recipe)
    copy = _OUTPUT_COPY[language]
    written = [
        name for name, audit in chart_audits.items() if audit["status"] == "written"
    ]
    data_only = [
        name
        for name, audit in chart_audits.items()
        if audit["status"] == "data_written"
    ]
    failed = [
        name
        for name, audit in chart_audits.items()
        if audit["status"] not in {"written", "data_written"}
    ]
    mappings = recipe["mappings"]
    lines = [
        f"# {copy['source_data_title']}",
        "",
        f"- {copy['source_file']}: `{recipe['source_file']}`",
        f"- {copy['x_metric']}: `{mappings['x_metric_column']}`",
        f"- {copy['y_metric']}: `{mappings['y_metric_column']}`",
        f"- {copy['size_metric']}: `{mappings['bubble_size_metric_column']}`",
        f"- {copy['dot_dimension']}: `{mappings['dot_dimension']}`",
        f"- {copy['charts_written']}: `{len(written)}`",
        f"- {copy['chart_candidates']}: `{len(data_only)}`",
        f"- {copy['attempts_not_written']}: `{len(failed)}`",
        "",
        f"## {copy['largest_dots']}",
        "",
    ]
    for row in summary_table.head(5).to_dicts():
        lines.append(f"- {row[mappings['dot_dimension']]}")
    lines.extend(
        [
            "",
            copy["pipeline_note"],
            "",
        ]
    )
    return "\n".join(lines)


def _format_value(value: Any) -> str:
    if isinstance(value, (int, float)):
        return f"{value:,.2f}"
    return str(value)


_OUTPUT_COPY: dict[str, dict[str, str]] = {
    "en": {
        "source_data_title": "Scatter & Bubble Source Data",
        "source_file": "Source file",
        "x_metric": "X metric",
        "y_metric": "Y metric",
        "size_metric": "Bubble size metric",
        "dot_dimension": "Dot dimension",
        "charts_written": "Legacy charts written",
        "chart_candidates": "Chart data candidates",
        "attempts_not_written": "Legacy chart attempts not written",
        "largest_dots": "Largest Dots",
        "pipeline_note": "Charts are generated by the vendored legacy charting pipeline. Failed attempts are listed in the audit rather than replaced by non-legacy redraws.",
        "report_title": "Scatter & Bubble Analysis",
        "comparison": "Charts compare `{x}` to `{y}`.",
        "comparison_docx": "Charts compare {x} to {y}.",
        "source_files": "Source Files",
        "charts": "Charts",
    },
    "es": {
        "source_data_title": "Datos fuente del análisis de dispersión y burbujas",
        "source_file": "Archivo fuente",
        "x_metric": "Métrica X",
        "y_metric": "Métrica Y",
        "size_metric": "Métrica de tamaño de burbuja",
        "dot_dimension": "Dimensión de los puntos",
        "charts_written": "Gráficos generados",
        "chart_candidates": "Candidatos con datos para gráficos",
        "attempts_not_written": "Intentos de gráficos no generados",
        "largest_dots": "Puntos principales",
        "pipeline_note": "Los gráficos se generan mediante el proceso de gráficos incorporado. Los intentos fallidos figuran en la auditoría y no se sustituyen por representaciones alternativas.",
        "report_title": "Análisis de dispersión y burbujas",
        "comparison": "Los gráficos comparan `{x}` con `{y}`.",
        "comparison_docx": "Los gráficos comparan {x} con {y}.",
        "source_files": "Archivos fuente",
        "charts": "Gráficos",
    },
}


def _output_language(recipe: dict[str, Any]) -> str:
    text = str(recipe.get("language") or "en").strip().lower().replace("_", "-")
    return "es" if text.split("-", 1)[0] == "es" else "en"


def write_client_report(
    recipe: dict[str, Any],
    summary_table: pl.DataFrame,
    chart_paths: list[str],
    output_dir: Path,
) -> tuple[list[str], dict[str, Any]]:
    """Write Markdown and DOCX client report."""

    copy = _OUTPUT_COPY[_output_language(recipe)]
    md_path = output_dir / "scatter_bubble_client_report.md"
    docx_path = output_dir / "scatter_bubble_client_report.docx"
    dot_dimension = str(recipe["mappings"]["dot_dimension"])
    size_metric = str(recipe["mappings"]["bubble_size_metric_column"])
    lines = [
        f"# {copy['report_title']}",
        "",
        copy["comparison"].format(
            x=recipe["mappings"]["x_metric_column"],
            y=recipe["mappings"]["y_metric_column"],
        ),
        "",
        f"## {copy['largest_dots']}",
        "",
    ]
    for row in summary_table.head(5).to_dicts():
        lines.append(
            f"- {row[dot_dimension]}: {_format_value(row[size_metric])} {size_metric}"
        )
    lines.extend(
        [
            "",
            f"## {copy['source_files']}",
            "",
            *[
                f"- `{Path(path).name}`"
                for path in chart_paths
                if Path(path).suffix.lower() == ".png"
            ],
            "",
        ]
    )
    md_path.write_text("\n".join(lines), encoding="utf-8")
    try:
        from docx import Document
        from docx.shared import Inches

        document = Document()
        document.add_heading(copy["report_title"], level=1)
        document.add_paragraph(
            copy["comparison_docx"].format(
                x=recipe["mappings"]["x_metric_column"],
                y=recipe["mappings"]["y_metric_column"],
            )
        )
        document.add_heading(copy["largest_dots"], level=2)
        for row in summary_table.head(5).to_dicts():
            document.add_paragraph(
                f"{row[dot_dimension]}: {_format_value(row[size_metric])} {size_metric}",
                style="List Bullet",
            )
        document.add_heading(copy["charts"], level=2)
        for chart_path in chart_paths:
            path = Path(chart_path)
            if path.suffix.lower() == ".png" and path.exists():
                document.add_picture(str(path), width=Inches(6.4))
        document.save(docx_path)
        status = "written"
        error = None
    except (
        ImportError,
        ModuleNotFoundError,
        OSError,
        ValueError,
        zipfile.BadZipFile,
    ) as exc:
        status = "not_written"
        error = str(exc)
    paths = [str(md_path)]
    if docx_path.exists():
        paths.append(str(docx_path))
    return paths, {
        "status": status,
        "markdown": md_path.name,
        "docx": docx_path.name if docx_path.exists() else None,
        "error": error,
    }


def _relative_path(path: Path, base: Path) -> str:
    try:
        return str(path.relative_to(base))
    except ValueError:
        return str(path)


def _artifact_kind(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".png", ".html"}:
        return "chart"
    if suffix in {".csv", ".xlsx"}:
        return "table"
    if suffix == ".json":
        return "context"
    if suffix in {".md", ".docx"}:
        return "report"
    return "artifact"


def write_chart_context_artifacts(
    chart_name: str,
    chart_context: dict[str, Any],
    output_dir: Path,
) -> tuple[list[str], dict[str, Any]]:
    """Write model-readable chart data captured from the legacy chart path."""

    safe_name = "".join(
        character if character.isalnum() or character in {"_", "-"} else "_"
        for character in chart_name
    )
    context_path = output_dir / f"{safe_name}_chart_context.json"
    table_path = output_dir / f"{safe_name}_chart_data.csv"
    write_json(context_path, chart_context)
    paths = [str(context_path)]
    data_frame = chart_context["data_frame"]
    rows = (
        data_frame["rows"]
        if isinstance(data_frame, dict) and "rows" in data_frame
        else []
    )
    table_status = "not_written_no_rows"
    if rows:
        pl.DataFrame(rows).write_csv(table_path)
        paths.append(str(table_path))
        table_status = "written"
    return paths, {
        "status": "written",
        "context_path": context_path.name,
        "table_path": table_path.name if table_path.exists() else None,
        "table_status": table_status,
        "source": chart_context["chart_data_source"],
    }


def _normalize_artifact_mode(artifact_mode: str) -> str:
    """Return a supported artifact mode or raise for invalid contract input."""

    normalized = str(artifact_mode or ARTIFACT_MODE_DATA_AND_RENDER).strip().lower()
    if normalized not in ARTIFACT_MODES:
        allowed = ", ".join(sorted(ARTIFACT_MODES))
        raise ValueError(f"Unsupported artifact_mode {artifact_mode!r}; use {allowed}.")
    return normalized


def _chart_artifact_id(spec: dict[str, Any]) -> str:
    """Return the chart artifact ID for generated chart data."""

    artifact_stem = Path(str(spec.get("artifact_name") or "")).stem
    if artifact_stem.endswith("_small_multiples"):
        return artifact_stem
    return str(spec.get("name") or artifact_stem)


def _spec_for_data_only_artifacts(spec: dict[str, Any]) -> dict[str, Any]:
    """Force chart-data capture for data-only artifact generation."""

    return {**spec, "capture_chart_data": True}


def _slim_data_only_context(chart_context: dict[str, Any]) -> dict[str, Any]:
    """Remove render-heavy payloads from data-only chart context."""

    return {
        key: value
        for key, value in chart_context.items()
        if key not in {"plotly_figures", "exports"}
    }


def run_scatter_bubble(
    input_path: Path,
    output_dir: Path,
    recipe_path: Path | None = None,
    *,
    language: str = "en",
    currency: str | None = None,
    artifact_mode: str = ARTIFACT_MODE_DATA_AND_RENDER,
) -> ScatterBubbleRunResult:
    """Run the full scatter-bubble workflow."""

    artifact_mode = _normalize_artifact_mode(artifact_mode)
    frame = read_table(input_path)
    existing_recipe = read_json(recipe_path)
    recipe = build_recipe(
        input_path, frame, language=language, existing_recipe=existing_recipe
    )
    recipe = preserve_recipe_filters(recipe, existing_recipe)
    recipe = preserve_recipe_cohorts(recipe, existing_recipe)
    if currency:
        recipe["options"]["currency"] = currency
    output_dir.mkdir(parents=True, exist_ok=True)
    try:
        frame, filter_audit = apply_recipe_filters(frame, recipe)
        recipe.setdefault("options", {})["recipe_filter_audit"] = filter_audit
        run_intake = write_run_intake(
            output_dir,
            input_path,
            recipe_path=recipe_path,
            recipe=recipe,
            source_row_count=frame.height,
        )
        canonical, preparation_audit = prepare_canonical_frame(frame, recipe)
        preparation_audit["recipe_filters"] = filter_audit
        current_period, previous_period = recipe_cohort_period_labels(
            recipe,
            default_current=CURRENT_PERIOD,
            default_previous="PY",
        )
        canonical, cohort_audit = apply_recipe_cohorts(
            canonical,
            recipe,
            period_column=CANONICAL_PERIOD,
            value_column=str(recipe["mappings"]["bubble_size_metric_column"]),
            current_period=current_period,
            previous_period=previous_period,
        )
        preparation_audit["recipe_cohorts"] = cohort_audit
        specs = build_chart_specs(canonical, recipe)
        prepared_data_cache = LegacyPreparedDataCache.empty()
        render_charts = artifact_mode != ARTIFACT_MODE_DATA_ONLY
        artifact_paths: list[str] = []
        chart_audits: dict[str, Any] = {}
        chart_context_artifacts: dict[str, Any] = {}
        for spec in specs:
            run_spec = (
                _spec_for_data_only_artifacts(spec)
                if artifact_mode == ARTIFACT_MODE_DATA_ONLY
                else spec
            )
            export = write_legacy_scatter_bubble_chart(
                canonical,
                recipe,
                output_dir,
                run_spec,
                prepared_data_cache=prepared_data_cache,
                render=render_charts,
            )
            artifact_paths.extend(export.paths)
            chart_audits[str(spec["name"])] = export.audit
            if export.chart_context:
                chart_context = export.chart_context
                if artifact_mode == ARTIFACT_MODE_DATA_ONLY:
                    chart_context = _slim_data_only_context(chart_context)
                context_paths, context_audit = write_chart_context_artifacts(
                    str(spec["name"]), chart_context, output_dir
                )
                artifact_paths.extend(context_paths)
                chart_context_artifacts[str(spec["name"])] = context_audit
                chart_audits[str(spec["name"])]["chart_context"] = context_audit
        summary_table = build_relationship_summary(canonical, recipe)
        canonical_path = output_dir / "scatter_bubble_canonical.csv"
        summary_path = output_dir / "scatter_bubble_summary.csv"
        canonical.write_csv(canonical_path)
        prepared_manifest_path = write_prepared_data_manifest(
            output_dir=output_dir,
            plugin="scatter-bubble-analysis",
            chart_family="scatter_bubble",
            source_file=input_path,
            prepared_path=canonical_path,
            frame=canonical,
            recipe=recipe,
            preparation_audit=preparation_audit,
        )
        summary_table.write_csv(summary_path)
        artifact_paths.extend(
            [str(canonical_path), str(prepared_manifest_path), str(summary_path)]
        )
        try:
            summary_table.write_excel(output_dir / "scatter_bubble_results.xlsx")
            xlsx_status = "written"
            artifact_paths.append(str(output_dir / "scatter_bubble_results.xlsx"))
        except (ImportError, ModuleNotFoundError, OSError, ValueError) as exc:
            xlsx_status = f"not_written: {exc}"
        context = {
            "schema_version": SCHEMA_VERSION,
            "source_file": str(input_path),
            "artifact_mode": artifact_mode,
            "mappings": recipe["mappings"],
            "options": recipe["options"],
            "prepared_data_manifest": prepared_manifest_path.name,
            "chart_audits": chart_audits,
            "chart_context_artifacts": chart_context_artifacts,
        }
        context_path = output_dir / "scatter_bubble_context.json"
        write_json(context_path, context)
        artifact_paths.append(str(context_path))
        summary_markdown = build_summary_markdown(recipe, summary_table, chart_audits)
        summary_md_path = output_dir / "scatter_bubble_summary.md"
        summary_md_path.write_text(summary_markdown, encoding="utf-8")
        artifact_paths.append(str(summary_md_path))
        report_paths, report_audit = write_client_report(
            recipe, summary_table, artifact_paths, output_dir
        )
        artifact_paths.extend(report_paths)
        used_recipe_path = output_dir / "used_recipe.json"
        write_json(used_recipe_path, recipe)
        artifact_paths.append(str(used_recipe_path))
        audit: dict[str, Any] = {
            "schema_version": SCHEMA_VERSION,
            "created_at": utc_now(),
            "input_file": str(input_path),
            "recipe": recipe,
            "legacy_runtime": {
                "preparation": preparation_audit,
                "chart_audits": chart_audits,
                "chart_context_artifacts": chart_context_artifacts,
                "client_report": report_audit,
                "artifact_mode": artifact_mode,
            },
            "checks": {
                "canonical_row_count": canonical.height,
                "summary_row_count": summary_table.height,
                "legacy_chart_attempt_count": len(chart_audits),
                "legacy_chart_written_count": sum(
                    1 for item in chart_audits.values() if item["status"] == "written"
                ),
                "legacy_chart_data_count": sum(
                    1
                    for item in chart_audits.values()
                    if item["status"] in {"written", "data_written"}
                ),
            },
            "outputs": {
                "scatter_bubble_canonical.csv": "written",
                "prepared_data_manifest.json": "written",
                "scatter_bubble_summary.csv": "written",
                "scatter_bubble_results.xlsx": xlsx_status,
                "scatter_bubble_context.json": "written",
                "scatter_bubble_summary.md": "written",
                "used_recipe.json": "written",
            },
        }
        for path in artifact_paths:
            audit["outputs"][_relative_path(Path(path), output_dir)] = "written"
        audit_path = output_dir / "scatter_bubble_audit.json"
        write_json(audit_path, audit)
        audit["outputs"]["scatter_bubble_audit.json"] = "written"
        review_session = write_review_session_artifacts(
            output_dir,
            input_path,
            run_id=run_intake.run_id,
            run_intake_path=run_intake.path,
            recipe_path=recipe_path,
            recipe=recipe,
            summary_rows=summary_table.to_dicts(),
            audit=audit,
        )
        audit["review_session"] = {
            "run_intake_path": review_session.run_intake_path.name,
            "review_payload_path": review_session.review_payload_path.name,
            "ui_decisions_path": review_session.ui_decisions_path.name,
            "review_html_path": review_session.review_html_path.name,
            "final_artifacts_path": review_session.final_artifacts_path.name,
            "review_item_count": review_session.review_item_count,
        }
        for path in (
            review_session.run_intake_path,
            review_session.review_payload_path,
            review_session.ui_decisions_path,
            review_session.review_html_path,
            review_session.final_artifacts_path,
        ):
            audit["outputs"][_relative_path(path, output_dir)] = "written"
        artifact_paths = [
            *artifact_paths,
            str(review_session.run_intake_path),
            str(review_session.review_payload_path),
            str(review_session.ui_decisions_path),
            str(review_session.review_html_path),
            str(review_session.final_artifacts_path),
        ]
        write_json(audit_path, audit)
        return ScatterBubbleRunResult(
            canonical_frame=canonical,
            audit=audit,
            summary_markdown=summary_markdown,
            artifact_paths=artifact_paths,
        )
    finally:
        cleanup_legacy_imports()


if __name__ == "__main__":
    raise SystemExit("Use run_scatter_bubble.py or inspect_inputs.py")
