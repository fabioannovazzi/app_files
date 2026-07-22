from __future__ import annotations

import argparse
import csv
import json
import logging
import re
import sys
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, Sequence

import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from .review_session import write_review_session_artifacts, write_run_intake
except ImportError:  # pragma: no cover - supports direct script imports
    import importlib.util

    _review_session_path = Path(__file__).resolve().parent / "review_session.py"
    _review_session_spec = importlib.util.spec_from_file_location(
        "mparanza_report_builder_review_session",
        _review_session_path,
    )
    assert _review_session_spec and _review_session_spec.loader
    _review_session = importlib.util.module_from_spec(_review_session_spec)
    sys.modules[_review_session_spec.name] = _review_session
    _review_session_spec.loader.exec_module(_review_session)
    write_review_session_artifacts = _review_session.write_review_session_artifacts
    write_run_intake = _review_session.write_run_intake

LOGGER = logging.getLogger(__name__)

SUPPORTED_LANGUAGES = ("it", "en", "fr", "de", "es")
SUPPORTED_DOCUMENT_LANGUAGES = ("auto", *SUPPORTED_LANGUAGES)
SUPPORTED_SUFFIXES = {".csv", ".pdf", ".xlsx", ".xlsm", ".zip"}
TEXT_SUFFIXES = {".pdf"}
WORKBOOK_SUFFIXES = {".xlsx", ".xlsm"}

REPORT_TYPES: dict[str, dict[str, Any]] = {
    "management_report": {
        "label": {
            "en": "Management report",
            "it": "Report gestionale",
            "fr": "Rapport de gestion",
            "de": "Managementbericht",
            "es": "Informe de gestión",
        },
        "sections": [
            "overview",
            "income_statement",
            "balance_sheet",
            "cash_flow",
            "budget",
            "debt",
            "investments",
            "taxes",
            "notes",
        ],
    },
    "local_government_review": {
        "label": {
            "en": "Local government review",
            "it": "Relazione ente locale",
            "fr": "Revue collectivite locale",
            "de": "Kommunaler Pruefbericht",
            "es": "Informe de entidad local",
        },
        "sections": [
            "overview",
            "fpv",
            "fcde",
            "debt",
            "cash",
            "taxes",
            "spending",
            "investments",
            "participations",
            "pnrr",
            "notes",
        ],
    },
    "annual_financial_statement": {
        "label": {
            "en": "Annual financial statement",
            "it": "Bilancio annuale",
            "fr": "Etats financiers annuels",
            "de": "Jahresabschluss",
            "es": "Estados financieros anuales",
        },
        "sections": [
            "overview",
            "balance_sheet",
            "income_statement",
            "cash_flow",
            "equity",
            "ratios",
            "segment",
            "debt",
            "capex",
            "notes",
        ],
    },
}

SECTION_TITLES: dict[str, dict[str, str]] = {
    "overview": {
        "en": "Overview",
        "it": "Sintesi",
        "fr": "Synthese",
        "de": "Ueberblick",
        "es": "Resumen",
    },
    "income_statement": {
        "en": "Income statement",
        "it": "Conto economico",
        "fr": "Compte de resultat",
        "de": "Gewinn- und Verlustrechnung",
        "es": "Cuenta de resultados",
    },
    "balance_sheet": {
        "en": "Balance sheet",
        "it": "Stato patrimoniale",
        "fr": "Bilan",
        "de": "Bilanz",
        "es": "Balance",
    },
    "cash_flow": {
        "en": "Cash flow",
        "it": "Rendiconto finanziario",
        "fr": "Flux de tresorerie",
        "de": "Cashflow",
        "es": "Flujo de caja",
    },
    "budget": {
        "en": "Budget",
        "it": "Budget",
        "fr": "Budget",
        "de": "Budget",
        "es": "Presupuesto",
    },
    "debt": {
        "en": "Debt",
        "it": "Debito",
        "fr": "Dette",
        "de": "Schulden",
        "es": "Deuda",
    },
    "investments": {
        "en": "Investments",
        "it": "Investimenti",
        "fr": "Investissements",
        "de": "Investitionen",
        "es": "Inversiones",
    },
    "taxes": {
        "en": "Taxes",
        "it": "Tributi",
        "fr": "Fiscalite",
        "de": "Steuern",
        "es": "Impuestos",
    },
    "notes": {
        "en": "Notes",
        "it": "Note",
        "fr": "Notes",
        "de": "Anmerkungen",
        "es": "Notas",
    },
    "fpv": {
        "en": "FPV",
        "it": "FPV",
        "fr": "FPV",
        "de": "FPV",
        "es": "FPV",
    },
    "fcde": {
        "en": "FCDE",
        "it": "FCDE",
        "fr": "FCDE",
        "de": "FCDE",
        "es": "FCDE",
    },
    "cash": {
        "en": "Cash",
        "it": "Cassa",
        "fr": "Tresorerie",
        "de": "Liquiditaet",
        "es": "Tesorería",
    },
    "spending": {
        "en": "Spending",
        "it": "Spesa",
        "fr": "Depenses",
        "de": "Ausgaben",
        "es": "Gasto",
    },
    "participations": {
        "en": "Participations",
        "it": "Partecipazioni",
        "fr": "Participations",
        "de": "Beteiligungen",
        "es": "Participaciones",
    },
    "pnrr": {
        "en": "PNRR",
        "it": "PNRR",
        "fr": "PNRR",
        "de": "PNRR",
        "es": "PNRR",
    },
    "equity": {
        "en": "Equity",
        "it": "Patrimonio netto",
        "fr": "Capitaux propres",
        "de": "Eigenkapital",
        "es": "Patrimonio neto",
    },
    "ratios": {
        "en": "Ratios",
        "it": "Indicatori",
        "fr": "Ratios",
        "de": "Kennzahlen",
        "es": "Ratios",
    },
    "segment": {
        "en": "Segment information",
        "it": "Informativa per settore",
        "fr": "Information sectorielle",
        "de": "Segmentinformationen",
        "es": "Información por segmentos",
    },
    "capex": {
        "en": "Capital expenditure",
        "it": "Investimenti tecnici",
        "fr": "Depenses d'investissement",
        "de": "Investitionsausgaben",
        "es": "Inversiones de capital",
    },
}

SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "overview": (
        "overview",
        "sintesi",
        "summary",
        "resume",
        "ueberblick",
        "resumen",
        "síntesis",
        "sintesis",
    ),
    "income_statement": (
        "income",
        "profit",
        "loss",
        "conto economico",
        "ricavi",
        "costi",
        "resultat",
        "guv",
        "cuenta de resultados",
        "pérdidas y ganancias",
        "perdidas y ganancias",
    ),
    "balance_sheet": (
        "balance",
        "stato patrimoniale",
        "attivo",
        "passivo",
        "bilan",
        "bilanz",
        "balance",
        "situación financiera",
        "situacion financiera",
    ),
    "cash_flow": (
        "cash flow",
        "rendiconto",
        "flussi",
        "tresorerie",
        "kapitalfluss",
        "flujo de caja",
        "flujos de efectivo",
    ),
    "budget": (
        "budget",
        "forecast",
        "prevision",
        "preventivo",
        "planung",
        "presupuesto",
        "previsión",
    ),
    "debt": (
        "debt",
        "debito",
        "mutui",
        "loans",
        "dette",
        "schulden",
        "deuda",
        "préstamos",
        "prestamos",
    ),
    "investments": (
        "investment",
        "investimenti",
        "capex",
        "immobilizzazioni",
        "investissements",
        "inversiones",
    ),
    "taxes": (
        "tax",
        "taxes",
        "imposte",
        "tributi",
        "fiscal",
        "steuern",
        "impuestos",
        "tributos",
    ),
    "notes": (
        "note",
        "notes",
        "comment",
        "commenti",
        "annexe",
        "anhang",
        "notas",
        "anexo",
    ),
    "fpv": ("fpv", "fondo pluriennale"),
    "fcde": ("fcde", "crediti dubbia", "doubtful", "creances douteuses"),
    "cash": (
        "cash",
        "cassa",
        "tesoreria",
        "banque",
        "liquiditaet",
        "tesorería",
        "efectivo",
    ),
    "spending": ("spesa", "spending", "depenses", "ausgaben", "gasto"),
    "participations": (
        "partecipazioni",
        "participations",
        "beteiligungen",
        "subsidiaries",
        "participaciones",
    ),
    "pnrr": ("pnrr", "rrf", "recovery"),
    "equity": (
        "equity",
        "patrimonio netto",
        "capitaux propres",
        "eigenkapital",
        "patrimonio neto",
    ),
    "ratios": ("ratio", "indicatori", "indici", "kennzahlen", "ratios"),
    "segment": ("segment", "settore", "sector", "secteur", "segmento"),
    "capex": (
        "capex",
        "capital expenditure",
        "investimenti tecnici",
        "inversiones de capital",
    ),
}

DOCX_COPY: dict[str, dict[str, str]] = {
    "entity": {
        "en": "Entity",
        "it": "Ente",
        "fr": "Entite",
        "de": "Einheit",
        "es": "Entidad",
    },
    "period": {
        "en": "Period",
        "it": "Periodo",
        "fr": "Periode",
        "de": "Zeitraum",
        "es": "Periodo",
    },
    "entity_pending": {
        "en": "Entity pending",
        "it": "Ente da definire",
        "fr": "Entite a definir",
        "de": "Einheit noch offen",
        "es": "Entidad pendiente",
    },
    "period_pending": {
        "en": "Period pending",
        "it": "Periodo da definire",
        "fr": "Periode a definir",
        "de": "Zeitraum noch offen",
        "es": "Periodo pendiente",
    },
    "executive_summary": {
        "en": "Executive summary",
        "it": "Sintesi",
        "fr": "Synthese",
        "de": "Zusammenfassung",
        "es": "Resumen ejecutivo",
    },
    "executive_summary_pending": {
        "en": "Codex executive summary pending.",
        "it": "Sintesi Codex in attesa.",
        "fr": "Synthese Codex en attente.",
        "de": "Codex-Zusammenfassung noch offen.",
        "es": "Resumen ejecutivo de Codex pendiente.",
    },
    "context": {
        "en": "Context",
        "it": "Contesto",
        "fr": "Contexte",
        "de": "Kontext",
        "es": "Contexto",
    },
    "source": {
        "en": "Source",
        "it": "Fonte",
        "fr": "Source",
        "de": "Quelle",
        "es": "Fuente",
    },
    "rows": {
        "en": "Rows",
        "it": "Righe",
        "fr": "Lignes",
        "de": "Zeilen",
        "es": "Filas",
    },
    "columns": {
        "en": "Columns",
        "it": "Colonne",
        "fr": "Colonnes",
        "de": "Spalten",
        "es": "Columnas",
    },
    "numeric_totals": {
        "en": "Deterministic numeric totals",
        "it": "Totali numerici deterministici",
        "fr": "Totaux numeriques deterministes",
        "de": "Deterministische numerische Summen",
        "es": "Totales numéricos deterministas",
    },
    "count": {
        "en": "count",
        "it": "conteggio",
        "fr": "nombre",
        "de": "Anzahl",
        "es": "recuento",
    },
    "sum": {
        "en": "sum",
        "it": "somma",
        "fr": "somme",
        "de": "Summe",
        "es": "suma",
    },
    "table_preview": {
        "en": "Table preview",
        "it": "Anteprima tabella",
        "fr": "Apercu du tableau",
        "de": "Tabellenvorschau",
        "es": "Vista previa de la tabla",
    },
    "unassigned": {
        "en": "No table assigned yet. Codex review pending for this section.",
        "it": "Nessuna tabella assegnata. Revisione Codex in attesa per questa sezione.",
        "fr": "Aucun tableau assigne. Revue Codex en attente pour cette section.",
        "de": "Noch keine Tabelle zugeordnet. Codex-Pruefung fuer diesen Abschnitt offen.",
        "es": "Todavía no hay una tabla asignada. La revisión de Codex está pendiente para esta sección.",
    },
    "codex_pending": {
        "en": "Codex review pending for this section.",
        "it": "Revisione Codex in attesa per questa sezione.",
        "fr": "Revue Codex en attente pour cette section.",
        "de": "Codex-Pruefung fuer diesen Abschnitt offen.",
        "es": "La revisión de Codex está pendiente para esta sección.",
    },
    "audit_appendix": {
        "en": "Audit appendix",
        "it": "Appendice audit",
        "fr": "Annexe d'audit",
        "de": "Audit-Anhang",
        "es": "Anexo de auditoría",
    },
    "report_status": {
        "en": "Report status",
        "it": "Stato report",
        "fr": "Statut du rapport",
        "de": "Berichtsstatus",
        "es": "Estado del informe",
    },
    "assigned_sections": {
        "en": "Assigned sections",
        "it": "Sezioni assegnate",
        "fr": "Sections assignees",
        "de": "Zugeordnete Abschnitte",
        "es": "Secciones asignadas",
    },
    "missing_sections": {
        "en": "Missing sections",
        "it": "Sezioni mancanti",
        "fr": "Sections manquantes",
        "de": "Fehlende Abschnitte",
        "es": "Secciones pendientes",
    },
    "model_api_calls": {
        "en": "Model API calls from scripts",
        "it": "Chiamate API modello dagli script",
        "fr": "Appels API modele par les scripts",
        "de": "Modell-API-Aufrufe aus Skripten",
        "es": "Llamadas a la API del modelo desde los scripts",
    },
    "draft": {
        "en": "Draft generated by deterministic scripts and Codex-guided narrative.",
        "it": "Bozza generata da script deterministici e narrativa guidata da Codex.",
        "fr": "Brouillon genere par scripts deterministes et narration guidee par Codex.",
        "de": "Entwurf durch deterministische Skripte und Codex-gefuehrte Narrative erstellt.",
        "es": "Borrador generado por scripts deterministas y narrativa guiada por Codex.",
    },
    "notes": {
        "en": "Notes",
        "it": "Note",
        "fr": "Notes",
        "de": "Anmerkungen",
        "es": "Notas",
    },
    "input_path": {
        "en": "Input path",
        "it": "Percorso di input",
        "fr": "Chemin d'entree",
        "de": "Eingabepfad",
        "es": "Ruta de entrada",
    },
    "tables_discovered": {
        "en": "Tables discovered",
        "it": "Tabelle rilevate",
        "fr": "Tableaux detectes",
        "de": "Erkannte Tabellen",
        "es": "Tablas detectadas",
    },
    "column": {
        "en": "Column",
        "it": "Colonna",
        "fr": "Colonne",
        "de": "Spalte",
        "es": "Columna",
    },
    "range": {
        "en": "Range",
        "it": "Intervallo",
        "fr": "Plage",
        "de": "Spannweite",
        "es": "Rango",
    },
    "generated_by": {
        "en": "Generated by the Build Report Codex plugin.",
        "it": "Generato dal plugin Codex Build Report.",
        "fr": "Genere par le plugin Codex Build Report.",
        "de": "Vom Codex-Plugin Build Report erstellt.",
        "es": "Generado por el plugin Codex Build Report.",
    },
}

AMOUNT_RE = re.compile(r"^\(?-?\d{1,3}(?:[.,\s]\d{3})*(?:[.,]\d+)?\)?$")
WORD_RE = re.compile(r"[a-z0-9]+")

__all__ = [
    "BuildResult",
    "InspectionResult",
    "add_common_args",
    "build_report",
    "configure_logging",
    "inspect_inputs",
    "normalize_language",
    "write_json",
]


@dataclass(frozen=True)
class InspectionResult:
    """Inspection outputs for the report-builder workflow."""

    inspection: dict[str, Any]
    suggested_recipe: dict[str, Any]


@dataclass(frozen=True)
class BuildResult:
    """Report build outputs and audit metadata."""

    analysis: dict[str, Any]
    audit: dict[str, Any]
    markdown_path: Path
    docx_path: Path
    review_session: dict[str, Any] | None = None


def configure_logging(verbose: bool = False) -> None:
    """Configure script logging without affecting imported use."""

    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(level=level, format="%(levelname)s: %(message)s")


def normalize_language(
    language: object | None,
    *,
    default: str = "en",
    allow_auto: bool = False,
) -> str:
    """Normalize a language tag to one supported plugin locale."""

    text = str(language or default).strip().lower().replace("_", "-")
    code = text.split("-", 1)[0]
    if allow_auto and code == "auto":
        return "auto"
    return code if code in SUPPORTED_LANGUAGES else default


def normalize_report_type(report_type: object | None) -> str:
    """Return a supported report type key."""

    value = str(report_type or "management_report").strip().lower().replace("-", "_")
    return value if value in REPORT_TYPES else "management_report"


def language_assumptions(
    recipe: dict[str, Any] | None = None,
    *,
    language: object | None = None,
    document_language: object | None = None,
) -> dict[str, str]:
    """Resolve working and source-document language assumptions."""

    recipe = recipe or {}
    working = normalize_language(language or recipe.get("language"), default="en")
    source = normalize_language(
        document_language or recipe.get("document_language") or "auto",
        default=working,
        allow_auto=True,
    )
    return {"language": working, "document_language": source}


def write_json(path: Path, payload: Any) -> None:
    """Write JSON with stable formatting."""

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def read_json(path: Path | None) -> dict[str, Any]:
    """Return a JSON object or an empty mapping when no file is provided."""

    if path is None:
        return {}
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"Recipe must be a JSON object: {path}")
    return payload


def clean_text(value: Any) -> str:
    """Normalize a cell value for display and matching."""

    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = str(value).replace("\u00a0", " ").replace("\u202f", " ")
    return re.sub(r"\s+", " ", text).strip()


def norm_label(value: Any) -> str:
    """Return an ASCII-ish lower-case label for heuristics."""

    text = unicodedata.normalize("NFKD", clean_text(value))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower()
    return re.sub(r"\s+", " ", text)


def safe_sheet_name(name: str, fallback: str) -> str:
    """Return a valid Excel worksheet name."""

    cleaned = re.sub(r"[\[\]:*?/\\]", " ", name).strip() or fallback
    return cleaned[:31]


def parse_amount(value: Any) -> Decimal | None:
    """Parse common financial number formats into Decimal."""

    if value is None or isinstance(value, bool):
        return None
    if isinstance(value, (int, float, Decimal)):
        try:
            return Decimal(str(value))
        except InvalidOperation:
            return None
    text = clean_text(value)
    if not text:
        return None
    compact = text.replace(" ", "")
    negative = compact.startswith("(") and compact.endswith(")")
    compact = compact.strip("()")
    compact = compact.replace("+", "")
    compact = compact.replace("EUR", "").replace("€", "")
    compact = compact.strip()
    if not AMOUNT_RE.match(compact):
        return None
    if "," in compact and "." in compact:
        if compact.rfind(",") > compact.rfind("."):
            compact = compact.replace(".", "").replace(",", ".")
        else:
            compact = compact.replace(",", "")
    elif "," in compact:
        compact = compact.replace(".", "").replace(",", ".")
    try:
        number = Decimal(compact)
    except InvalidOperation:
        return None
    return -number if negative else number


def row_nonempty_count(row: Sequence[Any]) -> int:
    """Count non-empty display values in a row."""

    return sum(1 for value in row if clean_text(value))


def trim_rows(rows: Sequence[Sequence[Any]]) -> list[list[Any]]:
    """Drop trailing empty rows and columns from a rectangular-ish table."""

    materialized = [list(row) for row in rows]
    while materialized and row_nonempty_count(materialized[-1]) == 0:
        materialized.pop()
    max_width = 0
    for row in materialized:
        for idx, value in enumerate(row, start=1):
            if clean_text(value):
                max_width = max(max_width, idx)
    if max_width == 0:
        return []
    return [
        row[:max_width] + [""] * max(0, max_width - len(row)) for row in materialized
    ]


def read_csv_rows(path: Path) -> list[list[Any]]:
    """Read a CSV file with deterministic dialect fallback."""

    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")

    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel
    return trim_rows(csv.reader(text.splitlines(), dialect))


def read_workbook_tables(path: Path) -> list[dict[str, Any]]:
    """Read visible worksheets from an Excel workbook."""

    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    tables: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        if sheet.sheet_state != "visible":
            continue
        rows = trim_rows(sheet.iter_rows(values_only=True))
        tables.append(
            {
                "kind": "worksheet",
                "source_file": path.name,
                "source_path": str(path),
                "sheet_name": sheet.title,
                "table_id": f"{path.name}::{sheet.title}",
                "rows": rows,
            }
        )
    return tables


def read_pdf_text_table(path: Path) -> dict[str, Any]:
    """Extract text lines from a readable PDF."""

    import pdfplumber

    lines: list[str] = []
    page_count = 0
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines.extend(line.strip() for line in text.splitlines() if line.strip())
    rows = [[line] for line in lines]
    return {
        "kind": "pdf_text",
        "source_file": path.name,
        "source_path": str(path),
        "sheet_name": "",
        "table_id": path.name,
        "rows": rows,
        "page_count": page_count,
    }


def extract_zip(path: Path, output_dir: Path) -> Path:
    """Extract a ZIP into output_dir/extracted_inputs with path traversal checks."""

    destination = output_dir / "extracted_inputs" / path.stem
    destination.mkdir(parents=True, exist_ok=True)
    root = destination.resolve()
    with zipfile.ZipFile(path) as archive:
        for member in archive.infolist():
            if member.is_dir():
                continue
            target = (destination / member.filename).resolve()
            if root not in target.parents and target != root:
                raise ValueError(f"Unsafe ZIP member path: {member.filename}")
            target.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(member) as source, target.open("wb") as handle:
                handle.write(source.read())
    return destination


def discover_input_files(input_path: Path, output_dir: Path) -> list[Path]:
    """Return supported files from a file, folder, or ZIP archive."""

    path = input_path.expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(path)
    if path.is_file() and path.suffix.lower() == ".zip":
        path = extract_zip(path, output_dir)
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED_SUFFIXES - {".zip"} else []
    files = [
        item
        for item in path.rglob("*")
        if item.is_file() and item.suffix.lower() in SUPPORTED_SUFFIXES - {".zip"}
    ]
    return sorted(files, key=lambda item: item.as_posix().lower())


def load_tables(input_path: Path, output_dir: Path) -> list[dict[str, Any]]:
    """Load all inspectable tables/text blocks from the input path."""

    tables: list[dict[str, Any]] = []
    for file_path in discover_input_files(input_path, output_dir):
        suffix = file_path.suffix.lower()
        try:
            if suffix in WORKBOOK_SUFFIXES:
                tables.extend(read_workbook_tables(file_path))
            elif suffix == ".csv":
                tables.append(
                    {
                        "kind": "csv",
                        "source_file": file_path.name,
                        "source_path": str(file_path),
                        "sheet_name": "",
                        "table_id": file_path.name,
                        "rows": read_csv_rows(file_path),
                    }
                )
            elif suffix in TEXT_SUFFIXES:
                tables.append(read_pdf_text_table(file_path))
        except (OSError, ValueError, zipfile.BadZipFile) as exc:
            LOGGER.warning("Could not inspect %s: %s", file_path, exc)
            tables.append(
                {
                    "kind": "error",
                    "source_file": file_path.name,
                    "source_path": str(file_path),
                    "sheet_name": "",
                    "table_id": file_path.name,
                    "rows": [],
                    "error": str(exc),
                }
            )
    return tables


def header_candidate_index(rows: Sequence[Sequence[Any]]) -> int | None:
    """Pick a likely header row using text density and position."""

    best_idx: int | None = None
    best_score = -1.0
    for idx, row in enumerate(rows[:20]):
        nonempty = row_nonempty_count(row)
        if nonempty == 0:
            continue
        labels = " ".join(norm_label(value) for value in row if clean_text(value))
        label_hits = sum(
            1
            for token in (
                "year",
                "period",
                "amount",
                "value",
                "totale",
                "importo",
                "descrizione",
                "budget",
                "actual",
                "saldo",
                "conto",
            )
            if token in labels
        )
        text_cells = sum(
            1 for value in row if clean_text(value) and parse_amount(value) is None
        )
        score = nonempty + label_hits * 2 + text_cells * 0.5 - idx * 0.1
        if score > best_score:
            best_idx = idx
            best_score = score
    return best_idx


def unique_headers(values: Sequence[Any], width: int) -> list[str]:
    """Return stable unique header labels."""

    headers: list[str] = []
    seen: dict[str, int] = {}
    for idx in range(width):
        base = clean_text(values[idx] if idx < len(values) else "")
        if not base:
            base = f"column_{idx + 1}"
        count = seen.get(base, 0)
        seen[base] = count + 1
        headers.append(base if count == 0 else f"{base}_{count + 1}")
    return headers


def rows_as_dicts(
    rows: Sequence[Sequence[Any]], header_idx: int | None
) -> list[dict[str, str]]:
    """Return display rows keyed by detected headers."""

    if not rows:
        return []
    width = max((len(row) for row in rows), default=0)
    if header_idx is None:
        headers = [f"column_{idx + 1}" for idx in range(width)]
        data_rows = rows
    else:
        headers = unique_headers(rows[header_idx], width)
        data_rows = rows[header_idx + 1 :]
    return [
        {
            headers[idx]: clean_text(row[idx] if idx < len(row) else "")
            for idx in range(width)
        }
        for row in data_rows
        if row_nonempty_count(row) > 0
    ]


def table_numeric_profile(
    rows: Sequence[Sequence[Any]], header_idx: int | None
) -> dict[str, Any]:
    """Summarize numeric columns without changing source values."""

    if not rows:
        return {"numeric_cells": 0, "numeric_columns": []}
    width = max((len(row) for row in rows), default=0)
    headers = (
        unique_headers(rows[header_idx], width)
        if header_idx is not None
        else [f"column_{idx + 1}" for idx in range(width)]
    )
    start_idx = header_idx + 1 if header_idx is not None else 0
    column_sums: list[dict[str, Any]] = []
    numeric_cells = 0
    for col_idx, header in enumerate(headers):
        values: list[Decimal] = []
        for row in rows[start_idx:]:
            value = parse_amount(row[col_idx] if col_idx < len(row) else None)
            if value is not None:
                values.append(value)
        numeric_cells += len(values)
        if values:
            total = sum(values, Decimal("0"))
            column_sums.append(
                {
                    "column": header,
                    "numeric_count": len(values),
                    "sum": float(total),
                    "min": float(min(values)),
                    "max": float(max(values)),
                }
            )
    return {"numeric_cells": numeric_cells, "numeric_columns": column_sums}


def suggest_section(table: dict[str, Any]) -> dict[str, Any]:
    """Suggest a report section for a table using transparent keyword matching."""

    rows = table.get("rows", [])
    preview_text = " ".join(
        clean_text(value)
        for row in rows[:12]
        for value in row[:10]
        if clean_text(value)
    )
    source_text = " ".join(
        [
            clean_text(table.get("source_file")),
            clean_text(table.get("sheet_name")),
            preview_text,
        ]
    )
    normalized = norm_label(source_text)
    best_section = ""
    best_score = 0
    for section, aliases in SECTION_ALIASES.items():
        score = sum(1 for alias in aliases if alias in normalized)
        if score > best_score:
            best_section = section
            best_score = score
    confidence = min(0.95, 0.35 + best_score * 0.2) if best_score else 0.0
    return {"section": best_section, "confidence": round(confidence, 2)}


def inspect_table(table: dict[str, Any]) -> dict[str, Any]:
    """Return a JSON-safe inspection record for one table/text block."""

    rows = table.get("rows", [])
    header_idx = header_candidate_index(rows)
    preview = rows_as_dicts(rows, header_idx)[:8]
    profile = table_numeric_profile(rows, header_idx)
    suggestion = suggest_section(table)
    return {
        "table_id": table.get("table_id", ""),
        "kind": table.get("kind", ""),
        "source_file": table.get("source_file", ""),
        "source_path": table.get("source_path", ""),
        "sheet_name": table.get("sheet_name", ""),
        "page_count": table.get("page_count", 0),
        "error": table.get("error", ""),
        "row_count": len([row for row in rows if row_nonempty_count(row) > 0]),
        "column_count": max((len(row) for row in rows), default=0),
        "header_row": header_idx + 1 if header_idx is not None else None,
        "numeric_cell_count": profile["numeric_cells"],
        "numeric_columns": profile["numeric_columns"][:12],
        "suggested_section": suggestion["section"],
        "suggestion_confidence": suggestion["confidence"],
        "preview_rows": preview,
    }


def section_title(section_key: str, language: str) -> str:
    """Return a locale-aware section title."""

    labels = SECTION_TITLES.get(section_key, {})
    return (
        labels.get(language)
        or labels.get("en")
        or section_key.replace("_", " ").title()
    )


def report_type_label(report_type: str, language: str) -> str:
    """Return a locale-aware report type label."""

    labels = REPORT_TYPES[report_type]["label"]
    return labels.get(language) or labels["en"]


def build_suggested_recipe(
    inspection: dict[str, Any],
    *,
    language: str,
    document_language: str,
    report_type: str,
) -> dict[str, Any]:
    """Build an editable recipe from inspection records."""

    assigned_tables: set[str] = set()
    sections: dict[str, Any] = {}
    table_records = inspection["tables"]
    for section_key in REPORT_TYPES[report_type]["sections"]:
        candidates = [
            table
            for table in table_records
            if table.get("suggested_section") == section_key
            and table.get("table_id") not in assigned_tables
        ]
        candidates.sort(
            key=lambda item: item.get("suggestion_confidence", 0), reverse=True
        )
        selected = candidates[0] if candidates else None
        assigned_table = selected["table_id"] if selected else ""
        if assigned_table:
            assigned_tables.add(assigned_table)
        sections[section_key] = {
            "title": section_title(section_key, language),
            "assigned_table": assigned_table,
            "codex_comment": "",
            "include_preview_rows": 8,
        }

    return {
        "version": 1,
        "language": language,
        "document_language": document_language,
        "report_type": report_type,
        "entity": "",
        "period": "",
        "executive_summary": "",
        "context_items": {},
        "sections": sections,
        "render": {
            "include_unassigned_tables": False,
            "include_table_previews": True,
        },
    }


def inspect_inputs(
    input_path: Path,
    output_dir: Path,
    *,
    language: object | None = None,
    document_language: object | None = None,
    report_type: object | None = None,
) -> InspectionResult:
    """Inspect report inputs and write inspection plus suggested recipe files."""

    output_dir.mkdir(parents=True, exist_ok=True)
    assumptions = language_assumptions(
        language=language,
        document_language=document_language,
    )
    report_key = normalize_report_type(report_type)
    raw_tables = load_tables(input_path, output_dir)
    tables = [inspect_table(table) for table in raw_tables]
    inspection = {
        "version": 1,
        "language": assumptions["language"],
        "document_language": assumptions["document_language"],
        "report_type": report_key,
        "report_type_label": report_type_label(report_key, assumptions["language"]),
        "input_path": str(input_path),
        "table_count": len(tables),
        "tables": tables,
        "supported_suffixes": sorted(SUPPORTED_SUFFIXES),
        "limitations": [
            "Scanned PDFs require OCR before deterministic extraction.",
            "Excel binary .xls files should be converted to .xlsx or CSV for this plugin version.",
        ],
    }
    recipe = build_suggested_recipe(
        inspection,
        language=assumptions["language"],
        document_language=assumptions["document_language"],
        report_type=report_key,
    )
    write_json(output_dir / "inspection.json", inspection)
    write_json(output_dir / "suggested_recipe.json", recipe)
    return InspectionResult(inspection=inspection, suggested_recipe=recipe)


def selected_sections(recipe: dict[str, Any]) -> dict[str, Any]:
    """Return recipe sections as a mapping."""

    sections = recipe.get("sections", {})
    if not isinstance(sections, dict):
        raise ValueError("Recipe sections must be a JSON object")
    return sections


def analysis_for_section(
    section_key: str,
    section_recipe: dict[str, Any],
    table_by_id: dict[str, dict[str, Any]],
) -> dict[str, Any]:
    """Build deterministic analysis for one report section."""

    table_id = clean_text(section_recipe.get("assigned_table"))
    table = table_by_id.get(table_id) if table_id else None
    if not table:
        return {
            "section": section_key,
            "title": clean_text(section_recipe.get("title")) or section_key,
            "assigned_table": table_id,
            "status": "unassigned",
            "row_count": 0,
            "column_count": 0,
            "numeric_columns": [],
            "preview_rows": [],
        }
    rows = table.get("rows", [])
    header_idx = header_candidate_index(rows)
    include_rows = int(section_recipe.get("include_preview_rows") or 8)
    profile = table_numeric_profile(rows, header_idx)
    return {
        "section": section_key,
        "title": clean_text(section_recipe.get("title")) or section_key,
        "assigned_table": table_id,
        "source_file": table.get("source_file", ""),
        "sheet_name": table.get("sheet_name", ""),
        "status": "assigned",
        "row_count": len([row for row in rows if row_nonempty_count(row) > 0]),
        "column_count": max((len(row) for row in rows), default=0),
        "header_row": header_idx + 1 if header_idx is not None else None,
        "numeric_columns": profile["numeric_columns"],
        "preview_rows": rows_as_dicts(rows, header_idx)[: max(0, include_rows)],
        "codex_comment": clean_text(section_recipe.get("codex_comment")),
    }


def markdown_table(rows: Sequence[dict[str, str]], *, max_rows: int = 8) -> str:
    """Render a compact Markdown table."""

    if not rows:
        return ""
    headers = list(rows[0].keys())[:8]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows[:max_rows]:
        values = [
            clean_text(row.get(header, "")).replace("|", "/") for header in headers
        ]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def render_markdown(recipe: dict[str, Any], analysis: dict[str, Any]) -> str:
    """Render a report draft in Markdown."""

    language = normalize_language(recipe.get("language"), default="en")
    report_type = normalize_report_type(recipe.get("report_type"))
    title = report_type_label(report_type, language)
    entity = clean_text(recipe.get("entity")) or docx_label("entity_pending", language)
    period = clean_text(recipe.get("period")) or docx_label("period_pending", language)
    executive_summary = clean_text(recipe.get("executive_summary"))
    if not executive_summary:
        executive_summary = docx_label("executive_summary_pending", language)

    lines = [
        f"# {title}",
        "",
        f"**{docx_label('entity', language)}:** {entity}",
        f"**{docx_label('period', language)}:** {period}",
        "",
        f"## {docx_label('executive_summary', language)}",
        "",
        executive_summary,
        "",
    ]

    context_items = recipe.get("context_items", {})
    if isinstance(context_items, dict) and context_items:
        lines.extend([f"## {docx_label('context', language)}", ""])
        for key, value in context_items.items():
            lines.append(f"- **{clean_text(key)}:** {clean_text(value)}")
        lines.append("")

    for section in analysis["sections"]:
        lines.extend([f"## {section['title']}", ""])
        if section["status"] != "assigned":
            lines.extend([docx_label("unassigned", language), ""])
            continue
        comment = section.get("codex_comment") or docx_label("codex_pending", language)
        lines.extend(
            [
                comment,
                "",
                f"{docx_label('source', language)}: {section.get('source_file', '')}"
                + (f" / {section['sheet_name']}" if section.get("sheet_name") else ""),
                f"{docx_label('rows', language)}: {section['row_count']} | "
                f"{docx_label('columns', language)}: {section['column_count']}",
            ]
        )
        numeric_columns = section.get("numeric_columns") or []
        if numeric_columns:
            lines.extend(["", f"{docx_label('numeric_totals', language)}:"])
            for column in numeric_columns[:8]:
                lines.append(
                    f"- {column['column']}: {docx_label('count', language)} "
                    f"{column['numeric_count']}, {docx_label('sum', language)} "
                    f"{column['sum']}"
                )
        preview_table = markdown_table(section.get("preview_rows") or [])
        if preview_table and recipe.get("render", {}).get(
            "include_table_previews", True
        ):
            lines.extend(["", preview_table])
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def docx_label(key: str, language: str) -> str:
    """Return a localized DOCX label."""

    labels = DOCX_COPY.get(key, {})
    return labels.get(language) or labels.get("en") or key.replace("_", " ").title()


def set_cell_shading(cell: Any, fill: str) -> None:
    """Set deterministic cell background shading."""

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, margin_twips: int = 80) -> None:
    """Set compact but readable cell margins."""

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_paragraph_font(
    paragraph: Any,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    """Apply simple font formatting to all runs in a paragraph."""

    for run in paragraph.runs:
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic


def set_docx_styles(document: Any) -> None:
    """Configure restrained report styles."""

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("334155")

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("243026")

    for style_name, size in (("Heading 1", 18), ("Heading 2", 13)):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("243026")


def add_small_note(document: Any, text: str) -> None:
    """Add a muted note paragraph."""

    paragraph = document.add_paragraph(text)
    set_paragraph_font(paragraph, size=9, color="667085")


def add_key_value_table(document: Any, rows: Sequence[tuple[str, Any]]) -> None:
    """Add a compact two-column metadata table."""

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = clean_text(label)
        cells[1].text = clean_text(value)
        set_cell_shading(cells[0], "E7ECE5")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, size=9.5)
        for paragraph in cells[0].paragraphs:
            set_paragraph_font(paragraph, bold=True, color="243026")


def add_dataframe_table(
    document: Any,
    rows: Sequence[dict[str, str]],
    *,
    max_rows: int = 8,
    max_columns: int = 6,
) -> None:
    """Add a real Word table for preview rows."""

    if not rows:
        return
    headers = list(rows[0].keys())[:max_columns]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = clean_text(header)
        set_cell_shading(header_cells[idx], "E7ECE5")
        set_cell_margins(header_cells[idx])
        for paragraph in header_cells[idx].paragraphs:
            set_paragraph_font(paragraph, size=8.5, bold=True, color="243026")

    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            value = clean_text(row.get(header, ""))
            cells[idx].text = value[:140] + ("..." if len(value) > 140 else "")
            set_cell_margins(cells[idx])
            for paragraph in cells[idx].paragraphs:
                set_paragraph_font(paragraph, size=8)


def add_numeric_totals_table(
    document: Any,
    numeric_columns: Sequence[dict[str, Any]],
    *,
    max_rows: int = 8,
    language: str = "en",
) -> None:
    """Add a small deterministic numeric summary table."""

    if not numeric_columns:
        return
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    headers = (
        docx_label("column", language),
        docx_label("count", language).capitalize(),
        docx_label("sum", language).capitalize(),
        docx_label("range", language),
    )
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        set_cell_shading(table.rows[0].cells[idx], "E7ECE5")
        set_cell_margins(table.rows[0].cells[idx])
        for paragraph in table.rows[0].cells[idx].paragraphs:
            set_paragraph_font(paragraph, size=8.5, bold=True, color="243026")
    for column in numeric_columns[:max_rows]:
        cells = table.add_row().cells
        cells[0].text = clean_text(column.get("column"))
        cells[1].text = clean_text(column.get("numeric_count"))
        cells[2].text = clean_text(column.get("sum"))
        cells[3].text = (
            f"{clean_text(column.get('min'))} - {clean_text(column.get('max'))}"
        )
        for cell in cells:
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, size=8)


def write_report_docx(
    recipe: dict[str, Any],
    analysis: dict[str, Any],
    audit: dict[str, Any],
    output_path: Path,
) -> None:
    """Write a styled DOCX report with real Word sections and tables."""

    language = normalize_language(recipe.get("language"), default="en")
    report_type = normalize_report_type(recipe.get("report_type"))
    title = report_type_label(report_type, language)
    entity = clean_text(recipe.get("entity")) or docx_label("entity_pending", language)
    period = clean_text(recipe.get("period")) or docx_label("period_pending", language)
    executive_summary = clean_text(recipe.get("executive_summary")) or docx_label(
        "codex_pending", language
    )

    document = Document()
    set_docx_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    document.core_properties.title = title
    document.core_properties.subject = f"{entity} - {period}"
    document.core_properties.comments = docx_label("generated_by", language)

    title_paragraph = document.add_paragraph(style="Title")
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = document.add_paragraph(docx_label("draft", language))
    set_paragraph_font(subtitle, size=10, color="667085")

    add_key_value_table(
        document,
        (
            (docx_label("entity", language), entity),
            (docx_label("period", language), period),
            (docx_label("report_status", language), audit.get("status", "draft")),
            (docx_label("model_api_calls", language), audit.get("model_api_calls", 0)),
        ),
    )

    document.add_heading(docx_label("executive_summary", language), level=1)
    document.add_paragraph(executive_summary)

    context_items = recipe.get("context_items", {})
    if isinstance(context_items, dict) and context_items:
        document.add_heading(docx_label("context", language), level=1)
        add_key_value_table(
            document,
            tuple(
                (clean_text(key), clean_text(value))
                for key, value in context_items.items()
            ),
        )

    for section_analysis in analysis.get("sections", []):
        document.add_heading(clean_text(section_analysis.get("title")), level=1)
        if section_analysis.get("status") != "assigned":
            paragraph = document.add_paragraph(docx_label("unassigned", language))
            set_paragraph_font(paragraph, italic=True, color="667085")
            continue

        comment = clean_text(section_analysis.get("codex_comment")) or docx_label(
            "codex_pending", language
        )
        document.add_paragraph(comment)
        source = clean_text(section_analysis.get("source_file"))
        sheet = clean_text(section_analysis.get("sheet_name"))
        source_value = source + (f" / {sheet}" if sheet else "")
        add_key_value_table(
            document,
            (
                (docx_label("source", language), source_value),
                (docx_label("rows", language), section_analysis.get("row_count", 0)),
                (
                    docx_label("columns", language),
                    section_analysis.get("column_count", 0),
                ),
            ),
        )

        numeric_columns = section_analysis.get("numeric_columns") or []
        if numeric_columns:
            document.add_heading(docx_label("numeric_totals", language), level=2)
            add_numeric_totals_table(document, numeric_columns, language=language)

        preview_rows = section_analysis.get("preview_rows") or []
        if preview_rows and recipe.get("render", {}).get(
            "include_table_previews", True
        ):
            document.add_heading(docx_label("table_preview", language), level=2)
            add_dataframe_table(document, preview_rows)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(docx_label("audit_appendix", language), level=1)
    add_key_value_table(
        document,
        (
            (docx_label("input_path", language), audit.get("input_path", "")),
            (docx_label("tables_discovered", language), audit.get("table_count", 0)),
            (
                docx_label("assigned_sections", language),
                audit.get("assigned_section_count", 0),
            ),
            (
                docx_label("missing_sections", language),
                audit.get("missing_section_count", 0),
            ),
            (docx_label("model_api_calls", language), audit.get("model_api_calls", 0)),
        ),
    )
    missing_sections = audit.get("missing_sections") or []
    if missing_sections:
        document.add_heading(docx_label("missing_sections", language), level=2)
        for missing in missing_sections:
            document.add_paragraph(clean_text(missing), style="List Bullet")
    notes = audit.get("notes") or []
    if notes:
        document.add_heading(docx_label("notes", language), level=2)
        for note in notes:
            add_small_note(document, clean_text(note))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def write_tables_workbook(output_path: Path, analysis: dict[str, Any]) -> None:
    """Write assigned table previews to an Excel workbook."""

    workbook = openpyxl.Workbook()
    default = workbook.active
    default.title = "summary"
    default.append(["section", "status", "assigned_table", "rows", "columns"])
    for section in analysis["sections"]:
        default.append(
            [
                section["section"],
                section["status"],
                section.get("assigned_table", ""),
                section.get("row_count", 0),
                section.get("column_count", 0),
            ]
        )
        rows = section.get("preview_rows") or []
        if rows:
            sheet = workbook.create_sheet(
                safe_sheet_name(
                    section["section"], f"section{len(workbook.worksheets)}"
                )
            )
            headers = list(rows[0].keys())
            sheet.append(headers)
            for row in rows:
                sheet.append([row.get(header, "") for header in headers])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)


def build_report(
    input_path: Path,
    output_dir: Path,
    *,
    recipe_path: Path | None = None,
    language: object | None = None,
    document_language: object | None = None,
    report_type: object | None = None,
) -> BuildResult:
    """Build report outputs from inspected files and an editable recipe."""

    output_dir.mkdir(parents=True, exist_ok=True)
    recipe = read_json(recipe_path)
    assumptions = language_assumptions(
        recipe,
        language=language,
        document_language=document_language,
    )
    if not isinstance(recipe.get("sections"), dict) or not recipe.get("sections"):
        inspection = inspect_inputs(
            input_path,
            output_dir,
            language=assumptions["language"],
            document_language=assumptions["document_language"],
            report_type=report_type,
        )
        recipe = inspection.suggested_recipe
    if language is not None:
        recipe["language"] = assumptions["language"]
    if document_language is not None:
        recipe["document_language"] = assumptions["document_language"]
    if report_type is not None:
        recipe["report_type"] = normalize_report_type(report_type)

    run_intake = write_run_intake(
        output_dir,
        input_path=input_path,
        recipe_path=recipe_path,
        language=str(recipe.get("language", assumptions["language"])),
        document_language=str(
            recipe.get("document_language", assumptions["document_language"])
        ),
        report_type=normalize_report_type(recipe.get("report_type")),
    )

    raw_tables = load_tables(input_path, output_dir)
    table_by_id = {clean_text(table.get("table_id")): table for table in raw_tables}
    sections_analysis = [
        analysis_for_section(section_key, section_recipe, table_by_id)
        for section_key, section_recipe in selected_sections(recipe).items()
    ]
    assigned_sections = [
        section for section in sections_analysis if section["status"] == "assigned"
    ]
    missing_sections = [
        section["section"]
        for section in sections_analysis
        if section["status"] != "assigned"
    ]
    analysis = {
        "version": 1,
        "language": recipe.get("language", assumptions["language"]),
        "document_language": recipe.get(
            "document_language", assumptions["document_language"]
        ),
        "report_type": normalize_report_type(recipe.get("report_type")),
        "entity": clean_text(recipe.get("entity")),
        "period": clean_text(recipe.get("period")),
        "sections": sections_analysis,
        "assigned_section_count": len(assigned_sections),
        "missing_sections": missing_sections,
    }
    table_inspection = [inspect_table(table) for table in raw_tables]
    report_language = str(recipe.get("language", assumptions["language"]))
    audit_notes = (
        [
            "El texto narrativo lo proporciona Codex en la receta, no los scripts auxiliares.",
            "Revise las secciones sin asignar y los comentarios pendientes de Codex antes del uso final.",
        ]
        if report_language == "es"
        else [
            "Narrative text is supplied by Codex in the recipe, not by helper scripts.",
            "Review unassigned sections and Codex-pending comments before final use.",
        ]
    )
    audit = {
        "version": 1,
        "status": "draft",
        "input_path": str(input_path),
        "table_count": len(raw_tables),
        "section_count": len(sections_analysis),
        "assigned_section_count": len(assigned_sections),
        "missing_section_count": len(missing_sections),
        "missing_sections": missing_sections,
        "codex_narrative_sections": sum(
            1 for section in sections_analysis if section.get("codex_comment")
        ),
        "model_api_calls": 0,
        "notes": audit_notes,
    }

    write_json(output_dir / "report_tables.json", {"tables": table_inspection})
    write_json(output_dir / "report_analysis.json", analysis)
    write_json(output_dir / "report_audit.json", audit)
    write_json(output_dir / "used_recipe.json", recipe)
    write_tables_workbook(output_dir / "report_tables.xlsx", analysis)
    markdown_text = render_markdown(recipe, analysis)
    markdown_path = output_dir / "report_draft.md"
    markdown_path.write_text(markdown_text, encoding="utf-8")
    docx_path = output_dir / "report.docx"
    write_report_docx(recipe, analysis, audit, docx_path)
    review_session = write_review_session_artifacts(
        output_dir,
        run_id=run_intake.run_id,
        run_intake_path=run_intake.path,
        analysis=analysis,
        audit=audit,
        recipe=recipe,
        paths={
            "report_tables": output_dir / "report_tables.json",
            "report_tables_xlsx": output_dir / "report_tables.xlsx",
            "report_analysis": output_dir / "report_analysis.json",
            "report_audit": output_dir / "report_audit.json",
            "used_recipe": output_dir / "used_recipe.json",
            "report_draft": markdown_path,
            "report_docx": docx_path,
        },
        tables=table_inspection,
    )
    audit["review_session"] = {
        "run_id": review_session.run_id,
        "run_intake": review_session.run_intake_path.name,
        "review_payload": review_session.review_payload_path.name,
        "ui_decisions": review_session.ui_decisions_path.name,
        "final_artifacts": review_session.final_artifacts_path.name,
        "review_item_count": review_session.review_item_count,
    }
    write_json(output_dir / "report_audit.json", audit)

    return BuildResult(
        analysis=analysis,
        audit=audit,
        markdown_path=markdown_path,
        docx_path=docx_path,
        review_session=audit["review_session"],
    )


def add_common_args(parser: argparse.ArgumentParser) -> None:
    """Add shared CLI arguments."""

    parser.add_argument("input_path", type=Path, help="Input file, folder, or ZIP")
    parser.add_argument(
        "--output-dir",
        type=Path,
        required=True,
        help="Directory where outputs are written",
    )
    parser.add_argument(
        "--language",
        default="en",
        choices=SUPPORTED_LANGUAGES,
        help="Working language for Codex and output labels",
    )
    parser.add_argument(
        "--document-language",
        default="auto",
        choices=SUPPORTED_DOCUMENT_LANGUAGES,
        help="Source-document language assumption",
    )
    parser.add_argument(
        "--report-type",
        default="management_report",
        choices=sorted(REPORT_TYPES),
        help="Report template to use",
    )
    parser.add_argument("--verbose", action="store_true", help="Enable debug logging")


def _json_default(value: Any) -> str:
    if isinstance(value, Path):
        return str(value)
    return str(value)


def print_json(payload: Any) -> None:
    """Print a JSON payload for CLI callers."""

    print(json.dumps(payload, ensure_ascii=False, indent=2, default=_json_default))
