#!/usr/bin/env python3
"""Package a source-reviewed INPS case into a professional-review draft."""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from acquisition_binding import (
    AcquisitionBindingError,
    build_acquisition_binding,
    compare_acquisition_bindings,
)
from case_core import (
    ensure_safe_output_dir,
    mark_private_file,
    prepare_private_directory,
    write_json,
    write_private_text,
)
from docx import Document
from docx.shared import Pt
from privacy_guard import privacy_issue, safe_identifier, safe_source_reference

__all__ = ["package_case", "main"]

LOGGER = logging.getLogger(__name__)
PLUGIN_ROOT = Path(__file__).resolve().parents[1]
VERDICTS = {
    "contradicted",
    "not_supported",
    "partially_supported",
    "supported",
    "uncertain",
}
SOURCE_TEMPORAL_ROLES = {
    "later_interpretive_authority",
    "period_rule",
    "research_cutoff_authority",
}
CLAIM_TYPES = {"calculation_basis", "case_application", "rule"}
ALLOWED_ACTIONS = [
    "accept",
    "reject",
    "edit",
    "mark_unclear",
    "request_more_documents",
    "skip",
]
PACKAGE_GENERATED_ARTIFACTS = (
    "applied_decisions.json",
    "blocked_case_note.md",
    "claims_review_normalized.json",
    "document_requests.md",
    "final_artifacts.json",
    "review_handoff.md",
    "review_payload.json",
    "revision_requirements.json",
    "studio_memo.docx",
    "studio_memo.md",
    "ui_decisions.json",
    "validation_audit.json",
)
CALCULATION_GENERATED_ARTIFACTS = (
    "calculation_audit.json",
    "calculation_results.csv",
    "calculation_results.json",
)


def _utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _load_object(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"JSON root must be an object: {path}")
    return payload


def _nonempty(value: Any) -> bool:
    return isinstance(value, str) and bool(value.strip())


def _list_of_strings(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def _privacy_safe_text(value: Any) -> Any:
    """Omit mechanically detectable identifiers from generated JSON artifacts."""

    return "[omitted_for_privacy]" if privacy_issue(value) else value


def _as_date(value: Any) -> bool:
    if not isinstance(value, str):
        return False
    try:
        date.fromisoformat(value)
    except ValueError:
        return False
    return True


def _as_datetime(value: Any) -> bool:
    if not isinstance(value, str):
        return False
    try:
        parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return False
    return parsed.tzinfo is not None


def _source_records_sha256(records: dict[str, Any]) -> str:
    source = dict(records)
    source.pop("validation", None)
    raw = json.dumps(
        source,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=str,
    ).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def _file_sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _clear_prior_package_artifacts(
    output_dir: Path, *, preserve_calculations: bool
) -> None:
    """Remove only outputs owned by this packaging step before a rerun."""

    artifact_names = list(PACKAGE_GENERATED_ARTIFACTS)
    if not preserve_calculations:
        artifact_names.extend(CALCULATION_GENERATED_ARTIFACTS)
    for name in artifact_names:
        artifact = output_dir / name
        if artifact.is_symlink() or artifact.is_file():
            artifact.unlink()
        elif artifact.exists():
            raise ValueError(f"package artifact path must be a file: {artifact}")


def _audit_case_validation(
    case_records_path: Path, case_records: dict[str, Any]
) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    validation = case_records.get("validation")
    if not isinstance(validation, dict) or validation.get("status") != "passed":
        return [
            {
                "code": "case_records_not_validated",
                "field": "case_records.validation",
                "message": "Package input must be the deterministic validated case record.",
            }
        ]
    expected_hash = str(validation.get("source_records_sha256", ""))
    actual_hash = _source_records_sha256(case_records)
    if not expected_hash or expected_hash != actual_hash:
        issues.append(
            {
                "code": "case_records_validation_hash_mismatch",
                "field": "case_records.validation.source_records_sha256",
                "message": "Validated case records changed after provenance validation.",
            }
        )
    expected_binding = validation.get("acquisition_binding")
    try:
        current_binding = build_acquisition_binding(
            case_records_path.parent / "file_inventory.json",
            case_records_path.parent / "run_intake.json",
        )
    except AcquisitionBindingError:
        issues.append(
            {
                "code": "unverifiable_acquisition_binding",
                "field": "case_records.validation.acquisition_binding",
                "message": "Inventory and acquisition posture must remain present and verifiable.",
            }
        )
        current_binding = None
    if current_binding is not None:
        issues.extend(compare_acquisition_bindings(expected_binding, current_binding))
    audit_path = case_records_path.parent / "case_records_audit.json"
    if not audit_path.is_file() or audit_path.is_symlink():
        issues.append(
            {
                "code": "missing_case_records_audit",
                "field": "case_records_audit.json",
                "message": "The deterministic case-record audit is required for packaging.",
            }
        )
        return issues
    audit = _load_object(audit_path)
    if audit.get("status") != "passed":
        issues.append(
            {
                "code": "case_records_audit_not_passed",
                "field": "case_records_audit.status",
                "message": "The case-record audit must pass before packaging.",
            }
        )
    if str(audit.get("source_records_sha256", "")) != expected_hash:
        issues.append(
            {
                "code": "case_records_audit_hash_mismatch",
                "field": "case_records_audit.source_records_sha256",
                "message": "The audit does not match the validated case records.",
            }
        )
    if audit.get("acquisition_binding") != expected_binding:
        issues.append(
            {
                "code": "case_records_audit_acquisition_binding_mismatch",
                "field": "case_records_audit.acquisition_binding",
                "message": "The case-record audit does not match acquisition provenance.",
            }
        )
    validated_output = str(audit.get("validated_output_path", ""))
    if (
        not validated_output
        or Path(validated_output).resolve() != case_records_path.resolve()
    ):
        issues.append(
            {
                "code": "case_records_audit_path_mismatch",
                "field": "case_records_audit.validated_output_path",
                "message": "The audit points to a different validated record file.",
            }
        )
    return issues


def _audit_calculation_provenance(
    calculations_path: Path,
    calculations: dict[str, Any],
    case_records_path: Path,
    claims_review_path: Path,
    output_dir: Path,
) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    if calculations_path.is_symlink():
        return [
            {
                "code": "unsafe_calculation_results_path",
                "field": "calculations",
                "message": "Calculation results must be a regular local file, not a symlink.",
            }
        ]
    expected_results_path = output_dir / "calculation_results.json"
    if calculations_path.resolve() != expected_results_path.resolve():
        issues.append(
            {
                "code": "calculation_artifacts_outside_package",
                "field": "calculations",
                "message": "Calculation results must be generated in the package output directory.",
            }
        )
    for name in (
        "calculation_results.json",
        "calculation_results.csv",
        "calculation_audit.json",
    ):
        artifact = output_dir / name
        if not artifact.is_file() or artifact.is_symlink():
            issues.append(
                {
                    "code": "missing_calculation_artifact",
                    "field": name,
                    "message": f"The calculation package artifact is missing or unsafe: {name}.",
                }
            )
    audit_path = calculations_path.parent / "calculation_audit.json"
    if not audit_path.is_file() or audit_path.is_symlink():
        return [
            {
                "code": "missing_calculation_audit",
                "field": "calculation_audit.json",
                "message": "Calculation results require the reconciler audit artifact.",
            }
        ]
    audit = _load_object(audit_path)
    if audit.get("status") != "passed":
        issues.append(
            {
                "code": "calculation_audit_not_passed",
                "field": "calculation_audit.status",
                "message": "Calculation audit must pass before results enter the memo.",
            }
        )
    if (
        str(audit.get("calculation_results_path", "")) == ""
        or Path(str(audit.get("calculation_results_path", ""))).resolve()
        != calculations_path.resolve()
    ):
        issues.append(
            {
                "code": "calculation_results_path_mismatch",
                "field": "calculation_audit.calculation_results_path",
                "message": "Calculation audit points to a different results file.",
            }
        )
    if str(audit.get("calculation_results_sha256", "")) != _file_sha256(
        calculations_path
    ):
        issues.append(
            {
                "code": "calculation_results_hash_mismatch",
                "field": "calculation_audit.calculation_results_sha256",
                "message": "Calculation results changed after deterministic reconciliation.",
            }
        )
    csv_path = output_dir / "calculation_results.csv"
    recorded_csv_path = str(audit.get("calculation_results_csv_path", ""))
    if not recorded_csv_path or Path(recorded_csv_path).resolve() != csv_path.resolve():
        issues.append(
            {
                "code": "calculation_results_csv_path_mismatch",
                "field": "calculation_audit.calculation_results_csv_path",
                "message": "Calculation audit points to a different CSV result file.",
            }
        )
    if csv_path.is_file() and not csv_path.is_symlink():
        if str(audit.get("calculation_results_csv_sha256", "")) != _file_sha256(
            csv_path
        ):
            issues.append(
                {
                    "code": "calculation_results_csv_hash_mismatch",
                    "field": "calculation_audit.calculation_results_csv_sha256",
                    "message": "Calculation CSV changed after deterministic reconciliation.",
                }
            )
    audit_inputs = audit.get("input_provenance")
    result_inputs = calculations.get("input_provenance")
    if not isinstance(audit_inputs, dict) or audit_inputs != result_inputs:
        issues.append(
            {
                "code": "calculation_input_provenance_mismatch",
                "field": "calculation_audit.input_provenance",
                "message": "Calculation result inputs do not match the audit.",
            }
        )
        return issues
    expected_inputs = {
        "case_records": case_records_path,
        "claims_review": claims_review_path,
    }
    for name, expected_path in expected_inputs.items():
        record = audit_inputs.get(name)
        if not isinstance(record, dict):
            issues.append(
                {
                    "code": "missing_calculation_input_provenance",
                    "field": f"calculation_audit.input_provenance.{name}",
                    "message": f"Calculation audit is missing {name} provenance.",
                }
            )
            continue
        recorded_path = str(record.get("path", ""))
        if (
            not recorded_path
            or Path(recorded_path).resolve() != expected_path.resolve()
        ):
            issues.append(
                {
                    "code": "calculation_input_path_mismatch",
                    "field": f"calculation_audit.input_provenance.{name}.path",
                    "message": f"Calculation {name} path does not match the package input.",
                }
            )
        if str(record.get("sha256", "")) != _file_sha256(expected_path):
            issues.append(
                {
                    "code": "calculation_input_hash_mismatch",
                    "field": f"calculation_audit.input_provenance.{name}.sha256",
                    "message": f"Calculation {name} changed after reconciliation.",
                }
            )
    recipe_record = audit_inputs.get("recipes")
    if not isinstance(recipe_record, dict):
        issues.append(
            {
                "code": "missing_calculation_recipe_provenance",
                "field": "calculation_audit.input_provenance.recipes",
                "message": "Calculation audit must identify the approved recipe file.",
            }
        )
    else:
        recipe_path = Path(str(recipe_record.get("path", "")))
        if (
            not str(recipe_record.get("path", ""))
            or not recipe_path.is_file()
            or recipe_path.is_symlink()
            or str(recipe_record.get("sha256", "")) != _file_sha256(recipe_path)
        ):
            issues.append(
                {
                    "code": "calculation_recipe_provenance_mismatch",
                    "field": "calculation_audit.input_provenance.recipes",
                    "message": "Approved recipe file is missing or changed after reconciliation.",
                }
            )
    return issues


def _audit_sources(
    value: Any, *, claim_field: str
) -> tuple[list[dict[str, Any]], list[str], list[dict[str, str]]]:
    issues: list[dict[str, str]] = []
    normalized: list[dict[str, Any]] = []
    references: list[str] = []
    if not isinstance(value, list) or not value:
        return (
            [],
            [],
            [
                {
                    "code": "missing_sources",
                    "field": f"{claim_field}.sources",
                    "message": "Every material claim needs one or more source records.",
                }
            ],
        )
    seen_ids: set[str] = set()
    for index, source in enumerate(value):
        field = f"{claim_field}.sources[{index}]"
        if not isinstance(source, dict):
            issues.append(
                {
                    "code": "invalid_source_record",
                    "field": field,
                    "message": "Source record must be an object.",
                }
            )
            continue
        source_id = str(source.get("source_id", "")).strip()
        source_id_safe = safe_identifier(source_id)
        if not source_id_safe or source_id in seen_ids:
            issues.append(
                {
                    "code": "invalid_or_duplicate_source_id",
                    "field": f"{field}.source_id",
                    "message": "source_id must be opaque, identifier-free, and unique within the claim.",
                }
            )
        else:
            seen_ids.add(source_id)
        reference = str(source.get("reference", "")).strip()
        if not reference:
            issues.append(
                {
                    "code": "missing_source_reference",
                    "field": f"{field}.reference",
                    "message": "Source reference must be non-empty.",
                }
            )
        elif not safe_source_reference(reference):
            issues.append(
                {
                    "code": "unsafe_source_reference",
                    "field": f"{field}.reference",
                    "message": "Source reference must omit identifiers and private, credentialed, or tokenized URLs.",
                }
            )
        else:
            references.append(reference)
        if source.get("temporal_role") not in SOURCE_TEMPORAL_ROLES:
            issues.append(
                {
                    "code": "invalid_source_temporal_role",
                    "field": f"{field}.temporal_role",
                    "message": "Source must identify its role in the temporal analysis.",
                }
            )
        if not _as_datetime(source.get("retrieved_at")):
            issues.append(
                {
                    "code": "missing_source_retrieval_time",
                    "field": f"{field}.retrieved_at",
                    "message": "Each source needs a retrieval date-time with timezone.",
                }
            )
        for name in ("version_note", "support_note"):
            if not _nonempty(source.get(name)):
                issues.append(
                    {
                        "code": f"missing_source_{name}",
                        "field": f"{field}.{name}",
                        "message": f"Each source needs a non-empty {name}.",
                    }
                )
            elif privacy_issue(source.get(name)):
                issues.append(
                    {
                        "code": "unsafe_source_note",
                        "field": f"{field}.{name}",
                        "message": "Source notes must omit raw identifiers and private or tokenized URLs.",
                    }
                )
        snapshot_hash = source.get("snapshot_sha256")
        if snapshot_hash is not None and (
            not isinstance(snapshot_hash, str)
            or len(snapshot_hash) != 64
            or any(character not in "0123456789abcdef" for character in snapshot_hash)
        ):
            issues.append(
                {
                    "code": "invalid_source_snapshot_hash",
                    "field": f"{field}.snapshot_sha256",
                    "message": "Snapshot hash must be a lowercase SHA-256 value when present.",
                }
            )
        normalized.append(
            {
                "source_id": (
                    source_id if source_id_safe else f"SRC-OMITTED-{index + 1:03d}"
                ),
                "reference": (
                    reference
                    if reference and safe_source_reference(reference)
                    else "[omitted_for_privacy]"
                ),
                "temporal_role": source.get("temporal_role"),
                "retrieved_at": source.get("retrieved_at"),
                "version_note": (
                    source.get("version_note")
                    if not privacy_issue(source.get("version_note"))
                    else "[omitted_for_privacy]"
                ),
                "support_note": (
                    source.get("support_note")
                    if not privacy_issue(source.get("support_note"))
                    else "[omitted_for_privacy]"
                ),
                "snapshot_sha256": snapshot_hash,
            }
        )
    return normalized, references, issues


def _audit_claims(
    claims_payload: dict[str, Any], fact_ids: set[str]
) -> tuple[list[dict[str, Any]], list[dict[str, str]]]:
    claims = claims_payload.get("claims")
    issues: list[dict[str, str]] = []
    normalized: list[dict[str, Any]] = []
    if not isinstance(claims, list) or not claims:
        return [], [
            {
                "code": "missing_claims",
                "field": "claims",
                "message": "At least one source-reviewed material claim is required.",
            }
        ]
    seen_ids: set[str] = set()
    for index, claim in enumerate(claims, start=1):
        field = f"claims[{index - 1}]"
        if not isinstance(claim, dict):
            issues.append(
                {
                    "code": "invalid_claim",
                    "field": field,
                    "message": "Claim must be an object.",
                }
            )
            continue
        claim_id = str(claim.get("claim_id") or f"CL-{index:03d}").strip()
        claim_id_safe = safe_identifier(claim_id)
        if not claim_id_safe:
            issues.append(
                {
                    "code": "unsafe_claim_id",
                    "field": f"{field}.claim_id",
                    "message": "claim_id must be an opaque identifier without personal data.",
                }
            )
            claim_id = f"CL-OMITTED-{index:03d}"
        if claim_id in seen_ids:
            issues.append(
                {
                    "code": "duplicate_claim_id",
                    "field": f"{field}.claim_id",
                    "message": "claim_id values must be unique.",
                }
            )
        seen_ids.add(claim_id)
        verdict = str(claim.get("verdict", ""))
        claim_type = str(claim.get("claim_type", ""))
        if not _nonempty(claim.get("claim_text")):
            issues.append(
                {
                    "code": "missing_claim_text",
                    "field": f"{field}.claim_text",
                    "message": "Material claim text is required.",
                }
            )
        elif privacy_issue(claim.get("claim_text")):
            issues.append(
                {
                    "code": "unsafe_claim_text",
                    "field": f"{field}.claim_text",
                    "message": "Claim text must omit raw identifiers and private or tokenized URLs.",
                }
            )
        if not _nonempty(claim.get("review_label")):
            issues.append(
                {
                    "code": "missing_safe_review_label",
                    "field": f"{field}.review_label",
                    "message": "A concise identifier-free review label is required.",
                }
            )
        elif privacy_issue(claim.get("review_label")):
            issues.append(
                {
                    "code": "unsafe_review_label",
                    "field": f"{field}.review_label",
                    "message": "Review labels must omit raw identity, tax codes, email, and private URLs.",
                }
            )
        if claim_type not in CLAIM_TYPES:
            issues.append(
                {
                    "code": "invalid_claim_type",
                    "field": f"{field}.claim_type",
                    "message": "Claim must be typed as rule, case_application, or calculation_basis.",
                }
            )
        if verdict not in VERDICTS:
            issues.append(
                {
                    "code": "invalid_verdict",
                    "field": f"{field}.verdict",
                    "message": f"Unsupported verdict: {verdict or '<empty>'}.",
                }
            )
        elif verdict != "supported":
            issues.append(
                {
                    "code": "material_claim_not_fully_supported",
                    "field": f"{field}.verdict",
                    "message": "Every material claim must be fully supported before the professional-review package is ready.",
                }
            )
        sources, source_refs, source_issues = _audit_sources(
            claim.get("sources"), claim_field=field
        )
        issues.extend(source_issues)
        if not _nonempty(claim.get("reasoning_review")):
            issues.append(
                {
                    "code": "missing_reasoning_review",
                    "field": f"{field}.reasoning_review",
                    "message": "Reasoning must be reviewed separately from source existence.",
                }
            )
        elif privacy_issue(claim.get("reasoning_review")):
            issues.append(
                {
                    "code": "unsafe_reasoning_review",
                    "field": f"{field}.reasoning_review",
                    "message": "Reasoning text must omit raw identifiers and private or tokenized URLs.",
                }
            )
        dependencies = _list_of_strings(
            claim.get("evidence_dependencies", claim.get("fact_refs", []))
        )
        if any(not safe_identifier(dependency) for dependency in dependencies):
            issues.append(
                {
                    "code": "unsafe_fact_dependency",
                    "field": f"{field}.evidence_dependencies",
                    "message": "Fact dependencies must use opaque identifier-free IDs.",
                }
            )
        if claim_type in {"case_application", "calculation_basis"} and not dependencies:
            issues.append(
                {
                    "code": "missing_case_fact_dependencies",
                    "field": f"{field}.evidence_dependencies",
                    "message": "Case-application and calculation-basis claims need validated fact dependencies.",
                }
            )
        unknown = sorted(set(dependencies) - fact_ids)
        if unknown:
            issues.append(
                {
                    "code": "unknown_fact_dependency",
                    "field": f"{field}.evidence_dependencies",
                    "message": f"Unknown fact ids: {', '.join(unknown)}.",
                }
            )
        period_scope = claim.get("period_scope")
        scope_status = (
            str(period_scope.get("status", ""))
            if isinstance(period_scope, dict)
            else ""
        )
        scope_start = (
            period_scope.get("start") if isinstance(period_scope, dict) else None
        )
        scope_end = period_scope.get("end") if isinstance(period_scope, dict) else None
        if scope_status not in {"confirmed", "open_ended", "unresolved"}:
            issues.append(
                {
                    "code": "missing_temporal_scope",
                    "field": f"{field}.period_scope",
                    "message": "Each claim needs a confirmed, open-ended, or unresolved temporal status.",
                }
            )
        if not isinstance(period_scope, dict) or not _nonempty(
            period_scope.get("note")
        ):
            issues.append(
                {
                    "code": "missing_temporal_scope_note",
                    "field": f"{field}.period_scope.note",
                    "message": "Temporal scope needs an explicit basis or unresolved-state note.",
                }
            )
        elif privacy_issue(period_scope.get("note")):
            issues.append(
                {
                    "code": "unsafe_temporal_scope_note",
                    "field": f"{field}.period_scope.note",
                    "message": "Temporal scope notes must omit raw identifiers and private URLs.",
                }
            )
        for name, boundary in (("start", scope_start), ("end", scope_end)):
            if boundary is not None and not _as_date(boundary):
                issues.append(
                    {
                        "code": "invalid_temporal_scope_boundary",
                        "field": f"{field}.period_scope.{name}",
                        "message": "Temporal boundaries must be ISO dates or null.",
                    }
                )
        if scope_status == "confirmed" and (
            not _as_date(scope_start) or not _as_date(scope_end)
        ):
            issues.append(
                {
                    "code": "unconfirmed_temporal_scope",
                    "field": f"{field}.period_scope",
                    "message": "A confirmed scope needs exact ISO start and end dates.",
                }
            )
        if verdict == "supported" and scope_status != "confirmed":
            issues.append(
                {
                    "code": "supported_claim_with_unresolved_temporal_scope",
                    "field": f"{field}.period_scope.status",
                    "message": "A supported claim requires a confirmed temporal scope.",
                }
            )
        if (
            _as_date(scope_start)
            and _as_date(scope_end)
            and date.fromisoformat(scope_start) > date.fromisoformat(scope_end)
        ):
            issues.append(
                {
                    "code": "invalid_temporal_scope_order",
                    "field": f"{field}.period_scope",
                    "message": "Temporal scope start must not be after its end.",
                }
            )
        if not _as_date(claim.get("research_cutoff_date")):
            issues.append(
                {
                    "code": "missing_research_cutoff",
                    "field": f"{field}.research_cutoff_date",
                    "message": "Each claim needs an ISO research cut-off date.",
                }
            )
        normalized_claim = {
            "claim_id": claim_id,
            "claim_text": _privacy_safe_text(claim.get("claim_text")),
            "review_label": _privacy_safe_text(claim.get("review_label")),
            "claim_type": claim_type,
            "verdict": verdict,
            "sources": sources,
            "source_refs": source_refs,
            "source_support": _privacy_safe_text(claim.get("source_support")),
            "reasoning_review": _privacy_safe_text(claim.get("reasoning_review")),
            "evidence_dependencies": [
                dependency for dependency in dependencies if safe_identifier(dependency)
            ],
            "period_scope": {
                "status": scope_status,
                "start": scope_start,
                "end": scope_end,
                "note": _privacy_safe_text(
                    period_scope.get("note") if isinstance(period_scope, dict) else None
                ),
            },
            "research_cutoff_date": claim.get("research_cutoff_date"),
            "professional_review_status": claim.get(
                "professional_review_status", "pending"
            ),
        }
        for name in ("later_authority_note", "proposed_fix", "uncertainty"):
            if name in claim:
                normalized_claim[name] = _privacy_safe_text(claim.get(name))
        normalized.append(normalized_claim)
    return normalized, issues


def _fact_rows(case_records: dict[str, Any]) -> list[dict[str, Any]]:
    return [fact for fact in case_records.get("facts", []) if isinstance(fact, dict)]


def _audit_fact_review_privacy(
    case_records: dict[str, Any],
) -> list[dict[str, str]]:
    """Protect review-visible identifiers without interpreting case meaning."""

    issues: list[dict[str, str]] = []
    for index, fact in enumerate(_fact_rows(case_records)):
        field = f"case_records.facts[{index}]"
        if not safe_identifier(fact.get("fact_id")):
            issues.append(
                {
                    "code": "unsafe_fact_id",
                    "field": f"{field}.fact_id",
                    "message": "Fact IDs must be opaque and identifier-free.",
                }
            )
        if privacy_issue(fact.get("review_label")):
            issues.append(
                {
                    "code": "unsafe_fact_review_label",
                    "field": f"{field}.review_label",
                    "message": "Fact review labels must omit raw identity, tax codes, email, and private URLs.",
                }
            )
    return issues


def _timeline_rows(case_records: dict[str, Any]) -> list[dict[str, Any]]:
    rows = [
        event for event in case_records.get("timeline", []) if isinstance(event, dict)
    ]
    return sorted(
        rows,
        key=lambda event: (str(event.get("date", "")), str(event.get("event_id", ""))),
    )


def _calculation_rows(calculations: dict[str, Any] | None) -> list[dict[str, Any]]:
    if not calculations:
        return []
    return [row for row in calculations.get("results", []) if isinstance(row, dict)]


def _missing_evidence(
    case_records: dict[str, Any], claims_payload: dict[str, Any]
) -> list[str]:
    items: list[str] = []
    for source in (
        case_records.get("open_questions", []),
        claims_payload.get("missing_evidence", []),
    ):
        if not isinstance(source, list):
            continue
        for item in source:
            if isinstance(item, dict):
                value = (
                    item.get("request")
                    or item.get("description")
                    or item.get("question")
                )
            else:
                value = item
            if _nonempty(value):
                text = str(value).strip()
                items.append(
                    "Request details omitted for privacy; consult the local case record."
                    if privacy_issue(text)
                    else text
                )
    return list(dict.fromkeys(items))


def _blocked_note_lines(
    case_records: dict[str, Any], issues: list[dict[str, str]], missing: list[str]
) -> list[str]:
    lines = [
        "# FASCICOLO BLOCCATO — NON È UN PARERE",
        "",
        f"**Quesito:** {_privacy_safe_text(case_records.get('professional_question', ''))}",
        "",
        "Il pacchetto conclusivo non è stato creato perché uno o più gate di validazione non sono soddisfatti.",
        "",
        "## Blocchi",
        "",
    ]
    lines.extend(
        f"- `{issue.get('code', 'validation_issue')}`: {issue.get('message', '')}"
        for issue in issues
    )
    lines.extend(["", "## Prossimi elementi richiesti", ""])
    if missing:
        lines.extend(f"- {item}" for item in missing)
    else:
        lines.append(
            "- Risolvere i blocchi indicati nell'audit e rieseguire la validazione."
        )
    return lines


def _memo_lines(
    case_records: dict[str, Any],
    claims: list[dict[str, Any]],
    calculations: list[dict[str, Any]],
    missing: list[str],
) -> list[str]:
    lines = [
        "# BOZZA PER REVISIONE PROFESSIONALE",
        "",
        "## Perimetro",
        "",
        f"**Quesito:** {case_records.get('professional_question', '')}",
        "",
        "Il documento organizza evidenze, ricerca e ricalcoli. Non costituisce un parere professionale. Può includere evidenze locali, esportazioni ufficiali o, solo dopo separata verifica del permesso del servizio, una cattura in sola lettura di una scheda INPS già autenticata; non gestisce credenziali, non attiva deleghe e non autorizza invii o adempimenti.",
        "",
        "## Fatti documentati",
        "",
    ]
    for fact in _fact_rows(case_records):
        lines.append(
            f"- **{fact.get('fact_id', '')}** [{fact.get('review_status', '')}]: {fact.get('statement', '')}"
        )
    lines.extend(["", "## Cronologia esplicita", ""])
    timeline = _timeline_rows(case_records)
    if timeline:
        for event in timeline:
            lines.append(
                f"- **{event.get('date', '')}** — {event.get('description', '')} "
                f"(fatti: {', '.join(map(str, event.get('source_fact_ids', [])))})"
            )
    else:
        lines.append("- Nessun evento cronologico esplicito registrato.")
    lines.extend(["", "## Conclusioni sottoposte a verifica delle fonti", ""])
    for claim in claims:
        lines.extend(
            [
                f"### {claim['claim_id']} — {claim.get('verdict', '')}",
                "",
                str(claim.get("claim_text", "")),
                "",
                f"**Supporto:** {claim.get('source_support', '')}",
                "",
                f"**Ragionamento:** {claim.get('reasoning_review', '')}",
                "",
                f"**Incertezza/correzione:** {claim.get('proposed_fix') or claim.get('uncertainty') or 'Nessuna indicata.'}",
                "",
                f"**Fonti:** {', '.join(claim.get('source_refs', []))}",
                "",
                f"**Fatti dipendenti:** {', '.join(claim.get('evidence_dependencies', [])) or 'Nessuno indicato.'}",
                "",
            ]
        )
    lines.extend(["## Ricalcoli approvati", ""])
    if calculations:
        for row in calculations:
            lines.append(
                f"- **{row.get('recipe_id', '')}** [{row.get('status', '')}]: "
                f"{row.get('description', '')} — {row.get('result', 'non calcolato')} {row.get('unit', '')}"
            )
    else:
        lines.append("- Nessun ricalcolo approvato incluso.")
    lines.extend(["", "## Evidenze o chiarimenti mancanti", ""])
    if missing:
        lines.extend(f"- {item}" for item in missing)
    else:
        lines.append(
            "- Nessuna richiesta aggiuntiva registrata; il professionista deve comunque confermare la completezza del fascicolo."
        )
    lines.extend(
        [
            "",
            "## Limiti e responsabilità",
            "",
            "- La selezione del quadro giuridico, delle fonti e delle conclusioni è stata svolta in modo model-led e deve essere verificata dal professionista.",
            "- Gli script hanno controllato soltanto forma, provenienza, riferimenti e aritmetica esplicitamente approvata.",
            "- Nessuna classificazione, aliquota, prescrizione o scadenza è stata scelta automaticamente.",
            "- Stato del documento: **pronto per revisione professionale**, non firmato e non depositabile.",
        ]
    )
    return lines


def _write_docx(path: Path, lines: Iterable[str]) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for line in lines:
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line.replace("**", ""))
    document.save(path)
    mark_private_file(path)


def _item(
    *,
    item_id: str,
    item_type: str,
    title: str,
    recommended_action: str,
    data: dict[str, Any],
    evidence: list[dict[str, Any]] | None = None,
    output_path: str | None = None,
) -> dict[str, Any]:
    titles = {
        "fact": "Fact record",
        "finding": "Finding record",
        "calculation": "Calculation record",
        "missing_evidence": "Missing evidence request",
        "authority": "Authority record",
        "audit_check": "Package validation audit",
        "artifact": "Package artifact",
    }
    summaries = {
        "fact": "Validated fact record",
        "finding": "Source-reviewed finding",
        "calculation": "Approved calculation record",
        "missing_evidence": "Missing evidence request",
        "authority": "Authority record",
        "audit_check": "Package validation audit",
        "artifact": "Package artifact",
    }
    return {
        "id": item_id,
        "item_type": item_type,
        "title": titles[item_type],
        "source_path": None,
        "output_path": output_path,
        "allowed_actions": ALLOWED_ACTIONS,
        "recommended_action": recommended_action,
        "evidence": [],
        "data": {**data, "summary": summaries[item_type]},
        "status": "needs_review",
    }


def _review_items(
    case_records: dict[str, Any],
    claims: list[dict[str, Any]],
    calculations: list[dict[str, Any]],
    missing: list[str],
    audit_status: str,
) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for fact_index, fact in enumerate(_fact_rows(case_records), start=1):
        status = str(fact.get("review_status", ""))
        raw_fact_id = fact.get("fact_id")
        fact_id = (
            str(raw_fact_id)
            if safe_identifier(raw_fact_id)
            else f"F-OMITTED-{fact_index:03d}"
        )
        items.append(
            _item(
                item_id=f"fact-{fact_id}",
                item_type="fact",
                title=f"Fact {fact_index}",
                recommended_action=(
                    "accept" if status == "confirmed" else "mark_unclear"
                ),
                data={
                    "fact_id": fact_id,
                    "summary": "Validated fact record",
                    "review_status": status,
                    "evidence_count": len(fact.get("evidence", [])),
                },
            )
        )
    for claim_index, claim in enumerate(claims, start=1):
        verdict = str(claim.get("verdict", ""))
        action = (
            "accept"
            if verdict == "supported"
            else (
                "reject"
                if verdict in {"not_supported", "contradicted"}
                else "mark_unclear"
            )
        )
        items.append(
            _item(
                item_id=f"finding-{claim['claim_id']}",
                item_type="finding",
                title=f"Finding {claim_index}",
                recommended_action=action,
                data={
                    "claim_id": claim["claim_id"],
                    "summary": "Source-reviewed finding",
                    "claim_type": claim.get("claim_type"),
                    "verdict": verdict,
                    "source_count": len(claim.get("source_refs", [])),
                    "professional_review_status": claim.get(
                        "professional_review_status"
                    ),
                },
                evidence=[
                    {"kind": "fact_dependency", "fact_id": fact_id}
                    for fact_id in claim.get("evidence_dependencies", [])
                    if safe_identifier(fact_id)
                ],
            )
        )
    for calculation_index, calculation in enumerate(calculations, start=1):
        raw_recipe_id = calculation.get("recipe_id")
        recipe_id = (
            str(raw_recipe_id)
            if safe_identifier(raw_recipe_id)
            else f"CALC-OMITTED-{calculation_index:03d}"
        )
        items.append(
            _item(
                item_id=f"calculation-{recipe_id}",
                item_type="calculation",
                title=f"Calculation {calculation_index}",
                recommended_action=(
                    "accept"
                    if audit_status == "passed"
                    and calculation.get("status") == "calculated"
                    else "mark_unclear"
                ),
                data={
                    "recipe_id": recipe_id,
                    "status": calculation.get("status"),
                    "summary": "Approved calculation record",
                },
            )
        )
    for index, request in enumerate(missing, start=1):
        items.append(
            _item(
                item_id=f"missing-evidence-{index:03d}",
                item_type="missing_evidence",
                title=f"Missing evidence request {index}",
                recommended_action="request_more_documents",
                data={"request_id": f"REQ-{index:03d}"},
            )
        )
    items.append(
        _item(
            item_id="audit-package",
            item_type="audit_check",
            title="Package validation audit",
            recommended_action="accept" if audit_status == "passed" else "reject",
            data={"status": audit_status},
            output_path="validation_audit.json",
        )
    )
    artifact_names = (
        (
            "studio_memo.md",
            "studio_memo.docx",
            "validation_audit.json",
            "review_handoff.md",
        )
        if audit_status == "passed"
        else ("blocked_case_note.md", "validation_audit.json", "review_handoff.md")
    )
    for name in artifact_names:
        items.append(
            _item(
                item_id=f"artifact-{name.replace('.', '-')}",
                item_type="artifact",
                title=f"Artifact {name}",
                recommended_action="accept",
                data={"path": name},
                output_path=name,
            )
        )
    return items


def _load_bound_run_intake(
    output_dir: Path, case_records: dict[str, Any]
) -> dict[str, Any]:
    path = output_dir / "run_intake.json"
    if not path.is_file() or path.is_symlink():
        raise ValueError(
            "run_intake.json from the inventory step is required; packaging cannot create local-only posture"
        )
    payload = _load_object(path)
    if (
        payload.get("plugin") != "previdenza-inps"
        or payload.get("workflow") != "previdenza-inps"
    ):
        raise ValueError("run_intake.json does not belong to this workflow")
    binding = (
        case_records.get("validation", {}).get("acquisition_binding", {})
        if isinstance(case_records.get("validation"), dict)
        else {}
    )
    if binding.get("run_id") and payload.get("run_id") != binding.get("run_id"):
        raise ValueError("run_intake.json does not match validated acquisition run")
    recorded_output = payload.get("output_dir")
    if (
        not isinstance(recorded_output, str)
        or Path(recorded_output).resolve() != output_dir.resolve()
    ):
        raise ValueError("run_intake.output_dir does not match the package directory")
    return payload


def _run_intake(
    output_dir: Path,
    case_records_path: Path,
    claims_review_path: Path,
    case_records: dict[str, Any],
    *,
    calculations_path: Path | None,
    package_status: str,
) -> dict[str, Any]:
    payload = _load_bound_run_intake(output_dir, case_records)
    primary_outputs = (
        ["studio_memo.md", "studio_memo.docx"]
        if package_status == "passed"
        else ["blocked_case_note.md"]
    )
    trace_outputs = primary_outputs + [
        "document_requests.md",
        "validation_audit.json",
        "claims_review_normalized.json",
        "review_payload.json",
        "ui_decisions.json",
        "final_artifacts.json",
        "review_handoff.md",
    ]
    if calculations_path is not None:
        trace_outputs.extend(
            name
            for name in (
                "calculation_results.json",
                "calculation_results.csv",
                "calculation_audit.json",
            )
            if (output_dir / name).is_file() and not (output_dir / name).is_symlink()
        )
    trace_inputs = ["validated_case_records", "claims_review"]
    if calculations_path is not None:
        trace_inputs.append("calculation_results")
    existing_trace = [
        entry
        for entry in payload.get("execution_trace", [])
        if isinstance(entry, dict) and entry.get("step_id") != "previdenza_inps_package"
    ]
    existing_trace.append(
        {
            "step_id": "previdenza_inps_package",
            "kind": "deterministic_packaging",
            "status": "passed" if package_status == "passed" else "blocked",
            "execution_location": "local_codex_workspace",
            "command": "python scripts/package_case.py",
            "inputs": trace_inputs,
            "outputs": trace_outputs,
        }
    )
    payload["execution_trace"] = existing_trace
    return payload


def package_case(
    case_records_path: Path,
    claims_review_path: Path,
    output_dir: Path,
    *,
    calculations_path: Path | None = None,
) -> dict[str, Any]:
    """Validate and package model-led findings without changing their meaning."""

    case_records = _load_object(case_records_path)
    claims_payload = _load_object(claims_review_path)
    calculations_payload = (
        _load_object(calculations_path) if calculations_path else None
    )
    _load_bound_run_intake(output_dir, case_records)
    fact_ids = {
        str(fact.get("fact_id"))
        for fact in _fact_rows(case_records)
        if fact.get("fact_id")
    }
    issues = _audit_case_validation(case_records_path, case_records)
    issues.extend(_audit_fact_review_privacy(case_records))
    claims, claim_issues = _audit_claims(claims_payload, fact_ids)
    issues.extend(claim_issues)
    calculations = _calculation_rows(calculations_payload)
    if calculations_path is not None and calculations_payload is not None:
        issues.extend(
            _audit_calculation_provenance(
                calculations_path,
                calculations_payload,
                case_records_path,
                claims_review_path,
                output_dir,
            )
        )
        if calculations_payload.get("status") != "passed" or not calculations:
            issues.append(
                {
                    "code": "calculation_results_not_passed",
                    "field": "calculations.status",
                    "message": "Calculation results must be passed and non-empty.",
                }
            )
    for calculation in calculations:
        if calculation.get("status") != "calculated":
            issues.append(
                {
                    "code": "calculation_not_run",
                    "field": f"calculations.{calculation.get('recipe_id', '')}",
                    "message": "A calculation recipe did not pass its mechanical gate.",
                }
            )
    status = "passed" if not issues else "validation_fail"
    audit = {
        "schema_version": "1.0",
        "plugin": "previdenza-inps",
        "status": status,
        "validated_at": _utc_now(),
        "claim_count": len(claims),
        "issue_count": len(issues),
        "issues": issues,
        "semantic_scope": "model_authored_not_reclassified_by_packager",
        "acquisition_binding": case_records.get("validation", {}).get(
            "acquisition_binding"
        ),
    }
    prepare_private_directory(output_dir)
    _clear_prior_package_artifacts(
        output_dir, preserve_calculations=calculations_path is not None
    )
    write_json(output_dir / "validation_audit.json", audit)
    normalized_claims = {
        "schema_version": "1.0",
        "plugin": "previdenza-inps",
        "language": claims_payload.get("language"),
        "claims": claims,
        "claim_count": len(claims),
        "privacy_posture": "allowlisted_fields_identifier_guarded",
    }
    write_json(output_dir / "claims_review_normalized.json", normalized_claims)

    missing = _missing_evidence(case_records, claims_payload)
    if status == "passed":
        memo_lines = _memo_lines(case_records, claims, calculations, missing)
        memo_path = output_dir / "studio_memo.md"
        write_private_text(memo_path, "\n".join(memo_lines).rstrip() + "\n")
        _write_docx(output_dir / "studio_memo.docx", memo_lines)
    else:
        blocked_lines = _blocked_note_lines(case_records, issues, missing)
        write_private_text(
            output_dir / "blocked_case_note.md",
            "\n".join(blocked_lines).rstrip() + "\n",
        )
    write_private_text(
        output_dir / "document_requests.md",
        "# Document and clarification requests\n\n"
        + (
            "\n".join(f"- {item}" for item in missing)
            if missing
            else "- None recorded."
        )
        + "\n",
    )

    run_intake = _run_intake(
        output_dir,
        case_records_path,
        claims_review_path,
        case_records,
        calculations_path=calculations_path,
        package_status=status,
    )
    run_id = str(run_intake["run_id"])
    items = _review_items(case_records, claims, calculations, missing, status)
    review_status = (
        "ready_for_professional_review" if status == "passed" else "validation_fail"
    )
    review_payload = {
        "schema_version": "1.0",
        "plugin": "previdenza-inps",
        "workflow": "previdenza-inps",
        "run_id": run_id,
        "review_type": "professional_case_review",
        "items": items,
        "item_count": len(items),
        "status": review_status,
        "privacy_notice": "Review payload omits document quotes and subject labels; inspect local artifacts for full evidence.",
    }
    write_json(output_dir / "review_payload.json", review_payload)
    review_payload_sha256 = _file_sha256(output_dir / "review_payload.json")
    write_json(
        output_dir / "ui_decisions.json",
        {
            "schema_version": "1.0",
            "plugin": "previdenza-inps",
            "workflow": "previdenza-inps",
            "run_id": run_id,
            "status": "pending",
            "decisions": [],
        },
    )
    outputs = [
        {"path": "document_requests.md", "kind": "md", "status": "written"},
        {"path": "validation_audit.json", "kind": "json", "status": "written"},
        {
            "path": "claims_review_normalized.json",
            "kind": "json",
            "status": "written",
        },
        {"path": "review_payload.json", "kind": "json", "status": "written"},
        {"path": "ui_decisions.json", "kind": "json", "status": "pending_review"},
        {
            "path": "review_handoff.md",
            "kind": "md",
            "status": "written",
            "required_text": [
                "validate_previdenza_inps_review",
                "render_previdenza_inps_review",
                "save_previdenza_inps_decisions",
                "apply_previdenza_inps_decisions",
            ],
        },
    ]
    if status == "passed":
        outputs[:0] = [
            {
                "path": "studio_memo.md",
                "kind": "md",
                "status": "written",
                "required_text": ["BOZZA PER REVISIONE PROFESSIONALE"],
            },
            {
                "path": "studio_memo.docx",
                "kind": "docx",
                "status": "written",
                "required_text": ["BOZZA PER REVISIONE PROFESSIONALE"],
            },
        ]
    else:
        outputs.insert(
            0,
            {
                "path": "blocked_case_note.md",
                "kind": "md",
                "status": "written",
                "required_text": ["FASCICOLO BLOCCATO", "NON È UN PARERE"],
            },
        )
    if calculations_path is not None:
        for name, kind in (
            ("calculation_results.json", "json"),
            ("calculation_results.csv", "csv"),
            ("calculation_audit.json", "json"),
        ):
            artifact = output_dir / name
            artifact_exists = artifact.is_file() and not artifact.is_symlink()
            outputs.append(
                {
                    "path": name,
                    "kind": kind,
                    "status": (
                        "written"
                        if status == "passed" and artifact_exists
                        else "blocked" if artifact_exists else "missing"
                    ),
                }
            )
    final_artifacts = {
        "schema_version": "1.0",
        "plugin": "previdenza-inps",
        "workflow": "previdenza-inps",
        "run_id": run_id,
        "status": review_status,
        "professional_review_required": True,
        "review_payload_sha256": review_payload_sha256,
        "acquisition_binding": case_records.get("validation", {}).get(
            "acquisition_binding"
        ),
        "outputs": outputs,
        "caveats": [
            "The package is a draft and does not constitute a filed or signed professional opinion."
        ],
        "blockers": issues,
        "next_actions": (
            [
                "Validate and render review_payload.json, then save and apply every professional review decision."
            ]
            if status == "passed"
            else [
                "Resolve every validation blocker, regenerate the affected artifacts, and rerun professional review."
            ]
        ),
        "review_status": review_status,
    }
    write_private_text(
        output_dir / "review_handoff.md",
        "# Previdenza INPS Review Handoff\n\n"
        "## Files\n\n"
        "- Intake: `run_intake.json`\n"
        "- Review queue: `review_payload.json`\n"
        "- Pending/saved decisions: `ui_decisions.json`\n"
        "- Applied decisions after review: `applied_decisions.json`\n"
        "- Final manifest: `final_artifacts.json`\n\n"
        "## Review sequence\n\n"
        "1. Validate `review_payload.json` with `validate_previdenza_inps_review`.\n"
        "2. Render it with `render_previdenza_inps_review`.\n"
        "3. Save reviewer actions with `save_previdenza_inps_decisions`.\n"
        "4. Apply actions with `apply_previdenza_inps_decisions`.\n\n"
        "The final memo remains a draft for professional review.\n",
    )
    write_json(output_dir / "final_artifacts.json", final_artifacts)
    write_json(output_dir / "run_intake.json", run_intake)
    return {
        "audit": audit,
        "review_payload": review_payload,
        "final_artifacts": final_artifacts,
    }


def main(argv: list[str] | None = None) -> int:
    """Package an assessed case and write the review handoff."""

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("case_records", type=Path)
    parser.add_argument("claims_review", type=Path)
    parser.add_argument("--calculations", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args(argv)
    try:
        output_dir = ensure_safe_output_dir(args.output_dir, plugin_root=PLUGIN_ROOT)
        result = package_case(
            args.case_records,
            args.claims_review,
            output_dir,
            calculations_path=args.calculations,
        )
    except (
        FileNotFoundError,
        json.JSONDecodeError,
        PermissionError,
        ValueError,
    ) as exc:
        LOGGER.error("%s", exc)
        return 1
    if result["audit"]["status"] != "passed":
        LOGGER.error(
            "Package written with %s validation issue(s).",
            result["audit"]["issue_count"],
        )
        return 1
    LOGGER.info("Case package is ready for professional review at %s", output_dir)
    return 0


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(message)s")
    raise SystemExit(main())
