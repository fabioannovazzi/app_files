"""Mechanical document inventory and extraction helpers for previdenza cases."""

from __future__ import annotations

import csv
import hashlib
import importlib
import io
import json
import logging
import mimetypes
import re
import stat
import sys
from dataclasses import dataclass
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any, Iterable
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from pypdf import PdfReader
from pypdf.errors import FileNotDecryptedError, PdfReadError

__all__ = [
    "ExtractionResult",
    "ensure_safe_output_dir",
    "extract_case_documents",
    "mark_private_file",
    "prepare_private_directory",
    "read_fragment_text",
    "write_json",
    "write_private_text",
]

LOGGER = logging.getLogger(__name__)
SUPPORTED_TEXT_SUFFIXES = {".csv", ".html", ".json", ".md", ".txt", ".xml"}
SUPPORTED_IMAGE_SUFFIXES = {".jpeg", ".jpg", ".png", ".tif", ".tiff"}
MAX_CELLS_PER_SHEET = 20_000
MAX_TEXT_CHARACTERS = 2_000_000
MIN_EMBEDDED_TEXT_CHARACTERS = 40
MIN_EMBEDDED_ALPHA_CHARACTERS = 20
SAFE_COMPONENT_RE = re.compile(r"[^A-Za-z0-9_.-]+")
OCR_VISUAL_CONFIRMATION_LIMITATION = "ocr_text_requires_visual_confirmation"
BROWSER_VISUAL_CONFIRMATION_LIMITATION = (
    "browser_capture_text_requires_visual_confirmation"
)
OCR_STATUS_LIMITATIONS = {
    "runtime_unavailable": "ocr_runtime_unavailable",
    "models_unavailable": "ocr_models_unavailable",
    "inference_failed": "ocr_inference_failed",
}


@dataclass(frozen=True)
class ExtractionResult:
    """Inventory artifacts written for one input folder."""

    inventory: dict[str, Any]
    inventory_path: Path
    inventory_csv_path: Path
    extraction_report_path: Path
    extracted_evidence_path: Path


def prepare_private_directory(path: Path) -> Path:
    """Create an access-restricted output directory and return it."""

    if path.is_symlink():
        raise PermissionError(f"private output directory cannot be a symlink: {path}")
    if path.exists():
        if not path.is_dir():
            raise NotADirectoryError(path)
        permissions = stat.S_IMODE(path.stat().st_mode)
        if permissions & 0o077:
            raise PermissionError(
                f"existing output directory must already be owner-only (0700): {path}"
            )
        return path
    path.mkdir(parents=True, mode=0o700, exist_ok=False)
    path.chmod(0o700)
    return path


def mark_private_file(path: Path) -> Path:
    """Restrict a written evidence artifact to its owning user."""

    path.chmod(0o600)
    return path


def write_private_text(path: Path, text: str) -> Path:
    """Write UTF-8 text with owner-only file permissions."""

    prepare_private_directory(path.parent)
    path.write_text(text, encoding="utf-8")
    return mark_private_file(path)


def write_json(path: Path, payload: dict[str, Any]) -> Path:
    """Write stable, owner-only UTF-8 JSON and return its path."""

    return write_private_text(
        path,
        json.dumps(payload, ensure_ascii=False, indent=2, default=str) + "\n",
    )


def _sha256_bytes(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def _find_git_root(start: Path) -> Path | None:
    for candidate in (start.resolve(), *start.resolve().parents):
        if (candidate / ".git").exists():
            return candidate
    return None


def ensure_safe_output_dir(output_dir: Path, *, plugin_root: Path) -> Path:
    """Reject run outputs inside the source Git workspace.

    The fixed boundary is justified by data separation and release integrity: user
    evidence must never become plugin source or a generated downloadable artifact.
    """

    expanded = output_dir.expanduser()
    if expanded.is_symlink():
        raise ValueError("output directory cannot be a symbolic link")
    resolved = expanded.resolve()
    git_root = _find_git_root(plugin_root)
    if git_root is not None and (resolved == git_root or git_root in resolved.parents):
        raise ValueError(
            f"output directory must be outside the Git workspace: {git_root}"
        )
    return prepare_private_directory(resolved)


def _safe_component(value: str) -> str:
    cleaned = SAFE_COMPONENT_RE.sub("-", value).strip("-._")
    return cleaned or "fragment"


def _bounded_text(text: str) -> tuple[str, list[str]]:
    normalized = text.replace("\x00", "").strip()
    if len(normalized) <= MAX_TEXT_CHARACTERS:
        return normalized, []
    return normalized[:MAX_TEXT_CHARACTERS], ["text_truncated_at_character_limit"]


def _text_fragment(text: str, *, kind: str, value: str | int) -> dict[str, Any]:
    bounded, limitations = _bounded_text(text)
    return {
        "locator": {"kind": kind, "value": value},
        "text": bounded,
        "limitations": limitations,
    }


def _embedded_text_is_sufficient(text: str) -> bool:
    """Reject mechanically weak PDF text layers before substantive review.

    The gate measures only extracted character coverage. It does not infer a
    document type, contribution regime, or legal meaning.
    """

    normalized = " ".join(text.split())
    alpha_count = sum(character.isalpha() for character in normalized)
    return len(normalized) >= MIN_EMBEDDED_TEXT_CHARACTERS and alpha_count >= max(
        MIN_EMBEDDED_ALPHA_CHARACTERS,
        len(normalized) // 8,
    )


def _vendor_module_candidates() -> list[Path]:
    """Return source, umbrella-package, and standalone vendor module roots."""

    plugin_root = Path(__file__).resolve().parents[1]
    candidates = [
        plugin_root.parent / "_shared" / "vendor" / "modules",
        plugin_root / "vendor" / "modules",
    ]
    if len(plugin_root.parents) > 1:
        candidates.insert(1, plugin_root.parents[1] / "vendor" / "modules")
    return candidates


def _load_ocr_adapter() -> Any | None:
    """Load Vera's optional OCR adapter without importing Paddle eagerly."""

    for candidate in _vendor_module_candidates():
        if not candidate.is_dir():
            continue
        if str(candidate) not in sys.path:
            sys.path.insert(0, str(candidate))
        break
    try:
        return importlib.import_module("vera_ocr")
    except ModuleNotFoundError as exc:
        if exc.name == "vera_ocr":
            return None
        raise


def _ocr_result_dict(result: Any) -> dict[str, Any]:
    """Normalize the shared adapter result for inventory serialization."""

    return {
        "text": str(getattr(result, "text", "")),
        "status": str(getattr(result, "status", "inference_failed")),
        "engine": str(getattr(result, "engine", "paddleocr")),
        "language": str(getattr(result, "language", "")),
        "line_count": int(getattr(result, "line_count", 0)),
        "warnings": list(getattr(result, "warnings", ())),
        "model_source": str(getattr(result, "model_source", "")),
        "network_used": bool(getattr(result, "network_used", False)),
        "runtime_versions": list(getattr(result, "runtime_versions", ())),
        "model_names": list(getattr(result, "model_names", ())),
        "model_revisions": list(getattr(result, "model_revisions", ())),
    }


def _run_local_ocr(
    image_bytes: bytes,
    *,
    language: str,
    cache_dir: Path | None,
    allow_model_download: bool,
    detection_model_dir: Path | None,
    recognition_model_dir: Path | None,
) -> dict[str, Any]:
    """Run the optional local Paddle adapter and return a non-throwing status."""

    adapter = _load_ocr_adapter()
    if adapter is None:
        return {
            "text": "",
            "status": "runtime_unavailable",
            "engine": "paddleocr",
            "language": language,
            "line_count": 0,
            "warnings": ["vera_ocr_adapter_unavailable"],
            "model_source": "",
            "network_used": False,
            "runtime_versions": [],
            "model_names": [],
            "model_revisions": [],
        }
    result = adapter.extract_text_from_image_bytes(
        image_bytes,
        language=language,
        cache_dir=cache_dir,
        allow_model_download=allow_model_download,
        detection_model_dir=detection_model_dir,
        recognition_model_dir=recognition_model_dir,
    )
    return _ocr_result_dict(result)


def _render_pdf_page(path: Path, page_index: int) -> bytes:
    """Render one PDF page to PNG bytes for local OCR."""

    fitz = importlib.import_module("fitz")
    document = fitz.open(path)
    try:
        page = document.load_page(page_index)
        pixmap = page.get_pixmap(dpi=300, alpha=False)
        return bytes(pixmap.tobytes("png"))
    finally:
        document.close()


def _ocr_limitation(status: str) -> str:
    return OCR_STATUS_LIMITATIONS.get(status, "ocr_no_text_detected")


def _apply_ocr_result(
    fragment: dict[str, Any], result: dict[str, Any]
) -> tuple[dict[str, Any], bool]:
    """Apply useful OCR text while retaining provenance and review limits."""

    ocr_metadata = {
        "engine": result.get("engine", "paddleocr"),
        "language": result.get("language", ""),
        "line_count": result.get("line_count", 0),
        "status": result.get("status", "inference_failed"),
        "warnings": result.get("warnings", []),
        "model_source": result.get("model_source", ""),
        "network_used": bool(result.get("network_used", False)),
        "runtime_versions": result.get("runtime_versions", []),
        "model_names": result.get("model_names", []),
        "model_revisions": result.get("model_revisions", []),
    }
    embedded_text = str(fragment.get("text", ""))
    if embedded_text:
        ocr_metadata["embedded_text_character_count"] = len(embedded_text)
        ocr_metadata["embedded_text_sha256"] = _sha256_bytes(
            embedded_text.encode("utf-8")
        )
    text = str(result.get("text", "")).strip()
    if result.get("status") != "ok" or not text:
        fragment["limitations"].append(_ocr_limitation(str(result.get("status", ""))))
        fragment["ocr"] = ocr_metadata
        return fragment, False
    replacement = _text_fragment(
        text,
        kind=str(fragment["locator"]["kind"]),
        value=fragment["locator"]["value"],
    )
    replacement["extraction_method"] = "paddle_ocr"
    replacement["limitations"] = list(
        dict.fromkeys(
            [
                *fragment["limitations"],
                *replacement["limitations"],
                OCR_VISUAL_CONFIRMATION_LIMITATION,
            ]
        )
    )
    replacement["ocr"] = ocr_metadata
    return replacement, True


def _extract_pdf(
    path: Path,
    *,
    enable_ocr: bool = True,
    ocr_language: str = "it",
    allow_ocr_model_download: bool = False,
    ocr_cache_dir: Path | None = None,
    ocr_detection_model_dir: Path | None = None,
    ocr_recognition_model_dir: Path | None = None,
) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    reader = PdfReader(str(path))
    limitations: list[str] = []
    if reader.is_encrypted:
        decrypted = reader.decrypt("")
        if not decrypted:
            return [], {"page_count": len(reader.pages)}, ["password_protected_pdf"]
    fragments: list[dict[str, Any]] = []
    ocr_attempted_pages = 0
    ocr_successful_pages = 0
    ocr_network_used = False
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except (FileNotDecryptedError, PdfReadError, KeyError, ValueError):
            text = ""
            limitations.append(f"page_{page_number}_text_extraction_failed")
        fragment = _text_fragment(text, kind="page", value=page_number)
        fragment["extraction_method"] = "embedded_text" if fragment["text"] else "none"
        embedded_text_sufficient = _embedded_text_is_sufficient(fragment["text"])
        if fragment["text"] and not embedded_text_sufficient:
            fragment["limitations"].append("embedded_text_below_ocr_quality_threshold")
        if not embedded_text_sufficient and enable_ocr:
            ocr_attempted_pages += 1
            try:
                image_bytes = _render_pdf_page(path, page_number - 1)
            except (
                ImportError,
                ModuleNotFoundError,
                OSError,
                RuntimeError,
                ValueError,
            ):
                fragment["limitations"].append("ocr_page_render_unavailable")
            else:
                result = _run_local_ocr(
                    image_bytes,
                    language=ocr_language,
                    cache_dir=ocr_cache_dir,
                    allow_model_download=allow_ocr_model_download,
                    detection_model_dir=ocr_detection_model_dir,
                    recognition_model_dir=ocr_recognition_model_dir,
                )
                ocr_network_used = ocr_network_used or bool(result.get("network_used"))
                fragment, ocr_succeeded = _apply_ocr_result(fragment, result)
                ocr_successful_pages += int(ocr_succeeded)
        if not fragment["text"]:
            fragment["limitations"].append("empty_text_possible_scan")
            if not enable_ocr:
                fragment["limitations"].append("ocr_disabled")
        elif not embedded_text_sufficient and not enable_ocr:
            fragment["limitations"].append("ocr_disabled")
        limitations.extend(
            f"page_{page_number}_{limitation}" for limitation in fragment["limitations"]
        )
        fragments.append(fragment)
    metadata = {
        "page_count": len(reader.pages),
        "ocr_enabled": enable_ocr,
        "ocr_language": ocr_language,
        "ocr_attempted_page_count": ocr_attempted_pages,
        "ocr_successful_page_count": ocr_successful_pages,
        "ocr_network_used": ocr_network_used,
    }
    return fragments, metadata, limitations


def _extract_image(
    path: Path,
    *,
    enable_ocr: bool,
    ocr_language: str,
    allow_ocr_model_download: bool,
    ocr_cache_dir: Path | None,
    ocr_detection_model_dir: Path | None,
    ocr_recognition_model_dir: Path | None,
) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    fragments: list[dict[str, Any]] = []
    attempted_pages = 0
    successful_pages = 0
    network_used = False

    def process_page(image_bytes: bytes, page_number: int) -> None:
        nonlocal attempted_pages, network_used, successful_pages
        fragment = _text_fragment("", kind="page", value=page_number)
        fragment["extraction_method"] = "none"
        if enable_ocr:
            attempted_pages += 1
            result = _run_local_ocr(
                image_bytes,
                language=ocr_language,
                cache_dir=ocr_cache_dir,
                allow_model_download=allow_ocr_model_download,
                detection_model_dir=ocr_detection_model_dir,
                recognition_model_dir=ocr_recognition_model_dir,
            )
            fragment, succeeded = _apply_ocr_result(fragment, result)
            successful_pages += int(succeeded)
            network_used = network_used or bool(result.get("network_used", False))
        else:
            fragment["limitations"].append("ocr_disabled")
        if not fragment["text"]:
            fragment["limitations"].append("no_extractable_text")
        fragments.append(fragment)

    if path.suffix.lower() not in {".tif", ".tiff"}:
        process_page(path.read_bytes(), 1)
    else:
        try:
            image_module = importlib.import_module("PIL.Image")
        except (ImportError, ModuleNotFoundError):
            fragment = _text_fragment("", kind="page", value=1)
            fragment["extraction_method"] = "none"
            fragment["limitations"].extend(
                ["image_frame_runtime_unavailable", "no_extractable_text"]
            )
            fragments.append(fragment)
        else:
            image_errors = (
                EOFError,
                OSError,
                ValueError,
                image_module.DecompressionBombError,
            )
            try:
                image = image_module.open(path)
            except image_errors:
                fragment = _text_fragment("", kind="page", value=1)
                fragment["extraction_method"] = "none"
                fragment["limitations"].extend(
                    ["image_frame_inventory_failed", "no_extractable_text"]
                )
                fragments.append(fragment)
            else:
                try:
                    frame_count = max(1, int(getattr(image, "n_frames", 1)))
                    for frame_index in range(frame_count):
                        try:
                            image.seek(frame_index)
                            frame = image.convert("RGB")
                            buffer = io.BytesIO()
                            frame.save(buffer, format="PNG")
                        except image_errors:
                            for missing_index in range(frame_index, frame_count):
                                fragment = _text_fragment(
                                    "", kind="page", value=missing_index + 1
                                )
                                fragment["extraction_method"] = "none"
                                fragment["limitations"].extend(
                                    [
                                        "image_frame_extraction_failed",
                                        "no_extractable_text",
                                    ]
                                )
                                fragments.append(fragment)
                            break
                        process_page(buffer.getvalue(), frame_index + 1)
                finally:
                    image.close()

    limitations = [
        f"page_{fragment['locator']['value']}_{item}"
        for fragment in fragments
        for item in fragment["limitations"]
    ]
    metadata = {
        "page_count": len(fragments),
        "ocr_enabled": enable_ocr,
        "ocr_language": ocr_language,
        "ocr_attempted_page_count": attempted_pages,
        "ocr_successful_page_count": successful_pages,
        "ocr_network_used": network_used,
    }
    return fragments, metadata, limitations


def _table_text(table: Any) -> str:
    rows: list[str] = []
    for row in table.rows:
        rows.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(rows)


def _extract_docx(path: Path) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(_table_text(table) for table in document.tables)
    fragment = _text_fragment("\n".join(parts), kind="document", value=1)
    if not fragment["text"]:
        fragment["limitations"].append("empty_text")
    return [fragment], {"paragraph_count": len(document.paragraphs)}, []


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _extract_workbook(
    path: Path,
) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    fragments: list[dict[str, Any]] = []
    limitations: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            rows: list[str] = []
            cells_seen = 0
            truncated = False
            for row in worksheet.iter_rows(values_only=True):
                remaining = MAX_CELLS_PER_SHEET - cells_seen
                if remaining <= 0:
                    truncated = True
                    break
                values = [_cell_text(value) for value in row[:remaining]]
                cells_seen += len(values)
                rows.append("\t".join(values).rstrip())
                if cells_seen >= MAX_CELLS_PER_SHEET:
                    truncated = True
                    break
            fragment = _text_fragment(
                "\n".join(rows), kind="sheet", value=worksheet.title
            )
            if truncated:
                fragment["limitations"].append("sheet_truncated_at_cell_limit")
                limitations.append(f"sheet_{worksheet.title}_truncated")
            fragments.append(fragment)
    finally:
        workbook.close()
    return fragments, {"sheet_names": workbook.sheetnames}, limitations


def _extract_plain_text(
    path: Path,
) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    if path.suffix.lower() == ".json":
        try:
            payload = json.loads(text)
        except json.JSONDecodeError:
            pass
        else:
            text = json.dumps(payload, ensure_ascii=False, indent=2, default=str)
    fragment = _text_fragment(text, kind="document", value=1)
    return [fragment], {}, []


def _email_part_text(part: Any) -> str:
    raw = part.get_payload(decode=True)
    if raw is None:
        payload = part.get_payload()
        return payload if isinstance(payload, str) else ""
    charset = part.get_content_charset() or "utf-8"
    try:
        return raw.decode(charset, errors="replace")
    except LookupError:
        return raw.decode("utf-8", errors="replace")


def _extract_email(
    path: Path,
) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    header_names = (
        "Date",
        "From",
        "To",
        "Cc",
        "Subject",
        "Message-ID",
        "In-Reply-To",
        "References",
    )
    headers = {
        name: str(message[name])
        for name in header_names
        if message.get(name) is not None
    }
    bodies: list[str] = []
    html_fallback: list[str] = []
    attachments: list[dict[str, Any]] = []
    for part in message.walk():
        if part.is_multipart():
            continue
        filename = part.get_filename()
        if part.get_content_disposition() == "attachment" or filename:
            raw = part.get_payload(decode=True) or b""
            attachments.append(
                {
                    "filename": filename or "unnamed-attachment",
                    "media_type": part.get_content_type(),
                    "size_bytes": len(raw),
                    "sha256": _sha256_bytes(raw),
                }
            )
            continue
        if part.get_content_type() == "text/plain":
            bodies.append(_email_part_text(part))
        elif part.get_content_type() == "text/html":
            html_fallback.append(_email_part_text(part))
    body = "\n\n".join(bodies or html_fallback)
    header_text = "\n".join(f"{name}: {value}" for name, value in headers.items())
    fragment = _text_fragment(
        f"{header_text}\n\n{body}".strip(), kind="message", value=1
    )
    limitations = ["email_attachments_inventory_only"] if attachments else []
    return (
        [fragment],
        {"email_headers": headers, "email_attachments": attachments},
        limitations,
    )


def _extract_file(
    path: Path,
    *,
    enable_ocr: bool,
    ocr_language: str,
    allow_ocr_model_download: bool,
    ocr_cache_dir: Path | None,
    ocr_detection_model_dir: Path | None,
    ocr_recognition_model_dir: Path | None,
) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(
            path,
            enable_ocr=enable_ocr,
            ocr_language=ocr_language,
            allow_ocr_model_download=allow_ocr_model_download,
            ocr_cache_dir=ocr_cache_dir,
            ocr_detection_model_dir=ocr_detection_model_dir,
            ocr_recognition_model_dir=ocr_recognition_model_dir,
        )
    if suffix in SUPPORTED_IMAGE_SUFFIXES:
        return _extract_image(
            path,
            enable_ocr=enable_ocr,
            ocr_language=ocr_language,
            allow_ocr_model_download=allow_ocr_model_download,
            ocr_cache_dir=ocr_cache_dir,
            ocr_detection_model_dir=ocr_detection_model_dir,
            ocr_recognition_model_dir=ocr_recognition_model_dir,
        )
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_workbook(path)
    if suffix == ".eml":
        return _extract_email(path)
    if suffix in SUPPORTED_TEXT_SUFFIXES:
        return _extract_plain_text(path)
    return [], {}, ["unsupported_file_type"]


def _fragment_name(document_id: str, locator: dict[str, Any]) -> str:
    kind = _safe_component(str(locator["kind"]))
    value = _safe_component(str(locator["value"]))
    return f"{document_id}__{kind}-{value}.txt"


def _iter_input_files(input_dir: Path) -> Iterable[Path]:
    return sorted(
        (
            path
            for path in input_dir.rglob("*")
            if (path.is_file() or path.is_symlink())
            and not any(part.startswith(".") for part in path.parts)
        ),
        key=lambda path: path.relative_to(input_dir).as_posix().casefold(),
    )


def _inventory_row(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "document_id": record["document_id"],
        "relative_path": record["relative_path"],
        "sha256": record["sha256"],
        "size_bytes": record["size_bytes"],
        "media_type": record["media_type"],
        "readability": record["readability"],
        "duplicate_of": record["duplicate_of"] or "",
        "fragment_count": record["fragment_count"],
        "limitations": "; ".join(record["limitations"]),
    }


def extract_case_documents(
    input_dir: Path,
    output_dir: Path,
    *,
    enable_ocr: bool = True,
    ocr_language: str = "it",
    allow_ocr_model_download: bool = False,
    ocr_cache_dir: Path | None = None,
    ocr_detection_model_dir: Path | None = None,
    ocr_recognition_model_dir: Path | None = None,
    visual_confirmation_methods: dict[Path, str] | None = None,
    excluded_paths: set[Path] | None = None,
    ocr_excluded_paths: set[Path] | None = None,
) -> ExtractionResult:
    """Inventory evidence; path exclusions can keep visual controls out of OCR."""

    input_dir = input_dir.expanduser().resolve()
    if not input_dir.is_dir():
        raise FileNotFoundError(f"input folder does not exist: {input_dir}")
    capture_methods = {
        path.expanduser().resolve(): method
        for path, method in (visual_confirmation_methods or {}).items()
    }
    excluded = {path.expanduser().resolve() for path in (excluded_paths or set())}
    ocr_excluded = {
        path.expanduser().resolve() for path in (ocr_excluded_paths or set())
    }
    prepare_private_directory(output_dir)
    extracted_dir = output_dir / "extracted"
    prepare_private_directory(extracted_dir)

    documents: list[dict[str, Any]] = []
    fragments_index: list[dict[str, Any]] = []
    hash_owner: dict[str, str] = {}
    extraction_errors: list[dict[str, str]] = []

    evidence_paths = [
        path for path in _iter_input_files(input_dir) if path.resolve() not in excluded
    ]
    for index, path in enumerate(evidence_paths, start=1):
        document_id = f"DOC-{index:03d}"
        relative_path = path.relative_to(input_dir).as_posix()
        is_symlink = path.is_symlink()
        raw = (
            path.readlink().as_posix().encode("utf-8")
            if is_symlink
            else path.read_bytes()
        )
        digest = _sha256_bytes(raw)
        duplicate_of = hash_owner.get(digest)
        if duplicate_of is None:
            hash_owner[digest] = document_id
        metadata: dict[str, Any] = {}
        limitations: list[str] = []
        fragments: list[dict[str, Any]] = []
        try:
            if is_symlink:
                limitations = ["symlink_not_followed"]
            else:
                fragments, metadata, limitations = _extract_file(
                    path,
                    enable_ocr=enable_ocr and path.resolve() not in ocr_excluded,
                    ocr_language=ocr_language,
                    allow_ocr_model_download=allow_ocr_model_download,
                    ocr_cache_dir=ocr_cache_dir,
                    ocr_detection_model_dir=ocr_detection_model_dir,
                    ocr_recognition_model_dir=ocr_recognition_model_dir,
                )
                capture_method = capture_methods.get(path.resolve())
                if capture_method:
                    # Browser text is mechanically linked to a captured page image,
                    # but only a human comparison can confirm its evidentiary use.
                    for fragment in fragments:
                        fragment["extraction_method"] = capture_method
                        fragment["limitations"] = sorted(
                            {
                                *fragment.get("limitations", []),
                                BROWSER_VISUAL_CONFIRMATION_LIMITATION,
                            }
                        )
                    limitations = sorted(
                        {*limitations, BROWSER_VISUAL_CONFIRMATION_LIMITATION}
                    )
        except (
            BadZipFile,
            FileNotDecryptedError,
            InvalidFileException,
            PackageNotFoundError,
            PdfReadError,
            PermissionError,
            ValueError,
        ) as exc:
            limitations = ["extraction_failed"]
            extraction_errors.append(
                {
                    "document_id": document_id,
                    "relative_path": relative_path,
                    "error_type": type(exc).__name__,
                }
            )
            LOGGER.warning(
                "Could not extract %s (%s)", relative_path, type(exc).__name__
            )

        readable_fragments = 0
        for fragment in fragments:
            locator = fragment["locator"]
            evidence_id = f"{document_id}#{locator['kind']}-{locator['value']}"
            fragment_path = extracted_dir / _fragment_name(document_id, locator)
            write_private_text(fragment_path, fragment["text"] + "\n")
            if fragment["text"].strip():
                readable_fragments += 1
            fragments_index.append(
                {
                    "evidence_id": evidence_id,
                    "document_id": document_id,
                    "locator": locator,
                    "text_path": fragment_path.relative_to(output_dir).as_posix(),
                    "text_sha256": _sha256_bytes(fragment["text"].encode("utf-8")),
                    "character_count": len(fragment["text"]),
                    "extraction_method": fragment.get(
                        "extraction_method", "native_text"
                    ),
                    "limitations": fragment["limitations"],
                    **({"ocr": fragment["ocr"]} if fragment.get("ocr") else {}),
                }
            )

        if "unsupported_file_type" in limitations:
            readability = "unsupported"
        elif readable_fragments:
            readability = "text_readable"
        elif fragments:
            readability = "no_extractable_text"
        else:
            readability = "unreadable"

        record = {
            "document_id": document_id,
            "relative_path": relative_path,
            "sha256": digest,
            "media_type": mimetypes.guess_type(path.name)[0]
            or "application/octet-stream",
            "suffix": path.suffix.lower(),
            "size_bytes": path.lstat().st_size,
            "readability": readability,
            "duplicate_of": duplicate_of,
            "fragment_count": len(fragments),
            "limitations": sorted(set(limitations)),
            **metadata,
        }
        documents.append(record)

    inventory = {
        "schema_version": "1.0",
        "plugin": "previdenza-inps",
        "input_dir": input_dir.as_posix(),
        "output_dir": output_dir.as_posix(),
        "document_count": len(documents),
        "readable_document_count": sum(
            document["readability"] == "text_readable" for document in documents
        ),
        "documents": documents,
        "evidence_fragments": fragments_index,
        "ocr": {
            "enabled": enable_ocr,
            "language": ocr_language,
            "model_download_allowed": allow_ocr_model_download,
            "attempted_page_count": sum(
                int(document.get("ocr_attempted_page_count", 0))
                for document in documents
            ),
            "successful_page_count": sum(
                int(document.get("ocr_successful_page_count", 0))
                for document in documents
            ),
            "network_used": any(
                bool(document.get("ocr_network_used")) for document in documents
            ),
            "visual_confirmation_required_fragment_count": sum(
                fragment.get("extraction_method") == "paddle_ocr"
                or fragment.get("extraction_method") == "browser_visible_text"
                or bool(
                    {
                        "embedded_text_below_ocr_quality_threshold",
                        OCR_VISUAL_CONFIRMATION_LIMITATION,
                        BROWSER_VISUAL_CONFIRMATION_LIMITATION,
                    }.intersection(map(str, fragment.get("limitations", [])))
                )
                for fragment in fragments_index
            ),
        },
        "semantic_classification": "not_performed",
    }
    inventory_path = write_json(output_dir / "file_inventory.json", inventory)

    inventory_csv_path = output_dir / "file_inventory.csv"
    fieldnames = (
        list(_inventory_row(documents[0]).keys())
        if documents
        else [
            "document_id",
            "relative_path",
            "sha256",
            "size_bytes",
            "media_type",
            "readability",
            "duplicate_of",
            "fragment_count",
            "limitations",
        ]
    )
    with inventory_csv_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(_inventory_row(record) for record in documents)
    mark_private_file(inventory_csv_path)

    extraction_report = {
        "schema_version": "1.0",
        "status": (
            "blocked_input"
            if not documents or not inventory["readable_document_count"]
            else (
                "partial_evidence"
                if any(record["limitations"] for record in documents)
                else "ready_for_fact_structuring"
            )
        ),
        "document_count": len(documents),
        "readable_document_count": inventory["readable_document_count"],
        "duplicate_count": sum(bool(record["duplicate_of"]) for record in documents),
        "unsupported_count": sum(
            record["readability"] == "unsupported" for record in documents
        ),
        "ocr_attempted_page_count": inventory["ocr"]["attempted_page_count"],
        "ocr_successful_page_count": inventory["ocr"]["successful_page_count"],
        "ocr_network_used": inventory["ocr"]["network_used"],
        "visual_confirmation_required_fragment_count": inventory["ocr"][
            "visual_confirmation_required_fragment_count"
        ],
        "extraction_errors": extraction_errors,
        "notes": [
            "Filename and document content were not assigned a legal or contribution regime.",
            "PaddleOCR text keeps its original page locator and requires visual human confirmation.",
            "Visible text captured from an authenticated browser tab requires comparison with its captured page image.",
            "OCR model downloads are disabled unless explicitly authorized for the run.",
            "Symbolic links are inventoried but never followed outside the case folder.",
        ],
    }
    extraction_report_path = write_json(
        output_dir / "extraction_report.json", extraction_report
    )

    extracted_evidence_path = output_dir / "extracted_evidence.md"
    lines = [
        "# Extracted evidence",
        "",
        "> Mechanical extraction only. No legal or contribution classification was performed.",
        "",
    ]
    fragments_by_document: dict[str, list[dict[str, Any]]] = {}
    for fragment in fragments_index:
        fragments_by_document.setdefault(fragment["document_id"], []).append(fragment)
    for document in documents:
        lines.extend(
            [
                f"## {document['document_id']} — {document['relative_path']}",
                "",
                f"Readability: `{document['readability']}`",
                "",
            ]
        )
        for fragment in fragments_by_document.get(document["document_id"], []):
            text = read_fragment_text(output_dir, fragment)
            lines.extend(
                [
                    f"### {fragment['evidence_id']}",
                    "",
                    f"Extraction method: `{fragment['extraction_method']}`",
                    "",
                ]
            )
            if fragment["limitations"]:
                lines.extend(
                    [
                        "Limitations: "
                        + ", ".join(
                            f"`{limitation}`" for limitation in fragment["limitations"]
                        ),
                        "",
                    ]
                )
            lines.extend([text, ""])
    write_private_text(extracted_evidence_path, "\n".join(lines).rstrip() + "\n")

    return ExtractionResult(
        inventory=inventory,
        inventory_path=inventory_path,
        inventory_csv_path=inventory_csv_path,
        extraction_report_path=extraction_report_path,
        extracted_evidence_path=extracted_evidence_path,
    )


def read_fragment_text(output_dir: Path, fragment: dict[str, Any]) -> str:
    """Read one inventory fragment from its output-relative path."""

    relative = Path(str(fragment["text_path"]))
    if relative.is_absolute() or ".." in relative.parts:
        raise ValueError(f"unsafe evidence fragment path: {relative}")
    return (output_dir / relative).read_text(encoding="utf-8").strip()
