"""Reviewed or clearly labelled draft root-cause variance report writer."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import polars as pl
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from ibcs_titles import build_ibcs_title, measure_line_segments
from PIL import Image, ImageDraw, ImageFont

__all__ = ["write_root_cause_client_report"]


MEASURE_COLUMNS = {
    "bridge_level",
    "bridge_dimensions",
    "variance_type",
    "variance_amount",
    "amount_baseline",
    "amount_comparison",
    "units_baseline",
    "units_comparison",
    "bridge_unique_value_weight",
}
TOTAL_VALUES = {"", "all", "__total", "total", "none", "null"}


def _language(recipe: dict[str, Any]) -> str:
    raw_language = str(recipe.get("language") or "en").lower().replace("_", "-")
    return raw_language.split("-", maxsplit=1)[0]


def _currency_note(recipe: dict[str, Any], labels: dict[str, str]) -> str:
    """Return the report note for the effective currency assumption."""

    currency = str(recipe.get("options", {}).get("currency") or "").strip()
    if currency:
        return labels["currency_note"].format(currency=currency)
    return labels["source_units"]


def _comparison_metadata(recipe: dict[str, Any], language: str) -> dict[str, str]:
    """Return report wording for the effective comparison."""

    mappings = recipe.get("mappings") or {}
    options = recipe.get("options") or {}
    baseline = str(mappings.get("baseline_period") or "baseline")
    comparison = str(mappings.get("comparison_period") or "comparison")
    basis = str(options.get("comparison_basis") or "")
    mode = str(options.get("period_comparison_mode") or "")
    baseline_upper = baseline.upper()
    comparison_upper = comparison.upper()
    if basis == "period":
        if language == "it":
            comparison_name = "periodo corrente"
            baseline_name = (
                "anno precedente"
                if mode in {"rolling_period", "year_to_date"}
                else "periodo precedente"
            )
        elif language == "es":
            comparison_name = "periodo actual"
            baseline_name = (
                "periodo del año anterior"
                if mode in {"rolling_period", "year_to_date"}
                else "periodo anterior"
            )
        else:
            comparison_name = "current period"
            baseline_name = (
                "prior-year period"
                if mode in {"rolling_period", "year_to_date"}
                else "prior period"
            )
    elif baseline_upper in {"PL", "PLAN"} and comparison_upper in {"AC", "ACTUAL"}:
        baseline_name = "Plan"
        comparison_name = "Real" if language == "es" else "Actual"
    else:
        baseline_name = baseline
        comparison_name = comparison
    separator = " frente a " if language == "es" else " vs "
    return {
        "baseline_label": baseline,
        "comparison_label": comparison,
        "baseline_name": baseline_name,
        "comparison_name": comparison_name,
        "comparison": (
            f"{comparison_name}{separator}{baseline_name} "
            f"({comparison}{separator}{baseline})"
        ),
    }


def _text(language: str) -> dict[str, str]:
    if language == "it":
        return {
            "title": "Analisi delle cause della varianza vendite",
            "draft_title": "Bozza di analisi delle varianze vendite",
            "draft_notice": (
                "Bozza di lavoro: il bridge di riferimento è selezionato con criteri "
                "meccanici e non costituisce una causa aziendale approvata."
            ),
            "accounting_controls": "Perimetro e controlli contabili",
            "accounting_field": "Controllo",
            "accounting_value": "Esito",
            "subtitle": "Actual vs Plan (AC vs PL)",
            "summary": "Sintesi",
            "source_data": "Dati di supporto principali",
            "reading_notes": "Note di lettura",
            "bridge_summary": "Bridge di riferimento",
            "product_line_drilldown": "Dettaglio della prima riga",
            "mixed_deep_dive": "Approfondimento misto",
            "chart_1": "Fonte 1 - Bridge di riferimento",
            "chart_2": "Fonte 2 - Dettaglio della prima riga",
            "chart_3": "Fonte 3 - Approfondimento per area e linea di prodotto",
            "chart_small_multiples": "Bridge standard per {dimension}",
            "chart_pvm_ladder": "Bridge Prezzo / Unità / Mix",
            "drilldown_findings": "Cosa emerge dai drilldown selezionati",
            "source_units": (
                "Gli importi sono in unità della sorgente: il file non "
                "fornisce una valuta esplicita."
            ),
            "currency_note": (
                "Gli importi sono presentati in {currency}; usare una valuta "
                "diversa solo se indicata dall'utente o dal file sorgente."
            ),
            "price_only": (
                "La varianza deriva solo dal prezzo: volume e mix sono pari a "
                "zero perché le unità {baseline_name} e {comparison_name} "
                "coincidono al livello di calcolo."
            ),
            "component_note": (
                "Il bridge standard è dominato da {dominant_type} "
                "({dominant_amount}); componenti principali: {components}."
            ),
            "residual_note": (
                "Le righe del bridge delle cause sono residuali: una riga "
                "successiva non va letta come totale assoluto della relativa "
                "dimensione."
            ),
            "source_caption": "Sintesi dei dati di supporto selezionati.",
            "drilldown_caption": (
                "Dettaglio dei contributi emersi dai drilldown selezionati."
            ),
            "chart_footer": (
                "Driver selezionati in sequenza; il saldo residuo è "
                "riconciliato in Altro."
            ),
            "chart_summary_title": "Bridge di sintesi delle cause",
            "chart_summary_subtitle": (
                "{comparison}, driver selezionati dall'analisi root-cause"
            ),
            "chart_drilldown_title": "Drilldown: dettaglio di {label}",
            "chart_drilldown_subtitle": (
                "{comparison}, dettaglio della riga selezionata"
            ),
            "chart_mixed_title": "Approfondimento: bridge delle cause misto",
            "chart_mixed_subtitle": "{comparison}, sequenza con dimensioni diverse",
            "summary_intro": (
                "La differenza tra {comparison_name} e {baseline_name} è {delta}. "
                "Il bridge di sintesi riconcilia "
                "il movimento con {driver_count} driver selezionati ({items}) "
                "e un residuo di {residual}."
            ),
            "draft_summary_intro": (
                "La differenza calcolata tra {comparison_name} e {baseline_name} è "
                "{delta}. Per questa bozza il motore mostra il bridge di riferimento "
                "più compatto fra quelli riconciliati: {driver_count} righe ({items}) "
                "e residuo {residual}. La scelta non è un giudizio sulle cause."
            ),
            "drilldown_intro": (
                "Il drilldown della riga principale dettaglia il contributo: "
                "{items}. Il contributo è quindi concentrato soprattutto su "
                "{top_label}."
            ),
            "draft_drilldown_intro": (
                "Il dettaglio della prima riga mostra i contributi calcolati: "
                "{items}. La rilevanza aziendale richiede conferma professionale."
            ),
            "mixed_intro": (
                "Una vista alternativa calcolata presenta la sequenza {items}. "
                "La vista lascia {residual} in residuo e deve essere letta "
                "secondo la logica residuale delle righe."
            ),
            "draft_mixed_intro": (
                "È disponibile anche una sequenza a dimensioni miste ({items}) con "
                "residuo {residual}; va valutata come spiegazione alternativa."
            ),
            "bridge_reading": "Mostra la chiusura matematica della sequenza.",
            "drilldown_reading": "Scompone la prima riga selezionata.",
            "mixed_reading": "Presenta una sequenza residuale alternativa.",
            "data_key": "Dato chiave",
            "reading": "Lettura",
            "residual": "residuo",
            "source_col": "Fonte",
            "analysis_area": "Area di analisi",
            "useful_reading": "Lettura utile",
            "chart_1_reading": (
                "La sequenza di riferimento ({items}) chiude il movimento "
                "{comparison} con residuo finale {residual}."
            ),
            "chart_2_reading": (
                "Il dettaglio della prima riga include {top_label} e gli altri "
                "contributi calcolati per linea prodotto."
            ),
            "chart_3_reading": (
                "La sequenza alternativa a dimensioni miste include {items}."
            ),
            "chart_small_multiples_reading": (
                "Il bridge standard per {dimension} presenta i contributi "
                "calcolati di maggiore valore assoluto: {items}."
            ),
            "chart_pvm_ladder_reading": (
                "La stessa varianza è letta a tre livelli: totale combinato, "
                "Prezzo separato da Unità e Mix, e Prezzo / Unità / Mix. "
                "Componenti principali: {items}."
            ),
            "chart_1_caption": (
                "Il bridge dimostra la chiusura matematica; la scelta della sequenza "
                "richiede revisione professionale."
            ),
            "chart_1_caption_approved": (
                "Il bridge dimostra la chiusura matematica; la scelta della sequenza "
                "è registrata nella revisione professionale."
            ),
            "chart_2_caption": (
                "Il dettaglio espone i contributi misurati della prima riga selezionata."
            ),
            "chart_3_caption": (
                "Questa vista va letta come sequenza residuale: le righe "
                "successive sono al netto delle righe precedenti."
            ),
            "chart_small_multiples_caption": (
                "Ogni pannello ripete il bridge compatto Prezzo / "
                "Unità e Mix / Saldo; la dimensione separa i pannelli."
            ),
            "chart_pvm_ladder_caption": (
                "La scala e i totali sono gli stessi in ogni pannello: cambia "
                "solo il livello di decomposizione della varianza."
            ),
            "pvm_ladder_source": "Lettura Prezzo / Unità / Mix",
            "pvm_ladder_reading": (
                "Confronta tre decomposizioni dello stesso movimento."
            ),
            "small_multiples_source": "Bridge standard in pannelli",
            "small_multiples_reading": (
                "Mostra lo stesso bridge standard per ciascun elemento della "
                "dimensione selezionata."
            ),
            "drilldown_balance": (
                "Il saldo include compensazioni fra contributi positivi e " "negativi."
            ),
            "drilldown_concentration": "Il contributo non dipende da un solo elemento.",
            "drilldown_composition": "Il residuo è concentrato su pochi contributi.",
            "top_contributions": "Principali contributi: {items}",
            "root_cause_bridge": "Bridge delle cause",
            "dimension": "dimensione",
            "dimension_product": "prodotto",
            "dimension_region": "area",
            "dimension_subregion": "sottoarea",
            "dimension_customer": "cliente",
            "dimension_channel": "canale",
            "sales": "Vendite",
            "total": "Totale",
            "variance_type_price": "Prezzo",
            "variance_type_units": "Unità",
            "variance_type_volume": "Volume",
            "variance_type_mix": "Mix",
            "variance_type_other": "Altro",
            "variance_type_balance": "Saldo",
            "variance_type_net_sales": "Vendite nette",
            "variance_type_discount": "Sconto",
            "variance_type_cogs": "Costo del venduto",
            "variance_type_gross_margin": "Margine lordo",
            "variance_type_units_and_mix": "Unità e Mix",
            "variance_type_price_and_units_and_mix": "Prezzo, Unità e Mix",
            "variance_type_price_and_volume_and_mix": "Prezzo, Volume e Mix",
        }
    if language == "es":
        return {
            "title": "Análisis de las causas de la varianza de ventas",
            "draft_title": "Borrador de análisis de variaciones de ventas",
            "draft_notice": (
                "Borrador de trabajo: el puente de referencia se selecciona con "
                "criterios mecánicos y no constituye una causa empresarial aprobada."
            ),
            "accounting_controls": "Perímetro y controles contables",
            "accounting_field": "Control",
            "accounting_value": "Resultado",
            "subtitle": "Real frente a Plan (AC frente a PL)",
            "summary": "Resumen",
            "source_data": "Datos de soporte principales",
            "reading_notes": "Notas de lectura",
            "bridge_summary": "Puente de resumen",
            "product_line_drilldown": "Desglose por línea de producto",
            "mixed_deep_dive": "Análisis detallado multidimensional",
            "chart_1": "Fuente 1 - Puente de referencia",
            "chart_2": "Fuente 2 - Detalle de la primera fila",
            "chart_3": "Fuente 3 - Detalle por área y línea de producto",
            "chart_small_multiples": "Puente estándar por {dimension}",
            "chart_pvm_ladder": "Puente de Precio / Unidades / Mix",
            "drilldown_findings": "Hallazgos de los desgloses seleccionados",
            "source_units": (
                "Los importes se muestran en las unidades de la fuente porque "
                "el archivo no indica una moneda explícita."
            ),
            "currency_note": (
                "Los importes se presentan en {currency}; utilice otra moneda "
                "solo si la indica el usuario o el archivo fuente."
            ),
            "price_only": (
                "El análisis determinista solo refleja el efecto precio: el "
                "volumen y el mix son cero porque las unidades de "
                "{baseline_name} y {comparison_name} coinciden en el nivel de "
                "cálculo."
            ),
            "component_note": (
                "El puente estándar está dominado por {dominant_type} "
                "({dominant_amount}); componentes principales: {components}."
            ),
            "residual_note": (
                "Las filas del puente de causa raíz son residuales: una fila "
                "posterior no representa el total independiente de esa dimensión."
            ),
            "source_caption": "Resumen de los datos de soporte seleccionados.",
            "drilldown_caption": (
                "Detalle de las contribuciones de los desgloses seleccionados."
            ),
            "chart_footer": (
                "Los impulsores seleccionados se muestran en secuencia; el "
                "saldo residual se concilia en Otros."
            ),
            "chart_summary_title": "Puente de resumen: {label}",
            "chart_summary_subtitle": (
                "{comparison}, impulsor seleccionado del análisis de causa raíz"
            ),
            "chart_drilldown_title": "Desglose: {label}",
            "chart_drilldown_subtitle": (
                "{comparison}, detalle de la fila seleccionada"
            ),
            "chart_mixed_title": "Puente multidimensional de causa raíz",
            "chart_mixed_subtitle": (
                "{comparison}, secuencia con dimensiones distintas"
            ),
            "summary_intro": (
                "La diferencia entre {comparison_name} y {baseline_name} es "
                "{delta}. El puente de resumen concilia el movimiento con "
                "{driver_count} impulsores seleccionados ({items}) y un residual "
                "de {residual}."
            ),
            "draft_summary_intro": (
                "La diferencia calculada entre {comparison_name} y {baseline_name} "
                "es {delta}. Este borrador muestra el puente de referencia mecánico "
                "con {driver_count} filas ({items}) y residual {residual}; no es una "
                "conclusión sobre las causas empresariales."
            ),
            "drilldown_intro": (
                "El desglose de la fila principal detalla la contribución: "
                "{items}. Por tanto, la contribución se concentra principalmente "
                "en {top_label}."
            ),
            "draft_drilldown_intro": (
                "El detalle de la primera fila muestra las contribuciones calculadas: "
                "{items}. Su relevancia empresarial requiere revisión profesional."
            ),
            "mixed_intro": (
                "Una vista alternativa calculada presenta la secuencia {items}. "
                "La vista deja {residual} como saldo residual y debe leerse "
                "según la lógica residual de las filas."
            ),
            "draft_mixed_intro": (
                "También existe una secuencia de dimensiones mixtas ({items}) con "
                "residual {residual}; debe evaluarse como explicación alternativa."
            ),
            "bridge_reading": "Muestra el cierre matemático de la secuencia.",
            "drilldown_reading": "Descompone la primera fila seleccionada.",
            "mixed_reading": "Presenta una secuencia residual alternativa.",
            "data_key": "Dato clave",
            "reading": "Lectura",
            "residual": "residual",
            "source_col": "Fuente",
            "analysis_area": "Área de análisis",
            "useful_reading": "Lectura útil",
            "chart_1_reading": (
                "La secuencia de referencia ({items}) cierra el movimiento "
                "{comparison} con un residual final de {residual}."
            ),
            "chart_2_reading": (
                "El detalle de la primera fila incluye {top_label} y las demás "
                "contribuciones calculadas por línea de producto."
            ),
            "chart_3_reading": (
                "La secuencia alternativa de dimensiones mixtas incluye {items}."
            ),
            "chart_small_multiples_reading": (
                "El puente estándar por {dimension} presenta las contribuciones "
                "calculadas de mayor valor absoluto: {items}."
            ),
            "chart_pvm_ladder_reading": (
                "La misma varianza se presenta en tres niveles: total combinado, "
                "Precio separado de Unidades y Mix, y Precio / Unidades / Mix. "
                "Componentes principales: {items}."
            ),
            "chart_1_caption": (
                "El puente demuestra el cierre matemático; la selección de la "
                "secuencia requiere revisión profesional."
            ),
            "chart_1_caption_approved": (
                "El puente demuestra el cierre matemático; la selección de la "
                "secuencia queda registrada en la revisión profesional."
            ),
            "chart_2_caption": (
                "El detalle muestra las contribuciones medidas de la primera fila."
            ),
            "chart_3_caption": (
                "Esta vista debe leerse como una secuencia residual: las filas "
                "posteriores son netas de las anteriores."
            ),
            "chart_small_multiples_caption": (
                "Cada panel repite el puente compacto de Precio / Unidades y Mix / "
                "Saldo; la dimensión separa los paneles."
            ),
            "chart_pvm_ladder_caption": (
                "La escala y los totales son iguales en todos los paneles: solo "
                "cambia el nivel de descomposición de la varianza."
            ),
            "pvm_ladder_source": "Lectura de Precio / Unidades / Mix",
            "pvm_ladder_reading": (
                "Compara tres descomposiciones del mismo movimiento."
            ),
            "small_multiples_source": "Múltiplos comparativos estándar",
            "small_multiples_reading": (
                "Muestra el mismo puente estándar para cada elemento de la "
                "dimensión seleccionada."
            ),
            "drilldown_balance": (
                "El saldo incluye compensaciones entre contribuciones positivas "
                "y negativas."
            ),
            "drilldown_concentration": (
                "La contribución no depende de un solo elemento."
            ),
            "drilldown_composition": (
                "El residual se concentra en pocas contribuciones."
            ),
            "top_contributions": "Contribuciones principales: {items}",
            "root_cause_bridge": "Puente de causa raíz",
            "dimension": "dimensión",
            "sales": "Ventas",
            "total": "Total",
            "variance_type_price": "Precio",
            "variance_type_units": "Unidades",
            "variance_type_volume": "Volumen",
            "variance_type_mix": "Mix",
            "variance_type_other": "Otros",
            "variance_type_balance": "Saldo",
            "variance_type_net_sales": "Ventas netas",
            "variance_type_discount": "Descuento",
            "variance_type_cogs": "Coste de ventas",
            "variance_type_gross_margin": "Margen bruto",
            "variance_type_units_and_mix": "Unidades y Mix",
            "variance_type_price_and_units_and_mix": "Precio, Unidades y Mix",
            "variance_type_price_and_volume_and_mix": "Precio, Volumen y Mix",
        }
    return {
        "title": "Sales Variance Root-Cause Analysis",
        "draft_title": "Draft Sales Variance Analysis",
        "draft_notice": (
            "Working draft: the reference bridge is selected mechanically and is "
            "not an approved business-cause conclusion."
        ),
        "accounting_controls": "Perimeter and accounting controls",
        "accounting_field": "Control",
        "accounting_value": "Result",
        "subtitle": "Actual vs Plan",
        "summary": "Summary",
        "source_data": "Key Source Data",
        "reading_notes": "Reading Notes",
        "bridge_summary": "Summary bridge",
        "product_line_drilldown": "Product-line drilldown",
        "mixed_deep_dive": "Mixed-dimension deep dive",
        "chart_1": "Source 1 - Reference Bridge",
        "chart_2": "Source 2 - First-row Detail",
        "chart_3": "Source 3 - Area And Product-Line Detail",
        "chart_small_multiples": "Standard bridge by {dimension}",
        "chart_pvm_ladder": "Price / Units / Mix bridge",
        "drilldown_findings": "Selected Drilldown Findings",
        "source_units": (
            "Amounts are shown in source units because the workbook does not "
            "provide an explicit currency."
        ),
        "currency_note": (
            "Amounts are presented in {currency}; use another currency only "
            "when the user or source file states it."
        ),
        "price_only": (
            "The deterministic run is price-only: volume and mix are zero "
            "because {baseline_name} and {comparison_name} units match at the "
            "calculation grain."
        ),
        "component_note": (
            "The standard bridge is dominated by {dominant_type} "
            "({dominant_amount}); main components: {components}."
        ),
        "residual_note": (
            "Root-cause bridge rows are residual rows: a later row is not the "
            "standalone total for that dimension."
        ),
        "source_caption": "Summary of selected source data.",
        "drilldown_caption": "Detail of contributions from selected drilldowns.",
        "chart_footer": (
            "Selected drivers are shown in sequence; the residual balance is "
            "reconciled to Other."
        ),
        "chart_summary_title": "Summary bridge: {label}",
        "chart_summary_subtitle": "{comparison}, selected root-cause driver",
        "chart_drilldown_title": "Drilldown: {label}",
        "chart_drilldown_subtitle": "{comparison}, selected-row detail",
        "chart_mixed_title": "Mixed root-cause bridge",
        "chart_mixed_subtitle": "{comparison}, mixed-dimension sequence",
        "summary_intro": (
            "The difference between {comparison_name} and {baseline_name} is "
            "{delta}. The summary bridge reconciles "
            "the movement with {driver_count} selected drivers ({items}) and "
            "residual of {residual}."
        ),
        "draft_summary_intro": (
            "The calculated difference between {comparison_name} and "
            "{baseline_name} is {delta}. This draft shows the mechanically compact "
            "reference bridge with {driver_count} rows ({items}) and residual "
            "{residual}; it is not a business-cause conclusion."
        ),
        "drilldown_intro": (
            "The main-row drilldown details the contribution: {items}. The "
            "contribution is therefore concentrated mainly in {top_label}."
        ),
        "draft_drilldown_intro": (
            "The first-row detail shows the calculated contributions: {items}. "
            "Their business relevance requires professional review."
        ),
        "mixed_intro": (
            "An alternative calculated view presents the sequence {items}. "
            "It leaves {residual} in residual balance and must be read using "
            "the rows' residual logic."
        ),
        "draft_mixed_intro": (
            "A mixed-dimension sequence is also available ({items}) with residual "
            "{residual}; it must be assessed as an alternative explanation."
        ),
        "bridge_reading": "Shows the mathematical closure of the sequence.",
        "drilldown_reading": "Breaks down the first selected row.",
        "mixed_reading": "Presents an alternative residual sequence.",
        "data_key": "Key data",
        "reading": "Reading",
        "residual": "residual",
        "source_col": "Source",
        "analysis_area": "Analysis area",
        "useful_reading": "Useful reading",
        "chart_1_reading": (
            "The reference sequence ({items}) closes the {comparison} "
            "movement with final residual {residual}."
        ),
        "chart_2_reading": (
            "The first-row detail includes {top_label} and the other calculated "
            "product-line contributions."
        ),
        "chart_3_reading": (
            "The alternative mixed-dimension sequence includes {items}."
        ),
        "chart_small_multiples_reading": (
            "The standard bridge by {dimension} presents the calculated "
            "contributions with the largest absolute values: {items}."
        ),
        "chart_pvm_ladder_reading": (
            "The same variance is read at three levels: combined total, Price "
            "separated from Units & Mix, and Price / Units / Mix. Main "
            "components: {items}."
        ),
        "chart_1_caption": (
            "The bridge demonstrates mathematical closure; sequence selection "
            "requires professional review."
        ),
        "chart_1_caption_approved": (
            "The bridge demonstrates mathematical closure; sequence selection "
            "is recorded in the professional review."
        ),
        "chart_2_caption": (
            "The detail exposes the measured contributions of the first selected row."
        ),
        "chart_3_caption": (
            "This view should be read as a residual sequence: later rows are "
            "net of prior rows."
        ),
        "chart_small_multiples_caption": (
            "Each panel repeats the compact Price / Units & Mix / "
            "Balance bridge; the dimension separates the panels."
        ),
        "chart_pvm_ladder_caption": (
            "Scale and totals are the same in every panel: only the variance "
            "decomposition depth changes."
        ),
        "pvm_ladder_source": "Price / Units / Mix read",
        "pvm_ladder_reading": "Compares three decompositions of the same movement.",
        "small_multiples_source": "Standard small multiples",
        "small_multiples_reading": (
            "Shows the same standard bridge for each selected dimension member."
        ),
        "drilldown_balance": (
            "The balance includes offsets between positive and negative "
            "contributions."
        ),
        "drilldown_concentration": "The contribution is not dependent on one item.",
        "drilldown_composition": "The residual is concentrated in a few items.",
        "top_contributions": "Main contributions: {items}",
    }


def _variance_type_label(value: Any, labels: dict[str, str]) -> str:
    """Return a localized display label without changing the calculation key."""

    raw_value = str(value or "").strip()
    label_keys = {
        "price": "variance_type_price",
        "units": "variance_type_units",
        "volume": "variance_type_volume",
        "mix": "variance_type_mix",
        "other": "variance_type_other",
        "balance": "variance_type_balance",
        "net sales": "variance_type_net_sales",
        "discount": "variance_type_discount",
        "cogs": "variance_type_cogs",
        "gross margin": "variance_type_gross_margin",
        "units & mix": "variance_type_units_and_mix",
        "price & units & mix": "variance_type_price_and_units_and_mix",
        "price & volume & mix": "variance_type_price_and_volume_and_mix",
    }
    label_key = label_keys.get(raw_value.casefold())
    return labels.get(label_key, raw_value) if label_key else raw_value


def _dimension_label(value: Any, labels: dict[str, str]) -> str:
    """Return a localized display label for common source dimensions."""

    raw_value = str(value or "").strip()
    label_key = f"dimension_{raw_value.casefold().replace(' ', '_')}"
    return labels.get(label_key, raw_value)


def _localized_selected_label(value: Any, labels: dict[str, str]) -> str:
    """Localize the generated variance-type suffix in a selected-row label."""

    raw_value = str(value or "").strip()
    prefix, separator, variance_type = raw_value.rpartition(" - ")
    if not separator:
        return labels.get("total", raw_value) if raw_value == "Total" else raw_value
    localized_type = _variance_type_label(variance_type, labels)
    return f"{prefix}{separator}{localized_type}"


def _is_total(value: Any) -> bool:
    if value is None:
        return True
    return str(value).strip().lower() in TOTAL_VALUES


def _format_amount(value: Any) -> str:
    amount = float(value or 0.0)
    sign = "-" if amount < 0 else "+"
    magnitude = abs(amount)
    if magnitude >= 1_000_000:
        return f"{sign}{magnitude / 1_000_000:.2f}M"
    if magnitude >= 1_000:
        return f"{sign}{magnitude / 1_000:.1f}K"
    return f"{sign}{magnitude:.0f}"


def _format_delta(value: Any) -> str:
    magnitude = abs(float(value or 0.0))
    if magnitude >= 1_000_000:
        return f"{magnitude / 1_000_000:.2f}M"
    if magnitude >= 1_000:
        return f"{magnitude / 1_000:.1f}K"
    return f"{magnitude:.0f}"


def _format_residual(value: Any) -> str:
    return _format_amount(value)


def _summary_float(row: dict[str, Any], key: str) -> float:
    try:
        return float(row.get(key) or 0.0)
    except (TypeError, ValueError):
        return 0.0


def _summary_bool(row: dict[str, Any], key: str) -> bool:
    return str(row.get(key)).strip().lower() == "true"


def _split_summary_values(value: Any) -> list[str]:
    return [item.strip() for item in str(value or "").split("|") if item.strip()]


def _collect_csv_scan(path: Path) -> pl.DataFrame:
    """Read generated CSV artifacts through a lazy scan and collect once."""

    lf = pl.scan_csv(path)
    try:
        return lf.collect(engine="streaming")
    except pl.exceptions.PolarsError:
        return lf.collect()


def _read_csv(path: Path) -> pl.DataFrame:
    if not path.exists():
        return pl.DataFrame()
    return _collect_csv_scan(path)


def _row_label(
    row: dict[str, Any], display_labels: dict[str, str] | None = None
) -> str:
    row_labels: list[str] = []
    for key, value in row.items():
        if key in MEASURE_COLUMNS or _is_total(value):
            continue
        row_labels.append(str(value))
    variance_type = row.get("variance_type")
    if variance_type and not _is_total(variance_type):
        row_labels.append(_variance_type_label(variance_type, display_labels or {}))
    if row_labels:
        return " / ".join(row_labels)
    return (display_labels or {}).get("total", "Total")


def _chart_path(output_dir: Path, value: Any) -> Path | None:
    raw = str(value or "")
    if not raw:
        return None
    path = Path(raw)
    if path.exists():
        return path
    candidate = output_dir / path.name
    return candidate if candidate.exists() else None


def _select_summary(
    summary_rows: list[dict[str, Any]],
    recipe: dict[str, Any],
) -> tuple[dict[str, Any] | None, str]:
    """Return an explicitly reviewed alternative or a mechanical draft reference."""

    if not summary_rows:
        return None, "none"
    root_cause_review = recipe.get("accounting_review", {}).get("root_cause_review", {})
    if root_cause_review.get("status") == "approved":
        selected = root_cause_review.get("selected_alternative")
        for row in summary_rows:
            if row.get("alternative_result") == selected:
                return row, "explicit_professional_or_model_review"
    return (
        min(
            summary_rows,
            key=lambda row: (
                abs(_summary_float(row, "other_residual")),
                int(row.get("row_count") or 999),
                int(row.get("alternative_result") or 999),
            ),
        ),
        "mechanical_draft_reference",
    )


def _first_dimensions(row: dict[str, Any]) -> list[str]:
    sequence = _split_summary_values(row.get("selected_sequence_bridge_dimensions"))
    if not sequence:
        return []
    return [item.strip() for item in sequence[0].split(",") if item.strip()]


def _select_mixed(summary_rows: list[dict[str, Any]]) -> dict[str, Any] | None:
    mixed = [
        row
        for row in summary_rows
        if _summary_bool(row, "selected_sequence_has_mixed_dimensions")
    ]
    if not mixed:
        return None
    readable = [row for row in mixed if len(_first_dimensions(row)) > 1]
    if readable:
        return min(readable, key=lambda row: int(row.get("alternative_result") or 999))
    return min(
        mixed,
        key=lambda row: (
            abs(_summary_float(row, "other_residual")),
            int(row.get("alternative_result") or 999),
        ),
    )


def _bridge_rows(output_dir: Path, alternative: int) -> list[dict[str, Any]]:
    frame = _read_csv(output_dir / f"root_cause_bridge_alt_{alternative}.csv")
    return frame.to_dicts() if not frame.is_empty() else []


def _drilldown_rows(
    output_dir: Path,
    alternative: int,
    selected_row: int,
) -> list[dict[str, Any]]:
    frame = _read_csv(
        output_dir
        / f"root_cause_bridge_alt_{alternative}_drilldown_row_{selected_row}.csv"
    )
    return frame.to_dicts() if not frame.is_empty() else []


def _selected_amounts(row: dict[str, Any]) -> list[float]:
    amounts: list[float] = []
    for item in _split_summary_values(row.get("selected_amounts")):
        try:
            amounts.append(float(item))
        except ValueError:
            continue
    return amounts


def _top_label_amounts(
    rows: list[dict[str, Any]],
    display_labels: dict[str, str] | None = None,
    limit: int = 3,
) -> list[tuple[str, float]]:
    sorted_rows = sorted(
        rows,
        key=lambda row: abs(float(row.get("variance_amount") or 0.0)),
        reverse=True,
    )
    return [
        (_row_label(row, display_labels), float(row.get("variance_amount") or 0.0))
        for row in sorted_rows[:limit]
    ]


def _format_label_amounts(items: list[tuple[str, float]], *, limit: int = 3) -> str:
    return ", ".join(
        f"{label} {_format_amount(amount)}" for label, amount in items[:limit]
    )


def _best_product_line_items(
    output_dir: Path,
    alternative: int,
    row_index: int,
    display_labels: dict[str, str] | None = None,
) -> list[tuple[str, float]]:
    rows = _drilldown_rows(output_dir, alternative, row_index)
    return _top_label_amounts(rows, display_labels)


def _top_positive_and_negative(
    rows: list[dict[str, Any]], display_labels: dict[str, str] | None = None
) -> list[tuple[str, float]]:
    positives = sorted(
        (
            (
                _row_label(row, display_labels),
                float(row.get("variance_amount") or 0.0),
            )
            for row in rows
            if float(row.get("variance_amount") or 0.0) > 0
        ),
        key=lambda item: item[1],
        reverse=True,
    )[:2]
    negatives = sorted(
        (
            (
                _row_label(row, display_labels),
                float(row.get("variance_amount") or 0.0),
            )
            for row in rows
            if float(row.get("variance_amount") or 0.0) < 0
        ),
        key=lambda item: abs(item[1]),
        reverse=True,
    )[:1]
    return [*positives, *negatives]


def _drilldown_finding_rows(
    output_dir: Path,
    mixed_row: dict[str, Any] | None,
    labels: dict[str, str],
) -> list[list[str]]:
    if mixed_row is None:
        return []
    alternative = int(mixed_row.get("alternative_result") or 0)
    bridge_rows = _bridge_rows(output_dir, alternative)
    findings: list[list[str]] = []
    for index, parent_row in enumerate(bridge_rows[:3], start=1):
        drilldown_rows = _drilldown_rows(output_dir, alternative, index)
        if not drilldown_rows:
            continue
        top_items = _top_positive_and_negative(drilldown_rows, labels)
        reading = (
            labels["drilldown_balance"]
            if any(amount < 0 for _, amount in top_items)
            else labels["drilldown_concentration"]
        )
        findings.append(
            [
                _row_label(parent_row, labels),
                reading,
                labels["top_contributions"].format(
                    items=_format_label_amounts(top_items, limit=3)
                ),
            ]
        )
    return findings


def _small_multiples_info(
    output_dir: Path,
    labels: dict[str, str],
) -> dict[str, Any] | None:
    """Return client-report metadata for the standard small-multiples chart."""

    image_path = output_dir / "waterfall_small_multiples.png"
    context_path = output_dir / "waterfall_small_multiples_context.json"
    if not image_path.exists() or not context_path.exists():
        return None
    try:
        context = json.loads(context_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if context.get("status") != "written":
        return None
    dimension = _dimension_label(
        context.get("dimension") or labels.get("dimension", "dimension"), labels
    )
    panels = [
        panel
        for panel in context.get("panels", [])
        if isinstance(panel, dict) and panel.get("dimension_value") is not None
    ]
    items: list[tuple[str, float]] = []
    for panel in panels[:3]:
        dominant = panel.get("dominant_component") or {}
        variance_type = str(dominant.get("variance_type") or "").strip()
        label = str(panel.get("dimension_value") or "").strip()
        if variance_type:
            label = f"{label} / {_variance_type_label(variance_type, labels)}"
        try:
            amount = float(dominant.get("variance_amount") or 0.0)
        except (TypeError, ValueError):
            amount = 0.0
        items.append((label, amount))
    item_text = _format_label_amounts(items) if items else dimension
    return {
        "dimension": dimension,
        "image_path": image_path,
        "item_text": item_text,
        "title": labels["chart_small_multiples"].format(dimension=dimension),
        "reading": labels["chart_small_multiples_reading"].format(
            dimension=dimension,
            items=item_text,
        ),
        "caption": labels["chart_small_multiples_caption"],
        "source_row": [
            labels["small_multiples_source"],
            labels["small_multiples_reading"],
            item_text,
        ],
    }


def _pvm_ladder_info(
    output_dir: Path,
    labels: dict[str, str],
) -> dict[str, Any] | None:
    """Return client-report metadata for the PVM decomposition ladder."""

    image_path = output_dir / "pvm_decomposition_ladder.png"
    context_path = output_dir / "pvm_decomposition_ladder_context.json"
    if not image_path.exists() or not context_path.exists():
        return None
    try:
        context = json.loads(context_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    level_three = next(
        (
            level
            for level in context.get("levels", [])
            if isinstance(level, dict) and int(level.get("level") or 0) == 3
        ),
        None,
    )
    components = []
    if isinstance(level_three, dict):
        components = [
            item for item in level_three.get("components", []) if isinstance(item, dict)
        ]
    ranked: list[tuple[str, float]] = []
    for component in components:
        try:
            amount = float(component.get("variance_amount") or 0.0)
        except (TypeError, ValueError):
            amount = 0.0
        if abs(amount) > 0:
            ranked.append(
                (
                    _variance_type_label(component.get("variance_type"), labels),
                    amount,
                )
            )
    ranked.sort(key=lambda item: abs(item[1]), reverse=True)
    item_text = _format_label_amounts(ranked, limit=3) if ranked else ""
    return {
        "image_path": image_path,
        "item_text": item_text,
        "title": labels["chart_pvm_ladder"],
        "reading": labels["chart_pvm_ladder_reading"].format(items=item_text),
        "caption": labels["chart_pvm_ladder_caption"],
        "source_row": [
            labels["pvm_ladder_source"],
            labels["pvm_ladder_reading"],
            item_text,
        ],
    }


def _standard_component_note(
    output_dir: Path,
    labels: dict[str, str],
    comparison: dict[str, str],
) -> str | None:
    """Return a data-driven note about standard variance components."""

    context_path = output_dir / "waterfall_small_multiples_context.json"
    if not context_path.exists():
        return None
    try:
        context = json.loads(context_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    totals: dict[str, float] = {}
    for panel in context.get("panels", []):
        if not isinstance(panel, dict):
            continue
        for component in panel.get("components", []):
            if not isinstance(component, dict):
                continue
            variance_type = str(component.get("variance_type") or "").strip()
            if not variance_type:
                continue
            try:
                amount = float(component.get("variance_amount") or 0.0)
            except (TypeError, ValueError):
                amount = 0.0
            totals[variance_type] = totals.get(variance_type, 0.0) + amount
    if not totals:
        return None
    price_amount = totals.get("Price", 0.0)
    non_price_amount = sum(
        abs(amount)
        for variance_type, amount in totals.items()
        if variance_type not in {"Price", "Other"}
    )
    other_amount = abs(totals.get("Other", 0.0))
    if abs(price_amount) > 0.000001 and non_price_amount <= 0.000001:
        return labels["price_only"].format(**comparison)
    ordered_components = sorted(
        totals.items(),
        key=lambda item: abs(item[1]),
        reverse=True,
    )
    dominant_type, dominant_amount = ordered_components[0]
    if other_amount <= 0.000001:
        ordered_components = [item for item in ordered_components if item[0] != "Other"]
    component_text = ", ".join(
        f"{_variance_type_label(variance_type, labels)} {_format_amount(amount)}"
        for variance_type, amount in ordered_components[:4]
    )
    return labels["component_note"].format(
        dominant_type=_variance_type_label(dominant_type, labels),
        dominant_amount=_format_amount(dominant_amount),
        components=component_text,
    )


def _load_font(size: int, *, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        (
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
            if bold
            else "/System/Library/Fonts/Supplemental/Arial.ttf"
        ),
        (
            "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf"
            if bold
            else "/System/Library/Fonts/Supplemental/Helvetica.ttf"
        ),
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def _draw_segmented_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    segments: tuple[tuple[str, bool], ...],
    *,
    fill: tuple[int, int, int],
    regular_font: ImageFont.ImageFont,
    bold_font: ImageFont.ImageFont,
) -> None:
    """Draw one title line with per-segment emphasis."""

    x, y = xy
    for text, emphasized in segments:
        if not text:
            continue
        font = bold_font if emphasized else regular_font
        draw.text((x, y), text, fill=fill, font=font)
        bbox = draw.textbbox((x, y), text, font=font)
        x += bbox[2] - bbox[0]


def _write_localized_chart(
    source: Path | None,
    target: Path,
    title: str,
    subtitle: str,
    footer: str,
    *,
    title_lines: list[str] | None = None,
) -> Path | None:
    if source is None or not source.exists():
        return None
    image = Image.open(source).convert("RGB")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, image.width, 128), fill="white")
    if title_lines:
        if title_lines:
            draw.text(
                (58, 24),
                title_lines[0],
                fill=(80, 85, 92),
                font=_load_font(18),
            )
        if len(title_lines) > 1:
            _draw_segmented_text(
                draw,
                (58, 49),
                measure_line_segments(title_lines[1]),
                fill=(34, 40, 49),
                regular_font=_load_font(18),
                bold_font=_load_font(18, bold=True),
            )
        if len(title_lines) > 2:
            draw.text(
                (58, 75),
                title_lines[2],
                fill=(80, 85, 92),
                font=_load_font(18),
            )
    else:
        draw.text((58, 40), title, fill=(34, 40, 49), font=_load_font(18, bold=True))
        draw.text((58, 82), subtitle, fill=(80, 85, 92), font=_load_font(18))
    note_top = image.height - 48
    draw.rectangle((0, note_top, image.width, image.height), fill="white")
    draw.text((58, note_top + 10), footer, fill=(105, 105, 105), font=_load_font(18))
    image.save(target)
    return target


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_docx_table(
    document: Document, headers: list[str], rows: list[list[str]]
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if len(headers) == 2:
        widths = [2700, 6660]
    elif len(headers) == 3:
        widths = [2000, 3000, 4360]
    else:
        base_width = 9360 // max(len(headers), 1)
        widths = [base_width] * len(headers)
        widths[-1] += 9360 - sum(widths)
    _set_table_geometry(table, widths)


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    """Apply the standard-business-brief fixed DXA table geometry."""

    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths)))
    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "120")
    table_layout = table_pr.find(qn("w:tblLayout"))
    if table_layout is None:
        table_layout = OxmlElement("w:tblLayout")
        table_pr.append(table_layout)
    table_layout.set(qn("w:type"), "fixed")

    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(width))
            margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(margins)
            for name, value in (
                ("top", 80),
                ("bottom", 80),
                ("start", 120),
                ("end", 120),
            ):
                margin = margins.find(qn(f"w:{name}"))
                if margin is None:
                    margin = OxmlElement(f"w:{name}")
                    margins.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def _add_docx_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(9)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 94)


def _add_docx_chart(
    document: Document,
    title: str,
    reading: str,
    image_path: Path | None,
    caption: str,
    reading_label: str,
    width_inches: float = 6.45,
) -> None:
    document.add_heading(title, level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{reading_label}: ").bold = True
    paragraph.add_run(reading)
    if image_path is not None and image_path.exists():
        document.add_picture(str(image_path), width=Inches(width_inches))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_docx_caption(document, caption)


def _accounting_rows(
    recipe: dict[str, Any],
    language: str,
) -> list[list[str]]:
    """Return localized display rows for explicit accounting-review controls."""

    readiness = recipe.get("accounting_readiness") or {}
    review = recipe.get("accounting_review") or {}
    fields = {
        "it": {
            "perimeter": "Perimetro",
            "tie_out": "Quadratura con la fonte",
            "bridge": "Chiusura del bridge",
            "convention": "Convenzione favorevole/sfavorevole",
            "materiality": "Materialità",
            "professional": "Revisione professionale",
            "root_cause": "Revisione delle alternative",
            "unresolved": "Elementi irrisolti",
            "none": "Nessuno",
            "baseline": "base",
            "comparison": "confronto",
            "calculated": "calcolato",
            "source": "fonte",
            "max_delta": "scostamento massimo",
            "reviewer": "revisore",
            "reviewed_at": "data",
            "alternative": "alternativa",
            "rationale": "motivazione",
        },
        "es": {
            "perimeter": "Perímetro",
            "tie_out": "Conciliación con la fuente",
            "bridge": "Cierre del puente",
            "convention": "Convención favorable/desfavorable",
            "materiality": "Materialidad",
            "professional": "Revisión profesional",
            "root_cause": "Revisión de alternativas",
            "unresolved": "Elementos pendientes",
            "none": "Ninguno",
            "baseline": "base",
            "comparison": "comparación",
            "calculated": "calculado",
            "source": "fuente",
            "max_delta": "diferencia máxima",
            "reviewer": "revisor",
            "reviewed_at": "fecha",
            "alternative": "alternativa",
            "rationale": "justificación",
        },
        "en": {
            "perimeter": "Perimeter",
            "tie_out": "Source tie-out",
            "bridge": "Bridge closure",
            "convention": "Favorable/adverse convention",
            "materiality": "Materiality",
            "professional": "Professional review",
            "root_cause": "Alternative review",
            "unresolved": "Unresolved items",
            "none": "None",
            "baseline": "baseline",
            "comparison": "comparison",
            "calculated": "calculated",
            "source": "source",
            "max_delta": "maximum delta",
            "reviewer": "reviewer",
            "reviewed_at": "date",
            "alternative": "alternative",
            "rationale": "rationale",
        },
    }.get(language, {})
    if not fields:
        fields = {
            "perimeter": "Perimeter",
            "tie_out": "Source tie-out",
            "bridge": "Bridge closure",
            "convention": "Favorable/adverse convention",
            "materiality": "Materiality",
            "professional": "Professional review",
            "root_cause": "Alternative review",
            "unresolved": "Unresolved items",
            "none": "None",
            "baseline": "baseline",
            "comparison": "comparison",
            "calculated": "calculated",
            "source": "source",
            "max_delta": "maximum delta",
            "reviewer": "reviewer",
            "reviewed_at": "date",
            "alternative": "alternative",
            "rationale": "rationale",
        }
    perimeter = readiness.get("perimeter") or review.get("perimeter") or {}
    tie_out = readiness.get("source_tie_out") or {}
    bridge = readiness.get("component_bridge") or {}
    convention = (
        readiness.get("favorable_adverse_convention")
        or review.get("favorable_adverse_convention")
        or {}
    )
    materiality = readiness.get("materiality") or review.get("materiality") or {}
    professional = readiness.get("professional_review") or {}
    root_cause = readiness.get("root_cause_review") or {}
    unresolved = [
        _localized_control_text(value, language)
        for value in readiness.get("unresolved_items", [])
        if str(value).strip()
    ]
    status_labels = {
        "it": {
            "pending": "in sospeso",
            "established": "confermato",
            "passed": "superato",
            "failed": "non superato",
            "not_established": "non disponibile",
            "applied": "applicata",
            "not_applied": "non applicata",
            "approved": "approvata",
        },
        "es": {
            "pending": "pendiente",
            "established": "confirmado",
            "passed": "superado",
            "failed": "no superado",
            "not_established": "no disponible",
            "applied": "aplicada",
            "not_applied": "no aplicada",
            "approved": "aprobada",
        },
    }.get(language, {})

    def status(value: Any, default: str) -> str:
        raw = str(value or default)
        return status_labels.get(raw, raw)

    professional_details = [status(professional.get("status"), "pending")]
    if professional.get("reviewed_by"):
        professional_details.append(
            f"{fields['reviewer']}: {professional.get('reviewed_by')}"
        )
    if professional.get("reviewed_at"):
        professional_details.append(
            f"{fields['reviewed_at']}: {professional.get('reviewed_at')}"
        )
    root_cause_details = [status(root_cause.get("status"), "pending")]
    if root_cause.get("selected_alternative") is not None:
        root_cause_details.append(
            f"{fields['alternative']}: {root_cause.get('selected_alternative')}"
        )
    if root_cause.get("rationale"):
        root_cause_details.append(
            f"{fields['rationale']}: {root_cause.get('rationale')}"
        )
    return [
        [
            fields["perimeter"],
            f"{status(perimeter.get('status'), 'pending')} — "
            f"{perimeter.get('description') or '-'}",
        ],
        [
            fields["tie_out"],
            (
                f"{status(tie_out.get('status'), 'not_established')}; "
                f"{fields['baseline']}: {fields['calculated']} "
                f"{tie_out.get('baseline_calculated_total')} / {fields['source']} "
                f"{tie_out.get('baseline_source_total')}; {fields['comparison']}: "
                f"{fields['calculated']} {tie_out.get('comparison_calculated_total')} / "
                f"{fields['source']} {tie_out.get('comparison_source_total')}"
            ),
        ],
        [
            fields["bridge"],
            f"{status(bridge.get('status'), 'not_established')}; "
            f"{fields['max_delta']} {bridge.get('max_abs_reconciliation_delta')}",
        ],
        [
            fields["convention"],
            f"{status(convention.get('status'), 'pending')} — "
            f"{convention.get('description') or '-'}",
        ],
        [
            fields["materiality"],
            f"{status(materiality.get('status'), 'pending')} — "
            f"{materiality.get('threshold') if materiality.get('threshold') is not None else '-'}; "
            f"{materiality.get('basis') or '-'}",
        ],
        [fields["professional"], "; ".join(professional_details)],
        [fields["root_cause"], "; ".join(root_cause_details)],
        [fields["unresolved"], "; ".join(unresolved) or fields["none"]],
    ]


def _localized_control_text(value: Any, language: str) -> str:
    """Translate stable accounting-control messages for report display."""

    raw = str(value)
    translations = {
        "it": {
            "Confirm the entity and consolidation perimeter.": (
                "Confermare il perimetro societario e di consolidamento."
            ),
            "Provide approved baseline and comparison source totals for tie-out.": (
                "Fornire i totali approvati della fonte per base e confronto."
            ),
            "Confirm the favorable/adverse sign convention.": (
                "Confermare la convenzione favorevole/sfavorevole."
            ),
            "Complete the applied materiality threshold and basis.": (
                "Completare la soglia e il criterio di materialità applicati."
            ),
            "Confirm materiality or explicitly record that it is not applied.": (
                "Confermare la materialità o registrare che non è applicata."
            ),
            "Resolve the failed source-total tie-out.": (
                "Risolvere la quadratura non superata con i totali della fonte."
            ),
            "Resolve the component-bridge reconciliation control.": (
                "Risolvere il controllo di riconciliazione del bridge."
            ),
        }
    }
    return translations.get(language, {}).get(raw, raw)


def _write_docx(
    output_path: Path,
    payload: dict[str, Any],
    labels: dict[str, str],
) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(
            "1F4D78" if style_name == "Heading 3" else "2E74B5"
        )
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    title = document.add_paragraph()
    title_run = title.add_run(payload["title"])
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(36, 48, 38)
    subtitle = document.add_paragraph(payload["subtitle"])
    subtitle.runs[0].font.color.rgb = RGBColor(91, 101, 94)
    document.add_heading(labels["summary"], level=1)
    for paragraph in payload["summary_paragraphs"]:
        document.add_paragraph(paragraph)
    document.add_heading(labels["accounting_controls"], level=1)
    _add_docx_table(
        document,
        [labels["accounting_field"], labels["accounting_value"]],
        payload["accounting_rows"],
    )
    document.add_heading(labels["source_data"], level=1)
    _add_docx_table(document, payload["source_headers"], payload["source_rows"])
    _add_docx_caption(document, labels["source_caption"])
    document.add_heading(labels["reading_notes"], level=2)
    for note in payload["notes"]:
        document.add_paragraph(note, style="List Bullet")
    if payload["chart_sections"]:
        document.add_page_break()
    for index, chart_section in enumerate(payload["chart_sections"]):
        if index > 0 and chart_section["page_break_before"]:
            document.add_page_break()
        _add_docx_chart(
            document,
            chart_section["title"],
            chart_section["reading"],
            chart_section["image_path"],
            chart_section["caption"],
            labels["reading"],
            float(chart_section.get("width_inches") or 6.45),
        )
    if payload["drilldown_findings"]:
        document.add_heading(labels["drilldown_findings"], level=1)
        _add_docx_table(
            document,
            [
                labels["analysis_area"],
                labels["useful_reading"],
                labels["source_col"],
            ],
            payload["drilldown_findings"],
        )
        _add_docx_caption(document, labels["drilldown_caption"])
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = payload["title"]
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(108, 118, 110)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        page_run = footer.add_run()
        page_run.font.size = Pt(9)
        page_run.font.color.rgb = RGBColor(108, 118, 110)
        field_begin = OxmlElement("w:fldChar")
        field_begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        field_separate = OxmlElement("w:fldChar")
        field_separate.set(qn("w:fldCharType"), "separate")
        field_end = OxmlElement("w:fldChar")
        field_end.set(qn("w:fldCharType"), "end")
        page_run._r.extend([field_begin, instruction, field_separate, field_end])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _write_markdown(
    output_path: Path,
    payload: dict[str, Any],
    labels: dict[str, str],
) -> None:
    lines = [
        f"# {payload['title']}",
        "",
        payload["subtitle"],
        "",
        f"## {labels['summary']}",
        "",
        *payload["summary_paragraphs"],
        "",
        f"## {labels['accounting_controls']}",
        "",
        f"| {labels['accounting_field']} | {labels['accounting_value']} |",
        "| --- | --- |",
        *(f"| {row[0]} | {row[1]} |" for row in payload["accounting_rows"]),
        "",
        f"## {labels['source_data']}",
        "",
        f"| {' | '.join(payload['source_headers'])} |",
        f"| {' | '.join(['---'] * len(payload['source_headers']))} |",
    ]
    lines.extend(f"| {' | '.join(row)} |" for row in payload["source_rows"])
    lines.extend(["", f"## {labels['reading_notes']}", ""])
    lines.extend(f"- {note}" for note in payload["notes"])
    for chart_section in payload["chart_sections"]:
        lines.extend(
            [
                "",
                f"## {chart_section['title']}",
                "",
                f"**{labels['reading']}**: {chart_section['reading']}",
                "",
            ]
        )
        image_path = chart_section["image_path"]
        if isinstance(image_path, Path) and image_path.exists():
            lines.append(f"![{chart_section['title']}]({image_path.name})")
            lines.append("")
        lines.append(chart_section["caption"])
    output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _build_payload(
    summary_rows: list[dict[str, Any]],
    recipe: dict[str, Any],
    output_dir: Path,
) -> dict[str, Any] | None:
    language = _language(recipe)
    labels = _text(language)
    comparison = _comparison_metadata(recipe, language)
    summary_row, selection_method = _select_summary(summary_rows, recipe)
    if summary_row is None:
        return None
    readiness = recipe.get("accounting_readiness") or {}
    approved_for_client_use = (
        readiness.get("client_report_status") == "approved_for_client_use"
        and selection_method == "explicit_professional_or_model_review"
    )
    pvm_ladder = _pvm_ladder_info(output_dir, labels)
    small_multiples = _small_multiples_info(output_dir, labels)
    summary_alt = int(summary_row.get("alternative_result") or 0)
    summary_labels = [
        _localized_selected_label(value, labels)
        for value in _split_summary_values(summary_row.get("selected_labels"))
    ]
    summary_label = (
        summary_labels[0]
        if summary_labels
        else labels.get("root_cause_bridge", "Root-cause bridge")
    )
    summary_amounts = _selected_amounts(summary_row)
    summary_amount = summary_amounts[0] if summary_amounts else 0.0
    selected_items = list(zip(summary_labels, summary_amounts))
    selected_text = _format_label_amounts(selected_items, limit=4)
    summary_key_text = (
        selected_text or f"{summary_label} {_format_amount(summary_amount)}"
    )
    driver_count = max(len(summary_labels), len(summary_amounts), 1)
    total_delta = sum(summary_amounts) + _summary_float(summary_row, "other_residual")
    product_items = _best_product_line_items(output_dir, summary_alt, 1, labels)
    top_product = product_items[0][0] if product_items else summary_label
    product_text = _format_label_amounts(product_items)
    mixed_row = _select_mixed(summary_rows)
    mixed_items: list[str] = []
    mixed_chart: Path | None = None
    mixed_residual = 0.0
    if mixed_row is not None:
        mixed_alt = int(mixed_row.get("alternative_result") or 0)
        mixed_bridge_rows = _bridge_rows(output_dir, mixed_alt)
        mixed_items = [_row_label(row, labels) for row in mixed_bridge_rows[:5]]
        mixed_residual = _summary_float(mixed_row, "other_residual")
        mixed_chart = _write_localized_chart(
            _chart_path(output_dir, mixed_row.get("chart_path")),
            output_dir / "root_cause_client_report_mixed_bridge.png",
            labels["chart_mixed_title"],
            labels["chart_mixed_subtitle"].format(**comparison),
            labels["chart_footer"],
            title_lines=build_ibcs_title(
                recipe,
                chart_kind="variable_root_cause",
            ).lines(),
        )
    summary_chart = _write_localized_chart(
        _chart_path(output_dir, summary_row.get("chart_path")),
        output_dir / "root_cause_client_report_summary_bridge.png",
        labels["chart_summary_title"].format(label=summary_label),
        labels["chart_summary_subtitle"].format(**comparison),
        labels["chart_footer"],
        title_lines=build_ibcs_title(
            recipe,
            chart_kind="root_cause",
        ).lines(),
    )
    drilldown_chart = _write_localized_chart(
        output_dir / f"root_cause_bridge_alt_{summary_alt}_drilldown_row_1.png",
        output_dir / "root_cause_client_report_drilldown.png",
        labels["chart_drilldown_title"].format(label=summary_label),
        labels["chart_drilldown_subtitle"].format(**comparison),
        labels["chart_footer"],
        title_lines=build_ibcs_title(
            recipe,
            chart_kind="root_cause_drilldown",
            selection_label=summary_label,
        ).lines(),
    )
    has_drilldown = bool(product_items) or drilldown_chart is not None
    mixed_text = ", ".join(mixed_items) if mixed_items else summary_label
    subtitle = (
        f"{Path(str(recipe.get('source_file') or '')).stem or labels.get('sales', 'Sales')} | "
        f"{comparison['comparison']}"
    )
    summary_template = (
        labels["summary_intro"]
        if approved_for_client_use
        else labels["draft_summary_intro"]
    )
    summary_paragraphs = [
        summary_template.format(
            delta=_format_amount(total_delta),
            driver_count=driver_count,
            items=summary_key_text,
            residual=_format_residual(summary_row.get("other_residual")),
            **comparison,
        )
    ]
    if has_drilldown:
        drilldown_template = (
            labels["drilldown_intro"]
            if approved_for_client_use
            else labels["draft_drilldown_intro"]
        )
        summary_paragraphs.append(
            drilldown_template.format(
                items=product_text or summary_label,
                top_label=top_product,
            )
        )
    if mixed_row is not None:
        mixed_template = (
            labels["mixed_intro"]
            if approved_for_client_use
            else labels["draft_mixed_intro"]
        )
        summary_paragraphs.append(
            mixed_template.format(
                items=mixed_text,
                residual=_format_residual(mixed_residual),
            )
        )
    source_rows = [
        [
            labels["bridge_summary"],
            labels["bridge_reading"],
            f"{summary_key_text}; {labels['residual']} "
            f"{_format_residual(summary_row.get('other_residual'))}",
        ]
    ]
    if pvm_ladder is not None:
        source_rows.append(pvm_ladder["source_row"])
    if small_multiples is not None:
        source_rows.append(small_multiples["source_row"])
    if has_drilldown:
        source_rows.append(
            [
                labels["product_line_drilldown"],
                labels["drilldown_reading"],
                product_text,
            ]
        )
    if mixed_row is not None:
        source_rows.append(
            [
                labels["mixed_deep_dive"],
                labels["mixed_reading"],
                mixed_text,
            ]
        )
    chart_sections = []
    if pvm_ladder is not None:
        chart_sections.append(
            {
                "title": pvm_ladder["title"],
                "reading": pvm_ladder["reading"],
                "image_path": pvm_ladder["image_path"],
                "caption": pvm_ladder["caption"],
                "page_break_before": False,
                "width_inches": 6.2,
            }
        )
    if small_multiples is not None:
        chart_sections.append(
            {
                "title": small_multiples["title"],
                "reading": small_multiples["reading"],
                "image_path": small_multiples["image_path"],
                "caption": small_multiples["caption"],
                "page_break_before": False,
                "width_inches": 5.9,
            }
        )
    chart_sections.append(
        {
            "title": labels["chart_1"],
            "reading": labels["chart_1_reading"].format(
                items=summary_key_text,
                residual=_format_residual(summary_row.get("other_residual")),
                **comparison,
            ),
            "image_path": summary_chart,
            "caption": labels[
                (
                    "chart_1_caption_approved"
                    if approved_for_client_use
                    else "chart_1_caption"
                )
            ],
            "page_break_before": True,
        }
    )
    if has_drilldown:
        chart_sections.append(
            {
                "title": labels["chart_2"],
                "reading": labels["chart_2_reading"].format(top_label=top_product),
                "image_path": drilldown_chart,
                "caption": labels["chart_2_caption"],
                "page_break_before": False,
            }
        )
    if mixed_row is not None:
        chart_sections.append(
            {
                "title": labels["chart_3"],
                "reading": labels["chart_3_reading"].format(items=mixed_text),
                "image_path": mixed_chart,
                "caption": labels["chart_3_caption"],
                "page_break_before": True,
            }
        )
    component_note = _standard_component_note(output_dir, labels, comparison)
    notes = [_currency_note(recipe, labels)]
    if not approved_for_client_use:
        notes.insert(0, labels["draft_notice"])
    if component_note:
        notes.append(component_note)
    notes.append(labels["residual_note"])
    payload = {
        "labels": labels,
        "title": (
            labels["title"] if approved_for_client_use else labels["draft_title"]
        ),
        "subtitle": subtitle,
        "summary_paragraphs": summary_paragraphs,
        "accounting_rows": _accounting_rows(recipe, language),
        "source_headers": [
            labels["source_col"],
            labels["reading"],
            labels["data_key"],
        ],
        "source_rows": source_rows,
        "notes": notes,
        "chart_sections": chart_sections,
        "drilldown_findings": _drilldown_finding_rows(output_dir, mixed_row, labels),
        "report_status": (
            "approved_for_client_use"
            if approved_for_client_use
            else "draft_pending_professional_review"
        ),
        "selection_method": selection_method,
    }
    return payload


def write_root_cause_client_report(
    *,
    summary_rows: list[dict[str, Any]],
    recipe: dict[str, Any],
    output_dir: Path,
) -> tuple[list[str], dict[str, Any]]:
    """Write an approved report or a visibly labelled professional-review draft."""

    payload = _build_payload(summary_rows, recipe, output_dir)
    if payload is None:
        return [], {"status": "not_written_no_summary_rows"}
    labels = payload["labels"]
    md_path = output_dir / "root_cause_client_report.md"
    docx_path = output_dir / "root_cause_client_report.docx"
    _write_markdown(md_path, payload, labels)
    _write_docx(docx_path, payload, labels)
    paths = [str(md_path), str(docx_path)]
    for chart_section in payload["chart_sections"]:
        chart = chart_section["image_path"]
        if isinstance(chart, Path) and chart.exists():
            paths.append(str(chart))
    chart_artifacts = [
        Path(path).name
        for path in paths
        if Path(path).suffix.lower() in {".png", ".jpg", ".jpeg"}
    ]
    return paths, {
        "status": payload["report_status"],
        "selection_method": payload["selection_method"],
        "markdown": md_path.name,
        "docx": docx_path.name,
        "selected_row_count": len(payload["source_rows"]),
        "drilldown_finding_count": len(payload["drilldown_findings"]),
        "chart_artifacts": chart_artifacts,
        "pvm_decomposition_ladder_chart_included": any(
            artifact == "pvm_decomposition_ladder.png" for artifact in chart_artifacts
        ),
        "small_multiples_chart_included": any(
            artifact == "waterfall_small_multiples.png" for artifact in chart_artifacts
        ),
    }
