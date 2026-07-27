from __future__ import annotations

import copy
import csv
import hashlib
import importlib.util
import json
import os
import re
import shutil
import stat
import subprocess
import sys
import zipfile
from decimal import Decimal, localcontext
from pathlib import Path
from typing import Any

import openpyxl
import pytest
from docx import Document

from scripts.validate_plugin_review_contract import validate_contract

ROOT = Path(__file__).resolve().parents[2]
SCRIPT_DIR = ROOT / "plugins" / "concordato-plan-review" / "scripts"
CORE_PATH = SCRIPT_DIR / "concordato_plan_core.py"
MCP_SERVER_PATH = ROOT / "plugins" / "concordato-plan-review" / "mcp" / "server.cjs"


def load_core() -> Any:
    if str(SCRIPT_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPT_DIR))
    spec = importlib.util.spec_from_file_location("concordato_plan_core", CORE_PATH)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _save_workbook(path: Path, rows: list[list[Any]]) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Dati"
    for row in rows:
        sheet.append(row)
    workbook.save(path)


def _call_mcp_server(
    messages: list[dict[str, object]],
    *,
    server_path: Path = MCP_SERVER_PATH,
    env: dict[str, str] | None = None,
) -> list[dict[str, object]]:
    node = shutil.which("node")
    if node is None:
        pytest.skip("Node.js is required to exercise the Concordato MCP server.")
    completed = subprocess.run(
        [node, str(server_path), "--stdio"],
        input="\n".join(json.dumps(message) for message in messages) + "\n",
        capture_output=True,
        text=True,
        check=True,
        timeout=10,
        env={**os.environ, **(env or {})},
    )
    return [json.loads(line) for line in completed.stdout.splitlines() if line.strip()]


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def _docx_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _seal_review_payload(content: dict[str, Any]) -> dict[str, Any]:
    digest = hashlib.sha256(
        json.dumps(
            content,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()
    return {**content, "content_sha256": digest}


def _reseal_decision_content(decision: dict[str, Any]) -> None:
    decision["content_sha256"] = hashlib.sha256(
        json.dumps(
            decision["content"],
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()


def _reseal_payload(payload: dict[str, Any]) -> None:
    content = dict(payload)
    content.pop("content_sha256", None)
    payload["content_sha256"] = hashlib.sha256(
        json.dumps(
            content,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()


def _refresh_receipt(receipt: dict[str, Any], root: Path) -> None:
    payload = (root / receipt["path"]).read_bytes()
    receipt["byte_count"] = len(payload)
    receipt["sha256"] = hashlib.sha256(payload).hexdigest()


def _write_json(path: Path, payload: Any) -> None:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _replace_nested(
    payload: dict[str, Any],
    path: tuple[str, ...],
    value: Any,
) -> None:
    target = payload
    for key in path[:-1]:
        target = target[key]
    target[path[-1]] = value


def _reviewed_source_recipe(
    core: Any,
    inspection_run: Any,
    roles: dict[str, str],
    *,
    excluded_candidate_ids: set[str] | None = None,
) -> dict[str, Any]:
    excluded = excluded_candidate_ids or set()
    return core.review_source_roles(
        inspection_run.inventory,
        roles,
        inspection_run.raw_amount_candidates,
        {
            core.candidate_id(candidate): (
                "excluded_non_amount"
                if core.candidate_id(candidate) in excluded
                else "candidate_amount"
            )
            for candidate in inspection_run.raw_amount_candidates
        },
        reviewer_ref="pytest.reviewer",
        reviewed_on="2026-07-24",
        reference_date=str(inspection_run.audit.get("reference_date") or ""),
        tolerance=str(inspection_run.audit.get("tolerance") or "0"),
    )


def _build_qualified_concordato_run(
    tmp_path: Path,
) -> tuple[Any, Path, Path]:
    core = load_core()
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["Debiti tributari", 100]],
    )
    _save_workbook(
        input_dir / "supporto.xlsx",
        [["Voce", "Saldo"], ["Debiti tributari", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {
            "piano.xlsx": "concordato_plan",
            "supporto.xlsx": "other_support",
        },
    )
    core.run_concordato_review(
        input_dir,
        output_dir,
        tolerance="0",
        recipe=recipe,
    )
    return core, input_dir, output_dir


def _build_multirow_qualified_concordato_run(
    tmp_path: Path,
) -> tuple[Any, Path]:
    core = load_core()
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [
            ["Voce", "Importo"],
            ["Debiti tributari", 100],
            ["Debiti fornitori", 200],
            ["Debiti previdenziali", 300],
            ["Assunzione prospettica", 400],
        ],
    )
    _save_workbook(
        input_dir / "supporto.xlsx",
        [
            ["Voce", "Saldo"],
            ["Debiti tributari", 100],
            ["Debiti fornitori", 200],
            ["Debiti previdenziali", 300],
        ],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {
            "piano.xlsx": "concordato_plan",
            "supporto.xlsx": "other_support",
        },
    )
    core.run_concordato_review(
        input_dir,
        output_dir,
        tolerance="0",
        recipe=recipe,
    )
    return core, output_dir


def test_parse_amount_token_handles_italian_amounts_and_dates() -> None:
    core = load_core()

    assert core.parse_amount_token("1.234.567,89") == Decimal("1234567.89")
    assert core.parse_amount_token("(2.500,00)") == Decimal("-2500.00")
    assert core.parse_amount_token("31.03.2026") is None
    assert core.parse_amount_token("1.234") is None


def test_extract_amount_candidates_preserves_full_italian_currency_token() -> None:
    core = load_core()

    candidates = core.extract_amount_candidates_from_text(
        "Debiti verso banche € 1.730.547,50 e dipendenti € 410.086,56.",
        source_file="piano.pdf",
        source_role="concordato_plan",
        location="page 1",
    )

    assert [(candidate.token, candidate.amount) for candidate in candidates] == [
        ("€ 1.730.547,50", Decimal("1730547.50")),
        ("€ 410.086,56", Decimal("410086.56")),
    ]


def test_exact_amount_match_includes_exact_cent_tolerance_boundary() -> None:
    core = load_core()
    candidates = [
        core.AmountCandidate(
            "plan.xlsx",
            "concordato_plan",
            "Plan!A1",
            Decimal("100.01"),
            "100,01",
            "trade payables",
        ),
        core.AmountCandidate(
            "support.xlsx",
            "trial_balance",
            "TB!B2",
            Decimal("100"),
            "100,00",
            "trade payables",
        ),
    ]

    matches = core.find_exact_amount_matches(candidates, tolerance="0.01")

    assert matches[0]["difference"] == "0.01"
    assert matches[0]["abs_difference"] == "0.01"
    assert matches[0]["plan_amount"] == "100.01"
    assert matches[0]["support_amount"] == "100"


def test_exact_amount_match_excludes_value_outside_exact_tolerance() -> None:
    core = load_core()
    candidates = [
        core.AmountCandidate(
            "plan.xlsx",
            "concordato_plan",
            "Plan!A1",
            Decimal("100.01"),
            "100,01",
            "trade payables",
        ),
        core.AmountCandidate(
            "support.xlsx",
            "trial_balance",
            "TB!B2",
            Decimal("100"),
            "100,00",
            "trade payables",
        ),
    ]

    matches = core.find_exact_amount_matches(candidates, tolerance="0.009")

    assert matches == []


def test_missing_source_role_receipt_blocks_numeric_parser_and_gates(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["Debiti tributari", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {"piano.xlsx": "concordato_plan"},
    )
    recipe.pop("source_role_decision")

    def forbidden_parser(*_args: Any, **_kwargs: Any) -> None:
        raise AssertionError("numeric parser ran before source authority closed")

    monkeypatch.setattr(core, "_workbook_candidates", forbidden_parser)

    run = core.run_concordato_review(
        input_dir,
        tmp_path / "rejected",
        tolerance="0",
        recipe=recipe,
    )

    assert run.audit["raw_amount_candidate_count"] == 0
    assert run.audit["amount_candidate_count"] == 0
    assert run.audit["candidate_match_count"] == 0
    assert run.audit["source_qualification_status"] == "unsupported_source_layout"
    assert run.audit["assurance_gates"]["gates"]["source"]["status"] == "failed"
    assert (
        run.audit["assurance_gates"]["gates"]["reconciliation"]["status"] == "blocked"
    )


def test_resealed_stale_source_role_receipt_blocks_numeric_parser(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["Debiti tributari", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {"piano.xlsx": "concordato_plan"},
    )
    source_decision = recipe["source_role_decision"]
    source_decision["content"]["source_roles"][0]["unit"] = "1000"
    _reseal_decision_content(source_decision)

    def forbidden_parser(*_args: Any, **_kwargs: Any) -> None:
        raise AssertionError("numeric parser ran under stale source-role authority")

    monkeypatch.setattr(core, "_workbook_candidates", forbidden_parser)

    run = core.run_concordato_review(
        input_dir,
        tmp_path / "rejected",
        tolerance="0",
        recipe=recipe,
    )

    assert run.audit["raw_amount_candidate_count"] == 0
    assert run.audit["amount_candidate_count"] == 0
    assert run.audit["candidate_match_count"] == 0
    assert run.audit["source_qualification_status"] == "unsupported_source_layout"
    assert "source-role declaration is stale" in run.audit["source_role_review_error"]


@pytest.mark.parametrize(
    ("field_path", "replacement"),
    [
        (("signed_difference_formula",), "support_amount - plan_amount"),
        (
            ("sign_convention", "positive"),
            "support_amount_exceeds_plan_amount",
        ),
        (("reference_period",), "2099-12-31"),
        (("tolerance",), "999"),
    ],
)
def test_resealed_stale_calculation_contract_blocks_parser_before_extraction(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    field_path: tuple[str, ...],
    replacement: str,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["Debiti tributari", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        reference_date="2026-03-31",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {"piano.xlsx": "concordato_plan"},
    )
    calculation_decision = recipe["calculation_decision"]
    _replace_nested(calculation_decision["content"], field_path, replacement)
    _reseal_decision_content(calculation_decision)

    def forbidden_parser(*_args: Any, **_kwargs: Any) -> None:
        raise AssertionError("numeric parser ran under stale calculation authority")

    monkeypatch.setattr(core, "_workbook_candidates", forbidden_parser)

    run = core.run_concordato_review(
        input_dir,
        tmp_path / "rejected",
        reference_date="2026-03-31",
        tolerance="0",
        recipe=recipe,
    )

    assert run.audit["raw_amount_candidate_count"] == 0
    assert run.audit["amount_candidate_count"] == 0
    assert run.audit["candidate_match_count"] == 0
    assert run.audit["calculation_review_status"] == "withheld"
    assert (
        run.audit["assurance_gates"]["gates"]["reconciliation"]["status"] == "blocked"
    )


def test_resealed_stale_candidate_perimeter_withholds_authoritative_rows(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["Debiti tributari", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {"piano.xlsx": "concordato_plan"},
    )
    calculation_decision = recipe["calculation_decision"]
    calculation_decision["content"]["candidate_perimeter"][0]["amount"] = "101"
    _reseal_decision_content(calculation_decision)

    run = core.run_concordato_review(
        input_dir,
        tmp_path / "rejected",
        tolerance="0",
        recipe=recipe,
    )

    assert run.audit["raw_amount_candidate_count"] == 1
    assert run.audit["amount_candidate_count"] == 0
    assert run.audit["candidate_match_count"] == 0
    assert run.audit["source_qualification_status"] == "unsupported_source_layout"
    assert "candidate perimeter" in run.audit["source_role_review_error"]


def test_calculation_receipt_binds_formula_sign_period_sources_and_candidates(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["Debiti tributari", 99]],
    )
    _save_workbook(
        input_dir / "supporto.xlsx",
        [["Voce", "Saldo"], ["Debiti tributari", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        reference_date="2026-03-31",
        tolerance="1",
    )

    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {
            "piano.xlsx": "concordato_plan",
            "supporto.xlsx": "other_support",
        },
    )

    content = recipe["calculation_decision"]["content"]
    assert content["formula_id"] == "plan_amount_minus_support_amount.v1"
    assert content["signed_difference_formula"] == ("plan_amount - support_amount")
    assert content["absolute_difference_formula"] == "abs(difference)"
    assert content["tolerance_comparison_formula"] == ("abs_difference <= tolerance")
    assert content["sign_convention"] == {
        "positive": "plan_amount_exceeds_support_amount",
        "zero": "plan_amount_equals_support_amount",
        "negative": "plan_amount_below_support_amount",
    }
    assert content["reference_period"] == "2026-03-31"
    assert content["tolerance"] == "1"
    assert len(content["source_perimeter"]) == 2
    assert len(content["candidate_perimeter"]) == 2
    assert {entry["candidate_id"] for entry in content["candidate_perimeter"]} == {
        core.candidate_id(candidate) for candidate in inspection.raw_amount_candidates
    }
    assert [
        (receipt["root_id"], receipt["path"])
        for receipt in content["implementation_receipts"]
    ] == [
        ("implementation", ".codex-plugin/plugin.json"),
        ("implementation", ".app.json"),
        ("implementation", ".mcp.json"),
        ("implementation", "assets/concordato-plan-review-widget.html"),
        ("implementation", "assets/icon.svg"),
        ("implementation", "assets/review-workbench-adapter.json"),
        ("implementation", "mcp/server.cjs"),
        ("implementation", "scripts/apply_review_edits.py"),
        ("implementation", "scripts/check_dependencies.py"),
        ("implementation", "scripts/concordato_plan_core.py"),
        ("implementation", "scripts/concordato_semantic.py"),
        ("implementation", "scripts/finalize_output_closure.py"),
        ("implementation", "scripts/implementation_bootstrap.py"),
        ("implementation", "scripts/output_closure.py"),
        ("implementation", "scripts/replay_assurance.py"),
        ("implementation", "scripts/review_case_model.py"),
        ("implementation", "scripts/review_session.py"),
        ("implementation", "scripts/review_source_roles.py"),
        ("implementation", "scripts/run_concordato_review.py"),
        ("assurance_implementation", "__init__.py"),
        ("assurance_implementation", "contracts.py"),
        ("assurance_implementation", "decisions.py"),
        ("assurance_implementation", "envelope.py"),
        ("assurance_implementation", "money.py"),
        ("assurance_implementation", "relationships.py"),
        ("assurance_implementation", "review_output_transaction.cjs"),
        ("assurance_implementation", "serialization.py"),
    ]
    assert "remain reviewer judgment" in content["judgment_boundary"]


@pytest.mark.parametrize(
    ("root_name", "relative_path"),
    [
        ("assurance", "money.py"),
        ("plugin", "scripts/output_closure.py"),
    ],
)
def test_transitive_implementation_mutation_blocks_candidate_authority(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    root_name: str,
    relative_path: str,
) -> None:
    core = load_core()
    plugin_copy = tmp_path / "concordato-plan-review"
    assurance_copy = tmp_path / "vera_assurance"
    shutil.copytree(
        ROOT / "plugins" / "concordato-plan-review",
        plugin_copy,
        ignore=shutil.ignore_patterns("__pycache__"),
    )
    shutil.copytree(
        ROOT / "plugins" / "_shared" / "vendor" / "modules" / "vera_assurance",
        assurance_copy,
        ignore=shutil.ignore_patterns("__pycache__"),
    )
    monkeypatch.setattr(core, "COMPONENT_ROOT", plugin_copy)
    monkeypatch.setattr(core, "SCRIPT_DIR", plugin_copy / "scripts")
    monkeypatch.setattr(core, "ASSURANCE_IMPLEMENTATION_ROOT", assurance_copy)
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(input_dir / "piano.xlsx", [["Importo"], [100]])
    _save_workbook(input_dir / "supporto.xlsx", [["Saldo"], [100]])
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {
            "piano.xlsx": "concordato_plan",
            "supporto.xlsx": "other_support",
        },
    )
    mutation_root = assurance_copy if root_name == "assurance" else plugin_copy
    mutation_path = mutation_root / relative_path
    mutation_path.write_bytes(
        mutation_path.read_bytes() + b"\n# receipt-boundary mutation\n"
    )

    run = core.run_concordato_review(
        input_dir,
        tmp_path / "qualified",
        tolerance="0",
        recipe=recipe,
    )

    assert run.audit["raw_amount_candidate_count"] == 0
    assert run.audit["amount_candidate_count"] == 0
    assert run.audit["candidate_match_count"] == 0
    assert run.audit["source_qualification_status"] == "unsupported_source_layout"
    assert "implementation" in run.audit["source_role_review_error"].lower()


def _copy_concordato_runtime(tmp_path: Path, name: str) -> tuple[Path, Path]:
    runtime_root = tmp_path / name
    plugin_copy = runtime_root / "concordato-plan-review"
    shared_copy = runtime_root / "_shared" / "vendor" / "modules" / "vera_assurance"
    shutil.copytree(
        ROOT / "plugins" / "concordato-plan-review",
        plugin_copy,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    shutil.copytree(
        ROOT / "plugins" / "_shared" / "vendor" / "modules" / "vera_assurance",
        shared_copy,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    return plugin_copy, shared_copy


@pytest.mark.parametrize(
    "attack_kind",
    ["empty_directory", "regular", "symlink", "hardlink", "fifo"],
)
def test_transitive_contract_rejects_every_unowned_implementation_entry(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    attack_kind: str,
) -> None:
    core = load_core()
    plugin_copy, shared_copy = _copy_concordato_runtime(tmp_path, attack_kind)
    monkeypatch.setattr(core, "COMPONENT_ROOT", plugin_copy)
    monkeypatch.setattr(core, "SCRIPT_DIR", plugin_copy / "scripts")
    monkeypatch.setattr(core, "ASSURANCE_IMPLEMENTATION_ROOT", shared_copy)
    input_dir = tmp_path / f"{attack_kind}-input"
    input_dir.mkdir()
    _save_workbook(input_dir / "piano.xlsx", [["Importo"], [100]])
    _save_workbook(input_dir / "supporto.xlsx", [["Saldo"], [100]])
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / f"{attack_kind}-inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {
            "piano.xlsx": "concordato_plan",
            "supporto.xlsx": "other_support",
        },
    )
    rogue = plugin_copy / "scripts" / "rogue"
    external = tmp_path / f"{attack_kind}-external"
    if attack_kind == "empty_directory":
        rogue.mkdir()
    elif attack_kind == "regular":
        rogue.write_bytes(b"unreceipted implementation")
    elif attack_kind == "symlink":
        external.write_bytes(b"external")
        rogue.symlink_to(external)
    elif attack_kind == "hardlink":
        external.write_bytes(b"external")
        rogue.hardlink_to(external)
    else:
        os.mkfifo(rogue)

    with pytest.raises(RuntimeError, match="implementation"):
        core.validate_implementation_contract(recipe["calculation_decision"])


def test_real_python_entry_rejects_timestamp_valid_unreceipted_bytecode(
    tmp_path: Path,
) -> None:
    plugin_copy, _ = _copy_concordato_runtime(tmp_path, "malicious-pyc")
    target = plugin_copy / "scripts" / "concordato_plan_core.py"
    source = target.read_bytes()
    source_stat = target.stat()
    marker = tmp_path / "malicious-pyc-executed.txt"
    malicious = (
        "from pathlib import Path as _AttackPath\n"
        f"_AttackPath({marker.as_posix()!r}).write_text("
        "'executed before validation\\n', encoding='utf-8')\n"
        f"exec(compile({source.decode('utf-8')!r}, {target.as_posix()!r}, "
        "'exec'), globals())\n"
    )
    code = compile(malicious, target.as_posix(), "exec")
    cache_prefix = sys.pycache_prefix
    try:
        sys.pycache_prefix = None
        cache_path = Path(importlib.util.cache_from_source(target.as_posix()))
    finally:
        sys.pycache_prefix = cache_prefix
    cache_path.parent.mkdir()
    cache_path.write_bytes(
        importlib._bootstrap_external._code_to_timestamp_pyc(
            code,
            int(source_stat.st_mtime),
            source_stat.st_size,
        )
    )

    completed = subprocess.run(
        [
            sys.executable,
            "-I",
            "-B",
            str(plugin_copy / "scripts" / "run_concordato_review.py"),
            "--help",
        ],
        cwd=plugin_copy,
        capture_output=True,
        text=True,
        check=False,
        timeout=30,
    )

    assert completed.returncode != 0
    assert "implementation" in completed.stderr.lower()
    assert not marker.exists()
    assert target.read_bytes() == source
    assert target.stat().st_size == source_stat.st_size
    assert target.stat().st_mtime_ns == source_stat.st_mtime_ns


@pytest.mark.parametrize("attack_kind", ["symlink", "hardlink", "fifo"])
def test_real_python_entry_rejects_unsafe_bootstrap_before_read(
    tmp_path: Path,
    attack_kind: str,
) -> None:
    plugin_copy, _ = _copy_concordato_runtime(
        tmp_path,
        f"bootstrap-{attack_kind}",
    )
    bootstrap = plugin_copy / "scripts" / "implementation_bootstrap.py"
    original = bootstrap.read_bytes()
    bootstrap.unlink()
    external = tmp_path / f"bootstrap-{attack_kind}-external.py"
    if attack_kind == "symlink":
        external.write_bytes(original)
        bootstrap.symlink_to(external)
    elif attack_kind == "hardlink":
        external.write_bytes(original)
        bootstrap.hardlink_to(external)
    else:
        os.mkfifo(bootstrap)

    completed = subprocess.run(
        [
            sys.executable,
            "-I",
            "-B",
            str(plugin_copy / "scripts" / "run_concordato_review.py"),
            "--help",
        ],
        cwd=plugin_copy,
        capture_output=True,
        text=True,
        check=False,
        timeout=30,
    )

    assert completed.returncode != 0
    assert "bootstrap is not a real file" in completed.stderr


def test_mcp_rejects_unowned_implementation_path_before_stdio(
    tmp_path: Path,
) -> None:
    plugin_copy, _ = _copy_concordato_runtime(tmp_path, "mcp-rogue")
    (plugin_copy / "scripts" / "__pycache__").mkdir()

    with pytest.raises(subprocess.CalledProcessError):
        _call_mcp_server(
            [
                {
                    "jsonrpc": "2.0",
                    "id": 1,
                    "method": "tools/list",
                    "params": {},
                }
            ],
            server_path=plugin_copy / "mcp" / "server.cjs",
        )


def test_resealed_candidate_perimeter_reorder_withholds_authoritative_rows(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["A", 100], ["B", 200]],
    )
    _save_workbook(
        input_dir / "supporto.xlsx",
        [["Voce", "Saldo"], ["A", 100], ["B", 200]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {
            "piano.xlsx": "concordato_plan",
            "supporto.xlsx": "other_support",
        },
    )
    perimeter = recipe["calculation_decision"]["content"]["candidate_perimeter"]
    perimeter.reverse()
    _reseal_decision_content(recipe["calculation_decision"])

    run = core.run_concordato_review(
        input_dir,
        tmp_path / "qualified",
        tolerance="0",
        recipe=recipe,
    )

    assert run.audit["raw_amount_candidate_count"] == 4
    assert run.audit["amount_candidate_count"] == 0
    assert run.audit["candidate_match_count"] == 0
    assert run.audit["source_qualification_status"] == "unsupported_source_layout"


def test_unreviewed_filename_roles_do_not_emit_candidate_matches(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(input_dir / "piano CP.xlsx", [["Importo"], [100]])
    _save_workbook(input_dir / "bilancio.xlsx", [["Saldo"], [100]])

    run = core.run_concordato_review(input_dir, tmp_path / "output")

    assert run.raw_amount_candidates
    assert run.amount_candidates == []
    assert run.exact_matches == []
    assert run.audit["source_qualification_status"] == "needs_review"
    assert (
        json.loads((run.output_dir / "assurance_gates.json").read_text())[
            "report_ready"
        ]
        is False
    )


def test_reviewed_candidate_dispositions_exclude_identifier_matches(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Valore"], ["Employee ID", 12345], ["Total debt", 100]],
    )
    _save_workbook(
        input_dir / "supporto.xlsx",
        [["Voce", "Valore"], ["Invoice ID", 12345], ["Total debt", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    excluded = {
        core.candidate_id(candidate)
        for candidate in inspection.raw_amount_candidates
        if " ID" in candidate.context
    }
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {
            "piano.xlsx": "concordato_plan",
            "supporto.xlsx": "other_support",
        },
        excluded_candidate_ids=excluded,
    )

    run = core.run_concordato_review(
        input_dir,
        tmp_path / "output",
        tolerance="0",
        recipe=recipe,
    )

    assert [row["plan_amount"] for row in run.exact_matches] == ["100"]
    assert all(
        candidate.amount != Decimal("12345") for candidate in run.amount_candidates
    )


def test_truncated_workbook_abstains_without_candidate_facts(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["A", 100], ["B", 200]],
    )

    run = core.run_concordato_review(
        input_dir,
        tmp_path / "output",
        max_rows_per_sheet=1,
    )

    assert run.raw_amount_candidates == []
    assert run.amount_candidates == []
    assert run.exact_matches == []
    assert run.audit["source_qualification_status"] == "unsupported_source_layout"
    assert "exceeds max_rows_per_sheet" in run.audit["extraction_errors"][0]["error"]


def test_binary_float_tolerance_is_rejected() -> None:
    core = load_core()

    with pytest.raises(ValueError, match="binary float"):
        core.find_exact_amount_matches([], tolerance=0.1 + 0.2)


def test_untrusted_ooxml_entity_is_rejected_without_emitting_candidates(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    safe_path = tmp_path / "safe.xlsx"
    malicious_path = input_dir / "piano.xlsx"
    _save_workbook(safe_path, [["Voce", "Importo"], ["Debiti", 100]])
    with zipfile.ZipFile(safe_path) as source_archive:
        with zipfile.ZipFile(malicious_path, "w") as target_archive:
            for member in source_archive.infolist():
                payload = source_archive.read(member.filename)
                if member.filename == "xl/workbook.xml":
                    xml_text = payload.decode("utf-8")
                    if "?>" in xml_text:
                        declaration, body = xml_text.split("?>", 1)
                        declaration += "?>"
                    else:
                        declaration, body = "", xml_text
                    payload = (
                        declaration
                        + '<!DOCTYPE workbook [<!ENTITY xxe SYSTEM "file:///etc/passwd">]>'
                        + body.replace('name="Dati"', 'name="&xxe;"', 1)
                    ).encode("utf-8")
                target_archive.writestr(member, payload)

    run = core.run_concordato_review(input_dir, output_dir, tolerance="0")

    assert run.raw_amount_candidates == []
    assert run.amount_candidates == []
    assert run.audit["source_qualification_status"] == "unsupported_source_layout"
    assert run.audit["extraction_errors"]


def test_source_swap_during_extraction_uses_captured_bytes(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    source = input_dir / "piano.xlsx"
    replacement = tmp_path / "replacement.xlsx"
    _save_workbook(source, [["Importo"], [100]])
    _save_workbook(replacement, [["Importo"], [999]])
    original_reader = core._workbook_candidates

    def swap_live_source(path: Path, **kwargs: Any) -> Any:
        original_bytes = path.read_bytes()
        path.write_bytes(replacement.read_bytes())
        try:
            return original_reader(path, **kwargs)
        finally:
            path.write_bytes(original_bytes)

    monkeypatch.setattr(core, "_workbook_candidates", swap_live_source)

    run = core.run_concordato_review(input_dir, tmp_path / "output")

    assert [candidate.amount for candidate in run.raw_amount_candidates] == [
        Decimal("100")
    ]
    assert source.read_bytes() != replacement.read_bytes()


def test_duplicate_basenames_keep_distinct_source_identities(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    (input_dir / "a").mkdir(parents=True)
    (input_dir / "b").mkdir(parents=True)
    _save_workbook(input_dir / "a" / "piano.xlsx", [["Importo"], [100]])
    _save_workbook(input_dir / "b" / "piano.xlsx", [["Importo"], [100]])

    run = core.run_concordato_review(input_dir, tmp_path / "output")

    assert {candidate.source_file for candidate in run.raw_amount_candidates} == {
        "a/piano.xlsx",
        "b/piano.xlsx",
    }
    assert (
        len({candidate.source_artifact_ref for candidate in run.raw_amount_candidates})
        == 2
    )
    assert (
        len({core.candidate_id(candidate) for candidate in run.raw_amount_candidates})
        == 2
    )


def test_run_concordato_review_writes_candidate_workpapers(tmp_path: Path) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "plan_material.xlsx",
        [
            ["Voce", "Importo"],
            ["Debiti tributari entro 12 mesi", 4124413.15],
            ["Assunzione prospettica non storica", 999999.99],
        ],
    )
    _save_workbook(
        input_dir / "accounting_support.xlsx",
        [
            ["Voce", "Saldo rettificato"],
            ["Debiti tributari entro 12 mesi", 4124413.15],
        ],
    )

    inspection_run = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        reference_date="2026-03-31",
        language="it",
        document_language="it",
        tolerance="0.01",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection_run,
        {
            "plan_material.xlsx": "concordato_plan",
            "accounting_support.xlsx": "other_support",
        },
    )
    run = core.run_concordato_review(
        input_dir,
        output_dir,
        reference_date="2026-03-31",
        language="it",
        document_language="it",
        tolerance="0.01",
        recipe=recipe,
    )

    audit = json.loads((output_dir / "run_audit.json").read_text(encoding="utf-8"))
    matches = (output_dir / "exact_amount_matches.csv").read_text(encoding="utf-8")

    assert run.audit["candidate_match_count"] == 1
    assert audit["tolerance"] == "0.01"
    assert audit["deterministic_boundary"].startswith("Inventory, extraction")
    assert "candidate_amount_match" in matches
    assert (output_dir / "concordato_tie_out_workpaper.xlsx").exists()
    assert (output_dir / "concordato_review_summary.docx").exists()
    assert (output_dir / "review_packet.md").exists()
    assert (output_dir / "run_intake.json").exists()
    assert (output_dir / "review_payload.json").exists()
    assert (output_dir / "ui_decisions.json").exists()
    assert (output_dir / "final_artifacts.json").exists()

    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    assert review_payload["plugin"] == "concordato-plan-review"
    assert review_payload["review_type"] == "concordato_preventivo_review"
    assert review_payload["item_count"] == len(review_payload["items"])
    item_types = {item["item_type"] for item in review_payload["items"]}
    assert {
        "source_inventory",
        "candidate_amount_match",
        "unmatched_plan_amount",
        "review_artifact",
        "codex_review_memo",
    } <= item_types
    assert review_payload["summary"]["candidate_match_count"] == 1
    assert review_payload["summary"]["unmatched_plan_amount_count"] == 1
    unmatched_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "unmatched_plan_amount"
    )
    assert unmatched_item["recommended_action"] == "request_more_documents"
    assert unmatched_item["data"]["requested_document"].startswith(
        "Support document or explanatory schedule for concordato plan amount "
    )
    assert (
        unmatched_item["data"]["required_document"]
        == unmatched_item["data"]["requested_document"]
    )
    assert unmatched_item["data"]["reason"] == (
        "No deterministic support amount matched this plan amount within tolerance."
    )
    assert unmatched_item["data"]["source_file"] == "plan_material.xlsx"
    assert unmatched_item["data"]["amount"] == "999,999.99"
    assert (
        unmatched_item["evidence"][0]["requested_document"]
        == unmatched_item["data"]["requested_document"]
    )

    ui_decisions = json.loads(
        (output_dir / "ui_decisions.json").read_text(encoding="utf-8")
    )
    assert ui_decisions["decision_source"] == "not_collected"
    assert ui_decisions["status"] == "pending_review"

    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    assert final_artifacts["run_id"] == review_payload["run_id"]
    assert final_artifacts["status"] == "written_pending_review"
    handoff_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "review_handoff.md"
    )
    handoff_text = (output_dir / "review_handoff.md").read_text(encoding="utf-8")
    assert handoff_output["required_text"] == [
        "Review Handoff",
        "review_payload.json",
        "ui_decisions.json",
        "applied_decisions.json",
        "final_artifacts.json",
    ]
    assert "render_concordato_plan_review" in handoff_text
    assert "apply_concordato_plan_decisions" in handoff_text
    review_packet_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "review_packet.md"
    )
    assert review_packet_output["required_text"] == [
        "# Pacchetto di revisione del concordato preventivo",
        "## Revisione professionale richiesta",
        "## Appendice deterministica di tie-out",
    ]
    workpaper_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "concordato_tie_out_workpaper.xlsx"
    )
    assert workpaper_output["required_sheets"] == [
        "Inventory",
        "Amount candidates",
        "Candidate matches",
    ]
    assert workpaper_output["required_sheet_headers"] == {
        "Inventory": [
            "path",
            "relative_path",
            "name",
            "suffix",
            "size_bytes",
            "supported",
            "suggested_role",
            "reviewed_role",
            "source_artifact_ref",
            "capture_status",
            "reviewed_currency",
            "reviewed_unit",
        ],
        "Amount candidates": [
            "candidate_id",
            "source_file",
            "source_artifact_ref",
            "source_role",
            "location",
            "amount",
            "currency",
            "unit",
            "token",
            "context",
        ],
        "Candidate matches": [
            "plan_source_file",
            "plan_source_artifact_ref",
            "plan_location",
            "plan_amount",
            "plan_currency",
            "plan_unit",
            "plan_context",
            "support_source_file",
            "support_source_artifact_ref",
            "support_role",
            "support_location",
            "support_amount",
            "support_currency",
            "support_unit",
            "support_context",
            "difference",
            "abs_difference",
            "tolerance",
            "within_tolerance",
            "calculation_formula_id",
            "context_token_overlap",
            "match_status",
        ],
    }
    first_candidate = run.amount_candidates[0]
    assert workpaper_output["required_cells"] == {
        "Inventory": {
            "B1": "relative_path",
            "B2": "accounting_support.xlsx",
            "C1": "name",
            "C2": "accounting_support.xlsx",
            "G1": "suggested_role",
            "G2": "unclassified",
        },
        "Amount candidates": {
            "B1": "source_file",
            "B2": first_candidate.source_file,
            "D1": "source_role",
            "D2": first_candidate.source_role,
            "E1": "location",
            "E2": first_candidate.location,
            "F1": "amount",
            "F2": core.decimal_text(first_candidate.amount),
        },
        "Candidate matches": {
            "A1": "plan_source_file",
            "A2": "plan_material.xlsx",
            "D1": "plan_amount",
            "D2": "4124413.15",
            "H1": "support_source_file",
            "H2": "accounting_support.xlsx",
            "L1": "support_amount",
            "L2": "4124413.15",
            "V1": "match_status",
            "V2": "candidate_amount_match",
        },
    }
    assert "required_sheet_headers" in workpaper_output["qa_checks"]
    assert "required_cells" in workpaper_output["qa_checks"]
    exact_matches_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "exact_amount_matches.csv"
    )
    assert exact_matches_output["row_count"] == run.audit["candidate_match_count"]
    assert exact_matches_output["required_columns"] == [
        "plan_amount",
        "support_amount",
        "difference",
        "match_status",
    ]
    contract_report = validate_contract(
        output_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()

    document = Document(output_dir / "concordato_review_summary.docx")
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "battono per importo" in document_text
    assert "non battono" in document_text


def test_numeric_evidence_ledger_closes_all_multirow_material_addresses(
    tmp_path: Path,
) -> None:
    core, output_dir = _build_multirow_qualified_concordato_run(tmp_path)
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )

    validated = core.validate_numeric_evidence_closure(output_dir, ledger)
    observed_addresses = {
        (address["path"], address["locator"])
        for entry in validated["entries"]
        for address in entry["output_addresses"]
    }

    assert (
        "exact_amount_matches.csv",
        "row=3;column=difference",
    ) in observed_addresses
    assert (
        "concordato_tie_out_workpaper.xlsx",
        "Candidate matches!P4",
    ) in observed_addresses
    assert (
        "concordato_review_summary.docx",
        "table=3;row=4;column=5",
    ) in observed_addresses


def test_numeric_evidence_replay_rejects_resealed_forged_formula_binding(
    tmp_path: Path,
) -> None:
    core, output_dir = _build_multirow_qualified_concordato_run(tmp_path)
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )
    ledger["formula_decision_ref"] = "decision.forged"
    for entry in ledger["entries"]:
        entry["calculation_decision_ref"] = "decision.forged"
    _reseal_payload(ledger)

    with pytest.raises(ValueError, match="formula decision binding is stale"):
        core.validate_numeric_evidence_closure(output_dir, ledger)


def test_numeric_evidence_replay_rejects_resealed_source_outside_perimeter(
    tmp_path: Path,
) -> None:
    core, output_dir = _build_multirow_qualified_concordato_run(tmp_path)
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )
    candidate_entry = next(
        entry for entry in ledger["entries"] if entry["kind"] == "candidate_amount"
    )
    candidate_entry["source_addresses"][0]["locator"] = "Dati!Z999"
    _reseal_payload(ledger)

    with pytest.raises(ValueError, match="outside the reviewed perimeter"):
        core.validate_numeric_evidence_closure(output_dir, ledger)


@pytest.mark.parametrize("kind", ["unmatched_plan_residual", "summary_count"])
def test_numeric_evidence_replay_rejects_resealed_derived_value_forgery(
    tmp_path: Path,
    kind: str,
) -> None:
    core, output_dir = _build_multirow_qualified_concordato_run(tmp_path)
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )
    entry = next(item for item in ledger["entries"] if item["kind"] == kind)
    entry["value"] = "999"
    _reseal_payload(ledger)

    with pytest.raises(ValueError, match="not independently derived"):
        core.validate_numeric_evidence_closure(output_dir, ledger)


def test_numeric_evidence_replay_rejects_fully_rehashed_rendered_difference(
    tmp_path: Path,
) -> None:
    core, output_dir = _build_multirow_qualified_concordato_run(tmp_path)
    ledger_path = output_dir / "numeric_evidence_ledger.json"
    ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
    difference = next(
        entry
        for entry in ledger["entries"]
        if entry["kind"] == "signed_difference"
        and any(
            address["path"] == "concordato_review_summary.docx"
            for address in entry["output_addresses"]
        )
    )

    csv_changes: list[tuple[int, str]] = []
    xlsx_changes: list[tuple[str, str]] = []
    docx_changes: list[tuple[int, int, int]] = []
    for address in difference["output_addresses"]:
        locator = address["locator"]
        if address["path"].endswith(".csv"):
            match = re.fullmatch(
                r"row=(\d+);column=([A-Za-z0-9_]+)",
                locator,
            )
            assert match
            csv_changes.append((int(match.group(1)), match.group(2)))
            address["rendered_value"] = "999"
        elif address["path"].endswith(".xlsx"):
            sheet, coordinate = locator.rsplit("!", 1)
            xlsx_changes.append((sheet, coordinate))
            address["rendered_value"] = "999"
        elif address["path"].endswith(".docx"):
            match = re.fullmatch(
                r"table=(\d+);row=(\d+);column=(\d+)",
                locator,
            )
            assert match
            docx_changes.append(tuple(int(value) for value in match.groups()))
            address["rendered_value"] = "999,00"

    csv_path = output_dir / "exact_amount_matches.csv"
    with csv_path.open(encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        fieldnames = list(reader.fieldnames or [])
        rows = list(reader)
    for row_index, column in csv_changes:
        rows[row_index - 1][column] = "999"
    with csv_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    workbook_path = output_dir / "concordato_tie_out_workpaper.xlsx"
    workbook = openpyxl.load_workbook(workbook_path)
    for sheet, coordinate in xlsx_changes:
        workbook[sheet][coordinate] = 999
    workbook.save(workbook_path)

    docx_path = output_dir / "concordato_review_summary.docx"
    document = Document(docx_path)
    for table_index, row_index, column_index in docx_changes:
        document.tables[table_index - 1].rows[row_index - 1].cells[
            column_index - 1
        ].text = "999,00"
    document.save(docx_path)
    _reseal_payload(ledger)
    _write_json(ledger_path, ledger)

    changed_paths = {
        "exact_amount_matches.csv",
        "concordato_tie_out_workpaper.xlsx",
        "concordato_review_summary.docx",
        "numeric_evidence_ledger.json",
    }
    envelope_path = output_dir / "assurance_envelope.json"
    envelope = json.loads(envelope_path.read_text(encoding="utf-8"))
    for receipt in envelope["artifact_receipts"]:
        if receipt["root_id"] == "run" and receipt["path"] in changed_paths:
            _refresh_receipt(receipt, output_dir)
    _reseal_payload(envelope)
    _write_json(envelope_path, envelope)

    review_path = output_dir / "review_payload.json"
    review_payload = json.loads(review_path.read_text(encoding="utf-8"))
    review_payload["assurance"]["envelope_content_sha256"] = envelope["content_sha256"]
    _reseal_payload(review_payload)
    _write_json(review_path, review_payload)

    final_path = output_dir / "final_artifacts.json"
    final_artifacts = json.loads(final_path.read_text(encoding="utf-8"))
    for output in final_artifacts["outputs"]:
        if output["path"] in {
            *changed_paths,
            "assurance_envelope.json",
        }:
            target = output_dir / output["path"]
            output["size_bytes"] = target.stat().st_size
            output["sha256"] = hashlib.sha256(target.read_bytes()).hexdigest()
    _write_json(final_path, final_artifacts)

    closure_path = output_dir / "workflow_output_closure.json"
    closure = json.loads(closure_path.read_text(encoding="utf-8"))
    closure["assurance_envelope_content_sha256"] = envelope["content_sha256"]
    for receipt in closure["artifact_receipts"]:
        if receipt["path"] in {
            *changed_paths,
            "assurance_envelope.json",
            "review_payload.json",
            "final_artifacts.json",
        }:
            _refresh_receipt(receipt, output_dir)
    _reseal_payload(closure)
    _write_json(closure_path, closure)

    from output_closure import validate_output_closure
    from replay_assurance import replay_assurance

    validate_output_closure(output_dir)
    with pytest.raises(ValueError, match="not independently derived"):
        replay_assurance(output_dir)


def test_numeric_evidence_replay_rejects_last_csv_match_row_mutation(
    tmp_path: Path,
) -> None:
    core, output_dir = _build_multirow_qualified_concordato_run(tmp_path)
    path = output_dir / "exact_amount_matches.csv"
    with path.open(encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        fieldnames = list(reader.fieldnames or [])
        rows = list(reader)
    rows[2]["difference"] = "1"
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    with pytest.raises(ValueError, match="does not match"):
        core.validate_numeric_evidence_closure(output_dir)


def test_numeric_evidence_replay_rejects_last_workbook_match_row_mutation(
    tmp_path: Path,
) -> None:
    core, output_dir = _build_multirow_qualified_concordato_run(tmp_path)
    path = output_dir / "concordato_tie_out_workpaper.xlsx"
    workbook = openpyxl.load_workbook(path)
    workbook["Candidate matches"]["P4"] = 1
    workbook.save(path)

    with pytest.raises(ValueError, match="does not match"):
        core.validate_numeric_evidence_closure(output_dir)


def test_numeric_evidence_replay_rejects_last_word_match_row_mutation(
    tmp_path: Path,
) -> None:
    core, output_dir = _build_multirow_qualified_concordato_run(tmp_path)
    path = output_dir / "concordato_review_summary.docx"
    document = Document(path)
    document.tables[2].rows[3].cells[4].text = "1,00"
    document.save(path)

    with pytest.raises(ValueError, match="does not match"):
        core.validate_numeric_evidence_closure(output_dir)


def test_exact_amount_match_preserves_high_precision_decimal_boundary() -> None:
    core = load_core()
    candidates = [
        core.AmountCandidate(
            "plan.xlsx",
            "concordato_plan",
            "Plan!A1",
            Decimal("100.000000000000000001"),
            "100.000000000000000001",
            "trade payables",
        ),
        core.AmountCandidate(
            "support.xlsx",
            "trial_balance",
            "TB!B2",
            Decimal("100"),
            "100",
            "trade payables",
        ),
    ]

    matches = core.find_exact_amount_matches(
        candidates,
        tolerance="0.000000000000000001",
    )

    assert matches[0]["difference"] == "0.000000000000000001"
    assert matches[0]["abs_difference"] == "0.000000000000000001"
    assert matches[0]["within_tolerance"] == "true"


def test_exact_amount_match_is_independent_of_global_decimal_precision() -> None:
    core = load_core()
    huge = Decimal("12345678901234567890123456789.01")
    candidates = [
        core.AmountCandidate(
            "plan.xlsx",
            "concordato_plan",
            "Plan!A1",
            huge,
            str(huge),
            "large exact amount",
        ),
        core.AmountCandidate(
            "support.xlsx",
            "trial_balance",
            "TB!B2",
            Decimal("0"),
            "0",
            "large exact amount",
        ),
    ]

    with localcontext() as context:
        context.prec = 6
        matches = core.find_exact_amount_matches(candidates, tolerance=str(huge))

    assert matches[0]["difference"] == str(huge)
    assert matches[0]["abs_difference"] == str(huge)
    assert matches[0]["within_tolerance"] == "true"


@pytest.mark.parametrize(
    "value",
    [
        Decimal("1.0000000000000000001"),
        Decimal("123456789012345678901234567890123456789"),
    ],
)
def test_exact_amount_match_rejects_values_outside_decimal_contract(
    value: Decimal,
) -> None:
    core = load_core()
    candidates = [
        core.AmountCandidate(
            "plan.xlsx",
            "concordato_plan",
            "Plan!A1",
            value,
            str(value),
            "over-bound amount",
        ),
        core.AmountCandidate(
            "support.xlsx",
            "trial_balance",
            "TB!B2",
            Decimal("0"),
            "0",
            "support",
        ),
    ]

    with pytest.raises(ValueError, match="decimal"):
        core.find_exact_amount_matches(candidates, tolerance="0")


def test_qualified_numeric_outputs_are_repeatable_for_same_reviewed_contract(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["Debiti tributari", 100]],
    )
    _save_workbook(
        input_dir / "supporto.xlsx",
        [["Voce", "Saldo"], ["Debiti tributari", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {
            "piano.xlsx": "concordato_plan",
            "supporto.xlsx": "other_support",
        },
    )
    first = tmp_path / "first"
    second = tmp_path / "second"

    core.run_concordato_review(
        input_dir,
        first,
        tolerance="0",
        recipe=recipe,
    )
    core.run_concordato_review(
        input_dir,
        second,
        tolerance="0",
        recipe=recipe,
    )

    assert (first / "amount_candidates.csv").read_bytes() == (
        second / "amount_candidates.csv"
    ).read_bytes()
    assert (first / "exact_amount_matches.csv").read_bytes() == (
        second / "exact_amount_matches.csv"
    ).read_bytes()
    assert json.loads((first / "numeric_evidence_ledger.json").read_text()) == (
        json.loads((second / "numeric_evidence_ledger.json").read_text())
    )


def test_workflow_output_closure_replays_exact_initial_allowlist(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    from output_closure import validate_output_closure

    closure = validate_output_closure(output_dir)
    physical = {
        path.relative_to(output_dir).as_posix()
        for path in output_dir.rglob("*")
        if path.is_file()
    }

    assert closure["phase"] == "initial_run_finalization"
    assert closure["previous_closure_content_sha256"] is None
    assert set(closure["declared_paths"]) == physical - {"workflow_output_closure.json"}


def test_workflow_output_closure_rejects_unexpected_regular_file(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    from output_closure import validate_output_closure

    (output_dir / "rogue-output.json").write_text(
        '{"unexpected":true}\n',
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="file sets do not match"):
        validate_output_closure(output_dir)


def test_workflow_output_closure_rejects_unexpected_empty_directory(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    from output_closure import validate_output_closure

    (output_dir / "rogue-empty").mkdir()

    with pytest.raises(ValueError, match="directory set does not close"):
        validate_output_closure(output_dir)


def test_workflow_output_closure_rejects_missing_declared_file(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    from output_closure import validate_output_closure

    (output_dir / "review_packet.md").unlink()

    with pytest.raises(ValueError, match="file sets do not match"):
        validate_output_closure(output_dir)


def test_workflow_output_closure_rejects_symlink(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    from output_closure import validate_output_closure

    (output_dir / "rogue-link").symlink_to(output_dir / "run_audit.json")

    with pytest.raises(ValueError, match="contains a symlink"):
        validate_output_closure(output_dir)


def test_workflow_output_closure_rejects_expected_path_hardlink(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    from output_closure import validate_output_closure

    target = output_dir / "review_packet.md"
    external = tmp_path / "external-review-packet.md"
    external.write_bytes(target.read_bytes())
    target.unlink()
    os.link(external, target)

    assert target.stat().st_nlink == 2
    with pytest.raises(ValueError, match="hard link"):
        validate_output_closure(output_dir)


def test_standalone_replay_rejects_rehashed_review_packet_narrative(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    packet_path = output_dir / "review_packet.md"
    packet_path.write_bytes(packet_path.read_bytes() + b"\nREHASHED-MUTATION\n")

    final_path = output_dir / "final_artifacts.json"
    final_artifacts = json.loads(final_path.read_text(encoding="utf-8"))
    packet_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "review_packet.md"
    )
    packet_output["size_bytes"] = packet_path.stat().st_size
    packet_output["sha256"] = hashlib.sha256(packet_path.read_bytes()).hexdigest()
    _write_json(final_path, final_artifacts)

    closure_path = output_dir / "workflow_output_closure.json"
    closure = json.loads(closure_path.read_text(encoding="utf-8"))
    for receipt in closure["artifact_receipts"]:
        if receipt["path"] in {"review_packet.md", "final_artifacts.json"}:
            _refresh_receipt(receipt, output_dir)
    _reseal_payload(closure)
    _write_json(closure_path, closure)

    from output_closure import validate_output_closure
    from replay_assurance import replay_assurance

    validate_output_closure(output_dir)
    with pytest.raises(ValueError, match="review_packet.md.*not independently"):
        replay_assurance(output_dir)


def test_standalone_replay_rejects_rehashed_review_handoff_narrative(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    handoff_path = output_dir / "review_handoff.md"
    handoff_path.write_bytes(handoff_path.read_bytes() + b"\nREHASHED-MUTATION\n")

    final_path = output_dir / "final_artifacts.json"
    final_artifacts = json.loads(final_path.read_text(encoding="utf-8"))
    handoff_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "review_handoff.md"
    )
    handoff_output["size_bytes"] = handoff_path.stat().st_size
    handoff_output["sha256"] = hashlib.sha256(handoff_path.read_bytes()).hexdigest()
    _write_json(final_path, final_artifacts)

    closure_path = output_dir / "workflow_output_closure.json"
    closure = json.loads(closure_path.read_text(encoding="utf-8"))
    for receipt in closure["artifact_receipts"]:
        if receipt["path"] in {"review_handoff.md", "final_artifacts.json"}:
            _refresh_receipt(receipt, output_dir)
    _reseal_payload(closure)
    _write_json(closure_path, closure)

    from output_closure import validate_output_closure
    from replay_assurance import replay_assurance

    validate_output_closure(output_dir)
    with pytest.raises(ValueError, match="review_handoff.md.*not independently"):
        replay_assurance(output_dir)


def test_standalone_replay_rejects_stale_run_audit_output_index(
    tmp_path: Path,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    audit_path = output_dir / "run_audit.json"
    audit_path.write_bytes(audit_path.read_bytes() + b"\n")
    closure_path = output_dir / "workflow_output_closure.json"
    closure = json.loads(closure_path.read_text(encoding="utf-8"))
    receipt = next(
        item
        for item in closure["artifact_receipts"]
        if item["path"] == "run_audit.json"
    )
    _refresh_receipt(receipt, output_dir)
    _reseal_payload(closure)
    _write_json(closure_path, closure)

    from output_closure import validate_output_closure
    from replay_assurance import replay_assurance

    validate_output_closure(output_dir)
    with pytest.raises(ValueError, match="indexed final artifact size is stale"):
        replay_assurance(output_dir)


@pytest.mark.parametrize(
    "relative_path",
    [
        "run_audit.json",
        "review_payload.json",
        "review_handoff.md",
        "final_artifacts.json",
    ],
)
def test_workflow_output_closure_rejects_late_artifact_mutation(
    tmp_path: Path,
    relative_path: str,
) -> None:
    load_core()
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    from output_closure import validate_output_closure

    with (output_dir / relative_path).open("ab") as handle:
        handle.write(b"\nlate-mutation")

    with pytest.raises(ValueError):
        validate_output_closure(output_dir)


def test_spanish_mcp_runtime_feedback_handoff_and_errors(tmp_path: Path) -> None:
    review_payload = _seal_review_payload(
        {
            "schema_version": "1.0",
            "plugin": "concordato-plan-review",
            "workflow": "concordato-plan-review",
            "run_id": "concordato-es-runtime",
            "language": "es",
            "review_type": "concordato_preventivo_review",
            "items": [
                {
                    "id": "source-es-1",
                    "item_type": "source_inventory",
                    "title": "plan.xlsx",
                    "allowed_actions": ["accept", "skip"],
                    "recommended_action": "accept",
                }
            ],
            "item_count": 1,
            "assurance": {"final_ready": False},
            "status": "ready_for_review",
        }
    )
    run_intake = {
        "run_id": review_payload["run_id"],
        "language": "spa",
    }
    decision = {"item_id": "source-es-1", "action": "accept"}
    messages = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "initialize",
            "params": {"_meta": {"locale": "es-ES"}},
        },
        {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/call",
            "params": {
                "name": "validate_concordato_plan_review",
                "arguments": {"review_payload": review_payload},
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 3,
            "method": "tools/call",
            "params": {
                "name": "save_concordato_plan_decisions",
                "arguments": {
                    "review_payload": review_payload,
                    "decisions": [decision],
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 4,
            "method": "tools/call",
            "params": {
                "name": "apply_concordato_plan_decisions",
                "arguments": {
                    "run_intake": run_intake,
                    "review_payload": review_payload,
                    "decisions": [decision],
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 5,
            "method": "tools/call",
            "params": {
                "name": "validate_concordato_plan_review",
                "arguments": {"review_payload": {**review_payload, "items": "invalid"}},
            },
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}
    validation = responses[2]["result"]["structuredContent"]
    saved = responses[3]["result"]["structuredContent"]
    applied = responses[4]["result"]["structuredContent"]
    invalid = responses[5]["result"]["structuredContent"]

    assert (
        "Use validate_concordato_plan_review antes"
        in responses[1]["result"]["instructions"]
    )
    assert validation["message"].startswith("Los datos de revisión")
    assert saved["message"].startswith("Las decisiones son válidas")
    assert applied["message"].startswith("Las decisiones aplicadas son válidas")
    assert applied["persisted"] is False
    assert applied["application_status"] == "review_applied_assurance_withheld"
    assert invalid["error"] == "review_payload.items debe ser una matriz"


def test_spanish_run_localizes_review_packet_workbook_and_contract(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "Empresa plan concordato.xlsx",
        [
            ["Concepto", "Importe"],
            ["Deuda tributaria", 250000.25],
        ],
    )

    inspection_run = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        reference_date="2026-01-31",
        language="es-ES",
        document_language="es-ES",
        tolerance="0.01",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection_run,
        {"Empresa plan concordato.xlsx": "concordato_plan"},
    )
    core.run_concordato_review(
        input_dir,
        output_dir,
        reference_date="2026-01-31",
        language="es-ES",
        document_language="es-ES",
        tolerance="0.01",
        recipe=recipe,
    )

    review_packet = (output_dir / "review_packet.md").read_text(encoding="utf-8")
    workbook = openpyxl.load_workbook(
        output_dir / "concordato_tie_out_workpaper.xlsx",
        data_only=True,
        read_only=True,
    )
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    review_handoff = (output_dir / "review_handoff.md").read_text(encoding="utf-8")
    document = Document(output_dir / "concordato_review_summary.docx")
    document_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    outputs_by_path = {output["path"]: output for output in final_artifacts["outputs"]}
    unmatched_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "unmatched_plan_amount"
    )
    artifact_item = next(
        item
        for item in review_payload["items"]
        if item["output_path"] == "concordato_review_summary.docx"
    )

    assert review_packet.startswith("# Paquete de revisión del concordato preventivo\n")
    assert "- Idioma de trabajo: `es`" in review_packet
    assert "## Anexo determinista de conciliación numérica" in review_packet
    assert "## Revisión profesional requerida" in review_packet
    assert workbook.sheetnames == [
        "Inventario",
        "Importes candidatos",
        "Coincidencias candidatas",
    ]
    assert workbook["Inventario"]["A1"].value == "path"
    assert workbook["Importes candidatos"]["D1"].value == "source_role"
    assert outputs_by_path["review_packet.md"]["required_text"] == [
        "# Paquete de revisión del concordato preventivo",
        "## Revisión profesional requerida",
        "## Anexo determinista de conciliación numérica",
    ]
    assert outputs_by_path["concordato_tie_out_workpaper.xlsx"]["required_sheets"] == [
        "Inventario",
        "Importes candidatos",
        "Coincidencias candidatas",
    ]
    assert set(
        outputs_by_path["concordato_tie_out_workpaper.xlsx"]["required_sheet_headers"]
    ) == {
        "Inventario",
        "Importes candidatos",
        "Coincidencias candidatas",
    }
    assert review_payload["language"] == "es"
    assert review_payload["document_language"] == "es"
    assert review_payload["columns"] == [
        {"field": "item_type", "label": "Tipo"},
        {"field": "title", "label": "Elemento"},
        {"field": "recommended_action", "label": "Acción sugerida"},
        {"field": "source_path", "label": "Fuente"},
        {"field": "output_path", "label": "Salida"},
        {"field": "status", "label": "Estado"},
    ]
    assert unmatched_item["recommended_action"] == "request_more_documents"
    assert unmatched_item["data"]["requested_document"].startswith(
        "Justificante o anexo explicativo para el importe del plan del concordato preventivo"
    )
    assert unmatched_item["data"]["reason"].startswith(
        "Ningún importe justificativo determinista"
    )
    assert unmatched_item["data"]["review_note"].startswith(
        "Ningún importe de origen coincide"
    )
    assert artifact_item["title"] == "Resumen numérico en Word"
    assert review_handoff.startswith(
        "# Entrega para revisión: Revisión del concordato preventivo\n"
    )
    assert "## Revisión en Codex" in review_handoff
    assert outputs_by_path["review_handoff.md"]["required_text"][0] == (
        "Entrega para revisión"
    )
    assert outputs_by_path["concordato_review_summary.docx"]["required_text"] == [
        "Anexo numérico del concordato preventivo",
        "Conclusión operativa",
        "Aspectos que deben explicarse en el memorando de revisión",
        "Archivos analizados",
    ]
    assert "Anexo numérico del concordato preventivo" in document_text
    assert "Conclusión operativa" in document_text
    assert "Archivos analizados" in document_text
    assert "Fuentes reconocidas" in document_text
    assert "Ejemplos de importes que no coinciden" in document_text
    assert "Aspectos que deben explicarse" in document_text
    assert "Appendice numerica del concordato preventivo" not in document_text
    assert "Conclusione operativa" not in document_text
    assert any(
        caveat.startswith("Las coincidencias exactas por importe")
        for caveat in final_artifacts["caveats"]
    )
    assert final_artifacts["next_actions"][2].startswith(
        "Revise procedimiento, acreedores, tratamiento, liquidez"
    )
    contract_report = validate_contract(
        output_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()


def test_concordato_request_more_documents_prefills_blocker_context(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "case_plan.xlsx",
        [
            ["Voce", "Importo"],
            ["Debiti tributari entro 12 mesi", 4124413.15],
            ["Assunzione prospettica non storica", 999999.99],
        ],
    )
    _save_workbook(
        input_dir / "supporting_ledger.xlsx",
        [
            ["Voce", "Saldo rettificato"],
            ["Debiti tributari entro 12 mesi", 4124413.15],
        ],
    )

    inspection_run = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        reference_date="2026-03-31",
        language="it",
        document_language="it",
        tolerance="0.01",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection_run,
        {
            "case_plan.xlsx": "concordato_plan",
            "supporting_ledger.xlsx": "other_support",
        },
    )
    core.run_concordato_review(
        input_dir,
        output_dir,
        reference_date="2026-03-31",
        language="it",
        document_language="it",
        tolerance="0.01",
        recipe=recipe,
    )
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    unmatched_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "unmatched_plan_amount"
    )

    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "tools/call",
            "params": {
                "name": "apply_concordato_plan_decisions",
                "arguments": {
                    "run_intake": run_intake,
                    "review_payload": review_payload,
                    "final_artifacts": final_artifacts,
                    "decisions": [
                        {
                            "item_id": unmatched_item["id"],
                            "action": "request_more_documents",
                            "reviewer_note": "Ask for support of the unmatched plan amount.",
                        }
                    ],
                    "decision_source": "pytest_unmatched_plan_request",
                    "reviewer": "pytest",
                },
            },
        }
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}
    payload = responses[1]["result"]["structuredContent"]
    applied = json.loads(
        (output_dir / "applied_decisions.json").read_text(encoding="utf-8")
    )
    updated_final = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    expected_document = unmatched_item["data"]["requested_document"]

    assert payload["ok"] is True
    assert payload["application_status"] == "blocked"
    assert applied["effects"][0]["requested_documents"] == [expected_document]
    assert applied["effects"][0]["followup_context"]["source_file"] == (
        unmatched_item["data"]["source_file"]
    )
    assert applied["effects"][0]["followup_context"]["source_table"] == (
        unmatched_item["data"]["source_table"]
    )
    assert applied["effects"][0]["followup_context"]["amount"] == "999,999.99"
    assert updated_final["blockers"][0]["requested_documents"] == [expected_document]
    assert updated_final["blockers"][0]["followup_context"]["reason"] == (
        "No deterministic support amount matched this plan amount within tolerance."
    )


def test_skill_and_scripts_keep_report_builder_out_of_the_workflow() -> None:
    skill_text = (
        ROOT
        / "plugins"
        / "concordato-plan-review"
        / "skills"
        / "concordato-plan-review"
        / "SKILL.md"
    ).read_text(encoding="utf-8")
    script_text = "\n".join(
        path.read_text(encoding="utf-8") for path in SCRIPT_DIR.glob("*.py")
    )

    assert "This is not a general report builder" in skill_text
    assert "scripts/check_dependencies.py" in skill_text
    assert "requirements" in skill_text.lower()
    assert "Never write run outputs inside this Git workspace" in skill_text
    assert "Keep the improvement note local" in skill_text
    assert "validate_concordato_plan_review" in skill_text
    assert "render_concordato_plan_review" in skill_text
    assert "report-builder" not in script_text
    assert "modules.llm" not in script_text
    assert "model_router" not in script_text


def test_static_page_exposes_concordato_specific_outputs() -> None:
    page = (
        ROOT / "static" / "shared" / "concordato-plan-review" / "index.html"
    ).read_text(encoding="utf-8")

    for snippet in (
        "Riesamina il caso, non soltanto i numeri del piano",
        "concordato_case_model.json",
        "creditor_treatment.csv",
        "sources_and_uses.csv",
        "liquidity_schedule.csv",
        "concordato_review_workpaper.xlsx",
        "concordato_preventivo_review_summary.docx",
        "concordato_tie_out_workpaper.xlsx",
        "../vera/index.html",
        "Concordato preventivo",
        "Vera",
    ):
        assert snippet in page


def test_concordato_mcp_server_validates_and_renders_review_payload() -> None:
    review_payload = _seal_review_payload(
        {
            "schema_version": "1.0",
            "plugin": "concordato-plan-review",
            "workflow": "concordato-plan-review",
            "run_id": "concordato-test-run",
            "review_type": "concordato_preventivo_review",
            "items": [
                {
                    "id": "source-1",
                    "item_type": "source_inventory",
                    "title": "piano.xlsx",
                    "source_path": "/tmp/piano.xlsx",
                    "output_path": None,
                    "allowed_actions": ["accept", "edit", "mark_unclear", "skip"],
                    "recommended_action": "accept",
                    "evidence": [],
                    "data": {"suggested_role": "concordato_plan"},
                    "status": "needs_review",
                },
                {
                    "id": "unmatched-plan-amount-1",
                    "item_type": "unmatched_plan_amount",
                    "title": "piano.xlsx Dati!B3 999,999.99",
                    "source_path": "piano.xlsx",
                    "output_path": "amount_candidates.csv",
                    "allowed_actions": [
                        "accept",
                        "edit",
                        "mark_unclear",
                        "request_more_documents",
                        "skip",
                    ],
                    "recommended_action": "request_more_documents",
                    "evidence": [{"kind": "plan_context", "text": "Assunzione"}],
                    "data": {
                        "amount": 999999.99,
                        "match_status": "no_candidate_amount_match",
                    },
                    "status": "needs_review",
                },
            ],
            "item_count": 2,
            "columns": [],
            "evidence": {},
            "allowed_actions": [
                "accept",
                "edit",
                "mark_unclear",
                "request_more_documents",
                "skip",
            ],
            "status": "ready_for_review",
            "assurance": {"final_ready": False},
            "summary": {
                "file_count": 1,
                "plan_amount_candidate_count": 1,
                "unmatched_plan_amount_count": 1,
            },
        }
    )
    messages: list[dict[str, object]] = [
        {"jsonrpc": "2.0", "id": 1, "method": "tools/list"},
        {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/call",
            "params": {
                "name": "validate_concordato_plan_review",
                "arguments": {"review_payload": review_payload},
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 3,
            "method": "tools/call",
            "params": {
                "name": "render_concordato_plan_review",
                "arguments": {"review_payload": review_payload},
            },
        },
        {"jsonrpc": "2.0", "id": 4, "method": "resources/list"},
        {
            "jsonrpc": "2.0",
            "id": 5,
            "method": "resources/read",
            "params": {"uri": "ui://widget/concordato-plan-review.html"},
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}

    tool_names = {tool["name"] for tool in responses[1]["result"]["tools"]}
    assert {
        "validate_concordato_plan_review",
        "render_concordato_plan_review",
    } <= tool_names
    validate_result = responses[2]["result"]["structuredContent"]
    assert validate_result["ok"] is True
    assert validate_result["item_count"] == 2
    render_result = responses[3]["result"]
    assert render_result["structuredContent"]["widget_type"] == "concordato_plan_review"
    assert (
        render_result["_meta"]["openai/outputTemplate"]
        == "ui://widget/concordato-plan-review.html"
    )
    resource_uris = {
        resource["uri"] for resource in responses[4]["result"]["resources"]
    }
    assert "ui://widget/concordato-plan-review.html" in resource_uris
    widget_html = responses[5]["result"]["contents"][0]["text"]
    assert "Concordato Preventivo Review" in widget_html


def test_concordato_mcp_apply_creates_codex_review_memo_from_edit(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "concordato"
    input_dir.mkdir()
    _save_workbook(
        input_dir / "piano.xlsx",
        [["Voce", "Importo"], ["Debiti tributari", 100]],
    )
    inspection = core.run_concordato_review(
        input_dir,
        tmp_path / "inspection",
        tolerance="0",
    )
    recipe = _reviewed_source_recipe(
        core,
        inspection,
        {"piano.xlsx": "concordato_plan"},
    )
    core.run_concordato_review(
        input_dir,
        output_dir,
        tolerance="0",
        recipe=recipe,
    )
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    memo_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "codex_review_memo"
    )
    memo_text = (
        "# Codex review memo\n\n"
        "Il piano batte per importo su una voce e richiede evidenza per "
        "l'assunzione prospettica."
    )
    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "tools/call",
            "params": {
                "name": "apply_concordato_plan_decisions",
                "arguments": {
                    "run_intake": run_intake,
                    "review_payload": review_payload,
                    "final_artifacts": final_artifacts,
                    "decisions": [
                        {
                            "item_id": memo_item["id"],
                            "action": "edit",
                            "edit_value": memo_text,
                            "reviewer_note": "Use this memo for handoff.",
                        }
                    ],
                },
            },
        }
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}

    payload = responses[1]["result"]["structuredContent"]
    assert payload["ok"] is True
    assert payload["application_status"] == "partial_review_applied"
    assert payload["target_update_count"] == 1
    assert payload["native_regeneration_count"] == 0
    assert payload["native_regenerated_count"] == 1
    assert payload["run_intake_path"] == str(output_dir / "run_intake.json")
    assert (output_dir / "codex_run_review.md").read_text(encoding="utf-8") == memo_text
    updated_docx_text = _docx_text(
        output_dir / "concordato_preventivo_review_summary.docx"
    )
    assert "Memo revisore Codex" in updated_docx_text
    assert "Il piano batte per importo su una voce" in updated_docx_text
    applied = json.loads(
        (output_dir / "applied_decisions.json").read_text(encoding="utf-8")
    )
    assert applied["target_update_paths"] == ["codex_run_review.md"]
    assert applied["effects"][0]["artifact_update"] == "target_artifact_created"
    assert applied["effects"][0]["promoted_from_revision"].startswith(
        "revisions/codex_run_review__codex-review-memo"
    )
    assert applied["effects"][0]["native_regeneration_status"] == "regenerated"
    assert applied["native_regenerated_paths"] == [
        "concordato_preventivo_review_summary.docx"
    ]
    updated_final = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    assert updated_final["final_ready"] is False
    assert updated_final["review_application"]["target_update_paths"] == [
        "codex_run_review.md"
    ]
    memo_output = next(
        output
        for output in updated_final["outputs"]
        if output["path"] == "codex_run_review.md"
    )
    assert memo_output["status"] == "updated_from_review"
    summary_output = next(
        output
        for output in updated_final["outputs"]
        if output["path"] == "concordato_preventivo_review_summary.docx"
    )
    assert summary_output["status"] == "updated_from_review"
    assert summary_output["native_regenerated"] is True
    assert "Memo revisore Codex" in summary_output["required_text"]
    assert (
        "Il piano batte per importo su una voce" in summary_output["required_text"][2]
    )
    assert updated_final["review_application"]["native_regenerated_paths"] == [
        "concordato_preventivo_review_summary.docx"
    ]
    assurance_envelope = json.loads(
        (output_dir / "assurance_envelope.json").read_text(encoding="utf-8")
    )
    assert "concordato_preventivo_review_summary.docx" not in {
        receipt["path"] for receipt in assurance_envelope["artifact_receipts"]
    }
    output_closure = json.loads(
        (output_dir / "workflow_output_closure.json").read_text(encoding="utf-8")
    )
    assert "concordato_preventivo_review_summary.docx" in {
        receipt["path"] for receipt in output_closure["artifact_receipts"]
    }
    assert output_closure["phase"] == "review_apply_finalization"
    updated_run_intake = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )
    review_apply_steps = [
        step
        for step in updated_run_intake["execution_trace"]
        if step["kind"] == "deterministic_review_apply"
    ]
    assert len(review_apply_steps) == 1
    assert {
        "applied_decisions.json",
        "codex_run_review.md",
        "concordato_preventivo_review_summary.docx",
        "final_artifacts.json",
        "ui_decisions.json",
    } <= set(review_apply_steps[0]["outputs"])
    contract_report = validate_contract(
        output_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()


def test_concordato_mcp_rejects_forged_review_payload_before_write(
    tmp_path: Path,
) -> None:
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    persisted = json.loads((output_dir / "review_payload.json").read_text())
    forged = json.loads(json.dumps(persisted))
    forged_item = next(
        item for item in forged["items"] if item["item_type"] == "codex_review_memo"
    )
    forged_item["output_path"] = "review_packet.md"
    original_packet = (output_dir / "review_packet.md").read_text()

    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {
                    "name": "apply_concordato_plan_decisions",
                    "arguments": {
                        "run_intake": run_intake,
                        "review_payload": forged,
                        "decisions": [
                            {
                                "item_id": forged_item["id"],
                                "action": "edit",
                                "edit_value": "FORGED",
                            }
                        ],
                    },
                },
            }
        ]
    )[0]["result"]["structuredContent"]

    assert response["ok"] is False
    assert "content_sha256 is stale" in response["error"]
    assert (output_dir / "review_packet.md").read_text() == original_packet
    assert not (output_dir / "applied_decisions.json").exists()


@pytest.mark.parametrize(
    "tool_name",
    [
        "validate_concordato_plan_review",
        "render_concordato_plan_review",
    ],
)
def test_concordato_mcp_read_only_tools_reject_stale_review_digest(
    tmp_path: Path,
    tool_name: str,
) -> None:
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    review_payload["items"][0]["title"] = "FORGED REVIEW TITLE"

    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {
                    "name": tool_name,
                    "arguments": {"review_payload": review_payload},
                },
            }
        ]
    )[0]["result"]["structuredContent"]

    assert response["ok"] is False
    assert "content_sha256 is stale" in response["error"]


@pytest.mark.parametrize(
    "tool_name",
    [
        "validate_concordato_plan_review",
        "render_concordato_plan_review",
    ],
)
def test_concordato_mcp_read_only_tools_replay_persisted_review_context(
    tmp_path: Path,
    tool_name: str,
) -> None:
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    persisted = json.loads((output_dir / "review_payload.json").read_text())
    forged_content = dict(persisted)
    forged_content.pop("content_sha256")
    forged_content["items"] = json.loads(json.dumps(persisted["items"]))
    forged_content["items"][0]["title"] = "RESEALED FORGED REVIEW TITLE"
    forged = _seal_review_payload(forged_content)

    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {
                    "name": tool_name,
                    "arguments": {
                        "run_intake": run_intake,
                        "review_payload": forged,
                    },
                },
            }
        ]
    )[0]["result"]["structuredContent"]

    assert response["ok"] is False
    assert "persisted review payload" in response["error"].lower()


def test_concordato_mcp_rejects_forged_final_artifacts_before_write(
    tmp_path: Path,
) -> None:
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    persisted_final = json.loads((output_dir / "final_artifacts.json").read_text())
    forged_final = json.loads(json.dumps(persisted_final))
    forged_final["assurance"]["final_ready"] = True
    forged_final["assurance"]["gate_register"]["report_ready"] = True
    forged_final["caveats"] = []
    forged_final["next_actions"] = ["Publish"]
    item = review_payload["items"][0]
    action = "skip" if "skip" in item["allowed_actions"] else "accept"

    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {
                    "name": "apply_concordato_plan_decisions",
                    "arguments": {
                        "run_intake": run_intake,
                        "review_payload": review_payload,
                        "final_artifacts": forged_final,
                        "decisions": [{"item_id": item["id"], "action": action}],
                    },
                },
            }
        ]
    )[0]["result"]["structuredContent"]

    assert response["ok"] is False
    assert "persisted final artifacts" in response["error"].lower()
    assert not (output_dir / "applied_decisions.json").exists()
    assert json.loads((output_dir / "final_artifacts.json").read_text()) == (
        persisted_final
    )


def test_concordato_mcp_rejects_tampered_numeric_output_before_write(
    tmp_path: Path,
) -> None:
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    matches_path = output_dir / "exact_amount_matches.csv"
    matches_path.write_text(
        matches_path.read_text().replace(",100,", ",999,", 1),
        encoding="utf-8",
    )
    item = review_payload["items"][0]
    action = "skip" if "skip" in item["allowed_actions"] else "accept"

    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {
                    "name": "apply_concordato_plan_decisions",
                    "arguments": {
                        "run_intake": run_intake,
                        "review_payload": review_payload,
                        "decisions": [{"item_id": item["id"], "action": action}],
                    },
                },
            }
        ]
    )[0]["result"]["structuredContent"]

    assert response["ok"] is False
    assert "assurance replay failed" in response["error"].lower()
    assert not (output_dir / "applied_decisions.json").exists()


def test_complete_concordato_review_remains_assurance_withheld(
    tmp_path: Path,
) -> None:
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    decisions = [
        {
            "item_id": item["id"],
            "action": "skip" if "skip" in item["allowed_actions"] else "accept",
        }
        for item in review_payload["items"]
    ]

    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {
                    "name": "apply_concordato_plan_decisions",
                    "arguments": {
                        "run_intake": run_intake,
                        "review_payload": review_payload,
                        "decisions": decisions,
                    },
                },
            }
        ]
    )[0]["result"]["structuredContent"]

    assert response["ok"] is True
    assert response["application_status"] == "review_applied_assurance_withheld"
    assert response["final_artifacts"]["final_ready"] is False
    assert (
        response["final_artifacts"]["assurance"]["gate_register"]["report_ready"]
        is False
    )


def _concordato_transaction_case(
    tmp_path: Path,
) -> tuple[Path, dict[str, Any]]:
    _, _, output_dir = _build_qualified_concordato_run(tmp_path)
    output_dir.chmod(0o750)
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    decisions = [
        {
            "item_id": item["id"],
            "action": "skip" if "skip" in item["allowed_actions"] else "accept",
        }
        for item in review_payload["items"]
    ]
    return output_dir, {
        "run_intake": run_intake,
        "review_payload": review_payload,
        "final_artifacts": final_artifacts,
        "decisions": decisions,
    }


def _concordato_memo_arguments(arguments: dict[str, Any]) -> dict[str, Any]:
    review_payload = arguments["review_payload"]
    memo_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "codex_review_memo"
    )
    return {
        **arguments,
        "decisions": [
            {
                "item_id": memo_item["id"],
                "action": "edit",
                "edit_value": (
                    "# Memo revisore\n\n"
                    "La corrispondenza per importo richiede giudizio professionale."
                ),
            }
        ],
    }


def _concordato_tree_image(root: Path) -> dict[str, tuple[Any, ...]]:
    image: dict[str, tuple[Any, ...]] = {}
    for entry in [root, *sorted(root.rglob("*"))]:
        observed = entry.lstat()
        relative = "." if entry == root else entry.relative_to(root).as_posix()
        mode = observed.st_mode & 0o7777
        if stat.S_ISREG(observed.st_mode):
            image[relative] = (
                "file",
                mode,
                observed.st_nlink,
                entry.read_bytes(),
            )
        elif stat.S_ISDIR(observed.st_mode):
            image[relative] = ("directory", mode)
        elif stat.S_ISLNK(observed.st_mode):
            image[relative] = ("symlink", mode, os.readlink(entry))
        else:
            image[relative] = ("special", stat.S_IFMT(observed.st_mode), mode)
    return image


def _concordato_faulted_server(
    tmp_path: Path,
    *,
    needle: str,
    replacement: str,
) -> Path:
    source = MCP_SERVER_PATH.read_text(encoding="utf-8")
    plugin_root_line = 'const PLUGIN_ROOT = path.resolve(__dirname, "..");'
    assert source.count(plugin_root_line) == 1
    assert source.count(needle) == 1
    source = source.replace(
        plugin_root_line,
        f"const PLUGIN_ROOT = {json.dumps(str(MCP_SERVER_PATH.parents[1]))};",
        1,
    )
    source = source.replace(needle, replacement, 1)
    server_path = tmp_path / "concordato-faulted-server.cjs"
    server_path.write_text(source, encoding="utf-8")
    return server_path


def _concordato_transaction_call(
    tool_name: str,
    arguments: dict[str, Any],
    *,
    server_path: Path = MCP_SERVER_PATH,
    env: dict[str, str] | None = None,
) -> dict[str, Any]:
    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {"name": tool_name, "arguments": arguments},
            }
        ],
        server_path=server_path,
        env=env,
    )[0]
    return response["result"]["structuredContent"]


def _concordato_phase_child(
    tmp_path: Path,
    *,
    mode: str,
) -> Path:
    script = tmp_path / f"concordato-child-{mode}.sh"
    branch = (
        'echo "/private/client/raw-stdout"\n'
        'echo "Traceback: /private/client/raw-stderr" >&2\n'
        "exit 17\n"
        if mode == "nonzero"
        else 'echo "not-json /private/client/raw-output"\nexit 0\n'
    )
    script.write_text(
        "#!/bin/sh\n"
        'if [ "$(basename "$3")" = "replay_assurance.py" ]; then\n'
        '  exec "$REAL_PYTHON" "$@"\n'
        "fi\n"
        f"{branch}",
        encoding="utf-8",
    )
    script.chmod(0o700)
    return script


def _concordato_post_child_tamper(tmp_path: Path) -> Path:
    script = tmp_path / "concordato-post-child-tamper.py"
    script.write_text(
        f"#!{sys.executable}\n"
        "import json\n"
        "import os\n"
        "import pathlib\n"
        "import subprocess\n"
        "import sys\n"
        "real_python = os.environ['REAL_PYTHON']\n"
        "script_index = next(\n"
        "    index for index, value in enumerate(sys.argv[1:], 1)\n"
        "    if value.endswith('.py')\n"
        ")\n"
        "script_name = pathlib.Path(sys.argv[script_index]).name\n"
        "if script_name == 'replay_assurance.py':\n"
        "    os.execv(real_python, [real_python, *sys.argv[1:]])\n"
        "completed = subprocess.run(\n"
        "    [real_python, *sys.argv[1:]],\n"
        "    capture_output=True,\n"
        "    text=True,\n"
        "    check=False,\n"
        ")\n"
        "if completed.returncode != 0:\n"
        "    sys.stdout.write(completed.stdout)\n"
        "    sys.stderr.write(completed.stderr)\n"
        "    raise SystemExit(completed.returncode)\n"
        "result = json.loads(completed.stdout.strip().splitlines()[-1])\n"
        "arguments = sys.argv[script_index + 1:]\n"
        "output_dir = pathlib.Path(arguments[arguments.index('--output-dir') + 1])\n"
        "mode = os.environ['CONCORDATO_TAMPER_MODE']\n"
        "if mode == 'rogue_effect_paths':\n"
        "    rogue_path = output_dir / 'rogue.json'\n"
        "    rogue_path.write_text('{\"forged\":true}\\n', encoding='utf-8')\n"
        "    applied_path = pathlib.Path(\n"
        "        arguments[arguments.index('--applied-decisions') + 1]\n"
        "    )\n"
        "    final_path = pathlib.Path(\n"
        "        arguments[arguments.index('--final-artifacts') + 1]\n"
        "    )\n"
        "    applied = json.loads(applied_path.read_text(encoding='utf-8'))\n"
        "    final_artifacts = json.loads(final_path.read_text(encoding='utf-8'))\n"
        "    applied['effects'].append(\n"
        "        {'item_id': 'forged-child-effect', 'action': 'edit', "
        "'artifact_update': 'native_regenerated'}\n"
        "    )\n"
        "    regenerated = list(applied.get('native_regenerated_paths') or [])\n"
        "    regenerated.append('rogue.json')\n"
        "    applied['native_regenerated_paths'] = regenerated\n"
        "    applied['native_regenerated_count'] = len(regenerated)\n"
        "    application = final_artifacts['review_application']\n"
        "    application['native_regenerated_paths'] = regenerated\n"
        "    application['native_regenerated_count'] = len(regenerated)\n"
        "    final_artifacts['outputs'].append(\n"
        "        {'path': 'rogue.json', 'kind': 'json', 'status': 'forged'}\n"
        "    )\n"
        "    applied_path.write_text(\n"
        "        json.dumps(applied, ensure_ascii=False, indent=2) + '\\n',\n"
        "        encoding='utf-8',\n"
        "    )\n"
        "    final_path.write_text(\n"
        "        json.dumps(final_artifacts, ensure_ascii=False, indent=2) + '\\n',\n"
        "        encoding='utf-8',\n"
        "    )\n"
        "    result['native_regenerated_paths'] = regenerated\n"
        "    result['applied_decisions'] = applied\n"
        "    result['final_artifacts'] = final_artifacts\n"
        "elif mode == 'forged_contract':\n"
        "    applied_path = pathlib.Path(\n"
        "        arguments[arguments.index('--applied-decisions') + 1]\n"
        "    )\n"
        "    final_path = pathlib.Path(\n"
        "        arguments[arguments.index('--final-artifacts') + 1]\n"
        "    )\n"
        "    applied = json.loads(applied_path.read_text(encoding='utf-8'))\n"
        "    final_artifacts = json.loads(final_path.read_text(encoding='utf-8'))\n"
        "    applied['effects'].append(\n"
        "        {'item_id': 'forged-child-effect', 'action': 'accept'}\n"
        "    )\n"
        "    regenerated = [\n"
        "        'concordato_preventivo_review_summary.docx',\n"
        "        'concordato_preventivo_review_summary.docx',\n"
        "    ]\n"
        "    applied['native_regenerated_paths'] = regenerated\n"
        "    applied['native_regenerated_count'] = 2\n"
        "    application = final_artifacts['review_application']\n"
        "    application['native_regenerated_paths'] = regenerated\n"
        "    application['native_regenerated_count'] = 2\n"
        "    applied_path.write_text(\n"
        "        json.dumps(applied, ensure_ascii=False, indent=2) + '\\n',\n"
        "        encoding='utf-8',\n"
        "    )\n"
        "    final_path.write_text(\n"
        "        json.dumps(final_artifacts, ensure_ascii=False, indent=2) + '\\n',\n"
        "        encoding='utf-8',\n"
        "    )\n"
        "    result['native_regenerated_paths'] = regenerated\n"
        "    result['applied_decisions'] = applied\n"
        "    result['final_artifacts'] = final_artifacts\n"
        "elif mode == 'run_intake':\n"
        "    run_path = output_dir / 'run_intake.json'\n"
        "    run_intake = json.loads(run_path.read_text(encoding='utf-8'))\n"
        "    run_intake['forged_child_field'] = True\n"
        "    run_path.write_text(\n"
        "        json.dumps(run_intake, ensure_ascii=False, indent=2) + '\\n',\n"
        "        encoding='utf-8',\n"
        "    )\n"
        "elif mode == 'review_handoff':\n"
        "    with (output_dir / 'review_handoff.md').open('a', encoding='utf-8') as handle:\n"
        "        handle.write('\\n/private/client/forged-handoff\\n')\n"
        "elif mode == 'docx_content':\n"
        "    summary_path = output_dir / 'concordato_preventivo_review_summary.docx'\n"
        "    backup_path = (\n"
        "        output_dir\n"
        "        / 'revisions'\n"
        "        / 'originals'\n"
        "        / 'concordato_preventivo_review_summary__codex-review-memo.docx'\n"
        "    )\n"
        "    summary_path.write_bytes(backup_path.read_bytes())\n"
        "    final_path = pathlib.Path(\n"
        "        arguments[arguments.index('--final-artifacts') + 1]\n"
        "    )\n"
        "    final_artifacts = json.loads(final_path.read_text(encoding='utf-8'))\n"
        "    summary_output = next(\n"
        "        output\n"
        "        for output in final_artifacts['outputs']\n"
        "        if output.get('path') == 'concordato_preventivo_review_summary.docx'\n"
        "    )\n"
        "    summary_output['size_bytes'] = summary_path.stat().st_size\n"
        "    final_path.write_text(\n"
        "        json.dumps(final_artifacts, ensure_ascii=False, indent=2) + '\\n',\n"
        "        encoding='utf-8',\n"
        "    )\n"
        "    result['final_artifacts'] = final_artifacts\n"
        "elif mode == 'docx_numeric':\n"
        "    from docx import Document\n"
        "    summary_path = output_dir / 'concordato_preventivo_review_summary.docx'\n"
        "    document = Document(summary_path)\n"
        "    document.tables[0].rows[1].cells[1].text = '999'\n"
        "    document.save(summary_path)\n"
        "    final_path = pathlib.Path(\n"
        "        arguments[arguments.index('--final-artifacts') + 1]\n"
        "    )\n"
        "    final_artifacts = json.loads(final_path.read_text(encoding='utf-8'))\n"
        "    summary_output = next(\n"
        "        output\n"
        "        for output in final_artifacts['outputs']\n"
        "        if output.get('path') == 'concordato_preventivo_review_summary.docx'\n"
        "    )\n"
        "    summary_output['size_bytes'] = summary_path.stat().st_size\n"
        "    final_path.write_text(\n"
        "        json.dumps(final_artifacts, ensure_ascii=False, indent=2) + '\\n',\n"
        "        encoding='utf-8',\n"
        "    )\n"
        "    result['final_artifacts'] = final_artifacts\n"
        "else:\n"
        "    raise SystemExit(f'unknown tamper mode: {mode}')\n"
        "sys.stdout.write(json.dumps(result, ensure_ascii=False) + '\\n')\n",
        encoding="utf-8",
    )
    script.chmod(0o700)
    return script


def _concordato_forged_successor_chain(tmp_path: Path) -> Path:
    script = tmp_path / "concordato-forged-successor-chain.py"
    script.write_text(
        f"#!{sys.executable}\n"
        "import hashlib\n"
        "import json\n"
        "import os\n"
        "import pathlib\n"
        "import subprocess\n"
        "import sys\n"
        "real_python = os.environ['REAL_PYTHON']\n"
        "script_index = next(\n"
        "    index for index, value in enumerate(sys.argv[1:], 1)\n"
        "    if value.endswith('.py')\n"
        ")\n"
        "script_name = pathlib.Path(sys.argv[script_index]).name\n"
        "if script_name != 'finalize_output_closure.py':\n"
        "    os.execv(real_python, [real_python, *sys.argv[1:]])\n"
        "completed = subprocess.run(\n"
        "    [real_python, *sys.argv[1:]],\n"
        "    capture_output=True,\n"
        "    text=True,\n"
        "    check=False,\n"
        ")\n"
        "if completed.returncode != 0:\n"
        "    sys.stdout.write(completed.stdout)\n"
        "    sys.stderr.write(completed.stderr)\n"
        "    raise SystemExit(completed.returncode)\n"
        "arguments = sys.argv[script_index + 1:]\n"
        "output_dir = pathlib.Path(arguments[arguments.index('--output-dir') + 1])\n"
        "closure_path = output_dir / 'workflow_output_closure.json'\n"
        "closure = json.loads(closure_path.read_text(encoding='utf-8'))\n"
        "closure['previous_closure_content_sha256'] = '0' * 64\n"
        "content = dict(closure)\n"
        "content.pop('content_sha256')\n"
        "digest = hashlib.sha256(\n"
        "    json.dumps(\n"
        "        content,\n"
        "        ensure_ascii=False,\n"
        "        sort_keys=True,\n"
        "        separators=(',', ':'),\n"
        "    ).encode('utf-8')\n"
        ").hexdigest()\n"
        "closure['content_sha256'] = digest\n"
        "closure_path.write_text(\n"
        "    json.dumps(closure, ensure_ascii=False, indent=2) + '\\n',\n"
        "    encoding='utf-8',\n"
        ")\n"
        "result = json.loads(completed.stdout.strip().splitlines()[-1])\n"
        "result['content_sha256'] = digest\n"
        "sys.stdout.write(json.dumps(result, ensure_ascii=False) + '\\n')\n",
        encoding="utf-8",
    )
    script.chmod(0o700)
    return script


def test_concordato_review_transaction_honest_apply_commits_without_residue(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    predecessor = json.loads(
        (output_dir / "workflow_output_closure.json").read_text(encoding="utf-8")
    )

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
    )

    assert result["ok"] is True
    assert result["persisted"] is True
    assert result["application_status"] == "review_applied_assurance_withheld"
    assert result["final_artifacts"]["final_ready"] is False
    assert output_dir.stat().st_mode & 0o7777 == 0o750
    closure = json.loads(
        (output_dir / "workflow_output_closure.json").read_text(encoding="utf-8")
    )
    assert closure["phase"] == "review_apply_finalization"
    assert closure["previous_closure_content_sha256"] == predecessor["content_sha256"]
    from replay_assurance import replay_assurance

    replay = replay_assurance(output_dir)
    assert replay["workflow_output_closure_content_sha256"] == closure["content_sha256"]
    assert replay["workflow_output_closure_phase"] == "review_apply_finalization"
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_edit_of_assurance_bound_markdown_writes_revision_only(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    review_payload = arguments["review_payload"]
    semantic_item = next(
        item for item in review_payload["items"] if item["id"] == "semantic-review"
    )
    semantic_path = output_dir / "concordato_semantic_review.md"
    predecessor_bytes = semantic_path.read_bytes()
    edit_text = "Reviewer revision; predecessor semantic report remains immutable."
    arguments = {
        **arguments,
        "decisions": [
            {
                "item_id": semantic_item["id"],
                "action": "edit",
                "edit_value": edit_text,
            }
        ],
    }

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
    )

    assert result["ok"] is True
    assert result["target_update_count"] == 0
    assert semantic_path.read_bytes() == predecessor_bytes
    applied = json.loads(
        (output_dir / "applied_decisions.json").read_text(encoding="utf-8")
    )
    effect = applied["effects"][0]
    assert effect["artifact_update"] == "revision_artifact_written"
    revision_path = output_dir / effect["revision_artifact"]
    assert revision_path.read_text(encoding="utf-8") == edit_text
    from replay_assurance import replay_assurance

    replay = replay_assurance(output_dir)
    assert replay["workflow_output_closure_phase"] == "review_apply_finalization"


def test_standalone_replay_rejects_rehashed_applied_state(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
    )
    assert result["ok"] is True

    applied_path = output_dir / "applied_decisions.json"
    applied = json.loads(applied_path.read_text(encoding="utf-8"))
    applied["blocker_count"] = 999
    _write_json(applied_path, applied)

    final_path = output_dir / "final_artifacts.json"
    final_artifacts = json.loads(final_path.read_text(encoding="utf-8"))
    final_artifacts["review_application"]["blocker_count"] = 999
    applied_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "applied_decisions.json"
    )
    applied_output["size_bytes"] = applied_path.stat().st_size
    applied_output["sha256"] = hashlib.sha256(applied_path.read_bytes()).hexdigest()
    _write_json(final_path, final_artifacts)

    closure_path = output_dir / "workflow_output_closure.json"
    closure = json.loads(closure_path.read_text(encoding="utf-8"))
    for receipt in closure["artifact_receipts"]:
        if receipt["path"] in {"applied_decisions.json", "final_artifacts.json"}:
            _refresh_receipt(receipt, output_dir)
    _reseal_payload(closure)
    _write_json(closure_path, closure)

    from output_closure import (
        validate_final_artifact_index,
        validate_output_closure,
    )
    from replay_assurance import replay_assurance

    validate_final_artifact_index(output_dir)
    validate_output_closure(output_dir)
    with pytest.raises(ValueError, match="blocker_count is stale"):
        replay_assurance(output_dir)


def test_concordato_review_transaction_honest_save_seals_successor(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    predecessor = json.loads(
        (output_dir / "workflow_output_closure.json").read_text(encoding="utf-8")
    )

    result = _concordato_transaction_call(
        "save_concordato_plan_decisions",
        arguments,
    )

    closure = json.loads(
        (output_dir / "workflow_output_closure.json").read_text(encoding="utf-8")
    )
    assert result["ok"] is True
    assert closure["phase"] == "review_save_finalization"
    assert closure["previous_closure_content_sha256"] == predecessor["content_sha256"]
    assert not (output_dir / "applied_decisions.json").exists()


def test_concordato_review_transaction_rejects_preexisting_unexpected_file(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    rogue = output_dir / "unexpected.bin"
    rogue.write_bytes(b"not-a-workflow-output")
    rogue.chmod(0o640)
    before = _concordato_tree_image(output_dir)

    result = _concordato_transaction_call(
        "save_concordato_plan_decisions",
        arguments,
    )

    assert result["ok"] is False
    assert "assurance replay failed" in result["error"].lower()
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_review_transaction_rejects_preexisting_empty_directory(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    (output_dir / "unexpected-empty").mkdir()
    before = _concordato_tree_image(output_dir)

    result = _concordato_transaction_call(
        "save_concordato_plan_decisions",
        arguments,
    )

    assert result["ok"] is False
    assert "assurance replay failed" in result["error"].lower()
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_review_transaction_rejects_forged_successor_chain(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    before = _concordato_tree_image(output_dir)
    wrapper = _concordato_forged_successor_chain(tmp_path)

    result = _concordato_transaction_call(
        "save_concordato_plan_decisions",
        arguments,
        env={
            "PYTHON": str(wrapper),
            "REAL_PYTHON": sys.executable,
        },
    )

    assert result["ok"] is False
    assert result["error"] == "Concordato successor assurance chain is stale."
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize(
    ("tool_name", "needle"),
    [
        (
            "save_concordato_plan_decisions",
            "      const workingResult = saveDecisionPayloadWrites(workingArgs);\n",
        ),
        (
            "apply_concordato_plan_decisions",
            "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n",
        ),
    ],
)
def test_concordato_review_transaction_late_failure_restores_bytes_and_modes(
    tmp_path: Path,
    tool_name: str,
    needle: str,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    before = _concordato_tree_image(output_dir)
    faulted = _concordato_faulted_server(
        tmp_path,
        needle=needle,
        replacement=(
            needle
            + '      throw new Error("/private/client/concordato-late-failure");\n'
        ),
    )

    result = _concordato_transaction_call(
        tool_name,
        arguments,
        server_path=faulted,
    )

    assert result["ok"] is False
    assert result["error"] == (
        "Concordato review "
        + ("save" if tool_name.startswith("save") else "apply")
        + " transaction failed safely."
    )
    assert "/private/client" not in result["error"]
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_review_transaction_rejects_forged_save_response_contract(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    before = _concordato_tree_image(output_dir)
    needle = "      const workingResult = saveDecisionPayloadWrites(workingArgs);\n"
    faulted = _concordato_faulted_server(
        tmp_path,
        needle=needle,
        replacement=(needle + """
      Object.assign(workingResult, {
        validation_type: "forged_save",
        run_id: "forged-run",
        decision_count: 777,
        item_count: 888,
        status: "forged_status",
        ui_decisions_path: "/private/client/forged-ui.json",
        message: "forged message",
      });
"""),
    )

    result = _concordato_transaction_call(
        "save_concordato_plan_decisions",
        arguments,
        server_path=faulted,
    )

    assert result == {
        "ok": False,
        "error": "Concordato saved decisions did not close.",
    }
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_review_transaction_rejects_forged_apply_response_contract(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    before = _concordato_tree_image(output_dir)
    needle = "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n"
    faulted = _concordato_faulted_server(
        tmp_path,
        needle=needle,
        replacement=(needle + """
      Object.assign(workingResult, {
        validation_type: "forged_apply",
        run_id: "forged-run",
        decision_count: 101,
        item_count: 102,
        blocker_count: 103,
        revision_count: 104,
        target_update_count: 105,
        structured_update_count: 106,
        native_regeneration_count: 107,
        native_regenerated_count: 108,
        application_status: "final_ready",
        ui_decisions_path: "/private/client/forged-ui.json",
        applied_decisions_path: "/private/client/forged-applied.json",
        final_artifacts_path: "/private/client/forged-final.json",
        run_intake_path: "/private/client/forged-intake.json",
        message: "forged message",
      });
"""),
    )

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
        server_path=faulted,
    )

    assert result == {
        "ok": False,
        "error": "Concordato response did not close.",
    }
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize("attack_kind", ["symlink", "hardlink", "fifo"])
def test_concordato_review_transaction_rejects_working_tree_poison(
    tmp_path: Path,
    attack_kind: str,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    before = _concordato_tree_image(output_dir)
    external = tmp_path / "external-target.bin"
    external.write_bytes(b"EXTERNAL-UNCHANGED")
    external.chmod(0o600)
    needle = "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n"
    injected = """
      const poisonPath = path.join(workingOutputDir, "ui_decisions.json");
      fs.unlinkSync(poisonPath);
      if (process.env.REVIEW_TX_ATTACK_KIND === "symlink") {
        fs.symlinkSync(process.env.REVIEW_TX_EXTERNAL, poisonPath);
      } else if (process.env.REVIEW_TX_ATTACK_KIND === "hardlink") {
        fs.linkSync(process.env.REVIEW_TX_EXTERNAL, poisonPath);
      } else {
        require("node:child_process").spawnSync("/usr/bin/mkfifo", [poisonPath]);
      }
"""
    faulted = _concordato_faulted_server(
        tmp_path,
        needle=needle,
        replacement=needle + injected,
    )

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
        server_path=faulted,
        env={
            "REVIEW_TX_ATTACK_KIND": attack_kind,
            "REVIEW_TX_EXTERNAL": str(external),
        },
    )

    assert result["ok"] is False
    assert "/" not in result["error"]
    assert external.read_bytes() == b"EXTERNAL-UNCHANGED"
    assert external.stat().st_mode & 0o7777 == 0o600
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_review_transaction_rejects_transaction_root_relocation_without_moving_canonical(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    before = _concordato_tree_image(output_dir)
    canonical_inode = output_dir.stat().st_ino
    needle = "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n"
    injected = """
      const transactionRoot = path.dirname(workingOutputDir);
      fs.renameSync(transactionRoot, `${transactionRoot}-moved`);
      throw new Error("/private/client/transaction-root-relocation");
"""
    faulted = _concordato_faulted_server(
        tmp_path,
        needle=needle,
        replacement=needle + injected,
    )

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
        server_path=faulted,
    )

    assert result["ok"] is False
    assert result["error"] == ("Concordato review apply transaction failed safely.")
    assert output_dir.stat().st_ino == canonical_inode
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))
    assert not list(tmp_path.glob(".generated-review-commit-*"))
    assert not list(tmp_path.glob(".generated-review-recovery-*"))


def test_concordato_review_transaction_restores_after_commit_deletion(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    before = _concordato_tree_image(output_dir)
    needle = "    committed = true;\n    const committedImage ="
    faulted = _concordato_faulted_server(
        tmp_path,
        needle=needle,
        replacement=(
            "    committed = true;\n"
            "    generatedReviewRemoveExactPath(resolvedOutputDir);\n"
            "    const committedImage ="
        ),
    )

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
        server_path=faulted,
    )

    assert result["ok"] is False
    assert "/" not in result["error"]
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_review_transaction_enforces_size_bound_before_mutation(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    closure_path = output_dir / "workflow_output_closure.json"
    oversized = output_dir / "oversized.bin"
    with oversized.open("wb") as stream:
        stream.truncate(128 * 1024 * 1024 + 1)
    root_inode = output_dir.stat().st_ino
    closure_inode = closure_path.stat().st_ino
    closure_bytes = closure_path.read_bytes()

    result = _concordato_transaction_call(
        "save_concordato_plan_decisions",
        arguments,
    )

    assert result["ok"] is False
    assert output_dir.stat().st_ino == root_inode
    assert closure_path.stat().st_ino == closure_inode
    assert closure_path.read_bytes() == closure_bytes
    assert oversized.stat().st_size == 128 * 1024 * 1024 + 1
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_child_receives_only_working_output_paths(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    arguments = _concordato_memo_arguments(arguments)
    child_log = tmp_path / "child-arguments.log"
    wrapper = tmp_path / "recording-python.sh"
    wrapper.write_text(
        "#!/bin/sh\n"
        'printf "%s\\n" "$@" >> "$REVIEW_TX_CHILD_LOG"\n'
        'exec "$REAL_PYTHON" "$@"\n',
        encoding="utf-8",
    )
    wrapper.chmod(0o700)

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
        env={
            "PYTHON": str(wrapper),
            "REAL_PYTHON": sys.executable,
            "REVIEW_TX_CHILD_LOG": str(child_log),
        },
    )

    child_arguments = child_log.read_text(encoding="utf-8")
    persisted_trace = (output_dir / "run_intake.json").read_text(encoding="utf-8")
    assert result["ok"] is True
    assert str(output_dir) not in child_arguments
    assert ".generated-review-transaction-" in child_arguments
    assert "/working" in child_arguments
    assert ".generated-review-transaction-" not in persisted_trace
    assert str(child_log) not in persisted_trace
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_child_start_failure_is_fixed_and_rolls_back(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    before = _concordato_tree_image(output_dir)
    non_executable = tmp_path / "python-directory"
    non_executable.mkdir()

    result = _concordato_transaction_call(
        "save_concordato_plan_decisions",
        arguments,
        env={"PYTHON": str(non_executable)},
    )

    assert result["ok"] is False
    assert result["error"] == "Concordato assurance replay could not start."
    assert str(non_executable) not in result["error"]
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize("tamper_kind", ["artifact_bytes", "envelope_digest"])
def test_concordato_parent_replay_rejects_stale_receipts_despite_exact_child_echo(
    tmp_path: Path,
    tamper_kind: str,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    review_payload = arguments["review_payload"]
    assurance = review_payload["assurance"]
    child_result = {
        "ok": True,
        "run_id": review_payload["run_id"],
        "review_payload_content_sha256": review_payload["content_sha256"],
        "assurance_envelope_content_sha256": assurance["envelope_content_sha256"],
        "report_ready": assurance["gate_register"]["report_ready"],
    }
    if tamper_kind == "artifact_bytes":
        artifact_path = output_dir / "exact_amount_matches.csv"
        artifact_path.write_bytes(
            artifact_path.read_bytes() + b"\nFORGED-RECEIPTED-BYTES\n"
        )
    else:
        envelope_path = output_dir / "assurance_envelope.json"
        envelope = json.loads(envelope_path.read_text(encoding="utf-8"))
        envelope["content_sha256"] = "0" * 64
        envelope_path.write_text(
            json.dumps(envelope, indent=2) + "\n",
            encoding="utf-8",
        )
    before = _concordato_tree_image(output_dir)
    child = tmp_path / f"exact-echo-replay-{tamper_kind}.py"
    child.write_text(
        f"#!{sys.executable}\n"
        "import json\n"
        f"print(json.dumps({child_result!r}))\n",
        encoding="utf-8",
    )
    child.chmod(0o700)

    result = _concordato_transaction_call(
        "save_concordato_plan_decisions",
        arguments,
        env={"PYTHON": str(child)},
    )

    assert result == {
        "ok": False,
        "error": "Concordato assurance replay returned an invalid result.",
    }
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize(
    ("mode", "expected_error"),
    [
        ("nonzero", "Concordato review application failed."),
        ("malformed", "Concordato review application returned an invalid result."),
    ],
)
def test_concordato_child_failure_after_persistence_is_fixed_and_rolls_back(
    tmp_path: Path,
    mode: str,
    expected_error: str,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    arguments = _concordato_memo_arguments(arguments)
    before = _concordato_tree_image(output_dir)
    child = _concordato_phase_child(tmp_path, mode=mode)

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
        env={"PYTHON": str(child), "REAL_PYTHON": sys.executable},
    )

    assert result["ok"] is False
    assert result["error"] == expected_error
    assert "/private/client" not in result["error"]
    assert "Traceback" not in result["error"]
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_concordato_child_cannot_forge_professional_readiness(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    arguments = _concordato_memo_arguments(arguments)
    before = _concordato_tree_image(output_dir)
    child = tmp_path / "forged-concordato-child.py"
    child.write_text(
        f"#!{sys.executable}\n"
        "import json\n"
        "import pathlib\n"
        "import sys\n"
        "script_index = next(\n"
        "    index for index, value in enumerate(sys.argv[1:], 1)\n"
        "    if value.endswith('.py')\n"
        ")\n"
        "if pathlib.Path(sys.argv[script_index]).name == 'replay_assurance.py':\n"
        f"    print(json.dumps({repr({'ok': True, 'run_id': arguments['review_payload']['run_id'], 'review_payload_content_sha256': arguments['review_payload']['content_sha256'], 'assurance_envelope_content_sha256': arguments['review_payload']['assurance']['envelope_content_sha256'], 'report_ready': False})}))\n"
        "    raise SystemExit(0)\n"
        "applied_path = pathlib.Path(sys.argv[sys.argv.index('--applied-decisions') + 1])\n"
        "final_path = pathlib.Path(sys.argv[sys.argv.index('--final-artifacts') + 1])\n"
        "applied = json.loads(applied_path.read_text())\n"
        "final = json.loads(final_path.read_text())\n"
        "applied['application_status'] = 'final_ready'\n"
        "final['status'] = 'final_ready'\n"
        "final['review_status'] = 'final_ready'\n"
        "final['final_ready'] = True\n"
        "final['review_application']['application_status'] = 'final_ready'\n"
        "applied_path.write_text(json.dumps(applied) + '\\n')\n"
        "final_path.write_text(json.dumps(final) + '\\n')\n"
        "print(json.dumps({'ok': True, 'updated_effect_count': 1, "
        "'native_regenerated_paths': [], 'backup_paths': [], "
        "'applied_decisions': applied, 'final_artifacts': final}))\n",
        encoding="utf-8",
    )
    child.chmod(0o700)

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
        env={"PYTHON": str(child)},
    )

    assert result["ok"] is False
    assert (
        result["error"] == "Concordato review application returned an invalid result."
    )
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize(
    "tamper_mode",
    [
        "rogue_effect_paths",
        "forged_contract",
        "run_intake",
        "review_handoff",
        "docx_content",
        "docx_numeric",
    ],
)
def test_concordato_child_cannot_self_authorize_or_tamper_parent_receipts(
    tmp_path: Path,
    tamper_mode: str,
) -> None:
    output_dir, arguments = _concordato_transaction_case(tmp_path)
    arguments = _concordato_memo_arguments(arguments)
    before = _concordato_tree_image(output_dir)
    child = _concordato_post_child_tamper(tmp_path)

    result = _concordato_transaction_call(
        "apply_concordato_plan_decisions",
        arguments,
        env={
            "PYTHON": str(child),
            "REAL_PYTHON": sys.executable,
            "CONCORDATO_TAMPER_MODE": tamper_mode,
        },
    )

    assert result["ok"] is False
    assert "/" not in result["error"]
    assert "\\" not in result["error"]
    assert "private" not in result["error"].lower()
    assert _concordato_tree_image(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))
