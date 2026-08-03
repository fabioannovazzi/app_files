from __future__ import annotations

import importlib.util
import json
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from scripts.validate_plugin_review_contract import validate_contract

ROOT = Path(__file__).resolve().parents[2]
PLUGIN_ROOT = ROOT / "plugins" / "deep-research-validator"
SCRIPTS_DIR = PLUGIN_ROOT / "scripts"
MCP_SERVER_PATH = PLUGIN_ROOT / "mcp" / "server.cjs"
VERA_PRODUCT_PAGE_LINK = "../vera/index.html"


def _running_customer_output(
    tmp_path: Path,
) -> tuple[Path, Path, dict[str, Any]]:
    ledger_path = ROOT / "plugins" / "studio-archive" / "scripts" / "client_ledger.py"
    module_name = "test_deep_research_customer_ledger"
    ledger = sys.modules.get(module_name)
    if ledger is None:
        spec = importlib.util.spec_from_file_location(module_name, ledger_path)
        assert spec is not None and spec.loader is not None
        ledger = importlib.util.module_from_spec(spec)
        sys.modules[module_name] = ledger
        spec.loader.exec_module(ledger)
    client_root = tmp_path / "Managed Customer"
    client_root.mkdir()
    client_id = "client_111111111111111111111111"
    ledger.create_client_manifest(client_root, client_id)
    engagement = ledger.create_engagement(client_root, client_id, "Test engagement")
    managed_claim = "The Italian VAT rule applies to the transaction."
    managed_inputs = {
        "document_inventory.json": {
            "source_name": "deep_research.md",
            "character_count": 80,
            "word_count": 12,
            "urls": ["https://example.com"],
        },
        "source_inventory.json": {
            "sources": [
                {
                    "kind": "url",
                    "source_id": "source-001",
                    "url": "https://example.com",
                    "status": "available",
                    "excerpt": managed_claim,
                }
            ]
        },
        "claims_review.json": _claims_review(
            [
                _claim_review(
                    managed_claim,
                    cited_passage="The Italian VAT rule applies",
                )
            ]
        ),
        "answer_contract.json": _answer_contract(),
        "local_source.txt": {
            "text": "The Italian VAT rule applies to the transaction."
        },
    }
    input_ids: list[str] = []
    for file_name, payload in managed_inputs.items():
        source = tmp_path / file_name
        source.write_text(json.dumps(payload) + "\n", encoding="utf-8")
        imported = ledger.import_document(
            client_root,
            client_id,
            engagement["engagement_id"],
            source,
            "source",
        )
        input_ids.append(imported["receipt"]["input_id"])
    prepared = ledger.prepare_run(
        client_root,
        client_id,
        engagement["engagement_id"],
        "deep-research-validator",
        "test-version",
        input_ids=input_ids,
    )
    running = ledger.start_run(
        client_root,
        engagement["engagement_id"],
        prepared["run"]["run_id"],
    )
    context = running["context"]
    return (
        Path(running["output_dir"]),
        Path(context["context_path"]),
        context,
    )


def _file_snapshot(root: Path) -> dict[str, bytes]:
    """Return the exact relative file content below a test root."""

    return {
        path.relative_to(root).as_posix(): path.read_bytes()
        for path in sorted(root.rglob("*"))
        if path.is_file()
    }


def _answer_contract(
    *,
    generation_route: str = "chatgpt_deep_research",
    document_type: str = "legal research report",
    validation_scope: str = "all_material_claims",
) -> dict[str, str]:
    return {
        "schema_version": "1.0",
        "question_domain": "legal",
        "generation_route": generation_route,
        "document_type": document_type,
        "purpose": "Answer the supplied legal question",
        "audience": "Professional reviewer",
        "output_language": "English",
        "jurisdiction_status": "confirmed",
        "jurisdiction": "Italian law",
        "evidence_display": "inline_citations",
        "validation_profile": "source_identity_support_reasoning_and_judgment",
        "validation_scope": validation_scope,
        "correction_policy": "correct_when_supported",
        "judgment_policy": "flag_for_professional_review",
    }


def _no_issue() -> list[dict[str, str]]:
    return [
        {
            "type": "none",
            "explanation": "No defect identified.",
            "treatment_action": "none",
            "treatment_status": "not_needed",
            "treatment_explanation": "No treatment is required.",
        }
    ]


def _claim_review(
    claim_text: str,
    *,
    claim_index: int = 1,
    source_ref: str = "source-001",
    cited_passage: str = "",
    support_status: str = "supported",
    support_analysis: str = "The source semantically supports the claim.",
    reasoning_status: str = "sound",
    reasoning_analysis: str = "The conclusion follows from the supported premise.",
    judgment_status: str = "not_judgment_dependent",
    judgment_analysis: str = "No additional professional judgment is required.",
    issues: list[dict[str, str]] | None = None,
    disposition_status: str = "retain",
    revised_claim: str = "",
    reviewer_action: str = "accept",
    proposed_fix: str = "",
) -> dict[str, object]:
    return {
        "claim_index": claim_index,
        "claim_text": claim_text,
        "claim_location": f"Section 1, claim {claim_index}",
        "materiality": "material",
        "source_checks": [
            {
                "source_ref": source_ref,
                "identity_status": "matches_cited_source",
                "identity_analysis": "The source reference identifies the cited authority.",
                "cited_passage": cited_passage,
            }
        ],
        "support": {"status": support_status, "analysis": support_analysis},
        "reasoning": {
            "status": reasoning_status,
            "analysis": reasoning_analysis,
            "supported_premises": ["The cited source establishes the premise."],
            "missing_premises": (
                [] if reasoning_status == "sound" else ["Application fact"]
            ),
        },
        "professional_judgment": {
            "status": judgment_status,
            "analysis": judgment_analysis,
            "factors": (
                []
                if judgment_status == "not_judgment_dependent"
                else ["Fact-sensitive application"]
            ),
            "alternative_interpretations": (
                []
                if judgment_status == "not_judgment_dependent"
                else ["A different application may be reasonable."]
            ),
        },
        "issues": issues or _no_issue(),
        "disposition": {
            "status": disposition_status,
            "analysis": "The claim disposition follows the recorded assessments.",
            "revised_claim": revised_claim,
        },
        "reviewer_action": reviewer_action,
        "proposed_fix": proposed_fix,
    }


def _claims_review(
    claims: list[dict[str, object]],
    *,
    language: str = "en",
    validated_document: str = "Validated text.",
    coverage_scope: str = "all_material_claims",
    document_revision_status: str = "not_required",
    overall_outcome: str | None = None,
) -> dict[str, object]:
    resolved_outcome = (
        overall_outcome
        or {
            "not_required": "no_material_defect_identified",
            "completed": "corrected",
            "required": "correction_required",
            "blocked": "not_reliable",
            "professional_review_required": "professional_review_required",
        }[document_revision_status]
    )
    return {
        "schema_version": "2.0",
        "language": language,
        "validation_objective": "question_to_validated_answer",
        "coverage_review": {
            "selection_method": "model_led_materiality_review",
            "scope": coverage_scope,
            "reviewed_sections": ["Full answer"],
            "omitted_sections": [],
            "limitations": (
                []
                if coverage_scope == "all_material_claims"
                else ["Review scope was limited."]
            ),
            "analysis": "The full answer was read and material claims were selected semantically.",
            "reviewer_action": (
                "accept" if coverage_scope == "all_material_claims" else "mark_unclear"
            ),
        },
        "contract_review": {
            "question_answered": {
                "status": "conforms",
                "analysis": "The question is answered.",
            },
            "document_type": {
                "status": "conforms",
                "analysis": "The document type conforms.",
            },
            "audience": {
                "status": "conforms",
                "analysis": "The answer suits the audience.",
            },
            "evidence_display": {
                "status": "conforms",
                "analysis": "The evidence display conforms.",
            },
            "issues": _no_issue(),
            "reviewer_action": "accept",
        },
        "claims": claims,
        "overall_assessment": {
            "outcome": resolved_outcome,
            "analysis": f"The recorded overall outcome is {resolved_outcome}.",
            "residual_uncertainties": [],
            "professional_review_items": [],
        },
        "document_revision": {
            "status": document_revision_status,
            "summary": (
                "No revision is required."
                if document_revision_status == "not_required"
                else "Revision remains required."
            ),
            "unresolved_changes": (
                []
                if document_revision_status == "not_required"
                else ["Regenerate the answer."]
            ),
        },
        "validated_document": validated_document,
    }


def load_script(module_name: str, script_name: str):
    script_path = SCRIPTS_DIR / script_name
    sys.path.insert(0, str(SCRIPTS_DIR))
    try:
        spec = importlib.util.spec_from_file_location(module_name, script_path)
        assert spec and spec.loader
        module = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = module
        spec.loader.exec_module(module)
        return module
    finally:
        sys.path.remove(str(SCRIPTS_DIR))


def _call_mcp_server(messages: list[dict[str, object]]) -> list[dict[str, object]]:
    node = shutil.which("node")
    if node is None:
        pytest.skip(
            "Node.js is required to exercise the Deep Research Validator MCP server."
        )
    completed = subprocess.run(
        [node, str(MCP_SERVER_PATH), "--stdio"],
        input="\n".join(json.dumps(message) for message in messages) + "\n",
        capture_output=True,
        text=True,
        check=True,
        timeout=10,
    )
    return [json.loads(line) for line in completed.stdout.splitlines() if line.strip()]


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def _docx_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_inspect_document_extracts_references_and_claim_candidates(
    tmp_path: Path,
) -> None:
    inspect_mod = load_script(
        "deep_research_validator_inspect_document",
        "inspect_document.py",
    )
    document = tmp_path / "report.md"
    document.write_text(
        "\n".join(
            [
                "# VAT report",
                "The Italian VAT rule applies to the transaction [1].",
                "The conclusion is supported by [Agenzia](https://example.com/source).",
                "[^1]: https://example.com/source",
            ]
        ),
        encoding="utf-8",
    )

    paths = inspect_mod.write_inspection(document, tmp_path / "out")
    inventory = json.loads(paths["document_inventory"].read_text(encoding="utf-8"))

    assert inventory["source_name"] == "report.md"
    assert inventory["headings"] == ["VAT report"]
    assert inventory["urls"] == ["https://example.com/source"]
    assert inventory["markdown_links"][0]["label"] == "Agenzia"
    assert inventory["footnotes"][0]["id"] == "1"
    assert inventory["mechanical_claim_candidates"]


def test_inspect_sources_can_skip_network_fetch(tmp_path: Path) -> None:
    inspect_sources = load_script(
        "deep_research_validator_inspect_sources",
        "inspect_sources.py",
    )
    inventory = tmp_path / "document_inventory.json"
    inventory.write_text(
        json.dumps(
            {"urls": ["https://example.com/a"], "footnotes": [], "markdown_links": []}
        ),
        encoding="utf-8",
    )

    paths = inspect_sources.write_source_inventory(
        inventory,
        tmp_path / "out",
        fetch_urls=False,
    )
    payload = json.loads(paths["source_inventory"].read_text(encoding="utf-8"))

    assert payload["url_count"] == 1
    assert payload["sources"][0]["status"] == "listed_not_fetched"


def test_inspect_sources_blocks_loopback_without_opening_it(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    inspect_mod = load_script(
        "deep_research_validator_inspect_sources_loopback",
        "inspect_sources.py",
    )
    inventory = tmp_path / "document_inventory.json"
    inventory.write_text(
        json.dumps(
            {
                "urls": ["http://127.0.0.1/private"],
                "footnotes": [],
                "markdown_links": [],
            }
        ),
        encoding="utf-8",
    )

    def fail_if_opened(*_args, **_kwargs):
        raise AssertionError("A loopback URL must not reach the HTTP opener.")

    monkeypatch.setattr(inspect_mod.urllib.request, "build_opener", fail_if_opened)

    payload = inspect_mod.inspect_sources(inventory)

    assert payload["sources"][0]["status"] == "blocked_non_public_destination"


def test_inspect_sources_revalidates_redirect_destination(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    inspect_mod = load_script(
        "deep_research_validator_inspect_sources_redirect",
        "inspect_sources.py",
    )
    inventory = tmp_path / "document_inventory.json"
    inventory.write_text(
        json.dumps(
            {
                "urls": ["https://public.example/source"],
                "footnotes": [],
                "markdown_links": [],
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(
        inspect_mod.socket,
        "getaddrinfo",
        lambda *_args, **_kwargs: [
            (
                inspect_mod.socket.AF_INET,
                inspect_mod.socket.SOCK_STREAM,
                6,
                "",
                ("93.184.216.34", 443),
            )
        ],
    )

    class RedirectingOpener:
        def __init__(self, handler) -> None:
            self.handler = handler

        def open(self, request, *, timeout: float):
            del timeout
            return self.handler.redirect_request(
                request,
                None,
                302,
                "Found",
                {},
                "http://127.0.0.1/admin",
            )

    monkeypatch.setattr(
        inspect_mod.urllib.request,
        "build_opener",
        lambda handler: RedirectingOpener(handler),
    )

    payload = inspect_mod.inspect_sources(inventory)

    assert payload["sources"][0]["status"] == "blocked_non_public_destination"


def test_inspect_sources_allows_public_https_destination(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    inspect_mod = load_script(
        "deep_research_validator_inspect_sources_public",
        "inspect_sources.py",
    )
    inventory = tmp_path / "document_inventory.json"
    inventory.write_text(
        json.dumps(
            {
                "urls": ["https://public.example/source"],
                "footnotes": [],
                "markdown_links": [],
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(
        inspect_mod.socket,
        "getaddrinfo",
        lambda *_args, **_kwargs: [
            (
                inspect_mod.socket.AF_INET,
                inspect_mod.socket.SOCK_STREAM,
                6,
                "",
                ("93.184.216.34", 443),
            )
        ],
    )

    class PublicResponse:
        status = 200

        def __enter__(self):
            return self

        def __exit__(self, *_args) -> None:
            return None

        def read(self, _limit: int) -> bytes:
            return (
                b"<html><body>"
                + b"Official public source content. " * 12
                + b"</body></html>"
            )

    class PublicOpener:
        def open(self, _request, *, timeout: float) -> PublicResponse:
            del timeout
            return PublicResponse()

    monkeypatch.setattr(
        inspect_mod.urllib.request,
        "build_opener",
        lambda _handler: PublicOpener(),
    )

    payload = inspect_mod.inspect_sources(inventory)

    assert payload["sources"][0]["status"] == "available"


def test_package_validation_writes_audit_and_package(tmp_path: Path) -> None:
    package_mod = load_script(
        "deep_research_validator_package_validation",
        "package_validation.py",
    )
    document_inventory = tmp_path / "document_inventory.json"
    source_inventory = tmp_path / "source_inventory.json"
    claims_review = tmp_path / "claims_review_draft.json"
    answer_contract = tmp_path / "answer_contract.json"
    document_inventory.write_text(
        json.dumps(
            {"character_count": 80, "word_count": 12, "urls": ["https://example.com"]}
        ),
        encoding="utf-8",
    )
    source_inventory.write_text(
        json.dumps(
            {
                "sources": [
                    {
                        "kind": "url",
                        "source_id": "source-001",
                        "url": "https://example.com",
                        "status": "available",
                        "excerpt": "The Italian VAT rule applies to the transaction.",
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    claims_review.write_text(
        json.dumps(
            _claims_review(
                [
                    _claim_review(
                        "The Italian VAT rule applies to the transaction.",
                        cited_passage="The Italian VAT rule applies",
                    )
                ]
            )
        ),
        encoding="utf-8",
    )
    answer_contract.write_text(
        json.dumps(_answer_contract()),
        encoding="utf-8",
    )
    client_run_id = "run_" + "c" * 24

    paths = package_mod.write_validation_package(
        document_inventory,
        source_inventory,
        claims_review,
        tmp_path / "out",
        answer_contract_path=answer_contract,
        client_run_id=client_run_id,
    )
    audit = json.loads(paths["validation_audit"].read_text(encoding="utf-8"))
    run_intake = json.loads(
        (tmp_path / "out" / "run_intake.json").read_text(encoding="utf-8")
    )
    review_payload = json.loads(
        (tmp_path / "out" / "review_payload.json").read_text(encoding="utf-8")
    )
    ui_decisions = json.loads(
        (tmp_path / "out" / "ui_decisions.json").read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        (tmp_path / "out" / "final_artifacts.json").read_text(encoding="utf-8")
    )

    assert audit["status"] == "record_complete"
    assert run_intake["run_id"] == client_run_id
    assert audit["delivery_readiness"] == "reviewed_answer_ready"
    assert audit["review_session"]["run_id"] == run_intake["run_id"]
    assert (
        audit["claim_observations"][0]["source_observations"][0][
            "exact_passage_presence"
        ]
        == "present"
    )
    assert (
        paths["validated_document"].read_text(encoding="utf-8").strip()
        == "Validated text."
    )
    assert "Answer Validation Record" in paths["validation_package"].read_text(
        encoding="utf-8"
    )
    assert review_payload["plugin"] == "deep-research-validator"
    assert review_payload["run_id"] == run_intake["run_id"]
    assert review_payload["review_type"] == "answer_validation_review"
    assert review_payload["item_count"] == len(review_payload["items"])
    item_types = {item["item_type"] for item in review_payload["items"]}
    assert {"supported_claim", "validation_artifact"} <= item_types
    claim_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "supported_claim"
    )
    claim_evidence = claim_item["evidence"][0]
    assert claim_evidence["kind"] == "answer_validation_assessment"
    assert claim_evidence["claim_text"] == (
        "The Italian VAT rule applies to the transaction."
    )
    assert claim_evidence["source_checks"][0]["cited_passage"] == (
        "The Italian VAT rule applies"
    )
    assert claim_evidence["support"]["status"] == "supported"
    assert claim_evidence["reasoning"]["status"] == "sound"
    assert claim_evidence["professional_judgment"]["status"] == (
        "not_judgment_dependent"
    )
    assert claim_evidence["issues"][0]["type"] == "none"
    assert claim_item["data"]["target_artifact"] == "claims_review.json"
    assert claim_item["data"]["target_records_key"] == "claims"
    assert claim_item["data"]["target_id_field"] == "claim_index"
    assert claim_item["data"]["target_record_id"] == "1"
    assert claim_item["data"]["target_field"] == "proposed_fix"
    assert review_payload["summary"]["record_integrity_status"] == ("record_complete")
    assert ui_decisions["status"] == "pending_review"
    assert final_artifacts["status"] == "written_pending_review"
    output_records = {output["path"]: output for output in final_artifacts["outputs"]}
    assert (
        output_records["validation_audit.json"]["size_bytes"]
        == paths["validation_audit"].stat().st_size
    )
    for output in output_records.values():
        if "size_bytes" in output:
            assert (
                output["size_bytes"]
                == (tmp_path / "out" / output["path"]).stat().st_size
            )
    handoff_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "review_handoff.md"
    )
    handoff_text = (tmp_path / "out" / "review_handoff.md").read_text(encoding="utf-8")
    assert handoff_output["required_text"] == [
        "Review Handoff",
        "review_payload.json",
        "ui_decisions.json",
        "applied_decisions.json",
        "final_artifacts.json",
    ]
    assert "render_deep_research_review" in handoff_text
    assert "apply_deep_research_decisions" in handoff_text
    package_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "validation_package.md"
    )
    assert package_output["required_text"] == [
        "# Answer Validation Record",
        "## Assurance Boundary",
        "## Answer Contract",
        "## Answer-Contract Review",
        "## Review Coverage",
        "## Document Inventory",
        "## Claim Assessments",
    ]
    contract_report = validate_contract(
        tmp_path / "out",
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()


def test_managed_package_writer_persists_only_run_relative_paths(
    tmp_path: Path,
) -> None:
    output_dir, context_path, context = _running_customer_output(tmp_path)
    sources_mod = load_script(
        "deep_research_validator_inspect_sources_managed_paths",
        "inspect_sources.py",
    )
    package_mod = load_script(
        "deep_research_validator_package_managed_paths",
        "package_validation.py",
    )
    inputs_by_name = {
        Path(binding["path"]).name: Path(binding["path"])
        for binding in context["input_bindings"]
    }
    generated_inventory = sources_mod.write_source_inventory(
        inputs_by_name["document_inventory.json"],
        output_dir,
        source_files=[inputs_by_name["local_source.txt"]],
        fetch_urls=False,
        run_root=Path(context["run_root"]),
    )["source_inventory"]

    paths = package_mod.write_validation_package(
        inputs_by_name["document_inventory.json"],
        generated_inventory,
        inputs_by_name["claims_review.json"],
        output_dir,
        answer_contract_path=inputs_by_name["answer_contract.json"],
        client_engagement=context,
        client_run_id=str(context["run_id"]),
    )
    run_intake = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )
    audit = json.loads(paths["validation_audit"].read_text(encoding="utf-8"))
    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    client_root = str(context["studio_client_folder"]["client_root"])

    assert run_intake["path_reference"] == "run_root_relative"
    assert run_intake["output_dir"] == "outputs"
    assert all(not Path(value).is_absolute() for value in run_intake["input_paths"])
    assert all(
        not Path(value).is_absolute()
        for key, value in audit["review_session"].items()
        if key.endswith("_path")
    )
    assert review_payload["path_reference"] == "run_root_relative"
    assert client_root not in json.dumps(
        {
            "run_intake": run_intake,
            "review_session": audit["review_session"],
            "source_artifacts": review_payload["source_artifacts"],
        }
    )
    assert all(
        client_root not in path.read_text(encoding="utf-8")
        for path in output_dir.glob("*.json")
    )
    assert client_root not in (output_dir / "validation_package.md").read_text(
        encoding="utf-8"
    )
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    claim_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "supported_claim"
    )
    decisions = [
        {
            "item_id": item["id"],
            "action": "edit" if item["id"] == claim_item["id"] else "accept",
            **(
                {"edit_value": "Limit the conclusion to the cited VAT rule."}
                if item["id"] == claim_item["id"]
                else {}
            ),
        }
        for item in review_payload["items"]
    ]
    old_client_root = Path(context["studio_client_folder"]["client_root"])
    renamed_client_root = tmp_path / "Renamed Customer"
    context_relative = context_path.relative_to(old_client_root)
    output_relative = output_dir.relative_to(old_client_root)
    old_client_root.rename(renamed_client_root)
    context_path = renamed_client_root / context_relative
    output_dir = renamed_client_root / output_relative

    responses = {
        response["id"]: response
        for response in _call_mcp_server(
            [
                {
                    "jsonrpc": "2.0",
                    "id": 1,
                    "method": "tools/call",
                    "params": {
                        "name": "save_deep_research_decisions",
                        "arguments": {
                            "run_intake": run_intake,
                            "client_engagement": context_path.as_posix(),
                            "review_payload": review_payload,
                            "decisions": decisions,
                        },
                    },
                },
                {
                    "jsonrpc": "2.0",
                    "id": 2,
                    "method": "tools/call",
                    "params": {
                        "name": "apply_deep_research_decisions",
                        "arguments": {
                            "run_intake": run_intake,
                            "client_engagement": context_path.as_posix(),
                            "review_payload": review_payload,
                            "final_artifacts": final_artifacts,
                            "decisions": decisions,
                        },
                    },
                },
            ]
        )
    }

    assert responses[1]["result"]["isError"] is False
    applied = responses[2]["result"]
    assert applied["isError"] is False, applied
    assert applied["structuredContent"]["structured_update_count"] == 1
    assert applied["structuredContent"]["run_intake_path"] == str(
        output_dir / "run_intake.json"
    )
    updated_claims = json.loads(
        (output_dir / "claims_review.json").read_text(encoding="utf-8")
    )
    assert updated_claims["claims"][0]["proposed_fix"] == (
        "Limit the conclusion to the cited VAT rule."
    )
    assert not old_client_root.exists()


def test_managed_package_writer_rejects_input_outside_run_without_writes(
    tmp_path: Path,
) -> None:
    output_dir, _context_path, context = _running_customer_output(tmp_path)
    package_mod = load_script(
        "deep_research_validator_package_managed_input_escape",
        "package_validation.py",
    )
    inputs_by_name = {
        Path(binding["path"]).name: Path(binding["path"])
        for binding in context["input_bindings"]
    }
    outside_input = tmp_path / "outside-answer-contract.json"
    outside_input.write_text(
        json.dumps(_answer_contract()) + "\n",
        encoding="utf-8",
    )
    before = _file_snapshot(tmp_path)

    with pytest.raises(ValueError, match="outside the current run"):
        package_mod.write_validation_package(
            inputs_by_name["document_inventory.json"],
            inputs_by_name["source_inventory.json"],
            inputs_by_name["claims_review.json"],
            output_dir,
            answer_contract_path=outside_input,
            client_engagement=context,
            client_run_id=str(context["run_id"]),
        )

    assert _file_snapshot(tmp_path) == before


def test_managed_source_inspection_rejects_unbound_file_without_writes(
    tmp_path: Path,
) -> None:
    output_dir, _context_path, context = _running_customer_output(tmp_path)
    sources_mod = load_script(
        "deep_research_validator_inspect_sources_unbound_file",
        "inspect_sources.py",
    )
    inputs_by_name = {
        Path(binding["path"]).name: Path(binding["path"])
        for binding in context["input_bindings"]
    }
    outside = tmp_path / "not-imported.txt"
    outside.write_text("Unbound source evidence.\n", encoding="utf-8")
    before = _file_snapshot(tmp_path)

    with pytest.raises(ValueError, match="outside the current customer run"):
        sources_mod.write_source_inventory(
            inputs_by_name["document_inventory.json"],
            output_dir,
            source_files=[outside],
            fetch_urls=False,
            run_root=Path(context["run_root"]),
        )

    assert _file_snapshot(tmp_path) == before


def test_package_validation_localizes_spanish_review_artifacts(tmp_path: Path) -> None:
    package_mod = load_script(
        "deep_research_validator_package_validation_es",
        "package_validation.py",
    )
    document_inventory = tmp_path / "document_inventory.json"
    source_inventory = tmp_path / "source_inventory.json"
    claims_review = tmp_path / "claims_review_draft.json"
    answer_contract = tmp_path / "answer_contract.json"
    document_inventory.write_text(
        json.dumps(
            {
                "source_name": "informe.md",
                "character_count": 90,
                "word_count": 14,
                "urls": ["https://example.com/fuente"],
            }
        ),
        encoding="utf-8",
    )
    source_inventory.write_text(
        json.dumps(
            {
                "sources": [
                    {
                        "kind": "url",
                        "source_id": "source-001",
                        "url": "https://example.com/fuente",
                        "status": "available",
                        "excerpt": "La norma del IVA se aplica a la operación.",
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    claims_review.write_text(
        json.dumps(
            _claims_review(
                [
                    _claim_review(
                        "La norma del IVA se aplica a la operación.",
                        cited_passage="La norma del IVA se aplica",
                    )
                ],
                language="es-ES",
                validated_document="Texto validado.",
            ),
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    answer_contract.write_text(
        json.dumps(_answer_contract()),
        encoding="utf-8",
    )

    paths = package_mod.write_validation_package(
        document_inventory,
        source_inventory,
        claims_review,
        tmp_path / "out",
        answer_contract_path=answer_contract,
    )

    package_text = paths["validation_package"].read_text(encoding="utf-8")
    run_intake = json.loads(
        (tmp_path / "out" / "run_intake.json").read_text(encoding="utf-8")
    )
    review_payload = json.loads(
        (tmp_path / "out" / "review_payload.json").read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        (tmp_path / "out" / "final_artifacts.json").read_text(encoding="utf-8")
    )
    handoff_text = (tmp_path / "out" / "review_handoff.md").read_text(encoding="utf-8")
    claim_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "supported_claim"
    )
    package_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "validation_package.md"
    )

    assert "# Registro de validación de la respuesta" in package_text
    assert "## Inventario del documento" in package_text
    assert "## Evaluaciones de las afirmaciones" in package_text
    assert "### Afirmación 1" in package_text
    assert run_intake["language"] == "es"
    assert "Codex debe ejecutar" in run_intake["dependency_check"]["note"]
    assert review_payload["language"] == "es"
    assert review_payload["columns"][1]["label"] == "Afirmación o artefacto"
    assert claim_item["title"].startswith("Afirmación 1:")
    assert claim_item["data"]["edit_hint"].startswith("Al editar esta afirmación")
    assert "Entrega para revisión" in handoff_text
    assert "Revisión en Codex" in handoff_text
    assert package_output["required_text"] == [
        "# Registro de validación de la respuesta",
        "## Límite de aseguramiento",
        "## Contrato de respuesta",
        "## Revisión del contrato de respuesta",
        "## Cobertura de la revisión",
        "## Inventario del documento",
        "## Evaluaciones de las afirmaciones",
    ]
    assert final_artifacts["caveats"][0].startswith("La identidad de la fuente")
    assert final_artifacts["next_actions"][0].startswith("Ejecute")
    contract_report = validate_contract(
        tmp_path / "out",
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()


def test_package_validation_flags_missing_review_fields(tmp_path: Path) -> None:
    package_mod = load_script(
        "deep_research_validator_package_validation_missing",
        "package_validation.py",
    )
    document_inventory = {"character_count": 0, "urls": []}
    source_inventory = {"sources": []}
    claims_review = {"claims": [{"claim_index": 1, "claim_text": "", "verdict": "bad"}]}

    audit = package_mod.build_audit(
        document_inventory,
        source_inventory,
        claims_review,
        _answer_contract(),
    )

    assert audit["status"] == "record_incomplete"
    assert "document_text_present" in audit["failed_checks"]
    assert "review_schema_version" in audit["failed_checks"]
    assert "coverage_review_complete" in audit["failed_checks"]
    assert "contract_review_complete" in audit["failed_checks"]
    assert "claim_assessments_complete" in audit["failed_checks"]
    assert "issue_treatments_complete" in audit["failed_checks"]


def test_quote_mismatch_does_not_override_semantic_support() -> None:
    package_mod = load_script(
        "deep_research_validator_semantic_boundary",
        "package_validation.py",
    )
    claims_review = _claims_review(
        [
            _claim_review(
                "This is a dog.",
                cited_passage="This is a terrier.",
                support_analysis=(
                    "A terrier is a type of dog, so the source semantically "
                    "supports the broader claim."
                ),
                reasoning_analysis=(
                    "The taxonomic inference is directional and valid."
                ),
            )
        ]
    )

    audit = package_mod.build_audit(
        {"character_count": 40, "urls": []},
        {
            "sources": [
                {
                    "source_id": "source-001",
                    "kind": "file",
                    "status": "available",
                    "excerpt": "The animal is a small terrier.",
                }
            ]
        },
        claims_review,
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["status"] == "record_complete"
    assert (
        audit["claim_observations"][0]["source_observations"][0][
            "exact_passage_presence"
        ]
        == "absent"
    )
    assert audit["support_attention_claim_indices"] == []


def test_audit_separates_reasoning_and_professional_judgment() -> None:
    package_mod = load_script(
        "deep_research_validator_reasoning_boundary",
        "package_validation.py",
    )
    issues = [
        {
            "type": "reasoning_gap",
            "explanation": "Application depends on unresolved facts.",
            "treatment_action": "state_uncertainty",
            "treatment_status": "applied",
            "treatment_explanation": "The answer states the missing application facts.",
        },
        {
            "type": "judgment_dependent",
            "explanation": "Application is fact-sensitive.",
            "treatment_action": "professional_review",
            "treatment_status": "professional_review_required",
            "treatment_explanation": "A legal professional must assess the facts.",
        },
    ]
    claims_review = _claims_review(
        [
            _claim_review(
                "The exception applies to this client.",
                support_status="partially_supported",
                support_analysis="The source establishes the exception but not its application.",
                reasoning_status="uncertain",
                reasoning_analysis="Application depends on unresolved facts.",
                judgment_status="professional_judgment_required",
                judgment_analysis="Materiality and application require legal review.",
                issues=issues,
                disposition_status="professional_review",
                reviewer_action="mark_unclear",
                proposed_fix="State the missing facts and avoid a final conclusion.",
            )
        ],
        validated_document="The exception may apply, subject to professional review.",
        document_revision_status="professional_review_required",
    )

    audit = package_mod.build_audit(
        {"character_count": 80, "urls": []},
        {
            "sources": [
                {
                    "source_id": "source-001",
                    "kind": "file",
                    "status": "available",
                    "excerpt": "The exception exists under specified conditions.",
                }
            ]
        },
        claims_review,
        _answer_contract(document_type="one-page legal letter"),
    )

    assert audit["status"] == "record_complete"
    assert audit["reasoning_attention_claim_indices"] == [1]
    assert audit["judgment_dependent_claim_indices"] == [1]
    assert (
        "does not certify legal correctness"
        in audit["assurance_boundary"]["record_integrity_meaning"]
    )
    assert audit["delivery_readiness"] == "professional_review_required"


def test_source_capture_preserves_passage_beyond_preview(tmp_path: Path) -> None:
    inspect_mod = load_script(
        "deep_research_validator_full_source_capture",
        "inspect_sources.py",
    )
    package_mod = load_script(
        "deep_research_validator_full_source_audit",
        "package_validation.py",
    )
    document_inventory_path = tmp_path / "document_inventory.json"
    document_inventory_path.write_text(
        json.dumps({"urls": [], "footnotes": [], "markdown_links": []}),
        encoding="utf-8",
    )
    source_file = tmp_path / "authority.txt"
    target_passage = "The decisive rule appears after the preview boundary."
    source_file.write_text(("x" * 1400) + target_passage, encoding="utf-8")
    output_dir = tmp_path / "out"

    paths = inspect_mod.write_source_inventory(
        document_inventory_path,
        output_dir,
        source_files=[source_file],
        fetch_urls=False,
    )
    source_inventory = json.loads(paths["source_inventory"].read_text(encoding="utf-8"))

    source = source_inventory["sources"][0]
    assert source["source_id"] == "source-001"
    assert target_passage not in source["excerpt"]
    assert source["capture_scope"] == "complete_local_text"
    assert target_passage in (output_dir / source["captured_text_path"]).read_text(
        encoding="utf-8"
    )

    audit = package_mod.build_audit(
        {"character_count": 40, "urls": []},
        source_inventory,
        _claims_review(
            [
                _claim_review(
                    "The decisive rule applies.",
                    cited_passage=target_passage,
                )
            ]
        ),
        _answer_contract(generation_route="codex_direct"),
        source_base_dir=output_dir,
    )

    observation = audit["claim_observations"][0]["source_observations"][0]
    assert observation["exact_passage_presence"] == "present"
    assert observation["observation_scope"] == "complete_local_text"


def test_exact_passage_is_not_searched_in_a_different_source() -> None:
    package_mod = load_script(
        "deep_research_validator_wrong_source_boundary",
        "package_validation.py",
    )
    wrong_source_issue = {
        "type": "wrong_source",
        "explanation": "The cited source is not the authority containing the passage.",
        "treatment_action": "replace_source",
        "treatment_status": "proposed",
        "treatment_explanation": "Replace the citation and reassess support.",
    }
    claim = _claim_review(
        "The rule applies.",
        cited_passage="The rule applies.",
        support_status="uncertain",
        support_analysis="Support cannot be assessed against the cited source.",
        issues=[wrong_source_issue],
        disposition_status="pending_source",
        reviewer_action="request_more_documents",
    )
    source_checks = claim["source_checks"]
    assert isinstance(source_checks, list)
    source_checks[0]["identity_status"] = "different_source"
    source_checks[0]["identity_analysis"] = "The passage belongs to source-002."

    audit = package_mod.build_audit(
        {"character_count": 40, "urls": []},
        {
            "sources": [
                {
                    "source_id": "source-001",
                    "status": "available",
                    "excerpt": "A different proposition appears here.",
                },
                {
                    "source_id": "source-002",
                    "status": "available",
                    "excerpt": "The rule applies.",
                },
            ]
        },
        _claims_review([claim], document_revision_status="required"),
        _answer_contract(generation_route="codex_direct"),
    )

    observations = audit["claim_observations"][0]["source_observations"]
    assert len(observations) == 1
    assert observations[0]["source_id"] == "source-001"
    assert observations[0]["exact_passage_presence"] == "absent"
    assert audit["source_identity_attention_claim_indices"] == [1]


def test_exact_overlap_does_not_override_semantic_contradiction() -> None:
    package_mod = load_script(
        "deep_research_validator_negation_boundary",
        "package_validation.py",
    )
    contradiction_issue = {
        "type": "source_contradiction",
        "explanation": "Negation reverses the proposition.",
        "treatment_action": "correct_claim",
        "treatment_status": "applied",
        "treatment_explanation": "The corrected answer preserves the negation.",
    }
    review = _claims_review(
        [
            _claim_review(
                "This is a dog.",
                cited_passage="This is not a dog.",
                support_status="contradicted",
                support_analysis="The source negates the claim.",
                issues=[contradiction_issue],
                disposition_status="revise",
                revised_claim="This is not a dog.",
                reviewer_action="edit",
            )
        ],
        validated_document="This is not a dog.",
        document_revision_status="completed",
    )

    audit = package_mod.build_audit(
        {"character_count": 40, "urls": []},
        {
            "sources": [
                {
                    "source_id": "source-001",
                    "status": "available",
                    "excerpt": "This is not a dog.",
                }
            ]
        },
        review,
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["status"] == "record_complete"
    assert audit["support_attention_claim_indices"] == [1]
    assert (
        audit["claim_observations"][0]["source_observations"][0][
            "exact_passage_presence"
        ]
        == "present"
    )
    assert review["claims"][0]["support"]["status"] == "contradicted"


@pytest.mark.parametrize(
    ("issue_type", "treatment_action", "treatment_status"),
    [
        ("source_unavailable", "obtain_source", "blocked"),
        ("source_not_identified", "identify_source", "proposed"),
        ("wrong_source", "replace_source", "proposed"),
        ("wrong_source_version", "replace_source", "proposed"),
        ("wrong_jurisdiction_or_period", "replace_source", "proposed"),
        ("missing_source_support", "add_support", "proposed"),
        ("partial_or_overbroad_support", "narrow_claim", "proposed"),
        ("source_contradiction", "correct_claim", "proposed"),
        ("qualification_or_scope_distortion", "restore_qualification", "proposed"),
        ("temporal_or_modality_distortion", "correct_time_or_modality", "proposed"),
        ("reasoning_gap", "add_reasoning", "proposed"),
        (
            "judgment_dependent",
            "professional_review",
            "professional_review_required",
        ),
        ("answer_contract_failure", "revise_answer_contract", "proposed"),
    ],
)
def test_issue_categories_require_explicit_treatment(
    issue_type: str,
    treatment_action: str,
    treatment_status: str,
) -> None:
    package_mod = load_script(
        f"deep_research_validator_issue_{issue_type}",
        "package_validation.py",
    )
    issue = {
        "type": issue_type,
        "explanation": f"The review identified {issue_type}.",
        "treatment_action": treatment_action,
        "treatment_status": treatment_status,
        "treatment_explanation": f"Apply {treatment_action} before delivery.",
    }
    claim = _claim_review(
        "A material claim.",
        issues=[issue],
        disposition_status=(
            "professional_review"
            if treatment_status == "professional_review_required"
            else "revise"
        ),
        reviewer_action="mark_unclear",
    )
    if treatment_status == "professional_review_required":
        judgment = claim["professional_judgment"]
        assert isinstance(judgment, dict)
        judgment["status"] = "professional_judgment_required"
        judgment["analysis"] = "Professional application remains necessary."
        judgment["factors"] = ["Application facts"]
        judgment["alternative_interpretations"] = ["Alternative application"]

    audit = package_mod.build_audit(
        {"character_count": 40, "urls": []},
        {
            "sources": [
                {
                    "source_id": "source-001",
                    "status": "available",
                    "excerpt": "A source passage.",
                }
            ]
        },
        _claims_review([claim], document_revision_status="required"),
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["status"] == "record_complete"
    assert audit["invalid_issue_indices"] == []


def test_issue_treatment_rejects_category_inappropriate_action() -> None:
    package_mod = load_script(
        "deep_research_validator_issue_action_contract",
        "package_validation.py",
    )
    claim = _claim_review(
        "A material claim.",
        issues=[
            {
                "type": "wrong_source",
                "explanation": "The cited authority is the wrong source.",
                "treatment_action": "add_reasoning",
                "treatment_status": "proposed",
                "treatment_explanation": "This action cannot repair source identity.",
            }
        ],
        disposition_status="revise",
        reviewer_action="edit",
    )

    audit = package_mod.build_audit(
        {"character_count": 40, "urls": []},
        {
            "sources": [
                {
                    "source_id": "source-001",
                    "status": "available",
                    "excerpt": "A source passage.",
                }
            ]
        },
        _claims_review([claim], document_revision_status="required"),
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["record_integrity_status"] == "record_incomplete"
    assert audit["invalid_issue_indices"] == [1]


def test_contract_failure_and_limited_coverage_have_distinct_readiness() -> None:
    package_mod = load_script(
        "deep_research_validator_contract_and_coverage",
        "package_validation.py",
    )
    source_inventory = {
        "sources": [
            {
                "source_id": "source-001",
                "status": "available",
                "excerpt": "A source passage.",
            }
        ]
    }
    contract_failure = _claims_review(
        [_claim_review("A material claim.")],
        document_revision_status="required",
    )
    contract_review = contract_failure["contract_review"]
    assert isinstance(contract_review, dict)
    contract_review["question_answered"] = {
        "status": "does_not_conform",
        "analysis": "The document answers a different question.",
    }
    contract_review["issues"] = [
        {
            "type": "answer_contract_failure",
            "explanation": "The answer addresses a different question.",
            "treatment_action": "revise_answer_contract",
            "treatment_status": "proposed",
            "treatment_explanation": "Rewrite the answer to the contracted question.",
        }
    ]
    contract_review["reviewer_action"] = "edit"

    contract_audit = package_mod.build_audit(
        {"character_count": 40, "urls": []},
        source_inventory,
        contract_failure,
        _answer_contract(generation_route="codex_direct"),
    )
    limited_audit = package_mod.build_audit(
        {"character_count": 40, "urls": []},
        source_inventory,
        _claims_review(
            [_claim_review("A material claim.")],
            coverage_scope="limited",
        ),
        _answer_contract(
            generation_route="codex_direct",
            validation_scope="limited",
        ),
    )

    assert contract_audit["contract_attention_dimensions"] == ["question_answered"]
    assert contract_audit["delivery_readiness"] == "revision_required"
    assert limited_audit["coverage_scope"] == "limited"
    assert limited_audit["delivery_readiness"] == "evidence_limited"


def test_one_page_letter_can_validate_all_material_claims() -> None:
    package_mod = load_script(
        "deep_research_validator_one_page_letter",
        "package_validation.py",
    )
    audit = package_mod.build_audit(
        {"character_count": 120, "urls": []},
        {
            "sources": [
                {
                    "source_id": "source-001",
                    "status": "available",
                    "excerpt": "The filing deadline is 30 days.",
                }
            ]
        },
        _claims_review(
            [
                _claim_review(
                    "The filing deadline is 30 days.",
                    cited_passage="The filing deadline is 30 days.",
                )
            ],
            validated_document="The filing deadline is 30 days.",
        ),
        _answer_contract(
            generation_route="codex_direct",
            document_type="one-page legal letter",
        ),
    )

    assert audit["record_integrity_status"] == "record_complete"
    assert audit["coverage_scope"] == "all_material_claims"
    assert audit["delivery_readiness"] == "reviewed_answer_ready"


def test_contradicted_claim_cannot_be_ready_without_treatment() -> None:
    package_mod = load_script(
        "deep_research_validator_contradiction_consistency",
        "package_validation.py",
    )
    review = _claims_review(
        [
            _claim_review(
                "This is a dog.",
                support_status="contradicted",
                support_analysis="The source says the opposite.",
            )
        ]
    )

    audit = package_mod.build_audit(
        {"character_count": 20, "urls": []},
        {"sources": [{"source_id": "source-001", "status": "available"}]},
        review,
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["record_integrity_status"] == "record_incomplete"
    assert audit["delivery_readiness"] == "review_record_incomplete"
    errors = {item["error"] for item in audit["consistency_errors"]}
    assert "attention_assessment_requires_issue_treatment" in errors
    assert "unsupported_or_contradicted_claim_cannot_be_retained" in errors


def test_unsound_reasoning_cannot_be_ready_when_claim_is_retained() -> None:
    package_mod = load_script(
        "deep_research_validator_reasoning_consistency",
        "package_validation.py",
    )
    review = _claims_review(
        [
            _claim_review(
                "The exception applies.",
                reasoning_status="unsound",
                reasoning_analysis="A required premise is missing.",
            )
        ]
    )

    audit = package_mod.build_audit(
        {"character_count": 20, "urls": []},
        {"sources": [{"source_id": "source-001", "status": "available"}]},
        review,
        _answer_contract(generation_route="codex_direct"),
    )

    errors = {item["error"] for item in audit["consistency_errors"]}
    assert audit["delivery_readiness"] == "review_record_incomplete"
    assert "unsound_reasoning_cannot_be_retained" in errors


def test_supported_claim_requires_an_identified_source_check() -> None:
    package_mod = load_script(
        "deep_research_validator_source_check_consistency",
        "package_validation.py",
    )
    claim = _claim_review("The filing period is 30 days.")
    claim["source_checks"] = []

    audit = package_mod.build_audit(
        {"character_count": 20, "urls": []},
        {"sources": []},
        _claims_review([claim]),
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["record_integrity_status"] == "record_incomplete"
    assert audit["consistency_errors"] == [
        {
            "scope": "claim",
            "claim_index": 1,
            "error": "source_check_required_for_support_assessment",
        }
    ]


def test_different_source_requires_an_issue_treatment() -> None:
    package_mod = load_script(
        "deep_research_validator_source_identity_consistency",
        "package_validation.py",
    )
    claim = _claim_review("The filing period is 30 days.")
    source_checks = claim["source_checks"]
    assert isinstance(source_checks, list)
    source_checks[0]["identity_status"] = "different_source"
    source_checks[0]["identity_analysis"] = "The passage is from another authority."

    audit = package_mod.build_audit(
        {"character_count": 20, "urls": []},
        {"sources": [{"source_id": "source-001", "status": "available"}]},
        _claims_review([claim]),
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["source_identity_attention_claim_indices"] == [1]
    assert audit["delivery_readiness"] == "review_record_incomplete"
    assert audit["consistency_errors"][0]["error"] == (
        "attention_assessment_requires_issue_treatment"
    )


def test_rejected_claim_is_never_reported_ready_for_delivery() -> None:
    package_mod = load_script(
        "deep_research_validator_reviewer_rejection",
        "package_validation.py",
    )
    review = _claims_review(
        [_claim_review("A material claim.", reviewer_action="reject")],
        document_revision_status="required",
    )

    audit = package_mod.build_audit(
        {"character_count": 20, "urls": []},
        {"sources": [{"source_id": "source-001", "status": "available"}]},
        review,
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["record_integrity_status"] == "record_complete"
    assert audit["rejected_claim_indices"] == [1]
    assert audit["delivery_readiness"] == "revision_required"


def test_contract_attention_requires_matching_failure_treatment() -> None:
    package_mod = load_script(
        "deep_research_validator_contract_consistency",
        "package_validation.py",
    )
    review = _claims_review(
        [_claim_review("A material claim.")],
        document_revision_status="required",
    )
    contract_review = review["contract_review"]
    assert isinstance(contract_review, dict)
    contract_review["question_answered"] = {
        "status": "does_not_conform",
        "analysis": "The answer addresses another question.",
    }

    audit = package_mod.build_audit(
        {"character_count": 20, "urls": []},
        {"sources": [{"source_id": "source-001", "status": "available"}]},
        review,
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["record_integrity_status"] == "record_incomplete"
    assert audit["consistency_errors"] == [
        {
            "scope": "contract_review",
            "error": "contract_attention_requires_failure_treatment",
        }
    ]


def test_contract_failure_treatment_requires_attention_status() -> None:
    package_mod = load_script(
        "deep_research_validator_contract_reverse_consistency",
        "package_validation.py",
    )
    review = _claims_review(
        [_claim_review("A material claim.")],
        document_revision_status="required",
    )
    contract_review = review["contract_review"]
    assert isinstance(contract_review, dict)
    contract_review["issues"] = [
        {
            "type": "answer_contract_failure",
            "explanation": "The contract requires revision.",
            "treatment_action": "revise_answer_contract",
            "treatment_status": "proposed",
            "treatment_explanation": "Revise the answer contract.",
        }
    ]

    audit = package_mod.build_audit(
        {"character_count": 20, "urls": []},
        {"sources": [{"source_id": "source-001", "status": "available"}]},
        review,
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["record_integrity_status"] == "record_incomplete"
    assert audit["consistency_errors"] == [
        {
            "scope": "contract_review",
            "error": "contract_failure_treatment_requires_attention_status",
        }
    ]


def test_professional_review_revision_cannot_claim_no_material_defect() -> None:
    package_mod = load_script(
        "deep_research_validator_revision_outcome_consistency",
        "package_validation.py",
    )
    review = _claims_review(
        [_claim_review("A material claim.")],
        document_revision_status="professional_review_required",
        overall_outcome="no_material_defect_identified",
    )

    audit = package_mod.build_audit(
        {"character_count": 20, "urls": []},
        {"sources": [{"source_id": "source-001", "status": "available"}]},
        review,
        _answer_contract(generation_route="codex_direct"),
    )

    assert audit["record_integrity_status"] == "record_incomplete"
    assert audit["consistency_errors"] == [
        {
            "scope": "document_revision",
            "error": "revision_status_conflicts_with_overall_outcome",
        }
    ]


def test_static_page_and_skill_match_plugin_contract() -> None:
    page = (
        ROOT / "static" / "shared" / "deep-research-validator" / "index.html"
    ).read_text(encoding="utf-8")
    skill = (PLUGIN_ROOT / "skills" / "deep-research-validator" / "SKILL.md").read_text(
        encoding="utf-8"
    )

    for snippet in (
        "Validate Deep Research",
        "Valida Deep Research",
        "Valider Deep Research",
        "Deep Research validieren",
        "One prompt to get started.",
        "Un solo prompt per iniziare.",
        "Un solo prompt para empezar.",
        "document_inventory.json",
        "source_inventory.json",
        "claims_review.json",
        "validation_audit.json",
        "validated_document.md",
        "validation_package.md",
        VERA_PRODUCT_PAGE_LINK,
        'id="vera-link"',
        "../vera/index.html?lang=it",
        'document.getElementById("vera-link").href = `../vera/index.html?lang=${lang}`',
        "/?lang=${lang}",
    ):
        assert snippet in page

    assert "must not make direct OpenAI API calls" in skill
    assert "Keep the improvement note local to chat or run artifacts." in skill
    assert "validate_deep_research_review" in skill
    assert "render_deep_research_review" in skill


def test_deep_research_mcp_server_validates_renders_and_applies_review_payload(
    tmp_path: Path,
) -> None:
    output_dir, context_path, context = _running_customer_output(tmp_path)
    client_run_id = str(context["run_id"])
    document_inventory_path = output_dir / "document_inventory.json"
    document_inventory_path.write_text(
        json.dumps(
            {
                "source_name": "deep_research.md",
                "character_count": 128,
                "word_count": 20,
                "urls": ["https://example.com"],
            }
        )
        + "\n",
        encoding="utf-8",
    )
    source_inventory_path = output_dir / "source_inventory.json"
    source_inventory_path.write_text(
        json.dumps(
            {
                "sources": [
                    {
                        "source_id": "source-001",
                        "url": "https://example.com",
                        "title": "Example source",
                        "status": "available",
                        "excerpt": "VAT rule applies in the cited transaction.",
                    }
                ]
            }
        )
        + "\n",
        encoding="utf-8",
    )
    claims_review_path = output_dir / "claims_review.json"
    claims_review_path.write_text(
        json.dumps(
            _claims_review(
                [
                    _claim_review(
                        "VAT rule applies.",
                        cited_passage="VAT rule applies",
                    )
                ]
            )
        )
        + "\n",
        encoding="utf-8",
    )
    (output_dir / "answer_contract.json").write_text(
        json.dumps(_answer_contract()) + "\n",
        encoding="utf-8",
    )
    (output_dir / "validation_audit.json").write_text(
        json.dumps(
            {
                "record_integrity_status": "record_complete",
                "delivery_readiness": "reviewed_answer_ready",
                "claim_count": 1,
                "source_count": 1,
            }
        )
        + "\n",
        encoding="utf-8",
    )
    (output_dir / "validation_package.md").write_text(
        "\n".join(
            [
                "# Answer Validation Record",
                "",
                "## Assurance Boundary",
                "",
                "## Answer Contract",
                "",
                "## Answer-Contract Review",
                "",
                "## Review Coverage",
                "",
                "## Document Inventory",
                "",
                "## Claim Assessments",
                "",
                "Proposed fix:",
                "",
            ]
        ),
        encoding="utf-8",
    )
    _write_docx(
        output_dir / "validated_document.docx",
        ["Reviewed answer", "Original answer text."],
    )
    run_intake = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": client_run_id,
        "created_at": "2026-01-01T00:00:00Z",
        "language": "en",
        "path_reference": "run_root_relative",
        "input_paths": [
            "outputs/document_inventory.json",
            "outputs/source_inventory.json",
            "outputs/claims_review.json",
            "outputs/answer_contract.json",
        ],
        "output_dir": "outputs",
        "inferred_task": "answer_validation_review_payload",
        "assumptions": {},
        "unresolved_questions": [],
        "dependency_check": {"status": "not_run"},
        "data_posture": {
            "local_files_read": [
                "outputs/document_inventory.json",
                "outputs/source_inventory.json",
                "outputs/claims_review.json",
                "outputs/answer_contract.json",
            ],
            "external_connectors_used": [],
            "upload_paths_used": [],
            "remote_sql_execution_used": False,
            "hosted_notebook_execution_used": False,
        },
    }
    review_payload = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": client_run_id,
        "source_paths": [
            "document_inventory.json",
            "source_inventory.json",
            "claims_review.json",
        ],
        "review_type": "answer_validation_review",
        "items": [
            {
                "id": "claim-1",
                "item_type": "supported_claim",
                "title": "Claim 1: VAT rule applies.",
                "output_path": "claims_review.json",
                "allowed_actions": ["accept", "edit", "mark_unclear", "skip"],
                "recommended_action": "accept",
                "evidence": [
                    {
                        "kind": "claim_review",
                        "claim_text": "VAT rule applies.",
                        "verdict": "supported",
                        "source_refs": ["https://example.com"],
                        "source_quote": "VAT rule applies",
                        "source_support": "Directly supported.",
                    }
                ],
                "data": {
                    "claim_index": 1,
                    "claim_text": "VAT rule applies.",
                    "verdict": "supported",
                    "target_artifact": "claims_review.json",
                    "target_records_key": "claims",
                    "target_id_field": "claim_index",
                    "target_record_id": "1",
                    "target_field": "proposed_fix",
                },
                "status": "needs_review",
            },
            {
                "id": "source-limit-1",
                "item_type": "source_limit",
                "title": "https://example.com/source",
                "output_path": "source_inventory.json",
                "allowed_actions": [
                    "accept",
                    "edit",
                    "mark_unclear",
                    "request_more_documents",
                    "skip",
                ],
                "recommended_action": "request_more_documents",
                "evidence": [
                    {
                        "kind": "source_availability",
                        "status": "listed_not_fetched",
                    }
                ],
                "data": {"status": "listed_not_fetched"},
                "status": "needs_review",
            },
        ],
        "item_count": 2,
        "columns": [],
        "source_artifacts": {},
        "evidence": {},
        "allowed_actions": [
            "accept",
            "reject",
            "edit",
            "mark_unclear",
            "request_more_documents",
            "skip",
        ],
        "status": "ready_for_review",
        "summary": {
            "record_integrity_status": "record_complete",
            "delivery_readiness": "reviewed_answer_ready",
            "claim_count": 1,
            "attention_claim_count": 0,
            "source_count": 1,
            "failed_check_count": 0,
        },
    }
    ui_decisions = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": client_run_id,
        "review_payload_path": "review_payload.json",
        "decisions": [],
        "decision_count": 0,
        "status": "pending_review",
    }
    final_artifacts = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": client_run_id,
        "outputs": [
            {
                "path": "claims_review.json",
                "kind": "json",
                "status": "written",
                "records_key": "claims",
                "row_count": 1,
                "required_columns": [
                    "claim_index",
                    "claim_text",
                    "verdict",
                    "proposed_fix",
                ],
            },
            {
                "path": "validation_audit.json",
                "kind": "json",
                "status": "written",
            },
            {
                "path": "validation_package.md",
                "kind": "md",
                "status": "written",
                "required_text": [
                    "# Answer Validation Record",
                    "## Assurance Boundary",
                    "## Answer Contract",
                    "## Answer-Contract Review",
                    "## Review Coverage",
                    "## Document Inventory",
                    "## Claim Assessments",
                ],
            },
            {
                "path": "validated_document.docx",
                "kind": "docx",
                "status": "written",
            },
        ],
        "caveats": [],
        "next_actions": [],
        "status": "written_pending_review",
    }
    (output_dir / "run_intake.json").write_text(
        json.dumps(run_intake, indent=2) + "\n",
        encoding="utf-8",
    )
    (output_dir / "review_payload.json").write_text(
        json.dumps(review_payload, indent=2) + "\n",
        encoding="utf-8",
    )
    (output_dir / "final_artifacts.json").write_text(
        json.dumps(final_artifacts, indent=2) + "\n",
        encoding="utf-8",
    )
    old_client_root = Path(context["studio_client_folder"]["client_root"])
    assert run_intake["path_reference"] == "run_root_relative"
    assert run_intake["output_dir"] == "outputs"
    assert all(not Path(value).is_absolute() for value in run_intake["input_paths"])
    renamed_client_root = tmp_path / "Renamed Customer"
    context_relative = context_path.relative_to(old_client_root)
    output_relative = output_dir.relative_to(old_client_root)
    old_client_root.rename(renamed_client_root)
    context_path = renamed_client_root / context_relative
    output_dir = renamed_client_root / output_relative
    claims_review_path = output_dir / "claims_review.json"
    assert not old_client_root.exists()
    messages: list[dict[str, object]] = [
        {"jsonrpc": "2.0", "id": 1, "method": "tools/list"},
        {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/call",
            "params": {
                "name": "validate_deep_research_review",
                "arguments": {"review_payload": review_payload},
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 3,
            "method": "tools/call",
            "params": {
                "name": "render_deep_research_review",
                "arguments": {
                    "run_intake": run_intake,
                    "client_engagement": context_path.as_posix(),
                    "review_payload": review_payload,
                    "ui_decisions": ui_decisions,
                    "final_artifacts": final_artifacts,
                },
            },
        },
        {"jsonrpc": "2.0", "id": 4, "method": "resources/list"},
        {
            "jsonrpc": "2.0",
            "id": 5,
            "method": "resources/read",
            "params": {"uri": "ui://widget/deep-research-review.html"},
        },
        {
            "jsonrpc": "2.0",
            "id": 6,
            "method": "tools/call",
            "params": {
                "name": "save_deep_research_decisions",
                "arguments": {
                    "run_intake": run_intake,
                    "client_engagement": context_path.as_posix(),
                    "review_payload": review_payload,
                    "ui_decisions": ui_decisions,
                    "decisions": [
                        {
                            "item_id": "claim-1",
                            "action": "edit",
                            "edit_value": (
                                "Narrow the claim to the cited VAT rule only."
                            ),
                        },
                        {
                            "item_id": "source-limit-1",
                            "action": "accept",
                        },
                    ],
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 7,
            "method": "tools/call",
            "params": {
                "name": "apply_deep_research_decisions",
                "arguments": {
                    "run_intake": run_intake,
                    "client_engagement": context_path.as_posix(),
                    "review_payload": review_payload,
                    "ui_decisions": ui_decisions,
                    "final_artifacts": final_artifacts,
                    "decisions": [
                        {
                            "item_id": "claim-1",
                            "action": "edit",
                            "edit_value": (
                                "Narrow the claim to the cited VAT rule only."
                            ),
                        },
                        {
                            "item_id": "source-limit-1",
                            "action": "accept",
                        },
                    ],
                },
            },
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}

    tool_names = {tool["name"] for tool in responses[1]["result"]["tools"]}
    assert {
        "validate_deep_research_review",
        "render_deep_research_review",
        "save_deep_research_decisions",
        "apply_deep_research_decisions",
    } <= tool_names
    validate_result = responses[2]["result"]["structuredContent"]
    assert validate_result["ok"] is True
    assert validate_result["item_count"] == 2
    render_result = responses[3]["result"]
    assert render_result["structuredContent"]["widget_type"] == "deep_research_review"
    assert (
        render_result["_meta"]["openai/outputTemplate"]
        == "ui://widget/deep-research-review.html"
    )
    resource_uris = {
        resource["uri"] for resource in responses[4]["result"]["resources"]
    }
    assert "ui://widget/deep-research-review.html" in resource_uris
    widget_html = responses[5]["result"]["contents"][0]["text"]
    assert "Answer Validation Review" in widget_html
    save_result = responses[6]["result"]["structuredContent"]
    assert save_result["ok"] is True
    assert save_result["persisted"] is True
    assert save_result["decision_count"] == 2
    apply_result = responses[7]["result"]["structuredContent"]
    assert apply_result["ok"] is True
    assert apply_result["run_intake_path"] == str(output_dir / "run_intake.json")
    assert apply_result["structured_update_count"] == 1
    assert apply_result["native_regeneration_count"] == 0
    assert apply_result["native_regenerated_count"] == 0
    assert apply_result["application_status"] == "revision_required"
    updated_claims = json.loads(claims_review_path.read_text(encoding="utf-8"))
    assert updated_claims["claims"][0]["proposed_fix"] == (
        "Narrow the claim to the cited VAT rule only."
    )
    assert updated_claims["document_revision"]["status"] == "required"
    assert updated_claims["validated_document"] == ""
    package_text = (output_dir / "validation_package.md").read_text(encoding="utf-8")
    assert "Narrow the claim to the cited VAT rule only." in package_text
    docx_text = _docx_text(output_dir / "validated_document.docx")
    assert "Original answer text." in docx_text
    assert "Narrow the claim to the cited VAT rule only." not in docx_text
    applied = json.loads((output_dir / "applied_decisions.json").read_text())
    assert applied["effects"][0]["structured_update"] == {
        "id_field": "claim_index",
        "record_id": "1",
        "target_field": "proposed_fix",
        "records_key": "claims",
        "updated_rows": 1,
    }
    assert applied["effects"][0]["downstream_regeneration_status"] == "regenerated"
    assert applied["effects"][0]["downstream_regenerated_paths"] == [
        "validation_audit.json",
        "validation_package.md",
    ]
    assert applied["effects"][0]["semantic_regeneration_required"] is True
    assert applied["effects"][0]["semantic_regeneration_status"] == "required"
    assert applied["downstream_regenerated_paths"] == [
        "validation_audit.json",
        "validation_package.md",
    ]
    assert applied["native_regeneration_count"] == 0
    assert applied["native_regenerated_count"] == 0
    assert applied["native_regenerated_paths"] == []
    assert applied["semantic_regeneration_count"] == 1
    final_after_apply = json.loads((output_dir / "final_artifacts.json").read_text())
    assert final_after_apply["status"] == "revision_required"
    claims_output = next(
        output
        for output in final_after_apply["outputs"]
        if output["path"] == "claims_review.json"
    )
    assert claims_output["status"] == "updated_from_review"
    assert claims_output["records_key"] == "claims"
    assert claims_output["required_columns"] == ["claim_index", "proposed_fix"]
    package_output = next(
        output
        for output in final_after_apply["outputs"]
        if output["path"] == "validation_package.md"
    )
    assert package_output["status"] == "updated_from_review"
    assert (
        "Narrow the claim to the cited VAT rule only."
        in package_output["required_text"]
    )
    docx_output = next(
        output
        for output in final_after_apply["outputs"]
        if output["path"] == "validated_document.docx"
    )
    assert docx_output["status"] == "superseded_requires_semantic_regeneration"
    assert final_after_apply["review_application"]["downstream_regenerated_paths"] == [
        "validation_audit.json",
        "validation_package.md",
    ]
    assert final_after_apply["review_application"]["native_regenerated_paths"] == []
    assert (
        final_after_apply["review_application"]["semantic_regeneration_status"]
        == "required"
    )
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_apply_steps = [
        step
        for step in run_intake["execution_trace"]
        if step["kind"] == "deterministic_review_apply"
    ]
    assert len(review_apply_steps) == 1
    assert {
        "applied_decisions.json",
        "claims_review.json",
        "final_artifacts.json",
        "validation_audit.json",
        "validation_package.md",
        "ui_decisions.json",
    } <= set(review_apply_steps[0]["outputs"])
    contract = validate_contract(
        output_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract.ok is True, contract.errors


def test_deep_research_mcp_localizes_spanish_runtime_and_handoff(
    tmp_path: Path,
) -> None:
    output_dir, context_path, context = _running_customer_output(tmp_path)
    client_run_id = str(context["run_id"])
    review_payload = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": client_run_id,
        "language": "es_ES",
        "review_type": "answer_validation_review",
        "item_count": 1,
        "items": [
            {
                "id": "artifact-1",
                "item_type": "validation_artifact",
                "title": "Paquete de validación",
                "output_path": "validation_package.md",
                "allowed_actions": ["accept", "mark_unclear", "skip"],
                "recommended_action": "accept",
                "status": "needs_review",
            }
        ],
    }
    run_intake = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": client_run_id,
        "language": "es",
        "path_reference": "run_root_relative",
        "output_dir": "outputs",
    }
    final_artifacts = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": client_run_id,
        "outputs": [],
        "caveats": [],
        "next_actions": [],
        "status": "written_pending_review",
    }
    decisions = [{"item_id": "artifact-1", "action": "accept"}]
    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "initialize",
            "params": {
                "protocolVersion": "2024-11-05",
                "_meta": {"locale": "es-ES"},
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/call",
            "params": {
                "name": "validate_deep_research_review",
                "arguments": {"review_payload": review_payload},
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 3,
            "method": "tools/call",
            "params": {
                "name": "validate_deep_research_review",
                "arguments": {
                    "review_payload": {**review_payload, "item_count": 2},
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 4,
            "method": "tools/call",
            "params": {
                "name": "save_deep_research_decisions",
                "arguments": {
                    "review_payload": review_payload,
                    "decisions": decisions,
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 5,
            "method": "tools/call",
            "params": {
                "name": "apply_deep_research_decisions",
                "arguments": {
                    "review_payload": review_payload,
                    "final_artifacts": final_artifacts,
                    "decisions": decisions,
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 6,
            "method": "tools/call",
            "params": {
                "name": "apply_deep_research_decisions",
                "arguments": {
                    "run_intake": run_intake,
                    "client_engagement": context_path.as_posix(),
                    "review_payload": review_payload,
                    "final_artifacts": final_artifacts,
                    "decisions": decisions,
                },
            },
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}
    initialized = responses[1]["result"]
    validated = responses[2]["result"]["structuredContent"]
    invalid = responses[3]["result"]["structuredContent"]
    saved_without_output = responses[4]["result"]["structuredContent"]
    applied_without_output = responses[5]["result"]["structuredContent"]
    applied = responses[6]["result"]["structuredContent"]

    assert "Ejecute validate_deep_research_review" in initialized["instructions"]
    assert "son válidos" in validated["message"]
    assert "debe coincidir" in invalid["error"]
    assert "no se ha escrito ningún archivo" in saved_without_output["message"]
    assert "no se ha escrito ningún archivo" in applied_without_output["message"]
    assert "Se ha aplicado 1 decisión" in applied["message"]
    assert applied["final_artifacts"]["next_actions"] == [
        "Utilice final_artifacts.json como galería revisada de artefactos para la entrega."
    ]
    handoff = (output_dir / "review_handoff.md").read_text(encoding="utf-8")
    assert "Entrega para revisión" in handoff
    assert "## Revisión en Codex" in handoff
    assert "<!-- review-contract: Review Handoff -->" in handoff
    assert "Validate the payload" not in handoff


@pytest.mark.parametrize(
    "tool_name",
    ["save_deep_research_decisions", "apply_deep_research_decisions"],
)
@pytest.mark.parametrize(
    "output_ref_kind", ["escape", "stale_absolute", "missing_context"]
)
def test_deep_research_mcp_managed_output_reference_is_rejected_without_writes(
    tmp_path: Path,
    tool_name: str,
    output_ref_kind: str,
) -> None:
    output_dir, context_path, context = _running_customer_output(tmp_path)
    old_client_root = Path(context["studio_client_folder"]["client_root"])
    if output_ref_kind == "stale_absolute":
        output_ref = output_dir.as_posix()
        renamed_client_root = tmp_path / "Renamed Customer"
        context_relative = context_path.relative_to(old_client_root)
        old_client_root.rename(renamed_client_root)
        context_path = renamed_client_root / context_relative
    elif output_ref_kind == "escape":
        output_ref = "../outside"
    else:
        output_ref = "outputs"
    run_id = str(context["run_id"])
    run_intake = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": run_id,
        "path_reference": "run_root_relative",
        "output_dir": output_ref,
    }
    review_payload = {
        "schema_version": "1.0",
        "plugin": "deep-research-validator",
        "workflow": "deep-research-validator",
        "run_id": run_id,
        "review_type": "answer_validation_review",
        "item_count": 1,
        "items": [
            {
                "id": "artifact-1",
                "item_type": "validation_artifact",
                "title": "Validation package",
                "output_path": "validation_package.md",
                "allowed_actions": ["accept", "mark_unclear", "skip"],
                "recommended_action": "accept",
                "status": "needs_review",
            }
        ],
    }
    arguments: dict[str, object] = {
        "run_intake": run_intake,
        "review_payload": review_payload,
        "decisions": [{"item_id": "artifact-1", "action": "accept"}],
    }
    if output_ref_kind != "missing_context":
        arguments["client_engagement"] = context_path.as_posix()
    if tool_name.startswith("apply_"):
        arguments["final_artifacts"] = {
            "schema_version": "1.0",
            "plugin": "deep-research-validator",
            "workflow": "deep-research-validator",
            "run_id": run_id,
            "outputs": [],
            "caveats": [],
            "next_actions": [],
            "status": "written_pending_review",
        }
    before = _file_snapshot(tmp_path)
    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "tools/call",
            "params": {"name": tool_name, "arguments": arguments},
        }
    ]

    response = _call_mcp_server(messages)[0]["result"]

    assert response["isError"] is True
    assert response["structuredContent"]["ok"] is False
    assert _file_snapshot(tmp_path) == before
    assert not (tmp_path / "outside").exists()
