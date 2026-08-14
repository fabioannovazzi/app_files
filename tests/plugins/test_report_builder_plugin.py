from __future__ import annotations

import copy
import hashlib
import importlib._bootstrap_external
import importlib.util
import json
import os
import re
import shutil
import stat
import subprocess
import sys
import threading
import time
import zipfile
from pathlib import Path
from typing import Any, Callable

import openpyxl
import pytest
from docx import Document
from pypdf import PdfWriter

from scripts.validate_plugin_review_contract import validate_contract

ROOT = Path(__file__).resolve().parents[2]
PLUGIN_ROOT = ROOT / "plugins" / "report-builder"
SCRIPT_DIR = PLUGIN_ROOT / "scripts"
CORE_PATH = SCRIPT_DIR / "report_builder_core.py"
MCP_SERVER_PATH = PLUGIN_ROOT / "mcp" / "server.cjs"


def _canonical_json_sha256(value: Any) -> str:
    return hashlib.sha256(
        json.dumps(
            value,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            allow_nan=False,
        ).encode("utf-8")
    ).hexdigest()


def load_core() -> Any:
    if str(SCRIPT_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPT_DIR))
    spec = importlib.util.spec_from_file_location("report_builder_core", CORE_PATH)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _load_customer_ledger() -> Any:
    path = ROOT / "plugins" / "studio-archive" / "scripts" / "client_ledger.py"
    module_name = "report_builder_customer_ledger"
    existing = sys.modules.get(module_name)
    if existing is not None:
        return existing
    spec = importlib.util.spec_from_file_location(module_name, path)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def _call_mcp_server_response(
    method: str, params: dict[str, Any] | None = None
) -> dict[str, Any]:
    if shutil.which("node") is None:
        pytest.skip("node is required for MCP server checks")
    request = {"jsonrpc": "2.0", "id": 1, "method": method, "params": params or {}}
    completed = subprocess.run(
        ["node", str(MCP_SERVER_PATH)],
        input=json.dumps(request) + "\n",
        capture_output=True,
        check=True,
        text=True,
    )
    responses = [
        json.loads(line)
        for line in completed.stdout.splitlines()
        if line.strip().startswith("{")
    ]
    assert responses
    return responses[-1]


def _call_mcp_server(
    method: str, params: dict[str, Any] | None = None
) -> dict[str, Any]:
    response = _call_mcp_server_response(method, params)
    assert "error" not in response
    return response["result"]


def test_spanish_docx_output_record_uses_spanish_required_text(
    tmp_path: Path,
) -> None:
    load_core()
    review_session = sys.modules["mparanza_report_builder_review_session"]
    (tmp_path / "report.docx").write_bytes(b"docx-placeholder")
    (tmp_path / "report_draft.md").write_text(
        "# Informe de gestión\n",
        encoding="utf-8",
    )
    analysis = {
        "language": "es",
        "sections": [{"title": "Resultados del periodo", "status": "assigned"}],
    }
    audit = {"missing_section_count": 1}

    outputs = review_session.build_output_records(tmp_path, audit, analysis)

    report = next(output for output in outputs if output["path"] == "report.docx")
    draft = next(output for output in outputs if output["path"] == "report_draft.md")
    assert report["required_text"] == [
        "Resumen ejecutivo",
        "Anexo de auditoría",
        "Estado del informe",
        "Llamadas a la API del modelo desde los scripts",
        "Secciones asignadas",
        "Ruta de entrada",
        "Tablas detectadas",
        "Secciones pendientes",
        "Resultados del periodo",
    ]
    assert draft["required_text"] == [
        "## Resumen ejecutivo",
        "## Resultados del periodo",
        "Fuente:",
        "Filas:",
    ]


def test_render_markdown_localizes_all_spanish_wrapper_copy() -> None:
    core = load_core()
    recipe = {
        "language": "es",
        "report_type": "management_report",
        "context_items": {"Moneda": "EUR"},
        "render": {"include_table_previews": False},
    }
    analysis = {
        "sections": [
            {
                "title": "Resultados",
                "status": "assigned",
                "source_file": "informe.xlsx",
                "sheet_name": "Resultados",
                "row_count": 3,
                "column_count": 2,
                "numeric_columns": [
                    {"column": "Importe", "numeric_count": 3, "sum": "250.00"}
                ],
                "preview_rows": [],
            },
            {
                "title": "Tesorería",
                "status": "unassigned",
                "numeric_columns": [],
                "preview_rows": [],
            },
        ]
    }

    markdown = core.render_markdown(recipe, analysis)

    assert markdown.startswith("# Informe de gestión")
    assert "**Entidad:** Entidad pendiente" in markdown
    assert "**Periodo:** Periodo pendiente" in markdown
    assert "## Resumen ejecutivo" in markdown
    assert "Resumen ejecutivo de Codex pendiente." in markdown
    assert "## Contexto" in markdown
    assert "La revisión de Codex está pendiente para esta sección." in markdown
    assert "Fuente: informe.xlsx / Resultados" in markdown
    assert "Filas: 3 | Columnas: 2" in markdown
    assert "Totales numéricos deterministas:" in markdown
    assert "Importe: suma 250.00" in markdown
    assert "recuento 3" not in markdown
    assert "Todavía no hay una tabla asignada." in markdown
    assert "Executive summary" not in markdown
    assert "Source:" not in markdown
    assert "Rows:" not in markdown


def test_numeric_profile_uses_exact_text_and_abstains_on_ambiguous_locale() -> None:
    core = load_core()
    rows = [
        ["Line", "Amount"],
        ["A", "100.10"],
        ["B", "200.20"],
        ["Ambiguous", "1.234"],
    ]

    profile = core.table_numeric_profile(rows, 0)

    assert profile["numeric_columns"] == [
        {
            "column": "Amount",
            "column_index": 2,
            "numeric_rows": [2, 3],
            "numeric_count": 2,
            "candidate_rows": [2, 3, 4],
            "sum": "300.3",
            "min": "100.1",
            "max": "200.2",
        }
    ]
    assert not any(
        isinstance(value, float)
        for column in profile["numeric_columns"]
        for value in column.values()
    )


def _save_workbook(path: Path) -> None:
    workbook = openpyxl.Workbook()
    income = workbook.active
    income.title = "Income Statement"
    income.append(["Line", "Actual", "Budget"])
    income.append(["Revenue", 1000, 950])
    income.append(["Costs", -620, -600])
    income.append(["Result", 380, 350])

    balance = workbook.create_sheet("Balance Sheet")
    balance.append(["Line", "Amount"])
    balance.append(["Assets", 2000])
    balance.append(["Equity", 900])
    balance.append(["Debt", 1100])

    cash = workbook.create_sheet("Cash Flow")
    cash.append(["Line", "Amount"])
    cash.append(["Operating cash", 250])
    cash.append(["Investing cash", -50])
    workbook.save(path)


def _numeric_review_args(
    inspection: dict[str, Any],
    table_id: str,
    included_columns: list[str],
    *,
    excluded_cell_rows: dict[str, set[int]] | None = None,
) -> dict[str, Any]:
    """Build explicit candidate/row dispositions for a known test fixture."""

    table = next(item for item in inspection["tables"] if item["table_id"] == table_id)
    candidates = [item["column"] for item in table["numeric_columns"]]
    inventory = {
        item["column"]: item["nonblank_cells"]
        for item in table["numeric_measure_cells"]
    }
    excluded_rows = excluded_cell_rows or {}
    return {
        "header_row": table["header_row"],
        "columns": included_columns,
        "excluded_columns": [
            column for column in candidates if column not in included_columns
        ],
        "cell_dispositions": {
            column: {
                cell["row"]: (
                    "exclude"
                    if cell["row"] in excluded_rows.get(column, set())
                    else "include"
                )
                for cell in inventory[column]
            }
            for column in included_columns
        },
        "sign_policy": "as_presented_v1",
    }


def _save_formula_workbook(path: Path, cached_value: str | None) -> None:
    """Create a formula workbook and optionally inject a cached OOXML value."""

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Budget"
    sheet.append(["line", "amount"])
    sheet.append(["A", 10])
    sheet.append(["B", 20])
    sheet.append(["Subtotal", "=SUM(B2:B3)"])
    workbook.save(path)
    if cached_value is None:
        return
    rewritten = path.with_name(f"{path.stem}-cached.xlsx")
    replaced = 0
    with (
        zipfile.ZipFile(path) as source,
        zipfile.ZipFile(
            rewritten,
            "w",
            compression=zipfile.ZIP_DEFLATED,
        ) as target,
    ):
        for member in source.infolist():
            data = source.read(member.filename)
            if member.filename == "xl/worksheets/sheet1.xml":
                data, replaced = re.subn(
                    rb'(<c r="B4"[^>]*><f>[^<]*</f><v>)[^<]*(</v>)',
                    rb"\g<1>" + cached_value.encode("ascii") + rb"\g<2>",
                    data,
                )
            target.writestr(member, data)
    assert replaced == 1
    rewritten.replace(path)


def _tree_snapshot(root: Path) -> dict[str, bytes]:
    return {
        path.relative_to(root).as_posix(): path.read_bytes()
        for path in root.rglob("*")
        if path.is_file()
    }


def _transaction_tree_state(
    root: Path,
) -> dict[str, tuple[str, bytes | str | None, int]]:
    state = {
        ".": (
            "directory",
            None,
            stat.S_IMODE(root.lstat().st_mode),
        )
    }
    for candidate in sorted(root.rglob("*")):
        entry = candidate.lstat()
        relative = candidate.relative_to(root).as_posix()
        mode = stat.S_IMODE(entry.st_mode)
        if stat.S_ISREG(entry.st_mode):
            state[relative] = ("file", candidate.read_bytes(), mode)
        elif stat.S_ISDIR(entry.st_mode):
            state[relative] = ("directory", None, mode)
        elif stat.S_ISLNK(entry.st_mode):
            state[relative] = ("symlink", os.readlink(candidate), mode)
        elif stat.S_ISFIFO(entry.st_mode):
            state[relative] = ("fifo", None, mode)
        else:
            state[relative] = ("special", None, mode)
    return state


def _install_review_transaction_commit_fault(
    monkeypatch: Any,
    tmp_path: Path,
    output_dir: Path,
    scenario: str,
) -> tuple[Path, Path]:
    preload = tmp_path / f"review-transaction-{scenario}.cjs"
    preload.write_text(
        """
"use strict";
const childProcess = require("node:child_process");
const fs = require("node:fs");
const path = require("node:path");
const originalRenameSync = fs.renameSync;
const canonical = path.resolve(process.env.REVIEW_TX_CANONICAL);
const external = path.resolve(process.env.REVIEW_TX_EXTERNAL);
const marker = path.resolve(process.env.REVIEW_TX_MARKER);
const scenario = process.env.REVIEW_TX_SCENARIO;
let triggered = false;
fs.renameSync = function reviewTransactionCommitFault(source, target) {
  const commitRoot = path.dirname(source);
  const isCommit =
    path.basename(source) === "candidate" &&
    path.basename(commitRoot).startsWith(
      ".generated-review-commit-",
    ) &&
    path.resolve(target) === canonical;
  if (!triggered && isCommit) {
    triggered = true;
    fs.writeFileSync(marker, "commit fault triggered\\n");
    const trustedBackup = path.join(commitRoot, "trusted-backup");
    if (scenario === "poison_file") {
      fs.writeFileSync(
        path.join(trustedBackup, "run_intake.json"),
        "POISONED SNAPSHOT\\n",
      );
    } else if (scenario === "delete_snapshot") {
      fs.rmSync(trustedBackup, { recursive: true, force: true });
    } else if (scenario === "symlink_snapshot") {
      fs.rmSync(trustedBackup, { recursive: true, force: true });
      fs.symlinkSync(external, trustedBackup, "dir");
    } else if (scenario === "fifo_snapshot") {
      fs.rmSync(trustedBackup, { recursive: true, force: true });
      fs.mkdirSync(trustedBackup, { mode: 0o700 });
      const fifo = path.join(trustedBackup, "poison.fifo");
      const created = childProcess.spawnSync("mkfifo", [fifo]);
      if (created.status !== 0) {
        throw new Error("FIFO probe setup failed");
      }
    } else if (scenario === "canonical_delete") {
      originalRenameSync.call(fs, source, target);
      fs.rmSync(target, { recursive: true, force: true });
      throw new Error(
        "/private/client/run injected late transaction failure",
      );
    } else {
      throw new Error("Unknown transaction fault scenario");
    }
    throw new Error(
      "/private/client/run injected late transaction failure",
    );
  }
  return originalRenameSync.call(fs, source, target);
};
""".lstrip(),
        encoding="utf-8",
    )
    external = tmp_path / f"external-{scenario}"
    external.mkdir()
    (external / "sentinel.bin").write_bytes(b"external unchanged")
    marker = tmp_path / f"commit-fault-{scenario}.marker"
    monkeypatch.setenv("NODE_OPTIONS", f"--require={preload}")
    monkeypatch.setenv("REVIEW_TX_CANONICAL", output_dir.as_posix())
    monkeypatch.setenv("REVIEW_TX_EXTERNAL", external.as_posix())
    monkeypatch.setenv("REVIEW_TX_MARKER", marker.as_posix())
    monkeypatch.setenv("REVIEW_TX_SCENARIO", scenario)
    return marker, external


def _install_review_transaction_root_relocation(
    monkeypatch: Any,
    tmp_path: Path,
) -> Path:
    preload = tmp_path / "review-transaction-root-relocation.cjs"
    preload.write_text(
        """
"use strict";
const fs = require("node:fs");
const Module = require("node:module");
const path = require("node:path");
const originalLoad = Module._load;
const marker = path.resolve(process.env.REVIEW_TX_MARKER);
let triggered = false;
Module._load = function reviewTransactionRootRelocation(
  request,
  parent,
  isMain,
) {
  const loaded = originalLoad.call(this, request, parent, isMain);
  if (
    typeof request !== "string" ||
    !request.endsWith("review_output_transaction.cjs") ||
    typeof loaded.withGeneratedReviewOutputTransaction !== "function"
  ) {
    return loaded;
  }
  return {
    ...loaded,
    withGeneratedReviewOutputTransaction(outputDir, operation, options) {
      return loaded.withGeneratedReviewOutputTransaction(
        outputDir,
        (context) => {
          const envelope = operation(context);
          if (!triggered && context.workingOutputDir) {
            triggered = true;
            const transactionRoot = path.dirname(context.workingOutputDir);
            fs.renameSync(
              transactionRoot,
              `${transactionRoot}-moved`,
            );
            fs.writeFileSync(marker, "transaction root relocated\\n");
          }
          return envelope;
        },
        options,
      );
    },
  };
};
""".lstrip(),
        encoding="utf-8",
    )
    marker = tmp_path / "transaction-root-relocation.marker"
    monkeypatch.setenv("NODE_OPTIONS", f"--require={preload}")
    monkeypatch.setenv("REVIEW_TX_MARKER", marker.as_posix())
    return marker


def _install_report_transaction_result_fault(
    monkeypatch: Any,
    tmp_path: Path,
    output_dir: Path,
    attack_kind: str,
) -> Path:
    preload = tmp_path / f"report-result-{attack_kind}.cjs"
    preload.write_text(
        """
"use strict";
const fs = require("node:fs");
const Module = require("node:module");
const path = require("node:path");
const originalLoad = Module._load;
const canonical = path.resolve(process.env.REVIEW_TX_CANONICAL);
const marker = path.resolve(process.env.REVIEW_TX_MARKER);
const attackKind = process.env.REVIEW_TX_RESULT_ATTACK;
let triggered = false;
Module._load = function reportTransactionResultFault(
  request,
  parent,
  isMain,
) {
  const loaded = originalLoad.call(this, request, parent, isMain);
  if (
    typeof request !== "string" ||
    !request.endsWith("review_output_transaction.cjs") ||
    typeof loaded.generatedReviewTransactionEnvelope !== "function"
  ) {
    return loaded;
  }
  const originalEnvelope = loaded.generatedReviewTransactionEnvelope;
  return {
    ...loaded,
    generatedReviewTransactionEnvelope(result, authorizedWritePaths) {
      const envelope = originalEnvelope(result, authorizedWritePaths);
      if (
        !triggered &&
        result &&
        (result.validation_type === "report_builder_decisions" ||
          result.validation_type === "report_builder_application")
      ) {
        triggered = true;
        fs.writeFileSync(marker, "result fault triggered\\n");
        if (attackKind === "forged_result") {
          envelope.result = { ...result, run_id: "forged-run-id" };
        } else if (attackKind === "forged_scalar_count") {
          envelope.result = { ...result, revision_count: 999 };
        } else if (
          attackKind === "unauthorized_source_mapping_path"
        ) {
          const parentDir = path.dirname(canonical);
          const transactionName = fs
            .readdirSync(parentDir)
            .find((name) =>
              name.startsWith(".generated-review-transaction-"),
            );
          if (!transactionName) {
            throw new Error("Staged transaction directory was not found");
          }
          fs.writeFileSync(
            path.join(
              parentDir,
              transactionName,
              "working",
              "rogue-source-mapping.bin",
            ),
            "unauthorized\\n",
          );
        } else if (attackKind === "tampered_staged_file") {
          const parentDir = path.dirname(canonical);
          const transactionName = fs
            .readdirSync(parentDir)
            .find((name) =>
              name.startsWith(".generated-review-transaction-"),
            );
          if (!transactionName) {
            throw new Error("Staged transaction directory was not found");
          }
          fs.writeFileSync(
            path.join(
              parentDir,
              transactionName,
              "working",
              "ui_decisions.json",
            ),
            '{"forged":true}\\n',
          );
        } else {
          throw new Error("Unknown result fault scenario");
        }
      }
      return envelope;
    },
  };
};
""".lstrip(),
        encoding="utf-8",
    )
    marker = tmp_path / f"result-fault-{attack_kind}.marker"
    monkeypatch.setenv("NODE_OPTIONS", f"--require={preload}")
    monkeypatch.setenv("REVIEW_TX_CANONICAL", output_dir.as_posix())
    monkeypatch.setenv("REVIEW_TX_MARKER", marker.as_posix())
    monkeypatch.setenv("REVIEW_TX_RESULT_ATTACK", attack_kind)
    return marker


def _report_transaction_fixture(
    tmp_path: Path,
) -> tuple[Path, dict[str, Any]]:
    source = tmp_path / "budget.csv"
    source.write_text("line,amount\nA,10\nB,20\n", encoding="utf-8")
    managed = _managed_report_run(tmp_path, source)
    output_dir = managed["output_dir"]
    run_id = managed["run_id"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    assert run_intake["path_reference"] == "run_root_relative"
    assert run_intake["output_dir"] == "outputs"
    assert run_intake["run_id"] == run_id
    assert review_payload["run_id"] == run_id
    section_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "report_section"
        and item["data"]["status"] == "assigned"
    )
    return output_dir, {
        "client_engagement": managed["context_path"].as_posix(),
        "run_intake": run_intake,
        "review_payload": review_payload,
        "final_artifacts": final_artifacts,
        "decisions": [
            {
                "item_id": section_item["id"],
                "action": "edit",
                "edit_value": "Transaction containment review.",
            }
        ],
    }


def _managed_report_run(
    tmp_path: Path,
    source: Path,
    *,
    recipe: Path | None = None,
    recipe_builder: Callable[[Any, Path, Path], Path] | None = None,
    report_type: str = "management_report",
    language: str = "en",
    document_language: str = "auto",
) -> dict[str, Any]:
    """Build a report in one running ledger run from exact imported receipts."""

    if recipe is not None and recipe_builder is not None:
        raise ValueError("recipe and recipe_builder are mutually exclusive")
    core = load_core()
    ledger = _load_customer_ledger()
    client_root = tmp_path / "Customer"
    client_root.mkdir()
    client_id = "client_555555555555555555555555"
    ledger.create_client_manifest(client_root, client_id)
    engagement = ledger.create_engagement(client_root, client_id, "Report review")
    imported = ledger.import_document(
        client_root,
        client_id,
        engagement["engagement_id"],
        source,
        "source",
    )
    imported_recipe = (
        ledger.import_document(
            client_root,
            client_id,
            engagement["engagement_id"],
            recipe,
            "source",
        )
        if recipe is not None
        else None
    )
    input_ids = [imported["receipt"]["input_id"]]
    if imported_recipe is not None:
        input_ids.append(imported_recipe["receipt"]["input_id"])
    prepared = ledger.prepare_run(
        client_root,
        client_id,
        engagement["engagement_id"],
        "report-builder",
        "test-version",
        input_ids=input_ids,
    )
    running = ledger.start_run(
        client_root,
        engagement["engagement_id"],
        prepared["run"]["run_id"],
    )
    bindings_by_id = {
        item["binding_id"]: item for item in running["context"]["input_bindings"]
    }
    input_path = Path(bindings_by_id[imported["receipt"]["input_id"]]["path"])
    recipe_path = (
        Path(bindings_by_id[imported_recipe["receipt"]["input_id"]]["path"])
        if imported_recipe is not None
        else None
    )
    output_dir = Path(running["output_dir"])
    run_id = str(running["context"]["run_id"])
    if recipe_builder is not None:
        generated_recipe = recipe_builder(core, input_path, output_dir)
        recipe_path = output_dir / "suggested_recipe.json"
        if generated_recipe.resolve() != recipe_path.resolve():
            recipe_path.write_bytes(generated_recipe.read_bytes())
            generated_recipe.unlink()
    result = core.build_report(
        input_path,
        output_dir,
        recipe_path=recipe_path,
        report_type=report_type,
        language=language,
        document_language=document_language,
        run_id=run_id,
        client_engagement=running["context"],
    )
    return {
        "core": core,
        "output_dir": output_dir,
        "source_path": input_path,
        "recipe_path": recipe_path,
        "run_id": run_id,
        "context": running["context"],
        "context_path": Path(running["context_path"]),
        "result": result,
    }


def _current_report_source(output_dir: Path) -> Path:
    """Return the exact imported source path bound by the private source index."""

    source_index = json.loads((output_dir / "source_index.json").read_text())
    source = source_index["sources"][0]
    source_root = Path(source["root_path"])
    if not source_root.is_absolute():
        source_root = output_dir.parent / source_root
    return source_root / source["receipt"]["path"]


@pytest.mark.parametrize(
    "tool_name",
    ["save_report_builder_decisions", "apply_report_builder_decisions"],
)
def test_managed_review_persistence_continues_after_customer_folder_rename(
    tmp_path: Path,
    tool_name: str,
) -> None:
    source = tmp_path / "budget.csv"
    source.write_text("line,amount\nA,10\nB,20\n", encoding="utf-8")
    managed = _managed_report_run(tmp_path, source)
    original_client_root = tmp_path / "Customer"
    renamed_client_root = tmp_path / "Renamed Customer"
    source_index = json.loads(
        (managed["output_dir"] / "source_index.json").read_text(encoding="utf-8")
    )
    assert source_index["sources"][0]["identity_key"].startswith("run_root:inputs/")
    original_root_text = original_client_root.resolve().as_posix()
    for artifact in managed["output_dir"].rglob("*"):
        if artifact.is_file() and artifact.suffix.lower() in {".json", ".md"}:
            assert original_root_text not in artifact.read_text(encoding="utf-8")
    relative_context = managed["context_path"].relative_to(original_client_root)
    relative_output = managed["output_dir"].relative_to(original_client_root)
    original_client_root.rename(renamed_client_root)
    context_path = renamed_client_root / relative_context
    output_dir = renamed_client_root / relative_output
    sys.modules["mparanza_report_builder_integrity"].validate_review_integrity(
        output_dir
    )
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    section_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "report_section"
        and item["data"]["status"] == "assigned"
    )

    result = _call_mcp_server(
        "tools/call",
        {
            "name": tool_name,
            "arguments": {
                "client_engagement": context_path.as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": section_item["id"],
                        "action": "accept",
                    }
                ],
            },
        },
    )["structuredContent"]

    assert result["ok"] is True, result
    assert result["persisted"] is True
    assert (output_dir / "ui_decisions.json").is_file()
    if tool_name == "apply_report_builder_decisions":
        assert (output_dir / "applied_decisions.json").is_file()
    renamed_root_text = renamed_client_root.resolve().as_posix()
    for artifact in output_dir.rglob("*"):
        if artifact.is_file() and artifact.suffix.lower() in {".json", ".md"}:
            assert renamed_root_text not in artifact.read_text(encoding="utf-8")


def test_managed_source_index_rejects_absolute_identity_key(tmp_path: Path) -> None:
    source = tmp_path / "budget.csv"
    source.write_text("line,amount\nA,10\nB,20\n", encoding="utf-8")
    managed = _managed_report_run(tmp_path, source)
    output_dir = managed["output_dir"]
    index_path = output_dir / "source_index.json"
    source_index = json.loads(index_path.read_text(encoding="utf-8"))
    source_index["sources"][0]["identity_key"] = managed["source_path"].as_posix()
    integrity = sys.modules["mparanza_report_builder_integrity"]
    content = {
        "schema_version": source_index["schema_version"],
        "sources": source_index["sources"],
        "archive_manifests": source_index["archive_manifests"],
        "archive_member_bindings": source_index["archive_member_bindings"],
    }
    source_index["content_sha256"] = integrity.canonical_json_sha256(content)
    index_path.write_text(
        json.dumps(source_index, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="identity is not portable"):
        integrity.validate_source_index(output_dir)


def test_managed_archive_source_identities_are_portable_and_container_bound(
    tmp_path: Path,
) -> None:
    archive_path = tmp_path / "budget.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr("nested/budget.csv", "line,amount\nA,10\nB,20\n")
    managed = _managed_report_run(tmp_path, archive_path)
    source_index = json.loads(
        (managed["output_dir"] / "source_index.json").read_text(encoding="utf-8")
    )
    sources = {source["artifact_id"]: source for source in source_index["sources"]}
    container_id = source_index["archive_manifests"][0]["container_artifact_id"]
    binding = source_index["archive_member_bindings"][0]
    container_identity = sources[container_id]["identity_key"]
    member_identity = sources[binding["member_artifact_id"]]["identity_key"]

    assert container_identity.startswith("run_root:inputs/")
    assert member_identity == (f"{container_identity}::{binding['member_path']}")
    assert (tmp_path / "Customer").resolve().as_posix() not in json.dumps(
        source_index,
        ensure_ascii=False,
    )


@pytest.mark.parametrize(
    "tool_name",
    ["save_report_builder_decisions", "apply_report_builder_decisions"],
)
def test_managed_review_persistence_rejects_output_escape_without_writes(
    tmp_path: Path,
    tool_name: str,
) -> None:
    source = tmp_path / "budget.csv"
    source.write_text("line,amount\nA,10\nB,20\n", encoding="utf-8")
    managed = _managed_report_run(tmp_path, source)
    output_dir = managed["output_dir"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    run_intake["output_dir"] = "../escaped-output"
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    section_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "report_section"
        and item["data"]["status"] == "assigned"
    )
    before = _tree_snapshot(output_dir)
    escaped_output = managed["context_path"].parent.parent / "escaped-output"

    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": tool_name,
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": section_item["id"],
                        "action": "accept",
                    }
                ],
            },
        },
    )

    failure = response["result"]["structuredContent"]
    assert failure["ok"] is False
    assert failure["error"] == (
        "Report Builder output reference leaves the customer run."
    )
    assert _tree_snapshot(output_dir) == before
    assert not escaped_output.exists()


def _reseal_report_builder_self_consistent(output_dir: Path) -> None:
    """Refresh all self-authored hashes without performing re-derivation."""

    final_path = output_dir / "final_artifacts.json"
    final_artifacts = json.loads(final_path.read_text(encoding="utf-8"))
    for output in final_artifacts["outputs"]:
        artifact = output_dir / output["path"]
        if not artifact.is_file():
            continue
        content = artifact.read_bytes()
        output["size_bytes"] = len(content)
        output["sha256"] = hashlib.sha256(content).hexdigest()
    final_path.write_text(
        json.dumps(final_artifacts, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    integrity_path = output_dir / "review_integrity.json"
    integrity = json.loads(integrity_path.read_text(encoding="utf-8"))
    for receipt in integrity["protected_files"]:
        artifact = output_dir / receipt["path"]
        content = artifact.read_bytes()
        receipt["byte_count"] = len(content)
        receipt["sha256"] = hashlib.sha256(content).hexdigest()
    for name in ("run_intake", "review_payload", "final_artifacts"):
        payload = json.loads((output_dir / f"{name}.json").read_text(encoding="utf-8"))
        integrity["payload_digests"][name] = hashlib.sha256(
            json.dumps(
                payload,
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
                allow_nan=False,
            ).encode("utf-8")
        ).hexdigest()
    for receipt in integrity["prepared_validation"]["rederived_artifacts"]:
        artifact = output_dir / receipt["path"]
        content = artifact.read_bytes()
        receipt["byte_count"] = len(content)
        receipt["sha256"] = hashlib.sha256(content).hexdigest()
    content = dict(integrity)
    content.pop("content_sha256")
    integrity["content_sha256"] = hashlib.sha256(
        json.dumps(
            content,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            allow_nan=False,
        ).encode("utf-8")
    ).hexdigest()
    integrity_path.write_text(
        json.dumps(integrity, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def test_official_reseal_rejects_extra_report_audit_claim(tmp_path: Path) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    audit_path = output_dir / "report_audit.json"
    audit = json.loads(audit_path.read_text(encoding="utf-8"))
    audit["forged_professional_claim"] = "The report was professionally approved."
    audit_path.write_text(
        json.dumps(audit, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    final_path = output_dir / "final_artifacts.json"
    final_artifacts = json.loads(final_path.read_text(encoding="utf-8"))
    audit_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "report_audit.json"
    )
    audit_bytes = audit_path.read_bytes()
    audit_output["size_bytes"] = len(audit_bytes)
    audit_output["sha256"] = hashlib.sha256(audit_bytes).hexdigest()
    final_path.write_text(
        json.dumps(final_artifacts, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    core = load_core()

    # Act / Assert
    with pytest.raises(
        ValueError,
        match="Report Builder audit projection is not exactly rederived",
    ):
        core.seal_review_integrity(
            output_dir,
            run_id=arguments["run_intake"]["run_id"],
        )


def test_python_integrity_rejects_current_source_hardlink(tmp_path: Path) -> None:
    # Arrange
    output_dir, _ = _report_transaction_fixture(tmp_path)
    source_path = _current_report_source(output_dir)
    external_path = tmp_path / "same-source-bytes.csv"
    external_path.write_bytes(source_path.read_bytes())
    source_path.unlink()
    source_path.hardlink_to(external_path)
    integrity = sys.modules["mparanza_report_builder_integrity"]

    # Act / Assert
    with pytest.raises(ValueError, match="ordinary single-link file"):
        integrity.validate_review_integrity(output_dir)


@pytest.mark.parametrize(
    "tool_name",
    [
        "validate_report_builder_review",
        "save_report_builder_decisions",
        "apply_report_builder_decisions",
    ],
)
@pytest.mark.parametrize("attack_kind", ["rogue_file", "empty_directory"])
def test_assured_review_rejects_unowned_physical_paths_without_mutation(
    tmp_path: Path,
    tool_name: str,
    attack_kind: str,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    if attack_kind == "rogue_file":
        (output_dir / "rogue-unowned.bin").write_bytes(b"foreign")
    else:
        (output_dir / "rogue-empty").mkdir()
    before = _transaction_tree_state(output_dir)

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {"name": tool_name, "arguments": arguments},
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert "authorization" in failure["error"].lower()
    assert _transaction_tree_state(output_dir) == before


@pytest.mark.parametrize("attack_kind", ["symlink", "hardlink", "fifo"])
def test_assured_review_rejects_nonordinary_owned_artifact(
    tmp_path: Path,
    attack_kind: str,
) -> None:
    # Arrange
    if attack_kind == "fifo" and sys.platform == "win32":
        pytest.skip("FIFO physical-output probe requires a POSIX host.")
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    target = output_dir / "report_draft.md"
    content = target.read_bytes()
    external = tmp_path / "report-draft-source.md"
    external.write_bytes(content)
    target.unlink()
    if attack_kind == "symlink":
        target.symlink_to(external)
    elif attack_kind == "hardlink":
        os.link(external, target)
    else:
        os.mkfifo(target)
    before = _transaction_tree_state(output_dir)

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": arguments,
        },
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == "Report Builder persisted review authorization failed."
    assert _transaction_tree_state(output_dir) == before


def test_build_rejects_preexisting_unowned_path_and_restores_prior_tree(
    tmp_path: Path,
) -> None:
    # Arrange
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text("line,amount\nA,10\n", encoding="utf-8")
    output_dir = tmp_path / "report"
    output_dir.mkdir()
    (output_dir / "rogue-unowned.bin").write_bytes(b"prior")
    before = _transaction_tree_state(output_dir)

    # Act / Assert
    with pytest.raises(ValueError, match="physical output set does not close"):
        core.build_report(
            input_path,
            output_dir,
            report_type="management_report",
        )
    assert _transaction_tree_state(output_dir) == before


def test_review_integrity_receipts_exact_transitive_implementation_set(
    tmp_path: Path,
) -> None:
    # Arrange / Act
    output_dir, _ = _report_transaction_fixture(tmp_path)
    integrity = json.loads(
        (output_dir / "review_integrity.json").read_text(encoding="utf-8")
    )
    references = integrity["implementation_artifact_refs"]
    receipts = integrity["implementation_receipts"]

    # Assert
    assert integrity["schema_version"] == "report_builder.review_integrity.v4"
    assert len(references) == 32
    assert [receipt["artifact_id"] for receipt in receipts] == references
    assert {
        "implementation.report_builder.scripts.report_builder_core.py",
        "implementation.report_builder.scripts.implementation_bootstrap.py",
        "implementation.report_builder.mcp.server.cjs",
        "implementation.report_builder..app.json",
        "implementation.report_builder..mcp.json",
        "implementation.report_builder.assets.icon.svg",
        "implementation.report_builder.assets.report-builder-review-widget.html",
        "implementation.report_builder.scripts.review_successor.py",
        "implementation.vera_assurance.serialization.py",
        "implementation.vera_assurance.review_output_transaction.cjs",
    } <= set(references)


@pytest.mark.parametrize("attack_kind", ["changed", "missing", "expanded", "reordered"])
def test_mcp_rejects_forged_transitive_implementation_receipts(
    tmp_path: Path,
    attack_kind: str,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    integrity_path = output_dir / "review_integrity.json"
    integrity = json.loads(integrity_path.read_text(encoding="utf-8"))
    if attack_kind == "changed":
        integrity["implementation_receipts"][0]["sha256"] = "0" * 64
    elif attack_kind == "missing":
        integrity["implementation_receipts"].pop()
        integrity["implementation_artifact_refs"].pop()
    elif attack_kind == "expanded":
        receipt = copy.deepcopy(integrity["implementation_receipts"][0])
        receipt["artifact_id"] = "implementation.report_builder.rogue.py"
        receipt["path"] = "rogue.py"
        integrity["implementation_receipts"].append(receipt)
        integrity["implementation_artifact_refs"].append(receipt["artifact_id"])
    else:
        integrity["implementation_receipts"] = list(
            reversed(integrity["implementation_receipts"])
        )
    content = dict(integrity)
    content.pop("content_sha256")
    integrity["content_sha256"] = hashlib.sha256(
        json.dumps(
            content,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            allow_nan=False,
        ).encode("utf-8")
    ).hexdigest()
    integrity_path.write_text(
        json.dumps(integrity, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": arguments,
        },
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == "Report Builder persisted review authorization failed."


@pytest.mark.parametrize(
    "relative_path",
    [
        "scripts/report_builder_core.py",
        "mcp/server.cjs",
        "assets/report-builder-review-widget.html",
        "shared:serialization.py",
    ],
)
def test_python_transitive_contract_rejects_copied_tree_mutation(
    tmp_path: Path,
    relative_path: str,
) -> None:
    # Arrange
    load_core()
    contract = sys.modules["implementation_contract"]
    references = [
        receipt["artifact_id"] for receipt in contract.build_implementation_receipts()
    ]
    receipts = contract.build_implementation_receipts()
    plugin_copy = tmp_path / "report-builder"
    shared_copy = tmp_path / "vera_assurance"
    shutil.copytree(
        PLUGIN_ROOT,
        plugin_copy,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    shutil.copytree(
        contract.SHARED_ROOT,
        shared_copy,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    if relative_path.startswith("shared:"):
        target = shared_copy / relative_path.removeprefix("shared:")
    else:
        target = plugin_copy / relative_path
    target.write_bytes(target.read_bytes() + b"\n# copied-tree mutation\n")

    # Act / Assert
    with pytest.raises(ValueError, match="does not match current bytes"):
        contract.validate_implementation_contract(
            references,
            receipts,
            artifact_roots={
                "implementation": plugin_copy,
                "assurance_implementation": shared_copy,
            },
        )


def _copy_report_builder_runtime(
    tmp_path: Path,
    name: str,
) -> tuple[Path, Path]:
    runtime_root = tmp_path / name
    plugin_copy = runtime_root / "report-builder"
    shared_copy = runtime_root / "_shared" / "vendor" / "modules" / "vera_assurance"
    shutil.copytree(
        PLUGIN_ROOT,
        plugin_copy,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    load_core()
    contract = sys.modules["implementation_contract"]
    shutil.copytree(
        contract.SHARED_ROOT,
        shared_copy,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    return plugin_copy, shared_copy


@pytest.mark.parametrize(
    "attack_kind",
    ["empty_directory", "regular", "symlink", "hardlink", "fifo"],
)
def test_python_transitive_contract_rejects_every_unowned_entry(
    tmp_path: Path,
    attack_kind: str,
) -> None:
    # Arrange
    load_core()
    contract = sys.modules["implementation_contract"]
    references = [
        receipt["artifact_id"] for receipt in contract.build_implementation_receipts()
    ]
    receipts = contract.build_implementation_receipts()
    plugin_copy, shared_copy = _copy_report_builder_runtime(
        tmp_path,
        attack_kind,
    )
    rogue = plugin_copy / "scripts" / "rogue"
    external = tmp_path / f"{attack_kind}-external"
    if attack_kind == "empty_directory":
        rogue.mkdir()
    elif attack_kind == "regular":
        rogue.write_bytes(b"unreceipted implementation")
    elif attack_kind == "symlink":
        external.write_bytes(b"external")
        rogue.symlink_to(external)
    elif attack_kind == "hardlink":
        external.write_bytes(b"external")
        rogue.hardlink_to(external)
    else:
        os.mkfifo(rogue)

    # Act / Assert
    with pytest.raises(RuntimeError, match="implementation"):
        contract.validate_implementation_contract(
            references,
            receipts,
            artifact_roots={
                "implementation": plugin_copy,
                "assurance_implementation": shared_copy,
            },
        )


def test_real_python_entry_rejects_timestamp_valid_unreceipted_bytecode(
    tmp_path: Path,
) -> None:
    # Arrange
    plugin_copy, _ = _copy_report_builder_runtime(tmp_path, "malicious-pyc")
    target = plugin_copy / "scripts" / "implementation_contract.py"
    source = target.read_bytes()
    source_stat = target.stat()
    marker = tmp_path / "malicious-pyc-executed.txt"
    malicious = (
        "from pathlib import Path as _AttackPath\n"
        f"_AttackPath({marker.as_posix()!r}).write_text("
        "'executed before validation\\n', encoding='utf-8')\n"
        f"exec(compile({source.decode('utf-8')!r}, {target.as_posix()!r}, "
        "'exec'), globals())\n"
    )
    code = compile(malicious, target.as_posix(), "exec")
    cache_prefix = sys.pycache_prefix
    try:
        sys.pycache_prefix = None
        cache_path = Path(importlib.util.cache_from_source(target.as_posix()))
    finally:
        sys.pycache_prefix = cache_prefix
    cache_path.parent.mkdir()
    cache_path.write_bytes(
        importlib._bootstrap_external._code_to_timestamp_pyc(
            code,
            int(source_stat.st_mtime),
            source_stat.st_size,
        )
    )

    # Act
    completed = subprocess.run(
        [
            sys.executable,
            "-I",
            "-B",
            str(plugin_copy / "scripts" / "build_report.py"),
            "--help",
        ],
        cwd=plugin_copy,
        capture_output=True,
        text=True,
        check=False,
        timeout=30,
    )

    # Assert
    assert completed.returncode != 0
    assert "implementation" in completed.stderr.lower()
    assert not marker.exists()
    assert target.read_bytes() == source
    assert target.stat().st_size == source_stat.st_size
    assert target.stat().st_mtime_ns == source_stat.st_mtime_ns


@pytest.mark.parametrize("attack_kind", ["symlink", "hardlink", "fifo"])
def test_real_python_entry_rejects_unsafe_bootstrap_before_read(
    tmp_path: Path,
    attack_kind: str,
) -> None:
    # Arrange
    plugin_copy, _ = _copy_report_builder_runtime(
        tmp_path,
        f"bootstrap-{attack_kind}",
    )
    bootstrap = plugin_copy / "scripts" / "implementation_bootstrap.py"
    original = bootstrap.read_bytes()
    bootstrap.unlink()
    external = tmp_path / f"bootstrap-{attack_kind}-external.py"
    if attack_kind == "symlink":
        external.write_bytes(original)
        bootstrap.symlink_to(external)
    elif attack_kind == "hardlink":
        external.write_bytes(original)
        bootstrap.hardlink_to(external)
    else:
        os.mkfifo(bootstrap)

    # Act
    completed = subprocess.run(
        [
            sys.executable,
            "-I",
            "-B",
            str(plugin_copy / "scripts" / "build_report.py"),
            "--help",
        ],
        cwd=plugin_copy,
        capture_output=True,
        text=True,
        check=False,
        timeout=30,
    )

    # Assert
    assert completed.returncode != 0
    assert "bootstrap is not a real file" in completed.stderr


def test_mcp_rejects_unowned_implementation_path_before_local_require(
    tmp_path: Path,
) -> None:
    # Arrange
    plugin_copy, _ = _copy_report_builder_runtime(tmp_path, "mcp-rogue")
    (plugin_copy / "scripts" / "__pycache__").mkdir()
    node = shutil.which("node")
    if node is None:
        pytest.skip("node is required for MCP server checks")

    # Act
    completed = subprocess.run(
        [node, str(plugin_copy / "mcp" / "server.cjs")],
        input=json.dumps(
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/list",
                "params": {},
            }
        )
        + "\n",
        capture_output=True,
        text=True,
        check=False,
        timeout=30,
    )

    # Assert
    assert completed.returncode != 0
    assert "implementation" in completed.stderr.lower()


@pytest.mark.parametrize(
    "artifact_name",
    [
        "report_analysis.json",
        "report_tables.json",
        "report_draft.md",
        "report.docx",
    ],
)
def test_prepared_replay_rejects_self_consistently_resealed_forgery(
    tmp_path: Path,
    artifact_name: str,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    artifact = output_dir / artifact_name
    if artifact_name == "report_analysis.json":
        payload = json.loads(artifact.read_text(encoding="utf-8"))
        payload["entity"] = "Forged entity"
        artifact.write_text(json.dumps(payload, indent=2) + "\n", encoding="utf-8")
    elif artifact_name == "report_tables.json":
        payload = json.loads(artifact.read_text(encoding="utf-8"))
        payload["tables"][0]["row_count"] += 1
        artifact.write_text(json.dumps(payload, indent=2) + "\n", encoding="utf-8")
    elif artifact_name == "report_draft.md":
        artifact.write_text(
            artifact.read_text(encoding="utf-8") + "\nForged conclusion.\n",
            encoding="utf-8",
        )
    else:
        artifact.write_bytes(artifact.read_bytes() + b"forged-docx")
    _reseal_report_builder_self_consistent(output_dir)

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": arguments,
        },
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == "Report Builder persisted review authorization failed."


def test_numeric_replay_rejects_self_consistently_resealed_ledger_forgery(
    tmp_path: Path,
) -> None:
    # Arrange
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    _save_workbook(input_path)
    managed = _managed_report_run(
        tmp_path,
        input_path,
        recipe_builder=lambda managed_core, source, work_dir: (
            _reviewed_numeric_recipe(managed_core, source, work_dir)
        ),
    )
    output_dir = managed["output_dir"]
    ledger_path = output_dir / "numeric_evidence_ledger.json"
    ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
    ledger["entries"][0]["value"] = "999999"
    if "content_sha256" in ledger:
        content = dict(ledger)
        content.pop("content_sha256")
        ledger["content_sha256"] = hashlib.sha256(
            json.dumps(
                content,
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
                allow_nan=False,
            ).encode("utf-8")
        ).hexdigest()
    ledger_path.write_text(
        json.dumps(ledger, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    _reseal_report_builder_self_consistent(output_dir)
    arguments = {
        "client_engagement": managed["context_path"].as_posix(),
        "run_intake": json.loads(
            (output_dir / "run_intake.json").read_text(encoding="utf-8")
        ),
        "review_payload": json.loads(
            (output_dir / "review_payload.json").read_text(encoding="utf-8")
        ),
        "final_artifacts": json.loads(
            (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
        ),
    }

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": arguments,
        },
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == "Report Builder persisted review authorization failed."


@pytest.mark.parametrize("attack_kind", ["applied_decision", "application_effect"])
def test_successor_replay_rejects_self_consistently_resealed_divergence(
    tmp_path: Path,
    attack_kind: str,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    applied_result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": arguments,
        },
    )["structuredContent"]
    assert applied_result["ok"] is True
    applied_path = output_dir / "applied_decisions.json"
    applied = json.loads(applied_path.read_text(encoding="utf-8"))
    if attack_kind == "applied_decision":
        applied["decisions"][0]["edit_value"] = "Forged successor decision."
    else:
        applied["effects"][0]["edit_value"] = "Forged successor effect."
    applied_path.write_text(
        json.dumps(applied, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    _reseal_report_builder_self_consistent(output_dir)

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "client_engagement": arguments["client_engagement"],
                "run_intake": json.loads(
                    (output_dir / "run_intake.json").read_text(encoding="utf-8")
                ),
                "review_payload": json.loads(
                    (output_dir / "review_payload.json").read_text(encoding="utf-8")
                ),
                "final_artifacts": json.loads(
                    (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
                ),
            },
        },
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == "Report Builder persisted review authorization failed."


def test_source_mapping_successor_accepts_a_new_review_round(
    tmp_path: Path,
) -> None:
    # Arrange
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    _save_workbook(input_path)
    managed = _managed_report_run(tmp_path, input_path)
    output_dir = managed["output_dir"]
    first_review = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    table_item = next(
        item
        for item in first_review["items"]
        if item["item_type"] == "table_evidence"
        and item["data"]["section"] == "income_statement"
    )
    first_result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": json.loads(
                    (output_dir / "run_intake.json").read_text(encoding="utf-8")
                ),
                "review_payload": first_review,
                "final_artifacts": json.loads(
                    (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
                ),
                "decisions": [
                    {
                        "item_id": table_item["id"],
                        "action": "edit",
                        "edit_value": "report.xlsx::Cash Flow",
                    }
                ],
            },
        },
    )["structuredContent"]
    assert first_result["ok"] is True
    predecessor_checkpoint = first_result["integrity_checkpoint"]
    second_review = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )

    # Act
    second_result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": json.loads(
                    (output_dir / "run_intake.json").read_text(encoding="utf-8")
                ),
                "review_payload": second_review,
                "final_artifacts": json.loads(
                    (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
                ),
                "expected_predecessor_checkpoint": predecessor_checkpoint,
                "decisions": _nonblocking_decisions(second_review),
            },
        },
    )["structuredContent"]
    validation = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": json.loads(
                    (output_dir / "run_intake.json").read_text(encoding="utf-8")
                ),
                "review_payload": json.loads(
                    (output_dir / "review_payload.json").read_text(encoding="utf-8")
                ),
                "final_artifacts": json.loads(
                    (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
                ),
                "expected_predecessor_checkpoint": predecessor_checkpoint,
            },
        },
    )["structuredContent"]

    # Assert
    assert second_result["ok"] is True, second_result
    assert validation["ok"] is True


@pytest.mark.parametrize(
    "checkpoint_value",
    [None, "0" * 64],
)
def test_later_review_requires_exact_external_predecessor_checkpoint_without_mutation(
    tmp_path: Path,
    checkpoint_value: str | None,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    first = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": arguments,
        },
    )["structuredContent"]
    assert first["ok"] is True
    current_review = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    current_item = next(
        item
        for item in current_review["items"]
        if item["id"] == arguments["decisions"][0]["item_id"]
    )
    later_arguments = {
        "run_intake": json.loads(
            (output_dir / "run_intake.json").read_text(encoding="utf-8")
        ),
        "review_payload": current_review,
        "final_artifacts": json.loads(
            (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
        ),
        "decisions": [
            {
                "item_id": current_item["id"],
                "action": "edit",
                "edit_value": "Later externally anchored review.",
            }
        ],
    }
    if checkpoint_value is not None:
        later_arguments["expected_predecessor_checkpoint"] = checkpoint_value
    before = _transaction_tree_state(output_dir)

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": later_arguments,
        },
    )

    # Assert
    failure = response["result"]["structuredContent"]
    assert failure["ok"] is False
    assert _transaction_tree_state(output_dir) == before


def test_successor_validation_requires_retained_predecessor_checkpoint(
    tmp_path: Path,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    first = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": arguments,
        },
    )["structuredContent"]
    predecessor_checkpoint = first["integrity_checkpoint"]
    current_review = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    second_arguments = {
        "client_engagement": arguments["client_engagement"],
        "run_intake": json.loads(
            (output_dir / "run_intake.json").read_text(encoding="utf-8")
        ),
        "review_payload": current_review,
        "final_artifacts": json.loads(
            (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
        ),
        "decisions": [
            {
                "item_id": arguments["decisions"][0]["item_id"],
                "action": "edit",
                "edit_value": "Checkpoint-backed successor.",
            }
        ],
        "expected_predecessor_checkpoint": predecessor_checkpoint,
    }
    second = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": second_arguments,
        },
    )["structuredContent"]
    assert second["ok"] is True
    validation_arguments = {
        "client_engagement": arguments["client_engagement"],
        "run_intake": json.loads(
            (output_dir / "run_intake.json").read_text(encoding="utf-8")
        ),
        "review_payload": json.loads(
            (output_dir / "review_payload.json").read_text(encoding="utf-8")
        ),
        "final_artifacts": json.loads(
            (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
        ),
    }

    # Act
    missing = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": validation_arguments,
        },
    )["structuredContent"]
    wrong = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                **validation_arguments,
                "expected_predecessor_checkpoint": "0" * 64,
            },
        },
    )["structuredContent"]
    accepted = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                **validation_arguments,
                "expected_predecessor_checkpoint": predecessor_checkpoint,
            },
        },
    )["structuredContent"]
    validator = SCRIPT_DIR / "validate_review_integrity.py"
    missing_cli = subprocess.run(
        [
            sys.executable,
            "-I",
            "-B",
            str(validator),
            "--output-dir",
            str(output_dir),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    accepted_cli = subprocess.run(
        [
            sys.executable,
            "-I",
            "-B",
            str(validator),
            "--output-dir",
            str(output_dir),
            "--expected-predecessor-checkpoint",
            predecessor_checkpoint,
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    # Assert
    assert missing["ok"] is False
    assert wrong["ok"] is False
    assert accepted["ok"] is True
    assert missing_cli.returncode != 0
    assert accepted_cli.returncode == 0, accepted_cli.stderr
    assert json.loads(accepted_cli.stdout)["ok"] is True


def test_alternative_honest_predecessor_cannot_replace_retained_checkpoint(
    tmp_path: Path,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    baseline = tmp_path / "baseline"
    shutil.copytree(output_dir, baseline)

    alternative_arguments = copy.deepcopy(arguments)
    alternative_arguments["decisions"][0][
        "edit_value"
    ] = "Alternative honest predecessor."
    alternative = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": alternative_arguments,
        },
    )["structuredContent"]
    assert alternative["ok"] is True
    alternative_state = {
        name: json.loads((output_dir / f"{name}.json").read_text(encoding="utf-8"))
        for name in (
            "run_intake",
            "review_payload",
            "ui_decisions",
            "applied_decisions",
            "final_artifacts",
            "review_integrity",
        )
    }
    alternative_output = tmp_path / "alternative"
    output_dir.rename(alternative_output)
    shutil.copytree(baseline, output_dir)

    genuine_arguments = copy.deepcopy(arguments)
    genuine_arguments["decisions"][0]["edit_value"] = "Genuine predecessor."
    genuine = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": genuine_arguments,
        },
    )["structuredContent"]
    genuine_checkpoint = genuine["integrity_checkpoint"]
    current_arguments = {
        "client_engagement": arguments["client_engagement"],
        "run_intake": json.loads(
            (output_dir / "run_intake.json").read_text(encoding="utf-8")
        ),
        "review_payload": json.loads(
            (output_dir / "review_payload.json").read_text(encoding="utf-8")
        ),
        "final_artifacts": json.loads(
            (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
        ),
        "decisions": [
            {
                "item_id": arguments["decisions"][0]["item_id"],
                "action": "edit",
                "edit_value": "Current successor.",
            }
        ],
        "expected_predecessor_checkpoint": genuine_checkpoint,
    }
    current = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": current_arguments,
        },
    )["structuredContent"]
    assert current["ok"] is True
    applied_path = output_dir / "applied_decisions.json"
    final_path = output_dir / "final_artifacts.json"
    applied = json.loads(applied_path.read_text(encoding="utf-8"))
    final_artifacts = json.loads(final_path.read_text(encoding="utf-8"))
    genuine_history_path = output_dir / applied["review_history_paths"][0]
    genuine_history = json.loads(genuine_history_path.read_text(encoding="utf-8"))
    genuine_history_path.rename(tmp_path / "genuine-history.json")
    forged_content = {
        "schema_version": "report_builder.review_history_entry.v2",
        "archived_at": genuine_history["archived_at"],
        "predecessor_checkpoint": alternative["integrity_checkpoint"],
        "predecessor_integrity": alternative_state["review_integrity"],
        "run_intake": alternative_state["run_intake"],
        "review_payload": alternative_state["review_payload"],
        "ui_decisions": alternative_state["ui_decisions"],
        "applied_decisions": alternative_state["applied_decisions"],
        "final_artifacts": alternative_state["final_artifacts"],
    }
    forged_digest = _canonical_json_sha256(forged_content)
    forged_relative_path = f"revisions/history/application__{forged_digest}.json"
    (output_dir / forged_relative_path).write_text(
        json.dumps(
            {
                **forged_content,
                "content_sha256": forged_digest,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    applied["predecessor_checkpoint"] = alternative["integrity_checkpoint"]
    applied["review_history_paths"] = [forged_relative_path]
    final_artifacts["review_application"]["predecessor_checkpoint"] = alternative[
        "integrity_checkpoint"
    ]
    final_artifacts["review_application"]["review_history_paths"] = [
        forged_relative_path
    ]
    applied_path.write_text(
        json.dumps(applied, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    final_path.write_text(
        json.dumps(final_artifacts, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    original_integrity = (output_dir / "review_integrity.json").read_bytes()
    core = load_core()

    # Act / Assert
    with pytest.raises(
        ValueError,
        match="predecessor checkpoint",
    ):
        core.seal_review_integrity(
            output_dir,
            run_id=arguments["run_intake"]["run_id"],
            expected_predecessor_checkpoint=genuine_checkpoint,
        )
    assert (output_dir / "review_integrity.json").read_bytes() == original_integrity
    mcp_validation = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "client_engagement": arguments["client_engagement"],
                "run_intake": json.loads(
                    (output_dir / "run_intake.json").read_text(encoding="utf-8")
                ),
                "review_payload": json.loads(
                    (output_dir / "review_payload.json").read_text(encoding="utf-8")
                ),
                "final_artifacts": final_artifacts,
                "expected_predecessor_checkpoint": genuine_checkpoint,
            },
        },
    )["structuredContent"]
    assert mcp_validation["ok"] is False


def test_complete_text_only_review_is_report_ready_but_not_published(
    tmp_path: Path,
) -> None:
    # Arrange
    core = load_core()
    input_path = tmp_path / "notes.csv"
    input_path.write_text(
        "topic,note\nOperations,Stable\nCustomers,Diversified\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    recipe = inspection.suggested_recipe
    for section in recipe["sections"].values():
        section["assigned_table"] = input_path.name
        section["codex_comment"] = "Reviewer-supported qualitative comment."
    recipe_path = tmp_path / "recipe.json"
    core.write_json(recipe_path, recipe)
    managed = _managed_report_run(tmp_path, input_path, recipe=recipe_path)
    output_dir = managed["output_dir"]
    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    initial_final = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )

    # Act
    result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": json.loads(
                    (output_dir / "run_intake.json").read_text(encoding="utf-8")
                ),
                "review_payload": review_payload,
                "final_artifacts": initial_final,
                "decisions": _nonblocking_decisions(review_payload),
            },
        },
    )["structuredContent"]
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )

    # Assert
    assert result["application_status"] == "final_ready"
    assert final_artifacts["report_ready"] is True
    assert final_artifacts["assurance"]["gates"]["source"]["status"] == "passed"
    assert final_artifacts["assurance"]["gates"]["preparation"]["status"] == "passed"
    assert (
        final_artifacts["assurance"]["gates"]["reconciliation"]["status"]
        == "not_applicable"
    )
    assert (
        final_artifacts["assurance"]["gates"]["semantic_review"]["status"] == "passed"
    )
    assert final_artifacts["assurance"]["gates"]["reporting"]["status"] == "passed"
    assert final_artifacts["assurance"]["gates"]["publication"]["status"] == "withheld"


def _install_report_builder_fake_authority_child(
    monkeypatch: Any,
    tmp_path: Path,
) -> Path:
    marker = tmp_path / "fake-report-builder-child-invoked"
    fake_python = tmp_path / "fake-report-builder-authority-python"
    fake_python.write_text(
        "\n".join(
            [
                "#!/usr/bin/env python3",
                "import json",
                "import os",
                "import sys",
                "from pathlib import Path",
                'if "-c" in sys.argv:',
                "    os.execv(sys.executable, [sys.executable, *sys.argv[1:]])",
                'Path(os.environ["RB_FAKE_CHILD_MARKER"]).write_text("invoked\\n")',
                'print(json.dumps({"ok": True}))',
                "",
            ]
        ),
        encoding="utf-8",
    )
    fake_python.chmod(0o700)
    monkeypatch.setenv("PYTHON", fake_python.as_posix())
    monkeypatch.setenv("RB_FAKE_CHILD_MARKER", marker.as_posix())
    return marker


def _report_source_mapping_transaction_fixture(
    tmp_path: Path,
) -> tuple[Path, dict[str, Any]]:
    core = load_core()
    source = tmp_path / "source-mapping.xlsx"
    _save_workbook(source)
    ledger = _load_customer_ledger()
    client_root = tmp_path / "Customer"
    client_root.mkdir()
    client_id = "client_777777777777777777777777"
    ledger.create_client_manifest(client_root, client_id)
    engagement = ledger.create_engagement(client_root, client_id, "Report mapping")
    imported = ledger.import_document(
        client_root,
        client_id,
        engagement["engagement_id"],
        source,
        "source",
    )
    prepared = ledger.prepare_run(
        client_root,
        client_id,
        engagement["engagement_id"],
        "report-builder",
        "test-version",
        input_ids=[imported["receipt"]["input_id"]],
    )
    running = ledger.start_run(
        client_root,
        engagement["engagement_id"],
        prepared["run"]["run_id"],
    )
    input_path = Path(running["context"]["input_bindings"][0]["path"])
    output_dir = Path(running["output_dir"])
    run_id = str(running["context"]["run_id"])
    core.build_report(
        input_path,
        output_dir,
        report_type="management_report",
        run_id=run_id,
        client_engagement=running["context"],
    )
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    assert run_intake["run_id"] == run_id
    assert review_payload["run_id"] == run_id
    table_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "table_evidence"
        and item["data"]["section"] == "income_statement"
    )
    return output_dir, {
        "client_engagement": Path(running["context_path"]).as_posix(),
        "run_intake": run_intake,
        "review_payload": review_payload,
        "final_artifacts": final_artifacts,
        "decisions": [
            {
                "item_id": table_item["id"],
                "action": "edit",
                "edit_value": "source-mapping.xlsx::Cash Flow",
            }
        ],
    }


def _reviewed_numeric_recipe(
    core: Any,
    input_path: Path,
    work_dir: Path,
    *,
    language: str = "en",
    document_language: str = "en",
    report_type: str = "management_report",
) -> Path:
    inspection = core.inspect_inputs(
        input_path,
        work_dir,
        language=language,
        document_language=document_language,
        report_type=report_type,
    )
    recipe = inspection.suggested_recipe
    tables_by_id = {
        table["table_id"]: table for table in inspection.inspection["tables"]
    }
    for section_key, section in recipe["sections"].items():
        table_id = section["assigned_table"]
        if not table_id:
            continue
        columns = [
            column["column"] for column in tables_by_id[table_id]["numeric_columns"]
        ]
        if not columns:
            continue
        recipe = core.review_numeric_measure_columns(
            inspection.inspection,
            recipe,
            section_key=section_key,
            **_numeric_review_args(
                inspection.inspection,
                table_id,
                columns,
            ),
            reviewer_ref="reviewer.pytest",
            reviewed_on="2026-07-24",
            numeric_locale="en",
            currency="EUR",
            unit="currency",
            scale="1",
            parse_policy="strict_all_nonblank_v1",
        )
    recipe_path = work_dir / "reviewed_recipe.json"
    core.write_json(recipe_path, recipe)
    return recipe_path


def _tamper_report_source(
    input_path: Path,
    output_dir: Path,
    arguments: dict[str, Any],
) -> dict[str, Any]:
    del output_dir
    input_path.write_text(
        "line,amount\nA,999\nB,20\n",
        encoding="utf-8",
    )
    return arguments


def _tamper_report_output(
    input_path: Path,
    output_dir: Path,
    arguments: dict[str, Any],
) -> dict[str, Any]:
    del input_path
    with (output_dir / "report_draft.md").open("a", encoding="utf-8") as handle:
        handle.write("\nforged output\n")
    return arguments


def _tamper_persisted_review_payload(
    input_path: Path,
    output_dir: Path,
    arguments: dict[str, Any],
) -> dict[str, Any]:
    del input_path
    persisted = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    persisted["forged"] = True
    (output_dir / "review_payload.json").write_text(
        json.dumps(persisted, indent=2) + "\n",
        encoding="utf-8",
    )
    return arguments


def _supply_stale_review_payload(
    input_path: Path,
    output_dir: Path,
    arguments: dict[str, Any],
) -> dict[str, Any]:
    del input_path, output_dir
    stale_arguments = copy.deepcopy(arguments)
    stale_arguments["review_payload"]["items"][0]["title"] = "Stale review item"
    return stale_arguments


def _nonblocking_decisions(review_payload: dict[str, Any]) -> list[dict[str, str]]:
    decisions = []
    for item in review_payload["items"]:
        allowed = item["allowed_actions"]
        action = "accept" if "accept" in allowed else "skip"
        decisions.append({"item_id": item["id"], "action": action})
    return decisions


def test_build_seals_exact_numeric_ledger_and_rejects_changed_rendered_cell(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    output_dir = tmp_path / "out"
    _save_workbook(input_path)
    recipe_path = _reviewed_numeric_recipe(
        core,
        input_path,
        tmp_path / "inspection",
    )

    core.build_report(input_path, output_dir, recipe_path=recipe_path)

    ledger_path = output_dir / "numeric_evidence_ledger.json"
    ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
    analysis = json.loads(
        (output_dir / "report_analysis.json").read_text(encoding="utf-8")
    )
    assert ledger["schema_version"] == "vera.numeric_evidence_ledger.v1"
    assert ledger["entries"]
    assert {
        output["artifact_ref"]
        for entry in ledger["entries"]
        for output in entry["outputs"]
    } == {
        "output.report_docx",
        "output.report_draft",
        "output.report_tables",
    }
    assert all(
        entry["source"]["value"] == entry["value"] for entry in ledger["entries"]
    )
    source_receipts = json.loads(
        (output_dir / "source_receipts.json").read_text(encoding="utf-8")
    )
    assert source_receipts["schema_version"] == "report_builder.source_receipts.v1"
    assert {source["artifact_id"] for source in source_receipts["sources"]} == {
        entry["source"]["artifact_ref"] for entry in ledger["entries"]
    }
    assert not any(
        isinstance(column[metric], float)
        for section in analysis["sections"]
        for column in section["numeric_columns"]
        for metric in ("sum", "min", "max")
    )

    workbook_path = output_dir / "report_tables.xlsx"
    workbook = openpyxl.load_workbook(workbook_path)
    workbook["numeric_evidence"]["D2"] = "999"
    workbook.save(workbook_path)

    with pytest.raises(ValueError, match="Excel evidence"):
        core.write_numeric_evidence_ledger(output_dir, analysis)


def test_numeric_ledger_replay_rejects_changed_source_bytes(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    output_dir = tmp_path / "out"
    _save_workbook(input_path)
    recipe_path = _reviewed_numeric_recipe(
        core,
        input_path,
        tmp_path / "inspection",
    )
    core.build_report(input_path, output_dir, recipe_path=recipe_path)
    analysis = json.loads(
        (output_dir / "report_analysis.json").read_text(encoding="utf-8")
    )

    workbook = openpyxl.load_workbook(input_path)
    workbook["Income Statement"]["B2"] = 999
    workbook.save(input_path)

    with pytest.raises(ValueError, match="does not match current bytes"):
        core.write_numeric_evidence_ledger(output_dir, analysis)


def test_source_capture_rejects_swapped_bytes_during_read(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    mutant_path = tmp_path / "mutant.xlsx"
    _save_workbook(input_path)
    _save_workbook(mutant_path)
    mutant = openpyxl.load_workbook(mutant_path)
    mutant["Income Statement"]["B2"] = 999
    mutant.save(mutant_path)
    mutant_bytes = mutant_path.read_bytes()
    original_read_bytes = Path.read_bytes

    def swapped_read_bytes(path: Path) -> bytes:
        if path.resolve() == input_path.resolve():
            return mutant_bytes
        return original_read_bytes(path)

    monkeypatch.setattr(Path, "read_bytes", swapped_read_bytes)

    tables = core.load_tables(input_path, tmp_path / "out")

    assert len(tables) == 1
    assert tables[0]["kind"] == "error"
    assert "changed while it was captured" in tables[0]["error"]


def test_duplicate_basenames_receive_distinct_source_bound_table_ids(
    tmp_path: Path,
) -> None:
    core = load_core()
    first = tmp_path / "first" / "report.xlsx"
    second = tmp_path / "second" / "report.xlsx"
    first.parent.mkdir()
    second.parent.mkdir()
    _save_workbook(first)
    _save_workbook(second)

    tables = core.load_tables(tmp_path, tmp_path / "out")

    table_ids = [table["table_id"] for table in tables]
    source_refs = {table["source_artifact_ref"] for table in tables}
    assert len(table_ids) == len(set(table_ids))
    assert all("@" in table_id for table_id in table_ids)
    assert len(source_refs) == 2


def test_unreviewed_numeric_identifier_columns_are_not_reported(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text(
        "expense,amount,account_id\nTravel,10.00,1001\nSupplies,20.00,1002\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = input_path.name
    recipe_path = tmp_path / "recipe.json"
    core.write_json(recipe_path, recipe)

    core.build_report(
        input_path,
        tmp_path / "report",
        recipe_path=recipe_path,
    )

    analysis = json.loads(
        (tmp_path / "report" / "report_analysis.json").read_text(encoding="utf-8")
    )
    budget = next(
        section for section in analysis["sections"] if section["section"] == "budget"
    )
    assert budget["numeric_measure_status"] == "needs_review"
    assert budget["numeric_columns"] == []
    assert {column["column"] for column in budget["numeric_measure_candidates"]} == {
        "amount",
        "account_id",
    }
    assert not (tmp_path / "report" / "numeric_evidence_ledger.json").exists()


def test_reviewed_measure_mapping_excludes_numeric_identifiers(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text(
        "expense,amount,account_id\nTravel,10.00,1001\nSupplies,20.00,1002\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = input_path.name
    recipe = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        **_numeric_review_args(
            inspection.inspection,
            input_path.name,
            ["amount"],
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency="EUR",
        unit="currency",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )
    recipe_path = tmp_path / "reviewed_recipe.json"
    core.write_json(recipe_path, recipe)

    core.build_report(
        input_path,
        tmp_path / "report",
        recipe_path=recipe_path,
    )

    ledger = json.loads(
        (tmp_path / "report" / "numeric_evidence_ledger.json").read_text(
            encoding="utf-8"
        )
    )
    assert len(ledger["entries"]) == 1
    assert ledger["entries"][0]["value"] == "30"
    assert "account_id" not in ledger["entries"][0]["source"]["locator"]


def test_report_period_change_invalidates_reviewed_numeric_measure(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "income_statement.csv"
    input_path.write_text(
        "Period,Revenue\n2024,10\n2025,20\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="annual_financial_statement",
    )
    recipe = inspection.suggested_recipe
    for section in recipe["sections"].values():
        section["assigned_table"] = ""
    recipe["period"] = "2025"
    recipe["sections"]["income_statement"]["assigned_table"] = input_path.name
    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="income_statement",
        **_numeric_review_args(
            inspection.inspection,
            input_path.name,
            ["Revenue"],
            excluded_cell_rows={"Revenue": {2}},
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency="USD",
        unit="currency",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )
    reviewed_path = tmp_path / "reviewed-recipe.json"
    core.write_json(reviewed_path, reviewed)
    core.build_report(
        input_path,
        tmp_path / "reviewed-report",
        recipe_path=reviewed_path,
    )
    assert (tmp_path / "reviewed-report" / "numeric_evidence_ledger.json").is_file()

    reviewed["period"] = "2024"
    stale_path = tmp_path / "stale-period-recipe.json"
    core.write_json(stale_path, reviewed)
    core.build_report(
        input_path,
        tmp_path / "stale-period-report",
        recipe_path=stale_path,
    )
    analysis = json.loads(
        (tmp_path / "stale-period-report" / "report_analysis.json").read_text(
            encoding="utf-8"
        )
    )
    income = next(
        section
        for section in analysis["sections"]
        if section["section"] == "income_statement"
    )

    assert analysis["period"] == "2024"
    assert income["numeric_measure_status"] == "needs_review"
    assert income["numeric_columns"] == []
    assert not (
        tmp_path / "stale-period-report" / "numeric_evidence_ledger.json"
    ).exists()


def test_numeric_decision_binds_locale_semantics_cells_and_ledger_reference(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text(
        'line;amount\nA;"EUR 1.234,50"\nB;"EUR 765,50"\n',
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = input_path.name

    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        **_numeric_review_args(
            inspection.inspection,
            input_path.name,
            ["amount"],
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="it",
        currency="EUR",
        unit="currency",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )
    decision = reviewed["sections"]["budget"]["numeric_measure_decision"]
    content = decision["content"]
    dispositions = content["column_dispositions"][0]["cells"]

    assert decision["decision_type"] == "numeric_measure_mapping"
    assert decision["decision_id"].startswith("decision.report_numeric_measures.")
    assert content["numeric_contract"] == {
        "policy_id": "strict_all_nonblank_v1",
        "locale": "it",
        "decimal_separator": ",",
        "thousands_separator": ".",
        "currency": "EUR",
        "unit": "currency",
        "scale": "1",
        "sign_policy": "as_presented_v1",
    }
    assert dispositions == [
        {
            "row": 2,
            "coordinate": "B2",
            "source_text": "EUR 1.234,50",
            "number_format": "",
            "formula": None,
            "formula_cached_value": None,
            "formula_cache_status": "not_formula",
            "status": "included",
            "currency_marker": "EUR",
            "ambiguous_currency_symbols": [],
            "canonical_value": "1234.5",
            "signed_value": "1234.5",
            "scaled_value": "1234.5",
        },
        {
            "row": 3,
            "coordinate": "B3",
            "source_text": "EUR 765,50",
            "number_format": "",
            "formula": None,
            "formula_cached_value": None,
            "formula_cache_status": "not_formula",
            "status": "included",
            "currency_marker": "EUR",
            "ambiguous_currency_symbols": [],
            "canonical_value": "765.5",
            "signed_value": "765.5",
            "scaled_value": "765.5",
        },
    ]

    recipe_path = tmp_path / "reviewed_recipe.json"
    core.write_json(recipe_path, reviewed)
    output_dir = tmp_path / "report"
    core.build_report(input_path, output_dir, recipe_path=recipe_path)
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )

    assert ledger["entries"][0]["value"] == "2000"
    assert ledger["entries"][0]["unit"] == "currency"
    assert ledger["entries"][0]["currency"] == "EUR"
    assert ledger["entries"][0]["decision_ref"] == decision["decision_id"]


@pytest.mark.parametrize(
    "mutation",
    ("wrong_type", "wrong_id", "future_review", "wrong_disposition"),
)
def test_build_rejects_forged_numeric_decision_with_valid_content_digest(
    tmp_path: Path,
    mutation: str,
) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text(
        "line,amount\nA,10.00\nB,20.00\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = input_path.name
    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        **_numeric_review_args(
            inspection.inspection,
            input_path.name,
            ["amount"],
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency="EUR",
        unit="currency",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )
    decision = reviewed["sections"]["budget"]["numeric_measure_decision"]
    forged_content = copy.deepcopy(decision["content"])
    decision_id = decision["decision_id"]
    decision_type = decision["decision_type"]
    reviewed_on = decision["reviewed_on"]
    if mutation == "wrong_type":
        decision_type = "different_numeric_decision"
    elif mutation == "wrong_id":
        decision_id = "decision.report_numeric_measures.forged"
    elif mutation == "future_review":
        reviewed_on = "9999-01-01"
    else:
        forged_content["column_dispositions"][0]["cells"][0]["source_text"] = "999"
    forged = core.build_reviewed_decision_receipt(
        decision_id=decision_id,
        decision_type=decision_type,
        status=decision["status"],
        reviewer_ref=decision["reviewer_ref"],
        reviewed_on=reviewed_on,
        adapter_id=decision["adapter_id"],
        adapter_version=decision["adapter_version"],
        source_artifact_refs=decision["source_artifact_refs"],
        content=forged_content,
    )
    reviewed["sections"]["budget"]["numeric_measure_decision"] = forged
    recipe_path = tmp_path / "forged_recipe.json"
    core.write_json(recipe_path, reviewed)

    output_dir = tmp_path / "report"
    core.build_report(input_path, output_dir, recipe_path=recipe_path)
    analysis = json.loads(
        (output_dir / "report_analysis.json").read_text(encoding="utf-8")
    )
    budget = next(
        section for section in analysis["sections"] if section["section"] == "budget"
    )

    assert budget["numeric_measure_status"] == "needs_review"
    assert budget["numeric_columns"] == []
    assert budget["numeric_measure_decision"] is None
    assert not (output_dir / "numeric_evidence_ledger.json").exists()


@pytest.mark.parametrize(
    "csv_text",
    (
        "line,amount\nA,EUR 10\nB,USD 20\n",
        "line,amount\nA,EUR 10\nB,not-a-number\n",
    ),
)
def test_numeric_review_fails_closed_for_mixed_or_unresolved_cells(
    tmp_path: Path,
    csv_text: str,
) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text(csv_text, encoding="utf-8")
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = input_path.name

    with pytest.raises(ValueError, match="unresolved nonblank rows"):
        core.review_numeric_measure_columns(
            inspection.inspection,
            recipe,
            section_key="budget",
            **_numeric_review_args(
                inspection.inspection,
                input_path.name,
                ["amount"],
            ),
            reviewer_ref="reviewer.pytest",
            reviewed_on="2026-07-24",
            numeric_locale="en",
            currency="EUR",
            unit="currency",
            scale="1",
            parse_policy="strict_all_nonblank_v1",
        )


def test_percentage_measure_requires_reviewed_unit_and_scale(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "rates.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Budget"
    sheet.append(["line", "rate"])
    sheet.append(["A", 0.1])
    sheet.append(["B", 0.2])
    sheet["B2"].number_format = "0%"
    sheet["B3"].number_format = "0%"
    workbook.save(input_path)
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = "rates.xlsx::Budget"
    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        **_numeric_review_args(
            inspection.inspection,
            "rates.xlsx::Budget",
            ["rate"],
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency=None,
        unit="percentage",
        scale="100",
        parse_policy="strict_all_nonblank_v1",
    )
    recipe_path = tmp_path / "reviewed_recipe.json"
    core.write_json(recipe_path, reviewed)

    output_dir = tmp_path / "report"
    core.build_report(input_path, output_dir, recipe_path=recipe_path)
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )

    assert ledger["entries"][0]["value"] == "30"
    assert ledger["entries"][0]["unit"] == "percentage"
    assert ledger["entries"][0]["currency"] is None


@pytest.mark.parametrize(
    ("cached_value", "expected_cache_status"),
    (("999", "present"), ("30", "present"), (None, "missing")),
    ids=("cached-wrong", "cached-correct", "uncached"),
)
def test_formula_cells_fail_closed_even_when_workbook_cache_looks_usable(
    tmp_path: Path,
    cached_value: str | None,
    expected_cache_status: str,
) -> None:
    core = load_core()
    input_path = tmp_path / "formula.xlsx"
    _save_formula_workbook(input_path, cached_value)
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    formula_cell = next(
        cell
        for column in table["numeric_measure_cells"]
        for cell in column["nonblank_cells"]
        if cell["coordinate"] == "B4"
    )
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]

    with pytest.raises(ValueError, match="unresolved nonblank rows"):
        core.review_numeric_measure_columns(
            inspection.inspection,
            recipe,
            section_key="budget",
            **_numeric_review_args(
                inspection.inspection,
                table["table_id"],
                ["amount"],
            ),
            reviewer_ref="reviewer.pytest",
            reviewed_on="2026-07-24",
            numeric_locale="en",
            currency="EUR",
            unit="currency",
            scale="1",
            parse_policy="strict_all_nonblank_v1",
        )

    assert formula_cell["formula"] == "=SUM(B2:B3)"
    assert formula_cell["formula_cache_status"] == expected_cache_status
    assert formula_cell["formula_cached_value"] == cached_value


def test_explicit_subtotal_exclusion_prevents_double_counting(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "formula.xlsx"
    _save_formula_workbook(input_path, "30")
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]
    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        **_numeric_review_args(
            inspection.inspection,
            table["table_id"],
            ["amount"],
            excluded_cell_rows={"amount": {4}},
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency="EUR",
        unit="currency",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )
    recipe_path = tmp_path / "recipe.json"
    core.write_json(recipe_path, reviewed)
    output_dir = tmp_path / "report"

    core.build_report(input_path, output_dir, recipe_path=recipe_path)
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )
    decision_cells = reviewed["sections"]["budget"]["numeric_measure_decision"][
        "content"
    ]["column_dispositions"][0]["cells"]

    assert ledger["entries"][0]["value"] == "30"
    assert next(cell for cell in decision_cells if cell["row"] == 4)["status"] == (
        "excluded_by_review"
    )


def test_uncached_formula_only_column_remains_a_disposable_candidate(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "formula-only.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Budget"
    sheet.append(["line", "amount"])
    sheet.append(["Calculated total", "=SUM(10,20)"])
    workbook.save(input_path)

    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]
    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        **_numeric_review_args(
            inspection.inspection,
            table["table_id"],
            [],
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency=None,
        unit="number",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )

    assert table["numeric_columns"][0]["column"] == "amount"
    assert table["numeric_columns"][0]["formula_cell_count"] == 1
    assert reviewed["sections"]["budget"]["excluded_numeric_candidate_columns"] == [
        "amount"
    ]


def test_all_numeric_candidates_can_be_explicitly_excluded(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "identifiers.csv"
    input_path.write_text("name,account_id\nA,1001\nB,1002\n", encoding="utf-8")
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]
    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        **_numeric_review_args(
            inspection.inspection,
            table["table_id"],
            [],
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency=None,
        unit="count",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )
    recipe_path = tmp_path / "recipe.json"
    core.write_json(recipe_path, reviewed)
    output_dir = tmp_path / "report"

    core.build_report(input_path, output_dir, recipe_path=recipe_path)
    analysis = json.loads(
        (output_dir / "report_analysis.json").read_text(encoding="utf-8")
    )
    section = next(item for item in analysis["sections"] if item["section"] == "budget")

    assert section["numeric_measure_status"] == "reviewed"
    assert section["numeric_columns"] == []
    assert not (output_dir / "numeric_evidence_ledger.json").exists()


def test_numeric_review_rejects_undisposed_candidate_column_or_cell(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text(
        "line,amount,account_id\nA,10,1001\nB,20,1002\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]

    with pytest.raises(ValueError, match="Every numeric candidate column"):
        core.review_numeric_measure_columns(
            inspection.inspection,
            recipe,
            section_key="budget",
            header_row=table["header_row"],
            columns=["amount"],
            excluded_columns=[],
            cell_dispositions={"amount": {2: "include", 3: "include"}},
            reviewer_ref="reviewer.pytest",
            reviewed_on="2026-07-24",
            numeric_locale="en",
            currency="EUR",
            unit="currency",
            scale="1",
            parse_policy="strict_all_nonblank_v1",
            sign_policy="as_presented_v1",
        )
    with pytest.raises(ValueError, match="cell dispositions do not close"):
        core.review_numeric_measure_columns(
            inspection.inspection,
            recipe,
            section_key="budget",
            header_row=table["header_row"],
            columns=["amount"],
            excluded_columns=["account_id"],
            cell_dispositions={"amount": {2: "include"}},
            reviewer_ref="reviewer.pytest",
            reviewed_on="2026-07-24",
            numeric_locale="en",
            currency="EUR",
            unit="currency",
            scale="1",
            parse_policy="strict_all_nonblank_v1",
            sign_policy="as_presented_v1",
        )


def test_numeric_review_requires_explicit_supported_sign_policy(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text("line,amount\nA,-10\nB,20\n", encoding="utf-8")
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]
    review_args = _numeric_review_args(
        inspection.inspection,
        table["table_id"],
        ["amount"],
    )
    review_args["sign_policy"] = ""

    with pytest.raises(ValueError, match="sign policy"):
        core.review_numeric_measure_columns(
            inspection.inspection,
            recipe,
            section_key="budget",
            **review_args,
            reviewer_ref="reviewer.pytest",
            reviewed_on="2026-07-24",
            numeric_locale="en",
            currency="EUR",
            unit="currency",
            scale="1",
            parse_policy="strict_all_nonblank_v1",
        )


def test_explicit_headerless_review_closes_every_row_and_rendered_output(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "headerless.csv"
    input_path.write_text(
        "Product A,North,100\nProduct B,South,200\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    assert table["header_row"] == 1
    assert [item["column"] for item in table["headerless_numeric_columns"]] == [
        "column_3"
    ]
    cells = next(
        item["nonblank_cells"]
        for item in table["headerless_numeric_measure_cells"]
        if item["column"] == "column_3"
    )
    recipe = inspection.suggested_recipe
    for section in recipe["sections"].values():
        section["assigned_table"] = ""
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]
    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        header_row=None,
        columns=["column_3"],
        excluded_columns=[],
        cell_dispositions={"column_3": {int(cell["row"]): "include" for cell in cells}},
        reviewer_ref="reviewer.headerless-pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency="USD",
        unit="currency",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
        sign_policy="as_presented_v1",
    )
    recipe_path = tmp_path / "reviewed_recipe.json"
    core.write_json(recipe_path, reviewed)
    output_dir = tmp_path / "report"

    core.build_report(input_path, output_dir, recipe_path=recipe_path)

    analysis = json.loads(
        (output_dir / "report_analysis.json").read_text(encoding="utf-8")
    )
    budget = next(
        section for section in analysis["sections"] if section["section"] == "budget"
    )
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )
    public_tables = json.loads(
        (output_dir / "report_tables.json").read_text(encoding="utf-8")
    )
    workbook = openpyxl.load_workbook(
        output_dir / "report_tables.xlsx",
        data_only=True,
    )
    markdown = (output_dir / "report_draft.md").read_text(encoding="utf-8")

    assert budget["header_row"] is None
    assert budget["numeric_columns"][0]["numeric_rows"] == [1, 2]
    assert budget["numeric_columns"][0]["sum"] == "300"
    assert ledger["entries"][0]["source"]["locator"] == "headerless.csv!C1,C2"
    assert "numeric_measure_cells" not in public_tables["tables"][0]
    assert "headerless_numeric_measure_cells" not in public_tables["tables"][0]
    assert [cell.value for cell in workbook["numeric_evidence"][1]] == [
        "evidence_id",
        "section",
        "column",
        "sum",
        "currency",
        "unit",
        "scale",
    ]
    assert "column_3: sum 300 | Currency: USD | Unit: currency | Scale: 1" in markdown


def test_header_heuristic_cannot_hide_the_only_numeric_candidate(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "ambiguous-header.csv"
    input_path.write_text(
        "Product A,North,100\nProduct B,South,N/A\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    recipe = inspection.suggested_recipe
    for section in recipe["sections"].values():
        section["assigned_table"] = ""
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]
    recipe_path = tmp_path / "recipe.json"
    core.write_json(recipe_path, recipe)

    core.build_report(input_path, tmp_path / "report", recipe_path=recipe_path)

    analysis = json.loads(
        (tmp_path / "report" / "report_analysis.json").read_text(encoding="utf-8")
    )
    budget = next(
        section for section in analysis["sections"] if section["section"] == "budget"
    )
    assert budget["numeric_measure_status"] == "needs_review"
    assert [item["column"] for item in budget["numeric_measure_candidates"]] == [
        "column_3"
    ]
    assert not (tmp_path / "report" / "numeric_evidence_ledger.json").exists()


def test_ambiguous_currency_symbol_rejects_non_currency_contract(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "amounts.csv"
    input_path.write_text(
        'line,amount\nA,"$1,000"\nB,"$250"\n',
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]

    with pytest.raises(ValueError, match="unresolved nonblank rows"):
        core.review_numeric_measure_columns(
            inspection.inspection,
            recipe,
            section_key="budget",
            **_numeric_review_args(
                inspection.inspection,
                table["table_id"],
                ["amount"],
            ),
            reviewer_ref="reviewer.non-currency-pytest",
            reviewed_on="2026-07-24",
            numeric_locale="en",
            currency=None,
            unit="number",
            scale="1",
            parse_policy="strict_all_nonblank_v1",
        )


@pytest.mark.parametrize(
    ("field", "value"),
    [
        ("context_key", "Revenue USD 999"),
        ("context_value", "Revenue was USD 999"),
        ("period", "Current period; revenue was USD 999"),
        ("entity", "Revenue USD 999"),
    ],
)
def test_narrative_numeric_boundary_rejects_structured_field_bypasses(
    field: str,
    value: str,
) -> None:
    core = load_core()
    recipe: dict[str, Any] = {
        "entity": "",
        "period": "",
        "executive_summary": "",
        "context_items": {},
        "sections": {},
    }
    if field == "context_key":
        recipe["context_items"] = {value: "supported"}
    elif field == "context_value":
        recipe["context_items"] = {"Revenue": value}
    else:
        recipe[field] = value

    with pytest.raises(ValueError):
        core.validate_narrative_numeric_boundary(recipe)


def test_numeric_review_cli_records_explicit_column_cell_and_sign_contract(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text(
        "line,amount,account_id\nA,10,1001\nB,20,1002\n",
        encoding="utf-8",
    )
    ledger_path = ROOT / "plugins" / "studio-archive" / "scripts" / "client_ledger.py"
    spec = importlib.util.spec_from_file_location(
        "report_builder_test_client_ledger",
        ledger_path,
    )
    assert spec and spec.loader
    ledger = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = ledger
    spec.loader.exec_module(ledger)
    client_root = tmp_path / "Studio" / "Report Client"
    client_root.mkdir(parents=True)
    client_id = "client_333333333333333333333333"
    ledger.create_client_manifest(client_root, client_id)
    engagement = ledger.create_engagement(client_root, client_id, "Report review")
    imported = ledger.import_document(
        client_root,
        client_id,
        engagement["engagement_id"],
        input_path,
        "source",
    )
    prepared = ledger.prepare_run(
        client_root,
        client_id,
        engagement["engagement_id"],
        "report-builder",
        "test-version",
        input_ids=[imported["receipt"]["input_id"]],
    )
    running = ledger.start_run(
        client_root,
        engagement["engagement_id"],
        prepared["run"]["run_id"],
    )
    input_path = Path(running["context"]["input_bindings"][0]["path"])
    inspection_dir = Path(running["output_dir"]) / "inspection"
    inspection = core.inspect_inputs(
        input_path,
        inspection_dir,
        report_type="management_report",
    )
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = input_path.name
    recipe_path = inspection_dir / "recipe.json"
    output_path = inspection_dir / "reviewed.json"
    expansion_path = inspection_dir / "model-context" / "budget-row.json"
    core.write_json(recipe_path, recipe)

    subprocess.run(
        [
            sys.executable,
            str(SCRIPT_DIR / "expand_model_context.py"),
            "--client-engagement",
            str(running["context_path"]),
            "--inspection-control",
            str(inspection_dir / "inspection_control.json"),
            "--output",
            str(expansion_path),
            "--table-id",
            str(inspection.inspection["tables"][0]["table_id"]),
            "--header-row",
            str(inspection.inspection["tables"][0]["header_row"]),
            "--columns",
            "line,amount",
            "--row-start",
            "2",
            "--row-limit",
            "1",
            "--purpose",
            "Review the first budget line.",
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    subprocess.run(
        [
            sys.executable,
            str(SCRIPT_DIR / "review_numeric_measures.py"),
            "--client-engagement",
            str(running["context_path"]),
            "--inspection",
            str(inspection_dir / "inspection.json"),
            "--recipe",
            str(recipe_path),
            "--output",
            str(output_path),
            "--section",
            "budget",
            "--header-row",
            str(inspection.inspection["tables"][0]["header_row"]),
            "--columns",
            "amount",
            "--exclude-columns",
            "account_id",
            "--cell-disposition",
            "amount:2:include",
            "--cell-disposition",
            "amount:3:include",
            "--reviewer-ref",
            "reviewer.pytest",
            "--reviewed-on",
            "2026-07-24",
            "--numeric-locale",
            "en",
            "--currency",
            "EUR",
            "--unit",
            "currency",
            "--scale",
            "1",
            "--parse-policy",
            "strict_all_nonblank_v1",
            "--sign-policy",
            "as_presented_v1",
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    reviewed = json.loads(output_path.read_text(encoding="utf-8"))
    content = reviewed["sections"]["budget"]["numeric_measure_decision"]["content"]
    expansion = json.loads(expansion_path.read_text(encoding="utf-8"))

    assert expansion["rows"][0]["cells"]["line"]["source_text"] == "A"
    assert expansion["context_receipt"]["purpose"] == ("Review the first budget line.")
    assert content["numeric_measure_columns"] == ["amount"]
    assert content["excluded_numeric_candidate_columns"] == ["account_id"]
    assert content["numeric_contract"]["sign_policy"] == "as_presented_v1"


def test_reviewed_office_outputs_are_byte_identical_across_replays(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    _save_workbook(input_path)
    recipe_path = _reviewed_numeric_recipe(
        core,
        input_path,
        tmp_path / "inspection",
    )

    core.build_report(input_path, tmp_path / "first", recipe_path=recipe_path)
    core.build_report(input_path, tmp_path / "second", recipe_path=recipe_path)

    assert (tmp_path / "first" / "report.docx").read_bytes() == (
        tmp_path / "second" / "report.docx"
    ).read_bytes()
    assert (tmp_path / "first" / "report_tables.xlsx").read_bytes() == (
        tmp_path / "second" / "report_tables.xlsx"
    ).read_bytes()


def test_spanish_build_localizes_review_artifacts_docx_and_strict_contract(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "informe.xlsx"
    output_dir = tmp_path / "salida"
    _save_workbook(input_path)
    recipe_path = _reviewed_numeric_recipe(
        core,
        input_path,
        tmp_path / "inspeccion",
        language="es",
        document_language="es",
        report_type="annual_financial_statement",
    )

    core.build_report(
        input_path,
        output_dir,
        recipe_path=recipe_path,
        language="es-ES",
        document_language="es-ES",
        report_type="annual_financial_statement",
    )

    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    review_handoff = (output_dir / "review_handoff.md").read_text(encoding="utf-8")
    document = Document(output_dir / "report.docx")
    document_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    missing_item = next(
        item
        for item in review_payload["items"]
        if item["id"].startswith("missing-section-")
    )
    artifact_item = next(
        item for item in review_payload["items"] if item["output_path"] == "report.docx"
    )
    outputs_by_path = {output["path"]: output for output in final_artifacts["outputs"]}

    assert review_payload["language"] == "es"
    assert review_payload["columns"] == [
        {"field": "item_type", "label": "Tipo"},
        {"field": "title", "label": "Elemento del informe"},
        {"field": "recommended_action", "label": "Acción sugerida"},
        {"field": "source_path", "label": "Fuente"},
        {"field": "output_path", "label": "Salida"},
        {"field": "status", "label": "Estado"},
    ]
    assert missing_item["title"].startswith("Falta la asignación de la sección:")
    assert missing_item["data"]["requested_document"].startswith(
        "Tabla de origen o soporte narrativo para la sección"
    )
    assert missing_item["data"]["reason"].startswith(
        "No hay ninguna tabla de origen determinista"
    )
    assert artifact_item["title"] == "Informe de Word"
    assert review_handoff.startswith("# Entrega para revisión: Generador de informes\n")
    assert "## Revisión en Codex" in review_handoff
    assert outputs_by_path["review_handoff.md"]["required_text"][0] == (
        "Entrega para revisión"
    )
    assert {
        "Resumen ejecutivo",
        "Anexo de auditoría",
        "Ruta de entrada",
        "Tablas detectadas",
    } <= set(outputs_by_path["report.docx"]["required_text"])
    assert "Resumen ejecutivo" in document_text
    assert "Anexo de auditoría" in document_text
    assert "Ruta de entrada" in document_text
    assert "Tablas detectadas" in document_text
    assert "Columna" in document_text
    assert "Suma" in document_text
    assert "Recuento" not in document_text
    assert "Rango" not in document_text
    assert "Input path" not in document_text
    assert "Tables discovered" not in document_text
    assert final_artifacts["caveats"][0].startswith("Codex sigue siendo responsable")
    assert final_artifacts["next_actions"][2].startswith("Use report.docx")
    contract_report = validate_contract(
        output_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()


def test_plugin_inspects_and_builds_report_without_model_calls(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    output_dir = tmp_path / "out"
    report_dir = output_dir / "report"
    _save_workbook(input_path)

    inspection = core.inspect_inputs(
        input_path,
        output_dir,
        language="en",
        document_language="auto",
        report_type="management_report",
    )
    recipe_path = output_dir / "suggested_recipe.json"
    recipe = json.loads(recipe_path.read_text(encoding="utf-8"))
    recipe["entity"] = "Example Ltd"
    recipe["period"] = "2025"
    recipe["executive_summary"] = "Codex reviewed the mapped tables."
    recipe["sections"]["income_statement"][
        "codex_comment"
    ] = "Revenue and result were reviewed against the income statement table."
    recipe_path.write_text(json.dumps(recipe, indent=2), encoding="utf-8")

    result = core.build_report(
        input_path,
        report_dir,
        recipe_path=recipe_path,
        language="en",
        document_language="auto",
        report_type="management_report",
    )

    inspection_payload = json.loads((output_dir / "inspection.json").read_text())
    analysis_payload = json.loads((report_dir / "report_analysis.json").read_text())
    audit_payload = json.loads((report_dir / "report_audit.json").read_text())
    run_intake = json.loads((report_dir / "run_intake.json").read_text())
    review_payload = json.loads((report_dir / "review_payload.json").read_text())
    ui_decisions = json.loads((report_dir / "ui_decisions.json").read_text())
    final_artifacts = json.loads((report_dir / "final_artifacts.json").read_text())
    draft = (report_dir / "report_draft.md").read_text(encoding="utf-8")
    document = Document(result.docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )

    assert inspection.inspection["table_count"] == 3
    assert inspection_payload["language"] == "en"
    assert recipe["sections"]["income_statement"]["assigned_table"]
    assert analysis_payload["assigned_section_count"] >= 3
    assert audit_payload["model_api_calls"] == 0
    assert "Revenue and result were reviewed" in draft
    assert result.docx_path.exists()
    assert "Management report" in paragraph_text
    assert "Executive summary" in paragraph_text
    assert "Audit appendix" in paragraph_text
    assert "Revenue" in table_text
    assert len(document.tables) >= 4
    assert (report_dir / "report_tables.json").exists()
    assert (report_dir / "report_tables.xlsx").exists()
    assert (report_dir / "used_recipe.json").exists()
    assert result.review_session == audit_payload["review_session"]
    assert audit_payload["review_session"]["run_id"] == run_intake["run_id"]
    assert review_payload["plugin"] == "report-builder"
    assert review_payload["workflow"] == "report-builder"
    assert review_payload["run_id"] == run_intake["run_id"]
    assert review_payload["review_type"] == "report_builder_review"
    assert review_payload["item_count"] == len(review_payload["items"])
    item_types = {item["item_type"] for item in review_payload["items"]}
    assert "report_section" in item_types
    assert "table_evidence" in item_types
    assert "report_artifact" in item_types
    assert review_payload["summary"]["assigned_section_count"] >= 3
    assert review_payload["summary"]["table_count"] == 3
    income_section_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "report_section"
        and item["data"]["section"] == "income_statement"
    )
    assert income_section_item["data"]["target_artifact"] == "report.docx"
    assert (
        income_section_item["data"]["target_path"]
        == "sections.income_statement.codex_comment"
    )
    assert income_section_item["data"]["target_field"] == "codex_comment"
    income_table_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "table_evidence"
        and item["data"]["section"] == "income_statement"
    )
    assert income_table_item["data"]["target_artifact"] == "report.docx"
    assert (
        income_table_item["data"]["target_path"]
        == "sections.income_statement.assigned_table"
    )
    assert income_table_item["data"]["target_field"] == "assigned_table"
    assert "report.xlsx::Cash Flow" in income_table_item["data"]["available_table_ids"]
    assert income_table_item["data"]["preview_rows"][0]["Line"] == "Revenue"
    assert income_table_item["evidence"][0]["preview_rows"][0]["Line"] == "Revenue"
    assert ui_decisions["status"] == "pending_review"
    assert ui_decisions["decision_count"] == 0
    assert final_artifacts["status"] == "written_pending_review"
    handoff_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "review_handoff.md"
    )
    handoff_text = (report_dir / "review_handoff.md").read_text(encoding="utf-8")
    assert handoff_output["required_text"] == [
        "Review Handoff",
        "review_payload.json",
        "ui_decisions.json",
        "applied_decisions.json",
        "final_artifacts.json",
    ]
    assert "render_report_builder_review" in handoff_text
    assert "apply_report_builder_decisions" in handoff_text
    report_draft_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "report_draft.md"
    )
    first_section_title = analysis_payload["sections"][0]["title"]
    assert "## Executive summary" in report_draft_output["required_text"]
    assert f"## {first_section_title}" in report_draft_output["required_text"]
    assert "Source:" in report_draft_output["required_text"]
    assert "Rows:" in report_draft_output["required_text"]
    assert "required_text" in report_draft_output["qa_checks"]
    report_docx_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "report.docx"
    )
    assert "Executive summary" in report_docx_output["required_text"]
    assert "Audit appendix" in report_docx_output["required_text"]
    assert "Report status" in report_docx_output["required_text"]
    assert "Model API calls from scripts" in report_docx_output["required_text"]
    assert first_section_title in report_docx_output["required_text"]
    assert "required_text" in report_docx_output["qa_checks"]
    report_tables_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "report_tables.xlsx"
    )
    first_section_row, first_section = next(
        (
            (index, section)
            for index, section in enumerate(analysis_payload["sections"], start=2)
            if section["assigned_table"]
        )
    )
    preview_sheet = first_section["section"]
    assert report_tables_output["required_sheets"] == ["summary", preview_sheet]
    assert report_tables_output["required_sheet_headers"] == {
        "summary": ["section", "status", "assigned_table", "rows", "columns"],
        preview_sheet: ["Line", "Actual", "Budget"],
    }
    assert report_tables_output["required_cells"] == {
        "summary": {
            f"A{first_section_row}": str(first_section["section"]),
            f"B{first_section_row}": str(first_section["status"]),
            f"C{first_section_row}": str(first_section["assigned_table"]),
            f"D{first_section_row}": str(first_section["row_count"]),
            f"E{first_section_row}": str(first_section["column_count"]),
        },
        preview_sheet: {
            "A1": "Line",
            "A2": "Revenue",
            "B1": "Actual",
            "B2": "[numeric source value withheld]",
            "C1": "Budget",
            "C2": "[numeric source value withheld]",
        },
    }
    assert "required_sheet_headers" in report_tables_output["qa_checks"]
    assert "required_cells" in report_tables_output["qa_checks"]
    report_tables_json_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "report_tables.json"
    )
    assert report_tables_json_output["records_key"] == "tables"
    assert report_tables_json_output["row_count"] == audit_payload["table_count"]
    assert report_tables_json_output["required_columns"] == [
        "table_id",
        "source_file",
        "row_count",
        "column_count",
    ]
    contract_report = validate_contract(
        report_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()


def test_inspection_writes_bounded_model_packet_and_private_full_control(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "customers.csv"
    rows = ["customer,amount"] + [
        f"Customer {index},{index * 10}" for index in range(1, 206)
    ]
    input_path.write_text("\n".join(rows) + "\n", encoding="utf-8")
    output_dir = tmp_path / "inspection"

    result = core.inspect_inputs(
        input_path,
        output_dir,
        report_type="management_report",
    )

    model_packet = json.loads((output_dir / "inspection.json").read_text())
    control_packet = json.loads((output_dir / "inspection_control.json").read_text())
    receipt = json.loads((output_dir / "model_context_receipt.json").read_text())
    assert "numeric_measure_cells" not in model_packet["tables"][0]
    assert "headerless_numeric_measure_cells" not in model_packet["tables"][0]
    assert "numeric_measure_cells" in control_packet["tables"][0]
    assert result.inspection == control_packet
    assert "Customer 20" not in (output_dir / "inspection.json").read_text()
    assert "Customer 20" in (output_dir / "inspection_control.json").read_text()
    assert model_packet["model_context"]["full_population_processed_locally"] is True
    assert (
        receipt["default_model_packet"]["sha256"]
        == hashlib.sha256((output_dir / "inspection.json").read_bytes()).hexdigest()
    )
    assert (
        receipt["private_control"]["sha256"]
        == hashlib.sha256(
            (output_dir / "inspection_control.json").read_bytes()
        ).hexdigest()
    )
    table = result.inspection["tables"][0]
    expanded_names = []
    for row_start, row_limit in ((2, 100), (102, 100), (202, 5)):
        packet = core.build_model_context_expansion(
            result.inspection,
            table_id=table["table_id"],
            header_row=table["header_row"],
            columns=["customer"],
            row_start=row_start,
            row_limit=row_limit,
            purpose=f"Review source rows starting at {row_start}.",
        )
        expanded_names.extend(
            row["cells"]["customer"]["source_text"] for row in packet["rows"]
        )
    assert expanded_names == [f"Customer {index}" for index in range(1, 206)]


def test_model_context_expansion_is_exact_bounded_and_receipted(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "customers.csv"
    input_path.write_text(
        "customer,amount,internal_note\n"
        "Alpha,10,ordinary\n"
        "Private Counterparty,20,sensitive evidence\n"
        "Gamma,30,ordinary\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="management_report",
    ).inspection
    table = inspection["tables"][0]

    packet = core.build_model_context_expansion(
        inspection,
        table_id=table["table_id"],
        header_row=table["header_row"],
        columns=["customer", "amount"],
        row_start=3,
        row_limit=1,
        purpose="Review one exceptional line for the budget narrative.",
    )

    assert packet["selection"]["columns"] == ["customer", "amount"]
    assert packet["rows"] == [
        {
            "source_row": 3,
            "cells": {
                "customer": {
                    "row": 3,
                    "coordinate": "A3",
                    "source_text": "Private Counterparty",
                    "number_format": "",
                    "formula": None,
                    "formula_cached_value": None,
                    "formula_cache_status": "not_formula",
                },
                "amount": {
                    "row": 3,
                    "coordinate": "B3",
                    "source_text": "20",
                    "number_format": "",
                    "formula": None,
                    "formula_cached_value": None,
                    "formula_cache_status": "not_formula",
                },
            },
        }
    ]
    unhashed = {key: value for key, value in packet.items() if key != "context_receipt"}
    assert packet["context_receipt"]["content_sha256"] == _canonical_json_sha256(
        unhashed
    )
    assert "sensitive evidence" not in json.dumps(packet)
    with pytest.raises(ValueError, match="row_limit"):
        core.build_model_context_expansion(
            inspection,
            table_id=table["table_id"],
            header_row=table["header_row"],
            columns=["customer"],
            row_start=2,
            row_limit=101,
            purpose="Attempt an over-broad disclosure.",
        )


def test_plugin_marks_unassigned_sections_for_codex_review(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    output_dir = tmp_path / "out"
    _save_workbook(input_path)

    result = core.build_report(
        input_path,
        output_dir,
        language="en",
        report_type="annual_financial_statement",
    )
    draft = result.markdown_path.read_text(encoding="utf-8")
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    missing_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "review_issue"
        and item["id"].startswith("missing-section-")
    )

    assert result.audit["missing_section_count"] > 0
    assert "Codex review pending for this section." in draft
    assert "request_more_documents" in missing_item["allowed_actions"]
    assert missing_item["data"]["requested_document"].startswith(
        "Source table or narrative support for report section "
    )
    assert (
        missing_item["data"]["required_document"]
        == missing_item["data"]["requested_document"]
    )
    assert missing_item["data"]["reason"] == (
        "No deterministic source table is mapped to this report section."
    )
    assert missing_item["data"]["source_table"] == "unassigned"
    assert missing_item["data"]["record_id"] == missing_item["data"]["section"]
    assert (
        missing_item["evidence"][0]["requested_document"]
        == missing_item["data"]["requested_document"]
    )


def test_report_builder_request_more_documents_prefills_blocker_context(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "report.xlsx"
    _save_workbook(input_path)
    managed = _managed_report_run(
        tmp_path,
        input_path,
        language="en",
        report_type="annual_financial_statement",
    )
    output_dir = managed["output_dir"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    missing_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "review_issue"
        and item["id"].startswith("missing-section-")
    )

    apply_result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": missing_item["id"],
                        "action": "request_more_documents",
                        "reviewer_note": "Ask the client for the missing schedule.",
                    }
                ],
                "decision_source": "pytest_missing_section_request",
                "reviewer": "pytest",
            },
        },
    )

    payload = apply_result["structuredContent"]
    applied = json.loads((output_dir / "applied_decisions.json").read_text())
    updated_final = json.loads((output_dir / "final_artifacts.json").read_text())
    expected_document = missing_item["data"]["requested_document"]

    assert payload["ok"] is True
    assert payload["application_status"] == "blocked"
    assert applied["effects"][0]["requested_documents"] == [expected_document]
    assert (
        applied["effects"][0]["followup_context"]["record_id"]
        == missing_item["data"]["section"]
    )
    assert applied["effects"][0]["followup_context"]["source_table"] == "unassigned"
    assert updated_final["blockers"][0]["requested_documents"] == [expected_document]
    assert updated_final["blockers"][0]["followup_context"]["reason"] == (
        "No deterministic source table is mapped to this report section."
    )


def test_report_builder_apply_decisions_regenerates_docx_for_section_edit(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    inspection_dir = tmp_path / "inspection"
    _save_workbook(input_path)

    inspection = core.inspect_inputs(
        input_path,
        inspection_dir,
        language="en",
        document_language="auto",
        report_type="management_report",
    )
    recipe_path = inspection_dir / "suggested_recipe.json"
    recipe = json.loads(recipe_path.read_text(encoding="utf-8"))
    recipe["entity"] = "Example Ltd"
    recipe["period"] = "2025"
    recipe["executive_summary"] = "Codex reviewed the mapped tables."
    recipe["sections"]["income_statement"][
        "codex_comment"
    ] = "Original income statement narrative."
    recipe_path.write_text(json.dumps(recipe, indent=2), encoding="utf-8")
    managed = _managed_report_run(
        tmp_path,
        input_path,
        recipe=recipe_path,
        language="en",
        document_language="auto",
        report_type="management_report",
    )
    report_dir = managed["output_dir"]
    result = managed["result"]
    assert inspection.inspection["table_count"] == 3
    review_payload = json.loads((report_dir / "review_payload.json").read_text())
    run_intake = json.loads((report_dir / "run_intake.json").read_text())
    final_artifacts = json.loads((report_dir / "final_artifacts.json").read_text())
    section_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "report_section"
        and item["data"]["section"] == "income_statement"
    )
    revised_text = "Reviewer-approved income statement narrative."
    assert revised_text not in "\n".join(
        paragraph.text for paragraph in Document(result.docx_path).paragraphs
    )

    apply_result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": section_item["id"],
                        "action": "edit",
                        "edit_value": revised_text,
                        "reviewer_note": "Replace section narrative in native report.",
                    }
                ],
                "decision_source": "pytest_docx_regeneration",
                "reviewer": "pytest",
            },
        },
    )

    payload = apply_result["structuredContent"]
    assert payload["ok"] is True
    assert payload["application_status"] == "partial_review_applied"
    assert payload["native_regeneration_count"] == 0
    assert payload["native_regenerated_count"] == 1
    assert set(payload["applied_decisions"]["native_regenerated_paths"]) >= {
        "report.docx",
        "report_draft.md",
        "used_recipe.json",
        "report_analysis.json",
    }

    updated_recipe = json.loads((report_dir / "used_recipe.json").read_text())
    updated_analysis = json.loads((report_dir / "report_analysis.json").read_text())
    updated_draft = (report_dir / "report_draft.md").read_text(encoding="utf-8")
    updated_docx_text = "\n".join(
        paragraph.text for paragraph in Document(report_dir / "report.docx").paragraphs
    )
    applied = json.loads((report_dir / "applied_decisions.json").read_text())
    updated_final = json.loads((report_dir / "final_artifacts.json").read_text())

    assert (
        updated_recipe["sections"]["income_statement"]["codex_comment"] == revised_text
    )
    assert (
        next(
            section
            for section in updated_analysis["sections"]
            if section["section"] == "income_statement"
        )["codex_comment"]
        == revised_text
    )
    assert revised_text in updated_draft
    assert revised_text in updated_docx_text
    assert applied["effects"][0]["artifact_update"] == "native_artifact_regenerated"
    assert applied["effects"][0]["native_regeneration_status"] == "regenerated"
    assert applied["effects"][0]["terminal_application"] is True
    assert applied["effects"][0]["application_receipt"]["target_path"] == (
        "sections.income_statement.codex_comment"
    )
    assert len(applied["effects"][0]["application_receipt"]["report_docx_sha256"]) == 64
    assert {
        receipt["path"]
        for receipt in applied["effects"][0]["application_receipt"][
            "regenerated_outputs"
        ]
    } >= {
        "used_recipe.json",
        "report_analysis.json",
        "report_draft.md",
        "report.docx",
    }
    assert applied["native_regeneration_count"] == 0
    assert applied["native_regenerated_count"] == 1
    assert applied["application_status"] == "partial_review_applied"
    assert updated_final["status"] == "partial_review_applied"
    report_output = next(
        output for output in updated_final["outputs"] if output["path"] == "report.docx"
    )
    assert report_output["status"] == "updated_from_review"
    assert report_output["native_regenerated"] is True
    assert "Regenerate native DOCX/XLSX/PDF outputs before final handoff." not in (
        updated_final["next_actions"]
    )
    assert (
        "Complete remaining review decisions before final handoff."
        in updated_final["next_actions"]
    )
    assert (report_dir / "revisions/originals").exists()


def test_report_builder_apply_decisions_regenerates_outputs_for_source_mapping_edit(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    inspection_dir = tmp_path / "inspection"
    _save_workbook(input_path)

    inspection = core.inspect_inputs(
        input_path,
        inspection_dir,
        language="en",
        document_language="auto",
        report_type="management_report",
    )
    recipe_path = inspection_dir / "suggested_recipe.json"
    recipe = json.loads(recipe_path.read_text(encoding="utf-8"))
    recipe["entity"] = "Example Ltd"
    recipe["period"] = "2025"
    recipe["executive_summary"] = "Codex reviewed the mapped tables."
    recipe["sections"]["income_statement"][
        "codex_comment"
    ] = "Income statement narrative follows the mapped source."
    recipe_path.write_text(json.dumps(recipe, indent=2), encoding="utf-8")
    managed = _managed_report_run(
        tmp_path,
        input_path,
        recipe=recipe_path,
        language="en",
        document_language="auto",
        report_type="management_report",
    )
    report_dir = managed["output_dir"]
    assert inspection.inspection["table_count"] == 3

    review_payload = json.loads((report_dir / "review_payload.json").read_text())
    run_intake = json.loads((report_dir / "run_intake.json").read_text())
    final_artifacts = json.loads((report_dir / "final_artifacts.json").read_text())
    table_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "table_evidence"
        and item["data"]["section"] == "income_statement"
    )
    revised_table_id = "report.xlsx::Cash Flow"
    assert revised_table_id in table_item["data"]["available_table_ids"]

    apply_result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": table_item["id"],
                        "action": "edit",
                        "edit_value": revised_table_id,
                        "reviewer_note": "Use the cash flow table for this section.",
                    }
                ],
                "decision_source": "pytest_source_mapping_regeneration",
                "reviewer": "pytest",
            },
        },
    )

    payload = apply_result["structuredContent"]
    updated_recipe = json.loads((report_dir / "used_recipe.json").read_text())
    updated_analysis = json.loads((report_dir / "report_analysis.json").read_text())
    updated_audit = json.loads((report_dir / "report_audit.json").read_text())
    updated_final = json.loads((report_dir / "final_artifacts.json").read_text())
    updated_draft = (report_dir / "report_draft.md").read_text(encoding="utf-8")
    workbook = openpyxl.load_workbook(report_dir / "report_tables.xlsx", data_only=True)
    income_section = next(
        section
        for section in updated_analysis["sections"]
        if section["section"] == "income_statement"
    )
    outputs_by_path = {
        output["path"]: output
        for output in updated_final["outputs"]
        if isinstance(output, dict)
    }

    assert payload["ok"] is True
    assert payload["application_status"] == "partial_review_applied"
    assert (
        updated_recipe["sections"]["income_statement"]["assigned_table"]
        == revised_table_id
    )
    assert income_section["assigned_table"] == revised_table_id
    assert income_section["sheet_name"] == "Cash Flow"
    assert income_section["row_count"] == 3
    assert "Source: report.xlsx / Cash Flow" in updated_draft
    assert updated_audit["review_native_regeneration"]["status"] == "regenerated"
    assert (
        "report_tables.xlsx" in payload["applied_decisions"]["native_regenerated_paths"]
    )
    assert workbook["summary"]["C3"].value == revised_table_id
    assert workbook["income_statement"]["A2"].value == "Operating cash"
    assert workbook["income_statement"]["B2"].value == "[numeric source value withheld]"
    assert outputs_by_path["report_tables.xlsx"]["status"] == "updated_from_review"
    assert outputs_by_path["report_tables.xlsx"]["native_regenerated"] is True
    assert outputs_by_path["report_tables.xlsx"]["required_sheets"] == [
        "summary",
        "income_statement",
    ]
    assert outputs_by_path["report_tables.xlsx"]["required_sheet_headers"] == {
        "summary": ["section", "status", "assigned_table", "rows", "columns"],
        "income_statement": ["Line", "Amount"],
    }
    assert outputs_by_path["report_tables.xlsx"]["required_cells"] == {
        "summary": {
            "A3": "income_statement",
            "B3": "assigned",
            "C3": revised_table_id,
            "D3": "3",
            "E3": "2",
        },
        "income_statement": {
            "A1": "Line",
            "A2": "Operating cash",
            "B1": "Amount",
            "B2": "[numeric source value withheld]",
        },
    }
    assert outputs_by_path["report_tables.json"]["row_count"] == 3
    assert (
        "report_tables.xlsx"
        in updated_final["review_application"]["native_regenerated_paths"]
    )
    contract_report = validate_contract(
        report_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()


def test_invalid_mapping_preflight_leaves_persisted_run_byte_identical(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "report.xlsx"
    _save_workbook(input_path)
    managed = _managed_report_run(tmp_path, input_path)
    output_dir = managed["output_dir"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    table_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "table_evidence"
    )
    before = _tree_snapshot(output_dir)

    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": table_item["id"],
                        "action": "edit",
                        "edit_value": "forged.xlsx::Not A Current Table",
                    }
                ],
            },
        },
    )
    failure = response["result"]["structuredContent"]

    assert failure["ok"] is False
    assert failure["error"] == "Report Builder review transaction failed safely."
    assert _tree_snapshot(output_dir) == before
    assert not (output_dir / "applied_decisions.json").exists()
    assert not (output_dir / "revisions").exists()
    assert not list(tmp_path.glob(".report-builder-transaction-*"))


def test_existing_application_lock_fails_closed_before_review_writes(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "budget.csv"
    input_path.write_text("line,amount\nA,10\nB,20\n", encoding="utf-8")
    managed = _managed_report_run(tmp_path, input_path)
    output_dir = managed["output_dir"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    section_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "report_section"
        and item["data"]["status"] == "assigned"
    )
    lock_path = output_dir / ".report-builder-application.lock"
    lock_path.write_text('{"pid":1}\n', encoding="utf-8")
    before = _tree_snapshot(output_dir)

    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": section_item["id"],
                        "action": "edit",
                        "edit_value": "This edit must wait for the active transaction.",
                    }
                ],
            },
        },
    )
    failure = response["result"]["structuredContent"]

    assert failure["ok"] is False
    assert "already in progress" in failure["error"]
    assert _tree_snapshot(output_dir) == before
    assert not (output_dir / "applied_decisions.json").exists()
    assert not list(tmp_path.glob(".report-builder-transaction-*"))


def test_source_mutation_during_application_rolls_back_every_output_write(
    tmp_path: Path,
) -> None:
    source = tmp_path / "budget.csv"
    source.write_text("line,amount\nA,10\nB,20\n", encoding="utf-8")
    managed = _managed_report_run(tmp_path, source)
    output_dir = managed["output_dir"]
    input_path = managed["source_path"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    section_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "report_section"
        and item["data"]["status"] == "assigned"
    )
    before = _tree_snapshot(output_dir)
    mutated = threading.Event()

    def mutate_after_first_transaction_write() -> None:
        deadline = time.monotonic() + 10
        while time.monotonic() < deadline:
            transaction_revisions = list(
                output_dir.parent.glob(
                    ".generated-review-transaction-*/working/revisions"
                )
            )
            if transaction_revisions:
                input_path.write_text(
                    "line,amount\nA,999\nB,20\n",
                    encoding="utf-8",
                )
                mutated.set()
                return
            time.sleep(0.001)

    mutator = threading.Thread(target=mutate_after_first_transaction_write)
    mutator.start()
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": section_item["id"],
                        "action": "edit",
                        "edit_value": "A transactionally applied narrative.",
                    }
                ],
            },
        },
    )
    mutator.join(timeout=10)
    failure = response["result"]["structuredContent"]

    assert mutated.is_set()
    assert failure["ok"] is False
    assert any(
        token in failure["error"].lower()
        for token in (
            "source",
            "receipt",
            "integrity",
            "stale",
            "transaction failed safely",
            "native regeneration failed",
        )
    ), failure
    assert _tree_snapshot(output_dir) == before
    assert not (output_dir / "applied_decisions.json").exists()
    assert not (output_dir / "revisions").exists()
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize(
    "tool_name",
    ["save_report_builder_decisions", "apply_report_builder_decisions"],
)
@pytest.mark.parametrize(
    "scenario",
    [
        "poison_file",
        "delete_snapshot",
        "symlink_snapshot",
        "fifo_snapshot",
        "canonical_delete",
    ],
)
def test_review_transaction_restores_trusted_bytes_and_modes_after_commit_fault(
    monkeypatch: Any,
    tmp_path: Path,
    tool_name: str,
    scenario: str,
) -> None:
    # Arrange
    if scenario == "fifo_snapshot" and sys.platform == "win32":
        pytest.skip("FIFO transaction probe requires a POSIX host.")
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    mode_probe_file = output_dir / "report_draft.md"
    mode_probe_file.chmod(0o640)
    output_dir.chmod(0o750)
    before = _transaction_tree_state(output_dir)
    marker, external = _install_review_transaction_commit_fault(
        monkeypatch,
        tmp_path,
        output_dir,
        scenario,
    )
    external_before = _transaction_tree_state(external)

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {"name": tool_name, "arguments": arguments},
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert marker.read_text(encoding="utf-8") == "commit fault triggered\n"
    assert failure["ok"] is False
    error = failure["error"]
    assert error == "Report Builder review transaction failed safely."
    assert len(error) <= 240
    assert "\n" not in error
    assert "/" not in error
    assert "\\" not in error
    assert "Traceback" not in error
    assert _transaction_tree_state(output_dir) == before
    assert _transaction_tree_state(external) == external_before
    assert not (output_dir / "poison.fifo").exists()
    assert not list(tmp_path.glob(".generated-review-transaction-*"))
    assert not list(tmp_path.glob(".generated-review-commit-*"))
    assert not list(tmp_path.glob(".generated-review-recovery-*"))


def test_review_transaction_rejects_root_relocation_without_moving_canonical(
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    before = _transaction_tree_state(output_dir)
    canonical_inode = output_dir.stat().st_ino
    marker = _install_review_transaction_root_relocation(
        monkeypatch,
        tmp_path,
    )

    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": arguments,
        },
    )

    assert marker.read_text(encoding="utf-8") == ("transaction root relocated\n")
    assert response["result"]["structuredContent"]["ok"] is False
    assert output_dir.stat().st_ino == canonical_inode
    assert _transaction_tree_state(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))
    assert not list(tmp_path.glob(".generated-review-commit-*"))
    assert not list(tmp_path.glob(".generated-review-recovery-*"))


@pytest.mark.parametrize(
    "tool_name",
    ["save_report_builder_decisions", "apply_report_builder_decisions"],
)
def test_review_transaction_bounds_fail_before_canonical_mutation(
    monkeypatch: Any,
    tmp_path: Path,
    tool_name: str,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    sentinel = output_dir / "run_intake.json"
    sentinel_before = (
        sentinel.read_bytes(),
        sentinel.lstat().st_ino,
        stat.S_IMODE(sentinel.lstat().st_mode),
    )
    oversized = output_dir / "oversized.bin"
    with oversized.open("wb") as handle:
        handle.truncate(128 * 1024 * 1024 + 1)
    oversized.chmod(0o640)
    oversized_before = oversized.lstat()
    marker, external = _install_review_transaction_commit_fault(
        monkeypatch,
        tmp_path,
        output_dir,
        "poison_file",
    )
    external_before = _transaction_tree_state(external)

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {"name": tool_name, "arguments": arguments},
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == "Report Builder review transaction failed safely."
    assert not marker.exists()
    assert (
        sentinel.read_bytes(),
        sentinel.lstat().st_ino,
        stat.S_IMODE(sentinel.lstat().st_mode),
    ) == sentinel_before
    oversized_after = oversized.lstat()
    assert oversized_after.st_ino == oversized_before.st_ino
    assert oversized_after.st_size == oversized_before.st_size
    assert stat.S_IMODE(oversized_after.st_mode) == 0o640
    assert _transaction_tree_state(external) == external_before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize(
    "tool_name",
    ["save_report_builder_decisions", "apply_report_builder_decisions"],
)
def test_review_transaction_honest_commit_preserves_unwritten_owned_modes(
    tmp_path: Path,
    tool_name: str,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    mode_probe_file = output_dir / "source_index.json"
    mode_probe_file.chmod(0o640)
    output_dir.chmod(0o750)
    before = _transaction_tree_state(output_dir)

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {"name": tool_name, "arguments": arguments},
    )
    result = response["result"]["structuredContent"]

    # Assert
    assert result["ok"] is True
    after = _transaction_tree_state(output_dir)
    assert after["."] == before["."]
    assert after["source_index.json"] == before["source_index.json"]
    expected_artifact = (
        "ui_decisions.json"
        if tool_name == "save_report_builder_decisions"
        else "applied_decisions.json"
    )
    assert (output_dir / expected_artifact).is_file()
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize(
    ("tool_name", "attack_kind"),
    [
        ("save_report_builder_decisions", "forged_result"),
        ("save_report_builder_decisions", "tampered_staged_file"),
        ("apply_report_builder_decisions", "forged_result"),
        ("apply_report_builder_decisions", "forged_scalar_count"),
        ("apply_report_builder_decisions", "tampered_staged_file"),
        (
            "apply_report_builder_decisions",
            "unauthorized_source_mapping_path",
        ),
    ],
)
def test_review_transaction_parent_hook_rejects_forged_result_or_staged_state(
    monkeypatch: Any,
    tmp_path: Path,
    tool_name: str,
    attack_kind: str,
) -> None:
    # Arrange
    fixture = (
        _report_source_mapping_transaction_fixture
        if attack_kind == "unauthorized_source_mapping_path"
        else _report_transaction_fixture
    )
    output_dir, arguments = fixture(tmp_path)
    before = _transaction_tree_state(output_dir)
    marker = _install_report_transaction_result_fault(
        monkeypatch,
        tmp_path,
        output_dir,
        attack_kind,
    )

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {"name": tool_name, "arguments": arguments},
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert marker.read_text(encoding="utf-8") == "result fault triggered\n"
    assert failure["ok"] is False
    assert failure["error"] == "Report Builder review transaction failed safely."
    assert "forged-run-id" not in failure["error"]
    assert tmp_path.as_posix() not in failure["error"]
    assert _transaction_tree_state(output_dir) == before
    assert not list(tmp_path.glob(".generated-review-transaction-*"))


def test_report_artifact_edit_is_rejected_without_exact_adapter() -> None:
    review_payload = {
        "schema_version": "1.0",
        "plugin": "report-builder",
        "workflow": "report-builder",
        "run_id": "unsupported-artifact-edit",
        "review_type": "report_builder_review",
        "items": [
            {
                "id": "artifact-1",
                "item_type": "report_artifact",
                "title": "Markdown report draft",
                "output_path": "report_draft.md",
                "allowed_actions": ["accept", "edit"],
                "recommended_action": "accept",
                "evidence": [],
                "data": {"path": "report_draft.md"},
                "status": "needs_review",
            }
        ],
        "item_count": 1,
        "status": "ready_for_review",
    }

    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "run_intake": {"run_id": review_payload["run_id"]},
                "review_payload": review_payload,
                "decisions": [
                    {
                        "item_id": "artifact-1",
                        "action": "edit",
                        "edit_value": "This must never replace the generated report.",
                    }
                ],
            },
        },
    )
    failure = response["result"]["structuredContent"]

    assert failure["ok"] is False
    assert "cannot edit a report artifact" in failure["error"]


@pytest.mark.parametrize(
    "tamper",
    (
        _tamper_report_source,
        _tamper_report_output,
        _tamper_persisted_review_payload,
        _supply_stale_review_payload,
    ),
    ids=(
        "source-bytes",
        "output-bytes",
        "persisted-review-payload",
        "supplied-review-payload",
    ),
)
def test_mcp_apply_rejects_stale_state_before_any_review_write(
    tmp_path: Path,
    tamper: Any,
) -> None:
    source = tmp_path / "budget.csv"
    source.write_text(
        "line,amount\nA,10\nB,20\n",
        encoding="utf-8",
    )
    managed = _managed_report_run(tmp_path, source)
    input_path = managed["source_path"]
    output_dir = managed["output_dir"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    section_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "report_section" and "edit" in item["allowed_actions"]
    )
    arguments = {
        "client_engagement": managed["context_path"].as_posix(),
        "run_intake": run_intake,
        "review_payload": review_payload,
        "final_artifacts": final_artifacts,
        "decisions": [
            {
                "item_id": section_item["id"],
                "action": "edit",
                "edit_value": "This edit must not be written.",
            }
        ],
    }
    ui_before = (output_dir / "ui_decisions.json").read_bytes()
    final_before = (output_dir / "final_artifacts.json").read_bytes()
    tampered_arguments = tamper(input_path, output_dir, arguments)

    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": tampered_arguments,
        },
    )
    failure = response["result"]["structuredContent"]

    assert failure["ok"] is False
    assert any(
        term in failure["error"].lower()
        for term in (
            "authorization",
            "integrity",
            "receipt",
            "stale",
            "does not match",
            "customer-folder",
        )
    )
    assert not (output_dir / "applied_decisions.json").exists()
    assert not (output_dir / "revisions").exists()
    assert (output_dir / "ui_decisions.json").read_bytes() == ui_before
    assert (output_dir / "final_artifacts.json").read_bytes() == final_before


def test_mcp_save_rejects_forged_persisted_identity_without_mutation(
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    forged = copy.deepcopy(arguments)
    forged_run_id = "forged-caller-save-run"
    forged["run_intake"]["run_id"] = forged_run_id
    forged["review_payload"]["run_id"] = forged_run_id
    forged["final_artifacts"]["run_id"] = forged_run_id
    before = _transaction_tree_state(output_dir)
    child_marker = _install_report_builder_fake_authority_child(
        monkeypatch,
        tmp_path,
    )

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "save_report_builder_decisions",
            "arguments": forged,
        },
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == (
        "Report Builder customer-run preflight returned an invalid result"
    )
    assert _transaction_tree_state(output_dir) == before
    assert not child_marker.exists()
    assert not list(tmp_path.glob(".generated-review-transaction-*"))
    assert not list(tmp_path.glob(".generated-review-commit-*"))
    assert not list(tmp_path.glob(".generated-review-recovery-*"))


def test_mcp_apply_rejects_forged_one_item_caller_state_without_mutation(
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    forged = copy.deepcopy(arguments)
    forged_run_id = "forged-caller-apply-fresh-run"
    forged_item_id = "forged-caller-fresh-report-section"
    forged["run_intake"]["run_id"] = forged_run_id
    forged["review_payload"]["run_id"] = forged_run_id
    forged["review_payload"]["items"] = [
        {
            "id": forged_item_id,
            "item_type": "report_section",
            "title": "Forged caller section",
            "allowed_actions": ["accept"],
            "recommended_action": "accept",
            "evidence": [],
            "data": {
                "section": "forged",
                "status": "assigned",
                "target_artifact": "report.docx",
                "target_path": "sections.forged.codex_comment",
            },
            "status": "needs_review",
        }
    ]
    forged["review_payload"]["item_count"] = 1
    forged["final_artifacts"]["run_id"] = forged_run_id
    forged["decisions"] = [{"item_id": forged_item_id, "action": "accept"}]
    before = _transaction_tree_state(output_dir)
    child_marker = _install_report_builder_fake_authority_child(
        monkeypatch,
        tmp_path,
    )

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": forged,
        },
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == (
        "Report Builder customer-run preflight returned an invalid result"
    )
    assert _transaction_tree_state(output_dir) == before
    assert not child_marker.exists()
    assert not list(tmp_path.glob(".generated-review-transaction-*"))
    assert not list(tmp_path.glob(".generated-review-commit-*"))
    assert not list(tmp_path.glob(".generated-review-recovery-*"))


def test_mcp_apply_rejects_current_source_receipt_staleness_without_mutation(
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    # Arrange
    output_dir, arguments = _report_transaction_fixture(tmp_path)
    source_path = _current_report_source(output_dir)
    source_path.write_text(
        "line,amount\nA,999\nB,20\n",
        encoding="utf-8",
    )
    before = _transaction_tree_state(output_dir)
    child_marker = _install_report_builder_fake_authority_child(
        monkeypatch,
        tmp_path,
    )

    # Act
    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": arguments,
        },
    )
    failure = response["result"]["structuredContent"]

    # Assert
    assert failure["ok"] is False
    assert failure["error"] == (
        "Report Builder persistence requires a running v2 customer-folder workflow run"
    )
    assert _transaction_tree_state(output_dir) == before
    assert not child_marker.exists()
    assert not list(tmp_path.glob(".generated-review-transaction-*"))
    assert not list(tmp_path.glob(".generated-review-commit-*"))
    assert not list(tmp_path.glob(".generated-review-recovery-*"))


def test_accepting_every_review_item_cannot_finalize_pending_numeric_measures(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "budget.csv"
    input_path.write_text(
        "line,amount,account_id\nA,10,1001\nB,20,1002\n",
        encoding="utf-8",
    )
    managed = _managed_report_run(tmp_path, input_path)
    output_dir = managed["output_dir"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    assert any(
        item["id"].startswith("numeric-measure-review-")
        for item in review_payload["items"]
    )

    result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": _nonblocking_decisions(review_payload),
            },
        },
    )["structuredContent"]
    updated_final = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    updated_applied = json.loads(
        (output_dir / "applied_decisions.json").read_text(encoding="utf-8")
    )

    assert updated_applied["decision_count"] == updated_applied["item_count"]
    assert result["application_status"] == "partial_review_applied"
    assert updated_final["status"] == "partial_review_applied"
    assert any(
        blocker.get("kind") == "numeric_measure_review"
        for blocker in updated_final["blockers"]
    )
    assert (
        "Use final_artifacts.json as the reviewed artifact gallery for handoff."
        not in updated_final["next_actions"]
    )
    assert not (output_dir / "numeric_evidence_ledger.json").exists()

    validation = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": json.loads(
                    (output_dir / "run_intake.json").read_text(encoding="utf-8")
                ),
                "review_payload": json.loads(
                    (output_dir / "review_payload.json").read_text(encoding="utf-8")
                ),
                "final_artifacts": updated_final,
            },
        },
    )
    assert validation["structuredContent"]["ok"] is True


def test_source_mapping_change_removes_stale_numeric_artifacts_and_references(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "report.xlsx"
    _save_workbook(input_path)

    def build_reviewed_recipe(core: Any, source: Path, work_dir: Path) -> Path:
        inspection = core.inspect_inputs(
            source,
            work_dir,
            report_type="management_report",
        )
        reviewed = core.review_numeric_measure_columns(
            inspection.inspection,
            inspection.suggested_recipe,
            section_key="income_statement",
            **_numeric_review_args(
                inspection.inspection,
                "report.xlsx::Income Statement",
                ["Actual", "Budget"],
            ),
            reviewer_ref="reviewer.pytest",
            reviewed_on="2026-07-24",
            numeric_locale="en",
            currency="EUR",
            unit="currency",
            scale="1",
            parse_policy="strict_all_nonblank_v1",
        )
        recipe_path = work_dir / "reviewed_recipe.json"
        core.write_json(recipe_path, reviewed)
        return recipe_path

    managed = _managed_report_run(
        tmp_path,
        input_path,
        recipe_builder=build_reviewed_recipe,
    )
    output_dir = managed["output_dir"]
    assert (output_dir / "numeric_evidence_ledger.json").is_file()
    assert (output_dir / "source_receipts.json").is_file()

    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    table_item = next(
        item
        for item in review_payload["items"]
        if item["item_type"] == "table_evidence"
        and item["data"]["section"] == "income_statement"
    )

    result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
                "decisions": [
                    {
                        "item_id": table_item["id"],
                        "action": "edit",
                        "edit_value": "report.xlsx::Cash Flow",
                    }
                ],
            },
        },
    )["structuredContent"]
    updated_analysis = json.loads(
        (output_dir / "report_analysis.json").read_text(encoding="utf-8")
    )
    updated_review = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    updated_final = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    income = next(
        section
        for section in updated_analysis["sections"]
        if section["section"] == "income_statement"
    )
    output_paths = {output["path"] for output in updated_final["outputs"]}

    assert result["application_status"] == "partial_review_applied"
    assert income["numeric_measure_status"] == "needs_review"
    assert income["numeric_columns"] == []
    assert updated_review["status"] == "ready_for_review_after_regeneration"
    assert "numeric_evidence" not in updated_review["source_artifacts"]
    assert "source_receipts" not in updated_review["source_artifacts"]
    assert not any(
        item["output_path"] in {"numeric_evidence_ledger.json", "source_receipts.json"}
        for item in updated_review["items"]
    )
    assert "numeric_evidence_ledger.json" not in output_paths
    assert "source_receipts.json" not in output_paths
    assert not (output_dir / "numeric_evidence_ledger.json").exists()
    assert not (output_dir / "source_receipts.json").exists()
    assert any(
        blocker.get("kind") == "source_mapping_review"
        for blocker in updated_final["blockers"]
    )

    validation = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": json.loads(
                    (output_dir / "run_intake.json").read_text(encoding="utf-8")
                ),
                "review_payload": updated_review,
                "final_artifacts": updated_final,
            },
        },
    )
    assert validation["structuredContent"]["ok"] is True


def test_zip_build_keeps_private_roots_and_raw_inputs_out_of_public_gallery(
    tmp_path: Path,
) -> None:
    core = load_core()
    archive_path = tmp_path / "inputs.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr(
            "schedules/budget.csv",
            "line,amount\nA,10\nB,20\n",
        )
    inspection_dir = tmp_path / "inspection"
    inspection = core.inspect_inputs(
        archive_path,
        inspection_dir,
        report_type="management_report",
    )
    table = inspection.inspection["tables"][0]
    recipe = inspection.suggested_recipe
    recipe["sections"]["budget"]["assigned_table"] = table["table_id"]
    reviewed = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key="budget",
        **_numeric_review_args(
            inspection.inspection,
            table["table_id"],
            ["amount"],
        ),
        reviewer_ref="reviewer.pytest",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency="EUR",
        unit="currency",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )
    recipe_path = tmp_path / "reviewed_recipe.json"
    core.write_json(recipe_path, reviewed)
    output_dir = tmp_path / "report"

    core.build_report(archive_path, output_dir, recipe_path=recipe_path)

    analysis = json.loads(
        (output_dir / "report_analysis.json").read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    budget = next(
        section for section in analysis["sections"] if section["section"] == "budget"
    )
    output_paths = {output["path"] for output in final_artifacts["outputs"]}
    public_payload_names = (
        "review_payload.json",
        "final_artifacts.json",
        "report_tables.json",
        "report_analysis.json",
        "report_audit.json",
        "used_recipe.json",
        "source_receipts.json",
    )
    public_payload_text = "\n".join(
        (output_dir / name).read_text(encoding="utf-8") for name in public_payload_names
    )

    assert budget["numeric_measure_status"] == "reviewed"
    assert budget["numeric_columns"][0]["sum"] == "30"
    assert output_paths <= {
        "report_tables.json",
        "report_tables.xlsx",
        "report_analysis.json",
        "report_draft.md",
        "report.docx",
        "report_audit.json",
        "used_recipe.json",
        "numeric_evidence_ledger.json",
        "source_receipts.json",
        "review_handoff.md",
    }
    assert not any(path.startswith("extracted_inputs/") for path in output_paths)
    assert "source_index.json" not in output_paths
    assert "review_integrity.json" not in output_paths
    assert (output_dir / "extracted_inputs").is_dir()
    assert (output_dir / "source_index.json").is_file()
    assert (output_dir / "review_integrity.json").is_file()
    assert tmp_path.as_posix() not in public_payload_text
    for output in final_artifacts["outputs"]:
        artifact_path = output_dir / output["path"]
        assert output["size_bytes"] == artifact_path.stat().st_size
        assert (
            output["sha256"] == hashlib.sha256(artifact_path.read_bytes()).hexdigest()
        )


def test_replaced_zip_inventory_contains_only_current_archive_members(
    tmp_path: Path,
) -> None:
    core = load_core()
    archive_path = tmp_path / "inputs.zip"
    output_dir = tmp_path / "inspection"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr("old.csv", "line,amount\nOld,10\n")
    first = core.inspect_inputs(
        archive_path,
        output_dir,
        report_type="management_report",
    )
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr("new.csv", "line,amount\nNew,20\n")

    second = core.inspect_inputs(
        archive_path,
        output_dir,
        report_type="management_report",
    )

    assert {table["source_file"] for table in first.inspection["tables"]} == {"old.csv"}
    assert {table["source_file"] for table in second.inspection["tables"]} == {
        "new.csv"
    }


def test_zip_rejects_duplicate_portable_canonical_member_paths(
    tmp_path: Path,
) -> None:
    core = load_core()
    archive_path = tmp_path / "duplicates.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr("Schedules/Budget.csv", "line,amount\nA,10\n")
        archive.writestr("schedules/budget.csv", "line,amount\nB,20\n")

    with pytest.raises(ValueError, match="duplicate canonical member paths"):
        core.inspect_inputs(
            archive_path,
            tmp_path / "inspection",
            report_type="management_report",
        )


def test_zip_review_integrity_binds_the_original_archive_bytes(tmp_path: Path) -> None:
    archive_path = tmp_path / "inputs.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr("budget.csv", "line,amount\nA,10\nB,20\n")
    managed = _managed_report_run(tmp_path, archive_path)
    output_dir = managed["output_dir"]
    run_intake = json.loads((output_dir / "run_intake.json").read_text())
    review_payload = json.loads((output_dir / "review_payload.json").read_text())
    final_artifacts = json.loads((output_dir / "final_artifacts.json").read_text())
    valid = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
            },
        },
    )
    assert valid["structuredContent"]["ok"] is True, valid["structuredContent"]
    before = _tree_snapshot(output_dir)
    source_index = json.loads((output_dir / "source_index.json").read_text())
    container_id = source_index["archive_manifests"][0]["container_artifact_id"]
    container = next(
        source
        for source in source_index["sources"]
        if source["artifact_id"] == container_id
    )
    container_root = Path(container["root_path"])
    if not container_root.is_absolute():
        container_root = output_dir.parent / container_root
    current_archive = container_root / container["receipt"]["path"]
    with zipfile.ZipFile(current_archive, "w") as archive:
        archive.writestr("budget.csv", "line,amount\nA,999\nB,20\n")

    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
            },
        },
    )
    failure = response["result"]["structuredContent"]

    assert failure["ok"] is False
    assert failure["error"] == "Report Builder persisted review authorization failed."
    assert _tree_snapshot(output_dir) == before


def test_rebuild_resets_prior_applied_state_and_revision_tree(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "budget.csv"
    input_path.write_text("line,amount\nA,10\nB,20\n", encoding="utf-8")
    output_dir = tmp_path / "report"
    core.build_report(input_path, output_dir, report_type="management_report")
    first_run = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )["run_id"]
    (output_dir / "applied_decisions.json").write_text(
        '{"run_id":"stale"}\n',
        encoding="utf-8",
    )
    revisions = output_dir / "revisions"
    revisions.mkdir()
    (revisions / "stale.txt").write_text("stale", encoding="utf-8")

    core.build_report(input_path, output_dir, report_type="management_report")
    second_run = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )["run_id"]
    ui_decisions = json.loads(
        (output_dir / "ui_decisions.json").read_text(encoding="utf-8")
    )

    assert second_run != first_run
    assert ui_decisions["run_id"] == second_run
    assert len(ui_decisions["review_payload_sha256"]) == 64
    assert not (output_dir / "applied_decisions.json").exists()
    assert not revisions.exists()


def test_skill_and_scripts_keep_codex_as_the_narrative_layer() -> None:
    skill_text = (
        ROOT / "plugins" / "report-builder" / "skills" / "report-builder" / "SKILL.md"
    ).read_text(encoding="utf-8")
    script_text = "\n".join(
        path.read_text(encoding="utf-8") for path in SCRIPT_DIR.glob("*.py")
    )

    assert "The user should not interact directly with CLI scripts" in skill_text
    assert "must not make direct OpenAI API calls" in skill_text
    assert "scripts/check_dependencies.py" in skill_text
    assert "it`, `en`, `fr`, `de`, and `es`" in skill_text
    assert "missing deterministic extraction script" in skill_text
    assert "Keep the improvement note local to chat or run artifacts." in skill_text
    assert "validate_report_builder_review" in skill_text
    assert "render_report_builder_review" in skill_text
    assert "ui://widget/report-builder-review.html" in skill_text
    assert "native Plan-mode" in skill_text
    assert "modules.llm" not in script_text
    assert "model_router" not in script_text
    assert "get_openai_client" not in script_text


def test_static_page_exposes_five_language_switch_and_prompts() -> None:
    page = (ROOT / "static" / "shared" / "report-builder" / "index.html").read_text(
        encoding="utf-8"
    )

    for snippet in (
        'data-lang="it"',
        'data-lang="en"',
        'data-lang="fr"',
        'data-lang="de"',
        'data-lang="es"',
        "Turn source tables into a reviewable Word report.",
        "Da tabelle sorgente a un report Word rivedibile.",
        "Transformer les tableaux source en rapport Word révisable.",
        "Quelltabellen in einen prüfbaren Word-Bericht verwandeln.",
        "Convierta tablas fuente en un informe Word revisable.",
        "Prepara una bozza DOCX da Excel, CSV e PDF leggibili.",
        "Ready prompts",
        "Prompt pronti",
        "File prodotti rivedibili",
        "Usa Genera report sui file in /percorso/report.",
        '"download.button": "Vera"',
    ):
        assert snippet in page


def test_mcp_review_server_validates_and_renders_report_payload() -> None:
    review_payload = {
        "schema_version": "1.0",
        "plugin": "report-builder",
        "workflow": "report-builder",
        "run_id": "report-builder-test-run",
        "review_type": "report_builder_review",
        "item_count": 3,
        "items": [
            {
                "id": "report-section-1",
                "item_type": "report_section",
                "title": "Income statement (assigned)",
                "source_path": "report.xlsx::Income Statement",
                "output_path": "report_draft.md",
                "allowed_actions": ["accept", "edit", "mark_unclear", "skip"],
                "recommended_action": "accept",
                "evidence": [{"kind": "section_status", "status": "assigned"}],
                "data": {},
                "status": "needs_review",
            },
            {
                "id": "table-evidence-1",
                "item_type": "table_evidence",
                "title": "Evidence table for Income statement",
                "source_path": "report.xlsx::Income Statement",
                "output_path": "report_tables.json",
                "allowed_actions": ["accept", "edit", "mark_unclear", "skip"],
                "recommended_action": "accept",
                "evidence": [{"kind": "table_evidence"}],
                "data": {},
                "status": "needs_review",
            },
            {
                "id": "artifact-1",
                "item_type": "report_artifact",
                "title": "Word report",
                "output_path": "report.docx",
                "allowed_actions": ["accept", "mark_unclear", "skip"],
                "recommended_action": "accept",
                "evidence": [{"kind": "artifact_status", "exists": True}],
                "data": {},
                "status": "needs_review",
            },
        ],
    }
    run_intake = {
        "schema_version": "1.0",
        "plugin": "report-builder",
        "workflow": "report-builder",
        "run_id": "report-builder-test-run",
    }
    ui_decisions = {
        "schema_version": "1.0",
        "plugin": "report-builder",
        "workflow": "report-builder",
        "run_id": "report-builder-test-run",
        "decisions": [],
        "status": "pending_review",
    }
    final_artifacts = {
        "schema_version": "1.0",
        "plugin": "report-builder",
        "workflow": "report-builder",
        "run_id": "report-builder-test-run",
        "outputs": [],
        "status": "written_pending_review",
    }

    tools = _call_mcp_server("tools/list")
    tool_names = {tool["name"] for tool in tools["tools"]}
    assert "validate_report_builder_review" in tool_names
    assert "render_report_builder_review" in tool_names

    validate_result = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "review_payload": review_payload,
                "run_intake": run_intake,
                "ui_decisions": ui_decisions,
                "final_artifacts": final_artifacts,
            },
        },
    )
    validation = json.loads(validate_result["content"][0]["text"])
    assert validation["ok"] is True
    assert validation["item_count"] == 3

    render_result = _call_mcp_server(
        "tools/call",
        {
            "name": "render_report_builder_review",
            "arguments": {
                "review_payload": review_payload,
                "run_intake": run_intake,
                "ui_decisions": ui_decisions,
                "final_artifacts": final_artifacts,
            },
        },
    )
    rendered = json.loads(render_result["content"][0]["text"])
    assert rendered["widget_type"] == "report_builder_review"
    assert (
        render_result["_meta"]["openai/outputTemplate"]
        == "ui://widget/report-builder-review.html"
    )

    resources = _call_mcp_server("resources/list")
    assert any(
        resource["uri"] == "ui://widget/report-builder-review.html"
        for resource in resources["resources"]
    )
    widget = _call_mcp_server(
        "resources/read", {"uri": "ui://widget/report-builder-review.html"}
    )
    assert "Build Report Review" in widget["contents"][0]["text"]


def test_in_memory_apply_cannot_finalize_numeric_review_issue() -> None:
    review_payload = {
        "schema_version": "1.0",
        "plugin": "report-builder",
        "workflow": "report-builder",
        "run_id": "report-builder-numeric-pending",
        "review_type": "report_builder_review",
        "items": [
            {
                "id": "numeric-measure-review-1",
                "item_type": "review_issue",
                "title": "Numeric measure review pending",
                "allowed_actions": ["mark_unclear", "skip"],
                "recommended_action": "mark_unclear",
                "evidence": [
                    {"kind": "numeric_measure_review_pending", "section": "budget"}
                ],
            }
        ],
        "item_count": 1,
        "status": "ready_for_review",
    }

    applied = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "run_intake": {"run_id": review_payload["run_id"]},
                "review_payload": review_payload,
                "decisions": [
                    {
                        "item_id": "numeric-measure-review-1",
                        "action": "skip",
                    }
                ],
            },
        },
    )["structuredContent"]

    assert applied["application_status"] == "partial_review_applied"
    assert applied["applied_decisions"]["numeric_measure_pending_review_count"] == 1
    assert applied["final_artifacts"]["blockers"] == [
        {
            "kind": "numeric_measure_review",
            "status": "needs_review",
            "pending_count": 1,
        }
    ]


def test_spanish_mcp_runtime_feedback_and_errors_without_persistence() -> None:
    review_payload = {
        "schema_version": "1.0",
        "plugin": "report-builder",
        "workflow": "report-builder",
        "run_id": "report-builder-es-runtime",
        "language": "es",
        "review_type": "report_builder_review",
        "items": [
            {
                "id": "section-es-1",
                "item_type": "report_section",
                "title": "Resultados",
                "allowed_actions": ["accept", "skip"],
                "recommended_action": "accept",
            }
        ],
        "item_count": 1,
        "status": "ready_for_review",
    }
    run_intake = {
        "run_id": review_payload["run_id"],
        "language": "es_ES",
    }
    decision = {"item_id": "section-es-1", "action": "accept"}

    initialized = _call_mcp_server("initialize", {"_meta": {"locale": "es-ES"}})
    validation_result = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {"review_payload": review_payload},
        },
    )
    save_result = _call_mcp_server(
        "tools/call",
        {
            "name": "save_report_builder_decisions",
            "arguments": {
                "review_payload": review_payload,
                "decisions": [decision],
            },
        },
    )
    apply_result = _call_mcp_server(
        "tools/call",
        {
            "name": "apply_report_builder_decisions",
            "arguments": {
                "run_intake": run_intake,
                "review_payload": review_payload,
                "decisions": [decision],
            },
        },
    )
    invalid_result = _call_mcp_server(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {"review_payload": {**review_payload, "items": "invalid"}},
        },
    )

    validation = validation_result["structuredContent"]
    saved = save_result["structuredContent"]
    applied = apply_result["structuredContent"]
    invalid = invalid_result["structuredContent"]

    assert "Use validate_report_builder_review antes" in initialized["instructions"]
    assert validation["message"].startswith("Los datos de revisión")
    assert saved["message"].startswith("Las decisiones son válidas")
    assert saved["persisted"] is False
    assert applied["message"].startswith("Las decisiones aplicadas son válidas")
    assert applied["persisted"] is False
    assert applied["final_artifacts"]["next_actions"][-1].startswith(
        "Use final_artifacts.json como galería"
    )
    assert invalid["error"] == "review_payload.items debe ser una matriz"


@pytest.mark.parametrize("failure_target", ["load_tables", "seal_review_integrity"])
def test_failed_rebuild_restores_exact_prior_output(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    failure_target: str,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    output_dir = tmp_path / "report"
    _save_workbook(input_path)
    recipe_path = _reviewed_numeric_recipe(core, input_path, tmp_path / "inspection")
    core.build_report(input_path, output_dir, recipe_path=recipe_path)
    (output_dir / "applied_decisions.json").write_text(
        '{"status":"reviewed"}\n',
        encoding="utf-8",
    )
    revisions = output_dir / "revisions"
    revisions.mkdir()
    (revisions / "accepted.md").write_text("accepted edit\n", encoding="utf-8")
    before = _tree_snapshot(output_dir)

    def fail(*args: Any, **kwargs: Any) -> None:
        del args, kwargs
        raise ValueError(f"injected {failure_target} failure")

    monkeypatch.setattr(core, failure_target, fail)

    with pytest.raises(ValueError, match=f"injected {failure_target} failure"):
        core.build_report(input_path, output_dir, recipe_path=recipe_path)

    assert _tree_snapshot(output_dir) == before


def test_build_rejects_symlink_output_without_mutating_target(tmp_path: Path) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    real_output = tmp_path / "real-output"
    linked_output = tmp_path / "linked-output"
    _save_workbook(input_path)
    real_output.mkdir()
    sentinel = real_output / "accepted-review.txt"
    sentinel.write_text("preserve me\n", encoding="utf-8")
    linked_output.symlink_to(real_output, target_is_directory=True)

    with pytest.raises(ValueError, match="cannot be a symbolic link"):
        core.build_report(input_path, linked_output)

    assert sentinel.read_text(encoding="utf-8") == "preserve me\n"
    assert sorted(path.name for path in real_output.iterdir()) == [sentinel.name]


def test_build_rejects_child_file_symlink_without_mutating_external_file(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    output_dir = tmp_path / "output"
    external = tmp_path / "external-analysis.json"
    _save_workbook(input_path)
    output_dir.mkdir()
    external.write_text('{"accepted":true}\n', encoding="utf-8")
    (output_dir / "report_analysis.json").symlink_to(external)
    before = external.read_bytes()

    with pytest.raises(ValueError, match="contains a symbolic link"):
        core.build_report(input_path, output_dir)

    assert external.read_bytes() == before
    assert (output_dir / "report_analysis.json").is_symlink()


def test_build_rejects_child_directory_symlink_without_mutating_external_tree(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    output_dir = tmp_path / "output"
    external = tmp_path / "external-revisions"
    _save_workbook(input_path)
    output_dir.mkdir()
    external.mkdir()
    (external / "accepted.md").write_text("accepted edit\n", encoding="utf-8")
    (output_dir / "revisions").symlink_to(external, target_is_directory=True)
    before = _tree_snapshot(external)

    with pytest.raises(ValueError, match="contains a symbolic link"):
        core.build_report(input_path, output_dir)

    assert _tree_snapshot(external) == before
    assert (output_dir / "revisions").is_symlink()


def test_build_rejects_child_hard_link_without_mutating_external_file(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "report.xlsx"
    output_dir = tmp_path / "output"
    external = tmp_path / "external-analysis.json"
    _save_workbook(input_path)
    output_dir.mkdir()
    external.write_text('{"accepted":true}\n', encoding="utf-8")
    (output_dir / "report_analysis.json").hardlink_to(external)
    before = external.read_bytes()

    with pytest.raises(ValueError, match="contains a hard-linked file"):
        core.build_report(input_path, output_dir)

    assert external.read_bytes() == before
    assert (output_dir / "report_analysis.json").read_bytes() == before


def test_zip_member_swap_between_extraction_and_capture_is_withheld(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    core = load_core()
    archive_path = tmp_path / "inputs.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr("nested/data.csv", "line,amount\nA,10\n")
    original_capture = core._capture_source
    swapped = False

    def swap_extracted_member(
        path: Path,
        *,
        identity_key: str | None = None,
    ) -> tuple[bytes, dict[str, Any]]:
        nonlocal swapped
        if (
            not swapped
            and path.suffix.lower() == ".csv"
            and "extracted_inputs" in path.parts
        ):
            path.write_text("line,amount\nA,999\n", encoding="utf-8")
            swapped = True
        return original_capture(path, identity_key=identity_key)

    monkeypatch.setattr(core, "_capture_source", swap_extracted_member)

    tables = core.load_tables(archive_path, tmp_path / "inspection")

    assert swapped is True
    assert len(tables) == 1
    assert tables[0]["kind"] == "error"
    assert "does not derive from the captured archive" in tables[0]["error"]
    assert not (tmp_path / "inspection" / "source_index.json").exists()


def test_all_nine_reviewed_measures_render_and_close_to_ledger(
    tmp_path: Path,
) -> None:
    core = load_core()
    input_path = tmp_path / "income_statement.csv"
    output_dir = tmp_path / "report"
    columns = [f"Measure {index}" for index in range(1, 10)]
    input_path.write_text(
        "Period," + ",".join(columns) + "\n"
        "Current," + ",".join(str(index) for index in range(1, 10)) + "\n"
        "Prior," + ",".join(str(index * 10) for index in range(1, 10)) + "\n",
        encoding="utf-8",
    )
    inspection = core.inspect_inputs(
        input_path,
        tmp_path / "inspection",
        report_type="annual_financial_statement",
    )
    recipe = inspection.suggested_recipe
    section_key = "income_statement"
    recipe["sections"][section_key]["assigned_table"] = input_path.name
    recipe = core.review_numeric_measure_columns(
        inspection.inspection,
        recipe,
        section_key=section_key,
        **_numeric_review_args(
            inspection.inspection,
            input_path.name,
            columns,
            excluded_cell_rows={column: {3} for column in columns},
        ),
        reviewer_ref="reviewer.nine-measure-test",
        reviewed_on="2026-07-24",
        numeric_locale="en",
        currency="USD",
        unit="currency",
        scale="1",
        parse_policy="strict_all_nonblank_v1",
    )
    recipe_path = tmp_path / "reviewed_recipe.json"
    core.write_json(recipe_path, recipe)

    core.build_report(input_path, output_dir, recipe_path=recipe_path)

    analysis = json.loads(
        (output_dir / "report_analysis.json").read_text(encoding="utf-8")
    )
    section = next(
        item for item in analysis["sections"] if item["section"] == section_key
    )
    ledger = json.loads(
        (output_dir / "numeric_evidence_ledger.json").read_text(encoding="utf-8")
    )
    markdown = (output_dir / "report_draft.md").read_text(encoding="utf-8")
    document = Document(output_dir / "report.docx")
    numeric_table = next(
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["Column", "Sum", "Currency", "Unit", "Scale"]
    )

    assert len(section["numeric_columns"]) == 9
    assert len(ledger["entries"]) == 9
    assert "Measure 9: sum 9 | Currency: USD | Unit: currency | Scale: 1" in markdown
    assert len(numeric_table.rows) == 10
    assert [cell.text for cell in numeric_table.rows[-1].cells] == [
        "Measure 9",
        "9",
        "USD",
        "currency",
        "1",
    ]


def test_build_rejects_direct_input_symlink(tmp_path: Path) -> None:
    core = load_core()
    source = tmp_path / "source.csv"
    linked_source = tmp_path / "linked.csv"
    source.write_text("line,amount\nA,10\n", encoding="utf-8")
    linked_source.symlink_to(source)

    with pytest.raises(ValueError, match="input path cannot be a symbolic link"):
        core.build_report(linked_source, tmp_path / "output")

    assert not (tmp_path / "output").exists()


def test_build_rejects_direct_input_hard_link(tmp_path: Path) -> None:
    core = load_core()
    source = tmp_path / "source.csv"
    linked_source = tmp_path / "linked.csv"
    source.write_text("line,amount\nA,10\n", encoding="utf-8")
    linked_source.hardlink_to(source)

    with pytest.raises(ValueError, match="input source cannot be hard-linked"):
        core.build_report(linked_source, tmp_path / "output")

    assert source.read_text(encoding="utf-8") == "line,amount\nA,10\n"
    assert not (tmp_path / "output").exists()


def test_non_text_pdf_is_a_non_dismissible_source_qualification_blocker(
    tmp_path: Path,
) -> None:
    core = load_core()
    source = tmp_path / "scanned.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with source.open("wb") as handle:
        writer.write(handle)
    output_dir = tmp_path / "report"

    core.build_report(source, output_dir)

    tables = json.loads((output_dir / "report_tables.json").read_text(encoding="utf-8"))
    review = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    source_blocker = next(
        item
        for item in review["items"]
        if item["id"].startswith("source-qualification-failure-")
    )
    assert tables["tables"][0]["kind"] == "error"
    assert "unsupported_source_layout" in tables["tables"][0]["error"]
    assert source_blocker["recommended_action"] == "request_more_documents"
    assert source_blocker["allowed_actions"] == [
        "mark_unclear",
        "request_more_documents",
    ]
    assert not (output_dir / "numeric_evidence_ledger.json").exists()


def test_mcp_child_failure_omits_traceback_and_absolute_paths(tmp_path: Path) -> None:
    source = tmp_path / "budget.csv"
    source.write_text("line,amount\nA,10\n", encoding="utf-8")
    managed = _managed_report_run(tmp_path, source)
    output_dir = managed["output_dir"]
    run_intake = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )
    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    (output_dir / "review_integrity.json").write_text("[]\n", encoding="utf-8")

    response = _call_mcp_server_response(
        "tools/call",
        {
            "name": "validate_report_builder_review",
            "arguments": {
                "client_engagement": managed["context_path"].as_posix(),
                "run_intake": run_intake,
                "review_payload": review_payload,
                "final_artifacts": final_artifacts,
            },
        },
    )

    failure = response["result"]["structuredContent"]
    assert failure["ok"] is False
    assert failure["error"] == "Report Builder persisted review authorization failed."
    assert "Traceback" not in failure["error"]
    assert tmp_path.as_posix() not in failure["error"]
    assert "\n" not in failure["error"]


def test_csv_formula_like_headers_are_literal_strings_in_generated_xlsx(
    tmp_path: Path,
) -> None:
    core = load_core()
    source = tmp_path / "income_statement.csv"
    source.write_text(
        "=2+2,Actual\n@command,10\n+command,20\n",
        encoding="utf-8",
    )
    output_dir = tmp_path / "report"

    core.build_report(source, output_dir)

    workbook = openpyxl.load_workbook(
        output_dir / "report_tables.xlsx",
        data_only=False,
    )
    preview = next(
        sheet
        for sheet in workbook.worksheets
        if sheet.title not in {"summary", "numeric_evidence"}
    )
    assert preview["A1"].value == "=2+2"
    assert preview["A1"].data_type == "s"
    assert preview["A2"].value == "[numeric source value withheld]"
    assert preview["A2"].data_type == "s"
    assert not any(cell.data_type == "f" for row in preview.iter_rows() for cell in row)
