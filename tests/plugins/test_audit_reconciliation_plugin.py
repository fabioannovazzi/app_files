from __future__ import annotations

import hashlib
import importlib._bootstrap_external
import importlib.util
import json
import os
import shutil
import stat
import subprocess
import sys
from pathlib import Path
from types import ModuleType
from typing import Any

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from scripts.validate_plugin_review_contract import validate_contract

sys.dont_write_bytecode = True
sys.pycache_prefix = "/dev/null/audit-reconciliation-tests"

ROOT = Path(__file__).resolve().parents[2]
PLUGIN_ROOT = ROOT / "plugins" / "audit-reconciliation"
SCRIPT_DIR = PLUGIN_ROOT / "scripts"
REVIEW_SERVER_PATH = SCRIPT_DIR / "review_server.py"
ASSURANCE_PATH = SCRIPT_DIR / "audit_assurance.py"
RECONCILIATION_WORKFLOW_PATH = SCRIPT_DIR / "reconciliation_workflow.py"
RAW_INPUT_RUNNER_PATH = SCRIPT_DIR / "raw_input_runner.py"
CHECK_DEPENDENCIES_PATH = SCRIPT_DIR / "check_dependencies.py"
MISSING_EVIDENCE_PATH = SCRIPT_DIR / "build_missing_evidence_requests.py"
MCP_SERVER_PATH = PLUGIN_ROOT / "mcp" / "server.cjs"


def _load_customer_ledger() -> ModuleType:
    path = ROOT / "plugins" / "studio-archive" / "scripts" / "client_ledger.py"
    module_name = "test_audit_reconciliation_customer_ledger"
    existing = sys.modules.get(module_name)
    if existing is not None:
        return existing
    spec = importlib.util.spec_from_file_location(module_name, path)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def _running_audit_output(tmp_path: Path) -> Path:
    ledger = _load_customer_ledger()
    client_root = tmp_path / "Audit Customer"
    client_root.mkdir()
    client_id = "client_333333333333333333333333"
    ledger.create_client_manifest(client_root, client_id)
    engagement = ledger.create_engagement(client_root, client_id, "Audit review")
    source = tmp_path / "audit-source.txt"
    source.write_text("audit source\n", encoding="utf-8")
    imported = ledger.import_document(
        client_root,
        client_id,
        engagement["engagement_id"],
        source,
        "source",
    )
    prepared = ledger.prepare_run(
        client_root,
        client_id,
        engagement["engagement_id"],
        "audit-reconciliation",
        "test-version",
        input_ids=[imported["receipt"]["input_id"]],
    )
    running = ledger.start_run(
        client_root,
        engagement["engagement_id"],
        prepared["run"]["run_id"],
    )
    return Path(running["output_dir"])


def _customer_run_id(output_dir: Path) -> str:
    candidate = output_dir.resolve()
    while candidate != candidate.parent:
        context_path = candidate / "context.json"
        if context_path.is_file():
            return str(json.loads(context_path.read_text(encoding="utf-8"))["run_id"])
        candidate = candidate.parent
    raise AssertionError("customer-run context is unavailable")


def _customer_context_path(output_dir: Path) -> Path:
    candidate = output_dir.resolve()
    while candidate != candidate.parent:
        context_path = candidate / "context.json"
        if context_path.is_file():
            return context_path
        candidate = candidate.parent
    raise AssertionError("customer-run context is unavailable")


def _rename_customer_output(output_dir: Path) -> tuple[Path, Path, Path]:
    context_path = _customer_context_path(output_dir)
    client_root = context_path.parents[5]
    renamed_root = client_root.with_name(f"{client_root.name} Renamed")
    relative_output = output_dir.relative_to(client_root)
    relative_context = context_path.relative_to(client_root)
    client_root.rename(renamed_root)
    return (
        renamed_root / relative_output,
        renamed_root / relative_context,
        output_dir,
    )


def load_review_session() -> Any:
    load_reconciliation_workflow()
    return sys.modules["review_session"]


def load_raw_input_runner() -> Any:
    if str(SCRIPT_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPT_DIR))
    spec = importlib.util.spec_from_file_location(
        "audit_reconciliation_raw_input_runner", RAW_INPUT_RUNNER_PATH
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def load_review_server() -> Any:
    if str(SCRIPT_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPT_DIR))
    spec = importlib.util.spec_from_file_location(
        "audit_reconciliation_review_server", REVIEW_SERVER_PATH
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def load_assurance() -> Any:
    if str(SCRIPT_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPT_DIR))
    spec = importlib.util.spec_from_file_location(
        "audit_reconciliation_plugin_assurance",
        ASSURANCE_PATH,
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def load_reconciliation_workflow() -> Any:
    if str(SCRIPT_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPT_DIR))
    spec = importlib.util.spec_from_file_location(
        "audit_reconciliation_plugin_workflow",
        RECONCILIATION_WORKFLOW_PATH,
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def load_check_dependencies() -> Any:
    if str(SCRIPT_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPT_DIR))
    spec = importlib.util.spec_from_file_location(
        "audit_reconciliation_check_dependencies", CHECK_DEPENDENCIES_PATH
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def load_missing_evidence_requests() -> Any:
    if str(SCRIPT_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPT_DIR))
    spec = importlib.util.spec_from_file_location(
        "audit_reconciliation_missing_evidence", MISSING_EVIDENCE_PATH
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def test_missing_evidence_request_workbook_is_spanish(tmp_path: Path) -> None:
    module = load_missing_evidence_requests()
    pack = module.build_missing_evidence_request_pack(
        [
            {
                "row_id": "row-1",
                "reconciliation_status": "probable_payment",
                "document": "FAC-001",
                "amount": "1250.00",
            }
        ],
        entity_name="Entidad Demo",
        counterparty_name="Cliente Demo",
        cutoff_date="2026-12-31",
        language="es",
    )
    output_path = module.write_missing_evidence_workbook(
        tmp_path / "solicitudes.xlsx", pack
    )

    workbook = load_workbook(output_path, read_only=True, data_only=True)
    assert pack.language == "es"
    assert pack.instructions[0]["principle"] == (
        "No vuelva a solicitar documentos ya adquiridos"
    )
    assert "instrucciones" in workbook.sheetnames
    assert "pagos_probables" in workbook.sheetnames
    assert workbook["pagos_probables"]["A1"].value == "id_línea_papel_trabajo"
    assert workbook["pagos_probables"]["J2"].value.startswith(
        "Confirmación de la asignación"
    )


def _call_mcp_server(
    messages: list[dict[str, object]],
    *,
    server_path: Path = MCP_SERVER_PATH,
    env: dict[str, str] | None = None,
) -> list[dict[str, object]]:
    node = shutil.which("node")
    if node is None:
        pytest.skip(
            "Node.js is required to exercise the Audit Reconciliation MCP server."
        )
    completed = subprocess.run(
        [node, str(server_path), "--stdio"],
        input="\n".join(json.dumps(message) for message in messages) + "\n",
        capture_output=True,
        text=True,
        check=True,
        timeout=10,
        env={**os.environ, **(env or {})},
    )
    return [json.loads(line) for line in completed.stdout.splitlines() if line.strip()]


def _copy_audit_mcp_runtime(tmp_path: Path) -> Path:
    plugin_copy = tmp_path / "audit-reconciliation"
    shutil.copytree(
        PLUGIN_ROOT,
        plugin_copy,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    shutil.copytree(
        ROOT / "plugins" / "_shared" / "vendor" / "modules" / "vera_assurance",
        tmp_path / "_shared" / "vendor" / "modules" / "vera_assurance",
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    return plugin_copy


def _audit_mcp_surface_message(surface: str) -> dict[str, object]:
    base: dict[str, object] = {"jsonrpc": "2.0", "id": 1, "method": surface}
    if surface == "initialize":
        base["params"] = {"protocolVersion": "2024-11-05"}
    elif surface == "resources/read":
        base["params"] = {"uri": "ui://widget/audit-reconciliation-review.html"}
    elif surface.startswith("tool:"):
        base["method"] = "tools/call"
        base["params"] = {
            "name": surface.removeprefix("tool:"),
            "arguments": {},
        }
    return base


def _write_workbook(
    path: Path,
    sheet_names: list[str],
    *,
    headers_by_sheet: dict[str, list[str]] | None = None,
    rows_by_sheet: dict[str, list[list[object]]] | None = None,
) -> None:
    workbook = Workbook()
    workbook.remove(workbook.active)
    for sheet_name in sheet_names:
        sheet = workbook.create_sheet(sheet_name)
        headers = (headers_by_sheet or {}).get(sheet_name, ["Campo", "Valore"])
        sheet.append(headers)
        rows = (rows_by_sheet or {}).get(sheet_name)
        if rows is None:
            rows = [["fixture" for _header in headers]]
        for row in rows:
            sheet.append(row)
    workbook.save(path)


def _write_docx(path: Path, title: str, required_text: list[str] | None = None) -> None:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("Fixture document for native artifact validation.")
    for fragment in required_text or []:
        document.add_paragraph(fragment)
    document.save(path)


def test_review_session_writes_audit_reconciliation_contract(tmp_path: Path) -> None:
    review_session = load_review_session()
    output_dir = tmp_path / "out"
    output_dir.mkdir()
    audit_workbook = output_dir / "riconciliazione_audit.xlsx"
    accountant_workbook = output_dir / "scheda_operativa_commercialista.xlsx"
    word_report = output_dir / "relazione_riconciliazione_audit.docx"
    missing_requests = output_dir / "richieste_mirate_evidenze.xlsx"
    _write_workbook(
        audit_workbook,
        [
            "Indice",
            "Assunzioni",
            "Problemi elaborazione fonti",
            "Dettaglio riconciliazione",
            "Sintesi",
            "Controlli",
            "Revisione Codex",
        ],
        headers_by_sheet={
            "Indice": ["Foglio", "Righe"],
            "Assunzioni": ["Campo", "Valore"],
            "Dettaglio riconciliazione": [
                "Documento",
                "Esito riconciliazione",
                "ID riga",
            ],
        },
        rows_by_sheet={
            "Indice": [
                ["Assunzioni", 1],
                ["Inventario fonti", 1],
                ["Problemi elaborazione fonti", 0],
                ["Righe normalizzate", 0],
                ["Dettaglio riconciliazione", 3],
            ],
            "Assunzioni": [["Currency", "EUR"]],
            "Dettaglio riconciliazione": [["INV-1", "closed", "row-1"]],
        },
    )
    _write_workbook(
        accountant_workbook,
        ["Legenda", "Scheda operativa", "Dettaglio riscontri"],
        headers_by_sheet={
            "Legenda": ["campo", "valore"],
            "Scheda operativa": [
                "id dettaglio",
                "partita",
                "stato riscontro",
                "azione richiesta",
            ],
            "Dettaglio riscontri": [
                "id dettaglio",
                "partita",
                "tipo evidenza",
                "riferimento fonte",
            ],
        },
        rows_by_sheet={
            "Legenda": [
                [
                    "Scopo",
                    (
                        "Scheda operativa riga-per-riga: data pagamento/incasso, "
                        "fonte, modalita, compensazione, stato, confidenza e "
                        "azione richiesta."
                    ),
                ],
                ["Righe", 3],
            ],
            "Scheda operativa": [
                ["R0001", "INV-1", "Chiuso", "Verificare fonte"],
            ],
        },
    )
    required_word_text = [
        "Sintesi esecutiva",
        "Perimetro e metodo",
        "Come leggere gli esiti",
        "Controlli automatici",
        "Revisione manuale Codex",
        "Limiti della procedura",
        "Rinvio al file Excel",
    ]
    _write_docx(word_report, "Relazione riconciliazione", required_word_text)
    _write_workbook(missing_requests, ["Richieste"])

    run_intake = review_session.write_run_intake(
        output_dir,
        assumptions={
            "scope_year": 2025,
            "cutoff_date": "2025-12-31",
            "currency": "EUR",
            "report_language": "it",
            "document_language": "it",
            "factoring_pro_soluto_closes_item": True,
        },
        source_inventory=[{"source_file": "open_items.xlsx"}],
        language="it",
        source_hint="open_items.xlsx",
    )
    result = {
        "excel_path": str(audit_workbook),
        "accountant_report_path": str(accountant_workbook),
        "word_path": str(word_report),
        "assumptions": {"currency": "EUR"},
        "reconciliation_rows": [
            {
                "record_id": "row-1",
                "document_no": "INV-1",
                "reconciliation_status": "closed",
            },
            {
                "record_id": "row-2",
                "document_no": "INV-2",
                "reconciliation_status": "needs_evidence",
            },
            {
                "record_id": "row-3",
                "document_no": "INV-3",
                "reconciliation_status": "unresolved",
            },
        ],
        "review_rows": [
            {
                "review_id": "review:closed",
                "record_id": "row-1",
                "document_no": "INV-1",
                "amount": "120.50",
                "deterministic_status": "closed",
                "deterministic_rule": "external_bank_match",
                "review_status": "PENDING",
                "review_selection_reason": "mandatory_closure_evidence",
                "review_instruction": "Verify source evidence.",
                "review_notes": "",
                "source_file": "open_items.xlsx",
                "source_row": "2",
            },
            {
                "review_id": "review:missing",
                "record_id": "row-2",
                "document_no": "INV-2",
                "amount": "88.00",
                "deterministic_status": "needs_evidence",
                "deterministic_rule": "payment_order_only",
                "review_status": "PENDING",
                "review_selection_reason": "risk_flag",
                "review_instruction": "Request missing evidence.",
                "review_notes": "",
            },
        ],
        "checks": [
            {"check": "row_count", "status": "PASS", "actual": 3, "expected": 3},
            {
                "check": "codex_review_complete",
                "status": "FAIL",
                "actual": "PENDING",
                "expected": "PASS",
            },
        ],
        "checks_pass": False,
        "bank_allocation_candidates": [{"record_id": "candidate-1"}],
        "account_rollforward_check": [
            {
                "account": "TOTAL",
                "account_name": "Conti confrontati",
                "status": "Difference",
                "opening_difference_journal_minus_ledger": "0.00",
                "closing_difference_journal_minus_ledger": "2236.67",
                "review_note": "Journal and ledger closing balances differ.",
            }
        ],
    }
    (output_dir / "codex_review_packet.json").write_text(
        json.dumps(result["review_rows"], ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    run_intake_payload = json.loads(run_intake.path.read_text(encoding="utf-8"))
    dependency_check = run_intake_payload["dependency_check"]

    session = review_session.write_review_session_artifacts(
        output_dir,
        run_id=run_intake.run_id,
        run_intake_path=run_intake.path,
        result=result,
        source_inventory=[{"source_file": "open_items.xlsx"}],
        missing_evidence_requests_path=missing_requests,
        language="it",
    )

    review_payload = json.loads(session.review_payload_path.read_text(encoding="utf-8"))
    ui_decisions = json.loads(session.ui_decisions_path.read_text(encoding="utf-8"))
    final_artifacts = json.loads(
        session.final_artifacts_path.read_text(encoding="utf-8")
    )
    review_html = session.review_html_path.read_text(encoding="utf-8")
    artifact_card = session.artifact_card_path.read_text(encoding="utf-8")

    assert review_payload["plugin"] == "audit-reconciliation"
    assert review_payload["run_id"] == run_intake.run_id
    assert review_payload["review_type"] == "audit_reconciliation_review"
    assert review_payload["item_count"] == len(review_payload["items"])
    item_types = {item["item_type"] for item in review_payload["items"]}
    assert {
        "closure_evidence_review",
        "missing_evidence_review",
        "check_exception",
        "workpaper_artifact",
        "report_artifact",
        "evidence_request_artifact",
    } <= item_types
    assert review_payload["summary"]["reconciliation_row_count"] == 3
    assert review_payload["summary"]["failed_check_count"] == 1
    assert review_payload["summary"]["checks_pass"] is False
    assert review_payload["summary"]["rollforward_exception_count"] == 1
    assert review_payload["summary"]["rollforward_exceptions"][0] == {
        "account": "TOTAL",
        "account_name": "Conti confrontati",
        "status": "Difference",
        "opening_difference": "0.00",
        "closing_difference": "2236.67",
        "review_note": "Journal and ledger closing balances differ.",
    }
    closed_item = next(
        item for item in review_payload["items"] if item["id"] == "review:closed"
    )
    assert closed_item["data"]["target_artifact"] == "codex_review_packet.json"
    assert closed_item["data"]["target_id_field"] == "review_id"
    assert closed_item["data"]["target_record_id"] == "review:closed"
    assert closed_item["data"]["target_field"] == "review_notes"
    assert dependency_check["status"] != "not_run_by_script"
    assert "checked_at" in dependency_check
    assert "requirement_files" in dependency_check or "note" in dependency_check
    assert ui_decisions["status"] == "pending_review"
    assert session.review_html_path.name == "review_ui.html"
    assert session.artifact_card_path.name == "artifact_card.md"
    assert "window.openai = { toolOutput:" in review_html
    assert run_intake.run_id in review_html
    assert '"item_count": 7' in review_html
    assert "Review safeguards" in review_html
    assert "review-safeguards" in review_html
    assert "renderReviewSafeguards" in review_html
    assert "Execution provenance" in review_html
    assert "execution-provenance" in review_html
    assert "local_codex_workspace" in review_html
    assert "scripts/review_server.py" in artifact_card
    assert "ui_decisions.json" in artifact_card
    assert "final_artifacts.json" in artifact_card
    assert final_artifacts["status"] == "written_pending_review"
    assert any(
        "Account roll-forward has exception rows" in caveat
        for caveat in final_artifacts["caveats"]
    )
    assert any(
        output["path"] == "richieste_mirate_evidenze.xlsx"
        for output in final_artifacts["outputs"]
    )
    packet_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "codex_review_packet.json"
    )
    assert packet_output["required_columns"] == ["review_id", "review_notes"]
    audit_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "riconciliazione_audit.xlsx"
    )
    accountant_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "scheda_operativa_commercialista.xlsx"
    )
    word_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "relazione_riconciliazione_audit.docx"
    )
    assert audit_output["artifact_role"] == "audit_workpaper"
    assert audit_output["required_sheets"] == [
        "Indice",
        "Assunzioni",
        "Problemi elaborazione fonti",
        "Dettaglio riconciliazione",
        "Sintesi",
        "Controlli",
        "Revisione Codex",
    ]
    assert "required_sheets" in audit_output["qa_checks"]
    assert audit_output["required_sheet_headers"] == {
        "Indice": ["Foglio", "Righe"],
        "Assunzioni": ["Campo", "Valore"],
    }
    assert audit_output["required_cells"] == {
        "Indice": {
            "A1": "Foglio",
            "B1": "Righe",
            "A2": "Assunzioni",
            "A6": "Dettaglio riconciliazione",
        },
        "Assunzioni": {
            "A1": "Campo",
            "B1": "Valore",
            "A2": "Currency",
            "B2": "EUR",
        },
        "Dettaglio riconciliazione": {
            "A1": "Documento",
            "A2": "INV-1",
            "C1": "ID riga",
            "C2": "row-1",
        },
    }
    assert "required_sheet_headers" in audit_output["qa_checks"]
    assert "required_cells" in audit_output["qa_checks"]
    assert accountant_output["artifact_role"] == "accountant_workbook"
    assert accountant_output["required_sheets"] == [
        "Legenda",
        "Scheda operativa",
        "Dettaglio riscontri",
    ]
    assert accountant_output["required_sheet_headers"] == {
        "Legenda": ["campo", "valore"],
        "Scheda operativa": [
            "id dettaglio",
            "partita",
            "stato riscontro",
            "azione richiesta",
        ],
        "Dettaglio riscontri": [
            "id dettaglio",
            "partita",
            "tipo evidenza",
            "riferimento fonte",
        ],
    }
    assert accountant_output["required_cells"] == {
        "Legenda": {
            "A1": "campo",
            "B1": "valore",
            "A3": "Righe",
            "B3": "3",
        },
        "Scheda operativa": {
            "A1": "id dettaglio",
            "B1": "partita",
            "A2": "R0001",
            "B2": "INV-1",
        },
        "Dettaglio riscontri": {"A1": "id dettaglio", "B1": "partita"},
    }
    assert "required_cells" in accountant_output["qa_checks"]
    assert word_output["artifact_role"] == "word_report"
    assert word_output["required_text"] == required_word_text
    assert "word_document_xml" in word_output["qa_checks"]
    assert "required_text" in word_output["qa_checks"]
    assert any(
        output["path"] == "review_ui.html" for output in final_artifacts["outputs"]
    )
    contract_report = validate_contract(
        output_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()
    assert any(
        output["path"] == "artifact_card.md" for output in final_artifacts["outputs"]
    )
    review_handoff = final_artifacts["review_handoff"]
    assert review_handoff["primary"] == "local_browser_server"
    assert review_handoff["status"] == "browser_review_required"
    assert review_handoff["required_before_final_delivery"] is True
    assert review_handoff["server"]["script"] == "scripts/review_server.py"
    assert review_handoff["server"]["host"] == "127.0.0.1"
    assert review_handoff["server"]["port"] == "auto"
    assert review_handoff["server"]["opens"] == "system_browser"
    assert review_handoff["server"]["required"] is True
    assert "scripts/review_server.py" in review_handoff["server"]["command"]
    assert review_handoff["server"]["writes"] == [
        "ui_decisions.json",
        "applied_decisions.json",
        "final_artifacts.json",
    ]
    assert review_handoff["artifact_card"] == {
        "path": "artifact_card.md",
        "required": True,
        "announce_to_user": True,
    }
    assert review_handoff["mcp"]["status"] == "optional_integrated_surface"
    assert review_handoff["mcp"]["tool_sequence"] == [
        "validate_audit_reconciliation_review",
        "render_audit_reconciliation_review",
    ]
    assert review_handoff["fallback"]["artifact"] == "review_ui.html"
    assert review_handoff["fallback"]["persistence"] == "copy_or_download_json"
    assert (
        "Open the browser review server with scripts/review_server.py"
        in final_artifacts["next_actions"][0]
    )
    assert (
        "Use the browser page to save or apply decisions"
        in final_artifacts["next_actions"][1]
    )


def _write_review_server_fixture(output_dir: Path) -> None:
    output_dir.mkdir(exist_ok=True)
    run_id = (
        _customer_run_id(output_dir)
        if (output_dir.parent / "context.json").is_file()
        else "audit-reconciliation-test-run"
    )
    (output_dir / "codex_review_packet.json").write_text(
        json.dumps(
            [
                {
                    "review_id": "review:closed",
                    "record_id": "row-1",
                    "document_no": "INV-1",
                    "review_notes": "",
                }
            ],
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    review_payload = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": run_id,
        "review_type": "audit_reconciliation_review",
        "source_paths": ["open_items.xlsx"],
        "items": [
            {
                "id": "review:closed",
                "item_type": "closure_evidence_review",
                "title": "INV-1 | 120.50 | closed",
                "source_path": "open_items.xlsx; row 2",
                "output_path": "codex_review_packet.json",
                "allowed_actions": ["accept", "edit", "mark_unclear", "skip"],
                "recommended_action": "accept",
                "evidence": [],
                "data": {
                    "target_artifact": "codex_review_packet.json",
                    "target_id_field": "review_id",
                    "target_record_id": "review:closed",
                    "target_field": "review_notes",
                },
                "status": "needs_review",
            },
            {
                "id": "review:missing",
                "item_type": "missing_evidence_review",
                "title": "INV-2 | 88.00 | needs evidence",
                "source_path": "open_items.xlsx; row 3",
                "output_path": "richieste_mirate_evidenze.xlsx",
                "allowed_actions": ["accept", "request_more_documents", "skip"],
                "recommended_action": "request_more_documents",
                "evidence": [],
                "data": {"target_artifact": "richieste_mirate_evidenze.xlsx"},
                "status": "needs_review",
            },
        ],
        "item_count": 2,
        "columns": [],
        "source_artifacts": {},
        "allowed_actions": [
            "accept",
            "reject",
            "edit",
            "mark_unclear",
            "request_more_documents",
            "skip",
        ],
        "status": "ready_for_review",
        "summary": {"review_item_count": 2},
    }
    run_intake = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": run_id,
        "created_at": "2026-01-01T00:00:00Z",
        "language": "it",
        "input_paths": ["open_items.xlsx"],
        "output_dir": output_dir.as_posix(),
        "inferred_task": "audit_reconciliation_review",
        "assumptions": {},
        "unresolved_questions": [],
        "dependency_check": {"status": "not_run"},
        "data_posture": {
            "local_files_read": ["open_items.xlsx"],
            "external_connectors_used": [],
            "upload_paths_used": [],
            "remote_sql_execution_used": False,
            "hosted_notebook_execution_used": False,
        },
        "execution_trace": [
            {
                "step_id": "audit_reconciliation_run",
                "kind": "deterministic_review_session",
                "status": "passed",
                "execution_location": "local_codex_workspace",
                "command": [
                    "python",
                    "plugins/audit-reconciliation/scripts/raw_input_runner.py",
                ],
                "inputs": ["open_items.xlsx"],
                "outputs": [
                    "review_payload.json",
                    "artifact_card.md",
                    "codex_review_packet.json",
                    "final_artifacts.json",
                ],
            }
        ],
    }
    ui_decisions = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": run_id,
        "decided_at": None,
        "decision_source": "not_collected",
        "review_payload_path": "review_payload.json",
        "decisions": [],
        "decision_count": 0,
        "status": "pending_review",
    }
    final_artifacts = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": run_id,
        "outputs": [{"path": "artifact_card.md", "kind": "md", "status": "written"}],
        "review_handoff": {
            "primary": "local_browser_server",
            "artifact_card": {"path": "artifact_card.md", "required": True},
        },
        "caveats": [],
        "next_actions": [],
        "status": "written_pending_review",
    }
    (output_dir / "artifact_card.md").write_text(
        "# Audit Reconciliation Review\n",
        encoding="utf-8",
    )
    for name, payload in {
        "run_intake.json": run_intake,
        "review_payload.json": review_payload,
        "ui_decisions.json": ui_decisions,
        "final_artifacts.json": final_artifacts,
    }.items():
        (output_dir / name).write_text(
            json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )


def test_review_server_saves_local_browser_decisions(tmp_path: Path) -> None:
    review_server = load_review_server()
    output_dir = _running_audit_output(tmp_path)
    _write_review_server_fixture(output_dir)

    result = review_server.save_decisions(
        output_dir,
        {
            "decisions": [
                {
                    "item_id": "review:closed",
                    "action": "accept",
                    "reviewer_note": "Evidence checked.",
                }
            ],
            "decision_source": "mcp_widget",
        },
    )

    saved = json.loads((output_dir / "ui_decisions.json").read_text(encoding="utf-8"))
    assert result["ok"] is True
    assert result["persisted"] is True
    assert saved["decision_source"] == "local_review_server"
    assert saved["decision_count"] == 1
    assert saved["item_count"] == 2
    assert saved["status"] == "partial_review"
    assert saved["decisions"][0]["status"] == "accepted"


def test_review_server_rejects_non_loopback_host(tmp_path: Path) -> None:
    review_server = load_review_server()
    output_dir = tmp_path / "out"
    _write_review_server_fixture(output_dir)

    with pytest.raises(ValueError, match="loopback"):
        review_server.serve_review(
            output_dir,
            host="0.0.0.0",
            open_browser=False,
        )


def test_review_server_renders_local_browser_bridge(tmp_path: Path) -> None:
    review_server = load_review_server()
    output_dir = tmp_path / "out"
    _write_review_server_fixture(output_dir)

    html = review_server.render_review_html(output_dir)

    assert "/api/call-tool" in html
    assert "async callTool" in html
    assert '"can_persist": true' in html
    assert '"fallback": "local_review_server"' in html
    assert "audit_reconciliation_review" in html
    assert "Review safeguards" in html
    assert "review-safeguards" in html
    assert "renderReviewSafeguards" in html
    assert "safeguardLocalExecution" in html
    assert "safeguardExternalRoute" in html
    assert "safeguardBoundedPayload" in html
    assert "safeguardDecisionPersistence" in html
    assert "safeguardFinalArtifacts" in html
    assert "Execution provenance" in html
    assert "execution-provenance" in html
    assert '"execution_location": "local_codex_workspace"' in html


def test_review_server_applies_local_browser_decisions(tmp_path: Path) -> None:
    review_server = load_review_server()
    output_dir = tmp_path / "out"
    _write_review_server_fixture(output_dir)

    result = review_server.apply_decisions(
        output_dir,
        {
            "decisions": [
                {
                    "item_id": "review:closed",
                    "action": "edit",
                    "edit_value": "Reviewer confirmed external bank support.",
                },
                {"item_id": "review:missing", "action": "accept"},
            ],
        },
    )

    applied = json.loads(
        (output_dir / "applied_decisions.json").read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    assert result["ok"] is True
    assert result["application_status"] == "blocked"
    assert result["run_intake_path"] == (output_dir / "run_intake.json").as_posix()
    assert applied["decision_source"] == "local_review_server"
    assert applied["decision_count"] == 2
    assert applied["blocker_count"] == 1
    assert applied["completion_blockers"][0]["kind"] == "assurance_replay_required"
    assert applied["structured_update_count"] == 1
    assert applied["structured_update_paths"] == ["codex_review_packet.json"]
    assert applied["effects"][0]["structured_update"] == {
        "id_field": "review_id",
        "record_id": "review:closed",
        "target_field": "review_notes",
        "records_key": "",
        "updated_rows": 1,
    }
    packet = json.loads(
        (output_dir / "codex_review_packet.json").read_text(encoding="utf-8")
    )
    assert packet[0]["review_notes"] == "Reviewer confirmed external bank support."
    assert final_artifacts["status"] == "blocked"
    assert final_artifacts["review_handoff"]["primary"] == "local_browser_server"
    assert final_artifacts["review_application"]["decision_count"] == 2
    assert final_artifacts["review_application"]["structured_update_count"] == 1
    packet_output = next(
        output
        for output in final_artifacts["outputs"]
        if output["path"] == "codex_review_packet.json"
    )
    assert packet_output["status"] == "updated_from_review"
    assert packet_output["required_columns"] == ["review_id", "review_notes"]
    assert any(
        output["path"] == "applied_decisions.json"
        for output in final_artifacts["outputs"]
    )
    run_intake = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )
    review_apply_steps = [
        step
        for step in run_intake["execution_trace"]
        if step["kind"] == "deterministic_review_apply"
    ]
    assert len(review_apply_steps) == 1
    assert {
        "applied_decisions.json",
        "codex_review_packet.json",
        "final_artifacts.json",
        "ui_decisions.json",
    } <= set(review_apply_steps[0]["outputs"])
    contract_report = validate_contract(
        output_dir,
        strict_data_posture=True,
        strict_execution_trace=True,
        strict_output_paths=True,
        strict_output_content=True,
    )
    assert contract_report.ok, contract_report.as_dict()


def test_review_server_failed_checks_block_final_ready(tmp_path: Path) -> None:
    review_server = load_review_server()
    output_dir = tmp_path / "out"
    _write_review_server_fixture(output_dir)
    review_payload_path = output_dir / "review_payload.json"
    review_payload = json.loads(review_payload_path.read_text(encoding="utf-8"))
    review_payload["summary"] = {
        "review_item_count": 2,
        "failed_check_count": 1,
        "checks_pass": False,
    }
    review_payload_path.write_text(
        json.dumps(review_payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    result = review_server.apply_decisions(
        output_dir,
        {
            "decisions": [
                {"item_id": "review:closed", "action": "accept"},
                {"item_id": "review:missing", "action": "accept"},
            ]
        },
    )

    assert result["application_status"] == "blocked"
    assert result["blocker_count"] == 2
    assert {
        blocker["kind"]
        for blocker in result["applied_decisions"]["completion_blockers"]
    } == {"failed_deterministic_checks", "assurance_replay_required"}
    assert result["final_artifacts"]["status"] == "blocked"


def test_review_server_skipped_required_review_blocks_final_ready(
    tmp_path: Path,
) -> None:
    review_server = load_review_server()
    output_dir = tmp_path / "out"
    _write_review_server_fixture(output_dir)

    result = review_server.apply_decisions(
        output_dir,
        {
            "decisions": [
                {"item_id": "review:closed", "action": "skip"},
                {"item_id": "review:missing", "action": "accept"},
            ]
        },
    )

    assert result["application_status"] == "blocked"
    assert result["applied_decisions"]["completion_blockers"][0]["kind"] == (
        "skipped_required_review"
    )


def test_check_dependencies_builds_run_intake_contract() -> None:
    check_dependencies = load_check_dependencies()

    result = check_dependencies.build_dependency_check(
        explicit_files=["requirements.txt"]
    )

    assert result["status"] in {"ok", "missing_dependencies"}
    assert result["command"] == (
        "python scripts/check_dependencies.py --requirements requirements.txt"
    )
    assert result["requirement_files"] == ["requirements.txt"]
    assert "checked_at" in result
    assert result["checked_count"] == len(result["checked"])
    assert result["missing_count"] == len(result["missing"])


def test_dependency_checker_runs_before_optional_third_party_imports() -> None:
    completed = subprocess.run(
        [
            sys.executable,
            "-S",
            CHECK_DEPENDENCIES_PATH.as_posix(),
            "--requirements",
            "requirements.txt",
        ],
        cwd=PLUGIN_ROOT,
        check=False,
        capture_output=True,
        text=True,
        env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"},
    )

    assert completed.returncode == 1
    assert completed.stdout.startswith("MISSING_DEPENDENCIES\n")
    assert "ModuleNotFoundError" not in completed.stderr


def test_raw_input_runner_rejects_git_workspace_output_dir(tmp_path: Path) -> None:
    raw_input_runner = load_raw_input_runner()
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    blocked_output_dir = ROOT / "out" / "audit-reconciliation-test"

    with pytest.raises(ValueError, match="outside the Git workspace"):
        raw_input_runner.extract_normalized_records(
            input_dir,
            output_dir=blocked_output_dir,
        )


def test_audit_reconciliation_mcp_server_validates_and_renders_review_payload() -> None:
    review_payload = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": "audit-reconciliation-test-run",
        "review_type": "audit_reconciliation_review",
        "items": [
            {
                "id": "review:closed",
                "item_type": "closure_evidence_review",
                "title": "INV-1 | 120.50 | closed",
                "source_path": "open_items.xlsx; row 2",
                "output_path": "codex_review_packet.json",
                "allowed_actions": ["accept", "edit", "mark_unclear", "skip"],
                "recommended_action": "accept",
                "evidence": [
                    {
                        "kind": "deterministic_classification",
                        "status": "closed",
                        "rule": "external_bank_match",
                    }
                ],
                "data": {"review_status": "PENDING"},
                "status": "needs_review",
            },
            {
                "id": "check-1",
                "item_type": "check_exception",
                "title": "codex_review_complete",
                "output_path": "run_manifest.json",
                "allowed_actions": ["accept", "reject", "mark_unclear", "skip"],
                "recommended_action": "reject",
                "evidence": [{"kind": "deterministic_check", "status": "FAIL"}],
                "data": {"status": "FAIL"},
                "status": "needs_review",
            },
        ],
        "item_count": 2,
        "columns": [],
        "source_artifacts": {},
        "allowed_actions": [
            "accept",
            "reject",
            "edit",
            "mark_unclear",
            "request_more_documents",
            "skip",
        ],
        "status": "ready_for_review",
        "summary": {
            "reconciliation_row_count": 3,
            "review_row_count": 1,
            "failed_check_count": 1,
            "reconciliation_status_counts": {"closed": 1, "unresolved": 2},
        },
    }
    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "initialize",
            "params": {"protocolVersion": "2024-11-05"},
        },
        {"jsonrpc": "2.0", "id": 2, "method": "tools/list"},
        {
            "jsonrpc": "2.0",
            "id": 3,
            "method": "tools/call",
            "params": {
                "name": "validate_audit_reconciliation_review",
                "arguments": {"review_payload": review_payload},
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 4,
            "method": "tools/call",
            "params": {
                "name": "render_audit_reconciliation_review",
                "arguments": {"review_payload": review_payload},
            },
        },
        {"jsonrpc": "2.0", "id": 5, "method": "resources/list"},
        {
            "jsonrpc": "2.0",
            "id": 6,
            "method": "resources/read",
            "params": {"uri": "ui://widget/audit-reconciliation-review.html"},
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}

    instructions = responses[1]["result"]["instructions"]
    assert "primary visible handoff is the local browser review server" in instructions
    assert "scripts/review_server.py" in instructions
    assert "artifact_card.md" in instructions
    assert "optional integrated Codex review surface" in instructions
    assert "review_payload_path" in instructions
    assert "openai/outputTemplate" in instructions
    assert "static fallback" in instructions
    tool_names = {tool["name"] for tool in responses[2]["result"]["tools"]}
    assert {
        "validate_audit_reconciliation_review",
        "render_audit_reconciliation_review",
    } <= tool_names
    validate_result = responses[3]["result"]["structuredContent"]
    assert validate_result["ok"] is True
    assert validate_result["item_count"] == 2
    assert "scripts/review_server.py" in validate_result["message"]
    assert "artifact_card.md" in validate_result["message"]
    render_result = responses[4]["result"]
    assert render_result["structuredContent"]["widget_type"] == (
        "audit_reconciliation_review"
    )
    assert (
        render_result["_meta"]["openai/outputTemplate"]
        == "ui://widget/audit-reconciliation-review.html"
    )
    assert render_result["_meta"]["openai/widgetAccessible"] is True
    assert render_result["_meta"]["ui"] == {
        "resourceUri": "ui://widget/audit-reconciliation-review.html",
        "visibility": ["model"],
    }
    resource_uris = {
        resource["uri"] for resource in responses[5]["result"]["resources"]
    }
    assert "ui://widget/audit-reconciliation-review.html" in resource_uris
    widget_html = responses[6]["result"]["contents"][0]["text"]
    assert "Audit Reconciliation Review" in widget_html


@pytest.mark.parametrize(
    "surface",
    [
        "initialize",
        "tools/list",
        "resources/list",
        "resources/read",
        "resources/templates/list",
        "prompts/list",
        "tool:validate_audit_reconciliation_review",
        "tool:render_audit_reconciliation_review",
        "tool:save_audit_reconciliation_decisions",
        "tool:apply_audit_reconciliation_decisions",
    ],
)
def test_audit_mcp_rejects_expanded_tree_before_every_public_surface(
    tmp_path: Path,
    surface: str,
) -> None:
    plugin_copy = _copy_audit_mcp_runtime(tmp_path)
    rogue_cache = plugin_copy / "mcp" / "__pycache__"
    rogue_cache.mkdir()
    (rogue_cache / "rogue.pyc").write_bytes(b"rogue")

    with pytest.raises(subprocess.CalledProcessError) as raised:
        _call_mcp_server(
            [_audit_mcp_surface_message(surface)],
            server_path=plugin_copy / "mcp" / "server.cjs",
        )

    assert raised.value.stdout == ""
    assert "exact 25-file contract" in raised.value.stderr


@pytest.mark.parametrize(
    "surface",
    [
        "initialize",
        "tools/list",
        "resources/list",
        "resources/read",
        "resources/templates/list",
        "prompts/list",
        "tool:validate_audit_reconciliation_review",
        "tool:render_audit_reconciliation_review",
        "tool:save_audit_reconciliation_decisions",
        "tool:apply_audit_reconciliation_decisions",
    ],
)
def test_audit_mcp_rejects_post_start_expansion_before_next_public_surface(
    tmp_path: Path,
    surface: str,
) -> None:
    node = shutil.which("node")
    if node is None:
        pytest.skip(
            "Node.js is required to exercise the Audit Reconciliation MCP server."
        )
    plugin_copy = _copy_audit_mcp_runtime(tmp_path)
    process = subprocess.Popen(
        [node, str(plugin_copy / "mcp" / "server.cjs"), "--stdio"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        bufsize=1,
    )
    assert process.stdin is not None
    assert process.stdout is not None
    assert process.stderr is not None
    try:
        initialize = _audit_mcp_surface_message("initialize")
        process.stdin.write(json.dumps(initialize) + "\n")
        process.stdin.flush()
        initialized = json.loads(process.stdout.readline())
        assert initialized["result"]["serverInfo"]["name"] == (
            "audit-reconciliation-widgets"
        )

        rogue_cache = plugin_copy / "mcp" / "__pycache__"
        rogue_cache.mkdir()
        (rogue_cache / "rogue.pyc").write_bytes(b"rogue")
        process.stdin.write(json.dumps(_audit_mcp_surface_message(surface)) + "\n")
        process.stdin.flush()
        process.stdin.close()
        returncode = process.wait(timeout=10)
        remaining_stdout = process.stdout.read()
        stderr = process.stderr.read()
    finally:
        if process.poll() is None:
            process.kill()
            process.wait(timeout=10)

    assert returncode != 0
    assert remaining_stdout == ""
    assert "exact 25-file contract" in stderr


def test_audit_reconciliation_mcp_server_localizes_spanish_runtime_feedback(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    run_id = _customer_run_id(output_dir)
    review_payload = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": run_id,
        "language": "spa",
        "review_type": "audit_reconciliation_review",
        "items": [
            {
                "id": "review:closed",
                "item_type": "closure_evidence_review",
                "title": "FAC-1 | 120,50 | cerrada",
                "allowed_actions": ["accept", "mark_unclear"],
                "recommended_action": "accept",
                "status": "needs_review",
            }
        ],
        "item_count": 1,
        "status": "ready_for_review",
    }
    decisions = [{"item_id": "review:closed", "action": "accept"}]
    run_intake = {
        "run_id": run_id,
        "output_dir": str(output_dir),
        "working_language": "es_ES",
    }
    final_artifacts = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": run_id,
        "outputs": [],
        "next_actions": [],
    }
    output_dir.mkdir(exist_ok=True)
    for name, payload in [
        ("run_intake.json", run_intake),
        ("review_payload.json", review_payload),
        ("final_artifacts.json", final_artifacts),
    ]:
        (output_dir / name).write_text(
            json.dumps(payload, indent=2) + "\n",
            encoding="utf-8",
        )
    invalid_payload = {**review_payload, "item_count": 2}
    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "initialize",
            "params": {
                "protocolVersion": "2024-11-05",
                "meta": {"locale": "es-ES"},
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/call",
            "params": {
                "name": "validate_audit_reconciliation_review",
                "arguments": {"review_payload": review_payload},
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 3,
            "method": "tools/call",
            "params": {
                "name": "save_audit_reconciliation_decisions",
                "arguments": {
                    "review_payload": review_payload,
                    "decisions": decisions,
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 4,
            "method": "tools/call",
            "params": {
                "name": "apply_audit_reconciliation_decisions",
                "arguments": {
                    "review_payload": review_payload,
                    "decisions": decisions,
                    "final_artifacts": final_artifacts,
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 5,
            "method": "tools/call",
            "params": {
                "name": "apply_audit_reconciliation_decisions",
                "arguments": {
                    "run_intake": run_intake,
                    "review_payload": review_payload,
                    "decisions": decisions,
                    "final_artifacts": final_artifacts,
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 6,
            "method": "tools/call",
            "params": {
                "name": "validate_audit_reconciliation_review",
                "arguments": {"review_payload": invalid_payload},
            },
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}

    assert "entrega visible principal" in responses[1]["result"]["instructions"]
    assert responses[2]["result"]["structuredContent"]["message"].startswith(
        "El payload de revisión"
    )
    assert (
        "No se proporcionó run_intake.output_dir"
        in responses[3]["result"]["structuredContent"]["message"]
    )
    no_output_apply = responses[4]["result"]["structuredContent"]
    assert "No se proporcionó run_intake.output_dir" in no_output_apply["message"]
    assert no_output_apply["application_status"] == "blocked"
    assert no_output_apply["final_artifacts"]["next_actions"][-1].startswith(
        "Resuelva las decisiones de revisión bloqueadas"
    )
    persisted_apply = responses[5]["result"]["structuredContent"]
    assert persisted_apply["message"].startswith("Se han aplicado 1 decisiones")
    assert persisted_apply["application_status"] == "blocked"
    assert persisted_apply["final_artifacts"]["next_actions"][-1].startswith(
        "Resuelva las decisiones de revisión bloqueadas"
    )
    handoff = (output_dir / "review_handoff.md").read_text(encoding="utf-8")
    assert "<!-- Review Handoff -->" in handoff
    assert "Entrega para revisión" in handoff
    assert "## Revisión en Codex" in handoff
    error_result = responses[6]["result"]
    assert error_result["isError"] is True
    assert error_result["structuredContent"]["error"].startswith(
        "No se pudo validar la solicitud:"
    )


def test_audit_reconciliation_mcp_server_accepts_local_review_paths(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    _write_review_server_fixture(output_dir)
    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "initialize",
            "params": {"protocolVersion": "2024-11-05"},
        },
        {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/call",
            "params": {
                "name": "validate_audit_reconciliation_review",
                "arguments": {
                    "review_payload_path": str(output_dir / "review_payload.json")
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 3,
            "method": "tools/call",
            "params": {
                "name": "render_audit_reconciliation_review",
                "arguments": {
                    "run_intake_path": str(output_dir / "run_intake.json"),
                    "review_payload_path": str(output_dir / "review_payload.json"),
                    "ui_decisions_path": str(output_dir / "ui_decisions.json"),
                    "final_artifacts_path": str(output_dir / "final_artifacts.json"),
                },
            },
        },
        {
            "jsonrpc": "2.0",
            "id": 4,
            "method": "tools/call",
            "params": {
                "name": "save_audit_reconciliation_decisions",
                "arguments": {
                    "run_intake_path": str(output_dir / "run_intake.json"),
                    "review_payload_path": str(output_dir / "review_payload.json"),
                    "decisions": [
                        {"item_id": "review:closed", "action": "accept"},
                    ],
                },
            },
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}

    validate_result = responses[2]["result"]["structuredContent"]
    assert validate_result["ok"] is True
    assert validate_result["item_count"] == 2
    assert "review_payload_path" in validate_result["message"]
    render_result = responses[3]["result"]["structuredContent"]
    assert render_result["review_payload"]["run_id"] == _customer_run_id(output_dir)
    assert render_result["decision_policy"]["can_persist"] is True
    save_result = responses[4]["result"]["structuredContent"]
    assert save_result["ok"] is True
    assert save_result["persisted"] is True
    saved = json.loads((output_dir / "ui_decisions.json").read_text(encoding="utf-8"))
    assert saved["decision_count"] == 1
    assert saved["decisions"][0]["item_id"] == "review:closed"


def test_audit_reconciliation_mcp_failed_check_blocks_final_ready() -> None:
    review_payload = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": "audit-reconciliation-failed-check",
        "review_type": "audit_reconciliation_review",
        "items": [
            {
                "id": "review:closed",
                "item_type": "closure_evidence_review",
                "title": "INV-1",
                "allowed_actions": ["accept", "skip"],
                "recommended_action": "accept",
                "status": "needs_review",
            }
        ],
        "item_count": 1,
        "status": "ready_for_review",
        "summary": {"failed_check_count": 1, "checks_pass": False},
    }
    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "initialize",
            "params": {"protocolVersion": "2024-11-05"},
        },
        {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/call",
            "params": {
                "name": "apply_audit_reconciliation_decisions",
                "arguments": {
                    "review_payload": review_payload,
                    "decisions": [{"item_id": "review:closed", "action": "accept"}],
                    "final_artifacts": {
                        "schema_version": "1.0",
                        "plugin": "audit-reconciliation",
                        "workflow": "audit-reconciliation",
                        "run_id": "audit-reconciliation-failed-check",
                        "outputs": [],
                    },
                },
            },
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}
    result = responses[2]["result"]["structuredContent"]

    assert result["application_status"] == "blocked"
    assert result["blocker_count"] == 2
    assert {
        blocker["kind"]
        for blocker in result["applied_decisions"]["completion_blockers"]
    } == {"failed_deterministic_checks", "assurance_replay_required"}
    assert result["final_artifacts"]["status"] == "blocked"


def test_audit_reconciliation_mcp_pending_required_review_blocks_final_ready() -> None:
    review_payload = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": "audit-reconciliation-pending-review",
        "review_type": "audit_reconciliation_review",
        "items": [
            {
                "id": "review:closed",
                "item_type": "closure_evidence_review",
                "title": "INV-1",
                "allowed_actions": ["accept", "skip"],
                "recommended_action": "accept",
                "status": "needs_review",
            },
            {
                "id": "review:manual",
                "item_type": "manual_review",
                "title": "INV-2",
                "allowed_actions": ["accept", "skip"],
                "recommended_action": "accept",
                "status": "needs_review",
            },
        ],
        "item_count": 2,
        "status": "ready_for_review",
        "summary": {"failed_check_count": 0, "checks_pass": True},
    }
    messages: list[dict[str, object]] = [
        {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "initialize",
            "params": {"protocolVersion": "2024-11-05"},
        },
        {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/call",
            "params": {
                "name": "apply_audit_reconciliation_decisions",
                "arguments": {
                    "review_payload": review_payload,
                    "decisions": [
                        {"item_id": "review:closed", "action": "accept"},
                    ],
                    "final_artifacts": {
                        "schema_version": "1.0",
                        "plugin": "audit-reconciliation",
                        "workflow": "audit-reconciliation",
                        "run_id": "audit-reconciliation-pending-review",
                        "outputs": [],
                    },
                },
            },
        },
    ]

    responses = {response["id"]: response for response in _call_mcp_server(messages)}
    result = responses[2]["result"]["structuredContent"]

    assert result["application_status"] == "blocked"
    assert result["blocker_count"] == 2
    assert {
        blocker["kind"]
        for blocker in result["applied_decisions"]["completion_blockers"]
    } == {"pending_required_review", "assurance_replay_required"}
    assert result["final_artifacts"]["status"] == "blocked"


def test_skill_mentions_browser_review_and_mcp_tools() -> None:
    skill_text = (
        PLUGIN_ROOT / "skills" / "audit-reconciliation" / "SKILL.md"
    ).read_text(encoding="utf-8")
    manifest = json.loads(
        (PLUGIN_ROOT / ".codex-plugin" / "plugin.json").read_text(encoding="utf-8")
    )

    assert manifest["mcpServers"] == "./.mcp.json"
    assert manifest["apps"] == "./.app.json"
    assert "scripts/review_server.py" in skill_text
    assert "artifact_card.md" in skill_text
    assert "local browser review server" in skill_text
    assert "review surface" in skill_text
    assert "validate_audit_reconciliation_review" in skill_text
    assert "render_audit_reconciliation_review" in skill_text
    assert "review_payload_path" in skill_text
    assert "MCP render is no longer the primary" in skill_text
    assert "Do not treat `review_ui.html`, Markdown summaries" in skill_text
    assert "ui_decisions.json" in skill_text


def _audit_transaction_case(
    output_dir: Path,
) -> dict[str, dict[str, Any] | list[dict[str, str]]]:
    output_dir.mkdir(mode=0o750, exist_ok=True)
    output_dir.chmod(0o750)
    nested = output_dir / "nested"
    nested.mkdir(mode=0o711)
    nested.chmod(0o711)
    sentinel = nested / "sentinel.bin"
    sentinel.write_bytes(b"\x00audit-original\xff")
    sentinel.chmod(0o640)
    run_id = _customer_run_id(output_dir)
    run_intake = {
        "run_id": run_id,
        "output_dir": str(output_dir),
        "working_language": "en",
        "execution_trace": [],
    }
    review_payload = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": run_id,
        "review_type": "audit_reconciliation_review",
        "items": [
            {
                "id": "review:closed",
                "item_type": "closure_evidence_review",
                "title": "Closed item",
                "allowed_actions": ["accept"],
                "recommended_action": "accept",
                "status": "needs_review",
            }
        ],
        "item_count": 1,
        "status": "ready_for_review",
        "summary": {"failed_check_count": 0, "checks_pass": True},
    }
    final_artifacts = {
        "schema_version": "1.0",
        "plugin": "audit-reconciliation",
        "workflow": "audit-reconciliation",
        "run_id": run_id,
        "outputs": [],
        "next_actions": [],
    }
    for name, payload in [
        ("run_intake.json", run_intake),
        ("review_payload.json", review_payload),
        ("final_artifacts.json", final_artifacts),
    ]:
        (output_dir / name).write_text(
            json.dumps(payload, indent=2) + "\n",
            encoding="utf-8",
        )
    return {
        "run_intake": run_intake,
        "review_payload": review_payload,
        "final_artifacts": final_artifacts,
        "decisions": [{"item_id": "review:closed", "action": "accept"}],
    }


def _audit_tree_image(root: Path) -> dict[str, tuple[Any, ...]]:
    image: dict[str, tuple[Any, ...]] = {}
    for entry in [root, *sorted(root.rglob("*"))]:
        observed = entry.lstat()
        relative = "." if entry == root else entry.relative_to(root).as_posix()
        mode = observed.st_mode & 0o7777
        if stat.S_ISREG(observed.st_mode):
            image[relative] = (
                "file",
                mode,
                observed.st_nlink,
                entry.read_bytes(),
            )
        elif stat.S_ISDIR(observed.st_mode):
            image[relative] = ("directory", mode)
        elif stat.S_ISLNK(observed.st_mode):
            image[relative] = ("symlink", mode, os.readlink(entry))
        else:
            image[relative] = ("special", stat.S_IFMT(observed.st_mode), mode)
    return image


@pytest.mark.parametrize(
    "attack",
    [
        "prepared_tamper",
        "root_empty_directory",
        "successor_predecessor_resealed",
    ],
)
def test_audit_mcp_terminal_replay_rejects_invalid_assured_tree(
    tmp_path: Path,
    attack: str,
) -> None:
    assurance = load_assurance()
    workflow = load_reconciliation_workflow()
    review_server = load_review_server()
    output_dir = _running_audit_output(tmp_path)
    open_items = [
        {
            "record_id": "open-1",
            "document_key": "INV-1|2026",
            "document_no": "INV-1",
            "document_date": "2026-01-01",
            "amount": "100.00",
            "currency": "EUR",
        }
    ]
    workflow_args = {
        "output_dir": output_dir,
        "run_id": _customer_run_id(output_dir),
        "open_items": open_items,
        "evidence_rows": [],
        "assumptions": {
            "scope_year": "2026",
            "amount_tolerance": "0",
            "assurance_run_date": "2026-07-25",
        },
        "review_rows": [
            {
                "record_id": "open-1",
                "review_status": "PENDING",
                "reviewer_ref": "",
                "reviewed_on": "",
            }
        ],
        "require_completed_review": False,
        "fail_on_check_errors": False,
        "language": "en",
    }
    workflow.build_reconciliation_artifacts(**workflow_args)
    expected_predecessor_checkpoint = json.loads(
        (output_dir / "assurance_receipts.json").read_text(encoding="utf-8")
    )["content_sha256"]
    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    decisions = [
        {"item_id": item["id"], "action": "accept"} for item in review_payload["items"]
    ]
    review_server.apply_decisions(
        output_dir,
        {
            "decisions": decisions,
            "expected_predecessor_checkpoint": expected_predecessor_checkpoint,
        },
    )
    authority = json.loads(
        (output_dir / "professional_review.json").read_text(encoding="utf-8")
    )
    workflow_args["review_rows"] = authority["records"]
    workflow_args["expected_predecessor_checkpoint"] = expected_predecessor_checkpoint
    workflow.build_reconciliation_artifacts(**workflow_args)
    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    decisions = [
        {"item_id": item["id"], "action": "accept"} for item in review_payload["items"]
    ]

    if attack == "prepared_tamper":
        prepared_path = output_dir / "prepared_records.json"
        prepared = json.loads(prepared_path.read_text(encoding="utf-8"))
        prepared["open_items"][0]["amount"] = "999.00"
        prepared_path.write_text(
            json.dumps(prepared, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        seal_path = output_dir / "assurance_receipts.json"
        seal = json.loads(seal_path.read_text(encoding="utf-8"))
        prepared_bytes = prepared_path.read_bytes()
        seal["prepared_receipt"]["byte_count"] = len(prepared_bytes)
        seal["prepared_receipt"]["sha256"] = hashlib.sha256(prepared_bytes).hexdigest()
        seal_content = {
            key: value for key, value in seal.items() if key != "content_sha256"
        }
        seal["content_sha256"] = assurance.canonical_json_sha256(seal_content)
        seal_path.write_text(
            json.dumps(seal, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
    elif attack == "root_empty_directory":
        (output_dir / "rogue-empty").mkdir()
    else:
        professional_path = output_dir / "professional_review.json"
        professional = json.loads(professional_path.read_text(encoding="utf-8"))
        professional["predecessor_assurance_sha256"] = "f" * 64
        professional_content = {
            key: value for key, value in professional.items() if key != "content_sha256"
        }
        professional["content_sha256"] = assurance.canonical_json_sha256(
            professional_content
        )
        professional_path.write_text(
            json.dumps(
                professional,
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
            )
            + "\n",
            encoding="utf-8",
        )
        seal_path = output_dir / "assurance_receipts.json"
        seal = json.loads(seal_path.read_text(encoding="utf-8"))
        seal["professional_review_authority"] = professional
        professional_bytes = professional_path.read_bytes()
        seal["professional_review_receipt"]["byte_count"] = len(professional_bytes)
        seal["professional_review_receipt"]["sha256"] = hashlib.sha256(
            professional_bytes
        ).hexdigest()
        seal_content = {
            key: value for key, value in seal.items() if key != "content_sha256"
        }
        seal["content_sha256"] = assurance.canonical_json_sha256(seal_content)
        seal_path.write_text(
            json.dumps(
                seal,
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
            )
            + "\n",
            encoding="utf-8",
        )
    before = _audit_tree_image(output_dir)
    run_intake = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )

    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {
                    "name": "apply_audit_reconciliation_decisions",
                    "arguments": {
                        "run_intake": run_intake,
                        "review_payload": review_payload,
                        "decisions": decisions,
                        "final_artifacts": final_artifacts,
                        "expected_predecessor_checkpoint": (
                            expected_predecessor_checkpoint
                        ),
                    },
                },
            }
        ]
    )[0]

    assert response["result"]["isError"] is True
    assert (
        "assurance replay failed"
        in response["result"]["structuredContent"]["error"].lower()
    )
    assert _audit_tree_image(output_dir) == before


@pytest.mark.parametrize(
    "module_name",
    ["audit_assurance", "reconciliation_helpers"],
)
def test_audit_mcp_rejects_timestamp_valid_local_bytecode_before_python_bridge_import(
    tmp_path: Path,
    module_name: str,
) -> None:
    workflow = load_reconciliation_workflow()
    output_dir = _running_audit_output(tmp_path)
    workflow.build_reconciliation_artifacts(
        output_dir=output_dir,
        open_items=[
            {
                "record_id": "open-1",
                "document_key": "INV-1|2026",
                "document_no": "INV-1",
                "document_date": "2026-01-01",
                "amount": "100.00",
                "currency": "EUR",
            }
        ],
        evidence_rows=[],
        assumptions={
            "scope_year": "2026",
            "amount_tolerance": "0",
            "assurance_run_date": "2026-07-25",
        },
        require_completed_review=False,
        fail_on_check_errors=False,
        language="en",
    )
    plugin_copy = tmp_path / "audit-reconciliation"
    shutil.copytree(
        PLUGIN_ROOT,
        plugin_copy,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    shutil.copytree(
        ROOT / "plugins" / "_shared" / "vendor" / "modules" / "vera_assurance",
        tmp_path / "_shared" / "vendor" / "modules" / "vera_assurance",
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", "*.pyo"),
    )
    target_source = plugin_copy / "scripts" / f"{module_name}.py"
    metadata_source = (
        target_source
        if target_source.exists()
        else plugin_copy / "scripts" / "retained_sources" / f"{module_name}.source"
    )
    cache_dir = target_source.parent / "__pycache__"
    cache_dir.mkdir()
    marker = tmp_path / f"{module_name}-mcp-bytecode-executed"
    source_stat = metadata_source.stat()
    malicious_code = compile(
        (
            "from pathlib import Path\n"
            f"Path({marker.as_posix()!r}).write_text('executed', encoding='utf-8')\n"
        ),
        target_source.as_posix(),
        "exec",
    )
    cache_path = cache_dir / f"{module_name}.{sys.implementation.cache_tag}.pyc"
    cache_path.write_bytes(
        importlib._bootstrap_external._code_to_timestamp_pyc(
            malicious_code,
            int(source_stat.st_mtime),
            source_stat.st_size,
        )
    )
    run_intake = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )
    review_payload = json.loads(
        (output_dir / "review_payload.json").read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
    )
    decisions = [
        {"item_id": item["id"], "action": "accept"} for item in review_payload["items"]
    ]
    before = _audit_tree_image(output_dir)

    with pytest.raises(subprocess.CalledProcessError) as raised:
        _call_mcp_server(
            [
                {
                    "jsonrpc": "2.0",
                    "id": 1,
                    "method": "tools/call",
                    "params": {
                        "name": "apply_audit_reconciliation_decisions",
                        "arguments": {
                            "run_intake": run_intake,
                            "review_payload": review_payload,
                            "decisions": decisions,
                            "final_artifacts": final_artifacts,
                        },
                    },
                }
            ],
            server_path=plugin_copy / "mcp" / "server.cjs",
            env={
                "PYTHONDONTWRITEBYTECODE": "",
                "PYTHONPYCACHEPREFIX": "",
            },
        )

    assert raised.value.stdout == ""
    assert "exact 25-file contract" in raised.value.stderr
    assert not marker.exists()
    assert _audit_tree_image(output_dir) == before


def test_audit_mcp_honest_successor_lifecycle_replays_retained_transition(
    tmp_path: Path,
) -> None:
    assurance = load_assurance()
    workflow = load_reconciliation_workflow()
    output_dir = _running_audit_output(tmp_path)
    workflow_args: dict[str, Any] = {
        "output_dir": output_dir,
        "run_id": _customer_run_id(output_dir),
        "open_items": [
            {
                "record_id": "open-1",
                "document_key": "INV-1|2026",
                "document_no": "INV-1",
                "document_date": "2026-01-01",
                "amount": "100.00",
                "currency": "EUR",
            }
        ],
        "evidence_rows": [],
        "assumptions": {
            "scope_year": "2026",
            "amount_tolerance": "0",
            "assurance_run_date": "2026-07-25",
        },
        "review_rows": [
            {
                "record_id": "open-1",
                "review_status": "PENDING",
                "reviewer_ref": "",
                "reviewed_on": "",
            }
        ],
        "require_completed_review": False,
        "fail_on_check_errors": False,
        "language": "en",
    }
    workflow.build_reconciliation_artifacts(**workflow_args)
    predecessor = json.loads(
        (output_dir / "assurance_receipts.json").read_text(encoding="utf-8")
    )

    def apply_current(
        checkpoint: str | None,
    ) -> dict[str, Any]:
        run_intake = json.loads(
            (output_dir / "run_intake.json").read_text(encoding="utf-8")
        )
        review_payload = json.loads(
            (output_dir / "review_payload.json").read_text(encoding="utf-8")
        )
        final_artifacts = json.loads(
            (output_dir / "final_artifacts.json").read_text(encoding="utf-8")
        )
        arguments: dict[str, Any] = {
            "run_intake": run_intake,
            "review_payload": review_payload,
            "decisions": [
                {
                    "item_id": item["id"],
                    "action": "accept",
                }
                for item in review_payload["items"]
            ],
            "final_artifacts": final_artifacts,
        }
        if checkpoint is not None:
            arguments["expected_predecessor_checkpoint"] = checkpoint
        response = _call_mcp_server(
            [
                {
                    "jsonrpc": "2.0",
                    "id": 1,
                    "method": "tools/call",
                    "params": {
                        "name": "apply_audit_reconciliation_decisions",
                        "arguments": arguments,
                    },
                }
            ]
        )[0]
        return response["result"]

    before = _audit_tree_image(output_dir)
    missing_checkpoint = apply_current(None)
    assert missing_checkpoint["isError"] is True
    assert "external expected predecessor checkpoint is required" in (
        missing_checkpoint["structuredContent"]["error"].lower()
    )
    assert _audit_tree_image(output_dir) == before

    wrong_checkpoint_value = (
        "f" * 64 if predecessor["content_sha256"] != "f" * 64 else "e" * 64
    )
    wrong_checkpoint = apply_current(wrong_checkpoint_value)
    assert wrong_checkpoint["isError"] is True
    assert "external expected predecessor checkpoint does not match" in (
        wrong_checkpoint["structuredContent"]["error"].lower()
    )
    assert _audit_tree_image(output_dir) == before

    first_result = apply_current(predecessor["content_sha256"])
    assert first_result.get("isError") is not True
    first = first_result["structuredContent"]
    assert first["application_status"] == "blocked"
    history_dir = (
        output_dir / "assurance_transition_history" / predecessor["content_sha256"]
    )
    assert (history_dir / "transition_receipt.json").is_file()
    authority = json.loads(
        (output_dir / "professional_review.json").read_text(encoding="utf-8")
    )
    workflow_args["review_rows"] = authority["records"]
    workflow_args["expected_predecessor_checkpoint"] = predecessor["content_sha256"]
    workflow.build_reconciliation_artifacts(**workflow_args)
    successor = assurance.validate_assurance_run(
        output_dir,
        expected_predecessor_checkpoint=predecessor["content_sha256"],
    )
    assert (
        successor["professional_review_authority"]["predecessor_assurance_sha256"]
        == predecessor["content_sha256"]
    )
    assert len(successor["review_transition_receipts"]) == 1

    second_result = apply_current(predecessor["content_sha256"])
    assert second_result.get("isError") is not True
    second = second_result["structuredContent"]

    assert second["application_status"] == "final_ready"
    assert (
        second["applied_decisions"]["professional_review"][
            "successor_assurance_replayed"
        ]
        is True
    )
    assert (
        assurance.validate_assurance_run(
            output_dir,
            expected_predecessor_checkpoint=predecessor["content_sha256"],
        )["gate_register"]["gates"]["publication"]["status"]
        == "withheld"
    )


def _audit_faulted_server(
    tmp_path: Path,
    *,
    needle: str,
    replacement: str,
) -> Path:
    source = MCP_SERVER_PATH.read_text(encoding="utf-8")
    plugin_root_line = 'const PLUGIN_ROOT = path.resolve(__dirname, "..");'
    assert source.count(plugin_root_line) == 1
    assert source.count(needle) == 1
    source = source.replace(
        plugin_root_line,
        f"const PLUGIN_ROOT = {json.dumps(str(MCP_SERVER_PATH.parents[1]))};",
        1,
    )
    source = source.replace(needle, replacement, 1)
    server_path = tmp_path / "audit-faulted-server.cjs"
    server_path.write_text(source, encoding="utf-8")
    return server_path


def _audit_transaction_call(
    tool_name: str,
    arguments: dict[str, Any],
    *,
    server_path: Path = MCP_SERVER_PATH,
    env: dict[str, str] | None = None,
) -> dict[str, Any]:
    response = _call_mcp_server(
        [
            {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "tools/call",
                "params": {"name": tool_name, "arguments": arguments},
            }
        ],
        server_path=server_path,
        env=env,
    )[0]
    return response["result"]["structuredContent"]


def _portable_audit_transaction_case(
    tmp_path: Path,
) -> tuple[Path, dict[str, Any]]:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    assurance = load_assurance()
    context = assurance.load_client_engagement_context_file(
        _customer_context_path(output_dir),
        expected_workflow_id="audit-reconciliation",
    )
    source_path = Path(context["input_bindings"][0]["path"])
    load_review_session().write_run_intake(
        output_dir,
        assumptions={"currency": "EUR"},
        source_paths=[source_path],
        language="en",
        client_engagement=context,
        run_id=str(context["run_id"]),
    )
    run_intake = json.loads(
        (output_dir / "run_intake.json").read_text(encoding="utf-8")
    )
    arguments["run_intake"] = run_intake
    return output_dir, arguments


def test_audit_review_save_and_apply_survive_customer_folder_rename(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _portable_audit_transaction_case(tmp_path)
    old_customer_root = _customer_context_path(output_dir).parents[5].as_posix()
    old_output = output_dir.as_posix()
    assert arguments["run_intake"]["path_reference"] == "run_root_relative"
    assert arguments["run_intake"]["output_dir"] == "outputs"
    assert all(
        not Path(value).is_absolute()
        for value in arguments["run_intake"]["input_paths"]
    )

    renamed_output, current_context, stale_output = _rename_customer_output(output_dir)
    arguments["client_engagement"] = current_context.as_posix()

    saved = _audit_transaction_call(
        "save_audit_reconciliation_decisions",
        arguments,
    )
    applied = _audit_transaction_call(
        "apply_audit_reconciliation_decisions",
        arguments,
    )

    assert saved["ok"] is True
    assert saved["persisted"] is True
    assert applied["ok"] is True
    assert applied["persisted"] is True
    assert (renamed_output / "ui_decisions.json").is_file()
    assert (renamed_output / "applied_decisions.json").is_file()
    assert not stale_output.exists()
    assert old_output not in arguments["run_intake"]["output_dir"]
    assert all(
        old_customer_root not in artifact.read_text(encoding="utf-8")
        for artifact in renamed_output.rglob("*")
        if artifact.is_file() and artifact.suffix.lower() in {".json", ".md"}
    )


def test_audit_review_rejects_run_root_escape_without_writing(
    tmp_path: Path,
) -> None:
    output_dir, arguments = _portable_audit_transaction_case(tmp_path)
    arguments["client_engagement"] = _customer_context_path(output_dir).as_posix()
    forged = json.loads(json.dumps(arguments))
    forged["run_intake"]["output_dir"] = "../outside"
    before = _audit_tree_image(output_dir)

    result = _audit_transaction_call(
        "save_audit_reconciliation_decisions",
        forged,
    )

    assert result["ok"] is False
    assert "leaves the customer run" in result["error"]
    assert _audit_tree_image(output_dir) == before
    assert not (output_dir.parent.parent / "outside").exists()


def test_audit_review_transaction_honest_apply_commits_without_residue(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)

    result = _audit_transaction_call(
        "apply_audit_reconciliation_decisions",
        arguments,
    )

    assert result["ok"] is True
    assert result["persisted"] is True
    assert result["application_status"] == "blocked"
    assert result["blocker_count"] == 1
    assert result["applied_decisions"]["completion_blockers"] == [
        {
            "kind": "assurance_replay_required",
            "detail": (
                "Reviewed decisions require native regeneration and a fresh "
                "successor assurance receipt replay before final readiness."
            ),
        }
    ]
    assert (output_dir / "applied_decisions.json").is_file()
    assert (output_dir / "nested" / "sentinel.bin").read_bytes() == (
        b"\x00audit-original\xff"
    )
    assert output_dir.stat().st_mode & 0o7777 == 0o750
    assert (output_dir / "nested").stat().st_mode & 0o7777 == 0o711
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


def test_audit_review_transaction_rejects_forged_caller_review_target(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    target = output_dir / "nested" / "controlled.txt"
    target.write_bytes(b"ORIGINAL-CANONICAL-BYTES")
    target.chmod(0o640)
    before = _audit_tree_image(output_dir)
    forged_arguments = json.loads(json.dumps(arguments))
    forged_item = forged_arguments["review_payload"]["items"][0]
    forged_item.update(
        {
            "title": "Forged writable item",
            "allowed_actions": ["edit"],
            "recommended_action": "edit",
            "data": {"target_artifact": "nested/controlled.txt"},
        }
    )
    forged_arguments["decisions"] = [
        {
            "item_id": forged_item["id"],
            "action": "edit",
            "edit_value": "FORGED-CALLER-WRITE",
        }
    ]

    result = _audit_transaction_call(
        "apply_audit_reconciliation_decisions",
        forged_arguments,
    )

    assert result == {
        "ok": False,
        "error": (
            "Caller review payload does not match the persisted "
            "Audit Reconciliation review payload."
        ),
    }
    assert target.read_bytes() == b"ORIGINAL-CANONICAL-BYTES"
    assert target.stat().st_mode & 0o7777 == 0o640
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize(
    ("tool_name", "context_name", "expected_error"),
    [
        (
            "save_audit_reconciliation_decisions",
            "run_intake",
            (
                "Caller run intake does not match the persisted "
                "Audit Reconciliation run intake."
            ),
        ),
        (
            "apply_audit_reconciliation_decisions",
            "run_intake",
            (
                "Caller run intake does not match the persisted "
                "Audit Reconciliation run intake."
            ),
        ),
        (
            "save_audit_reconciliation_decisions",
            "final_artifacts",
            (
                "Caller final artifacts do not match the persisted "
                "Audit Reconciliation final artifacts."
            ),
        ),
        (
            "apply_audit_reconciliation_decisions",
            "final_artifacts",
            (
                "Caller final artifacts do not match the persisted "
                "Audit Reconciliation final artifacts."
            ),
        ),
        (
            "save_audit_reconciliation_decisions",
            "ui_decisions",
            (
                "Caller UI decisions do not match the persisted "
                "Audit Reconciliation UI decisions."
            ),
        ),
        (
            "apply_audit_reconciliation_decisions",
            "ui_decisions",
            (
                "Caller UI decisions do not match the persisted "
                "Audit Reconciliation UI decisions."
            ),
        ),
    ],
)
def test_audit_review_transaction_rejects_forged_caller_context(
    tmp_path: Path,
    tool_name: str,
    context_name: str,
    expected_error: str,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    if context_name == "ui_decisions":
        persisted_ui_decisions = {
            "schema_version": "1.0",
            "plugin": "audit-reconciliation",
            "workflow": "audit-reconciliation",
            "run_id": "audit-transaction-run",
            "decisions": [],
            "status": "pending_review",
        }
        (output_dir / "ui_decisions.json").write_text(
            json.dumps(persisted_ui_decisions, indent=2) + "\n",
            encoding="utf-8",
        )
        arguments["ui_decisions"] = {
            **persisted_ui_decisions,
            "status": "forged",
        }
    else:
        arguments[context_name] = {
            **arguments[context_name],
            "forged_caller_field": True,
        }
    before = _audit_tree_image(output_dir)

    result = _audit_transaction_call(tool_name, arguments)

    assert result == {"ok": False, "error": expected_error}
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize(
    ("tool_name", "needle"),
    [
        (
            "save_audit_reconciliation_decisions",
            "      const workingResult = saveDecisionPayloadWrites(workingArgs);\n",
        ),
        (
            "apply_audit_reconciliation_decisions",
            "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n",
        ),
    ],
)
def test_audit_review_transaction_late_failure_restores_bytes_and_modes(
    tmp_path: Path,
    tool_name: str,
    needle: str,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    before = _audit_tree_image(output_dir)
    faulted = _audit_faulted_server(
        tmp_path,
        needle=needle,
        replacement=(
            needle + '      throw new Error("/private/client/audit-late-failure");\n'
        ),
    )

    result = _audit_transaction_call(
        tool_name,
        arguments,
        server_path=faulted,
    )

    assert result["ok"] is False
    assert result["error"] == (
        "Audit Reconciliation review "
        + ("save" if tool_name.startswith("save") else "apply")
        + " transaction failed safely."
    )
    assert "/private/client" not in result["error"]
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


def test_audit_review_transaction_rejects_forged_save_response_contract(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    before = _audit_tree_image(output_dir)
    needle = "      const workingResult = saveDecisionPayloadWrites(workingArgs);\n"
    faulted = _audit_faulted_server(
        tmp_path,
        needle=needle,
        replacement=(needle + """
      Object.assign(workingResult, {
        validation_type: "forged_save",
        run_id: "forged-run",
        decision_count: 777,
        item_count: 888,
        status: "forged_status",
        ui_decisions_path: "/private/client/forged-ui.json",
        message: "forged message",
      });
"""),
    )

    result = _audit_transaction_call(
        "save_audit_reconciliation_decisions",
        arguments,
        server_path=faulted,
    )

    assert result == {
        "ok": False,
        "error": "Audit Reconciliation saved decisions did not close.",
    }
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


def test_audit_review_transaction_rejects_forged_apply_response_contract(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    before = _audit_tree_image(output_dir)
    needle = "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n"
    faulted = _audit_faulted_server(
        tmp_path,
        needle=needle,
        replacement=(needle + """
      Object.assign(workingResult, {
        validation_type: "forged_apply",
        run_id: "forged-run",
        decision_count: 101,
        item_count: 102,
        blocker_count: 103,
        revision_count: 104,
        target_update_count: 105,
        structured_update_count: 106,
        native_regeneration_count: 107,
        native_regenerated_count: 108,
        application_status: "final_ready",
        ui_decisions_path: "/private/client/forged-ui.json",
        applied_decisions_path: "/private/client/forged-applied.json",
        final_artifacts_path: "/private/client/forged-final.json",
        run_intake_path: "/private/client/forged-intake.json",
        message: "forged message",
      });
"""),
    )

    result = _audit_transaction_call(
        "apply_audit_reconciliation_decisions",
        arguments,
        server_path=faulted,
    )

    assert result == {
        "ok": False,
        "error": "Audit Reconciliation response did not close.",
    }
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize("self_authorized_path", ["rogue.json", "ui_decisions.json"])
def test_audit_review_transaction_rejects_persisted_result_self_authorization(
    tmp_path: Path,
    self_authorized_path: str,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    before = _audit_tree_image(output_dir)
    needle = "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n"
    injected = """
      const selfAuthorizedPath = process.env.REVIEW_TX_SELF_AUTHORIZED_PATH;
      if (selfAuthorizedPath === "rogue.json") {
        const roguePath = path.join(workingOutputDir, selfAuthorizedPath);
        fs.writeFileSync(roguePath, "{\\"forged\\":true}\\n", "utf8");
      }
      const appliedPath = path.join(workingOutputDir, "applied_decisions.json");
      const finalPath = path.join(workingOutputDir, "final_artifacts.json");
      const applied = JSON.parse(fs.readFileSync(appliedPath, "utf8"));
      const finalArtifacts = JSON.parse(fs.readFileSync(finalPath, "utf8"));
      applied.native_regenerated_count = 1;
      applied.native_regenerated_paths = [selfAuthorizedPath];
      finalArtifacts.review_application.native_regenerated_count = 1;
      finalArtifacts.review_application.native_regenerated_paths = [selfAuthorizedPath];
      fs.writeFileSync(appliedPath, `${JSON.stringify(applied, null, 2)}\\n`, "utf8");
      fs.writeFileSync(finalPath, `${JSON.stringify(finalArtifacts, null, 2)}\\n`, "utf8");
      workingResult.native_regenerated_count = 1;
      workingResult.applied_decisions = applied;
      workingResult.final_artifacts = finalArtifacts;
"""
    faulted = _audit_faulted_server(
        tmp_path,
        needle=needle,
        replacement=needle + injected,
    )

    result = _audit_transaction_call(
        "apply_audit_reconciliation_decisions",
        arguments,
        server_path=faulted,
        env={"REVIEW_TX_SELF_AUTHORIZED_PATH": self_authorized_path},
    )

    assert result["ok"] is False
    assert "/" not in result["error"]
    assert "\\" not in result["error"]
    assert self_authorized_path not in result["error"]
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


@pytest.mark.parametrize("attack_kind", ["symlink", "hardlink", "fifo"])
def test_audit_review_transaction_rejects_working_tree_poison(
    tmp_path: Path,
    attack_kind: str,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    before = _audit_tree_image(output_dir)
    external = tmp_path / "external-target.bin"
    external.write_bytes(b"EXTERNAL-UNCHANGED")
    external.chmod(0o600)
    needle = "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n"
    injected = """
      const poisonPath = path.join(workingOutputDir, "ui_decisions.json");
      fs.unlinkSync(poisonPath);
      if (process.env.REVIEW_TX_ATTACK_KIND === "symlink") {
        fs.symlinkSync(process.env.REVIEW_TX_EXTERNAL, poisonPath);
      } else if (process.env.REVIEW_TX_ATTACK_KIND === "hardlink") {
        fs.linkSync(process.env.REVIEW_TX_EXTERNAL, poisonPath);
      } else {
        require("node:child_process").spawnSync("/usr/bin/mkfifo", [poisonPath]);
      }
"""
    faulted = _audit_faulted_server(
        tmp_path,
        needle=needle,
        replacement=needle + injected,
    )

    result = _audit_transaction_call(
        "apply_audit_reconciliation_decisions",
        arguments,
        server_path=faulted,
        env={
            "REVIEW_TX_ATTACK_KIND": attack_kind,
            "REVIEW_TX_EXTERNAL": str(external),
        },
    )

    assert result["ok"] is False
    assert "/" not in result["error"]
    assert external.read_bytes() == b"EXTERNAL-UNCHANGED"
    assert external.stat().st_mode & 0o7777 == 0o600
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


def test_audit_review_transaction_rejects_transaction_root_relocation_without_moving_canonical(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    before = _audit_tree_image(output_dir)
    canonical_inode = output_dir.stat().st_ino
    needle = "      const workingResult = applyDecisionPayloadWrites(workingArgs);\n"
    injected = """
      const transactionRoot = path.dirname(workingOutputDir);
      fs.renameSync(transactionRoot, `${transactionRoot}-moved`);
      throw new Error("/private/client/transaction-root-relocation");
"""
    faulted = _audit_faulted_server(
        tmp_path,
        needle=needle,
        replacement=needle + injected,
    )

    result = _audit_transaction_call(
        "apply_audit_reconciliation_decisions",
        arguments,
        server_path=faulted,
    )

    assert result["ok"] is False
    assert result["error"] == (
        "Audit Reconciliation review apply transaction failed safely."
    )
    assert output_dir.stat().st_ino == canonical_inode
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))
    assert not list(output_dir.parent.glob(".generated-review-commit-*"))
    assert not list(output_dir.parent.glob(".generated-review-recovery-*"))


def test_audit_review_transaction_restores_after_commit_deletion(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    before = _audit_tree_image(output_dir)
    needle = "    committed = true;\n    const committedImage ="
    faulted = _audit_faulted_server(
        tmp_path,
        needle=needle,
        replacement=(
            "    committed = true;\n"
            "    generatedReviewRemoveExactPath(resolvedOutputDir);\n"
            "    const committedImage ="
        ),
    )

    result = _audit_transaction_call(
        "apply_audit_reconciliation_decisions",
        arguments,
        server_path=faulted,
    )

    assert result["ok"] is False
    assert "/" not in result["error"]
    assert _audit_tree_image(output_dir) == before
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))


def test_audit_review_transaction_enforces_size_bound_before_mutation(
    tmp_path: Path,
) -> None:
    output_dir = _running_audit_output(tmp_path)
    arguments = _audit_transaction_case(output_dir)
    sentinel = output_dir / "nested" / "sentinel.bin"
    oversized = output_dir / "oversized.bin"
    with oversized.open("wb") as stream:
        stream.truncate(128 * 1024 * 1024 + 1)
    root_inode = output_dir.stat().st_ino
    sentinel_inode = sentinel.stat().st_ino
    sentinel_bytes = sentinel.read_bytes()

    result = _audit_transaction_call(
        "save_audit_reconciliation_decisions",
        arguments,
    )

    assert result["ok"] is False
    assert output_dir.stat().st_ino == root_inode
    assert sentinel.stat().st_ino == sentinel_inode
    assert sentinel.read_bytes() == sentinel_bytes
    assert oversized.stat().st_size == 128 * 1024 * 1024 + 1
    assert not list(output_dir.parent.glob(".generated-review-transaction-*"))
