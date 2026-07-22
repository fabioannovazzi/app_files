from __future__ import annotations

import importlib.util
import json
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from scripts.validate_plugin_review_contract import validate_contract

OUTPUTS = (
    Path(__file__).resolve().parents[2]
    / "plugins"
    / "audit-reconciliation"
    / "scripts"
    / "workpaper_outputs.py"
)
REVIEW_SESSION = OUTPUTS.with_name("review_session.py")


def load_outputs():
    spec = importlib.util.spec_from_file_location("audit_workpaper_outputs", OUTPUTS)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def load_review_session():
    script_dir = str(REVIEW_SESSION.parent)
    if script_dir not in sys.path:
        sys.path.insert(0, script_dir)
    spec = importlib.util.spec_from_file_location(
        "audit_workpaper_review_session", REVIEW_SESSION
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def document_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_build_audit_workbook_sheets_has_standard_tabs():
    outputs = load_outputs()
    sheets = outputs.build_audit_workbook_sheets(
        assumptions={"cutoff_date": "2023-12-31"},
        source_inventory=[{"source_file": "ledger.pdf"}],
        normalized_records=[{"record_id": "n1"}],
        reconciliation_rows=[
            {
                "record_id": "open-1",
                "amount": "100.00",
                "reconciliation_status": "closed",
                "rule_applied": "direct_external_or_documented",
            }
        ],
        checks=[{"check": "row_count", "status": "PASS"}],
        bank_allocation_candidates=[
            {
                "candidate_id": "candidate-1",
                "candidate_type": "unallocated_counterparty_bank_pool",
            }
        ],
        external_evidence_summary=[
            {
                "external_category": "TOTAL",
                "settlement_effect_signed_net_debit_minus_credit": "-180.00",
            }
        ],
        external_evidence_detail=[
            {
                "record_id": "bank-direct",
                "external_category": "direct_counterparty_bank_receipt",
            }
        ],
        ledger_balance_rows=[
            {"account": "TOTAL", "closing_balance_signed_debit_minus_credit": "550.00"}
        ],
        account_rollforward_check=[{"account": "TOTAL", "status": "PASS"}],
        aggregate_rollforward_summary=[
            {"account": "TOTAL", "closing_net_debit_minus_credit": "550.00"}
        ],
        aggregate_rollforward_rows=[{"record_id": "journal_rollforward:1"}],
        post_cutoff_candidates=[{"candidate_id": "post-cutoff-1"}],
        aging_summary=[{"aging_bucket": "181-365", "rows": 1}],
        evidence_concentration=[{"support_bucket": "bank", "rows": 1}],
        review_signals=[{"record_id": "open-1", "review_signal_rank": 1}],
        document_source_map=[{"document_key": "INV1|2023", "open_item_rows": 1}],
        reversal_candidates=[{"candidate_id": "reversal-1"}],
        cutoff_window_movements=[{"record_id": "cutoff-1"}],
        review_rows=[{"record_id": "open-1", "review_result": "PASS"}],
    )

    assert list(sheets) == outputs.STANDARD_SHEET_ORDER
    assert sheets["Summary"][0]["rows"] == 1
    assert sheets["Bank allocation candidates"][0]["candidate_id"] == "candidate-1"
    assert (
        sheets["External evidence aggregate"][0][
            "settlement_effect_signed_net_debit_minus_credit"
        ]
        == "-180.00"
    )
    assert sheets["External evidence detail"][0]["record_id"] == "bank-direct"
    assert (
        sheets["Ledger balance check"][0]["closing_balance_signed_debit_minus_credit"]
        == "550.00"
    )
    assert sheets["Account rollforward check"][0]["status"] == "PASS"
    assert (
        sheets["Journal rollforward"][0]["closing_net_debit_minus_credit"] == "550.00"
    )
    assert sheets["Journal detail"][0]["record_id"] == "journal_rollforward:1"
    assert sheets["Post-cutoff candidates"][0]["candidate_id"] == "post-cutoff-1"
    assert sheets["Open item aging"][0]["aging_bucket"] == "181-365"
    assert sheets["Evidence concentration"][0]["support_bucket"] == "bank"
    assert sheets["Review signals"][0]["review_signal_rank"] == 1
    assert sheets["Document source map"][0]["document_key"] == "INV1|2023"
    assert sheets["Reversal candidates"][0]["candidate_id"] == "reversal-1"
    assert sheets["Cutoff window movements"][0]["record_id"] == "cutoff-1"
    assert sheets["Index"][0]["Sheet"] == "Assumptions"


def test_write_excel_workpaper_creates_auditable_tabs(tmp_path):
    outputs = load_outputs()
    path = tmp_path / "audit.xlsx"
    sheets = outputs.build_audit_workbook_sheets(
        assumptions={"cutoff_date": "2023-12-31"},
        source_inventory=[],
        normalized_records=[],
        reconciliation_rows=[],
        checks=[],
        review_rows=[],
    )

    outputs.write_excel_workpaper(path, sheets)

    workbook = load_workbook(path)
    assert workbook.sheetnames == outputs.STANDARD_SHEET_ORDER
    assert workbook["Index"].freeze_panes == "A2"


def test_write_excel_workpaper_adds_italian_presentation_tabs_without_removing_raw(
    tmp_path,
):
    outputs = load_outputs()
    path = tmp_path / "audit.xlsx"
    sheets = outputs.build_audit_workbook_sheets(
        assumptions={"factoring_pro_soluto_closes_item": True},
        source_inventory=[{"source_role": "bank_statement"}],
        normalized_records=[],
        reconciliation_rows=[
            {
                "record_id": "open-1",
                "amount": "100.00",
                "reconciliation_status": "needs_evidence",
                "rule_applied": "internal_closure_without_external",
                "matched_evidence_type": "internal_closure",
            }
        ],
        checks=[],
        review_rows=[],
    )

    outputs.write_excel_workpaper(path, sheets, language="it")

    workbook = load_workbook(path)
    assert "Dettaglio riconciliazione" in workbook.sheetnames
    assert "Reconciliation detail" in workbook.sheetnames
    assert workbook["Reconciliation detail"].sheet_state == "hidden"
    localized_sheet = workbook["Dettaglio riconciliazione"]
    headers = [cell.value for cell in localized_sheet[1]]
    values = [cell.value for cell in localized_sheet[2]]
    assert "Esito riconciliazione" in headers
    assert "Regola applicata" in headers
    assert "Serve evidenza aggiuntiva" in values
    assert "Chiusura interna senza evidenza esterna" in values
    assert "needs_evidence" not in {str(value) for value in values}
    assert "internal_closure_without_external" not in {str(value) for value in values}
    assumption_values = [
        cell.value
        for row in workbook["Assunzioni"].iter_rows(values_only=False)
        for cell in row
    ]
    assert (
        "Factoring pro-soluto chiude con evidenza esterna allocata" in assumption_values
    )
    assert "factoring_pro_soluto_closes_item" not in {
        str(value) for value in assumption_values
    }


def test_spanish_workpapers_and_review_contract_are_fully_localized(tmp_path):
    outputs = load_outputs()
    excel_path = tmp_path / "conciliacion_auditoria.xlsx"
    word_path = tmp_path / "informe_conciliacion.docx"
    assumptions = {
        "scope_year": "2024",
        "cutoff_date": "2024-12-31",
        "currency": "EUR",
        "factoring_pro_soluto_closes_item": True,
        "post_cutoff_events_excluded": True,
    }
    reconciliation_rows = [
        {
            "record_id": "open-1",
            "document_no": "FAC-001",
            "amount": "100.00",
            "reconciliation_status": "needs_evidence",
            "rule_applied": "internal_closure_without_external",
            "matched_evidence_type": "internal_closure",
        }
    ]
    checks = [{"check": "review_packet", "status": "WARN"}]
    review_rows = [
        {
            "record_id": "open-1",
            "review_status": "PENDING",
            "source_file": "diario.xlsx",
            "source_page": 2,
            "source_row": 14,
        }
    ]
    sheets = outputs.build_audit_workbook_sheets(
        assumptions=assumptions,
        source_inventory=[{"source_role": "bank_statement"}],
        normalized_records=[],
        reconciliation_rows=reconciliation_rows,
        checks=checks,
        review_rows=review_rows,
        language="es",
    )

    outputs.write_excel_workpaper(excel_path, sheets, language="es-ES")
    outputs.write_word_report(
        word_path,
        title="Informe de conciliación de auditoría",
        metadata={"scope_year": "2024", "cutoff_date": "2024-12-31"},
        summary_rows=sheets["Summary"],
        assumptions=assumptions,
        next_steps=["Obtener la evidencia pendiente."],
        source_inventory=[{"source_role": "bank_statement"}],
        checks=checks,
        review_rows=review_rows,
        language="spa",
    )

    workbook = load_workbook(excel_path)
    assert "Índice" in workbook.sheetnames
    assert "Supuestos" in workbook.sheetnames
    assert "Detalle de conciliación" in workbook.sheetnames
    assert "Revisión Codex" in workbook.sheetnames
    assert workbook["Reconciliation detail"].sheet_state == "hidden"
    detail_values = {
        str(cell.value)
        for row in workbook["Detalle de conciliación"].iter_rows()
        for cell in row
        if cell.value is not None
    }
    assert "Estado de conciliación" in detail_values
    assert "Regla aplicada" in detail_values
    assert "Requiere evidencia adicional" in detail_values
    assert "Cierre interno sin evidencia externa" in detail_values
    assert "Asiento de cierre" in detail_values
    assert "needs_evidence" not in detail_values
    assert "internal_closure_without_external" not in detail_values
    assumption_values = {
        str(cell.value)
        for row in workbook["Supuestos"].iter_rows()
        for cell in row
        if cell.value is not None
    }
    assert "Año del alcance" in assumption_values
    assert (
        "El factoring sin recurso cierra con evidencia externa asignada"
        in assumption_values
    )
    assert "Sí" in assumption_values

    report = document_text(word_path)
    for expected in (
        "Resumen ejecutivo",
        "Papel de trabajo de conciliación determinista",
        "Alcance y método",
        "Cómo interpretar los resultados",
        "Requiere evidencia adicional",
        "Cierre interno sin evidencia externa",
        "Extracto bancario",
        "Controles automáticos",
        "Aviso",
        "Revisión manual de Codex",
        "Pendiente",
        "Limitaciones del procedimiento",
        "Referencia al archivo Excel",
    ):
        assert expected in report
    for internal_or_english in (
        "needs_evidence",
        "internal_closure_without_external",
        "bank_statement",
        "Deterministic reconciliation workpaper",
        "Scope and Method",
        "Procedure Limits",
    ):
        assert internal_or_english not in report

    review_session = load_review_session()
    run_intake = review_session.write_run_intake(
        tmp_path,
        assumptions=assumptions,
        language="spa",
        source_hint="auditoria-es",
    )
    result = {
        "excel_path": excel_path,
        "word_path": word_path,
        "assumptions": assumptions,
        "reconciliation_rows": reconciliation_rows,
        "checks": checks,
        "review_rows": review_rows,
        "checks_pass": False,
    }
    review_result = review_session.write_review_session_artifacts(
        tmp_path,
        run_id=run_intake.run_id,
        run_intake_path=run_intake.path,
        result=result,
        language="spa",
    )
    run_intake_payload = json.loads(run_intake.path.read_text(encoding="utf-8"))
    review_payload = json.loads(
        review_result.review_payload_path.read_text(encoding="utf-8")
    )
    final_artifacts = json.loads(
        review_result.final_artifacts_path.read_text(encoding="utf-8")
    )
    assert review_payload["columns"] == [
        {"field": "item_type", "label": "Tipo"},
        {"field": "title", "label": "Línea o artefacto"},
        {"field": "recommended_action", "label": "Acción sugerida"},
        {"field": "source_path", "label": "Fuente"},
        {"field": "output_path", "label": "Salida"},
        {"field": "status", "label": "Estado"},
    ]
    review_row_item = next(
        item for item in review_payload["items"] if item["id"] == "open-1"
    )
    assert review_row_item["source_path"] == "diario.xlsx; página 2; fila 14"
    assert review_row_item["data"]["edit_hint"].startswith(
        "Editar esta línea de revisión"
    )
    artifact_titles = {
        item["title"]
        for item in review_payload["items"]
        if item["item_type"] in {"workpaper_artifact", "report_artifact"}
    }
    assert artifact_titles == {
        "Libro de conciliación de auditoría",
        "Informe narrativo de conciliación",
    }
    assert run_intake_payload["data_posture"]["notes"] == [
        "Los scripts de conciliación leen las rutas locales de evidencias contables registradas en input_paths.",
        "Los datos de revisión muestran líneas de conciliación y referencias de evidencias acotadas para su revisión en la interfaz.",
        "De forma predeterminada no se utiliza ningún conector externo, ruta de carga, SQL remoto ni cuaderno alojado.",
    ]
    spanish_review_contract = json.dumps(
        {"run_intake": run_intake_payload, "review_payload": review_payload},
        ensure_ascii=False,
    )
    for english_copy in (
        "Row or artifact",
        "Suggested action",
        "Editing this review row",
        "Audit reconciliation workbook",
        "Narrative reconciliation report",
        "Review payloads expose bounded",
    ):
        assert english_copy not in spanish_review_contract
    artifact_card = review_result.artifact_card_path.read_text(encoding="utf-8")
    assert "# Ficha de artefactos de la conciliación de auditoría" in artifact_card
    assert "## Siguiente acción" in artifact_card
    assert "Audit Reconciliation Artifact Card" not in artifact_card
    assert "Los datos mostrados en el navegador están acotados" in " ".join(
        final_artifacts["caveats"]
    )
    assert "Abra el servidor de revisión" in final_artifacts["next_actions"][0]
    assert "Open the browser review server" not in json.dumps(
        final_artifacts, ensure_ascii=False
    )
    audit_output = next(
        output
        for output in final_artifacts["outputs"]
        if output.get("artifact_role") == "audit_workpaper"
    )
    word_output = next(
        output
        for output in final_artifacts["outputs"]
        if output.get("artifact_role") == "word_report"
    )
    assert audit_output["required_sheets"] == [
        "Índice",
        "Supuestos",
        "Detalle de conciliación",
        "Resumen",
        "Controles",
        "Revisión Codex",
    ]
    assert audit_output["required_sheet_headers"] == {
        "Índice": ["Hoja", "Líneas"],
        "Supuestos": ["Campo", "Valor"],
    }
    assert audit_output["required_cells"]["Supuestos"]["A2"] == "Año del alcance"
    assert word_output["required_text"][0] == "Resumen ejecutivo"
    assert word_output["required_text"][-1] == "Referencia al archivo Excel"
    contract_report = validate_contract(tmp_path, strict_output_content=True)
    artifact_errors = [
        error
        for error in contract_report.errors
        if excel_path.name in error or word_path.name in error
    ]
    assert artifact_errors == []


def test_write_word_report_defaults_to_italian_labels(tmp_path):
    outputs = load_outputs()
    path = tmp_path / "relazione.docx"

    outputs.write_word_report(
        path,
        title="Relazione di riconciliazione",
        metadata={"Periodo": "2023"},
        summary_rows=[{"Esito": "Chiuso", "Righe": 1}],
        assumptions={"cutoff_date": "2023-12-31"},
        next_steps=["Acquisire documentazione mancante."],
        narrative="Sintesi tecnica.",
    )

    text = document_text(path)
    assert "Conclusioni" in text
    assert "Assunzioni" in text
    assert "Prossimi passi" in text


def test_write_word_report_localizes_italian_internal_codes(tmp_path):
    outputs = load_outputs()
    path = tmp_path / "relazione.docx"

    outputs.write_word_report(
        path,
        title="Relazione di riconciliazione",
        metadata={},
        summary_rows=[
            {
                "reconciliation_status": "needs_evidence",
                "rule_applied": "internal_closure_without_external",
                "rows": 1,
                "amount": "100.00",
            }
        ],
        assumptions={
            "factoring_pro_soluto_closes_item": True,
            "post_cutoff_events_excluded": True,
        },
        next_steps=["Acquisire evidenze aggiuntive."],
        source_inventory=[{"source_role": "bank_statement"}],
        checks=[{"check": "review_packet", "status": "WARN"}],
        review_rows=[{"review_status": "PENDING"}],
    )

    text = document_text(path)
    assert "Serve evidenza aggiuntiva" in text
    assert "Chiusura interna senza evidenza esterna" in text
    assert "Factoring pro-soluto chiude con evidenza esterna allocata: Sì" in text
    assert "Eventi post cut-off esclusi dalla chiusura: Sì" in text
    assert "Estratto conto bancario" in text
    assert "needs_evidence" not in text
    assert "internal_closure_without_external" not in text
    assert "factoring_pro_soluto_closes_item" not in text
    assert "post_cutoff_events_excluded" not in text


def test_write_word_report_includes_deterministic_check_summaries(tmp_path):
    outputs = load_outputs()
    path = tmp_path / "relazione.docx"

    outputs.write_word_report(
        path,
        title="Relazione di riconciliazione",
        metadata={},
        summary_rows=[],
        assumptions={"cutoff_date": "2023-12-31"},
        next_steps=[],
        account_rollforward_check=[
            {
                "account": "TOTAL",
                "account_name": "Conti confrontati",
                "opening_difference_journal_minus_ledger": "0.00",
                "closing_difference_journal_minus_ledger": "8941820.33",
                "status": "DIFFERENCE",
                "review_note": "Il saldo ricostruito dal giornale non coincide con il saldo finale del mastro.",
            }
        ],
        post_cutoff_candidates=[
            {
                "document_no": "23FE01/000425",
                "document_date": "2023-01-31",
                "open_amount": "109484.65",
                "evidence_date": "2024-01-12",
                "evidence_type": "external_bank",
                "exact_amount_match": "YES",
                "evidence_source_file": "banca.pdf",
                "evidence_amount": "109484.65",
            }
        ],
        aging_summary=[
            {"aging_bucket": "181-365", "rows": 1, "amount_abs_total": "109484.65"}
        ],
        evidence_concentration=[
            {
                "support_bucket": "payment_order",
                "reconciliation_status": "needs_evidence",
                "rows": 1,
                "amount_abs_total": "109484.65",
                "share_of_abs_amount_percent": "100.00",
            }
        ],
        review_signals=[
            {
                "review_signal_rank": 1,
                "document_no": "23FE01/000425",
                "amount": "109484.65",
                "age_days_at_reference": 334,
                "reconciliation_status": "needs_evidence",
                "review_signals": "high_value; needs_evidence",
            }
        ],
        document_source_map=[
            {
                "document_no_examples": "23FE01/000425",
                "open_amount_total": "109484.65",
                "open_item_rows": 1,
                "ledger_rows": 1,
                "journal_rows": 1,
                "bank_rows": 0,
                "payment_order_rows": 1,
                "factoring_rows": 0,
                "review_note": "Documento presente in distinta/compensazione: verificare supporto di chiusura.",
            }
        ],
        reversal_candidates=[
            {
                "document_no": "23FE01/000425",
                "open_amount": "109484.65",
                "evidence_amount": "-109484.65",
                "evidence_type": "compensation",
                "evidence_date": "2023-12-15",
                "candidate_reasons": "opposite_sign_amount",
            }
        ],
        cutoff_window_movements=[
            {
                "cutoff_window_timing": "before_cutoff",
                "evidence_type": "compensation",
                "amount": "-109484.65",
            }
        ],
    )

    text = document_text(path)
    assert "Controllo saldi da mastro e giornale" in text
    assert "Differenza" in text
    assert "8,941,820.33" in text
    assert "Evidenze successive al cut-off" in text
    assert "Principali candidati" in text
    assert "23FE01/000425" in text
    assert "Analisi deterministiche aggiuntive" in text
    assert "Aging partite aperte" in text
    assert "Mappa documento-fonti" in text
    assert "Possibili storni, giroconti o compensazioni" in text
