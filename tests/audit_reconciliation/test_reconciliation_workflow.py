from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

SCRIPTS = (
    Path(__file__).resolve().parents[2] / "plugins" / "audit-reconciliation" / "scripts"
)
WORKFLOW = SCRIPTS / "reconciliation_workflow.py"


def load_workflow():
    sys.path.insert(0, str(SCRIPTS))
    spec = importlib.util.spec_from_file_location(
        "audit_reconciliation_workflow", WORKFLOW
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def document_text(path: str | Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_build_reconciliation_artifacts_writes_excel_and_word(tmp_path):
    workflow = load_workflow()
    result = workflow.build_reconciliation_artifacts(
        output_dir=tmp_path,
        open_items=[
            {
                "record_id": "open-1",
                "document_key": "INV1|2023",
                "document_date": "2023-03-01",
                "amount": "100.00",
            }
        ],
        evidence_rows=[
            {
                "record_id": "bank-1",
                "evidence_type": "external_bank",
                "document_key": "INV1|2023",
                "posting_date": "2023-12-15",
                "amount": "100.00",
                "source_file": "bank.pdf",
                "source_page": "1",
            }
        ],
        assumptions={"scope_year": "2023", "cutoff_date": "2023-12-31"},
        source_inventory=[{"source_file": "bank.pdf", "source_role": "bank_statement"}],
        ledger_balance_rows=[
            {"account": "TOTAL", "closing_balance_signed_debit_minus_credit": "100.00"}
        ],
        account_rollforward_check=[{"account": "TOTAL", "status": "PASS"}],
        aggregate_rollforward_summary=[
            {"account": "TOTAL", "closing_net_debit_minus_credit": "100.00"}
        ],
        aggregate_rollforward_rows=[{"record_id": "journal_rollforward:1"}],
        metadata={"Periodo": "2023"},
        narrative="Riconciliazione completata sui dati normalizzati.",
    )

    assert result["checks_pass"] is True
    assert result["reconciliation_rows"][0]["reconciliation_status"] == "closed"
    assert Path(result["excel_path"]).exists()
    assert Path(result["accountant_report_path"]).exists()
    assert Path(result["word_path"]).exists()
    assert Path(result["review_session"]["review_html_path"]).exists()
    assert Path(result["review_session"]["review_html_path"]).name == "review_ui.html"

    workbook = load_workbook(result["excel_path"])
    assert "Reconciliation detail" in workbook.sheetnames
    assert "Bank allocation candidates" in workbook.sheetnames
    assert "External evidence aggregate" in workbook.sheetnames
    assert "External evidence detail" in workbook.sheetnames
    assert "Ledger balance check" in workbook.sheetnames
    assert "Account rollforward check" in workbook.sheetnames
    assert "Journal rollforward" in workbook.sheetnames
    assert "Journal detail" in workbook.sheetnames
    assert "Post-cutoff candidates" in workbook.sheetnames
    assert "Open item aging" in workbook.sheetnames
    assert "Evidence concentration" in workbook.sheetnames
    assert "Review signals" in workbook.sheetnames
    assert "Document source map" in workbook.sheetnames
    assert "Reversal candidates" in workbook.sheetnames
    assert "Cutoff window movements" in workbook.sheetnames
    assert "Review" in workbook.sheetnames
    assert "bank_allocation_candidates" in result
    accountant_workbook = load_workbook(result["accountant_report_path"])
    assert "Scheda operativa" in accountant_workbook.sheetnames
    assert "Dettaglio riscontri" in accountant_workbook.sheetnames
    assert "external_evidence_summary" in result
    assert "external_evidence_detail" in result
    assert result["ledger_balance_rows"][0]["account"] == "TOTAL"
    assert result["account_rollforward_check"][0]["status"] == "PASS"
    assert result["aggregate_rollforward_summary"][0]["account"] == "TOTAL"
    assert result["aging_summary"][0]["aging_bucket"] == "181-365"
    assert result["evidence_concentration"][0]["support_bucket"] == "bank"
    assert result["document_source_map"][0]["bank_rows"] == 1
    assert result["cutoff_window_movements"][0]["record_id"] == "bank-1"
    assert any(
        check["check"] == "codex_review_packet_present" for check in result["checks"]
    )
    assert result["review_rows"][0]["review_status"] == "PENDING"
    text = document_text(result["word_path"])
    assert "Conclusioni" in text
    assert "Controllo saldi da mastro e giornale" in text
    assert "Analisi deterministiche aggiuntive" in text


def test_workflow_can_be_loaded_without_preconfigured_python_path(tmp_path):
    spec = importlib.util.spec_from_file_location("audit_workflow_direct", WORKFLOW)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)

    result = module.build_reconciliation_artifacts(
        output_dir=tmp_path,
        open_items=[],
        evidence_rows=[],
        assumptions={"scope_year": "2023"},
    )

    assert result["checks_pass"] is True
    assert Path(result["excel_path"]).exists()
    assert Path(result["word_path"]).exists()


def test_default_next_steps_mentions_unresolved_rows():
    workflow = load_workflow()
    steps = workflow.default_next_steps(
        [{"reconciliation_status": "unresolved"}],
        language="it",
    )

    assert any("non risolte" in step for step in steps)


def test_default_next_steps_uses_selected_language():
    workflow = load_workflow()
    steps = workflow.default_next_steps(
        [{"reconciliation_status": "needs_evidence"}],
        language="en_US",
    )

    assert any("Obtain the evidence" in step for step in steps)


def test_workflow_can_require_completed_review(tmp_path):
    workflow = load_workflow()

    try:
        workflow.build_reconciliation_artifacts(
            output_dir=tmp_path,
            open_items=[
                {
                    "record_id": "open-1",
                    "document_key": "INV1|2023",
                    "document_date": "2023-03-01",
                    "amount": "100.00",
                }
            ],
            evidence_rows=[],
            assumptions={"scope_year": "2023"},
            require_completed_review=True,
        )
    except ValueError as exc:
        assert "codex_review_completed" in str(exc)
    else:
        raise AssertionError(
            "expected pending Codex review to fail when completion is required"
        )
