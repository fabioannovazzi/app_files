from __future__ import annotations

import argparse
import copy
import csv
import hashlib
import json
import logging
import re
import shutil
import stat
import sys
import tempfile
import unicodedata
import zipfile
from collections import Counter
from collections.abc import Mapping
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any, Sequence

import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from .report_builder_integrity import (
        load_source_index,
        resolve_source_record_path,
        seal_review_integrity,
        source_identity_key,
        validate_source_index,
        write_source_index,
    )
    from .review_session import (
        refresh_final_artifacts,
        write_review_session_artifacts,
        write_run_intake,
    )
except ImportError:  # pragma: no cover - supports direct script imports
    import importlib.util

    _integrity_path = Path(__file__).resolve().parent / "report_builder_integrity.py"
    _integrity_spec = importlib.util.spec_from_file_location(
        "mparanza_report_builder_integrity",
        _integrity_path,
    )
    if _integrity_spec is None or _integrity_spec.loader is None:
        raise ImportError("Could not load Report Builder integrity module")
    _integrity = importlib.util.module_from_spec(_integrity_spec)
    sys.modules[_integrity_spec.name] = _integrity
    _integrity_spec.loader.exec_module(_integrity)
    load_source_index = _integrity.load_source_index
    resolve_source_record_path = _integrity.resolve_source_record_path
    seal_review_integrity = _integrity.seal_review_integrity
    source_identity_key = _integrity.source_identity_key
    validate_source_index = _integrity.validate_source_index
    write_source_index = _integrity.write_source_index

    _review_session_path = Path(__file__).resolve().parent / "review_session.py"
    _review_session_spec = importlib.util.spec_from_file_location(
        "mparanza_report_builder_review_session",
        _review_session_path,
    )
    if _review_session_spec is None or _review_session_spec.loader is None:
        raise ImportError("Could not load Report Builder review-session module")
    _review_session = importlib.util.module_from_spec(_review_session_spec)
    sys.modules[_review_session_spec.name] = _review_session
    _review_session_spec.loader.exec_module(_review_session)
    refresh_final_artifacts = _review_session.refresh_final_artifacts
    write_review_session_artifacts = _review_session.write_review_session_artifacts
    write_run_intake = _review_session.write_run_intake


def _ensure_vendor_import_path() -> None:
    """Expose component-local or repository-shared Vera vendor modules."""

    component_root = Path(__file__).resolve().parents[1]
    candidates = (
        component_root / "vendor" / "modules",
        component_root.parent / "_shared" / "vendor" / "modules",
    )
    for candidate in candidates:
        if candidate.is_dir() and str(candidate) not in sys.path:
            sys.path.insert(0, str(candidate))


_ensure_vendor_import_path()
from vera_assurance import (  # noqa: E402
    MoneyValidationError,
    artifact_receipt,
    build_numeric_evidence_ledger,
    build_reviewed_decision_receipt,
    decimal_text,
    file_snapshot,
    parse_canonical_decimal,
    parse_localized_decimal,
    validate_artifact_receipt,
    validate_reviewed_decision_receipt,
)

LOGGER = logging.getLogger(__name__)

SUPPORTED_LANGUAGES = ("it", "en", "fr", "de", "es")
SUPPORTED_DOCUMENT_LANGUAGES = ("auto", *SUPPORTED_LANGUAGES)
SUPPORTED_SUFFIXES = {".csv", ".pdf", ".xlsx", ".xlsm", ".zip"}
TEXT_SUFFIXES = {".pdf"}
WORKBOOK_SUFFIXES = {".xlsx", ".xlsm"}
NUMERIC_MEASURE_ADAPTER_ID = "report_table_numeric_profile"
NUMERIC_MEASURE_ADAPTER_VERSION = "v4"
NUMERIC_MEASURE_DECISION_TYPE = "numeric_measure_mapping"
NUMERIC_PARSE_POLICY = "strict_all_nonblank_v1"
NUMERIC_SIGN_POLICIES = {"as_presented_v1", "invert_v1"}
MODEL_CONTEXT_PREVIEW_ROWS = 8
MODEL_CONTEXT_MAX_EXPANSION_ROWS = 100
MODEL_CONTEXT_MAX_EXPANSION_COLUMNS = 16
NUMERIC_LOCALE_SEPARATORS: dict[str, tuple[str, str]] = {
    "en": (".", ","),
    "it": (",", "."),
    "fr": (",", "."),
    "de": (",", "."),
    "es": (",", "."),
}
NUMERIC_UNITS = {"currency", "number", "count", "ratio", "percentage"}
_CURRENCY_CODES = ("EUR", "USD", "GBP", "CHF", "CAD", "AUD", "JPY")
_CURRENCY_SYMBOLS = {"€": "EUR", "£": "GBP"}
_AMBIGUOUS_CURRENCY_SYMBOLS = {"$"}

REPORT_TYPES: dict[str, dict[str, Any]] = {
    "management_report": {
        "label": {
            "en": "Management report",
            "it": "Report gestionale",
            "fr": "Rapport de gestion",
            "de": "Managementbericht",
            "es": "Informe de gestión",
        },
        "sections": [
            "overview",
            "income_statement",
            "balance_sheet",
            "cash_flow",
            "budget",
            "debt",
            "investments",
            "taxes",
            "notes",
        ],
    },
    "local_government_review": {
        "label": {
            "en": "Local government review",
            "it": "Relazione ente locale",
            "fr": "Revue collectivite locale",
            "de": "Kommunaler Pruefbericht",
            "es": "Informe de entidad local",
        },
        "sections": [
            "overview",
            "fpv",
            "fcde",
            "debt",
            "cash",
            "taxes",
            "spending",
            "investments",
            "participations",
            "pnrr",
            "notes",
        ],
    },
    "annual_financial_statement": {
        "label": {
            "en": "Annual financial statement",
            "it": "Bilancio annuale",
            "fr": "Etats financiers annuels",
            "de": "Jahresabschluss",
            "es": "Estados financieros anuales",
        },
        "sections": [
            "overview",
            "balance_sheet",
            "income_statement",
            "cash_flow",
            "equity",
            "ratios",
            "segment",
            "debt",
            "capex",
            "notes",
        ],
    },
}

SECTION_TITLES: dict[str, dict[str, str]] = {
    "overview": {
        "en": "Overview",
        "it": "Sintesi",
        "fr": "Synthese",
        "de": "Ueberblick",
        "es": "Resumen",
    },
    "income_statement": {
        "en": "Income statement",
        "it": "Conto economico",
        "fr": "Compte de resultat",
        "de": "Gewinn- und Verlustrechnung",
        "es": "Cuenta de resultados",
    },
    "balance_sheet": {
        "en": "Balance sheet",
        "it": "Stato patrimoniale",
        "fr": "Bilan",
        "de": "Bilanz",
        "es": "Balance",
    },
    "cash_flow": {
        "en": "Cash flow",
        "it": "Rendiconto finanziario",
        "fr": "Flux de tresorerie",
        "de": "Cashflow",
        "es": "Flujo de caja",
    },
    "budget": {
        "en": "Budget",
        "it": "Budget",
        "fr": "Budget",
        "de": "Budget",
        "es": "Presupuesto",
    },
    "debt": {
        "en": "Debt",
        "it": "Debito",
        "fr": "Dette",
        "de": "Schulden",
        "es": "Deuda",
    },
    "investments": {
        "en": "Investments",
        "it": "Investimenti",
        "fr": "Investissements",
        "de": "Investitionen",
        "es": "Inversiones",
    },
    "taxes": {
        "en": "Taxes",
        "it": "Tributi",
        "fr": "Fiscalite",
        "de": "Steuern",
        "es": "Impuestos",
    },
    "notes": {
        "en": "Notes",
        "it": "Note",
        "fr": "Notes",
        "de": "Anmerkungen",
        "es": "Notas",
    },
    "fpv": {
        "en": "FPV",
        "it": "FPV",
        "fr": "FPV",
        "de": "FPV",
        "es": "FPV",
    },
    "fcde": {
        "en": "FCDE",
        "it": "FCDE",
        "fr": "FCDE",
        "de": "FCDE",
        "es": "FCDE",
    },
    "cash": {
        "en": "Cash",
        "it": "Cassa",
        "fr": "Tresorerie",
        "de": "Liquiditaet",
        "es": "Tesorería",
    },
    "spending": {
        "en": "Spending",
        "it": "Spesa",
        "fr": "Depenses",
        "de": "Ausgaben",
        "es": "Gasto",
    },
    "participations": {
        "en": "Participations",
        "it": "Partecipazioni",
        "fr": "Participations",
        "de": "Beteiligungen",
        "es": "Participaciones",
    },
    "pnrr": {
        "en": "PNRR",
        "it": "PNRR",
        "fr": "PNRR",
        "de": "PNRR",
        "es": "PNRR",
    },
    "equity": {
        "en": "Equity",
        "it": "Patrimonio netto",
        "fr": "Capitaux propres",
        "de": "Eigenkapital",
        "es": "Patrimonio neto",
    },
    "ratios": {
        "en": "Ratios",
        "it": "Indicatori",
        "fr": "Ratios",
        "de": "Kennzahlen",
        "es": "Ratios",
    },
    "segment": {
        "en": "Segment information",
        "it": "Informativa per settore",
        "fr": "Information sectorielle",
        "de": "Segmentinformationen",
        "es": "Información por segmentos",
    },
    "capex": {
        "en": "Capital expenditure",
        "it": "Investimenti tecnici",
        "fr": "Depenses d'investissement",
        "de": "Investitionsausgaben",
        "es": "Inversiones de capital",
    },
}

SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "overview": (
        "overview",
        "sintesi",
        "summary",
        "resume",
        "ueberblick",
        "resumen",
        "síntesis",
        "sintesis",
    ),
    "income_statement": (
        "income",
        "profit",
        "loss",
        "conto economico",
        "ricavi",
        "costi",
        "resultat",
        "guv",
        "cuenta de resultados",
        "pérdidas y ganancias",
        "perdidas y ganancias",
    ),
    "balance_sheet": (
        "balance",
        "stato patrimoniale",
        "attivo",
        "passivo",
        "bilan",
        "bilanz",
        "balance",
        "situación financiera",
        "situacion financiera",
    ),
    "cash_flow": (
        "cash flow",
        "rendiconto",
        "flussi",
        "tresorerie",
        "kapitalfluss",
        "flujo de caja",
        "flujos de efectivo",
    ),
    "budget": (
        "budget",
        "forecast",
        "prevision",
        "preventivo",
        "planung",
        "presupuesto",
        "previsión",
    ),
    "debt": (
        "debt",
        "debito",
        "mutui",
        "loans",
        "dette",
        "schulden",
        "deuda",
        "préstamos",
        "prestamos",
    ),
    "investments": (
        "investment",
        "investimenti",
        "capex",
        "immobilizzazioni",
        "investissements",
        "inversiones",
    ),
    "taxes": (
        "tax",
        "taxes",
        "imposte",
        "tributi",
        "fiscal",
        "steuern",
        "impuestos",
        "tributos",
    ),
    "notes": (
        "note",
        "notes",
        "comment",
        "commenti",
        "annexe",
        "anhang",
        "notas",
        "anexo",
    ),
    "fpv": ("fpv", "fondo pluriennale"),
    "fcde": ("fcde", "crediti dubbia", "doubtful", "creances douteuses"),
    "cash": (
        "cash",
        "cassa",
        "tesoreria",
        "banque",
        "liquiditaet",
        "tesorería",
        "efectivo",
    ),
    "spending": ("spesa", "spending", "depenses", "ausgaben", "gasto"),
    "participations": (
        "partecipazioni",
        "participations",
        "beteiligungen",
        "subsidiaries",
        "participaciones",
    ),
    "pnrr": ("pnrr", "rrf", "recovery"),
    "equity": (
        "equity",
        "patrimonio netto",
        "capitaux propres",
        "eigenkapital",
        "patrimonio neto",
    ),
    "ratios": ("ratio", "indicatori", "indici", "kennzahlen", "ratios"),
    "segment": ("segment", "settore", "sector", "secteur", "segmento"),
    "capex": (
        "capex",
        "capital expenditure",
        "investimenti tecnici",
        "inversiones de capital",
    ),
}

DOCX_COPY: dict[str, dict[str, str]] = {
    "entity": {
        "en": "Entity",
        "it": "Ente",
        "fr": "Entite",
        "de": "Einheit",
        "es": "Entidad",
    },
    "period": {
        "en": "Period",
        "it": "Periodo",
        "fr": "Periode",
        "de": "Zeitraum",
        "es": "Periodo",
    },
    "entity_pending": {
        "en": "Entity pending",
        "it": "Ente da definire",
        "fr": "Entite a definir",
        "de": "Einheit noch offen",
        "es": "Entidad pendiente",
    },
    "period_pending": {
        "en": "Period pending",
        "it": "Periodo da definire",
        "fr": "Periode a definir",
        "de": "Zeitraum noch offen",
        "es": "Periodo pendiente",
    },
    "executive_summary": {
        "en": "Executive summary",
        "it": "Sintesi",
        "fr": "Synthese",
        "de": "Zusammenfassung",
        "es": "Resumen ejecutivo",
    },
    "executive_summary_pending": {
        "en": "Claude executive summary pending.",
        "it": "Sintesi Claude in attesa.",
        "fr": "Synthese Claude en attente.",
        "de": "Claude-Zusammenfassung noch offen.",
        "es": "Resumen ejecutivo de Claude pendiente.",
    },
    "context": {
        "en": "Context",
        "it": "Contesto",
        "fr": "Contexte",
        "de": "Kontext",
        "es": "Contexto",
    },
    "source": {
        "en": "Source",
        "it": "Fonte",
        "fr": "Source",
        "de": "Quelle",
        "es": "Fuente",
    },
    "rows": {
        "en": "Rows",
        "it": "Righe",
        "fr": "Lignes",
        "de": "Zeilen",
        "es": "Filas",
    },
    "columns": {
        "en": "Columns",
        "it": "Colonne",
        "fr": "Colonnes",
        "de": "Spalten",
        "es": "Columnas",
    },
    "numeric_totals": {
        "en": "Deterministic numeric totals",
        "it": "Totali numerici deterministici",
        "fr": "Totaux numeriques deterministes",
        "de": "Deterministische numerische Summen",
        "es": "Totales numéricos deterministas",
    },
    "count": {
        "en": "count",
        "it": "conteggio",
        "fr": "nombre",
        "de": "Anzahl",
        "es": "recuento",
    },
    "sum": {
        "en": "sum",
        "it": "somma",
        "fr": "somme",
        "de": "Summe",
        "es": "suma",
    },
    "currency": {
        "en": "Currency",
        "it": "Valuta",
        "fr": "Devise",
        "de": "Waehrung",
        "es": "Moneda",
    },
    "unit": {
        "en": "Unit",
        "it": "Unita",
        "fr": "Unite",
        "de": "Einheit",
        "es": "Unidad",
    },
    "scale": {
        "en": "Scale",
        "it": "Scala",
        "fr": "Echelle",
        "de": "Skalierung",
        "es": "Escala",
    },
    "table_preview": {
        "en": "Table preview",
        "it": "Anteprima tabella",
        "fr": "Apercu du tableau",
        "de": "Tabellenvorschau",
        "es": "Vista previa de la tabla",
    },
    "unassigned": {
        "en": "No table assigned yet. professional review pending for this section.",
        "it": "Nessuna tabella assegnata. Revisione Claude in attesa per questa sezione.",
        "fr": "Aucun tableau assigne. Revue Claude en attente pour cette section.",
        "de": "Noch keine Tabelle zugeordnet. Claude-Pruefung fuer diesen Abschnitt offen.",
        "es": "Todavía no hay una tabla asignada. La revisión de Claude está pendiente para esta sección.",
    },
    "codex_pending": {
        "en": "professional review pending for this section.",
        "it": "Revisione Claude in attesa per questa sezione.",
        "fr": "Revue Claude en attente pour cette section.",
        "de": "Claude-Pruefung fuer diesen Abschnitt offen.",
        "es": "La revisión de Claude está pendiente para esta sección.",
    },
    "audit_appendix": {
        "en": "Audit appendix",
        "it": "Appendice audit",
        "fr": "Annexe d'audit",
        "de": "Audit-Anhang",
        "es": "Anexo de auditoría",
    },
    "report_status": {
        "en": "Report status",
        "it": "Stato report",
        "fr": "Statut du rapport",
        "de": "Berichtsstatus",
        "es": "Estado del informe",
    },
    "assigned_sections": {
        "en": "Assigned sections",
        "it": "Sezioni assegnate",
        "fr": "Sections assignees",
        "de": "Zugeordnete Abschnitte",
        "es": "Secciones asignadas",
    },
    "missing_sections": {
        "en": "Missing sections",
        "it": "Sezioni mancanti",
        "fr": "Sections manquantes",
        "de": "Fehlende Abschnitte",
        "es": "Secciones pendientes",
    },
    "model_api_calls": {
        "en": "Model API calls from scripts",
        "it": "Chiamate API modello dagli script",
        "fr": "Appels API modele par les scripts",
        "de": "Modell-API-Aufrufe aus Skripten",
        "es": "Llamadas a la API del modelo desde los scripts",
    },
    "draft": {
        "en": "Draft generated by deterministic scripts and Claude-guided narrative.",
        "it": "Bozza generata da script deterministici e narrativa guidata da Claude.",
        "fr": "Brouillon genere par scripts deterministes et narration guidee par Claude.",
        "de": "Entwurf durch deterministische Skripte und Claude-gefuehrte Narrative erstellt.",
        "es": "Borrador generado por scripts deterministas y narrativa guiada por Claude.",
    },
    "notes": {
        "en": "Notes",
        "it": "Note",
        "fr": "Notes",
        "de": "Anmerkungen",
        "es": "Notas",
    },
    "input_path": {
        "en": "Input path",
        "it": "Percorso di input",
        "fr": "Chemin d'entree",
        "de": "Eingabepfad",
        "es": "Ruta de entrada",
    },
    "tables_discovered": {
        "en": "Tables discovered",
        "it": "Tabelle rilevate",
        "fr": "Tableaux detectes",
        "de": "Erkannte Tabellen",
        "es": "Tablas detectadas",
    },
    "column": {
        "en": "Column",
        "it": "Colonna",
        "fr": "Colonne",
        "de": "Spalte",
        "es": "Columna",
    },
    "range": {
        "en": "Range",
        "it": "Intervallo",
        "fr": "Plage",
        "de": "Spannweite",
        "es": "Rango",
    },
    "generated_by": {
        "en": "Generated by the Build Report Claude plugin.",
        "it": "Generato dal plugin Claude Build Report.",
        "fr": "Genere par le plugin Claude Build Report.",
        "de": "Vom Claude-Plugin Build Report erstellt.",
        "es": "Generado por el plugin Claude Build Report.",
    },
}

WORD_RE = re.compile(r"[a-z0-9]+")

__all__ = [
    "BuildResult",
    "InspectionResult",
    "add_common_args",
    "build_report",
    "configure_logging",
    "build_model_context_expansion",
    "inspect_inputs",
    "load_indexed_tables",
    "normalize_language",
    "review_numeric_measure_columns",
    "validate_narrative_numeric_boundary",
    "write_numeric_evidence_ledger",
    "write_json",
]


@dataclass(frozen=True)
class InspectionResult:
    """Inspection outputs for the report-builder workflow."""

    inspection: dict[str, Any]
    suggested_recipe: dict[str, Any]


@dataclass(frozen=True)
class BuildResult:
    """Report build outputs and audit metadata."""

    analysis: dict[str, Any]
    audit: dict[str, Any]
    markdown_path: Path
    docx_path: Path
    review_session: dict[str, Any] | None = None


def configure_logging(verbose: bool = False) -> None:
    """Configure script logging without affecting imported use."""

    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(level=level, format="%(levelname)s: %(message)s")


def normalize_language(
    language: object | None,
    *,
    default: str = "en",
    allow_auto: bool = False,
) -> str:
    """Normalize a language tag to one supported plugin locale."""

    text = str(language or default).strip().lower().replace("_", "-")
    code = text.split("-", 1)[0]
    if allow_auto and code == "auto":
        return "auto"
    return code if code in SUPPORTED_LANGUAGES else default


def normalize_report_type(report_type: object | None) -> str:
    """Return a supported report type key."""

    value = str(report_type or "management_report").strip().lower().replace("-", "_")
    return value if value in REPORT_TYPES else "management_report"


def language_assumptions(
    recipe: dict[str, Any] | None = None,
    *,
    language: object | None = None,
    document_language: object | None = None,
) -> dict[str, str]:
    """Resolve working and source-document language assumptions."""

    recipe = recipe or {}
    working = normalize_language(language or recipe.get("language"), default="en")
    source = normalize_language(
        document_language or recipe.get("document_language") or "auto",
        default=working,
        allow_auto=True,
    )
    return {"language": working, "document_language": source}


def write_json(path: Path, payload: Any) -> None:
    """Write JSON with stable formatting."""

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def _stabilize_office_package(path: Path) -> None:
    """Normalize OOXML metadata, entry order, and ZIP timestamps."""

    source = Path(path)
    with zipfile.ZipFile(source) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    stable_entries: list[tuple[zipfile.ZipInfo, bytes]] = []
    for info, data in entries:
        if info.filename == "docProps/core.xml":
            for tag in (b"created", b"modified"):
                pattern = (
                    rb"(<dcterms:" + tag + rb"\b[^>]*>)[^<]*(</dcterms:" + tag + rb">)"
                )
                data = re.sub(
                    pattern,
                    rb"\g<1>2000-01-01T00:00:00Z\g<2>",
                    data,
                )
        stable_info = zipfile.ZipInfo(
            filename=info.filename,
            date_time=(1980, 1, 1, 0, 0, 0),
        )
        stable_info.compress_type = info.compress_type
        stable_info.comment = info.comment
        stable_info.internal_attr = info.internal_attr
        stable_info.external_attr = info.external_attr
        stable_info.create_system = info.create_system
        stable_entries.append((stable_info, data))
    temporary = source.with_name(f".{source.name}.stable")
    with zipfile.ZipFile(
        temporary,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as archive:
        for info, data in sorted(stable_entries, key=lambda item: item[0].filename):
            archive.writestr(info, data)
    temporary.replace(source)


def read_json(path: Path | None) -> dict[str, Any]:
    """Return a JSON object or an empty mapping when no file is provided."""

    if path is None:
        return {}
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"Recipe must be a JSON object: {path}")
    return payload


def clean_text(value: Any) -> str:
    """Normalize a cell value for display and matching."""

    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = str(value).replace("\u00a0", " ").replace("\u202f", " ")
    return re.sub(r"\s+", " ", text).strip()


def norm_label(value: Any) -> str:
    """Return an ASCII-ish lower-case label for heuristics."""

    text = unicodedata.normalize("NFKD", clean_text(value))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower()
    return re.sub(r"\s+", " ", text)


def safe_sheet_name(name: str, fallback: str) -> str:
    """Return a valid Excel worksheet name."""

    cleaned = re.sub(r"[\[\]:*?/\\]", " ", name).strip() or fallback
    return cleaned[:31]


def parse_amount(value: Any) -> Decimal | None:
    """Parse unambiguous financial numbers into exact Decimal values."""

    if value is None or isinstance(value, bool):
        return None
    try:
        return parse_localized_decimal(
            value,
            label="report numeric cell",
            allow_float=isinstance(value, float),
        )
    except MoneyValidationError:
        return None


def _looks_numeric_candidate(value: Any) -> bool:
    """Surface numeric-looking syntax for review without assigning semantics."""

    if value is None or isinstance(value, bool):
        return False
    if isinstance(value, (int, float, Decimal)):
        return True
    text = clean_text(value).upper()
    if not text or not any(character.isdigit() for character in text):
        return False
    for code in _CURRENCY_CODES:
        text = re.sub(rf"\b{code}\b", "", text)
    for symbol in {*_CURRENCY_SYMBOLS, *_AMBIGUOUS_CURRENCY_SYMBOLS}:
        text = text.replace(symbol, "")
    text = re.sub(r"\b(?:CR|DR)\b", "", text)
    return re.sub(r"[\d\s.,'’+\-()%]", "", text) == ""


def row_nonempty_count(row: Sequence[Any]) -> int:
    """Count non-empty display values in a row."""

    return sum(1 for value in row if clean_text(value))


def trim_rows(rows: Sequence[Sequence[Any]]) -> list[list[Any]]:
    """Drop trailing empty rows and columns from a rectangular-ish table."""

    materialized = [list(row) for row in rows]
    while materialized and row_nonempty_count(materialized[-1]) == 0:
        materialized.pop()
    max_width = 0
    for row in materialized:
        for idx, value in enumerate(row, start=1):
            if clean_text(value):
                max_width = max(max_width, idx)
    if max_width == 0:
        return []
    return [
        row[:max_width] + [""] * max(0, max_width - len(row)) for row in materialized
    ]


def read_csv_rows(
    path: Path,
    *,
    source_bytes: bytes | None = None,
) -> list[list[Any]]:
    """Read a CSV file with deterministic dialect fallback."""

    raw = path.read_bytes() if source_bytes is None else source_bytes
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")

    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel
    return trim_rows(csv.reader(text.splitlines(), dialect))


def read_workbook_tables(
    path: Path,
    *,
    source_bytes: bytes | None = None,
) -> list[dict[str, Any]]:
    """Read formula and cached-value views from one immutable workbook snapshot."""

    captured = path.read_bytes() if source_bytes is None else source_bytes
    cached_workbook = openpyxl.load_workbook(
        BytesIO(captured),
        data_only=True,
        read_only=True,
    )
    formula_workbook = openpyxl.load_workbook(
        BytesIO(captured),
        data_only=False,
        read_only=True,
    )
    tables: list[dict[str, Any]] = []
    cached_sheets = {sheet.title: sheet for sheet in cached_workbook.worksheets}
    for formula_sheet in formula_workbook.worksheets:
        if formula_sheet.sheet_state != "visible":
            continue
        cached_sheet = cached_sheets.get(formula_sheet.title)
        if cached_sheet is None:
            raise ValueError(
                f"Workbook cached view is missing worksheet: {formula_sheet.title}"
            )
        formula_rows = [list(row) for row in formula_sheet.iter_rows(values_only=False)]
        cached_rows = [list(row) for row in cached_sheet.iter_rows(values_only=False)]
        height = max(len(formula_rows), len(cached_rows))
        width = max(
            (len(row) for row in [*formula_rows, *cached_rows]),
            default=0,
        )
        rows: list[list[Any]] = []
        formula_cells: dict[str, dict[str, Any]] = {}
        cell_formats: dict[str, str] = {}
        for row_offset in range(height):
            display_row: list[Any] = []
            for column_offset in range(width):
                formula_cell = (
                    formula_rows[row_offset][column_offset]
                    if row_offset < len(formula_rows)
                    and column_offset < len(formula_rows[row_offset])
                    else None
                )
                cached_cell = (
                    cached_rows[row_offset][column_offset]
                    if row_offset < len(cached_rows)
                    and column_offset < len(cached_rows[row_offset])
                    else None
                )
                formula_value = formula_cell.value if formula_cell is not None else None
                cached_value = cached_cell.value if cached_cell is not None else None
                is_formula = isinstance(
                    formula_value, str
                ) and formula_value.startswith("=")
                # Formula text remains the visible source value; an unverified
                # workbook cache is evidence metadata, never a reported literal.
                display_value = formula_value if is_formula else cached_value
                display_row.append(display_value)
                coordinate_key = f"{row_offset + 1}:{column_offset + 1}"
                number_format = clean_text(
                    formula_cell.number_format if formula_cell is not None else ""
                )
                if clean_text(display_value) or is_formula:
                    cell_formats[coordinate_key] = number_format
                if is_formula:
                    formula_cells[coordinate_key] = {
                        "coordinate": (
                            formula_cell.coordinate
                            if formula_cell is not None
                            else f"{_excel_column_name(column_offset + 1)}{row_offset + 1}"
                        ),
                        "formula": formula_value,
                        "cached_value": clean_text(cached_value),
                        "cache_status": (
                            "present" if cached_value is not None else "missing"
                        ),
                    }
            rows.append(display_row)
        rows = trim_rows(rows)
        cell_formats = {
            key: value
            for key, value in cell_formats.items()
            if int(key.split(":", 1)[0]) <= len(rows)
        }
        tables.append(
            {
                "kind": "worksheet",
                "source_file": path.name,
                "source_path": path.name,
                "sheet_name": formula_sheet.title,
                "table_id": f"{path.name}::{formula_sheet.title}",
                "rows": rows,
                "cell_formats": cell_formats,
                "formula_cells": formula_cells,
            }
        )
    cached_workbook.close()
    formula_workbook.close()
    return tables


def read_pdf_text_table(
    path: Path,
    *,
    source_bytes: bytes | None = None,
) -> dict[str, Any]:
    """Extract text lines from a readable PDF."""

    import pdfplumber

    lines: list[str] = []
    page_count = 0
    pdf_source: Path | BytesIO = path if source_bytes is None else BytesIO(source_bytes)
    with pdfplumber.open(pdf_source) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines.extend(line.strip() for line in text.splitlines() if line.strip())
    if not lines:
        raise ValueError(
            "unsupported_source_layout: PDF has no extractable text; OCR is required"
        )
    rows = [[line] for line in lines]
    return {
        "kind": "pdf_text",
        "source_file": path.name,
        "source_path": path.name,
        "sheet_name": "",
        "table_id": path.name,
        "rows": rows,
        "page_count": page_count,
    }


def _canonical_zip_member_path(member_name: str) -> str:
    """Return a portable member identity or reject an unsafe archive path."""

    normalized = unicodedata.normalize("NFC", member_name.replace("\\", "/"))
    member_path = PurePosixPath(normalized)
    if (
        not normalized
        or normalized.startswith("/")
        or member_path.is_absolute()
        or any(part in {"", ".", ".."} for part in member_path.parts)
    ):
        raise ValueError(f"Unsafe ZIP member path: {member_name}")
    return member_path.as_posix()


def _zip_destination(
    path: Path,
    output_dir: Path,
    *,
    source_bytes: bytes | None = None,
) -> Path:
    """Bind an extraction directory to the current archive bytes."""

    digest = hashlib.sha256(
        path.read_bytes() if source_bytes is None else source_bytes
    ).hexdigest()
    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "-", path.stem).strip("-._") or "archive"
    return output_dir / "extracted_inputs" / f"{safe_stem}-{digest[:24]}"


def _zip_member_manifest(source_bytes: bytes) -> list[dict[str, Any]]:
    """Inventory every canonical member from the exact captured archive bytes."""

    manifest: list[dict[str, Any]] = []
    with zipfile.ZipFile(BytesIO(source_bytes)) as archive:
        seen: dict[str, str] = {}
        for member in archive.infolist():
            if member.is_dir():
                continue
            canonical = _canonical_zip_member_path(member.filename)
            portable_identity = canonical.casefold()
            if portable_identity in seen:
                raise ValueError(
                    "ZIP contains duplicate canonical member paths: "
                    f"{seen[portable_identity]} and {member.filename}"
                )
            if (member.external_attr >> 16) & 0o170000 == 0o120000:
                raise ValueError(
                    f"ZIP symbolic links are not supported: {member.filename}"
                )
            digest = hashlib.sha256()
            byte_count = 0
            with archive.open(member) as source:
                for chunk in iter(lambda: source.read(1024 * 1024), b""):
                    byte_count += len(chunk)
                    digest.update(chunk)
            if byte_count != member.file_size:
                raise ValueError(f"ZIP member size changed while read: {canonical}")
            seen[portable_identity] = member.filename
            manifest.append(
                {
                    "path": canonical,
                    "byte_count": byte_count,
                    "sha256": digest.hexdigest(),
                }
            )
    return sorted(manifest, key=lambda item: str(item["path"]).casefold())


def extract_zip(
    path: Path,
    output_dir: Path,
    *,
    source_bytes: bytes | None = None,
) -> Path:
    """Materialize exactly the current ZIP member set into an isolated snapshot."""

    captured = path.read_bytes() if source_bytes is None else source_bytes
    manifest = _zip_member_manifest(captured)
    expected_by_path = {str(item["path"]): item for item in manifest}
    destination = _zip_destination(path, output_dir, source_bytes=captured)
    destination.parent.mkdir(parents=True, exist_ok=True)
    staging = Path(
        tempfile.mkdtemp(
            prefix=f".{destination.name}.",
            dir=destination.parent,
        )
    )
    committed = False
    try:
        with zipfile.ZipFile(BytesIO(captured)) as archive:
            members: list[tuple[zipfile.ZipInfo, str]] = []
            for member in archive.infolist():
                if member.is_dir():
                    continue
                canonical = _canonical_zip_member_path(member.filename)
                members.append((member, canonical))
            for member, canonical in members:
                target = staging.joinpath(*PurePosixPath(canonical).parts)
                target.parent.mkdir(parents=True, exist_ok=True)
                digest = hashlib.sha256()
                byte_count = 0
                with archive.open(member) as source, target.open("wb") as handle:
                    for chunk in iter(lambda: source.read(1024 * 1024), b""):
                        byte_count += len(chunk)
                        digest.update(chunk)
                        handle.write(chunk)
                expected = expected_by_path[canonical]
                if (
                    byte_count != expected["byte_count"]
                    or digest.hexdigest() != expected["sha256"]
                ):
                    raise ValueError(f"ZIP member changed while extracted: {canonical}")
        if destination.exists():
            shutil.rmtree(destination)
        staging.replace(destination)
        committed = True
    finally:
        if not committed:
            shutil.rmtree(staging, ignore_errors=True)
    return destination


def discover_input_files(
    input_path: Path,
    output_dir: Path,
    *,
    zip_source_bytes: bytes | None = None,
) -> list[Path]:
    """Return supported files from a file, folder, or ZIP archive."""

    requested_path = input_path.expanduser()
    if requested_path.is_symlink():
        raise ValueError("Report Builder input path cannot be a symbolic link")
    requested_metadata = requested_path.lstat()
    if (
        stat.S_ISREG(requested_metadata.st_mode)
        and requested_metadata.st_nlink > 1
        and requested_path.suffix.lower() in SUPPORTED_SUFFIXES
    ):
        raise ValueError("Report Builder input source cannot be hard-linked")
    path = requested_path.resolve()
    if not path.exists():
        raise FileNotFoundError(path)
    if path.is_file() and path.suffix.lower() == ".zip":
        path = extract_zip(path, output_dir, source_bytes=zip_source_bytes)
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED_SUFFIXES - {".zip"} else []
    files: list[Path] = []
    pending = [path]
    while pending:
        directory = pending.pop()
        for item in directory.iterdir():
            metadata = item.lstat()
            if stat.S_ISLNK(metadata.st_mode):
                relative = item.relative_to(path).as_posix()
                raise ValueError(
                    "Report Builder input directory contains a symbolic link: "
                    f"{relative}"
                )
            if stat.S_ISDIR(metadata.st_mode):
                pending.append(item)
                continue
            if stat.S_ISREG(metadata.st_mode):
                if (
                    metadata.st_nlink > 1
                    and item.suffix.lower() in SUPPORTED_SUFFIXES - {".zip"}
                ):
                    relative = item.relative_to(path).as_posix()
                    raise ValueError(
                        "Report Builder input directory contains a hard-linked "
                        f"source: {relative}"
                    )
                if item.suffix.lower() in SUPPORTED_SUFFIXES - {".zip"}:
                    files.append(item)
                continue
            if item.suffix.lower() in SUPPORTED_SUFFIXES - {".zip"}:
                relative = item.relative_to(path).as_posix()
                raise ValueError(
                    "Report Builder input directory contains an unsupported file "
                    f"type: {relative}"
                )
    return sorted(files, key=lambda item: item.as_posix().lower())


def _capture_source(
    path: Path,
    *,
    identity_key: str | None = None,
) -> tuple[bytes, dict[str, Any]]:
    """Capture immutable bytes and a stable receipt for one source file."""

    source = path.expanduser().resolve()
    source_identity = identity_key or source.as_posix()
    path_digest = hashlib.sha256(source_identity.encode("utf-8")).hexdigest()
    before_count, before_digest = file_snapshot(source)
    artifact_id = f"source.{path_digest}.{before_digest}"
    root_id = f"source_{path_digest}"
    before_receipt = artifact_receipt(
        source.parent,
        source,
        artifact_id=artifact_id,
        root_id=root_id,
        role="source",
    )
    captured = source.read_bytes()
    after_receipt = artifact_receipt(
        source.parent,
        source,
        artifact_id=artifact_id,
        root_id=root_id,
        role="source",
    )
    captured_digest = hashlib.sha256(captured).hexdigest()
    if (
        before_receipt != after_receipt
        or before_receipt["byte_count"] != before_count
        or before_receipt["sha256"] != before_digest
        or len(captured) != before_count
        or captured_digest != before_digest
    ):
        raise ValueError(f"Source changed while it was captured: {source}")
    return captured, {
        "identity_key": source_identity,
        "root_path": str(source.parent),
        "receipt": before_receipt,
    }


def _tables_from_captured_source(
    path: Path,
    source_bytes: bytes,
) -> list[dict[str, Any]]:
    """Parse one source only from its captured immutable bytes."""

    suffix = path.suffix.lower()
    if suffix in WORKBOOK_SUFFIXES:
        return read_workbook_tables(path, source_bytes=source_bytes)
    if suffix == ".csv":
        return [
            {
                "kind": "csv",
                "source_file": path.name,
                "source_path": path.name,
                "sheet_name": "",
                "table_id": path.name,
                "rows": read_csv_rows(path, source_bytes=source_bytes),
            }
        ]
    if suffix in TEXT_SUFFIXES:
        return [read_pdf_text_table(path, source_bytes=source_bytes)]
    return []


def _disambiguate_table_ids(tables: list[dict[str, Any]]) -> None:
    """Keep familiar IDs when unique and bind collisions to source identity."""

    counts = Counter(clean_text(table.get("table_id")) for table in tables)
    for table in tables:
        table_id = clean_text(table.get("table_id"))
        if counts[table_id] <= 1:
            continue
        source_ref = clean_text(table.get("source_artifact_ref"))
        path_fragment = source_ref.split(".")[1][:16] if source_ref else "unknown"
        sheet = clean_text(table.get("sheet_name")) or "table"
        table["table_id"] = f"{table.get('source_file')}@{path_fragment}::{sheet}"


def load_tables(input_path: Path, output_dir: Path) -> list[dict[str, Any]]:
    """Load all inspectable tables/text blocks from the input path."""

    tables: list[dict[str, Any]] = []
    input_source = input_path.expanduser().resolve()
    archive_source_receipt: dict[str, Any] | None = None
    archive_bytes: bytes | None = None
    archive_member_manifest: list[dict[str, Any]] = []
    if input_source.is_file() and input_source.suffix.lower() == ".zip":
        archive_bytes, archive_source_receipt = _capture_source(
            input_source,
            identity_key=source_identity_key(output_dir, input_source),
        )
        archive_member_manifest = _zip_member_manifest(archive_bytes)
    archive_root = (
        _zip_destination(
            input_source,
            output_dir,
            source_bytes=archive_bytes,
        ).resolve()
        if archive_bytes is not None
        else None
    )
    for file_path in discover_input_files(
        input_path,
        output_dir,
        zip_source_bytes=archive_bytes,
    ):
        source_metadata: dict[str, Any] = {}
        try:
            source_identity = source_identity_key(output_dir, file_path)
            archive_member_binding: dict[str, Any] | None = None
            if archive_root is not None and archive_source_receipt is not None:
                member_path = file_path.resolve().relative_to(archive_root).as_posix()
                source_identity = (
                    f"{archive_source_receipt['identity_key']}::{member_path}"
                )
            source_bytes, source_receipt = _capture_source(
                file_path,
                identity_key=source_identity,
            )
            if archive_root is not None and archive_source_receipt is not None:
                expected_member = next(
                    (
                        item
                        for item in archive_member_manifest
                        if item["path"] == member_path
                    ),
                    None,
                )
                if expected_member is None:
                    raise ValueError(
                        f"Extracted ZIP member is absent from manifest: {member_path}"
                    )
                member_receipt = source_receipt["receipt"]
                if (
                    member_receipt["byte_count"] != expected_member["byte_count"]
                    or member_receipt["sha256"] != expected_member["sha256"]
                ):
                    raise ValueError(
                        "Extracted ZIP member does not derive from the captured "
                        f"archive: {member_path}"
                    )
                archive_member_binding = {
                    "container_artifact_id": archive_source_receipt["receipt"][
                        "artifact_id"
                    ],
                    "member_path": member_path,
                    "member_artifact_id": member_receipt["artifact_id"],
                    "byte_count": expected_member["byte_count"],
                    "sha256": expected_member["sha256"],
                }
            source_metadata = {
                "source_artifact_ref": source_receipt["receipt"]["artifact_id"],
                "source_receipt": source_receipt,
                **(
                    {"container_source_receipt": archive_source_receipt}
                    if archive_source_receipt is not None
                    else {}
                ),
                **(
                    {
                        "container_member_manifest": archive_member_manifest,
                        "archive_member_binding": archive_member_binding,
                    }
                    if archive_member_binding is not None
                    else {}
                ),
            }
            parsed_tables = _tables_from_captured_source(file_path, source_bytes)
            for table in parsed_tables:
                table.update(source_metadata)
            tables.extend(parsed_tables)
        except (OSError, ValueError, zipfile.BadZipFile) as exc:
            LOGGER.warning("Could not inspect %s: %s", file_path, exc)
            public_error = str(exc).replace(
                file_path.expanduser().resolve().as_posix(),
                file_path.name,
            )
            tables.append(
                {
                    "kind": "error",
                    "source_file": file_path.name,
                    "source_path": file_path.name,
                    "sheet_name": "",
                    "table_id": file_path.name,
                    "rows": [],
                    "error": public_error,
                    **source_metadata,
                }
            )
    _disambiguate_table_ids(tables)
    if any(isinstance(table.get("source_receipt"), dict) for table in tables):
        write_source_index(output_dir, tables)
    return tables


def load_indexed_tables(
    output_dir: Path,
    *,
    persist_source_index: bool = True,
) -> list[dict[str, Any]]:
    """Reload exactly the receipted sources used by the reviewed generation."""

    source_index = validate_source_index(output_dir)
    tables: list[dict[str, Any]] = []
    for source_record in source_index["sources"]:
        if not isinstance(source_record, dict):
            raise ValueError("Malformed Report Builder source index")
        root_path = source_record.get("root_path")
        identity_key = source_record.get("identity_key")
        receipt = source_record.get("receipt")
        if (
            not isinstance(root_path, str)
            or not isinstance(identity_key, str)
            or not identity_key
            or not isinstance(receipt, dict)
        ):
            raise ValueError("Malformed Report Builder source record")
        source_path = resolve_source_record_path(output_dir, source_record)
        captured, current_receipt = _capture_source(
            source_path,
            identity_key=identity_key,
        )
        if (
            current_receipt["identity_key"] != source_record["identity_key"]
            or current_receipt["receipt"] != source_record["receipt"]
        ):
            raise ValueError(
                "Source receipt does not match the reviewed generation: "
                f"{source_record.get('artifact_id')}"
            )
        source_metadata = {
            "source_artifact_ref": receipt["artifact_id"],
            "source_receipt": source_record,
        }
        try:
            parsed_tables = _tables_from_captured_source(source_path, captured)
            for table in parsed_tables:
                table.update(source_metadata)
            tables.extend(parsed_tables)
        except (OSError, ValueError, zipfile.BadZipFile) as exc:
            LOGGER.warning("Could not replay %s: %s", source_path, exc)
            public_error = str(exc).replace(
                source_path.expanduser().resolve().as_posix(),
                source_path.name,
            )
            tables.append(
                {
                    "kind": "error",
                    "source_file": source_path.name,
                    "source_path": source_path.name,
                    "sheet_name": "",
                    "table_id": source_path.name,
                    "rows": [],
                    "error": public_error,
                    **source_metadata,
                }
            )
    _disambiguate_table_ids(tables)
    if persist_source_index:
        write_source_index(output_dir, tables)
    return tables


def header_candidate_index(rows: Sequence[Sequence[Any]]) -> int | None:
    """Pick a text-supported header row without consuming numeric data."""

    best_idx: int | None = None
    best_score = -1.0
    for idx, row in enumerate(rows[:20]):
        nonempty = row_nonempty_count(row)
        if nonempty == 0:
            continue
        labels = " ".join(norm_label(value) for value in row if clean_text(value))
        label_hits = sum(
            1
            for token in (
                "year",
                "period",
                "amount",
                "value",
                "totale",
                "importo",
                "descrizione",
                "budget",
                "actual",
                "saldo",
                "conto",
                "metric",
                "line",
                "item",
                "description",
                "date",
                "currency",
                "unit",
                "account",
            )
            if token in labels
        )
        text_cells = sum(
            1 for value in row if clean_text(value) and parse_amount(value) is None
        )
        if text_cells < 2 and not (text_cells >= 1 and label_hits > 0):
            continue
        if any(
            _looks_numeric_candidate(value)
            for earlier_row in rows[:idx]
            for value in earlier_row
            if clean_text(value)
        ):
            continue
        score = nonempty + label_hits * 2 + text_cells * 0.5 - idx * 0.1
        if score > best_score:
            best_idx = idx
            best_score = score
    return best_idx


def unique_headers(values: Sequence[Any], width: int) -> list[str]:
    """Return stable unique header labels."""

    headers: list[str] = []
    seen: dict[str, int] = {}
    for idx in range(width):
        base = clean_text(values[idx] if idx < len(values) else "")
        if not base:
            base = f"column_{idx + 1}"
        count = seen.get(base, 0)
        seen[base] = count + 1
        headers.append(base if count == 0 else f"{base}_{count + 1}")
    return headers


def rows_as_dicts(
    rows: Sequence[Sequence[Any]], header_idx: int | None
) -> list[dict[str, str]]:
    """Return display rows keyed by detected headers."""

    if not rows:
        return []
    width = max((len(row) for row in rows), default=0)
    if header_idx is None:
        headers = [f"column_{idx + 1}" for idx in range(width)]
        data_rows = rows
    else:
        headers = unique_headers(rows[header_idx], width)
        data_rows = rows[header_idx + 1 :]
    return [
        {
            headers[idx]: clean_text(row[idx] if idx < len(row) else "")
            for idx in range(width)
        }
        for row in data_rows
        if row_nonempty_count(row) > 0
    ]


def _redact_numeric_preview_rows(
    rows: Sequence[dict[str, str]],
) -> list[dict[str, str]]:
    """Keep source context while withholding unledgered numeric cell values."""

    redacted: list[dict[str, str]] = []
    for row in rows:
        redacted.append(
            {
                header: (
                    "[numeric source value withheld]"
                    if re.search(r"\d", clean_text(value))
                    or clean_text(value).startswith(("=", "+", "-", "@"))
                    else clean_text(value)
                )
                for header, value in row.items()
            }
        )
    return redacted


def table_numeric_profile(
    rows: Sequence[Sequence[Any]],
    header_idx: int | None,
    *,
    formula_cells: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Inventory literal and formula numeric candidates without semantic selection."""

    if not rows:
        return {"numeric_cells": 0, "numeric_columns": []}
    width = max((len(row) for row in rows), default=0)
    headers = (
        unique_headers(rows[header_idx], width)
        if header_idx is not None
        else [f"column_{idx + 1}" for idx in range(width)]
    )
    start_idx = header_idx + 1 if header_idx is not None else 0
    formulas = formula_cells or {}
    column_sums: list[dict[str, Any]] = []
    numeric_cells = 0
    for col_idx, header in enumerate(headers):
        values: list[Decimal] = []
        numeric_rows: list[int] = []
        candidate_rows: list[int] = []
        for row_idx, row in enumerate(rows[start_idx:], start=start_idx + 1):
            source_value = row[col_idx] if col_idx < len(row) else None
            if _looks_numeric_candidate(source_value):
                candidate_rows.append(row_idx)
            value = parse_amount(source_value)
            if value is not None:
                values.append(value)
                numeric_rows.append(row_idx)
        formula_rows = [
            row_idx
            for row_idx in range(start_idx + 1, len(rows) + 1)
            if f"{row_idx}:{col_idx + 1}" in formulas
        ]
        numeric_cells += len(
            set(candidate_rows) | set(numeric_rows) | set(formula_rows)
        )
        if candidate_rows or values or formula_rows:
            total = sum(values, Decimal("0"))
            column_sums.append(
                {
                    "column": header,
                    "column_index": col_idx + 1,
                    "numeric_rows": numeric_rows,
                    "numeric_count": len(values),
                    "sum": decimal_text(total) if values else None,
                    "min": decimal_text(min(values)) if values else None,
                    "max": decimal_text(max(values)) if values else None,
                    **(
                        {
                            "formula_rows": formula_rows,
                            "formula_cell_count": len(formula_rows),
                        }
                        if formula_rows
                        else {}
                    ),
                    **(
                        {"candidate_rows": candidate_rows}
                        if candidate_rows != numeric_rows
                        else {}
                    ),
                }
            )
    return {"numeric_cells": numeric_cells, "numeric_columns": column_sums}


def _numeric_candidate_columns(
    columns: Sequence[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Remove unreviewed values while preserving candidate-cell diagnostics."""

    return [
        {
            key: value
            for key, value in column.items()
            if key not in {"sum", "min", "max"}
        }
        for column in columns
        if isinstance(column, dict)
    ]


def _numeric_cell_inventory(
    rows: Sequence[Sequence[Any]],
    header_idx: int | None,
    *,
    cell_formats: dict[str, Any] | None = None,
    formula_cells: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    """Bind each nonblank cell to its coordinate, format, and formula/cache state."""

    if not rows:
        return []
    width = max((len(row) for row in rows), default=0)
    headers = (
        unique_headers(rows[header_idx], width)
        if header_idx is not None
        else [f"column_{idx + 1}" for idx in range(width)]
    )
    start_idx = header_idx + 1 if header_idx is not None else 0
    formats = cell_formats or {}
    formulas = formula_cells or {}
    inventory: list[dict[str, Any]] = []
    for col_idx, header in enumerate(headers):
        cells: list[dict[str, Any]] = []
        for row_idx, row in enumerate(rows[start_idx:], start=start_idx + 1):
            value = row[col_idx] if col_idx < len(row) else None
            source_text = clean_text(value)
            if not source_text:
                continue
            coordinate_key = f"{row_idx}:{col_idx + 1}"
            formula_metadata = formulas.get(coordinate_key)
            if formula_metadata is not None and not isinstance(
                formula_metadata, Mapping
            ):
                raise ValueError("Malformed workbook formula metadata")
            cells.append(
                {
                    "row": row_idx,
                    "coordinate": (
                        clean_text(formula_metadata.get("coordinate"))
                        if isinstance(formula_metadata, Mapping)
                        else f"{_excel_column_name(col_idx + 1)}{row_idx}"
                    ),
                    "source_text": source_text,
                    "number_format": clean_text(formats.get(coordinate_key)),
                    "formula": (
                        clean_text(formula_metadata.get("formula"))
                        if isinstance(formula_metadata, Mapping)
                        else None
                    ),
                    "formula_cached_value": (
                        clean_text(formula_metadata.get("cached_value")) or None
                        if isinstance(formula_metadata, Mapping)
                        else None
                    ),
                    "formula_cache_status": (
                        clean_text(formula_metadata.get("cache_status"))
                        if isinstance(formula_metadata, Mapping)
                        else "not_formula"
                    ),
                }
            )
        inventory.append(
            {
                "column": header,
                "column_index": col_idx + 1,
                "nonblank_cells": cells,
            }
        )
    return inventory


def _currency_markers(source_text: str, number_format: str) -> list[str]:
    """Return unambiguous currency markers without inferring from bare `$`."""

    combined = f"{source_text} {number_format}".upper()
    markers = {code for code in _CURRENCY_CODES if re.search(rf"\b{code}\b", combined)}
    markers.update(
        currency for symbol, currency in _CURRENCY_SYMBOLS.items() if symbol in combined
    )
    return sorted(markers)


def _ambiguous_currency_symbols(source_text: str, number_format: str) -> list[str]:
    """Return symbols whose ISO currency must come from reviewed context."""

    combined = f"{source_text} {number_format}"
    return sorted(
        symbol for symbol in _AMBIGUOUS_CURRENCY_SYMBOLS if symbol in combined
    )


def _numeric_contract(
    *,
    numeric_locale: object,
    currency: object | None,
    unit: object,
    scale: object,
    parse_policy: object,
    sign_policy: object,
) -> dict[str, Any]:
    """Validate reviewed numeric semantics with an explicit parser contract."""

    locale = normalize_language(numeric_locale, default="")
    if locale not in NUMERIC_LOCALE_SEPARATORS:
        raise ValueError(
            f"Reviewed numeric locale must be one of {sorted(NUMERIC_LOCALE_SEPARATORS)}"
        )
    policy = clean_text(parse_policy)
    if policy != NUMERIC_PARSE_POLICY:
        raise ValueError(f"Unsupported numeric parse policy: {policy}")
    normalized_sign_policy = clean_text(sign_policy)
    if normalized_sign_policy not in NUMERIC_SIGN_POLICIES:
        raise ValueError(f"Unsupported numeric sign policy: {normalized_sign_policy}")
    normalized_unit = clean_text(unit).lower()
    if normalized_unit not in NUMERIC_UNITS:
        raise ValueError(f"Unsupported numeric unit: {normalized_unit}")
    normalized_currency = clean_text(currency).upper() or None
    if normalized_unit == "currency":
        if (
            normalized_currency is None
            or re.fullmatch(r"[A-Z]{3}", normalized_currency) is None
        ):
            raise ValueError("Currency measures require a reviewed ISO currency")
    elif normalized_currency is not None:
        raise ValueError("Non-currency measures cannot declare a currency")
    scale_text = clean_text(scale)
    try:
        scale_value = parse_canonical_decimal(scale_text, label="numeric scale")
    except MoneyValidationError as exc:
        raise ValueError(str(exc)) from exc
    if scale_value <= 0:
        raise ValueError("Numeric scale must be greater than zero")
    decimal_separator, thousands_separator = NUMERIC_LOCALE_SEPARATORS[locale]
    return {
        "policy_id": policy,
        "locale": locale,
        "decimal_separator": decimal_separator,
        "thousands_separator": thousands_separator,
        "currency": normalized_currency,
        "unit": normalized_unit,
        "scale": decimal_text(scale_value),
        "sign_policy": normalized_sign_policy,
    }


def _signed_numeric_value(value: Decimal, sign_policy: str) -> Decimal:
    """Apply only the exact reviewer-selected, mechanically replayable sign rule."""

    if sign_policy == "as_presented_v1":
        return value
    if sign_policy == "invert_v1":
        return -value
    raise ValueError(f"Unsupported numeric sign policy: {sign_policy}")


def _strict_reviewed_numeric_profile(
    inventory: Sequence[dict[str, Any]],
    columns: Sequence[str],
    contract: dict[str, Any],
    cell_reviews: Mapping[str, Mapping[int, str]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Replay an explicit include/exclude disposition for every selected cell."""

    by_name = {
        clean_text(item.get("column")): item
        for item in inventory
        if isinstance(item, dict) and clean_text(item.get("column"))
    }
    unknown = set(columns) - set(by_name)
    if unknown:
        raise ValueError(f"Unknown numeric measure columns: {sorted(unknown)}")
    if set(cell_reviews) != set(columns):
        raise ValueError(
            "Selected numeric columns require one explicit cell-disposition map each"
        )
    scale_value = parse_canonical_decimal(contract["scale"], label="numeric scale")
    reviewed_columns: list[dict[str, Any]] = []
    dispositions: list[dict[str, Any]] = []
    observed_currencies: set[str] = set()
    for name in columns:
        item = by_name[name]
        cells = item.get("nonblank_cells")
        if not isinstance(cells, list) or not cells:
            raise ValueError(f"Reviewed measure column has no nonblank cells: {name}")
        raw_cell_review = cell_reviews.get(name)
        if not isinstance(raw_cell_review, Mapping):
            raise ValueError(f"Reviewed measure cell dispositions are missing: {name}")
        if any(
            not isinstance(row, int) or isinstance(row, bool) for row in raw_cell_review
        ):
            raise ValueError(
                f"Reviewed measure cell disposition rows are invalid: {name}"
            )
        expected_rows = {
            int(cell["row"])
            for cell in cells
            if isinstance(cell, dict)
            and isinstance(cell.get("row"), int)
            and not isinstance(cell.get("row"), bool)
        }
        reviewed_rows = set(raw_cell_review)
        if reviewed_rows != expected_rows:
            raise ValueError(
                f"Reviewed measure cell dispositions do not close for {name}: "
                f"expected {sorted(expected_rows)}, got {sorted(reviewed_rows)}"
            )
        if any(
            disposition not in {"include", "exclude"}
            for disposition in raw_cell_review.values()
        ):
            raise ValueError(
                f"Reviewed measure cell dispositions are invalid for {name}"
            )
        values: list[Decimal] = []
        numeric_rows: list[int] = []
        cell_dispositions: list[dict[str, Any]] = []
        unresolved_rows: list[int] = []
        for cell in cells:
            if not isinstance(cell, dict):
                raise ValueError(f"Malformed numeric cell inventory for {name}")
            row_number = cell.get("row")
            source_text = clean_text(cell.get("source_text"))
            number_format = clean_text(cell.get("number_format"))
            coordinate = clean_text(cell.get("coordinate"))
            formula = clean_text(cell.get("formula"))
            formula_cached_value = clean_text(cell.get("formula_cached_value"))
            formula_cache_status = clean_text(cell.get("formula_cache_status"))
            if not isinstance(row_number, int) or isinstance(row_number, bool):
                raise ValueError(f"Malformed numeric source row for {name}")
            disposition = raw_cell_review[row_number]
            base_disposition = {
                "row": row_number,
                "coordinate": coordinate,
                "source_text": source_text,
                "number_format": number_format,
                "formula": formula or None,
                "formula_cached_value": formula_cached_value or None,
                "formula_cache_status": formula_cache_status,
            }
            if disposition == "exclude":
                cell_dispositions.append(
                    {
                        **base_disposition,
                        "status": "excluded_by_review",
                        "reason": "reviewed_cell_exclusion",
                    }
                )
                continue
            markers = _currency_markers(source_text, number_format)
            ambiguous_symbols = _ambiguous_currency_symbols(
                source_text,
                number_format,
            )
            observed_currencies.update(markers)
            percent_formatted = "%" in number_format
            reason = ""
            try:
                if formula:
                    raise MoneyValidationError(
                        "formula cells require a separately verified recalculation policy"
                    )
                value = parse_localized_decimal(
                    source_text,
                    label=f"{name} row {row_number}",
                    decimal_separator=contract["decimal_separator"],
                    thousands_separator=contract["thousands_separator"],
                )
                if len(markers) > 1:
                    raise MoneyValidationError("cell has multiple currency markers")
                if contract["unit"] == "currency":
                    if markers and markers != [contract["currency"]]:
                        raise MoneyValidationError(
                            "cell currency does not match reviewed currency"
                        )
                    if percent_formatted:
                        raise MoneyValidationError(
                            "currency measure uses a percentage number format"
                        )
                elif markers or ambiguous_symbols:
                    raise MoneyValidationError(
                        "non-currency measure contains a currency marker or symbol"
                    )
                if percent_formatted and contract["unit"] != "percentage":
                    raise MoneyValidationError(
                        "percentage number format requires percentage unit"
                    )
            except MoneyValidationError as exc:
                value = None
                reason = str(exc)
            if value is None:
                unresolved_rows.append(row_number)
                cell_dispositions.append(
                    {
                        **base_disposition,
                        "status": "unresolved",
                        "reason": reason,
                    }
                )
                continue
            signed = _signed_numeric_value(value, contract["sign_policy"])
            scaled = signed * scale_value
            values.append(scaled)
            numeric_rows.append(row_number)
            cell_dispositions.append(
                {
                    **base_disposition,
                    "status": "included",
                    "currency_marker": markers[0] if markers else None,
                    "ambiguous_currency_symbols": ambiguous_symbols,
                    "canonical_value": decimal_text(value),
                    "signed_value": decimal_text(signed),
                    "scaled_value": decimal_text(scaled),
                }
            )
        if unresolved_rows:
            raise ValueError(
                f"Reviewed measure column {name} has unresolved nonblank rows: "
                f"{unresolved_rows}"
            )
        if contract["unit"] == "currency" and len(observed_currencies) > 1:
            raise ValueError(
                "Reviewed currency measure contains mixed explicit currencies: "
                f"{sorted(observed_currencies)}"
            )
        if not values:
            raise ValueError(f"Reviewed measure column has no included cells: {name}")
        total = sum(values, Decimal("0"))
        reviewed_columns.append(
            {
                "column": name,
                "column_index": int(item["column_index"]),
                "numeric_rows": numeric_rows,
                "numeric_count": len(values),
                "sum": decimal_text(total),
                "min": decimal_text(min(values)),
                "max": decimal_text(max(values)),
                "currency": contract["currency"],
                "unit": contract["unit"],
                "scale": contract["scale"],
                "sign_policy": contract["sign_policy"],
            }
        )
        dispositions.append(
            {
                "column": name,
                "column_index": int(item["column_index"]),
                "nonblank_cell_count": len(cell_dispositions),
                "cells": cell_dispositions,
            }
        )
    return reviewed_columns, dispositions


def _excluded_numeric_column_dispositions(
    inventory: Sequence[dict[str, Any]],
    columns: Sequence[str],
) -> list[dict[str, Any]]:
    """Expand reviewed column exclusions to exact cell-level dispositions."""

    by_name = {
        clean_text(item.get("column")): item
        for item in inventory
        if isinstance(item, dict) and clean_text(item.get("column"))
    }
    result: list[dict[str, Any]] = []
    for name in columns:
        item = by_name.get(name)
        if not isinstance(item, dict):
            raise ValueError(f"Unknown excluded numeric candidate column: {name}")
        cells = item.get("nonblank_cells")
        if not isinstance(cells, list) or not cells:
            raise ValueError(f"Excluded numeric candidate has no cells: {name}")
        result.append(
            {
                "column": name,
                "column_index": int(item["column_index"]),
                "status": "excluded_by_review",
                "nonblank_cell_count": len(cells),
                "cells": [
                    {
                        "row": int(cell["row"]),
                        "coordinate": clean_text(cell.get("coordinate")),
                        "source_text": clean_text(cell.get("source_text")),
                        "number_format": clean_text(cell.get("number_format")),
                        "formula": clean_text(cell.get("formula")) or None,
                        "formula_cached_value": (
                            clean_text(cell.get("formula_cached_value")) or None
                        ),
                        "formula_cache_status": clean_text(
                            cell.get("formula_cache_status")
                        ),
                        "status": "excluded_by_review",
                        "reason": "reviewed_column_exclusion",
                    }
                    for cell in cells
                    if isinstance(cell, dict)
                ],
            }
        )
    return result


def _numeric_decision_id(section_key: str, source_artifact_ref: str) -> str:
    return (
        "decision.report_numeric_measures."
        + hashlib.sha256(
            f"{section_key}:{source_artifact_ref}".encode("utf-8")
        ).hexdigest()
    )


def suggest_section(table: dict[str, Any]) -> dict[str, Any]:
    """Suggest a report section for a table using transparent keyword matching."""

    rows = table.get("rows", [])
    preview_text = " ".join(
        clean_text(value)
        for row in rows[:12]
        for value in row[:10]
        if clean_text(value)
    )
    source_text = " ".join(
        [
            clean_text(table.get("source_file")),
            clean_text(table.get("sheet_name")),
            preview_text,
        ]
    )
    normalized = norm_label(source_text)
    best_section = ""
    best_score = 0
    for section, aliases in SECTION_ALIASES.items():
        score = sum(1 for alias in aliases if alias in normalized)
        if score > best_score:
            best_section = section
            best_score = score
    confidence = min(0.95, 0.35 + best_score * 0.2) if best_score else 0.0
    return {"section": best_section, "confidence": round(confidence, 2)}


def inspect_table(table: dict[str, Any]) -> dict[str, Any]:
    """Return a JSON-safe inspection record for one table/text block."""

    rows = table.get("rows", [])
    header_idx = header_candidate_index(rows)
    preview = _redact_numeric_preview_rows(
        rows_as_dicts(rows, header_idx)[:MODEL_CONTEXT_PREVIEW_ROWS]
    )
    profile = table_numeric_profile(
        rows,
        header_idx,
        formula_cells=table.get("formula_cells"),
    )
    numeric_candidates = _numeric_candidate_columns(profile["numeric_columns"])
    headerless_profile = table_numeric_profile(
        rows,
        None,
        formula_cells=table.get("formula_cells"),
    )
    suggestion = suggest_section(table)
    return {
        "table_id": table.get("table_id", ""),
        "kind": table.get("kind", ""),
        "source_file": table.get("source_file", ""),
        "source_path": table.get("source_file", ""),
        "source_artifact_ref": table.get("source_artifact_ref", ""),
        "sheet_name": table.get("sheet_name", ""),
        "page_count": table.get("page_count", 0),
        "error": table.get("error", ""),
        "row_count": len([row for row in rows if row_nonempty_count(row) > 0]),
        "column_count": max((len(row) for row in rows), default=0),
        "header_row": header_idx + 1 if header_idx is not None else None,
        "numeric_cell_count": profile["numeric_cells"],
        "numeric_columns": numeric_candidates,
        "numeric_measure_cells": _numeric_cell_inventory(
            rows,
            header_idx,
            cell_formats=table.get("cell_formats"),
            formula_cells=table.get("formula_cells"),
        ),
        "header_review_options": {
            "detected_header_row": header_idx + 1 if header_idx is not None else None,
            "supported_choices": [
                *([header_idx + 1] if header_idx is not None else []),
                None,
            ],
        },
        "headerless_numeric_columns": _numeric_candidate_columns(
            headerless_profile["numeric_columns"]
        ),
        "headerless_numeric_measure_cells": _numeric_cell_inventory(
            rows,
            None,
            cell_formats=table.get("cell_formats"),
            formula_cells=table.get("formula_cells"),
        ),
        "suggested_section": suggestion["section"],
        "suggestion_confidence": suggestion["confidence"],
        "preview_rows": preview,
    }


def _public_table_inspection(table: dict[str, Any]) -> dict[str, Any]:
    """Remove full-population cell values from a model-visible table inventory."""

    public = copy.deepcopy(table)
    public.pop("numeric_measure_cells", None)
    public.pop("headerless_numeric_measure_cells", None)
    return public


def _model_inspection_packet(inspection: Mapping[str, Any]) -> dict[str, Any]:
    """Build the bounded default model packet from private deterministic control.

    The selection and omission rules are fixed because packet shape and disclosure
    bounds are mechanically verifiable privacy controls.  They do not decide which
    evidence is professionally relevant; that judgment remains with the model and
    reviewer through explicit targeted expansion.
    """

    packet = copy.deepcopy(dict(inspection))
    packet["tables"] = [
        _public_table_inspection(table)
        for table in inspection.get("tables", [])
        if isinstance(table, dict)
    ]
    packet["model_context"] = {
        "packet_role": "bounded_default",
        "full_population_processed_locally": True,
        "full_cell_inventory_model_visible": False,
        "preview_rows_per_table": MODEL_CONTEXT_PREVIEW_ROWS,
        "numeric_values_in_preview": "redacted",
        "targeted_expansion": {
            "command": "scripts/expand_model_context.py",
            "selection": "one table, exact columns, and one source-row range",
            "max_rows_per_packet": MODEL_CONTEXT_MAX_EXPANSION_ROWS,
            "max_columns_per_packet": MODEL_CONTEXT_MAX_EXPANSION_COLUMNS,
            "repeatable": True,
        },
    }
    return packet


def _packet_sha256(value: Mapping[str, Any]) -> str:
    """Hash canonical JSON packet content without making semantic judgments."""

    return hashlib.sha256(
        json.dumps(
            value,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()


def build_model_context_expansion(
    inspection_control: Mapping[str, Any],
    *,
    table_id: str,
    header_row: int | None,
    columns: Sequence[str],
    row_start: int,
    row_limit: int,
    purpose: str,
) -> dict[str, Any]:
    """Return one bounded, explicit table slice from private inspection control.

    Bounds, exact selectors, and receipt generation are deterministic security and
    auditability controls.  The caller supplies the purpose and selected evidence;
    this helper does not classify relevance or decide professional sufficiency.
    """

    selected_table_id = clean_text(table_id)
    selected_purpose = clean_text(purpose)
    selected_columns = [clean_text(item) for item in columns if clean_text(item)]
    if not selected_table_id:
        raise ValueError("Model-context expansion requires a table_id")
    if not selected_purpose or len(selected_purpose) > 240:
        raise ValueError(
            "Model-context expansion purpose must contain 1 to 240 characters"
        )
    if isinstance(row_start, bool) or row_start < 1:
        raise ValueError("Model-context expansion row_start must be positive")
    if (
        isinstance(row_limit, bool)
        or row_limit < 1
        or row_limit > MODEL_CONTEXT_MAX_EXPANSION_ROWS
    ):
        raise ValueError(
            "Model-context expansion row_limit must be between 1 and "
            f"{MODEL_CONTEXT_MAX_EXPANSION_ROWS}"
        )
    if (
        not selected_columns
        or len(selected_columns) > MODEL_CONTEXT_MAX_EXPANSION_COLUMNS
    ):
        raise ValueError(
            "Model-context expansion requires 1 to "
            f"{MODEL_CONTEXT_MAX_EXPANSION_COLUMNS} exact columns"
        )
    if len(selected_columns) != len(set(selected_columns)):
        raise ValueError("Model-context expansion columns must be unique")

    table = next(
        (
            item
            for item in inspection_control.get("tables", [])
            if isinstance(item, dict)
            and clean_text(item.get("table_id")) == selected_table_id
        ),
        None,
    )
    if table is None:
        raise ValueError("Model-context expansion references an unknown table_id")
    detected_header_row = table.get("header_row")
    if header_row is None:
        inventory = table.get("headerless_numeric_measure_cells")
        first_data_row = 1
    elif header_row == detected_header_row:
        inventory = table.get("numeric_measure_cells")
        first_data_row = header_row + 1
    else:
        raise ValueError(
            "Model-context expansion header_row must be the detected row or none"
        )
    if not isinstance(inventory, list):
        raise ValueError("Inspection control lacks the full cell inventory")
    inventory_by_column = {
        clean_text(item.get("column")): item
        for item in inventory
        if isinstance(item, dict) and clean_text(item.get("column"))
    }
    unknown = set(selected_columns) - set(inventory_by_column)
    if unknown:
        raise ValueError(
            f"Model-context expansion references unknown columns: {sorted(unknown)}"
        )
    if row_start < first_data_row:
        raise ValueError(
            f"Model-context expansion row_start must be at least {first_data_row}"
        )
    available_rows = [
        int(cell["row"])
        for column in selected_columns
        for cell in inventory_by_column[column].get("nonblank_cells", [])
        if isinstance(cell, dict) and isinstance(cell.get("row"), int)
    ]
    if not available_rows or row_start > max(available_rows):
        raise ValueError("Model-context expansion row_start is outside the table")
    row_end = min(row_start + row_limit - 1, max(available_rows))
    cells_by_column: dict[str, dict[int, dict[str, Any]]] = {}
    for column in selected_columns:
        raw_cells = inventory_by_column[column].get("nonblank_cells", [])
        if not isinstance(raw_cells, list):
            raise ValueError("Inspection control contains malformed cell inventory")
        cells_by_column[column] = {
            int(cell["row"]): copy.deepcopy(cell)
            for cell in raw_cells
            if isinstance(cell, dict)
            and isinstance(cell.get("row"), int)
            and row_start <= int(cell["row"]) <= row_end
        }
    packet: dict[str, Any] = {
        "schema_version": "report-builder.model-context.v1",
        "workflow": "report-builder",
        "packet_role": "targeted_expansion",
        "purpose": selected_purpose,
        "table": {
            "table_id": selected_table_id,
            "kind": clean_text(table.get("kind")),
            "source_file": clean_text(table.get("source_file")),
            "source_artifact_ref": clean_text(table.get("source_artifact_ref")),
            "sheet_name": clean_text(table.get("sheet_name")),
        },
        "selection": {
            "header_row": header_row,
            "columns": selected_columns,
            "row_start": row_start,
            "row_end": row_end,
            "requested_row_limit": row_limit,
            "disclosed_row_count": row_end - row_start + 1,
            "max_rows_per_packet": MODEL_CONTEXT_MAX_EXPANSION_ROWS,
            "max_columns_per_packet": MODEL_CONTEXT_MAX_EXPANSION_COLUMNS,
        },
        "rows": [
            {
                "source_row": row,
                "cells": {
                    column: cells_by_column[column].get(row)
                    for column in selected_columns
                },
            }
            for row in range(row_start, row_end + 1)
        ],
    }
    packet["context_receipt"] = {
        "schema_version": "vera.model_context_receipt.v1",
        "content_sha256": _packet_sha256(packet),
        "table_id": selected_table_id,
        "header_row": header_row,
        "columns": selected_columns,
        "row_start": row_start,
        "row_end": row_end,
        "purpose": selected_purpose,
    }
    return packet


def section_title(section_key: str, language: str) -> str:
    """Return a locale-aware section title."""

    labels = SECTION_TITLES.get(section_key, {})
    return (
        labels.get(language)
        or labels.get("en")
        or section_key.replace("_", " ").title()
    )


def report_type_label(report_type: str, language: str) -> str:
    """Return a locale-aware report type label."""

    labels = REPORT_TYPES[report_type]["label"]
    return labels.get(language) or labels["en"]


def build_suggested_recipe(
    inspection: dict[str, Any],
    *,
    language: str,
    document_language: str,
    report_type: str,
) -> dict[str, Any]:
    """Build an editable recipe from inspection records."""

    assigned_tables: set[str] = set()
    sections: dict[str, Any] = {}
    table_records = inspection["tables"]
    for section_key in REPORT_TYPES[report_type]["sections"]:
        candidates = [
            table
            for table in table_records
            if table.get("suggested_section") == section_key
            and table.get("table_id") not in assigned_tables
            and table.get("kind") != "error"
            and not clean_text(table.get("error"))
            and int(table.get("row_count") or 0) > 0
        ]
        candidates.sort(
            key=lambda item: item.get("suggestion_confidence", 0), reverse=True
        )
        selected = candidates[0] if candidates else None
        assigned_table = selected["table_id"] if selected else ""
        if assigned_table:
            assigned_tables.add(assigned_table)
        sections[section_key] = {
            "title": section_title(section_key, language),
            "assigned_table": assigned_table,
            "codex_comment": "",
            "include_preview_rows": 8,
            "numeric_measure_columns": [],
            "excluded_numeric_candidate_columns": [],
            "numeric_measure_decision": None,
        }

    return {
        "version": 1,
        "language": language,
        "document_language": document_language,
        "report_type": report_type,
        "entity": "",
        "period": "",
        "executive_summary": "",
        "context_items": {},
        "sections": sections,
        "render": {
            "include_unassigned_tables": False,
            "include_table_previews": True,
        },
    }


def review_numeric_measure_columns(
    inspection: dict[str, Any],
    recipe: dict[str, Any],
    *,
    section_key: str,
    header_row: int | None,
    columns: Sequence[str],
    excluded_columns: Sequence[str],
    cell_dispositions: Mapping[str, Mapping[int, str]],
    reviewer_ref: str,
    reviewed_on: str,
    numeric_locale: str,
    currency: str | None,
    unit: str,
    scale: str,
    parse_policy: str,
    sign_policy: str,
) -> dict[str, Any]:
    """Bind every numeric candidate and cell to an exact reviewed disposition."""

    sections = selected_sections(recipe)
    section = sections.get(section_key)
    if not isinstance(section, dict):
        raise ValueError(f"Unknown report section: {section_key}")
    table_id = clean_text(section.get("assigned_table"))
    table = next(
        (
            item
            for item in inspection.get("tables", [])
            if isinstance(item, dict) and clean_text(item.get("table_id")) == table_id
        ),
        None,
    )
    if table is None:
        raise ValueError(f"Assigned table is not present in inspection: {table_id}")
    source_artifact_ref = clean_text(table.get("source_artifact_ref"))
    if not source_artifact_ref:
        raise ValueError("Assigned table has no stable source-artifact identity")
    detected_header_row = table.get("header_row")
    if header_row is not None and (
        not isinstance(header_row, int)
        or isinstance(header_row, bool)
        or header_row < 1
    ):
        raise ValueError("Reviewed header row must be a positive integer or none")
    if header_row is None:
        inventory = table.get("headerless_numeric_measure_cells")
        candidate_records = table.get("headerless_numeric_columns")
    elif header_row == detected_header_row:
        inventory = table.get("numeric_measure_cells")
        candidate_records = table.get("numeric_columns")
    else:
        raise ValueError(
            "Reviewed header row must be the detected candidate or none; "
            "prepare a bounded adapter for another header layout"
        )
    if not isinstance(inventory, list) or not isinstance(candidate_records, list):
        raise ValueError("Inspection lacks full numeric cell inventory")
    available_names = [
        clean_text(item.get("column"))
        for item in inventory
        if isinstance(item, dict) and clean_text(item.get("column"))
    ]
    numeric_candidates = [
        clean_text(item.get("column"))
        for item in candidate_records or []
        if isinstance(item, dict) and clean_text(item.get("column"))
    ]
    requested = [clean_text(column) for column in columns if clean_text(column)]
    excluded = [clean_text(column) for column in excluded_columns if clean_text(column)]
    if len(requested) != len(set(requested)) or len(excluded) != len(set(excluded)):
        raise ValueError("Reviewed numeric candidate columns must be unique")
    if set(requested) & set(excluded):
        raise ValueError(
            "Numeric candidate columns cannot be both included and excluded"
        )
    unknown = (set(requested) | set(excluded)) - set(numeric_candidates)
    if unknown:
        raise ValueError(f"Unknown numeric candidate columns: {sorted(unknown)}")
    if set(requested) | set(excluded) != set(numeric_candidates):
        undisposed = set(numeric_candidates) - set(requested) - set(excluded)
        raise ValueError(
            "Every numeric candidate column requires an include/exclude disposition: "
            f"{sorted(undisposed)}"
        )
    if set(cell_dispositions) != set(requested):
        raise ValueError(
            "Every included numeric column requires explicit dispositions for all cells"
        )
    normalized_columns = [name for name in available_names if name in requested]
    normalized_excluded = [name for name in available_names if name in excluded]
    contract = _numeric_contract(
        numeric_locale=numeric_locale,
        currency=currency,
        unit=unit,
        scale=scale,
        parse_policy=parse_policy,
        sign_policy=sign_policy,
    )
    _, included_dispositions = _strict_reviewed_numeric_profile(
        inventory,
        normalized_columns,
        contract,
        cell_dispositions,
    )
    excluded_dispositions = _excluded_numeric_column_dispositions(
        inventory,
        normalized_excluded,
    )
    dispositions_by_name = {
        clean_text(item.get("column")): item
        for item in [*included_dispositions, *excluded_dispositions]
    }
    dispositions = [
        dispositions_by_name[name]
        for name in numeric_candidates
        if name in dispositions_by_name
    ]
    try:
        reviewed_date = date.fromisoformat(reviewed_on)
    except ValueError as exc:
        raise ValueError("reviewed_on must be an ISO date") from exc
    if reviewed_date > date.today():
        raise ValueError("reviewed_on cannot be in the future")
    content = {
        "section": section_key,
        "report_period": clean_text(recipe.get("period")),
        "header_row": header_row,
        "source_artifact_ref": source_artifact_ref,
        "table_id": table_id,
        "source_file": clean_text(table.get("source_file")),
        "sheet_name": clean_text(table.get("sheet_name")),
        "numeric_measure_columns": normalized_columns,
        "excluded_numeric_candidate_columns": normalized_excluded,
        "numeric_contract": contract,
        "column_dispositions": dispositions,
    }
    decision = build_reviewed_decision_receipt(
        decision_id=_numeric_decision_id(section_key, source_artifact_ref),
        decision_type=NUMERIC_MEASURE_DECISION_TYPE,
        status="reviewed",
        reviewer_ref=reviewer_ref,
        reviewed_on=reviewed_on,
        adapter_id=NUMERIC_MEASURE_ADAPTER_ID,
        adapter_version=NUMERIC_MEASURE_ADAPTER_VERSION,
        source_artifact_refs=[source_artifact_ref],
        content=content,
    )
    updated = copy.deepcopy(recipe)
    updated_section = updated["sections"][section_key]
    updated_section["numeric_measure_columns"] = normalized_columns
    updated_section["excluded_numeric_candidate_columns"] = normalized_excluded
    updated_section["numeric_measure_decision"] = decision
    return updated


def inspect_inputs(
    input_path: Path,
    output_dir: Path,
    *,
    language: object | None = None,
    document_language: object | None = None,
    report_type: object | None = None,
) -> InspectionResult:
    """Inspect all report inputs locally and write bounded model/control packets."""

    output_dir.mkdir(parents=True, exist_ok=True)
    assumptions = language_assumptions(
        language=language,
        document_language=document_language,
    )
    report_key = normalize_report_type(report_type)
    raw_tables = load_tables(input_path, output_dir)
    tables = [inspect_table(table) for table in raw_tables]
    inspection = {
        "version": 1,
        "language": assumptions["language"],
        "document_language": assumptions["document_language"],
        "report_type": report_key,
        "report_type_label": report_type_label(report_key, assumptions["language"]),
        "input_path": input_path.name,
        "table_count": len(tables),
        "tables": tables,
        "supported_suffixes": sorted(SUPPORTED_SUFFIXES),
        "limitations": [
            "Scanned PDFs require OCR before deterministic extraction.",
            "Excel binary .xls files should be converted to .xlsx or CSV for this plugin version.",
            (
                "Workbook formulas and cached values are inventoried separately; "
                "formula cells cannot be included in measures without a verified "
                "recalculation/export adapter."
            ),
        ],
    }
    recipe = build_suggested_recipe(
        inspection,
        language=assumptions["language"],
        document_language=assumptions["document_language"],
        report_type=report_key,
    )
    model_inspection = _model_inspection_packet(inspection)
    control_path = output_dir / "inspection_control.json"
    model_path = output_dir / "inspection.json"
    write_json(control_path, inspection)
    write_json(model_path, model_inspection)
    write_json(
        output_dir / "model_context_receipt.json",
        {
            "schema_version": "vera.model_context_receipt.v1",
            "workflow": "report-builder",
            "full_population_processed_locally": True,
            "default_model_packet": artifact_receipt(
                output_dir,
                model_path,
                artifact_id="model_context.default_inspection",
                role="model_context",
                media_type="application/json",
            ),
            "private_control": artifact_receipt(
                output_dir,
                control_path,
                artifact_id="control.full_inspection",
                role="private_control",
                media_type="application/json",
            ),
            "bounds": {
                "preview_rows_per_table": MODEL_CONTEXT_PREVIEW_ROWS,
                "max_expansion_rows": MODEL_CONTEXT_MAX_EXPANSION_ROWS,
                "max_expansion_columns": MODEL_CONTEXT_MAX_EXPANSION_COLUMNS,
            },
        },
    )
    write_json(output_dir / "suggested_recipe.json", recipe)
    return InspectionResult(inspection=inspection, suggested_recipe=recipe)


def selected_sections(recipe: dict[str, Any]) -> dict[str, Any]:
    """Return recipe sections as a mapping."""

    sections = recipe.get("sections", {})
    if not isinstance(sections, dict):
        raise ValueError("Recipe sections must be a JSON object")
    return sections


def validate_narrative_numeric_boundary(recipe: dict[str, Any]) -> None:
    """Reject free-form numerals that have no numeric-evidence ledger address."""

    entity = clean_text(recipe.get("entity"))
    if re.search(r"\d", entity):
        raise ValueError(
            "Entity contains digits that cannot be distinguished from an "
            "unledgered numeric claim; use a digit-free reviewed display name."
        )
    period = clean_text(recipe.get("period"))
    canonical_period = re.fullmatch(
        r"(?:"
        r"(?:FY\s*)?\d{4}"
        r"|Q[1-4]\s+\d{4}"
        r"|\d{4}-\d{2}-\d{2}"
        r"|(?:Year|Period|Quarter)\s+ended\s+\d{4}-\d{2}-\d{2}"
        r"|\d{4}-\d{2}-\d{2}\s+(?:to|through)\s+\d{4}-\d{2}-\d{2}"
        r")",
        period,
        flags=re.IGNORECASE,
    )
    if re.search(r"\d", period) and canonical_period is None:
        raise ValueError(
            "Reporting period contains numeric text outside the supported "
            "canonical period formats."
        )
    narrative_fields: list[tuple[str, str]] = [
        ("executive_summary", clean_text(recipe.get("executive_summary"))),
    ]
    context_items = recipe.get("context_items")
    if isinstance(context_items, dict):
        for key, value in context_items.items():
            key_text = clean_text(key)
            narrative_fields.extend(
                [
                    (f"context_items.{key_text}.key", key_text),
                    (f"context_items.{key_text}.value", clean_text(value)),
                ]
            )
    for section_key, section in selected_sections(recipe).items():
        if not isinstance(section, dict):
            continue
        narrative_fields.extend(
            [
                (
                    f"sections.{section_key}.title",
                    clean_text(section.get("title")),
                ),
                (
                    f"sections.{section_key}.codex_comment",
                    clean_text(section.get("codex_comment")),
                ),
            ]
        )
    offenders = [field for field, text in narrative_fields if re.search(r"\d", text)]
    if offenders:
        raise ValueError(
            "Free-form report narrative contains unledgered numeric tokens: "
            f"{offenders}. Remove the numerals and rely on reviewed numeric "
            "measures until a claim-basis reference contract is available."
        )


def _reviewed_numeric_profile(
    section_key: str,
    section_recipe: dict[str, Any],
    table: dict[str, Any],
    *,
    report_period: str,
    header_row: int | None,
    candidates: Sequence[dict[str, Any]],
) -> tuple[list[dict[str, Any]], str, dict[str, Any] | None, str]:
    """Return only source-bound measure columns approved in the recipe."""

    if table.get("kind") not in {"worksheet", "csv"}:
        return [], "not_applicable", None, ""
    raw_columns = section_recipe.get("numeric_measure_columns")
    raw_excluded_columns = section_recipe.get("excluded_numeric_candidate_columns")
    decision = section_recipe.get("numeric_measure_decision")
    if not isinstance(decision, dict):
        if not candidates:
            return [], "not_applicable", None, ""
        return (
            [],
            "needs_review",
            None,
            "Numeric-looking columns are candidates until their measure role is reviewed.",
        )
    if not isinstance(raw_columns, list) or not isinstance(raw_excluded_columns, list):
        return [], "needs_review", None, "Numeric candidate dispositions are malformed."
    columns = [clean_text(item) for item in raw_columns]
    excluded_columns = [clean_text(item) for item in raw_excluded_columns]
    if (
        any(not item for item in [*columns, *excluded_columns])
        or len(columns) != len(set(columns))
        or len(excluded_columns) != len(set(excluded_columns))
        or set(columns) & set(excluded_columns)
    ):
        return [], "needs_review", None, "Reviewed measure columns are malformed."
    source_artifact_ref = clean_text(table.get("source_artifact_ref"))
    try:
        validated = validate_reviewed_decision_receipt(
            decision,
            expected_source_artifact_refs=[source_artifact_ref],
            expected_adapter_id=NUMERIC_MEASURE_ADAPTER_ID,
            expected_adapter_version=NUMERIC_MEASURE_ADAPTER_VERSION,
            require_reviewed=True,
        )
        if validated["decision_type"] != NUMERIC_MEASURE_DECISION_TYPE:
            raise ValueError("Numeric-measure decision type is stale")
        if validated["decision_id"] != _numeric_decision_id(
            section_key, source_artifact_ref
        ):
            raise ValueError("Numeric-measure decision identity is stale")
        if date.fromisoformat(validated["reviewed_on"]) > date.today():
            raise ValueError("Numeric-measure review date is in the future")
        content = validated["content"]
        reviewed_header_row = content.get("header_row")
        if reviewed_header_row is not None and (
            not isinstance(reviewed_header_row, int)
            or isinstance(reviewed_header_row, bool)
            or reviewed_header_row < 1
            or reviewed_header_row > len(table.get("rows", []))
        ):
            raise ValueError("Reviewed header row is invalid or stale")
        reviewed_header_index = (
            reviewed_header_row - 1 if isinstance(reviewed_header_row, int) else None
        )
        raw_contract = content.get("numeric_contract")
        if not isinstance(raw_contract, dict):
            raise ValueError("Numeric-measure contract is missing")
        contract = _numeric_contract(
            numeric_locale=raw_contract.get("locale"),
            currency=raw_contract.get("currency"),
            unit=raw_contract.get("unit"),
            scale=raw_contract.get("scale"),
            parse_policy=raw_contract.get("policy_id"),
            sign_policy=raw_contract.get("sign_policy"),
        )
        if raw_contract != contract:
            raise ValueError("Numeric-measure parser contract is not canonical")
        inventory = _numeric_cell_inventory(
            table.get("rows", []),
            reviewed_header_index,
            cell_formats=table.get("cell_formats"),
            formula_cells=table.get("formula_cells"),
        )
        available_names = [
            clean_text(item.get("column"))
            for item in inventory
            if isinstance(item, dict) and clean_text(item.get("column"))
        ]
        unknown = set(columns) - set(available_names)
        if unknown:
            raise ValueError(f"Reviewed measure columns are stale: {sorted(unknown)}")
        current_profile = table_numeric_profile(
            table.get("rows", []),
            reviewed_header_index,
            formula_cells=table.get("formula_cells"),
        )
        current_candidates = _numeric_candidate_columns(
            current_profile["numeric_columns"]
        )
        candidate_names = [
            clean_text(candidate.get("column"))
            for candidate in current_candidates
            if isinstance(candidate, dict) and clean_text(candidate.get("column"))
        ]
        if set(columns) | set(excluded_columns) != set(candidate_names):
            raise ValueError("Numeric candidate column dispositions are incomplete")
        normalized_columns = [name for name in available_names if name in columns]
        normalized_excluded = [
            name for name in available_names if name in excluded_columns
        ]
        recorded_dispositions = content.get("column_dispositions")
        if not isinstance(recorded_dispositions, list):
            raise ValueError("Numeric cell dispositions are missing")
        recorded_by_name = {
            clean_text(item.get("column")): item
            for item in recorded_dispositions
            if isinstance(item, dict) and clean_text(item.get("column"))
        }
        cell_reviews: dict[str, dict[int, str]] = {}
        for name in normalized_columns:
            disposition = recorded_by_name.get(name)
            cells = disposition.get("cells") if isinstance(disposition, dict) else None
            if not isinstance(cells, list):
                raise ValueError(f"Numeric cell dispositions are missing: {name}")
            cell_reviews[name] = {}
            for cell in cells:
                if not isinstance(cell, dict):
                    raise ValueError(f"Malformed numeric cell disposition: {name}")
                row_number = cell.get("row")
                status = clean_text(cell.get("status"))
                if not isinstance(row_number, int) or isinstance(row_number, bool):
                    raise ValueError(f"Malformed numeric cell row: {name}")
                if status == "included":
                    cell_reviews[name][row_number] = "include"
                elif (
                    status == "excluded_by_review"
                    and cell.get("reason") == "reviewed_cell_exclusion"
                ):
                    cell_reviews[name][row_number] = "exclude"
                else:
                    raise ValueError(f"Invalid numeric cell disposition: {name}")
        selected, dispositions = _strict_reviewed_numeric_profile(
            inventory,
            normalized_columns,
            contract,
            cell_reviews,
        )
        excluded_dispositions = _excluded_numeric_column_dispositions(
            inventory,
            normalized_excluded,
        )
        dispositions_by_name = {
            clean_text(item.get("column")): item
            for item in [*dispositions, *excluded_dispositions]
        }
        all_dispositions = [
            dispositions_by_name[name]
            for name in candidate_names
            if name in dispositions_by_name
        ]
        expected_content = {
            "section": section_key,
            "report_period": report_period,
            "source_artifact_ref": source_artifact_ref,
            "table_id": clean_text(table.get("table_id")),
            "source_file": clean_text(table.get("source_file")),
            "sheet_name": clean_text(table.get("sheet_name")),
            "header_row": reviewed_header_row,
            "numeric_measure_columns": normalized_columns,
            "excluded_numeric_candidate_columns": normalized_excluded,
            "numeric_contract": contract,
            "column_dispositions": all_dispositions,
        }
        if content != expected_content:
            raise ValueError(
                "Numeric-measure review content does not match the current table"
            )
    except (ValueError, TypeError):
        return (
            [],
            "needs_review",
            None,
            "Numeric-measure review receipt is missing, invalid, or stale.",
        )
    return selected, "reviewed", validated, ""


def analysis_for_section(
    section_key: str,
    section_recipe: dict[str, Any],
    table_by_id: dict[str, dict[str, Any]],
    *,
    report_period: str,
) -> dict[str, Any]:
    """Build deterministic analysis for one report section."""

    table_id = clean_text(section_recipe.get("assigned_table"))
    table = table_by_id.get(table_id) if table_id else None
    if not table:
        return {
            "section": section_key,
            "title": clean_text(section_recipe.get("title")) or section_key,
            "assigned_table": table_id,
            "status": "unassigned",
            "row_count": 0,
            "column_count": 0,
            "numeric_columns": [],
            "numeric_measure_candidates": [],
            "numeric_measure_status": "not_applicable",
            "numeric_measure_decision": None,
            "numeric_measure_limitation": "",
            "preview_rows": [],
        }
    rows = table.get("rows", [])
    if (
        table.get("kind") == "error"
        or clean_text(table.get("error"))
        or not any(row_nonempty_count(row) > 0 for row in rows)
    ):
        return {
            "section": section_key,
            "title": clean_text(section_recipe.get("title")) or section_key,
            "assigned_table": table_id,
            "source_file": table.get("source_file", ""),
            "source_path": table.get("source_file", ""),
            "source_artifact_ref": table.get("source_artifact_ref", ""),
            "sheet_name": table.get("sheet_name", ""),
            "status": "unsupported_source_layout",
            "row_count": 0,
            "column_count": 0,
            "numeric_columns": [],
            "numeric_measure_candidates": [],
            "numeric_measure_status": "not_applicable",
            "numeric_measure_decision": None,
            "numeric_measure_limitation": (
                clean_text(table.get("error"))
                or "Source table contains no reportable rows."
            ),
            "preview_rows": [],
            "codex_comment": clean_text(section_recipe.get("codex_comment")),
        }
    header_idx = header_candidate_index(rows)
    include_rows = int(section_recipe.get("include_preview_rows") or 8)
    profile = (
        table_numeric_profile(
            rows,
            header_idx,
            formula_cells=table.get("formula_cells"),
        )
        if table.get("kind") in {"worksheet", "csv"}
        else {"numeric_cells": 0, "numeric_columns": []}
    )
    numeric_candidates = _numeric_candidate_columns(profile["numeric_columns"])
    headerless_profile = (
        table_numeric_profile(
            rows,
            None,
            formula_cells=table.get("formula_cells"),
        )
        if table.get("kind") in {"worksheet", "csv"}
        else {"numeric_cells": 0, "numeric_columns": []}
    )
    headerless_candidates = _numeric_candidate_columns(
        headerless_profile["numeric_columns"]
    )
    pending_candidates = numeric_candidates or headerless_candidates
    header_row = header_idx + 1 if header_idx is not None else None
    (
        numeric_columns,
        numeric_measure_status,
        numeric_measure_decision,
        numeric_measure_limitation,
    ) = _reviewed_numeric_profile(
        section_key,
        section_recipe,
        table,
        report_period=report_period,
        header_row=header_row,
        candidates=pending_candidates,
    )
    effective_header_row = header_row
    if isinstance(numeric_measure_decision, dict):
        decision_content = numeric_measure_decision.get("content")
        if isinstance(decision_content, dict):
            effective_header_row = decision_content.get("header_row")
    return {
        "section": section_key,
        "title": clean_text(section_recipe.get("title")) or section_key,
        "assigned_table": table_id,
        "source_file": table.get("source_file", ""),
        "source_path": table.get("source_file", ""),
        "source_artifact_ref": table.get("source_artifact_ref", ""),
        "sheet_name": table.get("sheet_name", ""),
        "status": "assigned",
        "row_count": len([row for row in rows if row_nonempty_count(row) > 0]),
        "column_count": max((len(row) for row in rows), default=0),
        "header_row": effective_header_row,
        "numeric_columns": numeric_columns,
        "numeric_measure_candidates": pending_candidates,
        "numeric_header_review_options": {
            "detected_header_row": header_row,
            "detected_header_candidates": numeric_candidates,
            "headerless_candidates": headerless_candidates,
        },
        "numeric_measure_status": numeric_measure_status,
        "numeric_measure_decision": numeric_measure_decision,
        "numeric_measure_limitation": numeric_measure_limitation,
        "preview_rows": _redact_numeric_preview_rows(
            rows_as_dicts(
                rows,
                (
                    effective_header_row - 1
                    if isinstance(effective_header_row, int)
                    else None
                ),
            )[: max(0, include_rows)]
        ),
        "codex_comment": clean_text(section_recipe.get("codex_comment")),
    }


def markdown_table(rows: Sequence[dict[str, str]], *, max_rows: int = 8) -> str:
    """Render a compact Markdown table."""

    if not rows:
        return ""
    headers = list(rows[0].keys())[:8]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows[:max_rows]:
        values = [
            clean_text(row.get(header, "")).replace("|", "/") for header in headers
        ]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def render_markdown(recipe: dict[str, Any], analysis: dict[str, Any]) -> str:
    """Render a report draft in Markdown."""

    language = normalize_language(recipe.get("language"), default="en")
    report_type = normalize_report_type(recipe.get("report_type"))
    title = report_type_label(report_type, language)
    entity = clean_text(recipe.get("entity")) or docx_label("entity_pending", language)
    period = clean_text(recipe.get("period")) or docx_label("period_pending", language)
    executive_summary = clean_text(recipe.get("executive_summary"))
    if not executive_summary:
        executive_summary = docx_label("executive_summary_pending", language)

    lines = [
        f"# {title}",
        "",
        f"**{docx_label('entity', language)}:** {entity}",
        f"**{docx_label('period', language)}:** {period}",
        "",
        f"## {docx_label('executive_summary', language)}",
        "",
        executive_summary,
        "",
    ]

    context_items = recipe.get("context_items", {})
    if isinstance(context_items, dict) and context_items:
        lines.extend([f"## {docx_label('context', language)}", ""])
        for key, value in context_items.items():
            lines.append(f"- **{clean_text(key)}:** {clean_text(value)}")
        lines.append("")

    for section in analysis["sections"]:
        lines.extend([f"## {section['title']}", ""])
        if section["status"] != "assigned":
            lines.extend([docx_label("unassigned", language), ""])
            continue
        comment = section.get("codex_comment") or docx_label("codex_pending", language)
        lines.extend(
            [
                comment,
                "",
                f"{docx_label('source', language)}: {section.get('source_file', '')}"
                + (f" / {section['sheet_name']}" if section.get("sheet_name") else ""),
                f"{docx_label('rows', language)}: {section['row_count']} | "
                f"{docx_label('columns', language)}: {section['column_count']}",
            ]
        )
        numeric_columns = section.get("numeric_columns") or []
        if numeric_columns:
            lines.extend(["", f"{docx_label('numeric_totals', language)}:"])
            for column in numeric_columns:
                currency = clean_text(column.get("currency")) or "none"
                lines.append(
                    f"- {column['column']}: {docx_label('sum', language)} "
                    f"{column['sum']} | {docx_label('currency', language)}: "
                    f"{currency} | {docx_label('unit', language)}: "
                    f"{clean_text(column.get('unit'))} | "
                    f"{docx_label('scale', language)}: "
                    f"{clean_text(column.get('scale'))}"
                )
        preview_table = markdown_table(section.get("preview_rows") or [])
        if preview_table and recipe.get("render", {}).get(
            "include_table_previews", True
        ):
            lines.extend(["", preview_table])
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def docx_label(key: str, language: str) -> str:
    """Return a localized DOCX label."""

    labels = DOCX_COPY.get(key, {})
    return labels.get(language) or labels.get("en") or key.replace("_", " ").title()


def set_cell_shading(cell: Any, fill: str) -> None:
    """Set deterministic cell background shading."""

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, margin_twips: int = 80) -> None:
    """Set compact but readable cell margins."""

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_paragraph_font(
    paragraph: Any,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    """Apply simple font formatting to all runs in a paragraph."""

    for run in paragraph.runs:
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic


def set_docx_styles(document: Any) -> None:
    """Configure restrained report styles."""

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("334155")

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("243026")

    for style_name, size in (("Heading 1", 18), ("Heading 2", 13)):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("243026")


def add_small_note(document: Any, text: str) -> None:
    """Add a muted note paragraph."""

    paragraph = document.add_paragraph(text)
    set_paragraph_font(paragraph, size=9, color="667085")


def add_key_value_table(document: Any, rows: Sequence[tuple[str, Any]]) -> None:
    """Add a compact two-column metadata table."""

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = clean_text(label)
        cells[1].text = clean_text(value)
        set_cell_shading(cells[0], "E7ECE5")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, size=9.5)
        for paragraph in cells[0].paragraphs:
            set_paragraph_font(paragraph, bold=True, color="243026")


def add_dataframe_table(
    document: Any,
    rows: Sequence[dict[str, str]],
    *,
    max_rows: int = 8,
    max_columns: int = 6,
) -> None:
    """Add a real Word table for preview rows."""

    if not rows:
        return
    headers = list(rows[0].keys())[:max_columns]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = clean_text(header)
        set_cell_shading(header_cells[idx], "E7ECE5")
        set_cell_margins(header_cells[idx])
        for paragraph in header_cells[idx].paragraphs:
            set_paragraph_font(paragraph, size=8.5, bold=True, color="243026")

    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            value = clean_text(row.get(header, ""))
            cells[idx].text = value[:140] + ("..." if len(value) > 140 else "")
            set_cell_margins(cells[idx])
            for paragraph in cells[idx].paragraphs:
                set_paragraph_font(paragraph, size=8)


def add_numeric_totals_table(
    document: Any,
    numeric_columns: Sequence[dict[str, Any]],
    *,
    language: str = "en",
) -> None:
    """Add every reviewed, ledger-backed numeric total."""

    if not numeric_columns:
        return
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    headers = (
        docx_label("column", language),
        docx_label("sum", language).capitalize(),
        docx_label("currency", language),
        docx_label("unit", language),
        docx_label("scale", language),
    )
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        set_cell_shading(table.rows[0].cells[idx], "E7ECE5")
        set_cell_margins(table.rows[0].cells[idx])
        for paragraph in table.rows[0].cells[idx].paragraphs:
            set_paragraph_font(paragraph, size=8.5, bold=True, color="243026")
    for column in numeric_columns:
        cells = table.add_row().cells
        cells[0].text = clean_text(column.get("column"))
        cells[1].text = clean_text(column.get("sum"))
        cells[2].text = clean_text(column.get("currency")) or "none"
        cells[3].text = clean_text(column.get("unit"))
        cells[4].text = clean_text(column.get("scale"))
        for cell in cells:
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, size=8)


def write_report_docx(
    recipe: dict[str, Any],
    analysis: dict[str, Any],
    audit: dict[str, Any],
    output_path: Path,
) -> None:
    """Write a styled DOCX report with real Word sections and tables."""

    language = normalize_language(recipe.get("language"), default="en")
    report_type = normalize_report_type(recipe.get("report_type"))
    title = report_type_label(report_type, language)
    entity = clean_text(recipe.get("entity")) or docx_label("entity_pending", language)
    period = clean_text(recipe.get("period")) or docx_label("period_pending", language)
    executive_summary = clean_text(recipe.get("executive_summary")) or docx_label(
        "codex_pending", language
    )

    document = Document()
    set_docx_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    document.core_properties.title = title
    document.core_properties.subject = f"{entity} - {period}"
    document.core_properties.comments = docx_label("generated_by", language)

    title_paragraph = document.add_paragraph(style="Title")
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = document.add_paragraph(docx_label("draft", language))
    set_paragraph_font(subtitle, size=10, color="667085")

    add_key_value_table(
        document,
        (
            (docx_label("entity", language), entity),
            (docx_label("period", language), period),
            (docx_label("report_status", language), audit.get("status", "draft")),
            (docx_label("model_api_calls", language), audit.get("model_api_calls", 0)),
        ),
    )

    document.add_heading(docx_label("executive_summary", language), level=1)
    document.add_paragraph(executive_summary)

    context_items = recipe.get("context_items", {})
    if isinstance(context_items, dict) and context_items:
        document.add_heading(docx_label("context", language), level=1)
        add_key_value_table(
            document,
            tuple(
                (clean_text(key), clean_text(value))
                for key, value in context_items.items()
            ),
        )

    for section_analysis in analysis.get("sections", []):
        document.add_heading(clean_text(section_analysis.get("title")), level=1)
        if section_analysis.get("status") != "assigned":
            paragraph = document.add_paragraph(docx_label("unassigned", language))
            set_paragraph_font(paragraph, italic=True, color="667085")
            continue

        comment = clean_text(section_analysis.get("codex_comment")) or docx_label(
            "codex_pending", language
        )
        document.add_paragraph(comment)
        source = clean_text(section_analysis.get("source_file"))
        sheet = clean_text(section_analysis.get("sheet_name"))
        source_value = source + (f" / {sheet}" if sheet else "")
        add_key_value_table(
            document,
            (
                (docx_label("source", language), source_value),
                (docx_label("rows", language), section_analysis.get("row_count", 0)),
                (
                    docx_label("columns", language),
                    section_analysis.get("column_count", 0),
                ),
            ),
        )

        numeric_columns = section_analysis.get("numeric_columns") or []
        if numeric_columns:
            document.add_heading(docx_label("numeric_totals", language), level=2)
            add_numeric_totals_table(document, numeric_columns, language=language)

        preview_rows = section_analysis.get("preview_rows") or []
        if preview_rows and recipe.get("render", {}).get(
            "include_table_previews", True
        ):
            document.add_heading(docx_label("table_preview", language), level=2)
            add_dataframe_table(document, preview_rows)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(docx_label("audit_appendix", language), level=1)
    add_key_value_table(
        document,
        (
            (docx_label("input_path", language), audit.get("input_path", "")),
            (docx_label("tables_discovered", language), audit.get("table_count", 0)),
            (
                docx_label("assigned_sections", language),
                audit.get("assigned_section_count", 0),
            ),
            (
                docx_label("missing_sections", language),
                audit.get("missing_section_count", 0),
            ),
            (docx_label("model_api_calls", language), audit.get("model_api_calls", 0)),
        ),
    )
    missing_sections = audit.get("missing_sections") or []
    if missing_sections:
        document.add_heading(docx_label("missing_sections", language), level=2)
        for missing in missing_sections:
            document.add_paragraph(clean_text(missing), style="List Bullet")
    notes = audit.get("notes") or []
    if notes:
        document.add_heading(docx_label("notes", language), level=2)
        for note in notes:
            add_small_note(document, clean_text(note))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    _stabilize_office_package(output_path)


def _excel_column_name(index: int) -> str:
    """Return a one-based Excel column name."""

    name = ""
    while index > 0:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


def _numeric_evidence_rows(analysis: dict[str, Any]) -> list[dict[str, Any]]:
    """Return the numeric summaries that appear in every rendered report."""

    evidence: list[dict[str, Any]] = []
    for section_index, section in enumerate(analysis.get("sections", [])):
        if not isinstance(section, dict) or section.get("status") != "assigned":
            continue
        source_artifact_ref = clean_text(section.get("source_artifact_ref"))
        numeric_decision = section.get("numeric_measure_decision")
        decision_ref = (
            clean_text(numeric_decision.get("decision_id"))
            if isinstance(numeric_decision, dict)
            else ""
        )
        numeric_contract = (
            numeric_decision.get("content", {}).get("numeric_contract")
            if isinstance(numeric_decision, dict)
            and isinstance(numeric_decision.get("content"), dict)
            else None
        )
        source_name = clean_text(section.get("sheet_name")) or clean_text(
            section.get("assigned_table")
        )
        for column_offset, column in enumerate(section.get("numeric_columns") or []):
            if not isinstance(column, dict):
                continue
            value = clean_text(column.get("sum"))
            if not value:
                continue
            column_index = int(column.get("column_index") or column_offset + 1)
            physical_rows = [
                int(row_number)
                for row_number in column.get("numeric_rows", [])
                if isinstance(row_number, int) and not isinstance(row_number, bool)
            ]
            evidence_id = (
                f"numeric.section_{section_index + 1:03d}."
                f"column_{column_offset + 1:03d}.sum"
            )
            cell_column = _excel_column_name(column_index)
            source_cells = ",".join(
                f"{cell_column}{row_number}" for row_number in physical_rows
            )
            limitations = []
            if not source_artifact_ref:
                limitations.append("source_receipt_unresolved")
            if not decision_ref:
                limitations.append("numeric_measure_decision_unresolved")
            evidence.append(
                {
                    "evidence_id": evidence_id,
                    "section_index": section_index,
                    "section": clean_text(section.get("section")),
                    "column_offset": column_offset,
                    "column": clean_text(column.get("column")),
                    "value": value,
                    "currency": column.get("currency"),
                    "unit": clean_text(column.get("unit")),
                    "scale": clean_text(column.get("scale")),
                    "decision_ref": decision_ref,
                    "numeric_contract": numeric_contract,
                    "source_artifact_ref": source_artifact_ref,
                    "source_table_id": clean_text(section.get("assigned_table")),
                    "source_sheet": clean_text(section.get("sheet_name")),
                    "source_header_row": section.get("header_row"),
                    "source_locator": (
                        f"{source_name}!{source_cells}"
                        if source_cells
                        else f"{source_name}!column:{cell_column}"
                    ),
                    "prepared_locator": (
                        f"/sections/{section_index}/numeric_columns/"
                        f"{column_offset}/sum"
                    ),
                    "column_index": column_index,
                    "numeric_rows": physical_rows,
                    "numeric_count": int(column.get("numeric_count") or 0),
                    "limitations": limitations,
                }
            )
    return evidence


def write_tables_workbook(output_path: Path, analysis: dict[str, Any]) -> None:
    """Write assigned table previews to an Excel workbook."""

    def append_safe_row(sheet: Any, values: Sequence[Any]) -> None:
        row_index = (
            1
            if sheet.max_row == 1
            and sheet.max_column == 1
            and sheet.cell(row=1, column=1).value is None
            else sheet.max_row + 1
        )
        for column_index, value in enumerate(values, start=1):
            cell = sheet.cell(row=row_index, column=column_index)
            cell.value = value
            if isinstance(value, str):
                cell.data_type = "s"
                cell.number_format = "@"

    workbook = openpyxl.Workbook()
    default = workbook.active
    default.title = "summary"
    append_safe_row(
        default,
        ["section", "status", "assigned_table", "rows", "columns"],
    )
    for section in analysis["sections"]:
        append_safe_row(
            default,
            [
                section["section"],
                section["status"],
                section.get("assigned_table", ""),
                section.get("row_count", 0),
                section.get("column_count", 0),
            ],
        )
        rows = section.get("preview_rows") or []
        if rows:
            sheet = workbook.create_sheet(
                safe_sheet_name(
                    section["section"], f"section{len(workbook.worksheets)}"
                )
            )
            headers = list(rows[0].keys())
            append_safe_row(sheet, headers)
            for row in rows:
                append_safe_row(
                    sheet,
                    [row.get(header, "") for header in headers],
                )
    evidence_sheet = workbook.create_sheet("numeric_evidence")
    append_safe_row(
        evidence_sheet,
        ["evidence_id", "section", "column", "sum", "currency", "unit", "scale"],
    )
    for evidence in _numeric_evidence_rows(analysis):
        append_safe_row(
            evidence_sheet,
            [
                evidence["evidence_id"],
                evidence["section"],
                evidence["column"],
                evidence["value"],
                evidence["currency"] or "none",
                evidence["unit"],
                evidence["scale"],
            ],
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)
    _stabilize_office_package(output_path)


def _markdown_numeric_locations(
    path: Path,
    evidence_rows: Sequence[dict[str, Any]],
    *,
    language: str,
) -> dict[str, tuple[str, str]]:
    """Reopen Markdown and locate each exact rendered numeric value."""

    lines = path.read_text(encoding="utf-8").splitlines()
    locations: dict[str, tuple[str, str]] = {}
    cursor = 0
    for evidence in evidence_rows:
        expected = (
            f"- {evidence['column']}: {docx_label('sum', language)} "
            f"{evidence['value']} | {docx_label('currency', language)}: "
            f"{evidence['currency'] or 'none'} | "
            f"{docx_label('unit', language)}: {evidence['unit']} | "
            f"{docx_label('scale', language)}: {evidence['scale']}"
        )
        match = next(
            (index for index in range(cursor, len(lines)) if lines[index] == expected),
            None,
        )
        if match is None:
            raise ValueError(
                "Numeric output closure failed for Markdown evidence "
                f"{evidence['evidence_id']}"
            )
        locations[str(evidence["evidence_id"])] = (
            f"line:{match + 1}",
            str(evidence["value"]),
        )
        cursor = match + 1
    return locations


def _docx_numeric_locations(
    path: Path,
    evidence_rows: Sequence[dict[str, Any]],
    *,
    language: str,
) -> dict[str, tuple[str, str]]:
    """Reopen Word output and locate each numeric-summary cell."""

    document = Document(path)
    rendered: list[tuple[int, int, str, str, str, str, str]] = []
    expected_headers = (
        docx_label("column", language),
        docx_label("sum", language).capitalize(),
        docx_label("currency", language),
        docx_label("unit", language),
        docx_label("scale", language),
    )
    for table_index, table in enumerate(document.tables, start=1):
        if not table.rows or len(table.rows[0].cells) != 5:
            continue
        header = tuple(clean_text(cell.text) for cell in table.rows[0].cells)
        if header != expected_headers:
            continue
        for row_index, row in enumerate(table.rows[1:], start=2):
            rendered.append(
                (
                    table_index,
                    row_index,
                    clean_text(row.cells[0].text),
                    clean_text(row.cells[1].text),
                    clean_text(row.cells[2].text),
                    clean_text(row.cells[3].text),
                    clean_text(row.cells[4].text),
                )
            )
    if len(rendered) != len(evidence_rows):
        raise ValueError(
            "Numeric output closure failed because Word numeric rows do not close"
        )
    locations: dict[str, tuple[str, str]] = {}
    for evidence, (
        table_index,
        row_index,
        column,
        value,
        currency,
        unit,
        scale,
    ) in zip(evidence_rows, rendered, strict=True):
        if (
            column != evidence["column"]
            or value != evidence["value"]
            or currency != (evidence["currency"] or "none")
            or unit != evidence["unit"]
            or scale != evidence["scale"]
        ):
            raise ValueError(
                "Numeric output closure failed for Word evidence "
                f"{evidence['evidence_id']}"
            )
        locations[str(evidence["evidence_id"])] = (
            f"table:{table_index}/row:{row_index}/cell:2",
            value,
        )
    return locations


def _workbook_numeric_locations(
    path: Path,
    evidence_rows: Sequence[dict[str, Any]],
) -> dict[str, tuple[str, str]]:
    """Reopen Excel output and locate each exact numeric-summary cell."""

    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    if "numeric_evidence" not in workbook.sheetnames:
        raise ValueError("Numeric output closure failed: Excel evidence sheet missing")
    sheet = workbook["numeric_evidence"]
    headers = [clean_text(cell.value) for cell in sheet[1]]
    if headers != [
        "evidence_id",
        "section",
        "column",
        "sum",
        "currency",
        "unit",
        "scale",
    ]:
        raise ValueError(
            "Numeric output closure failed: Excel evidence headers changed"
        )
    rendered = [
        (
            clean_text(row[0].value),
            clean_text(row[2].value),
            clean_text(row[3].value),
            clean_text(row[4].value),
            clean_text(row[5].value),
            clean_text(row[6].value),
            row[3].coordinate,
        )
        for row in sheet.iter_rows(min_row=2)
        if clean_text(row[0].value)
    ]
    if len(rendered) != len(evidence_rows):
        raise ValueError(
            "Numeric output closure failed because Excel numeric rows do not close"
        )
    locations: dict[str, tuple[str, str]] = {}
    for evidence, (
        evidence_id,
        column,
        value,
        currency,
        unit,
        scale,
        coordinate,
    ) in zip(evidence_rows, rendered, strict=True):
        if (
            evidence_id != evidence["evidence_id"]
            or column != evidence["column"]
            or value != evidence["value"]
            or currency != (evidence["currency"] or "none")
            or unit != evidence["unit"]
            or scale != evidence["scale"]
        ):
            raise ValueError(
                "Numeric output closure failed for Excel evidence "
                f"{evidence['evidence_id']}"
            )
        locations[evidence_id] = (f"numeric_evidence!{coordinate}", value)
    return locations


def _replay_source_numeric_value(
    evidence: dict[str, Any],
    sources_by_id: dict[str, dict[str, Any]],
    output_dir: Path,
) -> str:
    """Reopen one receipted source and recompute the referenced column total."""

    source_receipt = sources_by_id.get(str(evidence["source_artifact_ref"]))
    if not isinstance(source_receipt, dict):
        raise ValueError(
            f"Numeric source receipt is missing for {evidence['evidence_id']}"
        )
    root_path = source_receipt.get("root_path")
    receipt = source_receipt.get("receipt")
    if not isinstance(root_path, str) or not isinstance(receipt, dict):
        raise ValueError(
            f"Numeric source receipt is malformed for {evidence['evidence_id']}"
        )
    source_path = resolve_source_record_path(output_dir, source_receipt)
    validated = validate_artifact_receipt(source_path.parent, receipt)
    if validated["artifact_id"] != evidence["source_artifact_ref"]:
        raise ValueError(
            f"Numeric source identity is stale for {evidence['evidence_id']}"
        )
    identity_key = source_receipt.get("identity_key")
    if not isinstance(identity_key, str) or not identity_key:
        raise ValueError(
            f"Numeric source identity is missing for {evidence['evidence_id']}"
        )
    captured, current_receipt = _capture_source(
        source_path,
        identity_key=identity_key,
    )
    if (
        current_receipt["identity_key"] != source_receipt["identity_key"]
        or current_receipt["receipt"] != source_receipt["receipt"]
    ):
        raise ValueError(
            f"Numeric source receipt changed for {evidence['evidence_id']}"
        )
    tables = _tables_from_captured_source(source_path, captured)
    source_sheet = clean_text(evidence.get("source_sheet"))
    candidates = [
        table for table in tables if clean_text(table.get("sheet_name")) == source_sheet
    ]
    if len(candidates) != 1:
        raise ValueError(
            f"Numeric source table is not unique for {evidence['evidence_id']}"
        )
    rows = candidates[0].get("rows", [])
    expected_header_row = evidence.get("source_header_row")
    if expected_header_row is not None and (
        not isinstance(expected_header_row, int)
        or isinstance(expected_header_row, bool)
        or expected_header_row < 1
        or expected_header_row > len(rows)
    ):
        raise ValueError(
            f"Reviewed numeric source header is invalid for {evidence['evidence_id']}"
        )
    header_idx = (
        expected_header_row - 1 if isinstance(expected_header_row, int) else None
    )
    inventory = _numeric_cell_inventory(
        rows,
        header_idx,
        cell_formats=candidates[0].get("cell_formats"),
        formula_cells=candidates[0].get("formula_cells"),
    )
    contract = evidence.get("numeric_contract")
    if not isinstance(contract, dict):
        raise ValueError(f"Numeric contract is missing for {evidence['evidence_id']}")
    matching_inventory = [
        item
        for item in inventory
        if item["column_index"] == evidence["column_index"]
        and item["column"] == evidence["column"]
    ]
    if len(matching_inventory) != 1:
        raise ValueError(
            f"Numeric source column is not unique for {evidence['evidence_id']}"
        )
    matching_columns, _ = _strict_reviewed_numeric_profile(
        matching_inventory,
        [str(evidence["column"])],
        contract,
        {
            str(evidence["column"]): {
                int(cell["row"]): (
                    "include"
                    if int(cell["row"]) in set(evidence["numeric_rows"])
                    else "exclude"
                )
                for cell in matching_inventory[0].get("nonblank_cells", [])
                if isinstance(cell, dict)
                and isinstance(cell.get("row"), int)
                and not isinstance(cell.get("row"), bool)
            }
        },
    )
    if len(matching_columns) != 1:
        raise ValueError(
            f"Numeric source column is not unique for {evidence['evidence_id']}"
        )
    column = matching_columns[0]
    if (
        column["numeric_rows"] != evidence["numeric_rows"]
        or column["numeric_count"] != evidence["numeric_count"]
        or column["sum"] != evidence["value"]
    ):
        raise ValueError(f"Numeric source closure failed for {evidence['evidence_id']}")
    return str(column["sum"])


def _write_source_receipts(
    output_dir: Path,
    evidence_rows: Sequence[dict[str, Any]],
    sources_by_id: dict[str, dict[str, Any]],
) -> Path:
    """Persist the unique source receipts referenced by numeric evidence."""

    for evidence in evidence_rows:
        artifact_id = str(evidence["source_artifact_ref"])
        source_receipt = sources_by_id.get(artifact_id)
        if not isinstance(source_receipt, dict):
            raise ValueError(
                f"Numeric source receipt is missing for {evidence['evidence_id']}"
            )
        receipt = source_receipt.get("receipt")
        if not isinstance(receipt, dict):
            raise ValueError(
                f"Numeric source receipt is malformed for {evidence['evidence_id']}"
            )
    path = output_dir / "source_receipts.json"
    write_json(
        path,
        {
            "schema_version": "report_builder.source_receipts.v1",
            "sources": [
                {
                    "artifact_id": artifact_id,
                    "receipt": dict(sources_by_id[artifact_id]["receipt"]),
                }
                for artifact_id in sorted(
                    {str(evidence["source_artifact_ref"]) for evidence in evidence_rows}
                )
            ],
        },
    )
    return path


def write_numeric_evidence_ledger(
    output_dir: Path,
    analysis: dict[str, Any],
    *,
    source_context_dir: Path | None = None,
) -> dict[str, Any] | None:
    """Reopen rendered outputs and seal exact numeric source-to-output closure."""

    output_dir = Path(output_dir)
    ledger_path = output_dir / "numeric_evidence_ledger.json"
    prepared_path = output_dir / "report_analysis.json"
    prepared = read_json(prepared_path)
    if prepared != analysis:
        raise ValueError("Numeric output closure failed: prepared analysis is stale")
    evidence_rows = _numeric_evidence_rows(prepared)
    if not evidence_rows:
        ledger_path.unlink(missing_ok=True)
        (output_dir / "source_receipts.json").unlink(missing_ok=True)
        return None

    source_resolution_dir = Path(source_context_dir or output_dir)
    source_index = validate_source_index(
        output_dir,
        source_context_dir=source_resolution_dir,
    )
    sources_by_id = {
        str(source["artifact_id"]): source
        for source in source_index["sources"]
        if isinstance(source, dict)
    }
    _write_source_receipts(output_dir, evidence_rows, sources_by_id)
    language = normalize_language(prepared.get("language"), default="en")
    workbook_locations = _workbook_numeric_locations(
        output_dir / "report_tables.xlsx",
        evidence_rows,
    )
    markdown_locations = _markdown_numeric_locations(
        output_dir / "report_draft.md",
        evidence_rows,
        language=language,
    )
    docx_locations = _docx_numeric_locations(
        output_dir / "report.docx",
        evidence_rows,
        language=language,
    )
    entries = []
    for evidence in evidence_rows:
        evidence_id = str(evidence["evidence_id"])
        source_value = _replay_source_numeric_value(
            evidence,
            sources_by_id,
            source_resolution_dir,
        )
        workbook_locator, workbook_value = workbook_locations[evidence_id]
        markdown_locator, markdown_value = markdown_locations[evidence_id]
        docx_locator, docx_value = docx_locations[evidence_id]
        entries.append(
            {
                "evidence_id": evidence_id,
                "value": evidence["value"],
                "unit": evidence["unit"],
                "currency": evidence["currency"],
                "source": {
                    "artifact_ref": evidence["source_artifact_ref"],
                    "locator": evidence["source_locator"],
                    "value": source_value,
                },
                "prepared": {
                    "artifact_ref": "prepared.report_analysis",
                    "locator": evidence["prepared_locator"],
                    "value": evidence["value"],
                },
                "outputs": [
                    {
                        "artifact_ref": "output.report_tables",
                        "locator": workbook_locator,
                        "value": workbook_value,
                    },
                    {
                        "artifact_ref": "output.report_draft",
                        "locator": markdown_locator,
                        "value": markdown_value,
                    },
                    {
                        "artifact_ref": "output.report_docx",
                        "locator": docx_locator,
                        "value": docx_value,
                    },
                ],
                "calculation_ref": ("calculation.reviewed_signed_scaled_cell_sum.v2"),
                "decision_ref": evidence["decision_ref"],
                "limitations": evidence["limitations"],
            }
        )
    ledger = build_numeric_evidence_ledger(
        entries,
        ledger_id="report_builder.numeric_outputs",
    )
    write_json(ledger_path, ledger)
    return ledger


def _build_report_in_place(
    input_path: Path,
    output_dir: Path,
    *,
    recipe_path: Path | None = None,
    language: object | None = None,
    document_language: object | None = None,
    report_type: object | None = None,
    run_id: str | None = None,
    client_engagement: Mapping[str, Any] | None = None,
) -> BuildResult:
    """Build report outputs from inspected files and an editable recipe."""

    output_dir.mkdir(parents=True, exist_ok=True)
    # A rebuild is a new review run; prior reviewer application state must not carry.
    (output_dir / "applied_decisions.json").unlink(missing_ok=True)
    revisions_dir = output_dir / "revisions"
    if revisions_dir.exists():
        shutil.rmtree(revisions_dir)
    recipe = read_json(recipe_path)
    assumptions = language_assumptions(
        recipe,
        language=language,
        document_language=document_language,
    )
    if not isinstance(recipe.get("sections"), dict) or not recipe.get("sections"):
        inspection = inspect_inputs(
            input_path,
            output_dir,
            language=assumptions["language"],
            document_language=assumptions["document_language"],
            report_type=report_type,
        )
        recipe = inspection.suggested_recipe
    if language is not None:
        recipe["language"] = assumptions["language"]
    if document_language is not None:
        recipe["document_language"] = assumptions["document_language"]
    if report_type is not None:
        recipe["report_type"] = normalize_report_type(report_type)
    validate_narrative_numeric_boundary(recipe)

    run_intake = write_run_intake(
        output_dir,
        input_path=input_path,
        recipe_path=recipe_path,
        language=str(recipe.get("language", assumptions["language"])),
        document_language=str(
            recipe.get("document_language", assumptions["document_language"])
        ),
        report_type=normalize_report_type(recipe.get("report_type")),
        run_id=run_id,
        client_engagement=client_engagement,
    )

    raw_tables = load_tables(input_path, output_dir)
    table_by_id = {clean_text(table.get("table_id")): table for table in raw_tables}
    sections_analysis = [
        analysis_for_section(
            section_key,
            section_recipe,
            table_by_id,
            report_period=clean_text(recipe.get("period")),
        )
        for section_key, section_recipe in selected_sections(recipe).items()
    ]
    assigned_sections = [
        section for section in sections_analysis if section["status"] == "assigned"
    ]
    missing_sections = [
        section["section"]
        for section in sections_analysis
        if section["status"] != "assigned"
    ]
    numeric_measure_pending_sections = [
        section["section"]
        for section in sections_analysis
        if section.get("numeric_measure_status") == "needs_review"
    ]
    analysis = {
        "version": 1,
        "language": recipe.get("language", assumptions["language"]),
        "document_language": recipe.get(
            "document_language", assumptions["document_language"]
        ),
        "report_type": normalize_report_type(recipe.get("report_type")),
        "entity": clean_text(recipe.get("entity")),
        "period": clean_text(recipe.get("period")),
        "sections": sections_analysis,
        "assigned_section_count": len(assigned_sections),
        "missing_sections": missing_sections,
        "numeric_measure_pending_sections": numeric_measure_pending_sections,
    }
    table_inspection = [
        _public_table_inspection(inspect_table(table)) for table in raw_tables
    ]
    report_language = str(recipe.get("language", assumptions["language"]))
    audit_notes = (
        [
            "El texto narrativo lo proporciona Claude en la receta, no los scripts auxiliares.",
            "Revise las secciones sin asignar y los comentarios pendientes de Claude antes del uso final.",
        ]
        if report_language == "es"
        else [
            "Narrative text is supplied by Claude in the recipe, not by helper scripts.",
            "Review unassigned sections and Claude-pending comments before final use.",
        ]
    )
    if numeric_measure_pending_sections:
        audit_notes.append(
            (
                "Las columnas con apariencia numérica permanecen excluidas de los totales hasta que se revise su función como medidas."
                if report_language == "es"
                else "Numeric-looking columns remain excluded from totals until their measure role is reviewed."
            )
        )
    audit = {
        "version": 1,
        "status": "draft",
        "input_path": input_path.name,
        "table_count": len(raw_tables),
        "section_count": len(sections_analysis),
        "assigned_section_count": len(assigned_sections),
        "missing_section_count": len(missing_sections),
        "missing_sections": missing_sections,
        "numeric_measure_pending_section_count": len(numeric_measure_pending_sections),
        "numeric_measure_pending_sections": numeric_measure_pending_sections,
        "codex_narrative_sections": sum(
            1 for section in sections_analysis if section.get("codex_comment")
        ),
        "model_api_calls": 0,
        "notes": audit_notes,
    }

    write_json(output_dir / "report_tables.json", {"tables": table_inspection})
    write_json(output_dir / "report_analysis.json", analysis)
    write_json(output_dir / "report_audit.json", audit)
    write_json(output_dir / "used_recipe.json", recipe)
    write_tables_workbook(output_dir / "report_tables.xlsx", analysis)
    markdown_text = render_markdown(recipe, analysis)
    markdown_path = output_dir / "report_draft.md"
    markdown_path.write_text(markdown_text, encoding="utf-8")
    docx_path = output_dir / "report.docx"
    write_report_docx(recipe, analysis, audit, docx_path)
    numeric_evidence = write_numeric_evidence_ledger(output_dir, analysis)
    numeric_evidence_path = output_dir / "numeric_evidence_ledger.json"
    source_receipts_path = output_dir / "source_receipts.json"
    review_session = write_review_session_artifacts(
        output_dir,
        run_id=run_intake.run_id,
        run_intake_path=run_intake.path,
        analysis=analysis,
        audit=audit,
        recipe=recipe,
        paths={
            "report_tables": output_dir / "report_tables.json",
            "report_tables_xlsx": output_dir / "report_tables.xlsx",
            "report_analysis": output_dir / "report_analysis.json",
            "report_audit": output_dir / "report_audit.json",
            "used_recipe": output_dir / "used_recipe.json",
            "report_draft": markdown_path,
            "report_docx": docx_path,
            **(
                {
                    "numeric_evidence": numeric_evidence_path,
                    "source_receipts": source_receipts_path,
                }
                if numeric_evidence is not None
                else {}
            ),
        },
        tables=table_inspection,
    )
    audit["review_session"] = {
        "run_id": review_session.run_id,
        "run_intake": review_session.run_intake_path.name,
        "review_payload": review_session.review_payload_path.name,
        "ui_decisions": review_session.ui_decisions_path.name,
        "final_artifacts": review_session.final_artifacts_path.name,
        "review_item_count": review_session.review_item_count,
    }
    write_json(output_dir / "report_audit.json", audit)
    refresh_final_artifacts(output_dir, audit=audit, analysis=analysis)
    seal_review_integrity(output_dir, run_id=run_intake.run_id)

    return BuildResult(
        analysis=analysis,
        audit=audit,
        markdown_path=markdown_path,
        docx_path=docx_path,
        review_session=audit["review_session"],
    )


def _reject_unsafe_output_entries(output_dir: Path) -> None:
    """Reject linked or special prior state before any output path is opened."""

    pending = [output_dir]
    while pending:
        directory = pending.pop()
        for child in directory.iterdir():
            metadata = child.lstat()
            if stat.S_ISLNK(metadata.st_mode):
                relative = child.relative_to(output_dir).as_posix()
                raise ValueError(
                    "Report Builder output directory contains a symbolic link: "
                    f"{relative}"
                )
            if stat.S_ISREG(metadata.st_mode):
                if metadata.st_nlink > 1:
                    relative = child.relative_to(output_dir).as_posix()
                    raise ValueError(
                        "Report Builder output directory contains a hard-linked "
                        f"file: {relative}"
                    )
                continue
            if stat.S_ISDIR(metadata.st_mode):
                pending.append(child)
                continue
            relative = child.relative_to(output_dir).as_posix()
            raise ValueError(
                "Report Builder output directory contains an unsupported file "
                f"type: {relative}"
            )


def build_report(
    input_path: Path,
    output_dir: Path,
    *,
    recipe_path: Path | None = None,
    language: object | None = None,
    document_language: object | None = None,
    report_type: object | None = None,
    run_id: str | None = None,
    client_engagement: Mapping[str, Any] | None = None,
) -> BuildResult:
    """Build atomically, restoring the exact prior run after any failure."""

    requested_target = output_dir.expanduser()
    if requested_target.is_symlink():
        raise ValueError("Report Builder output path cannot be a symbolic link.")
    target = requested_target.resolve()
    if target.exists() and not target.is_dir():
        raise ValueError("Report Builder output path must be a directory.")
    if target.exists():
        _reject_unsafe_output_entries(target)
    target.parent.mkdir(parents=True, exist_ok=True)
    transaction_root = Path(
        tempfile.mkdtemp(
            prefix=f".{target.name}.report-build-transaction.",
            dir=target.parent,
        )
    )
    snapshot = transaction_root / "prior-output"
    existed = target.exists()
    if existed:
        shutil.copytree(target, snapshot, symlinks=True)
    completed = False
    try:
        result = _build_report_in_place(
            input_path,
            target,
            recipe_path=recipe_path,
            language=language,
            document_language=document_language,
            report_type=report_type,
            run_id=run_id,
            client_engagement=client_engagement,
        )
        completed = True
        return result
    finally:
        if not completed:
            if target.exists():
                shutil.rmtree(target)
            if existed:
                shutil.copytree(snapshot, target, symlinks=True)
        shutil.rmtree(transaction_root, ignore_errors=True)


def add_common_args(parser: argparse.ArgumentParser) -> None:
    """Add shared CLI arguments."""

    parser.add_argument("input_path", type=Path, help="Input file, folder, or ZIP")
    parser.add_argument(
        "--output-dir",
        type=Path,
        required=True,
        help="Directory where outputs are written",
    )
    parser.add_argument(
        "--language",
        default="en",
        choices=SUPPORTED_LANGUAGES,
        help="Working language for Claude and output labels",
    )
    parser.add_argument(
        "--document-language",
        default="auto",
        choices=SUPPORTED_DOCUMENT_LANGUAGES,
        help="Source-document language assumption",
    )
    parser.add_argument(
        "--report-type",
        default="management_report",
        choices=sorted(REPORT_TYPES),
        help="Report template to use",
    )
    parser.add_argument("--verbose", action="store_true", help="Enable debug logging")


def _json_default(value: Any) -> str:
    if isinstance(value, Path):
        return str(value)
    return str(value)


def print_json(payload: Any) -> None:
    """Print a JSON payload for CLI callers."""

    print(json.dumps(payload, ensure_ascii=False, indent=2, default=_json_default))
