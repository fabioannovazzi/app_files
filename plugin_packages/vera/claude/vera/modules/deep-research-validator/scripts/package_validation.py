"""Package Claude-written answer-validation outputs."""

from __future__ import annotations

import argparse
import json
import re
import sys
import tempfile
from pathlib import Path
from typing import Any

try:
    import pypandoc  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - optional dependency
    pypandoc = None  # type: ignore[assignment]

try:
    from docx import Document  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore[assignment]

try:
    from .review_session import (
        synchronize_final_artifact_sizes,
        write_review_session_artifacts,
        write_run_intake,
    )
except ImportError:  # pragma: no cover - supports direct script imports
    import importlib.util

    _review_session_path = Path(__file__).resolve().parent / "review_session.py"
    _review_session_spec = importlib.util.spec_from_file_location(
        "mparanza_deep_research_validator_review_session",
        _review_session_path,
    )
    assert _review_session_spec and _review_session_spec.loader
    _review_session = importlib.util.module_from_spec(_review_session_spec)
    sys.modules[_review_session_spec.name] = _review_session
    _review_session_spec.loader.exec_module(_review_session)
    synchronize_final_artifact_sizes = _review_session.synchronize_final_artifact_sizes
    write_review_session_artifacts = _review_session.write_review_session_artifacts
    write_run_intake = _review_session.write_run_intake

__all__ = [
    "build_audit",
    "render_validation_package",
    "try_write_docx",
    "validate_answer_contract",
    "write_validation_package",
]

ANSWER_CONTRACT_REQUIRED_FIELDS = (
    "schema_version",
    "question_domain",
    "generation_route",
    "document_type",
    "purpose",
    "audience",
    "output_language",
    "jurisdiction_status",
    "jurisdiction",
    "evidence_display",
    "validation_profile",
    "validation_scope",
    "correction_policy",
    "judgment_policy",
)
ANSWER_CONTRACT_ENUMS = {
    "question_domain": {"legal", "tax", "compliance", "mixed"},
    "generation_route": {
        "chatgpt_deep_research",
        "codex_direct",
        "external_document",
    },
    "jurisdiction_status": {
        "confirmed",
        "assumed",
        "unresolved",
        "not_applicable",
    },
    "evidence_display": {
        "inline_citations",
        "footnotes",
        "source_record_only",
        "mixed",
        "not_specified",
    },
    "validation_profile": {"source_identity_support_reasoning_and_judgment"},
    "validation_scope": {
        "all_material_claims",
        "selected_material_claims",
        "limited",
    },
    "correction_policy": {"correct_when_supported", "review_only"},
    "judgment_policy": {"flag_for_professional_review"},
}
REVIEW_SCHEMA_VERSION = "2.0"
ALLOWED_SUPPORT_STATUSES = {
    "supported",
    "partially_supported",
    "not_supported",
    "contradicted",
    "uncertain",
}
ALLOWED_REASONING_STATUSES = {
    "sound",
    "partially_sound",
    "unsound",
    "uncertain",
    "not_applicable",
}
ALLOWED_JUDGMENT_STATUSES = {
    "not_judgment_dependent",
    "professional_judgment_required",
    "contested",
    "uncertain",
}
ALLOWED_SOURCE_IDENTITY_STATUSES = {
    "matches_cited_source",
    "different_source",
    "uncertain",
    "not_assessed",
}
ALLOWED_ISSUE_TYPES = {
    "none",
    "source_unavailable",
    "source_not_identified",
    "wrong_source",
    "wrong_source_version",
    "wrong_jurisdiction_or_period",
    "missing_source_support",
    "partial_or_overbroad_support",
    "source_contradiction",
    "qualification_or_scope_distortion",
    "temporal_or_modality_distortion",
    "reasoning_gap",
    "judgment_dependent",
    "answer_contract_failure",
}
ALLOWED_TREATMENT_ACTIONS = {
    "none",
    "obtain_source",
    "identify_source",
    "replace_source",
    "add_support",
    "narrow_claim",
    "correct_claim",
    "restore_qualification",
    "correct_time_or_modality",
    "add_reasoning",
    "state_uncertainty",
    "remove_claim",
    "professional_review",
    "revise_answer_contract",
}
ISSUE_TREATMENT_ACTIONS = {
    "none": {"none"},
    "source_unavailable": {
        "obtain_source",
        "state_uncertainty",
        "remove_claim",
        "professional_review",
    },
    "source_not_identified": {
        "identify_source",
        "state_uncertainty",
        "remove_claim",
    },
    "wrong_source": {"replace_source", "state_uncertainty", "remove_claim"},
    "wrong_source_version": {
        "replace_source",
        "state_uncertainty",
        "remove_claim",
    },
    "wrong_jurisdiction_or_period": {
        "replace_source",
        "state_uncertainty",
        "remove_claim",
    },
    "missing_source_support": {
        "add_support",
        "narrow_claim",
        "state_uncertainty",
        "remove_claim",
    },
    "partial_or_overbroad_support": {
        "add_support",
        "narrow_claim",
        "restore_qualification",
        "state_uncertainty",
    },
    "source_contradiction": {"correct_claim", "remove_claim", "state_uncertainty"},
    "qualification_or_scope_distortion": {
        "restore_qualification",
        "narrow_claim",
        "correct_claim",
    },
    "temporal_or_modality_distortion": {
        "correct_time_or_modality",
        "narrow_claim",
        "state_uncertainty",
    },
    "reasoning_gap": {
        "add_reasoning",
        "narrow_claim",
        "state_uncertainty",
        "remove_claim",
    },
    "judgment_dependent": {"professional_review", "state_uncertainty"},
    "answer_contract_failure": {"revise_answer_contract"},
}
ALLOWED_TREATMENT_STATUSES = {
    "not_needed",
    "proposed",
    "applied",
    "blocked",
    "professional_review_required",
}
ALLOWED_DISPOSITIONS = {
    "retain",
    "revise",
    "remove",
    "caveat",
    "pending_source",
    "professional_review",
}
ALLOWED_REVIEWER_ACTIONS = {
    "accept",
    "reject",
    "edit",
    "mark_unclear",
    "request_more_documents",
}
ALLOWED_COVERAGE_SCOPES = {
    "all_material_claims",
    "selected_material_claims",
    "limited",
}
ALLOWED_CONTRACT_CONFORMANCE_STATUSES = {
    "conforms",
    "partially_conforms",
    "does_not_conform",
    "uncertain",
    "not_reviewed",
}
CONTRACT_REVIEW_DIMENSIONS = (
    "question_answered",
    "document_type",
    "audience",
    "evidence_display",
)
ALLOWED_OVERALL_OUTCOMES = {
    "no_material_defect_identified",
    "corrected",
    "correction_required",
    "evidence_limited",
    "professional_review_required",
    "not_reliable",
    "uncertain",
}
ALLOWED_DOCUMENT_REVISION_STATUSES = {
    "not_required",
    "completed",
    "required",
    "blocked",
    "professional_review_required",
}


def _language_code(value: object | None) -> str:
    text = str(value or "en").strip().lower().replace("_", "-")
    return "es" if text.startswith("es") else "en"


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _write_json(path: Path, payload: dict[str, Any]) -> None:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _clean_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _source_aliases(source: dict[str, Any]) -> set[str]:
    aliases: set[str] = set()
    for field in (
        "source_id",
        "url",
        "requested_url",
        "final_url",
        "path",
        "origin_path",
        "name",
    ):
        value = _clean_text(source.get(field))
        if value:
            aliases.add(value.casefold())
    return aliases


def _resolve_source(
    source_ref: str,
    sources: list[dict[str, Any]],
) -> tuple[str, dict[str, Any] | None]:
    """Resolve only exact identifiers; source relevance remains model-led."""

    target = _clean_text(source_ref).casefold()
    if not target:
        return "missing_reference", None
    matches = [source for source in sources if target in _source_aliases(source)]
    if len(matches) == 1:
        return "resolved", matches[0]
    if len(matches) > 1:
        return "ambiguous", None
    return "not_in_inventory", None


def _captured_source_text(
    source: dict[str, Any],
    source_base_dir: Path | None,
) -> tuple[str, str]:
    """Read a bounded captured source snapshot or fall back to its excerpt.

    Relative capture paths are accepted only inside the source-inventory
    directory. This fixed path rule is deterministic for filesystem safety; it
    does not assess source meaning or authority.
    """

    relative = _clean_text(source.get("captured_text_path"))
    if relative and source_base_dir is not None:
        candidate = (source_base_dir / relative).resolve()
        base = source_base_dir.resolve()
        if candidate.is_relative_to(base) and candidate.is_file():
            return candidate.read_text(encoding="utf-8", errors="ignore"), _clean_text(
                source.get("capture_scope") or "captured_text"
            )
    return str(source.get("excerpt") or ""), "inventory_excerpt"


def _exact_passage_observation(
    source_check: dict[str, Any],
    sources: list[dict[str, Any]],
    source_base_dir: Path | None,
) -> dict[str, Any]:
    source_ref = _clean_text(source_check.get("source_ref"))
    resolution, source = _resolve_source(source_ref, sources)
    passage = _clean_text(source_check.get("cited_passage"))
    observation: dict[str, Any] = {
        "source_ref": source_ref,
        "resolution_status": resolution,
        "source_id": source.get("source_id") if source else None,
        "access_status": source.get("status") if source else "not_observed",
        "exact_passage_presence": "not_tested",
        "observation_scope": "none",
        "meaning": (
            "Exact presence is a mechanical observation only and does not decide "
            "semantic support."
        ),
    }
    if source is None or not passage:
        return observation
    source_text, scope = _captured_source_text(source, source_base_dir)
    observation["observation_scope"] = scope
    if not source_text.strip():
        return observation
    target = _clean_text(passage).casefold()
    haystack = _clean_text(source_text).casefold()
    observation["exact_passage_presence"] = (
        "present" if target in haystack else "absent"
    )
    return observation


def validate_answer_contract(answer_contract: dict[str, Any]) -> dict[str, Any]:
    """Validate the explicit answer handoff without making semantic choices."""

    missing_fields = [
        field
        for field in ANSWER_CONTRACT_REQUIRED_FIELDS
        if not isinstance(answer_contract.get(field), str)
        or (
            field != "jurisdiction"
            and not str(answer_contract.get(field) or "").strip()
        )
    ]
    invalid_fields = [
        field
        for field, allowed in ANSWER_CONTRACT_ENUMS.items()
        if str(answer_contract.get(field) or "").strip() not in allowed
    ]
    jurisdiction_status = str(answer_contract.get("jurisdiction_status") or "").strip()
    jurisdiction = str(answer_contract.get("jurisdiction") or "").strip()
    if jurisdiction_status in {"confirmed", "assumed"} and not jurisdiction:
        invalid_fields.append("jurisdiction")
    invalid_fields = list(dict.fromkeys(invalid_fields))
    return {
        "status": "pass" if not missing_fields and not invalid_fields else "fail",
        "missing_fields": missing_fields,
        "invalid_fields": invalid_fields,
        "policy": (
            "shape_validation_only; document type, route, jurisdiction, and "
            "validation posture are model-led or user-confirmed"
        ),
    }


def _claim_index(value: object, fallback: int) -> int:
    try:
        parsed = int(value or fallback)
    except (TypeError, ValueError):
        return fallback
    return parsed if parsed > 0 else fallback


def _string_list(value: object) -> bool:
    return isinstance(value, list) and all(isinstance(item, str) for item in value)


def _assessment_is_valid(value: object, allowed_statuses: set[str]) -> bool:
    return (
        isinstance(value, dict)
        and _clean_text(value.get("status")) in allowed_statuses
        and bool(_clean_text(value.get("analysis")))
    )


def _reasoning_is_valid(value: object) -> bool:
    if not isinstance(value, dict):
        return False
    return (
        _assessment_is_valid(value, ALLOWED_REASONING_STATUSES)
        and _string_list(value.get("supported_premises"))
        and _string_list(value.get("missing_premises"))
    )


def _judgment_is_valid(value: object) -> bool:
    if not isinstance(value, dict):
        return False
    return (
        _assessment_is_valid(value, ALLOWED_JUDGMENT_STATUSES)
        and _string_list(value.get("factors"))
        and _string_list(value.get("alternative_interpretations"))
    )


def _issues_are_valid(value: object, *, allowed_types: set[str]) -> bool:
    if not isinstance(value, list) or not value:
        return False
    issue_types: list[str] = []
    for issue in value:
        if not isinstance(issue, dict):
            return False
        issue_type = _clean_text(issue.get("type"))
        treatment_action = _clean_text(issue.get("treatment_action"))
        treatment_status = _clean_text(issue.get("treatment_status"))
        if (
            issue_type not in allowed_types
            or not _clean_text(issue.get("explanation"))
            or treatment_action not in ALLOWED_TREATMENT_ACTIONS
            or treatment_action not in ISSUE_TREATMENT_ACTIONS.get(issue_type, set())
            or treatment_status not in ALLOWED_TREATMENT_STATUSES
            or not _clean_text(issue.get("treatment_explanation"))
        ):
            return False
        if issue_type == "none":
            if treatment_action != "none" or treatment_status != "not_needed":
                return False
        elif treatment_action == "none" or treatment_status == "not_needed":
            return False
        issue_types.append(issue_type)
    return not ("none" in issue_types and len(issue_types) != 1)


def _non_none_issues(value: object) -> list[dict[str, Any]]:
    """Return recorded issue treatments other than the explicit no-issue marker."""

    if not isinstance(value, list):
        return []
    return [
        issue
        for issue in value
        if isinstance(issue, dict) and _clean_text(issue.get("type")) != "none"
    ]


def _claim_consistency_errors(claim: dict[str, Any]) -> list[str]:
    """Return mechanically provable contradictions within one review record.

    This function does not decide whether a source supports a claim or whether
    legal reasoning is sound. It only prevents the validator from treating a
    record as complete when the reviewer's own coded assessments conflict.
    """

    errors: list[str] = []
    source_checks = claim.get("source_checks")
    support = claim.get("support") if isinstance(claim.get("support"), dict) else {}
    reasoning = (
        claim.get("reasoning") if isinstance(claim.get("reasoning"), dict) else {}
    )
    judgment = (
        claim.get("professional_judgment")
        if isinstance(claim.get("professional_judgment"), dict)
        else {}
    )
    disposition = (
        claim.get("disposition") if isinstance(claim.get("disposition"), dict) else {}
    )
    support_status = _clean_text(support.get("status"))
    reasoning_status = _clean_text(reasoning.get("status"))
    judgment_status = _clean_text(judgment.get("status"))
    disposition_status = _clean_text(disposition.get("status"))
    non_none_issues = _non_none_issues(claim.get("issues"))

    if support_status in {"supported", "partially_supported", "contradicted"} and not (
        isinstance(source_checks, list) and source_checks
    ):
        errors.append("source_check_required_for_support_assessment")
    identity_attention = isinstance(source_checks, list) and any(
        isinstance(source_check, dict)
        and _clean_text(source_check.get("identity_status")) != "matches_cited_source"
        for source_check in source_checks
    )
    semantic_attention = support_status != "supported" or reasoning_status in {
        "partially_sound",
        "unsound",
        "uncertain",
    }
    judgment_attention = judgment_status in {
        "professional_judgment_required",
        "contested",
        "uncertain",
    }
    if (identity_attention or semantic_attention or judgment_attention) and not (
        non_none_issues
    ):
        errors.append("attention_assessment_requires_issue_treatment")
    if support_status in {"not_supported", "contradicted"} and (
        disposition_status == "retain"
    ):
        errors.append("unsupported_or_contradicted_claim_cannot_be_retained")
    if reasoning_status == "unsound" and disposition_status == "retain":
        errors.append("unsound_reasoning_cannot_be_retained")
    return errors


def _contract_consistency_errors(contract_review: object) -> list[str]:
    """Return contradictions between conformance statuses and contract issues."""

    if not isinstance(contract_review, dict):
        return []
    attention = any(
        isinstance(contract_review.get(dimension), dict)
        and _clean_text(contract_review[dimension].get("status")) != "conforms"
        for dimension in CONTRACT_REVIEW_DIMENSIONS
    )
    issue_types = {
        _clean_text(issue.get("type"))
        for issue in contract_review.get("issues", [])
        if isinstance(issue, dict)
    }
    errors: list[str] = []
    if attention and "answer_contract_failure" not in issue_types:
        errors.append("contract_attention_requires_failure_treatment")
    if not attention and "answer_contract_failure" in issue_types:
        errors.append("contract_failure_treatment_requires_attention_status")
    return errors


def _delivery_readiness(
    *,
    record_integrity_status: str,
    revision_status: str,
    judgment_dependent: bool,
    evidence_limited: bool,
    contract_attention: bool,
    pending_treatments: bool,
    blocked_treatments: bool,
    coverage_limited: bool,
    overall_outcome: str,
) -> str:
    """Aggregate explicit review statuses; never infer substantive correctness."""

    if record_integrity_status != "record_complete":
        return "review_record_incomplete"
    if revision_status == "blocked" or blocked_treatments:
        return "blocked"
    if overall_outcome == "not_reliable":
        return "not_reliable"
    if (
        revision_status == "required"
        or pending_treatments
        or contract_attention
        or overall_outcome == "correction_required"
    ):
        return "revision_required"
    if (
        judgment_dependent
        or revision_status == "professional_review_required"
        or overall_outcome in {"professional_review_required", "uncertain"}
    ):
        return "professional_review_required"
    if evidence_limited or coverage_limited or overall_outcome == "evidence_limited":
        return "evidence_limited"
    return "reviewed_answer_ready"


def build_audit(
    document_inventory: dict[str, Any],
    source_inventory: dict[str, Any],
    claims_review: dict[str, Any],
    answer_contract: dict[str, Any],
    *,
    source_base_dir: Path | None = None,
) -> dict[str, Any]:
    """Audit review-record completeness without deciding substantive meaning.

    Fixed checks are justified here by the stable review contract: they verify
    required fields, allowed codes, exact source identifiers, and exact passage
    presence. All legal relevance, support, reasoning, and judgment statuses are
    authored by the model or professional reviewer and are never overridden.
    """

    claims = claims_review.get("claims", [])
    if not isinstance(claims, list):
        claims = []
    sources = source_inventory.get("sources", [])
    if not isinstance(sources, list):
        sources = []

    failed_checks: list[str] = []
    if int(document_inventory.get("character_count", 0) or 0) <= 0:
        failed_checks.append("document_text_present")
    if not claims:
        failed_checks.append("claims_review_present")
    answer_contract_audit = validate_answer_contract(answer_contract)
    if answer_contract_audit["status"] != "pass":
        failed_checks.append("answer_contract")

    if _clean_text(claims_review.get("schema_version")) != REVIEW_SCHEMA_VERSION:
        failed_checks.append("review_schema_version")
    if not _clean_text(claims_review.get("validation_objective")):
        failed_checks.append("validation_objective_present")

    coverage = claims_review.get("coverage_review")
    coverage_valid = isinstance(coverage, dict)
    if coverage_valid:
        coverage_valid = (
            _clean_text(coverage.get("selection_method"))
            == "model_led_materiality_review"
            and _clean_text(coverage.get("scope")) in ALLOWED_COVERAGE_SCOPES
            and bool(_clean_text(coverage.get("analysis")))
            and all(
                isinstance(coverage.get(field), list)
                and all(isinstance(item, str) for item in coverage.get(field, []))
                for field in ("reviewed_sections", "omitted_sections", "limitations")
            )
            and _clean_text(coverage.get("reviewer_action")) in ALLOWED_REVIEWER_ACTIONS
        )
    if not coverage_valid:
        failed_checks.append("coverage_review_complete")
    coverage_scope = (
        _clean_text(coverage.get("scope")) if isinstance(coverage, dict) else ""
    )
    coverage_reviewer_rejected = (
        isinstance(coverage, dict)
        and _clean_text(coverage.get("reviewer_action")) == "reject"
    )
    contract_scope = _clean_text(answer_contract.get("validation_scope"))
    if coverage_valid and coverage_scope != contract_scope:
        failed_checks.append("coverage_matches_answer_contract")

    contract_review = claims_review.get("contract_review")
    contract_review_valid = isinstance(contract_review, dict)
    contract_attention: list[str] = []
    if contract_review_valid:
        for dimension in CONTRACT_REVIEW_DIMENSIONS:
            assessment = contract_review.get(dimension)
            if not (
                isinstance(assessment, dict)
                and _clean_text(assessment.get("status"))
                in ALLOWED_CONTRACT_CONFORMANCE_STATUSES
                and bool(_clean_text(assessment.get("analysis")))
            ):
                contract_review_valid = False
                continue
            if _clean_text(assessment.get("status")) != "conforms":
                contract_attention.append(dimension)
        if not _issues_are_valid(
            contract_review.get("issues"),
            allowed_types={"none", "answer_contract_failure"},
        ):
            contract_review_valid = False
        if _clean_text(contract_review.get("reviewer_action")) not in (
            ALLOWED_REVIEWER_ACTIONS
        ):
            contract_review_valid = False
    if not contract_review_valid:
        failed_checks.append("contract_review_complete")
    contract_reviewer_rejected = (
        isinstance(contract_review, dict)
        and _clean_text(contract_review.get("reviewer_action")) == "reject"
    )
    consistency_errors: list[dict[str, Any]] = []
    for error in _contract_consistency_errors(contract_review):
        consistency_errors.append({"scope": "contract_review", "error": error})

    invalid_claim_indices: list[int] = []
    invalid_source_check_indices: list[int] = []
    invalid_issue_indices: list[int] = []
    invalid_disposition_indices: list[int] = []
    invalid_reviewer_action_indices: list[int] = []
    support_attention_claims: list[int] = []
    reasoning_attention_claims: list[int] = []
    judgment_dependent_claims: list[int] = []
    source_identity_attention_claims: list[int] = []
    pending_treatment_claims: list[int] = []
    blocked_treatment_claims: list[int] = []
    evidence_limited_claims: list[int] = []
    claim_observations: list[dict[str, Any]] = []

    for position, raw_claim in enumerate(claims, start=1):
        if not isinstance(raw_claim, dict):
            invalid_claim_indices.append(position)
            continue
        claim = raw_claim
        claim_index = _claim_index(claim.get("claim_index"), position)
        claim_valid = (
            bool(_clean_text(claim.get("claim_text")))
            and bool(_clean_text(claim.get("claim_location")))
            and _clean_text(claim.get("materiality")) in {"material", "supporting"}
            and isinstance(claim.get("proposed_fix"), str)
        )

        source_checks = claim.get("source_checks")
        source_checks_valid = isinstance(source_checks, list)
        source_observations: list[dict[str, Any]] = []
        if source_checks_valid:
            for source_check in source_checks:
                if not (
                    isinstance(source_check, dict)
                    and bool(_clean_text(source_check.get("source_ref")))
                    and _clean_text(source_check.get("identity_status"))
                    in ALLOWED_SOURCE_IDENTITY_STATUSES
                    and bool(_clean_text(source_check.get("identity_analysis")))
                    and isinstance(source_check.get("cited_passage"), str)
                ):
                    source_checks_valid = False
                    continue
                if _clean_text(source_check.get("identity_status")) != (
                    "matches_cited_source"
                ):
                    source_identity_attention_claims.append(claim_index)
                observation = _exact_passage_observation(
                    source_check,
                    sources,
                    source_base_dir,
                )
                source_observations.append(observation)
                if observation["access_status"] not in {"available", "not_observed"}:
                    evidence_limited_claims.append(claim_index)
                if observation["resolution_status"] != "resolved":
                    evidence_limited_claims.append(claim_index)
        if not source_checks_valid:
            invalid_source_check_indices.append(claim_index)

        support = claim.get("support")
        support_valid = _assessment_is_valid(support, ALLOWED_SUPPORT_STATUSES)
        if support_valid and _clean_text(support.get("status")) != "supported":
            support_attention_claims.append(claim_index)
        reasoning = claim.get("reasoning")
        reasoning_valid = _reasoning_is_valid(reasoning)
        if reasoning_valid and _clean_text(reasoning.get("status")) in {
            "partially_sound",
            "unsound",
            "uncertain",
        }:
            reasoning_attention_claims.append(claim_index)
        judgment = claim.get("professional_judgment")
        judgment_valid = _judgment_is_valid(judgment)
        if judgment_valid and _clean_text(judgment.get("status")) in {
            "professional_judgment_required",
            "contested",
            "uncertain",
        }:
            judgment_dependent_claims.append(claim_index)

        issues = claim.get("issues")
        issues_valid = _issues_are_valid(issues, allowed_types=ALLOWED_ISSUE_TYPES)
        if not issues_valid:
            invalid_issue_indices.append(claim_index)
        elif any(
            _clean_text(issue.get("treatment_status")) == "proposed"
            for issue in issues
            if isinstance(issue, dict)
        ):
            pending_treatment_claims.append(claim_index)
        if isinstance(issues, list) and any(
            isinstance(issue, dict)
            and _clean_text(issue.get("treatment_status")) == "blocked"
            for issue in issues
        ):
            blocked_treatment_claims.append(claim_index)
        if isinstance(issues, list) and any(
            isinstance(issue, dict)
            and _clean_text(issue.get("treatment_status"))
            == "professional_review_required"
            for issue in issues
        ):
            judgment_dependent_claims.append(claim_index)
        if isinstance(issues, list) and any(
            isinstance(issue, dict)
            and _clean_text(issue.get("type"))
            in {"source_unavailable", "source_not_identified"}
            for issue in issues
        ):
            evidence_limited_claims.append(claim_index)

        disposition = claim.get("disposition")
        if not (
            isinstance(disposition, dict)
            and _clean_text(disposition.get("status")) in ALLOWED_DISPOSITIONS
            and bool(_clean_text(disposition.get("analysis")))
            and isinstance(disposition.get("revised_claim"), str)
        ):
            invalid_disposition_indices.append(claim_index)
        if _clean_text(claim.get("reviewer_action")) not in ALLOWED_REVIEWER_ACTIONS:
            invalid_reviewer_action_indices.append(claim_index)

        if not (claim_valid and support_valid and reasoning_valid and judgment_valid):
            invalid_claim_indices.append(claim_index)
        for error in _claim_consistency_errors(claim):
            consistency_errors.append(
                {"scope": "claim", "claim_index": claim_index, "error": error}
            )
        claim_observations.append(
            {
                "claim_index": claim_index,
                "source_observations": source_observations,
            }
        )

    if invalid_claim_indices:
        failed_checks.append("claim_assessments_complete")
    if invalid_source_check_indices:
        failed_checks.append("source_identity_assessments_complete")
    if invalid_issue_indices:
        failed_checks.append("issue_treatments_complete")
    if invalid_disposition_indices:
        failed_checks.append("claim_dispositions_complete")
    if invalid_reviewer_action_indices:
        failed_checks.append("reviewer_actions_valid")

    rejected_claim_indices = [
        _claim_index(claim.get("claim_index"), position)
        for position, claim in enumerate(claims, start=1)
        if isinstance(claim, dict)
        and _clean_text(claim.get("reviewer_action")) == "reject"
    ]

    overall = claims_review.get("overall_assessment")
    overall_valid = (
        isinstance(overall, dict)
        and _clean_text(overall.get("outcome")) in ALLOWED_OVERALL_OUTCOMES
        and bool(_clean_text(overall.get("analysis")))
        and _string_list(overall.get("residual_uncertainties"))
        and _string_list(overall.get("professional_review_items"))
    )
    if not overall_valid:
        failed_checks.append("overall_assessment_complete")

    document_revision = claims_review.get("document_revision")
    revision_status = (
        _clean_text(document_revision.get("status"))
        if isinstance(document_revision, dict)
        else ""
    )
    revision_valid = (
        isinstance(document_revision, dict)
        and revision_status in ALLOWED_DOCUMENT_REVISION_STATUSES
        and bool(_clean_text(document_revision.get("summary")))
        and _string_list(document_revision.get("unresolved_changes"))
    )
    validated_document = str(claims_review.get("validated_document") or "").strip()
    if (
        revision_status
        in {
            "not_required",
            "completed",
            "professional_review_required",
        }
        and not validated_document
    ):
        revision_valid = False
    if not revision_valid:
        failed_checks.append("document_revision_complete")

    overall_outcome = _clean_text(overall.get("outcome")) if overall_valid else ""
    attention_recorded = any(
        (
            support_attention_claims,
            reasoning_attention_claims,
            judgment_dependent_claims,
            source_identity_attention_claims,
            pending_treatment_claims,
            blocked_treatment_claims,
            contract_attention,
            rejected_claim_indices,
            ["coverage"] if coverage_reviewer_rejected else [],
            ["contract"] if contract_reviewer_rejected else [],
        )
    )
    if attention_recorded and overall_outcome == "no_material_defect_identified":
        consistency_errors.append(
            {
                "scope": "overall_assessment",
                "error": "recorded_attention_conflicts_with_no_material_defect",
            }
        )
    allowed_outcomes_by_revision = {
        "not_required": {
            "no_material_defect_identified",
            "evidence_limited",
            "uncertain",
        },
        "completed": {"corrected"},
        "required": {
            "correction_required",
            "evidence_limited",
            "professional_review_required",
            "not_reliable",
            "uncertain",
        },
        "blocked": {"evidence_limited", "not_reliable", "uncertain"},
        "professional_review_required": {
            "professional_review_required",
            "uncertain",
        },
    }
    if (
        revision_status in allowed_outcomes_by_revision
        and overall_outcome
        and overall_outcome not in allowed_outcomes_by_revision[revision_status]
    ):
        consistency_errors.append(
            {
                "scope": "document_revision",
                "error": "revision_status_conflicts_with_overall_outcome",
            }
        )
    if revision_status == "not_required" and (
        pending_treatment_claims
        or blocked_treatment_claims
        or contract_attention
        or rejected_claim_indices
        or coverage_reviewer_rejected
        or contract_reviewer_rejected
    ):
        consistency_errors.append(
            {
                "scope": "document_revision",
                "error": "unresolved_treatment_conflicts_with_no_revision",
            }
        )
    if consistency_errors:
        failed_checks.append("review_state_consistent")

    failed_checks = list(dict.fromkeys(failed_checks))
    record_integrity_status = (
        "record_complete" if not failed_checks else "record_incomplete"
    )
    delivery_readiness = _delivery_readiness(
        record_integrity_status=record_integrity_status,
        revision_status=revision_status,
        judgment_dependent=bool(judgment_dependent_claims),
        evidence_limited=bool(evidence_limited_claims),
        contract_attention=bool(contract_attention or contract_reviewer_rejected),
        pending_treatments=bool(
            pending_treatment_claims
            or rejected_claim_indices
            or coverage_reviewer_rejected
        ),
        blocked_treatments=bool(blocked_treatment_claims),
        coverage_limited=coverage_scope != "all_material_claims",
        overall_outcome=overall_outcome,
    )
    return {
        "status": record_integrity_status,
        "record_integrity_status": record_integrity_status,
        "delivery_readiness": delivery_readiness,
        "failed_checks": failed_checks,
        "claim_count": len(claims),
        "support_attention_claim_indices": sorted(set(support_attention_claims)),
        "reasoning_attention_claim_indices": reasoning_attention_claims,
        "judgment_dependent_claim_indices": sorted(set(judgment_dependent_claims)),
        "source_identity_attention_claim_indices": sorted(
            set(source_identity_attention_claims)
        ),
        "evidence_limited_claim_indices": sorted(set(evidence_limited_claims)),
        "pending_treatment_claim_indices": sorted(set(pending_treatment_claims)),
        "blocked_treatment_claim_indices": sorted(set(blocked_treatment_claims)),
        "invalid_claim_indices": sorted(set(invalid_claim_indices)),
        "invalid_source_check_indices": sorted(set(invalid_source_check_indices)),
        "invalid_issue_indices": sorted(set(invalid_issue_indices)),
        "invalid_disposition_indices": sorted(set(invalid_disposition_indices)),
        "invalid_reviewer_action_indices": sorted(set(invalid_reviewer_action_indices)),
        "rejected_claim_indices": sorted(set(rejected_claim_indices)),
        "coverage_reviewer_rejected": coverage_reviewer_rejected,
        "contract_reviewer_rejected": contract_reviewer_rejected,
        "consistency_errors": consistency_errors,
        "contract_attention_dimensions": contract_attention,
        "coverage_scope": coverage_scope,
        "claim_observations": claim_observations,
        "source_count": len(sources),
        "document_url_count": len(document_inventory.get("urls", []) or []),
        "answer_contract": answer_contract,
        "answer_contract_audit": answer_contract_audit,
        "assurance_boundary": {
            "mechanically_observed": (
                "document and captured-source availability, exact identifier "
                "resolution, exact passage presence in the cited source snapshot, "
                "review-record shape, and contradictions among explicit review "
                "status fields"
            ),
            "semantically_assessed": (
                "source identity and authority, claim meaning, entailment, "
                "contradiction, qualification, scope, time, and modality"
            ),
            "reasoning_assessed": (
                "whether the conclusion follows from supported premises and "
                "whether intermediate premises are missing"
            ),
            "judgment_dependent": (
                "legal applicability, materiality, competing interpretations, "
                "professional choices, and uncertain outcomes"
            ),
            "record_integrity_meaning": (
                "record_complete means the required assessments and treatments "
                "were recorded; it does not certify legal correctness or replace "
                "professional review"
            ),
        },
    }


def _package_markdown(
    document_inventory: dict[str, Any],
    source_inventory: dict[str, Any],
    claims_review: dict[str, Any],
    audit: dict[str, Any],
    validated_document: str,
) -> str:
    spanish = _language_code(claims_review.get("language")) == "es"
    claims = claims_review.get("claims", [])
    claim_lines: list[str] = []
    if isinstance(claims, list):
        for position, claim in enumerate(claims, start=1):
            if not isinstance(claim, dict):
                continue
            support = (
                claim.get("support") if isinstance(claim.get("support"), dict) else {}
            )
            reasoning = (
                claim.get("reasoning")
                if isinstance(claim.get("reasoning"), dict)
                else {}
            )
            judgment = (
                claim.get("professional_judgment")
                if isinstance(claim.get("professional_judgment"), dict)
                else {}
            )
            disposition = (
                claim.get("disposition")
                if isinstance(claim.get("disposition"), dict)
                else {}
            )
            source_checks = claim.get("source_checks")
            issues = claim.get("issues")
            claim_index = claim.get("claim_index", position)
            if spanish:
                lines = [
                    f"### Afirmación {claim_index}",
                    f"**Texto:** {_clean_text(claim.get('claim_text'))}",
                    f"**Ubicación:** {_clean_text(claim.get('claim_location'))}",
                    f"**Materialidad:** {_clean_text(claim.get('materiality'))}",
                    "#### Identidad y acceso a las fuentes",
                    json.dumps(source_checks, ensure_ascii=False, indent=2),
                    f"#### Respaldo semántico — {_clean_text(support.get('status'))}",
                    _clean_text(support.get("analysis")),
                    f"#### Razonamiento — {_clean_text(reasoning.get('status'))}",
                    _clean_text(reasoning.get("analysis")),
                    f"#### Juicio profesional — {_clean_text(judgment.get('status'))}",
                    _clean_text(judgment.get("analysis")),
                    "#### Incidencias y tratamiento",
                    json.dumps(issues, ensure_ascii=False, indent=2),
                    f"#### Disposición — {_clean_text(disposition.get('status'))}",
                    _clean_text(disposition.get("analysis")),
                    f"**Corrección propuesta:** {_clean_text(claim.get('proposed_fix'))}",
                ]
            else:
                lines = [
                    f"### Claim {claim_index}",
                    f"**Text:** {_clean_text(claim.get('claim_text'))}",
                    f"**Location:** {_clean_text(claim.get('claim_location'))}",
                    f"**Materiality:** {_clean_text(claim.get('materiality'))}",
                    "#### Source identity and access",
                    json.dumps(source_checks, ensure_ascii=False, indent=2),
                    f"#### Semantic source support — {_clean_text(support.get('status'))}",
                    _clean_text(support.get("analysis")),
                    f"#### Reasoning — {_clean_text(reasoning.get('status'))}",
                    _clean_text(reasoning.get("analysis")),
                    f"#### Professional judgment — {_clean_text(judgment.get('status'))}",
                    _clean_text(judgment.get("analysis")),
                    "#### Issues and treatment",
                    json.dumps(issues, ensure_ascii=False, indent=2),
                    f"#### Disposition — {_clean_text(disposition.get('status'))}",
                    _clean_text(disposition.get("analysis")),
                    f"**Proposed fix:** {_clean_text(claim.get('proposed_fix'))}",
                ]
            claim_lines.append("\n\n".join(lines).strip())
    if spanish:
        sections = [
            "# Registro de validación de la respuesta",
            f"Integridad del registro: {audit.get('record_integrity_status')}",
            f"Preparación para entrega: {audit.get('delivery_readiness')}",
            f"Afirmaciones revisadas: {audit.get('claim_count')}",
            f"Fuentes examinadas: {audit.get('source_count')}",
            "## Límite de aseguramiento",
            json.dumps(audit.get("assurance_boundary"), ensure_ascii=False, indent=2),
            "## Contrato de respuesta",
            json.dumps(audit.get("answer_contract"), ensure_ascii=False, indent=2),
            "## Revisión del contrato de respuesta",
            json.dumps(
                claims_review.get("contract_review"), ensure_ascii=False, indent=2
            ),
            "## Cobertura de la revisión",
            json.dumps(
                claims_review.get("coverage_review"), ensure_ascii=False, indent=2
            ),
            "## Inventario del documento",
            f"Palabras: {document_inventory.get('word_count', 0)}",
            f"URL: {document_inventory.get('urls', [])}",
            "## Inventario de fuentes",
            json.dumps(source_inventory, ensure_ascii=False, indent=2),
            "## Observaciones mecánicas de las fuentes",
            json.dumps(audit.get("claim_observations"), ensure_ascii=False, indent=2),
            "## Evaluaciones de las afirmaciones",
            (
                "\n\n".join(claim_lines)
                if claim_lines
                else "No se revisaron afirmaciones."
            ),
            "## Evaluación general",
            json.dumps(
                claims_review.get("overall_assessment"), ensure_ascii=False, indent=2
            ),
            "## Estado de revisión del documento",
            json.dumps(
                claims_review.get("document_revision"), ensure_ascii=False, indent=2
            ),
        ]
    else:
        sections = [
            "# Answer Validation Record",
            f"Record integrity: {audit.get('record_integrity_status')}",
            f"Delivery readiness: {audit.get('delivery_readiness')}",
            f"Claims reviewed: {audit.get('claim_count')}",
            f"Sources inspected: {audit.get('source_count')}",
            "## Assurance Boundary",
            json.dumps(audit.get("assurance_boundary"), ensure_ascii=False, indent=2),
            "## Answer Contract",
            json.dumps(audit.get("answer_contract"), ensure_ascii=False, indent=2),
            "## Answer-Contract Review",
            json.dumps(
                claims_review.get("contract_review"), ensure_ascii=False, indent=2
            ),
            "## Review Coverage",
            json.dumps(
                claims_review.get("coverage_review"), ensure_ascii=False, indent=2
            ),
            "## Document Inventory",
            f"Words: {document_inventory.get('word_count', 0)}",
            f"URLs: {document_inventory.get('urls', [])}",
            "## Source Inventory",
            json.dumps(source_inventory, ensure_ascii=False, indent=2),
            "## Mechanical Source Observations",
            json.dumps(audit.get("claim_observations"), ensure_ascii=False, indent=2),
            "## Claim Assessments",
            "\n\n".join(claim_lines) if claim_lines else "No claims reviewed.",
            "## Overall Assessment",
            json.dumps(
                claims_review.get("overall_assessment"), ensure_ascii=False, indent=2
            ),
            "## Document Revision Status",
            json.dumps(
                claims_review.get("document_revision"), ensure_ascii=False, indent=2
            ),
        ]
    if validated_document.strip():
        sections.extend(
            [
                (
                    "## Respuesta revisada o corregida"
                    if spanish
                    else "## Reviewed or Corrected Answer"
                ),
                validated_document.strip(),
            ]
        )
    return "\n\n".join(sections).strip() + "\n"


def render_validation_package(
    document_inventory: dict[str, Any],
    source_inventory: dict[str, Any],
    claims_review: dict[str, Any],
    audit: dict[str, Any],
    validated_document: str,
) -> str:
    """Render package Markdown from already-reviewed validation records."""

    return _package_markdown(
        document_inventory,
        source_inventory,
        claims_review,
        audit,
        validated_document,
    )


def _write_docx_fallback(markdown_text: str, output_path: Path) -> bool:
    if Document is None:
        return False
    document = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading_match = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading_match:
            document.add_heading(
                heading_match.group(2).strip(),
                level=min(len(heading_match.group(1)), 4),
            )
            continue
        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue
        document.add_paragraph(line)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return True


def try_write_docx(markdown_text: str, output_path: Path) -> bool:
    """Write a DOCX from Markdown using local deterministic renderers only."""

    if pypandoc is None:
        return _write_docx_fallback(markdown_text, output_path)
    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmpfile:
            temp_path = Path(tmpfile.name)
        pypandoc.convert_text(
            markdown_text,
            to="docx",
            format="md",
            outputfile=str(temp_path),
        )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(temp_path.read_bytes())
        return True
    except (OSError, RuntimeError):
        return _write_docx_fallback(markdown_text, output_path)
    finally:
        if temp_path and temp_path.exists():
            temp_path.unlink()


def write_validation_package(
    document_inventory_path: Path,
    source_inventory_path: Path,
    claims_review_path: Path,
    output_dir: Path,
    *,
    answer_contract_path: Path,
    validated_document_path: Path | None = None,
    write_docx: bool = False,
) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    document_inventory = _read_json(document_inventory_path)
    source_inventory = _read_json(source_inventory_path)
    claims_review = _read_json(claims_review_path)
    answer_contract = _read_json(answer_contract_path)
    run_intake = write_run_intake(
        output_dir,
        document_inventory_path=document_inventory_path,
        source_inventory_path=source_inventory_path,
        claims_review_path=claims_review_path,
        answer_contract_path=answer_contract_path,
        document_inventory=document_inventory,
        source_inventory=source_inventory,
        claims_review=claims_review,
        answer_contract=answer_contract,
    )
    validated_document = (
        validated_document_path.read_text(encoding="utf-8")
        if validated_document_path and validated_document_path.exists()
        else str(claims_review.get("validated_document", "") or "")
    )
    if (
        validated_document.strip()
        and not str(claims_review.get("validated_document") or "").strip()
    ):
        claims_review = dict(claims_review)
        claims_review["validated_document"] = validated_document.strip()

    audit = build_audit(
        document_inventory,
        source_inventory,
        claims_review,
        answer_contract,
        source_base_dir=source_inventory_path.parent,
    )
    answer_contract_out_path = output_dir / "answer_contract.json"
    audit_path = output_dir / "validation_audit.json"
    review_out_path = output_dir / "claims_review.json"
    package_path = output_dir / "validation_package.md"
    _write_json(audit_path, audit)
    _write_json(review_out_path, claims_review)
    _write_json(answer_contract_out_path, answer_contract)
    package_text = render_validation_package(
        document_inventory,
        source_inventory,
        claims_review,
        audit,
        validated_document,
    )
    package_path.write_text(package_text, encoding="utf-8")

    paths = {
        "answer_contract": answer_contract_out_path,
        "claims_review": review_out_path,
        "validation_audit": audit_path,
        "validation_package": package_path,
    }
    if validated_document.strip():
        validated_md_path = output_dir / "validated_document.md"
        validated_md_path.write_text(
            validated_document.strip() + "\n",
            encoding="utf-8",
        )
        paths["validated_document"] = validated_md_path
    if write_docx and validated_document.strip():
        docx_path = output_dir / "validated_document.docx"
        if try_write_docx(validated_document, docx_path):
            paths["validated_document_docx"] = docx_path
    review_session = write_review_session_artifacts(
        output_dir,
        run_id=run_intake.run_id,
        run_intake_path=run_intake.path,
        document_inventory_path=document_inventory_path,
        source_inventory_path=source_inventory_path,
        claims_review_path=claims_review_path,
        answer_contract_path=answer_contract_path,
        document_inventory=document_inventory,
        source_inventory=source_inventory,
        claims_review=claims_review,
        answer_contract=answer_contract,
        audit=audit,
        paths=paths,
    )
    audit["review_session"] = {
        "run_id": review_session.run_id,
        "run_intake_path": str(review_session.run_intake_path),
        "review_payload_path": str(review_session.review_payload_path),
        "ui_decisions_path": str(review_session.ui_decisions_path),
        "final_artifacts_path": str(review_session.final_artifacts_path),
        "review_item_count": review_session.review_item_count,
    }
    _write_json(audit_path, audit)
    synchronize_final_artifact_sizes(review_session.final_artifacts_path)
    return paths


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document_inventory", type=Path)
    parser.add_argument("source_inventory", type=Path)
    parser.add_argument("claims_review", type=Path)
    parser.add_argument(
        "--answer-contract-file",
        type=Path,
        required=True,
        help="JSON answer contract produced during question intake.",
    )
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--validated-document", type=Path)
    parser.add_argument("--docx", action="store_true")
    args = parser.parse_args()
    write_validation_package(
        args.document_inventory,
        args.source_inventory,
        args.claims_review,
        args.output_dir,
        answer_contract_path=args.answer_contract_file,
        validated_document_path=args.validated_document,
        write_docx=args.docx,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
